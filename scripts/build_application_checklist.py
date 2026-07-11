#!/usr/bin/env python3
"""Build a one-page job-specific application checklist."""

from __future__ import annotations

import re
import sys
import zipfile
from datetime import date, datetime
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import build_cover_letter
import build_interview_cheat_sheet as cheat
import build_resume
import business_context
import resume_analysis
from utils import assert_no_template_leakage, enforce_prose_quality


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
DEBRIEF_HISTORY = PROJECT_ROOT / "jobs" / "debrief_history.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"
FONT_NAME = "Carlito"
DEBRIEF_DELIMITER = "POST-INTERVIEW DEBRIEF CAPTURED"


def docx_text(path: Path) -> str:
    return "\n".join(cheat.paragraph_texts(path))


def term_occurrences(text: str, term: str) -> int:
    if " " in term:
        return len(re.findall(re.escape(term), text, flags=re.I))
    return len(re.findall(rf"\b{re.escape(term)}\b", text, flags=re.I))


def top_keywords(job_description: str, limit: int = 12) -> list[str]:
    keywords = build_resume.audit_keywords(job_description)
    return sorted(
        keywords,
        key=lambda keyword: (
            term_occurrences(job_description, keyword),
            len(keyword.split()),
            len(keyword),
            keyword,
        ),
        reverse=True,
    )[:limit]


def keyword_status(keyword: str, resume_text: str) -> str:
    return "COVERED" if term_occurrences(resume_text, keyword) else "MISSING"


def source_resume_label(path: Path) -> str:
    if path == build_resume.PRESALES_CSM_RESUME:
        return "Pre-Sales/CSM"
    if path == build_resume.IMPLEMENTATION_RESUME:
        return "Implementation"
    return path.stem


def latest_tailored_resume(job_description: str) -> Path | None:
    candidates = build_resume.matching_output_files(OUTPUT_DIR, job_description, "Resume.docx")
    return candidates[0] if candidates else None


def analysis_resume_path(job_description: str, selected_resume: Path) -> Path:
    return latest_tailored_resume(job_description) or selected_resume


def analysis_basis_label(analysis_resume: Path, selected_resume: Path) -> str:
    if analysis_resume == selected_resume:
        return f"Source resume ({source_resume_label(selected_resume)})"
    return f"Tailored resume ({analysis_resume.name})"


def fit_classification(resume_path: Path, job_description: str) -> str:
    with TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        with zipfile.ZipFile(resume_path) as archive:
            archive.extract("word/document.xml", temp_root)
        document_xml = temp_root / "word" / "document.xml"
        document_text = build_resume.visible_text(document_xml)
        alignment_grade = str(build_resume.alignment_score_report(job_description, document_text).get("grade", "")).strip() or None
        status, notes = build_resume.final_fit_audit(
            document_xml,
            job_description,
            alignment_grade=alignment_grade,
        )
    if notes:
        return f"{status} ({len(notes)} note{'s' if len(notes) != 1 else ''})"
    return status


def fit_snapshot_status(job_description: str, analysis_resume: Path, selected_resume: Path, resume_text: str) -> str:
    if analysis_resume == selected_resume:
        report = build_resume.alignment_score_report(job_description, resume_text)
        return str(report.get("grade", "")).strip()
    return fit_classification(analysis_resume, job_description)


def cover_opening_pattern_name(company_name: str, role_title: str, job_description: str) -> str:
    for _label, condition, pattern_fn in build_cover_letter._opening_pattern_options(company_name, role_title, job_description):
        if condition(company_name, role_title, job_description):
            return pattern_fn.__name__
    return "_pyramid_opening"


def lead_story_line(profile: build_resume.JobProblemProfile, job_description: str, resume_text: str) -> str:
    stories = cheat.hero_stories(profile, job_description, resume_text)
    if not stories:
        return "No resume-supported story available. Review source resume evidence before interview prep."
    story = stories[0]
    return f"{story.title}: {story.result}"


def field_value(entry: str, label: str) -> str:
    match = re.search(rf"(?im)^\s*{re.escape(label)}:\s*(.*?)\s*$", entry)
    return match.group(1).strip() if match else ""


def section_value(entry: str, section_label: str) -> str:
    lines = entry.splitlines()
    section_start = ""
    wanted = section_label.lower().rstrip(":") + ":"
    collected: list[str] = []
    for line in lines:
        cleaned = line.strip()
        if not section_start:
            if cleaned.lower() == wanted:
                section_start = wanted
            continue
        if not cleaned:
            break
        collected.append(cleaned)
    return " ".join(collected).strip()


def entry_sort_date(entry: str) -> datetime:
    first_line = entry.splitlines()[0] if entry.splitlines() else ""
    match = re.search(r"(\d{4}-\d{2}-\d{2})(?:\s+(\d{2}:\d{2}:\d{2}))?", first_line)
    if match:
        raw = f"{match.group(1)} {match.group(2) or '00:00:00'}"
        try:
            return datetime.strptime(raw, "%Y-%m-%d %H:%M:%S")
        except ValueError:
            pass
    interview_date = field_value(entry, "Interview date")
    try:
        return datetime.strptime(interview_date, "%Y-%m-%d")
    except ValueError:
        return datetime.min


def find_debrief_entries_for_company(company_name: str, debrief_path: Path) -> list[str]:
    if not debrief_path.exists():
        return []
    text = debrief_path.read_text(encoding="utf-8-sig")
    entries = [
        f"{DEBRIEF_DELIMITER}{entry}".strip()
        for entry in text.split(DEBRIEF_DELIMITER)
        if entry.strip()
    ]
    company_lower = company_name.lower()
    matches = [
        entry
        for entry in entries
        if company_lower in field_value(entry, "Company name").lower()
    ]
    return sorted(matches, key=entry_sort_date, reverse=True)


def repeated_debrief_themes(entries: list[str]) -> list[str]:
    counts: dict[str, int] = {}
    for entry in entries:
        for label in ("Stories that generated follow-up questions", "Unexpected questions"):
            value = section_value(entry, label)
            for part in re.split(r"[.;]\s+|\n", value):
                cleaned = re.sub(r"\s+", " ", part).strip(" -")
                if len(cleaned) < 8 or cleaned.lower() == "none supplied":
                    continue
                counts[cleaned] = counts.get(cleaned, 0) + 1
    return [theme for theme, count in counts.items() if count > 1][:3]


def add_previous_round_intelligence(document: Document, entries: list[str]) -> None:
    add_heading(document, "Previous Round Intelligence")
    if not entries:
        paragraph = document.add_paragraph(
            "No prior interview history found for this company. Run python tasks.py debrief after each round to capture intelligence."
        )
        format_paragraph(paragraph, size=8.5)
        return

    most_recent = entries[0]
    summary = (
        f"Most recent round: {field_value(most_recent, 'Round number') or 'not supplied'} | "
        f"Outcome: {field_value(most_recent, 'Outcome') or 'not supplied'}"
    )
    paragraph = document.add_paragraph(summary)
    format_paragraph(paragraph, size=8.5, bold=True)
    for label in (
        "Stories that generated follow-up questions",
        "Unexpected questions",
        "Specific interviewer language about the role",
    ):
        value = section_value(most_recent, label) or "None supplied."
        paragraph = document.add_paragraph(f"{label}: {value}")
        format_paragraph(paragraph, size=8.2)

    if len(entries) >= 2:
        themes = repeated_debrief_themes(entries)
        paragraph = document.add_paragraph(
            "Pattern Across Rounds: " + ("; ".join(themes) if themes else "No repeated question or story themes detected yet.")
        )
        format_paragraph(paragraph, size=8.2)


def add_business_context_readiness(document: Document, job_description: str, profile: build_resume.JobProblemProfile) -> None:
    context = business_context.extract_business_context(job_description)
    questions = business_context.business_interview_questions(job_description, limit=1)
    add_heading(document, "Business Context Readiness")
    rows = (
        ("Company business", context.business_model or "Research needed"),
        ("Product/customer", ", ".join(item for item in (context.product_or_service, context.customer_type) if item) or "Research needed"),
        ("Role problem", profile.core_problem),
        ("Operational risk", context.operational_complexity or "Confirm in interview"),
        ("Technical/compliance", ", ".join(context.technical_stack + context.compliance_signals) or "Not detected"),
        ("Best question", questions[0].question if questions else "Ask what outcome this hire must improve first."),
    )
    for label, value in rows:
        paragraph = document.add_paragraph(f"{label}: {value}")
        format_paragraph(paragraph, size=8.4)


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0


def format_paragraph(paragraph, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    format_paragraph(paragraph, size=10.5, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))


def add_compact_bullets(document: Document, values: tuple[str, ...] | list[str], empty_text: str = "None") -> None:
    items = values if values else [empty_text]
    for value in items:
        paragraph = document.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Inches(0.14)
        paragraph.paragraph_format.first_line_indent = Inches(-0.14)
        paragraph.add_run(f"- {value}")
        format_paragraph(paragraph, size=8.7)


def build_application_checklist() -> Path:
    job_description = build_resume.validate_inputs()
    company_name = build_resume.extract_output_name(job_description)
    output_target_name = build_resume.extract_output_target_name(job_description)
    role_title = build_resume.extract_job_title(job_description) or "Target Role"
    selected_resume = build_resume.choose_resume(job_description)
    analysis_resume = analysis_resume_path(job_description, selected_resume)
    resume_text = build_resume.docx_visible_text_from_path(analysis_resume)
    profile = build_resume.job_problem_profile(job_description, resume_text)
    readiness = None
    if analysis_resume != selected_resume:
        readiness = build_resume.resume_readiness_for_output(
            job_description,
            analysis_resume,
            source_resume_text=build_resume.docx_visible_text_from_path(selected_resume),
            audit_status=resume_analysis.output_audit_state(analysis_resume),
        )
    fit_status = fit_snapshot_status(job_description, analysis_resume, selected_resume, resume_text)
    alignment_report = build_resume.alignment_score_report(job_description, resume_text)
    gate_decision, _gate_actions = build_resume.alignment_gate_decision(
        int(alignment_report["total_score"]),
        alignment_report,
        job_description,
        company_name,
    )
    story_line = lead_story_line(profile, job_description, resume_text)
    opening_pattern = cover_opening_pattern_name(company_name, role_title, job_description)
    risks = list(build_resume.poor_fit_requirements(job_description, resume_text))[:3]
    debrief_entries = find_debrief_entries_for_company(company_name, DEBRIEF_HISTORY)

    document = Document()
    set_document_style(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{company_name} - {role_title}")
    format_paragraph(title, size=14, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Application Checklist | {date.today().isoformat()}")
    format_paragraph(subtitle, size=9)

    add_heading(document, "Fit Snapshot")
    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ("Detected Lane", "Fit Classification", "Analysis Basis")
    values = (profile.lane_label, fit_status, analysis_basis_label(analysis_resume, selected_resume))
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        for paragraph in cell.paragraphs:
            format_paragraph(paragraph, size=8.5, bold=True)
    for idx, value in enumerate(values):
        cell = table.cell(1, idx)
        cell.text = value
        for paragraph in cell.paragraphs:
            format_paragraph(paragraph, size=8.2)
    gate_paragraph = document.add_paragraph(
        f"Gate Decision: {gate_decision} | Alignment Score: {alignment_report['total_score']}/100"
    )
    format_paragraph(gate_paragraph, size=8.2)
    positioning_statement = build_resume.generate_positioning_statement(profile, job_description)
    positioning_paragraph = document.add_paragraph(f"Positioning Statement: {positioning_statement}")
    format_paragraph(positioning_paragraph, size=8.2)

    add_heading(document, "Top Keywords")
    keyword_table = document.add_table(rows=1, cols=2)
    keyword_table.style = "Table Grid"
    keyword_table.cell(0, 0).text = "Keyword"
    keyword_table.cell(0, 1).text = "Status"
    for keyword in top_keywords(job_description):
        row = keyword_table.add_row().cells
        row[0].text = keyword
        row[1].text = keyword_status(keyword, resume_text)
    for row in keyword_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, size=7.8, bold=row == keyword_table.rows[0])

    add_heading(document, "Evidence Areas")
    paragraph = document.add_paragraph()
    paragraph.add_run("Direct Matches").bold = True
    format_paragraph(paragraph, size=8.8)
    add_compact_bullets(document, profile.direct_matches)
    paragraph = document.add_paragraph()
    paragraph.add_run("Adjacent Matches").bold = True
    format_paragraph(paragraph, size=8.8)
    add_compact_bullets(document, profile.adjacent_matches)
    paragraph = document.add_paragraph()
    paragraph.add_run("Unsupported Requirements").bold = True
    format_paragraph(paragraph, size=8.8)
    add_compact_bullets(document, profile.unsupported_requirements)

    add_heading(document, "Story To Lead With")
    paragraph = document.add_paragraph(story_line)
    format_paragraph(paragraph, size=8.8)
    enforce_prose_quality(story_line, "checklist_narrative", label="Checklist story lead", mode="warn")

    add_heading(document, "Cover Approach")
    paragraph = document.add_paragraph(f"Opening pattern: {opening_pattern}")
    format_paragraph(paragraph, size=8.8)
    enforce_prose_quality(f"Opening pattern: {opening_pattern}", "checklist_narrative", label="Checklist cover approach", mode="warn")

    if readiness and readiness.hard_blockers:
        add_heading(document, "Resume Bridge Gaps")
        add_compact_bullets(
            document,
            [build_resume.resume_gap_summary_line(gap) for gap in readiness.hard_blockers[:3]],
        )

    add_business_context_readiness(document, job_description, profile)

    add_heading(document, "Three Risks")
    add_compact_bullets(document, risks, "No major fit risks detected")

    add_previous_round_intelligence(document, debrief_entries)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output_path = OUTPUT_DIR / f"Christian Estrada - {output_target_name} Application Checklist.docx"
    visible_text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    assert_no_template_leakage(visible_text)
    document.save(output_path)
    print(f"Application checklist created: {output_path}")
    print(f"Prior debrief entries incorporated: {len(debrief_entries)}")
    return output_path


def main() -> int:
    build_application_checklist()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
