#!/usr/bin/env python3
"""Build Christian's Phase 5 career operating plan."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import _bootstrap

_bootstrap.ensure_script_path()
_bootstrap.configure_fresh_pycache()

import interview_intelligence


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"
FONT = "Aptos"
ACCENT = RGBColor(0x1F, 0x3B, 0x5C)
LIGHT_FILL = "EAF1F8"
HEADER_FILL = "1F3B5C"


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _format_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(10)
    for style_name, size in (("Heading 1", 17), ("Heading 2", 12), ("Heading 3", 10.5)):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def _add_title(document: Document, plan: interview_intelligence.CareerOperatingPlan) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Christian Estrada - Career Operating Plan")
    run.font.name = FONT
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.add_run(f"Phase 5 operating rhythm | {plan.plan_date.isoformat()}").font.size = Pt(9)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def _add_role_lists(document: Document, plan: interview_intelligence.CareerOperatingPlan) -> None:
    document.add_heading("Target Roles", level=2)
    document.add_heading("Near-term realistic roles", level=3)
    for role in plan.near_term_roles:
        _add_bullet(document, role)
    document.add_heading("Stretch / north-star role", level=3)
    for role in plan.stretch_roles:
        _add_bullet(document, role)


def _add_modes(document: Document, plan: interview_intelligence.CareerOperatingPlan) -> None:
    document.add_heading("Operating Modes", level=2)
    for mode in plan.modes:
        document.add_heading(mode.name, level=3)
        document.add_paragraph(mode.focus)
        for action in mode.actions:
            _add_bullet(document, action)


def _add_gap_table(
    document: Document,
    title: str,
    gaps: Sequence[interview_intelligence.CareerPlanGap],
) -> None:
    document.add_heading(title, level=2)
    for gap in gaps:
        document.add_heading(gap.label, level=3)
        document.add_paragraph(gap.safe_description)
        _add_bullet(document, f"Study / prep reference: {'; '.join(gap.track_references)}")
        _add_bullet(document, f"Next action: {gap.action}")
    document.add_paragraph()


def _add_checkpoints(document: Document, plan: interview_intelligence.CareerOperatingPlan) -> None:
    document.add_heading("Review Rhythm", level=2)
    for checkpoint in plan.checkpoints:
        _add_bullet(document, checkpoint)
    document.add_heading("Study References", level=2)
    for reference in plan.study_references:
        _add_bullet(document, reference)


def document_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join((*paragraphs, *table_cells))


def render_career_operating_plan(plan: interview_intelligence.CareerOperatingPlan) -> Document:
    document = Document()
    _format_document(document)
    _add_title(document, plan)
    document.add_paragraph(
        "Use this as the operating layer that connects interview readiness, target roles, safe gap language, Study tracks, and review cadence."
    )
    _add_role_lists(document, plan)
    _add_modes(document, plan)
    _add_gap_table(document, "Development Areas To Track", plan.development_gaps)
    _add_gap_table(document, "Stretch-Role Gap Map", plan.stretch_role_gaps)
    _add_checkpoints(document, plan)
    return document


def build_career_operating_plan_docx(*, output_dir: Path = OUTPUT_DIR) -> Path:
    plan = interview_intelligence.build_career_plan()
    output_dir.mkdir(exist_ok=True)
    output = output_dir / f"Christian Estrada - Career Operating Plan {date.today().isoformat()}.docx"
    document = render_career_operating_plan(plan)
    rendered_text = document_text(document)
    interview_intelligence.assert_safe_generated_text(rendered_text, interview_intelligence.load_self_inventory())
    document.save(output)
    print(f"Career operating plan created: {output}")
    return output


def main() -> None:
    build_career_operating_plan_docx()


if __name__ == "__main__":
    main()
