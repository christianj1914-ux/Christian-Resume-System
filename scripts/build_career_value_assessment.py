#!/usr/bin/env python3
"""Build the Four Career Value Driver Assessment."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"
DRIVERS = ("Domain Expertise", "Execution Ability", "Energy and Attitude", "Relationships")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a Four Career Value Driver Assessment.")
    parser.add_argument("--sample", action="store_true", help="Generate a sample/template assessment without interactive input.")
    return parser.parse_args()


def safe_input(prompt: str) -> str:
    try:
        return input(prompt)
    except (KeyboardInterrupt, EOFError):
        raise SystemExit("Assessment capture canceled.")


def prompt_required(label: str) -> str:
    while True:
        value = safe_input(f"{label}: ").strip()
        if value:
            return value
        print(f"{label} is required.")


def collect_data(sample: bool = False) -> list[dict[str, str]]:
    rows = []
    for driver in DRIVERS:
        if sample:
            rows.append(
                {
                    "driver": driver,
                    "rating": "[1-10]",
                    "example": "[specific example]",
                    "gap": "[where this is not fully demonstrated yet]",
                    "action": "[one action to strengthen it]",
                }
            )
            continue
        print()
        print(driver)
        rows.append(
            {
                "driver": driver,
                "rating": prompt_required("Rating 1-10"),
                "example": prompt_required("Specific example supporting the rating"),
                "gap": prompt_required("Where this is not fully demonstrated yet"),
                "action": prompt_required("Action to strengthen it"),
            }
        )
    return rows


def build_assessment(sample: bool = False) -> Path:
    rows = collect_data(sample)
    OUTPUT_DIR.mkdir(exist_ok=True)
    suffix = " TEMPLATE" if sample else ""
    output = OUTPUT_DIR / f"Christian Estrada - Career Value Assessment{suffix} {datetime.now().strftime('%Y-%m-%d')}.docx"
    doc = Document()
    doc.add_heading("Four Career Value Driver Assessment", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for index, header in enumerate(("Driver", "Rating", "Example", "Current Gap", "Action")):
        table.cell(0, index).text = header
    for row_data in rows:
        row = table.add_row().cells
        row[0].text = row_data["driver"]
        row[1].text = row_data["rating"]
        row[2].text = row_data["example"]
        row[3].text = row_data["gap"]
        row[4].text = row_data["action"]
    doc.save(output)
    print(f"Career value assessment created: {output}")
    return output


def main() -> None:
    args = parse_args()
    build_assessment(args.sample)


if __name__ == "__main__":
    main()
