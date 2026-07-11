#!/usr/bin/env python3
"""
Build Christian Estrada's tailored DOCX cover letter.

Rules enforced from AGENTS.md:
- use the generated resume output and job description as source material
- save Word documents only to /output
- never create placeholders or PDFs
- never use LinkedIn page content as source material
- keep the letter one page, concise, and resume-supported
"""

from __future__ import annotations

import argparse
import _bootstrap

_bootstrap.ensure_script_path()

import contextlib
import io
import json
import os
import re
import sys
from collections import Counter
from dataclasses import asdict, dataclass, replace
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_resume
import business_context
import evidence_engine
import interview_context
import job_context_archive
import proof_text
import prose_engine
import question_prep
import resume_analysis
from build_resume import CONTACT_EMAIL
from config.language_rules import CLICHE_PATTERNS, PLACEHOLDER_PATTERNS, PROMPT_LEAK_PATTERNS
from config.paths import COMPANY_RESEARCH, INTERVIEW_NOTES, JOB_DESCRIPTION, JOBS_DIR, OUTPUT_DIR, PROJECT_ROOT, SCRATCH_DIR
from modules.employer_playbooks import consulting_bigfour
import render_checks
from text_safety import fix_indefinite_articles, sentence_splice_issues, with_indefinite_article as shared_with_indefinite_article
from utils import assert_no_template_leakage, clean_source_text, debug_print, enforce_prose_quality, fail, join_answer_sentences, optional_text, prose_quality_report, read_text


RESUME_FONT = "Carlito"
NAME_BLUE = RGBColor(31, 78, 121)
LINKEDIN_URL = build_resume.LINKEDIN_URL
VISIBLE_LINKEDIN_URL = re.sub(r"^https?://(?:www\.)?", "", LINKEDIN_URL, flags=re.I).rstrip("/")
MAX_BODY_WORDS = 540
STANDARD_MIN_LETTER_WORDS = 80
STANDARD_MAX_LETTER_WORDS = 170
MAX_STANDARD_BODY_SENTENCES = 6
LONG_MIN_LETTER_WORDS = 180
LONG_MAX_LETTER_WORDS = 240
HEADER_LINE_SPACING = 1.15
BODY_LINE_SPACING = 1.15
BODY_FONT_SIZE = 12
NAME_FONT_SIZE = 22
BULLET_LEFT_INDENT_INCHES = 0.50
BULLET_HANGING_INDENT_INCHES = -0.25
BULLET_TEXT_TAB_INCHES = 0.50
HEADER_TO_DATE_GAP_PT = 18
BODY_PARAGRAPH_GAP_PT = 0
SIGNATURE_GAP_PT = 0
STANDARD_COVER_MODE = "standard"
LONG_COVER_MODE = "long"
DEFAULT_COVER_MODE = STANDARD_COVER_MODE
LEGACY_DEFAULT_COVER_MODE = "default"
LEGACY_CONCISE_COVER_MODE = "concise"
COVER_TRACE_DIR = SCRATCH_DIR / "cover_letter_traces"


def fail(message: str) -> None:
    print(f"ERROR: {message}", file=sys.stderr)
    raise SystemExit(message)

FIRM_PROFILE_REGISTRY: dict[str, dict[str, object]] = {
    "bain": {
        "names": (
            "bain",
            "bain & company",
            "bain and company",
        ),
        "cover_style": "bain",
        "proof_emphasis": "client-facing consulting delivery, executive alignment, measurable outcomes",
        "support_style": "learning",
    },
    "state_farm": {
        "names": (
            "state farm",
            "state farm insurance",
            "state farm mutual automobile insurance company",
        ),
        "cover_style": "state_farm_process",
        "proof_emphasis": "claims process improvement, Lean Six Sigma, operational data, service quality, and customer experience",
        "role_title_patterns": (r"\bprocess engineer\b",),
    },
}

COVER_PROFILE_RULES: tuple[dict[str, object], ...] = (
    {
        "key": "bain",
        "registered_key": "bain",
        "opening_method": "",
        "proof_mode": "bain",
    },
    {
        "key": "state_farm",
        "registered_key": "state_farm",
        "opening_method": (
            "My fit is practical: I map the current state, use data to separate root cause from symptoms, test the fix, "
            "and bring the people doing the work along so the improvement lasts."
        ),
        "proof_mode": "state_farm",
    },
    {
        "key": "construction_engineering",
        "employer_context_key": "construction_engineering",
        "opening_method": (
            "I map the workflow, protect the financial calendar, configure to the client's real constraints, "
            "and keep adoption moving after go-live."
        ),
        "proof_mode": "construction_engineering",
    },
)

_FIRM_NAME_ABBREVIATIONS = {
    "co": "company",
    "corp": "corporation",
    "intl": "international",
    "mfg": "manufacturing",
    "svcs": "services",
    "tech": "technology",
}
_FIRM_NAME_STOPWORDS = {
    "and",
    "the",
    "company",
    "co",
    "corporation",
    "corp",
    "inc",
    "incorporated",
    "llc",
    "llp",
    "lp",
    "ltd",
    "limited",
    "plc",
}


def _normalized_firm_tokens(value: str) -> tuple[str, ...]:
    cleaned = value.lower().replace("&", " and ")
    cleaned = re.sub(r"'s\b", "", cleaned)
    cleaned = re.sub(r"[^a-z0-9]+", " ", cleaned)
    tokens = []
    for token in cleaned.split():
        expanded = _FIRM_NAME_ABBREVIATIONS.get(token, token)
        if expanded not in _FIRM_NAME_STOPWORDS:
            tokens.append(expanded)
    return tuple(tokens)


def matches_firm_name(company_name: str, *names: str) -> bool:
    """Match firm names after punctuation stripping, abbreviation handling, and partial token comparison."""
    company_tokens = set(_normalized_firm_tokens(company_name))
    if not company_tokens:
        return False
    for name in names:
        name_tokens = set(_normalized_firm_tokens(name))
        if name_tokens and (name_tokens <= company_tokens or company_tokens <= name_tokens):
            return True
    return False


def _registered_firm_profile(
    company_name: str,
    role_title: str = "",
    job_description: str = "",
    company_research: str = "",
    interview_notes: str = "",
) -> tuple[str, dict[str, object]] | None:
    for key, profile in FIRM_PROFILE_REGISTRY.items():
        match_condition = profile.get("match_condition")
        if callable(match_condition):
            if not match_condition(company_name, role_title, job_description, company_research, interview_notes):
                continue
        else:
            names = tuple(str(name) for name in profile.get("names", (key,)))
            if not matches_firm_name(company_name, *names):
                continue
        role_patterns = tuple(str(pattern) for pattern in profile.get("role_title_patterns", ()))
        if role_patterns and not any(re.search(pattern, role_title, re.I) for pattern in role_patterns):
            continue
        return key, profile
    return None


def _uses_firm_profile(company_name: str, profile_key: str, role_title: str = "", job_description: str = "") -> bool:
    match = _registered_firm_profile(company_name, role_title, job_description)
    return bool(match and match[0] == profile_key)


def active_cover_profile(
    company_name: str,
    role_title: str,
    job_description: str,
) -> dict[str, object] | None:
    registered = _registered_firm_profile(company_name, role_title, job_description)
    employer_context = build_resume.primary_employer_context(job_description)
    employer_context_key = str(employer_context.get("key", "")) if employer_context else ""
    for rule in COVER_PROFILE_RULES:
        registered_key = str(rule.get("registered_key", "")).strip()
        if registered_key and registered and registered[0] == registered_key:
            return rule
        context_key = str(rule.get("employer_context_key", "")).strip()
        if context_key and employer_context_key == context_key:
            return rule
    return None


def repair_cover_paragraph(text: str, label: str) -> str:
    cleaned = re.sub(r"\s+", " ", fix_indefinite_articles(text or "")).strip()
    issues = sentence_splice_issues(cleaned)
    if issues:
        fail(f"{label} contains unresolved grammar fragments: {', '.join(issues)}")
    return cleaned


def relevant_supplied_context(company_name: str, role_title: str = "") -> str:
    context = interview_context.load_company_context(
        JOBS_DIR,
        company_name,
        role_title,
        company_research_path=COMPANY_RESEARCH,
        global_interview_notes_path=INTERVIEW_NOTES,
    )
    return context.supplied_context


def is_education_assessment_context(job_description: str) -> bool:
    context = business_context.extract_business_context(job_description)
    if context.industry == "education / assessment":
        return True
    return bool(
        re.search(
            r"\b(?:school assessment|assessment item|assessment and learning|measurement and learning|"
            r"learner-facing|instructional|k-12|psychometric|constructed-response|technology-enhanced items|tei)\b",
            job_description,
            re.I,
        )
    )


def concrete_company_context_line(text: str) -> str:
    concrete_signals = (
        "launch", "rollout", "migration", "cutover", "integration", "product", "platform", "jira", "azure",
        "salesforce", "sap", "coupa", "power bi", "sql", "ticket", "issue", "workflow", "support",
        "claims", "customer", "adoption", "reporting", "analytics",
    )
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line or re.match(r"^(POST-INTERVIEW NOTE|INTERVIEW NOTES FILE:)", line, re.I):
            continue
        if len(line.split()) < 5:
            continue
        if any(signal in line.lower() for signal in concrete_signals):
            return line
    return ""


def normalize_company_context_detail(line: str, company_name: str) -> str:
    detail = re.sub(r"\s+", " ", line).strip().rstrip(".")
    detail = re.sub(rf"^{re.escape(company_name)}\s*(?:-\s*[^,.:]+)?\s*,\s*", "", detail, flags=re.I)
    replacements = (
        (r"^we would be supporting\b", "the team would support"),
        (r"^we would support\b", "the team would support"),
        (r"^we would be using\b", "the team would use"),
        (r"^we would use\b", "the team would use"),
        (r"^internally using\b", "the team would use"),
        (r"^this role would support\b", "the role would support"),
        (r"^this role supports\b", "the role supports"),
    )
    for pattern, replacement in replacements:
        detail = re.sub(pattern, replacement, detail, flags=re.I)
    if detail and detail[0].isupper():
        detail = detail[0].lower() + detail[1:]
    return detail


def company_specific_context_sentence(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    company_context_text: str = "",
) -> str:
    detail_line = concrete_company_context_line(company_context_text)
    if not detail_line:
        return ""
    detail = normalize_company_context_detail(detail_line, company_name)
    if not detail:
        return ""
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_key = effective_lane_key(role_title, job_description, profile)
    focus_by_lane = {
        "presales_solution": "solution judgment",
        "customer_success": "adoption and issue-management discipline",
        "change_enablement": "adoption discipline",
        "analytics_operations": "data and decision discipline",
        "implementation_delivery": "implementation discipline",
        "process_improvement": "process-improvement discipline",
        "corporate_strategy": "structured decision discipline",
    }
    focus = focus_by_lane.get(lane_key, "cross-functional delivery discipline")
    return f"The detail that {detail} makes the work concrete and points to the kind of {focus} the role needs."



@dataclass(frozen=True)
class CoverLetterResult:
    company_name: str
    role_title: str
    resume_docx: Path
    output_docx: Path
    bullets_used: int
    audit_status: str
    specificity_warnings: list[str]
    cover_warnings: list[str]
    preflight_warnings: list[str]
    mode: str


@dataclass(frozen=True)
class CoverLetterSignals:
    company_mission: str
    role_core_function: str
    top_accomplishment: str
    fit_bridge: str
    jd_skill_terms: tuple[str, ...]
    ambiguity_process: str
    jd_test_environments: tuple[str, ...]
    communication_metric: str
    partner_functions: tuple[str, ...]
    jd_pain_area: str


@dataclass(frozen=True)
class CoverLetterDraft:
    salutation: str
    body_paragraphs: tuple[str, ...]
    signals: CoverLetterSignals
    paragraph_shape: int
    proof_mapped_terms: tuple[str, ...]
    close_mapped_terms: tuple[str, ...]
    mode: str
    repair_rule_ids: tuple[str, ...] = ()


@dataclass(frozen=True)
class CoverLetterPlan:
    company_name: str
    role_title: str
    prose_role_title: str
    job_description: str
    mode: str
    draft_style: str
    lane_key: str
    sections: tuple[tuple[str, str], ...]
    signals: CoverLetterSignals
    salutation: str
    body_paragraphs: tuple[str, ...]
    opening: str
    proof: str
    bridge: str
    workflow: str
    communication: str
    closing: str
    proof_mapped_terms: tuple[str, ...]
    close_mapped_terms: tuple[str, ...]
    warnings: tuple[str, ...]
    blockers: tuple[str, ...]
    selection_debug: dict[str, object]


@dataclass(frozen=True)
class EvidenceBullet:
    text: str
    evidence: tuple[str, ...]
    signals: tuple[str, ...]


@dataclass(frozen=True)
class CoverResponseBank:
    active_responses: tuple[question_prep.QualificationsResponse, ...]
    company_interest: str
    relevant_experience: str
    unique_qualifications: str
    communication: str
    proof_bullets: tuple[str, ...]


@dataclass(frozen=True)
class CoverSentenceCandidate:
    text: str
    source: str


def normalize_cover_mode(mode: str) -> str:
    normalized = (mode or DEFAULT_COVER_MODE).strip().lower()
    if normalized in {STANDARD_COVER_MODE, LEGACY_CONCISE_COVER_MODE, "short"}:
        return STANDARD_COVER_MODE
    if normalized in {LONG_COVER_MODE, LEGACY_DEFAULT_COVER_MODE, "full", "extended"}:
        return LONG_COVER_MODE
    if normalized not in {STANDARD_COVER_MODE, LONG_COVER_MODE}:
        fail(f"unsupported cover-letter mode: {mode}")
    return normalized


def cover_letter_word_range(mode: str) -> tuple[int, int]:
    normalized = normalize_cover_mode(mode)
    if normalized == LONG_COVER_MODE:
        return LONG_MIN_LETTER_WORDS, LONG_MAX_LETTER_WORDS
    return STANDARD_MIN_LETTER_WORDS, STANDARD_MAX_LETTER_WORDS


def sentence_list(text: str) -> list[str]:
    return build_resume.summary_sentences(re.sub(r"\s+", " ", text).strip())


def limit_sentences(text: str, max_sentences: int) -> str:
    if max_sentences <= 0:
        return ""
    sentences = sentence_list(text)
    if not sentences:
        return re.sub(r"\s+", " ", text).strip()
    return " ".join(sentences[:max_sentences]).strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w']+\b", text))


def safe_sentence(text: str) -> str:
    cleaned = normalize_cover_dash_punctuation(re.sub(r"\s+", " ", text or "")).strip()
    if not cleaned:
        return ""
    cleaned = re.sub(r"^That includes\b", "My experience includes", cleaned, flags=re.I)
    cleaned = re.sub(r"^That background\b", "My background", cleaned, flags=re.I)
    cleaned = re.sub(r"^That experience\b", "My experience", cleaned, flags=re.I)
    cleaned = re.sub(r"^That work\b", "This work", cleaned, flags=re.I)
    cleaned = re.sub(r"^I bring\s+i bring\b", "I bring", cleaned, flags=re.I)
    if cleaned[-1] not in ".!?":
        cleaned += "."
    return cleaned


def response_sentences(text: str) -> list[str]:
    return [safe_sentence(sentence) for sentence in sentence_list(text) if safe_sentence(sentence)]


def response_sentence_for_category(
    responses: tuple[question_prep.QualificationsResponse, ...],
    category: str,
    *,
    sentence_index: int = 0,
) -> str:
    for response in responses:
        if question_prep.question_category(response.prompt) != category:
            continue
        sentences = response_sentences(response.answer)
        if sentences:
            return sentences[min(sentence_index, len(sentences) - 1)]
    return ""


def dedupe_sentences(sentences: list[str]) -> list[str]:
    deduped: list[str] = []
    seen: set[str] = set()
    for sentence in sentences:
        normalized = re.sub(r"[^a-z0-9]+", " ", sentence.lower()).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(sentence)
    return deduped


def short_company_interest_sentence(
    bank: CoverResponseBank,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
) -> str:
    answer = bank.company_interest or question_prep.cover_company_interest_answer(job_description, resume_text)
    sentences = response_sentences(answer)
    if sentences:
        first = sentences[0]
        first = re.sub(
            rf"^I am interested in joining\s+{re.escape(company_name)}\s+because\s+",
            "",
            first,
            flags=re.I,
        )
        if re.match(
            rf"^I am interested in the\s+.+?\s+role at\s+{re.escape(company_name)}\b",
            first,
            re.I,
        ):
            return safe_sentence(first)
        first = re.sub(
            rf"^I am interested in the\s+.+?\s+role at\s+{re.escape(company_name)}\s+because\s+",
            "",
            first,
            flags=re.I,
        )
        return safe_sentence(f"I am interested in the {safe_role_title(role_title)} role at {safe_company_name(company_name)} because {first[:1].lower()}{first[1:]}")
    return safe_sentence(
        f"I am interested in the {safe_role_title(role_title)} role at {safe_company_name(company_name)} because it sits close to {extract_pain_area(job_description, effective_lane_key(role_title, job_description, build_resume.job_problem_profile(job_description, resume_text)), (), ())}."
    )


def build_cover_response_bank(
    job_description: str,
    resume_text: str,
    *,
    company_name: str = "",
    role_title: str = "",
    application_responses: tuple[question_prep.QualificationsResponse, ...] = (),
) -> CoverResponseBank:
    _, snapshot, _approved_resume_text = question_prep.selected_resume_snapshot(job_description)
    responses = application_responses or ()
    categorized = {
        question_prep.question_category(response.prompt): response.answer
        for response in responses
    }
    evidence_items = safe_selected_evidence_items_ordered(job_description, resume_text)
    proof_bullets = tuple(
        safe_sentence(proof_text.sanitize_proof_sentence(item.text, max_words=24))
        for item in evidence_items
    )
    return CoverResponseBank(
        active_responses=responses,
        company_interest=categorized.get("company_interest") or question_prep.cover_company_interest_answer(
            job_description,
            resume_text,
            company_name=company_name,
            role_title=role_title,
        ),
        relevant_experience=categorized.get("relevant_experience") or question_prep.cover_relevant_experience_answer(job_description, snapshot, resume_text),
        unique_qualifications=categorized.get("unique_qualifications") or question_prep.cover_unique_qualifications_answer(job_description, resume_text),
        communication=categorized.get("communication") or question_prep.cover_communication_answer(),
        proof_bullets=tuple(sentence for sentence in proof_bullets if sentence),
    )


def dedupe_cover_sentence_candidates(sentences: list[str], *, limit: int = 3) -> list[str]:
    deduped: list[str] = []
    for sentence in sentences:
        if any(cover_sentences_near_duplicate(sentence, existing) for existing in deduped):
            continue
        deduped.append(sentence)
        if len(deduped) >= limit:
            break
    return deduped


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def sentence_starts_with_lowercase_fragment(text: str) -> bool:
    normalized = normalize_spaces(text)
    return bool(normalized) and normalized[0].islower()


def cover_lane_terms(lane_key: str) -> tuple[str, ...]:
    return {
        "presales_solution": ("discovery", "solution", "buyer", "demo", "requirements", "stakeholder", "presentation"),
        "customer_success": ("account", "retention", "adoption", "customer", "qbr", "renewal", "risk"),
        "implementation_delivery": ("implementation", "migration", "go-live", "testing", "training", "integration", "delivery"),
        "change_enablement": ("change", "adoption", "training", "stakeholder", "launch", "reinforcement"),
        "analytics_operations": ("analytics", "reporting", "dashboard", "data", "measurement", "sql", "workflow"),
        "corporate_strategy": ("analysis", "recommendation", "decision", "strategy", "risk", "revenue", "director", "vp", "client"),
    }.get(lane_key, ("workflow", "stakeholder", "delivery"))


def cover_lane_term_hits(text: str, lane_key: str) -> int:
    lowered = text.lower()
    return sum(1 for term in cover_lane_terms(lane_key) if term in lowered)


def non_target_implementation_penalty(text: str, lane_key: str) -> int:
    if lane_key == "implementation_delivery":
        return 0
    lowered = text.lower()
    implementation_terms = ("migration", "go-live", "uat", "testing", "integration", "warehouse", "cutover")
    implementation_hits = sum(1 for term in implementation_terms if term in lowered)
    if implementation_hits == 0:
        return 0
    if cover_lane_term_hits(text, lane_key) > 0:
        return max(0, implementation_hits - 1)
    return implementation_hits + 1


def opening_support_candidates(bank: CoverResponseBank) -> list[CoverSentenceCandidate]:
    candidates: list[CoverSentenceCandidate] = []
    for sentence in response_sentences(bank.communication):
        candidates.append(CoverSentenceCandidate(text=sentence, source="communication"))
    for sentence in response_sentences(bank.unique_qualifications):
        candidates.append(CoverSentenceCandidate(text=sentence, source="unique_qualifications"))
    for sentence in response_sentences(bank.relevant_experience):
        if re.match(r"^I bring approximately\b", sentence, re.I):
            continue
        candidates.append(CoverSentenceCandidate(text=sentence, source="relevant_experience"))
    return candidates


def opening_support_fallback_sentence(lane_key: str) -> str:
    fallback_by_lane = {
        "presales_solution": "I have spent much of my career making complex system work easier to scope, explain, and hand off cleanly.",
        "customer_success": "I have spent much of my career helping customer-facing work stay clear, accountable, and usable once the pressure rises.",
        "implementation_delivery": "I have spent much of my career keeping implementation work clear across scope, testing, training, and go-live readiness.",
        "change_enablement": "I have spent much of my career helping teams move from stakeholder alignment to real adoption and day-to-day follow-through.",
        "analytics_operations": "I have spent much of my career turning messy operational questions into reporting, workflow clarity, and decisions people can trust.",
        "corporate_strategy": "I have spent much of my career turning half-defined business questions into structured analysis, stakeholder alignment, and practical next steps.",
    }
    return safe_sentence(fallback_by_lane.get(lane_key, fallback_by_lane["implementation_delivery"]))


def select_opening_support_sentence(
    candidates: list[CoverSentenceCandidate],
    *,
    opening_sentence: str,
    company_name: str,
    role_title: str,
    job_description: str,
    lane_key: str,
) -> tuple[str, list[dict[str, object]]]:
    debug_rows: list[dict[str, object]] = []
    ranked: list[tuple[int, int, str, dict[str, object]]] = []
    for index, candidate in enumerate(candidates):
        cleaned = proof_text.sanitize_proof_sentence(
            candidate.text,
            max_words=24,
            artifact="cover_letter_opening",
        )
        row = {
            "purpose": "opening_support",
            "source": candidate.source,
            "original_text": candidate.text,
            "text": cleaned or candidate.text,
            "selected": False,
            "score": 0,
            "rejection_reason": "",
        }
        if not cleaned:
            row["rejection_reason"] = "sanitized_empty"
        elif sentence_starts_with_lowercase_fragment(cleaned):
            row["rejection_reason"] = "lowercase_fragment"
        elif cover_sentences_near_duplicate(cleaned, opening_sentence):
            row["rejection_reason"] = "duplicate_with_opening"
        elif cover_sentence_is_generic(cleaned, company_name, role_title, job_description):
            row["rejection_reason"] = "generic_opening_support"
        elif re.match(r"^I bring approximately\b", cleaned, re.I):
            row["rejection_reason"] = "generic_experience_summary"
        else:
            score = (
                cover_sentence_score(cleaned, company_name, role_title, job_description)
                + (cover_lane_term_hits(cleaned, lane_key) * 3)
                + (3 if paragraph_has_fast_proof(cleaned) else 0)
            )
            row["score"] = score
            ranked.append((score, -index, cleaned, row))
        debug_rows.append(row)
    ranked.sort(reverse=True)
    if ranked:
        selected_text = ranked[0][2]
        ranked[0][3]["selected"] = True
        return selected_text, debug_rows
    fallback = opening_support_fallback_sentence(lane_key)
    debug_rows.append(
        {
            "purpose": "opening_support",
            "source": "fallback",
            "original_text": fallback,
            "text": fallback,
            "selected": True,
            "score": 0,
            "rejection_reason": "fallback_used",
        }
    )
    return fallback, debug_rows


def cover_safe_active_response_sentences(
    response: question_prep.QualificationsResponse,
    *,
    lane_key: str,
    job_description: str,
) -> list[str]:
    category = question_prep.question_category(response.prompt)
    if category == "customer_profile":
        template_by_lane = {
            "presales_solution": "I have worked across B2B manufacturing and enterprise software customers through 80+ client engagements where discovery, stakeholder communication, and solution handoff all had to stay clear.",
            "customer_success": "I have managed customer-facing work where executive relationships, adoption follow-through, and workflow clarity all mattered to the account outcome.",
            "implementation_delivery": "I have coordinated multi-stakeholder customer environments where operations, finance, and technical teams all needed a clear path from scope through go-live.",
            "change_enablement": "I have coordinated operating environments where stakeholder alignment, training, and adoption had to hold across multiple teams and workflows.",
            "analytics_operations": "I have built operational reporting and decision-support tools leaders could actually use across the business.",
            "corporate_strategy": "I have structured complex business problems so cross-functional stakeholders could see the tradeoffs and move toward a practical next step.",
        }
        return [safe_sentence(template_by_lane.get(lane_key, template_by_lane["implementation_delivery"]))]
    if category == "implementation_volume":
        template_by_lane = {
            "presales_solution": "I have handled concurrent workstreams across 80+ client engagements and five-site operations where discovery, stakeholder communication, and handoff discipline all had to stay visible at the same time.",
            "implementation_delivery": "I have coordinated concurrent implementation work across 80+ client engagements and five-site operations where scope, testing, and go-live readiness all had to stay visible at the same time.",
        }
        return [safe_sentence(template_by_lane.get(lane_key, template_by_lane["implementation_delivery"]))]
    if category == "implementation_success":
        template_by_lane = {
            "presales_solution": "I measure success by whether discovery holds up in delivery, stakeholders stay aligned, and the handoff leaves no hidden workflow gaps.",
            "implementation_delivery": "I measure success by go-live readiness, stakeholder confidence, user adoption, and whether the workflow holds up after handoff.",
        }
        return [safe_sentence(template_by_lane.get(lane_key, template_by_lane["implementation_delivery"]))]
    if category == "ai_passion":
        if re.search(r"\b(?:ai|artificial intelligence|automation|agent|model)\b", job_description, re.I):
            return [safe_sentence("I use AI to speed research, structure documentation, and improve analysis while keeping judgment and accountability with the final decision.")]
        return []
    if category == "ambiguity_delivery":
        template_by_lane = {
            "presales_solution": "I have turned ambiguous customer questions into clearer scope, next steps, and handoff plans across complex enterprise work.",
            "corporate_strategy": "I have turned ambiguous cross-functional work into clearer structure, tradeoffs, and decision paths teams could actually use.",
            "implementation_delivery": "I have turned ambiguous implementation work into clearer scope, operating rhythm, and launch readiness across complex environments.",
        }
        return [safe_sentence(template_by_lane.get(lane_key, template_by_lane["implementation_delivery"]))]
    if category == "expectation_gap":
        template_by_lane = {
            "presales_solution": "I have handled expectation gaps by explaining technical tradeoffs clearly and giving stakeholders a workable path forward before trust eroded.",
            "customer_success": "I have protected customer trust by resetting expectations early, explaining tradeoffs clearly, and keeping the next step workable.",
            "implementation_delivery": "I have protected delivery trust by surfacing constraints early, documenting tradeoffs, and steering stakeholders toward workable alternatives.",
        }
        return [safe_sentence(template_by_lane.get(lane_key, template_by_lane["implementation_delivery"]))]
    if category == "generic_bridge":
        return []
    return response_sentences(response.answer)


def standard_cover_proof_candidates(
    bank: CoverResponseBank,
    *,
    signals: CoverLetterSignals,
    lane_key: str,
    job_description: str,
) -> list[CoverSentenceCandidate]:
    candidates: list[CoverSentenceCandidate] = []
    if signals.top_accomplishment:
        candidates.append(CoverSentenceCandidate(text=signals.top_accomplishment, source="signals.top_accomplishment"))
    for sentence in bank.proof_bullets:
        candidates.append(CoverSentenceCandidate(text=sentence, source="proof_bullet"))
    for sentence in response_sentences(bank.relevant_experience):
        if re.match(r"^I bring approximately\b", sentence, re.I):
            continue
        candidates.append(CoverSentenceCandidate(text=sentence, source="relevant_experience"))
    for sentence in response_sentences(bank.unique_qualifications):
        candidates.append(CoverSentenceCandidate(text=sentence, source="unique_qualifications"))
    for sentence in response_sentences(bank.communication):
        candidates.append(CoverSentenceCandidate(text=sentence, source="communication"))
    for response in bank.active_responses:
        category = question_prep.question_category(response.prompt)
        if category in {"relevant_experience", "unique_qualifications", "communication", "company_interest"}:
            continue
        for sentence in cover_safe_active_response_sentences(
            response,
            lane_key=lane_key,
            job_description=job_description,
        ):
            candidates.append(CoverSentenceCandidate(text=sentence, source=f"application_response:{category}"))
    if lane_key == "presales_solution":
        presales_bonus = "I also bring pre-sales discovery and product demonstration experience grounded in implementation reality."
        candidates.append(CoverSentenceCandidate(text=presales_bonus, source="presales_bonus"))
    return candidates


def selected_response_bank_proof_rows(proof_debug: list[dict[str, object]]) -> list[dict[str, object]]:
    response_sources = {"relevant_experience", "unique_qualifications", "communication"}
    return [
        row
        for row in proof_debug
        if row.get("selected") and row.get("source") in response_sources
    ]


def strongest_response_bank_proof_row(
    proof_debug: list[dict[str, object]],
    *,
    opening_sentences: tuple[str, ...],
    selected_proof_sentences: list[str],
) -> dict[str, object] | None:
    response_sources = {"relevant_experience", "unique_qualifications", "communication"}
    source_priority = {
        "relevant_experience": 3,
        "communication": 2,
        "unique_qualifications": 1,
    }
    candidates: list[tuple[int, int, int, str, dict[str, object]]] = []
    for row in proof_debug:
        source = str(row.get("source") or "")
        if source not in response_sources:
            continue
        if row.get("selected"):
            continue
        if row.get("rejection_reason"):
            continue
        cleaned = str(row.get("text") or "").strip()
        if not cleaned:
            continue
        if any(
            cover_sentences_near_duplicate(cleaned, existing)
            for existing in [*opening_sentences, *selected_proof_sentences]
        ):
            continue
        candidates.append(
            (
                1 if re.search(r"\d|\$|%", cleaned) else 0,
                source_priority.get(source, 0),
                int(row.get("score") or 0),
                cleaned,
                row,
            )
        )
    if not candidates:
        return None
    candidates.sort(reverse=True)
    return candidates[0][4]


def preserve_non_response_proof_sentence(
    proof_debug: list[dict[str, object]],
    selected_proof_sentences: list[str],
) -> str:
    for preferred_source in ("signals.top_accomplishment", "proof_bullet"):
        for row in proof_debug:
            if row.get("selected") and row.get("source") == preferred_source:
                cleaned = str(row.get("text") or "").strip()
                if cleaned:
                    return cleaned
    for row in proof_debug:
        if row.get("selected"):
            cleaned = str(row.get("text") or "").strip()
            if cleaned:
                return cleaned
    return selected_proof_sentences[0] if selected_proof_sentences else ""


def cover_ownership_language_score(text: str) -> int:
    lowered = text.lower()
    score = 0
    strong_phrases = (
        "i owned",
        "i led",
        "i was responsible for",
        "i coordinated",
        "i drove",
        "i built",
        "i delivered",
        "i stabilized",
    )
    soft_phrases = (
        "i supported",
        "i helped",
        "i assisted",
        "i was involved in",
        "we worked on",
    )
    score += sum(4 for phrase in strong_phrases if phrase in lowered)
    score -= sum(3 for phrase in soft_phrases if phrase in lowered)
    return score


def proof_selection_score(text: str, lane_key: str, job_description: str) -> int:
    metric_hits, action_hits, keyword_hits, negative_length_penalty = proof_text.proof_candidate_score(text, job_description)
    lane_bonus = cover_lane_term_hits(text, lane_key) * 5
    proof_bonus = 6 if paragraph_has_fast_proof(text) else 0
    company_bonus = 2 if re.search(r"\b(?:East West Manufacturing|Aptean|Home Depot)\b", text, re.I) else 0
    ownership_bonus = cover_ownership_language_score(text)
    penalty = non_target_implementation_penalty(text, lane_key) * 3
    return (metric_hits * 8) + (action_hits * 4) + (keyword_hits * 3) + lane_bonus + proof_bonus + company_bonus + ownership_bonus + negative_length_penalty - penalty
 

def proof_fallback_sentences(lane_key: str) -> list[str]:
    fallback_by_lane = {
        "presales_solution": "I have led 60+ executive workshops and QBRs, supported 80+ client engagements, and kept discovery, stakeholder communication, and handoff quality grounded in delivery reality.",
        "customer_success": "I have led 60+ executive workshops and QBRs and helped stabilize $1M+ in at-risk annual revenue while keeping adoption, executive communication, and escalation clarity visible.",
        "implementation_delivery": "I have owned five-site operations and built 200+ reporting tools that made status, workflow gaps, and next-step decisions easier to track through delivery.",
        "change_enablement": "I have led role-based training and built rollout materials that kept stakeholder alignment and adoption visible through change.",
        "analytics_operations": "I have built 200+ reporting tools across five-site operations where leaders needed clearer workflow visibility and decision support.",
        "corporate_strategy": "I have led 60+ executive workshops and built 200+ reporting tools that turned messy operating questions into clearer decisions and practical next steps.",
    }
    return [safe_sentence(fallback_by_lane.get(lane_key, fallback_by_lane["implementation_delivery"]))]


def select_proof_sentences(
    candidates: list[CoverSentenceCandidate],
    *,
    opening_sentences: tuple[str, ...],
    company_name: str,
    role_title: str,
    job_description: str,
    lane_key: str,
    limit: int = 2,
) -> tuple[list[str], list[dict[str, object]]]:
    debug_rows: list[dict[str, object]] = []
    ranked: list[tuple[int, int, str, dict[str, object]]] = []
    for index, candidate in enumerate(candidates):
        cleaned = proof_text.sanitize_proof_sentence(
            candidate.text,
            max_words=32,
            artifact="cover_letter_proof",
        )
        row = {
            "purpose": "proof",
            "source": candidate.source,
            "original_text": candidate.text,
            "text": cleaned or candidate.text,
            "selected": False,
            "score": 0,
            "rejection_reason": "",
        }
        if not cleaned:
            row["rejection_reason"] = "sanitized_empty"
        elif sentence_starts_with_lowercase_fragment(cleaned):
            row["rejection_reason"] = "lowercase_fragment"
        elif re.match(r"^I bring (?:approximately|a mix of)\b", cleaned, re.I):
            row["rejection_reason"] = "generic_experience_summary"
        elif re.match(r"^My experience includes\b", cleaned, re.I):
            row["rejection_reason"] = "generic_proof"
        elif any(cover_sentences_near_duplicate(cleaned, sentence) for sentence in opening_sentences):
            row["rejection_reason"] = "duplicate_with_opening"
        elif cover_sentence_is_generic(cleaned, company_name, role_title, job_description) and not paragraph_has_fast_proof(cleaned):
            row["rejection_reason"] = "generic_proof"
        else:
            candidate_prose = prose_quality_report(cleaned, "cover_letter_proof")
            row["prose_failures"] = list(candidate_prose.get("failures", ()))
            if row["prose_failures"]:
                row["rejection_reason"] = "prose_failure"
                debug_rows.append(row)
                continue
            source_bonus = {
                "signals.top_accomplishment": 10,
                "proof_bullet": 8,
                "relevant_experience": 5,
                "unique_qualifications": 4,
                "communication": 3,
            }.get(candidate.source, 0)
            score = proof_selection_score(cleaned, lane_key, job_description) + source_bonus
            row["score"] = score
            ranked.append((score, -index, cleaned, row))
        debug_rows.append(row)

    ranked.sort(reverse=True)
    selected: list[str] = []
    for _score, _neg_index, cleaned, row in ranked:
        if any(cover_sentences_near_duplicate(cleaned, existing) for existing in selected):
            row["rejection_reason"] = "duplicate_with_selected_proof"
            continue
        if selected and not re.search(r"\bI\b", selected[-1]) and re.match(r"^I\b", cleaned):
            row["rejection_reason"] = "abrupt_first_person_switch"
            continue
        row["selected"] = True
        selected.append(cleaned)
        if len(selected) >= limit:
            break
    if selected and not any(paragraph_has_fast_proof(sentence) for sentence in selected):
        for fallback in proof_fallback_sentences(lane_key):
            if any(cover_sentences_near_duplicate(fallback, existing) for existing in selected):
                continue
            selected = [fallback]
            debug_rows.append(
                {
                    "purpose": "proof",
                    "source": "fallback",
                    "original_text": fallback,
                    "text": fallback,
                    "selected": True,
                    "score": 0,
                    "rejection_reason": "fallback_used_for_concrete_proof",
                }
            )
            break
        selected = selected[:limit]
    if selected:
        return selected, debug_rows

    fallback = proof_fallback_sentences(lane_key)
    debug_rows.append(
        {
            "purpose": "proof",
            "source": "fallback",
            "original_text": fallback[0],
            "text": fallback[0],
            "selected": True,
            "score": 0,
            "rejection_reason": "fallback_used",
        }
    )
    return fallback, debug_rows


def short_fit_sentence(
    bank: CoverResponseBank,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
) -> tuple[str, list[dict[str, object]]]:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    opening_sentence = short_company_interest_sentence(bank, company_name, role_title, job_description, resume_text)
    return select_opening_support_sentence(
        opening_support_candidates(bank),
        opening_sentence=opening_sentence,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
        lane_key=profile.primary_lane,
    )


def concise_bridge_sentence(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    supported_area = {
        "implementation_delivery": "the implementation side of this work",
        "change_enablement": "the change and adoption side of this work",
        "analytics_operations": "the analytics and reporting side of this work",
        "customer_success": "the customer-facing delivery side of this work",
        "presales_solution": "the solution and discovery side of this work",
    }.get(profile.primary_lane, "the core work in this role")
    return safe_sentence(
        f"My closest match is {supported_area}, and I would ramp honestly on the more specialized pieces."
    )


def concise_close_sentence(company_name: str, role_title: str, signals: CoverLetterSignals) -> str:
    pain_area = safe_connector_fragment(signals.jd_pain_area, max_words=12) or "highest-priority work"
    role_fragment = safe_role_title(role_title)
    specialty_term = next(
        (
            safe_connector_fragment(term, max_words=5)
            for term in signals.jd_skill_terms
            if safe_connector_fragment(term, max_words=5)
        ),
        "",
    )
    if specialty_term and specialty_term.lower() not in pain_area.lower():
        return safe_sentence(
            f"I would welcome the chance to discuss how I could support {company_possessive(company_name)} {specialty_term} work around {pain_area}."
        )
    return safe_sentence(
        f"I would welcome the chance to discuss how I could support {company_possessive(company_name)} {pain_area} in the {role_fragment} role."
    )


def compose_question_driven_standard_paragraphs(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    signals: CoverLetterSignals,
    *,
    mode: str = STANDARD_COVER_MODE,
    original_role_title: str = "",
    application_responses: tuple[question_prep.QualificationsResponse, ...] = (),
    force_bridge: bool = False,
) -> tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...], dict[str, object]]:
    normalized_mode = normalize_cover_mode(mode)
    bank = build_cover_response_bank(
        job_description,
        resume_text,
        company_name=company_name,
        role_title=original_role_title or role_title,
        application_responses=application_responses,
    )
    profile = build_resume.job_problem_profile(job_description, resume_text)
    opening_sentence = short_company_interest_sentence(bank, company_name, role_title, job_description, resume_text)
    opening_support_sentence, opening_debug = short_fit_sentence(
        bank,
        company_name,
        role_title,
        job_description,
        resume_text,
    )
    opening_candidates = [
        opening_sentence,
        opening_support_sentence,
    ]
    if normalized_mode == LONG_COVER_MODE:
        for sentence in response_sentences(bank.unique_qualifications):
            if any(cover_sentences_near_duplicate(sentence, existing) for existing in opening_candidates):
                continue
            opening_candidates.append(sentence)
            break
    opening_sentences = tuple(
        dedupe_sentences(
            opening_candidates
        )[: (3 if normalized_mode == LONG_COVER_MODE else 2)]
    )
    proof_limit = 4 if normalized_mode == LONG_COVER_MODE else 2
    proof_sentences, proof_debug = select_proof_sentences(
        standard_cover_proof_candidates(
            bank,
            signals=signals,
            lane_key=profile.primary_lane,
            job_description=job_description,
        ),
        opening_sentences=opening_sentences,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
        lane_key=profile.primary_lane,
        limit=proof_limit,
    )
    if not any(re.search(r"\d|\$|%", sentence) for sentence in proof_sentences):
        quantified_fallback = next(
            (
                sentence
                for sentence in [*proof_fallback_sentences(profile.primary_lane), *response_sentences(bank.relevant_experience), *response_sentences(bank.communication)]
                if re.search(r"\d|\$|%", sentence)
            ),
            "",
        )
        if quantified_fallback:
            proof_sentences = [quantified_fallback, *proof_sentences[1:]] if proof_sentences else [quantified_fallback]
    if normalized_mode == STANDARD_COVER_MODE and bank.active_responses and not selected_response_bank_proof_rows(proof_debug):
        strongest_response_row = strongest_response_bank_proof_row(
            proof_debug,
            opening_sentences=opening_sentences,
            selected_proof_sentences=proof_sentences,
        )
        if strongest_response_row:
            response_sentence = str(strongest_response_row.get("text") or "").strip()
            preserved_sentence = preserve_non_response_proof_sentence(proof_debug, proof_sentences)
            proof_sentences = dedupe_cover_sentence_candidates(
                [preserved_sentence, response_sentence],
                limit=2,
            )
            strongest_response_row["selected"] = True
            if preserved_sentence != response_sentence:
                proof_debug.append(
                    {
                        "purpose": "proof",
                        "source": "response_bank_override",
                        "original_text": response_sentence,
                        "text": response_sentence,
                        "selected": True,
                        "score": strongest_response_row.get("score", 0),
                        "rejection_reason": "standard_mode_response_bank_required",
                    }
                )
    close_sentence = concise_close_sentence(company_name, role_title, signals)
    closing_sentences: list[str] = []
    if force_bridge or normalized_mode == LONG_COVER_MODE:
        bridge_sentence = concise_bridge_sentence(company_name, role_title, job_description, resume_text)
        if bridge_sentence:
            closing_sentences.append(bridge_sentence)
    if normalized_mode == LONG_COVER_MODE:
        for sentence in response_sentences(bank.communication):
            if any(cover_sentences_near_duplicate(sentence, existing) for existing in [*opening_sentences, *proof_sentences, *closing_sentences]):
                continue
            closing_sentences.append(sentence)
            break
    closing_sentences.append(close_sentence)
    if normalized_mode == LONG_COVER_MODE:
        support_sentences: list[str] = []
        for sentence in [*proof_sentences[2:], *response_sentences(bank.relevant_experience), *response_sentences(bank.communication)]:
            if any(
                cover_sentences_near_duplicate(sentence, existing)
                for existing in [*opening_sentences, *proof_sentences[:2], *closing_sentences, *support_sentences]
            ):
                continue
            support_sentences.append(sentence)
            if len(support_sentences) >= 2:
                break
        if len(support_sentences) < 2:
            fallback_support_pool = [
                *response_sentences(bank.unique_qualifications),
                *response_sentences(bank.company_interest),
                signals.communication_metric,
                signals.top_accomplishment,
            ]
            for sentence in fallback_support_pool:
                cleaned = sentence.strip()
                if not cleaned:
                    continue
                if any(
                    cover_sentences_near_duplicate(cleaned, existing)
                    for existing in [*opening_sentences, *proof_sentences[:2], *closing_sentences, *support_sentences]
                ):
                    continue
                support_sentences.append(cleaned)
                if len(support_sentences) >= 2:
                    break
        long_paragraphs_list = [
            opening_sentences[0] if opening_sentences else "",
            " ".join(opening_sentences[1:]).strip(),
            " ".join(proof_sentences[:2]).strip(),
            " ".join(support_sentences).strip(),
            " ".join(sentence for sentence in closing_sentences if sentence).strip(),
        ]
        if cover_letter_document_word_count(preferred_salutation(company_name), [paragraph for paragraph in long_paragraphs_list if paragraph]) < LONG_MIN_LETTER_WORDS:
            augmentation = "That toolkit helps me move from raw data to a usable next step quickly."
            if long_paragraphs_list[3]:
                long_paragraphs_list[3] = f"{long_paragraphs_list[3]} {augmentation}".strip()
            else:
                long_paragraphs_list[3] = augmentation
        long_paragraphs = tuple(long_paragraphs_list)
        paragraphs = tuple(paragraph for paragraph in long_paragraphs if paragraph)
        paragraph_purposes = [
            {"paragraph": 1, "purpose": "company_and_role_opening", "sentences": list(opening_sentences[:1])},
            {"paragraph": 2, "purpose": "fit_context", "sentences": list(opening_sentences[1:])},
            {"paragraph": 3, "purpose": "proof", "sentences": list(proof_sentences[:2])},
            {"paragraph": 4, "purpose": "supporting_proof", "sentences": list(support_sentences)},
            {"paragraph": 5, "purpose": "bridge_and_close", "sentences": list(closing_sentences)},
        ]
    else:
        paragraphs = (
            " ".join(opening_sentences).strip(),
            " ".join(proof_sentences).strip(),
            " ".join(sentence for sentence in closing_sentences if sentence).strip(),
        )
        paragraph_purposes = [
            {"paragraph": 1, "purpose": "company_and_role_opening", "sentences": list(opening_sentences)},
            {"paragraph": 2, "purpose": "proof", "sentences": list(proof_sentences)},
            {"paragraph": 3, "purpose": "bridge_and_close", "sentences": list(closing_sentences)},
        ]
    proof_terms = tuple(term for term in signals.jd_skill_terms if any(term.lower() in paragraph.lower() for paragraph in paragraphs[:2]))
    close_terms = tuple(term for term in signals.jd_skill_terms if term.lower() in close_sentence.lower())
    selection_debug = {
        "paragraph_purposes": paragraph_purposes,
        "opening_support_candidates": opening_debug,
        "proof_candidates": proof_debug,
        "selected_opening_sentences": list(opening_sentences),
        "selected_proof_sentences": list(proof_sentences),
        "selected_closing_sentences": list(closing_sentences),
    }
    return tuple(paragraph for paragraph in paragraphs if paragraph), proof_terms[:4], close_terms[:4], selection_debug


def require_file(path: Path, label: str) -> None:
    if not path.is_file():
        fail(f"{label} not found: {path}")


def clean_role_title(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" .:-")
    value = value.replace("Lead Solutioning", "Lead of Solution Consultancy")
    return value[:90]


def normalized_phrase_in_text(text: str, phrase: str) -> bool:
    normalized_text = re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()
    normalized_phrase = re.sub(r"[^a-z0-9]+", " ", (phrase or "").lower()).strip()
    return bool(normalized_text and normalized_phrase and normalized_phrase in normalized_text)


def extract_role_title(job_description: str) -> str | None:
    patterns = (
        r"(?im)^\s*(?:role|job title|position)\s*[:\-]\s*(.+?)\s*$",
        r"(?i)\bAs the\s+([^,\n.]+?)(?:,|\syou will\b)",
        r"(?i)\bThe Role\s*\n\s*([^.\n]{5,90})",
    )
    for pattern in patterns:
        match = re.search(pattern, job_description)
        if match:
            candidate = clean_role_title(match.group(1))
            if candidate and candidate.lower() not in {"job description", "the role"}:
                return candidate

    title = build_resume.extract_job_title(job_description)
    if title:
        return title
    return None


_COVER_SECTION_HEADERS = {
    "about the role": "intro",
    "about": "intro",
    "about us": "intro",
    "job summary": "intro",
    "position summary": "intro",
    "role overview": "intro",
    "summary": "intro",
    "overview": "intro",
    "a typical day in the life includes": "what_youll_do",
    "typical day in the life includes": "what_youll_do",
    "typical day in the life": "what_youll_do",
    "a typical day includes": "what_youll_do",
    "what you'll do": "what_youll_do",
    "what you will do": "what_youll_do",
    "what you do": "what_youll_do",
    "responsibilities": "what_youll_do",
    "what we value": "values",
    "our values": "values",
    "what we offer": "values",
    "benefits": "values",
    "why join us": "values",
    "who you are": "who_you_are",
    "qualifications": "who_you_are",
    "requirements": "who_you_are",
    "basic qualifications": "who_you_are",
    "preferred qualifications": "who_you_are",
    "minimum qualifications": "who_you_are",
    "knowledge skills and abilities": "who_you_are",
    "skills and experience": "who_you_are",
    "what you need": "who_you_are",
    "what you will need": "who_you_are",
    "what youll need": "who_you_are",
    "must haves": "who_you_are",
    "nice to have": "who_you_are",
    "nice to haves": "who_you_are",
    "what would be nice to have": "who_you_are",
}
_NORMALIZED_COVER_SECTION_HEADERS = {
    re.sub(r"[^a-z0-9]+", " ", key.lower()).strip(): value
    for key, value in _COVER_SECTION_HEADERS.items()
}

_MISSION_SKIP_PATTERNS = (
    r"\bculture\b",
    r"\bteam members?\b",
    r"\bchapter\b",
    r"\bauthentic selves\b",
    r"\bcelebrate our successes\b",
    r"^\s*use your\b",
    r"^\s*in this role\b",
    r"^\s*as a[n]?\s+[A-Z][A-Za-z0-9&/ -]{1,80},\s+you(?:'|’)ll\b",
    r"^\s*as a[n]?\s+[^.]{1,120},\s+you will\b",
    r"^\s*this position\b",
    r"^\s*this role focuses\b",
    r"^\s*a typical day in the life includes\b",
    r"^\s*what you(?:'|’)ll do\b",
    r"^\s*what you will do\b",
    r"^\s*what you need\b",
    r"^\s*what you will need\b",
    r"^\s*responsibilities\b",
    r"^\s*requirements\b",
    r"^\s*qualifications\b",
    r"^\s*preferred qualifications\b",
    r"^\s*minimum qualifications\b",
    r"^\s*who you are\b",
    r"^\s*what you will need\b",
    r"^\s*nice to have\b",
    r"^\s*nice to haves\b",
    r"^\s*what would be nice to have\b",
    r"\breports into the\b",
    r"\bwork remotely\b",
    r"^\s*we(?:'|’)re looking\b",
    r"\bis looking for\b",
    r"\bapply today\b",
    r"\bjoin us immediately\b",
)

_MISSION_PROMO_PATTERNS = (
    r"\balong for the ride\b",
    r"\bspecial place\b",
    r"\ball-?around great people\b",
    r"\bcareer can take off\b",
    r"\bjoin us in (?:our|the) journey\b",
    r"\bwould(?:n't| not) it be amazing\b",
    r"\bif you(?:'|’)re willing to do more\b",
    r"\bunique combination of brilliance\b",
    r"\bevery door\b",
    r"\bwork,\s*play and grow\b",
    r"\bbring your brains\b",
    r"\bwinning team\b",
    r"\bpassion and customer focus\b",
    r"\bis open(?:\.|$)\b",
)

_PARTNER_FUNCTION_PATTERNS = (
    "content",
    "product",
    "marketing",
    "technology",
    "cx",
    "customer experience",
    "research",
    "engineering",
    "operations",
    "support",
    "finance",
    "sales",
)

_COVER_ANALYTICS_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("AI-assisted", r"\bai-assisted\b|\bai-guided\b"),
    ("retention", r"\bretention\b"),
    ("subscription", r"\bsubscription(?:-based)?\b"),
    ("lifecycle", r"\blifecycle\b"),
    ("cohort analysis", r"\bcohort analysis\b"),
    ("churn", r"\bchurn\b"),
    ("reactivation", r"\breactivation\b"),
    ("ad hoc analysis", r"\bad hoc analysis\b|\bad hoc\b"),
    ("test design", r"\btest design\b"),
    ("measurement", r"\bmeasurement\b"),
    ("validation", r"\bvalidation\b|\bvalidity\b|\bdefensibility\b|\binterpretability\b"),
    ("assessment quality", r"\bfairness\b|\baccessibility\b|\binstructional appropriateness\b"),
    ("learning systems", r"\blearning systems?\b|\blearner-facing\b|\binstructional\b"),
    ("testing strategy", r"\btesting strategy\b"),
    ("SQL", r"\bsql\b"),
    ("Excel", r"\bexcel\b"),
    ("Looker", r"\blooker\b"),
    ("Snowflake", r"\bsnowflake\b"),
    ("MixPanel", r"\bmixpanel\b"),
    ("Segment", r"\bsegment\b"),
    ("member experience", r"\bmember experience\b"),
    ("member loyalty", r"\bmember loyalty\b"),
)

_COVER_SUPPORT_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("customer experience", r"\bcustomer experience\b|\bcx\b"),
    ("retention", r"\bretention\b"),
    ("loyalty", r"\bloyalty\b"),
    ("escalations", r"\bescalations?\b"),
    ("membership", r"\bmembership\b"),
    ("billing", r"\bbilling\b"),
    ("AI-assisted", r"\bai-assisted\b|\bai-driven\b"),
    ("real-time", r"\breal[- ]time\b"),
    ("Zendesk", r"\bzendesk\b"),
    ("phone and email", r"\bphone\b|\bemail\b"),
    ("quality", r"\bquality\b"),
)

_COVER_CS_ENTERPRISE_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("adoption", r"\badoption\b"),
    ("QBR", r"\bqbr\b|\bquarterly business review\b"),
    ("account health", r"\baccount health\b|\bat-risk accounts?\b"),
    ("retention", r"\bretention\b|\bchurn\b|\brenewal\b"),
    ("stakeholder communication", r"\bstakeholder\b|\bcommunicat(?:e|ing|ion)\b"),
    ("expansion", r"\bexpansion\b|\bupsell\b|\bcross-sell\b"),
    ("decision support", r"\bdecisions?\b|\bvisibility\b|\binsights?\b"),
)

_COVER_STRATEGY_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("strategy", r"\bstrateg(?:y|ic)\b"),
    ("transformation", r"\btransformation\b"),
    ("client service", r"\bclient service\b|\bclients?\b"),
    ("executive alignment", r"\bexecutive\b|\bsenior-level clients?\b"),
    ("recommendations", r"\brecommendation(?:s)?\b|\bfindings\b"),
    ("requirements gathering", r"\brequirements gathering\b|\bstrategy sessions?\b"),
    ("process improvement", r"\bprocess(?:es)?\b|\broot cause\b|\bgap analysis\b"),
)

_COVER_CHANGE_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("change management", r"\bchange management\b|\bchange initiatives?\b|\bchange adoption\b"),
    ("organizational design", r"\borganiz(?:ation|ational) design\b|\brestructuring\b|\bfuture state\b|\btransition planning\b"),
    ("leadership development", r"\bleadership development\b|\bcoaching support\b"),
    ("team effectiveness", r"\bteam effectiveness\b|\bteam charters?\b|\bworking sessions?\b"),
    ("workforce transformation", r"\bworkforce transformation\b|\borganizational transformation\b"),
    ("training", r"\btraining materials?\b|\btraining\b|\bworkshops?\b"),
)

_COVER_IMPLEMENTATION_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("requirements documentation", r"\brequirements documentation\b|\bdocument(?:ing)? project requirements\b|\bdistilling, documenting, and tracking requirements\b"),
    ("project management", r"\bproject management\b|\bproject manager\b|\bdeputy project manager\b"),
    ("stakeholder communication", r"\bclient-facing\b|\bnon-technical clients?\b|\bclient relationships?\b|\bverbally brief\b|\bstand-ups?\b"),
    ("user stories", r"\buser stories\b|\bacceptance criteria\b"),
    ("status reporting", r"\brequirements status\b|\bactivity status\b|\bstatus\b"),
    ("agile delivery", r"\bscrum master\b|\bproduct owner\b|\bagile\b|\bjira\b|\bconfluence\b"),
)

_GENERIC_COVER_TERM_PATTERNS: tuple[tuple[str, str], ...] = (
    ("data analysis", r"\bdata analytics?\b|\banalytical\b"),
    ("stakeholder communication", r"\bcommunicat(?:e|ing|ion)\b"),
    ("cross-functional delivery", r"\bcross-functional\b"),
    ("workflow", r"\bworkflow\b"),
    ("reporting", r"\breporting\b"),
    ("testing", r"\btesting\b"),
    ("decision support", r"\bdecisions?\b|\binsights?\b"),
)

APPROVED_COVER_ACRONYMS = {
    "ERP", "CRM", "SQL", "API", "SOW", "KPI", "QBR", "NLP", "SMS", "BI", "EAC", "ETC",
    "CPQ", "PMI", "SDLC", "SFDC", "UAT",
}
WEAK_OPENER_SIGNALS = (
    "stands out because",
    "this role",
    "relevant to this role",
    "target role",
    "this job description",
    "i am writing to",
    "i am applying for",
    "i would like to express",
)
CONTRACTION_PATTERN = re.compile(
    r"\b(?:can't|couldn't|didn't|doesn't|don't|hadn't|hasn't|haven't|he'd|he'll|he's|here's|how's|"
    r"i'd|i'll|i'm|i've|isn't|it's|let's|mustn't|shan't|she'd|she'll|she's|shouldn't|that's|there's|"
    r"they'd|they'll|they're|they've|we'd|we'll|we're|we've|weren't|what's|who's|won't|wouldn't|you'd|"
    r"you'll|you're|you've)\b",
    re.I,
)
ABSTRACT_COVER_PATTERNS: tuple[tuple[str, str], ...] = (
    (r"\brole moves\b", "cover letter uses the abstract line 'role moves' instead of direct fit language"),
    (r"\bThe work needs\b", "cover letter uses the abstract line 'The work needs' instead of direct proof"),
    (r"\bTogether, that background shows\b", "cover letter uses the canned line 'Together, that background shows'"),
    (r"\bWhere the role reaches beyond\b", "cover letter uses the canned bridge 'Where the role reaches beyond'"),
    (
        r"(^|[.!?]\s+)At [A-Z][A-Za-z0-9&' .-]+,\s+[^.]{0,120}\brole supporting\b",
        "cover letter opens with role-summary narration instead of a direct company-and-role statement",
    ),
)
JD_ARTIFACT_PATTERN = re.compile(
    r"(?im)^\s*(?:specific\s+)?"
    r"(?:responsibilities|requirements|qualifications|activities|overview|"
    r"what you(?:'|’)ll do|what you will do|who you are|preferred qualifications)"
    r"\s*[:\-–]?\s*$",
)

_GENERIC_MISSION_PATTERNS = (
    r"\bprofessional services for b2b customers or client accounts\b",
    r"\bsoftware platform for b2b customers or client accounts\b",
    r"\bis solving a concrete customer problem\b",
)


def _cleanup_short_fragment(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").replace("\ufeff", " ")).strip(" .,:;|-")
    if re.search(r"\b[a-z][A-Z]+\w*", cleaned):
        words = []
        for word in cleaned.split():
            if word.isupper() and len(word) <= 4:
                words.append(word)
            elif re.search(r"[A-Za-z]", word):
                words.append(word[:1].upper() + word[1:].lower())
            else:
                words.append(word)
        cleaned = " ".join(words)
    return cleaned


def has_bullet_artifact(text: str) -> bool:
    return bool(re.search(r"(?:^|\n)[•\*\-–]\s+\w", text or ""))


def cover_allowed_acronyms(
    role_title: str = "",
    job_description: str = "",
    company_name: str = "",
) -> set[str]:
    allowed = {token.upper() for token in APPROVED_COVER_ACRONYMS}
    for token in re.findall(r"\b[A-Z]{2,6}\b", role_title or ""):
        allowed.add(token.upper())
    for token in re.findall(r"\b[A-Z]{2,6}\b", company_name or ""):
        allowed.add(token.upper())
    token_counts = Counter(token.upper() for token in re.findall(r"\b[A-Z]{2,6}\b", job_description or ""))
    for token, count in token_counts.items():
        if count >= 2:
            allowed.add(token)
    return allowed


def sanitize_for_spoken_text(text: str) -> str:
    if not text:
        return ""
    kept_lines: list[str] = []
    for raw_line in text.replace("\ufeff", " ").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue
        if re.match(r"^[•\-\*–]", line):
            continue
        if normalize_cover_section_header(line.strip(" .:-")):
            continue
        if re.fullmatch(
            r"(?:specific\s+)?(?:responsibilities|requirements|qualifications|overview|activities|"
            r"about the role|job summary|position summary|role overview|what you(?:'|’)ll do|what you will do|"
            r"what you need|what you will need|what would be nice to have|who you are|preferred qualifications|"
            r"minimum qualifications|nice to haves?|must haves?)[:.]?",
            line,
            re.I,
        ):
            continue
        if re.fullmatch(r"[A-Z][A-Z &/:\-]{4,}", line):
            continue
        kept_lines.append(line)
    cleaned = " ".join(kept_lines)
    cleaned = normalize_cover_dash_punctuation(cleaned)
    cleaned = re.sub(r"(?:^|\s)[•\*\-–]\s+\w[^•\*\-–]*", " ", cleaned)

    def replace_all_caps(match: re.Match[str]) -> str:
        token = match.group(0)
        return token if token.upper() in APPROVED_COVER_ACRONYMS else " "

    cleaned = re.sub(r"\b[A-Z]{3,}\b", replace_all_caps, cleaned)
    cleaned = re.sub(r"\b[a-z][A-Z]+\w*", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .,:;|-")
    if len(cleaned) < 10:
        return ""
    return cleaned


def normalize_cover_dash_punctuation(text: str) -> str:
    if not text:
        return ""
    cleaned = text.replace("\u2014", " - ").replace("\u2013", " - ")
    cleaned = re.sub(r"[ \t]*--[ \t]*", ", ", cleaned)
    cleaned = re.sub(r",[ \t]*,", ", ", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r" *\n *", "\n", cleaned)
    return cleaned.strip()


def mission_candidate_has_recruiting_fluff(text: str) -> bool:
    cleaned = normalize_cover_dash_punctuation(text)
    return any(re.search(pattern, cleaned, re.I) for pattern in _MISSION_PROMO_PATTERNS)


def trusted_company_description(company_name: str, text: str) -> str:
    if not company_name or not text:
        return ""
    match = re.search(
        rf"\b(?:thousands?\s+of\s+)?(?:customers?|clients?|organizations?|teams?)[^.!?\n]*?\btrust\s+{re.escape(company_name)}\s+to\s+([^.!?\n]+)",
        text,
        re.I,
    )
    if not match:
        return ""
    clause = normalize_cover_dash_punctuation(match.group(1))
    clause = re.split(r"\b(?:a concept we call|so that)\b", clause, maxsplit=1, flags=re.I)[0].rstrip(" ,;:-")
    if word_count(clause) < 4:
        return ""
    return ensure_sentence(f"{company_name} helps customers {clause}")


def safe_company_name(raw: str) -> str:
    fallback = _cleanup_short_fragment(raw)
    if not fallback or len(fallback) > 60:
        return "your organization"
    if has_bullet_artifact(fallback) or re.search(r"\bRESPONSIBILITIES\b|\bREQUIREMENTS\b|\bQUALIFICATIONS\b", fallback, re.I):
        return "your organization"
    cleaned = sanitize_for_spoken_text(raw)
    if cleaned and len(cleaned) <= 60:
        return cleaned
    return fallback


def safe_role_title(raw: str) -> str:
    fallback = _cleanup_short_fragment(raw)
    if not fallback:
        return "this role"
    if len(fallback) > 80 or len(fallback.split()) > 12:
        shortened = build_resume.primary_header_title(fallback)
        if shortened and len(shortened) <= 60:
            return shortened
        return "this role"
    if has_bullet_artifact(fallback) or re.search(r"\b[A-Z]{4,}\b", fallback):
        return "this role"
    if 2 <= len(fallback.split()) <= 12:
        return fallback
    cleaned = sanitize_for_spoken_text(raw)
    if cleaned and len(cleaned) <= 80 and not re.search(r"\b[A-Z]{4,}\b", cleaned):
        return cleaned
    return fallback


def safe_core_problem(raw: str) -> str:
    fallback = _cleanup_short_fragment(raw)
    if not fallback or len(fallback) > 200:
        return ""
    if has_bullet_artifact(fallback) or re.fullmatch(
        r"(?:specific\s+)?(?:responsibilities|requirements|qualifications|overview|activities)\s*[:\-–]?",
        fallback,
        re.I,
    ):
        return ""
    cleaned = sanitize_for_spoken_text(raw)
    if cleaned and len(cleaned) <= 200:
        return cleaned
    return fallback


def sentence_safe_role_core_problem(raw: str) -> str:
    phrase = safe_core_problem(raw) or raw
    if phrase.count(",") >= 2:
        trimmed = phrase.split(",", 1)[0].strip()
        if trimmed:
            return trimmed
    return phrase


def cover_letter_prose_check_text(body_text: str) -> str:
    stripped = re.sub(r"^.*?Dear\s+[^,\n]+,\s*", "", body_text, flags=re.S).strip()
    return stripped or body_text


def normalize_cover_section_header(value: str) -> str:
    cleaned = re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()
    return _NORMALIZED_COVER_SECTION_HEADERS.get(cleaned, "")


def split_cover_section_line(value: str) -> tuple[str, str]:
    direct_header = normalize_cover_section_header(value)
    if direct_header:
        return direct_header, ""
    for header_text in sorted(_COVER_SECTION_HEADERS, key=len, reverse=True):
        header_pattern = re.sub(r"\s+", r"\\s+", re.escape(header_text))
        match = re.match(
            rf"^\s*(?:specific\s+)?{header_pattern}\b\s*[:\-–]?\s*(.*)$",
            value,
            re.I,
        )
        if match:
            return _COVER_SECTION_HEADERS[header_text], match.group(1).strip()
    return "", ""


def looks_like_cover_section_label(value: str) -> bool:
    cleaned = re.sub(r"\s+", " ", sanitize_for_spoken_text(value)).strip(" .:-")
    if not cleaned:
        return False
    if normalize_cover_section_header(cleaned):
        return True
    return bool(
        re.fullmatch(
            r"(?:specific\s+)?(?:summary|overview|description|about us|about the role|job summary|position summary|role overview|"
            r"what you(?:'|’)ll do|what you will do|what you need|what you will need|what would be nice to have|"
            r"responsibilities|requirements|qualifications|preferred qualifications|minimum qualifications|"
            r"who you are|must haves?|nice to haves?|knowledge skills and abilities|skills and experience|"
            r"what we value|our values|what we offer|benefits|why join us)",
            cleaned,
            re.I,
        )
    )


def normalize_cover_role_title_for_prose(value: str) -> str:
    cleaned = clean_role_title(value)
    cleaned = re.sub(r"\s*/\s*", " and ", cleaned)
    cleaned = re.sub(r",\s*", " ", cleaned)
    cleaned = re.sub(r"\bSr\b\.?", "Senior", cleaned, flags=re.I)
    if cleaned and (len(cleaned) > 60 or len(cleaned.split()) > 8 or " - " in cleaned):
        shortened = build_resume.primary_header_title(cleaned)
        if shortened:
            cleaned = shortened
    return re.sub(r"\s+", " ", cleaned).strip()


def cover_letter_jd_sections(job_description: str) -> dict[str, str]:
    scoped = build_resume.role_requirement_text(job_description)
    sections: dict[str, list[str]] = {
        "intro": [],
        "what_youll_do": [],
        "who_you_are": [],
        "values": [],
    }
    current_section = "intro"
    for raw_line in scoped.splitlines():
        line = raw_line.strip()
        if not line:
            if sections[current_section] and sections[current_section][-1] != "":
                sections[current_section].append("")
            continue
        cleaned_line = sanitize_for_spoken_text(line)
        normalized_header, remainder = split_cover_section_line(cleaned_line or line)
        if normalized_header:
            current_section = normalized_header
            cleaned_remainder = sanitize_for_spoken_text(remainder)
            if cleaned_remainder:
                sections[current_section].append(cleaned_remainder)
            continue
        if not cleaned_line or looks_like_cover_section_label(cleaned_line):
            continue
        sections[current_section].append(cleaned_line)
    return {
        key: re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()
        for key, lines in sections.items()
    }


def cover_section_sentences(text: str) -> list[str]:
    if not text:
        return []
    cleaned = sanitize_for_spoken_text(text)
    if not cleaned:
        return []
    return sentence_list(cleaned)


def rewrite_company_pronouns(text: str, company_name: str) -> str:
    rewritten = normalize_cover_dash_punctuation(text.strip())
    leading_replacements = (
        (r"^\s*We make\b", f"{company_name} makes"),
        (r"^\s*We help\b", f"{company_name} helps"),
        (r"^\s*We deliver\b", f"{company_name} delivers"),
        (r"^\s*We improve\b", f"{company_name} improves"),
        (r"^\s*We provide\b", f"{company_name} provides"),
        (r"^\s*We build\b", f"{company_name} builds"),
        (r"^\s*We create\b", f"{company_name} creates"),
        (r"^\s*We enable\b", f"{company_name} enables"),
        (r"^\s*We support\b", f"{company_name} supports"),
        (r"^\s*We are\b", f"{company_name} is"),
        (r"^\s*We['’]re\b", f"{company_name} is"),
    )
    for pattern, replacement in leading_replacements:
        if re.search(pattern, rewritten, re.I):
            rewritten = re.sub(pattern, replacement, rewritten, flags=re.I)
            break
    rewritten = re.sub(r"^\s*We\b", company_name, rewritten, flags=re.I)
    rewritten = re.sub(r"^\s*Our\b", "Its", rewritten, flags=re.I)
    rewritten = re.sub(r"\bwe['’]re\b", f"{company_name} is", rewritten, flags=re.I)
    rewritten = re.sub(r"\bwe\b", company_name, rewritten, flags=re.I)
    rewritten = re.sub(r"\bour\b", "its", rewritten, flags=re.I)
    rewritten = re.sub(r"\bours\b", "its", rewritten, flags=re.I)
    rewritten = re.sub(
        rf"^{re.escape(company_name)} makes\b(.*)\band deliver\b",
        rf"{company_name} makes\1and delivers",
        rewritten,
        flags=re.I,
    )
    if company_name:
        rewritten = re.sub(
            rf"^\s*At\s+{re.escape(company_name)},\s+{re.escape(company_name)}\b",
            company_name,
            rewritten,
            flags=re.I,
        )
    return re.sub(r"\s+", " ", rewritten).strip()


def compact_cover_sentence(text: str, *, max_words: int = 22) -> str:
    cleaned = normalize_cover_dash_punctuation(re.sub(r"\s+", " ", text).strip().rstrip("."))
    if word_count(cleaned) <= max_words:
        return f"{cleaned}."
    sentences = sentence_list(cleaned)
    if sentences:
        first_sentence = ensure_sentence(sentences[0].strip().rstrip("."))
        if first_sentence and word_count(first_sentence.rstrip(".")) <= max_words:
            return first_sentence
    chunks = [chunk.strip() for chunk in re.split(r"[;:]", cleaned) if chunk.strip()]
    if chunks and word_count(chunks[0]) <= max_words:
        return f"{chunks[0].rstrip('.') }.".replace(" .", ".")
    words = cleaned.split()
    trimmed = " ".join(words[:max_words]).rstrip(",;:-")
    for pattern in (
        r",?\s+by\s+[^,;:.]+$",
        r",?\s+using\s+[^,;:.]+$",
        r",?\s+while\s+[^,;:.]+$",
        r",?\s+including\s+[^,;:.]+$",
        r",?\s+with\s+[^,;:.]+$",
    ):
        tightened = re.sub(pattern, "", trimmed, flags=re.I).rstrip(",;:- ")
        if tightened != trimmed and word_count(tightened) >= 6:
            trimmed = tightened
            break
    return f"{trimmed}."


def compress_cover_evidence_text(text: str) -> str:
    cleaned = proof_text.rewrite_dense_proof_patterns(re.sub(r"\s+", " ", (text or "")).strip())
    replacements = (
        (
            r"\bconverted complex implementation and data-migration requirements into statements of work and functional requirements across 80\+ international manufacturing client engagements, clarifying scope, milestones, and cost baselines before build work began\b",
            "converted complex implementation and data-migration requirements into statements of work and functional requirements across 80+ international manufacturing client engagements before build work began",
        ),
        (
            r"\bcomplex implementation, data migration, integration, and customization needs\b",
            "complex implementation and data-migration requirements",
        ),
        (
            r"\bimplementation readiness, scope alignment, sandbox testing, user acceptance validation, targeted future-state system training, and administrative release tasks\b",
            "implementation readiness, testing, training, and release coordination",
        ),
        (
            r"\bsupporting data migration, validation, training, and adoption\b",
            "supporting migration readiness, user validation, and adoption",
        ),
        (
            r"\bstabilized high-risk accounts across a \$6M\+ book of business, including more than one million dollars in at-risk annual revenue, by diagnosing root causes, consolidating ownership, and driving resolution\b",
            "stabilized high-risk accounts across a $6M+ book of business, including $1M+ in at-risk annual revenue, by consolidating ownership and driving resolution",
        ),
    )
    for pattern, replacement in replacements:
        cleaned = re.sub(pattern, replacement, cleaned, flags=re.I)
    return cleaned


def cover_support_sentence_is_usable(text: str, *, min_words: int = 6) -> bool:
    cleaned = re.sub(r"\s+", " ", (text or "")).strip()
    if word_count(cleaned) < min_words:
        return False
    if has_bullet_artifact(cleaned):
        return False
    if re.fullmatch(
        r"(?:responsibilities|requirements|qualifications|overview|who you are|what you(?:'|’)ll do|what you will do|what you will need)[:.]?",
        cleaned,
        re.I,
    ):
        return False
    if re.search(r"\bas the primary [^.]{0,140}\bacross a\.?$", cleaned, re.I):
        return False
    if re.search(r"\b(?:a|an|the|and|or|across|through|with|for|to|of|as|into|from|by)\.?$", cleaned, re.I):
        return False
    return True


def safe_connector_fragment(text: str, *, max_words: int = 10) -> str:
    cleaned = _cleanup_short_fragment(text)
    if not cleaned:
        return ""
    cleaned = re.sub(r"^(?:needs?|needing)\s+", "", cleaned, flags=re.I)
    if word_count(cleaned) > max_words:
        return ""
    if has_bullet_artifact(cleaned):
        return ""
    if re.search(r"\b(?:responsibilities|requirements|qualifications|overview)\b", cleaned, re.I):
        return ""
    if re.search(r"[;:]", cleaned):
        return ""
    if re.search(r"^(?:and|or|with|across|through|for|to|of|on|in)\b", cleaned, re.I):
        return ""
    if re.search(r"\b(?:building|delivering|translating|turning|protecting|keeping|aligning)\b", cleaned, re.I):
        return ""
    return cleaned


def with_indefinite_article(phrase: str) -> str:
    cleaned = re.sub(r"\s+", " ", phrase).strip()
    if not cleaned or re.match(r"^(?:a|an|the)\b", cleaned, re.I):
        return cleaned
    return shared_with_indefinite_article(cleaned)


def company_possessive(company_name: str) -> str:
    cleaned = re.sub(r"\s+", " ", company_name).strip()
    if not cleaned:
        return ""
    if cleaned.lower().endswith("s"):
        return f"{cleaned}'"
    return f"{cleaned}'s"


def preferred_salutation(company_name: str) -> str:
    cleaned = re.sub(r"\s+", " ", company_name).strip()
    cleaned = re.sub(r"\b(?:Inc\.?|LLC|Ltd\.?|Corporation|Corp\.?|Company|Co\.)\b", "", cleaned, flags=re.I).strip(" ,")
    awkward = (
        not cleaned
        or len(cleaned.split()) > 4
        or bool(re.search(r"[/|@]", cleaned))
        or len(cleaned) > 30
    )
    return "Dear Hiring Manager," if awkward else f"Dear {cleaned} Team,"


def comma_join(items: tuple[str, ...] | list[str]) -> str:
    normalized = [item.strip() for item in items if item and item.strip()]
    if not normalized:
        return ""
    if len(normalized) == 1:
        return normalized[0]
    if len(normalized) == 2:
        return f"{normalized[0]} and {normalized[1]}"
    return f"{', '.join(normalized[:-1])}, and {normalized[-1]}"


def display_cover_term(value: str) -> str:
    cleaned = re.sub(r"\s+", " ", value).strip()
    return "CX" if cleaned.lower() == "cx" else cleaned


def display_partner_functions(items: tuple[str, ...] | list[str]) -> tuple[str, ...]:
    return tuple(display_cover_term(item) for item in items if item and item.strip())


def partner_team_phrase(items: tuple[str, ...] | list[str]) -> str:
    displayed = display_partner_functions(items)
    if not displayed:
        return ""
    return f"teams across {comma_join(displayed)}"


def support_escalation_role(role_title: str, job_description: str) -> bool:
    combined = f"{role_title}\n{job_description}".lower()
    return bool(
        re.search(r"\b(customer experience|customer support|support specialist|support agent|member services|member support|cx)\b", combined)
        and re.search(r"\b(escalation|retention|billing|membership|zendesk|phone|email|vip|partner|human touch|support)\b", combined)
    )


def extract_partner_functions(job_description: str, sections: dict[str, str]) -> tuple[str, ...]:
    combined = "\n".join(
        part for part in (sections.get("what_youll_do", ""), sections.get("who_you_are", ""), job_description) if part
    )
    found: list[str] = []
    canonical_labels = {
        "customer experience": "cx",
    }
    for term in _PARTNER_FUNCTION_PATTERNS:
        label = canonical_labels.get(term, term)
        if re.search(rf"\b{re.escape(term)}\b", combined, re.I) and label not in found:
            found.append(label)
    return tuple(found[:4])


def extract_test_environments(job_description: str, sections: dict[str, str]) -> tuple[str, ...]:
    source_lines = [
        line.strip()
        for part in (sections.get("what_youll_do", ""), sections.get("intro", ""))
        for line in part.splitlines()
        if line.strip()
    ]
    environments: list[str] = []
    allowed_pattern = (
        r"\b(?:mobile|member services|email|app|product|marketing|cx|research|qa|uat|sandbox|training|rollout|deployment|lifecycle|"
        r"clients?|countries|regions|sites|frontline|executives?)\b"
    )
    blocked_pattern = (
        r"\b(?:deliverables|collaborate|cross-functional|what you'll do|what you will do|"
        r"support testing|production lifecycle|global clients?|product owners?|scrum masters?|technical leads?|"
        r"collaborating closely|ensure delivery success)\b"
    )
    for line in source_lines:
        for match in re.finditer(r"\(e\.g\.,\s*([^)]+)\)", line, re.I):
            raw_parts = re.split(r",|&| and ", match.group(1))
            for part in raw_parts:
                cleaned = re.sub(r"\s+", " ", part).strip(" .")
                if not cleaned or re.search(blocked_pattern, cleaned, re.I):
                    continue
                if not re.search(allowed_pattern, cleaned, re.I):
                    continue
                if cleaned.lower() not in {item.lower() for item in environments}:
                    environments.append(cleaned)
        for match in re.finditer(r"\bacross\s+([^.;\n]+)", line, re.I):
            span = match.group(1)
            if not re.search(allowed_pattern, span, re.I):
                continue
            raw_parts = re.split(r",|&| and ", span)
            for part in raw_parts:
                cleaned = re.sub(r"^(?:the|new)\s+", "", part.strip(), flags=re.I)
                cleaned = re.sub(r"^[A-Za-z][A-Za-z0-9&.\-]*['’]s\s+", "", cleaned)
                cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
                if not cleaned or re.search(blocked_pattern, cleaned, re.I):
                    continue
                if not re.search(allowed_pattern, cleaned, re.I):
                    continue
                if re.search(r"\bteams?\b$", cleaned, re.I):
                    continue
                if word_count(cleaned) > 7 and not re.search(r"\d|qa|uat|sandbox", cleaned, re.I):
                    continue
                if cleaned.lower() not in {item.lower() for item in environments}:
                    environments.append(cleaned)
    return tuple(environments[:4])


def extract_company_mission(company_name: str, sections: dict[str, str], job_description: str) -> str:
    company_name = safe_company_name(company_name)
    metadata_header_pattern = r"^\s*(?:position title|department|agency|office|directorate|announcement number|series|grade|security clearance)\b|^\s*\d+\."
    requirement_line_pattern = (
        r"\b(?:bachelor(?:'|’)s|master(?:'|’)s|degree|years? of experience|client-facing experience|"
        r"what technical skills|qualifications are necessary|capability signals|non-negotiables|bonus points)\b"
    )
    alias_names = [company_name]
    if company_name:
        first_token = company_name.split()[0].strip(",.:;")
        if first_token and first_token.lower() != company_name.lower():
            alias_names.append(first_token)
        acronym = re.sub(r"[^A-Z]", "", company_name)
        if len(acronym) >= 2 and acronym.lower() != company_name.lower():
            alias_names.append(acronym)
    alias_names = list(dict.fromkeys(name for name in alias_names if name))
    intro_raw = sections.get("intro", "")
    raw_intro_lines = [
        line.strip()
        for line in intro_raw.splitlines()
        if line.strip()
    ]
    intro_sentences: list[str] = []
    buffer = ""
    for raw_line in raw_intro_lines:
        line = sanitize_for_spoken_text(raw_line)
        if not line:
            continue
        if buffer and not re.search(r"[.!?]$", buffer) and not looks_like_cover_section_label(line):
            buffer = f"{buffer} {line}".strip()
            continue
        if buffer:
            intro_sentences.append(buffer)
        buffer = line
    if buffer:
        intro_sentences.append(buffer)
    company_descriptor_patterns = tuple(
        pattern
        for name in alias_names
        for pattern in (
            rf"\b{re.escape(name)}\b\s+(?:is|makes|delivers|builds|provides|supports|helps|enables|serves|powers|drives|offers)\b",
            rf"\b{re.escape(name)}\b\s+has\s+become\b",
            rf"\b{re.escape(name)}['’]s\b",
        )
    )
    direct_mission_patterns = tuple(
        pattern
        for name in alias_names
        for pattern in (
            rf"(?:At\s+)?{re.escape(name)}[\s\S]*?\b(?:company mission|mission) is\b[\s\S]*?(?:\.|$)",
            rf"\b{re.escape(name)}\b\s+has\s+become\b[\s\S]*?(?:\.|$)",
            rf"\b{re.escape(name)}\b\s+(?:makes|is|delivers|builds|provides|supports|helps|enables|serves|powers|drives|offers)\b[\s\S]*?(?:\.|$)",
        )
    )
    direct_candidate = ""
    for pattern in direct_mission_patterns:
        direct_match = re.search(pattern, intro_raw, re.I)
        if not direct_match:
            continue
        direct_text = re.sub(r"\s+", " ", direct_match.group(0).strip())
        direct_text = re.sub(r"^\s*At\s+", "", direct_text, flags=re.I)
        occurrences = list(re.finditer(re.escape(company_name), direct_text, re.I))
        if len(occurrences) >= 2:
            direct_text = direct_text[:occurrences[1].start()].rstrip(" ,;:-")
        direct_text = re.sub(
            rf"^\s*{re.escape(company_name)},\s+our company mission is\b",
            f"{company_possessive(company_name)} mission is",
            direct_text,
            flags=re.I,
        )
        if not mission_candidate_has_recruiting_fluff(direct_text):
            direct_candidate = direct_text
            break

    def fallback_mission() -> str:
        if (
            re.search(r"\b(?:department|agency|administration|command|service|office|bureau)\b", company_name, re.I)
            or re.search(r"^\s*position title:", job_description, re.I)
            or re.search(r"\b(?:qualification standards|GS-\d+|knowledge, skills and abilities|KSA(?:'s)?|USAJOBS)\b", job_description, re.I)
        ):
            return ""
        trusted_description = trusted_company_description(company_name, intro_raw) or trusted_company_description(company_name, job_description)
        if trusted_description:
            return trusted_description
        context = business_context.extract_business_context(job_description)
        if context.industry == "education / assessment":
            return f"{company_name} supports assessment and learning systems for schools, educators, and learners."
        if context.product_or_service and context.customer_type:
            product_or_service = context.product_or_service
            if (
                product_or_service.lower() == "professional services"
                and context.customer_type.lower() == "b2b customers or client accounts"
            ):
                return ""
            if (
                product_or_service.lower() == "software platform"
                and context.customer_type.lower() == "b2b customers or client accounts"
            ):
                return ""
            if product_or_service.lower().endswith("services"):
                return f"{company_name} delivers {product_or_service} for {context.customer_type}."
            return f"{company_name} delivers {with_indefinite_article(product_or_service)} for {context.customer_type}."
        if context.product_or_service:
            product_or_service = context.product_or_service
            if product_or_service.lower() == "professional services":
                return ""
            if product_or_service.lower() == "software platform":
                return ""
            if product_or_service.lower().endswith("services"):
                return f"{company_name} delivers {product_or_service}."
            return f"{company_name} delivers {with_indefinite_article(product_or_service)}."
        return ""

    intro_lines = [
        line
        for line in intro_sentences
        if line
        and not re.search(r"\b(?:reports to|department:|location:)\b", line, re.I)
    ]
    candidates = [
        line
        for line in intro_lines
        if not looks_like_cover_section_label(line)
        if not any(re.search(pattern, line, re.I) for pattern in _MISSION_SKIP_PATTERNS)
        and not re.search(metadata_header_pattern, line, re.I)
        and not re.search(requirement_line_pattern, line, re.I)
        and not mission_candidate_has_recruiting_fluff(line)
        and not re.search(r"\b(reports to|department:|location:|you will)\b", line, re.I)
        and not re.fullmatch(
            r"(?:summary|overview|description|about the role|responsibilities|requirements|qualifications|who you are|what you(?:'|’)ll do|what you will do|what you will need)[:.]?",
            line.strip(),
            re.I,
        )
    ]
    company_descriptor_lines = [
        line
        for line in candidates
        if any(re.search(pattern, line, re.I) for pattern in company_descriptor_patterns)
    ]
    business_descriptor_lines = [
        line
        for line in company_descriptor_lines
        if re.search(
            r"\b(software|cloud|platform|product|products|service|services|organizations?|customers?|clients?|"
            r"digital transformation|business cloud|industry specific markets?)\b",
            line,
            re.I,
        )
    ]
    mission_lines = [
        line
        for line in candidates
        if re.search(r"\b(company mission|mission is)\b", line, re.I)
    ]
    preferred = [
        line
        for line in candidates
        if re.search(r"\b(?:mission|make|help|deliver|improving|improve|enable|tailored|product|service|platform)\b", line, re.I)
    ]
    candidate = direct_candidate or (mission_lines[0] if mission_lines else "")
    if not candidate:
        candidate = next((line for line in business_descriptor_lines if not re.search(r"\bwas born with the mission\b", line, re.I)), "")
    if not candidate:
        candidate = next((line for line in company_descriptor_lines if not re.search(r"\bwas born with the mission\b", line, re.I)), "")
    if not candidate:
        candidate = next((line for line in preferred if not re.search(r"\bwas born with the mission\b", line, re.I)), "")
    if not candidate:
        candidates = [
            sentence
            for sentence in intro_sentences
            if not looks_like_cover_section_label(sentence)
            if not any(re.search(pattern, sentence, re.I) for pattern in _MISSION_SKIP_PATTERNS)
            and not re.search(metadata_header_pattern, sentence, re.I)
            and not re.search(requirement_line_pattern, sentence, re.I)
            and not mission_candidate_has_recruiting_fluff(sentence)
            and not re.search(r"\b(reports to|department:|location:|you will)\b", sentence, re.I)
            and not re.fullmatch(
                r"(?:summary|overview|description|about the role|responsibilities|requirements|qualifications|who you are|what you(?:'|’)ll do|what you will do|what you will need)[:.]?",
                sentence.strip(),
                re.I,
            )
        ]
        company_descriptor_lines = [
            sentence
            for sentence in candidates
            if any(re.search(pattern, sentence, re.I) for pattern in company_descriptor_patterns)
        ]
        business_descriptor_lines = [
            sentence
            for sentence in company_descriptor_lines
            if re.search(
                r"\b(software|cloud|platform|product|products|service|services|organizations?|customers?|clients?|"
                r"digital transformation|business cloud|industry specific markets?)\b",
                sentence,
                re.I,
            )
        ]
        mission_lines = [
            sentence
            for sentence in candidates
            if re.search(r"\b(company mission|mission is)\b", sentence, re.I)
        ]
        preferred = [
            sentence
            for sentence in candidates
            if re.search(r"\b(?:mission|make|help|deliver|improving|improve|enable|tailored|product|service|platform)\b", sentence, re.I)
        ]
        candidate = mission_lines[0] if mission_lines else ""
        if not candidate:
            candidate = next((sentence for sentence in business_descriptor_lines if not re.search(r"\bwas born with the mission\b", sentence, re.I)), "")
        if not candidate:
            candidate = next((sentence for sentence in company_descriptor_lines if not re.search(r"\bwas born with the mission\b", sentence, re.I)), "")
        if not candidate:
            candidate = next((sentence for sentence in preferred if not re.search(r"\bwas born with the mission\b", sentence, re.I)), "")
    if not candidate:
        candidate = preferred[0] if preferred else (candidates[0] if candidates else "")
    if not candidate:
        jd_lines = [
            sanitize_for_spoken_text(line)
            for line in job_description.splitlines()
        ]
        jd_candidates = [
            line
            for line in jd_lines
            if line
            and company_name
            and any(re.search(rf"\b{re.escape(name)}\b", line, re.I) for name in alias_names)
            and not looks_like_cover_section_label(line)
            and not any(re.search(pattern, line, re.I) for pattern in _MISSION_SKIP_PATTERNS)
            and not re.search(metadata_header_pattern, line, re.I)
            and not re.search(requirement_line_pattern, line, re.I)
            and not mission_candidate_has_recruiting_fluff(line)
            and not re.search(r"\b(reports to|department:|location:|you will)\b", line, re.I)
            and not re.search(
                r"\b(?:is seeking|looking for|job title|travel required|clearance required|job family|reports to|location)\b",
                line,
                re.I,
            )
        ]
        jd_descriptor_lines = [
            line
            for line in jd_candidates
            if any(re.search(pattern, line, re.I) for pattern in company_descriptor_patterns)
        ]
        jd_business_descriptor_lines = [
            line
            for line in jd_descriptor_lines
            if re.search(
                r"\b(software|cloud|platform|product|products|service|services|organizations?|customers?|clients?|"
                r"digital transformation|business cloud|industry specific markets?)\b",
                line,
                re.I,
            )
        ]
        jd_preferred = [
            line
            for line in jd_candidates
            if re.search(
                r"\b(?:deliver|delivering|drive|drives|help|helps|support|supports|provide|provides|build|builds|"
                r"enable|enables|technology|advisory|managed services|platform|future|impact)\b",
                line,
                re.I,
            )
        ]
        candidate = (
            jd_business_descriptor_lines[0]
            if jd_business_descriptor_lines
            else (jd_descriptor_lines[0] if jd_descriptor_lines else (jd_preferred[0] if jd_preferred else (jd_candidates[0] if jd_candidates else "")))
        )
    if not candidate:
        candidate = fallback_mission()
    candidate_is_company_descriptor = bool(
        company_name
        and any(re.search(pattern, candidate, re.I) for pattern in company_descriptor_patterns)
    )
    if candidate and intro_sentences and not re.search(r"[.!?]$", candidate) and not candidate_is_company_descriptor:
        expanded = next(
            (
                sentence
                for sentence in intro_sentences
                if candidate in sentence and len(sentence.split()) > len(candidate.split()) + 2
            ),
            "",
        )
        if expanded:
            candidate = expanded
    rewritten_candidate = rewrite_company_pronouns(candidate, company_name)
    if re.search(r"\bis\s+(?:innovators?|changemakers?|collaborators?|pioneers)\b", rewritten_candidate, re.I):
        rewritten_candidate = fallback_mission()
    rewritten_candidate = re.sub(rf"^\s*About\s+{re.escape(company_name)}[:.]?\s*", "", rewritten_candidate, flags=re.I)
    rewritten_candidate = re.sub(r"^\s*About\s+[A-Z][A-Za-z0-9&.' -]{2,40}[:.]?\s*", "", rewritten_candidate, flags=re.I)
    rewritten_candidate = re.sub(r"^\s*(?:summary|overview|description)[:.]?\s*", "", rewritten_candidate, flags=re.I)
    possessive_name = company_possessive(company_name)
    rewritten_candidate = re.sub(r"^\s*Together,\s+its\b", possessive_name, rewritten_candidate, flags=re.I)
    rewritten_candidate = re.sub(r"^\s*Its\b", possessive_name, rewritten_candidate, flags=re.I)
    rewritten_candidate = rewritten_candidate.lstrip("\ufeff")
    rewritten_candidate = re.sub(r"^\s*Company\s*:\s*[^.]+\.?\s*", "", rewritten_candidate, flags=re.I)
    if re.search(r"\brole overview\b|\b(?:is|are) seeking\b|^\s*this role\b", rewritten_candidate, re.I):
        rewritten_candidate = fallback_mission()
    if re.search(r"\bis looking for\b", rewritten_candidate, re.I):
        rewritten_candidate = fallback_mission()
    if re.search(metadata_header_pattern, rewritten_candidate, re.I):
        rewritten_candidate = fallback_mission()
    if (
        rewritten_candidate
        and company_name
        and not re.match(rf"^\s*{re.escape(company_name)}\b", rewritten_candidate, re.I)
        and re.match(r"^\s*(design|lead|own|drive|build|support|manage|coordinate|implement|develop|deliver)\b", rewritten_candidate, re.I)
    ):
        rewritten_candidate = fallback_mission()
    if company_name and not re.search(re.escape(company_name), rewritten_candidate, re.I):
        fallback = fallback_mission()
        if fallback and len(fallback.split()) >= 4:
            rewritten_candidate = fallback
    if mission_candidate_has_recruiting_fluff(rewritten_candidate):
        rewritten_candidate = fallback_mission()
    if any(re.search(pattern, rewritten_candidate, re.I) for pattern in _GENERIC_MISSION_PATTERNS):
        rewritten_candidate = ""
    return compact_cover_sentence(rewritten_candidate, max_words=20) if rewritten_candidate else ""


def extract_role_core_function(
    company_name: str,
    role_title: str,
    job_description: str,
    sections: dict[str, str],
) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    intro = sections.get("intro", "")
    role_lines = [
        sanitize_for_spoken_text(line)
        for part in (sections.get("what_youll_do", ""), intro)
        for line in part.splitlines()
        if sanitize_for_spoken_text(line) and not looks_like_cover_section_label(line)
    ]
    role_text = "\n".join(part for part in (sections.get("what_youll_do", ""), intro) if part)
    role_text_spoken = sanitize_for_spoken_text(role_text)
    responsibility_match = re.search(
        r"\b(?:responsible for|serve as .*? responsible for|lead responsible for)\s+(.+?)(?:\.|$|\bThis role\b|\bThe position\b|\bSuccess requires\b)",
        role_text_spoken,
        re.I,
    )
    if responsibility_match:
        phrase = safe_core_problem(responsibility_match.group(1))
        if 3 <= word_count(phrase) <= 22:
            return phrase
    for line in role_lines:
        responsibility_match = re.search(r"^(?:be\s+)?responsible for\s+(.+)$", line, re.I)
        if not responsibility_match:
            responsibility_match = re.search(r"\bresponsible for\s+(.+)$", line, re.I)
        if not responsibility_match:
            continue
        phrase = safe_core_problem(responsibility_match.group(1))
        if 3 <= word_count(phrase) <= 16:
            return phrase
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    lowered = role_text.lower()
    if lane_key == "analytics_operations":
        if is_education_assessment_context(job_description):
            return "supporting AI-assisted assessment and learning workflows with stronger validation, quality, and continuous improvement"
        if "subscription retention" in lowered or "member loyalty" in lowered:
            return "turning member behavior and subscription performance into clearer retention decisions"
        return "turning operational data into decisions the team can act on"
    if lane_key == "corporate_strategy":
        if re.search(r"\b(executive objectives|roadmap|product strategy|product accountability|stakeholder workstreams)\b", lowered):
            return "aligning executive objectives, roadmap choices, and stakeholder workstreams so delivery stays on track"
        if re.search(r"\b(program turnaround|modernization|financial management)\b", lowered):
            return "keeping modernization work aligned across strategy, oversight, and delivery"
        return "turning ambiguous client problems into structured decisions that hold up in execution"
    if lane_key == "customer_success":
        if support_escalation_role(role_title, job_description) or re.search(r"\b(escalation|billing|membership|phone|email|zendesk|vip|partner|human touch)\b", lowered):
            if re.search(r"\b(ai-assisted|ai-driven|human touch)\b", lowered):
                return "handling retention handoffs and complex member issues with sound judgment"
            return "handling complex member issues and escalation risk with sound judgment"
        return "protecting customer health, adoption, and renewal confidence"
    if lane_key == "change_enablement":
        if re.search(r"\b(organiz(?:ation|ational) design|restructuring|future state|transition planning)\b", lowered):
            return "turning organizational change into clearer roles, transition plans, and durable follow-through"
        if re.search(r"\b(leadership development|team effectiveness|coaching|workshops?|training)\b", lowered):
            return "turning transformation goals into leadership alignment, team effectiveness, and adoption"
        return "turning change initiatives into stakeholder clarity, adoption, and measurable follow-through"
    if lane_key == "presales_solution":
        return "turning buyer questions into credible solution decisions"
    if lane_key == "implementation_delivery":
        if re.search(r"\brequirements documentation\b|\buser stories\b|\bacceptance criteria\b|\brequirements status\b", lowered):
            return "turning requirements documentation, status reporting, and user-story coordination into client-ready delivery"
        if re.search(r"\bproject manager\b|\bdeputy project manager\b|\bscrum master\b|\bproduct owner\b", lowered):
            return "turning multi-workstream requirements, project coordination, and stakeholder updates into reliable delivery"
        if re.search(r"\b(scoping|requirements gathering|statement[s]? of work|sow|integration|integrations|data migration|api)\b", lowered):
            return "turning technical scoping, integrations, and data-migration work into successful delivery"
        return "moving implementation work from ambiguity to adoption"
    if lane_key == "process_improvement":
        return "turning root-cause analysis into measurable process gains"
    return safe_core_problem(f"solving the core problem behind the {role_title} role") or "solving the team's core problem"


def extract_cover_letter_terms(job_description: str, lane_key: str, sections: dict[str, str]) -> tuple[str, ...]:
    ordered_sources = (
        sections.get("who_you_are", ""),
        sections.get("what_youll_do", ""),
        sections.get("intro", ""),
        job_description,
    )
    preferred_patterns: tuple[tuple[str, str], ...] = ()
    if lane_key == "analytics_operations":
        preferred_patterns = _COVER_ANALYTICS_TERM_PATTERNS
    elif lane_key == "change_enablement":
        preferred_patterns = _COVER_CHANGE_TERM_PATTERNS
    elif lane_key == "implementation_delivery":
        preferred_patterns = _COVER_IMPLEMENTATION_TERM_PATTERNS
    elif lane_key == "corporate_strategy":
        preferred_patterns = _COVER_STRATEGY_TERM_PATTERNS
    elif lane_key == "customer_success":
        if re.search(r"\b(customer experience|support|escalation|billing|membership|zendesk|phone|email|vip|human touch)\b", job_description, re.I):
            preferred_patterns = _COVER_SUPPORT_TERM_PATTERNS
        else:
            preferred_patterns = _COVER_CS_ENTERPRISE_TERM_PATTERNS
    patterns = (*preferred_patterns, *_GENERIC_COVER_TERM_PATTERNS)
    found: list[str] = []
    for source in ordered_sources:
        for label, pattern in patterns:
            if label in found:
                continue
            if re.search(pattern, source, re.I):
                found.append(label)
            if len(found) >= 5:
                return tuple(found)
    if not found:
        fallbacks = [term for term in build_resume.visible_role_specialties(job_description) if term]
        found.extend(fallbacks[:3])
    return tuple(found[:5])


def extract_top_accomplishment(job_description: str, resume_text: str) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    role_title = build_resume.extract_job_title(job_description) or ""
    lane_key = effective_lane_key(role_title, job_description, profile)
    support_role = support_escalation_role(role_title, job_description)
    if lane_key == "analytics_operations" and is_education_assessment_context(job_description):
        def education_measurement_line_score(line: str) -> tuple[int, int, int, int, int]:
            return (
                1 if re.search(r"\d|%|\$", line) else 0,
                1 if re.search(r"\b(dashboard|reporting|kpi|sql|etl|power bi|crystal reports)\b", line, re.I) else 0,
                1 if re.search(r"\b(validat(?:e|ed|ion)?|quality|decision(?:s|-making)?)\b", line, re.I) else 0,
                1 if re.search(r"\b(ai-assisted|workflow|continuous improvement)\b", line, re.I) else 0,
                -len(line),
            )

        resume_lines = build_resume.experience_bullet_texts_from_text(resume_text)
        if not resume_lines:
            resume_lines = [
                re.sub(r"\s+", " ", line).strip()
                for line in resume_text.splitlines()
                if re.sub(r"\s+", " ", line).strip()
            ]
        preferred_lines = sorted(
            (
                line
                for line in resume_lines
                if re.search(
                    r"\b(dashboard|reporting|kpi|sql|etl|power bi|crystal reports|validation|decision)\b",
                    line,
                    re.I,
                )
            ),
            key=education_measurement_line_score,
            reverse=True,
        )
        if preferred_lines:
            accomplishment = preferred_lines[0].strip().rstrip(".")
            accomplishment = re.sub(r"\bmore than 200\b", "200+", accomplishment, flags=re.I)
            return accomplishment[:1].upper() + accomplishment[1:] + "."
    items = safe_selected_evidence_items_ordered(job_description, resume_text)
    if items and lane_key == "analytics_operations" and is_education_assessment_context(job_description):
        ranked = sorted(
            items,
            key=lambda item: (
                1 if re.search(r"\d|%|\$", item.text) else 0,
                1 if any(term in " ".join(item.signals).lower() for term in ("analytics", "reporting", "data", "dashboard", "kpi", "business intelligence", "sql")) else 0,
                1 if any(term in " ".join(item.signals).lower() for term in ("ai", "automation", "workflow")) else 0,
                1 if re.search(r"\b(validat(?:e|ed|ion)?|quality|decision(?:s|-making)?)\b", item.text, re.I) else 0,
                -len(item.text),
            ),
            reverse=True,
        )
        accomplishment = ranked[0].text if ranked else ""
        accomplishment = re.sub(r"\bmore than 200\b", "200+", accomplishment, flags=re.I)
        accomplishment = compress_cover_evidence_text(accomplishment)
        if accomplishment.count(",") >= 3:
            accomplishment = compact_cover_sentence(accomplishment, max_words=22).rstrip(".")
        return accomplishment[:1].upper() + accomplishment[1:] + ("." if accomplishment and not accomplishment.endswith(".") else "")
    if items:
        def score(item: EvidenceBullet) -> tuple[int, int, int]:
            signal_text = " ".join(item.signals).lower()
            lane_bonus = 0
            if lane_key == "analytics_operations" and any(term in signal_text for term in ("analytics", "reporting", "sql", "dashboard", "kpi", "decision")):
                lane_bonus = 6
            if lane_key == "implementation_delivery" and any(
                term in signal_text
                for term in (
                    "implementation",
                    "requirements",
                    "technical scoping",
                    "statement of work",
                    "sow",
                    "integration",
                    "data migration",
                    "testing",
                    "delivery",
                    "uat",
                    "solution architecture",
                )
            ):
                lane_bonus = max(lane_bonus, 8)
            if support_role and any(term in signal_text for term in ("retention", "revenue", "escalation", "risk", "account", "customer success")):
                lane_bonus = max(lane_bonus, 7)
            elif support_role and any(term in signal_text for term in ("support operations", "customer operations", "customer support", "customer experience", "service quality")):
                lane_bonus = max(lane_bonus, 5)
            quant_bonus = 3 if re.search(r"\d|%|\$", item.text) else 0
            return (lane_bonus + signal_score(job_description, item.signals) + quant_bonus, len(item.signals), -len(item.text))
        ranked = sorted(items, key=score, reverse=True)
        accomplishment = ranked[0].text if ranked else ""
        accomplishment = re.sub(r"\bmore than 200\b", "200+", accomplishment, flags=re.I)
        accomplishment = compress_cover_evidence_text(accomplishment)
        if accomplishment.count(",") >= 3:
            accomplishment = compact_cover_sentence(accomplishment, max_words=22).rstrip(".")
        return accomplishment[:1].upper() + accomplishment[1:] + ("." if accomplishment and not accomplishment.endswith(".") else "")

    resume_lines = [
        re.sub(r"\s+", " ", line).strip()
        for line in resume_text.splitlines()
        if re.sub(r"\s+", " ", line).strip()
    ]
    verb_bonus_pattern = r"\b(?:built|enabled|improved|reduced|stabilized|launched|led|protected|delivered)\b"
    analytics_pattern = r"\b(?:analytics?|reporting|sql|dashboard|dashboards|kpi|bi|cohort|retention|subscription|decision)\b"
    action_first_lines = [
        line
        for line in resume_lines
        if re.search(rf"^(?:{verb_bonus_pattern[2:-2]})\b", line, re.I)
    ]
    candidate_lines = action_first_lines or resume_lines
    if lane_key == "analytics_operations":
        analytics_candidate_lines = [
            line
            for line in candidate_lines
            if re.search(analytics_pattern, line, re.I)
        ]
        if analytics_candidate_lines:
            candidate_lines = analytics_candidate_lines
    ranked_lines = sorted(
        candidate_lines,
        key=lambda line: (
            8 if re.search(rf"^(?:{verb_bonus_pattern[2:-2]})\b", line, re.I) else 0,
            4 if re.search(r"\d|%|\$", line) else 0,
            5 if lane_key == "analytics_operations" and re.search(analytics_pattern, line, re.I) else 0,
            -6 if re.search(r"\b\d+\+ years\b", line, re.I) else 0,
            len(re.findall(analytics_pattern, line, re.I)),
            -len(line),
        ),
        reverse=True,
    )
    accomplishment = ranked_lines[0] if ranked_lines else ""
    accomplishment = re.sub(r"\bmore than 200\b", "200+", accomplishment, flags=re.I)
    accomplishment = compress_cover_evidence_text(accomplishment)
    if accomplishment.count(",") >= 3:
        accomplishment = compact_cover_sentence(accomplishment, max_words=22).rstrip(".")
    return accomplishment[:1].upper() + accomplishment[1:] + ("." if accomplishment and not accomplishment.endswith(".") else "")


def communication_metric_matches_avoid_text(candidate: str, avoid_text: str) -> bool:
    if not avoid_text.strip():
        return False
    if (
        re.search(r"\bworkshops?\b", candidate, re.I)
        and re.search(r"\bworkshops?\b", avoid_text, re.I)
        and re.search(r"\bqbrs?\b", candidate, re.I)
        and re.search(r"\bqbrs?\b", avoid_text, re.I)
    ):
        return True
    normalized_candidate = build_resume.normalize_compare(candidate)
    normalized_avoid = build_resume.normalize_compare(avoid_text)
    replacements = (
        (r"\bmore than (\d+)\b", r"\1+"),
        (r"\bfacilitated\b", "led"),
    )
    for pattern, replacement in replacements:
        normalized_candidate = re.sub(pattern, replacement, normalized_candidate, flags=re.I)
        normalized_avoid = re.sub(pattern, replacement, normalized_avoid, flags=re.I)
    return normalized_candidate == normalized_avoid or cover_sentences_near_duplicate(candidate, avoid_text)


def finalize_communication_metric_candidate(candidate: str, avoid_text: str = "") -> str:
    if not candidate.strip():
        return ""
    line = candidate.strip().rstrip(".")
    line = re.sub(r"\bmore than 60\b", "60+", line, flags=re.I)
    line = re.sub(r"\bfacilitated\b", "Led", line, flags=re.I)
    line = re.sub(
        r"^Built 200\+ dashboards, KPI reporting tools, and curated reporting views, supported 80\+ client engagements, and led ",
        "Led ",
        line,
        flags=re.I,
    )
    compacted = compact_cover_sentence(line, max_words=20)
    compacted = re.sub(r"\b(?:that|which|including|using)\.$", ".", compacted, flags=re.I)
    compacted = re.sub(r",\s*\.$", ".", compacted)
    compacted = re.sub(r"\s+\.", ".", compacted)
    if communication_metric_matches_avoid_text(compacted, avoid_text):
        return ""
    if not cover_support_sentence_is_usable(compacted):
        return ""
    return compacted


def extract_communication_metric(
    job_description: str,
    resume_text: str,
    avoid_text: str = "",
    *,
    lane_key: str | None = None,
    role_title: str = "",
) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    resolved_lane_key = lane_key or effective_lane_key(
        role_title or (build_resume.extract_job_title(job_description) or ""),
        job_description,
        profile,
    )
    if is_education_assessment_context(job_description):
        candidate = "Led 60+ executive workshops and QBRs across the Americas, Europe, and Asia."
        finalized = finalize_communication_metric_candidate(candidate, avoid_text)
        if finalized:
            return finalized

    if (
        resolved_lane_key == "customer_success"
        and re.search(r"\b(customer experience|customer support|support team|escalation|membership|billing|zendesk|phone|email|vip|human touch)\b", job_description, re.I)
    ):
        if re.search(r"\bservice gaps\b", resume_text, re.I) and re.search(r"\bLivePerson\b", resume_text, re.I):
            candidate = "Translated LivePerson and Salesforce interaction data into support insights that surfaced service gaps and recurring customer questions."
            finalized = finalize_communication_metric_candidate(candidate, avoid_text)
            if finalized:
                return finalized
        if re.search(r"\bSMS support\b", resume_text, re.I) or re.search(r"\bchat/text\b", resume_text, re.I):
            candidate = "Supported a zero-to-one SMS support channel by surfacing customer friction themes and workflow gaps."
            finalized = finalize_communication_metric_candidate(candidate, avoid_text)
            if finalized:
                return finalized

    if re.search(r"60\+[^.\n]{0,120}\bworkshops?\b[^.\n]{0,120}\bqbrs?\b", resume_text, re.I):
        geography = ""
        if (
            re.search(r"\b(?:the Americas|North America)\b", resume_text, re.I)
            and re.search(r"\bEurope\b", resume_text, re.I)
            and re.search(r"\bAsia\b", resume_text, re.I)
        ):
            geography = " across the Americas, Europe, and Asia"
        candidate = f"Led 60+ executive workshops and QBRs{geography}."
        finalized = finalize_communication_metric_candidate(candidate, avoid_text)
        if finalized:
            return finalized

    items = safe_selected_evidence_items_ordered(job_description, resume_text)
    ranked = sorted(
        items,
        key=lambda item: (
            1 if resolved_lane_key == "change_enablement" and any(term in " ".join(item.signals).lower() for term in ("training", "adoption", "stakeholder", "operations")) else 0,
            1 if any(term in " ".join(item.signals).lower() for term in ("executive", "workshop", "stakeholder", "business outcomes", "decision")) else 0,
            1 if re.search(r"\d|%|\$", item.text) else 0,
            signal_score(job_description, item.signals),
        ),
        reverse=True,
    )
    for item in ranked:
        finalized = finalize_communication_metric_candidate(item.text[:1].upper() + item.text[1:], avoid_text)
        if finalized:
            return finalized

    resume_lines = [
        re.sub(r"\s+", " ", line).strip()
        for line in resume_text.splitlines()
        if re.sub(r"\s+", " ", line).strip()
    ]
    ranked_lines = sorted(
        (
            line
            for line in resume_lines
            if re.search(r"\b(executive|workshop|qbr|stakeholder|written|verbal|presenting|communicat)\b", line, re.I)
        ),
        key=lambda line: (
            1 if re.search(r"\d|%|\$", line) else 0,
            1 if re.search(r"\b(executive|workshop|qbr|presenting|written|verbal)\b", line, re.I) else 0,
            len(line),
        ),
        reverse=True,
    )
    if ranked_lines:
        selected_line = next(
            (
                line
                for line in ranked_lines
                if not communication_metric_matches_avoid_text(line, avoid_text)
            ),
            "",
        )
        if not selected_line:
            return ""
        finalized = finalize_communication_metric_candidate(selected_line, avoid_text)
        if finalized:
            return finalized
    if re.search(r"\b80\+\b[^.\n]{0,120}\bclient engagements?\b", resume_text, re.I):
        return finalize_communication_metric_candidate(
            "Supported 80+ client engagements with structured stakeholder communication.",
            avoid_text,
        )
    return ""


def extract_ambiguity_process(job_description: str, lane_key: str | None = None) -> str:
    resolved_lane_key = lane_key or build_resume.job_problem_profile(job_description).primary_lane
    if resolved_lane_key == "analytics_operations":
        return "trace the workflow, validate the data sources, define the right question, and document the pattern so the analysis can be reused"
    if resolved_lane_key == "customer_success":
        if re.search(r"\b(customer experience|customer support|support team|escalation|membership|billing|zendesk|phone|email|vip|human touch)\b", job_description, re.I):
            return "trace the customer workflow, validate the risk, and turn the issue into a clear next step"
        return "trace the customer workflow, validate the signal behind the risk, and turn the issue into a plan owners can act on"
    if resolved_lane_key == "change_enablement":
        return "trace the stakeholder workflow, validate what has to change, and document the ownership, communication, and reinforcement needed for adoption"
    if resolved_lane_key == "presales_solution":
        return "trace the buyer workflow, validate the constraint, and document the solution path before the demo outruns delivery reality"
    if resolved_lane_key == "implementation_delivery":
        return "validate the workflow, dependencies, and testing path early so go-live risk does not compound"
    if resolved_lane_key == "process_improvement":
        return "trace the current-state workflow, validate the root cause, and document the fix so the new standard can stick"
    return "trace the workflow, validate the sources, and document the pattern so the work becomes easier to act on"


def extract_pain_area(job_description: str, lane_key: str, terms: tuple[str, ...], environments: tuple[str, ...]) -> str:
    lowered = job_description.lower()
    if lane_key == "customer_success":
        if re.search(r"\b(escalation|membership|billing)\b", lowered) and re.search(r"\bretention|loyalty\b", lowered):
            return "complex member issues, escalation risk, and retention pressure"
        if re.search(r"\bescalation|support\b", lowered):
            return "complex support issues and escalation risk"
        if "retention" in lowered and "loyalty" in lowered:
            return "retention and loyalty issues"
    if lane_key == "analytics_operations":
        if is_education_assessment_context(job_description):
            return "AI-assisted workflow validation, quality monitoring, and continuous improvement"
        if re.search(r"\b(customer support|support operations)\b", lowered) and re.search(r"\bworkflow|workflows|reporting|testing\b", lowered):
            return "customer support workflows, reporting visibility, and testing follow-through"
        if "retention" in lowered and "lifecycle" in lowered:
            return "retention and lifecycle data"
        if "retention" in lowered:
            return "retention friction"
        if "subscription" in lowered:
            return "subscription decisions"
    if lane_key == "change_enablement":
        if re.search(r"\b(organiz(?:ation|ational) design|restructuring|future state|transition planning)\b", lowered):
            return "organizational design, transition planning, and change adoption"
        if re.search(r"\b(leadership development|team effectiveness|coaching|training|workshops?)\b", lowered):
            return "leadership development, team effectiveness, and stakeholder adoption"
    if lane_key == "implementation_delivery":
        if re.search(r"\b(integration|integrations|api|data migration|migrations)\b", lowered) and re.search(r"\b(scoping|requirements|testing|delivery)\b", lowered):
            return "technical scoping, integrations, data migration, and delivery risk"
        if re.search(r"\b(integration|integrations|api)\b", lowered):
            return "integrations and delivery risk"
        if re.search(r"\b(schedule|scope|milestones|risks|dependencies|status reports?|workstreams?)\b", lowered):
            return "schedule control, stakeholder alignment, and delivery risk"
    if environments:
        return f"{comma_join(environments)} priorities"
    if terms:
        return comma_join(terms[:2])
    return build_resume.role_specialty_phrase(job_description, "the team's highest-priority work")


def build_cover_letter_signals(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
) -> CoverLetterSignals:
    sections = cover_letter_jd_sections(job_description)
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_key = effective_lane_key(role_title, job_description, profile)
    test_environments = extract_test_environments(job_description, sections)
    jd_skill_terms = extract_cover_letter_terms(job_description, lane_key, sections)
    top_accomplishment = extract_top_accomplishment(job_description, resume_text)
    try:
        fit_bridge = anticipated_fit_bridge(company_name, role_title, job_description, resume_text)
    except SystemExit:
        fit_bridge = ""
    return CoverLetterSignals(
        company_mission=extract_company_mission(company_name, sections, job_description),
        role_core_function=extract_role_core_function(company_name, role_title, job_description, sections),
        top_accomplishment=top_accomplishment,
        fit_bridge=fit_bridge,
        jd_skill_terms=jd_skill_terms,
        ambiguity_process=extract_ambiguity_process(job_description, lane_key),
        jd_test_environments=test_environments,
        communication_metric=extract_communication_metric(
            job_description,
            resume_text,
            avoid_text=top_accomplishment,
            lane_key=lane_key,
            role_title=role_title,
        ),
        partner_functions=extract_partner_functions(job_description, sections),
        jd_pain_area=extract_pain_area(job_description, lane_key, jd_skill_terms, test_environments),
    )


def cover_letter_plan_warnings(sections: dict[str, str], signals: CoverLetterSignals, mode: str) -> tuple[str, ...]:
    warnings: list[str] = []
    normalized_mode = normalize_cover_mode(mode)
    if not sections.get("intro") and not signals.company_mission:
        warnings.append("No clean company-intro line detected; the opening used a role-context fallback.")
    if not signals.jd_skill_terms:
        warnings.append("No specialty terms were extracted from the job description; the proof paragraph used broader lane evidence.")
    if normalized_mode == LONG_COVER_MODE and not (signals.partner_functions or signals.jd_test_environments):
        warnings.append("The JD did not expose partner-team or test-environment detail; the long cover used more general workflow language.")
    return tuple(warnings)


def cover_letter_plan_blockers(role_title: str, signals: CoverLetterSignals) -> tuple[str, ...]:
    blockers: list[str] = []
    if not signals.top_accomplishment.strip():
        blockers.append("No reliable proof sentence could be extracted from the resume for the cover letter.")
    if looks_like_cover_section_label(signals.company_mission):
        blockers.append("Company mission extraction resolved to a raw JD section header instead of company context.")
    if looks_like_cover_section_label(signals.role_core_function):
        blockers.append("Role-core-function extraction resolved to a raw JD section header instead of real work.")
    if looks_like_cover_section_label(signals.top_accomplishment) or has_bullet_artifact(signals.top_accomplishment):
        blockers.append("Top accomplishment extraction produced raw JD or note artifacts instead of clean proof.")
    return tuple(blockers)


def expand_short_cover_opening(opening: str, signals: CoverLetterSignals) -> str:
    if word_count(opening) >= 30:
        return opening
    signal_text = " ".join(signals.jd_skill_terms).lower()
    if any(term in signal_text for term in ("requirements", "testing", "implementation", "integration", "configuration")):
        supplement = "The role depends on clear client requirements, system setup, and follow-through staying aligned."
    elif any(term in signal_text for term in ("sql", "analytics", "reporting", "measurement", "retention", "lifecycle")):
        supplement = "The role depends on clean measurement, usable reporting, and decisions people can trust."
    else:
        supplement = "The role depends on clear ownership, stakeholder alignment, and follow-through."
    expanded = f"{opening} {ensure_sentence(supplement)}".strip()
    if word_count(expanded) >= 30:
        return expanded
    if any(term in signal_text for term in ("buyer", "discovery", "demo", "solution", "presales")):
        extra = "It also needs recommendations buyers and internal teams can trust once the conversation turns practical."
    elif any(term in signal_text for term in ("customer", "retention", "adoption", "renewal", "escalation")):
        extra = "The role only works when customer context, internal alignment, and next steps stay clear at the same time."
    else:
        extra = "The role only works when the context is clear, the decision path is usable, and ownership stays visible."
    return f"{expanded} {ensure_sentence(extra)}".strip()


def build_cover_letter_plan(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
    application_responses: tuple[question_prep.QualificationsResponse, ...] = (),
    force_bridge: bool = False,
) -> CoverLetterPlan:
    normalized_mode = normalize_cover_mode(mode)
    sections = cover_letter_jd_sections(job_description)
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_key = effective_lane_key(role_title, job_description, profile)
    prose_role_title = normalize_cover_role_title_for_prose(role_title)
    signals = build_cover_letter_signals(company_name, role_title, job_description, resume_text)
    salutation = preferred_salutation(company_name)
    body_paragraphs, proof_terms, close_terms, selection_debug = compose_question_driven_standard_paragraphs(
        company_name,
        prose_role_title or role_title,
        job_description,
        resume_text,
        signals,
        mode=normalized_mode,
        original_role_title=role_title,
        application_responses=application_responses,
        force_bridge=force_bridge or normalized_mode == LONG_COVER_MODE,
    )
    if not build_resume.jd_explicitly_requires_erp(job_description):
        body_paragraphs = tuple(
            question_prep.scrub_cover_answer_for_job(job_description, paragraph)
            for paragraph in body_paragraphs
        )
    opening = body_paragraphs[0] if body_paragraphs else ""
    proof = body_paragraphs[1] if len(body_paragraphs) > 1 else ""
    bridge = body_paragraphs[2] if len(body_paragraphs) > 2 else ""
    workflow = ""
    communication = ""
    closing = ""
    draft_style = "question_driven_standard"

    return CoverLetterPlan(
        company_name=company_name,
        role_title=role_title,
        prose_role_title=prose_role_title,
        job_description=job_description,
        mode=normalized_mode,
        draft_style=draft_style,
        lane_key=lane_key,
        sections=tuple((key, value) for key, value in sections.items()),
        signals=signals,
        salutation=salutation,
        body_paragraphs=body_paragraphs,
        opening=opening,
        proof=proof,
        bridge=bridge,
        workflow=workflow,
        communication=communication,
        closing=closing,
        proof_mapped_terms=proof_terms,
        close_mapped_terms=close_terms,
        warnings=cover_letter_plan_warnings(sections, signals, normalized_mode),
        blockers=cover_letter_plan_blockers(role_title, signals),
        selection_debug=selection_debug,
    )


def compose_cover_letter_question_driven_from_plan(plan: CoverLetterPlan) -> CoverLetterDraft:
    if plan.blockers:
        fail("Cover-letter planning failed before prose generation: " + "; ".join(plan.blockers))

    if plan.draft_style == "question_driven_standard":
        paragraphs = fit_standard_cover_word_budget(
            plan.salutation,
            plan.body_paragraphs,
            plan.company_name,
            plan.prose_role_title,
            plan.job_description,
            mode=plan.mode,
        )
        return CoverLetterDraft(
            salutation=plan.salutation,
            body_paragraphs=paragraphs,
            signals=plan.signals,
            paragraph_shape=len(paragraphs),
            proof_mapped_terms=plan.proof_mapped_terms,
            close_mapped_terms=plan.close_mapped_terms,
            mode=plan.mode,
        )

    if plan.mode == LONG_COVER_MODE:
        paragraphs = (
            plan.opening,
            plan.proof,
            plan.workflow,
            plan.communication,
            plan.closing,
        )
        return CoverLetterDraft(
            salutation=plan.salutation,
            body_paragraphs=paragraphs,
            signals=plan.signals,
            paragraph_shape=5,
            proof_mapped_terms=plan.proof_mapped_terms,
            close_mapped_terms=plan.close_mapped_terms,
            mode=LONG_COVER_MODE,
        )

    four_paragraphs = (
        plan.opening,
        plan.proof,
        " ".join(part for part in (plan.bridge, plan.workflow) if part).strip(),
        f"{plan.communication} {plan.closing}".strip(),
    )
    bridge_required = any(
        warning.startswith("Resume bridge gaps to address honestly:")
        for warning in plan.warnings
    )
    compact_third_paragraph = " ".join(
        part
        for part in ((plan.bridge if bridge_required else ""), plan.workflow, plan.communication, plan.closing)
        if part
    ).strip()
    three_paragraphs = (
        plan.opening,
        plan.proof,
        compact_third_paragraph,
    )

    preferred_min, preferred_max = preferred_cover_word_range(STANDARD_COVER_MODE)
    four_words = cover_letter_document_word_count(plan.salutation, four_paragraphs)
    three_words = cover_letter_document_word_count(plan.salutation, three_paragraphs)
    has_distinct_middle = bool(
        plan.signals.jd_test_environments or plan.signals.partner_functions or plan.signals.communication_metric
    )
    if bridge_required:
        chosen = four_paragraphs
        shape = 4
    elif has_distinct_middle and preferred_min <= four_words <= preferred_max:
        chosen = four_paragraphs
        shape = 4
    else:
        chosen = three_paragraphs
        shape = 3
    chosen_words = cover_letter_document_word_count(plan.salutation, chosen)
    if chosen_words < preferred_min and four_words <= STANDARD_MAX_LETTER_WORDS:
        chosen = four_paragraphs
        shape = 4
    if cover_letter_document_word_count(plan.salutation, chosen) > STANDARD_MAX_LETTER_WORDS:
        chosen = three_paragraphs
        shape = 3
    chosen = fit_standard_cover_word_budget(
        plan.salutation,
        chosen,
        plan.company_name,
        plan.prose_role_title,
        plan.job_description,
    )
    shape = len(chosen)

    return CoverLetterDraft(
        salutation=plan.salutation,
        body_paragraphs=tuple(chosen),
        signals=plan.signals,
        paragraph_shape=shape,
        proof_mapped_terms=plan.proof_mapped_terms,
        close_mapped_terms=plan.close_mapped_terms,
        mode=STANDARD_COVER_MODE,
    )


def compose_cover_letter_from_plan(plan: CoverLetterPlan) -> CoverLetterDraft:
    return compose_cover_letter_question_driven_from_plan(plan)


def cover_letter_banned_phrase_issues(text: str) -> list[str]:
    banned_phrases = (
        "my closest match is",
        "i would ramp honestly",
        "the role sits close to",
        "it is a strong match with how i have worked",
        "note:",
        "relevant to this role",
        "tailored to this role",
        "i am writing to apply",
        "i am excited to apply",
        "please find my application",
    )
    lowered = text.lower()
    return [f"cover letter contains banned phrase '{phrase}'" for phrase in banned_phrases if phrase in lowered]


def cover_letter_signals_from_brief(brief: question_prep.PositioningBrief) -> CoverLetterSignals:
    top_accomplishment = brief.selected_proof_sentences[0] if brief.selected_proof_sentences else brief.role_problem_phrase
    communication_metric = brief.top_proof_anchors[0] if brief.top_proof_anchors else ""
    jd_terms = tuple(brief.strongest_direct_proofs[:3])
    return CoverLetterSignals(
        company_mission=brief.company_specific_fact or brief.mission_or_context or brief.role_problem_phrase,
        role_core_function=brief.role_problem_phrase,
        top_accomplishment=top_accomplishment,
        fit_bridge=brief.strongest_bridge_theme or brief.role_problem_phrase,
        jd_skill_terms=jd_terms,
        ambiguity_process=brief.selected_proof_sentences[1] if len(brief.selected_proof_sentences) > 1 else "",
        jd_test_environments=tuple(brief.top_proof_anchors[:2]),
        communication_metric=communication_metric,
        partner_functions=tuple(brief.strongest_direct_proofs[1:3]),
        jd_pain_area=brief.role_problem_phrase,
    )


def proof_first_gap_sentence(brief: question_prep.PositioningBrief) -> str:
    if not brief.gap_honesty_boundary:
        return ""
    core_problem_phrase = sentence_safe_role_core_problem(brief.role_problem_phrase)
    gap_text = re.sub(r"^The area where I am building rather than established is\s+", "", brief.gap_honesty_boundary, flags=re.I).rstrip(".")
    return ensure_sentence(
        f"I can create immediate value around {core_problem_phrase} while continuing to build depth in {gap_text}."
    )


def friendly_direct_proof_phrase(brief: question_prep.PositioningBrief) -> str:
    phrase_map = {
        "Change Adoption and Enablement": "role-based enablement and adoption work",
        "Client-Side ERP Ownership and Finance Partnership": "client-side ERP ownership with finance partnership",
        "Implementation Delivery": "implementation delivery",
        "Project-Based ERP Delivery": "project-based ERP delivery",
        "Customer and Revenue Outcomes": "customer outcome ownership",
        "Customer Success and Retention": "customer success and retention work",
        "Analytics and Decision Support": "analytics and decision-support work",
    }
    phrases = [
        phrase_map.get(proof, proof.lower())
        for proof in brief.strongest_direct_proofs[:2]
        if phrase_map.get(proof, proof.lower())
    ]
    if not phrases:
        return "direct implementation proof"
    if len(phrases) == 1:
        return phrases[0]
    if phrases[0] == "client-side ERP ownership with finance partnership":
        return f"client-side ERP ownership, finance partnership, and {phrases[1]}"
    return f"{phrases[0]} and {phrases[1]}"


def jd_concrete_hook(job_description: str) -> str:
    candidates: list[tuple[int, int, str]] = []
    for raw_line in build_resume.role_requirement_text(job_description).splitlines():
        line = re.sub(r"^\s*(?:[-*•]|\d+[.)])\s*", "", raw_line).strip()
        line = re.sub(r"\s+", " ", line)
        if len(line.split()) < 6:
            continue
        lowered = line.lower()
        if any(
            phrase in lowered
            for phrase in (
                "is looking for",
                "we pride ourselves",
                "we want to talk",
                "most successful team members",
                "came out the other side",
                "what makes a great fit",
            )
        ):
            continue
        if re.match(r"^(?:company|job title|what we're looking for|what makes a great fit here)\b", lowered):
            continue
        score = 0
        if re.match(r"^(?:lead|gather|manage|deliver|develop|troubleshoot|stay|collaborate|configure|support)\b", lowered):
            score += 8
        if any(signal in lowered for signal in ("acumatica", "construction", "engineering", "project-based")):
            score += 7
        if any(signal in lowered for signal in ("go-live", "configuration", "requirements", "training", "implementation", "integrations", "smartsheet")):
            score += 6
        if "client" in lowered:
            score += 3
        if score <= 0:
            continue
        cleaned = re.sub(
            r"^(?:you(?:'ll| will)?|the role(?: will)?|this role(?: will)?|responsibilities include|responsible for)\s+",
            "",
            line,
            flags=re.I,
        ).rstrip(".")
        if cleaned:
            candidates.append((score, -len(cleaned), cleaned))
    if candidates:
        best = sorted(candidates, reverse=True)[0][2]
        return best[:1].lower() + best[1:] if best and best[0].isupper() else best
    lowered_job = job_description.lower()
    if "acumatica" in lowered_job and any(term in lowered_job for term in ("construction", "engineering")):
        return "configure Acumatica for engineering and construction clients and keep implementations on track through go-live"
    return build_resume.natural_problem_phrase(build_resume.job_problem_profile(job_description))


def proof_first_opening_paragraph(
    brief: question_prep.PositioningBrief,
    company_name: str,
    role_title: str,
) -> str:
    concrete_hook = jd_concrete_hook(read_text(JOB_DESCRIPTION))
    fallback_context_sentence = ensure_sentence(
        f"{company_name} is hiring {with_indefinite_article(role_title)} to {concrete_hook}."
    )
    context_sentence = (
        question_prep.ensure_company_named(brief.company_specific_fact, company_name)
        or question_prep.ensure_company_named(brief.mission_or_context, company_name)
        or fallback_context_sentence
    )
    if context_sentence == fallback_context_sentence:
        role_sentence = ensure_sentence(
            "The work only lands when client requirements, configuration choices, and follow-through stay aligned through go-live."
        )
    else:
        role_sentence = ensure_sentence(
            f"The work centers on {concrete_hook}."
        )
    motivation_sentence = ensure_sentence(brief.personal_reason_to_care)
    if re.search(r"\b(mission|nonprofit|values)\b", brief.employer_type, re.I):
        opening = join_answer_sentences(context_sentence, motivation_sentence, role_sentence)
    else:
        opening = join_answer_sentences(context_sentence, role_sentence, motivation_sentence)
    if word_count(opening) > 80:
        # The opening paragraph QC gate caps this at 80 words. Drop the motivation
        # sentence first since it is the least essential of the three when the
        # company-context and role-framing sentences alone already run long.
        opening = join_answer_sentences(context_sentence, role_sentence)
    return repair_cover_paragraph(opening, "cover opening")


def proof_first_proof_paragraph(brief: question_prep.PositioningBrief) -> str:
    base_sentence = brief.selected_proof_sentences[0] if brief.selected_proof_sentences else ensure_sentence(
        f"The background already includes {brief.strongest_direct_proofs[0]} in complex, cross-functional environments."
    )
    sentences = [ensure_sentence(base_sentence)]
    if brief.top_proof_anchors:
        anchor = brief.top_proof_anchors[0]
        if anchor.lower() not in base_sentence.lower():
            sentences.append(ensure_sentence(f"The scope of that work included {anchor}."))
    return join_answer_sentences(*sentences)


def proof_first_support_paragraph(brief: question_prep.PositioningBrief) -> str:
    if len(brief.selected_proof_sentences) > 1:
        return repair_cover_paragraph(
            join_answer_sentences(
            brief.selected_proof_sentences[1],
            brief.selected_proof_sentences[2] if len(brief.selected_proof_sentences) > 2 else "",
            ),
            "cover support paragraph",
        )
    if brief.strongest_bridge_theme:
        return repair_cover_paragraph(
            ensure_sentence(
            f"A useful adjacent strength is {brief.strongest_bridge_theme}, which keeps the work grounded when priorities, systems, and stakeholders all have to move together."
            ),
            "cover support paragraph",
        )
    core_problem_phrase = sentence_safe_role_core_problem(brief.role_problem_phrase)
    return repair_cover_paragraph(
        ensure_sentence(
            f"The work is strongest when {core_problem_phrase} is paired with clear ownership, communication, and follow-through."
        ),
        "cover support paragraph",
    )


def proof_first_close_paragraph(
    brief: question_prep.PositioningBrief,
    company_name: str,
    role_title: str,
    *,
    include_gap: bool = False,
) -> str:
    direct_proof_phrase = friendly_direct_proof_phrase(brief)
    sentences = [
        ensure_sentence(
            f"My background lines up directly with the {direct_proof_phrase} this {role_title} role calls for."
        ),
        ensure_sentence(
            f"I would welcome the chance to discuss where it could help {company_name} keep implementations moving with more confidence."
        ),
    ]
    if include_gap and brief.gap_honesty_boundary:
        sentences.insert(1, proof_first_gap_sentence(brief))
    return repair_cover_paragraph(join_answer_sentences(*sentences), "cover close paragraph")


def build_cover_letter_proof_first(
    brief: question_prep.PositioningBrief,
    mode: str,
    company_name: str,
    role_title: str,
    resume_audit_state: str,
    output_docx: Path,
    *,
    force_bridge: bool = False,
    job_description: str = "",
) -> CoverLetterPlan:
    normalized_mode = normalize_cover_mode(mode)
    opening = proof_first_opening_paragraph(brief, company_name, role_title)
    proof = proof_first_proof_paragraph(brief)
    support = proof_first_support_paragraph(brief)
    close = proof_first_close_paragraph(
        brief,
        company_name,
        role_title,
        include_gap=force_bridge or resume_audit_state in {"BRIDGE", "FAIL", "POOR"},
    )
    body_paragraphs = (opening, proof, support, close) if normalized_mode == LONG_COVER_MODE else (opening, proof, close)
    letter_text = "\n".join(body_paragraphs)
    banned_issues = cover_letter_banned_phrase_issues(letter_text)
    if banned_issues:
        fail("; ".join(banned_issues))
    signals = cover_letter_signals_from_brief(brief)
    return CoverLetterPlan(
        company_name=company_name,
        role_title=role_title,
        prose_role_title=role_title,
        job_description=job_description,
        mode=normalized_mode,
        draft_style="proof_first",
        lane_key=brief.primary_lane,
        sections=(),
        signals=signals,
        salutation=preferred_salutation(company_name),
        body_paragraphs=tuple(body_paragraphs),
        opening=opening,
        proof=proof,
        bridge=support,
        workflow=support if normalized_mode == LONG_COVER_MODE else "",
        communication="",
        closing=close,
        proof_mapped_terms=tuple(brief.top_proof_anchors[:2] or brief.strongest_direct_proofs[:2]),
        close_mapped_terms=(brief.role_problem_phrase,),
        warnings=(),
        blockers=(),
        selection_debug={
            "resume_audit_state": resume_audit_state,
            "proof_sentences": list(brief.selected_proof_sentences),
            "proof_anchors": list(brief.top_proof_anchors),
            "employer_type": brief.employer_type,
        },
    )


def compose_cover_letter_proof_first_from_brief(plan: CoverLetterPlan) -> CoverLetterDraft:
    # Fit the word budget for every mode, not just standard. fit_standard_cover_word_budget
    # is mode-aware (it reads max_words from cover_letter_word_range(mode)), so calling it
    # unconditionally lets long-mode letters get trimmed too instead of going straight to the
    # hard word-count gate in cover_letter_qc_issues/cover_letter_text_issues with no chance
    # to self-correct first.
    paragraphs = fit_standard_cover_word_budget(
        plan.salutation,
        plan.body_paragraphs,
        plan.company_name,
        plan.prose_role_title,
        plan.job_description,
        mode=plan.mode,
    )
    return CoverLetterDraft(
        salutation=plan.salutation,
        body_paragraphs=tuple(paragraphs),
        signals=plan.signals,
        paragraph_shape=len(paragraphs),
        proof_mapped_terms=plan.proof_mapped_terms,
        close_mapped_terms=plan.close_mapped_terms,
        mode=plan.mode,
    )


def write_cover_letter_trace(
    plan: CoverLetterPlan,
    draft: CoverLetterDraft,
    output_docx: Path,
    *,
    body_text: str = "",
    specificity_warnings: list[str] | None = None,
    cover_warnings: list[str] | None = None,
    preflight_warnings: list[str] | None = None,
    prose_report: dict[str, object] | None = None,
    failure: str = "",
) -> Path:
    COVER_TRACE_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", output_docx.stem).strip("_") or "cover_letter"
    trace_path = COVER_TRACE_DIR / f"{timestamp}_{safe_name}.json"
    snapshot_id = os.environ.get(job_context_archive.SNAPSHOT_ID_ENV, "").strip()
    snapshot_metadata = job_context_archive.current_snapshot_metadata() if snapshot_id else {}
    payload = {
        "output_docx": str(output_docx),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "job_context_snapshot": {
            "snapshot_id": snapshot_id,
            "metadata": snapshot_metadata,
        },
        "plan": asdict(plan),
        "draft": asdict(draft),
        "body_text": body_text,
        "specificity_warnings": specificity_warnings or [],
        "cover_warnings": cover_warnings or [],
        "preflight_warnings": preflight_warnings or [],
        "prose_report": prose_report or {},
        "trace_summary": {
            "snapshot_id": snapshot_id,
            "fallback_used": any("fallback" in warning.lower() for warning in plan.warnings),
            "chosen_company_fact": plan.signals.company_mission,
            "chosen_role_fact": plan.signals.role_core_function,
            "chosen_proof_bullet": plan.signals.top_accomplishment,
            "blocked_validation_rule": re.sub(r"^ERROR:\s*", "", failure).split(":", 1)[0].strip() if failure else "",
        },
        "failure_details": trace_failure_details(body_text, failure) if failure else {},
        "failure": failure,
    }
    trace_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return trace_path


def paragraph_texts(docx_path: Path) -> list[str]:
    document = Document(str(docx_path))
    texts = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if text:
            texts.append(text)
    return texts


def visible_resume_text(docx_path: Path) -> str:
    return "\n".join(paragraph_texts(docx_path))


def find_resume_output(job_description: str) -> Path:
    matches = build_resume.matching_output_files(OUTPUT_DIR, job_description, "Resume.docx")
    output_target_name = build_resume.extract_output_target_name(job_description)
    if matches:
        selected = matches[0]
        print(
            "Cover letter lookup: selected resume "
            f"'{selected.name}' (audit: {resume_analysis.output_audit_state(selected)})"
        )
        return selected

    def safe_mtime(path: Path) -> float:
        try:
            return path.stat().st_mtime
        except OSError:
            return 0.0

    print(f"Cover letter lookup: searching for resume matching '{output_target_name}'")
    output_docx_files = list(OUTPUT_DIR.glob("*.docx")) if OUTPUT_DIR.exists() else []
    print(f"Cover letter lookup: OUTPUT_DIR contains {len(output_docx_files)} .docx file(s)")
    recent = sorted(output_docx_files, key=safe_mtime, reverse=True)[:5]
    for path in recent:
        print(f"  Found: {path.name}")
    fail(
        "matching resume output not found for "
        f"{output_target_name}; run scripts/build_resume.py first"
    )


def contains_all(text: str, fragments: tuple[str, ...]) -> bool:
    normalized = text.lower()
    return all(fragment.lower() in normalized for fragment in fragments)


def signal_score(job_description: str, signals: tuple[str, ...]) -> int:
    normalized = job_description.lower()
    return sum(1 for signal in signals if signal.lower() in normalized)


def role_specific_signal_bonus(job_description: str, signals: tuple[str, ...]) -> int:
    normalized = job_description.lower()
    signal_text = " ".join(signals).lower()
    bonus = 0
    if re.search(r"\b(customer experience|customer support|support team|support role|escalation|membership|billing|zendesk|phone|email|vip|human touch)\b", normalized):
        if any(term in signal_text for term in ("retention", "revenue", "escalation", "risk", "account", "customer success")):
            bonus += 5
        if any(term in signal_text for term in ("support operations", "customer operations", "ai", "workflow automation", "chatbot")):
            bonus += 3
    if re.search(r"\b(data analyst|analytics|retention|lifecycle|sql|reporting)\b", normalized):
        if any(term in signal_text for term in ("analytics", "reporting", "data", "dashboard", "kpi", "business intelligence")):
            bonus += 4
    return bonus


def _select_from_candidates(
    candidates: tuple[EvidenceBullet, ...],
    job_description: str,
    resume_text: str,
    count: int,
    deemphasize_erp: bool,
) -> list[EvidenceBullet]:
    supported = [item for item in candidates if contains_all(resume_text, item.evidence)]
    if deemphasize_erp and not build_resume.jd_explicitly_requires_erp(job_description) and build_resume.should_deemphasize_erp_for_role(job_description):
        non_erp_supported = [
            item for item in supported
            if not re.search(r"\berp\b|enterprise resource planning", item.text, re.I)
        ]
        if len(non_erp_supported) >= count:
            supported = non_erp_supported
    lane_key = build_resume.job_problem_profile(job_description, resume_text).primary_lane
    lane_signal = {
        "implementation_delivery": {"implementation", "go-live", "requirements", "configuration", "migration"},
        "customer_success": {"retention", "renewal", "risk", "customer success", "qbr"},
        "presales_solution": {"discovery", "solution", "pre-sales", "demo"},
        "analytics_operations": {"analytics", "reporting", "dashboard", "kpi"},
        "change_enablement": {"change", "adoption", "training", "enablement"},
    }.get(lane_key, set())

    def lane_bonus(item: EvidenceBullet) -> int:
        return 3 if set(signal.lower() for signal in item.signals) & lane_signal else 0

    supported.sort(
        key=lambda item: (
            lane_bonus(item) + role_specific_signal_bonus(job_description, item.signals) + signal_score(job_description, item.signals),
        ),
        reverse=True,
    )
    if lane_signal:
        for index, item in enumerate(supported[:count]):
            if set(signal.lower() for signal in item.signals) & lane_signal:
                supported.insert(0, supported.pop(index))
                break
    if len(supported) < count:
        fail("not enough resume-supported proof points available for a complete cover letter")
    return supported[:count]


def canonical_cover_candidates() -> tuple[EvidenceBullet, ...]:
    return tuple(
        EvidenceBullet(
            text=item.claim,
            evidence=evidence_engine.canonical_evidence_support_terms(item),
            signals=item.lane_signals,
        )
        for item in evidence_engine.commercial_canonical_evidence()
    )


def evidence_bullets(job_description: str, resume_text: str) -> list[str]:
    return [item.text for item in _select_from_candidates(canonical_cover_candidates(), job_description, resume_text, 3, False)]


def selected_evidence_items(job_description: str, resume_text: str) -> list[EvidenceBullet]:
    return _select_from_candidates(canonical_cover_candidates(), job_description, resume_text, 4, True)


def concise_outcome_phrase(outcomes: tuple[str, ...], core_problem: str) -> str:
    if not outcomes:
        return "clearer delivery and adoption"
    if len(outcomes) == 1:
        return outcomes[0]
    if len(outcomes) == 2:
        return f"{outcomes[0]} and {outcomes[1]}"
    return f"{outcomes[0]}, {outcomes[1]}, and {outcomes[2]}"


def enforce_result_first_ordering(text: str) -> str:
    """
    Apply Result + Metric + Context ordering to a single proof point sentence when
    it clearly starts with context instead of outcome.
    """
    context_openers = (r"^(By |Through |When |During |After |While |Using |Across |In order to )",)
    for pattern in context_openers:
        if re.match(pattern, text, re.I):
            parts = text.split(",", 2)
            if len(parts) >= 2:
                last = parts[-1].strip()
                if re.search(r"\b(reduced|increased|improved|stabilized|delivered|built|enabled|cut)\b", last, re.I):
                    result_clause = last.rstrip(".")
                    context_clause = ", ".join(parts[:-1])
                    return f"{result_clause[0].upper()}{result_clause[1:]} by {context_clause.lower()}."
    return text


def selected_evidence_items_ordered(job_description: str, resume_text: str) -> list[EvidenceBullet]:
    """Wrapper that enforces Result + Metric + Context ordering on evidence items."""
    reordered: list[EvidenceBullet] = []
    for item in selected_evidence_items(job_description, resume_text):
        reordered.append(
            EvidenceBullet(
                text=enforce_result_first_ordering(item.text),
                evidence=item.evidence,
                signals=item.signals,
            )
        )
    return reordered


def safe_selected_evidence_items_ordered(job_description: str, resume_text: str) -> list[EvidenceBullet]:
    try:
        with contextlib.redirect_stderr(io.StringIO()):
            return selected_evidence_items_ordered(job_description, resume_text)
    except SystemExit:
        return []


def is_consulting_context(company_name: str, job_description: str) -> bool:
    context = build_resume.primary_employer_context(job_description)
    if context and str(context["key"]) == "consulting":
        return True
    role_title = extract_role_title(job_description) or ""
    firm_or_title = f"{company_name} {role_title}"
    return bool(re.search(r"\b(bain|mckinsey|bcg|deloitte|kpmg|ey|pwc|accenture|consulting firm|management consulting|strategy consulting|advisory consultant|advisory manager)\b", firm_or_title, re.I))


def bain_opening(role_title: str, company_name: str = "Bain") -> str:
    return consulting_bigfour.bigfour_cover_opening(company_name, role_title)


def is_broad_operator_context(job_description: str) -> bool:
    return build_resume.is_startup_or_broad_operator_role(job_description)


def is_early_stage_context(job_description: str) -> bool:
    hard_startup_signals = build_resume.jd_mentions(
        job_description,
        "yc-backed",
        "y combinator",
        "yc ",
        "early-stage",
        "early stage",
        "seed",
        "series a",
        "series b",
        "startup",
    )
    builder_signals = build_resume.jd_mentions(
        job_description,
        "category that didn't exist",
        "defining a category",
        "define a category",
        "from scratch",
        "build the function",
        "own a critical function",
    )
    operator_scale_signals = build_resume.jd_mentions(
        job_description,
        "high-growth",
        "maintain the implementation playbook",
        "entrepreneurial setting",
        "wearing multiple hats",
    )
    return hard_startup_signals or (builder_signals and operator_scale_signals)


def startup_stage_tension(company_name: str, job_description: str) -> str:
    compliance_workflow_context = (
        build_resume.jd_mentions(job_description, "compliance")
        and build_resume.jd_mentions(job_description, "legal", "marketing")
        and not build_resume.jd_mentions(
            job_description,
            "supply chain",
            "logistics",
            "warehouse",
            "transportation",
            "trade compliance",
            "order management",
            "manufacturing",
            "erp",
            "wms",
            "tms",
        )
    )
    if compliance_workflow_context:
        if build_resume.jd_mentions(job_description, "category that didn't exist", "defining a category", "define a category"):
            return "creating a new compliance category while helping enterprise legal, marketing, and technical teams adopt a workflow that may not already have a playbook"
        return "helping legal, marketing, and technical teams turn compliance complexity into a workflow people can actually adopt"
    if build_resume.jd_mentions(job_description, "ai-first", "ai first", "ai-powered", "ai powered"):
        return "turning an AI-first product into repeatable customer outcomes before every implementation pattern is fully settled"
    if build_resume.jd_mentions(job_description, "category that didn't exist", "defining a category", "define a category"):
        return "building a new category while giving enterprise customers enough structure to adopt something unfamiliar"
    if build_resume.jd_mentions(job_description, "signed contract", "go-live", "implementation playbook", "warm handoff"):
        return "turning signed enterprise customers into successful go-lives while the implementation motion is still being defined"
    return "moving from early traction to a repeatable operating rhythm without losing the speed and ownership that made the company work"



def broad_operator_opening(company_name: str, role_title: str, job_description: str) -> str:
    # Retained as a compatibility wrapper; the pattern selector now owns opening logic.
    return _situation_opening(company_name, role_title, job_description)


# Cover letter opening pattern library. Each pattern makes a different rhetorical move:
# situation, tension, belief, direct fit, or mission impact.
def _situation_opening(company_name: str, role_title: str, job_description: str) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    if is_early_stage_context(job_description):
        tension = startup_stage_tension(company_name, job_description)
        return (
            f"{company_name} is scaling through a stage where {tension}. "
            f"The {role_title} role exists to bring structure to that kind of operating problem, "
            "which is why my background fits here: ambiguous requirements, multiple stakeholder groups, "
            "technical questions that shape adoption, and the need to build the process while delivering the work."
        )
    specialty = build_resume.role_specialty_phrase(job_description, "complex enterprise work")
    return (
        f"{company_name} needs {specialty} to become a repeatable operating capability, not just a project. "
        f"The {role_title} role exists to close that gap. My background fits because I have done that work in environments "
        "where the requirements were ambiguous, the stakeholders did not always agree, and the outcome had to be measurable "
        "before the next decision could be made."
    )


def _tension_opening(company_name: str, role_title: str, job_description: str) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    lane_tensions = {
        "presales_solution": (
            "The hardest part of technical sales is not the demo. It is earning enough trust before the demo that the buyer "
            f"believes the solution can actually work inside their environment. The {role_title} role at "
            f"{company_name} is designed to manage, and it is the part of this work where my background is most directly relevant."
        ),
        "implementation_delivery": (
            f"The hardest part of implementation delivery is that scope, data, training, and stakeholder alignment all have to work together "
            f"for go-live to mean something. The {role_title} role at {company_name} appears to sit in that exact pressure point, "
            "and that is where my background is most directly useful."
        ),
        "customer_success": (
            f"At {company_name}, the {role_title} role begins after the sale and before renewal confidence is secure across a complex book of business. "
            "Account health, executive trust, mitigation planning, and growth all have to strengthen while customers are still deciding how much value the platform is creating. "
            "That is where my background is strongest, stabilizing complex accounts, rebuilding trust, and turning adoption into retained revenue."
        ),
        "change_enablement": (
            "Change adoption fails most often not during the announcement but during the month after, when the excitement fades "
            f"and the old habits are still easier. The {role_title} role at {company_name} sits right there: between the decision "
            "and the durable change. My background is most useful in that space."
        ),
        "analytics_operations": (
            "Reporting is only useful if it changes what someone decides to do. "
            f"The {role_title} role at {company_name} appears to exist because data is available but decisions are still slow, "
            "inconsistent, or based on the wrong signals. This is the kind of problem I have built systems, dashboards, "
            "and operating rhythms around."
        ),
        "process_improvement": (
            "Better process ideas only matter if the root cause is real, the improvement can be measured, and the people doing the work "
            f"adopt the new standard. The {role_title} role at {company_name} appears to need that practical link between analysis, "
            "change, and durable execution."
        ),
        "corporate_strategy": (
            "The most useful strategy work is not the analysis itself but the moment when the recommendation becomes something "
            f"a team can actually execute. The {role_title} role at {company_name} appears to sit at that junction: structured "
            "thinking that has to become a plan people can act on. My background connects most directly at that handoff point."
        ),
    }
    tension = lane_tensions.get(lane_key)
    if tension:
        return tension
    specialty = build_resume.role_specialty_phrase(job_description, "complex system and customer work")
    return (
        f"The {role_title} role at {company_name} lives in a specific operating tension: {specialty} has to become reliable "
        "enough to support decisions, adoption, and measurable business outcomes. This is familiar territory for me."
    )


def _belief_opening(company_name: str, role_title: str, job_description: str) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    if is_consulting_context(company_name, job_description):
        return (
            "The best consulting work lives past the recommendation. Someone has to stay in the room long enough to make the "
            "outcome real. That means clarifying the problem, aligning the people who own the work, validating the path, and measuring what changes. "
            f"This is the standard I have held myself to across ten years of enterprise delivery, and it is the reason the "
            f"{role_title} role at {company_name} is the right next conversation."
        )
    if is_broad_operator_context(job_description):
        return (
            "The most useful thing a technical operator can do is not solve the problem in front of them. It is build the system "
            f"that makes the problem solvable without them. This is the operating philosophy I would bring to the {role_title} "
            f"role at {company_name}: document the repeatable pattern, align the owners, and leave the team with a process it can "
            "keep using after the acute pressure is gone."
        )
    lens = build_resume.primary_story_lens(job_description)
    if lens:
        identity = str(lens.get("identity", "practical delivery and measurable outcomes"))
        return (
            f"The work I find most meaningful is {identity}: the kind where the result depends on getting different people to "
            f"trust the same process and keep using it after the launch moment passes. This is what the {role_title} role at "
            f"{company_name} appears to require, and it is where my background is most directly relevant."
        )
    specialty = build_resume.role_specialty_phrase(job_description, "complex enterprise work")
    return (
        "The standard I hold my work to is simple: does the output still work after I am not in the room? "
        f"At {company_name}, the {role_title} role appears to need exactly that kind of durable execution: {specialty} that "
        "becomes a process teams can trust, measure, and build on."
    )


def _direct_opening(company_name: str, role_title: str, job_description: str) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    specialty = build_resume.role_specialty_phrase(job_description, "enterprise systems and delivery work")
    supply_chain_context = bool(re.search(r"\b(supply chain|logistics?|transportation|warehousing|wms|tms)\b", job_description, re.I))
    lane_directs = {
        "analytics_operations": (
            f"{company_name} needs a {role_title} who can turn retention, lifecycle, and operational data into decisions the team can act on across member experience and subscription performance. "
            "That is the intersection of reporting, analysis, and stakeholder decision support I have spent the past decade building."
        ),
        "implementation_delivery": (
            f"The {role_title} role at {company_name} comes down to turning ambiguous requirements into working "
            f"{specialty} while keeping delivery moving when scope shifts. "
            "That is most of what I have done for the past ten years."
        ),
        "presales_solution": (
            (
                f"{company_name} needs a {role_title} who can turn complex customer requirements into a supply chain solution the buyer trusts "
                "before the deal closes. "
                if supply_chain_context
                else f"{company_name} needs a {role_title} who can turn complex customer requirements into a solution the buyer trusts before the deal closes. "
            )
            + "My background combines solution consulting, discovery, and hands-on ERP implementation experience because I can keep the conversation grounded in what the team can actually deliver."
        ),
        "customer_success": (
            f"{company_name} needs a {role_title} who can protect account health, guide adoption, and turn business reviews "
            "and escalations into clearer paths to renewal and expansion. This is the customer-facing work I have done across "
            "80+ international accounts across a $6M+ book of business, including more than one million dollars in at-risk "
            "revenue stabilized."
        ),
        "change_enablement": (
            f"{company_name} needs a {role_title} who can translate change into role clarity, adoption, and measurable follow-through after launch. "
            "My background fits that work because I have paired implementation delivery with training, documentation, and stakeholder alignment."
        ),
        "process_improvement": (
            f"{company_name} needs a {role_title} who can separate root cause from noise and turn analysis into measurable process improvement. "
            "That is the through-line in my reporting, workflow redesign, and cross-functional implementation work."
        ),
        "corporate_strategy": (
            f"{company_name} needs a {role_title} who can turn analysis into a plan leaders can act on. "
            "My background is strongest where the recommendation, the owners, and the execution path all have to become clear quickly."
        ),
    }
    direct = lane_directs.get(lane_key)
    if direct:
        return direct
    return (
        f"{company_name} needs a {role_title} who can take ambiguous {specialty} problems from unclear to structured to measurable. "
        "Ten years of implementation, customer success, and technical operations work is where that practical delivery capability comes from."
    )


def _mission_opening(company_name: str, role_title: str, job_description: str) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    role_lower = role_title.lower()
    if "strategy" in role_lower and "operations" not in role_lower:
        return (
            f"{company_name}'s mission creates a strategy question with real operating consequences: which opportunities should scale, "
            "which signals should guide the decision, and how do teams keep the Fellow experience at the center while the model grows? "
            f"That is the useful tension inside the {role_title} role, and it connects directly to the work I have done turning ambiguous "
            "business problems into operating plans, measurement rhythms, and stakeholder alignment."
        )
    if "operations" in role_lower and "transformation" not in role_lower:
        return (
            f"{company_name}'s mission has an operating test behind it: the work has to be clear enough for teams to execute, "
            "measurable enough for leaders to trust, and durable enough to improve outcomes beyond a single launch. "
            f"The {role_title} role sits in that bridge between ambition and execution, which is where my background is most useful."
        )
    if "operations" in role_lower:
        return (
            f"{company_name} is doing work where transformation only matters if the operating system can carry it. "
            "Programs, data, training, and stakeholder behavior all have to move together so more students can use the opportunity in front of them. "
            f"That is the practical challenge I see in the {role_title} role, and it is the kind of work my background is built around."
        )
    return (
        f"{company_name} is doing work that depends on something harder than a good program: it depends on whether students can "
        f"actually use the opportunity in front of them. The {role_title} role sits where systems, people, and outcomes have to "
        "move together, which is where my background is most useful: build the structure, validate the data, align the stakeholders, "
        "and make the process repeatable enough to matter at scale."
    )


def _pyramid_opening(company_name: str, role_title: str, job_description: str) -> str:
    """
    Pyramid Principle applied to the cover letter opening: conclusion first,
    then three supporting fit areas, with evidence handled in the proof paragraph.
    """
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    specialty = build_resume.role_specialty_phrase(job_description, "enterprise systems and customer delivery")
    level1_by_lane = {
        "presales_solution": (
            f"{company_name} needs a {role_title} who can turn buyer uncertainty into solution confidence before the deal closes - "
            "and that is the work I have done for more than ten years."
        ),
        "customer_success": (
            f"{company_name} needs a {role_title} who can protect account health, guide adoption, and convert escalations into retained trust - "
            "and that is the work I have done across 80+ international accounts."
        ),
        "implementation_delivery": (
            f"{company_name} needs a {role_title} who can take {specialty} work from ambiguous requirements to production-ready outcomes "
            "without losing stakeholder alignment - and that is the pattern I have repeated across more than ten years of enterprise delivery."
        ),
        "change_enablement": (
            f"{company_name} needs a {role_title} who can make change stick after the announcement - someone who stays in the room long enough "
            "to convert resistance into adoption. My background is most directly useful in that adoption gap."
        ),
        "analytics_operations": (
            f"{company_name} needs a {role_title} who can convert operational data into decisions leaders actually act on - not just reports they acknowledge. "
            "This is the work behind 200+ dashboards and KPI tools I have built."
        ),
        "process_improvement": (
            f"{company_name} needs a {role_title} who can find waste, prove root cause, and turn redesign into measurable efficiency, quality, or service gains. "
            "That is the pattern behind my process automation, reporting, and cross-functional implementation work."
        ),
    }
    conclusion = level1_by_lane.get(
        lane_key,
        f"{company_name} needs a {role_title} who can structure ambiguous {specialty} problems and keep the work tied to measurable outcomes. "
        "Ten years of implementation, customer success, and technical operations work is where that capability comes from.",
    )
    sub_points = build_resume.comma_series(tuple(list(build_resume.role_fit_checklist(profile, job_description))[:3]))
    return f"{conclusion} The three areas where that fits this role most directly are {sub_points}."


def _registered_firm_opening(company_name: str, role_title: str, job_description: str) -> str:
    registered_profile = _registered_firm_profile(company_name, role_title, job_description)
    if not registered_profile:
        return _pyramid_opening(company_name, role_title, job_description)
    _, profile_config = registered_profile
    cover_style = str(profile_config.get("cover_style", ""))
    if cover_style == "bain":
        return bain_opening(role_title, company_name)
    if cover_style == "state_farm_process":
        return (
            "A strong claims process is not just an internal efficiency exercise; it changes how quickly, clearly, and reliably a customer gets help. "
            f"That is what makes the {role_title} role at {company_name} compelling to me because it sits where Lean Six Sigma process improvement, claims process engineering, operational data, service quality, and customer experience all have to work together."
        )
    if cover_style == "belief":
        return _belief_opening(company_name, role_title, job_description)
    if cover_style == "tension":
        return _tension_opening(company_name, role_title, job_description)
    if cover_style == "direct":
        return _direct_opening(company_name, role_title, job_description)
    if cover_style == "situation":
        return _situation_opening(company_name, role_title, job_description)
    if cover_style == "mission":
        return _mission_opening(company_name, role_title, job_description)
    if cover_style == "pyramid":
        return _pyramid_opening(company_name, role_title, job_description)
    return _pyramid_opening(company_name, role_title, job_description)


def _detected_firm_opening(company_name: str, role_title: str, job_description: str) -> str:
    firm_profile = build_resume.detect_company_profile(company_name, job_description)
    if firm_profile:
        cover_style = str(firm_profile.get("cover_style", ""))
        if cover_style == "belief":
            return _belief_opening(company_name, role_title, job_description)
        if cover_style == "tension":
            return _tension_opening(company_name, role_title, job_description)
        if cover_style == "direct":
            return _direct_opening(company_name, role_title, job_description)
        if cover_style == "situation":
            return _situation_opening(company_name, role_title, job_description)
        if cover_style == "mission":
            return _mission_opening(company_name, role_title, job_description)
    return _pyramid_opening(company_name, role_title, job_description)


def _opening_pattern_options(
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
):
    lens = build_resume.primary_story_lens(job_description)
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    normalized_mode = normalize_cover_mode(mode)

    if normalized_mode == STANDARD_COVER_MODE:
        return [
            (
                "registered_firm_profile",
                lambda company, role, description: bool(_registered_firm_profile(company, role, description)),
                _registered_firm_opening,
            ),
            (
                "nonprofit_education_mission",
                lambda company, role, description: bool(lens and str(lens.get("key")) == "nonprofit_education"),
                _mission_opening,
            ),
            (
                "analytics_direct",
                lambda company, role, description: lane_key == "analytics_operations",
                _direct_opening,
            ),
            (
                "presales_direct",
                lambda company, role, description: lane_key == "presales_solution",
                _direct_opening,
            ),
            (
                "customer_success_direct",
                lambda company, role, description: lane_key == "customer_success",
                _direct_opening,
            ),
            (
                "change_direct",
                lambda company, role, description: lane_key == "change_enablement",
                _direct_opening,
            ),
            (
                "implementation_direct",
                lambda company, role, description: lane_key in {"implementation_delivery", "process_improvement"},
                _direct_opening,
            ),
            (
                "consulting_direct",
                lambda company, role, description: is_consulting_context(company, description),
                _direct_opening,
            ),
            (
                "default_direct",
                lambda company, role, description: True,
                _direct_opening,
            ),
        ]

    return [
        (
            "registered_firm_profile",
            lambda company, role, description: bool(_registered_firm_profile(company, role, description)),
            _registered_firm_opening,
        ),
        (
            "detected_firm_profile",
            # stub: implement detect_company_profile in build_resume.py before enabling this block.
            lambda company, role, description: False,
            _detected_firm_opening,
        ),
        (
            "nonprofit_education_mission",
            lambda company, role, description: bool(lens and str(lens.get("key")) == "nonprofit_education"),
            _mission_opening,
        ),
        (
            "consulting_belief",
            lambda company, role, description: is_consulting_context(company, description),
            _belief_opening,
        ),
        (
            "presales_direct",
            lambda company, role, description: lane_key == "presales_solution",
            _direct_opening,
        ),
        (
            "customer_success_tension",
            lambda company, role, description: lane_key == "customer_success",
            _tension_opening,
        ),
        (
            "early_stage_situation",
            lambda company, role, description: is_early_stage_context(description),
            _situation_opening,
        ),
        (
            "implementation_process_pyramid",
            lambda company, role, description: lane_key in {"implementation_delivery", "process_improvement"},
            _pyramid_opening,
        ),
        (
            "analytics_direct",
            lambda company, role, description: lane_key == "analytics_operations",
            _direct_opening,
        ),
        (
            "change_strategy_belief_or_tension",
            lambda company, role, description: lane_key in {"change_enablement", "corporate_strategy"},
            lambda company, role, description: (
                _belief_opening(company, role, description)
                if lens
                else _tension_opening(company, role, description)
            ),
        ),
        (
            "broad_operator_situation",
            lambda company, role, description: is_broad_operator_context(description),
            _situation_opening,
        ),
        (
            "default_pyramid",
            lambda company, role, description: True,
            _pyramid_opening,
        ),
    ]

def _select_opening_pattern_details(
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> tuple[str, str]:
    for label, condition, pattern_fn in _opening_pattern_options(company_name, role_title, job_description, mode=mode):
        if condition(company_name, role_title, job_description):
            return label, pattern_fn(company_name, role_title, job_description)
    return "default_pyramid", _pyramid_opening(company_name, role_title, job_description)


def _select_opening_pattern(
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> str:
    label, opening = _select_opening_pattern_details(company_name, role_title, job_description, mode=mode)
    debug_print(f"[cover] opening pattern selected: {label}", flag="DEBUG_COVER_LETTER")
    return opening


def selected_opening_pattern(company_name: str, role_title: str, job_description: str, *, mode: str = DEFAULT_COVER_MODE) -> str:
    try:
        return _select_opening_pattern(company_name, role_title, job_description, mode=mode)
    except TypeError as error:
        if "mode" not in str(error):
            raise
        return _select_opening_pattern(company_name, role_title, job_description)


def opening_paragraph(
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> str:
    return selected_opening_pattern(company_name, role_title, job_description, mode=mode)


def opening_reads_generic(opening: str, company_name: str, role_title: str, job_description: str) -> bool:
    generic_patterns = (
        r"\bexciting opportunity\b",
        r"\bgreat fit\b",
        r"\bstrong fit\b",
        r"\bi am excited\b",
        r"\binterested in\b",
        r"\bdrawn to\b",
        r"\bimpressed by\b",
    )
    if any(re.search(pattern, opening, re.I) for pattern in generic_patterns):
        return True

    opening_lower = opening.lower()
    context_terms = (
        "implementation", "adoption", "process", "data", "analytics", "customer",
        "client", "product", "platform", "workflow", "delivery", "claims",
        "transformation", "training", "renewal", "reporting", "go-live",
        "change", "stakeholder", "organizational", "leadership", "team",
    )
    company_in_opening = bool(company_name and re.search(re.escape(company_name), opening, re.I))
    role_in_opening = normalized_phrase_in_text(opening, role_title)
    has_context = any(term in opening_lower for term in context_terms)
    return not (company_in_opening and role_in_opening and has_context)


def moment_in_time_context(job_description: str, research_text: str = "") -> str:
    combined = f"{job_description}\n{research_text}".lower()
    if any(signal in combined for signal in ("acquisition", "integration", "migration", "cutover", "consolidation")):
        return "The role sits where system change and user adoption have to move together."
    if any(signal in combined for signal in ("scale", "scaling", "growth", "expansion", "hypergrowth", "rapidly growing")):
        return "The role sits where growth and process discipline have to scale together."
    if any(signal in combined for signal in ("turnaround", "stabilize", "stabilization", "transformation", "rebuild", "modernization")):
        return "The role sits where execution and operating clarity both need to tighten."
    return ""


def opening_keyword_hit_count(paragraph: str, job_description: str) -> int:
    if not job_description.strip():
        return 0
    return build_resume.keyword_hits(paragraph, build_resume.audit_keywords(job_description))


def opening_quality_problem(paragraph: str, company_name: str, role_title: str, job_description: str) -> str | None:
    cleaned = re.sub(r"\s+", " ", paragraph).strip()
    lowered = cleaned.lower()
    if cleaned != fix_indefinite_articles(cleaned):
        return "cover letter QC failed: opening paragraph has an article-agreement issue"
    for signal in WEAK_OPENER_SIGNALS:
        if signal in lowered:
            return f"cover letter QC failed: opening paragraph contains a weak or generic pattern: '{signal}'"
    for pattern, message in ABSTRACT_COVER_PATTERNS:
        if re.search(pattern, cleaned, re.I):
            return f"cover letter QC failed: {message}"
    opening_words = word_count(cleaned)
    if opening_words < 16 or opening_words > 80:
        return (
            f"cover letter QC failed: opening paragraph is {opening_words} words; expected 16-80"
        )
    company_in_opening = bool(company_name and re.search(re.escape(company_name), cleaned, re.I))
    role_in_opening = normalized_phrase_in_text(cleaned, role_title)
    has_context = opening_keyword_hit_count(cleaned, job_description) >= 2
    if not ((company_in_opening and role_in_opening) or (company_in_opening and has_context)):
        return (
            "cover letter QC failed: opening paragraph must be direct, concrete, and grounded in the company and role context"
        )
    role_noun = role_title.split()[-1].lower() if role_title else ""
    core_tokens = [word for word in re.findall(r"[a-z]+", lowered) if len(word) > 4]
    if role_noun and lowered.count(role_noun) >= 2:
        return "cover letter QC failed: opening restates the role title instead of a concrete hook"
    if core_tokens and len(set(core_tokens)) < max(6, len(core_tokens) // 2):
        return "cover letter QC failed: opening is circular or low-content"
    return None


def opening_method_paragraph(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    company_context_text: str = "",
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    normalized_mode = normalize_cover_mode(mode)
    opening = selected_opening_pattern(company_name, role_title, job_description, mode=normalized_mode)
    if normalized_mode == STANDARD_COVER_MODE:
        opening_sentences = sentence_list(smooth_cover_letter_text(opening))
        compact_opening = " ".join(opening_sentences[:1]).strip()
        if opening_reads_generic(compact_opening, company_name, role_title, job_description):
            compact_opening = " ".join(sentence_list(_direct_opening(company_name, role_title, job_description))[:1]).strip()
        if company_name and not re.search(re.escape(company_name), compact_opening, re.I):
            compact_opening = " ".join(sentence_list(_direct_opening(company_name, role_title, job_description))[:1]).strip()
        if role_title and not normalized_phrase_in_text(compact_opening, role_title):
            compact_opening = " ".join(sentence_list(_pyramid_opening(company_name, role_title, job_description))[:1]).strip()
        problem = opening_quality_problem(compact_opening, company_name, role_title, job_description)
        if problem:
            fail(problem)
        return compact_opening or opening

    context_sentence = company_specific_context_sentence(
        company_name,
        role_title,
        job_description,
        resume_text,
        company_context_text,
    )
    opening_core = " ".join(part for part in (opening, context_sentence) if part)
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_key = effective_lane_key(role_title, job_description, profile)
    cover_profile = active_cover_profile(company_name, role_title, job_description)
    if cover_profile:
        method = str(cover_profile.get("opening_method", "")).strip()
        if not method:
            return repair_cover_paragraph(opening_core, "cover opening paragraph")
        return repair_cover_paragraph(f"{opening_core} {method}", "cover opening paragraph")

    lens = build_resume.primary_story_lens(job_description)
    if lens and str(lens.get("key")) == "nonprofit_education":
        return opening_core

    method_sentences = {
        "presales_solution": (
            "I start by understanding the buyer's real constraint, then shape the recommendation around the workflow, data, and implementation realities behind it."
        ),
        "customer_success": (
            "I start by reading the health signals early because business reviews only matter when they turn into decisions about value, risk, and next steps."
        ),
        "change_enablement": (
            "I start with the behavior that has to change, then build the training, feedback, and ownership needed to make the change stick after launch."
        ),
        "analytics_operations": (
            "I start with the decision a leader needs to make, then work backward to the reporting and data checks that support it."
        ),
        "implementation_delivery": (
            "I keep delivery tied to adoption from the beginning so risks surface early and go-live leads to sustained use instead of a handoff."
        ),
        "corporate_strategy": (
            "I am strongest where analysis has to become an implementation plan and leaders need the tradeoffs made clear before a recommendation can matter."
        ),
    }

    if is_consulting_context(company_name, job_description):
        method = (
            "My approach to that kind of work is hands-on. I enter the ambiguity, structure the problem, earn trust quickly, "
            "and translate the recommendation into decisions, owners, and measurable next steps."
        )
        return f"{opening_core} {method}"

    if is_broad_operator_context(job_description):
        if is_early_stage_context(job_description):
            return opening_core
        method = (
            "My lane is technical operations: enough systems and data depth to diagnose issues, align owners, "
            "document decisions, and make improvements stick."
        )
        return f"{opening_core} {method}"

    method = method_sentences.get(lane_key, "")
    paragraph = f"{opening_core} {method}".strip() if method else opening_core
    problem = opening_quality_problem(paragraph, company_name, role_title, job_description)
    if problem:
        fail(problem)
    return repair_cover_paragraph(paragraph, "cover opening paragraph")

def proof_paragraph(company_name: str, job_description: str, resume_text: str) -> str:
    """
    Apply McKinsey SCR framework to the proof paragraph:
    Situation = Christian's relevant experience context;
    Complication = the problem this role exists to solve;
    Resolution = proof that maps to the complication.
    """
    role_title = build_resume.extract_job_title(job_description) or ""
    cover_profile = active_cover_profile(company_name, role_title, job_description)
    proof_mode = str(cover_profile.get("proof_mode", "")).strip() if cover_profile else ""
    if proof_mode == "bain":
        return (
            "Across more than 80 international manufacturing clients, the same client-side problems kept returning: ambiguous requirements, "
            "stakeholders with competing priorities, and operating decisions that could not wait for a perfect answer. I facilitated more than "
            "60 executive workshops and QBRs, built 200+ reporting tools that turned raw exports into usable operating decisions, and helped "
            "stabilize accounts inside a broader six-million-dollar client portfolio, including more than one million dollars in at-risk annual "
            "revenue, by clarifying ownership and resolving what the client actually needed."
        )
    if proof_mode == "state_farm":
        return (
            "The closest proof is a workflow improvement where I reduced manual inventory work by 78% and lowered discrepancies by 22%. "
            "The result came from the same discipline this role appears to require: understanding the current process, finding the failure point, building an auditable workflow, validating it with users, and measuring the result after rollout. "
            "I have also built 200+ KPI and reporting tools that turned operational data into clearer leadership decisions, and facilitated 60+ executive workshops and QBRs where the work was not just presenting information, but aligning people around what should change next. "
            "For State Farm, the relevance is claims process improvement that can be explained clearly, prioritized with cost and benefit in mind, and adopted by the teams doing the work."
        )
    if proof_mode == "construction_engineering":
        return repair_cover_paragraph(
            "The closest proof is client-side ERP ownership across five sites and 150+ users, including direct partnership with plant controllers, accounting managers, and the CFO on month-end close, audit readiness, and migration risk. "
            "I also protected migration readiness through ETL validation, SQL checks, control reviews, and cutover coordination, and stood up a new warehouse in the ERP while passing Amazon Robotics certification before go-live. "
            f"For {company_name}, that means finance-aware configuration, client-side credibility, and implementation decisions that hold up once the work reaches real projects and live operations.",
            "cover proof paragraph",
        )

    evidence = selected_evidence_items_ordered(job_description, resume_text)
    first, second, third = evidence[:3]
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_key = effective_lane_key(role_title, job_description, profile)
    situation = f"{first.text[0].upper()}{first.text[1:]}."

    complication_by_lane = {
        "presales_solution": (
            "Buyers decide whether to trust the solution before the implementation results exist, "
            "so discovery quality matters as much as demo polish."
        ),
        "customer_success": (
            "Adoption risk and renewal risk usually build quietly between business reviews, "
            "not loudly at renewal time."
        ),
        "change_enablement": (
            "Most resistance does not show up during the announcement; "
            "it shows up the month after, when the old habits are still easier."
        ),
        "implementation_delivery": (
            "Scope, data, training, and stakeholder alignment all have to work together "
            "for go-live to mean something."
        ),
        "analytics_operations": (
            "Reporting is only useful if it changes what someone decides to do."
        ),
        "process_improvement": (
            "Better ideas only matter if the root cause is real, "
            "the improvement is measurable, and the people doing the work actually adopt the new standard."
        ),
    }
    complication = complication_by_lane.get(
        lane_key,
        f"The role becomes harder when {build_resume.natural_problem_phrase(profile)} requires both technical depth and the stakeholder trust to act on the recommendation.",
    )
    if lane_key == "presales_solution":
        return (
            f"I have {first.text}. "
            f"I also {second.text}, and {third.text}. "
            f"For {company_name}, that means sharper discovery, clearer solution design, and recommendations that stay credible once implementation begins."
        )
    if lane_key == "customer_success":
        resolution = (
            f"I also {second.text}, and {third.text}. For {company_name}, that means steadier adoption, clearer renewal "
            "risk visibility, stronger executive alignment, and more grounded expansion conversations across complex accounts."
        )
    else:
        specialty = build_resume.role_specialty_phrase(job_description, "the role's core work")
        resolution = (
            f"I also {second.text}, and {third.text}. For {company_name}, that means {specialty} "
            f"that leads to {concise_outcome_phrase(profile.outcomes, build_resume.natural_problem_phrase(profile))}."
        )

    lens = build_resume.primary_story_lens(job_description)
    if lens and str(lens.get("key")) == "nonprofit_education":
        relevance = (
            f"For {company_name}, the relevance is the ability to make transformation practical: convert broad priorities into project plans, "
            "measurement habits, stakeholder alignment, and repeatable systems that scale with the mission."
        )
        return f"{situation} I also {second.text}, and {third.text}. {relevance}"

    return repair_cover_paragraph(f"{situation} {complication} {resolution}", "cover proof paragraph")


def learning_paragraph(company_name: str, job_description: str, resume_text: str) -> str:
    has_mckinsey = "McKinsey Forward" in resume_text
    specialty = build_resume.role_specialty_phrase(job_description, "the work")
    role_title = build_resume.extract_job_title(job_description) or ""
    if _uses_firm_profile(company_name, "bain", role_title, job_description):
        return (
            "Building reporting and analytics tools across a broad client portfolio reinforced a consulting lesson I trust: enter the ambiguity, "
            "structure the problem quickly, earn confidence early, and leave the client with something measurable."
        )
    if is_consulting_context(company_name, job_description) and has_mckinsey:
        return (
            "I am also a deliberate learner. Completing the McKinsey Forward program while leading global systems operations reflects a consistent pattern in how I develop: "
            "I build structured frameworks around the work I am doing so I can do it better, explain it more clearly, and transfer it to others. "
            f"That habit is one reason {specialty} work appeals to me."
        )
    lens = build_resume.primary_story_lens(job_description)
    if lens:
        return str(lens["cover_lesson"])
    if has_mckinsey:
        return (
            "I am also a deliberate learner. Completing the McKinsey Forward program while leading global systems operations reflects how I tend to develop: "
            "I build structure around complex work so I can explain it clearly, improve it, and help other people adopt it."
        )
    return (
        "I am deliberate about how I learn and communicate. When the work is unfamiliar or ambiguous, I build a practical framework around it, test that framework with stakeholders, "
        "and use it to make decisions clearer for the people depending on the outcome."
    )



def anticipated_fit_bridge(company_name: str, role_title: str, job_description: str, resume_text: str) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_key = effective_lane_key(role_title, job_description, profile)
    combined = f"{role_title}\n{job_description}"
    lowered = combined.lower()

    lens = build_resume.primary_story_lens(job_description)
    # Context flags computed once and used throughout branch selection to avoid redundant pattern matching.
    has_commercial_signals = bool(re.search(r"\b(expansion|upsell|cross-sell|commercial|revenue|renewal|retention|quota)\b", lowered))
    has_direct_report_signals = bool(re.search(r"\b(direct reports|people manager|manage a team|lead a team|team leadership)\b", lowered))
    has_new_toolset_signals = bool(re.search(r"\b(learn|new system|platform|tool|technical depth|hands-on)\b", lowered))
    is_nonprofit_education = bool(lens and str(lens.get("key")) == "nonprofit_education")
    is_consulting = is_consulting_context(company_name, job_description)

    if is_nonprofit_education:
        return (
            "My background is strongest in practical mission execution across teams that need a process adopted, the data trusted, and momentum sustained after launch. "
            "That mix is useful when the goal is measurable progress toward stronger opportunities for Fellows."
        )

    if is_consulting and lane_key == "corporate_strategy":
        return (
            "I have spent much of my career in the messy middle where strong ideas either gain traction or stall through requirements, stakeholder tradeoffs, data quality, adoption, and measurement. "
            "Client work benefits from that operating perspective."
        )

    if lane_key == "customer_success":
        if has_commercial_signals:
            return (
                "The commercial side of my background comes from using discovery, executive reviews, adoption signals, and risk recovery to protect revenue, rebuild trust, and identify where a customer can get more value."
            )
        return (
            "The customer-facing side of my background matters as much as the technical side because I have managed complex client portfolios, "
            "stabilized high-risk accounts, and translated product or workflow issues into clearer adoption plans and stronger customer trust."
        )

    if lane_key == "presales_solution":
        return (
            "My background brings implementation-side discovery, executive alignment, and solution judgment, which helps keep technical recommendations credible once delivery begins."
        )

    if lane_key == "implementation_delivery":
        return (
            "The strongest pattern in my background is client-facing implementation leadership. "
            "I am strongest in work that depends on clean scoping and requirements translation, disciplined testing, and steady follow-through so integrations and data migration stay credible from discovery through delivery."
        )

    if lane_key == "corporate_strategy":
        return (
            "My strategy value is operator-backed. I am strongest where analysis has to turn into execution, "
            "and where leaders need the facts, tradeoffs, risks, and adoption path made clear before a recommendation can matter."
        )

    if lane_key == "process_improvement":
        return (
            "I start process improvement work with a current-state map and use data to separate root cause from symptoms. "
            "Then I validate the fix with users and measure whether quality, speed, or service actually changed."
        )

    if has_direct_report_signals:
        return (
            "Where the role calls for leadership, my experience has been program-director style leadership through aligning owners without relying on title authority, "
            "setting cadence, clarifying decisions, and keeping cross-functional teams moving through complex delivery work."
        )

    if profile.unsupported_requirements or has_new_toolset_signals:
        return (
            "My value in a new toolset or environment is ramp speed with structure. "
            "I trace the workflow, validate the data, and document the operating pattern so early ambiguity becomes a repeatable path others can use."
        )

    if lane_key == "change_enablement":
        return (
            "The part of change work I understand best is the human operating gap after a decision is made. People need clarity, proof, training, feedback loops, "
            "and visible ownership before a new way of working becomes durable."
        )

    if lane_key == "analytics_operations":
        return (
            "My analytics value is not just building reports; it is making the decision behind the report clearer, confirming the source data can be trusted, "
            "and giving leaders a practical way to act on what they see."
        )

    return ""


def gap_address_paragraph(company_name: str, role_title: str, job_description: str, resume_text: str) -> str:
    gaps = build_resume.poor_fit_requirements(job_description, resume_text)
    if not gaps:
        return ""
    primary_gap = gaps[0].strip().rstrip(".")
    sanitized_gap = re.sub(r"^(?:lack of|missing|no)\s+", "", primary_gap, flags=re.I).strip()
    profile = build_resume.job_problem_profile(job_description, resume_text)
    specialty = build_resume.role_specialty_phrase(job_description, build_resume.natural_problem_phrase(profile))
    paragraph = " ".join(
        [
            f"The strongest fit signal for {company_name} is that my background already covers the operating pattern behind {specialty} through structured learning, stakeholder alignment, validation, and measurable follow-through.",
            f"If the {role_title} scope extends into {sanitized_gap}, I would handle that adjacent work the same way I have handled it before by getting close to the workflow, making the decision path explicit, and showing progress early.",
            f"That is the bridge I would bring to {company_name} through practical ramp speed, credible communication, and outcomes teams can use.",
        ]
    )
    defensive_openers = (
        "while i have not",
        "although i do not",
        "i know i lack",
        "i may not have",
    )
    for opener in defensive_openers:
        if paragraph.lower().startswith(opener):
            paragraph = paragraph[len(opener):].lstrip(" ,.-")
            break
    return paragraph


def readiness_bridge_sentence(
    readiness: build_resume.ResumeReadiness,
    company_name: str,
    role_title: str,
    job_description: str,
) -> str:
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    labels = {build_resume.normalize_compare(gap.label) for gap in readiness.hard_blockers}
    if any(
        label in labels
        for label in ("budget", "eac", "etc", "financial forecasting", "financial ownership", "earned value", "p l")
    ):
        return (
            "If the role extends into budget ownership and EAC or ETC tracking, I would be transparent about the adjacent nature of my experience and bring the same disciplined delivery control I have used in complex scope and stakeholder work."
        )
    if any("direct people leadership" in label or "people manager enablement" in label for label in labels):
        return (
            "If the role extends into direct people leadership, I would be transparent about the adjacent nature of my experience and bring the same owner-alignment, coaching, and operating-rhythm discipline I have used across cross-functional delivery work."
        )
    specialty = {
        "analytics_operations": "analytics and decision-support work",
        "change_enablement": "change-adoption work",
        "corporate_strategy": "strategy and transformation work",
        "customer_success": "customer-success ownership",
        "implementation_delivery": "implementation delivery",
        "presales_solution": "solution consulting",
        "process_improvement": "process-improvement work",
    }.get(
        lane_key,
        build_resume.role_specialty_phrase(job_description, role_title or "the role"),
    )
    return (
        f"If the role calls for deeper experience in {specialty} than I have owned directly, I would address that gap honestly, ramp quickly, and bring the same structured follow-through I have used in adjacent delivery work."
    )


def first_90_days_cover_sentence(role_title: str, job_description: str, resume_text: str = "") -> str:
    lowered = role_title.lower()
    senior_markers = ("senior", "lead", "principal", "director", "head", "manager")
    if not any(marker in lowered for marker in senior_markers):
        return ""
    profile = build_resume.job_problem_profile(job_description, resume_text)
    lane_sentences = {
        "presales_solution": "In the first 90 days, I would focus on learning the buyer friction points, tightening discovery quality, and making sure technical proof stays connected to commercial confidence.",
        "customer_success": "In the first 90 days, I would focus on account health signals, adoption risk, and an executive-review rhythm tied to value rather than status.",
        "change_enablement": "In the first 90 days, I would focus on where change is stalling, who needs clarity, and which adoption behaviors would create visible momentum fastest.",
        "analytics_operations": "In the first 90 days, I would focus on the decisions leaders most need to make, the data trust issues behind them, and the reporting changes that would improve action fastest.",
        "implementation_delivery": "In the first 90 days, I would focus on the implementation risks that surface early, the stakeholder decisions that control pace, and the adoption checkpoints that keep launch momentum intact.",
    }
    return lane_sentences.get(
        profile.primary_lane,
        "In the first 90 days, I would focus on where execution is most vulnerable, which decisions need tighter ownership, and what early win would build momentum fastest.",
    )


def closing_paragraph(company_name: str, role_title: str, job_description: str = "") -> str:
    """Write a direct closing that names the problem worth discussing next."""
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    specialty = build_resume.role_specialty_phrase(job_description, "the role's core work")
    lens = build_resume.primary_story_lens(job_description)

    if lens and str(lens.get("key")) == "nonprofit_education":
        return (
            f"I would welcome a conversation about where {company_name} most needs stronger reporting, stakeholder alignment, or implementation follow-through "
            f"in the {role_title} role."
        )
    if is_early_stage_context(job_description):
        return (
            f"I would welcome a conversation about where {company_name} most needs faster customer delivery, clearer documentation, or a repeatable operating playbook as the team scales."
        )

    lane_closes = {
        "presales_solution": (
            f"I would welcome a conversation about the buyer questions, demo risks, or late-stage technical objections creating the most drag in {company_name}'s current pipeline for the {role_title} role."
        ),
        "customer_success": (
            f"I would welcome a conversation about which customers at {company_name} most need steadier adoption, clearer executive alignment, or stronger recovery plans before renewal pressure builds."
        ),
        "change_enablement": (
            f"I would welcome a conversation about where {company_name} most needs adoption momentum, role clarity, or cleaner follow-through after launch in the {role_title} role."
        ),
        "analytics_operations": (
            f"I would welcome a conversation about which reporting gaps or decision bottlenecks would create the fastest operational lift for {company_name} in the {role_title} role."
        ),
        "implementation_delivery": (
            f"I would welcome a conversation about which delivery risks, data dependencies, or stakeholder decisions the {role_title} role needs to stabilize first at {company_name}."
        ),
        "process_improvement": (
            f"I would welcome a conversation about which process problems at {company_name} have lingered longest and would benefit most from root-cause work, measurable fixes, and user adoption discipline."
        ),
    }
    return lane_closes.get(
        lane_key,
        f"I would welcome a conversation about which {specialty} priorities at {company_name} would benefit most from clearer scope, faster decisions, and measurable follow-through in the {role_title} role.",
    )


def smooth_cover_letter_text(
    text: str,
    warnings: list[str] | None = None,
    *,
    allowed_acronyms: set[str] | None = None,
) -> str:
    """Remove stiff, colon-led transitions so the letter reads like prose."""
    replacements = (
        ("My fit is practical:", "My fit is practical because"),
        ("My approach is discovery-first:", "My approach starts with discovery."),
        ("My approach is account-structure-first:", "My approach starts with account structure."),
        ("My approach is decision-first:", "My approach starts with the decision."),
        ("My approach is operator-backed strategy:", "My approach is operator-backed strategy."),
        ("My lane is technical operations:", "My lane is technical operations with"),
        ("The experience behind that is concrete:", "The experience behind that is concrete because"),
        ("That result came from the same discipline this role appears to require:", "That result came from the same discipline this role appears to require, including"),
        ("The complication is that:", "The complication is that"),
        ("The strongest bridge in my background is value-led growth:", "The strongest bridge in my background is value-led growth through"),
        ("The customer-facing side of my background matters as much as the technical side:", "The customer-facing side of my background matters as much as the technical side because"),
        ("The bridge from implementation into solution consulting is practical:", "The bridge from implementation into solution consulting is practical because"),
        ("The bridge from systems and reporting work into process improvement is practical:", "The bridge from systems and reporting work into process improvement is practical because"),
        ("The strategy angle I bring is operator-backed:", "The strategy angle I bring is operator-backed because"),
        ("Where the role calls for leadership, my experience has been program-director style leadership:", "Where the role calls for leadership, my experience has been program-director style leadership through"),
        ("Where the toolset or environment is new, my value is the way I learn it: trace", "Where the toolset or environment is new, my value is the way I learn it. I trace"),
        ("The part of change work I understand best is the human operating gap after a decision is made:", "The part of change work I understand best is the human operating gap after a decision is made."),
        ("My approach to that kind of work is hands-on: enter", "My approach to that kind of work is hands-on. I enter"),
        ("My analytics value is not just building reports;", "My analytics value goes beyond building reports because it"),
        ("process-improvement rhythm: define", "process-improvement rhythm that helps teams define"),
        ("What I bring to", "What I bring to"),
    )
    cleaned = normalize_cover_dash_punctuation(text)
    for old, new in replacements:
        cleaned = cleaned.replace(old, new)
    cleaned = re.sub(
        r"\bis doing something genuinely difficult:\s+",
        "is doing something genuinely difficult by ",
        cleaned,
    )
    cleaned = re.sub(
        r",\s*which is why my background fits here:\s+",
        ", and my background fits because it has centered on ",
        cleaned,
    )
    cleaned = re.sub(
        r"\bvalue-led growth:\s+using\b",
        "value-led growth through",
        cleaned,
        flags=re.I,
    )
    cleaned = re.sub(
        r"\b([A-Z][A-Za-z]{2,40}):\s+",
        lambda match: f"{match.group(1)}. ",
        cleaned,
    )
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That is\b", r"\1This is", cleaned)
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That was\b", r"\1This was", cleaned)
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That result\b", r"\1The result", cleaned)
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That work\b", r"\1The work", cleaned)
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That habit\b", r"\1This habit", cleaned)
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That helps\b", r"\1This helps", cleaned)
    cleaned = re.sub(r"(^|(?<=[.!?])\s+)That matters\b", r"\1This matters", cleaned)
    cleaned = re.sub(
        r"([.!?]\s+)([a-z])",
        lambda match: match.group(1) + match.group(2).upper(),
        cleaned,
    )
    bullet_matches = re.findall(r"(?:^|\n)([•\*\-–]\s+\w[^•\*\-–]*)", cleaned)
    for fragment in bullet_matches:
        if warnings is not None:
            warnings.append(f"possible JD artifact found: {fragment.strip()}. Review before submitting.")
    cleaned = re.sub(r"(?:^|\n)[•\*\-–]\s+\w[^•\*\-–]*", " ", cleaned)

    mixed_case_hits = re.findall(r"\b[a-z][A-Z]+\w*\b", cleaned)
    for fragment in mixed_case_hits:
        if warnings is not None:
            warnings.append(f"possible JD artifact found: {fragment}. Review before submitting.")
    cleaned = re.sub(
        r"\b[a-z][A-Z]+\w*\b",
        lambda match: _cleanup_short_fragment(match.group(0)),
        cleaned,
    )

    allowed_tokens = {token.upper() for token in APPROVED_COVER_ACRONYMS}
    if allowed_acronyms:
        allowed_tokens.update(token.upper() for token in allowed_acronyms)
    all_caps_hits = [
        fragment
        for fragment in re.findall(r"\b[A-Z]{3,}\b", cleaned)
        if fragment.upper() not in allowed_tokens
    ]
    for fragment in all_caps_hits:
        if warnings is not None:
            warnings.append(f"possible JD artifact found: {fragment}. Review before submitting.")
    cleaned = re.sub(
        r"\b[A-Z]{3,}\b",
        lambda match: match.group(0) if match.group(0).upper() in allowed_tokens else " ",
        cleaned,
    )
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    remaining_colons = cleaned.count(":")
    if remaining_colons > 2:
        debug_print(
            f"cover letter smooth pass: {remaining_colons} colons remain after smoothing",
            file=sys.stderr,
            flag="DEBUG_COVER_LETTER",
        )
    return cleaned


def effective_lane_key(role_title: str, job_description: str, profile: build_resume.JobProblemProfile) -> str:
    return build_resume.effective_lane_key(role_title, job_description, profile)


def challenge_forecast_sentence(company_name: str, role_title: str, job_description: str) -> str:
    profile = build_resume.job_problem_profile(job_description)
    lane_key = effective_lane_key(role_title, job_description, profile)
    specialty = build_resume.role_specialty_phrase(job_description, "enterprise software delivery")
    forecasts = {
        "change_enablement": (
            f"The business challenge appears to be helping {company_name}'s teams adopt new ways of working "
            f"around {specialty} without losing trust, momentum, or operational clarity."
        ),
        "presales_solution": (
            f"The business challenge appears to be helping {company_name}'s buyers understand value, risk, "
            f"and implementation fit in {specialty} before major decisions are made."
        ),
        "customer_success": (
            f"The business challenge appears to be helping {company_name}'s customers see measurable value "
            f"from {specialty} while reducing adoption, renewal, and escalation risk."
        ),
        "implementation_delivery": (
            f"The business challenge appears to be moving complex {specialty} work from requirements to adoption "
            "without letting scope, timing, or stakeholder alignment drift."
        ),
        "analytics_operations": (
            f"The business challenge appears to be turning operational data into clearer decisions, "
            "better priorities, and measurable process improvement."
        ),
    }
    return forecasts.get(
        lane_key,
        f"The business challenge appears to be turning complex work into clearer decisions and measurable progress for {company_name}.",
    )


def company_context_sentence(job_description: str) -> str:
    context = build_resume.primary_employer_context(job_description)
    specialty = build_resume.role_specialty_phrase(job_description, "")
    values = build_resume.visible_values_phrase(job_description)
    if context and str(context["key"]) == "consulting":
        if specialty and values:
            return (
                f"The firm context also points to client-service judgment in {specialty}, with visible emphasis on {values}."
            )
        if specialty:
            return f"The firm context also points to client-service judgment and practical recommendations in {specialty}."
        return "The firm context also points to client-service judgment, structured discovery, and practical recommendations."
    if specialty and values:
        return f"The posting also emphasizes {specialty} in an environment that values {values}."
    if specialty:
        return f"The posting also points to a need for practical depth in {specialty}."
    if values:
        return f"The posting also emphasizes {values}, which matters when work crosses teams and customer expectations."
    return ""


def opening_theme(lane_key: str) -> str:
    themes = {
        "change_enablement": "change adoption, training, executive workshops, and implementation support",
        "customer_success": "customer adoption, risk reduction, account recovery, and post-go-live outcomes",
        "analytics_operations": "KPI reporting, dashboard development, process improvement, and stakeholder decision support",
        "presales_solution": "discovery, solution design, implementation realism, and customer-facing technical guidance",
        "implementation_delivery": "requirements, configuration, data migration, testing, go-live readiness, and adoption",
    }
    return themes.get(
        lane_key,
        "solution consulting, implementation leadership, customer outcomes, and cross-functional delivery",
    )


def legacy_closing_paragraph(company_name: str, job_description: str = "") -> str:
    profile = build_resume.job_problem_profile(job_description)
    specialty = build_resume.role_specialty_phrase(job_description, "")
    values = build_resume.visible_values_phrase(job_description)
    context_goal = f"{specialty} priorities" if specialty else "customer and delivery goals"
    if values:
        context_goal = f"{context_goal} while supporting {values}"
    if profile.primary_lane == "change_enablement":
        return (
            f"I would welcome the opportunity to discuss how my background in software delivery, training, "
            f"executive workshops, and adoption support can help {company_name} turn change into practical results."
        )
    if profile.primary_lane == "customer_success":
        return (
            f"I would welcome the opportunity to discuss how my background in customer adoption, executive "
            f"stakeholder engagement, account recovery, and measurable implementation outcomes can support "
            f"{company_name}'s {context_goal}."
        )
    if profile.primary_lane == "analytics_operations":
        return (
            f"I would welcome the opportunity to discuss how my background in analytics, reporting, workflow "
            f"improvement, and stakeholder decision support can support {company_name}'s {context_goal}."
        )
    return (
        f"I would welcome the opportunity to discuss how my background in software implementation, "
        f"solution consulting, executive stakeholder engagement, and measurable customer outcomes "
        f"can support {company_name}'s {context_goal}."
    )


def set_default_style(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = RESUME_FONT
    normal.font.size = Pt(BODY_FONT_SIZE)
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float | None = None, color=None):
    run = paragraph.add_run(text)
    run.font.name = RESUME_FONT
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), RESUME_FONT)
    font.set(qn("w:hAnsi"), RESUME_FONT)
    run_properties.append(font)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(BODY_FONT_SIZE * 2))
    run_properties.append(size)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run_element.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    after: int = 0,
    before: int = 0,
    align=None,
    line_spacing: float = BODY_LINE_SPACING,
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align
    if text:
        add_run(paragraph, text)
    return paragraph


def add_blank_line(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    add_run(paragraph, "", size=BODY_FONT_SIZE)


def add_header(document: Document, resume_texts: list[str]) -> None:
    title = add_paragraph(document, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=HEADER_LINE_SPACING)
    add_run(title, "Christian Estrada", bold=True, size=NAME_FONT_SIZE, color=NAME_BLUE)

    contact_line = next(
        (text for text in resume_texts[:5] if CONTACT_EMAIL in text),
        f"Atlanta, GA  |  770-710-4216  |  {CONTACT_EMAIL}  |  {VISIBLE_LINKEDIN_URL}",
    )
    contact = add_paragraph(
        document,
        after=HEADER_TO_DATE_GAP_PT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=HEADER_LINE_SPACING,
    )
    add_run(contact, contact_line, size=BODY_FONT_SIZE)


_COVER_SENTENCE_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "because", "by", "for", "from", "has", "have", "i",
    "in", "into", "is", "it", "its", "my", "of", "on", "or", "our", "that", "the", "their", "this",
    "to", "was", "we", "with",
}


def cover_sentence_tokens(text: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[a-z0-9$%+]+", text.lower())
        if token not in _COVER_SENTENCE_STOPWORDS and len(token) > 2
    }


def cover_sentences_near_duplicate(left: str, right: str, *, same_paragraph: bool = False) -> bool:
    left_tokens = cover_sentence_tokens(left)
    right_tokens = cover_sentence_tokens(right)
    if not left_tokens or not right_tokens:
        return False
    overlap = len(left_tokens & right_tokens)
    similarity = overlap / max(len(left_tokens | right_tokens), 1)
    return similarity >= 0.72 or (
        not same_paragraph
        and overlap >= 4
        and bool(re.search(r"\d|%|\$", left))
        and bool(re.search(r"\d|%|\$", right))
    )


def cover_sentence_is_thesis(sentence: str) -> bool:
    return bool(
        re.search(
            r"^(?:The (?:best|most useful|hardest)|The work I find most meaningful|The standard I hold|My approach|My lane|I am strongest where|Reporting is only useful if)\b",
            sentence,
            re.I,
        )
    )


def cover_sentence_has_context(sentence: str, company_name: str, role_title: str, job_description: str) -> bool:
    if company_name and re.search(re.escape(company_name), sentence, re.I):
        return True
    if normalized_phrase_in_text(sentence, role_title):
        return True
    specialty_terms = tuple(term.lower() for term in build_resume.visible_role_specialties(job_description)[:6])
    sentence_lower = sentence.lower()
    if any(term and term in sentence_lower for term in specialty_terms):
        return True
    profile = build_resume.job_problem_profile(job_description)
    lane_terms = {
        "analytics_operations": ("analytics", "retention", "lifecycle", "reporting", "decision", "workflow", "data", "process", "measurement", "assessment", "learning", "validation", "quality"),
        "customer_success": ("renewal", "adoption", "retention", "account", "customer", "value", "risk"),
        "implementation_delivery": ("implementation", "go-live", "delivery", "migration", "adoption", "workflow", "stakeholder"),
        "presales_solution": ("buyer", "discovery", "solution", "demo", "requirements"),
        "change_enablement": ("change", "adoption", "training", "workflow", "launch", "behavior"),
        "process_improvement": ("process", "root cause", "workflow", "service", "quality", "measure"),
        "corporate_strategy": ("strategy", "decision", "tradeoff", "recommendation", "plan"),
    }
    return any(term in sentence_lower for term in lane_terms.get(profile.primary_lane, ()))


def cover_sentence_is_generic(sentence: str, company_name: str, role_title: str, job_description: str) -> bool:
    generic_patterns = (
        r"\bI am also a deliberate learner\b",
        r"\bThis is familiar territory\b",
        r"\bThat mix helps\b",
        r"\bMy background fits\b",
        r"\bI am strongest where\b",
        r"\bI build structure around complex work\b",
        r"\bI translate technical work into operating decisions\b",
    )
    if any(re.search(pattern, sentence, re.I) for pattern in generic_patterns):
        return True
    return not (
        paragraph_has_fast_proof(sentence)
        or cover_sentence_has_context(sentence, company_name, role_title, job_description)
    )


def cover_sentence_score(sentence: str, company_name: str, role_title: str, job_description: str) -> int:
    score = 0
    if paragraph_has_fast_proof(sentence):
        score += 6
    if cover_sentence_has_context(sentence, company_name, role_title, job_description):
        score += 3
    # paragraph_has_fast_proof() is binary, so it can't tell a sentence with one number
    # from one with five. Add the actual signal density (scale/workflow/business-impact
    # markers) so the trimmer in fit_standard_cover_word_budget() doesn't tie-break toward
    # cutting the sentence carrying the most concrete proof just because it's also the
    # longest. This keeps the trimmer's notion of "valuable sentence" aligned with
    # lead_burial_check()'s objective_context_signal_count() gate downstream.
    score += build_resume.objective_context_signal_count(sentence)
    if "because" in sentence.lower():
        score += 1
    if cover_sentence_is_thesis(sentence):
        score -= 4
    if cover_sentence_is_generic(sentence, company_name, role_title, job_description):
        score -= 5
    return score


def select_cover_sentences(
    text: str,
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    limit: int,
    allow_thesis: bool = False,
) -> list[str]:
    sentences = [smooth_cover_letter_text(sentence) for sentence in sentence_list(text)]
    ranked: list[tuple[int, int, str]] = []
    for index, sentence in enumerate(sentences):
        if not sentence:
            continue
        if not allow_thesis and cover_sentence_is_thesis(sentence):
            continue
        if cover_sentence_is_generic(sentence, company_name, role_title, job_description):
            continue
        ranked.append((cover_sentence_score(sentence, company_name, role_title, job_description), index, sentence))
    ranked.sort(key=lambda item: (item[0], -item[1]), reverse=True)

    selected: list[tuple[int, str]] = []
    for _, index, sentence in ranked:
        if any(cover_sentences_near_duplicate(sentence, existing) for _, existing in selected):
            continue
        selected.append((index, sentence))
        if len(selected) >= limit:
            break
    selected.sort(key=lambda item: item[0])
    return [sentence for _, sentence in selected]


def preferred_cover_word_range(mode: str) -> tuple[int, int]:
    normalized = normalize_cover_mode(mode)
    if normalized == LONG_COVER_MODE:
        return 180, 240
    return 110, 145


def cover_letter_document_word_count(salutation: str, body_paragraphs: tuple[str, ...] | list[str]) -> int:
    return word_count(
        "\n".join(
            [
                salutation,
                *body_paragraphs,
                "Thank you for your time and consideration,",
                "Christian Estrada",
            ]
        )
    )


def ensure_sentence(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return ""
    if cleaned[-1] not in ".!?":
        cleaned += "."
    return cleaned


def mapped_terms_in_text(terms: tuple[str, ...], text: str) -> tuple[str, ...]:
    found: list[str] = []
    lowered = text.lower()
    for term in terms:
        if term.lower() in lowered and term not in found:
            found.append(term)
    return tuple(found)


def normalized_company_mission(company_name: str, mission_text: str) -> str:
    company_name = safe_company_name(company_name)
    mission = ensure_sentence(normalize_cover_dash_punctuation(sanitize_for_spoken_text(mission_text) or mission_text))
    if word_count(mission) < 3 or re.fullmatch(r"(?:at\s+[a-z0-9&.' -]+,\s*)?[. ]*", mission, re.I):
        mission = ""
    if company_name:
        mission = re.sub(
            rf"^\s*At\s+{re.escape(company_name)},\s+{re.escape(company_name)}\b",
            company_name,
            mission,
            flags=re.I,
        )
    mission = re.sub(
        rf"^\s*At\s+{re.escape(company_name)},\s+its company mission is to\b",
        f"{company_name} aims to",
        mission,
        flags=re.I,
    )
    mission = re.sub(
        r"^\s*Its company mission is to\b",
        f"{company_name} aims to",
        mission,
        flags=re.I,
    )
    mission = re.sub(
        r"^\s*Our company mission is to\b",
        f"{company_name} aims to",
        mission,
        flags=re.I,
    )
    if word_count(mission) > 20 or re.search(r"\b(?:to help|while|without|so that)\.?$", mission, re.I):
        for splitter in (r"\bto help\b", r"\bwhile\b", r"\bwithout\b", r"\bso that\b"):
            parts = re.split(splitter, mission, maxsplit=1, flags=re.I)
            if len(parts) != 2:
                continue
            trimmed = ensure_sentence(parts[0].rstrip(" ,;:-"))
            if word_count(trimmed) >= 6:
                mission = trimmed
                break
    mission = mission.replace("human-grade pet food, tailored for each dog’s nutritional needs, and delivers directly to its customers’ doors", "human-grade meals and delivers them to customers' doors")
    mission = mission.replace("human-grade pet food, tailored for each dog's nutritional needs, and delivers directly to its customers' doors", "human-grade meals and delivers them to customers' doors")
    mission = mission.replace("human-grade pet food, tailored to each dog, and delivers to customers' doors", "human-grade meals and delivers them to customers' doors")
    mission = mission.replace("tailored for each dog’s nutritional needs", "tailored to each dog")
    mission = mission.replace("tailored for each dog's nutritional needs", "tailored to each dog")
    mission = mission.replace("directly to its customers’ doors", "to customers' doors")
    mission = mission.replace("directly to its customers' doors", "to customers' doors")
    if mission and company_name and not re.search(re.escape(company_name), mission, re.I):
        mission = f"At {company_name}, {mission[:1].lower() + mission[1:] if mission[:1].isupper() else mission}"
    if re.match(
        rf"^(?:At\s+{re.escape(company_name)},\s+)?(?:needs|focuses|centers|protects|turns|depends)\b",
        mission,
        re.I,
    ):
        mission = ""
    return mission


def standard_mission_paragraph(company_name: str, role_title: str, signals: CoverLetterSignals) -> str:
    company_name = safe_company_name(company_name)
    role_title = safe_role_title(role_title)
    mission = normalized_company_mission(company_name, signals.company_mission)
    role_clause = safe_core_problem(re.sub(r"^building\s+", "", signals.role_core_function, flags=re.I)) or "the team's highest-priority work"
    if not mission:
        if company_name:
            role_sentence = ensure_sentence(f"At {company_name}, the {role_title} role centers on {role_clause}")
        else:
            role_sentence = ensure_sentence(f"The {role_title} role centers on {role_clause}")
        support_sentence = "The role needs clear decisions, aligned owners, and follow-through that holds once the work is live."
        return f"{role_sentence} {support_sentence}".strip()
    elif re.match(r"^building\s+", signals.role_core_function, re.I):
        role_sentence = ensure_sentence(
            f"The {role_title} role protects {role_clause}"
        )
    elif re.match(r"^(?:the|a|an|this|that|its|[a-z]+['’]s)\b", role_clause, re.I):
        role_sentence = ensure_sentence(
            f"The {role_title} role centers on {role_clause}"
        )
    elif gerund_match := re.match(r"^(turning|moving|handling|protecting|keeping|aligning|supporting)\s+(.+)$", role_clause, re.I):
        verb = gerund_match.group(1).lower()
        remainder = gerund_match.group(2).strip()
        finite_verb = {
            "aligning": "aligns",
            "handling": "handles",
            "keeping": "keeps",
            "moving": "moves",
            "protecting": "protects",
            "supporting": "supports",
            "turning": "turns",
        }.get(verb, "supports")
        role_sentence = ensure_sentence(
            f"The {role_title} role {finite_verb} {remainder}"
        )
    else:
        role_sentence = ensure_sentence(
            f"The {role_title} role depends on {role_clause}"
        )
    return " ".join(part for part in (mission, role_sentence) if part).strip()


def proof_reframing_paragraph(company_name: str, signals: CoverLetterSignals) -> tuple[str, tuple[str, ...]]:
    company_name = safe_company_name(company_name)
    mapped_terms = signals.jd_skill_terms[:4]
    if mapped_terms and mapped_terms[0] == "customer experience":
        mapped_terms = ("customer experience", "retention", "decision support")
    if mapped_terms == ("testing",) and "technical scoping" in signals.role_core_function.lower():
        mapped_terms = ("technical scoping", "integrations", "data migration")
    if word_count(comma_join(mapped_terms)) > 6:
        mapped_terms = signals.jd_skill_terms[:3]
    term_text = comma_join(mapped_terms)
    proof_text = signals.top_accomplishment.strip().rstrip(".")
    proof_text = re.sub(r"^Enabled real-time executive decision-making by building\s+", "Built ", proof_text, flags=re.I)
    proof_text = re.sub(r"^Built more than 200\b", "Built 200+", proof_text, flags=re.I)
    proof_text = re.sub(r"\bthat replaced raw exports and verbal status updates\b", "that replaced raw exports", proof_text, flags=re.I)
    proof_text = re.sub(
        r"including more than one million dollars in at-risk annual revenue, by diagnosing root causes, consolidating ownership, and driving resolution",
        "including $1M+ in at-risk annual revenue, by consolidating ownership and driving resolution",
        proof_text,
        flags=re.I,
    )
    if term_text:
        paragraph = " ".join(
            (
                ensure_sentence(proof_text),
                ensure_sentence(f"For {company_name}, it shows practical experience with {term_text}"),
            )
        ).strip()
    else:
        paragraph = " ".join(
            (
                ensure_sentence(proof_text),
                ensure_sentence(f"For {company_name}, it shows practical experience turning analysis into usable decisions"),
            )
        ).strip()
    supplemental_metric = signals.communication_metric.strip().rstrip(".")
    if (
        supplemental_metric
        and paragraph_has_fast_proof(supplemental_metric)
        and (word_count(paragraph) < 40 or not paragraph_has_fast_proof(paragraph))
        and not communication_metric_matches_avoid_text(supplemental_metric, proof_text)
    ):
        if term_text:
            paragraph = " ".join(
                (
                    ensure_sentence(proof_text),
                    ensure_sentence(
                        f"{supplemental_metric}. That experience also shows practical experience with {term_text} at the pace the role requires"
                    ),
                )
            ).strip()
        else:
            paragraph = " ".join(
                (
                    ensure_sentence(proof_text),
                    ensure_sentence(f"{supplemental_metric}. It also shows how I keep analysis tied to usable decisions"),
                )
            ).strip()
    return paragraph, mapped_terms_in_text(mapped_terms, paragraph)


def workflow_paragraph(signals: CoverLetterSignals) -> str:
    compact_process = re.sub(r", define the right question,\s*", ", ", signals.ambiguity_process, flags=re.I)
    compact_process = re.sub(r"\s+so the analysis can be reused$", "", compact_process, flags=re.I)
    first = ensure_sentence(f"When a business question arrives half-formed, I {compact_process}")
    env_text = comma_join(signals.jd_test_environments)
    if env_text and word_count(env_text) <= 6 and len(signals.jd_test_environments) <= 2:
        first = ensure_sentence(
            f"When a business question arrives half-formed, I {compact_process} so analysis stays useful across {env_text}"
        )
    return first


def communication_paragraph(company_name: str, signals: CoverLetterSignals, avoid_text: str = "") -> str:
    first = signals.communication_metric.strip().rstrip(".")
    if first and communication_metric_matches_avoid_text(first, avoid_text):
        first = ""
    partner_text = partner_team_phrase(signals.partner_functions)
    if partner_text:
        if "support insights" in first.lower() or "sms support" in first.lower():
            parts = []
            if first:
                parts.append(ensure_sentence(first))
            parts.append("This keeps product, CX, and support aligned through clear stakeholder communication.")
            return " ".join(parts).strip()
        if not first:
            return ensure_sentence(f"This keeps {partner_text} aligned around clearer decisions")
        return " ".join(
            (
                ensure_sentence(first),
                ensure_sentence(f"This keeps {partner_text} aligned around clearer decisions"),
            )
        ).strip()
    if first:
        return " ".join(
            (
                ensure_sentence(first),
                ensure_sentence("This keeps the work focused on decisions, owners, and next steps instead of status-heavy reporting"),
            )
        ).strip()
    return ensure_sentence("I keep analysis tied to decisions, owners, and next steps instead of status-heavy reporting")


def soft_close_paragraph(company_name: str, signals: CoverLetterSignals) -> tuple[str, tuple[str, ...]]:
    pain_area = safe_connector_fragment(signals.jd_pain_area) or "highest-priority work"
    close = ensure_sentence(
        f"I would welcome the chance to discuss how I could support {company_possessive(company_name)} {pain_area} on the team"
    )
    close_terms = list(signals.jd_test_environments[:2])
    for term in signals.jd_skill_terms:
        if term.lower() in close.lower() and term not in close_terms:
            close_terms.append(term)
    return close, tuple(close_terms[:4])


def tighten_standard_close_sentence(text: str) -> str:
    updated = text.strip().rstrip(".")
    replacements = (
        ("needs clearer visibility and faster early wins across", "needs faster early wins across"),
        ("needs faster early wins across", "needs early wins across"),
        ("I would welcome a conversation about where", "I would welcome a conversation about how"),
        ("A conversation about where", "A conversation about how"),
    )
    for original, replacement in replacements:
        updated = updated.replace(original, replacement)
    return ensure_sentence(updated)


def standard_sentence_trim_priority(
    sentence: str,
    company_name: str,
    role_title: str,
    job_description: str,
) -> tuple[int, int, int]:
    lowered = sentence.lower().strip()
    if lowered.startswith("i translate "):
        base_priority = 0
    elif lowered.startswith("led ") and "workshop" in lowered:
        base_priority = 1
    elif lowered.startswith("if the role extends into"):
        base_priority = 4
    elif lowered.startswith("when a business question arrives"):
        base_priority = 3
    else:
        base_priority = 2
    return (
        base_priority,
        cover_sentence_score(sentence, company_name, role_title, job_description),
        -word_count(sentence),
    )


def fit_standard_cover_word_budget(
    salutation: str,
    body_paragraphs: tuple[str, ...] | list[str],
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    mode: str = STANDARD_COVER_MODE,
) -> tuple[str, ...]:
    normalized_mode = normalize_cover_mode(mode)
    max_words = cover_letter_word_range(normalized_mode)[1]
    allowed_acronyms = cover_allowed_acronyms(
        role_title=role_title,
        job_description=job_description,
        company_name=company_name,
    )
    paragraphs = [
        smooth_cover_letter_text(paragraph, allowed_acronyms=allowed_acronyms).strip()
        for paragraph in body_paragraphs
        if paragraph.strip()
    ]
    sentence_total = sum(len(sentence_list(paragraph) or [paragraph]) for paragraph in paragraphs)
    if (
        cover_letter_document_word_count(salutation, paragraphs) <= max_words
        and (
            normalized_mode != STANDARD_COVER_MODE
            or sentence_total <= MAX_STANDARD_BODY_SENTENCES
        )
    ):
        return tuple(paragraphs)

    sentence_groups = [sentence_list(paragraph) or [paragraph] for paragraph in paragraphs]

    def rebuilt_paragraphs() -> list[str]:
        return [
            " ".join(sentence.strip() for sentence in sentences if sentence.strip()).strip()
            for sentences in sentence_groups
            if any(sentence.strip() for sentence in sentences)
        ]

    def sentence_total_after_trim() -> int:
        return sum(len(sentences) for sentences in sentence_groups)

    if sentence_groups and sentence_groups[-1]:
        sentence_groups[-1][-1] = tighten_standard_close_sentence(sentence_groups[-1][-1])
        paragraphs = rebuilt_paragraphs()
        if (
            cover_letter_document_word_count(salutation, paragraphs) <= max_words
            and (
                normalized_mode != STANDARD_COVER_MODE
                or sentence_total_after_trim() <= MAX_STANDARD_BODY_SENTENCES
            )
        ):
            return tuple(paragraphs)

    # Trim the single lowest-value sentence anywhere in the letter on each pass,
    # rather than fully draining one paragraph before ever considering another.
    # The opening paragraph used to be completely off-limits to this function
    # (only paragraphs 1..len-2 plus the close paragraph were ever touched), on
    # the assumption that the opening would always stay short. Now that the
    # opening can legitimately carry a real company-fact sentence plus a role
    # sentence plus a motivation sentence (up to its own 80-word cap), it can be
    # the most expensive paragraph in the letter, and shielding it entirely
    # forced every cut onto the proof paragraph, sometimes stripping out every
    # concrete-proof sentence just to find a handful of spare words. Each
    # paragraph still protects the sentence that should never move: the first
    # sentence everywhere except the closing paragraph, where the last sentence
    # (the call to action) is protected instead.
    last_index = len(sentence_groups) - 1
    while sentence_groups and (
        cover_letter_document_word_count(salutation, rebuilt_paragraphs()) > max_words
        or (
            normalized_mode == STANDARD_COVER_MODE
            and sentence_total_after_trim() > MAX_STANDARD_BODY_SENTENCES
        )
    ):
        candidates: list[tuple[int, int]] = []
        for paragraph_index, sentences in enumerate(sentence_groups):
            if len(sentences) <= 1:
                continue
            if paragraph_index == last_index:
                candidates.extend((paragraph_index, index) for index in range(len(sentences) - 1))
            elif paragraph_index == 0:
                # The opening paragraph's first sentence (company/role context) and second
                # sentence (the role-specific sentence naming role_title) are both load-bearing:
                # opening_quality_problem() requires the role title to appear in the first
                # paragraph, and that title only ever shows up in the role sentence. Only a
                # third sentence (e.g. a motivation aside) is safe to trim here.
                candidates.extend((paragraph_index, index) for index in range(2, len(sentences)))
            else:
                candidates.extend((paragraph_index, index) for index in range(1, len(sentences)))
        if not candidates:
            break
        paragraph_index, sentence_index = min(
            candidates,
            key=lambda pair: standard_sentence_trim_priority(
                sentence_groups[pair[0]][pair[1]],
                company_name,
                role_title,
                job_description,
            ),
        )
        del sentence_groups[paragraph_index][sentence_index]

    return tuple(rebuilt_paragraphs())


def compose_standard_cover_letter(
    company_name: str,
    role_title: str,
    job_description: str,
    signals: CoverLetterSignals,
) -> CoverLetterDraft:
    salutation = preferred_salutation(company_name)
    mission = standard_mission_paragraph(company_name, role_title, signals)
    proof_paragraph_text, proof_terms = proof_reframing_paragraph(company_name, signals)
    fit_bridge = ensure_sentence(signals.fit_bridge) if signals.fit_bridge else ""
    workflow = workflow_paragraph(signals)
    communication = communication_paragraph(company_name, signals, avoid_text=proof_paragraph_text)
    close, close_terms = soft_close_paragraph(company_name, signals)

    four_paragraphs = (
        mission,
        proof_paragraph_text,
        " ".join(part for part in (fit_bridge, workflow) if part).strip(),
        f"{communication} {close}".strip(),
    )
    compact_third_paragraph = f"{workflow} {communication} {close}".strip()
    three_paragraphs = (
        mission,
        proof_paragraph_text,
        compact_third_paragraph,
    )

    preferred_min, preferred_max = preferred_cover_word_range(STANDARD_COVER_MODE)
    four_words = cover_letter_document_word_count(salutation, four_paragraphs)
    three_words = cover_letter_document_word_count(salutation, three_paragraphs)
    has_distinct_middle = bool(signals.jd_test_environments or signals.partner_functions or signals.communication_metric)
    if has_distinct_middle and preferred_min <= four_words <= preferred_max:
        chosen = four_paragraphs
        shape = 4
    else:
        chosen = three_paragraphs
        shape = 3
    chosen_words = cover_letter_document_word_count(salutation, chosen)
    if chosen_words < preferred_min and four_words <= STANDARD_MAX_LETTER_WORDS:
        chosen = four_paragraphs
        shape = 4
    if cover_letter_document_word_count(salutation, chosen) > STANDARD_MAX_LETTER_WORDS:
        chosen = three_paragraphs
        shape = 3
    chosen = fit_standard_cover_word_budget(
        salutation,
        chosen,
        company_name,
        role_title,
        job_description,
    )
    shape = len(chosen)

    return CoverLetterDraft(
        salutation=salutation,
        body_paragraphs=tuple(chosen),
        signals=signals,
        paragraph_shape=shape,
        proof_mapped_terms=proof_terms,
        close_mapped_terms=close_terms,
        mode=STANDARD_COVER_MODE,
    )


def compose_cover_letter_draft(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
    application_responses: tuple[question_prep.QualificationsResponse, ...] = (),
) -> CoverLetterDraft:
    plan = build_cover_letter_plan(
        company_name,
        role_title,
        job_description,
        resume_text,
        mode=mode,
        application_responses=application_responses,
    )
    return compose_cover_letter_from_plan(plan)


def standard_body_paragraphs(
    opening: str,
    proof: str,
    support_paragraphs: list[str],
    gap_paragraph: str,
    closing: str,
    *,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    min_words, max_words = cover_letter_word_range(STANDARD_COVER_MODE)
    allowed_acronyms = cover_allowed_acronyms(
        role_title=role_title,
        job_description=job_description,
        company_name=company_name,
    )
    opening_sentences = select_cover_sentences(
        opening,
        company_name,
        role_title,
        job_description,
        limit=2,
    )
    if not opening_sentences:
        fallback_opening = " ".join(sentence_list(_direct_opening(company_name, role_title, job_description))[:1]).strip()
        opening_sentences = [smooth_cover_letter_text(fallback_opening, allowed_acronyms=allowed_acronyms)]
    opening_paragraph = " ".join(opening_sentences[:1]).strip()

    proof_sentence_limit = 1 if support_paragraphs else 2
    middle_sentences = select_cover_sentences(
        proof,
        company_name,
        role_title,
        job_description,
        limit=proof_sentence_limit,
    )
    extras: list[str] = []
    for paragraph in support_paragraphs:
        extras.extend(
            select_cover_sentences(
                paragraph,
                company_name,
                role_title,
                job_description,
                limit=1,
            )
        )
    if gap_paragraph:
        extras.extend(
            select_cover_sentences(
                gap_paragraph,
                company_name,
                role_title,
                job_description,
                limit=1,
            )
        )

    closing_sentences = sentence_list(smooth_cover_letter_text(closing, allowed_acronyms=allowed_acronyms))
    closing_paragraph = " ".join(closing_sentences[:1]).strip() or smooth_cover_letter_text(closing, allowed_acronyms=allowed_acronyms)

    deduped_middle: list[str] = []
    for sentence in [*middle_sentences, *extras]:
        if any(cover_sentences_near_duplicate(sentence, existing) for existing in deduped_middle):
            continue
        deduped_middle.append(sentence)
    if not deduped_middle:
        fallback_middle = " ".join(sentence_list(smooth_cover_letter_text(proof, allowed_acronyms=allowed_acronyms))[:1]).strip()
        if fallback_middle:
            deduped_middle = [fallback_middle]

    middle_paragraph = " ".join(deduped_middle[:2]).strip()
    body_paragraphs = [opening_paragraph, middle_paragraph, closing_paragraph]

    deferred_extras = deduped_middle[2:]
    if len(opening_sentences) > 1:
        deferred_extras.extend(opening_sentences[1:])

    letter_text = "\n".join(["Dear Hiring Manager,", *body_paragraphs, "Thank you,", "Christian Estrada"])
    while word_count(letter_text) < min_words and deferred_extras:
        next_sentence = deferred_extras.pop(0)
        if next_sentence and next_sentence not in middle_paragraph:
            middle_paragraph = f"{middle_paragraph} {next_sentence}".strip()
            body_paragraphs[1] = middle_paragraph
            letter_text = "\n".join(["Dear Hiring Manager,", *body_paragraphs, "Thank you,", "Christian Estrada"])

    while word_count(letter_text) > max_words and len(sentence_list(middle_paragraph)) > 1:
        middle_paragraph = " ".join(sentence_list(middle_paragraph)[:-1]).strip()
        body_paragraphs[1] = middle_paragraph
        letter_text = "\n".join(["Dear Hiring Manager,", *body_paragraphs, "Thank you,", "Christian Estrada"])

    return [paragraph for paragraph in body_paragraphs if paragraph]


def finalize_body_paragraphs(
    opening: str,
    proof: str,
    support_paragraphs: list[str],
    gap_paragraph: str,
    closing: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> list[str]:
    min_words, max_words = cover_letter_word_range(mode)
    body_paragraphs: list[str] = [opening, proof]
    if support_paragraphs:
        body_paragraphs.extend(support_paragraphs)
    if gap_paragraph:
        candidate_paragraphs = body_paragraphs + [gap_paragraph, closing]
        candidate_text = "\n".join(["Dear Hiring Manager,", *candidate_paragraphs, "Thank you,", "Christian Estrada"])
        candidate_word_total = word_count(candidate_text)
        if min_words <= candidate_word_total <= max_words:
            body_paragraphs.append(gap_paragraph)
    body_paragraphs.append(closing)
    if len(body_paragraphs) > 4:
        body_paragraphs = body_paragraphs[:2] + [" ".join(body_paragraphs[2:-1])] + [body_paragraphs[-1]]
    return body_paragraphs


def concise_body_paragraphs(
    opening: str,
    proof: str,
    support_paragraphs: list[str],
    gap_paragraph: str,
    closing: str,
    *,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
) -> list[str]:
    return standard_body_paragraphs(
        opening,
        proof,
        support_paragraphs,
        gap_paragraph,
        closing,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
    )


def repair_cover_letter_draft_structure(
    draft: CoverLetterDraft,
    *,
    company_name: str,
    role_title: str,
    job_description: str,
) -> CoverLetterDraft:
    """Apply deterministic composer/validator contract repairs before DOCX creation."""

    paragraphs = list(draft.body_paragraphs)
    if not paragraphs:
        return draft
    opening = paragraphs[0]
    if word_count(opening) < 16 or company_name.lower() not in opening.lower() or role_title.lower() not in opening.lower():
        context = build_resume.natural_problem_phrase(build_resume.job_problem_profile(job_description))
        opening = normalize_spaces(
            f"The {role_title} role at {company_name} centers on {context}. {opening}"
        )
    paragraphs[0] = opening

    closing = paragraphs[-1]
    if not re.search(r"\b(?:conversation|discuss|welcome)\b", closing, re.I):
        closing = normalize_spaces(
            closing.rstrip(". ")
            + f". I would welcome the chance to discuss how this experience could support {company_name}."
        )
    paragraphs[-1] = closing

    repaired = paragraphs
    for _pass in range(3):
        next_pass = [prose_engine.repair_text(paragraph, "cover").text for paragraph in repaired]
        if next_pass == repaired:
            break
        repaired = next_pass
    unresolved = tuple(
        dict.fromkeys(
            finding.rule_id
            for paragraph in repaired
            for finding in prose_engine.validate_text(paragraph, "cover")
            if finding.severity == "fail"
        )
    )
    return replace(
        draft,
        body_paragraphs=tuple(repaired),
        repair_rule_ids=unresolved,
    )


def draft_output_path(output_docx: Path) -> Path:
    stem = output_docx.stem
    if " DRAFT" not in stem:
        stem = re.sub(r"( (?:Long )?Cover Letter)$", r" DRAFT\1", stem)
        if stem == output_docx.stem:
            stem += " DRAFT"
    return output_docx.with_name(stem + output_docx.suffix)


def cover_failure_rule_ids(failure_text: str) -> tuple[str, ...]:
    mappings = (
        ("COVER_OPENING_LENGTH", r"opening paragraph is .*words"),
        ("COVER_COMPANY_CONTEXT", r"first paragraph must name a company-specific role context"),
        ("COVER_FORWARD_CLOSE", r"closing must be forward-looking"),
        ("COVER_SEMICOLON_DENSITY", r"too many semicolons"),
        ("COVER_PROSE_QUALITY", r"prose-quality|prose quality"),
        ("COVER_DUPLICATE_PROOF", r"duplicated proof|repeated sentence"),
    )
    found = [rule_id for rule_id, pattern in mappings if re.search(pattern, failure_text, re.I)]
    return tuple(found or ("COVER_UNREPAIRED_VALIDATION",))


def build_document(
    company_name: str,
    role_title: str,
    job_description: str,
    resume_docx: Path,
    output_docx: Path,
    *,
    mode: str = DEFAULT_COVER_MODE,
    readiness: build_resume.ResumeReadiness | None = None,
    application_responses: tuple[question_prep.QualificationsResponse, ...] = (),
    force_bridge: bool = False,
) -> tuple[int, list[str], list[str], list[str], Path, bool]:
    normalized_mode = normalize_cover_mode(mode)
    resume_texts = paragraph_texts(resume_docx)
    resume_text = "\n".join(resume_texts)
    evidence = safe_selected_evidence_items_ordered(job_description, resume_text)
    brief = question_prep.active_positioning_brief(job_description, resume_text)
    plan = build_cover_letter_proof_first(
        brief=brief,
        mode=normalized_mode,
        company_name=company_name,
        role_title=role_title,
        resume_audit_state=resume_analysis.output_audit_state(resume_docx),
        output_docx=output_docx,
        force_bridge=force_bridge,
        job_description=job_description,
    )
    if readiness and readiness.hard_blockers:
        warning_note = "Resume bridge gaps to address honestly: " + build_resume.resume_gap_blocker_message(readiness)
        plan = replace(
            plan,
            warnings=tuple([*plan.warnings, warning_note]),
        )
    draft = repair_cover_letter_draft_structure(
        compose_cover_letter_proof_first_from_brief(plan),
        company_name=company_name,
        role_title=plan.prose_role_title or role_title,
        job_description=job_description,
    )
    validation_role_title = plan.prose_role_title or role_title
    allowed_acronyms = cover_allowed_acronyms(
        role_title=role_title,
        job_description=job_description,
        company_name=company_name,
    )

    document = Document()
    set_default_style(document)
    add_header(document, resume_texts)
    cover_warnings: list[str] = list(plan.warnings)

    add_paragraph(document, date.today().strftime("%B %-d, %Y") if sys.platform != "win32" else date.today().strftime("%B %#d, %Y"))
    add_blank_line(document)
    add_paragraph(document, draft.salutation)
    add_blank_line(document)

    for index, paragraph_text_value in enumerate(draft.body_paragraphs):
        if index:
            add_blank_line(document)
        add_paragraph(
            document,
            smooth_cover_letter_text(
                paragraph_text_value,
                warnings=cover_warnings,
                allowed_acronyms=allowed_acronyms,
            ),
        )
    add_blank_line(document)
    add_paragraph(document, "Thank you for your time and consideration,", after=SIGNATURE_GAP_PT)
    add_paragraph(document, "Christian Estrada")

    if not build_resume.jd_explicitly_requires_erp(job_description):
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                cleaned = question_prep.scrub_cover_answer_for_job(job_description, run.text)
                if cleaned != run.text:
                    run.text = cleaned
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    prose_check_body = cover_letter_prose_check_text(body_text)
    specificity_warnings: list[str] = []
    preflight_warnings: list[str] = []
    prose_report: dict[str, object] = {}
    trace_path: Path | None = None
    validation_failures: list[str] = []
    if draft.repair_rule_ids:
        validation_failures.append(
            "Shared prose repair did not converge. Rule IDs: " + ", ".join(draft.repair_rule_ids)
        )
    try:
        try:
            specificity_warnings, cover_warnings = validate_cover_letter_text(
                body_text,
                job_description,
                company_name,
                mode=normalized_mode,
                existing_warnings=cover_warnings,
            )
        except SystemExit as error:
            validation_failures.append(str(error).strip() or "Cover letter text validation failed.")
        preflight_warnings = cover_letter_preflight(
            body_text,
            company_name,
            validation_role_title,
            job_description,
            mode=normalized_mode,
        )
        try:
            prose_report = enforce_prose_quality(
                prose_check_body,
                "cover_letter_full",
                label="Cover letter",
                mode="fail",
                check_template_leakage=True,
            )
        except SystemExit as error:
            validation_failures.append(str(error).strip() or "Cover letter failed prose-quality checks.")
            prose_report = prose_quality_report(prose_check_body, "cover_letter_full")
        cover_warnings.extend(str(warning) for warning in prose_report.get("warnings", ()))
        cover_warnings = dedupe_warnings(cover_warnings)
        for warning in cover_warnings:
            print(f"COVER LETTER WARNING: {warning}")
        for warning in preflight_warnings:
            print(f"COVER LETTER PREFLIGHT: {warning}")
        try:
            assert_cover_letter_qc(
                body_text,
                company_name,
                validation_role_title,
                job_description,
                mode=normalized_mode,
            )
        except SystemExit as error:
            validation_failures.append(str(error).strip() or "Cover letter QC failed.")
        if validation_failures:
            raise SystemExit("; ".join(item for item in validation_failures if item))
    except SystemExit as error:
        failure_text = str(error).strip()
        if not failure_text or failure_text == "1":
            prose_failures = [str(item) for item in prose_report.get("failures", ())]
            if prose_failures:
                failure_text = "Cover letter failed prose-quality checks: " + "; ".join(prose_failures)
            else:
                failure_text = "Cover letter build failed."
        rule_ids = cover_failure_rule_ids(failure_text)
        failure_text = f"Rule IDs: {', '.join(rule_ids)}. {failure_text}"
        actual_output = draft_output_path(output_docx)
        if document.paragraphs:
            banner = document.paragraphs[0].insert_paragraph_before("DRAFT - REQUIRES HUMAN REVIEW")
            banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in banner.runs:
                run.bold = True
        actual_output.parent.mkdir(exist_ok=True)
        document.save(str(actual_output))
        feedback_path = actual_output.with_name(actual_output.stem + " Build Feedback.txt")
        feedback_path.write_text(
            "This draft did not pass all sendable-document checks.\n\n" + failure_text + "\n",
            encoding="utf-8",
        )
        trace_path = write_cover_letter_trace(
            plan,
            draft,
            actual_output,
            body_text=body_text,
            specificity_warnings=specificity_warnings,
            cover_warnings=cover_warnings,
            preflight_warnings=preflight_warnings,
            prose_report=prose_report,
            failure=failure_text,
        )
        print(f"COVER LETTER TRACE: {trace_path}")
        return min(len(evidence), 3), specificity_warnings, cover_warnings, preflight_warnings, actual_output, True
    output_docx.parent.mkdir(exist_ok=True)
    document.save(str(output_docx))
    trace_path = write_cover_letter_trace(
        plan,
        draft,
        output_docx,
        body_text=body_text,
        specificity_warnings=specificity_warnings,
        cover_warnings=cover_warnings,
        preflight_warnings=preflight_warnings,
        prose_report=prose_report,
    )
    print(f"COVER LETTER TRACE: {trace_path}")
    return min(len(evidence), 3), specificity_warnings, cover_warnings, preflight_warnings, output_docx, False


def validate_cover_letter_specificity(
    text: str,
    company_name: str,
    job_description: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> list[str]:
    """Return warnings when the cover letter is not specific enough to the company or role."""
    problems: list[str] = []
    normalized_mode = normalize_cover_mode(mode)
    if company_name:
        company_count = len(re.findall(re.escape(company_name), text, re.I))
        minimum_company_mentions = 2 if normalized_mode == STANDARD_COVER_MODE else 3
        if company_count < minimum_company_mentions:
            problems.append(
                f"Cover letter names {company_name} only {company_count} time(s). It should appear at least {minimum_company_mentions} times to feel specific rather than templated."
            )
    if normalized_mode == LONG_COVER_MODE and "because" not in text.lower():
        problems.append(
            "Cover letter never uses the word 'because.' Coaching note: use it explicitly because it forces a concrete reason."
        )
    sections = cover_letter_jd_sections(job_description)
    profile = build_resume.job_problem_profile(job_description)
    role_title = extract_role_title(job_description) or build_resume.extract_job_title(job_description) or ""
    lane_key = effective_lane_key(role_title, job_description, profile)
    specialty_terms = list(extract_cover_letter_terms(job_description, lane_key, sections))
    specialty_hits = sum(1 for term in specialty_terms if term.lower() in text.lower())
    if specialty_terms and specialty_hits == 0:
        problems.append(
            f"Cover letter does not reference any of the role's specialty areas: {', '.join(specialty_terms[:3])}. Add at least one to pass the specificity test."
        )
    body_paragraphs = cover_letter_body_paragraphs(text)
    first_paragraph = body_paragraphs[0] if body_paragraphs else text
    generic_opener_patterns = (
        r"\bexciting opportunity\b",
        r"\bi am excited\b",
        r"\binterested in (?:the|this) role\b",
        r"\bdrawn to\b",
        r"\bimpressed by\b",
        r"\bindustry leader\b",
        r"\binnovative company\b",
    )
    for pattern in generic_opener_patterns:
        match = re.search(pattern, first_paragraph, re.I)
        if match:
            problems.append(
                f"Offer blocker: generic opener phrase detected ('{match.group(0)}'). Replace it with a company-specific business problem or operating context."
            )
            break
    generic_closes = (
        "i look forward to hearing from you",
        "i look forward to the opportunity",
        "please find my resume attached",
    )
    for phrase in generic_closes:
        if phrase in text.lower():
            problems.append(
                f"Generic closing phrase detected: '{phrase}'. Replace with a specific conversation ask tied to the role's problem."
            )
    body_paragraphs = cover_letter_body_paragraphs(text)
    proof_dense_paragraphs = sum(1 for paragraph in body_paragraphs if paragraph_has_fast_proof(paragraph))
    minimum_proof_paragraphs = 1 if normalized_mode == STANDARD_COVER_MODE else 2
    if proof_dense_paragraphs < minimum_proof_paragraphs:
        problems.append(
            "Executive presence note: metric/scope density is thin for a senior skim. Add one more concrete proof marker so the letter sounds more senior and grounded."
        )
    context_warnings = business_context.business_context_audit(text, job_description, "cover letter")
    if normalized_mode == STANDARD_COVER_MODE:
        context_warnings = [
            warning
            for warning in context_warnings
            if "too little objective business context" not in warning.lower()
        ]
    problems.extend(context_warnings)
    return problems


def cover_letter_body_paragraphs(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    start = next((index for index, line in enumerate(lines) if line.lower().startswith("dear ")), -1)
    end = next((index for index, line in enumerate(lines) if line.lower().startswith("thank you")), len(lines))
    if start < 0 or end <= start:
        return []
    return lines[start + 1:end]


def paragraph_has_fast_proof(paragraph: str) -> bool:
    if re.search(r"\d|%|\$", paragraph):
        return True
    has_named_scope = bool(
        re.search(
            r"\b(?:users|sites|clients|accounts|workshops|qbrs|dashboards|reports|revenue|portfolio|systems)\b",
            paragraph,
            re.I,
        )
    )
    has_named_company = bool(
        re.search(
            rf"\b(?:{re.escape(build_resume.COMPANY_EAST_WEST)}|{re.escape(build_resume.COMPANY_APTEAN)}|{re.escape(build_resume.COMPANY_HOME_DEPOT)})\b",
            paragraph,
            re.I,
        )
    )
    has_result = bool(
        re.search(
            r"\b(?:reduced|improved|accelerated|stabilized|delivered|built|launched|enabled|protected|converted|resolved|prevented|saved|streamlined|retained|grew|cut|validated|created|designed|drove)\b",
            paragraph,
            re.I,
        )
    )
    return has_result and (has_named_scope or has_named_company)


def paragraph_has_context_or_proof(paragraph: str, company_name: str, role_title: str, job_description: str) -> bool:
    return paragraph_has_fast_proof(paragraph) or cover_sentence_has_context(
        paragraph,
        company_name,
        role_title,
        job_description,
    )


def cover_letter_duplicate_sentence_problem(body_paragraphs: list[str]) -> str | None:
    details = cover_letter_duplicate_sentence_details(body_paragraphs)
    if details:
        return (
            "cover letter QC failed: duplicated proof or repeated sentence logic detected. "
            "Keep each sentence doing different work."
        )
    return None


def cover_letter_duplicate_sentence_details(body_paragraphs: list[str]) -> dict[str, object]:
    sentences: list[str] = []
    sentence_origins: list[tuple[int, str]] = []
    for paragraph_index, paragraph in enumerate(body_paragraphs, start=1):
        for sentence in sentence_list(paragraph):
            sentence_origins.append((paragraph_index, sentence))
            sentences.append(sentence)
    for index, sentence in enumerate(sentences):
        for compare_index, comparison in enumerate(sentences[index + 1:], start=index + 1):
            same_paragraph = sentence_origins[index][0] == sentence_origins[compare_index][0]
            if cover_sentences_near_duplicate(sentence, comparison, same_paragraph=same_paragraph):
                return {
                    "left_sentence": sentence,
                    "right_sentence": comparison,
                    "left_paragraph": sentence_origins[index][0],
                    "right_paragraph": sentence_origins[compare_index][0],
                }
    return {}


def lowercase_proof_problem(body_paragraphs: list[str]) -> str | None:
    if len(body_paragraphs) < 2:
        return None
    proof_paragraph = body_paragraphs[1].strip()
    if sentence_starts_with_lowercase_fragment(proof_paragraph):
        return "cover letter QC failed: proof paragraph begins with a malformed lowercase fragment"
    return None


def trace_failure_details(text: str, failure: str) -> dict[str, object]:
    details: dict[str, object] = {}
    body_paragraphs = cover_letter_body_paragraphs(text)
    if "duplicated proof or repeated sentence logic detected" in failure:
        duplicate_details = cover_letter_duplicate_sentence_details(body_paragraphs)
        if duplicate_details:
            details["duplicate_sentence_pair"] = duplicate_details
    if "malformed lowercase fragment" in failure:
        details["proof_paragraph"] = body_paragraphs[1] if len(body_paragraphs) > 1 else ""
    return details


def cover_letter_thesis_sentence_count(body_paragraphs: list[str]) -> int:
    return sum(
        1
        for paragraph in body_paragraphs
        for sentence in sentence_list(paragraph)
        if cover_sentence_is_thesis(sentence)
    )


def lead_burial_check(text: str) -> str | None:
    body_paragraphs = cover_letter_body_paragraphs(text)
    if not body_paragraphs:
        return "cover letter QC failed: body paragraphs could not be located"

    first_proof_index = next(
        (index for index, paragraph in enumerate(body_paragraphs) if paragraph_has_fast_proof(paragraph)),
        None,
    )
    if first_proof_index is None:
        return (
            "cover letter QC failed: no paragraph surfaces concrete proof. Add a metric, scope marker, or clear outcome early in the letter."
        )
    if first_proof_index > 1:
        return (
            "cover letter QC failed: structural lead burial detected. Surface the first concrete proof by the second body paragraph, not later."
        )
    first_two = " ".join(body_paragraphs[:2])
    if build_resume.objective_context_signal_count(first_two) < 3:
        return (
            "cover letter QC failed: the opening buries concrete context. The first two body paragraphs should quickly show scale, workflow, or business impact."
        )
    return None


def dedupe_warnings(warnings: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for warning in warnings:
        cleaned = warning.strip()
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        ordered.append(cleaned)
    return ordered


def cover_letter_preflight(
    text: str,
    company_name: str,
    role_title: str,
    job_description: str,
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> list[str]:
    issues: list[str] = []
    body_paragraphs = cover_letter_body_paragraphs(text)
    opening = body_paragraphs[0] if body_paragraphs else ""
    proof = body_paragraphs[1] if len(body_paragraphs) > 1 else ""
    if opening and not (
        (company_name and re.search(re.escape(company_name), opening, re.I))
        or normalized_phrase_in_text(opening, role_title)
    ):
        issues.append("Opening does not mention the company or role.")
    if proof and word_count(proof) < 18:
        issues.append("Proof paragraph is too short to carry evidence weight.")
    min_words, max_words = cover_letter_word_range(mode)
    nonempty_lines = [line.strip() for line in text.splitlines() if line.strip()]
    letter_start = next((index for index, line in enumerate(nonempty_lines) if line.lower().startswith("dear ")), 0)
    total_words = word_count("\n".join(nonempty_lines[letter_start:]))
    if total_words < min_words or total_words > max_words:
        issues.append(f"Cover letter is {total_words} words. Target is {min_words}-{max_words}.")
    for pattern in PROMPT_LEAK_PATTERNS:
        match = re.search(pattern, text, re.I)
        if match:
            issues.append(f"Prompt-leak phrase found: {match.group(0)}.")
    for pattern in CLICHE_PATTERNS:
        match = re.search(pattern, text, re.I)
        if match:
            issues.append(f"Cliche found: {match.group(0)}.")
    if re.search(r"(?:^|\n)[•\-\*]\s+\w", "\n".join(body_paragraphs), re.I):
        issues.append("Bullet characters found in cover letter body.")
    if build_resume.keyword_hits(text, build_resume.audit_keywords(job_description)) < 4:
        issues.append("Cover letter has fewer than 4 job-description keyword hits.")
    return dedupe_warnings(issues)


def assert_cover_letter_qc(
    text: str,
    company_name: str,
    role_title: str,
    job_description: str = "",
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> None:
    """Enforce Master Appendix I cover letter quality controls before saving."""
    issues = cover_letter_qc_issues(text, company_name, role_title, job_description, mode=mode)
    if issues:
        fail("; ".join(issues))


def cover_letter_qc_issues(
    text: str,
    company_name: str,
    role_title: str,
    job_description: str = "",
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> list[str]:
    min_words, max_words = cover_letter_word_range(mode)
    body_paragraphs = cover_letter_body_paragraphs(text)
    issues: list[str] = []
    if not body_paragraphs:
        return ["cover letter QC failed: body paragraphs could not be located"]

    first_paragraph = body_paragraphs[0].strip()
    opening_problem = opening_quality_problem(first_paragraph, company_name, role_title, job_description)
    if opening_problem:
        issues.append(opening_problem)
    if first_paragraph.lower().startswith("i am writing to express my interest"):
        issues.append("cover letter QC failed: opening uses the banned phrase 'I am writing to express my interest'")
    if (
        normalize_cover_mode(mode) == STANDARD_COVER_MODE
        and company_name
        and first_paragraph.lower().startswith(f"{company_name.lower()} needs")
    ):
        issues.append("cover letter QC failed: standard cover letter should open with a mission or value hook, not 'Company needs'")

    generic_praise_patterns = (
        r"\b(?:admire|impressed by|drawn to)\s+(?:your|the)\s+(?:company|organization|mission|commitment)\b",
        r"\bindustry leader\b",
        r"\binnovative company\b",
        r"\bexciting opportunity\b",
    )
    if any(re.search(pattern, first_paragraph, re.I) for pattern in generic_praise_patterns):
        issues.append("cover letter QC failed: first paragraph uses generic praise instead of a specific research finding")

    first_paragraph_lower = first_paragraph.lower()
    context_terms = (
        "need", "needs", "mission", "customer", "client", "product", "platform",
        "implementation", "adoption", "process", "data", "analytics", "claims",
        "practice", "students", "manufacturing", "healthcare", "software",
        "growth", "operating", "workflow", "delivery", "transformation",
        "change", "stakeholder", "organizational", "leadership", "team",
        "program", "grader", "grading", "volunteer", "coaching", "education", "student",
    )
    company_in_opening = bool(company_name and re.search(re.escape(company_name), first_paragraph, re.I))
    role_in_opening = normalized_phrase_in_text(first_paragraph, role_title)
    has_context = any(term in first_paragraph_lower for term in context_terms) or opening_keyword_hit_count(first_paragraph, job_description) >= 2
    if not (company_in_opening and role_in_opening and has_context):
        issues.append("cover letter QC failed: first paragraph must name a company-specific role context, not generic praise")
    buried_lead_problem = lead_burial_check(text)
    if buried_lead_problem:
        issues.append(buried_lead_problem)
    duplicate_problem = cover_letter_duplicate_sentence_problem(body_paragraphs)
    if duplicate_problem:
        issues.append(duplicate_problem)
    lowercase_problem = lowercase_proof_problem(body_paragraphs)
    if lowercase_problem:
        issues.append(lowercase_problem)
    thesis_count = cover_letter_thesis_sentence_count(body_paragraphs)
    if thesis_count > 1:
        issues.append("cover letter QC failed: body contains more than one thesis sentence. Keep the argument direct and proof-first.")
    if len(body_paragraphs) > 1 and not paragraph_has_fast_proof(body_paragraphs[1]):
        issues.append("cover letter QC failed: proof paragraph must surface concrete proof instead of a generic fit summary.")
    for index, paragraph in enumerate(body_paragraphs, 1):
        if not paragraph_has_context_or_proof(paragraph, company_name, role_title, job_description):
            issues.append(
                f"cover letter QC failed: body paragraph {index} does not carry company context or concrete proof."
            )

    nonempty_lines = [line.strip() for line in text.splitlines() if line.strip()]
    letter_start = next((index for index, line in enumerate(nonempty_lines) if line.lower().startswith("dear ")), 0)
    letter_word_count = word_count("\n".join(nonempty_lines[letter_start:]))
    if letter_word_count < min_words or letter_word_count > max_words:
        issues.append(
            f"cover letter QC failed: total length is {letter_word_count} words; "
            f"expected {min_words}-{max_words}"
        )

    closing = body_paragraphs[-1].strip()
    if not re.search(r"\b(conversation|discuss|welcome)\b", closing, re.I):
        issues.append("cover letter QC failed: closing must be forward-looking and include conversation, discuss, or welcome")

    if re.search(r"\bi hope to hear from you\b", text, re.I):
        issues.append("cover letter QC failed: closing uses the banned phrase 'I hope to hear from you'")

    passive_close_patterns = (
        r"\bi look forward to hearing from you\.?$",
        r"\bplease find my resume attached\.?$",
        r"\bthank you for your consideration\.?$",
        r"\bthank you for your time and consideration\.?$",
    )
    if any(re.search(pattern, closing, re.I) for pattern in passive_close_patterns):
        issues.append("cover letter QC failed: body ends with a passive close instead of a forward-looking conversation ask")
    return issues


def validate_cover_letter_shape(text: str, *, mode: str = DEFAULT_COVER_MODE) -> None:
    issues = cover_letter_shape_issues(text, mode=mode)
    if issues:
        fail("; ".join(issues))


def cover_letter_shape_issues(text: str, *, mode: str = DEFAULT_COVER_MODE) -> list[str]:
    normalized_mode = normalize_cover_mode(mode)
    body_paragraphs = cover_letter_body_paragraphs(text)
    max_body_paragraphs = 3 if normalized_mode == STANDARD_COVER_MODE else 5
    min_body_paragraphs = 2 if normalized_mode == STANDARD_COVER_MODE else 5
    issues: list[str] = []
    if len(body_paragraphs) < min_body_paragraphs:
        issues.append(
            f"cover letter has too few body paragraphs; keep it to at least {min_body_paragraphs} natural paragraphs before the short closing"
        )
    if len(body_paragraphs) > max_body_paragraphs:
        issues.append(
            f"cover letter has too many body paragraphs; keep it to {max_body_paragraphs} natural paragraphs before the short closing"
        )
    body_text = "\n".join(body_paragraphs)
    colon_count = body_text.count(":")
    if normalized_mode == STANDARD_COVER_MODE and colon_count > 0:
        issues.append("cover letter body uses a colon in standard mode; keep the default letter in clean prose")
    if normalized_mode == LONG_COVER_MODE and colon_count > 1:
        issues.append("cover letter body uses too many colons and may read like segmented AI text")
    semicolon_count = body_text.count(";")
    max_semicolons = 0 if normalized_mode == STANDARD_COVER_MODE else 1
    if semicolon_count > max_semicolons:
        issues.append("cover letter body uses too many semicolons and reads like a stitched list instead of prose")
    if re.search(r"(^|[.!?]\s+)That\b", body_text):
        issues.append("cover letter contains a sentence starting with 'That'; rewrite with simpler, more direct phrasing")
    if normalized_mode == STANDARD_COVER_MODE:
        sentence_total = sum(len(sentence_list(paragraph)) for paragraph in body_paragraphs)
        if sentence_total < 4 or sentence_total > 6:
            issues.append(
                f"cover letter has {sentence_total} body sentences; keep the standard letter to 4-6 sentences total"
            )
    return issues


def validate_cover_letter_text(
    text: str,
    job_description: str | None = None,
    company_name: str = "",
    *,
    mode: str = DEFAULT_COVER_MODE,
    existing_warnings: list[str] | None = None,
) -> tuple[list[str], list[str]]:
    issues = cover_letter_text_issues(
        text,
        job_description=job_description,
        company_name=company_name,
        mode=mode,
    )
    if issues:
        fail("; ".join(issues))

    raw_text = text
    text = normalize_cover_dash_punctuation(text)
    active_job_description = job_description if job_description is not None else read_text(JOB_DESCRIPTION)
    problems = validate_cover_letter_specificity(text, company_name, active_job_description, mode=mode)
    cover_warnings = list(existing_warnings or [])
    collective_cover_patterns = (
        r"\bwe have\b",
        r"\bour team\b",
        r"\bwe delivered\b",
        r"\bwe built\b",
    )
    collective_hits = sum(1 for pattern in collective_cover_patterns if re.search(pattern, text, re.I))
    if collective_hits >= 3:
        problems.append(
            "Cover letter uses collective 'we' language three or more times. Rewrite to name what Christian personally did, not what the team did."
        )
    for paragraph in cover_letter_body_paragraphs(text):
        sentences = re.split(r"(?<=[.!?])\s+", paragraph)
        for index in range(1, len(sentences)):
            previous = sentences[index - 1].strip()
            current = sentences[index].strip()
            if previous and current:
                prev_has_i = bool(re.search(r"\bI\b", previous))
                if not prev_has_i and re.match(r"^I\b", current):
                    if re.match(r"^I would welcome the chance to discuss\b", current, re.I):
                        continue
                    cover_warnings.append(
                        "abrupt first-person switch detected between sentences. Review for natural flow: "
                        f"'...{previous[-60:]}' -> '{current[:60]}...'"
                    )
    return dedupe_warnings(problems), dedupe_warnings(cover_warnings)


def cover_letter_text_issues(
    text: str,
    job_description: str | None = None,
    company_name: str = "",
    *,
    mode: str = DEFAULT_COVER_MODE,
) -> list[str]:
    raw_text = text
    issues: list[str] = []
    if "--" in raw_text:
        issues.append("cover letter contains a double-dash; use standard punctuation instead")
    if re.search(r"\bI bring approximately\b", raw_text, re.I):
        issues.append("cover letter contains the generic experience summary 'I bring approximately'; replace it with direct proof")
    contraction_match = CONTRACTION_PATTERN.search(raw_text)
    if contraction_match:
        issues.append(f"cover letter contains a contraction: {contraction_match.group(0)}")
    for pattern, message in ABSTRACT_COVER_PATTERNS:
        if re.search(pattern, raw_text, re.I):
            issues.append(message)
    for pattern, label in build_resume.RESPONSIBILITY_SOFTENING_PATTERNS:
        if re.search(pattern, raw_text, re.I):
            issues.append(f"cover letter downplays responsibility with the banned phrase '{label}'")
    text = normalize_cover_dash_punctuation(text)
    lowered = text.lower()
    linkedin_text = re.sub(r"^https?://(?:www\.)?", "", build_resume.LINKEDIN_URL.lower()).rstrip("/")
    if linkedin_text not in lowered:
        issues.append("LinkedIn URL missing from cover letter contact line")
    try:
        assert_no_template_leakage(text)
    except SystemExit as error:
        issues.append(str(error))
    for pattern in PLACEHOLDER_PATTERNS:
        if re.search(pattern, text, re.I):
            issues.append(f"placeholder text detected in cover letter: {pattern}")
    if re.search(r"(?m)^\s*[•·–]", text):
        issues.append("cover letter contains bullet points")
    jd_artifacts = JD_ARTIFACT_PATTERN.findall(text)
    if jd_artifacts:
        issues.append(
            f"Cover letter contains raw JD artifacts: {jd_artifacts[:3]}. "
            "The opening generation injected job description section text."
        )
    if re.search(r"\bThe pattern across my work\b", text) or re.search(r"\bThat pattern is\b", text):
        issues.append("cover letter contains repetitive pattern framing")
    if re.search(r"\bThe business context is\b", text):
        issues.append("cover letter contains leaked internal business-context debugging text")
    if re.search(r"\bThe experience I bring\b", text):
        issues.append("cover letter contains the canned proof opener 'The experience I bring'")
    if re.search(r"\bthe consistent pattern was the same\b", text, re.I):
        issues.append("cover letter contains the redundant phrase 'the consistent pattern was the same'")
    if re.search(r"\bAlong the way\b", text):
        issues.append("cover letter contains the weak transition 'Along the way'")
    if re.search(r"\bpractical execution depth\b", text, re.I):
        issues.append("cover letter contains the vague phrase 'practical execution depth'")
    if re.search(r"\bI want to do more of\b", text, re.I):
        issues.append("cover letter contains the explicit aspiration phrase 'I want to do more of' instead of direct evidence")
    if re.search(r"\breputation\b.*\bmatters because\b", text, re.I):
        issues.append("cover letter opens with a reputation-first framing instead of the role's concrete problem")
    if re.search(r"\bWhat I bring to\b.*\bSo what that means\b", text, re.I | re.S):
        issues.append("cover letter contains the 'What I bring / So what that means' close formula")
    if "My work has consistently centered on turning unclear needs into practical operating outcomes" in text:
        issues.append("cover letter contains generic operating-outcomes thesis sentence")
    if build_resume.contains_ai_writing_word(text):
        issues.append("cover letter contains banned AI-writing words")
    issues.extend(cover_letter_shape_issues(text, mode=mode))
    active_job_description = job_description if job_description is not None else read_text(JOB_DESCRIPTION)
    try:
        build_resume.assert_no_erp_language_for_non_erp_role(text, active_job_description, "cover letter")
    except SystemExit as error:
        issues.append(str(error))
    try:
        build_resume.assert_no_unsupported_platform_action_claims(text, "cover letter")
    except SystemExit as error:
        issues.append(str(error))
    body_paragraphs = cover_letter_body_paragraphs(text)
    if body_paragraphs:
        opening_conflicts = business_context.unsupported_target_context_warnings(
            body_paragraphs[0],
            active_job_description,
            "cover letter opening",
        )
        if opening_conflicts:
            issues.append("cover letter opening drifts into unsupported target-company context: " + "; ".join(opening_conflicts))
    return dedupe_warnings(issues)


def build_cover_letter_for_inputs(
    *,
    job_description: str,
    resume_docx: Path,
    output_docx: Path,
    company_name: str,
    role_title: str = "",
    mode: str = DEFAULT_COVER_MODE,
) -> CoverLetterResult:
    normalized_mode = normalize_cover_mode(mode)
    require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    if not job_description.strip():
        fail("job description is empty; refusing to create a placeholder cover letter")

    supplied_context = relevant_supplied_context(company_name, role_title or "")
    registered_profile = _registered_firm_profile(company_name, role_title or "", job_description, supplied_context)
    # stub: implement detect_company_profile in build_resume.py before enabling this block.
    firm_profile = build_resume.detect_company_profile(company_name, job_description) if False else None
    if registered_profile:
        firm_key, profile_config = registered_profile
        debug_print(f"Firm profile matched: {firm_key}", flag="DEBUG_COVER_LETTER")
        debug_print(f"Cover style note: {profile_config.get('cover_style', '')}", flag="DEBUG_COVER_LETTER")
        debug_print(f"Proof emphasis: {profile_config.get('proof_emphasis', '')}", flag="DEBUG_COVER_LETTER")
    elif firm_profile:
        debug_print(f"Firm profile matched: {firm_profile['key']}", flag="DEBUG_COVER_LETTER")
        debug_print(f"Cover style note: {firm_profile.get('cover_note', '')}", flag="DEBUG_COVER_LETTER")
        debug_print(f"Avoid: {firm_profile.get('avoid', '')}", flag="DEBUG_COVER_LETTER")

    audit_status = "PASS"
    final_role_title = clean_role_title(role_title)
    role_warning = ""
    resume_audit_state = resume_analysis.output_audit_state(resume_docx)
    source_resume_text = build_resume.docx_visible_text_from_path(build_resume.choose_resume(job_description))
    readiness = build_resume.resume_readiness_for_output(
        job_description,
        resume_docx,
        source_resume_text=source_resume_text,
        audit_status=resume_audit_state,
    )
    if not final_role_title:
        final_role_title = build_resume.extract_job_title(job_description) or "the role"
        role_warning = "Role title was not extracted cleanly for the cover letter; using a fallback role label."
    if resume_audit_state != "PASS":
        audit_status = resume_audit_state
    application_responses: tuple[question_prep.QualificationsResponse, ...] = ()
    if normalized_mode == STANDARD_COVER_MODE:
        application_responses = question_prep.active_application_question_responses(
            job_description,
            require_prompts=True,
        )
    bullets_used, specificity_warnings, cover_warnings, preflight_warnings, actual_output_docx, is_draft = build_document(
        company_name,
        final_role_title,
        job_description,
        resume_docx,
        output_docx,
        mode=normalized_mode,
        readiness=readiness,
        application_responses=application_responses,
        force_bridge=bool(readiness and readiness.hard_blockers) or resume_audit_state in {"BRIDGE", "FAIL", "POOR"},
    )
    if role_warning:
        cover_warnings = dedupe_warnings([role_warning, *cover_warnings])
    render_checks.render_docx(actual_output_docx)
    if is_draft:
        audit_status = "DRAFT"

    return CoverLetterResult(
        company_name,
        final_role_title,
        resume_docx,
        actual_output_docx,
        bullets_used,
        audit_status,
        specificity_warnings,
        cover_warnings,
        preflight_warnings,
        normalized_mode,
    )


def build_cover_letter(mode: str = DEFAULT_COVER_MODE) -> CoverLetterResult:
    normalized_mode = normalize_cover_mode(mode)
    require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    require_file(JOB_DESCRIPTION, "job description")

    job_description = read_text(JOB_DESCRIPTION)
    if not job_description:
        fail("jobs/job_description.txt is empty; refusing to create a placeholder cover letter")

    company_name = build_resume.extract_output_name(job_description)
    output_target_name = build_resume.extract_output_target_name(job_description)
    role_title = extract_role_title(job_description)
    resume_docx = find_resume_output(job_description)
    resume_audit_state = resume_analysis.output_audit_state(resume_docx)
    mode_suffix = " Long Cover Letter" if normalized_mode == LONG_COVER_MODE else " Cover Letter"
    output_docx = OUTPUT_DIR / f"Christian Estrada - {output_target_name}{mode_suffix}.docx"
    if resume_audit_state == "FAIL":
        output_docx = OUTPUT_DIR / f"Christian Estrada - {output_target_name} FAIL{mode_suffix}.docx"
    elif resume_audit_state == "POOR":
        output_docx = OUTPUT_DIR / f"Christian Estrada - {output_target_name} POOR{mode_suffix}.docx"
    elif resume_audit_state == "BRIDGE":
        output_docx = OUTPUT_DIR / f"Christian Estrada - {output_target_name} BRIDGE{mode_suffix}.docx"
    return build_cover_letter_for_inputs(
        job_description=job_description,
        resume_docx=resume_docx,
        output_docx=output_docx,
        company_name=company_name,
        role_title=role_title or "",
        mode=normalized_mode,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Christian Estrada's tailored cover letter.")
    parser.add_argument(
        "--mode",
        default=DEFAULT_COVER_MODE,
        help="Use 'standard' for the default 80-170 word cover letter or 'long' for the explicit 180-240 word version. Legacy aliases such as 'concise' and 'full' still work.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    result = build_cover_letter(args.mode)
    print(f"Company: {result.company_name}")
    print(f"Role: {result.role_title}")
    print(f"Resume source: {result.resume_docx}")
    print(f"Output DOCX: {result.output_docx}")
    print(f"Proof points used: {result.bullets_used}")
    print(f"Final audit: {result.audit_status}")
    print(f"Mode: {result.mode}")
    for warning in result.specificity_warnings:
        print(f"SPECIFICITY WARNING: {warning}")


if __name__ == "__main__":
    main()
