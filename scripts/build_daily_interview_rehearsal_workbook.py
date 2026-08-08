"""Build the expanded daily interview rehearsal workbook from the approved story bank.

The workbook deliberately has two different reading modes:

* story pages for active practice, with coaching, follow-up answers, and lane variants;
* a clean reference section at the end with only the stories and spoken follow-up answers.

The three generator-only alternate cards are not read from the story bank and therefore
cannot leak into the numbered rehearsal material.
"""

from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import interview_intelligence


PROJECT_ROOT = Path(__file__).resolve().parents[1]
STORY_BANK = PROJECT_ROOT / "interview_prep" / "Christian Estrada - Project Delivery Interview Stories.md"
OUTPUT = PROJECT_ROOT / "Study" / "Daily_Interview_Rehearsal_Workbook.docx"

BLUE = "1F4E79"
MID_BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
GRAY = "5B6573"
LIGHT_GRAY = "E7EBEF"
GREEN = "2E6B4F"


FOLLOW_UPS: dict[int, tuple[str, ...]] = {
    1: (
        "Who actually owned the decision when the four parties disagreed?",
        "What would you do differently if you ran it again?",
        "How did you know the new workflow was genuinely compliant and not just working?",
    ),
    2: (
        "Which workstream nearly slipped, and what did you do?",
        "How did you decide training came last?",
        "What did go-live day actually look like?",
    ),
    3: (
        "How did you measure the 78 and the 22 percent?",
        "Who resisted the change, and how did you handle it?",
        "What broke after you shipped it?",
    ),
    4: (
        "What was the customer actually angry about underneath the escalation?",
        "How did you keep product and development engaged?",
        "Which of those accounts stayed, and which did not?",
    ),
    5: (
        "How did you decide what a good SMS interaction looked like?",
        "What did you measure to know it was working?",
        "What would you build differently now?",
    ),
    6: (
        "What specifically did you change about the training on site?",
        "How did you know they were actually ready and not just trained?",
        "What did you get wrong the first time?",
    ),
    7: (
        "What did you say no to, and who pushed back?",
        "How did you know the data foundation was the true critical path?",
        "What would have happened if you had trained earlier?",
    ),
    8: (
        "How did you show them it was adoption and not the product without sounding defensive?",
        "What did the weekly cadence actually cover?",
        "Were they right about anything?",
    ),
    9: (
        "Why did you assume the source data was clean?",
        "Who did you tell, and how fast?",
        "What exactly is the checkpoint you added?",
    ),
    10: (
        "What did the feedback sound like word for word?",
        "How did you practice it?",
        "Where do you still catch yourself running long?",
    ),
    11: (
        "What was the hardest part of the migration?",
        "How did you handle a site that resisted the new process?",
        "What would you do differently on the next one?",
        "Why did the role end?",
    ),
    12: (
        "Give me an example of two clients who needed opposite approaches.",
        "What is the most common reason implementations break?",
        "How does the vendor view differ from the customer view?",
    ),
    13: (
        "Give me a report you refused to build as asked.",
        "How did you validate the source data?",
        "Which dashboard actually changed a decision?",
    ),
    14: (
        "Tell me about a QBR that went badly.",
        "How do you deliver bad news to an executive?",
        "What do you do when the room disagrees in front of you?",
    ),
    15: (
        "How did the CEO react when you told him?",
        "How did you keep the customer engaged for 13 months?",
        "Should the original scope have caught this?",
    ),
    16: (
        "Who pushed back on holding the date?",
        "How did you decide it was severe enough to hold?",
        "What did you change in testing afterward?",
    ),
    17: (
        "How did you frame the cost so it did not sound like an IT wish list?",
        "What if he had said no?",
        "How did you verify the environment was actually ready?",
    ),
    18: (
        "What would have saved that account?",
        "When did you know it was lost?",
        "What does proactive actually look like in your week?",
    ),
    19: (
        "Why were people using spreadsheets in the first place?",
        "How did you get them to trust the system view?",
        "Did the spreadsheets come back?",
    ),
    20: (
        "What makes a requirement backlog-ready in your definition?",
        "How did you build the UAT scenarios?",
        "What slipped through anyway?",
    ),
    21: (
        "What did each side turn out to be protecting?",
        "Who was more right?",
        "How did you make the tradeoff visible without picking a side?",
    ),
    22: (
        "What did you do in your first two weeks?",
        "How did you handle a client question you could not answer?",
        "What is your ramp plan for this role?",
    ),
}


# These are deliberately short spoken examples, not scripts. They stay inside the
# approved anchor facts and answer the follow-up directly before stopping.
FOLLOW_UP_ANSWERS: dict[int, tuple[str, ...]] = {
    1: (
        "I owned the coordination and decision structure even though the four groups did not report to me. When they disagreed, I brought the issue back to scope, compliance, timing, and the party that owned the downstream risk.",
        "I would validate the bank and vendor dependencies even earlier and make the decision log visible from day one. The overall five-month structure worked, but earlier escalation would have reduced waiting time.",
        "We validated the workflow against the agreed compliance requirements, tested the file flow with the relevant parties, and deployed an auditable process rather than just confirming that one transaction worked.",
    ),
    2: (
        "The data and product-family foundation was the part that could have slipped and taken everything else with it. I locked that first, ran independent GL work in parallel, and protected the final training window for the real configuration.",
        "Training came last because training on an unfinished configuration creates false confidence. I prepared the teams early, but held the formal training close enough to cutover that they learned the system they would actually use.",
        "Go-live was the convergence point: product families, GL accounts, BOMs, and site training all had to be ready together. I stayed focused on the critical path so the systems and the people were ready at the same time.",
    ),
    3: (
        "I compared the manual-work steps before and after the automation and tracked the inventory discrepancies against the prior process. The result was 78 percent less manual work and 22 percent fewer discrepancies, with an audit trail to explain the change.",
        "The natural resistance was concern about changing a process people knew, even though it was error-prone. I showed the control points, kept the workflow auditable, and connected the change to less rework rather than presenting automation as change for its own sake.",
        "The durable improvement was the control structure: automated adjustments, repeatable list maintenance, and an audit trail. The point was not that nothing could ever fail; it was that failures became visible and explainable instead of hidden in manual entry.",
    ),
    4: (
        "Underneath the escalation was a mix of adoption gaps, integration problems, and stalled custom work. The customer was reacting to the experience of not seeing progress, so I addressed ownership and follow-through rather than arguing about whether the product was technically capable.",
        "I created a weekly recovery cadence with clear owners and brought product and development the specific integration or customization issue, its customer impact, and the next decision needed. That made the work actionable instead of another general escalation.",
        "The accounts moved from at-risk to retained, but I would not claim every relationship has the same outcome. The lesson is to distinguish the accounts that can be recovered through ownership and adoption work from the ones where the underlying fit is genuinely wrong.",
    ),
    5: (
        "I started with the customer journey: how the text should open, route, and close, and what should happen when automation could not resolve the need. That gave the channel a usable workflow before we added the greetings, closings, chatbot logic, and monitoring.",
        "I looked at early channel usage and whether the workflow could be repeated consistently. At that stage the evidence was about a working, documented channel and usable interaction flow, not an unsupported claim about a mature volume or revenue result.",
        "I would define the escalation and measurement plan even earlier, especially for handoffs from automation to people. The original pilot established the channel and its repeatable setup; a scaled version would need stronger service-level and outcome tracking.",
    ),
    6: (
        "I adapted the onboarding and training to the local teams instead of treating the central document as the training itself. I went on site in El Paso and Juarez and made the compliance, financial, and process steps understandable in the way those teams worked.",
        "Readiness meant more than attendance. I checked whether people could use the process, understood the steps that affected their work, and were ready for the cutover; that is why I describe the sites as go-live ready, not merely trained.",
        "The first assumption I corrected was that one remote training approach would transfer cleanly across sites. The change was to make adoption local and observable rather than assuming the same material would land everywhere.",
    ),
    7: (
        "I said no to changes that would have put the hard date at risk, and I explained the dependency rather than making it a personal preference. The pushback became easier to handle once the team could see which request would move the critical path.",
        "The product families and BOMs were the data foundation for the rest of the launch. If they were wrong or late, the downstream setup and training would be built on a moving target, so I sequenced around that dependency.",
        "Training earlier on the wrong configuration would have created rework and false confidence. I used the earlier period for preparation and held the hands-on training until the configuration was stable enough to support a clean cutover.",
    ),
    8: (
        "I used the customer's own usage and workflow evidence to separate adoption gaps from product behavior. I acknowledged the places where the product or custom work was genuinely part of the problem, then focused the conversation on the shared fixes instead of defending the product.",
        "The cadence covered the active adoption gap, the stalled integration or custom-work item, the owner for each action, and what would be different by the next meeting. That made the recovery visible and gave the customer a reason to believe the plan.",
        "Yes. They were right that the experience was not working for them. The useful distinction was that their conclusion about the root cause was incomplete, not that their frustration was wrong; I could accept the frustration while diagnosing the fix.",
    ),
    9: (
        "I assumed the source was cleaner than it was because the early inspection did not surface the mapping problem. The miss was mine: I changed the process afterward so the assumption could not carry all the way to cutover again.",
        "I raised it as soon as I confirmed the impact, paused the affected cutover step, and told the people who needed to know what was wrong, what was contained, and what validation would happen before production. I did not wait for a perfect explanation before escalating the risk.",
        "The checkpoint was an explicit mapping-validation step before cutover, followed by a re-run against the affected data. It turns the source-data assumption into a testable gate instead of something discovered during a later inspection.",
    ),
    10: (
        "The feedback was that my updates ran too long and buried the point. That was accurate, so I changed the structure rather than defending the amount of detail.",
        "I practiced leading with the outcome first, then offering detail only when the audience needed it. I used status meetings and client calls as repetitions until the shorter structure became automatic.",
        "I still watch for it when the work is complex or I care deeply about the context. My control is to state the result, give one proof point, and pause so the interviewer can choose whether to go deeper.",
    ),
    11: (
        "The hardest part was aligning operations, finance, and engineering across five sites and two countries, not configuring one screen. I treated the work as an alignment and adoption problem and kept the migration, training, and decision points connected.",
        "I did not assume resistance meant the site was unwilling. I went on site, adapted the onboarding and training, and made the compliance and financial steps understandable locally so the team could use the process rather than merely receive it.",
        "I would make the cross-site decision log and readiness checks even more explicit earlier. The core approach would stay the same: align the owners, validate the data, and make adoption observable before cutover.",
        "My role wrapped when the migration finished and the team consolidated. I keep that answer neutral and return to the implementation work.",
    ),
    12: (
        "One client might need a legacy on-premise approach with more attention to constraints and sequencing, while another could use the cloud configuration and a faster standard path. I adapted the implementation to the environment rather than forcing one playbook across both.",
        "The most common break is not one software feature; it is a mismatch between the assumed process, the actual environment, and who owns the next decision. My vendor-side experience taught me to surface those gaps early.",
        "The vendor view sees patterns across many configurations and can spot recurring failure points. The customer view feels the operational impact every day, so combining both perspectives helps me translate a solution into something the users can actually adopt.",
    ),
    13: (
        "I would push back on a request for a dashboard that only listed more numbers without clarifying the decision it needed to support. I would ask what action should change, then build the smallest validated view that made that action visible.",
        "I validated the source against the business definition, checked the data behavior in the underlying system, and used the tools appropriate to the audience, including SQL, Crystal Reports, or Power BI. The validation had to answer whether the number was trustworthy enough to drive the decision.",
        "The strongest examples are the dashboards and KPI tools that improved visibility into operational performance, customer experience metrics, or process gaps. The proof is not that a chart existed; it is that the view made the next decision clearer and more usable.",
    ),
    14: (
        "When a QBR went badly, I brought the conversation back to the outcome, the visible risk, and the decision required rather than trying to smooth over the tension. I let the tradeoff be explicit and left the room with owners and next steps.",
        "I lead with the business impact and the choice the executive needs to make, then give the shortest proof that supports it. I do not bury bad news in process detail because that reduces confidence rather than protecting it.",
        "I name the disagreement and translate what each person is protecting: outcome, workflow, risk, or decision rights. Then I restate the shared objective and make the tradeoff visible so the room can decide without turning the discussion into a personality contest.",
    ),
    15: (
        "The CEO needed to hear that the blocker was infrastructure readiness, not a vague statement that the product was failing. I framed the technical constraint in terms of the user's operational outcome and the decision needed to move forward.",
        "I kept the customer engaged by making the thirteen-month path visible: what was known, what needed validation, what decision was next, and how the scope would be protected. Long delivery only works when the customer can see progress and ownership.",
        "The original scope should have surfaced the environment risk earlier. The lesson was to ask the infrastructure and readiness questions during discovery, not treat them as implementation detail that can wait until the build.",
    ),
    16: (
        "The people protecting the date pushed back because holding it had a visible business cost. I explained that a known UAT defect in a critical path was a larger customer and delivery risk than moving the date, and I brought the decision back to impact rather than preference.",
        "It was severe enough to hold because the defect affected the readiness of the release, not just a cosmetic detail. I would rather make the risk explicit before go-live than ask the customer to discover it in production.",
        "I added stronger validation around the affected behavior and treated UAT as a decision gate rather than a formality. The improvement was to make the defect pattern visible earlier and give the team a clearer release decision.",
    ),
    17: (
        "I framed the cost as protection for the customer's delivery outcome: without the infrastructure readiness work, the promised system behavior would remain at risk. That made it a business continuity and customer-trust conversation, not an IT wish list.",
        "If the CEO had said no, I would have made the residual risk and the tradeoff explicit, documented the decision, and offered the safest viable path within the constraint. The goal is honest decision-making, not winning the argument.",
        "I verified readiness by checking the environment against the dependency that was blocking the outcome, then confirming the fix with the people who owned the technical and business sides. A verbal assurance alone would not have been enough.",
    ),
    18: (
        "Earlier intervention on the adoption and integration signals might have created a different outcome. I still separate that lesson from the fact that the account was technically supported; the relationship needed proactive health ownership, not only reactive issue resolution.",
        "I knew it was lost when the customer no longer had enough trust or momentum for the recovery plan to change the relationship. That is why proactive work has to start before the escalation becomes the first clear signal.",
        "Proactive means reviewing adoption and risk signals, naming the next action before the customer asks, and keeping ownership visible. It is a weekly habit of looking for drift, not a special meeting after the account is already angry.",
    ),
    19: (
        "People used spreadsheets because they were fast, familiar, and filled a gap in ownership and visibility. The solution was not to shame the workaround; it was to make the system view clearer, more reliable, and easier to use.",
        "I connected the system view to the actual decision and validated the data with the people who relied on it. Trust grew when the system answered the operational question without making them reconstruct the answer somewhere else.",
        "Some spreadsheet behavior can return if the system stops serving the workflow. The durable fix is clear ownership, usable visibility, and follow-through so the system remains the place where the work and next step are visible.",
    ),
    20: (
        "Backlog-ready means the business need is clear, the acceptance condition is explicit, the owner and dependencies are known, and the team can tell what 'done' means. It is a bridge from an ambiguous request to work that can actually be built and validated.",
        "I built UAT scenarios from the intended business behavior, the important workflow paths, and the failure conditions that would matter after release. That connected the requirement to how a customer or user would know the change worked.",
        "Some issues can still slip through because real environments are messy. The response is to capture the learning, tighten the acceptance or validation step, and make ownership clear rather than pretending the first release proved everything.",
    ),
    21: (
        "Operations was protecting a workable process and finance was protecting control and accuracy; engineering was protecting feasibility and system integrity. Once those interests were visible, the disagreement was easier to solve than when it looked like two groups simply wanted different things.",
        "I would not frame it as one side being universally right. The right decision depends on the outcome, the control requirement, the user impact, and what the system can support without creating a larger risk.",
        "I made the tradeoff visible by stating the shared objective, the options, what each option protected, and the cost or risk of each. That let the owners choose with the facts in front of them rather than making me the person who picked a side.",
    ),
    22: (
        "In the first two weeks I learned through live client problems: I mapped the product, listened for the workflow and terminology, and used real implementation questions to build fluency. That approach helped me become useful without pretending I knew the product before I did.",
        "I would say what I know, what I do not know yet, and how I will get to a reliable answer. Then I would confirm the answer with the right product or technical owner and close the loop with the client instead of guessing.",
        "My ramp plan is to learn the product and customer workflow together, build a question and terminology map, shadow the highest-value conversations, and take ownership of increasingly complete work. I would measure ramp by the quality of my decisions and follow-through, not just by how much documentation I read.",
    ),
}


LANE_FRAMES = {
    "Implementation and Delivery": "sequence, ownership, validation, and go-live readiness",
    "Customer Success and Account Management": "adoption, customer trust, risk ownership, and value realization",
    "Analytics and Operations": "data quality, process control, decision usefulness, and measurable improvement",
    "Solutions Consulting and Pre-Sales": "discovery, diagnosis, tradeoffs, and translating technical detail into a customer decision",
    "Change Enablement and Process Improvement": "behavior change, usable workflows, repeatability, and adoption that holds",
}


STORY_COMPETENCIES: dict[int, tuple[str, ...]] = {
    1: ("Implementation and integration", "Stakeholder alignment", "Project delivery"),
    2: ("Project delivery", "Implementation and integration", "Stakeholder alignment"),
    3: ("Process improvement", "Data and analytics"),
    4: ("Customer relationship building", "Stakeholder alignment"),
    5: ("AI adoption", "Process improvement", "Adaptability / fast ramp"),
    6: ("Stakeholder alignment", "Customer relationship building", "Implementation and integration"),
    7: ("Project delivery", "Requirements translation"),
    8: ("Customer relationship building", "Discovery", "Stakeholder alignment"),
    9: ("Implementation and integration", "Project delivery", "Data and analytics"),
    10: ("Adaptability / fast ramp", "Stakeholder alignment"),
    11: ("Implementation and integration", "Project delivery", "Requirements translation", "Stakeholder alignment"),
    12: ("Adaptability / fast ramp", "Implementation and integration", "Customer relationship building"),
    13: ("Data and analytics", "Requirements translation", "Discovery"),
    14: ("Stakeholder alignment", "Customer relationship building", "Requirements translation"),
    15: ("Discovery", "Stakeholder alignment", "Project delivery"),
    16: ("Project delivery", "Implementation and integration"),
    17: ("Discovery", "Stakeholder alignment", "Technical fluency gap"),
    18: ("Customer relationship building", "Discovery"),
    19: ("Process improvement", "Data and analytics", "Adaptability / fast ramp"),
    20: ("Requirements translation", "Project delivery", "Process improvement"),
    21: ("Stakeholder alignment", "Discovery"),
    22: ("Adaptability / fast ramp", "Discovery", "Implementation and integration"),
}


FIVE_TELLS: tuple[tuple[str, str], ...] = (
    ("Buried outcome", "The result did not arrive by sentence two"),
    ("Stream of consciousness", "Narrated the timeline instead of leading with the claim"),
    ("Hedging", '"I guess", "kind of", "just"'),
    ("Warm-up wandering", "Setup before the point"),
    ("Volunteering salary", "Raised compensation unprompted"),
)


LANE_LOOP_SEQUENCES: dict[str, tuple[str, ...]] = {
    "Implementation and Delivery": ("Story 11", "Story 7", "Story 9", "Story 16", "Story 20", "Story 22"),
    "Customer Success and Account Management": ("Story 4", "Story 8", "Story 18", "Story 14", "Story 19", "Story 12"),
    "Analytics and Operations": ("Story 3", "Story 13", "Story 19", "Story 21", "Story 20", "Story 1"),
    "Solutions Consulting and Pre-Sales": ("Story 12", "Story 17", "Story 15", "Story 14", "Story 13", "Story 22"),
    "Change Enablement and Process Improvement": ("Story 6", "Story 3", "Story 10", "Story 21", "Story 5", "Story 9"),
}


LANE_BANK_HEADINGS: dict[str, str] = {
    "Implementation and Delivery": "Implementation and delivery consultant",
    "Customer Success and Account Management": "Customer Success and account management",
    "Analytics and Operations": "Analytics and operations",
    "Solutions Consulting and Pre-Sales": "Solutions consulting and pre-sales",
    "Change Enablement and Process Improvement": "Change enablement and process improvement",
}


LANE_CLOSES: dict[str, str] = {
    "Implementation and Delivery": "Close by asking what delivery risk needs the most ownership in the first 90 days.",
    "Customer Success and Account Management": "Close by asking which customer outcome would make this hire successful in the first 90 days.",
    "Analytics and Operations": "Close by asking which operational decision or metric most needs clearer visibility.",
    "Solutions Consulting and Pre-Sales": "Close by asking where discovery most often changes the solution conversation.",
    "Change Enablement and Process Improvement": "Close by asking which adoption or process change matters most after implementation.",
}


LANE_SELF_REVIEW = (
    "Did you lead with the claim every time?",
    "Did you exceed two full-mode answers?",
    "How many tells did you hear on playback?",
    "Did you close with a question that made them think?",
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = LIGHT_GRAY, size: str = "6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text: str, anchor: str, *, color: str = BLUE, underline: bool = True) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def _clean_bank_text(value: str) -> str:
    cleaned = re.sub(r"\*\*", "", value)
    cleaned = re.sub(r"(?m)^\s*(?:---|___)\s*$", "", cleaned)
    return re.sub(r"\n{2,}", "\n", cleaned).strip()


def parse_prep_lines(text: str) -> dict[int, str]:
    section_match = re.search(
        r"^# PREP lines for Stories 1 through 11\s*(.*?)(?=^# Role-tailored lead-ins by lane|\Z)",
        text,
        re.M | re.S,
    )
    if not section_match:
        return {}
    section = section_match.group(1)
    entries = re.finditer(
        r"\*\*Story (\d+), [^\n]+?\.\*\*\s*(.*?)(?=\n\*\*Story \d+,|\Z)",
        section,
        re.S,
    )
    return {int(match.group(1)): _clean_bank_text(match.group(2)) for match in entries}


def parse_lane_lead_ins(text: str) -> dict[str, str]:
    matches = list(re.finditer(r"^## (.+)$", text, re.M))
    lead_ins: dict[str, str] = {}
    for lane, heading in LANE_BANK_HEADINGS.items():
        match = next((item for item in matches if item.group(1).strip() == heading), None)
        if match is None:
            raise RuntimeError(f"Story bank is missing the documented lane heading: {heading}")
        next_heading = next((item for item in matches if item.start() > match.start()), None)
        block = text[match.end() : next_heading.start() if next_heading else len(text)]
        lead_match = re.search(r'^\*\*Lead-in:\*\* "(.+)"$', block, re.M)
        if lead_match is None:
            raise RuntimeError(f"Story bank is missing a Lead-in line for: {heading}")
        lead_ins[lane] = lead_match.group(1)
    return lead_ins


def parse_story_bank() -> list[dict[str, object]]:
    text = STORY_BANK.read_text(encoding="utf-8")
    separate_prep = parse_prep_lines(text)
    matches = list(re.finditer(r"^## Story (\d+): (.+)$", text, re.M))
    stories: list[dict[str, object]] = []
    for index, match in enumerate(matches):
        block_start = match.end()
        next_heading = re.search(r"^## ", text[block_start:], re.M)
        block_end = block_start + next_heading.start() if next_heading else len(text)
        block = text[block_start:block_end]
        anchor_match = re.search(r"Anchor facts:\s*(.+?)(?=\n### Full mode)", block, re.S)
        anchor = re.sub(r"\s+", " ", anchor_match.group(1)).strip() if anchor_match else ""

        def mode(name: str, next_names: tuple[str, ...]) -> str:
            pattern = r"### " + re.escape(name) + r".*?\n(.*?)(?=\n#{1,3}\s|\Z)"
            found = re.search(pattern, block, re.S)
            if not found:
                return ""
            return _clean_bank_text("\n".join(found.group(1).splitlines()))

        full = mode("Full mode", ("CART mode", "Short mode", "PREP mode"))
        cart = mode("CART mode", ("Short mode", "PREP mode"))
        short = mode("Short mode", ("PREP mode",))
        number = int(match.group(1))
        prep = mode("PREP mode", ()) or separate_prep.get(number, "")
        if not all((full, cart, short, prep)):
            missing = [label for label, value in (("Full", full), ("CART", cart), ("Short", short), ("PREP", prep)) if not value]
            raise RuntimeError(f"Story {number} has empty mode(s): {', '.join(missing)}")
        stories.append(
            {
                "number": number,
                "title": match.group(2).strip(),
                "anchor": anchor,
                "full": full,
                "cart": cart,
                "short": short,
                "prep": prep,
                "competencies": STORY_COMPETENCIES[number],
                "followups": FOLLOW_UPS[number],
                "answers": FOLLOW_UP_ANSWERS[number],
            }
        )
    if [int(s["number"]) for s in stories] != list(range(1, 23)):
        raise RuntimeError("Story bank must contain exactly Stories 1 through 22")
    taxonomy_names = set(interview_intelligence.COMPETENCY_TAXONOMY)
    unknown = sorted({name for story in stories for name in story["competencies"] if name not in taxonomy_names})
    if unknown:
        raise RuntimeError(f"Workbook competency mapping contains unknown taxonomy names: {unknown}")
    return stories


def set_run_font(run, *, size: float = 10.5, bold: bool = False, color: str = "000000", italic: bool = False) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04
    for name, size, color, before, after in (
        ("Heading 1", 17, BLUE, 12, 6),
        ("Heading 2", 13.2, MID_BLUE, 9, 4),
        ("Heading 3", 11.2, BLUE, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Small Note" not in styles:
        st = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = normal
        st.font.size = Pt(8.8)
        st.font.color.rgb = RGBColor.from_string(GRAY)
        st.paragraph_format.space_after = Pt(3)
    if "Answer" not in styles:
        st = styles.add_style("Answer", WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = normal
        st.paragraph_format.left_indent = Inches(0.18)
        st.paragraph_format.space_after = Pt(5)
    if "Question" not in styles:
        st = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = normal
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
        st.paragraph_format.space_before = Pt(4)
        st.paragraph_format.space_after = Pt(2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Daily Interview Rehearsal Workbook  |  ")
    set_run_font(run, size=8.3, color=GRAY)
    add_page_field(footer)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Daily Interview Rehearsal Workbook")
    set_run_font(r, size=22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(3)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("22 stories | follow-up answers | lane variants | clean reference")
    set_run_font(r, size=10.5, color=GRAY)
    p.paragraph_format.space_after = Pt(12)


def add_section_band(document: Document, text: str, *, bookmark: str | None = None, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size=17 if level == 1 else 13.2 if level == 2 else 11.2, bold=True, color=BLUE if level != 2 else MID_BLUE)
    if bookmark:
        add_bookmark(p, bookmark, 1000 + len(document.paragraphs))
    return p


def add_labeled_paragraph(document: Document, label: str, text: str, *, style: str = "Normal") -> None:
    p = document.add_paragraph(style=style)
    lead = p.add_run(label + " ")
    set_run_font(lead, size=10.2, bold=True, color=BLUE)
    body = p.add_run(text)
    set_run_font(body, size=10.2)


def add_bullet(document: Document, text: str, *, color: str = "000000") -> None:
    p = document.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=10.0, color=color)


def add_story_index(document: Document, stories: list[dict[str, object]]) -> None:
    add_section_band(document, "Clickable Story Index", bookmark="story_index")
    p = document.add_paragraph(style="Small Note")
    r = p.add_run("Click any story to jump to its practice page. Each story page links back here.")
    set_run_font(r, size=8.8, color=GRAY, italic=True)
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    for start in range(0, len(stories), 2):
        row = table.add_row().cells
        for col in range(2):
            if start + col >= len(stories):
                row[col].text = ""
                continue
            story = stories[start + col]
            set_cell_shading(row[col], PALE_BLUE if (start // 2) % 2 == 0 else "FFFFFF")
            set_cell_margins(row[col])
            p = row[col].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            add_internal_link(p, f"Story {story['number']}: {story['title']}", f"story_{story['number']}")
            for run in p.runs:
                set_run_font(run, size=9.1, color=BLUE)
    document.add_paragraph()
    quick = document.add_paragraph(style="Small Note")
    add_internal_link(quick, "Jump to clean 22-story reference", "clean_reference")
    add_internal_link(quick, "  |  Jump to lane mock loops", "lane_loops")


def add_daily_core(document: Document) -> None:
    add_section_band(document, "Part 1: Daily Core", bookmark="daily_core")
    p = document.add_paragraph("Run this five-minute core every day. The stories get rotated; the spine stays fixed.")
    for r in p.runs:
        set_run_font(r, size=10.2)
    core = (
        ("The spine", "Claim. Proof. Bridge. Say the result first, give one concrete detail, tie it to the role, and stop."),
        ("Your differentiator", "I have implemented the same kind of software from both sides of the table. At the vendor I delivered to 80-plus clients, and on the customer side I led a five-site ERP migration end to end."),
        ("The umbrella pitch", "I take ambiguous cross-functional work and turn it into structured delivery that gets adopted. My best work is where business, operations, and technology have to move together."),
        ("The method sentence", "Discovery, then scope and requirements, then stakeholder alignment, then milestones with checkpoints, then build and validate, then cutover and adoption."),
        ("The layoff line", "My role wrapped when the migration finished and the team consolidated."),
        ("The philosophy-first answer", "I run a structured discovery-to-go-live approach. I understand the real problem, define scope, align the owners, set checkpoints, validate before cutover, and stay through adoption."),
        ("The compress rule", "If the outcome is not clear by sentence two, restart shorter. Count buried outcomes, stream of consciousness, hedging, warm-up words, and tactical detours."),
    )
    for label, text in core:
        add_labeled_paragraph(document, label + ":", text)
    document.add_paragraph()


def add_rotation(document: Document, stories: list[dict[str, object]]) -> None:
    add_section_band(document, "Part 2: 11-Day Rotation", bookmark="rotation")
    add_labeled_paragraph(document, "Rotation:", "Two stories per day. Pair one hard-evidence story with one human or behavioral story. Repeat after Day 11.")
    pairs = [(11, 12), (1, 10), (2, 9), (3, 8), (4, 6), (5, 7), (13, 18), (14, 21), (15, 19), (16, 20), (17, 22)]
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    headers = ("Day", "Story A", "Story B")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=9.2, bold=True, color="FFFFFF")
    by_num = {int(s["number"]): s for s in stories}
    for day, (a, b) in enumerate(pairs, 1):
        cells = table.add_row().cells
        for c, text in zip(cells, (str(day), f"{a}. {by_num[a]['title']}", f"{b}. {by_num[b]['title']}")):
            c.text = text
            set_cell_margins(c)
            if day % 2 == 0:
                set_cell_shading(c, PALE_BLUE)
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9.0, color=BLUE if c is cells[0] else "000000", bold=c is cells[0])
    document.add_page_break()


def add_mode(document: Document, label: str, text: str) -> None:
    if not text:
        return
    p = document.add_paragraph(style="Heading 3")
    r = p.add_run(label)
    set_run_font(r, size=11.2, bold=True, color=BLUE)
    for part in re.split(r"\n+", text):
        if part.strip():
            p = document.add_paragraph(part.strip())
            for run in p.runs:
                set_run_font(run, size=9.8)


def add_recall_prompt(document: Document) -> None:
    p = document.add_paragraph(style="Small Note")
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE_BLUE)
    p_pr.append(shading)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    lead = p.add_run("Covered-page recall: ")
    set_run_font(lead, size=9.6, bold=True, color=BLUE)
    body = p.add_run("Before reading below, say the claim, one proof detail, and the result in about 20 seconds.")
    set_run_font(body, size=9.4, italic=True, color=GRAY)


def add_competencies(document: Document, story: dict[str, object]) -> None:
    competencies = ", ".join(str(item) for item in story["competencies"])
    add_labeled_paragraph(document, "Competencies tested:", competencies)


def add_rep_scoring_table(document: Document) -> None:
    add_section_band(document, "Rep Score", level=3)
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_borders(table)
    headers = ("Tell", "Meaning", "Count")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=70, bottom=70)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.9, bold=True, color="FFFFFF")
    for tell, meaning in FIVE_TELLS:
        cells = table.add_row().cells
        for cell, text in zip(cells, (tell, meaning, "____")):
            cell.text = text
            set_cell_margins(cell, top=65, bottom=65)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=8.6, bold=cell is cells[0])
    footer = table.add_row().cells
    footer[0].merge(footer[1])
    footer[0].text = "Time: ______ seconds"
    footer[2].text = "Clean pass: □"
    for cell in footer:
        set_cell_margins(cell, top=70, bottom=70)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.8, bold=True, color=BLUE)


def add_notes_block(document: Document) -> None:
    add_section_band(document, "Notes", level=3)
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table, color=LIGHT_GRAY, size="6")
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
    cell.text = "Write observations here:\n\n\n"
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=9.0, color=GRAY)


def add_lane_loop_questions(document: Document) -> None:
    questions = (
        "Tell me about yourself.",
        "Walk me through your most relevant project.",
        "Tell me about a failure.",
        "Tell me about a disagreement.",
        "Why this role, and what would you do in your first 90 days?",
        "What questions do you have for me?",
    )
    add_section_band(document, "Six-question interview run", level=3)
    for index, question in enumerate(questions, 1):
        add_labeled_paragraph(document, f"{index}.", question)


def add_lane_self_review(document: Document) -> None:
    add_section_band(document, "After-loop self-review", level=3)
    for prompt in LANE_SELF_REVIEW:
        add_bullet(document, prompt)


def add_lane_variants(document: Document, story: dict[str, object]) -> None:
    add_section_band(document, "Lane Variants", level=3)
    short = str(story["short"]).strip() or str(story["full"]).splitlines()[0]
    for lane, frame in LANE_FRAMES.items():
        p = document.add_paragraph(style="Answer")
        lead = p.add_run(lane + ": ")
        set_run_font(lead, size=9.8, bold=True, color=GREEN)
        rest = p.add_run(f"Use this story to show {frame}. Keep the proof anchored in: {short}")
        set_run_font(rest, size=9.5)


def add_followup_answers(document: Document, story: dict[str, object]) -> None:
    add_section_band(document, "Follow-Up Questions and Example Answers", level=3)
    p = document.add_paragraph(style="Small Note")
    r = p.add_run("These are spoken examples. Answer the question directly; do not restart the whole story.")
    set_run_font(r, size=8.8, color=GRAY, italic=True)
    questions = story["followups"]
    answers = story["answers"]
    for question, answer in zip(questions, answers):
        p = document.add_paragraph(style="Question")
        r = p.add_run("Q. " + question)
        set_run_font(r, size=9.9, bold=True, color=BLUE)
        p = document.add_paragraph(style="Answer")
        r = p.add_run("A. " + answer)
        set_run_font(r, size=9.6)


def add_story_page(document: Document, story: dict[str, object]) -> None:
    number = int(story["number"])
    heading = document.add_paragraph(style="Heading 1")
    add_bookmark(heading, f"story_{number}", 2000 + number)
    r = heading.add_run(f"Story {number}: {story['title']}")
    set_run_font(r, size=17, bold=True, color=BLUE)
    back = document.add_paragraph(style="Small Note")
    add_internal_link(back, "Back to clickable story index", "story_index")
    for run in back.runs:
        set_run_font(run, size=8.8, color=BLUE)
    add_recall_prompt(document)
    add_labeled_paragraph(document, "What this story proves:", str(story["anchor"]))
    add_labeled_paragraph(document, "Core answer:", str(story["short"]).strip() or str(story["full"]).splitlines()[0])
    add_mode(document, "PREP mode", str(story["prep"]))
    add_mode(document, "Short mode", str(story["short"]))
    add_mode(document, "Full mode", str(story["full"]))
    add_mode(document, "CART mode", str(story["cart"]))
    add_competencies(document, story)
    add_followup_answers(document, story)
    add_lane_variants(document, story)
    add_rep_scoring_table(document)
    add_notes_block(document)
    document.add_page_break()


def add_lane_loops(document: Document, bank_lead_ins: dict[str, str]) -> None:
    add_section_band(document, "Part 3: Lane Mock Loops", bookmark="lane_loops")
    for lane, sequence in LANE_LOOP_SEQUENCES.items():
        add_section_band(document, lane, level=2)
        add_labeled_paragraph(document, "Open with:", bank_lead_ins[lane])
        add_labeled_paragraph(document, "Run:", " → ".join(sequence))
        add_labeled_paragraph(document, "Hiring-manager lens:", LANE_FRAMES[lane].capitalize() + ".")
        add_lane_loop_questions(document)
        add_lane_self_review(document)
        add_labeled_paragraph(document, "Close:", LANE_CLOSES[lane])


def add_competency_coverage_map(document: Document, stories: list[dict[str, object]]) -> None:
    add_section_band(document, "Competency Coverage Map", bookmark="competency_map")
    by_competency: dict[str, list[int]] = {name: [] for name in interview_intelligence.COMPETENCY_TAXONOMY}
    for story in stories:
        for competency in story["competencies"]:
            by_competency[str(competency)].append(int(story["number"]))
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Competency", "Stories", "What to prove")):
        cell.text = text
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=75, bottom=75)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.9, bold=True, color="FFFFFF")
    for competency, metadata in interview_intelligence.COMPETENCY_TAXONOMY.items():
        cells = table.add_row().cells
        cells[0].text = competency
        cells[1].text = ", ".join(f"Story {number}" for number in by_competency[competency])
        cells[2].text = ", ".join(metadata["triggers"][:3])
        for cell in cells:
            set_cell_margins(cell, top=65, bottom=65)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=8.4, bold=cell is cells[0])


def add_rep_log(document: Document) -> None:
    add_section_band(document, "Rep Log", bookmark="rep_log")
    add_labeled_paragraph(document, "Use:", "Record the day, lane, time, tell count, and whether the answer passed cleanly.")
    table = document.add_table(rows=1, cols=7)
    table.autofit = False
    set_table_borders(table)
    headers = ("Date", "Stories", "Lane", "Time", "Tell count", "Clean pass", "Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, top=75, bottom=75)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=8.5, bold=True, color="FFFFFF")
    for _ in range(12):
        cells = table.add_row().cells
        for cell in cells:
            cell.text = ""
            set_cell_margins(cell, top=105, bottom=105)


def add_clean_reference(document: Document, stories: list[dict[str, object]]) -> None:
    add_section_band(document, "Part 4: Clean 22-Story Reference", bookmark="clean_reference")
    p = document.add_paragraph()
    r = p.add_run("This section is a straight-through reference: story answer, then follow-up answers. No practice instructions.")
    set_run_font(r, size=9.2, color=GRAY, italic=True)
    for story in stories:
        p = document.add_paragraph(style="Heading 2")
        r = p.add_run(f"Story {story['number']}: {story['title']}")
        set_run_font(r, size=13.2, bold=True, color=MID_BLUE)
        p = document.add_paragraph()
        r = p.add_run("Answer: ")
        set_run_font(r, size=9.8, bold=True, color=BLUE)
        r = p.add_run(str(story["full"]).split("\n")[0].strip() or str(story["short"]))
        set_run_font(r, size=9.6)
        for question, answer in zip(story["followups"], story["answers"]):
            p = document.add_paragraph()
            r = p.add_run("Q: " + question)
            set_run_font(r, size=9.5, bold=True, color=BLUE)
            p = document.add_paragraph()
            r = p.add_run("A: " + answer)
            set_run_font(r, size=9.3)


def build() -> Path:
    stories = parse_story_bank()
    bank_lead_ins = parse_lane_lead_ins(STORY_BANK.read_text(encoding="utf-8"))
    document = Document()
    configure_document(document)
    add_title(document)
    add_labeled_paragraph(document, "Purpose:", "Keep the generalized stories fresh every day so each job-specific guide becomes a different emphasis, not a different identity.")
    add_story_index(document, stories)
    add_daily_core(document)
    add_rotation(document, stories)
    for story in stories:
        add_story_page(document, story)
    add_lane_loops(document, bank_lead_ins)
    add_competency_coverage_map(document, stories)
    add_rep_log(document)
    document.add_page_break()
    add_clean_reference(document, stories)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    print(f"Created {build()}")
