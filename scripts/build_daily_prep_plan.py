#!/usr/bin/env python3
"""Build Christian's daily interview prep plan."""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path
from typing import Sequence

from docx import Document

import _bootstrap

_bootstrap.ensure_script_path()
_bootstrap.configure_fresh_pycache()

import interview_intelligence


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"


def _read_active_job_description() -> str:
    if not JOB_DESCRIPTION.exists():
        return ""
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip()


def _mode_label(mode: str) -> str:
    return mode.replace("_", " ").title().replace("Job Search", "Job-Search")


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def document_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def render_daily_prep_plan(plan: interview_intelligence.DailyPrepPlan) -> Document:
    document = Document()
    document.add_heading("Daily Prep Plan", level=1)
    document.add_paragraph(f"Mode: {_mode_label(plan.mode)}")
    document.add_paragraph(f"Date: {plan.plan_date.isoformat()}")
    document.add_paragraph(plan.emphasis)

    document.add_heading("Today's Reps", level=2)
    for rep in plan.reps:
        document.add_heading(f"{rep.title} ({rep.duration_minutes} minutes, {rep.weight})", level=3)
        for line in rep.instructions:
            add_bullet(document, line)
        if rep.proof_reference:
            document.add_paragraph(f"Reference: {rep.proof_reference}")

    document.add_heading("Completion Log", level=2)
    document.add_paragraph(plan.completion_prompt)
    document.add_paragraph("Columns: date, mode, reps_done, hedge_count, self_rated_clarity.")
    if plan.question_bank_checklist:
        document.add_heading("Question Bank Checklist", level=2)
        for item in plan.question_bank_checklist:
            add_bullet(document, item)
    return document


def build_daily_prep_plan_docx(
    *,
    mode: str = "job_search",
    output_dir: Path = OUTPUT_DIR,
    job_description: str | None = None,
) -> Path:
    plan = interview_intelligence.build_daily_prep_plan(
        mode,
        job_description=_read_active_job_description() if job_description is None else job_description,
    )
    output_dir.mkdir(exist_ok=True)
    output = output_dir / f"Christian Estrada - Daily Prep Plan - {_mode_label(plan.mode)} - {date.today().isoformat()}.docx"
    document = render_daily_prep_plan(plan)
    rendered_text = document_text(document)
    interview_intelligence.assert_safe_generated_text(rendered_text, interview_intelligence.load_self_inventory())
    document.save(output)
    print(f"Daily prep plan created: {output}")
    return output


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a daily interview prep plan.")
    parser.add_argument("--mode", choices=interview_intelligence.DAILY_PREP_MODES, default="job_search")
    parser.add_argument("--log-complete", action="store_true", help="Append a completed practice entry to scratch/prep_log.csv.")
    parser.add_argument("--reps-done", type=int, default=0)
    parser.add_argument("--hedge-count", type=int, default=0)
    parser.add_argument("--self-rated-clarity", type=int, default=0)
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> None:
    args = parse_args(argv)
    build_daily_prep_plan_docx(mode=args.mode)
    if args.log_complete:
        log_path = interview_intelligence.append_daily_prep_log(
            args.mode,
            args.reps_done,
            args.hedge_count,
            args.self_rated_clarity,
        )
        print(f"Daily prep log appended: {log_path}")


if __name__ == "__main__":
    main()
