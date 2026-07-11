"""Build an optional detailed interview guide with sample answers.

This is intentionally separate from the standard interview cheat sheet so the
normal resume run stays lightweight. Use this only when an interview is booked
and a longer prep document is useful.

Sources:
- jobs/job_description.txt
- the generated resume for the current company
- optional jobs/company_research.txt
- optional jobs/interview_notes.txt
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Mapping, Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import build_cover_letter
import build_debrief_analysis as debrief_analysis
import build_interview_cheat_sheet as cheat
import build_resume
import business_context
import interview_context
import question_prep
import prose_engine
import resume_analysis
from config.language_rules import PLACEHOLDER_PATTERNS, remove_approved_bracketed_metadata
from modules.employer_playbooks import consulting_bigfour
from modules.employer_playbooks import state_farm as state_farm_playbook
from modules.employer_playbooks.state_farm import (
    PrepInsights,
    add_state_farm_full_workbook,
    is_state_farm_active,
    state_farm_prep_insights,
    validate_state_farm_workbook_text,
)
import render_checks
from utils import assert_company_name_in_source, assert_no_template_leakage, clean_source_text, enforce_prose_quality, fail, optional_text, read_text


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
COMPANY_RESEARCH = PROJECT_ROOT / "jobs" / "company_research.txt"
INTERVIEW_NOTES = PROJECT_ROOT / "jobs" / "interview_notes.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"

FONT = "Carlito"
NAME_BLUE = RGBColor(31, 78, 121)
SECTION_BLUE = RGBColor(31, 78, 121)
SUBSECTION_BLUE = RGBColor(47, 84, 150)
TITLE_SIZE = 22
SUBTITLE_SIZE = 12
SECTION_SIZE = 15
BODY_SIZE = 12
SMALL_SIZE = 10.5

# Visual styling palette, matched to the approved Claude-built phone-screen guide
# reference design (dark navy banners and headers, mid-blue section bars, light
# blue answer boxes, pale yellow tip callouts, banded navy-header tables).
BANNER_NAVY_HEX = "1F3864"
SECTION_BAR_HEX = "2E74B5"
CARD_HEADER_NAVY_HEX = "1F3864"
ANSWER_BOX_HEX = "DCE6F1"
TIP_BOX_HEX = "FFF2CC"
TABLE_HEADER_NAVY_HEX = "1F3864"
TABLE_ROW_ALT_HEX = "F2F2F2"
WHITE_HEX = "FFFFFF"
SOFT_WHITE_HEX = "D9E2F0"


@dataclass(frozen=True)
class DetailedGuideResult:
    company_name: str
    role_title: str
    resume_docx: Path
    output_docx: Path


@dataclass(frozen=True)
class StoryAnswerParts:
    spoken: str
    coaching_note: str = ""


NAMED_ERP_PLATFORM_REPLACEMENTS = (
    (r"\bAptean Intuitive\b", "enterprise platform"),
    (r"\bAptean Encompix\b", "software platform"),
    (r"\bEpicor Kinetic\b", "successor platform"),
    (r"\bMicrosoft Dynamics 365\b", "enterprise platform"),
    (r"\bOracle ERP\b", "enterprise platform"),
)


def iter_all_paragraphs(document: Document):
    """Yields every paragraph in the document, including paragraphs nested
    inside table cells. Several sections (the title banner, section bars,
    subsection headers, Q&A answer/tip boxes, and the Story Anchor System
    table) are rendered as single-cell or banded tables for visual styling, so
    a plain document.paragraphs walk silently skips their text. Anything that
    scrubs, validates, or audits the document's full rendered content must use
    this instead of document.paragraphs directly."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)


def document_text(document: Document) -> str:
    """Full visible text of the document, including table-cell content. See
    iter_all_paragraphs() for why this is required instead of joining
    document.paragraphs directly."""
    return "\n".join(paragraph.text for paragraph in iter_all_paragraphs(document))


def scrub_document_for_job_language(document: Document, job_description: str) -> None:
    if build_resume.jd_explicitly_requires_erp(job_description):
        return
    for paragraph in iter_all_paragraphs(document):
        for run in paragraph.runs:
            cleaned = build_resume.scrub_erp_language_for_non_erp_text(run.text, job_description)
            for pattern, replacement in NAMED_ERP_PLATFORM_REPLACEMENTS:
                cleaned = re.sub(pattern, replacement, cleaned, flags=re.I)
            if cleaned != run.text:
                run.text = cleaned



def context_mentions_company(text: str, company_name: str) -> bool:
    return interview_context.context_mentions_company(text, company_name)


def relevant_company_research(text: str, company_name: str) -> str:
    return interview_context.relevant_company_context(text, company_name)


def safe_lines(text: str) -> list[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]


def set_default_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    style = document.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)


# ---------------------------------------------------------------------------
# Low-level OOXML styling primitives. python-docx has no first-class API for
# cell/paragraph shading or cell margins, so these reach into the underlying
# XML directly. Every higher-level helper below (banners, section bars, Q&A
# cards, banded tables) is built on top of these three functions.
# ---------------------------------------------------------------------------


def _content_width(document: Document) -> int:
    section = document.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, *, top: int = 110, bottom: int = 110, left: int = 160, right: int = 160) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _no_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tblPr.append(borders)


def _thin_table_borders(table, hex_color: str = "BFBFBF") -> None:
    """Adds a thin uniform border/gridline to a table, matching the reference
    guide's data tables (e.g. call pacing, answer length guide), as distinct
    from the borderless colored boxes used for banners and headers."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), hex_color)
        borders.append(node)
    tblPr.append(borders)


def _repeat_header_row(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trPr.append(node)


def _prevent_row_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    trPr.append(node)


def _set_table_width(table, width: int) -> None:
    """Sets the table-level width explicitly (w:tblW) in addition to column and
    cell widths. python-docx's column/cell width setters alone are not always
    honored by Word's layout engine; the explicit tblW makes the fixed width
    authoritative."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _full_width_box_table(document: Document) -> tuple:
    """A borderless, full-content-width single-cell table used as a colored box."""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = _content_width(document)
    table.columns[0].width = width
    cell = table.rows[0].cells[0]
    cell.width = width
    _no_table_borders(table)
    _set_table_width(table, width)
    return table, cell


def _set_cell_text(cell, lines: Sequence[tuple], *, base_size: float = BODY_SIZE) -> None:
    """lines is a sequence of (text, kwargs) where kwargs may set bold/italic/color/size."""
    cell.paragraphs[0].text = ""
    for index, (text, kwargs) in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(kwargs.get("size", base_size))
        run.bold = kwargs.get("bold", False)
        run.italic = kwargs.get("italic", False)
        color = kwargs.get("color")
        if color is not None:
            run.font.color.rgb = color


def _spacer(document: Document, *, points: float = 8) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.space_before = Pt(0)


def add_title(document: Document, company_name: str, role_title: str) -> None:
    """Full-width dark navy banner: bold white name/title line, italic light
    subtitle line with company, role, and date. Mirrors the reference guide's
    header banner."""
    table, cell = _full_width_box_table(document)
    _shade_cell(cell, BANNER_NAVY_HEX)
    _set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    date_format = "%B %#d, %Y" if sys.platform == "win32" else "%B %-d, %Y"
    generated = date.today().strftime(date_format)
    _set_cell_text(
        cell,
        [
            ("CHRISTIAN ESTRADA — DETAILED INTERVIEW GUIDE", {"bold": True, "size": TITLE_SIZE - 4, "color": RGBColor(0xFF, 0xFF, 0xFF)}),
            (f"{company_name}  |  {role_title}  |  {generated}", {"italic": True, "size": SUBTITLE_SIZE, "color": RGBColor(0xD9, 0xE2, 0xF0)}),
        ],
        base_size=BODY_SIZE,
    )
    _spacer(document, points=10)


def add_section(document: Document, title: str) -> None:
    """Full-width mid-blue section bar with bold white uppercase text, matching
    the reference guide's section dividers (e.g. "PRE-CALL ROUTINE")."""
    table, cell = _full_width_box_table(document)
    _shade_cell(cell, SECTION_BAR_HEX)
    _set_cell_margins(cell, top=80, bottom=80, left=160, right=160)
    _set_cell_text(cell, [(title.upper(), {"bold": True, "color": RGBColor(0xFF, 0xFF, 0xFF), "size": SECTION_SIZE - 2})])
    _spacer(document, points=6)


def add_subsection(document: Document, title: str) -> None:
    """Compact dark-navy header bar used for individual prompts/questions inside
    a section, matching the reference guide's per-question header treatment."""
    table, cell = _full_width_box_table(document)
    _shade_cell(cell, CARD_HEADER_NAVY_HEX)
    _set_cell_margins(cell, top=60, bottom=60, left=160, right=160)
    _set_cell_text(cell, [(title, {"bold": True, "color": RGBColor(0xFF, 0xFF, 0xFF), "size": BODY_SIZE})])
    _spacer(document, points=4)


def add_body(document: Document, text: str, *, small: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SMALL_SIZE if small else BODY_SIZE)


def add_answer_box(document: Document, text: str, *, label: str = "SAY THIS:") -> None:
    """Light blue answer box that follows a navy subsection/question header,
    matching the reference guide's "SAY THIS:" answer card body."""
    table, cell = _full_width_box_table(document)
    _shade_cell(cell, ANSWER_BOX_HEX)
    _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    spoken = prose_engine.spoken_register(text).text
    ledger = getattr(document, "_codex_spoken_sentence_ledger", set())
    kept_sentences: list[str] = []
    for sentence in re.split(r"(?<=[.!?])\s+", spoken):
        key = build_resume.normalize_compare(sentence)
        if len(key.split()) > 6 and key in ledger:
            continue
        if len(key.split()) > 6:
            ledger.add(key)
        kept_sentences.append(sentence)
    setattr(document, "_codex_spoken_sentence_ledger", ledger)
    spoken = " ".join(kept_sentences) or spoken
    lines = []
    if label:
        lines.append((label, {"bold": True, "color": RGBColor(0x1F, 0x3B, 0x5C), "size": SMALL_SIZE}))
    lines.append((spoken, {"size": BODY_SIZE}))
    _set_cell_text(cell, lines)
    if label and len(cell.paragraphs) > 1:
        cell.paragraphs[0].paragraph_format.keep_with_next = True
    _spacer(document, points=4)


def add_tip_box(document: Document, text: str, *, label: str = "TIP:") -> None:
    """Pale yellow callout box for coaching tips, matching the reference guide's
    tip footer under each Q&A card."""
    table, cell = _full_width_box_table(document)
    _shade_cell(cell, TIP_BOX_HEX)
    _set_cell_margins(cell, top=60, bottom=60, left=160, right=160)
    _set_cell_text(cell, [(f"{label} {text}" if label else text, {"italic": True, "size": SMALL_SIZE})])
    _spacer(document, points=6)


def add_qa_card(document: Document, prompt: str, answer: str, *, tip: str = "") -> None:
    """Full prompt/answer card: navy question header, light blue answer box,
    optional pale yellow tip footer. Used for the application-question and
    recent-interview-question prep lists, the sections most directly comparable
    to the reference guide's Q&A cards."""
    add_subsection(document, prompt)
    add_answer_box(document, answer)
    if tip.strip():
        add_tip_box(document, tip.strip())


def add_bullet(document: Document, text: str, *, small: bool = False) -> None:
    paragraph = document.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Pt(20)
    paragraph.paragraph_format.first_line_indent = Pt(-10)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    bullet = paragraph.add_run("- ")
    bullet.font.name = FONT
    bullet.font.size = Pt(SMALL_SIZE if small else BODY_SIZE)
    body = paragraph.add_run(text)
    body.font.name = FONT
    body.font.size = Pt(SMALL_SIZE if small else BODY_SIZE)


def add_banded_table(document: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    """Full-width table with a navy header row and alternating light-gray row
    banding, matching the reference guide's table style (e.g. the call-pacing
    and answer-length tables)."""
    if not rows:
        return
    column_count = len(headers)
    table = document.add_table(rows=1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = _content_width(document)
    column_width = int(width / column_count)
    for column in table.columns:
        column.width = column_width
    _thin_table_borders(table)
    _set_table_width(table, width)
    _repeat_header_row(table.rows[0])
    _prevent_row_split(table.rows[0])
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        cell.width = column_width
        _shade_cell(cell, TABLE_HEADER_NAVY_HEX)
        _set_cell_margins(cell, top=70, bottom=70, left=120, right=120)
        _set_cell_text(cell, [(header.upper(), {"bold": True, "color": RGBColor(0xFF, 0xFF, 0xFF), "size": SMALL_SIZE})])
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        _prevent_row_split(row)
        row_cells = row.cells
        fill = TABLE_ROW_ALT_HEX if row_index % 2 == 1 else WHITE_HEX
        for cell, value in zip(row_cells, row_values):
            cell.width = column_width
            _shade_cell(cell, fill)
            _set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
            _set_cell_text(cell, [(value, {"size": SMALL_SIZE})])
    _spacer(document, points=8)


def add_story_answer(document: Document, answer: StoryAnswerParts, *, prefix: str = "") -> None:
    spoken = prose_engine.spoken_register(answer.spoken.strip()).text
    if spoken:
        add_body(document, f"{prefix}{spoken}" if prefix else spoken)
    if answer.coaching_note.strip():
        add_body(document, answer.coaching_note.strip(), small=True)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def validate_text(
    text: str,
    *,
    company_name: str = "",
    role_title: str = "",
    jobs_dir: Path = PROJECT_ROOT / "jobs",
) -> None:
    validation_text = cheat.strip_external_question_prompts_for_validation(
        text,
        company_name,
        role_title,
        jobs_dir=jobs_dir,
    )
    if "--" in validation_text:
        fail("detailed interview guide contains double dashes")
    assert_no_template_leakage(validation_text)
    placeholder_scan_text = remove_approved_bracketed_metadata(validation_text)
    for pattern in PLACEHOLDER_PATTERNS:
        if re.search(pattern, placeholder_scan_text, re.I):
            fail(f"placeholder text detected in detailed interview guide: {pattern}")
    if build_resume.contains_ai_writing_word(validation_text):
        fail("detailed interview guide contains AI-writing disclosure language")
    forbidden_fragments = (
        "The reason I think my background fits",
        "That matters here because this role requires",
        "process waste, operational variation, data-backed root cause analysis, and measurable efficiency improvement",
    )
    for fragment in forbidden_fragments:
        if fragment.lower() in validation_text.lower():
            fail(f"detailed interview guide contains stale template language: {fragment}")


def notes_context(company_research: str, interview_notes: str) -> list[str]:
    def cleaned_note_lines(text: str) -> list[str]:
        cleaned: list[str] = []
        for line in safe_lines(text):
            stripped = re.sub(r"^[*-]\s*", "", line).strip()
            if stripped:
                cleaned.append(stripped)
        return cleaned

    lines: list[str] = []
    if company_research:
        lines.append("Company research supplied by Christian:")
        lines.extend(cleaned_note_lines(company_research)[:14])
    if interview_notes:
        lines.append("Interview or recruiter notes supplied by Christian:")
        lines.extend(cleaned_note_lines(interview_notes)[:14])
    if not lines:
        lines.append("No company research file found. For a richer guide later, add notes to jobs/company_research.txt or jobs/interview_notes.txt before running this script.")
    return lines


def add_application_question_prep_section(document: Document, job_description: str, resume_text: str = "") -> bool:
    responses = question_prep.active_application_question_responses(job_description)
    probes = question_prep.element_probe_responses(job_description, resume_text)
    if not responses and not probes:
        return False
    add_section(document, "Application / Supplemental Questions To Be Ready For")
    for response in responses:
        add_qa_card(document, response.prompt, response.answer)
    if probes:
        add_subsection(document, "Likely Requirement-Probe Questions")
        for response in probes:
            add_qa_card(document, response.prompt, response.answer)
    return True


def add_recent_interview_question_prep_section(
    document: Document,
    job_description: str,
    company_name: str,
    role_title: str,
    *,
    jobs_dir: Path = PROJECT_ROOT / "jobs",
    profile: build_resume.JobProblemProfile | None = None,
    stories: list[cheat.StoryCard] | None = None,
    resume_text: str = "",
    interview_notes: str = "",
) -> bool:
    items = question_prep.recent_interviewer_question_prep_items(
        job_description,
        company_name,
        role_title,
        jobs_dir=jobs_dir,
    )
    if not items:
        return False
    if not resume_text:
        _, _, resume_text = question_prep.selected_resume_snapshot(job_description)
    if profile is None:
        profile = cheat.adjusted_profile_for_role(
            build_resume.job_problem_profile(job_description, resume_text), role_title, job_description
        )
    if stories is None:
        stories = cheat.supported_story_bank(resume_text)
    add_section(document, "Recent Interview Questions To Be Ready For")
    used_titles: set[str] = set()
    seen_factual_answers: set[str] = set()
    for item in items:
        if item.category in question_prep.NON_STORY_INTERVIEWER_CATEGORIES:
            # Factual-only prompts (true reason a role ended, a specific
            # interviewer's own product) - a story would either be unrelated
            # or would force an invented fact. Use the honest fill-in script
            # as the spoken answer and keep the original coaching note as
            # the supporting tip instead of rendering it as the answer.
            spoken = question_prep.interviewer_question_factual_script(item.prompt, job_description, resume_text)
            normalized_spoken = build_resume.normalize_compare(spoken)
            tip = "" if normalized_spoken in seen_factual_answers else item.answer_angle
            seen_factual_answers.add(normalized_spoken)
            add_qa_card(document, item.prompt, spoken, tip=tip)
            continue
        story = cheat.likely_question_story(cheat.InterviewQuestion(item.prompt, item.answer_angle), stories, used_titles)
        used_titles.add(story.title)
        answer = story_sample_answer(story, profile, company_name, role_title, job_description, interview_notes, resume_text)
        add_qa_card(document, item.prompt, answer.spoken, tip=answer.coaching_note)
    return True


def text_has(text: str, *terms: str) -> bool:
    lowered = text.lower()
    return any(term.lower() in lowered for term in terms)



def concise_series(items: tuple[str, ...], limit: int = 3) -> str:
    selected = tuple(item for item in items if item)[:limit]
    return build_resume.comma_series(selected) if selected else "the role's priorities"


def build_prep_insights(
    company_name: str,
    role_title: str,
    job_description: str,
    company_research: str,
    interview_notes: str,
    profile: build_resume.JobProblemProfile,
) -> PrepInsights:
    combined = f"{job_description}\n{company_research}\n{interview_notes}"
    lowered = combined.lower()

    situation: list[str] = []
    scorecard: list[str] = []
    pushbacks: list[tuple[str, str]] = []
    selling: list[str] = []
    questions: list[tuple[str, str]] = []
    smart_questions: list[str] = []
    interviewer: list[str] = []
    brevity: list[str] = []

    if is_state_farm_active(company_name, role_title, job_description, company_research, interview_notes):
        return state_farm_prep_insights()
    if consulting_bigfour.is_bigfour_consulting_active(company_name, role_title, job_description, company_research, interview_notes):
        return consulting_bigfour.bigfour_prep_insights()

    if text_has(combined, "acquired", "acquisition", "post-merger", "integrated into"):
        situation.append(
            "This role may involve integration work after a business change: the interview will likely test whether Christian can stabilize daily operations while helping teams align process, data, and ownership."
        )
        scorecard.append("Can protect day-to-day operations while supporting process alignment after an acquisition.")
        selling.append("Lead with East West: five-site system ownership, Epicor migration support, cutover coordination, and cross-functional adoption under operating pressure.")
        smart_questions.append("Where are the biggest process, data, or ownership differences still showing up between the current environment and the future operating model?")

    if text_has(combined, "legacy", "25+ year", "unicode", "pick", "cobol", "rpg", "foxpro"):
        situation.append(
            "They likely need someone comfortable learning an old, lightly documented system rather than someone expecting a modern, clean SaaS admin environment."
        )
        scorecard.append("Can learn a legacy ERP quickly, document it as he learns, and avoid creating risk while knowledge transfers from the SME.")
        selling.append("Frame Aptean Intuitive as the proof: you owned a mission-critical manufacturing ERP, diagnosed issues independently, and kept operations moving while supporting migration to Epicor Kinetic.")
        pushbacks.append((
            "You have not used this exact legacy stack.",
            "That is true; I would not pretend otherwise. What I do bring is the pattern this role needs: I have inherited complex ERP environments, learned the workflow by tracing real transactions, documented the system as I went, and used SQL/ETL-style validation to reduce risk. My first goal would be to learn from the SME, map the highest-risk flows, and become useful without disrupting production."
        ))
        questions.append((
            "How would you learn a legacy system if documentation is limited?",
            "I would start with the highest-volume and highest-risk workflows, then shadow the SME while tracing real examples end to end: order, inventory, service flow, finance touchpoint, CRM feed, and data movement. I would document the workflow, known exceptions, owner, data source, and validation check as I learn so the team gains usable documentation instead of just transferring knowledge into my head."
        ))
        smart_questions.append("Which workflows would you want documented first because they carry the most operational or reporting risk?")

    if text_has(combined, "400 users", "800 users"):
        situation.append(
            "There is meaningful user scale and uncertainty, which means the role requires patience, documentation, and change discipline."
        )
        scorecard.append("Can support users at scale while preparing for a possible system migration or broader ERP transformation.")

    if text_has(combined, "etl", "extracting", "validating", "azure", "data warehouse", "salesforce"):
        situation.append(
            "Data flow matters as much as screen-level support: weak source data can become bad reporting, bad customer visibility, or bad integration decisions."
        )
        scorecard.append("Can validate data across source systems, ETL processes, CRM, data warehouse, and reporting workflows.")
        selling.append("Use the Aptean Intuitive to Epicor Kinetic migration bullet: extracting, validating, transforming, and loading ERP/database records through ETL tools and SQL validation.")
        questions.append((
            "What would you ask about Azure data storage or ERP hosting?",
            "I would ask how source data currently lands in CRM, warehouse, and reporting environments, what is batch versus near-real-time, which tables or extracts are most trusted, and where reconciliation breaks down. Then I would identify a few validation checks between source systems, ETL output, and downstream reporting so the team can separate system issues from data-quality issues."
        ))
        smart_questions.append("How does source data currently move into CRM, warehouse, or reporting tools, and where do reconciliations or trust issues usually appear?")

    if text_has(combined, "documentation", "training materials", "writing", "documenting", "demos"):
        situation.append(
            "Documentation is not a side task here; it is a risk-control mechanism because they are dependent on legacy knowledge and cross-training."
        )
        scorecard.append("Can create documentation and training materials that make legacy workflows easier for users and support teams to repeat.")
        selling.append(f"Use {build_resume.COMPANY_APTEAN} and {build_resume.COMPANY_EAST_WEST} examples: role-based training, adoption guides, executive workshops, QBRs, product demos, requirements docs, and workflow documentation.")
        questions.append((
            "Tell me about your documentation and training experience.",
            f"At {build_resume.COMPANY_EAST_WEST}, I created role-based training and workflow materials for operations, finance, and engineering users as we improved ERP adoption and supported the Epicor migration. At {build_resume.COMPANY_APTEAN}, I created training documentation, adoption guides, product updates, demos, SOWs, and requirements documents for manufacturing clients. The common thread is that I write documentation so someone can use it under pressure, not just so a document exists."
        ))

    if text_has(combined, "mining", "heavy equipment", "remanufacturing", "aftermarket", "equipment uptime", "total cost of ownership", "parts warehouses"):
        situation.append(
            "In this operating environment, system and data quality connect directly to customer uptime, inventory availability, service responsiveness, and total cost of ownership."
        )
        scorecard.append("Can connect ERP administration to operational reliability, customer responsiveness, and service quality, not just tickets.")
        questions.append((
            "What do you know about this operating environment?",
            "What stands out is that heavy operations depend on uptime, parts visibility, service response, and reliable data. For this role, that matters because system quality directly affects inventory visibility, pricing, reporting, and customer responsiveness."
        ))
        smart_questions.append("In a remanufacturing and parts environment, which operating metrics matter most for the ERP team: inventory accuracy, rebuild status, fulfillment speed, pricing, reporting, or customer response time?")

    if text_has(combined, "brevity", "technical program management", "post-merger integration", "electronics and communications engineering"):
        interviewer.append(
            "The interviewer likely values concise, technically grounded answers that connect systems work to execution. Lead with the answer, give one proof point, then pause."
        )
        interviewer.append(
            "His background suggests he may listen for structure: what you would do first, how you would reduce risk, how you communicate status, and how you prevent surprises."
        )
        brevity.extend((
            "Use a 20-second answer first: answer, proof point, business relevance.",
            "After each answer, pause and offer: 'I can go deeper on the technical side or the stakeholder side if helpful.'",
            "Avoid narrating your whole resume. Say the operating problem, the action, and the control/checkpoint."
        ))
        pushbacks.append((
            "Your answers may be too detailed for the interviewer's preferred style.",
            "Use this reset: 'The short version is: I would stabilize first, document as I learn, and use data validation to reduce migration risk. The relevant proof is my Aptean Intuitive ownership and Epicor migration support. I can expand on either side if useful.'"
        ))

    if text_has(combined, "contract", "extension", "6+ months"):
        pushbacks.append((
            "This is a contract role. Are you comfortable with that?",
            "Yes. I understand the immediate need is practical impact: learn the environment, support the SME, document critical workflows, and reduce risk while the ERP future state is evaluated. I am comfortable being measured on useful progress quickly."
        ))

    if not situation:
        lane_situations = {
            "presales_solution": (
                "The real evaluation is likely not whether Christian knows the product. It is whether he can earn buyer trust fast, "
                "ask the right discovery questions, and keep the solution narrative credible through implementation realism."
            ),
            "customer_success": (
                "The hidden question in this interview is likely: will this person catch account risk before it becomes churn risk? "
                "Lead with the $1M+ stabilization story and the QBR methodology."
            ),
            "change_enablement": (
                "The core concern is whether Christian can move people who did not ask to change. Lead with resistance, adoption signals, "
                "training, and what happened after launch."
            ),
            "analytics_operations": (
                "The evaluation is likely about judgment, not just tools. Can Christian ask the right question before building the report? "
                "Lead with the decision behind the dashboard."
            ),
        }
        situation.append(lane_situations.get(
            profile.primary_lane,
            f"The role appears to center on {profile.core_problem}. Prepare to show how your systems, customer, and implementation experience maps to that specific business problem."
        ))
    if not scorecard:
        lane_scorecards = {
            "presales_solution": "Can run credible discovery, build a solution narrative buyers trust, and keep the recommendation grounded in implementation reality.",
            "customer_success": "Can catch account risk early, structure executive conversations around business value, and protect renewal and expansion paths.",
            "change_enablement": "Can move stakeholders from resistance to adoption through training, communications, and visible measurement.",
            "analytics_operations": "Can turn messy data into decisions leaders act on, not just reports they acknowledge.",
        }
        scorecard.append(lane_scorecards.get(
            profile.primary_lane,
            "Can understand the business problem fast, structure ambiguous work, and produce measurable progress stakeholders can see."
        ))
    if not selling:
        lane_selling = {
            "presales_solution": "Lead with 80+ client engagements, 60+ executive workshops, discovery/SOW work, and implementation realism.",
            "customer_success": "Lead with $1M+ account risk stabilized, 80+ client portfolio, QBRs, adoption, and customer recovery.",
            "change_enablement": "Lead with role-based training, five-site adoption, executive workshops, and stakeholder alignment under change pressure.",
            "analytics_operations": "Lead with 200+ SQL-based reporting tools, dashboard adoption, and the decision each tool enabled.",
        }
        selling.append(lane_selling.get(profile.primary_lane, "Use the strongest proof points: enterprise system ownership, 80+ client portfolio, 200+ dashboards, 60+ workshops/QBRs, and $1M+ account recovery."))
    if not smart_questions:
        smart_questions.extend(cheat.questions_to_ask(company_name, profile, job_description)[:4])
    if not brevity:
        brevity.extend((
            "Lead with the direct answer before the story.",
            "Use one story at a time, then bridge back to the role.",
            "Stop after the role bridge and let the interviewer pull for detail."
        ))

    # Preserve order while removing duplicates.
    def unique(items: list[str]) -> tuple[str, ...]:
        seen: set[str] = set()
        out: list[str] = []
        for item in items:
            key = re.sub(r"\s+", " ", item.lower()).strip()
            if key and key not in seen:
                seen.add(key)
                out.append(item)
        return tuple(out)

    def unique_pairs(items: list[tuple[str, str]]) -> tuple[tuple[str, str], ...]:
        seen: set[str] = set()
        out: list[tuple[str, str]] = []
        for question, answer in items:
            key = question.lower().strip()
            if key not in seen:
                seen.add(key)
                out.append((question, answer))
        return tuple(out)

    return PrepInsights(
        situation_read=unique(situation),
        interviewer_read=unique(interviewer),
        likely_scorecard=unique(scorecard),
        pushbacks=unique_pairs(pushbacks),
        selling_angles=unique(selling),
        anticipated_questions=unique_pairs(questions),
        smart_questions=unique(smart_questions),
        brevity_rules=unique(brevity),
    )


def verified_company_research_points(company_name: str, job_description: str, company_research: str, interview_notes: str) -> tuple[tuple[str, str, str], ...]:
    combined = f"{company_name}\n{job_description}\n{company_research}\n{interview_notes}"
    if text_has(combined, "state farm"):
        return (
            (
                "State Farm scale and mission",
                "Use State Farm's scale as the business reason process engineering matters: claims process quality affects customers, service reliability, cost, speed, and trust at a very large insurer. Bridge this to Christian's process-improvement proof rather than reciting trivia.",
                "State Farm About Us: https://www.statefarm.com/about-us",
            ),
            (
                "Good Neighbor positioning",
                "State Farm's public identity centers on being a good neighbor. Interview bridge: process improvement is not only efficiency work; in claims, cleaner processes help customers experience faster, clearer, more reliable help during stressful moments.",
                "State Farm About Us: https://www.statefarm.com/about-us",
            ),
            (
                "Community and workforce investment",
                "If asked why State Farm, connect the supplied notes about workplace awards, tuition reimbursement, and Pathways-style community investment to Christian's preference for companies that build durable systems and invest in people.",
                "State Farm Careers / About pages: https://www.statefarm.com/careers and https://www.statefarm.com/about-us",
            ),
        )
    return ()


def add_verified_research_section(document: Document, research_points: tuple[tuple[str, str, str], ...]) -> None:
    if not research_points:
        return
    add_section(document, "Verified Company Research To Use")
    add_body(document, "Use these as conversation context. Do not recite every fact; pick one fact, explain why it matters operationally, then ask a question that shows business judgment.")
    for topic, point, source in research_points:
        add_subsection(document, topic)
        add_body(document, point)
        add_bullet(document, f"Source: {source}", small=True)


def add_insight_brief(document: Document, insights: PrepInsights) -> None:
    add_section(document, "Interpretive Interview Brief")
    add_subsection(document, "What The Notes Really Imply")
    for line in insights.situation_read[:8]:
        add_bullet(document, line)
    if insights.interviewer_read:
        add_subsection(document, "Interviewer Read")
        for line in insights.interviewer_read[:5]:
            add_bullet(document, line)
    add_subsection(document, "Opening Five Minutes")
    opening_lines = [
        "Use the first five minutes to diagnose pace, skepticism, warmth, and depth before committing to a long answer.",
        "If the interviewer sounds structured, lead with a clean point and one proof line. If they sound conversational, stay concise but let the tone breathe.",
        "Notice what creates the first follow-up question. That is usually the business pressure point to keep answering toward.",
    ]
    for line in opening_lines:
        add_bullet(document, line)
    add_subsection(document, "Likely Evaluation Scorecard")
    for line in insights.likely_scorecard[:8]:
        add_bullet(document, line)
    add_subsection(document, "How To Sell Yourself Briefly")
    for line in insights.selling_angles[:8]:
        add_bullet(document, line)
    add_subsection(document, "Brevity Rules For This Interview")
    for line in insights.brevity_rules[:5]:
        add_bullet(document, line)


def add_pushback_section(document: Document, insights: PrepInsights) -> None:
    add_section(document, "Likely Pushbacks And Short Answers")
    for concern, answer in insights.pushbacks[:10]:
        add_subsection(document, concern)
        add_body(document, answer)


def add_anticipated_question_section(document: Document, insights: PrepInsights) -> None:
    add_section(document, "Anticipated Questions From Notes")
    for question, answer in insights.anticipated_questions[:12]:
        add_subsection(document, question)
        add_body(document, answer)


def add_business_context_question_section(
    document: Document,
    job_description: str,
    supplied_context: str,
    stories: list[cheat.StoryCard],
    *,
    profile: build_resume.JobProblemProfile | None = None,
    company_name: str = "",
    role_title: str = "",
    resume_text: str = "",
    interview_notes: str = "",
) -> None:
    questions = business_context.business_interview_questions(job_description, supplied_context, limit=8)
    if not questions:
        return
    if not resume_text:
        _, _, resume_text = question_prep.selected_resume_snapshot(job_description)
    if profile is None:
        profile = cheat.adjusted_profile_for_role(
            build_resume.job_problem_profile(job_description, resume_text), role_title, job_description
        )
    add_section(document, "Business-Context Interview Questions")
    used_titles: set[str] = set()
    for item in questions:
        story = cheat.likely_question_story(cheat.InterviewQuestion(item.question, item.answer_angle), stories, used_titles)
        used_titles.add(story.title)
        answer = story_sample_answer(story, profile, company_name, role_title, job_description, interview_notes, resume_text)
        tip_parts = [
            f"Hidden business concern: {item.hidden_concern}",
            f"Ask back: {item.ask_back}",
        ]
        if answer.coaching_note.strip():
            tip_parts.append(answer.coaching_note.strip())
        add_qa_card(document, item.question, answer.spoken, tip=cheat.join_answer_sentences(*tip_parts))


def company_story_positioning(
    company_name: str,
    role_title: str,
    job_description: str,
    profile: build_resume.JobProblemProfile,
    company_research: str,
) -> str:
    if is_state_farm_active(company_name, role_title, job_description, company_research):
        return (
            f"For {company_name}'s {role_title} role, lead with process engineering rather than resume chronology. "
            "The business problem is claims operations reliability: reducing waste, improving service and quality, using data to find root cause, and creating communication materials that help leaders make decisions. "
            "Christian's strongest story is practical operational improvement: he mapped broken workflows, reduced manual inventory work by 78%, cut discrepancies by 22%, built 200+ reporting tools, and repeatedly brought cross-functional users through adoption. "
            "The message should be: I can spot the process issue, prove it with metrics, design the fix, and make the new process usable for the people doing the work."
        )
    lens = build_resume.primary_story_lens(job_description)
    values = build_resume.visible_values_phrase(job_description)
    specialty = build_resume.role_specialty_phrase(job_description, "the role's core work")
    lens_text = (
        f"The role appears to center on {lens['identity']}: {lens['business_problem']}"
        if lens
        else f"The role appears to center on {profile.core_problem}"
    )
    priority_phrase = next(iter(build_resume.jd_priority_phrases(job_description)), "")
    jd_text = f" The posting specifically emphasizes '{priority_phrase}', so echo that naturally when explaining the fit." if priority_phrase else ""
    value_text = f" The posting's visible values point to {values}." if values else ""
    research_text = " Use the supplied company research as conversation context and rehearsal material, not as wording to recite." if company_research else ""
    return (
        f"For {company_name}'s {role_title} role, lead with the business problem before the resume. "
        f"{lens_text} in {specialty}.{jd_text}{value_text} {research_text} "
        "Christian's best positioning is practical: he enters ambiguous system, customer, or reporting problems, "
        "creates structure, aligns the people involved, and turns the work into adoption, visibility, risk reduction, or measurable operating improvement."
    )


def detailed_pitch(company_name: str, role_title: str, profile: build_resume.JobProblemProfile) -> str:
    if profile.primary_lane == "process_improvement":
        return (
            f"The main pattern in my career is that I have usually been the person who finds the process gap other people have learned to live with. "
            f"At {build_resume.COMPANY_HOME_DEPOT}, that showed up as reporting frameworks when leaders were working from manual data and fragmented workflow signals. "
            f"At {build_resume.COMPANY_APTEAN}, it showed up as learning operational environments through the problems inside them, not just documentation. "
            f"At {build_resume.COMPANY_EAST_WEST}, it showed up as a recurring discrepancy nobody had root-caused until I mapped the workflow and rebuilt it with validated checkpoints. "
            f"For {company_name}'s {role_title} role, that is the same operating pattern applied to claims: define the process problem, validate the data, test the fix, document the new standard, and help the people doing the work adopt it. "
            "The proof points I would anchor on are 78% less manual work, 22% fewer discrepancies, 200+ KPI/reporting tools, and 60+ executive workshops and QBRs."
        )
    return (
        f"I help teams turn complex system, customer, and adoption problems into practical outcomes. For {company_name}'s {role_title} role, "
        "the three things I would bring are structured discovery, business-facing technical delivery, and calm stakeholder alignment. "
        "My experience sits across enterprise systems, SaaS, customer success, reporting, and implementation work: I have supported an 80+ international client portfolio, "
        "owned enterprise operations across five sites and more than 150 users, built 200+ dashboards and KPI tools, facilitated 60+ executive workshops and QBRs, "
        "and stabilized at-risk accounts tied to more than $1M in annual revenue. The pattern is consistent: clarify the problem, make the tradeoffs visible, "
        "validate the path with the people affected by it, and leave the team with a process they can keep using."
    )


def human_pivot_paragraph(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    notes_text: str = "",
) -> str:
    motivation = cheat.human_motivation_sentence(profile, notes_text)
    lane_consequence = {
        "analytics_operations": (
            "That pulled me toward building the reporting layer carefully, because a dashboard nobody trusts does not change any decisions."
        ),
        "implementation_delivery": (
            "That pushed me to stay involved past go-live, because the hardest problems usually show up when real users hit the live workflow."
        ),
        "customer_success": (
            "That meant I spent a lot of time on the adoption side, because value conversations land much better when the customer already trusts the experience."
        ),
        "presales_solution": (
            "That made me a better listener in discovery, because I knew what a weak requirements conversation produced on the delivery side."
        ),
        "change_enablement": (
            "That moved me toward investing more time in the middle of a rollout than the kickoff, because that is where resistance either gets addressed or hardens."
        ),
        "process_improvement": (
            "That made me more attentive to the people living inside the process, because a fix that looks clean on paper still fails if the day-to-day work gets harder."
        ),
        "corporate_strategy": (
            "That shaped how I think about strategy work, because a smart recommendation that never becomes usable operating behavior does not create much value."
        ),
    }
    bridge = cheat.interview_role_bridge_sentence(
        profile,
        company_name,
        role_title,
        job_description,
        resume_text,
    )
    return cheat.join_answer_sentences(
        motivation,
        lane_consequence.get(
            profile.primary_lane,
            "That shaped how I approach the work, always keeping the human layer visible alongside the technical one."
        ),
        bridge,
    )


def build_extended_tmay_sections(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    notes_text: str = "",
) -> list[tuple[str, str]]:
    stories = cheat.hero_stories(profile, job_description, resume_text)
    top_story = stories[0] if stories else None
    proof = ""
    if top_story:
        proof_parts: list[str] = []
        if top_story.title and top_story.hook:
            opener = cheat.story_opener_by_type(set(top_story.story_types), top_story.title)
            proof_parts.append(
                f"{opener[:1].upper() + opener[1:]} "
                f"{cheat.lower_clause(top_story.hook)}"
            )
        if top_story.level3_trait:
            proof_parts.append(cheat.spoken_level3_trait_sentence(top_story.level3_trait))
        if top_story.evidence:
            proof_parts.append(cheat.story_evidence_sentence(top_story.evidence))
        if top_story.result:
            proof_parts.append(cheat.story_result_sentence(top_story.result))
        proof = cheat.join_answer_sentences(*proof_parts) if proof_parts else ""
    role_bridge = cheat.join_answer_sentences(
        cheat.interview_role_bridge_sentence(
            profile,
            company_name,
            role_title,
            job_description,
            resume_text,
        ),
    )
    return [
        ("Opening Hook", cheat.natural_voice_opening(profile)),
        ("Career Arc", cheat.pitch_career_arc_sentence(profile, job_description)),
        ("Why I Care", human_pivot_paragraph(profile, company_name, role_title, job_description, resume_text, notes_text)),
        ("Proof Beat", proof or detailed_pitch(company_name, role_title, profile)),
        ("Why This Role", role_bridge),
        ("Conversation Pivot", cheat.natural_voice_closing(profile, role_title)),
    ]


def story_coaching_note(human_line: str, role_bridge: str = "") -> str:
    parts: list[str] = []
    if human_line.strip():
        parts.append(f"Human layer to weave in naturally: {human_line.strip()}")
    if role_bridge.strip():
        parts.append(f"Role bridge: {role_bridge.strip().rstrip('.')}")
    if not parts:
        return ""
    return f"Coaching note: {cheat.join_answer_sentences(*parts)}"


def add_extended_tmay_section(
    document: Document,
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    notes_text: str = "",
) -> None:
    add_section(document, "Tell Me About Yourself: Extended Version")
    add_body(
        document,
        "Use this when the interviewer is engaged and the opening gives Christian room to go beyond the shorter elevator speech. Rehearse the structure enough that it stays natural under redirects and follow-up questions."
    )
    for label, paragraph in build_extended_tmay_sections(
        profile,
        company_name,
        role_title,
        job_description,
        resume_text,
        notes_text,
    ):
        add_subsection(document, label)
        add_body(document, paragraph)


def add_story_anchor_system_section(
    document: Document,
    stories: Sequence[cheat.StoryCard],
    profile: build_resume.JobProblemProfile,
) -> None:
    add_section(document, "Story Anchor System")
    add_bullet(document, "Use the 3-point spine: hook, number or scale anchor, result.")
    add_bullet(document, "Rehearse the first line and the final business-result line until both land cleanly without looking down.")
    add_bullet(document, "If the middle gets loose, go back to the anchor phrase and the result instead of restarting the whole answer.")
    table_rows = []
    for index, card in enumerate(stories[:5], start=1):
        anchor_text, is_numeric = cheat.story_scale_anchor(card)
        anchor_label = "Number or scale" if is_numeric else "Scale (qualitative)"
        table_rows.append((f"Story {index}: {card.title}", f"{anchor_label}: {anchor_text}", card.result))
    add_banded_table(document, ("Story", "Anchor", "Result"), table_rows)
    add_subsection(document, "Coaching Notes")
    add_bullet(document, "The Human Element is the part that should sound freshest in every interview. Keep the idea stable, but let the wording flex naturally.")
    add_bullet(document, "If the interviewer leans in on the Proof Beat, stay there and go deeper before rushing to the bridge.")
    add_bullet(document, "When the opening needs to stay shorter, collapse this into the 60-second version instead of forcing the full sequence.")


def why_company_answer(
    company_name: str,
    role_title: str,
    job_description: str,
    profile: build_resume.JobProblemProfile,
    company_research: str,
) -> str:
    if is_state_farm_active(company_name, role_title, job_description, company_research):
        return (
            f"I am interested in {company_name} because process work matters differently in insurance: a cleaner claims process can change how quickly and clearly a customer gets help. "
            "I would avoid trying to manufacture a unique reason; the authentic reason is specific enough. State Farm's scale, good-neighbor customer promise, and investment in people matter to me because this kind of process work has real customer stakes and a real learning curve. "
            "The fit on my side is that I have spent my career turning messy workflows, data gaps, and adoption problems into measurable operating improvements. "
            f"For the {role_title} role, I would bring current-state mapping, data-backed root-cause analysis, stakeholder alignment, documentation, training, and a habit of measuring whether the change actually improved service, quality, or efficiency."
        )
    lens = build_resume.primary_story_lens(job_description)
    values = build_resume.visible_values_phrase(job_description)
    specialty = build_resume.role_specialty_phrase(job_description, "this work")
    if lens:
        opening = f"What interests me about {company_name} is that this role appears to sit where {lens['business_problem']} matters."
    else:
        opening = f"What interests me about {company_name} is the chance to do practical {specialty} work where execution and adoption matter."
    values_sentence = f" I also noticed the emphasis on {values}, which matters to me because those values only count when the work is usable." if values else ""
    research_sentence = " The company research I reviewed gives me more context for why that work matters right now." if company_research else ""
    return (
        f"{opening}{values_sentence}{research_sentence} My background connects because I have repeatedly worked in the space between business needs, systems, data, and adoption. "
        f"For this {role_title} role, I would want to bring direct value: understand the real operating problem, structure the path forward, and make the outcome measurable."
    )


def role_challenge_answer(company_name: str, role_title: str, profile: build_resume.JobProblemProfile, job_description: str) -> str:
    if is_state_farm_active(company_name, role_title, job_description):
        return (
            "The biggest challenge I would expect is improving claims processes without accidentally optimizing only one metric. "
            "In claims operations, speed, cost, service, quality, compliance, and employee adoption all interact. My first step would be to clarify the process outcome State Farm cares most about, map the current state with the people doing the work, validate the data, and separate root cause from symptoms. "
            "Then I would recommend a prioritized improvement with a cost/benefit view, pilot it, define success metrics, and build the communication rhythm so stakeholders understand what changed and why."
        )
    specialty = build_resume.role_specialty_phrase(job_description, "the role's core work")
    return (
        f"The biggest challenge I would expect in this role is turning {profile.core_problem} in {specialty} into practical progress without losing stakeholder trust. "
        "My first step would be to clarify the outcome, validate where the facts are thin, and map the highest-risk workflow with the people closest to it. "
        "From there I would create a simple decision rhythm, keep ownership visible, and look for an early improvement that proves momentum."
    )


def six_offer_blocker_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    specialty = build_resume.role_specialty_phrase(job_description, profile.core_problem)
    return [
        f"Offer blocker 1 - vague fit story: if Christian cannot explain why the {role_title} role at {company_name} matches his pattern of work in two sentences, the interview can end as 'interesting but unclear fit.'",
        f"Offer blocker 2 - generic proof: if answers stay at the level of responsibilities instead of measurable outcomes in {specialty}, the team may assume the execution depth is thinner than the resume suggests.",
        "Offer blocker 3 - weak ownership language: if stories sound like 'we did everything' instead of making Christian's personal judgment and actions visible, confidence drops quickly.",
        "Offer blocker 4 - no business translation: if technical or process answers never connect back to customer impact, adoption, risk, or decision quality, the value can sound too narrow.",
        "Offer blocker 5 - generic why-company answer: if interest sounds copied from any other employer, it weakens trust even when the experience is strong.",
        "Offer blocker 6 - passive close: if the interview ends without explicit interest, a sharp question, and a clear fit recap, the team may remember capability but not conviction.",
    ]


def four_trust_questions_audit(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    specialty = build_resume.role_specialty_phrase(job_description, profile.core_problem)
    return [
        f"Trust question 1 - Would I trust Christian to understand the real problem quickly? He needs to show structured discovery and a story that mirrors {profile.core_problem} in {specialty}.",
        "Trust question 2 - Would I trust Christian's judgment when the facts are incomplete? Strong answers should show what he noticed early, how he validated the risk, and how he avoided surprises.",
        f"Trust question 3 - Would I trust Christian with customers, executives, or cross-functional stakeholders at {company_name}? Use examples that show calm translation, direct ownership, and clear follow-through under pressure.",
        f"Trust question 4 - Would I trust Christian to create measurable progress in the {role_title} role? End answers with what changed: adoption, delivery pace, risk control, decision quality, or customer trust.",
    ]


def executive_presence_signals(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    specialty = build_resume.role_specialty_phrase(job_description, profile.core_problem)
    return [
        f"Open with the answer, not the autobiography. In the {role_title} interview, executive presence means naming the point first, then the proof.",
        f"Translate complexity into a decision. The strongest answers turn {profile.core_problem} in {specialty} into a clear risk, tradeoff, or next step the team can act on.",
        "Sound declarative, not hedged. Cut phrases like 'I think,' 'I would probably,' or 'kind of' when Christian already has the proof.",
        "Keep one metric, scope marker, or concrete operating detail in major answers so confidence sounds earned rather than performed.",
        f"Close by tying the answer back to what {company_name} needs most: adoption, risk control, decision quality, customer trust, or delivery pace.",
    ]


def story_quality_audit(card: cheat.StoryCard, profile: build_resume.JobProblemProfile) -> list[str]:
    """Audit the working sequence: Hook -> Noticing -> Action -> Result -> Bridge -> Calibration."""
    base_lines = [
        "HOOK check: Does the first sentence create stakes or tension instead of starting with chronology?",
        "TONAL LEAD-BURIAL check: If the most human or honest line shows up late, move it closer to the hook so the story sounds lived-in rather than sanitized.",
        f"NOTICING check: What does the interviewer learn about Christian's judgment? Use a concrete observation, not a generic trait. Current noticing cue: {cheat.spoken_level3_trait_sentence(card.level3_trait)}",
        "ACTION check: Does the story use 'I' for Christian's specific ownership and show the steps he personally drove?",
        f"RESULT check: Is there a quantified or concrete outcome, and is it stated with confidence? Result available: {card.result}",
        f"BRIDGE check: Does the ending connect this exact story to the role's process, data, stakeholder, or customer problem rather than using a canned sentence about {profile.core_problem}?",
        "TIMING: Hook = 10 seconds, Noticing = 10-15 seconds, Action = 35-50 seconds, Result/Bridge = 20 seconds. Stop after the calibration question.",
    ]
    i_count = len(re.findall(r"\bI\b", card.evidence))
    we_count = len(re.findall(r"\bwe\b", card.evidence, re.I))
    if we_count > i_count:
        base_lines.insert(
            2,
            "OWNERSHIP check: This story uses 'we' more than 'I' in the evidence. Name the team context once, then use 'I' for every specific action and decision.",
        )
    story_types = set(card.story_types)
    type_specific_lines: list[str] = []
    if "Challenge and Failure" in story_types:
        type_specific_lines.append(
            "RECOVERY check: Does the story name a specific behavior that changed after the failure, not just a lesson learned? 'I now do X differently' is stronger than 'I learned the importance of X.'"
        )
    if "Persuasion" in story_types or "Opposing Views" in story_types:
        type_specific_lines.append(
            "EMPATHY check: Does the story show what the other person was protecting or afraid of, not just what Christian argued? The persuasion lands when the interviewer understands why the other person was hard to move."
        )
    if "Managing and Leading" in story_types:
        type_specific_lines.append(
            "INFLUENCE check: Does the story show influence without formal authority? If Christian had direct reports in this story, name the specific action he took that would have worked even without the title."
        )
    if "Rapid Learning" in story_types:
        type_specific_lines.append(
            "RAMP check: Does the story show the specific method used to get up to speed, not just the fact that it happened quickly? Name the first thing learned, the first question asked, or the first decision made."
        )
    if "Teamwork" in story_types:
        type_specific_lines.append(
            "DIFFERENCE check: Does the story make clear what was different about the other person's background, style, or incentive, and how that difference was bridged rather than overcome?"
        )
    if len(type_specific_lines) > 2:
        base_lines = base_lines[: max(0, 8 - len(type_specific_lines))]
    return (base_lines + type_specific_lines)[:8]


def story_selection_decision_table(
    stories: list[cheat.StoryCard],
    profile: build_resume.JobProblemProfile,
    job_description: str,
) -> list[str]:
    """
    Story selection decision table from the coaching session.
    Relevancy > Recency > Impact > Multi-skill.
    """
    lines = [
        "Use this table when unsure which story to tell. Relevancy always beats recency.",
        "Criteria: (1) Mirrors the JD's specific problem, (2) Shows the exact skill they need, (3) Has a metric or scope marker, (4) Demonstrates how Christian thinks, not just what he did.",
    ]
    for card in stories[:6]:
        relevance_score = cheat.signal_score(job_description, card.signals)
        has_metric = any(ch.isdigit() for ch in card.result) or "$" in card.result or "%" in card.result
        multi_skill = len(card.story_types) >= 3
        rating = "STRONG" if relevance_score >= 3 and has_metric else ("GOOD" if relevance_score >= 2 else "USE AS BACKUP")
        lines.append(
            f"{rating}: {card.title} | JD signal match: {relevance_score} | Has metric: {'yes' if has_metric else 'no'} | Multi-skill: {'yes' if multi_skill else 'no'} | Best for: {', '.join(card.story_types[:2])}"
        )
    lines.append(
        "Primary recommendation: use the STRONG-rated story for the first behavioral question. Save GOOD-rated stories for follow-ups. Use backups only if asked for a second example."
    )
    return lines


KEYWORD_DISPLAY_OVERRIDES = {
    "ai": "AI",
    "api": "API",
    "bi": "BI",
    "crm": "CRM",
    "etl": "ETL",
    "kpi": "KPI",
    "qbr": "QBR",
    "sql": "SQL",
    "uat": "UAT",
}


KEYWORD_QUESTION_DETAILS: dict[str, tuple[str, str, str]] = {
    "account growth": (
        "How do you identify account growth opportunities without forcing a sales conversation?",
        "Start with business value, adoption gaps, stakeholder goals, and supported expansion discovery. Avoid claiming direct quota ownership.",
        "What signal tells you the customer is ready for a larger conversation?",
    ),
    "adoption": (
        "How do you drive adoption after a customer goes live?",
        "Explain how you confirm the workflow is usable, watch early behavior, use training or QBRs, and adjust the plan when adoption is weak.",
        "What would you measure first if adoption started to stall?",
    ),
    "analytics": (
        "How do you help customers turn analytics into real decisions?",
        "Begin with the decision the customer needs to make, then explain the KPI, data source, validation step, and action threshold.",
        "How do you keep analytics from becoming reporting for its own sake?",
    ),
    "automation": (
        "Where have you used automation to make a workflow cleaner or more reliable?",
        "Focus on the manual step, the risk it created, the control or workflow improvement, and the result people could see.",
        "How did you make sure the automated step did not create a new quality issue?",
    ),
    "business intelligence": (
        "What does business intelligence mean in practical customer or operations work?",
        "Frame it as decision support: reliable data, useful segmentation, clear ownership, and reporting that changes an action.",
        "How do you decide whether a dashboard is actually useful?",
    ),
    "continuous improvement": (
        "Tell me about a time you improved a process after seeing the same issue repeat.",
        "Use the notice-action-result sequence: pattern, root cause, change, validation, and what improved after the change.",
        "How did you know the improvement would hold?",
    ),
    "crm": (
        "How have you used CRM information to manage customer relationships or risk?",
        "Talk about account context, health signals, open issues, executive conversations, and follow-up discipline.",
        "What CRM signal would make you escalate an account?",
    ),
    "customer relationship": (
        "How do you build a customer relationship after onboarding without becoming only a support contact?",
        "Show how you connect the relationship to business goals, value realization, risk visibility, and a useful operating rhythm.",
        "How do you recover trust when the customer is frustrated?",
    ),
    "customer success": (
        "How do you drive customer success when adoption risk is not obvious yet?",
        "Use account health, stakeholder signals, QBRs, value realization, and risk recovery. Keep the answer commercially aware without inventing quota ownership.",
        "Which early warning signs would you look for?",
    ),
    "customer experience": (
        "How do you improve customer experience when the issue starts inside an internal workflow?",
        "Connect the customer outcome to the process, data, handoff, or adoption issue that creates the experience problem.",
        "How would you prove the customer experience actually improved?",
    ),
    "dashboards": (
        "Tell me about a dashboard or report that changed how people worked.",
        "Lead with the decision or workflow problem, then explain the metric, validation, stakeholder use, and outcome.",
        "How did you keep the dashboard from becoming noise?",
    ),
    "data": (
        "How do you validate whether customer or operations data can be trusted?",
        "Explain source checks, SQL or ETL validation when supported, reconciliation, stakeholder confirmation, and what you do before making a recommendation.",
        "What do you do when two data sources disagree?",
    ),
    "etl": (
        "What is your experience with ETL or data movement issues?",
        "Connect extraction, transformation, validation, migration, reporting, cutover coordination, and customer or operations impact.",
        "How do you explain a technical data issue to a nontechnical stakeholder?",
    ),
    "implementation": (
        "How do you manage implementation risk while keeping customers confident?",
        "Cover scope, requirements, milestones, data validation, testing, training, go-live, post-go-live support, and a clear escalation rhythm.",
        "What risk would you try to catch before go-live?",
    ),
    "kpi": (
        "How do you decide which KPIs matter?",
        "Tie the KPI to a decision, behavior, customer outcome, or operating risk. Avoid vanity metrics.",
        "What is an example of a KPI that can mislead people?",
    ),
    "onboarding": (
        "What does a strong onboarding experience look like?",
        "Frame onboarding as a controlled launch: expectations, data readiness, stakeholder ownership, training, early wins, and a feedback rhythm.",
        "How do you know onboarding worked after the handoff?",
    ),
    "operations": (
        "Tell me about a time operations needs and customer needs had to be balanced.",
        "Use a story where speed, control, reporting, finance, or customer impact had to be weighed clearly.",
        "How do you make the tradeoff visible without creating conflict?",
    ),
    "process improvement": (
        "Walk me through how you approach process improvement from problem to result.",
        "Use current-state mapping, root cause, data validation, stakeholder input, action plan, pilot or implementation, and measurement.",
        "How do you separate root cause from symptoms?",
    ),
    "project management": (
        "How do you keep a complex customer or implementation project on track?",
        "Talk about scope, owners, milestones, dependencies, risk checks, communication rhythm, and what you escalate early.",
        "What do you do when the plan is technically on track but stakeholders are losing confidence?",
    ),
    "reporting": (
        "Tell me about a time reporting changed a customer or business decision.",
        "Start with the decision, describe the reporting gap, show how you validated the data, and explain what changed afterward.",
        "How did you know the report was trusted?",
    ),
    "root cause": (
        "How do you find root cause when people disagree about the problem?",
        "Describe how you separate symptoms from cause using workflow evidence, data checks, and stakeholder input.",
        "What do you do if the first root-cause theory is wrong?",
    ),
    "sql": (
        "How have you used SQL in customer, implementation, or operations work?",
        "Keep it practical: validation, extraction, reporting, troubleshooting, dashboards, and reducing uncertainty before a decision.",
        "How do you explain SQL findings to someone who only wants the business answer?",
    ),
    "strategy": (
        "How do you handle strategy conversations with customers or stakeholders?",
        "Anchor the conversation in business goals, adoption gaps, value realization, risk, and the next practical decision.",
        "How do you keep a strategy conversation from becoming too abstract?",
    ),
    "technical account manager": (
        "How would you operate as a Technical Account Manager when the account has both relationship risk and technical delivery risk?",
        "Balance customer trust, account health, issue ownership, product or data context, escalation rhythm, and measurable value.",
        "How would you decide what to escalate versus what to solve directly?",
    ),
    "training": (
        "How do you make training useful after a system or process change?",
        "Connect training to the actual workflow, user questions, feedback, adoption, and measurable behavior after launch.",
        "How do you adjust when users are still confused after training?",
    ),
    "workshops": (
        "How do you run workshops or QBRs so they lead to decisions?",
        "Show how you set the purpose, bring the right data, manage different audiences, and leave with owners and next steps.",
        "How do you handle a workshop where stakeholders disagree?",
    ),
}


KEYWORD_STORY_HINTS: tuple[tuple[tuple[str, ...], tuple[str, ...]], ...] = (
    (("reporting", "analytics", "dashboard", "dashboards", "business intelligence", "kpi", "sql"), ("dashboard", "decision visibility")),
    (("etl", "data", "implementation", "onboarding", "uat", "migration"), ("lifecycle delivery", "product learning", "validation")),
    (("customer success", "customer relationship", "adoption", "renewal", "expansion", "crm", "qbr"), ("account", "$1m", "workshop", "qbr")),
    (("operations", "process", "process improvement", "root cause", "continuous improvement"), ("inventory", "operations versus finance", "finance alignment")),
    (("training", "change", "stakeholder", "workshops"), ("workshop", "qbr", "product learning")),
)


KEYWORD_PRIORITY_BONUS = {
    "adoption": 8,
    "analytics": 8,
    "crm": 5,
    "data": 8,
    "etl": 10,
    "implementation": 7,
    "kpi": 10,
    "reporting": 8,
    "sql": 10,
    "strategy": 6,
}


def keyword_display(keyword: str) -> str:
    normalized = build_resume.normalize_compare(keyword)
    if normalized in KEYWORD_DISPLAY_OVERRIDES:
        return KEYWORD_DISPLAY_OVERRIDES[normalized]
    words = keyword.split()
    displayed: list[str] = []
    for word in words:
        clean = build_resume.normalize_compare(word)
        displayed.append(KEYWORD_DISPLAY_OVERRIDES.get(clean, word.capitalize()))
    return " ".join(displayed)


def keyword_question_detail(keyword: str, profile: build_resume.JobProblemProfile) -> tuple[str, str, str]:
    normalized = build_resume.normalize_compare(keyword)
    if normalized in KEYWORD_QUESTION_DETAILS:
        return KEYWORD_QUESTION_DETAILS[normalized]
    display = keyword_display(keyword)
    return (
        f"How would you apply {display} in this role?",
        f"Use one specific story. Name what you noticed, explain the action, state the result, and connect it to {profile.core_problem}.",
        f"What would you measure to know the {display} work succeeded?",
    )


def keyword_present(text: str, keyword: str) -> bool:
    return bool(re.search(build_resume.keyword_regex(keyword.lower()), text.lower(), re.I))


def _phrase_contains(haystack: str, needle: str) -> bool:
    """True if needle appears in haystack as a whole-word phrase (not just a raw substring)."""
    return bool(re.search(rf"\b{re.escape(needle)}\b", haystack))


def important_resume_keywords(job_description: str, resume_text: str, max_keywords: int = 14) -> list[str]:
    candidates = set(build_resume.audit_keywords(job_description))
    candidates |= set(build_resume.SUMMARY_PLACEMENT_TERMS)
    candidates |= build_resume.jd_color_priority_terms(job_description)
    candidates |= set(KEYWORD_QUESTION_DETAILS)

    scored: list[tuple[int, int, str]] = []
    early_resume = resume_text[:1500]
    for raw_keyword in candidates:
        keyword = build_resume.normalize_compare(raw_keyword)
        if not keyword or len(keyword) < 3:
            continue
        if build_resume.is_generic_soft_keyword(keyword):
            continue
        if not keyword_present(job_description, keyword) or not keyword_present(resume_text, keyword):
            continue
        if not (
            build_resume.is_keyword_color_candidate(keyword, job_description)
            or keyword in build_resume.SUMMARY_PLACEMENT_TERMS
            or keyword in KEYWORD_QUESTION_DETAILS
        ):
            continue
        score = 0
        score += build_resume.keyword_hits(job_description, {keyword}) * 4
        score += build_resume.keyword_hits(resume_text, {keyword}) * 2
        if keyword_present(early_resume, keyword):
            score += 4
        if keyword in KEYWORD_QUESTION_DETAILS:
            score += 5
        if keyword in build_resume.SUMMARY_PLACEMENT_TERMS:
            score += 2
        if keyword in build_resume.jd_color_priority_terms(job_description):
            score += 4
        if " " in keyword:
            score += 3
        score += KEYWORD_PRIORITY_BONUS.get(keyword, 0)
        scored.append((score, len(keyword), keyword))

    scored.sort(reverse=True)
    available = {keyword for _, _, keyword in scored}
    priority_order = (
        "sql", "kpi", "etl", "reporting", "adoption", "data", "analytics",
        "implementation", "strategy", "operations", "dashboards",
        "customer relationship", "customer success", "technical account manager",
    )
    selected: list[str] = []
    for keyword in priority_order:
        if keyword in available:
            selected.append(keyword)
        if len(selected) >= max_keywords:
            return selected
    for _, _, keyword in scored:
        if keyword in selected:
            continue
        # Skip this keyword if it's just a sub-phrase of something already selected
        # (e.g. drop "global external" when "global external manufacturing" is already in).
        if any(keyword != existing and _phrase_contains(existing, keyword) for existing in selected):
            continue
        # Drop any already-selected keyword that is a sub-phrase of this longer, more specific one.
        superseded = [existing for existing in selected if existing != keyword and _phrase_contains(keyword, existing)]
        if superseded:
            selected = [existing for existing in selected if existing not in superseded]
        selected.append(keyword)
        if len(selected) >= max_keywords:
            break
    return selected


def best_story_for_keyword(
    keyword: str,
    stories: list[cheat.StoryCard],
    profile: build_resume.JobProblemProfile,
) -> cheat.StoryCard | None:
    if not stories:
        return None
    normalized = build_resume.normalize_compare(keyword)
    for keyword_terms, title_terms in KEYWORD_STORY_HINTS:
        if any(term in normalized for term in keyword_terms):
            for story in stories:
                title = story.title.lower()
                if any(term in title for term in title_terms):
                    return story
    question, angle, _ = keyword_question_detail(keyword, profile)
    return max(stories, key=lambda story: cheat.signal_score(f"{keyword} {question} {angle}", story.signals))


def keyword_sample_answer(
    keyword: str,
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    story: cheat.StoryCard | None,
) -> str:
    display = keyword_display(keyword)
    if story:
        return cheat.interview_join(
            f"For {display}, I focus on the customer workflow, the delivery risk, and the outcome the team needs to protect",
            cheat.spoken_caar_answer(story, profile),
        )
    return (
        f"For {display}, I would answer directly and stay grounded in supported experience. "
        f"My approach would be to clarify how the term shows up in the customer's workflow, identify the risk or decision behind it, validate the path with the right stakeholders, and connect the work back to {profile.core_problem}."
    )


def add_keyword_question_bank(
    document: Document,
    profile: build_resume.JobProblemProfile,
    job_description: str,
    resume_text: str,
    hero_stories: list[cheat.StoryCard],
    company_name: str = "",
    role_title: str = "",
    max_keywords: int = 8,
) -> None:
    keyword_limit = 10 if build_resume.is_consulting_job_description(job_description) else max_keywords
    keywords = important_resume_keywords(job_description, resume_text, keyword_limit)
    if not keywords:
        return
    add_section(document, "KEYWORD ANSWER REFERENCE")
    add_body(
        document,
        "Use this section between interviews to practice keyword-specific answers. The primary story bank and behavioral scripts above are higher priority for live preparation.",
    )
    for keyword in keywords:
        question, angle, follow_up = keyword_question_detail(keyword, profile)
        story = best_story_for_keyword(keyword, hero_stories, profile)
        add_subsection(document, keyword_display(keyword))
        add_bullet(document, f"Likely question: {question}")
        add_bullet(document, f"Underlying intent: {angle}")
        if story:
            add_bullet(document, f"Best story: {story.title} - {story.result}")
        add_body(document, f"Sample answer: {keyword_sample_answer(keyword, profile, company_name or 'the company', role_title or 'the role', story)}")
        add_bullet(document, f"Follow-up to prepare for: {follow_up}")







def story_sample_answer(
    card: cheat.StoryCard,
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str = "",
    interview_notes: str = "",
    resume_text: str = "",
) -> StoryAnswerParts:
    human_line = cheat.story_human_connection_line(
        card,
        profile,
        company_name,
        role_title,
        job_description,
        resume_text,
    )
    if is_state_farm_active(company_name, role_title, job_description):
        spoken = cheat.join_answer_sentences(
            f"{cheat.story_opener_by_type(set(card.story_types), card.title).capitalize()} where {cheat.lower_clause(card.hook)}",
            f"The key early signal was {cheat.lower_clause(state_farm_playbook.state_farm_story_noticing(card))}",
            cheat.story_evidence_sentence(card.evidence),
            cheat.story_result_sentence(card.result),
            state_farm_playbook.state_farm_story_bridge(card),
            state_farm_playbook.state_farm_calibration_question(card),
        )
        return StoryAnswerParts(
            spoken=spoken,
            coaching_note=story_coaching_note(human_line, profile.core_problem),
        )
    elif consulting_bigfour.is_bigfour_consulting_active(company_name, role_title, job_description, interview_notes=interview_notes):
        spoken = cheat.join_answer_sentences(
            f"{cheat.story_opener_by_type(set(card.story_types), card.title).capitalize()} where {cheat.lower_clause(card.hook)}",
            cheat.spoken_level3_trait_sentence(card.level3_trait),
            cheat.story_evidence_sentence(card.evidence),
            cheat.story_result_sentence(card.result),
            consulting_bigfour.bigfour_story_bridge(card),
            consulting_bigfour.bigfour_calibration_question(card),
        )
        return StoryAnswerParts(
            spoken=spoken,
            coaching_note=story_coaching_note(human_line, profile.core_problem),
        )
    return StoryAnswerParts(
        spoken=cheat.spoken_story_answer(card, profile, company_name, role_title, job_description),
        coaching_note=story_coaching_note(human_line, profile.core_problem),
    )


def behavioral_sample_answers(
    profile: build_resume.JobProblemProfile,
    stories: list[cheat.StoryCard],
    company_name: str,
    role_title: str,
    job_description: str = "",
    interview_notes: str = "",
    resume_text: str = "",
) -> list[tuple[str, StoryAnswerParts]]:
    scripts = cheat.behavioral_answer_scripts(
        profile,
        stories,
        job_description,
        company_name,
        role_title,
        resume_text,
        interview_notes,
    )
    cheat.validate_behavioral_answer_bank(scripts)
    answers: list[tuple[str, StoryAnswerParts]] = []
    for item in scripts:
        related = next((story for story in stories if story.title in item.answer), None)
        if related:
            answer = story_sample_answer(related, profile, company_name, role_title, job_description, interview_notes, resume_text)
        else:
            answer = StoryAnswerParts(spoken=item.answer)
        answers.append((item.prompt, answer))
    return answers


def add_company_fit_answer_bank(
    document: Document,
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    stories: list[cheat.StoryCard],
    supplied_context: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> None:
    add_section(document, "Company Fit And Common Questions")
    add_body(
        document,
        "Use this section after the rehearsal core. It keeps the company fit language, the most likely direct questions, and the keyword-specific answers in one place without repeating the full phone-screen playbook.",
    )
    add_subsection(document, "Company Background From The Posting")
    for line in cheat.company_background_lines(job_description, profile, company_name, role_title):
        add_bullet(document, line)
    supplied_background = cheat.supplied_company_background_lines(company_name, supplied_context)
    if supplied_background:
        add_subsection(document, "Supplied Company Intelligence")
        for line in supplied_background[:12]:
            add_bullet(document, line)
    add_subsection(document, "Why You, Why Now, Why Company")
    for item in cheat.positioning_answers(profile, company_name, role_title, job_description):
        add_bullet(document, f"{item.prompt}: {item.answer}")
    add_subsection(document, "Most Common Questions With Model Answers")
    for item in cheat.common_interview_answers(
        profile,
        company_name,
        role_title,
        job_description,
        stories,
        supplied_context=supplied_context,
        resume_text=resume_text,
        notes_text=notes_text,
    ):
        add_bullet(document, item.prompt)
        add_body(document, item.answer)
    keyword_answers = cheat.keyword_ready_answers(
        profile,
        job_description,
        stories,
        resume_text=resume_text,
        supplied_context=supplied_context,
    )
    if keyword_answers:
        add_subsection(document, "Keyword Questions With Model Answers")
        for item in keyword_answers:
            add_bullet(document, item.prompt)
            add_body(document, item.answer)


def add_reflection_prompt_bank(
    document: Document,
    profile: build_resume.JobProblemProfile,
    stories: list[cheat.StoryCard],
    company_name: str,
    role_title: str,
) -> None:
    add_section(document, "Pre-Interview Reflection Prompts")
    add_body(
        document,
        "Use these prompts to turn preparation into a sharper interview voice. The goal is to make Christian's value obvious, memorable, and easy for the interviewer to repeat in debrief.",
    )
    add_subsection(document, "Top Five Career Wins")
    for story in stories[:5]:
        add_bullet(document, f"{story.title}: {story.result}")
    add_subsection(document, "Traits This Role Likely Rewards")
    for line in (
        f"Outcome ownership: show how Christian turns {profile.core_problem} into measurable progress.",
        "Translation skill: show how technical, customer, and business stakeholders get aligned around the same next step.",
        "Judgment under ambiguity: show what Christian noticed before acting, not only the action itself.",
        "Commercial or business awareness: connect adoption, reporting, risk, or implementation quality to value for the customer or company.",
        "Learning speed: show how Christian becomes useful in unfamiliar products or industries without overclaiming domain ownership.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Debrief Words To Earn")
    add_body(
        document,
        f"Target debrief sentence: Christian came across as clear, accountable, and useful. He understood {company_name}'s problem, gave specific proof, and sounded genuinely interested in the {role_title} role."
    )
    add_subsection(document, "Uncomfortable Questions To Prepare")
    for prompt in (
        "Where is your background less direct than the job description?",
        "Have you owned the exact tool, industry, or technical stack we use?",
        "How do you balance being customer-facing with owning commercial outcomes?",
        "Why do you want this role instead of a more traditional implementation, CSM, or systems role?",
        "What would make you fail in the first 60 to 90 days?",
    ):
        add_bullet(document, prompt)
    add_body(
        document,
        "Answer pattern for uncomfortable questions: acknowledge the truth, name the transferable pattern, give one proof point, and explain how Christian would close the gap quickly."
    )
    add_subsection(document, "Ownership Audit Per Story")
    for card in stories[:5]:
        i_count = len(re.findall(r"\bI\b", card.evidence))
        we_count = len(re.findall(r"\bwe\b", card.evidence, re.I))
        flag = " - REVIEW: more 'we' than 'I'" if we_count > i_count else " - ownership language OK"
        add_bullet(document, f"{card.title}{flag}")
    add_subsection(document, "Silent-Watch Checklist")
    for line in (
        "Mute the recording and watch the first ten seconds only. Does Christian look like he already knows the point, or like he is searching for the opening?",
        "Check posture and eye line. The goal is steady, useful, and direct, not overperformed or apologetic.",
        "Watch for visible drift after the main point lands. If the answer looks finished before the mouth stops moving, the close is too long.",
        "Notice whether the result line looks confident or rushed. The visual rhythm should slow down slightly when the proof arrives.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Sound-Watch Checklist")
    for line in (
        "Listen without video and score sentence one: did the answer land the point, or did it warm up too long before saying anything concrete?",
        "Mark hedge words and filler such as 'I think,' 'kind of,' 'probably,' 'basically,' or repeated 'um.' Remove the first avoidable hedge before the next take.",
        "Check whether the answer keeps one clean arc: point, proof, relevance. If it loops backward, the middle is too loose.",
        "Listen for whether the last sentence sounds finished and memorable, not like Christian is still explaining after the answer is already clear.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Reusable Recording Audit")
    for line in cheat.recording_audit_lines(profile, role_title):
        add_bullet(document, line)
    add_subsection(document, "Lane-Specific Recording Focus")
    for line in lane_specific_recording_focus(profile, role_title):
        add_bullet(document, line)


def lane_specific_recording_focus(profile: build_resume.JobProblemProfile, role_title: str) -> list[str]:
    focus_by_lane = {
        "implementation_delivery": [
            f"For the {role_title} lane, make sure recordings show calm control of scope, risk, handoffs, and adoption rather than generic project language.",
            "Listen for whether the answer proves go-live realism: who needed alignment, what risk was surfaced early, and what changed because Christian stayed on it.",
            "If the story sounds like task coordination only, add the customer, workflow, or business consequence.",
        ],
        "presales_solution": [
            f"For the {role_title} lane, recordings should sound buyer-aware and commercially useful, not like product narration.",
            "Check whether Christian explains what he discovered before describing what he showed. Discovery should lead the demo, not the other way around.",
            "If the answer could fit any seller, add the implementation realism or stakeholder tradeoff that made the recommendation credible.",
        ],
        "customer_success": [
            f"For the {role_title} lane, recordings should sound like ownership of value, risk, and trust after the sale, not generic relationship management.",
            "Listen for whether the answer makes account health visible through signals, decisions, and follow-through rather than broad customer-care language.",
            "If the answer never reaches renewal risk, adoption, executive alignment, or value realization, the commercial edge is too soft.",
        ],
        "change_enablement": [
            f"For the {role_title} lane, recordings should prove behavior change, stakeholder clarity, and adoption follow-through, not just communications activity.",
            "Check whether Christian names the human operating gap before describing training or rollout mechanics.",
            "If the answer sounds like change theater, add the behavior that shifted and the proof that it held.",
        ],
        "analytics_operations": [
            f"For the {role_title} lane, recordings should connect data work to a decision, not just a dashboard.",
            "Listen for whether Christian names the business question, the trust issue in the data, and the action the reporting made easier.",
            "If the story ends at analysis, add what changed in the workflow, forecast, or operating choice.",
        ],
        "corporate_strategy": [
            f"For the {role_title} lane, recordings should sound structured and decision-ready, not purely conceptual.",
            "Check whether Christian frames the problem, the tradeoffs, the recommendation, and the implementation consequence in that order.",
            "If the answer sounds smart but not useful, add the owner, risk, or next-step logic.",
        ],
    }
    return focus_by_lane.get(
        profile.primary_lane,
        [
            f"For the {role_title} lane, recordings should make Christian sound like someone who has already handled similar complexity, not someone describing adjacent experience from a distance.",
            "Listen for whether the answer names the problem, the proof, and the business consequence clearly enough to repeat in debrief.",
            "If the answer feels generic, add one sharper operating detail and one measurable result.",
        ],
    )


def add_story_page(
    document: Document,
    card: cheat.StoryCard,
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str = "",
    interview_notes: str = "",
    resume_text: str = "",
    round_records: Sequence[Mapping[str, object]] = (),
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> None:
    answer = story_sample_answer(card, profile, company_name, role_title, job_description, interview_notes, resume_text)
    enforce_prose_quality(
        answer.spoken,
        "interview_story_answer",
        label=f"Detailed guide story answer ({card.title})",
        mode="warn",
        check_template_leakage=False,
    )
    add_section(document, card.title)
    add_subsection(document, "Sample Answer")
    add_story_answer(document, answer)
    add_subsection(document, "Best Uses And Proof Notes")
    add_bullet(document, f"Best for: {', '.join(card.story_types)}")
    add_bullet(document, f"Hook: {card.hook}")
    add_bullet(document, f"Three takeaways: {'; '.join(card.takeaways)}")
    add_bullet(document, f"Level 3 trait to show: {cheat.spoken_level3_trait_sentence(card.level3_trait)}")
    add_bullet(document, f"Evidence: {card.evidence}")
    add_bullet(document, f"Result: {card.result}")
    add_subsection(document, "Anchor Facts")
    add_bullet(document, cheat.story_anchor_fact_line(card, profile))
    add_subsection(document, "Adaptation Drills")
    for line in cheat.story_adaptation_drill_lines(card, profile, company_name, role_title, job_description):
        add_bullet(document, line)
    add_subsection(document, "Delivery Notes")
    for line in cheat.story_delivery_note_lines(card, profile, round_records, global_round_records):
        add_bullet(document, line)
    add_subsection(document, "Follow-Up Drill")
    for line in story_quality_audit(card, profile):
        add_bullet(document, line, small=True)



def add_final_round_strategy_section(
    document: Document,
    company_name: str,
    role_title: str,
    job_description: str,
    profile: build_resume.JobProblemProfile,
    insights: PrepInsights,
) -> None:
    add_section(document, "Final-Round Conversion Strategy")
    add_body(
        document,
        "At later-stage interviews, assume the baseline question has changed from whether Christian can do the work to whether he is the safest, clearest, most useful person to hire. The strategy is to mirror the company's problem before pitching himself, then offer a small amount of practical value in the conversation."
    )
    add_subsection(document, "Go Two Levels Deeper")
    add_bullet(document, f"Level 1 is the job description. Level 2 is what {company_name} is probably trying to fix in the next 6 to 12 months: {profile.core_problem}.")
    add_bullet(document, "Before the interview, identify the likely business goal, the likely operating challenge, and one practical initiative Christian could own early.")
    if insights.situation_read:
        add_bullet(document, f"Use the strongest supplied-note interpretation: {insights.situation_read[0]}")
    add_subsection(document, "Run The Swap Test")
    add_bullet(document, f"If an answer would still make sense after replacing {company_name} with another company, add one company-specific detail from the posting, interview notes, or verified research.")
    add_bullet(document, "Do not over-personalize. One specific detail plus one relevant proof point is usually stronger than a long speech.")
    add_subsection(document, "Differentiation Test")
    positioning = build_resume.obvious_choice_positioning(profile, job_description)
    add_bullet(document, positioning.get("short_line", ""))
    add_bullet(document, positioning.get("sentence", ""))
    add_bullet(document, "If another candidate could say that line without changing a word, add one stronger proof marker such as scale, stakeholder level, workflow complexity, or measurable outcome.")
    add_subsection(document, "Value Validation Project Ideas")
    for idea in value_validation_project_ideas(company_name, role_title, profile):
        add_bullet(document, idea)
    add_subsection(document, "Mirror Before You Pitch")
    add_body(
        document,
        f"A useful opening pattern: 'Based on what I read, it seems like the {role_title} role is less about owning tasks in isolation and more about turning {profile.core_problem} into progress people can actually use. The closest pattern in my background is...'"
    )
    add_subsection(document, "Think Like You Have Been Hired")
    add_bullet(document, "Ask at least one question that assumes responsibility, such as: 'If we fast-forward one year, what would this hire need to have changed for the team to say this was clearly the right decision?'")
    add_bullet(document, "Ask one risk question, such as: 'What usually slows this work down here: stakeholder alignment, data quality, competing priorities, adoption, or something else?'")
    add_subsection(document, "Build An Internal Advocate")
    add_bullet(document, "In the thank-you note, recap one specific point from the conversation, restate why the role still interests Christian, and add one useful thought that connects his background to the team's problem.")
    add_bullet(document, "If they raise a concern, answer it directly in the follow-up with humility and proof. The goal is to make it easy for someone in the room to explain his fit to others.")


def value_validation_project_ideas(company_name: str, role_title: str, profile: build_resume.JobProblemProfile) -> tuple[str, ...]:
    if profile.primary_lane == "customer_success":
        return (
            "Bring a 30-day account-health readout concept: what signals would show adoption risk, stakeholder risk, business-value risk, and renewal risk.",
            "Offer a simple QBR structure: customer goal, usage/adoption evidence, business impact, open risks, and next expansion or value-realization opportunity.",
            "Prepare a brief customer recovery play: diagnose root cause, reset success criteria, align owners, confirm next milestone, and document the value path.",
        )
    if profile.primary_lane == "presales_solution":
        return (
            "Bring a discovery map: buyer problem, current workflow, stakeholders, technical constraints, decision criteria, and proof needed before purchase.",
            "Offer a demo hypothesis: what Christian would validate before showing product value, and how he would connect features to business outcomes.",
            "Prepare a buyer-risk map: where trust, integration, data, adoption, or executive alignment could stall the deal.",
        )
    if profile.primary_lane in {"change_enablement", "corporate_strategy"}:
        return (
            "Bring a one-page change-risk map: impacted groups, visible resistance, operating risk, adoption metric, and first decision needed.",
            "Offer a first-90-days structure: listen, map stakeholders, identify friction, prioritize the highest-value process, pilot, measure, and scale.",
            "Prepare a decision memo outline: problem, facts, options, tradeoffs, recommendation, risks, and implementation checkpoints.",
        )
    if profile.primary_lane == "analytics_operations":
        return (
            "Bring a metrics-tree concept: business goal, operating metric, source data, owner, refresh cadence, and decision the report should support.",
            "Offer a dashboard quality checklist: trusted source, clear owner, action threshold, exception logic, and adoption feedback loop.",
            "Prepare a data-confidence question set: what data leaders trust today, where manual work still exists, and which decision needs faster visibility.",
        )
    return (
        "Bring a rollout-risk map: stakeholders, milestones, technical dependencies, training needs, validation checkpoints, and adoption risks.",
        "Offer a practical implementation playbook: discovery, requirements, configuration, testing, training, go-live, hypercare, and handoff.",
        "Prepare a status-cadence example: milestone, owner, blocker, decision needed, customer impact, and next checkpoint.",
    )


def add_hidden_assessment_section(
    document: Document,
    company_name: str,
    role_title: str,
    profile: build_resume.JobProblemProfile,
    job_description: str,
    stories: list[cheat.StoryCard],
) -> None:
    add_section(document, "What They Are Really Asking")
    add_body(
        document,
        "Use this grid to answer the hidden test behind common interview questions. Keep the answer brief, proof-based, and easy for the interviewer to remember in debrief.",
    )
    for card in cheat.question_intent_framework(profile, company_name, role_title, job_description, stories):
        add_subsection(document, card.prompt)
        add_bullet(document, f"What they are evaluating: {card.hidden_assessment}")
        add_bullet(document, f"Bad-answer trap: {card.bad_answer_trap}")
        add_bullet(document, f"Best story angle: {card.story_angle}")
        add_bullet(document, f"When to clarify first: {card.clarify_when}")
        add_bullet(document, f"Ideal answer length: {card.ideal_length}")











def add_general_answer_operating_system(document: Document, profile: build_resume.JobProblemProfile, framework_selection: cheat.AnswerFrameworkSelection) -> None:
    add_section(document, "Answer Operating System")
    add_subsection(document, "Answer Framework Hierarchy")
    for line in cheat.answer_framework_section_lines(framework_selection):
        add_bullet(document, line)
    add_subsection(document, "One Framework, Three Modes")
    for line in (
        "Do not reach for an acronym during the interview. Reach for the practiced sequence.",
        f"Primary full mode, 60-90 seconds: {framework_selection.label}.",
        "Brevity mode, 20-30 seconds: Context -> What I Noticed -> Result -> Bridge.",
        "Short answer mode, 15-20 seconds: Point -> Reason -> Example -> Point restatement.",
        "STAR, CAR, SAR, CAAR, CART, and HERO are prep labels. The live answer should sound like practiced judgment, not an acronym recital.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Two Question Types")
    for line in cheat.response_calibration_lines():
        add_bullet(document, line)
    add_subsection(document, "Story-Specific Bridge Rule")
    add_body(
        document,
        f"Every story bridge must name the specific proof the story gives for {profile.core_problem}: process gap, customer recovery, decision-quality improvement, ramp method, stakeholder tradeoff, validation checkpoint, or audience translation. Do not repeat one generic bridge across multiple stories."
    )
    add_subsection(document, "Calibration Question Rule")
    add_body(document, "End only the most important stories with a calibration question, and make each question story-specific. Repeating the same closer makes the answer sound over-rehearsed instead of adaptive.")













def add_thank_you_strategy_section(document: Document, company_name: str, role_title: str) -> None:
    add_section(document, "Thank-You Note Strategy")
    add_bullet(document, "Send it quickly. The note should feel personal, not like a second cover letter.")
    add_bullet(document, "Structure: thank them, name one specific conversation point, connect that point to Christian's proof, restate interest, and leave room for this role or another strong-fit opportunity if appropriate.")
    add_bullet(document, "Match the tone from the first five minutes: if the interview felt structured, keep the email tighter; if it felt relational, let one line carry more warmth without turning generic.")
    add_subsection(document, "Reusable Pattern")
    add_body(
        document,
        f"Thank you for taking the time to speak with me about the {role_title} role. I appreciated the conversation, especially the discussion around the team's most important implementation priorities. The part that stood out to me was the need for someone who can ask the right questions early, communicate risks clearly, and turn complex work into something teams can actually use. This is the part of my background I would be most excited to bring to {company_name}."
    )
    add_body(
        document,
        "If a concern came up in the interview, add one calm sentence that acknowledges it and reframes the fit with proof. Keep it direct and brief."
    )
    add_subsection(document, "Post-Interview Memory Capture")
    for line in (
        "Within 30 minutes, write down the exact words they used for the role's biggest problem, success metrics, team friction, customer pain, and any concern about fit.",
        "Mark which stories created follow-up questions or visible interest. Use those as primary stories in the next round.",
        "Capture any unexpected question and build a cleaner answer before the next conversation.",
        "Add any insider company language to jobs\\interview_notes.txt or jobs\\company_research.txt before rerunning future interview materials.",
    ):
        add_bullet(document, line)




















def build_document(company_name: str, role_title: str, job_description: str, resume_docx: Path, output_docx: Path) -> None:
    resume_text = "\n".join(cheat.paragraph_texts(resume_docx))
    profile = cheat.adjusted_profile_for_role(
        build_resume.job_problem_profile(job_description, resume_text),
        role_title,
        job_description,
    )
    jobs_dir = PROJECT_ROOT / "jobs"
    context_bundle = interview_context.load_company_context(
        jobs_dir,
        company_name,
        role_title,
        company_research_path=COMPANY_RESEARCH,
        global_interview_notes_path=INTERVIEW_NOTES,
        mode="compact",
    )
    global_round_records = interview_context.global_coaching_fallback_records(jobs_dir, company_name, role_title)
    company_research = context_bundle.company_research
    interview_notes = context_bundle.interview_notes
    supplied_context = context_bundle.supplied_context
    framework_selection = cheat.answer_framework_selection(job_description, interview_notes)
    insights = build_prep_insights(company_name, role_title, job_description, company_research, interview_notes, profile)
    verified_research = verified_company_research_points(company_name, job_description, company_research, interview_notes)
    stories = cheat.supported_story_bank(resume_text)
    debrief_summary = debrief_analysis.analyze_entries(context_bundle.round_records, company_name)
    hero_stories = cheat.hero_stories(profile, job_description, resume_text)
    if len(stories) < 6:
        fail("not enough resume-supported stories available for a detailed interview guide")
    if len(hero_stories) < 3:
        hero_stories = stories[:5]
    hero_stories = debrief_analysis.reorder_story_cards(hero_stories, debrief_summary)
    state_farm_mode = is_state_farm_active(company_name, role_title, job_description, company_research, interview_notes)

    document = Document()
    set_default_style(document)
    add_title(document, company_name, role_title)

    firm_callout = cheat.company_profile_interview_callout(company_name, job_description, profile)
    if firm_callout:
        add_section(document, "Firm-Specific Interview Profile")
        for line in firm_callout:
            add_bullet(document, line)

    add_section(document, "How To Use This Guide")
    add_bullet(document, "This is the long-form interview prep document. Use it when an interview is scheduled, not for every resume generation run.")
    add_bullet(document, "Keyword Answer Reference is at the back of the guide. Use it for targeted practice after reviewing the story bank.")
    add_section(document, "Answer Framework Hierarchy")
    for line in cheat.answer_framework_section_lines(framework_selection):
        add_bullet(document, line)
    add_section(document, "Rehearsal Method")
    for line in cheat.rehearsal_foundation_lines(role_title, context_bundle.round_records, global_round_records):
        add_bullet(document, line)
    add_section(document, "Anti-Filler And Length Control")
    for line in cheat.answer_mode_lines():
        add_bullet(document, line)
    for title, description, fix in cheat.communication_audit_reference(job_description, interview_notes)[:3]:
        add_bullet(document, f"{title}: {description} {fix}")
    add_story_anchor_system_section(document, hero_stories, profile)
    risk_label, risk_warning = cheat.candidate_archetype_assessment(profile, job_description, resume_text, supplied_context, interview_notes)
    add_section(document, "Self-Assessment: Your Likely Interview Risk")
    add_bullet(document, risk_label)
    add_bullet(document, risk_warning)
    add_section(document, "Top Recurring Answer Risks")
    for line in cheat.top_answer_risk_lines(profile, company_name, role_title, context_bundle.round_records, global_round_records):
        add_bullet(document, line)
    diagnosis_lines = cheat.latest_positioning_diagnosis_lines(context_bundle.round_records, global_round_records)
    if diagnosis_lines:
        add_section(document, "Latest Positioning Diagnosis")
        for line in diagnosis_lines:
            add_bullet(document, line)
    rewrite_lines = cheat.ownership_language_rewrite_lines(context_bundle.round_records, global_round_records)
    if rewrite_lines:
        add_section(document, "Ownership And Consultative Rewrites")
        for line in rewrite_lines:
            add_bullet(document, line)
    add_section(document, "Best Example To Use First")
    for line in cheat.best_example_to_use_first_lines(profile, context_bundle.round_records, global_round_records):
        add_bullet(document, line)
    add_section(document, "Six Offer Blockers To Avoid")
    for line in six_offer_blocker_lines(profile, company_name, role_title, job_description):
        add_bullet(document, line)
    add_section(document, "Executive Evaluation: Four Trust Questions")
    for line in four_trust_questions_audit(profile, company_name, role_title, job_description):
        add_bullet(document, line)
    add_section(document, "Executive Presence Signals")
    for line in executive_presence_signals(profile, company_name, role_title, job_description):
        add_bullet(document, line)
    add_section(document, "Executive Presence Corrections")
    for line in cheat.executive_presence_correction_lines(company_name, role_title, context_bundle.round_records, global_round_records):
        add_bullet(document, line)
    post_round_lines = list(cheat.post_round_intelligence_lines(supplied_context))
    if debrief_summary.most_common_question:
        post_round_lines.append(
            f"Recurring question across prior rounds ({debrief_summary.recurring_question_count}x): {debrief_summary.most_common_question}"
        )
    if debrief_summary.top_story_title:
        post_round_lines.append(
            f"Story that has drawn the strongest follow-up so far: {debrief_summary.top_story_title}. Keep it first-tier and easy to retell under pressure."
        )
    if post_round_lines:
        add_section(document, "Post-Round Intelligence To Prepare")
        for line in post_round_lines:
            add_bullet(document, line)
    if debrief_summary.top_coaching_signals:
        add_section(document, "Recurring Delivery Habits")
        for item in debrief_summary.top_coaching_signals:
            add_bullet(document, item)
    add_extended_tmay_section(document, profile, company_name, role_title, job_description, resume_text, interview_notes)
    add_hidden_assessment_section(document, company_name, role_title, profile, job_description, stories)
    if state_farm_mode:
        add_state_farm_full_workbook(document, profile, company_name, role_title, job_description, resume_text, hero_stories, add_keyword_question_bank)
    else:
        add_company_fit_answer_bank(
            document,
            profile,
            company_name,
            role_title,
            job_description,
            stories,
            supplied_context,
            resume_text,
            interview_notes,
        )

        add_reflection_prompt_bank(document, profile, hero_stories, company_name, role_title)

        add_section(document, "Role Challenge Forecast")
        for line in cheat.role_challenge_forecast(profile, company_name, role_title, job_description):
            add_bullet(document, line)

        add_subsection(document, "First 90-Day Approach")
        for line in cheat.first_90_day_approach(profile):
            add_bullet(document, line)
        add_subsection(document, "Gaps To Manage Honestly")
        for line in cheat.role_specific_gaps(profile, job_description):
            add_bullet(document, line)

        add_page_break(document)
        add_section(document, "Primary Story Bank With Sample Answers")
        for index, card in enumerate(hero_stories[:5]):
            if index:
                add_page_break(document)
            add_story_page(
                document,
                card,
                profile,
                company_name,
                role_title,
                job_description,
                interview_notes,
                resume_text,
                context_bundle.round_records,
                global_round_records,
            )

        add_page_break(document)
        add_section(document, "Additional Behavioral Answers")
        for prompt, answer in behavioral_sample_answers(profile, stories, company_name, role_title, job_description, interview_notes, resume_text):
            enforce_prose_quality(
                answer.spoken,
                "interview_story_answer",
                label=f"Detailed guide behavioral answer ({prompt[:50]})",
                mode="warn",
                check_template_leakage=False,
            )
            add_subsection(document, prompt)
            add_story_answer(document, answer)

        add_page_break(document)
        add_section(document, "Likely Interview Questions")
        used_likely_story_titles: set[str] = set()
        for item in cheat.likely_questions(profile, job_description):
            add_subsection(document, item.question)
            if item.question.lower().startswith("tell me about yourself"):
                add_body(
                    document,
                    cheat.ninety_second_pitch(
                        profile,
                        company_name,
                        role_title,
                        job_description,
                        resume_text,
                        interview_notes,
                    ),
                )
                add_body(document, f"Follow-up angle: {item.angle}", small=True)
                continue
            matching_story = cheat.likely_question_story(item, stories, used_likely_story_titles)
            used_likely_story_titles.add(matching_story.title)
            add_story_answer(
                document,
                story_sample_answer(matching_story, profile, company_name, role_title, job_description, interview_notes, resume_text),
                prefix="Model answer: ",
            )
            add_body(document, f"Follow-up angle: {item.angle}", small=True)
        add_application_question_prep_section(document, job_description, resume_text)
        add_recent_interview_question_prep_section(
            document,
            job_description,
            company_name,
            role_title,
            jobs_dir=jobs_dir,
            profile=profile,
            stories=stories,
            resume_text=resume_text,
            interview_notes=interview_notes,
        )

        add_section(document, "Three Supported Proof Themes")
        for line in cheat.three_supported_proof_theme_lines(profile, stories):
            add_bullet(document, line)

        add_page_break(document)
        add_section(document, "Answer Mechanics Reference")
        add_general_answer_operating_system(document, profile, framework_selection)

        add_section(document, "Story Selection Decision Table")
        add_body(document, "Use this as a quick index when an interviewer asks for a different example.")
        for line in story_selection_decision_table(hero_stories, profile, job_description):
            add_bullet(document, line)

        add_pushback_section(document, insights)
        add_anticipated_question_section(document, insights)
        add_business_context_question_section(
            document,
            job_description,
            supplied_context,
            stories,
            profile=profile,
            company_name=company_name,
            role_title=role_title,
            resume_text=resume_text,
            interview_notes=interview_notes,
        )

        add_section(document, "QUESTIONS TO ASK AND HOW TO CLOSE")
        add_subsection(document, "Strategic Question Filter")
        for line in cheat.strategic_question_filter_lines():
            add_bullet(document, line)
        add_subsection(document, "Stronger Consultative Questions To Ask")
        for line in cheat.consultative_question_drill_lines(profile, company_name, role_title):
            add_bullet(document, line)
        add_subsection(document, "Structural Diagnostic Question To Lead With")
        add_bullet(document, cheat.lane_structural_diagnostic_question(profile, company_name, role_title, job_description))
        supplied_questions, supplied_changes = cheat.reframe_questions_to_positive(
            cheat.supplied_smart_questions(supplied_context)
        )
        if supplied_changes:
            cheat.log_reframed_questions("detailed supplied questions", supplied_changes)
        if supplied_questions:
            add_subsection(document, "Highest-Value Supplied Questions")
            for question in supplied_questions[:12]:
                add_bullet(document, question)
        add_subsection(document, "Highest-Value Questions From The Notes")
        reframed_smart_questions, smart_changes = cheat.reframe_questions_to_positive(list(insights.smart_questions[:10]))
        if smart_changes:
            cheat.log_reframed_questions("detailed smart questions", smart_changes)
        for question in reframed_smart_questions:
            if question in supplied_questions:
                continue
            add_bullet(document, question)
        add_subsection(document, "General Backup Questions")
        final_questions: list[str] = []
        reframed_final_questions, final_changes = cheat.reframe_questions_to_positive(
            cheat.questions_to_ask(company_name, profile, job_description, supplied_context)
        )
        if final_changes:
            cheat.log_reframed_questions("detailed backup questions", final_changes)
        for question in reframed_final_questions:
            final_questions.append(question)
            add_bullet(document, question)
        warnings = cheat.interview_question_quality_warnings(list(supplied_questions) + reframed_smart_questions + final_questions)
        if warnings:
            add_subsection(document, "Question Quality Audit")
            for warning in warnings:
                add_bullet(document, warning)
        add_subsection(document, "Closing Mechanics")
        add_subsection(document, "Fit Summary Before The Close")
        for line in cheat.closing_fit_summary(profile):
            add_bullet(document, line)
        add_subsection(document, "Closing The Interview Drill")
        for line in cheat.closing_interview_drill_lines(company_name, role_title, profile, job_description, supplied_context):
            add_bullet(document, line)
        for title, lines in cheat.closing_mechanics(company_name, role_title, profile, job_description, supplied_context):
            add_subsection(document, title)
            for line in lines:
                add_bullet(document, line)
        add_thank_you_strategy_section(document, company_name, role_title)
        add_subsection(document, "Final Reminder")
        for line in cheat.answer_do_dont(job_description):
            add_bullet(document, line, small=True)

        add_page_break(document)
        add_keyword_question_bank(document, profile, job_description, resume_text, hero_stories, company_name, role_title)

    body = document_text(document)
    scrub_document_for_job_language(document, job_description)
    body = document_text(document)
    build_resume.assert_no_erp_language_for_non_erp_role(body, job_description, "detailed interview guide")
    validate_text(body, company_name=company_name, role_title=role_title)
    output_docx.parent.mkdir(exist_ok=True)
    document.save(str(output_docx))


def build_detailed_interview_guide() -> DetailedGuideResult:
    build_resume.require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    build_resume.require_file(JOB_DESCRIPTION, "job description")
    job_description = read_text(JOB_DESCRIPTION)
    if not job_description:
        fail("job description is empty")

    company_name = build_resume.extract_output_name(job_description)
    output_target_name = build_resume.extract_output_target_name(job_description)
    role_title = build_cover_letter.extract_role_title(job_description) or "Role"
    resume_docx = build_cover_letter.find_resume_output(job_description)
    resume_audit_state = resume_analysis.output_audit_state(resume_docx)
    output_name = f"Christian Estrada - {output_target_name} Detailed Interview Guide.docx"
    if resume_audit_state == "FAIL":
        output_name = f"Christian Estrada - {output_target_name} FAIL Detailed Interview Guide.docx"
    elif resume_audit_state == "POOR":
        output_name = f"Christian Estrada - {output_target_name} POOR Detailed Interview Guide.docx"
    elif resume_audit_state == "BRIDGE":
        output_name = f"Christian Estrada - {output_target_name} BRIDGE Detailed Interview Guide.docx"
    output_docx = OUTPUT_DIR / output_name
    return build_detailed_interview_guide_for_inputs(
        job_description=job_description,
        resume_docx=resume_docx,
        output_docx=output_docx,
        company_name=company_name,
        role_title=role_title,
    )


def build_detailed_interview_guide_for_inputs(
    *,
    job_description: str,
    resume_docx: Path,
    output_docx: Path,
    company_name: str,
    role_title: str,
) -> DetailedGuideResult:
    build_resume.require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    if not job_description.strip():
        fail("job description is empty")
    if not role_title or role_title == "Role":
        fail("could not determine role title; refusing to create a placeholder detailed interview guide")
    assert_company_name_in_source(company_name, job_description, label="detailed interview guide")
    prompt_state = question_prep.load_application_prompt_state()
    question_issues = question_prep.application_question_context_issues(job_description, prompt_state, workflow="commercial")
    actual_output = question_prep.application_question_draft_path(output_docx) if question_issues else output_docx
    with prose_engine.collect_spoken_repair_issues() as spoken_issues:
        build_document(company_name, role_title, job_description, resume_docx, actual_output)
    review_issues = tuple(dict.fromkeys((*question_issues, *spoken_issues)))
    if spoken_issues and not question_issues:
        draft_output = question_prep.application_question_draft_path(output_docx)
        actual_output.replace(draft_output)
        actual_output = draft_output
    if review_issues:
        question_prep.mark_docx_as_draft(actual_output, review_issues)
    render_checks.render_docx(actual_output)
    return DetailedGuideResult(company_name, role_title, resume_docx, actual_output)


def main() -> None:
    result = build_detailed_interview_guide()
    print(f"Company: {result.company_name}")
    print(f"Role: {result.role_title}")
    print(f"Resume source: {result.resume_docx}")
    print(f"Output DOCX: {result.output_docx}")


if __name__ == "__main__":
    main()
