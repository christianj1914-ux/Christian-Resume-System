#!/usr/bin/env python3
"""Build Christian Estrada's tailored federal DOCX resume."""

from __future__ import annotations

import _bootstrap

_bootstrap.ensure_script_path()

import argparse
import json
import re
import shutil
import uuid
from collections.abc import Sequence
from dataclasses import dataclass, replace
from datetime import date
from pathlib import Path

import build_standard_qualifications_statement
import question_prep
import render_checks
import evidence_engine
import requirement_engine
import prose_engine
import resume_analysis
import workflow_step_runner
from build_resume import _check_job_description_quality, docx_visible_text_from_path, strip_linkedin_job_board_boilerplate
from config.paths import (
    FEDERAL_ESSAY_SOURCE,
    FEDERAL_JOB_DESCRIPTION,
    FEDERAL_RESUME_SOURCE,
    OUTPUT_DIR,
    PROJECT_ROOT,
    SCRATCH_DIR,
)
from config.evidence_rules import GRADE_EQUIVALENCE_TEMPLATE
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from resume_analysis import (
    JobProblemProfile,
    audit_keyword_sort_key,
    audit_keywords,
    extract_company_name,
    extract_job_title,
    job_problem_profile,
    keyword_hits,
    keyword_set,
)
from resume_content import rewrite_supported_text, strengthen_outcome_framing
from resume_format import rendered_page_count, unpack_docx
from config.summary_contracts import FEDERAL_SUMMARY_WORD_RANGE
from utils import enforce_prose_quality, fail, has_great_eight_signal, read_text


FEDERAL_REQUIRED_SECTIONS = (
    "Professional Summary",
    "Technical Skills",
    "Work Experience",
    "Education",
    "Professional Development",
)
TARGET_PAGE_COUNT = 2
MAX_QUALIFICATIONS_PAGE_COUNT = 3
DEFAULT_QUESTION_CHAR_LIMIT = 4000
MIN_SUMMARY_WORDS, MAX_SUMMARY_WORDS = FEDERAL_SUMMARY_WORD_RANGE
NAME_FONT_SIZE = 14
SECTION_FONT_SIZE = 10
ROLE_FONT_SIZE = 10
MIN_BODY_FONT_SIZE = 10.0
HIGH_PRIORITY_BUCKET_THRESHOLD = 80


@dataclass(frozen=True)
class FederalRole:
    title: str
    start: str
    end: str
    hours_per_week: str
    salary: str
    company: str
    location: str
    supervisor: str
    supervisor_phone: str
    company_summary: str
    job_summary: str
    bullets: tuple[str, ...]


@dataclass(frozen=True)
class FederalEducation:
    degree: str
    date: str
    school: str
    details: str


@dataclass(frozen=True)
class FederalVolunteerExperience:
    title: str
    organization: str
    description: str


@dataclass(frozen=True)
class FederalContact:
    name: str
    location: str
    email: str
    phone: str
    citizenship: str
    veterans_preference: str
    clearance: str
    availability: str


@dataclass(frozen=True)
class FederalSource:
    contact: FederalContact
    technical_skills: tuple[str, ...]
    roles: tuple[FederalRole, ...]
    education: tuple[FederalEducation, ...]
    professional_development: tuple[str, ...]
    volunteer_experience: tuple[FederalVolunteerExperience, ...] = ()


@dataclass(frozen=True)
class FederalLayoutProfile:
    name: str
    font_size: float
    total_bullets: int
    company_summary_roles: int
    job_summary_roles: int
    max_bullets_by_role: tuple[int, ...]
    margins: tuple[float, float, float, float]
    blank_after_sections: bool
    blank_after_roles: bool
    show_duties_label: bool = False


@dataclass(frozen=True)
class FederalQualificationLayout:
    name: str
    font_size: float
    margins: tuple[float, float, float, float]
    blank_after_sections: bool


@dataclass(frozen=True)
class FederalRequirementBucket:
    label: str
    kind: str
    text: str
    priority: int
    clusters: tuple[str, ...]
    # Buckets can be split for readable evidence without multiplying their
    # source-list's audit weight or warning volume.
    weight_group_id: str = ""


@dataclass(frozen=True)
class FederalRequirementCoverage:
    bucket: FederalRequirementBucket
    alignment: str
    source_refs: tuple[str, ...]
    selected_refs: tuple[str, ...]
    narrative: str


@dataclass(frozen=True)
class FederalAudit:
    buckets: tuple[FederalRequirementCoverage, ...]
    cluster_weights: tuple[tuple[str, int], ...]
    keyword_targets: tuple[str, ...]
    warnings: tuple[str, ...]
    requirements: tuple[requirement_engine.RequirementElement, ...] = ()
    element_matches: tuple[evidence_engine.EvidenceMatch, ...] = ()
    target_context: requirement_engine.TargetContext | None = None
    coverage_report: evidence_engine.CoverageReport | None = None


@dataclass(frozen=True)
class FederalCompetencyPrompt:
    name: str
    description: str


@dataclass(frozen=True)
class FederalKSAItem:
    name: str
    text: str
    clusters: tuple[str, ...]


@dataclass(frozen=True)
class FederalApplicationQuestionSet:
    char_limit: int
    competencies: tuple[FederalCompetencyPrompt, ...]
    specialized_items: tuple[str, ...]
    ksas: tuple[FederalKSAItem, ...]

    @property
    def has_supplemental_questions(self) -> bool:
        return bool(self.competencies or self.specialized_items or self.ksas)


@dataclass(frozen=True)
class FederalQualificationsContent:
    questions: FederalApplicationQuestionSet
    standard_essay_answers: tuple[tuple[str, str], ...]
    active_question_responses: tuple[question_prep.QualificationsResponse, ...]
    recent_interviewer_scripts: tuple[tuple[str, str], ...]
    question_context_issues: tuple[str, ...]
    parse_diagnostics: tuple[requirement_engine.FederalParseDiagnostic, ...]
    question_one_paragraphs: tuple[str, ...] = ()
    question_one_labels: tuple[tuple[str, str], ...] = ()
    question_one_count: int = 0
    question_two_labels: tuple[tuple[str, str], ...] = ()
    question_two_count: int = 0


@dataclass(frozen=True)
class FederalEssayPrompt:
    key: str
    question: str
    answer: str


@dataclass(frozen=True)
class FederalBulletCandidate:
    role_index: int
    bullet_index: int
    score: int
    text: str
    matched_clusters: tuple[str, ...]
    matched_bucket_labels: tuple[str, ...]


@dataclass(frozen=True)
class FederalPlanCandidate:
    layout: FederalLayoutProfile
    bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...]
    audit: FederalAudit
    page_count: int | None
    quality_score: int


@dataclass(frozen=True)
class FederalQualificationPlan:
    layout: FederalQualificationLayout
    page_count: int | None


@dataclass(frozen=True)
class FederalResumePlan:
    summary: str
    audit: FederalAudit
    bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...]
    resume_layout: FederalLayoutProfile
    resume_page_count: int | None
    qualifications_layout: FederalQualificationLayout
    qualifications_page_count: int | None
    qualifications_content: FederalQualificationsContent


@dataclass(frozen=True)
class FederalBuildResult:
    output_docx: Path
    qualifications_docx: Path
    company_name: str
    page_count: int | None
    qualification_page_count: int | None
    layout_name: str
    selected_bullets: int


FEDERAL_LAYOUT_PROFILES = (
    FederalLayoutProfile(
        name="expanded",
        font_size=10.0,
        total_bullets=13,
        company_summary_roles=3,
        job_summary_roles=4,
        max_bullets_by_role=(5, 3, 2, 2, 1),
        margins=(0.42, 0.42, 0.5, 0.5),
        blank_after_sections=True,
        blank_after_roles=True,
    ),
    FederalLayoutProfile(
        name="balanced",
        font_size=10.0,
        total_bullets=12,
        company_summary_roles=3,
        job_summary_roles=3,
        max_bullets_by_role=(5, 3, 2, 1, 1),
        margins=(0.38, 0.38, 0.48, 0.48),
        blank_after_sections=True,
        blank_after_roles=False,
    ),
    FederalLayoutProfile(
        name="tight",
        font_size=10.0,
        total_bullets=11,
        company_summary_roles=2,
        job_summary_roles=2,
        max_bullets_by_role=(4, 3, 2, 1, 1),
        margins=(0.34, 0.34, 0.44, 0.44),
        blank_after_sections=False,
        blank_after_roles=False,
    ),
)

FEDERAL_QUALIFICATION_LAYOUTS = (
    FederalQualificationLayout(
        name="balanced",
        font_size=10.0,
        margins=(0.42, 0.42, 0.5, 0.5),
        blank_after_sections=True,
    ),
    FederalQualificationLayout(
        name="tight",
        font_size=10.0,
        margins=(0.34, 0.34, 0.44, 0.44),
        blank_after_sections=False,
    ),
)

FEDERAL_LANE_OPENERS = {
    "implementation_delivery": "implementation, enterprise systems, and technology delivery",
    "customer_success": "customer-facing implementation, enterprise systems, and service delivery",
    "presales_solution": "solution consulting, enterprise systems, and implementation delivery",
    "analytics_operations": "analytics, reporting, and enterprise systems delivery",
    "change_enablement": "change enablement, adoption, and enterprise systems delivery",
    "process_improvement": "process improvement, enterprise systems, and operational delivery",
    "corporate_strategy": "enterprise advisory, technology planning, and program delivery",
}

FEDERAL_OUTPUT_SKIP_LINES = {
    "duties",
    "qualifications",
    "how you will be evaluated",
    "attention to detail",
    "customer service",
    "oral communication",
    "problem solving",
    "accountability",
    "decision making",
    "flexibility",
    "integrity/honesty",
    "interpersonal skills",
    "learning",
    "reading comprehension",
    "reasoning",
    "self-management",
    "stress tolerance",
    "teamwork",
}

FEDERAL_IGNORE_SECTION_HEADINGS = {
    "overview",
    "summary",
    "key requirements",
    "benefits",
    "how to apply",
    "required documents",
    "conditions of employment",
    "security clearance",
    "drug test",
    "telework eligible",
    "remote job",
    "travel required",
    "appointment type",
    "work schedule",
    "relocation expenses reimbursed",
    "job family",
    "announcement number",
    "control number",
}

FEDERAL_SECTION_KIND_PATTERNS = (
    ("selective_factor", r"^selective factor\b[:\-]?\s*(.*)$"),
    ("specialized_experience", r"^specialized experience\b[:\-]?\s*(.*)$"),
    ("basic_qualifications", r"^(basic qualifications|requirements)\b[:\-]?\s*(.*)$"),
    ("qualifications", r"^qualifications\b[:\-]?\s*(.*)$"),
    ("evaluation", r"^how you will be evaluated\b[:\-]?\s*(.*)$"),
    ("duties", r"^duties\b[:\-]?\s*(.*)$"),
)

FEDERAL_KEYWORD_STOP_PHRASES = {
    "specialist",
    "specialist ai",
    "please",
    "receive",
    "responsibility",
    "detail",
    "environment",
    "government",
    "appointment",
    "process",
    "federal service",
    "receive training",
    "providing quality",
    "please refer",
}

FEDERAL_BUCKET_STOP_PHRASES = (
    "the experience may have been gained",
    "in addition to specialized experience",
    "must be u.s. citizen",
    "all newly appointed employees",
    "must successfully complete a background investigation",
    "public trust - background investigation",
    "all new hires will be required",
    "complete a declaration for federal employment",
    "have your salary sent",
    "if you are a male applicant",
    "go through a personal identity verification",
    "may be required obtain and use",
    "may be required file a confidential financial disclosure report",
    "may be required undergo an income tax verification",
    "treasury, departmental offices",
    "the u.s. department of the treasury has",
)

FEDERAL_COMPETENCY_PREFIXES = (
    "attention to detail:",
    "attention to detail -",
    "customer service:",
    "customer service -",
    "decision making:",
    "decision making -",
    "information management:",
    "information management -",
    "interpersonal skills:",
    "interpersonal skills -",
    "oral communication:",
    "oral communication -",
    "problem solving:",
    "problem solving -",
    "teamwork:",
    "teamwork -",
    "technical competence:",
    "technical competence -",
    "technical competence â€“",
)

FEDERAL_COMPETENCY_NAMES = (
    "Attention to Detail",
    "Customer Service",
    "Decision Making",
    "Information Management",
    "Interpersonal Skills",
    "Oral Communication",
    "Problem Solving",
    "Teamwork",
    "Technical Competence",
)

FEDERAL_COMPETENCY_NAME_RE = re.compile(
    r"^("
    + "|".join(re.escape(name) for name in FEDERAL_COMPETENCY_NAMES)
    + r")\s*(?:[-:\u2013\u2014]|â€“)\s*(.+)$",
    re.I,
)

FEDERAL_SPECIALIZED_DUTY_MARKERS: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "Duty 1",
        (
            "cross-functional project",
            "agile",
            "risk",
            "issues",
            "contracts",
            "change requests",
            "schedule",
            "budget",
            "organizational policies",
            "incremental software",
            "system enhancements",
            "statements of work",
            "functional requirements documents",
            "vendor agreements",
            "milestone schedules",
            "scope baselines",
        ),
    ),
    (
        "Duty 2",
        (
            "technical integration",
            "system modernization",
            "continuous improvement",
            "architecture",
            "cloud",
            "system interfaces",
            "cybersecurity",
            "technical debt",
            "release predictability",
            "zero-trust",
            "least-privilege",
            "access controls",
            "incident-response readiness",
            "etl validation",
            "user access reviews",
            "unplanned downtime",
            "inventory scrap",
        ),
    ),
    (
        "Duty 3",
        (
            "risk and issue",
            "risk assumptions",
            "blockers",
            "escalated",
            "change-control",
            "agile release",
            "program timelines",
            "deliverables",
            "internal control testing",
            "go-live readiness",
            "recovery testing",
            "defect triage",
            "validate resolutions",
        ),
    ),
)

FEDERAL_KSA_PREFIXES = (
    "expert knowledge",
    "mastery of",
    "expert written skill",
    "expert skill in",
    "strong interpersonal skills",
)

FEDERAL_KSA_STOP_PHRASES = (
    "demonstration of these ksas",
    "your possession of the ksas",
    "best qualified",
    "well qualified",
    "qualified =",
    "applicants applying for the gs-15 grade level must meet the following requirements",
)

FEDERAL_SOURCE_BRIDGE_REWRITES = {
    (
        "Accelerated documentation, reporting, SQL troubleshooting, and operational analysis through Codex, "
        "Claude, and AI-assisted tools that reduced manual effort across recurring systems and implementation processes."
    ): (
        "Implemented Codex, Claude, and AI-assisted workflows for documentation, reporting, SQL troubleshooting, "
        "and operational analysis, reducing manual effort across recurring systems, testing, and implementation processes."
    ),
    (
        "Helped build a zero-to-one internal SMS support channel as a founding pilot team member, configuring "
        "LivePerson LiveEngage chat/text workflows, automated greetings and closings, and AI-assisted chatbot logic "
        "for a customer communication channel that did not previously exist."
    ): (
        "Implemented LivePerson LiveEngage chat and text workflows, automated greetings and closings, and AI-assisted "
        "chatbot logic as a founding pilot team member for a new internal SMS support channel operating in a live "
        "customer-support environment."
    ),
    (
        "Built reporting frameworks and operational dashboards using Salesforce CRM and BI tools to surface customer "
        "interaction trends, service gaps, and efficiency opportunities, enabling data-driven service strategy "
        "decisions for senior eCommerce operations leadership."
    ): (
        "Built Salesforce CRM reporting frameworks and BI dashboards that surfaced customer interaction trends, "
        "service gaps, and efficiency opportunities, enabling data-driven service strategy decisions for senior "
        "eCommerce operations leadership."
    ),
}

FEDERAL_REQUIREMENT_CLUSTERS: dict[str, dict[str, object]] = {
    "implementation_delivery": {
        "label": "implementation and deployment",
        "job_signals": (
            "implementation",
            "implement",
            "deploy",
            "deployment",
            "go-live",
            "cutover",
            "rollout",
            "launch",
            "test environment",
            "testing",
            "uat",
            "prototype",
            "production",
        ),
        "direct_terms": (
            "full-lifecycle",
            "implementation",
            "go-live",
            "cutover",
            "requirements analysis",
            "requirements definition",
            "system configuration",
            "user acceptance validation",
            "integration testing",
            "production deployment",
            "go-live readiness",
        ),
        "adjacent_terms": (
            "stabilized",
            "workflow",
            "channel adoption",
            "pilot",
            "operational analysis",
        ),
        "summary_phrase": "full-lifecycle implementation, testing, deployment, and go-live coordination",
        "qualification_phrase": "full-lifecycle implementation delivery, testing, deployment, and go-live coordination",
        "skill_patterns": ("azure devops", "jira", "ms project"),
        "keyword_terms": ("implementation", "deployment", "testing", "go-live readiness"),
    },
    "data_migration": {
        "label": "data migration and validation",
        "job_signals": ("data migration", "etl", "data conversion", "validation", "integration"),
        "direct_terms": (
            "data extraction",
            "transformation",
            "etl",
            "data migration",
            "validation",
            "sql",
            "integrations",
            "conversion",
        ),
        "adjacent_terms": ("reporting", "crystal reports", "power bi"),
        "summary_phrase": "data migration, ETL validation, SQL reporting, and integration support",
        "qualification_phrase": "data migration planning, ETL validation, SQL-based reporting, and integration support",
        "skill_patterns": ("sql", "power bi", "power query", "etl"),
        "keyword_terms": ("data migration", "etl validation", "integration", "sql"),
    },
    "executive_alignment": {
        "label": "executive stakeholder alignment",
        "job_signals": ("executive", "senior executive", "director", "briefing", "stakeholder", "advisory"),
        "direct_terms": (
            "vp- and director-level",
            "director-level",
            "executive workshops",
            "quarterly business reviews",
            "qbrs",
            "stakeholders",
            "technical advisory",
            "investment decisions",
        ),
        "adjacent_terms": ("business outcome language", "customer-facing", "recommendations"),
        "summary_phrase": "VP-, director-, and C-suite stakeholder alignment on technology priorities and operating tradeoffs",
        "qualification_phrase": "VP-, director-, and C-suite alignment on technology priorities, progress, and business tradeoffs",
        "skill_patterns": ("ms project", "ms visio"),
        "keyword_terms": ("executive briefings", "stakeholder alignment", "technical advisory"),
    },
    "governance_risk": {
        "label": "governance, security, and risk controls",
        "job_signals": (
            "risk",
            "control",
            "governance",
            "audit",
            "cybersecurity",
            "security",
            "devsecops",
            "compliance",
            "access control",
        ),
        "direct_terms": (
            "least-privilege",
            "role-based access",
            "access controls",
            "audit readiness",
            "internal control testing",
            "permission frameworks",
            "risk documentation",
            "security and access validation",
            "compliance posture",
        ),
        "adjacent_terms": ("disaster recovery", "operational risks", "risk assumptions", "secure deployment"),
        "summary_phrase": "access-control governance, audit-readiness controls, and risk-aware delivery support",
        "qualification_phrase": "access-control governance, audit-readiness controls, and risk-aware implementation support",
        "skill_patterns": ("servicenow", "splunk"),
        "keyword_terms": ("governance", "risk management", "access controls", "audit readiness"),
    },
    "reporting_analytics": {
        "label": "reporting, analytics, and decision support",
        "job_signals": ("analytics", "dashboard", "reporting", "power bi", "business intelligence", "kpi", "analysis"),
        "direct_terms": (
            "dashboard",
            "business intelligence",
            "reporting",
            "sql-based",
            "power bi",
            "crystal reports",
            "forecast",
            "operational dashboards",
        ),
        "adjacent_terms": ("service gaps", "pipeline demand", "decision support"),
        "summary_phrase": "SQL-based reporting, KPI dashboards, and decision-support analytics",
        "qualification_phrase": "SQL-based reporting, KPI dashboards, operational analytics, and decision-support tooling",
        "skill_patterns": ("sql", "power bi", "crystal reports", "excel power query"),
        "keyword_terms": ("analytics", "reporting", "dashboards", "decision support"),
    },
    "change_adoption": {
        "label": "change adoption and training",
        "job_signals": ("change", "adoption", "training", "enablement", "communications", "readiness", "onboarding"),
        "direct_terms": (
            "training programs",
            "onboarding materials",
            "change communications",
            "adoption confidence",
            "enablement resources",
            "training documentation",
            "onboarding",
        ),
        "adjacent_terms": ("reduced post-go-live support volume", "adoption guides", "customer-facing success communications"),
        "summary_phrase": "training, onboarding, enablement, and adoption support across business and technical teams",
        "qualification_phrase": "training, onboarding, enablement, and adoption support across business and technical stakeholders",
        "skill_patterns": ("ms project", "ms visio"),
        "keyword_terms": ("change management", "adoption", "training", "readiness"),
    },
    "acquisition_support": {
        "label": "acquisition and procurement support",
        "job_signals": ("acquisition", "procurement", "contract", "vendor", "statement of work", "sow", "frd"),
        "direct_terms": (
            "statements of work",
            "functional requirements documents",
            "vendor agreements",
            "scope baselines",
            "milestone schedules",
            "contract execution",
            "pricing",
        ),
        "adjacent_terms": ("vendor evaluation", "technology acquisitions"),
        "summary_phrase": "requirements-ready documentation, vendor evaluation, and acquisition support",
        "qualification_phrase": "requirements-ready documentation, vendor evaluation, and acquisition support for technology initiatives",
        "skill_patterns": ("ms project", "jira"),
        "keyword_terms": ("acquisition support", "vendor evaluation", "requirements documentation"),
    },
    "customer_service_delivery": {
        "label": "customer and service delivery",
        "job_signals": ("customer", "client", "service", "support", "case", "account", "user"),
        "direct_terms": (
            "customer-facing",
            "client accounts",
            "support channel",
            "customer communication",
            "service strategy",
            "support volume",
            "case ownership",
            "customer service data",
        ),
        "adjacent_terms": ("b2b", "direct customer", "renewal confidence", "customer lifecycle"),
        "summary_phrase": "customer-facing delivery, service quality improvement, and user support operations",
        "qualification_phrase": "customer-facing delivery, service quality improvement, and user support operations",
        "skill_patterns": ("zendesk", "salesforce crm", "servicenow"),
        "keyword_terms": ("customer service", "support operations", "service delivery"),
    },
    "operations_improvement": {
        "label": "operations and process improvement",
        "job_signals": ("operations", "process", "efficiency", "workflow", "improvement", "root cause", "operational"),
        "direct_terms": (
            "root-cause analysis",
            "automated corrective workflow",
            "operational analysis",
            "efficiency opportunities",
            "processing",
            "workflow patterns",
            "operational risks",
        ),
        "adjacent_terms": ("inventory processing", "service quality", "replication"),
        "summary_phrase": "workflow improvement, root-cause analysis, and measurable operational efficiency gains",
        "qualification_phrase": "workflow improvement, root-cause analysis, and measurable operational efficiency gains",
        "skill_patterns": ("excel power query", "power bi"),
        "keyword_terms": ("process improvement", "workflow", "operational efficiency"),
    },
    "ai_workflow": {
        "label": "AI workflow support",
        "job_signals": (
            "ai",
            "artificial intelligence",
            "ai-enabled",
            "chatbot",
            "conversational ai",
            "nlp",
            "llm",
            "generative ai",
        ),
        "direct_terms": (
            "codex",
            "claude",
            "ai-assisted",
            "liveperson",
            "liveengage",
            "chatbot",
            "conversational ai",
            "nlp",
            "workflow automation",
        ),
        "adjacent_terms": ("automated greetings", "automated closings", "text workflows", "operational analysis"),
        "summary_phrase": "AI-assisted workflow support, conversational AI configuration, and automation-ready documentation",
        "qualification_phrase": "AI-assisted workflow support, conversational AI configuration, and automation-ready documentation",
        "skill_patterns": ("codex", "claude", "workflow automation", "generative ai"),
        "keyword_terms": ("AI-enabled systems", "conversational AI", "workflow automation"),
    },
    "cloud_modernization": {
        "label": "modernization and migration support",
        "job_signals": ("cloud-native", "modernization", "modernize", "migration", "platform transition", "saas"),
        "direct_terms": ("system modernization", "platform migration", "final cutover", "go-live readiness"),
        "adjacent_terms": ("enterprise platform", "concurrent workstreams", "cross-site cutover", "stabilized"),
        "summary_phrase": "platform modernization planning, migration support, and cross-functional cutover coordination",
        "qualification_phrase": "platform modernization planning, migration support, and cross-functional cutover coordination",
        "skill_patterns": ("azure devops", "jira"),
        "keyword_terms": ("modernization", "migration", "cross-functional delivery"),
    },
    "agile_delivery": {
        "label": "Agile and cross-functional delivery",
        "job_signals": ("agile", "sprint", "scrum", "backlog", "iteration", "product owner", "product manager", "roadmap", "user research"),
        "direct_terms": (
            "azure devops",
            "jira",
            "product management and product ownership",
            "agile development team",
            "de facto product owner",
            "backlog-ready requirements",
        ),
        "adjacent_terms": (
            "concurrent workstreams",
            "triage defects",
            "validate resolutions",
            "milestone schedules",
            "roadmap",
            "stakeholder requirements gathering",
            "sprint priorities",
            "feedback sessions with end users",
        ),
        "summary_phrase": "Agile delivery partnership with Product Management and Product Ownership, backlog-ready requirements, and cross-functional coordination",
        "qualification_phrase": "Agile delivery partnership with Product Management and Product Ownership, backlog-ready requirements, and cross-functional coordination",
        "skill_patterns": ("azure devops", "jira"),
        "keyword_terms": ("agile delivery", "product ownership", "cross-functional coordination", "issue triage"),
    },
    "technical_support": {
        "label": "technical support and systems administration",
        "job_signals": ("support", "systems administration", "active directory", "windows services", "troubleshooting"),
        "direct_terms": (
            "technical product support",
            "systems administration",
            "active directory",
            "windows services",
            "troubleshooting",
            "issue triage",
            "enterprise application access",
        ),
        "adjacent_terms": ("service issues", "customer-impacting technical workflows"),
        "summary_phrase": "technical support, systems administration, and issue-resolution coordination",
        "qualification_phrase": "technical support, systems administration, and issue-resolution coordination",
        "skill_patterns": ("servicenow", "zendesk", "splunk"),
        "keyword_terms": ("technical support", "systems administration", "troubleshooting"),
    },
}

FEDERAL_DEFAULT_CLUSTERS_BY_LANE = {
    "implementation_delivery": ("implementation_delivery", "data_migration", "executive_alignment", "governance_risk"),
    "customer_success": ("customer_service_delivery", "implementation_delivery", "change_adoption", "executive_alignment"),
    "presales_solution": ("executive_alignment", "implementation_delivery", "acquisition_support", "customer_service_delivery"),
    "analytics_operations": ("reporting_analytics", "operations_improvement", "executive_alignment", "governance_risk"),
    "change_enablement": ("change_adoption", "implementation_delivery", "executive_alignment", "customer_service_delivery"),
    "process_improvement": ("operations_improvement", "reporting_analytics", "implementation_delivery", "governance_risk"),
    "corporate_strategy": ("executive_alignment", "governance_risk", "reporting_analytics", "implementation_delivery"),
}


def require_file(path: Path, label: str) -> None:
    if not path.is_file():
        fail(f"{label} not found: {path}")


def validate_inputs(job_description_text: str | None = None) -> str:
    require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    require_file(FEDERAL_RESUME_SOURCE, "Federal resume source")
    require_file(FEDERAL_ESSAY_SOURCE, "Federal essay source")
    if job_description_text is None:
        require_file(FEDERAL_JOB_DESCRIPTION, "Federal job description")

    raw_job_description = read_text(FEDERAL_JOB_DESCRIPTION) if job_description_text is None else job_description_text.strip()
    if not raw_job_description:
        fail("jobs/federal_job_description.txt is empty; refusing to create a placeholder or partial federal resume")
    warning_count = _check_job_description_quality(raw_job_description)
    print(f"Federal JD quality check: {warning_count} warning(s).")
    job_description, removed_lines = strip_linkedin_job_board_boilerplate(raw_job_description)
    if removed_lines:
        print(f"Federal JD NOTE: removed {removed_lines} LinkedIn/job-board boilerplate line(s) before targeting.")
    return job_description


def load_federal_source() -> FederalSource:
    require_file(FEDERAL_RESUME_SOURCE, "Federal resume source")
    payload = json.loads(FEDERAL_RESUME_SOURCE.read_text(encoding="utf-8"))
    contact_payload = payload["contact"]
    contact = FederalContact(
        name=contact_payload["name"],
        location=contact_payload["location"],
        email=contact_payload["email"],
        phone=contact_payload["phone"],
        citizenship=contact_payload["citizenship"],
        veterans_preference=contact_payload["veterans_preference"],
        clearance=contact_payload["clearance"],
        availability=contact_payload["availability"],
    )
    confirmed_by_role = evidence_engine.confirmed_bullets_by_role(FEDERAL_RESUME_SOURCE)
    roles = tuple(
        FederalRole(
            title=role["title"],
            start=role["start"],
            end=role["end"],
            hours_per_week=role["hours_per_week"],
            salary=role["salary"],
            company=role["company"],
            location=role["location"],
            supervisor=role["supervisor"],
            supervisor_phone=role["supervisor_phone"],
            company_summary=role["company_summary"],
            job_summary=role["job_summary"],
            bullets=tuple(dict.fromkeys((*role["bullets"], *confirmed_by_role.get((role["company"], role["title"]), ())))),
        )
        for role in payload["roles"]
    )
    education = tuple(
        FederalEducation(
            degree=item["degree"],
            date=item["date"],
            school=item["school"],
            details=item["details"],
        )
        for item in payload["education"]
    )
    volunteer_experience = tuple(
        FederalVolunteerExperience(
            title=item["title"],
            organization=item["organization"],
            description=item["description"],
        )
        for item in payload.get("volunteer_experience", ())
    )
    source = FederalSource(
        contact=contact,
        technical_skills=tuple(payload["technical_skills"]),
        roles=roles,
        education=education,
        professional_development=tuple(payload["professional_development"]),
        volunteer_experience=volunteer_experience,
    )
    validate_federal_source(source)
    return source


def load_federal_standard_essays() -> tuple[FederalEssayPrompt, ...]:
    require_file(FEDERAL_ESSAY_SOURCE, "Federal essay source")
    payload = json.loads(FEDERAL_ESSAY_SOURCE.read_text(encoding="utf-8"))
    prompts = tuple(
        FederalEssayPrompt(
            key=item["key"],
            question=normalize_spaces(item["question"]),
            answer=normalize_spaces(item["answer"]),
        )
        for item in payload
    )
    if len(prompts) < 4:
        fail("Federal essay source must include the full standard essay-question bank.")
    required_keys = {
        "constitutional_commitment",
        "government_efficiency",
        "executive_orders",
        "work_ethic",
    }
    found_keys = {prompt.key for prompt in prompts}
    missing = sorted(required_keys - found_keys)
    if missing:
        fail(f"Federal essay source is missing expected prompt keys: {', '.join(missing)}")
    return prompts


FEDERAL_MONTH_YEAR_PATTERN = re.compile(r"^[A-Za-z]+\.?\s+\d{4}$")
FEDERAL_MONTHS = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}


def federal_month_ordinal(value: str, *, present_as: int | None = None) -> int | None:
    normalized = value.strip().lower().replace(".", "")
    if normalized == "present":
        return present_as
    parts = normalized.split()
    if len(parts) != 2 or not parts[1].isdigit():
        return None
    month = FEDERAL_MONTHS.get(parts[0])
    if month is None:
        return None
    return int(parts[1]) * 12 + month


def federal_roles_overlap(left: FederalRole, right: FederalRole) -> bool:
    left_start = federal_month_ordinal(left.start)
    left_end = federal_month_ordinal(left.end, present_as=9999 * 12)
    right_start = federal_month_ordinal(right.start)
    right_end = federal_month_ordinal(right.end, present_as=9999 * 12)
    if None in (left_start, left_end, right_start, right_end):
        return False
    return left_start <= right_end and right_start <= left_end


def federal_salary_value(role: FederalRole, roles: Sequence[FederalRole]) -> str:
    if SALARY_FIGURE_PATTERN.search(role.salary or ""):
        return role.salary.strip()
    for candidate in roles:
        if candidate is role:
            continue
        if candidate.company != role.company:
            continue
        if not federal_roles_overlap(role, candidate):
            continue
        if SALARY_FIGURE_PATTERN.search(candidate.salary or ""):
            return candidate.salary.strip()
    return role.salary.strip()


def validate_federal_source(source: FederalSource) -> None:
    # USAJobs requires every position to fully document job title, duties,
    # month/year start and end dates, AND hours worked per week. A full year
    # of qualifying experience is credited at 35-40 hours per week; part-time
    # roles are credited on actual time spent, but only if hours per week and
    # duties are clearly stated. Missing or malformed data here risks the
    # whole application being scored as not meeting time-in-grade or
    # specialized-experience requirements. See USAJobs "How You Will Be
    # Evaluated" language for the exact wording this enforces.
    if len(source.roles) < 4:
        fail("Federal resume source must include at least four work-experience roles.")
    if not source.technical_skills:
        fail("Federal resume source must include at least one technical skill.")
    for role in source.roles:
        role_label = f"{role.title or '(missing title)'} at {role.company or '(missing company)'}"
        if not role.title.strip():
            fail(f"Federal role is missing a job title: {role_label}")
        if not role.bullets:
            fail(f"Federal role is missing bullets (duties): {role_label}")
        if not role.company_summary:
            fail(f"Federal role is missing a company summary: {role_label}")
        if not role.job_summary:
            fail(f"Federal role is missing a job summary (duties): {role_label}")
        if not role.hours_per_week.strip() or not role.hours_per_week.strip().isdigit():
            fail(
                "Federal role is missing hours worked per week, or it is not a plain number: "
                f"{role_label} has hours_per_week='{role.hours_per_week}'. USAJobs requires hours "
                "per week on every position to credit experience toward a grade level."
            )
        for date_label, date_value in (("start", role.start), ("end", role.end)):
            normalized_date = date_value.strip()
            if normalized_date.lower() == "present":
                continue
            if not FEDERAL_MONTH_YEAR_PATTERN.match(normalized_date):
                fail(
                    f"Federal role {date_label} date must be month and year (e.g. 'March 2023'): "
                    f"{role_label} has {date_label}='{date_value}'."
                )


def extract_federal_agency_name(job_description: str) -> str | None:
    return resume_analysis.extract_federal_agency_name(job_description)


def extract_federal_role_title(job_description: str) -> str | None:
    return resume_analysis.extract_federal_official_title(job_description)


def extract_federal_output_name(job_description: str) -> str:
    return resume_analysis.extract_target_output_label(job_description, workflow="federal")


def source_visible_text(source: FederalSource) -> str:
    parts: list[str] = []
    parts.extend(source.technical_skills)
    for role in source.roles:
        parts.extend((role.company_summary, role.job_summary, *role.bullets))
    for item in source.education:
        parts.extend((item.degree, item.school, item.details))
    parts.extend(source.professional_development)
    for item in source.volunteer_experience:
        parts.extend((item.title, item.organization, item.description))
    return "\n".join(parts)


def source_work_text(source: FederalSource) -> str:
    parts: list[str] = []
    for role in source.roles:
        parts.extend((role.company_summary, role.job_summary, *role.bullets))
    return "\n".join(parts)


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def dedupe_preserve_order(items: list[str] | tuple[str, ...]) -> tuple[str, ...]:
    ordered: list[str] = []
    seen: set[str] = set()
    for item in items:
        cleaned = normalize_spaces(item)
        if not cleaned:
            continue
        lowered = cleaned.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        ordered.append(cleaned)
    return tuple(ordered)


def join_phrases(items: tuple[str, ...] | list[str]) -> str:
    cleaned = [item.strip() for item in items if item and item.strip()]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    if len(cleaned) == 2:
        return f"{cleaned[0]} and {cleaned[1]}"
    return f"{', '.join(cleaned[:-1])}, and {cleaned[-1]}"


def title_case_lane_text(profile_key: str) -> str:
    return FEDERAL_LANE_OPENERS.get(profile_key, "enterprise IT, systems, and program delivery")


def federal_ai_focus(job_description: str) -> bool:
    lowered = job_description.lower()
    return any(phrase_present(lowered, signal) for signal in FEDERAL_REQUIREMENT_CLUSTERS["ai_workflow"]["job_signals"])


def supported_ai_summary_fragments(source: FederalSource) -> tuple[str, ...]:
    visible = source_work_text(source).lower()
    fragments: list[str] = []
    if "codex" in visible and "claude" in visible:
        fragments.append("Codex and Claude workflow support")
    elif "codex" in visible:
        fragments.append("Codex-assisted workflow support")
    elif "claude" in visible:
        fragments.append("Claude-assisted workflow support")
    if "ai-assisted analysis" in visible or "ai-assisted tools" in visible:
        fragments.append("AI-assisted analysis")
    if "liveperson" in visible and "chatbot" in visible:
        fragments.append("LivePerson chatbot workflow configuration")
    elif "conversational ai" in visible:
        fragments.append("conversational AI workflow support")
    if not fragments:
        fragments.append("AI-assisted workflow support")
    return tuple(fragments[:3])


def phrase_present(text: str, phrase: str) -> bool:
    cleaned = phrase.strip().lower()
    if not cleaned:
        return False
    if re.fullmatch(r"[a-z0-9]+", cleaned):
        return re.search(rf"\b{re.escape(cleaned)}\b", text, re.I) is not None
    return cleaned in text.lower()


def count_phrase_hits(text: str, phrases: tuple[str, ...]) -> int:
    lowered = text.lower()
    return sum(1 for phrase in phrases if phrase_present(lowered, phrase))


def strip_bullet_prefix(text: str) -> str:
    return re.sub(r"^\s*(?:[-*•]+|\d+[.)])\s*", "", text).strip()


def visible_job_lines(job_description: str) -> list[str]:
    lines: list[str] = []
    for raw_line in job_description.splitlines():
        line = normalize_spaces(strip_bullet_prefix(raw_line))
        if not line:
            continue
        if re.fullmatch(r"[-]{5,}", line):
            continue
        lowered = line.lower().rstrip(":")
        if lowered in FEDERAL_OUTPUT_SKIP_LINES or lowered in FEDERAL_IGNORE_SECTION_HEADINGS:
            continue
        lines.append(line)
    return lines


def job_paragraphs(job_description: str) -> tuple[str, ...]:
    paragraphs: list[str] = []
    for block in re.split(r"(?:\r?\n\s*){2,}", job_description):
        cleaned = normalize_spaces(strip_bullet_prefix(block))
        if cleaned:
            paragraphs.append(cleaned)
    return dedupe_preserve_order(paragraphs)


def extract_application_char_limit(job_description: str) -> int:
    patterns = (
        r"(?i)limit your answer to\s+(\d{3,5})\s+characters",
        r"(?i)response will be limited to\s+(\d{3,5})\s+characters",
        r"(?i)(\d{3,5})\s+characters left",
    )
    for pattern in patterns:
        match = re.search(pattern, job_description)
        if match:
            return int(match.group(1))
    return DEFAULT_QUESTION_CHAR_LIMIT


def extract_it_competencies(job_description: str) -> tuple[FederalCompetencyPrompt, ...]:
    competencies: list[FederalCompetencyPrompt] = []
    seen: set[str] = set()
    canonical_names = {name.lower(): name for name in FEDERAL_COMPETENCY_NAMES}
    for raw_line in job_description.splitlines():
        line = normalize_spaces(strip_bullet_prefix(raw_line))
        if not line:
            continue
        match = FEDERAL_COMPETENCY_NAME_RE.match(line)
        if not match:
            continue
        name = canonical_names.get(match.group(1).lower(), normalize_spaces(match.group(1)))
        lowered = name.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        competencies.append(FederalCompetencyPrompt(name=name, description=normalize_spaces(match.group(2))))
    return tuple(competencies)


def extract_specialized_experience_items(job_description: str) -> tuple[str, ...]:
    items: list[str] = []
    collecting = False
    for raw_line in job_description.splitlines():
        line = normalize_spaces(strip_bullet_prefix(raw_line))
        if not line:
            continue
        lowered = line.lower()
        if lowered.startswith("specialized experience requirements"):
            collecting = True
            continue
        if collecting and lowered.startswith(("how you will be evaluated", "demonstration of these ksas", "your possession of the ksas", "best qualified", "well qualified", "qualified =")):
            break
        if not collecting:
            continue
        if lowered.startswith("experience "):
            items.append(line.rstrip("."))
    return dedupe_preserve_order(items)


def short_ksa_name(text: str) -> str:
    lowered = text.lower()
    split_tokens = (
        " sufficient to ",
        " serving as ",
        " are necessary to ",
        " thereby enabling ",
        " including formal strategic white papers",
    )
    for token in split_tokens:
        index = lowered.find(token)
        if index > 0:
            return text[:index].strip(" .:;,")
    return re.split(r"[.;]", text, maxsplit=1)[0].strip(" .:;,")


def extract_ksa_items(job_description: str, default_clusters: tuple[str, ...]) -> tuple[FederalKSAItem, ...]:
    items: list[FederalKSAItem] = []
    collecting = False
    seen: set[str] = set()
    for raw_line in job_description.splitlines():
        paragraph = normalize_spaces(strip_bullet_prefix(raw_line))
        if not paragraph:
            continue
        lowered = paragraph.lower()
        if "knowledge, skills, and abilities" in lowered and "assign a score" in lowered:
            collecting = True
            continue
        if not collecting and lowered.startswith("a panel of subject matter experts"):
            collecting = True
            continue
        if not collecting:
            continue
        if any(lowered.startswith(prefix) for prefix in FEDERAL_KSA_STOP_PHRASES):
            break
        if not any(lowered.startswith(prefix) for prefix in FEDERAL_KSA_PREFIXES):
            continue
        name = short_ksa_name(paragraph)
        key = name.lower()
        if key in seen:
            continue
        seen.add(key)
        items.append(
            FederalKSAItem(
                name=name,
                text=paragraph,
                clusters=infer_clusters(paragraph, default_clusters),
            )
        )
    return tuple(items)


def federal_application_questions(job_description: str, profile_name: str) -> FederalApplicationQuestionSet:
    default_clusters = FEDERAL_DEFAULT_CLUSTERS_BY_LANE.get(profile_name, ("implementation_delivery",))
    return FederalApplicationQuestionSet(
        char_limit=extract_application_char_limit(job_description),
        competencies=extract_it_competencies(job_description),
        specialized_items=extract_specialized_experience_items(job_description),
        ksas=extract_ksa_items(job_description, default_clusters),
    )


def requirement_priority(kind: str, label: str) -> int:
    if kind == "selective_factor":
        return 100
    if kind == "specialized_experience":
        return 96
    if kind == "core_experience":
        return 69
    if kind == "gs_level":
        match = re.search(r"gs[- ]?(\d{1,2})", label.lower())
        if match:
            return 72 + int(match.group(1))
        return 88
    if kind == "basic_qualifications":
        return 86
    if kind == "qualifications":
        return 82
    if kind == "evaluation":
        return 78
    if kind == "duties":
        return 74
    return 68


def inline_bucket_from_line(line: str) -> tuple[str, str, str] | None:
    normalized = line.strip()
    lowered = normalized.lower()
    gs_match = re.match(r"^(gs[- ]?\d{1,2})\s*[:\-]?\s*(.*)$", lowered, re.I)
    if gs_match:
        label = gs_match.group(1).upper().replace(" ", "-")
        return label, "gs_level", normalized[len(gs_match.group(1)):].strip(" :-")
    gs_sentence_match = re.match(r"^for the (gs[- ]?\d{1,2})\b[,\s]*(.*)$", lowered, re.I)
    if gs_sentence_match:
        label = gs_sentence_match.group(1).upper().replace(" ", "-")
        return label, "gs_level", normalized
    for kind, pattern in FEDERAL_SECTION_KIND_PATTERNS:
        match = re.match(pattern, lowered, re.I)
        if not match:
            continue
        label = {
            "selective_factor": "Selective Factor",
            "specialized_experience": "Specialized Experience",
            "basic_qualifications": "Basic Qualifications",
            "qualifications": "Qualifications",
            "evaluation": "How You Will Be Evaluated",
            "duties": "Duties",
        }[kind]
        trailing = match.group(match.lastindex or 1).strip(" :-") if match.lastindex else ""
        return label, kind, trailing
    return None


def infer_clusters(text: str, default_clusters: tuple[str, ...]) -> tuple[str, ...]:
    lowered = text.lower()
    scored: list[tuple[int, str]] = []
    for cluster_key, metadata in FEDERAL_REQUIREMENT_CLUSTERS.items():
        score = count_phrase_hits(lowered, metadata["job_signals"])
        score += count_phrase_hits(lowered, metadata.get("keyword_terms", ()))
        if cluster_key in default_clusters:
            score += 1
        if score > 0:
            scored.append((score, cluster_key))
    scored.sort(key=lambda item: (-item[0], item[1]))
    if not scored:
        return default_clusters[:3] if default_clusters else ("implementation_delivery",)
    ordered: list[str] = []
    for _score, cluster_key in scored:
        if cluster_key not in ordered:
            ordered.append(cluster_key)
    return tuple(ordered[:4])


def parse_requirement_buckets(
    job_description: str,
    profile_name: str,
    *,
    selected_grade: str = "",
) -> tuple[FederalRequirementBucket, ...]:
    default_clusters = FEDERAL_DEFAULT_CLUSTERS_BY_LANE.get(profile_name, ("implementation_delivery",))
    sections = requirement_engine.parse_federal_requirement_sections(job_description)
    buckets: list[FederalRequirementBucket] = [
        FederalRequirementBucket(
            label=section.label,
            kind=section.kind,
            text=section.text,
            priority=requirement_priority(section.kind, section.label),
            clusters=infer_clusters(section.text, default_clusters),
            weight_group_id=section.weight_group_id,
        )
        for section in sections
        if section.kind != "specialized_experience" or not selected_grade or section.grade == selected_grade
    ]
    if not buckets:
        fallback_text = "Relevant federal IT delivery, systems, and stakeholder-facing experience."
        buckets = [
            FederalRequirementBucket(
                label="Core Experience",
                kind="core_experience",
                text=fallback_text,
                priority=69,
                clusters=infer_clusters(fallback_text, default_clusters),
                weight_group_id="core_experience",
            )
        ]
    deduped: list[FederalRequirementBucket] = []
    seen: set[tuple[str, str]] = set()
    for bucket in buckets:
        key = (bucket.label.lower(), bucket.text.lower())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(bucket)
    return tuple(deduped)


def bullet_reference(source: FederalSource, role_index: int, bullet_index: int) -> str:
    return f"{source.roles[role_index].company} [{bullet_index + 1}]"


def cluster_alignment(source: FederalSource, cluster_key: str) -> str:
    metadata = FEDERAL_REQUIREMENT_CLUSTERS[cluster_key]
    work_text = source_work_text(source).lower()
    full_text = source_visible_text(source).lower()
    direct_hits = count_phrase_hits(work_text, metadata["direct_terms"])
    adjacent_hits = count_phrase_hits(work_text, metadata["adjacent_terms"])
    adjacent_hits += count_phrase_hits(full_text, metadata["skill_patterns"])
    if direct_hits > 0:
        return "Direct"
    if adjacent_hits > 0:
        return "Adjacent"
    return "Unsupported"


def candidate_cluster_strength(text: str, cluster_key: str) -> int:
    metadata = FEDERAL_REQUIREMENT_CLUSTERS[cluster_key]
    lowered = text.lower()
    direct_hits = count_phrase_hits(lowered, metadata["direct_terms"])
    if direct_hits:
        return 2
    adjacent_hits = count_phrase_hits(lowered, metadata["adjacent_terms"])
    if adjacent_hits:
        return 1
    return 0


def candidate_cluster_units(text: str, cluster_key: str) -> int:
    metadata = FEDERAL_REQUIREMENT_CLUSTERS[cluster_key]
    lowered = text.lower()
    direct_hits = count_phrase_hits(lowered, metadata["direct_terms"])
    if direct_hits:
        return direct_hits * 2
    adjacent_hits = count_phrase_hits(lowered, metadata["adjacent_terms"])
    return adjacent_hits


def cluster_source_refs(source: FederalSource, cluster_key: str, *, limit: int = 3) -> tuple[str, ...]:
    scored: list[tuple[int, str]] = []
    for role_index, role in enumerate(source.roles):
        for bullet_index, bullet in enumerate(role.bullets):
            units = candidate_cluster_units(bullet, cluster_key)
            if not units:
                continue
            score = units * 10 + max(0, 5 - role_index)
            scored.append((score, bullet_reference(source, role_index, bullet_index)))
    scored.sort(key=lambda item: (-item[0], item[1]))
    refs: list[str] = []
    for _score, ref in scored:
        if ref not in refs:
            refs.append(ref)
        if len(refs) >= limit:
            break
    return tuple(refs)


def coverage_alignment(cluster_alignments: tuple[str, ...]) -> str:
    if cluster_alignments and all(item == "Direct" for item in cluster_alignments):
        return "Direct"
    if any(item in {"Direct", "Adjacent"} for item in cluster_alignments):
        return "Adjacent"
    return "Unsupported"


def bucket_narrative(bucket: FederalRequirementBucket) -> str:
    phrases = [
        FEDERAL_REQUIREMENT_CLUSTERS[cluster]["qualification_phrase"]
        for cluster in bucket.clusters
        if cluster in FEDERAL_REQUIREMENT_CLUSTERS
    ]
    return join_phrases(phrases[:3]) or bucket.text


def federal_keyword_targets(job_description: str, buckets: tuple[FederalRequirementBucket, ...], cluster_weights: tuple[tuple[str, int], ...]) -> tuple[str, ...]:
    targets: list[str] = []
    for keyword in audit_keywords(job_description):
        cleaned = normalize_spaces(keyword).strip(" .")
        lowered = cleaned.lower()
        if not cleaned:
            continue
        if lowered in FEDERAL_KEYWORD_STOP_PHRASES:
            continue
        if len(cleaned.split()) == 1 and len(cleaned) < 5:
            continue
        if len(cleaned.split()) == 1 and lowered not in {"cloud-native", "agile", "governance"}:
            continue
        if cleaned and cleaned.lower() not in {item.lower() for item in targets}:
            targets.append(cleaned)
    for cluster_key, _weight in cluster_weights:
        metadata = FEDERAL_REQUIREMENT_CLUSTERS[cluster_key]
        for term in metadata["keyword_terms"]:
            if term.lower() not in {item.lower() for item in targets}:
                targets.append(term)
    for bucket in buckets:
        if bucket.kind in {"selective_factor", "specialized_experience"}:
            normalized = bucket.label.lower()
            if normalized not in {item.lower() for item in targets}:
                targets.append(bucket.label)
    return tuple(targets[:16])


def federal_requirement_audit(
    source: FederalSource,
    job_description: str,
    *,
    target_grade: str = "",
    target_context: requirement_engine.TargetContext | None = None,
) -> FederalAudit:
    profile = job_problem_profile(job_description, source_visible_text(source))
    target_context = target_context or requirement_engine.build_target_context(
        job_description,
        workflow="federal",
        target_grade=target_grade,
    )
    buckets = parse_requirement_buckets(
        job_description,
        profile.primary_lane,
        selected_grade=target_context.target_grade,
    )
    cluster_weight_map: dict[str, int] = {}
    coverages: list[FederalRequirementCoverage] = []
    initial_warnings: list[str] = []
    default_clusters = FEDERAL_DEFAULT_CLUSTERS_BY_LANE.get(profile.primary_lane, ("implementation_delivery",))
    grouped_text: dict[str, list[str]] = {}
    for bucket in buckets:
        grouped_text.setdefault(bucket.weight_group_id or bucket.label, []).append(bucket.text)
    group_clusters = {
        group_id: infer_clusters(" ".join(texts), default_clusters)
        for group_id, texts in grouped_text.items()
    }

    warned_groups: set[str] = set()
    weighted_groups: set[tuple[str, str]] = set()
    for bucket in buckets:
        alignments = tuple(cluster_alignment(source, cluster) for cluster in bucket.clusters)
        alignment = coverage_alignment(alignments)
        refs: list[str] = []
        for cluster in bucket.clusters:
            for ref in cluster_source_refs(source, cluster):
                if ref not in refs:
                    refs.append(ref)
        weight_group = bucket.weight_group_id or bucket.label
        for cluster in group_clusters[weight_group]:
            weight_key = (weight_group, cluster)
            if weight_key not in weighted_groups:
                weighted_groups.add(weight_key)
                cluster_weight_map[cluster] = cluster_weight_map.get(cluster, 0) + bucket.priority
        warning_group = bucket.weight_group_id or bucket.label
        if (
            bucket.priority >= HIGH_PRIORITY_BUCKET_THRESHOLD
            and alignment == "Unsupported"
            and warning_group not in warned_groups
        ):
            warned_groups.add(warning_group)
            initial_warnings.append(f"{bucket.label} is not well supported by the current federal source.")
        coverages.append(
            FederalRequirementCoverage(
                bucket=bucket,
                alignment=alignment,
                source_refs=tuple(refs[:3]),
                selected_refs=(),
                narrative=bucket_narrative(bucket),
            )
        )

    # program_delivery deliberately has no lane default in this change.  The
    # fallback makes implementation_delivery the meaningful winner of a tie.
    default_rank = {cluster: index for index, cluster in enumerate(FEDERAL_DEFAULT_CLUSTERS_BY_LANE.get(profile.primary_lane, ("implementation_delivery",)))}
    sorted_cluster_weights = tuple(
        sorted(
            cluster_weight_map.items(),
            key=lambda item: (-item[1], default_rank.get(item[0], float("inf")), item[0]),
        )
    )
    keywords = federal_keyword_targets(job_description, buckets, sorted_cluster_weights)
    records = evidence_engine.load_evidence_catalog(FEDERAL_RESUME_SOURCE)
    element_matches = evidence_engine.match_requirements(target_context.requirements, records)
    unsupported_by_group: dict[str, list[requirement_engine.RequirementElement]] = {}
    for element, match in zip(target_context.requirements, element_matches):
        if match.status == requirement_engine.RequirementStatus.UNSUPPORTED:
            group_id = element.requirement_group_id or element.element_id
            unsupported_by_group.setdefault(group_id, []).append(element)
    for group_id, elements in unsupported_by_group.items():
        label = "Specialized Experience " + group_id.rsplit("_", 1)[-1] if group_id.startswith("specialized_experience_") else elements[0].section.replace("_", " ").title()
        examples = "; ".join(element.text[:120] for element in elements[:2])
        initial_warnings.append(
            f"Unsupported federal requirements in {label} ({len(elements)}): {examples}"
        )
    return FederalAudit(
        buckets=tuple(coverages),
        cluster_weights=sorted_cluster_weights,
        keyword_targets=keywords,
        warnings=tuple(initial_warnings),
        requirements=target_context.requirements,
        element_matches=element_matches,
        target_context=target_context,
    )


def trim_summary(summary: str) -> str:
    words = summary.split()
    if len(words) <= MAX_SUMMARY_WORDS:
        return summary
    trimmed = " ".join(words[:MAX_SUMMARY_WORDS]).rstrip(" ,;:")
    if not trimmed.endswith("."):
        trimmed += "."
    return trimmed


def summary_scope_phrase() -> str:
    return (
        "Scope includes five-site systems ownership for 150+ users, 80+ international client engagements, "
        "200+ SQL and business-intelligence reporting tools, 60+ executive workshops and QBRs, and measurable "
        "improvements in audit readiness, adoption, service quality, and operational efficiency."
    )


def print_prose_warnings(
    label: str,
    text: str,
    artifact: str,
    *,
    fail_on_findings: bool = False,
    summary_word_range: tuple[int, int] | None = None,
) -> None:
    enforce_prose_quality(
        text,
        artifact,
        label=label,
        mode="fail" if fail_on_findings else "warn",
        summary_word_range=summary_word_range,
    )


def federal_experience_year_range(source: FederalSource) -> str:
    years: list[int] = []
    for role in source.roles:
        for value in (role.start, role.end):
            match = re.search(r"\b(20\d{2}|19\d{2})\b", value)
            if match:
                years.append(int(match.group(1)))
    if not years:
        return "10+"
    earliest = min(years)
    elapsed = max(1, date.today().year - earliest)
    rounded = max(2, elapsed - (elapsed % 2))
    return f"{rounded}+"


def federal_functional_label(
    job_description: str,
    profile: JobProblemProfile,
) -> str:
    lowered = job_description.lower()
    if (
        profile.primary_lane == "analytics_operations"
        and re.search(r"\b(data analytics|business intelligence|reporting|dashboard)\b", lowered)
    ):
        return "Federal analytics and data systems"
    if (
        profile.primary_lane == "change_enablement"
        and re.search(r"\b(change management|organizational change|adoption)\b", lowered)
    ):
        return "Federal program and change management"
    opening_window = lowered[:200]
    if profile.primary_lane == "implementation_delivery" and "it" not in opening_window:
        return "Federal implementation and delivery"
    return "Federal IT and enterprise systems"


def federal_environment_list(source: FederalSource) -> str:
    labels: list[str] = []
    mapping = {
        "east west": "multi-site manufacturing",
        "aptean": "SaaS",
        "home depot": "eCommerce",
        "aderant": "legal technology",
    }
    for role in source.roles:
        company_lower = role.company.lower()
        for key, label in mapping.items():
            if key in company_lower and label not in labels:
                labels.append(label)
    if not labels:
        labels = ["multi-site manufacturing", "SaaS", "eCommerce", "legal technology"]
    return prose_engine.human_series(labels[:4])


def build_gs14_summary(source: FederalSource, job_description: str, audit: FederalAudit | None = None) -> str:
    profile = job_problem_profile(job_description, source_visible_text(source))
    active_audit = audit or federal_requirement_audit(source, job_description)
    sentence_one = prose_engine.sentence(
        f"Enterprise technology leader with {federal_experience_year_range(source)} years successfully modernizing complex enterprise systems "
        f"across {federal_environment_list(source)} environments"
    )
    sentence_two = prose_engine.sentence(
        "Recent operational scope includes platforms for five sites and 150+ users, 200+ SQL/BI reporting tools, "
        "80+ international client engagements, plus 60+ global executive sessions"
    )
    sentence_three = prose_engine.sentence(
        "Translates complex requirements into secure delivery decisions for VP- and director-level stakeholders, improving "
        "user adoption, service quality, reporting integrity, and operational continuity during high-stakes modernization work"
    )
    if federal_ai_focus(job_description):
        sentence_three = prose_engine.sentence(
            "Uses Codex and Claude for AI-assisted documentation, reporting, and analysis while translating complex requirements "
            "into secure delivery decisions for executive stakeholders across high-stakes implementation environments"
        )
    summary = normalize_spaces(f"{sentence_one} {sentence_two} {sentence_three}")
    summary = trim_summary(summary)
    word_count = len(re.findall(r"\b[\w+.#'-]+\b", summary))
    if word_count < MIN_SUMMARY_WORDS:
        fail("Federal Professional Summary generated below minimum word count.")
    repaired = prose_engine.repair_text(summary, "federal_summary")
    if not repaired.converged:
        fail("Federal Professional Summary failed shared structural validation: " + ", ".join(item.rule_id for item in repaired.findings))
    summary = repaired.text
    print_prose_warnings(
        "Federal Professional Summary",
        summary,
        "resume_summary",
        fail_on_findings=True,
        summary_word_range=FEDERAL_SUMMARY_WORD_RANGE,
    )
    return summary


def bridge_hard_rewrite(text: str) -> str:
    return FEDERAL_SOURCE_BRIDGE_REWRITES.get(text, text)


def federal_keyword_bonus(text: str, keyword_targets: tuple[str, ...]) -> int:
    lowered = text.lower()
    score = 0
    for keyword in keyword_targets:
        if len(keyword) < 3:
            continue
        if " " in keyword:
            if keyword.lower() in lowered:
                score += 4
        elif re.search(rf"\b{re.escape(keyword.lower())}\b", lowered):
            score += 1
    return score


def score_text(text: str, keywords: set[str], profile_key: str, role_index: int, job_description: str) -> int:
    audit = federal_requirement_audit(load_federal_source(), job_description)
    lowered = text.lower()
    score = federal_keyword_bonus(text, tuple(sorted(set(keywords) | set(audit.keyword_targets))))
    for cluster_key, weight in audit.cluster_weights:
        units = candidate_cluster_units(text, cluster_key)
        if units:
            score += weight * units
    if has_great_eight_signal(text):
        score += 6
    if re.search(r"\b(?:vp|director|executive|enterprise|go-live|migration|workshop|dashboard|qbr|sql|risk)\b", lowered, re.I):
        score += 4
    lane_defaults = FEDERAL_DEFAULT_CLUSTERS_BY_LANE.get(profile_key, ())
    score += sum(2 for cluster in lane_defaults if candidate_cluster_strength(text, cluster))
    score += federal_specialized_duty_bonus(text)
    score += max(0, 5 - role_index)
    return score


def has_federal_questionnaire_markers(job_description: str) -> bool:
    lowered = job_description.lower()
    return (
        "knowledge, skills, and abilities" in lowered
        or ("attention to detail" in lowered and "customer service" in lowered and "oral communication" in lowered)
    )


def supplemental_question_bonus(text: str, job_description: str) -> int:
    if not has_federal_questionnaire_markers(job_description):
        return 0
    lowered = text.lower()
    score = 0
    weighted_patterns = (
        (r"\b(?:statement(?:s)? of work|functional requirements documents?|vendor agreements?|documentation|training guides?|adoption guides?)\b", 4),
        (r"\b(?:executive workshops?|quarterly business reviews?|qbrs?|stakeholders?|vp-|director-level|c-suite)\b", 4),
        (r"\b(?:customer|client|support|service quality|adoption|renewal confidence|user)\b", 3),
        (r"\b(?:validation|testing|access controls?|audit readiness|risk documentation|role-based access)\b", 3),
        (r"\b(?:analysis|root-cause|dashboard|reporting|forecast|operational risks?|recommendations)\b", 3),
    )
    for pattern, points in weighted_patterns:
        if re.search(pattern, lowered, re.I):
            score += points
    return score


def federal_specialized_duty_hits(text: str) -> tuple[str, ...]:
    lowered = text.lower()
    hits: list[str] = []
    for duty_label, markers in FEDERAL_SPECIALIZED_DUTY_MARKERS:
        if any(marker.lower() in lowered for marker in markers):
            hits.append(duty_label)
    return tuple(hits)


def federal_specialized_duty_bonus(text: str) -> int:
    hits = federal_specialized_duty_hits(text)
    if not hits:
        return 0
    bonus = len(hits) * 90
    if "Duty 2" in hits:
        bonus += 35
    return bonus


def tailor_text(text: str, job_description: str, active_clusters: tuple[str, ...] | None = None) -> str:
    rewritten = bridge_hard_rewrite(text)
    rewritten = rewrite_supported_text(rewritten, job_description)
    strengthened = strengthen_outcome_framing(rewritten, job_description)
    if active_clusters and "implementation_delivery" in active_clusters and "pilot" in strengthened.lower() and "implemented" not in strengthened.lower():
        strengthened = strengthened.replace("Helped build", "Implemented")
    return normalize_spaces(strengthened)


def selected_bullet_candidates_by_role(
    source: FederalSource,
    job_description: str,
    profile_name: str,
    keywords: set[str],
    layout: FederalLayoutProfile,
    audit: FederalAudit | None = None,
) -> tuple[tuple[FederalBulletCandidate, ...], ...]:
    active_audit = audit or federal_requirement_audit(source, job_description)
    catalog = evidence_engine.load_evidence_catalog(FEDERAL_RESUME_SOURCE)
    direct_record_ids = {
        evidence_id
        for match in active_audit.element_matches
        if match.status == requirement_engine.RequirementStatus.DIRECT
        for evidence_id in match.evidence_ids
    }
    required_confirmed_text = {
        record.rewrite_templates[0]
        for record in catalog
        if record.provenance == "christian-direct-confirmation"
        and (
            record.record_id in direct_record_ids
            or any(term.lower() in job_description.lower() for term in record.allowed_terminology)
        )
    }
    active_clusters = tuple(cluster for cluster, _weight in active_audit.cluster_weights)
    keyword_targets = tuple(sorted(set(keywords) | set(active_audit.keyword_targets)))
    candidates_by_role: list[list[FederalBulletCandidate]] = []

    for role_index, role in enumerate(source.roles):
        seen: set[str] = set()
        scored_candidates: list[FederalBulletCandidate] = []
        for bullet_index, bullet in enumerate(role.bullets):
            tailored = tailor_text(bullet, job_description, active_clusters)
            lowered = tailored.lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            matched_clusters = tuple(
                cluster_key for cluster_key in active_clusters if candidate_cluster_strength(tailored, cluster_key)
            )
            matched_bucket_labels = tuple(
                coverage.bucket.label
                for coverage in active_audit.buckets
                if any(cluster in matched_clusters for cluster in coverage.bucket.clusters)
            )
            cluster_score = 0
            for cluster_key, weight in active_audit.cluster_weights:
                units = candidate_cluster_units(tailored, cluster_key)
                if units:
                    cluster_score += weight * units
            score = federal_keyword_bonus(tailored, keyword_targets) + cluster_score + max(0, 5 - role_index)
            score += supplemental_question_bonus(tailored, job_description)
            if has_great_eight_signal(tailored):
                score += 6
            if re.search(r"\b(?:vp|director|executive|dashboard|sql|go-live|access|audit|workshop|risk|chatbot|training|migration)\b", tailored, re.I):
                score += 4
            score += federal_specialized_duty_bonus(tailored)
            if bullet in required_confirmed_text:
                score += 1000
            scored_candidates.append(
                FederalBulletCandidate(
                    role_index=role_index,
                    bullet_index=bullet_index,
                    score=score,
                    text=tailored,
                    matched_clusters=matched_clusters,
                    matched_bucket_labels=matched_bucket_labels,
                )
            )
        scored_candidates.sort(key=lambda item: (-item.score, item.bullet_index))
        candidates_by_role.append(scored_candidates)

    selected_by_role: list[list[FederalBulletCandidate]] = []
    selected_keys: set[tuple[int, int]] = set()
    covered_clusters: set[str] = set()

    for role_index, candidates in enumerate(candidates_by_role):
        if not candidates:
            selected_by_role.append([])
            continue
        chosen = candidates[0]
        selected_by_role.append([chosen])
        selected_keys.add((chosen.role_index, chosen.bullet_index))
        covered_clusters.update(chosen.matched_clusters)

    covered_duties = {
        duty
        for group in selected_by_role
        for candidate in group
        for duty in federal_specialized_duty_hits(candidate.text)
    }
    for duty_label, _markers in FEDERAL_SPECIALIZED_DUTY_MARKERS:
        if duty_label in covered_duties:
            continue
        best_candidate: FederalBulletCandidate | None = None
        for role_index, candidates in enumerate(candidates_by_role):
            max_for_role = layout.max_bullets_by_role[min(role_index, len(layout.max_bullets_by_role) - 1)]
            if len(selected_by_role[role_index]) >= max_for_role:
                continue
            for candidate in candidates:
                if (candidate.role_index, candidate.bullet_index) in selected_keys:
                    continue
                if duty_label not in federal_specialized_duty_hits(candidate.text):
                    continue
                if best_candidate is None or (candidate.score, -candidate.role_index, -candidate.bullet_index) > (
                    best_candidate.score,
                    -best_candidate.role_index,
                    -best_candidate.bullet_index,
                ):
                    best_candidate = candidate
        if best_candidate is None:
            continue
        selected_by_role[best_candidate.role_index].append(best_candidate)
        selected_keys.add((best_candidate.role_index, best_candidate.bullet_index))
        covered_clusters.update(best_candidate.matched_clusters)
        covered_duties.update(federal_specialized_duty_hits(best_candidate.text))

    cluster_priority = [cluster for cluster, _weight in active_audit.cluster_weights]
    for cluster_key in cluster_priority:
        if cluster_key in covered_clusters:
            continue
        best_candidate: FederalBulletCandidate | None = None
        for role_index, candidates in enumerate(candidates_by_role):
            max_for_role = layout.max_bullets_by_role[min(role_index, len(layout.max_bullets_by_role) - 1)]
            if len(selected_by_role[role_index]) >= max_for_role:
                continue
            for candidate in candidates:
                if (candidate.role_index, candidate.bullet_index) in selected_keys:
                    continue
                if cluster_key not in candidate.matched_clusters:
                    continue
                if best_candidate is None or (candidate.score, -candidate.role_index, -candidate.bullet_index) > (
                    best_candidate.score,
                    -best_candidate.role_index,
                    -best_candidate.bullet_index,
                ):
                    best_candidate = candidate
        if best_candidate is None:
            continue
        selected_by_role[best_candidate.role_index].append(best_candidate)
        selected_keys.add((best_candidate.role_index, best_candidate.bullet_index))
        covered_clusters.update(best_candidate.matched_clusters)

    remaining_candidates: list[FederalBulletCandidate] = []
    for candidates in candidates_by_role:
        for candidate in candidates:
            if (candidate.role_index, candidate.bullet_index) not in selected_keys:
                remaining_candidates.append(candidate)
    remaining_candidates.sort(key=lambda item: (-item.score, item.role_index, item.bullet_index))

    total_selected = sum(len(items) for items in selected_by_role)
    while total_selected < layout.total_bullets and remaining_candidates:
        candidate = remaining_candidates.pop(0)
        role_index = candidate.role_index
        max_for_role = layout.max_bullets_by_role[min(role_index, len(layout.max_bullets_by_role) - 1)]
        if len(selected_by_role[role_index]) >= max_for_role:
            continue
        selected_by_role[role_index].append(candidate)
        selected_keys.add((candidate.role_index, candidate.bullet_index))
        total_selected += 1

    final_groups: list[tuple[FederalBulletCandidate, ...]] = []
    for candidates in selected_by_role:
        ordered = tuple(sorted(candidates, key=lambda item: (-item.score, item.bullet_index)))
        final_groups.append(ordered)
    return tuple(final_groups)


def selected_bullets_by_role(
    source: FederalSource,
    job_description: str,
    profile_name: str,
    keywords: set[str],
    layout: FederalLayoutProfile,
) -> tuple[tuple[str, ...], ...]:
    groups = selected_bullet_candidates_by_role(source, job_description, profile_name, keywords, layout)
    return tuple(tuple(candidate.text for candidate in group) for group in groups)


def selected_reference_set(source: FederalSource, bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...]) -> set[str]:
    refs: set[str] = set()
    for group in bullet_groups:
        for candidate in group:
            refs.add(bullet_reference(source, candidate.role_index, candidate.bullet_index))
    return refs


def apply_selection_visibility(
    source: FederalSource,
    audit: FederalAudit,
    bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...],
) -> FederalAudit:
    selected_refs = selected_reference_set(source, bullet_groups)
    warnings = list(audit.warnings)
    updated_coverages: list[FederalRequirementCoverage] = []

    for coverage in audit.buckets:
        visible_refs = tuple(ref for ref in coverage.source_refs if ref in selected_refs)
        if coverage.bucket.priority >= HIGH_PRIORITY_BUCKET_THRESHOLD and not visible_refs:
            if coverage.alignment == "Direct":
                warnings.append(f"{coverage.bucket.label} is supported in the source but not explicit enough in the selected 2-page resume.")
            elif coverage.alignment == "Adjacent":
                warnings.append(f"{coverage.bucket.label} is only weakly supported and not clearly visible in the selected 2-page resume.")
        updated_coverages.append(replace(coverage, selected_refs=visible_refs))

    visible_work_text = "\n".join(
        (
            *(role.job_summary for role in source.roles),
            *(candidate.text for group in bullet_groups for candidate in group),
        )
    )
    competencies: tuple[str, ...] = ()
    if audit.target_context:
        competencies = tuple(dict.fromkeys((*audit.target_context.minimum_competencies, *audit.target_context.assessed_competencies)))
    coverage_report = evidence_engine.build_coverage_report(
        audit.requirements,
        audit.element_matches,
        visible_work_text,
        competencies=competencies,
    )
    if coverage_report.missing_direct:
        missing_labels = [
            element.text for element in audit.requirements if element.element_id in set(coverage_report.missing_direct)
        ]
        fail("Federal direct-requirement coverage gate failed: " + " | ".join(missing_labels))
    if coverage_report.never_claim_hits:
        fail("Federal unsupported-claim gate failed: " + ", ".join(coverage_report.never_claim_hits))
    for item in coverage_report.competency_misses:
        warnings.append(f"Competency weaving check: {item} is not explicit in selected work experience.")
    warnings.extend(f"Evidence confirmation question: {question}" for question in coverage_report.intake_questions)

    deduped_warnings: list[str] = []
    for warning in warnings:
        if warning not in deduped_warnings:
            deduped_warnings.append(warning)
    return replace(
        audit,
        buckets=tuple(updated_coverages),
        warnings=tuple(deduped_warnings),
        coverage_report=coverage_report,
    )


def resume_candidate_quality(audit: FederalAudit, bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...]) -> int:
    visible_priority = 0
    for coverage in audit.buckets:
        if coverage.selected_refs:
            visible_priority += coverage.bucket.priority
            if coverage.alignment == "Direct":
                visible_priority += coverage.bucket.priority // 2
    selected_score = sum(candidate.score for group in bullet_groups for candidate in group)
    warning_penalty = len(audit.warnings) * 25
    return visible_priority + selected_score - warning_penalty


def maybe_add_blank_paragraph(document: Document, *, size: float, enabled: bool, centered: bool = False) -> None:
    if not enabled:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0 if not centered else 1)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run("")
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_section_heading(document: Document, text: str, *, font_size: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)


def add_body_paragraph(document: Document, text: str, *, size: float, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_labeled_paragraph(document: Document, label: str, text: str, *, size: float, italic_label: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    label_run = paragraph.add_run(f"{label.rstrip(' .:')}: ")
    label_run.bold = True
    label_run.italic = italic_label
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(size)
    body_run = paragraph.add_run(text)
    body_run.font.name = "Calibri"
    body_run.font.size = Pt(size)


def add_bullet(document: Document, text: str, *, size: float) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.08)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def reorder_skills(skills: tuple[str, ...], job_description: str, audit: FederalAudit | None = None) -> tuple[str, ...]:
    active_audit = audit
    if active_audit is None:
        source = load_federal_source()
        active_audit = federal_requirement_audit(source, job_description)
    cluster_weights = dict(active_audit.cluster_weights)
    scored: list[tuple[int, str]] = []
    lowered_job = job_description.lower()

    for skill in skills:
        skill_lower = skill.lower()
        score = federal_keyword_bonus(skill, active_audit.keyword_targets)
        if skill_lower in lowered_job:
            score += 4
        for cluster_key, weight in cluster_weights.items():
            metadata = FEDERAL_REQUIREMENT_CLUSTERS[cluster_key]
            if count_phrase_hits(skill_lower, metadata["skill_patterns"]):
                score += weight
        scored.append((score, skill))
    scored.sort(key=lambda item: (-item[0], item[1]))
    return tuple(skill for _score, skill in scored)


def build_document(
    source: FederalSource,
    job_description: str,
    layout: FederalLayoutProfile,
    summary: str,
    bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...],
    audit: FederalAudit,
) -> Document:
    document = Document()
    section = document.sections[0]
    top, bottom, left, right = layout.margins
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(layout.font_size)

    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_before = Pt(0)
    name_paragraph.paragraph_format.space_after = Pt(0)
    name_paragraph.paragraph_format.line_spacing = 1
    name_run = name_paragraph.add_run(source.contact.name)
    name_run.bold = True
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(NAME_FONT_SIZE)

    contact_line = f"{source.contact.location}  |  {source.contact.email}  |  {source.contact.phone}"
    contact_paragraph = document.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.paragraph_format.space_before = Pt(0)
    contact_paragraph.paragraph_format.space_after = Pt(0)
    contact_paragraph.paragraph_format.line_spacing = 1
    contact_run = contact_paragraph.add_run(contact_line)
    contact_run.font.name = "Calibri"
    contact_run.font.size = Pt(layout.font_size)

    status_line = (
        f"Citizenship: {source.contact.citizenship}  |  Veteran's Preference: {source.contact.veterans_preference}  |  "
        f"Clearance: {source.contact.clearance}  |  Availability: {source.contact.availability}"
    )
    status_paragraph = document.add_paragraph()
    status_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_paragraph.paragraph_format.space_before = Pt(0)
    status_paragraph.paragraph_format.space_after = Pt(1)
    status_paragraph.paragraph_format.line_spacing = 1
    status_run = status_paragraph.add_run(status_line)
    status_run.font.name = "Calibri"
    status_run.font.size = Pt(layout.font_size)
    maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections, centered=True)

    add_section_heading(document, "Professional Summary", font_size=SECTION_FONT_SIZE)
    add_body_paragraph(document, summary, size=layout.font_size)
    maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    add_section_heading(document, "Technical Skills", font_size=SECTION_FONT_SIZE)
    add_body_paragraph(document, "  |  ".join(reorder_skills(source.technical_skills, job_description, audit)), size=layout.font_size)
    maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    add_section_heading(document, "Work Experience", font_size=SECTION_FONT_SIZE)
    shown_company_summaries = 0
    seen_companies: set[str] = set()
    active_clusters = tuple(cluster for cluster, _weight in audit.cluster_weights)
    for role_index, role in enumerate(source.roles):
        title_parts = [role.title, f"{role.start} - {role.end}", f"{role.hours_per_week} Hours Per Week"]
        salary_value = federal_salary_value(role, source.roles)
        if salary_value and salary_value != "$0":
            title_parts.append(salary_value)
        title_line = "  |  ".join(title_parts)
        add_body_paragraph(document, title_line, size=ROLE_FONT_SIZE, bold=True)
        company_line = f"{role.company}, {role.location}  |  Supervisor: {role.supervisor}, {role.supervisor_phone}"
        add_body_paragraph(document, company_line, size=layout.font_size)
        if role.company not in seen_companies and shown_company_summaries < layout.company_summary_roles:
            add_body_paragraph(document, role.company_summary, size=layout.font_size, italic=True)
            seen_companies.add(role.company)
            shown_company_summaries += 1
        if role_index < layout.job_summary_roles:
            role_summary = tailor_text(role.job_summary, job_description, active_clusters)
            if role_index == 0 and audit.target_context and audit.target_context.equivalent_grade:
                source_scope = " ".join((role.job_summary, *role.bullets))
                site_scope = "five sites" if re.search(r"\bfive[- ]site|\bfive sites\b", source_scope, re.I) else "multiple sites"
                user_match = re.search(r"\b150\+ users\b", source_scope, re.I)
                user_scope = user_match.group(0) if user_match else "a broad user base"
                role_summary = normalize_spaces(
                    role_summary
                    + " "
                    + GRADE_EQUIVALENCE_TEMPLATE.format(site_scope=site_scope, user_scope=user_scope)
                )
            add_body_paragraph(document, role_summary, size=layout.font_size, italic=True)
        if layout.show_duties_label:
            add_body_paragraph(document, "Duties & Responsibilities", size=layout.font_size, bold=True)
        for candidate in bullet_groups[role_index]:
            add_bullet(document, candidate.text, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_roles)

    add_section_heading(document, "Education", font_size=SECTION_FONT_SIZE)
    for item in source.education:
        add_body_paragraph(
            document,
            f"{item.degree}  |  {item.date}  |  {item.school}",
            size=layout.font_size,
            bold=True,
        )
        add_body_paragraph(document, item.details, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_roles)

    add_section_heading(document, "Professional Development", font_size=SECTION_FONT_SIZE)
    add_body_paragraph(document, "  |  ".join(source.professional_development), size=layout.font_size)
    if source.volunteer_experience:
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)
        add_section_heading(document, "Volunteer Experience", font_size=SECTION_FONT_SIZE)
        for item in source.volunteer_experience:
            add_body_paragraph(
                document,
                f"{item.title}  |  {item.organization}",
                size=layout.font_size,
                bold=True,
            )
            add_body_paragraph(document, item.description, size=layout.font_size)
    return document


def company_names_from_refs(refs: tuple[str, ...]) -> tuple[str, ...]:
    companies: list[str] = []
    for ref in refs:
        company = ref.split("[", 1)[0].strip()
        if company and company not in companies:
            companies.append(company)
    return tuple(companies)


def qualifications_paragraph(coverage: FederalRequirementCoverage) -> str:
    companies = company_names_from_refs(coverage.source_refs)
    company_text = ""
    if companies:
        company_text = f" across {join_phrases(list(companies)[:3])}"
    prefix = "Demonstrated experience includes" if coverage.alignment == "Direct" else "Relevant experience includes"
    return f"{prefix} {coverage.narrative}{company_text}."


def question_answer_length(paragraphs: tuple[str, ...], labeled_items: tuple[tuple[str, str], ...]) -> int:
    parts = list(paragraphs)
    parts.extend(f"{label}: {text}" for label, text in labeled_items)
    return len("\n\n".join(parts))


def specialized_experience_paragraphs(job_description: str) -> tuple[str, ...]:
    paragraphs = [
        (
            "In my recent role as Enterprise Systems Manager at East West Manufacturing, I served as the senior "
            "enterprise systems subject matter expert for a global five-site manufacturing operation supporting 150+ users "
            "across operations, finance, engineering, and supply chain. I advised VP- and director-level leaders on "
            "modernization priorities, vendor tradeoffs, and implementation risk, then led requirements analysis, vendor "
            "evaluation, ETL validation, user access reviews, control testing, go-live readiness, and cross-site cutover "
            "coordination. For this GS-14 posting, that role demonstrates at least GS-13-equivalent specialized experience "
            "across two of the three duty areas: cross-functional implementation leadership and technical integration/system "
            "modernization, with risk and change-control work supporting the third."
        ),
        (
            "The same role required me to align business needs to technology solutions and produce acquisition-ready "
            "documentation. I developed Statements of Work, Functional Requirements Documents, vendor agreements, milestone "
            "schedules, and risk assumptions for platform enhancements, integrations, and technology acquisitions. At Aptean, "
            "I owned full-lifecycle delivery for 80+ international manufacturing clients, facilitated 60+ executive workshops "
            "and QBRs, wrote SOWs and FRDs, and coordinated product, development, implementation, and support teams to "
            "stabilize risk, protect scope, and improve adoption."
        ),
    ]
    if federal_ai_focus(job_description):
        paragraphs = list(paragraphs)
        paragraphs[1] = (
            "The same role required me to align business needs to technology solutions and produce acquisition-ready "
            "documentation. I developed Statements of Work, Functional Requirements Documents, vendor agreements, milestone "
            "schedules, and risk assumptions for platform enhancements, system integrations, and technology acquisitions, and "
            "I used Codex, Claude, and other AI-assisted tools to accelerate documentation, reporting, SQL troubleshooting, "
            "and operational analysis. At Aptean, I owned full-lifecycle delivery for 80+ international manufacturing clients, "
            "facilitated 60+ executive workshops and QBRs, wrote SOWs and FRDs, and coordinated product, development, "
            "implementation, and support teams to stabilize risk and improve adoption."
        )
    return tuple(paragraphs)


def competency_response(name: str) -> str:
    normalized = name.lower()
    if normalized == "attention to detail":
        return (
            "I review requirements, ETL results, testing outcomes, access controls, and risk documentation before go-live "
            "or contract execution, supported by ITIL 4 and ServiceNow NextGen training."
        )
    if normalized == "customer service":
        return (
            "At Aptean and The Home Depot, I worked directly with customers and internal users, balancing technical "
            "constraints with service quality, adoption, and business outcomes while stabilizing high-risk accounts and "
            "improving support workflows."
        )
    if normalized == "decision making":
        return (
            "At East West, I advised VP- and director-level stakeholders on modernization priorities, vendor tradeoffs, "
            "implementation risk, and platform decisions affecting five sites and 150+ users."
        )
    if normalized == "information management":
        return (
            "I built and maintained information-management systems through 200+ SQL-based dashboards, reporting data models, "
            "Crystal Reports, Power BI tools, and inventory reporting used for decisions affecting $20M+ in daily inventory."
        )
    if normalized == "interpersonal skills":
        return (
            "I have worked across executives, customers, end users, vendors, and technical teams, including 60+ executive "
            "workshops and QBRs at Aptean and cross-functional coordination at East West."
        )
    if normalized == "oral communication":
        return (
            "I translated technical issues into plain business language for executives, facilitated workshops and QBRs, "
            "delivered training, and aligned stakeholders around scope, risk, and next steps."
        )
    if normalized == "problem solving":
        return (
            "Across East West, Aptean, and The Home Depot, I turned process gaps, customer trends, operational risks, and "
            "system constraints into corrective workflows, plans, reporting tools, and support-channel improvements."
        )
    if normalized == "teamwork":
        return (
            "I coordinated East West vendors and business users, partnered with Aptean product, development, implementation, "
            "and support teams, and helped a Home Depot pilot team launch a new SMS support channel."
        )
    if normalized == "technical competence":
        return (
            "My technical experience includes ERP modernization, SQL and BI reporting, ETL validation, access-control "
            "governance, recovery testing, Azure DevOps/TFS, Jira, ServiceNow, Salesforce, Power BI, and LivePerson workflows."
        )
    return (
        "My resume shows repeated experience translating complex technical work into practical decisions, documentation, and "
        "cross-functional delivery support."
    )


def generic_ksa_response(item: FederalKSAItem) -> str:
    clusters = item.clusters
    if "executive_alignment" in clusters:
        return (
            "My resume shows repeated executive-facing technology work, including VP-, director-, and C-suite alignment on "
            "technology priorities, implementation risk, and business tradeoffs. I have translated complex technical issues "
            "into clear recommendations, next steps, and decision-ready updates across manufacturing, SaaS, and support environments."
        )
    if "acquisition_support" in clusters:
        return (
            "My written work has included Statements of Work, Functional Requirements Documents, vendor agreements, risk "
            "documentation, training materials, and adoption guides used to drive enterprise technology decisions. That "
            "experience reflects disciplined documentation that turns technical complexity into clear scope, milestones, and recommendations."
        )
    if "reporting_analytics" in clusters or "operations_improvement" in clusters:
        return (
            "I have led analysis-heavy work that turned operational data, system issues, and workflow gaps into reporting, "
            "recommendations, and measurable process improvements. That includes dashboard development, root-cause analysis, "
            "risk identification, and decision support for leaders."
        )
    if "customer_service_delivery" in clusters or "change_adoption" in clusters:
        return (
            "My resume reflects repeated work at the intersection of customers, end users, and technical delivery teams. "
            "I have supported adoption, service quality, training, and stakeholder communication in ways that helped new "
            "processes or systems gain traction."
        )
    return (
        "My resume demonstrates senior technology delivery work that combined planning, documentation, stakeholder alignment, "
        "and measurable business impact across multiple environments."
    )


def ksa_response(item: FederalKSAItem) -> str:
    lowered = item.name.lower()
    if "written skill" in lowered or "documentation" in lowered or "white paper" in lowered:
        return (
            "My resume reflects extensive written documentation used by executives, customers, and delivery teams to make "
            "decisions. At East West and Aptean, I wrote Statements of Work, Functional Requirements Documents, vendor "
            "agreements, risk documentation, training guides, adoption materials, and executive-facing recommendation or status "
            "materials that translated complex implementation issues into clear scope, milestones, tradeoffs, and next steps."
        )
    if "planning" in lowered or "coordinating" in lowered or "projects" in lowered or "studies" in lowered:
        return (
            "I have led broad-scope initiatives that required planning, analysis, coordination, and recommendations across "
            "multiple business functions. At East West, I managed platform migration work plus new warehouse and robotics "
            "program setup across operations, finance, and engineering, while at The Home Depot I built reporting and analyzed "
            "service and pipeline trends to surface operational risks and decision points for leaders."
        )
    if "interpersonal" in lowered:
        return (
            "My resume shows repeated work at the boundary between executives, technical teams, and end users. I facilitated "
            "60+ executive workshops and quarterly business reviews at Aptean, coordinated vendors and cross-functional "
            "stakeholders at East West without direct supervisory authority, and supported workflow adoption at The Home Depot, "
            "all of which required translating technical work into clear business decisions and driving adoption across "
            "different audiences."
        )
    if "theories" in lowered or "principles" in lowered or "standards" in lowered:
        return (
            "My resume shows the closest equivalent in private-sector enterprise systems leadership. As Enterprise Systems "
            "Manager at East West Manufacturing, I served as the senior technical authority for a five-site operation, advised "
            "VP- and director-level leaders on modernization priorities, vendor tradeoffs, security and access controls, and "
            "implementation risk, and connected technical decisions to business impact. Earlier roles at Aptean and Aderant "
            "also required me to translate complex system issues for senior stakeholders and coordinate practical modernization "
            "decisions across technical and business teams."
        )
    if "artificial intelligence" in lowered or re.search(r"\bai\b", lowered, re.I):
        return (
            "My AI-related experience is strongest in applied business use rather than formal model development. At East West "
            "Manufacturing, I used Codex, Claude, and AI-assisted tools to accelerate documentation, reporting, SQL "
            "troubleshooting, and operational analysis, and at The Home Depot I helped configure LivePerson chat and text "
            "workflows, automated greetings and closings, and AI-assisted chatbot logic for a new SMS support channel. Those "
            "experiences show I can evaluate AI-enabled workflows, explain them to stakeholders, and apply them where they "
            "improve service, speed, or decision quality."
        )
    return generic_ksa_response(item)


def build_question_one_answer(questions: FederalApplicationQuestionSet, job_description: str) -> tuple[tuple[str, ...], tuple[tuple[str, str], ...], int]:
    paragraphs = specialized_experience_paragraphs(job_description)
    competencies = questions.competencies or tuple(FederalCompetencyPrompt(name=name, description="") for name in FEDERAL_COMPETENCY_NAMES)
    labeled_items = tuple((prompt.name, competency_response(prompt.name)) for prompt in competencies)
    return paragraphs, labeled_items, question_answer_length(paragraphs, labeled_items)


def build_question_two_answer(questions: FederalApplicationQuestionSet) -> tuple[tuple[tuple[str, str], ...], int]:
    labeled_items = tuple((item.name, ksa_response(item)) for item in questions.ksas)
    count = question_answer_length((), labeled_items)
    return labeled_items, count


def federal_role_focus_area(job_description: str, profile_name: str) -> str:
    lowered = job_description.lower()
    if federal_ai_focus(job_description):
        return "AI-enabled workflows, trustworthy reporting, and risk-aware modernization"
    if re.search(r"\b(?:cybersecurity|zero[- ]trust|incident response|access controls?|least-privilege|security)\b", lowered, re.I):
        return "secure modernization, incident-response readiness, and access-control discipline"
    if re.search(r"\b(?:data|analytics|reporting|dashboard|measurement|evidence)\b", lowered, re.I):
        return "data quality, reporting integrity, and operational decision support"
    if re.search(r"\b(?:acquisition|procurement|vendor|contract)\b", lowered, re.I):
        return "acquisition-ready documentation, vendor coordination, and implementation oversight"
    if profile_name == "implementation_delivery":
        return "enterprise implementation, workflow reliability, and measurable delivery outcomes"
    if profile_name == "change_enablement":
        return "adoption planning, stakeholder communication, and durable process change"
    if profile_name == "process_improvement":
        return "process discipline, root-cause analysis, and measurable efficiency gains"
    if profile_name == "analytics_operations":
        return "data quality, reporting integrity, and operational decision support"
    return "trustworthy systems, measurable service improvement, and mission-focused execution"


def federal_role_reference(job_description: str) -> str:
    role_title = extract_federal_role_title(job_description)
    if role_title:
        return f"the {role_title} role"
    return "this role"


def polished_standard_essay_answer(prompt: FederalEssayPrompt, job_description: str, profile_name: str) -> str:
    base = prompt.answer
    focus_area = federal_role_focus_area(job_description, profile_name)
    role_reference = federal_role_reference(job_description)
    if prompt.key == "constitutional_commitment":
        return normalize_spaces(
            f"{base} Work like that is what draws me to {role_reference}, "
            f"especially where the work depends on {focus_area}."
        )
    if prompt.key == "government_efficiency":
        polished = base.replace("As aforementioned, ", "")
        return normalize_spaces(
            f"{polished} In {role_reference}, I would apply that same approach to {focus_area} so the agency can operate "
            "with less waste, clearer reporting, and stronger execution discipline."
        )
    if prompt.key == "executive_orders":
        polished = base.replace(
            "The 2 most relevant Executive Orders seem to be",
            "Two Executive Orders that often align closely with federal technology work are",
        )
        return normalize_spaces(
            f"I would start by aligning to the policy priorities named in the posting and the agency mission. {polished} "
            f"If hired, I would translate those priorities into practical execution around {focus_area}."
        )
    if prompt.key == "work_ethic":
        polished = base.replace(
            "where I am able to often work on an island all on my own, but able to effectively collaborate with teams across different countries and time zones as well.",
            "where I often take ownership independently while still collaborating effectively with teams across countries and time zones.",
        )
        return normalize_spaces(
            f"{polished} The same blend of self-direction and follow-through would help me serve effectively in a position "
            f"that depends on {focus_area}."
        )
    return normalize_spaces(base)


def build_standard_federal_essay_answers(source: FederalSource, job_description: str) -> tuple[tuple[str, str], ...]:
    profile = job_problem_profile(job_description, source_visible_text(source))
    prompts = load_federal_standard_essays()
    return tuple(
        (prompt.question, polished_standard_essay_answer(prompt, job_description, profile.primary_lane))
        for prompt in prompts
    )


def additional_application_question_responses(
    job_description: str,
    questions: FederalApplicationQuestionSet,
    standard_essay_answers: tuple[tuple[str, str], ...],
) -> tuple[question_prep.QualificationsResponse, ...]:
    responses = question_prep.active_application_question_responses(
        job_description,
        excluded_categories=("company_interest",),
    )
    seen_prompts = tuple(
        [
            *(prompt.name for prompt in questions.competencies),
            *questions.specialized_items,
            *(item.name for item in questions.ksas),
            *(question for question, _answer in standard_essay_answers),
        ]
    )
    return question_prep.dedupe_question_responses(responses, seen_prompts)


def element_qualification_entries(audit: FederalAudit) -> tuple[tuple[str, str], ...]:
    if not audit.requirements or not audit.element_matches:
        return ()
    catalog = {record.record_id: record for record in evidence_engine.load_evidence_catalog(FEDERAL_RESUME_SOURCE)}
    matches = {match.element_id: match for match in audit.element_matches}
    entries: list[tuple[str, str]] = []
    for index, element in enumerate(audit.requirements, start=1):
        match = matches[element.element_id]
        evidence = [catalog[evidence_id] for evidence_id in match.evidence_ids if evidence_id in catalog]
        if match.status == requirement_engine.RequirementStatus.DIRECT and evidence:
            record = evidence[0]
            response = f"Direct - {record.employer}, {record.role}: {record.rewrite_templates[0]}"
        elif match.status == requirement_engine.RequirementStatus.ADJACENT and evidence:
            record = evidence[0]
            response = f"Adjacent - {record.employer}, {record.role}: {record.text} The unsupported portion is not claimed."
        elif match.status == requirement_engine.RequirementStatus.TRANSFERABLE and evidence:
            record = evidence[0]
            response = f"Transferable - {record.employer}, {record.role}: {record.text}"
        else:
            response = "Unsupported - No directly supported experience is claimed for this element."
        entries.append((f"Element {index}: {element.text}", normalize_spaces(response)))
    return tuple(entries)


def prepare_federal_qualifications_content(
    source: FederalSource,
    job_description: str,
    audit: FederalAudit,
    *,
    question_context_issues: tuple[str, ...] = (),
) -> FederalQualificationsContent:
    profile = job_problem_profile(job_description, source_visible_text(source))
    questions = federal_application_questions(job_description, profile.primary_lane)
    standard_essay_answers = build_standard_federal_essay_answers(source, job_description)
    active_responses = additional_application_question_responses(
        job_description,
        questions,
        standard_essay_answers,
    )
    question_one_paragraphs: tuple[str, ...] = ()
    question_one_labels: tuple[tuple[str, str], ...] = ()
    question_one_count = 0
    question_two_labels: tuple[tuple[str, str], ...] = ()
    question_two_count = 0
    if questions.has_supplemental_questions:
        question_one_paragraphs, question_one_labels, question_one_count = build_question_one_answer(
            questions,
            job_description,
        )
        question_two_labels, question_two_count = build_question_two_answer(questions)

    recent_items = question_prep.recent_interviewer_question_prep_items(
        job_description,
        (audit.target_context.agency or audit.target_context.company) if audit.target_context else "",
        audit.target_context.official_title if audit.target_context else (extract_federal_role_title(job_description) or ""),
    )
    recent_scripts = (
        build_standard_qualifications_statement.build_recent_interviewer_scripts(
            recent_items,
            job_description,
            source_visible_text(source),
            (audit.target_context.agency or audit.target_context.company) if audit.target_context else "",
            audit.target_context.official_title if audit.target_context else (extract_federal_role_title(job_description) or ""),
        )
        if recent_items
        else ()
    )

    # Validate the sendable content once before any layout render. Excluded
    # categories have already been removed before answer generation.
    for paragraph in question_one_paragraphs:
        print_prose_warnings("Federal Question 1", paragraph, "generic")
    for label, text in (*question_one_labels, *question_two_labels):
        print_prose_warnings(f"Federal Qualifications - {label}", text, "generic")
    for _question, answer in standard_essay_answers:
        print_prose_warnings("Federal Standard Essay Response", answer, "generic")
    for response in active_responses:
        if response.warning:
            print(response.warning)
        print_prose_warnings(f"Federal Additional Question - {response.prompt}", response.answer, "generic")
    for prompt, spoken in recent_scripts:
        print_prose_warnings(f"Federal Recent Interview Question - {prompt}", spoken, "generic")

    return FederalQualificationsContent(
        questions=questions,
        standard_essay_answers=standard_essay_answers,
        active_question_responses=active_responses,
        recent_interviewer_scripts=recent_scripts,
        question_context_issues=question_context_issues,
        parse_diagnostics=(audit.target_context.parse_diagnostics if audit.target_context else ()),
        question_one_paragraphs=question_one_paragraphs,
        question_one_labels=question_one_labels,
        question_one_count=question_one_count,
        question_two_labels=question_two_labels,
        question_two_count=question_two_count,
    )


def build_qualifications_document(
    source: FederalSource,
    job_description: str,
    layout: FederalQualificationLayout,
    audit: FederalAudit,
    content: FederalQualificationsContent | None = None,
) -> Document:
    content = content or prepare_federal_qualifications_content(source, job_description, audit)
    document = Document()
    section = document.sections[0]
    top, bottom, left, right = layout.margins
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(layout.font_size)

    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_before = Pt(0)
    name_paragraph.paragraph_format.space_after = Pt(0)
    name_paragraph.paragraph_format.line_spacing = 1
    name_run = name_paragraph.add_run(source.contact.name)
    name_run.bold = True
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(NAME_FONT_SIZE)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(1)
    title_paragraph.paragraph_format.line_spacing = 1
    title_run = title_paragraph.add_run("Federal Qualifications Statement")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(10)

    role_target = (
        audit.target_context.output_label
        if audit.target_context and audit.target_context.output_label
        else extract_federal_output_name(job_description)
    )
    add_body_paragraph(document, f"Target Position: {role_target}", size=layout.font_size, italic=True)
    maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    element_entries = element_qualification_entries(audit)
    if element_entries:
        add_section_heading(document, "Specialized Experience Element Mapping", font_size=SECTION_FONT_SIZE)
        for label, response in element_entries:
            add_labeled_paragraph(document, label, response, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    questions = content.questions
    standard_essay_answers = content.standard_essay_answers
    if questions.has_supplemental_questions:
        add_section_heading(document, "Question 1 - Specialized Experience", font_size=SECTION_FONT_SIZE)
        add_body_paragraph(
            document,
            f"Paste-ready response: {content.question_one_count:,} / {questions.char_limit:,} characters",
            size=layout.font_size,
            italic=True,
        )
        for paragraph in content.question_one_paragraphs:
            add_body_paragraph(document, paragraph, size=layout.font_size)
        for label, text in content.question_one_labels:
            add_labeled_paragraph(document, label, text, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

        if content.question_two_labels:
            add_section_heading(document, "Question 2 - Required KSAs", font_size=SECTION_FONT_SIZE)
            add_body_paragraph(
                document,
                f"Paste-ready response: {content.question_two_count:,} / {questions.char_limit:,} characters",
                size=layout.font_size,
                italic=True,
            )
            for label, text in content.question_two_labels:
                add_labeled_paragraph(document, label, text, size=layout.font_size)
            maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    if standard_essay_answers:
        add_section_heading(document, "Standard Federal Essay Responses", font_size=SECTION_FONT_SIZE)
        for question, answer in standard_essay_answers:
            add_body_paragraph(document, question, size=layout.font_size, bold=True)
            add_body_paragraph(document, answer, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    if not questions.has_supplemental_questions:
        for coverage in audit.buckets:
            if coverage.alignment == "Unsupported":
                continue
            add_section_heading(document, coverage.bucket.label, font_size=SECTION_FONT_SIZE)
            paragraph = qualifications_paragraph(coverage)
            print_prose_warnings(f"Federal Qualifications - {coverage.bucket.label}", paragraph, "generic")
            add_body_paragraph(document, paragraph, size=layout.font_size)
            maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    additional_responses = content.active_question_responses
    if additional_responses:
        add_section_heading(document, "Additional Application Questions", font_size=SECTION_FONT_SIZE)
        for response in additional_responses:
            add_body_paragraph(document, response.prompt, size=layout.font_size, bold=True)
            add_body_paragraph(document, response.answer, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)

    if content.recent_interviewer_scripts:
        add_section_heading(document, "Recent Interview Questions To Be Ready For", font_size=SECTION_FONT_SIZE)
        for prompt, spoken in content.recent_interviewer_scripts:
            add_body_paragraph(document, prompt, size=layout.font_size, bold=True)
            add_body_paragraph(document, spoken, size=layout.font_size)
        maybe_add_blank_paragraph(document, size=layout.font_size, enabled=layout.blank_after_sections)
    return document


def page_count_for_docx(docx_path: Path, temp_root: Path) -> int | None:
    unpack_root = temp_root / f"page_count_{uuid.uuid4().hex}"
    unpack_root.mkdir(parents=True, exist_ok=False)
    try:
        unpack_docx(docx_path, unpack_root)
        document_xml = unpack_root / "word" / "document.xml"
        return rendered_page_count(docx_path, temp_root, document_xml)
    finally:
        shutil.rmtree(unpack_root, ignore_errors=True)


def federal_page_count_label(page_count: int | None) -> str:
    return "unverified (renderer unavailable)" if page_count is None else str(page_count)


SALARY_FIGURE_PATTERN = re.compile(r"\$\d[\d,]*(?:\.\d{2})?")
HOURS_PER_WEEK_PATTERN = re.compile(r"\b\d{1,3}\s+Hours Per Week\b", re.I)


def federal_rendered_position_blocks(visible: str) -> list[dict[str, str]]:
    lines = [line.strip() for line in visible.splitlines() if line.strip()]
    positions: list[dict[str, str]] = []
    for index, line in enumerate(lines):
        if "Supervisor:" not in line:
            continue
        header = lines[index - 1] if index > 0 else ""
        employer = line.split(",", 1)[0].strip() or "Rendered position"
        title = header.split("|", 1)[0].strip() or employer
        positions.append(
            {
                "employer": employer,
                "title": title,
                "header": header,
                "supervisor": line,
                "text": f"{header}\n{line}",
            }
        )
    return positions


def federal_plain_text_validation(docx_path: Path) -> dict[str, object]:
    visible = docx_visible_text_from_path(docx_path)
    blockers: list[str] = []
    warnings: list[str] = []
    for section in FEDERAL_REQUIRED_SECTIONS:
        if section.lower() not in visible.lower():
            blockers.append(f"Missing required federal section: {section}.")
    if not re.search(r"\b[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}\b", visible, re.I):
        blockers.append("No email address detected in federal ATS plain text.")
    if not re.search(r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*-\s*(?:Present|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b", visible, re.I):
        blockers.append("No role date ranges detected in federal ATS plain text.")
    if "Supervisor:" not in visible:
        blockers.append("No supervisor lines detected in federal ATS plain text.")
    for position in federal_rendered_position_blocks(visible):
        if not HOURS_PER_WEEK_PATTERN.search(position["text"]):
            blockers.append(
                f"Federal eligibility blocker: {position['employer']} is missing Hours Per Week in the rendered resume."
            )
        if not SALARY_FIGURE_PATTERN.search(position["text"]):
            blockers.append(
                f"Federal eligibility blocker: {position['employer']} is missing salary in the rendered resume."
            )
    word_count = len(re.findall(r"\b[\w+.#'-]+\b", visible))
    if word_count < 450 or word_count > 1800:
        warnings.append(f"Federal ATS plain-text word count looks unusual ({word_count} words).")
    return {"blockers": blockers, "warnings": warnings, "word_count": word_count}


def _employer_text_block(visible: str, employer: str, following_employers: tuple[str, ...]) -> str:
    start_match = re.search(rf"(?m)^{re.escape(employer)},[^\n]*\|\s*Supervisor:", visible)
    start = start_match.start() if start_match else visible.find(employer)
    if start < 0:
        return ""
    end_positions = []
    for candidate in following_employers:
        match = re.search(rf"(?m)^{re.escape(candidate)},[^\n]*\|\s*Supervisor:", visible[start + len(employer):])
        if match:
            end_positions.append(start + len(employer) + match.start())
    end_positions = [position for position in end_positions if position >= 0]
    end = min(end_positions) if end_positions else len(visible)
    return visible[start:end]


VERSION_CONTROL_EVIDENCE_PATTERN = re.compile(
    r"\b(?:azure devops|tfs|git|github|version control)\b",
    re.I,
)


def assert_version_control_employer_scope(aptean_text: str, non_aptean_work_text: str) -> None:
    """Keep every version-control claim inside the Aptean role block."""

    if VERSION_CONTROL_EVIDENCE_PATTERN.search(aptean_text):
        if not all(term in aptean_text.lower() for term in ("azure devops/tfs", "git/github")):
            fail("Federal terminology injection failed: version-control tools are not explicit in the Aptean block.")
    if VERSION_CONTROL_EVIDENCE_PATTERN.search(non_aptean_work_text):
        fail("Federal provenance gate failed: version-control evidence appeared outside Aptean.")


def assert_final_federal_coverage(output_docx: Path, job_description: str, audit: FederalAudit) -> None:
    visible = docx_visible_text_from_path(output_docx)
    report = evidence_engine.build_coverage_report(
        audit.requirements,
        audit.element_matches,
        visible,
        competencies=tuple(
            dict.fromkeys(
                (*((audit.target_context.minimum_competencies if audit.target_context else ())),
                 *((audit.target_context.assessed_competencies if audit.target_context else ())))
            )
        ),
    )
    if report.missing_direct:
        fail("Final federal coverage gate failed for element IDs: " + ", ".join(report.missing_direct))
    if report.never_claim_hits:
        fail("Final federal never-claim gate failed: " + ", ".join(report.never_claim_hits))

    east_west = _employer_text_block(visible, "East West Manufacturing", ("Aptean", "The Home Depot", "Aderant"))
    aptean = _employer_text_block(visible, "Aptean", ("The Home Depot", "Aderant"))
    home_depot = _employer_text_block(visible, "The Home Depot", ("Aderant",))
    aderant = _employer_text_block(visible, "Aderant", ())
    lowered_jd = job_description.lower()
    if "stored procedures" in lowered_jd and not all(term in east_west.lower() for term in ("views", "stored procedures", "data models")):
        fail("Federal terminology injection failed: SQL objects are not explicit in the East West block.")
    non_aptean_work = "\n".join((east_west, home_depot, aderant))
    if "version control" in lowered_jd and not VERSION_CONTROL_EVIDENCE_PATTERN.search(aptean):
        fail("Federal terminology injection failed: version-control evidence is missing from the Aptean block.")
    assert_version_control_employer_scope(aptean, non_aptean_work)
    if "contact center" in lowered_jd and not all(term in home_depot.lower() for term in ("contact center", "service level", "customer interaction data")):
        fail("Federal terminology injection failed: contact-center terminology is not explicit in the Home Depot block.")
    if "recovery" in lowered_jd and "recovery testing" not in aderant.lower():
        fail("Federal terminology injection failed: capped recovery-testing evidence is not explicit in the Aderant block.")
    if audit.target_context and audit.target_context.equivalent_grade and "senior technical authority" not in east_west.lower():
        fail("Federal grade-equivalence framing is missing from the most recent role.")


def resume_plan(
    temp_root: Path,
    source: FederalSource,
    job_description: str,
    *,
    target_grade: str = "",
    question_context_issues: tuple[str, ...] = (),
) -> FederalResumePlan:
    profile = job_problem_profile(job_description, source_visible_text(source))
    keywords = keyword_set(job_description)
    target_context = requirement_engine.build_target_context(
        job_description,
        workflow="federal",
        target_grade=target_grade,
    )
    audit = federal_requirement_audit(
        source,
        job_description,
        target_grade=target_grade,
        target_context=target_context,
    )
    summary = build_gs14_summary(source, job_description, audit)

    resume_candidates: list[FederalPlanCandidate] = []
    for layout in FEDERAL_LAYOUT_PROFILES:
        if layout.font_size < MIN_BODY_FONT_SIZE:
            fail("Federal layout profile below 10pt minimum is not allowed.")
        bullet_groups = selected_bullet_candidates_by_role(source, job_description, profile.primary_lane, keywords, layout, audit)
        visible_audit = apply_selection_visibility(source, audit, bullet_groups)
        candidate_docx = temp_root / f"resume_{layout.name}.docx"
        document = build_document(source, job_description, layout, summary, bullet_groups, visible_audit)
        document.save(str(candidate_docx))
        page_count = page_count_for_docx(candidate_docx, temp_root)
        selected_bullets = sum(len(group) for group in bullet_groups)
        quality_score = resume_candidate_quality(visible_audit, bullet_groups)
        page_label = federal_page_count_label(page_count)
        print(
            f"Federal fit check: layout={layout.name} font={layout.font_size:.1f} bullets={selected_bullets} "
            f"pages={page_label} quality={quality_score}"
        )
        resume_candidates.append(
            FederalPlanCandidate(
                layout=layout,
                bullet_groups=bullet_groups,
                audit=visible_audit,
                page_count=page_count,
                quality_score=quality_score,
            )
        )

    exact_resume_candidates = [candidate for candidate in resume_candidates if candidate.page_count == TARGET_PAGE_COUNT]
    unknown_resume_candidates = [candidate for candidate in resume_candidates if candidate.page_count is None]
    if exact_resume_candidates:
        chosen_resume = max(exact_resume_candidates, key=lambda candidate: (candidate.quality_score, sum(len(group) for group in candidate.bullet_groups)))
    elif unknown_resume_candidates:
        chosen_resume = max(unknown_resume_candidates, key=lambda candidate: candidate.quality_score)
    else:
        fail("Federal resume did not resolve to exactly two pages at the 10pt minimum across available layouts.")

    qualifications_content = prepare_federal_qualifications_content(
        source,
        job_description,
        chosen_resume.audit,
        question_context_issues=question_context_issues,
    )
    qualification_candidates: list[FederalQualificationPlan] = []
    for layout in FEDERAL_QUALIFICATION_LAYOUTS:
        candidate_docx = temp_root / f"qualifications_{layout.name}.docx"
        document = build_qualifications_document(
            source,
            job_description,
            layout,
            chosen_resume.audit,
            qualifications_content,
        )
        document.save(str(candidate_docx))
        page_count = page_count_for_docx(candidate_docx, temp_root)
        page_label = federal_page_count_label(page_count)
        print(f"Federal qualifications check: layout={layout.name} font={layout.font_size:.1f} pages={page_label}")
        qualification_candidates.append(FederalQualificationPlan(layout=layout, page_count=page_count))

    compact_qualification_candidates = [
        candidate
        for candidate in qualification_candidates
        if candidate.page_count is not None and candidate.page_count <= MAX_QUALIFICATIONS_PAGE_COUNT
    ]
    unknown_qualification_candidates = [candidate for candidate in qualification_candidates if candidate.page_count is None]
    if compact_qualification_candidates:
        chosen_qualification = min(compact_qualification_candidates, key=lambda candidate: candidate.page_count or MAX_QUALIFICATIONS_PAGE_COUNT)
    elif unknown_qualification_candidates:
        chosen_qualification = unknown_qualification_candidates[0]
    else:
        chosen_qualification = min(
            qualification_candidates,
            key=lambda candidate: candidate.page_count or (MAX_QUALIFICATIONS_PAGE_COUNT + 1),
        )
        print(
            "FEDERAL QUALIFICATIONS WARNING: supplemental-question response exceeded the preferred page budget; "
            "using the most compact qualifications layout available."
        )

    plan_warnings = list(chosen_resume.audit.warnings)
    if chosen_resume.page_count is None:
        plan_warnings.append(
            "Federal page-count warning: selected the resume candidate with an unverified page count because the renderer was unavailable."
        )
    if chosen_qualification.page_count is None:
        plan_warnings.append(
            "Federal qualifications page-count warning: selected a qualifications layout with an unverified page count because the renderer was unavailable."
        )

    return FederalResumePlan(
        summary=summary,
        audit=replace(chosen_resume.audit, warnings=tuple(plan_warnings)),
        bullet_groups=chosen_resume.bullet_groups,
        resume_layout=chosen_resume.layout,
        resume_page_count=chosen_resume.page_count,
        qualifications_layout=chosen_qualification.layout,
        qualifications_page_count=chosen_qualification.page_count,
        qualifications_content=qualifications_content,
    )


def selected_bullet_reference_lines(source: FederalSource, bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...]) -> tuple[str, ...]:
    lines: list[str] = []
    for role_index, group in enumerate(bullet_groups):
        refs = ", ".join(bullet_reference(source, candidate.role_index, candidate.bullet_index) for candidate in group)
        lines.append(f"{source.roles[role_index].company}: {refs or 'none'}")
    return tuple(lines)


def selected_duty_visibility_lines(source: FederalSource, bullet_groups: tuple[tuple[FederalBulletCandidate, ...], ...]) -> tuple[str, ...]:
    references_by_duty: dict[str, list[str]] = {label: [] for label, _markers in FEDERAL_SPECIALIZED_DUTY_MARKERS}
    for group in bullet_groups:
        for candidate in group:
            ref = bullet_reference(source, candidate.role_index, candidate.bullet_index)
            for duty in federal_specialized_duty_hits(candidate.text):
                if ref not in references_by_duty[duty]:
                    references_by_duty[duty].append(ref)
    lines: list[str] = []
    for duty_label, _markers in FEDERAL_SPECIALIZED_DUTY_MARKERS:
        refs = references_by_duty[duty_label]
        if refs:
            lines.append(f"{duty_label}: surfaced via {', '.join(refs[:4])}")
        else:
            lines.append(f"{duty_label}: not surfaced in selected two-page bullets")
    return tuple(lines)


def requirement_report_lines(audit: FederalAudit) -> tuple[str, ...]:
    lines: list[str] = []
    if audit.requirements and audit.element_matches:
        lines.append("Requirement-element coverage:")
        match_by_id = {match.element_id: match for match in audit.element_matches}
        for element in audit.requirements:
            match = match_by_id[element.element_id]
            evidence = ", ".join(match.evidence_ids) or "honest non-claim"
            lines.append(f"  {element.element_id} [{match.status.value}] {element.text}")
            lines.append(f"    Evidence: {evidence}")
        if audit.coverage_report:
            lines.append("Competency weaving:")
            lines.append("  Visible: " + (", ".join(audit.coverage_report.competency_hits) or "none"))
            lines.append("  Not explicit: " + (", ".join(audit.coverage_report.competency_misses) or "none"))
            if audit.coverage_report.intake_questions:
                lines.append("Evidence confirmation questions:")
                lines.extend(f"  - {question}" for question in audit.coverage_report.intake_questions)
    for coverage in audit.buckets:
        evidence = ", ".join(coverage.source_refs) or "none detected"
        visible = ", ".join(coverage.selected_refs) or "not explicit in selected resume bullets"
        lines.append(
            f"{coverage.bucket.label} [{coverage.alignment}] -> {coverage.bucket.text}"
        )
        lines.append(f"  Evidence: {evidence}")
        lines.append(f"  Visible in resume: {visible}")
    return tuple(lines)


def federal_draft_channels(
    target_context: requirement_engine.TargetContext,
    question_context_issues: tuple[str, ...] = (),
) -> tuple[tuple[str, ...], tuple[str, ...], bool]:
    question_reasons = tuple(question_context_issues)
    parse_reasons = tuple(
        diagnostic.message
        for diagnostic in target_context.parse_diagnostics
        if diagnostic.requires_draft
    )
    return question_reasons, parse_reasons, bool(question_reasons or parse_reasons)


def federal_output_paths(output_label: str, *, is_draft: bool) -> tuple[Path, Path]:
    draft_marker = " DRAFT" if is_draft else ""
    return (
        OUTPUT_DIR / f"Christian Estrada - {output_label}{draft_marker} Federal Resume.docx",
        OUTPUT_DIR / f"Christian Estrada - {output_label}{draft_marker} Federal Qualifications Statement.docx",
    )


def validate_staged_federal_documents(
    resume_candidate: Path,
    qualifications_candidate: Path,
    job_description: str,
    audit: FederalAudit,
    temp_root: Path,
    *,
    resume_render_label: str,
    qualifications_render_label: str,
) -> tuple[int | None, int | None]:
    ats_report = federal_plain_text_validation(resume_candidate)
    for warning in ats_report["warnings"]:
        print(f"FEDERAL ATS WARNING: {warning}")
    if ats_report["blockers"]:
        fail("Federal ATS plain-text validation failed:\n  " + "\n  ".join(str(item) for item in ats_report["blockers"]))
    assert_final_federal_coverage(resume_candidate, job_description, audit)
    for warning in audit.warnings:
        print(f"FEDERAL FIT WARNING: {warning}")
    final_page_count = page_count_for_docx(resume_candidate, temp_root)
    qualifications_page_count = page_count_for_docx(qualifications_candidate, temp_root)
    if final_page_count is not None and final_page_count != TARGET_PAGE_COUNT:
        fail(f"Final federal resume page-count validation failed: expected 2, found {final_page_count}.")
    if qualifications_page_count is not None and qualifications_page_count > MAX_QUALIFICATIONS_PAGE_COUNT:
        fail(
            "Final federal qualifications page-count validation failed: "
            f"maximum {MAX_QUALIFICATIONS_PAGE_COUNT}, found {qualifications_page_count}."
        )
    resume_render = render_checks.render_docx(resume_candidate, label=resume_render_label)
    qualifications_render = render_checks.render_docx(
        qualifications_candidate,
        label=qualifications_render_label,
    )
    if render_checks.RENDER_AVAILABLE and (resume_render is None or qualifications_render is None):
        fail("Federal render validation failed before publication.")
    return final_page_count, qualifications_page_count


def build_federal_resume(*, target_grade: str = "") -> FederalBuildResult:
    source = load_federal_source()
    job_description = validate_inputs()

    OUTPUT_DIR.mkdir(exist_ok=True)
    SCRATCH_DIR.mkdir(exist_ok=True)
    prompt_state = question_prep.load_application_prompt_state()
    question_context_issues = tuple(
        question_prep.application_question_context_issues(
            job_description,
            prompt_state,
            workflow="federal",
        )
    )

    temp_root = SCRATCH_DIR / f"christian_federal_resume_{uuid.uuid4().hex}"
    temp_root.mkdir(parents=True, exist_ok=False)
    try:
        plan = resume_plan(
            temp_root,
            source,
            job_description,
            target_grade=target_grade,
            question_context_issues=question_context_issues,
        )
        target_context = plan.audit.target_context
        if target_context is None:
            fail("Federal target context was not available after planning.")
        _question_draft_reasons, parse_draft_reasons, is_draft = federal_draft_channels(
            target_context,
            question_context_issues,
        )
        company_name = target_context.output_label or extract_federal_output_name(job_description)
        output_docx, qualifications_docx = federal_output_paths(company_name, is_draft=is_draft)

        resume_document = build_document(
            source,
            job_description,
            plan.resume_layout,
            plan.summary,
            plan.bullet_groups,
            plan.audit,
        )
        resume_candidate = temp_root / output_docx.name
        resume_document.save(str(resume_candidate))

        qualifications_document = build_qualifications_document(
            source,
            job_description,
            plan.qualifications_layout,
            plan.audit,
            plan.qualifications_content,
        )
        qualifications_candidate = temp_root / qualifications_docx.name
        qualifications_document.save(str(qualifications_candidate))
        if question_context_issues:
            question_prep.mark_docx_as_draft(qualifications_candidate, question_context_issues)

        final_page_count, qualifications_page_count = validate_staged_federal_documents(
            resume_candidate,
            qualifications_candidate,
            job_description,
            plan.audit,
            temp_root,
            resume_render_label=output_docx.stem,
            qualifications_render_label=qualifications_docx.stem,
        )

        print(
            "FEDERAL PARSE STATUS: "
            f"duty_grade={target_context.duty_grade or 'unparsed'} "
            f"selected_grade={target_context.target_grade or 'unparsed'} "
            f"available_grades={','.join(target_context.available_grades) or 'none'} "
            f"selected_requirements={len(target_context.requirements)} "
            f"verified={target_context.verified}"
        )
        for diagnostic in target_context.parse_diagnostics:
            level = "WARNING" if diagnostic.requires_draft else "INFO"
            print(f"FEDERAL PARSE {level} [{diagnostic.code}]: {diagnostic.message}")

        workflow_step_runner.publish_document_set(
            (
                (resume_candidate, output_docx),
                (qualifications_candidate, qualifications_docx),
            ),
            quarantine_root=SCRATCH_DIR / "run_logs" / "quarantine",
        )
    finally:
        shutil.rmtree(temp_root, ignore_errors=True)

    return FederalBuildResult(
        output_docx=output_docx,
        qualifications_docx=qualifications_docx,
        company_name=company_name,
        page_count=final_page_count,
        qualification_page_count=qualifications_page_count,
        layout_name=plan.resume_layout.name,
        selected_bullets=sum(len(group) for group in plan.bullet_groups),
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Christian Estrada's tailored federal DOCX resume.")
    parser.add_argument("--no-pdf", action="store_true", help="Accepted for clarity; PDFs are never created.")
    parser.add_argument(
        "--target-grade",
        type=federal_grade_argument,
        default="",
        metavar="GS-XX",
        help="Select one parsed federal qualification grade; defaults to the highest listed grade.",
    )
    return parser.parse_args()


def federal_grade_argument(value: str) -> str:
    if not value:
        return ""
    normalized = requirement_engine.normalize_federal_grade(value)
    if not normalized:
        raise argparse.ArgumentTypeError("target grade must use GS-XX syntax, for example GS-11")
    return normalized


def main() -> None:
    args = parse_args()
    result = build_federal_resume(target_grade=args.target_grade)
    print(f"Company: {result.company_name}")
    print(f"Output DOCX: {result.output_docx}")
    print(f"Qualifications DOCX: {result.qualifications_docx}")
    print(f"Page count: {result.page_count}")
    print(f"Qualifications page count: {result.qualification_page_count}")
    print(f"Layout: {result.layout_name}")
    print(f"Selected bullets: {result.selected_bullets}")


if __name__ == "__main__":
    main()
