#!/usr/bin/env python3
"""Generate a role-specific first 90 days plan."""

from __future__ import annotations

from pathlib import Path

from docx import Document

import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"


def read_job() -> str:
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip() if JOB_DESCRIPTION.exists() else ""


def latest_resume(job_description: str) -> Path | None:
    candidates = resume_analysis.matching_output_files(OUTPUT_DIR, job_description, "Resume.docx")
    return candidates[0] if candidates else None


def actions_for_lane(lane: str) -> tuple[list[str], list[str], list[str]]:
    first_30 = [
        "Map the current customer journey, implementation stages, handoffs, and decision owners.",
        "Review active projects, stalled accounts, open risks, and recent escalations.",
        "Learn the product workflow deeply enough to explain it in customer language.",
        "Meet Product, Support, Sales, Customer Success, and implementation stakeholders.",
        "Identify the reporting, data, and communication rhythms the team already trusts.",
        "Document the most common blockers and where ownership is unclear.",
    ]
    days_31_60 = [
        "Own a defined implementation or improvement workstream with visible milestones.",
        "Create a simple risk/status view that makes blockers, next steps, and owners easy to see.",
        "Improve one customer-facing handoff, checklist, or training asset.",
        "Validate data, testing, and adoption risks before they reach the customer as surprises.",
        "Share early observations with the manager and confirm priorities before changing process.",
    ]
    days_61_90 = [
        "Turn the strongest early improvement into a repeatable playbook or operating rhythm.",
        "Show measurable movement in implementation pace, adoption clarity, customer confidence, or issue resolution.",
        "Present a concise recommendation for the next process or customer-experience improvement.",
        "Build trust as someone who can own details while keeping the business outcome visible.",
    ]
    if lane == "analytics_operations":
        days_31_60.append("Build or improve one dashboard/report that helps leaders act faster.")
    if lane == "customer_success":
        days_31_60.append("Review account health signals and identify renewal or adoption risks before they escalate.")
    if lane == "presales_solution":
        days_31_60.append("Shadow discovery and demo workflows, then improve one reusable buyer-value story.")
    if lane == "process_improvement":
        days_31_60.append("Run a small current-state/root-cause review and pilot one controlled workflow improvement.")
    return first_30, days_31_60, days_61_90


def build_plan() -> Path:
    job_description = read_job()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")
    company = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role = resume_analysis.extract_job_title(job_description) or "Target Role"
    profile = resume_analysis.job_problem_profile(job_description)
    resume = latest_resume(job_description)
    first_30, days_31_60, days_61_90 = actions_for_lane(profile.primary_lane)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {output_target_name} First 90 Days.docx"
    doc = Document()
    doc.add_heading(f"First 90 Days - {company}", level=1)
    doc.add_paragraph(f"Role: {role}")
    doc.add_paragraph(f"Lane: {profile.lane_label}")
    doc.add_paragraph(f"Reference resume: {resume.name if resume else 'No matching generated resume found'}")

    for title, actions in (
        ("Days 1 to 30: Listen and Learn", first_30),
        ("Days 31 to 60: Own and Improve", days_31_60),
        ("Days 61 to 90: Prove and Scale", days_61_90),
    ):
        doc.add_heading(title, level=2)
        for action in actions:
            doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("Success Measures", level=2)
    for measure in profile.outcomes or ("implementation quality", "adoption", "risk reduction"):
        doc.add_paragraph(str(measure), style="List Bullet")

    doc.save(output)
    print(f"First 90 Days plan created: {output}")
    return output


def main() -> None:
    build_plan()


if __name__ == "__main__":
    main()
