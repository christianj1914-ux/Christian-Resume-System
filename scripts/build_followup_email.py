#!/usr/bin/env python3
"""Generate application follow-up email variants after a quiet period."""

from __future__ import annotations

import csv
from datetime import date, datetime
from pathlib import Path

from docx import Document

import interview_context
import job_search_guidance as guidance
import resume_analysis
from utils import assert_no_template_leakage, enforce_prose_quality


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
APPLICATIONS_CSV = PROJECT_ROOT / "scratch" / "applications.csv"
OUTPUT_DIR = PROJECT_ROOT / "output"
COMPANY_RESEARCH = PROJECT_ROOT / "jobs" / "company_research.txt"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig") if path.exists() else ""


def tracker_rows() -> list[dict[str, str]]:
    if not APPLICATIONS_CSV.exists():
        return []
    with APPLICATIONS_CSV.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def application_timing(company_name: str) -> tuple[str, str]:
    for row in reversed(tracker_rows()):
        if company_name.lower() in row.get("company", "").lower():
            applied_date = row.get("applied_date", "")
            if applied_date:
                try:
                    applied = datetime.strptime(applied_date, "%Y-%m-%d").date()
                    days = (date.today() - applied).days
                    return applied_date, f"{days} day(s) since applying"
                except ValueError:
                    return applied_date, "applied date is not in YYYY-MM-DD format"
    return "[add applied date]", "application date not found in tracker"


def supplied_context(company_name: str) -> str:
    return interview_context.company_research_context(
        PROJECT_ROOT / "jobs",
        company_name,
        company_research_path=COMPANY_RESEARCH,
    )


def followup_email_variants(
    company_name: str,
    role: str,
    profile: resume_analysis.JobProblemProfile,
    news_line: str = "",
) -> list[tuple[str, str]]:
    variants = [
        (
            "Brief Follow-Up",
            f"Subject: Following up on {role}\n\n"
            f"Hello,\n\nI wanted to follow up on my application for the {role} role at {company_name}. "
            "I remain interested in the opportunity and would welcome any update on next steps when the team has one.\n\n"
            "Best,\nChristian Estrada",
        ),
        (
            "Value-Forward Update",
            f"Subject: {role} application follow-up\n\n"
            f"Hello,\n\nI am checking in on my application for the {role} role at {company_name}. "
            f"My background continues to look like a strong match because the work connects to {profile.core_problem}, "
            "customer-facing delivery, workflow improvement, and measurable operating outcomes. "
            "I would appreciate any update on the hiring timeline when convenient.\n\nBest,\nChristian Estrada",
        ),
    ]
    if news_line:
        variants.append(
            (
                "News-Aware Follow-Up",
                f"Subject: Quick follow-up on the {role} role\n\n"
                f"Hello,\n\nI wanted to follow up on my application for the {role} role at {company_name}. "
                f"I also noticed {news_line}, which made the role feel even more relevant to the kind of implementation, reporting, and stakeholder-alignment work I have been doing. "
                "If useful, I can share additional detail on the implementation, reporting, and stakeholder-alignment work most relevant to the role.\n\nBest,\nChristian Estrada",
            )
        )
    return variants


def build_followup() -> Path:
    job_description = read_text(JOB_DESCRIPTION).strip()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")
    company = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role = resume_analysis.extract_job_title(job_description) or "the role"
    profile = resume_analysis.job_problem_profile(job_description)
    applied_date, timing_note = application_timing(company)
    company_context = supplied_context(company)
    news_line = interview_context.recent_company_news_line(company_context, company)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {output_target_name} Follow-Up Email.docx"
    doc = Document()
    doc.add_heading(f"Follow-Up Email - {company}", level=1)
    doc.add_paragraph(f"Role: {role}")
    doc.add_paragraph(f"Applied date: {applied_date}")
    doc.add_paragraph(f"Timing: {timing_note}")
    doc.add_paragraph(f"Detected lane: {profile.lane_label}")

    doc.add_heading("Timing Rules", level=2)
    for line in guidance.follow_up_timing_lines():
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Email Rules", level=2)
    for line in guidance.concise_email_rules():
        doc.add_paragraph(line, style="List Bullet")

    for title, body in followup_email_variants(company, role, profile, news_line):
        enforce_prose_quality(body, "followup_email_body", label=f"Follow-up email body ({title})")
        doc.add_heading(title, level=2)
        for paragraph in body.split("\n\n"):
            doc.add_paragraph(paragraph)

    if news_line:
        doc.add_heading("Verified Company Update Used", level=2)
        doc.add_paragraph(news_line)
    visible_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    assert_no_template_leakage(visible_text)
    doc.save(output)
    print(f"Follow-up email guide created: {output}")
    return output


def main() -> None:
    build_followup()


if __name__ == "__main__":
    main()
