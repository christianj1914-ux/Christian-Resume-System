#!/usr/bin/env python3
"""
Generate a standing general advice document from coaching notes.
Run when coaching insights change. This is not job-specific; it is a career
operating manual for Christian.
"""

from __future__ import annotations

import csv
from collections import Counter
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_interview_cheat_sheet as cheat
import build_resume
import job_search_guidance as guidance
import track_applications
from utils import optional_text


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
APPLICATIONS_CSV = PROJECT_ROOT / "scratch" / "applications.csv"
FONT = "Carlito"
BLUE = RGBColor(31, 78, 121)


def set_run_font(run, size: int = 10, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, 13 if level == 1 else 11, bold=True, color=BLUE)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, 10)
    paragraph.paragraph_format.space_after = Pt(2)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, 10)
    paragraph.paragraph_format.space_after = Pt(6)


def selected_coaching_lines(lane_key: str, job_description: str) -> list[str]:
    lines = cheat.answer_do_dont(job_description)
    if lane_key == "presales_solution":
        wanted = ("Lead top-down", "because", "Interest must be explicit")
        selected = [line for line in lines if any(term.lower() in line.lower() for term in wanted)]
    elif lane_key == "customer_success":
        wanted = ("Account structure", "calibration", "Brevity")
        selected = [line for line in lines if any(term.lower() in line.lower() for term in wanted)]
    else:
        selected = lines[:3]
    return (selected or lines[:3])[:3]


def add_active_job_context(document: Document) -> None:
    job_description = optional_text(JOB_DESCRIPTION).strip()
    if not job_description:
        add_body(
            document,
            "Add a job description to jobs/job_description.txt and rerun this script to include active job context.",
        )
        return

    selected_resume = build_resume.choose_resume(job_description)
    resume_text = "\n".join(cheat.paragraph_texts(selected_resume))
    profile = build_resume.job_problem_profile(job_description, resume_text)
    keywords = sorted(
        build_resume.audit_keywords(job_description),
        key=lambda keyword: (
            job_description.lower().count(keyword.lower()),
            len(keyword.split()),
            len(keyword),
            keyword,
        ),
        reverse=True,
    )[:8]
    stories = cheat.hero_stories(profile, job_description, resume_text)

    add_heading(document, "Active Job Context")
    add_heading(document, "Current Target", level=2)
    add_body(
        document,
        f"Detected lane: {profile.lane_label}. Core problem: {profile.core_problem}. Source resume: {selected_resume.name}.",
    )

    add_heading(document, "Top Keywords to Emphasize", level=2)
    for keyword in keywords:
        add_bullet(document, keyword)

    add_heading(document, "Recommended Lead Story", level=2)
    if stories:
        add_body(document, f"{stories[0].title}: {stories[0].result}")
    else:
        add_body(document, "No resume-supported lead story detected for this job description.")

    add_heading(document, "Lane-Specific Coaching", level=2)
    for line in selected_coaching_lines(profile.primary_lane, job_description):
        add_bullet(document, line)


def application_rows() -> list[dict[str, str]]:
    if not APPLICATIONS_CSV.exists():
        return []
    with APPLICATIONS_CSV.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def add_current_search_performance(document: Document) -> None:
    rows = application_rows()
    if len(rows) < 3:
        return
    status_counts = Counter(row.get("current_status", "draft") or "draft" for row in rows)
    lane_counts = Counter(track_applications.tracker_lane_label(row) or "Unknown" for row in rows)
    advanced_statuses = {"phone_screen", "interview", "final_round", "offer"}
    advanced = sum(1 for row in rows if row.get("current_status") in advanced_statuses)

    add_heading(document, "Current Search Performance")
    add_body(document, f"Tracked applications: {len(rows)}. Advanced past application stage: {advanced}.")
    add_heading(document, "Pipeline Breakdown", level=2)
    add_body(document, " | ".join(f"{status}: {count}" for status, count in sorted(status_counts.items())))
    add_heading(document, "Lane Mix", level=2)
    for lane, count in lane_counts.most_common():
        add_bullet(document, f"{lane}: {count}")
    if advanced == 0:
        add_body(document, "Current signal: no tracked application has advanced yet; review targeting and top-third proof before increasing volume.")
    else:
        add_body(document, "Current signal: compare advanced applications against applied-only applications to identify stronger lanes, company types, and proof points.")


def general_advice_sections() -> dict[str, list[str]]:
    sections = guidance.general_advice_sections()
    sections["Resume Strategy"] = [
        "Master resume is 3-4 pages and acts as the career database; the employer-facing resume should stay targeted and tight.",
        *sections["Resume Strategy"],
        "Company context lines explain what each employer does and its scale; they broaden the fit story without inventing anything.",
        "Bullet formula: Result + Metric + Context. The first bullet per role should behave like a hook, not a duty list.",
        "Relevancy beats recency. The best supported story wins regardless of how old it is.",
        "Two pages is fine. Cramming onto one page hurts readability and often hurts proof density.",
        "Boring and ordinary is right for structure. Content beats design every time.",
    ]
    sections["Job Search Cadence"] = [
        "Apply early when possible, but do not confuse urgency with volume for its own sake.",
        "Track inputs and outcomes separately so targeting issues do not get hidden behind activity.",
        "If results are flat, change one variable at a time: target role, resume positioning, outreach strategy, or interview delivery.",
        "Use alerts, shortlists, and direct outreach to create a repeatable weekly rhythm instead of rebuilding the search from scratch each morning.",
    ]
    return sections


def build_general_advice() -> Path:
    document = Document()
    for section in document.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)
    document.styles["Normal"].font.name = FONT

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Christian Estrada - Career Operating Manual")
    set_run_font(run, 16, bold=True, color=BLUE)

    add_body(
        document,
        "Standing guidance for resumes, interviews, networking, job-search cadence, and concise business communication.",
    )

    for section_title, bullets in general_advice_sections().items():
        add_heading(document, section_title)
        for item in bullets:
            add_bullet(document, item)

    add_active_job_context(document)
    add_current_search_performance(document)

    output_path = OUTPUT_DIR / "Christian Estrada - Career Operating Manual.docx"
    output_path.parent.mkdir(exist_ok=True)
    document.save(str(output_path))
    return output_path


def main() -> None:
    output_path = build_general_advice()
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
