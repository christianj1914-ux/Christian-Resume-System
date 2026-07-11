#!/usr/bin/env python3
"""Build an internal interview guide for the active target role."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build an internal interview guide.")
    parser.add_argument("--sample", action="store_true", help="Use sample internal context for testing.")
    return parser.parse_args()


def read_job() -> str:
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip() if JOB_DESCRIPTION.exists() else ""


def ask(label: str, default: str = "") -> str:
    try:
        value = input(f"{label}: ").strip()
    except (KeyboardInterrupt, EOFError):
        raise SystemExit("Internal interview guide canceled.")
    return value or default


def bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def sample_context() -> tuple[str, str]:
    return (
        "Current implementation/customer operations role with cross-functional delivery ownership",
        "Warm but not direct; hiring manager knows Christian's work through shared stakeholders",
    )


def lane_strategy(profile: resume_analysis.JobProblemProfile) -> list[str]:
    strategies = [
        f"Translate current performance into the new lane: {profile.lane_label}.",
        "Name the business problems already solved internally before making the move sound like a preference.",
        "Show respect for the current team by explaining transition coverage, documentation, and handoff discipline.",
        "Use internal knowledge as leverage: systems, stakeholders, customers, process gaps, and operating rhythms.",
        "Ask sharper questions than an external candidate because Christian can reference the company's actual workflows.",
        "Close on continuity: faster ramp, lower hiring risk, and a clearer path from current credibility to future impact.",
    ]
    if profile.primary_lane == "customer_success":
        strategies.append("Emphasize account health, adoption, retention risk, and customer trust as the through-line.")
    if profile.primary_lane == "analytics_operations":
        strategies.append("Emphasize reporting discipline, decision quality, process visibility, and cleaner operating handoffs.")
    if profile.primary_lane == "presales_solution":
        strategies.append("Emphasize discovery, buyer problem framing, product fluency, and credible handoff from pre-sale to delivery.")
    return strategies


def build_internal_guide(sample: bool = False) -> Path:
    job_description = read_job()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")

    company = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role = resume_analysis.extract_job_title(job_description) or "Target Role"
    profile = resume_analysis.job_problem_profile(job_description)
    if sample:
        current_role, relationship = sample_context()
    else:
        current_role = ask("Current role/team", "Current internal role/team")
        relationship = ask("Existing relationship with the hiring manager", "Relationship not yet established")

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {output_target_name} Internal Interview Guide.docx"
    doc = Document()
    doc.add_heading(f"Internal Interview Guide - {company}", level=1)
    doc.add_paragraph(f"Target role: {role}")
    doc.add_paragraph(f"Current role/team: {current_role}")
    doc.add_paragraph(f"Hiring-manager relationship: {relationship}")
    doc.add_paragraph(f"Positioning lane: {profile.lane_label}")

    doc.add_heading("Six Internal Interview Strategies", level=2)
    bullet_list(doc, lane_strategy(profile)[:6])

    doc.add_heading("Current-Team Transition Message", level=2)
    doc.add_paragraph(
        "I care about leaving my current team stronger, not creating a gap. If selected, I would document active work, identify ownership for open items, and build a clean handoff plan so the move protects the business while letting me contribute at the next level."
    )

    doc.add_heading("Internal Credibility Story", level=2)
    doc.add_paragraph(
        f"The strongest internal positioning is that Christian already understands the company environment and can bring proven delivery habits into {role}. "
        "Lead with examples where he clarified ownership, improved reporting, stabilized customer or stakeholder confidence, and turned ambiguous operational needs into usable process."
    )

    doc.add_heading("Relationship and Politics Plan", level=2)
    bullet_list(
        doc,
        [
            "Ask the current manager for advice before asking for advocacy, unless timing or confidentiality makes that risky.",
            "Frame the move as business contribution and growth, not dissatisfaction with the current team.",
            "Identify one sponsor, one neutral stakeholder, and one likely concern before the interview.",
            "Prepare a clean answer for why the move makes sense now: scope, impact, lane fit, and readiness.",
            "Avoid insider criticism. Convert every observed gap into a constructive improvement opportunity.",
        ],
    )

    doc.add_heading("Questions to Ask", level=2)
    questions = [
        f"What would success in the first 90 days look like for someone moving into {role} from inside the company?",
        "Where have internal transfers succeeded or struggled on this team?",
        "Which internal relationships would matter most for ramping quickly?",
        "What current workflow, customer, reporting, or handoff problem would you want this role to improve first?",
        "What concerns would you want me to address now so you can evaluate me clearly?",
    ]
    bullet_list(doc, questions)

    doc.add_heading("Closing Script", level=2)
    doc.add_paragraph(
        f"I am interested in this role because it would let me turn company knowledge and proven delivery habits into broader impact. "
        f"My current work has helped me understand the operating rhythms, stakeholders, and customer problems behind the work, and I would bring that context into {role} with a faster ramp and a clear transition plan. "
        "Is there anything about my background or the internal move that I should clarify before you make a decision?"
    )

    doc.add_heading("Prep Checklist", level=2)
    bullet_list(
        doc,
        [
            "Confirm whether the current manager should be told before the interview.",
            "Prepare two internal examples with situation, action, result, and business relevance.",
            "Write a transition plan for current responsibilities.",
            "Review the job description for unsupported requirements and prepare honest bridge language.",
            "Prepare one constructive internal observation that shows judgment without sounding critical.",
        ],
    )

    doc.save(output)
    print(f"Internal interview guide created: {output}")
    return output


def main() -> None:
    args = parse_args()
    build_internal_guide(args.sample)


if __name__ == "__main__":
    main()
