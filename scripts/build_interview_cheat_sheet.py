"""Build a Word-only interview cheat sheet from the current job description.

Rules:
- use the current job description and generated resume as source material
- save Word documents only to /output
- never create placeholders or PDFs
- never use LinkedIn page content as source material
- avoid double dashes and unsupported claims
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Mapping, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor

import build_cover_letter
import build_debrief_analysis as debrief_analysis
import build_resume
import business_context
import interview_context
import proof_text
import question_prep
import prose_engine
import resume_analysis
from text_safety import neutralize_conflicting_region_lists
from build_resume import COMPANY_APTEAN, COMPANY_EAST_WEST, COMPANY_HOME_DEPOT
from config.job_profiles import TARGETING_LANES
from config.language_rules import PLACEHOLDER_PATTERNS, remove_approved_bracketed_metadata
from config.paths import COMPANY_RESEARCH, INTERVIEW_NOTES, JOB_DESCRIPTION, OUTPUT_DIR, PROJECT_ROOT
import render_checks
from utils import (
    GREAT_EIGHT_OUTCOMES,
    assert_company_name_in_source,
    assert_no_template_leakage,
    clean_source_text,
    debug_print,
    enforce_prose_quality,
    fail,
    has_great_eight_signal,
    join_answer_sentences,
    optional_text,
    read_text,
)


RESUME_FONT = "Carlito"
NAME_BLUE = RGBColor(31, 78, 121)
SECTION_BLUE = RGBColor(31, 78, 121)
TITLE_SIZE = 18
SUBTITLE_SIZE = 10
SECTION_SIZE = 10
BODY_SIZE = 9.2

@dataclass(frozen=True)
class CheatSheetResult:
    company_name: str
    role_title: str
    resume_docx: Path
    output_docx: Path


@dataclass(frozen=True)
class StoryCard:
    title: str
    story_types: tuple[str, ...]
    hook: str
    takeaways: tuple[str, str, str]
    evidence: str
    level3_trait: str
    result: str
    outcome: str
    evidence_terms: tuple[str, ...]
    signals: tuple[str, ...]


@dataclass(frozen=True)
class InterviewQuestion:
    question: str
    angle: str


@dataclass(frozen=True)
class PreparedAnswer:
    prompt: str
    answer: str


@dataclass(frozen=True)
class QuestionIntentCard:
    prompt: str
    hidden_assessment: str
    bad_answer_trap: str
    story_angle: str
    clarify_when: str
    ideal_length: str


@dataclass(frozen=True)
class AnswerFrameworkSelection:
    primary_framework: str
    label: str
    reason: str
    guidance: tuple[str, ...]


@dataclass(frozen=True)
class InterviewPitchParts:
    opening: str
    mission: str
    proof: str
    workflow: str
    communication: str
    bridge: str
    pain_area: str
    skill_terms: tuple[str, ...]


def lower_clause(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip().rstrip(".")
    if not cleaned:
        return ""
    return cleaned[:1].lower() + cleaned[1:] if cleaned[:1].isupper() else cleaned


_ACTION_FRAGMENT_STARTS = {
    "built",
    "configured",
    "coordinated",
    "created",
    "delivered",
    "developed",
    "drove",
    "enabled",
    "facilitated",
    "improved",
    "increased",
    "kept",
    "launched",
    "led",
    "maintained",
    "managed",
    "mapped",
    "negotiated",
    "owned",
    "protected",
    "reduced",
    "stabilized",
    "supported",
    "tracked",
    "translated",
    "turned",
    "used",
    "validated",
}


def starts_with_action_fragment(text: str) -> bool:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return False
    first = re.sub(r"[^A-Za-z'-]", "", cleaned.split()[0]).lower()
    return bool(first) and (first in _ACTION_FRAGMENT_STARTS or first.endswith("ed"))


def story_evidence_sentence(text: str) -> str:
    cleaned = neutralize_conflicting_region_lists(re.sub(r"\s+", " ", text).strip().rstrip("."))
    if not cleaned:
        return ""
    if spoken_word_count(cleaned) > 28 and ":" in cleaned:
        lead, detail = (part.strip() for part in cleaned.split(":", 1))
        lead_sentence = story_evidence_sentence(lead)
        detail = re.sub(r"^a\s+", "It was a ", detail, flags=re.I)
        if re.search(r"\s+across coordination with\s+", detail, re.I):
            work, partners = re.split(r"\s+across coordination with\s+", detail, maxsplit=1, flags=re.I)
            return interview_join(lead_sentence, work, f"I coordinated with {partners}")
        return interview_join(lead_sentence, detail)
    lowered = cleaned.lower()
    if lowered.startswith("i "):
        return cleaned
    if starts_with_action_fragment(cleaned):
        return f"I {lower_clause(cleaned)}"
    return f"My role was to {lower_clause(cleaned)}"


def tighten_story_result_text(text: str) -> str:
    cleaned = proof_text.rewrite_dense_proof_patterns(
        neutralize_conflicting_region_lists(re.sub(r"\s+", " ", text).strip().rstrip("."))
    )
    replacements = (
        (
            r"\bkept core manufacturing and finance operations running across North America and Asia while improving the ERP system through training, testing, and release readiness\b",
            "kept core operations and finance workflows stable across North America and Asia during ERP improvement work",
        ),
        (
            r"\bimproved post-go-live follow-through, clearer issue ownership, and more reliable coordination across customer-facing teams\b",
            "improved post-go-live ownership and coordination across customer-facing teams",
        ),
        (
            r"\bhelped manufacturing clients across the Americas, Europe, and Asia move through full lifecycle ERP implementation with clearer scope and lower delivery risk\b",
            "helped international manufacturing clients move through ERP implementation with clearer scope and lower delivery risk",
        ),
    )
    for pattern, replacement in replacements:
        cleaned = re.sub(pattern, replacement, cleaned, flags=re.I)
    return cleaned


def story_result_sentence(text: str) -> str:
    cleaned = tighten_story_result_text(text)
    if not cleaned:
        return ""
    lowered = cleaned.lower()
    if lowered.startswith(("the result was ", "the work ", "it ", "this ")):
        return cleaned
    if starts_with_action_fragment(cleaned):
        return f"The work {lower_clause(cleaned)}"
    return f"The result was {lower_clause(cleaned)}"


def story_where_result_clause(text: str) -> str:
    cleaned = tighten_story_result_text(text)
    if not cleaned:
        return ""
    if starts_with_action_fragment(cleaned):
        return f"where the work {lower_clause(cleaned)}"
    return f"where the result was {lower_clause(cleaned)}"


def story_opener_by_type(
    story_types: set[str],
    story_title: str,
    company: str = "",
) -> str:
    """Natural, varied lead-ins into a story's hook sentence.

    These intentionally never announce the story as "an example" or "the
    clearest case" and never speak the story's internal title out loud -
    real spoken answers just start describing the situation. story_title is
    kept as a parameter for call-site compatibility but is unused here.
    """
    company_clause = f"At {company}, " if company else ""
    if "Challenge and Failure" in story_types:
        return f"{company_clause}there was a stretch where"
    if "Persuasion" in story_types or "Opposing Views" in story_types:
        return f"{company_clause}there was a moment where"
    if "Rapid Learning" in story_types:
        return f"{company_clause}early on, there was a point where"
    if "Managing and Leading" in story_types or "Teamwork" in story_types:
        return f"{company_clause}there was a project where"
    if "Analysis and Decision" in story_types:
        return f"{company_clause}there was a decision point where"
    if "Individual Achievement" in story_types:
        return f"{company_clause}there was a piece of work where"
    return f"{company_clause}there was a time where"


def story_natural_reference(card: StoryCard, company: str = "") -> str:
    """Natural, spoken-sounding reference to a story for short-answer
    contexts (the phone-screen script, elevator pitches, quick proof lines).

    States the situation and result as plain sentences instead of announcing
    the story as "a good/best example" or naming its internal title out loud.
    """
    opener = story_opener_by_type(set(card.story_types), card.title, company)
    spoken_opener = opener[:1].upper() + opener[1:]
    parts = [f"{spoken_opener} {lower_clause(card.hook)}"]
    if card.result:
        parts.append(story_result_sentence(card.result))
    return join_answer_sentences(*parts)


def spoken_level3_trait_sentence(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip().rstrip(".")
    if not cleaned:
        return ""
    lowered = cleaned.lower()
    if lowered.startswith("show what was noticed:"):
        detail = cleaned.split(":", 1)[1].strip()
        segments = re.split(r"\.\s*show what was done:\s*", detail, maxsplit=1, flags=re.I)
        noticed = segments[0]
        noticed_sentence = f"The key early signal was that {lower_clause(noticed)}"
        if len(segments) == 2 and segments[1]:
            return interview_join(noticed_sentence, story_evidence_sentence(segments[1]))
        return noticed_sentence
    if lowered.startswith("show what was noticed in the room:"):
        detail = lower_clause(cleaned.split(":", 1)[1].strip())
        return f"The key early signal was that {detail}"
    if lowered.startswith("show the constraint that made this hard:"):
        detail = cleaned.split(":", 1)[1].strip()
        detail = re.sub(r"\s+[—–]\s+", ". ", detail)
        detail = re.sub(r"\.\s+([a-z])", lambda match: f". {match.group(1).upper()}", detail)
        return interview_join("The constraint was clear", detail)
    if lowered.startswith("show how "):
        detail = lower_clause(cleaned[9:].strip())
        return f"The key early signal was that {detail}"
    if lowered.startswith("show the changed behavior:"):
        detail = lower_clause(cleaned.split(":", 1)[1].strip())
        return f"What changed afterward was {detail}"
    if lowered.startswith("show what was done:"):
        detail = cleaned.split(":", 1)[1].strip()
        return story_evidence_sentence(detail)
    if lowered.startswith("show "):
        detail = lower_clause(cleaned[5:].strip())
        return f"The key early signal was that {detail}"
    return f"The key early signal was {lower_clause(cleaned)}"


def sentence_texts(text: str) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return []
    return build_resume.summary_sentences(cleaned)


BRIEF_RECRUITER_WORD_RANGE = (25, 55)
STANDARD_SPOKEN_WORD_RANGE = (95, 140)
CAREER_WALKTHROUGH_WORD_RANGE = (120, 170)
SPOKEN_SENTENCE_TARGET = (15, 22)


def interview_join(*parts: str) -> str:
    """Join deliberately ordered interview sentences without changing shared prose helpers."""
    return join_answer_sentences(*parts)


def spoken_answer_metrics(text: str) -> dict[str, float | int]:
    sentences = sentence_texts(text)
    lengths = [spoken_word_count(sentence) for sentence in sentences if sentence.strip()]
    return {
        "words": spoken_word_count(text),
        "sentences": len(lengths),
        "average_sentence_words": (sum(lengths) / len(lengths)) if lengths else 0.0,
        "longest_sentence_words": max(lengths, default=0),
    }


def assert_spoken_answer_budget(
    label: str,
    text: str,
    *,
    minimum_words: int,
    maximum_words: int,
    minimum_average_sentence_words: int = SPOKEN_SENTENCE_TARGET[0],
    maximum_average_sentence_words: int = SPOKEN_SENTENCE_TARGET[1],
    maximum_sentence_words: int = 32,
) -> None:
    metrics = spoken_answer_metrics(text)
    words = int(metrics["words"])
    average = float(metrics["average_sentence_words"])
    longest = int(metrics["longest_sentence_words"])
    if words < minimum_words or words > maximum_words:
        fail(f"{label} must be between {minimum_words} and {maximum_words} spoken words; got {words}")
    if average < minimum_average_sentence_words or average > maximum_average_sentence_words:
        fail(
            f"{label} should average {minimum_average_sentence_words} to "
            f"{maximum_average_sentence_words} words per sentence; got {average:.1f}"
        )
    if longest > maximum_sentence_words:
        fail(f"{label} contains a {longest}-word sentence; maximum is {maximum_sentence_words}")


def take_sentences(text: str, max_sentences: int = 1) -> str:
    if max_sentences <= 0:
        return ""
    sentences = sentence_texts(text)
    return " ".join(sentences[:max_sentences]).strip()


def take_proof_sentence(text: str) -> str:
    sentences = sentence_texts(text)
    return next((sentence for sentence in sentences if PROOF_SIGNAL_RE.search(sentence)), sentences[0] if sentences else "")


def interview_workflow_sentence(
    profile: build_resume.JobProblemProfile,
    job_description: str = "",
) -> str:
    specialty = build_resume.role_specialty_phrase(job_description, candidate_problem_phrase(profile))
    lane_lines = {
        "analytics_operations": (
            "My usual pattern is to start with the decision that needs to be made, then work backward to the reporting, data checks, and workflow changes that make that decision usable."
        ),
        "implementation_delivery": (
            "My usual pattern is to map the workflow first, surface delivery risk early, and stay close to the people using the process so go-live is not the finish line."
        ),
        "customer_success": (
            "My usual pattern is to read risk early, tighten ownership, and turn customer friction into a next step both the account team and the customer can act on."
        ),
        "presales_solution": (
            "My usual pattern is to understand the buyer's real constraint first, then connect the recommendation to the workflow and implementation reality behind it."
        ),
        "change_enablement": (
            "My usual pattern is to start with the behavior that needs to change, then build the training, feedback, and reinforcement around it."
        ),
        "process_improvement": (
            f"My usual pattern is to map where {specialty} breaks down, test the fix with the people living in it, and measure whether the work actually gets easier."
        ),
        "corporate_strategy": (
            "My usual pattern is to make the tradeoffs visible early, then turn the analysis into a path leaders can actually execute."
        ),
    }
    return lane_lines.get(
        profile.primary_lane,
        "My usual pattern is to make the problem concrete quickly, keep the right people aligned, and move the work toward a usable decision."
    )


def interview_translation_sentence(profile: build_resume.JobProblemProfile) -> str:
    lane_lines = {
        "analytics_operations": "I am also used to translating the data story into language leaders can trust quickly enough to act on.",
        "implementation_delivery": "I am also used to translating technical work into clear status, risk, and ownership so the next decision does not get lost in project language.",
        "customer_success": "I am also used to translating issues into business impact so the conversation stays focused on value, trust, and next steps.",
        "presales_solution": "I am also used to translating technical tradeoffs into buyer language without losing the implementation reality behind them.",
        "change_enablement": "I am also used to translating change into practical day-to-day language so people know what is different and what to do next.",
        "process_improvement": "I am also used to translating process findings into clear tradeoffs, owners, and measurable next steps.",
        "corporate_strategy": "I am also used to translating analysis into decision language leaders can actually use.",
    }
    return lane_lines.get(
        profile.primary_lane,
        "I am also used to translating technical and operational work into clear language people can act on."
    )


def interview_role_bridge_sentence(
    profile: build_resume.JobProblemProfile,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
    resume_text: str = "",
) -> str:
    bridge = lane_interview_bridge(profile)
    specialty = build_resume.role_specialty_phrase(job_description, candidate_problem_phrase(profile))
    story_signal = ""
    if re.sub(r"\s+", " ", resume_text).strip():
        stories = hero_stories(profile, job_description, resume_text)
        top_story = stories[0] if stories else None
        if top_story and top_story.result:
            story_signal = story_result_sentence(top_story.result)
    lead = (
        f"For {company_name}, the {role_title} role is a strong fit because it needs someone who can {bridge}."
        if company_name and role_title
        else f"The fit is strong because the role needs someone who can {bridge}."
    )
    if story_signal:
        return join_answer_sentences(
            lead,
            "The same follow-through shows up in the strongest proof.",
            story_signal,
        )
    return join_answer_sentences(
        lead,
        f"It is the same follow-through pattern I have used in {specialty} work.",
    )


def interview_pitch_parts(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
) -> InterviewPitchParts:
    proof = lane_interview_proof_line(profile)
    if re.sub(r"\s+", " ", resume_text).strip():
        stories = hero_stories(profile, job_description, resume_text)
        top_story = stories[0] if stories else None
        if top_story and top_story.result:
            proof = story_natural_reference(top_story)
    pain_area = concise_pitch_pain_area(candidate_problem_phrase(profile))
    return InterviewPitchParts(
        opening=pitch_opening_line(profile, job_description),
        mission=pitch_career_arc_sentence(profile, job_description),
        proof=proof,
        workflow=interview_workflow_sentence(profile, job_description),
        communication=interview_translation_sentence(profile),
        bridge=interview_role_bridge_sentence(profile, company_name, role_title, job_description, resume_text),
        pain_area=pain_area,
        skill_terms=tuple(list(build_resume.jd_priority_phrases(job_description))[:4]),
    )


def motivation_note_line(notes_text: str) -> str:
    motivation_patterns = (
        r"\bdrawn to\b",
        r"\bcare about\b",
        r"\bmotivates me\b",
        r"\blove about\b",
        r"\bwhat keeps me\b",
        r"\bwhy i\b",
        r"\bwhat matters to me\b",
        r"\bmeaningful to me\b",
        r"\bpersonally\b",
    )
    for line in note_lines(notes_text):
        if not re.search(r"\b(i|me|my)\b", line, re.I):
            continue
        if any(re.search(pattern, line, re.I) for pattern in motivation_patterns):
            cleaned = line.strip().rstrip(".")
            if len(cleaned) >= 20:
                return f"{cleaned}."
    return ""


def human_motivation_sentence(
    profile: build_resume.JobProblemProfile,
    notes_text: str = "",
) -> str:
    supplied_line = motivation_note_line(notes_text)
    if supplied_line:
        return supplied_line
    lane_sentences = {
        "analytics_operations": (
            "What keeps me engaged is the moment when a business question stops being guesswork and becomes clear enough for someone to act on with confidence."
        ),
        "implementation_delivery": (
            "The part I keep coming back to is when a system stops being a project plan and starts working for the people who have to live in the workflow every day."
        ),
        "customer_success": (
            "What matters to me in customer work is the shift from a customer managing issues to a customer feeling confident enough to get real value from the platform."
        ),
        "presales_solution": (
            "I like the discovery side of this work because the best solution conversations start by getting honest about the workflow before anyone jumps to features."
        ),
        "change_enablement": (
            "What I have learned is that resistance is usually not about the tool itself. It is about whether people feel clear, prepared, and heard."
        ),
        "process_improvement": (
            "What keeps me engaged is finding the friction people have learned to tolerate and replacing it with a process that actually makes the day easier."
        ),
        "corporate_strategy": (
            "What keeps me engaged is making an ambiguous problem concrete enough that a team can act on it instead of continuing to debate it."
        ),
    }
    return lane_sentences.get(
        profile.primary_lane,
        "What keeps me engaged is working on problems where the technical answer is only half the solution and the human follow-through determines whether the work sticks."
    )


def pitch_opening_line(profile: build_resume.JobProblemProfile, job_description: str = "") -> str:
    return value_compression_opening(profile, job_description)


def pitch_career_arc_sentence(profile: build_resume.JobProblemProfile, job_description: str = "") -> str:
    return career_through_line(profile, job_description)


def pitch_role_target_line(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    pitch_parts: InterviewPitchParts | None = None,
) -> str:
    if company_name and role_title and pitch_parts and pitch_parts.pain_area:
        pain_area = concise_pitch_pain_area(pitch_parts.pain_area)
        return (
            f"For {company_name}, the {role_title} role stands out because it needs someone who can make "
            f"{pain_area} clearer and more usable."
        )
    if company_name and role_title:
        return (
            f"For the {role_title} role at {company_name}, the value I would bring is direct: "
            f"clearer workflow, tighter ownership, and outcomes people can use."
        )
    return (
        f"The value I bring is direct: clearer workflow, tighter ownership, and outcomes people can use in "
        f"{candidate_problem_phrase(profile)} work."
    )


def pitch_invitation_line(company_name: str, role_title: str) -> str:
    if company_name and role_title:
        return (
            f"Happy to go deeper on the most useful angle, whether that is the implementation work itself "
            f"or the fit for the {role_title} role at {company_name}."
        )
    return "Happy to go deeper on the most useful angle."


def natural_voice_opening(profile: build_resume.JobProblemProfile) -> str:
    lane_openings = {
        "implementation_delivery": "My background is in systems and implementation work, and the consistent thread has been making complex workflows usable",
        "customer_success": "My background is in systems and customer-facing work, and the consistent thread has been turning adoption risk into clearer outcomes",
        "presales_solution": "My background is in systems and solution work, and the consistent thread has been helping people make practical technology decisions",
        "analytics_operations": "My background is in systems and operational reporting, and the consistent thread has been turning messy information into usable decisions",
        "change_enablement": "My background is in systems and cross-functional change, and the consistent thread has been helping new processes hold up in practice",
        "process_improvement": "My background is in systems and operational improvement, and the consistent thread has been finding the process gap behind recurring problems",
        "corporate_strategy": "My background is in systems and cross-functional problem solving, and the consistent thread has been turning ambiguity into practical operating plans",
    }
    return lane_openings.get(
        profile.primary_lane,
        "My background is in systems and operational problem solving, and the consistent thread has been making complex work easier to use",
    )


def natural_voice_closing(profile: build_resume.JobProblemProfile, role_title: str = "") -> str:
    role_phrase = f" this kind of {role_title} work" if role_title else " this kind of work"
    identity = {
        "implementation_delivery": "working where systems, implementation, and day-to-day operations meet",
        "customer_success": "working where customer trust, adoption, and operational follow-through meet",
        "presales_solution": "working where discovery, technical judgment, and customer decisions meet",
        "analytics_operations": "working where systems, data, and operational decisions meet",
        "change_enablement": "working where systems, people, and lasting process change meet",
        "process_improvement": "finding the operational problem underneath the visible system issue",
        "corporate_strategy": "turning complex cross-functional questions into decisions people can act on",
    }.get(profile.primary_lane, "working where systems and operations meet")
    return interview_join(
        f"What ties my experience together is {identity}",
        f"It is also what makes{role_phrase} appealing to me",
    )


def candidate_problem_phrase(profile: build_resume.JobProblemProfile) -> str:
    return {
        "implementation_delivery": "complex ERP implementations moving cleanly from discovery through go-live and adoption",
        "customer_success": "customer adoption, risk, and renewal work staying steady under pressure",
        "presales_solution": "buyer questions turning into workable solution decisions",
        "analytics_operations": "systems, data, and workflow questions turning into decisions people can use",
        "change_enablement": "systems and process changes turning into real adoption",
        "process_improvement": "operational friction turning into measurable process improvement",
        "corporate_strategy": "complex cross-functional questions turning into decisions people can act on",
    }.get(profile.primary_lane, build_resume.natural_problem_phrase(profile))


def concise_pitch_pain_area(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text).strip()
    replacements = (
        (
            r"\btechnical scoping, integrations, data migration, and delivery risk\b",
            "technical scoping and delivery risk",
        ),
        (
            r"\btraining, stakeholder alignment, and adoption follow-through\b",
            "training and adoption follow-through",
        ),
    )
    for pattern, replacement in replacements:
        cleaned = re.sub(pattern, replacement, cleaned, flags=re.I)
    return cleaned


PROOF_SIGNAL_RE = re.compile(
    r"(?:\d|\$\d|%|\b(?:clients|users|sites|dashboards|reports|workshops|qbrs|revenue|"
    r"example|proof|built|led|owned|supported|reduced|improved|stabilized|facilitated|"
    r"delivered|enabled|protected)\b)",
    re.I,
)
INSTRUCTIONAL_OPENING_RE = re.compile(
    r"^(?:use|choose|answer like|lead with|do not|when asked|this is the same question)\b",
    re.I,
)
DEFERRED_OPENING_RE = re.compile(
    r"\b(?:let me|first i(?:'|’)ll|first i will|before i|get into|to give you context|"
    r"the background|one example)\b",
    re.I,
)
DIRECT_IDENTITY_OPENING_RE = re.compile(
    r"^(?:my background is|my experience is|my strongest area is|i(?:'|â€™)?d say my|i would say my)\b",
    re.I,
)
WARMUP_OPENING_RE = re.compile(
    r"^(?:sure|well|thanks for asking|that(?:'|â€™)?s a great question|good question)[.!?, ]*$",
    re.I,
)


def spoken_word_count(text: str) -> int:
    return len(re.findall(r"\b[\w+.#'-]+\b", text))


def pitch_single_sentence_summary(
    profile: build_resume.JobProblemProfile,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
) -> str:
    if company_name and role_title:
        bridge = lane_interview_bridge(profile)
        return f"For {company_name}, I would bring the ability to {bridge}."
    opening = pitch_opening_line(profile, job_description).strip()
    return opening if opening.endswith(".") else f"{opening}."


def claim_then_prove_issues(
    text: str,
    *,
    require_proof: bool = True,
    max_first_sentence_words: int = 28,
) -> list[str]:
    issues: list[str] = []
    sentences = sentence_texts(text)
    if not sentences:
        return ["answer is empty"]
    first = sentences[0]
    if spoken_word_count(first) > max_first_sentence_words:
        issues.append("opening sentence is too long")
    if first.count(",") > 3:
        issues.append("opening sentence stacks too many clauses")
    if WARMUP_OPENING_RE.match(first):
        issues.append("opening sentence is conversational filler instead of the answer")
    if DEFERRED_OPENING_RE.search(first) and not DIRECT_IDENTITY_OPENING_RE.search(first):
        issues.append("opening sentence delays the point instead of landing the claim first")
    if require_proof:
        if len(sentences) < 2:
            issues.append("answer needs a proof sentence after the opening claim")
        else:
            trailing = " ".join(sentences[1:])
            if not PROOF_SIGNAL_RE.search(trailing):
                issues.append("answer needs proof after the opening claim")
    return issues


def assert_claim_then_prove_answer(
    label: str,
    text: str,
    *,
    min_words: int = 18,
    require_proof: bool = True,
    max_first_sentence_words: int = 28,
) -> None:
    if spoken_word_count(text) < min_words:
        fail(f"{label} is too short to sound like a spoken answer")
    issues = claim_then_prove_issues(
        text,
        require_proof=require_proof,
        max_first_sentence_words=max_first_sentence_words,
    )
    if issues:
        fail(f"{label} failed claim-then-prove validation: {'; '.join(issues)}")


def assert_full_spoken_answer(label: str, text: str, *, min_words: int = 35) -> None:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if spoken_word_count(cleaned) < min_words:
        fail(f"{label} is too short to sound like a full spoken answer")
    if INSTRUCTIONAL_OPENING_RE.search(cleaned):
        fail(f"{label} still reads like coaching instructions instead of a spoken answer")


def pitch_variants(
    profile: build_resume.JobProblemProfile,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> dict[str, str]:
    pitch_parts = interview_pitch_parts(profile, company_name, role_title, job_description, resume_text)
    claim = natural_voice_opening(profile)
    career_arc = pitch_career_arc_sentence(profile, job_description)
    human = human_motivation_sentence(profile, notes_text)
    proof_candidate = pitch_parts.proof if pitch_parts and pitch_parts.proof else ""
    proof = proof_candidate if proof_candidate and PROOF_SIGNAL_RE.search(proof_candidate) else lane_interview_proof_line(profile)
    concise_role_target = (
        f"For {company_name}, that gives me a practical foundation for the {role_title} work"
        if company_name and role_title
        else "That gives me a practical foundation for this work"
    )
    identity_close = natural_voice_closing(profile, role_title)
    thirty_second = interview_join(claim, proof)
    if spoken_word_count(thirty_second) < 40:
        thirty_second = interview_join(claim, proof, concise_role_target)
    elif spoken_word_count(thirty_second) > 80:
        thirty_second = interview_join(claim, take_proof_sentence(proof), concise_role_target)

    sixty_second = interview_join(claim, career_arc, proof, identity_close)
    if spoken_word_count(sixty_second) < STANDARD_SPOKEN_WORD_RANGE[0]:
        sixty_second = interview_join(claim, career_arc, proof, concise_role_target, identity_close)
    elif spoken_word_count(sixty_second) > STANDARD_SPOKEN_WORD_RANGE[1]:
        sixty_second = interview_join(claim, career_arc, take_proof_sentence(proof), concise_role_target, identity_close)

    ninety_second = interview_join(
        claim,
        human,
        proof,
        interview_workflow_sentence(profile, job_description),
        concise_role_target,
        identity_close,
    )
    if spoken_word_count(ninety_second) < 120:
        ninety_second = interview_join(claim, career_arc, human, proof, concise_role_target, identity_close)
    elif spoken_word_count(ninety_second) > 170:
        ninety_second = interview_join(
            claim,
            human,
            take_proof_sentence(proof),
            interview_workflow_sentence(profile, job_description),
            concise_role_target,
            identity_close,
        )
    variants = {
        "15_second": claim,
        "30_second": thirty_second,
        "60_second": sixty_second,
        "90_second": ninety_second,
    }
    assert_claim_then_prove_answer(
        "15-second pitch",
        variants["15_second"],
        min_words=10,
        require_proof=False,
        max_first_sentence_words=28,
    )
    assert_claim_then_prove_answer("30-second pitch", variants["30_second"], min_words=18)
    assert_claim_then_prove_answer("60-second pitch", variants["60_second"], min_words=30)
    assert_claim_then_prove_answer("90-second pitch", variants["90_second"], min_words=45)
    assert_spoken_answer_budget(
        "15-second pitch",
        variants["15_second"],
        minimum_words=10,
        maximum_words=25,
        maximum_sentence_words=28,
    )
    assert_spoken_answer_budget(
        "30-second pitch",
        variants["30_second"],
        minimum_words=40,
        maximum_words=80,
    )
    assert_spoken_answer_budget(
        "60-second pitch",
        variants["60_second"],
        minimum_words=STANDARD_SPOKEN_WORD_RANGE[0],
        maximum_words=STANDARD_SPOKEN_WORD_RANGE[1],
    )
    assert_spoken_answer_budget(
        "90-second pitch",
        variants["90_second"],
        minimum_words=120,
        maximum_words=170,
    )
    return variants


def human_layer_reference_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    notes_text: str = "",
) -> list[str]:
    pitch_parts = interview_pitch_parts(profile, company_name, role_title, job_description, resume_text)
    lines = [f"Human line to keep in reserve: {human_motivation_sentence(profile, notes_text)}"]
    if pitch_parts and pitch_parts.communication:
        lines.append(f"Translation line for longer answers: {pitch_parts.communication}")
    if pitch_parts and pitch_parts.workflow:
        lines.append(f"Operating line for longer answers: {pitch_parts.workflow}")
    lines.append(
        "Long-answer pattern: after the result, add one sentence about what changed for the people living in the workflow before you bridge back to the role."
    )
    return lines


def story_human_connection_line(
    card: StoryCard,
    profile: build_resume.JobProblemProfile,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
    resume_text: str = "",
) -> str:
    pitch_parts = interview_pitch_parts(profile, company_name, role_title, job_description, resume_text)
    pain_area = pitch_parts.pain_area if pitch_parts and pitch_parts.pain_area else candidate_problem_phrase(profile)
    story_types = set(card.story_types)
    signal_text = " ".join(card.signals).lower()
    theme_key = story_theme_key(card)
    if "Challenge and Failure" in story_types:
        return "The human side of that story was rebuilding confidence after something had gone sideways, not just fixing the process."
    if "Persuasion" in story_types or "Opposing Views" in story_types:
        return "The human side was understanding what each group was protecting so the next step felt credible to everyone involved."
    if any(term in signal_text for term in ("messaging", "workflow", "automation", "chatbot", "nlp")):
        return "The human side of that work was making a new channel feel understandable and consistent for the people using it and the customers receiving it."
    if theme_key == "dashboards":
        return f"The human side of that work was making {pain_area} clear enough that people could act instead of debating the inputs."
    if theme_key == "learning" or "Rapid Learning" in story_types:
        return "The human side was becoming useful fast enough that the team could trust my judgment while I was still ramping."
    if "Analysis and Decision" in story_types:
        return f"The human side of that work was making {pain_area} clear enough that people could act instead of debating the inputs."
    if "Managing and Leading" in story_types or "Teamwork" in story_types:
        return "The human side was giving different stakeholders a shared path forward instead of letting each group stay in its own version of the problem."
    return f"The human side of that work was making {pain_area} easier for people to live with day to day, not just technically complete."


def lane_interview_proof_line(profile: build_resume.JobProblemProfile) -> str:
    if profile.primary_lane == "customer_success":
        return "I have supported 80+ client engagements, led 60+ executive workshops and QBRs, and helped stabilize $1M+ in at-risk annual revenue."
    if profile.primary_lane == "analytics_operations":
        return "I have built 200+ dashboards and KPI reporting tools, supported 80+ client engagements, and led 60+ executive workshops and QBRs."
    if profile.primary_lane == "presales_solution":
        return "I have supported 80+ client engagements, led 60+ executive workshops and QBRs, and kept recommendations credible after implementation."
    if profile.primary_lane == "change_enablement":
        return "I have built role-based training, clearer reporting, and stakeholder rhythms that made system change easier to adopt."
    return "I have supported 80+ client engagements, built 200+ reporting tools, and worked across five sites and 150+ users."


def lane_interview_bridge(profile: build_resume.JobProblemProfile) -> str:
    if profile.primary_lane == "analytics_operations":
        return "make the data useful, keep the workflow clear, and turn analysis into decisions leaders can act on"
    if profile.primary_lane == "customer_success":
        return "make risk visible early, keep ownership clear, and turn the work into retained trust and expansion confidence"
    if profile.primary_lane == "presales_solution":
        return "run sharper discovery, make tradeoffs visible early, and keep the recommendation credible after the sale"
    if profile.primary_lane == "change_enablement":
        return "make the workflow usable, reduce resistance, and leave the team with adoption that sticks"
    if profile.primary_lane == "implementation_delivery":
        return "lead technical scoping, surface delivery risk early, and turn requirements into workable implementation plans"
    return "make the workflow clear, keep the right people aligned, and turn the result into measurable progress"


def scrub_document_for_job_language(document: Document, job_description: str) -> None:
    if document is None:
        fail("interview document is None; cannot scrub language")
    if not hasattr(document, 'paragraphs'):
        fail("interview document is missing paragraphs attribute; document structure is invalid")
    if not job_description or not isinstance(job_description, str):
        fail("job description is empty or invalid; cannot perform language scrubbing")
    
    try:
        if build_resume.jd_explicitly_requires_erp(job_description):
            return
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                cleaned = build_resume.scrub_erp_language_for_non_erp_text(run.text, job_description)
                if cleaned != run.text:
                    run.text = cleaned
    except Exception as e:
        fail(f"error while scrubbing interview document for job language: {type(e).__name__}: {e}")



def note_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in clean_source_text(text).splitlines():
        line = re.sub(r"^[\-*•#\d.\s]+", "", raw).strip()
        line = line.strip('"').strip()
        if len(line) >= 8:
            lines.append(line)
    return lines


def supplied_company_background_lines(company_name: str, supplied_text: str) -> list[str]:
    lines: list[str] = []
    wanted_prefixes = (
        "What it is", "Founder/CEO", "Stage", "Model", "Mission/Philosophy",
        "Primary verticals", "Sweet spot", "Core Product Capabilities",
        "Competitive Differentiators", "Where the Product Is Heading",
        "What to Know About the TAM Role", "CEO", "Founded", "Industry",
        "Company size", "Specialties",
    )
    wanted_terms = (
        "HIPAA", "EMR", "EHR", "natural language", "one-click analysis",
        "reverse ETL", "predictive analytics", "multi-location", "healthcare",
        "aesthetics", "data fragmentation", "LTV", "CAC", "ROI", "onboarding",
        "standardization", "AI analytics", "business intelligence", "600+",
    )
    for line in note_lines(supplied_text):
        if company_name.lower() not in supplied_text.lower() and "corral" in line.lower():
            continue
        if line.lower() in {"industry", "company size", "specialties", "locations", "primary", "headquarters", "website"}:
            continue
        if line.endswith(":") and len(line.split()) <= 6:
            continue
        if line.startswith(wanted_prefixes) or any(term.lower() in line.lower() for term in wanted_terms):
            if len(line) > 210:
                line = line[:207].rstrip() + "..."
            if line not in lines:
                lines.append(line)
        if len(lines) >= 10:
            break
    return lines


def post_round_intelligence_lines(supplied_text: str) -> list[str]:
    cleaned = clean_source_text(supplied_text)
    lowered = cleaned.lower()
    lines: list[str] = []
    if re.search(r"\bspectra\s+path\b", cleaned, re.I):
        lines.append("Company intelligence to use: product-specific implementation work came up in notes. Ask how workflow, data, and issue volume differ across customer use cases.")
    if re.search(r"\bjira\b", cleaned, re.I):
        lines.append("Tool signal to use: Jira came up in notes. Prepare to explain issue tracking by owner, business impact, priority, dependency, and customer-facing update.")
    if re.search(r"\bazure\b", cleaned, re.I):
        lines.append("Tool signal to use: Azure came up in notes. Keep the bridge honest: cloud-adjacent collaboration, documentation, ticket context, and technical handoffs, not unsupported Azure administration.")
    if re.search(r"\b(?:hipaa|phi|protected health information)\b", lowered, re.I):
        lines.append("Compliance signal to prepare: for HIPAA/PHI, do not claim direct compliance ownership. Bridge to documentation discipline, validation, access awareness, escalation judgment, and regulated-workflow care.")
    if re.search(r"\bexcel\b", cleaned, re.I):
        lines.append("Round feedback to prepare: Excel comfort should connect to reporting, data validation, exports, analysis, and the 200+ dashboard/reporting-tools proof point.")
    if re.search(r"\btrain(?:ing)? others\b|role-based|workshops?|qbr", cleaned, re.I):
        lines.append("Round feedback to prepare: training answers should use role-based enablement, 60+ workshops/QBRs, customer adoption, and clear workflow instruction.")
    if re.search(r"\btechnical projects?\b|integrations?|api|data migration|sql", cleaned, re.I):
        lines.append("Round feedback to prepare: technical-project answers should use implementation lifecycle, data migration, SQL validation, integrations, issue tracking, and cross-functional technical handoffs.")
    return lines[:8]


def supplied_smart_questions(supplied_text: str) -> list[str]:
    questions: list[str] = []
    for match in re.finditer(r'"([^"]+\?)"', clean_source_text(supplied_text)):
        question = match.group(1).strip()
        if question not in questions:
            questions.append(question)
    if len(questions) < 8:
        for line in note_lines(supplied_text):
            if line.endswith("?") and len(line) <= 180 and line not in questions:
                questions.append(line)
    weak_patterns = re.compile(r"\b(culture like|day in the life|work-life balance)\b", re.I)
    return [question for question in questions if not weak_patterns.search(question)][:14]


NEGATIVE_QUESTION_PATTERNS = (
    (r"\bwhat are the biggest challenges\b", "What does the team need to improve most in the next year?"),
    (r"\bwhat are the main problems\b", "What is the one problem you would most want this hire to solve?"),
    (r"\bwhat is not working\b", "What would working well look like in that area?"),
    (r"\bwhy do people leave\b", "What makes people who have been successful here want to stay?"),
    (r"\bwhat went wrong\b", "What did the team learn from that experience?"),
    (r"\bwhat are the weaknesses\b", "Where does the team see the most opportunity to improve?"),
    (r"\bwhat is lacking\b", "What capability would make the biggest difference right now?"),
    (r"\bwhy hasn't\b", "What would need to be true for that to change?"),
    (r"\bwhy is this role open\b", "What created the opportunity for this role right now?"),
    (r"\bwhat do employees complain about\b", "What does the team most want to see improve?"),
    (r"\bwhat keeps you up at night\b", "What is the one problem you are most focused on solving right now?"),
)


def audit_question_framing(question: str) -> tuple[bool, str]:
    for pattern, reframe in NEGATIVE_QUESTION_PATTERNS:
        if re.search(pattern, question, re.I):
            return True, reframe
    return False, ""


def reframe_questions_to_positive(questions: list[str]) -> tuple[list[str], list[tuple[str, str]]]:
    cleaned: list[str] = []
    changes: list[tuple[str, str]] = []
    for question in questions:
        needs_reframe, reframe = audit_question_framing(question)
        if needs_reframe:
            cleaned.append(reframe)
            changes.append((question, reframe))
        else:
            cleaned.append(question)
    return cleaned, changes


def log_reframed_questions(section: str, changes: list[tuple[str, str]]) -> None:
    for original, reframe in changes:
        debug_print(
            f"[questions] {section}: reframed '{original[:60]}' -> '{reframe[:60]}'",
            flag="DEBUG_INTERVIEW_GUIDE",
        )


def interview_question_quality_warnings(questions: list[str]) -> list[str]:
    warnings = list(business_context.question_quality_warnings(questions))
    for question in questions:
        needs_reframe, reframe = audit_question_framing(question)
        if needs_reframe:
            warnings.append(
                f"Negative framing: '{question[:60]}' -> Suggested: '{reframe[:60]}'"
            )
    return warnings


def paragraph_texts(docx_path: Path) -> list[str]:
    document = Document(str(docx_path))
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def set_default_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(30)
    section.bottom_margin = Pt(30)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    style = document.styles["Normal"]
    style.font.name = RESUME_FONT
    style.font.size = Pt(BODY_SIZE)


def add_title(document: Document, company_name: str, role_title: str) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Christian Estrada")
    run.bold = True
    run.font.name = RESUME_FONT
    run.font.size = Pt(TITLE_SIZE)
    run.font.color.rgb = NAME_BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Interview Cheat Sheet | {company_name} | {role_title}")
    run.font.name = RESUME_FONT
    run.font.size = Pt(SUBTITLE_SIZE)

    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = generated.add_run(date.today().strftime("%B %#d, %Y") if sys.platform == "win32" else date.today().strftime("%B %-d, %Y"))
    run.font.name = RESUME_FONT
    run.font.size = Pt(BODY_SIZE)


def add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.name = RESUME_FONT
    run.font.size = Pt(SECTION_SIZE)
    run.font.color.rgb = SECTION_BLUE


def add_subsection(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = RESUME_FONT
    run.font.size = Pt(BODY_SIZE)
    run.font.color.rgb = SECTION_BLUE


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = RESUME_FONT
    run.font.size = Pt(BODY_SIZE)


def add_body_lines(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(line)
        run.font.name = RESUME_FONT
        run.font.size = Pt(BODY_SIZE)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Pt(14)
    paragraph.paragraph_format.first_line_indent = Pt(-7)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(chr(8226) + " ")
    run.font.name = RESUME_FONT
    run.font.size = Pt(BODY_SIZE)
    body = paragraph.add_run(text)
    body.font.name = RESUME_FONT
    body.font.size = Pt(BODY_SIZE)


def mentions(text: str, *terms: str) -> bool:
    lowered = text.lower()
    for term in terms:
        normalized = term.lower()
        if len(normalized) <= 3 and re.search(rf"\b{re.escape(normalized)}\b", lowered):
            return True
        if len(normalized) > 3 and normalized in lowered:
            return True
    return False


def contains_all(text: str, fragments: tuple[str, ...]) -> bool:
    lowered = text.lower()
    return all(fragment.lower() in lowered for fragment in fragments)


def detect_panel_context(job_description: str, interview_notes: str) -> bool:
    combined_text = f"{job_description}\n{interview_notes}".lower()
    panel_signals = (
        "panel interview",
        "multiple interviewers",
        "interview panel",
        "group interview",
        "hiring committee",
        "panel round",
    )
    if any(signal in combined_text for signal in panel_signals):
        return True

    title_keywords = re.compile(
        r"\b("
        r"hiring manager|manager|director|senior director|vp|vice president|"
        r"head|lead|team lead|recruiter|hr|human resources|"
        r"product|operations|engineering|finance|sales|marketing|"
        r"customer success|implementation|support|analytics|data"
        r")\b",
        re.I,
    )
    name_with_title = re.compile(
        r"^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*(?:,|-|\u2013|\u2014|\||:|\()\s*(.+?)\)?$"
    )
    panelists: set[str] = set()
    for line in interview_notes.splitlines():
        match = name_with_title.match(line.strip())
        if match and title_keywords.search(match.group(2)):
            panelists.add(match.group(1).strip())
    return len(panelists) > 1


def panel_composition_type(job_description: str, interview_notes: str) -> str:
    notes_lower = interview_notes.lower()
    combined_text = f"{job_description}\n{interview_notes}"
    if "teammates" in notes_lower or "peer review" in notes_lower:
        return "all_peers"
    if "hiring manager" in notes_lower and "team members" in notes_lower:
        return "manager_plus_reports"
    if re.search(r"\bhr\b|\bhuman resources\b", notes_lower, re.I) and re.search(r"\bsenior leadership\b|\bvp\b", notes_lower, re.I):
        return "hr_plus_leadership"
    department_patterns = (
        r"\bhr\b",
        r"\bhuman resources\b",
        r"\bfinance\b",
        r"\boperations\b",
        r"\bengineering\b",
        r"\bproduct\b",
        r"\bsales\b",
        r"\bmarketing\b",
        r"\bcustomer success\b",
        r"\bimplementation\b",
        r"\bsupport\b",
        r"\blegal\b",
        r"\bit\b",
        r"\bdata\b",
        r"\banalytics\b",
        r"\brevenue\b",
        r"\bpeople\b",
        r"\brecruiting\b",
    )
    department_count = sum(1 for pattern in department_patterns if re.search(pattern, combined_text, re.I))
    if department_count >= 2:
        return "cross_functional"
    return "cross_functional"


def detect_late_stage_context(interview_notes: str, job_description: str) -> bool:
    interview_note_signals = (
        "offer",
        "compensation",
        "final round",
        "start date",
        "background check",
        "references",
    )
    job_description_signals = interview_note_signals + (
        "total compensation",
        "total rewards",
        "salary transparency",
        "pay range",
        "compensation range",
        "equity",
    )
    detected = mentions(interview_notes, *interview_note_signals) or mentions(
        job_description,
        *job_description_signals,
    )
    if detected:
        debug_print(
            "Negotiation prep section included: late-stage signals detected in job description or interview notes.",
            flag="DEBUG_INTERVIEW_GUIDE",
        )
    return detected


def panel_composition_strategy(composition_type: str) -> str:
    strategies = {
        "all_peers": "All-peers panel: earn credibility as a future collaborator. Keep answers practical, avoid sounding above the work, and show how Christian shares information, removes friction, and helps peers succeed.",
        "manager_plus_reports": "Manager plus direct reports: give the manager confidence in ownership while showing respect for the people closest to the day-to-day work. Balance executive-ready outcomes with team-level listening.",
        "cross_functional": "Cross-functional panel: translate every answer across business, technical, customer, and operational concerns. Name the shared outcome, then show how Christian keeps different groups aligned.",
        "hr_plus_leadership": "HR plus leadership: stay concise on motivation, fit, and risk while using executive-ready proof. Show why the move makes sense, how Christian learns quickly, and what measurable value he can bring.",
    }
    return strategies.get(composition_type, strategies["cross_functional"])


def panel_three_reason_why_company(company_name: str, role_title: str) -> str:
    return (
        "Use three reasons. One: name the industry or technology value Christian has seen firsthand, grounded in resume-supported experience. "
        f"Two: connect {company_name}'s differentiator to a specific recent signal from the job description, company research, or interview notes, such as a leadership hire, product update, or growth rate. "
        f"Three: name a culture observation with a specific detail tied to {company_name}, the {role_title} role, or the panel conversation."
    )


def why_company_three_reason_answer(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    job_description: str,
    supplied_context: str = "",
) -> str:
    industry_reference = (profile.lane_label or "customer-facing systems work").lower()
    specialty = build_resume.role_specialty_phrase(job_description, "")
    if not specialty:
        specialty = "customer-facing systems and adoption work"
    role_specialties = build_resume.visible_role_specialties(job_description)
    if len(role_specialties) >= 3:
        specialties_phrase = f"{role_specialties[0]}, {role_specialties[1]}, and {role_specialties[2]}"
    elif len(role_specialties) == 2:
        specialties_phrase = f"{role_specialties[0]} and {role_specialties[1]}"
    elif role_specialties:
        specialties_phrase = role_specialties[0]
    else:
        specialties_phrase = ""
    company_differentiator = f"its focus on {specialties_phrase}" if specialties_phrase else "the way the role connects customer problems, technology, and business outcomes"
    culture_observation = build_resume.visible_values_phrase(job_description)
    if not culture_observation:
        culture_observation = "the customer-focused and collaborative language in the posting"
    customer_benefit = {
        "customer_success": "customers get measurable value after launch",
        "presales_solution": "customers understand the right solution before they commit",
        "implementation_delivery": "customers move from requirements to adoption with less delivery risk",
        "analytics_operations": "teams turn operating data into clearer decisions",
        "change_enablement": "teams adopt new workflows with less disruption",
        "process_improvement": "teams improve cycle time, quality, and service outcomes",
    }.get(profile.primary_lane, "customers turn complex system work into practical outcomes")
    supplied_lines = supplied_company_background_lines(company_name, supplied_context)
    if supplied_lines:
        recent_signal = re.sub(r"\s+", " ", supplied_lines[0]).strip()
        if len(recent_signal) > 120:
            recent_signal = recent_signal[:117].rstrip() + "..."
    else:
        recent_signal = ""
    recent_signal_clause = f" and {recent_signal}" if recent_signal else ""
    return (
        f"There are three reasons I was excited to see this role at {company_name} open up. "
        f"One, throughout my time in {industry_reference}, I have seen firsthand how valuable {specialty} can be "
        f"at helping {customer_benefit}, which is what initially drew my interest. "
        f"Two, I am particularly interested in {company_name} because of {company_differentiator}{recent_signal_clause}. "
        f"Three, from what I can see about your culture, such as {culture_observation}, {company_name} appears to have the values I am looking for."
    )


def add_panel_interview_section(document: Document, company_name: str, role_title: str, composition_type: str) -> None:
    add_section(document, "Panel Interview Playbook")

    add_subsection(document, "Panel composition read")
    add_bullet(document, panel_composition_strategy(composition_type))

    add_subsection(document, "Panelist research checklist")
    for item in (
        "LinkedIn tenure at the company",
        "Multiple departments or roles held inside the company",
        "Industry context and the two or three closest competitors",
        "One likely pain point, operating risk, or solution hypothesis worth testing in the interview",
        "Industry switches before joining the company",
        "Education background and field of study",
        "Volunteer work or community involvement",
        "Recent LinkedIn posts",
        "Published articles or shared thought leadership",
        "Mutual connections or shared professional communities",
    ):
        add_bullet(document, item)

    add_subsection(document, "Eye contact distribution technique")
    add_body(
        document,
        "Answer the questioner directly for 20-30 seconds, then pan to each panelist while finishing the answer. Never neglect junior members. Lock in, make the point, and move on.",
    )

    add_subsection(document, "Three-reason answer: Why do you want to work here?")
    add_body(document, panel_three_reason_why_company(company_name, role_title))

    add_subsection(document, "Individual thank-you email rule")
    add_body(
        document,
        "Send one unique email per panelist within 24 hours. Four-sentence template: Thank them for their time and name the role. Reference one specific moment from their question or comment. Connect that moment to one resume-supported proof point. Close with continued interest and appreciation.",
    )

    add_subsection(document, "Water and note-taking pacing technique")
    add_body(
        document,
        "Use note-taking and water drinking as pacing mechanisms before answering or between answer sections. Panelists will not interrupt either action.",
    )

    add_subsection(document, "Mandatory closing question")
    add_body(document, 'Ask before leaving: "Where do we go from here?"')


def add_virtual_panel_adaptation_table(document: Document) -> None:
    add_subsection(document, "Virtual panel adaptation table")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Panel dynamic"
    header_cells[1].text = "Video adaptation"
    rows = (
        ("Room eye contact", "Use the camera lens as the substitute for room panning after acknowledging the questioner by name."),
        ("Opening presence", "Address each person by name when logging on, then let the interviewer lead the order."),
        ("Visible gestures", "Keep gestures in frame between waist and neck so emphasis reads clearly on screen."),
        ("Note-taking", "Say briefly that Christian is taking a note, look down intentionally, then return to the camera before answering."),
    )
    for panel_dynamic, video_adaptation in rows:
        cells = table.add_row().cells
        cells[0].text = panel_dynamic
        cells[1].text = video_adaptation
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.name = RESUME_FONT
                    run.font.size = Pt(BODY_SIZE)


def signal_score(job_description: str, signals: tuple[str, ...]) -> int:
    lowered = job_description.lower()
    return sum(1 for signal in signals if signal.lower() in lowered)


def fit_label(profile: build_resume.JobProblemProfile, resume_docx: Path) -> str:
    audit_state = resume_analysis.output_audit_state(resume_docx)
    if audit_state == "POOR":
        return "Poor Fit"
    if audit_state == "FAIL":
        return "Stretch Fit"
    if audit_state == "BRIDGE":
        return "Bridge Fit"
    if len(profile.direct_matches) >= 4 and not profile.unsupported_requirements:
        return "Strong Fit"
    if len(profile.direct_matches) >= 2:
        return "Adjacent Fit"
    return "Stretch Fit"


def adjusted_profile_for_role(
    profile: build_resume.JobProblemProfile,
    role_title: str,
    job_description: str,
) -> build_resume.JobProblemProfile:
    lane_key = build_resume.effective_lane_key(role_title, job_description, profile)
    return build_resume.adjust_profile_for_lane(profile, lane_key)


def four_value_factors(profile: build_resume.JobProblemProfile, job_description: str) -> list[str]:
    """
    The 4 value factors from the career coaching session.
    Positioning should lead with Christian's strongest two: domain expertise and execution.
    """
    specialty = build_resume.role_specialty_phrase(job_description, "enterprise systems and implementation")
    return [
        f"DOMAIN EXPERTISE: 10+ years in {specialty}. Lead with this because it is Christian's clearest differentiator from generalists and recent grads applying to the same roles.",
        "EXECUTION ABILITY: 80+ client engagements, five-site system ownership, and 200+ reporting tools delivered. Proof that Christian gets things done, not just discusses them.",
        "ENERGY AND ATTITUDE: Surface this through delivery, not claims. Show genuine interest in the company's problem. Smile. Be specific about why this role and this employer.",
        "RELATIONSHIPS: 60+ executive workshops and QBRs, plus $1M+ account recovery through stakeholder alignment. The proof is in cross-functional outcomes, not title-level connection.",
        "Positioning tip: in the first 6 months of any new role, the biggest advantage is asking 'newbie' questions without judgment. Signal that you will do this intentionally because it shows self-awareness and learning agility.",
    ]


def career_through_line(profile: build_resume.JobProblemProfile, job_description: str) -> str:
    through_lines = {
        "implementation_delivery": "The consistent theme across my career is turning ambiguous system and implementation problems into usable operating outcomes that hold up after I leave the room.",
        "customer_success": "The consistent theme across my career is closing the gap between what customers were promised and what they actually experienced after go-live.",
        "presales_solution": "The consistent theme across my career is helping buyers make complex technology decisions they can actually defend to their leadership.",
        "analytics_operations": "The consistent theme across my career is converting messy operational data into decisions that change how teams work, not just what they see.",
        "change_enablement": "The consistent theme across my career is making change land after the announcement, not just during it.",
        "process_improvement": "The consistent theme across my career is finding the structural gap in a process that everyone has learned to work around, and building the fix that makes it stop recurring.",
        "corporate_strategy": "The consistent theme across my career is turning ambiguous business problems and competing stakeholder priorities into structured recommendations, clear operating plans, and measurable outcomes that teams can act on without losing sight of the original goal.",
    }
    return through_lines.get(
        profile.primary_lane,
        "The consistent theme across my career is turning complex systems, customer needs, and operational data into practical outcomes people can use.",
    )


def pitch_for_profile(
    profile: build_resume.JobProblemProfile,
    job_description: str = "",
    company_name: str = "",
    role_title: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> str:
    return pitch_variants(
        profile,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
        resume_text=resume_text,
        notes_text=notes_text,
    )["30_second"]


def sixty_second_pitch(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> str:
    return pitch_variants(
        profile,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
        resume_text=resume_text,
        notes_text=notes_text,
    )["60_second"]


def value_compression_opening(profile: build_resume.JobProblemProfile, job_description: str = "") -> str:
    lane_openings = {
        "presales_solution": "I help buyers move from technical uncertainty to credible solution confidence without losing implementation realism.",
        "customer_success": "I help customers move from adoption risk and unclear value to steadier outcomes, stronger trust, and clearer renewal confidence.",
        "change_enablement": "I help teams move from change friction and stakeholder ambiguity to practical adoption people can actually sustain.",
        "analytics_operations": "I help leaders move from noisy data and workflow confusion to clearer decisions teams can act on.",
        "implementation_delivery": "I help teams move from messy implementation work to usable systems, clearer ownership, and stronger adoption after launch.",
        "process_improvement": "I help teams move from recurring workflow friction and manual work to clearer controls, faster execution, and measurable process improvement.",
        "corporate_strategy": "I help leaders move from ambiguous priorities to structured decisions and operating plans teams can actually use.",
    }
    return lane_openings.get(
        profile.primary_lane,
        "I help teams move from ambiguous work to practical outcomes people can use."
    )


def read_the_room_opening(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    specialty = build_resume.role_specialty_phrase(job_description, candidate_problem_phrase(profile))
    answers = [
        "Read the room before you sell: the first five minutes tell you whether this interviewer wants speed, warmth, structure, skepticism, or detail.",
        f"For {company_name}'s {role_title} role, listen for whether the opening tone is more diagnostic or more relational, then match that pace without losing clarity.",
        f"If they start broad, stay concise and business-first. If they start detailed, use one concrete example from {specialty} work earlier than usual.",
        "Watch what gets a follow-up: customer risk, stakeholder friction, data trust, delivery pace, or adoption. That is the lane to stay in for the next answer.",
    ]
    return answers


def spoken_positioning_variants(profile: build_resume.JobProblemProfile, job_description: str = "") -> list[str]:
    statement = build_resume.generate_positioning_statement(profile, job_description)
    role_anchor = {
        "presales_solution": "the team needs buyer-facing clarity before a technical recommendation lands",
        "customer_success": "the team needs post-sale ownership tied to adoption and renewal confidence",
        "change_enablement": "the team needs change work translated into practical adoption",
        "analytics_operations": "the team needs data and workflow issues turned into usable decisions",
        "implementation_delivery": "the team needs customer-facing delivery that holds up after go-live",
    }.get(profile.primary_lane, "the team needs practical problem solving with measurable follow-through")
    return [
        f"Direct: {statement}",
        f"Conversational: The thread through my background is solving this kind of problem in a practical way, especially when {role_anchor}.",
        f"Executive: If the role needs someone who can stay close to the work and still keep the outcome clear, that is where I am strongest.",
    ]


def why_role(profile: build_resume.JobProblemProfile, company_name: str, role_title: str) -> str:
    if profile.primary_lane == "presales_solution":
        return (
            f"{company_name}'s {role_title} role is compelling because it sits at the intersection of discovery, product depth, "
            "business value, and customer trust. The role needs someone who can ask precise questions, build a credible solution "
            "narrative, tailor demonstrations, and keep Enterprise sales cycles moving without losing the customer's real pain."
        )
    if profile.primary_lane == "customer_success":
        return (
            f"{company_name}'s {role_title} role is compelling because it rewards product-oriented customer success work: adoption, "
            "risk reduction, customer education, and scalable process improvement. It matches the way I have supported complex "
            "ERP customers across implementation, post-go-live optimization, and account stabilization."
        )
    return (
        f"{company_name}'s {role_title} role is compelling because it requires both customer-facing judgment and hands-on execution. "
        "It matches the way I have led discovery, configuration alignment, testing, training, and go-live readiness for complex "
        "enterprise software environments."
    )


def company_context_lines(job_description: str, profile: build_resume.JobProblemProfile, company_name: str) -> list[str]:
    context = build_resume.primary_employer_context(job_description)
    specialty = build_resume.role_specialty_phrase(job_description, "")
    values = build_resume.visible_values_phrase(job_description)
    lens = build_resume.primary_story_lens(job_description)
    lines: list[str] = []
    if lens:
        lines.append(f"Story lens: {lens['identity']} - {lens['business_problem']}.")
        lines.append(f"Christian's matching story: {lens['candidate_story']}.")
        lines.append(f"How to sound in the interview: {lens['interview_lens']}")
    if context:
        context_key = str(context["key"])
        context_summary = str(context["summary"])
        lines.append(f"Employer context: {context_summary}")
        if context_key == "consulting":
            lines.append(
                "Firm/practice lens: speak like a consultant - clarify the client problem, show structured discovery, name tradeoffs, and bridge recommendations to implementation reality."
            )
        elif context_key == "saas":
            lines.append(
                "Company lens: speak to adoption, customer lifecycle, product feedback, retention risk, and implementation quality."
            )
        elif context_key == "manufacturing":
            lines.append(
                "Industry lens: connect ERP, inventory, supply chain, finance operations, reporting, and cross-site adoption."
            )
        elif context_key == "financial_services":
            lines.append(
                "Industry lens: emphasize risk-aware delivery, controls, stakeholder governance, trustworthy reporting, and careful validation."
            )
        elif context_key == "customer_experience":
            lines.append(
                "Industry lens: emphasize workflow design, customer engagement, messaging/contact-center operations, adoption, and analytics."
            )
        elif context_key == "healthcare":
            lines.append(
                "Industry lens: emphasize workflow clarity, adoption, risk control, reporting, and service quality."
            )
    if specialty:
        lines.append(f"Role specialty to reference: {specialty}. Tie examples to supported experience, not unsupported domain ownership.")
    if values:
        lines.append(f"Visible values to echo carefully: {values}. Translate these into proof through workshops, training, account recovery, and cross-functional delivery.")
    lines.append(
        f"Career strategy bridge: position this as a move toward broader {profile.lane_label.lower()} work where Christian can combine technical systems depth with client-facing judgment."
    )
    return lines


def company_background_lines(job_description: str, profile: build_resume.JobProblemProfile, company_name: str, role_title: str) -> list[str]:
    jd_lower = job_description.lower()
    context = build_resume.primary_employer_context(job_description)
    specialty = build_resume.role_specialty_phrase(job_description, "customer implementation")
    lines: list[str] = []
    if context:
        lines.append(f"Company background from the posting: {context['summary']}")
    else:
        lines.append(f"Company background from the posting: {company_name} is hiring this role to turn {specialty} into reliable customer outcomes.")
    objective_context = build_resume.objective_business_context_sentence(job_description)
    if objective_context:
        lines.append(objective_context)
    if "cloud" in jd_lower or "platform" in jd_lower:
        lines.append(f"Role context: the work appears tied to a cloud or platform product, so answers should connect configuration, data, integrations, training, and adoption.")
    if build_resume.jd_mentions(job_description, "healthcare", "logistics", "compliance"):
        lines.append("Industry lens: use healthcare, logistics, and compliance language carefully. Speak to workflow accuracy, documentation, validation, service quality, and risk control without claiming unsupported domain ownership.")
    if build_resume.jd_mentions(job_description, "5-10", "5 to 10", "implementation projects", "20-30", "20 to 30", "active customer accounts"):
        lines.append("Scale signal from the posting: be ready to explain how you manage several implementations and active accounts at once without losing issue ownership or customer communication.")
    if build_resume.jd_mentions(job_description, "api", "automated file transfer", "sql", "jira", "azure devops", "json", "xml", "edi", "boomi"):
        lines.append("Technical context: expect questions about how you work around APIs, file transfers, SQL/data validation, issue tracking, integrations, and handoffs with technical teams.")
    if build_resume.jd_mentions(job_description, "punchout", "eprocurement", "e-procurement", "edi", "eorder api", "api opening"):
        lines.append("Digital order-channel bridge: do not overclaim direct PunchOut or eProcurement ownership. Tie the fit to ERP integration support, requirements translation, UAT, data validation, release coordination, and customer-workflow judgment, then explain that this is the base for ramping quickly into PunchOut and eOrder API work.")
    lines.append(f"Interview goal: make the conversation about how Christian can help {company_name} reduce delivery risk, protect customer trust, and make the {role_title} work feel organized from discovery through adoption.")
    return lines


def personal_fit_lines(profile: build_resume.JobProblemProfile, resume_text: str, job_description: str) -> list[str]:
    lines = [
        f"Best fit lane: {profile.lane_label}.",
        f"Core business problem to speak to: {candidate_problem_phrase(profile)}.",
        f"Implementation proof: {COMPANY_APTEAN} supports discovery, requirements definition, configuration, data migration, testing, go-live, and post-go-live support for 80+ international manufacturing clients.",
        "Customer proof: the $1M+ account stabilization story shows calm issue ownership, cross-functional coordination, and trust recovery when a customer is at risk.",
        "Training and adoption proof: 60+ workshops and QBRs support executive communication, enablement, and practical customer education.",
        "Data proof: 200+ dashboards/reporting tools plus SQL-supported validation and ETL/data transformation support technical troubleshooting and decision visibility.",
    ]
    if build_resume.jd_mentions(job_description, "healthcare", "logistics", "compliance"):
        lines.append("Industry bridge: frame manufacturing and enterprise software work as experience with operational accuracy, workflow discipline, traceability, and customer-facing implementation in complex environments.")
    if build_resume.jd_mentions(job_description, "jira", "azure devops", "engineering", "qa", "product", "support"):
        lines.append("Cross-functional bridge: use examples where Christian coordinated Product, Support, operations, finance, and customer stakeholders around issues, testing, and adoption.")
    if build_resume.jd_mentions(job_description, "punchout", "eprocurement", "e-procurement", "edi", "api"):
        lines.append("Integration bridge: if PunchOut, eProcurement, EDI, or eOrder APIs come up, position Christian as strong on ERP/data integration coordination, UAT, issue triage, and adoption even where the exact channel tooling is part of the ramp curve.")
    return lines


def thirty_minute_hiring_official_strategy(profile: build_resume.JobProblemProfile, company_name: str, role_title: str) -> list[str]:
    return [
        "Opening energy: make the first 30 seconds easy to engage with. Sound glad to be there, direct, and ready to talk about the work.",
        f"Fit translation: translate Christian's experience into {company_name}'s environment instead of reciting resume history. Use their vocabulary for {candidate_problem_phrase(profile)}.",
        "Power story: lead with one story that proves readiness for the role. Do not scatter five examples when one strong example can carry the interview.",
        "Strategic questions: ask about success, obstacles, customer value, and where this person needs to drive change. The questions should show how Christian would think after hire.",
        f"Close: summarize the fit for {role_title}, name the strongest proof point, and ask whether that matches what the team needs most.",
    ]


def is_early_stage_context(job_description: str) -> bool:
    return build_resume.jd_mentions(
        job_description,
        "yc-backed",
        "y combinator",
        "yc ",
        "early-stage",
        "early stage",
        "seed",
        "series a",
        "series b",
        "category that didn't exist",
        "defining a category",
        "define a category",
        "from scratch",
        "high-growth",
        "build the function",
        "own a critical function",
        "maintain the implementation playbook",
    )


def startup_interview_lines(job_description: str, supplied_text: str) -> list[str]:
    combined = f"{job_description}\n{supplied_text}".lower()
    startup_signals = (
        "startup",
        "build",
        "builder",
        "scale",
        "scaling",
        "fast-paced",
        "autonomy",
        "ownership",
        "multiple hats",
        "ambiguity",
        "tight team",
    )
    matched_signals = {
        signal
        for signal in startup_signals
        if re.search(rf"\b{re.escape(signal)}\b", combined)
    }
    if len(matched_signals) < 2:
        return []
    enterprise_counter_signal = bool(
        re.search(
            r"\b(fortune 500|fortune500|publicly traded|nasdaq|nyse|global operations|enterprise scale|thousands of employees|worldwide|multinational|established leader)\b",
            combined,
            re.I,
        )
    )
    if enterprise_counter_signal and len(matched_signals) < 3:
        return []
    opener = (
        "Startup fit: speak like a builder who can own outcomes, work across blurry lanes, and create process without waiting for a perfect structure."
        if is_early_stage_context(job_description)
        else "Operator fit: speak like someone who can own outcomes across blurry lanes without needing a perfect structure first, which is different from startup-mode but shares the same ownership emphasis."
    )
    return [
        opener,
        "Large-company bridge: frame bigger-company experience as useful pattern recognition - what scalable systems look like and what bureaucracy to avoid.",
        "Extra-work question: do not sound hesitant. A grounded answer is that startups have seasons where the team has to lean in, and Christian is comfortable doing that when the work matters and the goals are clear.",
        "Red flags to avoid: heavy structure needs, narrow job boundaries, work-life-balance questions, risk aversion, or sounding like training has to be extensive before contribution starts.",
    ]


def tam_interview_lines(job_description: str, supplied_text: str) -> list[str]:
    combined = f"{job_description}\n{supplied_text}".lower()
    if not re.search(r"\b(technical account manager|tam|customer success|customer lifecycle|net revenue retention|nrr|account health|renewal|expansion)\b", combined):
        return []
    return [
        "TAM frame: position the role as business-case driven customer ownership, not only support, project delivery, or relationship management.",
        "Three pillars to show: keep the customer stable, improve day-2 operations, and transfer product/data knowledge so the customer becomes more capable.",
        "AI-era customer success: emphasize business outcomes, proactive insight, customer challenge, product feedback, and commercial awareness instead of activity tracking or polite follow-up.",
        "Best Christian bridge: implementation depth plus account recovery, QBR/workshop communication, SQL/reporting fluency, and supported expansion or renewal-risk conversations.",
    ]


def strategic_question_filter_lines() -> list[str]:
    return [
        "Use questions to position thinking, not just gather information.",
        "Strong topics: success outcomes, obstacles, business value, customer pain, team handoffs, data quality, adoption, and where the role needs to push change.",
        f"Use at least one diagnostic question such as: \"{biggest_gap_question()}\"",
        "Research before the interview should go beyond the company website: understand the industry, key competitors, and one reasoned pain-point or solution hypothesis you can test in conversation.",
        "Avoid early weak questions: culture, day in the life, work-life balance, or anything that makes the conversation about perks before contribution.",
    ]


def biggest_gap_question() -> str:
    return "What is the biggest gap between the experience you are seeing and what you actually need for this role?"


def results_based_story_positioning(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    stories: list[StoryCard],
) -> list[str]:
    top_stories = stories[:5]
    wins = [f"{story.title} - {story.result}" for story in top_stories]
    debrief_words = {
        "customer_success": "commercially aware, trusted, outcome-focused",
        "presales_solution": "credible, consultative, practical",
        "implementation_delivery": "organized, accountable, adoption-minded",
        "analytics_operations": "clear, analytical, business-focused",
        "change_enablement": "calm, persuasive, practical",
        "process_improvement": "structured, measurable, steady",
    }.get(profile.primary_lane, "clear, accountable, practical")
    lines = [
        "Tell-me-about-yourself rule: do not walk through the resume chronologically. Lead with the biggest relevant wins and connect them to the role.",
        f"Results-based summary: Christian's strongest career pattern is turning {candidate_problem_phrase(profile)} into usable outcomes through implementation, data, customer communication, and adoption.",
        "Career-story compression rule: if the story were written down, it should fit on one page with one clear through-line, the strongest proof, and why this role is the right next step.",
        "Top career wins to keep ready:",
    ]
    lines.extend(wins)
    lines.append(f"Debrief target: the interview team should describe Christian as {debrief_words}.")
    lines.append("Two-minute guardrail: answer the question, give one example, state the result, and bridge back to the role before adding extra detail.")
    lines.append(f"Explicit interest close: based on what I have learned, I am very interested in the {role_title} role at {company_name}, especially because the work connects directly to the customer and business problems I have been solving.")
    return lines


def memorable_candidate_lines(profile: build_resume.JobProblemProfile, company_name: str) -> list[str]:
    obvious_choice = build_resume.obvious_choice_positioning(profile)
    short_line = str(obvious_choice.get("short_line", "")).strip()
    return [
        short_line or "Obvious-choice frame: make it easy for them to see the same kind of problem, stakeholders, and outcome pattern Christian has already handled.",
        "Match the interviewer's energy: direct with direct, thoughtful with thoughtful, conversational with conversational.",
        "Answer as a business problem solver: name the problem, what changed, and what improved because of the work.",
        "Listen for pain points such as growth, missed goals, data quality, turnover, slow decisions, chaotic handoffs, or unclear ownership, then weave that language into the next answer.",
        f"Teach one useful idea if the opening appears natural: for {company_name}, connect Christian's experience to a practical way to improve {candidate_problem_phrase(profile)}.",
        "Make interest obvious. Do not assume they know. Say it directly, connect it to their needs, and close with intention.",
    ]


def diagnose_before_selling_pivot_question(profile: build_resume.JobProblemProfile) -> str:
    lane_questions = {
        "presales_solution": "Before I answer that fully, is the bigger issue for the team discovery quality, buyer confidence, or implementation fit after the sale?",
        "customer_success": "Before I answer that fully, is the bigger issue adoption, renewal risk, or executive alignment?",
        "change_enablement": "Before I answer that fully, is the bigger issue stakeholder alignment, adoption behavior, or execution follow-through?",
        "analytics_operations": "Before I answer that fully, is the bigger issue data trust, decision speed, or workflow visibility?",
        "implementation_delivery": "Before I answer that fully, is the bigger issue implementation pace, adoption after go-live, or cross-functional decision-making?",
    }
    return lane_questions.get(
        profile.primary_lane,
        "Before I answer that fully, is the bigger issue speed, adoption, or decision quality for the team right now?",
    )


def diagnose_before_selling_framework(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    pivot_question = diagnose_before_selling_pivot_question(profile)
    specialty = build_resume.role_specialty_phrase(job_description, candidate_problem_phrase(profile))
    return [
        "Diagnose before selling: do not reach for your best proof until you know what problem the interviewer is actually trying to solve.",
        f"Listen for what is being protected in {company_name}'s {role_title} role: pace, adoption, risk, customer trust, data quality, stakeholder alignment, or decision quality.",
        f"Use the pivot question when the prompt is broad: {pivot_question}",
        f"Once the issue is clearer, answer with one example from {specialty} work, then bridge back to the team's likely pressure point.",
    ]


def lane_structural_diagnostic_question(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> str:
    questions = {
        "presales_solution": f"In your strongest deals at {company_name}, what usually decides momentum first: discovery quality, technical proof, or confidence in implementation fit?",
        "customer_success": f"When accounts drift at {company_name}, is the bigger issue adoption depth, executive alignment, or renewal risk visibility?",
        "change_enablement": f"When change efforts stall in a {role_title} context, is the main issue sponsor alignment, manager translation, or frontline adoption behavior?",
        "analytics_operations": f"When reporting fails to change decisions at {company_name}, is the main problem data trust, unclear ownership, or weak actionability?",
        "implementation_delivery": f"When implementations slow down at {company_name}, is the bigger issue scope clarity, stakeholder alignment, or adoption after go-live?",
    }
    return questions.get(
        profile.primary_lane,
        f"For {role_title}, what usually creates the biggest drag first: unclear priorities, weak ownership, or low adoption confidence?",
    )


def bold_diagnostic_questions(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    structural = lane_structural_diagnostic_question(profile, company_name, role_title, job_description)
    return [
        structural,
        f"If I joined {company_name}, where would you most want a new hire to diagnose before proposing a fix?",
        "What friction is visible to leadership right now that the team closest to the work experiences differently?",
        "Which problem looks straightforward from the outside but becomes more complex once someone is actually in the role?",
    ]


def preloaded_questions(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    structural = lane_structural_diagnostic_question(profile, company_name, role_title, job_description)
    second = f"If someone is truly strong in the first 90 days of the {role_title} role, what starts to look clearer or move faster for the team?"
    third = {
        "presales_solution": "When a deal is headed in the right direction here, what starts to change first in the buyer conversation?",
        "customer_success": f"When customer relationships are healthiest at {company_name}, what behaviors or signals show up earlier than most people expect?",
        "change_enablement": "When adoption is gaining traction here, what changes first: the manager conversation, the frontline behavior, or the operating rhythm?",
        "analytics_operations": "When reporting is actually helping decisions here, what changes first in the way leaders or teams act on the information?",
        "implementation_delivery": f"When implementations are healthiest at {company_name}, what becomes visible early: cleaner scope decisions, faster issue ownership, or stronger adoption momentum?",
    }.get(
        profile.primary_lane,
        f"For {role_title}, what is the earliest signal that the work is moving in the right direction?"
    )
    return [structural, second, third]


def power_story_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    supplied_text: str,
    stories: list[StoryCard],
) -> list[str]:
    combined = f"{job_description}\n{supplied_text}"
    story = max(stories, key=lambda item: signal_score(combined, item.signals))
    return [
        f"Lead story: {story.title}.",
        f"Why this story: it is the cleanest proof that Christian can translate {candidate_problem_phrase(profile)} into a customer or business result for {company_name}.",
        f"Context: {story.hook}",
        f"Action: {story.evidence}",
        f"Result: {story.result}",
        f"Bridge: connect the story to {role_title} by saying this is the same operating pattern - understand the customer problem, make the work measurable, coordinate the owners, and turn the result into adoption or growth.",
    ]


def interest_answer_for_lane(profile: build_resume.JobProblemProfile, company_name: str, role_title: str) -> str:
    lane_answers = {
        "presales_solution": (
            f"I am very interested in the {role_title} role because it sits where discovery, product credibility, and customer trust meet. "
            f"What stands out about {company_name} is the chance to help buyers understand a solution clearly enough to make a confident decision, while still keeping implementation reality in view."
        ),
        "customer_success": (
            f"I am very interested in the {role_title} role because it centers on adoption, account health, and customer value after the sale. "
            f"That is the kind of work I want more of: helping {company_name}'s customers turn software investment into measurable outcomes and retained trust."
        ),
        "analytics_operations": (
            f"I am very interested in the {role_title} role because it connects operating data to decisions. "
            f"The work sounds like a strong fit for the way I have used reporting, process visibility, and stakeholder alignment to help teams act with more clarity."
        ),
        "change_enablement": (
            f"I am very interested in the {role_title} role because it is about making change usable after the announcement. "
            "That fits the part of my background I want to keep building: helping teams move through adoption, training, stakeholder alignment, and practical workflow change."
        ),
        "process_improvement": (
            f"I am very interested in the {role_title} role because it appears to need someone who can find the structural cause behind recurring process friction. "
            "That is close to the work I enjoy most: mapping the current state, making the issue measurable, piloting the fix, and helping the team keep using it."
        ),
        "corporate_strategy": (
            f"I am very interested in the {role_title} role because it connects analysis to decisions that teams actually have to execute. "
            f"What appeals to me is helping {company_name} turn ambiguous priorities into a practical operating plan, with enough structure that stakeholders can act on it."
        ),
    }
    return lane_answers.get(
        profile.primary_lane,
        f"I am very interested in this role. What stands out is that {company_name} needs someone who can connect customer problems, data, implementation, and adoption rather than treating those as separate lanes. It is the kind of work I want to do more of, and my background gives me practical examples of turning unclear customer or system problems into outcomes people can use.",
    )


def education_credentials_answer(resume_text: str) -> str:
    lowered = resume_text.lower()
    has_bachelors = bool(re.search(r"\bbachelor(?:'s|s)?\b", lowered))
    has_masters = bool(re.search(r"\bmaster(?:'s|s)?\b", lowered))
    has_itil = "itil" in lowered
    if has_bachelors and has_masters:
        opening = "I have both a bachelor's and a master's degree in Information Systems"
    elif has_masters:
        opening = "My formal education includes a master's degree in Information Systems"
    elif has_bachelors:
        opening = "My formal education includes a bachelor's degree in Information Systems"
    else:
        opening = "My education gave me a formal foundation in information systems and business technology"
    credential = "I also completed ITIL Foundation training" if has_itil else "I have continued building that foundation through professional development"
    return interview_join(
        opening,
        credential,
        "The practical value has been learning how to connect technical decisions with the way people and operations actually work",
    )


def chronological_career_walkthrough(resume_text: str) -> str:
    text = interview_join(
        "My career started in a corporate technology environment at Home Depot, where I supported customer-facing systems and learned how technical work affects the customer experience",
        f"From there, I moved into ERP consulting at {COMPANY_APTEAN}, supporting manufacturing clients from discovery and requirements through configuration and installation",
        "I stayed involved through go-live and post-launch support, learning how to translate operating needs into system behavior",
        f"Most recently, at {COMPANY_EAST_WEST}, I owned the ERP environment across inventory, scheduling, reporting, and integrations",
        "I also supported dashboards and the cross-functional work needed to keep those workflows reliable",
        "I partnered with engineering, planning, purchasing, supply chain, and finance, so I saw how one system issue could affect several parts of the business",
        "What ties the progression together is that I have spent my career where systems, customers, and operations meet",
    )
    assert_spoken_answer_budget(
        "Career walkthrough",
        text,
        minimum_words=CAREER_WALKTHROUGH_WORD_RANGE[0],
        maximum_words=CAREER_WALKTHROUGH_WORD_RANGE[1],
        minimum_average_sentence_words=14,
    )
    return text


def product_learning_answer() -> str:
    return interview_join(
        "I learn a new product in three layers",
        "First, I learn the core workflows, modules, implementation method, and the customer use cases the product handles most often",
        "Second, I map those features to the operational problems they solve, including what changes in visibility, execution, reporting, or decision quality",
        "Third, I study real customer scenarios, especially common implementation issues, adoption challenges, and the points where teams usually need guidance",
        "The approach helps me move beyond feature knowledge and understand how the product is actually used in practice",
    )


def honest_gap_framing(gap_area: str, strength_area: str, adjacent_capability: str, partner_group: str) -> str:
    return interview_join(
        f"I'd be direct and say {gap_area} is not my deepest area",
        f"My depth is stronger in {strength_area}",
        f"What I do understand well is {adjacent_capability}",
        f"When the work moves deeper into {gap_area}, I partner with {partner_group} and stay responsible for keeping the operational and system implications aligned",
        "I would rather be precise about that boundary than overstate it, because that makes the collaboration and the final solution more reliable",
    )


def manufacturing_depth_answer() -> str:
    return interview_join(
        "My manufacturing experience is strongest in the ERP and operational workflow layer",
        "I have supported item setup, inventory adjustments, order and batch processing, scheduling, shop floor workflows, reporting, and integrations",
        "My role was to own or troubleshoot how those processes behaved in the system and coordinate with the teams affected by them",
        "Transaction quality mattered because it consistently influenced inventory accuracy, planning reliability, production visibility, and downstream financial reporting",
        "I would not position myself as a plant engineer, but I am comfortable in manufacturing environments where the ERP must reflect real material and order flow",
    )


def technical_depth_boundary_answer() -> str:
    return honest_gap_framing(
        "controls engineering or schematic interpretation",
        "ERP, process flow, transaction integrity, and operational systems support",
        "how shop-floor activity and technical processes need to be represented accurately in the business system",
        "engineering and controls specialists",
    )


def _company_note_lines(supplied_context: str) -> list[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in supplied_context.splitlines() if line.strip()]


def compensation_logistics_answer(supplied_context: str) -> str:
    lines = _company_note_lines(supplied_context)
    candidate_lines = [
        line
        for line in lines
        if re.search(r"\b(?:i|christian)\b", line, re.I)
        and not re.search(r"\b(?:recruiter|company|budget|approved range)\b", line, re.I)
    ]
    compensation_line = next(
        (
            line
            for line in candidate_lines
            if re.search(r"\b(?:target(?:ing|ed)?|expect(?:ing|ed)?|looking for|compensation|salary)\b", line, re.I)
            and re.search(r"(?:\$\s?\d[\d,]*(?:k|K)?|\b(?:low|mid|high)[- ]six figures\b)", line, re.I)
        ),
        "",
    )
    travel_line = next(
        (
            line
            for line in candidate_lines
            if re.search(r"\btravel\b", line, re.I)
            and re.search(r"\b(?:comfortable|fine|open|available|can)\b", line, re.I)
        ),
        "",
    )
    availability_line = next(
        (
            line
            for line in candidate_lines
            if re.search(r"\b(?:available|availability|start|move quickly)\b", line, re.I)
        ),
        "",
    )
    amount_match = re.search(
        r"(?:\$\s?\d[\d,]*(?:k|K)?(?:\s*(?:to|-|through)\s*\$?\s?\d[\d,]*(?:k|K)?)?|\b(?:low|mid|high)[- ]six figures\b)",
        compensation_line,
        re.I,
    )
    if amount_match:
        opening = f"I'm targeting {amount_match.group(0)}, though I am flexible depending on the full scope and total package"
    else:
        opening = "I'd like to understand the approved range and total package, then I can respond with the right context"
    details: list[str] = [opening]
    if travel_line:
        travel_match = re.search(r"(?:about\s+)?\d+\s+days?\s*(?:per|/)\s*month", travel_line, re.I)
        details.append(
            f"I am comfortable with {travel_match.group(0)} of travel"
            if travel_match
            else "I am comfortable with the travel expectations described for the role"
        )
    if availability_line:
        details.append("My timing is flexible, and I can move quickly if the process is a mutual fit")
    if len(details) == 1:
        details.append("I am also happy to discuss travel and timing directly so we can confirm the practical fit early")
    return interview_join(*details)


def common_interview_answers(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    stories: list[StoryCard],
    panel_context: bool = False,
    supplied_context: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> list[PreparedAnswer]:
    achievement = story_for_type(stories, "Individual Achievement") or stories[0]
    persuasion = story_for_type(stories, "Persuasion") or achievement
    analysis = story_for_type(stories, "Analysis and Decision") or achievement
    rapid = story_for_type(stories, "Rapid Learning") or achievement
    teamwork = story_for_type(stories, "Teamwork") or persuasion
    questions = [
        PreparedAnswer(
            "Diagnose-before-selling pivot question",
            diagnose_before_selling_pivot_question(profile),
        ),
        PreparedAnswer(
            "Tell me about yourself",
            sixty_second_pitch(profile, company_name, role_title, job_description, resume_text, notes_text),
        ),
        PreparedAnswer(
            f"Why {company_name}?",
            why_company_three_reason_answer(profile, company_name, job_description, supplied_context),
        ),
        PreparedAnswer(
            f"Why this {role_title} role?",
            interview_join(
                "What stands out to me is the mix of customer-facing problem solving and practical delivery",
                f"At {COMPANY_APTEAN}, I supported requirements, configuration, data migration, testing, go-live, and post-launch support for manufacturing clients",
                "The experience taught me implementation quality depends on both the system setup and how well the workflow works for the people using it",
                f"The {role_title} role lets me keep building at that intersection, which is where I have done my strongest work",
            ),
        ),
        PreparedAnswer(
            "Why should we hire you?",
            join_answer_sentences(
                "Hire me if this role needs someone who can combine customer-facing judgment with hands-on delivery discipline",
                "I can manage ambiguity, keep issues visible, work with technical and nontechnical stakeholders, and connect the work to adoption instead of treating launch as the finish line",
                story_natural_reference(achievement),
            ),
        ),
        PreparedAnswer(
            "How interested are you in this role?",
            interest_answer_for_lane(profile, company_name, role_title),
        ),
        PreparedAnswer(
            "How would you manage several implementations at once?",
            "I would keep a simple operating rhythm: clear milestone plans, visible risks and issues, named owners, customer-facing status, and documented next steps after each interaction. The key is not to remember everything manually. The key is to make the work visible enough that scope, data, testing, and stakeholder blockers are caught early.",
        ),
        PreparedAnswer(
            "How do you handle a difficult implementation issue?",
            interview_join(
                "I start by separating the visible symptom from the actual failure point",
                "Then I identify who owns each part of the issue, what evidence is still missing, and what the customer needs to hear next",
                story_natural_reference(persuasion),
                "The goal is one recovery path with clear ownership, not a customer experience split across several teams",
            ),
        ),
        PreparedAnswer(
            "How do you approach data migration or integration risk?",
            interview_join(
                "I treat data and integration work as a validation problem, not only a technical task",
                "I want source clarity, field-level checks, test cases, exception handling, and a signoff path before go-live",
                story_natural_reference(analysis),
                "That keeps the technical handoff connected to the operational result the customer expects",
            ),
        ),
        PreparedAnswer(
            "How do you train customers and drive adoption?",
            "I make training role-based and practical. Customers need to understand the workflow they will use, the decisions they need to make, and where to go when something breaks. My strongest proof is facilitating 60+ workshops and QBRs, where the goal was to make complex systems understandable enough for executives, operators, and customer teams to act on.",
        ),
        PreparedAnswer(
            "How do you work with Product, Engineering, QA, or Support?",
            interview_join(
                "I translate the issue into business impact, a clear reproduction path, priority, and the decision needed",
                "The mistake in cross-functional work is assuming every group defines urgency the same way",
                story_natural_reference(teamwork),
                "I try to make the tradeoff visible and leave every group with a concrete next step",
            ),
        ),
        PreparedAnswer(
            "What is a gap we should know about?",
            interview_join(
                "I'd be direct about any domain area where my experience is adjacent rather than deep",
                "My strongest depth is in implementation judgment, workflow learning, data validation, training, and customer issue ownership",
                "When I need to close an edge, I get close to the workflow, validate my understanding with the right specialists, and make progress visible early",
                story_natural_reference(rapid),
                "The method lets me learn quickly without pretending I know more than the evidence supports",
            ),
        ),
        PreparedAnswer(
            "Tell me about your education and credentials.",
            education_credentials_answer(resume_text),
        ),
        PreparedAnswer(
            "Walk me through your career.",
            chronological_career_walkthrough(resume_text),
        ),
        PreparedAnswer(
            "How would you learn a new product quickly?",
            product_learning_answer(),
        ),
        PreparedAnswer(
            "What are your compensation expectations and practical constraints?",
            compensation_logistics_answer(supplied_context),
        ),
    ]
    answer_by_prompt = {item.prompt: item.answer for item in questions}
    assert_spoken_answer_budget(
        "Education and credentials answer",
        answer_by_prompt["Tell me about your education and credentials."],
        minimum_words=BRIEF_RECRUITER_WORD_RANGE[0],
        maximum_words=BRIEF_RECRUITER_WORD_RANGE[1],
        minimum_average_sentence_words=10,
    )
    assert_spoken_answer_budget(
        "Product learning answer",
        answer_by_prompt["How would you learn a new product quickly?"],
        minimum_words=70,
        maximum_words=105,
    )
    assert_spoken_answer_budget(
        "Compensation and logistics answer",
        answer_by_prompt["What are your compensation expectations and practical constraints?"],
        minimum_words=BRIEF_RECRUITER_WORD_RANGE[0],
        maximum_words=BRIEF_RECRUITER_WORD_RANGE[1],
        minimum_average_sentence_words=10,
    )
    return questions


def keyword_ready_answers(
    profile: build_resume.JobProblemProfile,
    job_description: str,
    stories: list[StoryCard],
    resume_text: str = "",
    supplied_context: str = "",
) -> list[PreparedAnswer]:
    keyword_prompts: list[tuple[tuple[str, ...], str, str]] = [
        (("manufacturing", "shop floor", "production planning", "bom"), "Manufacturing Experience", "Walk me through your manufacturing systems experience."),
        (("electrical schematics", "control systems", "controls", "plc", "scada"), "Technical Depth Boundary", "How deep is your experience with controls, schematics, or plant-floor systems?"),
        (("implementation", "go-live", "configuration"), "Implementation", "How do you run an implementation from discovery through go-live?"),
        (("data migration", "etl", "sql", "database"), "Data Migration And SQL", "How do you keep data work accurate during implementation?"),
        (("integration", "api", "file transfer", "json", "xml", "edi", "boomi"), "Integrations", "How comfortable are you working around integrations and technical handoffs?"),
        (("marketing analytics", "ga4", "paid media", "crm data", "ltv", "cac"), "Marketing Analytics", "How would you help customers turn marketing analytics into better business decisions?"),
        (("net revenue retention", "nrr", "renewal", "expansion", "growth"), "NRR And Growth", "How do you connect customer success work to expansion and retention?"),
        (("ai tools", "ai", "ask", "natural language", "automation"), "AI-Enabled Customer Work", "How are you using AI in your current workflow, and how would you help customers trust AI-driven insights?"),
        (("training", "adoption", "enablement"), "Training And Adoption", "How do you make sure customers actually use what was implemented?"),
        (("jira", "azure devops", "issue tracking"), "Issue Tracking", "How do you manage open issues across teams?"),
        (("customer accounts", "active customer", "customer communication"), "Customer Accounts", "How do you keep customers confident when several accounts or projects are active?"),
        (("healthcare", "logistics", "compliance"), "Healthcare, Logistics, And Compliance", "How would you handle a regulated or compliance-sensitive customer environment?"),
    ]
    answers: list[PreparedAnswer] = []
    for terms, label, question in keyword_prompts:
        if not build_resume.jd_mentions(job_description, *terms):
            continue
        story = max(stories, key=lambda item: signal_score(" ".join(terms), item.signals)) if stories else None
        if label == "Manufacturing Experience":
            answer = manufacturing_depth_answer()
        elif label == "Technical Depth Boundary":
            answer = technical_depth_boundary_answer()
        elif label == "Implementation":
            answer = f"I would start with discovery and success criteria, translate those into configuration and milestone plans, then manage testing, training, and go-live readiness with visible risks and owners. My {COMPANY_APTEAN} work supports that pattern across requirements, configuration, data migration, testing, go-live, and post-go-live support."
        elif label == "Data Migration And SQL":
            answer = "I would focus on validation: source data, field mapping, exception checks, test records, signoff, and a clear cutover plan. My background supports SQL-based validation, ETL/data transformation, database record extraction, and dashboard/reporting work, so I can work with technical teams without turning the customer conversation into jargon."
        elif label == "Integrations":
            answer = "My value around integrations is the implementation behavior that usually determines success: clarify the workflow, document the expected handoff, track errors, validate test cases, and keep the customer informed. I am strongest when technical integration work has to stay connected to business impact, next steps, and customer trust."
        elif label == "Marketing Analytics":
            answer = "I would start with the business decision behind the report: acquisition source, conversion path, lifetime value, retention, or margin. My background supports the translation layer because I have built 200+ dashboards and reporting tools, worked with customer and transactional data, and helped stakeholders move from raw exports to decisions they could act on."
        elif label == "NRR And Growth":
            answer = "I would treat expansion as the result of trusted onboarding and visible value. My supported lane is renewal-risk management, expansion discovery, customer health, QBRs, and stabilizing high-risk accounts. The proof is account recovery plus the ability to turn implementation success into the next value conversation."
        elif label == "AI-Enabled Customer Work":
            answer = "I use AI as a work accelerator, not a substitute for judgment. My supported experience includes AI-assisted documentation, analysis, workflow support, and conversational messaging configuration. With customers, I would focus on trust: validate the data, explain the recommendation clearly, and show how the insight connects to a business action."
        elif label == "Training And Adoption":
            answer = "I would train around roles and decisions, not just features. The customer needs to know what changes in the workflow, what a good outcome looks like, and how to handle exceptions. My 60+ workshops and QBRs support that kind of practical enablement."
        elif label == "Issue Tracking":
            answer = "I would keep issues visible by owner, priority, business impact, due date, dependency, and customer-facing update. The point is to prevent hidden blockers from becoming trust problems. My $1M+ account stabilization story is the best example of creating order around fragmented issues."
        elif label == "Customer Accounts":
            answer = "I would use a consistent account rhythm: priorities, risks, open issues, next milestone, and customer expectation. The account-health side of my background fits this because I have supported adoption, QBRs, renewal-risk conversations, and high-risk customer recovery."
        else:
            answer = "I would treat regulated or service-sensitive environments with extra care around documentation, validation, workflow clarity, and signoff. My closest proof is operational discipline in complex customer and manufacturing environments, where the work only succeeds if the process is trusted and the handoffs are clear."
        if story and label not in {"Manufacturing Experience", "Technical Depth Boundary"}:
            answer = interview_join(answer, story_natural_reference(story))
        if label in {"Manufacturing Experience", "Technical Depth Boundary"}:
            assert_spoken_answer_budget(
                f"{label} answer",
                answer,
                minimum_words=STANDARD_SPOKEN_WORD_RANGE[0],
                maximum_words=STANDARD_SPOKEN_WORD_RANGE[1],
            )
        answers.append(PreparedAnswer(f"{label}: {question}", answer))
    return answers[:9]


def fit_snapshot(profile: build_resume.JobProblemProfile, resume_docx: Path) -> list[str]:
    supported = ", ".join(profile.direct_matches[:4]) if profile.direct_matches else "limited direct overlap"
    adjacent = ", ".join(profile.adjacent_matches[:3]) if profile.adjacent_matches else "no major adjacent areas flagged"
    watchouts = ", ".join(profile.unsupported_requirements[:3]) if profile.unsupported_requirements else "no major unsupported requirement areas detected"
    return [
        f"Fit: {fit_label(profile, resume_docx)}",
        f"Target lane: {profile.lane_label}",
        f"Core problem to speak to: {candidate_problem_phrase(profile)}",
        f"Audience to reference: {profile.audience}",
        f"Best supported themes: {supported}",
        f"Bridge themes to foreground with proof: {adjacent}",
        f"Watchouts: {watchouts}",
    ]


def qualification_test_lines(job_description: str, resume_text: str) -> list[str]:
    """
    Extract JD requirements and map each to PASS / PARTIAL / GAP for internal QA.
    """
    profile = build_resume.job_problem_profile(job_description, resume_text)
    direct = list(profile.direct_matches[:6])
    adjacent = list(profile.adjacent_matches[:4])
    unsupported = list(profile.unsupported_requirements[:4])
    return [
        "Evidence check: map the job's requirement bullets to what Christian can honestly support.",
        f"PASS ({len(direct)} direct evidence areas): {', '.join(direct) if direct else 'review manually'}",
        f"PARTIAL ({len(adjacent)} bridge areas - foreground with proof): {', '.join(adjacent) if adjacent else 'none flagged'}",
        f"GAP ({len(unsupported)} unsupported areas - do not claim): {', '.join(unsupported) if unsupported else 'none flagged'}",
        "If GAP items are in the Requirements section, keep them as honest interview watchouts instead of claims.",
    ]


def recruiter_three_stage_audit(profile: build_resume.JobProblemProfile, job_description: str, resume_text: str) -> list[str]:
    """
    3-stage recruiter read model from the coaching notes.
    Stage 1: 2-3 sec skim; Stage 2: 10 sec scan; Stage 3: full proof read.
    """
    specialty = build_resume.role_specialty_phrase(job_description, "the role's core work")
    return [
        f"STAGE 1 (2-3 sec skim): Is Christian's target role lane visible at the very top? Win condition: '{profile.lane_label}' or equivalent is the first thing a recruiter sees.",
        f"STAGE 2 (10 sec scan): Do the skills section and first bullet of each role contain the JD's exact language? Win condition: '{specialty}' terms appear in the first bullet of the most recent role, not buried in bullet 5 or 6.",
        "STAGE 3 (full read): Does every bullet answer 'What changed because Christian was there?' Win condition: Result + Metric + Context on every line. No duty-only bullets. At least 3 bullets that become interview stories.",
        "Context sub-check: Has Christian done similar roles in the same industry? If yes, lead with it. If no, the Professional Summary must bridge the gap explicitly.",
        "Competence sub-check: Are the specific abilities the JD needs proven with numbers, scope markers, or go-live/adoption evidence? Soft skill claims without proof fail Stage 3.",
    ]


def role_positioning(profile: build_resume.JobProblemProfile, company_name: str, role_title: str) -> list[str]:
    if profile.primary_lane == "presales_solution":
        return [
            f"{company_name} appears to need a {role_title} who can reduce buyer uncertainty.",
            "Lead with discovery, demo tailoring, executive credibility, and solution fit.",
            "Keep answers tied to business value, customer trust, and implementation realism.",
        ]
    if profile.primary_lane == "customer_success":
        return [
            f"{company_name} appears to need a {role_title} who can protect adoption and account health.",
            "Lead with risk reduction, executive alignment, post-go-live improvement, and customer education.",
            "Keep answers tied to workflow stability, measurable outcomes, and stakeholder trust.",
        ]
    if profile.primary_lane == "analytics_operations":
        return [
            f"{company_name} appears to need a {role_title} who can turn messy operating data into clearer decisions.",
            "Lead with KPI visibility, reporting, process gaps, and action-oriented recommendations.",
            "Keep answers tied to decision support, business clarity, and measurable improvement.",
        ]
    if profile.primary_lane == "change_enablement":
        return [
            f"{company_name} appears to need a {role_title} who can help change stick after rollout decisions are made.",
            "Lead with adoption risk, stakeholder alignment, training, and reinforcement mechanisms.",
            "Keep answers tied to resistance reduction, readiness, and sustained usage.",
        ]
    if profile.primary_lane == "process_improvement":
        return [
            f"{company_name} appears to need a {role_title} who can improve complex processes without losing quality, service, or stakeholder trust.",
            "Lead with current-state mapping, root cause, Lean Six Sigma-style structure, operational metrics, pilots, and measurable before/after results.",
            "Keep answers tied to claims/customer experience, efficiency, quality, communication materials, and adoption by the people doing the work.",
        ]
    return [
        f"{company_name} appears to need a {role_title} who can guide complex delivery work without losing stakeholder confidence.",
        "Lead with discovery, scope alignment, testing, go-live readiness, and post-launch stabilization.",
        "Keep answers tied to delivery risk, customer workflows, and practical execution.",
    ]


def positioning_answers(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    panel_context: bool = False,
) -> list[PreparedAnswer]:
    specialty = build_resume.role_specialty_phrase(job_description, "the role's specialty area")
    values = build_resume.visible_values_phrase(job_description)
    company_context = build_resume.primary_employer_context(job_description)
    story_lens = build_resume.primary_story_lens(job_description)
    if story_lens:
        why_company = (
            f"{company_name} appears to need someone who understands {story_lens['identity']}: "
            f"{story_lens['business_problem']}. I would connect that to {story_lens['candidate_story']} and keep the answer grounded in measurable outcomes."
        )
    elif company_context and str(company_context["key"]) == "consulting":
        why_company = (
            f"{company_name} appears to need someone who can bring client-service judgment to {specialty}: clarify the client problem, "
            "shape practical recommendations, and keep implementation risk visible."
        )
    elif values:
        why_company = (
            f"{company_name} appears to value {values}, and the role connects that to {specialty}. "
            "I would make that concrete by listening for friction, aligning stakeholders, and turning the work into visible outcomes."
        )
    else:
        why_company = (
            f"{company_name} appears to need someone who can connect {specialty} to practical delivery. "
            "I would study where the team is seeing friction first, then use discovery, metrics, stakeholder alignment, and adoption planning to make the work concrete."
        )
    if panel_context:
        why_company = panel_three_reason_why_company(company_name, role_title)
    if profile.primary_lane == "process_improvement":
        return [
            PreparedAnswer(
                "Why you",
                f"Because this role needs someone who can turn ambiguous {specialty} problems into measurable process improvement. My strongest proof is reducing manual work by 78%, lowering discrepancies by 22%, building 200+ reporting tools, facilitating 60+ executive workshops and QBRs, and aligning users across five sites and 150+ people.",
            ),
            PreparedAnswer(
                "Why now",
                f"This is the right next step because my background has moved from hands-on systems and reporting ownership into broader {profile.lane_label.lower()}, stakeholder alignment, and measurable operational improvement.",
            ),
            PreparedAnswer(
                "Why this company",
                why_company,
            ),
            PreparedAnswer(
                "Why this role is logical",
                f"The {role_title} role is a logical next step because it sits close to work I have already done: current-state diagnosis, process redesign, operational reporting, communication materials, training, adoption, and executive-ready recommendations. The stretch is deeper claims context, and I would manage that with fast workflow learning, strong data validation, and SME feedback.",
            ),
        ]

    return [
        PreparedAnswer(
            "Why you",
            f"Because this role needs someone who can turn ambiguous {specialty} problems into practical operating improvements. My strongest proof is enterprise system ownership across five sites and 150+ users, 80+ client implementations, 200+ reporting tools, 60+ executive workshops, and $1M+ in account risk stabilized.",
        ),
        PreparedAnswer(
            "Why now",
            f"This is the right next step because my background has moved from hands-on systems ownership into broader {profile.lane_label.lower()}, stakeholder alignment, and customer outcomes. I can bring both the execution detail and the business judgment needed when the work is complex and cross-functional.",
        ),
        PreparedAnswer(
            "Why this company",
            why_company,
        ),
        PreparedAnswer(
            "Why this role is logical",
            f"The {role_title} role is a logical next step because it sits close to the work I have already done: requirements, scope, systems delivery, customer recovery, reporting, training, and executive communication. The stretch is deeper {specialty} context, and I would manage that with fast learning, strong discovery, and early validation.",
        ),
    ]


def ninety_second_pitch(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> str:
    return pitch_variants(
        profile,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
        resume_text=resume_text,
        notes_text=notes_text,
    )["90_second"]


def role_challenge_forecast(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
) -> list[str]:
    outcomes = ", ".join(profile.outcomes)
    specialty = build_resume.role_specialty_phrase(job_description, "the role's core work")
    values = build_resume.visible_values_phrase(job_description)
    context = build_resume.primary_employer_context(job_description)
    if context and str(context.get("key")) in {"early_stage", "saas"}:
        business_problem = (
            f"Likely pressure: {company_name} needs the {role_title} to turn {candidate_problem_phrase(profile)} "
            "into repeatable customer outcomes before every operating pattern is fully settled."
        )
    elif profile.primary_lane == "corporate_strategy":
        business_problem = (
            f"Likely pressure: {company_name} needs the {role_title} to turn analysis in {specialty} "
            "into an implementation plan before the recommendation can matter."
        )
    elif build_resume.is_consulting_job_description(job_description):
        business_problem = (
            f"Likely pressure: {company_name} needs the {role_title} to enter ambiguous client situations, "
            f"structure the problem quickly, and produce {specialty} recommendations that can survive implementation."
        )
    else:
        business_problem = (
            f"Likely pressure: {company_name} needs the {role_title} to convert {candidate_problem_phrase(profile)} "
            f"in {specialty} into {outcomes}."
        )
    lane_proof_lead = {
        "presales_solution": "80+ client implementations, 60+ executive workshops, and $1M+ in account risk stabilized",
        "customer_success": "$1M+ in account recovery, 80+ client portfolio, and 60+ QBRs and business reviews",
        "analytics_operations": "200+ dashboards, SQL-based KPI tools, and process improvement across five-site operations",
        "change_enablement": "change adoption programs, role-based training, 60+ executive workshops, and 150+ user deployments",
        "process_improvement": "78% manual-work reduction, 22% discrepancy reduction, 200+ KPI/reporting tools, and cross-functional workflow standardization",
        "implementation_delivery": "five-site enterprise system ownership, 80+ client implementations, ETL migration support, and go-live readiness",
    }
    proof = lane_proof_lead.get(profile.primary_lane, "enterprise system ownership, 80+ client implementations, dashboards, workshops, and account recovery")
    values_line = (
        f"Values lens: {values} should be demonstrated through concrete examples, not personality claims."
        if values
        else "Values lens: use the role context and supported evidence, not invented culture fit."
    )
    return [
        business_problem,
        "First 90-day challenge: learn the business context fast, find where stakeholders disagree, and separate real delivery risk from noise.",
        f"Stakeholder risk: {profile.audience} may each define success differently; surface that conflict early.",
        values_line,
        "Success metrics: faster decisions, cleaner scope, stronger adoption, reduced escalation, visible progress against business outcomes.",
        "What could fail: unclear requirements, over-customization, weak validation, or a solution that works technically but does not get adopted.",
        f"Christian's map: {proof}.",
    ]


def first_90_day_approach(profile: build_resume.JobProblemProfile) -> list[str]:
    return [
        "Days 1-30: listen for business outcomes, stakeholder incentives, current friction, decision rights, and the metrics leadership already trusts.",
        "Days 31-60: map the highest-risk workflows, clarify scope, identify adoption blockers, and build a simple operating rhythm for decisions and escalations.",
        "Days 61-90: deliver visible improvements, validate results with users and leaders, document repeatable practices, and confirm the next set of priorities.",
    ]


def expanded_story_bank() -> list[StoryCard]:
    return [
        StoryCard(
            title="Inventory adjustment system",
            story_types=("Individual Achievement", "Analysis and Decision", "Ambiguous Problem"),
            hook="The challenge was a high-volume inventory adjustment process that was too manual and exposed the operation to avoidable discrepancies.",
            takeaways=("Structured the messy workflow before building", "Validated the fix against operational reality", "Turned the work into measurable business improvement"),
            evidence="Built a large-scale inventory adjustment system in Aptean Intuitive ERP for high-volume operations.",
            level3_trait="Show what was noticed: repeated manual touches were creating delay and discrepancy risk, so the workflow was mapped, tested, and tightened before broader use.",
            result="Reduced manual work for the adjustment process by 78% and lowered inventory adjustment discrepancies by 22%.",
            outcome="Use this for process improvement, structured problem solving, and practical systems execution.",
            evidence_terms=("78%", "22%", "inventory adjustment"),
            signals=("inventory", "process", "optimization", "efficiency", "operations", "analysis"),
        ),
        StoryCard(
            title="Aptean rapid product learning",
            story_types=("Rapid Learning", "Challenge and Failure", "Individual Achievement"),
            hook="The challenge was becoming credible quickly across a complex ERP product and a broad international client base.",
            takeaways=("Learned through client problems, not abstract study", "Built a repeatable discovery rhythm", "Converted product complexity into clearer customer decisions"),
            evidence="Managed 80+ international client engagements through 12 full-lifecycle ERP implementations, carrying up to four concurrent deliveries from discovery through data migration, UAT, and post-go-live support.",
            level3_trait="Show how unfamiliar workflows were broken into requirements, risks, decision owners, and next actions until the client could move forward.",
            result="Delivered 12 full-lifecycle ERP implementations and managed up to four at a time, becoming a customer-facing implementation consultant and pre-sales resource across complex ERP delivery work.",
            outcome="Use this when asked about learning quickly, ambiguity, or becoming useful before every answer is known.",
            evidence_terms=("12 full-lifecycle", "4 concurrent", "80+ international"),
            signals=("learning", "rapid", "implementation", "requirements", "customer", "erp"),
        ),
        StoryCard(
            title="$1M+ account stabilization",
            story_types=("Persuasion", "Challenge and Failure", "Customer Disagreement"),
            hook="The challenge was recovering customer trust after integration, customization, or unresolved workflow issues put accounts at risk.",
            takeaways=("Created one accountable path through the issue", "Listened for the real business pain behind the escalation", "Kept product, development, and customer stakeholders focused on resolution"),
            evidence="Consolidated case ownership, led structured working sessions, and coordinated product and development teams around complex failures.",
            level3_trait="Show what was noticed in the room: the customer needed ownership and a credible recovery path more than another status update.",
            result="Stabilized at-risk accounts representing more than one million dollars in annual revenue.",
            outcome="Use this for customer trust, escalation recovery, and influencing without authority.",
            evidence_terms=("annual revenue", "integration"),
            signals=("risk", "escalation", "retention", "revenue", "integration", "customer success"),
        ),
        StoryCard(
            title="200+ dashboards and decision visibility",
            story_types=("Analysis and Decision", "Individual Achievement", "Ambiguous Problem"),
            hook="The challenge was that leaders needed clearer visibility into performance, workflow friction, and trend signals.",
            takeaways=("Clarified the decision the report needed to support", "Translated operational questions into usable metrics", "Made the output practical for leaders and operators"),
            evidence="Built 200+ dashboards, KPI reports, and analytics tools using SQL, Crystal Reports, and Power BI.",
            level3_trait="Show how the question behind the data was clarified before building the report.",
            result="Improved visibility into operational performance, customer experience metrics, and process gaps.",
            outcome="Use this for data-driven decision-making, analytical structure, and business-minded reporting.",
            evidence_terms=("200+", "KPI", "Power BI"),
            signals=("analytics", "dashboard", "kpi", "reporting", "data", "visibility"),
        ),
        StoryCard(
            title="60+ workshops and QBRs",
            story_types=("Managing and Leading", "Persuasion", "Teamwork"),
            hook="The challenge was keeping executives and delivery stakeholders aligned when each group cared about different outcomes.",
            takeaways=("Read each stakeholder group differently", "Made tradeoffs visible", "Kept the conversation tied to business objectives"),
            evidence="Facilitated 60+ executive workshops and quarterly business reviews focused on roadmap alignment, adoption needs, and business priorities.",
            level3_trait="Show what was noticed: executives needed confidence in outcomes, operators needed workflow clarity, and delivery teams needed decision rights.",
            result="Maintained executive confidence throughout multi-phase delivery programs.",
            outcome="Use this for leadership, executive communication, and working with people from different backgrounds.",
            evidence_terms=("60+ executive workshops", "QBR"),
            signals=("executive", "stakeholder", "qbr", "alignment", "roadmap", "leadership"),
        ),
        StoryCard(
            title="East West ERP ownership",
            story_types=("Managing and Leading", "Ambiguous Problem", "Teamwork"),
            hook="The challenge was owning a mission-critical ERP environment across multiple sites without losing adoption, data, or operational trust.",
            takeaways=("Put structure around ambiguous needs", "Balanced operations, finance, and engineering priorities", "Protected adoption through training and validation"),
            evidence="Owned ERP strategy, administration, and continuous improvement across five sites and more than 150 users.",
            level3_trait="Show how each group was heard differently before requirements and recommendations were finalized.",
            result="Kept core operations and finance workflows running across a global footprint while improving the ERP system through training, testing, and release readiness.",
            outcome="Use this as the main story for role scope, stakeholder complexity, and practical ownership.",
            evidence_terms=("five sites", "150", "enterprise system"),
            signals=("implementation", "go-live", "delivery", "testing", "global", "stakeholder"),
        ),
        StoryCard(
            title="East West Salesforce visibility",
            story_types=("Analysis and Decision", "Teamwork", "Process Improvement"),
            hook="The challenge was giving business teams clearer visibility into requests and customer or project activity during migration work without letting follow-up disappear into spreadsheet tracking.",
            takeaways=("Kept the operating view inside the system", "Connected CRM visibility to ERP and reporting work", "Made owner and next step easier to see across teams"),
            evidence="Used Salesforce alongside ERP data and SQL-backed reporting to track requests, surface owner and next step, and give business teams clearer visibility into customer and project activity during migration and post-go-live support.",
            level3_trait="Show what was noticed: when teams updated side trackers instead of the system, the real blocker was not effort but visibility and ownership.",
            result="Improved cross-functional coordination and reduced manual status chasing during migration and post-go-live support.",
            outcome="Use this for Salesforce adoption, digital mindset, and explaining why system-based workflow visibility is stronger than spreadsheet-driven coordination.",
            evidence_terms=("East West", "Salesforce"),
            signals=("salesforce", "crm", "visibility", "reporting", "workflow", "adoption", "digital", "operations"),
        ),
        StoryCard(
            title="Salesforce backlog and release coordination",
            story_types=("Analysis and Decision", "Managing and Leading", "Teamwork"),
            hook="The challenge was keeping customer-facing CRM and digital-workflow changes useful for customers and support teams instead of turning them into another spreadsheet-driven status exercise.",
            takeaways=("Translated noisy requests into backlog-ready work", "Used testing and release discipline to protect adoption", "Kept the workflow visible in system rather than in side trackers"),
            evidence="Turned business needs into backlog-ready requirements, coordinated UAT, and validated releases across Salesforce customer and marketing workflows.",
            level3_trait="Show what was noticed: when ownership lived in email threads or spreadsheets, follow-up got blurry, so the work had to move into clearer CRM workflows, test scenarios, and next-step tracking.",
            result="Improved post-go-live follow-through, clearer issue ownership, and more reliable coordination across customer-facing teams.",
            outcome="Use this for Salesforce product ownership, backlog management, UAT, release coordination, and explaining why structured CRM workflows beat spreadsheet tracking.",
            evidence_terms=("Service Cloud", "Marketing Cloud"),
            signals=("salesforce", "crm", "digital", "backlog", "uat", "release", "product", "adoption", "workflow", "customer experience"),
        ),
        StoryCard(
            title="LivePerson messaging workflows",
            story_types=("Individual Achievement", "Analysis and Decision", "Rapid Learning"),
            hook="The challenge was helping customer-facing teams adopt a new communication channel with consistent messaging workflows.",
            takeaways=("Learned the workflow through real customer interactions", "Configured repeatable messaging steps", "Used trend monitoring to improve the operating model"),
            evidence=f"Configured LivePerson LiveEngage chat and text workflows, automated greetings and closings, and supported the {COMPANY_HOME_DEPOT} SMS texting pilot.",
            level3_trait="Show how customer interaction trends shaped the workflow language and channel adoption support.",
            result="Improved consistency in customer-facing eCommerce communication workflows.",
            outcome="Use this carefully for practical automation, messaging workflows, conversational AI, or channel adoption.",
            evidence_terms=("LivePerson LiveEngage", "automated greetings and closings"),
            signals=("automation", "ai", "chatbot", "messaging", "workflow", "nlp"),
        ),
        StoryCard(
            title="Aptean lifecycle delivery",
            story_types=("Individual Achievement", "Managing and Leading", "Ambiguous Problem"),
            hook="The challenge was guiding clients from ambiguous business needs into practical ERP scope, delivery, and adoption.",
            takeaways=("Started with discovery before solutioning", "Converted requirements into scope and milestones", "Stayed with clients through go-live and hypercare"),
            evidence="Led discovery, requirements definition, configuration, data migration, integration, testing, go-live, and post-go-live support.",
            level3_trait="Show how vague asks were translated into SOWs, functional requirements, test plans, and delivery checkpoints.",
            result="Helped international clients move through full lifecycle ERP implementation with clearer scope and lower delivery risk.",
            outcome="Use this for implementation, consulting delivery, and structuring ambiguous work.",
            evidence_terms=("requirements", "implementation", "adoption"),
            signals=("discovery", "requirements", "solution", "design", "implementation", "consulting"),
        ),
        StoryCard(
            title="Operations versus finance alignment",
            story_types=("Persuasion", "Teamwork", "Opposing Views"),
            hook="The challenge was balancing operational needs, finance controls, cost, timeline, and system impact when stakeholders did not naturally want the same thing.",
            takeaways=("Listened for the constraint behind each position", "Made tradeoffs explicit", "Recommended the option that protected the business outcome"),
            evidence="Led cross-functional discovery, surfaced the tradeoffs, and negotiated priorities with vendors and internal stakeholders.",
            level3_trait="Show what was noticed: one group was optimizing speed, another was protecting control, and the answer had to make the tradeoff visible.",
            result="Balanced cost, timeline, and operational impact across competing stakeholder interests.",
            outcome="Use this for opposing views, difficult stakeholders, and influence without authority.",
            evidence_terms=("finance", "engineering", "stakeholder"),
            signals=("opposing", "disagree", "stakeholder", "finance", "operations", "persuasion"),
        ),
        StoryCard(
            title="Failure lesson and stronger validation",
            story_types=("Challenge and Failure", "Analysis and Decision", "Rapid Learning"),
            hook="The lesson is that unclear requirements or weak validation can turn a solvable system issue into a larger adoption problem.",
            takeaways=("Own the miss without over-explaining", "Show the control that changed afterward", "Connect the lesson to better delivery risk management"),
            evidence="Led go-live readiness, sandbox testing, user acceptance validation, issue triage, and release readiness across ERP work.",
            level3_trait="Show the changed behavior: clearer requirements, stronger validation checkpoints, more explicit rollback planning, and earlier stakeholder signoff.",
            result="Reduced production disruption, downstream defects, and implementation risk across concurrent program tracks.",
            outcome="Use this for failure questions. Keep it honest, calm, and focused on what changed.",
            evidence_terms=("validation", "cutover coordination"),
            signals=("failure", "mistake", "learn", "testing", "validation", "risk"),
        ),
        StoryCard(
            title="Customer loss and proactive success lesson",
            story_types=("Challenge and Failure", "Customer Disagreement", "Persuasion"),
            hook="The challenge was recovering a customer relationship after inheriting a broken implementation — and still losing the account despite resolving every technical issue.",
            takeaways=(
                "Owned the relationship directly rather than managing it through escalation",
                "Negotiated feature acceleration to rebuild trust faster than a standard roadmap allowed",
                "Learned that waiting for a customer to raise concerns means the decision is already made",
            ),
            evidence="Took ownership of an at-risk account at a manufacturing ERP company where an incorrectly configured integration had eroded trust; met with the customer president weekly and worked across product and development to accelerate key roadmap items from a six-month timeline into a two-month beta release.",
            level3_trait="Show what was noticed: the customer needed a clear owner and a credible path forward, not another status update. Then show what changed: reaching out before a customer has a reason to complain, not after, because by the time someone raises their hand the decision is often already made.",
            result="Resolved every technical issue the customer had raised. The account still churned when the customer chose a cheaper competitor. The loss clarified the proactive customer success model applied in every engagement afterward.",
            outcome="Use this for failure questions, customer churn questions, or any question about proactive account management and what was learned from a loss. It is the strongest story for roles where customer health ownership is an explicit expectation.",
            evidence_terms=("at-risk annual revenue",),
            signals=("failure", "churn", "loss", "customer", "proactive", "retention", "account", "escalation", "discovery", "executive", "consulting", "transformation", "enablement"),
        ),
        StoryCard(
            title="13-month modernization complexity",
            story_types=("Ambiguous Problem", "Managing and Leading", "Persuasion"),
            hook="The challenge was discovering mid-engagement that a customer who wanted to modernize their ERP was running infrastructure too outdated to support any modern software at all.",
            takeaways=(
                "Surfaced a constraint the customer had not anticipated and could not work around",
                "Aligned CEO and upper management on real costs before any software work could begin",
                "Kept the engagement alive through a 13-month delivery when the scope was set for four to seven months",
            ),
            evidence="Led a full ERP modernization engagement where requirements gathering revealed tens of thousands of dollars in required hardware upgrades before implementation could begin. I delivered a satisfied customer and billable customization work through an engagement that ran nearly three times the standard timeline.",
            level3_trait="Show what was noticed: the customer asked for software, but the real constraint was infrastructure. Show what was done: named the problem directly to leadership instead of softening it, managed expectations through a significantly longer delivery, and kept the customer confident enough to stay.",
            result="Delivered a satisfied customer after a 13-month engagement scoped at four to seven months. The extended timeline opened billable customization work that would not have existed in a standard delivery.",
            outcome="Use this for most complex implementation, stakeholder alignment under pressure, managing scope surprises, or expectations management with executive audiences who did not anticipate the real cost or timeline of the work they asked for.",
            evidence_terms=("implementation", "go-live"),
            signals=("complex", "implementation", "timeline", "stakeholder", "executive", "scope", "modernization", "discovery", "consulting", "transformation"),
        ),
        StoryCard(
            title="UAT defect catch before go-live",
            story_types=("Challenge and Failure", "Persuasion", "Analysis and Decision"),
            hook="The challenge was discovering during UAT that a defect would break a live client workflow if it reached production.",
            takeaways=(
                "Named the go-live risk directly instead of softening it",
                "Coordinated root-cause work quickly across development and product partners",
                "Protected the client outcome even when that meant slowing the timeline",
            ),
            evidence="Identified a critical defect during user acceptance testing, led triage with development, validated the fix, and withheld go-live approval until the workflow was safe.",
            level3_trait="Show what was noticed: the real risk was not a bug count but the production impact on a live client process, so the conversation stayed anchored on business harm, validation, and release readiness instead of schedule pressure.",
            result="Prevented a production issue that would have disrupted live client operations after go-live.",
            outcome="Use this for delivery risk management, quality validation, cross-functional coordination, or any question about making a difficult go-live call.",
            evidence_terms=("UAT", "user acceptance", "go-live"),
            signals=("uat", "testing", "risk", "delivery", "client management", "validation", "defect", "go-live"),
        ),
        StoryCard(
            title="CEO hardware scoping conversation",
            story_types=("Persuasion", "Customer Disagreement", "Managing and Leading"),
            hook="The challenge was getting an executive sponsor to treat outdated infrastructure as a business risk before ERP work could move forward.",
            takeaways=(
                "Diagnosed the real blocker before talking solutions",
                "Framed hardware upgrades as implementation risk rather than IT preference",
                "Kept leadership, vendors, and technical teams aligned on readiness",
            ),
            evidence="Scoped server and hardware requirements with leadership, vendors, and IT teams to confirm compatibility, capacity, security, and upgrade readiness before ERP deployment.",
            level3_trait="Show what was noticed: leadership thought the project was a software decision, but the real constraint was infrastructure readiness, so the discussion had to shift from features to business exposure if the environment failed under live load.",
            result="Secured infrastructure-readiness alignment early enough to prevent post-deployment performance failures.",
            outcome="Use this for executive persuasion, technical scoping, stakeholder alignment, or surfacing hidden delivery risk before go-live.",
            evidence_terms=("hardware", "infrastructure", "upgrade readiness"),
            signals=("persuasion", "stakeholder", "executive", "hardware", "implementation", "risk", "scope", "infrastructure"),
        ),
        StoryCard(
            title="Amazon Robotics warehouse certification",
            story_types=("Individual Achievement", "Managing and Leading", "Ambiguous Problem"),
            hook="The challenge was standing up an entirely new warehouse facility in the ERP and passing Amazon's compliance and certification requirements before anything could go live.",
            takeaways=(
                "Managed a multi-stakeholder cross-functional process where every configuration decision had a downstream certification consequence",
                "Coordinated across internal finance, operations, and external compliance simultaneously",
                "Delivered on Amazon's timeline without flexibility to learn by doing",
            ),
            evidence="Led the Amazon Robotics warehouse setup at East West Manufacturing: a six-month compliance and certification process that required configuring every product family, bill of materials, and component structure in the ERP across coordination with the CFO, plant controllers at every site, manufacturers, vendors, and Amazon's compliance team.",
            level3_trait="Show the constraint that made this hard: Amazon's certification requirements meant there was no room for iteration — everything had to be right before it could go live, and every configuration decision upstream affected compliance downstream.",
            result="Achieved full Amazon Robotics certification and delivered a live operational warehouse environment.",
            outcome="Use this for high-stakes cross-functional delivery, compliance-constrained implementation, or any question about the most complex project managed. It is the strongest manufacturing execution proof with named external stakeholders and no room for error.",
            evidence_terms=("Amazon Robotics",),
            signals=("manufacturing", "implementation", "compliance", "delivery", "executive", "stakeholder", "go-live", "complex"),
        ),
    ]


def hero_stories(profile: build_resume.JobProblemProfile, job_description: str, resume_text: str) -> list[StoryCard]:
    supported = [card for card in expanded_story_bank() if contains_all(resume_text, card.evidence_terms)]
    if not supported:
        return []

    lane_bonus_terms = {
        "presales_solution": {"discovery", "solution", "pre-sales", "executive"},
        "customer_success": {"risk", "retention", "customer success", "executive"},
        "analytics_operations": {"analytics", "dashboard", "kpi", "data"},
        "change_enablement": {"training", "adoption", "stakeholder", "workflow"},
        "process_improvement": {"process", "optimization", "efficiency", "analysis", "workflow", "operations", "data"},
        "implementation_delivery": {"implementation", "go-live", "delivery", "testing"},
        "corporate_strategy": {"consulting", "executive", "discovery", "stakeholder", "analysis", "transformation", "client"},
    }
    lane_bonus = lane_bonus_terms.get(profile.primary_lane, set())

    def score(card: StoryCard) -> int:
        bonus = sum(1 for signal in card.signals if signal in lane_bonus)
        return signal_score(job_description, card.signals) + bonus

    supported.sort(key=score, reverse=True)
    return supported[:5]


def supported_story_bank(resume_text: str) -> list[StoryCard]:
    return [card for card in expanded_story_bank() if contains_all(resume_text, card.evidence_terms)]


def story_for_type(stories: list[StoryCard], story_type: str) -> StoryCard | None:
    if story_type == "Challenge and Failure":
        # Prefer the churn/loss story when available: it carries a specific, actionable lesson
        # about proactive customer success that resonates more broadly than a go-live testing
        # miss, and it performed well in actual interviews (Plataine, June 2026).
        preferred = next((card for card in stories if card.title == "Customer loss and proactive success lesson"), None)
        if preferred:
            return preferred
        preferred = next((card for card in stories if card.title == "Failure lesson and stronger validation"), None)
        if preferred:
            return preferred
    return next((card for card in stories if story_type in card.story_types), None)


def story_theme_key(card: StoryCard) -> str:
    lowered = card.title.lower()
    if "inventory" in lowered:
        return "inventory"
    if "account" in lowered or "$1m" in lowered:
        return "account"
    if "dashboard" in lowered or "decision visibility" in lowered:
        return "dashboards"
    if "rapid" in lowered or "product learning" in lowered:
        return "learning"
    if "operations versus finance" in lowered or "finance alignment" in lowered:
        return "ops_finance"
    if "failure" in lowered or "validation" in lowered:
        return "failure"
    if "workshop" in lowered or "qbr" in lowered:
        return "workshops"
    if "customer loss" in lowered or "proactive success" in lowered:
        return "customer_loss"
    if "13-month" in lowered or "modernization complexity" in lowered:
        return "modernization_scope"
    if "amazon robotics" in lowered or "warehouse certification" in lowered:
        return "amazon_robotics"
    if "erp ownership" in lowered:
        return "erp_ownership"
    if "salesforce visibility" in lowered:
        return "crm_visibility"
    if "backlog" in lowered or "release coordination" in lowered:
        return "backlog_release"
    if "liveperson" in lowered or "messaging workflows" in lowered:
        return "messaging_automation"
    if "lifecycle delivery" in lowered:
        return "lifecycle_delivery"
    return "default"


def story_specific_bridge(card: StoryCard, profile: build_resume.JobProblemProfile) -> str:
    key = story_theme_key(card)
    bridges = {
        "inventory": "Bridge: this is the process-improvement proof: map the actual workflow, find the structural gap, validate the fix with users, pilot it, and measure whether the work actually changed.",
        "account": "Bridge: this is the trust-recovery proof: when an experience is breaking down, the answer is accountable ownership and a credible path forward, not a better status cadence.",
        "dashboards": "Bridge: this is the decision-quality proof: define the business decision before touching the data, validate the source, and segment the view so the next action is obvious.",
        "learning": "Bridge: this is the ramp proof: learn through the live workflow and the people doing it, not through documentation alone.",
        "ops_finance": "Bridge: this is the stakeholder-tradeoff proof: surface what each group is protecting before designing a process that has to satisfy all of them.",
        "failure": "Bridge: this is the quality-control proof: build SME validation, acceptance criteria, and checkpoints into the process before go-live, not after the risk reaches users.",
        "workshops": "Bridge: this is the translation proof: turn one process goal into the decision, workflow, or risk language each audience needs to act.",
        "erp_ownership": "Bridge: this is the systems-ownership proof: keep a mission-critical platform stable for every stakeholder group while still pushing through training, testing, and release improvements instead of just maintaining the status quo.",
        "crm_visibility": "Bridge: this is the visibility proof: when work lives in side trackers instead of the system of record, the fix is moving ownership and next steps back into the platform everyone already uses.",
        "backlog_release": "Bridge: this is the release-discipline proof: turn noisy, ambiguous requests into backlog-ready work with real UAT and validation so adoption survives past go-live.",
        "messaging_automation": "Bridge: this is the channel-adoption proof: learn the new workflow through live customer interactions first, then standardize the steps so the whole team can repeat it.",
        "lifecycle_delivery": "Bridge: this is the discovery-to-delivery proof: turn an ambiguous ask into defined scope, milestones, and checkpoints, then stay through go-live so adoption actually holds.",
        "customer_loss": "Bridge: this is the proactive account ownership proof: success in a high-touch customer role means identifying risk before the customer names it, because by the time they raise their hand the decision may already be made.",
        "modernization_scope": "Bridge: this is the scoping realism proof: the most dangerous assumption in a complex implementation is that the customer's environment matches what they believe it to be; surface constraints early, name the real cost directly, and hold expectations across a longer-than-expected delivery.",
        "amazon_robotics": "Bridge: this is the compliance-constrained delivery proof: when a customer or partner has non-negotiable certification requirements, there is no room to learn by doing — every configuration decision upstream has to account for what it unlocks or blocks downstream.",
    }
    return bridges.get(
        key,
        f"Bridge: connect this specific outcome to {candidate_problem_phrase(profile)} by naming the workflow, decision, stakeholder, risk, or customer problem it proves.",
    )


def response_calibration_lines() -> list[str]:
    return [
        "Classify the question first: behavioral stories use the primary framework named for this interview; future approach questions use philosophy-first, then one supporting proof point and a role bridge.",
        "Use PREP for short direct answers: Point, Reason, Example, Point restatement. This is for 15-20 second yes/no or quick-fit questions.",
        "Use CART for consulting, advisory, senior, VP, executive, strategic recommendation, or board-level stakeholder contexts; state the impact first in Context.",
        "Calibrate length before answering: true/false = 15 seconds, multiple choice = about 30 seconds, fill-in-the-blank = 20-30 seconds, essay or behavioral = 60-90 seconds.",
        "Structural audit: by sentence two, the listener should already know the stakes, Christian's role, or the concrete proof. If the answer is still warming up, restart and land the point sooner.",
        "Tonal audit: sound calm, factual, and reflective. Do not oversell, over-apologize, or overperform the wording; let the answer sound like a real memory backed by rehearsal.",
        "Tell them the time, do not build them the clock. If a yes/no answer reaches sentence five, land the result and stop.",
    ]


def has_executive_round_signal(job_description: str, interview_notes: str = "") -> bool:
    combined = f"{job_description}\n{interview_notes}"
    return bool(
        re.search(
            r"\b(?:vp|svp|evp|vice president|senior vice president|executive|c-suite|c suite|chief\s+\w+|ceo|cfo|coo|cto|cio|board|board-level|senior director|senior leader|senior leadership|senior executive|senior stakeholder|senior management|managing director|partner|principal)\b",
            combined,
            re.I,
        )
    )


def has_cart_strategy_signal(job_description: str, interview_notes: str = "") -> bool:
    return build_resume.jd_mentions(
        f"{job_description}\n{interview_notes}",
        "strategic",
        "recommendation",
        "recommendations",
        "executive presence",
        "stakeholder alignment at the board level",
        "stakeholder alignment at board level",
        "board-level stakeholder alignment",
    )


def answer_framework_selection(job_description: str, interview_notes: str = "") -> AnswerFrameworkSelection:
    reasons: list[str] = []
    if build_resume.is_consulting_job_description(job_description):
        reasons.append("the job description signals a consulting or advisory context")
    if has_executive_round_signal(job_description, interview_notes):
        reasons.append("the job description or interview notes signal senior, VP, board, or executive-level evaluation")
    if has_cart_strategy_signal(job_description, interview_notes):
        reasons.append("the role language emphasizes strategy, recommendations, executive presence, or board-level stakeholder alignment")

    if reasons:
        return AnswerFrameworkSelection(
            primary_framework="CART",
            label="CART - Context, Approach, Role, Transformation",
            reason="; ".join(reasons),
            guidance=(
                "Lead with impact in Context by stating the result first, then explain the situation only as much as needed.",
                "Use Approach to show the structure of the thinking, including what Christian noticed before acting.",
                "Use Role to make Christian's personal contribution clear without overstating formal authority.",
                "Use Transformation to close with what changed, why it mattered, and how the lesson applies to this role.",
            ),
        )

    return AnswerFrameworkSelection(
        primary_framework="HOOK_NOTICING",
        label="Hook, Noticing, Action, Result, Bridge, Calibration",
        reason="the role reads as a non-consulting, non-senior implementation, customer success, analytics, or change enablement interview",
        guidance=(
            "Use Hook to name the business problem or delivery challenge before the task details.",
            "Use Noticing to show judgment: what Christian saw, diagnosed, or understood before acting.",
            "Use Action and Result to prove the work with supported scope, metrics, and outcomes.",
            "Use Bridge and Calibration to connect the story back to the role and invite the interviewer to confirm the priority.",
        ),
    )


def answer_framework_section_lines(selection: AnswerFrameworkSelection) -> list[str]:
    lines = [
        f"Primary framework: {selection.label}.",
        f"Why this framework: {selection.reason}.",
    ]
    lines.extend(selection.guidance)
    lines.extend([
        "Use PREP - Point, Reason, Example, Point restatement - for yes/no, factual, or under-20-second questions.",
        "Use philosophy-first when the question asks how Christian handles something: state the repeatable process first, then give one supporting story as evidence.",
    ])
    return lines


def candidate_archetype_assessment(
    profile: build_resume.JobProblemProfile,
    job_description: str,
    resume_text: str,
    supplied_context: str = "",
    interview_notes: str = "",
) -> tuple[str, str]:
    combined = f"{job_description}\n{interview_notes}\n{supplied_context}"

    final_round = bool(re.search(r"\b(?:final round|final interview|final-stage|final stage|last round|round\s*(?:3|4|5|6|7|8|9)|third round|fourth round|fifth round)\b", combined, re.I))
    round_three_or_higher = bool(re.search(r"\b(?:round\s*(?:3|4|5|6|7|8|9)|third round|fourth round|fifth round)\b", combined, re.I))
    executive_interview = bool(re.search(r"\b(?:executive interview|executive round|vp interview|vice president interview|c-suite|c suite|board-level|board level)\b", combined, re.I))
    panel_round = detect_panel_context(job_description, interview_notes) or mentions(interview_notes, "panel")
    technical_track = (
        profile.primary_lane == "analytics_operations"
        or mentions(job_description, "engineering", "data", "analytics", "technical", "sql", "api", "etl", "data warehouse", "reporting", "dashboard")
        or mentions(supplied_context, "technical projects", "excel", "jira", "azure")
        or bool(re.search(r"\b(?:engineer|developer|analyst|technical consultant|technical account|data analyst|business intelligence)\b", job_description, re.I))
    )
    if mentions(supplied_context, "excel", "training others", "technical projects"):
        return (
            "Round-specific proof risk",
            "The last round tested Excel comfort, training others, and technical projects. Prepare one concise answer for each before repeating broad implementation stories.",
        )

    assessments = (
        (
            "Passive Candidate risk",
            final_round or executive_interview,
            "You are evaluating them as much as they are evaluating you. Walk in with that energy. Use the hesitancy close and commitment lock before you leave.",
        ),
        (
            "Nervous Candidate risk",
            round_three_or_higher or mentions(job_description, "high-stakes", "executive presence"),
            "Run the full pre-interview warm-up: power pose for two minutes, five diaphragmatic breaths, mantra aloud. Physiological preparation is not optional at this stage.",
        ),
        (
            "Over-Preparer risk",
            panel_round or final_round,
            "In panel or final rounds, brittle word-for-word delivery is detectable. Rehearse five anchors per story until the structure is automatic, then let the wording flex naturally.",
        ),
        (
            "Technician risk",
            technical_track,
            "Every technical answer must land as a business outcome. Before stating any technical detail, say what it enabled or improved for the customer, team, or operation.",
        ),
        (
            "Under-Seller risk",
            mentions(job_description, "quota", "revenue", "ownership", "accountability"),
            "In revenue and ownership roles, use 'I' for every action you personally drove. Audit each story: did I actually do this? Then own it fully.",
        ),
        (
            "Rambler risk",
            mentions(job_description, "concise", "executive communication", "C-suite", "board-level"),
            "This audience expects short, structured answers. Use PREP for direct questions. Stop after the result. Let them pull for more detail.",
        ),
    )
    for label, condition, warning in assessments:
        if condition:
            return label, warning
    return (
        "Rambler risk",
        "This audience expects short, structured answers. Use PREP for direct questions. Stop after the result. Let them pull for more detail.",
    )


def role_specific_track_guidance(
    profile: build_resume.JobProblemProfile,
    role_title: str,
    job_description: str,
) -> tuple[str, tuple[str, ...]] | None:
    lane_hint = profile.primary_lane
    executive_track = bool(
        re.search(r"\b(?:vp|director|head of|chief|svp|president)\b", role_title, re.I)
        or build_resume.jd_mentions(job_description, "leadership philosophy", "P&L", "board", "organizational design")
    )
    sales_track = build_resume.jd_mentions(
        job_description,
        "quota",
        "pipeline",
        "demo",
        "discovery",
        "account executive",
        "sales",
        "closing",
        "ARR",
        "revenue",
    )
    pm_track = build_resume.jd_mentions(
        job_description,
        "product manager",
        "product management",
        "roadmap",
        "prioritization",
        "product sense",
        "feature",
        "user story",
    )
    technical_track = build_resume.jd_mentions(
        job_description,
        "software engineer",
        "engineering",
        "data scientist",
        "machine learning",
        "system design",
        "SQL",
        "python",
        "backend",
        "frontend",
    )

    if executive_track:
        return (
            "Executive Track",
            (
                "By this round, competence is assumed. The evaluation is: do they want to work with you every day? Show genuine enthusiasm and strategic depth.",
                "Prepare a fully quantified Hero Story with financial translation: what did it mean in dollars, time, or scale to the organization?",
                "Leadership philosophy question: state your belief, give one specific behavior, tell a brief story, and describe how your thinking has evolved.",
                "Match the interviewer archetype: Talker wants to drive the conversation; ask questions about their role, their journey, their views. Listener wants to hear your stories.",
                "Build relationships, not performances. The executive who likes you as a person is doing the most important evaluation in the room.",
            ),
        )
    if sales_track:
        return (
            "Sales Track",
            (
                "Your ability to sell yourself is meta-evidence of your ability to sell their product. The interview IS the demo.",
                "Prepare your biggest deal story: size, how you found it, stakeholders, what almost killed it, how you navigated it, timeline, and what you learned.",
                "Know your quota attainment number and be prepared to be specific and honest. Never round up.",
                "If there is a mock call component, use 70 percent questions and listening, 30 percent talking. Ask before pitching.",
                "For 'walk me through your sales process,' cover: prospecting, qualification, discovery, value articulation, objection handling, closing, and post-sale.",
            ),
        )
    if pm_track:
        return (
            "PM Track",
            (
                "Prioritization questions: always clarify the strategic goal first. Is this role optimizing for growth, retention, monetization, or something else?",
                "Product design questions: user first, then success metrics, then solutions, then one prioritized recommendation with explicit tradeoffs.",
                "Metrics questions: name the primary metric, the guardrail metric, and one leading indicator. Explain what failure looks like.",
                "RICE scoring is your prioritization framework: Reach times Impact times Confidence divided by Effort.",
                "For build vs. buy vs. partner: first determine whether this capability is a competitive moat or a commodity function. That drives everything.",
            ),
        )
    if technical_track:
        return (
            "Technical Track",
            (
                "Technical skill plus poor communication loses to technical skill plus strong communication. Narrate your thinking out loud throughout.",
                "System design questions: start with clarifying questions on scale and constraints before drawing anything.",
                "For behavioral questions in technical roles, translate every technical action into a business or user outcome.",
                "Disagreement with a technical decision story: show independent thinking plus collaborative resolution, not stubbornness.",
                "Handling non-technical stakeholders question: show how you translate technical constraints into business decisions without condescension.",
            ),
        )
    return None


def six_story_type_lines(stories: list[StoryCard]) -> list[str]:
    story_types = (
        "Individual Achievement",
        "Managing and Leading",
        "Persuasion",
        "Analysis and Decision",
        "Challenge and Failure",
        "Teamwork",
    )
    lines: list[str] = []
    for story_type in story_types:
        card = story_for_type(stories, story_type)
        if card:
            lines.append(f"{story_type}: Use {card.title}. Representative prompt: {representative_prompt(story_type)} What makes it land: {card.outcome}")
        else:
            lines.append(
                f"{story_type}: No resume-supported story available for this type. "
                "Build one before the interview using the career events most relevant to this role."
            )
    return lines


def representative_prompt(story_type: str) -> str:
    prompts = {
        "Individual Achievement": "Tell me about a professional accomplishment you are proud of.",
        "Managing and Leading": "Tell me about a time you led people through a difficult situation.",
        "Persuasion": "Tell me about a time you had to convince someone who disagreed with you.",
        "Analysis and Decision": "Tell me about a difficult decision you made with limited information.",
        "Challenge and Failure": "Describe a failure or tough lesson and how you responded.",
        "Teamwork": "Describe a time you worked with someone whose style or background was different from yours.",
    }
    return prompts.get(story_type, "Tell me about a relevant example.")


def calibration_question(card: StoryCard, profile: build_resume.JobProblemProfile) -> str:
    key = story_theme_key(card)
    if key == "inventory":
        return "Is finding the structural root cause, rather than building better error detection, the kind of thinking this team is looking for?"
    if key == "account":
        return "Is cross-functional coordination without formal authority a regular part of how this role operates?"
    if key == "dashboards":
        return "Is defining the business decision before touching the data a discipline the team is still building, or is that infrastructure mature?"
    if key == "learning":
        return "Is a 30-to-60 day ramp where I learn the operating environment before recommending anything the right expectation?"
    if key == "ops_finance":
        return "Is navigating competing-priority tension, where different teams define success differently, a regular part of this role?"
    if key == "failure":
        return "Is building SME validation into process changes before go-live a strong discipline here already, or an area to strengthen?"
    if key == "workshops":
        return "How often does this role need to translate the same change differently for leaders, operators, and frontline teams?"
    if "Managing and Leading" in card.story_types or "Teamwork" in card.story_types:
        return "Is cross-functional coordination and stakeholder alignment a regular part of how this role will need to operate?"
    if "Challenge and Failure" in card.story_types:
        return "Is managing delivery risk and building better checkpoints part of what this role needs to own?"
    if "Persuasion" in card.story_types or "Opposing Views" in card.story_types:
        return "Is earning alignment across groups who see the problem differently a core part of how this role succeeds?"
    return f"Is that the kind of challenge you expect the {profile.lane_label.lower()} work to require here?"


def caar_answer(card: StoryCard, profile: build_resume.JobProblemProfile) -> list[str]:
    """Working story sequence: Hook -> Noticing -> Action -> Result -> Bridge -> Calibration."""
    takeaways = "; ".join(card.takeaways)
    return [
        f"Hook: {card.hook}",
        f"Noticing: {card.level3_trait}",
        f"Action: {card.evidence}",
        f"Result: {card.result}",
        story_specific_bridge(card, profile),
        f"Calibration: {calibration_question(card, profile)}",
        f"Three takeaways: {takeaways}.",
    ]


def cart_answer(card: StoryCard, profile: build_resume.JobProblemProfile) -> list[str]:
    """Senior or consulting sequence: Context -> Approach -> Role -> Transformation."""
    takeaways = "; ".join(card.takeaways)
    return [
        f"CONTEXT (impact first): {card.result} The situation mattered because the underlying challenge was this: {card.hook} I chose to act because the work connected directly to {candidate_problem_phrase(profile)} and needed clear ownership.",
        f"APPROACH: My approach was to first identify the real business risk, then structure the work around the people and process affected, and finally validate the outcome before moving forward. Here are the three things I did: {takeaways}.",
        f"ROLE: I personally worked through each step: I discovered what mattered most, I decided where structure or validation was needed, and I acted through this concrete work: {card.evidence} What I noticed that others could miss was this: {card.level3_trait}",
        f"TRANSFORMATION: The primary result was: {card.result} The downstream benefit was stronger {candidate_problem_phrase(profile)}, and the work created a repeatable practice for future decisions, stakeholder alignment, and risk control. {story_specific_bridge(card, profile).replace('Bridge: ', '')}",
    ]


def should_use_cart(company_name: str, role_title: str, job_description: str) -> bool:
    return bool(
        build_resume.is_consulting_job_description(job_description)
        or re.search(r"\b(?:vp|director|head of|senior director)\b", role_title, re.I)
        or build_resume.jd_mentions(
            job_description,
            "executive stakeholders",
            "board",
            "C-suite",
            "leadership alignment",
        )
    )


def uses_star_answer_framework(company_name: str, job_description: str) -> bool:
    return build_resume.jd_mentions(job_description, "star method", "situation task action result")


def pyramid_answer(card: StoryCard, profile: build_resume.JobProblemProfile) -> list[str]:
    takeaways = "; ".join(card.takeaways)
    return [
        f"Direct answer: Yes. The way I would approach that is to structure the problem first, then use evidence and stakeholder alignment to move toward a practical outcome.",
        f"Three takeaways: {takeaways}.",
        f"Vignette 1: {card.hook} {card.level3_trait}",
        f"Vignette 2: {card.evidence}",
        f"Result: {card.result}",
        f"Role bridge: {story_specific_bridge(card, profile).replace('Bridge: ', '')}",
        f"Calibration question: {calibration_question(card, profile)}",
    ]


def spoken_caar_answer(card: StoryCard, profile: build_resume.JobProblemProfile) -> str:
    opener = story_opener_by_type(set(card.story_types), card.title)
    parts = [
        f"{opener[:1].upper() + opener[1:]} {lower_clause(card.hook)}",
    ]
    if card.level3_trait:
        parts.append(spoken_level3_trait_sentence(card.level3_trait))
    if card.evidence:
        parts.append(story_evidence_sentence(card.evidence))
    if card.result:
        parts.append(story_result_sentence(card.result))
    bridge = story_specific_bridge(card, profile).replace("Bridge: ", "")
    if bridge and spoken_word_count(interview_join(*parts, bridge)) <= STANDARD_SPOKEN_WORD_RANGE[1]:
        parts.append(bridge[:1].upper() + bridge[1:])
    return interview_join(*parts)


def spoken_cart_answer(card: StoryCard, profile: build_resume.JobProblemProfile) -> str:
    opener = story_opener_by_type(set(card.story_types), card.title)
    parts = [
        f"{opener[:1].upper() + opener[1:]} {lower_clause(card.hook)}",
        cart_takeaway_sentence(card),
    ]
    if card.evidence:
        parts.append(story_evidence_sentence(card.evidence))
    if card.level3_trait:
        parts.append(spoken_level3_trait_sentence(card.level3_trait))
    if card.result:
        parts.append(story_result_sentence(card.result))
    bridge = story_specific_bridge(card, profile).replace("Bridge: ", "")
    if bridge and spoken_word_count(interview_join(*parts, bridge)) <= STANDARD_SPOKEN_WORD_RANGE[1]:
        parts.append(bridge[:1].upper() + bridge[1:])
    return interview_join(*parts)


def cart_takeaway_sentence(card: StoryCard) -> str:
    steps = [lower_clause(item) for item in card.takeaways if item]
    if not steps:
        return "I kept the business risk visible while moving the work forward"
    if len(steps) == 1:
        return f"My approach was to {steps[0]} while keeping the business risk visible"
    if len(steps) == 2:
        return f"First, I {steps[0]}. Then I {steps[1]} while keeping the business risk visible"
    return (
        f"First, I {steps[0]}. Then I {steps[1]}. Finally, I {steps[2]} while keeping the business risk visible"
    )


def spoken_pyramid_answer(card: StoryCard, profile: build_resume.JobProblemProfile) -> str:
    opener = story_opener_by_type(set(card.story_types), card.title)
    parts = [
        "Yes, I have handled that kind of situation by structuring the problem first and then using evidence and stakeholder alignment to move toward a practical outcome",
        f"{opener[:1].upper() + opener[1:]} {lower_clause(card.hook)}",
    ]
    if card.level3_trait:
        parts.append(spoken_level3_trait_sentence(card.level3_trait))
    if card.evidence:
        parts.append(story_evidence_sentence(card.evidence))
    if card.result:
        parts.append(story_result_sentence(card.result))
    bridge = story_specific_bridge(card, profile).replace("Bridge: ", "")
    if bridge and spoken_word_count(interview_join(*parts, bridge)) <= STANDARD_SPOKEN_WORD_RANGE[1]:
        parts.append(bridge[:1].upper() + bridge[1:])
    return interview_join(*parts)


def spoken_story_answer(
    card: StoryCard,
    profile: build_resume.JobProblemProfile,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
) -> str:
    if uses_star_answer_framework(company_name, job_description):
        answer = spoken_pyramid_answer(card, profile)
    elif should_use_cart(company_name, role_title, job_description):
        answer = spoken_cart_answer(card, profile)
    else:
        answer = spoken_caar_answer(card, profile)
    answer = prose_engine.spoken_register(answer).text
    assert_full_spoken_answer(f"{card.title} story answer", answer, min_words=40)
    return answer


def great_eight_story_audit(stories: list[StoryCard]) -> dict[str, str]:
    audit_note = (
        "AUDIT: This story result does not clearly connect to a business outcome from the Great Eight. "
        "Add a sentence that names what the business gained: revenue, cost, efficiency, quality, risk, compliance, customer experience, or team performance."
    )
    return {
        card.title: audit_note
        for card in stories
        if not has_great_eight_signal(card.result)
    }


def story_bank_bullets(stories: list[StoryCard]) -> list[str]:
    audit_notes = great_eight_story_audit(stories)
    lines: list[str] = []
    for card in stories:
        line = f"{card.title}: Types: {', '.join(card.story_types)}. Takeaways: {'; '.join(card.takeaways)}. Result: {card.result}"
        if card.title in audit_notes:
            line = f"{line} {audit_notes[card.title]}"
        lines.append(line)
    return lines


def likely_gaps(job_description: str) -> list[str]:
    gaps: list[str] = []
    if mentions(job_description, "talent acquisition", "recruiting", "ats", "applicant tracking"):
        gaps.append(
            "Talent acquisition domain: be direct that ATS and recruiting workflow exposure is adjacent rather than primary. Bridge to configurable enterprise software, workflow discovery, reporting, adoption, and stakeholder training."
        )
    if mentions(job_description, "$100k", "100k arr", "6+ month", "sales cycles"):
        gaps.append(
            "Enterprise SaaS deal-cycle scale: explain that direct ARR ownership was not always visible, but pre-sales discovery, demos, executive workshops, scope alignment, and $1M+ risk recovery are directly supported."
        )
    if mentions(job_description, "demo infrastructure", "demo engineering"):
        gaps.append(
            "Demo engineering: emphasize tailored demos, configuration thinking, reporting, and solution design. Avoid claiming ownership of formal demo infrastructure unless asked about transferable experience."
        )
    if mentions(job_description, "python", "llm", "agent", "api", "jwt", "endpoint"):
        gaps.append(
            "Advanced engineering or AI-agent ownership: keep the answer honest. Supported experience includes Codex-assisted automation, LivePerson LiveEngage workflows, reporting, and workflow automation, not production engineering ownership."
        )
    if not gaps:
        gaps.append(
            "Main gap to manage: avoid sounding too broad. Anchor each answer to the job's problem, then give one concrete proof point and one result."
        )
    return gaps


def role_specific_gaps(profile: build_resume.JobProblemProfile, job_description: str) -> list[str]:
    gaps = likely_gaps(job_description)
    if profile.unsupported_requirements:
        gaps.insert(
            0,
            "Unsupported requirement areas: keep the edge brief, then redirect to the closest resume-backed pattern and proof instead of stretching the claim."
        )
    if profile.primary_lane == "presales_solution":
        gaps.append(
            "If asked about quota or closing ownership, clarify that the strongest support is in discovery, solution design, demos, and technical buyer confidence rather than direct quota carrying."
        )
    elif profile.primary_lane == "customer_success":
        gaps.append(
            "If asked about renewals or expansion ownership, focus on account stabilization, executive reviews, adoption support, and revenue-risk protection."
        )
    elif profile.primary_lane == "implementation_delivery":
        gaps.append(
            "If asked about formal program leadership, stay grounded in implementation ownership, cross-functional coordination, testing, and go-live support rather than claiming broad PMO authority."
        )
    return list(dict.fromkeys(gaps))


def recruiter_reminders() -> list[str]:
    return [
        "Value first, proof second, chronology never.",
        "Classify the question before answering: past experience, future approach, or short direct answer.",
        "Answer in 45 to 90 seconds unless they ask for more. Short answers invite follow-up.",
        "Use PREP for short yes/no or quick-fit questions: Point, Reason, Example, Point restatement.",
        "Start with the business problem, then evidence, result, and why it matters for this role.",
        "Use one example at a time. Do not stack every relevant experience into one answer.",
        "Use the word 'because' in all why-answers. It forces a concrete reason.",
        "End sentences with a downward tone - sounds confident. Upward pitch sounds uncertain.",
        "Eliminate fillers before the interview: 'as well,' 'you know,' 'also.'",
        "Pause before answering. Speed is not competence - an appropriate response at the right time is.",
        "When a question is complex, paraphrase it before answering: 'So what you are asking is...' This buys time and shows listening.",
        "For on-demand video, use Result -> Method -> Bridge. Do not start with long context because there is no live interviewer to pull you back.",
        "After stating a major metric, hold eye contact through the camera lens for two full seconds before continuing.",
        "If you lose your place, use the recovery phrase: 'Let me come back to that from a slightly different angle,' then land the result.",
        "VIDEO SETUP: Camera at eye level, torso in frame. Light source behind the camera facing you.",
        "VIDEO SETUP: Clean professional background. Test audio, video, and connection before the call.",
        "VIDEO SETUP: Business casual attire. No hats, t-shirts, tank tops, pajamas, or distracting motion in the background.",
        "VIDEO SETUP: Maintain eye contact through the camera lens, not the screen.",
        "VIDEO SETUP: Use gestures in the truth plane - waist to neck. Sit tall, lean slightly forward.",
        "VIDEO SETUP: Practice 5 minutes of speaking directly to camera daily before the interview.",
        "LISTENING: Nod, use facial expressions, brief verbal affirmations. Keep hands visible.",
        "Ownership audit: before every story answer, confirm every action uses 'I' not 'we'. Name the team context in one sentence, then switch to 'I' for every decision and action. 'We delivered' is invisible. 'I owned the go-live readiness and the team executed' is a hire.",
        "LISTENING: Paraphrase before responding when a question is complex - it buys time and shows you heard it.",
    ]


def recruiter_translation_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
) -> list[str]:
    proof_by_lane = {
        "presales_solution": "discovery, buyer translation, implementation realism, and customer-facing trust",
        "customer_success": "adoption, account-risk control, executive communication, and customer recovery",
        "analytics_operations": "reporting clarity, workflow diagnosis, validation discipline, and decision support",
        "change_enablement": "stakeholder translation, role-based adoption, and measurable follow-through",
    }
    proof_phrase = proof_by_lane.get(
        profile.primary_lane,
        "implementation discipline, customer-facing ownership, reporting clarity, and adoption follow-through",
    )
    return [
        "Recruiters are translators and screeners. Give them one sentence they can repeat upward instead of hoping they infer the bridge from job titles alone.",
        f"Problem-solution-proof headline: I help teams turn {candidate_problem_phrase(profile)} into usable outcomes through {proof_phrase}.",
        "Keep the proof recruiter-friendly: client count, dashboards, implementations, workshops, sites, users, or customer-recovery scope are easier to carry upward than tool lists alone.",
        f"If asked for a fast summary of fit for the {role_title} role at {company_name}, keep it to one clear headline plus one proof example.",
    ]


def phone_screen_first_round_playbook(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> list[str]:
    next_step_phrase = {
        "presales_solution": "discovery, buyer-facing solution work, and implementation-aware customer conversations",
        "customer_success": "adoption, customer health, and measurable value after go-live",
        "analytics_operations": "workflow visibility, reporting clarity, and decision support",
        "change_enablement": "adoption, stakeholder translation, and behavior change work",
    }.get(
        profile.primary_lane,
        f"{candidate_problem_phrase(profile)}, customer-facing ownership, and practical delivery work",
    )
    return [
        "30-minute pacing: 2-3 minutes rapport, 4-5 minutes Tell Me About Yourself, 12-15 minutes fit questions, 5 minutes for Christian's questions, and 2-3 minutes for the close and next steps.",
        "Target most first-round answers at about one minute. If the recruiter or hiring manager wants more detail, let them pull for more.",
        "Research beyond the website: know the industry, a few relevant competitors, and one likely pain point or solution hypothesis you can test in the conversation.",
        "Bring a compact career narrative that could fit on one page: what Christian has done, the proof behind it, and why this role is the right next move.",
        "How are you?: Doing well, thank you. I was looking forward to learning more about the role and what the team needs most.",
        f"What do you know about {company_name}?: I understand the team is hiring this {role_title} role to make {candidate_problem_phrase(profile)} more reliable, and that is a strong fit with the work I have already been doing.",
        f"Tell me about yourself: {pitch_for_profile(profile, job_description, company_name, role_title, resume_text, notes_text)}",
        f"What are you looking for next?: I am looking for more {next_step_phrase}, because that is the part of my background I want to keep building.",
        f"Why are you leaving?: I have learned a lot in my current lane, and I am now looking for a role with more direct exposure to {candidate_problem_phrase(profile)}. That is why this opportunity stands out.",
        f"Why this role?: The {role_title} role stands out because it needs someone who can understand the workflow, keep issues visible, and turn the work into adoption or decision quality instead of stopping at activity.",
        "Compensation early in the process: I would like to understand the approved range and total package first, then I can react to it with the right context.",
    ]


def neutral_exit_narrative_lines(
    profile: build_resume.JobProblemProfile,
    role_title: str,
) -> list[str]:
    return [
        "Keep the exit story short, positive, and forward-looking: what Christian learned, what he wants more of next, and why this role fits that move.",
        "Do not vent about a manager, the company, politics, or an internal frustration that the new employer cannot fix.",
        f"Safe answer pattern: I have learned a lot in my current work, and I am now looking for more direct ownership around {candidate_problem_phrase(profile)}. The {role_title} role is attractive because it gives me more of that operating pattern.",
        "If the move is tied to reorganization, contract timing, or limited growth, state it once without drama and return to fit.",
    ]


def interview_honesty_and_privacy_lines() -> list[str]:
    return [
        "Be exact on dates, titles, tools, education, certifications, time off, and the real depth of any platform or domain.",
        "If experience is adjacent rather than direct, lead with the supported work pattern first and name the exact edge only when the interviewer needs that level of detail.",
        "Be honest about other interviews, start-date timing, and upcoming commitments without turning the answer into a negotiation speech.",
        "Family plans, politics, religion, workplace drama, and current salary can stay private or be handled briefly without apology.",
        "If a question becomes too personal, redirect calmly to availability, work style, or role fit.",
    ]


def closing_interview_drill_lines(
    company_name: str,
    role_title: str,
    profile: build_resume.JobProblemProfile,
    job_description: str = "",
    supplied_context: str = "",
) -> list[str]:
    desire_focus = closing_desire_focus(profile, job_description, supplied_context)
    return [
        "Close in three moves: explicit interest, one specific role or project hook, and a clear next-step signal.",
        f"Sample close: I am very interested in the {role_title} role at {company_name}, especially because of {desire_focus}. Based on our conversation, it sounds like the team needs someone who can bring structure to {candidate_problem_phrase(profile)}, and that is where my background is strongest. I would welcome the next step.",
        "If the interviewer named a visible pressure point, repeat it once in the close so the final impression sounds specific rather than rehearsed.",
        f"High-value objection check: \"{biggest_gap_question()} I would rather understand it now so I can address it directly.\"",
        "Keep the spoken close under about 20 seconds, then move into timeline and follow-up questions.",
    ]


def supported_theme_story(
    stories: list[StoryCard],
    terms: tuple[str, ...],
    used_titles: set[str],
) -> StoryCard:
    if not stories:
        fail("supported_theme_story() requires at least one story")
    scored: list[tuple[int, StoryCard]] = []
    for story in stories:
        story_text = " ".join(
            (
                story.title,
                story.hook,
                story.evidence,
                story.result,
                story.outcome,
                " ".join(story.story_types),
                " ".join(story.signals),
            )
        ).lower()
        score = sum(story_text.count(term.lower()) for term in terms)
        if story.title in used_titles:
            score -= 3
        scored.append((score, story))
    selected = max(scored, key=lambda item: item[0])[1]
    used_titles.add(selected.title)
    return selected


def three_supported_proof_theme_lines(
    profile: build_resume.JobProblemProfile,
    stories: list[StoryCard],
) -> list[str]:
    theme_specs = {
        "presales_solution": (
            ("Discovery quality", ("discovery", "requirements", "buyer", "solution")),
            ("Implementation realism", ("implementation", "go-live", "configuration", "testing")),
            ("Executive trust", ("customer", "executive", "persuasion", "alignment")),
        ),
        "customer_success": (
            ("Customer trust under pressure", ("customer", "recovery", "risk", "trust")),
            ("Adoption discipline", ("adoption", "training", "workshop", "enablement")),
            ("Commercial awareness", ("renewal", "account", "growth", "value")),
        ),
        "analytics_operations": (
            ("Decision clarity", ("dashboard", "reporting", "decision", "visibility")),
            ("Validation discipline", ("sql", "validation", "data", "accuracy")),
            ("Workflow diagnosis", ("process", "workflow", "improve", "operations")),
        ),
        "change_enablement": (
            ("Stakeholder translation", ("stakeholder", "alignment", "communication", "persuasion")),
            ("Behavior change", ("adoption", "training", "enablement", "role-based")),
            ("Follow-through", ("result", "implementation", "ownership", "sustain")),
        ),
    }
    default_specs = (
        ("Implementation control", ("implementation", "go-live", "risk", "delivery")),
        ("Stakeholder alignment", ("stakeholder", "customer", "alignment", "workshop")),
        ("Adoption and measurable follow-through", ("adoption", "training", "reporting", "result")),
    )
    lines = [
        "If an interviewer asks for three words, give three supported proof themes instead of three adjectives. Each theme needs one concrete example behind it.",
    ]
    used_titles: set[str] = set()
    for label, terms in theme_specs.get(profile.primary_lane, default_specs):
        story = supported_theme_story(stories, terms, used_titles)
        lines.append(f"{label}: {story.title} - {story.result}")
    return lines


def likely_questions(profile: build_resume.JobProblemProfile, job_description: str) -> list[InterviewQuestion]:
    questions = [
        InterviewQuestion(
            question="Tell me about yourself.",
            angle="Use a results-based career summary, not chronology. Lead with the biggest wins that prove the role fit."
        ),
        InterviewQuestion(
            question="Why are you interested in this role?",
            angle="Connect the role's core problem to your strongest direct evidence, not to a generic interest in the company."
        ),
        InterviewQuestion(
            question="What is the most similar work you have done before?",
            angle="Choose one story that matches the role lane most directly and walk it through Hook, Evidence, Result, and Outcome."
        ),
        InterviewQuestion(
            question="What are the top wins from your background that matter here?",
            angle="Name three quantified wins, then explain the pattern they prove for this employer."
        ),
        InterviewQuestion(
            question="What would you want this interview team to remember about you?",
            angle="Give three debrief words tied to evidence, then connect them to the role's business problem."
        ),
        InterviewQuestion(
            question="How do you handle competing stakeholder priorities?",
            angle="Use the executive workshops story and show how you create alignment, surface tradeoffs, and keep momentum."
        ),
        InterviewQuestion(
            question="Describe a high-risk situation you had to stabilize.",
            angle="Use the $1M+ account recovery story and keep the answer focused on ownership, structure, and restored trust."
        ),
    ]
    if profile.primary_lane == "presales_solution":
        questions.extend(
            [
                InterviewQuestion(
                    question="How do you run discovery before a demo or technical recommendation?",
                    angle="Talk about translating requirements into a practical solution path, then tailoring the narrative to the buyer's real pain."
                ),
                InterviewQuestion(
                    question="How do you make complex product value credible to executive buyers?",
                    angle="Anchor to business outcomes, workflow impact, and implementation realism rather than feature volume."
                ),
            ]
        )
    elif profile.primary_lane == "customer_success":
        questions.extend(
            [
                InterviewQuestion(
                    question="How do you spot adoption risk before it becomes churn risk?",
                    angle="Discuss signal tracking, structured check-ins, executive alignment, and practical next steps."
                ),
                InterviewQuestion(
                    question="How do you manage a difficult customer escalation?",
                    angle="Show calm ownership, cross-functional coordination, and recovery of trust through follow-through."
                ),
            ]
        )
    elif profile.primary_lane == "analytics_operations":
        questions.extend(
            [
                InterviewQuestion(
                    question="How do you turn raw data into action for business leaders?",
                    angle="Use the dashboards story and explain how you move from reporting to decision support."
                ),
                InterviewQuestion(
                    question="How do you validate that a metric or report is useful?",
                    angle="Talk about business context first, then usability, accuracy, and what decision the output enables."
                ),
            ]
        )
    else:
        questions.extend(
            [
                InterviewQuestion(
                    question="How do you keep implementations on track when scope and risk start moving?",
                    angle="Talk through discovery, checkpoints, issue triage, and stakeholder communication."
                ),
                InterviewQuestion(
                    question="How do you define readiness for testing or go-live?",
                    angle="Frame readiness as alignment across configuration, data, users, defects, and stakeholder confidence."
                ),
            ]
        )

    if mentions(job_description, "ai", "automation", "chatbot", "nlp"):
        questions.append(
            InterviewQuestion(
                question="What is your actual hands-on automation or AI experience?",
                angle="Keep it honest: practical workflow automation, LivePerson messaging configuration, and Codex-assisted documentation are supported; enterprise AI strategy ownership is not."
            )
        )
    if mentions(job_description, "sql", "snowflake", "teradata", "data warehouse"):
        questions.append(
            InterviewQuestion(
                question="How deep is your reporting or SQL background?",
                angle="Lead with 200+ dashboards and reporting tools, then stay precise about supported platforms and avoid claiming unsupported warehouse depth."
            )
        )
    return questions[:8]


def likely_question_story(item: InterviewQuestion, stories: list[StoryCard], used_titles: set[str] | None = None) -> StoryCard:
    used_titles = used_titles or set()
    prompt = f"{item.question} {item.angle}".lower()
    lifecycle_terms = ("implementation", "go-live", "configuration", "readiness", "lifecycle", "rollout")
    enhancement_terms = ("upgrade", "customization", "service pack", "already live", "enhancement")
    hints: tuple[tuple[tuple[str, ...], tuple[str, ...]], ...] = (
        (("data", "metric", "report", "excel", "sql", "validate"), ("Analysis and Decision", "KPI", "dashboard", "reporting")),
        (("risk", "stabilize", "difficult", "escalation", "customer"), ("Challenge and Failure", "Persuasion", "account stabilization", "recovery")),
        (("train", "training", "adoption", "enablement"), ("Teamwork", "workshop", "QBR", "adoption")),
        (("discovery", "demo", "solution", "buyer"), ("Persuasion", "pre-sales", "discovery", "solution")),
        (("implementation", "go-live", "configuration", "readiness"), ("Individual Achievement", "implementation", "go-live", "Aptean")),
        (("technical", "project", "integration", "api", "migration"), ("Rapid Learning", "ERP", "migration", "technical")),
        (("stakeholder", "competing", "priority", "influence"), ("Persuasion", "Teamwork", "alignment")),
        (("gap", "learn", "new", "comfort"), ("Rapid Learning", "Challenge and Failure", "learning")),
    )
    scored: list[tuple[int, StoryCard]] = []
    for story in stories:
        score = signal_score(prompt, story.signals)
        story_text = " ".join((story.title, story.hook, " ".join(story.story_types), " ".join(story.signals))).lower()
        for question_terms, story_terms in hints:
            if any(term in prompt for term in question_terms):
                if any(term.lower() in story_text for term in story_terms):
                    score += 8
        if any(term in prompt for term in lifecycle_terms):
            if "aptean" in story_text:
                score += 10
            if "east west" in story_text:
                score -= 4
        if any(term in prompt for term in enhancement_terms):
            if "east west" in story_text:
                score += 10
        if story.title in used_titles:
            # Large enough to lose to any unused story regardless of keyword
            # bonuses above, so a story only repeats within the same question
            # list when every other story has truly been used already. A
            # small penalty here let the same story answer two different
            # questions verbatim in the same interview, which is exactly the
            # "sounds rehearsed" failure mode this tool exists to avoid.
            score -= 1000
        scored.append((score, story))
    return max(scored, key=lambda item_score: item_score[0])[1]


def question_story_angle(
    prompt: str,
    angle: str,
    stories: list[StoryCard],
    fallback: str,
) -> str:
    if not stories:
        return fallback
    story = likely_question_story(InterviewQuestion(prompt, angle), stories)
    lead = next((takeaway for takeaway in story.takeaways if takeaway), story.result)
    return f"Use {story.title}: {lead}"


def question_intent_framework(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    stories: list[StoryCard],
) -> list[QuestionIntentCard]:
    return [
        QuestionIntentCard(
            prompt="Why are you leaving / why now?",
            hidden_assessment="Positive transition logic, realism about the move, and whether Christian sounds likely to accept the job terms.",
            bad_answer_trap="Complaining about the current employer, naming a fixable internal issue, or sounding vague about what the next role should provide.",
            story_angle="Use the forward-looking fit story: what work Christian wants more of, why this role provides it, and one proof point that makes the pivot credible.",
            clarify_when="Ask a clarifying question only if they bundle this together with why the company or role is attractive.",
            ideal_length="15-20 seconds",
        ),
        QuestionIntentCard(
            prompt="Why this company?",
            hidden_assessment="Research depth, mutual fit, and whether Christian understands the operating problem behind the posting.",
            bad_answer_trap="Generic mission praise, homepage language, or naming benefits without a role-specific reason.",
            story_angle=question_story_angle(
                "Why this company?",
                "Use a story that proves research depth, operating relevance, and role-fit logic.",
                stories,
                f"Use one proof point that connects directly to {candidate_problem_phrase(profile)}.",
            ),
            clarify_when="If the role context is still vague, ask one question about the team's current priorities before giving a longer answer.",
            ideal_length="20-30 seconds",
        ),
        QuestionIntentCard(
            prompt="What value do you offer / why should we hire you?",
            hidden_assessment="Whether Christian can summarize differentiated value in plain English instead of listing attributes.",
            bad_answer_trap="A strengths list with no proof, a chronology recap, or language that sounds interchangeable with other candidates.",
            story_angle=question_story_angle(
                "What value do you offer?",
                "Use a result-first proof point that shows differentiated value.",
                stories,
                "Use one problem, one action pattern, and one measurable result.",
            ),
            clarify_when="If the job priorities are not yet clear, ask which parts of the role matter most so the answer stays relevant.",
            ideal_length="20-30 seconds",
        ),
        QuestionIntentCard(
            prompt="How will this role be better for you than your current one?",
            hidden_assessment="Whether the company can actually provide a better fit and whether Christian is moving toward a clear next chapter.",
            bad_answer_trap="Lifestyle-only reasons, abstract growth language, or implying the current employer is failing Christian.",
            story_angle="Frame the move around scope, customer or business problem, and proof that Christian already performs the adjacent pattern well.",
            clarify_when="Ask about the most important responsibilities first if the role scope still sounds broad.",
            ideal_length="20-30 seconds",
        ),
        QuestionIntentCard(
            prompt="What would you do in the first 90 days?",
            hidden_assessment="Operating rhythm, assumption discipline, and whether Christian can become useful quickly without overpromising.",
            bad_answer_trap="Grand transformation promises, generic enthusiasm, or jumping into tool specifics before defining success.",
            story_angle="Use a listen-map-own-improve pattern tied to the lane: workflow, stakeholders, risk checkpoints, proof, and a first visible win.",
            clarify_when="Ask one clarifying question if responsibilities, ownership, or customer scope are still unclear.",
            ideal_length="45-60 seconds",
        ),
        QuestionIntentCard(
            prompt="Where do you see yourself in five years?",
            hidden_assessment="Realistic ambition, growth runway, and whether the role can logically fit Christian's trajectory.",
            bad_answer_trap="Title-chasing, talking past the role, or signaling that the team is only a short stop on the way somewhere else.",
            story_angle="Tie the answer to capability growth: deeper ownership of the problems this role already solves rather than a vanity title ladder.",
            clarify_when="No clarification needed unless they explicitly ask about title path instead of capability path.",
            ideal_length="20-30 seconds",
        ),
        QuestionIntentCard(
            prompt="Tell me about a disagreement or influence challenge.",
            hidden_assessment="Listening, judgment under tension, and whether Christian can create progress without making the other person the villain.",
            bad_answer_trap="Casting the other side as unreasonable, skipping what Christian noticed, or describing influence as repeated insistence.",
            story_angle=question_story_angle(
                "Tell me about a disagreement or influence challenge.",
                "Use a persuasion or stakeholder-alignment story.",
                stories,
                "Use a story where Christian surfaced the constraint behind the disagreement and created a usable next step.",
            ),
            clarify_when="Ask whether they mean a peer, manager, customer, or cross-functional disagreement if the audience matters to the answer.",
            ideal_length="45-60 seconds",
        ),
        QuestionIntentCard(
            prompt="Tell me about a time you solved or implemented something complex.",
            hidden_assessment="Initiative, structure, and whether Christian can move from diagnosis to execution and measurable outcome.",
            bad_answer_trap="A task list with no architecture, no tradeoffs, and no result that matters outside the project plan.",
            story_angle=question_story_angle(
                "Tell me about a time you solved or implemented something complex.",
                "Use a complex implementation or process-improvement story.",
                stories,
                "Use a story with a visible beginning, constraint set, execution path, and business result.",
            ),
            clarify_when="No clarification needed unless they define 'complex' in a domain Christian wants to narrow, such as customer, technical, or process complexity.",
            ideal_length="60-90 seconds",
        ),
        QuestionIntentCard(
            prompt="Tell me about a time something went wrong or you failed.",
            hidden_assessment="Resilience, preparation, and whether Christian learns fast enough to lower repeat risk.",
            bad_answer_trap="Choosing a negligence story, minimizing the issue, or telling a story with no specific change afterward.",
            story_angle=question_story_angle(
                "Tell me about a time something went wrong or you failed.",
                "Use a challenge, failure, or recovery story.",
                stories,
                "Use a story where validation tightened after the miss and the lesson is concrete.",
            ),
            clarify_when="Ask whether they want a failure, a conflict, or a recovery example if the prompt stays broad.",
            ideal_length="45-60 seconds",
        ),
        QuestionIntentCard(
            prompt="How do you educate yourself or learn fast?",
            hidden_assessment="Self-directed learning, resourcefulness, and how Christian becomes useful in unfamiliar tools or domains.",
            bad_answer_trap="Saying 'I Google it' with no structure, or leaning entirely on others to do the learning for him.",
            story_angle=question_story_angle(
                "How do you educate yourself or learn fast?",
                "Use a rapid-learning story.",
                stories,
                "Use one example where Christian built a short learning plan and became useful quickly.",
            ),
            clarify_when="Ask whether they mean learning a tool, an industry, or a workflow if the answer would change materially.",
            ideal_length="20-30 seconds",
        ),
        QuestionIntentCard(
            prompt="How would your coworkers describe you?",
            hidden_assessment="Self-awareness and whether Christian knows the repeatable strengths other people actually experience.",
            bad_answer_trap="Listing weaknesses when not asked, using empty adjectives, or naming strengths with no supporting behavior.",
            story_angle="Use a strength that can be defended with repeated behavior: clarity, ownership, calm under pressure, or useful follow-through.",
            clarify_when="No clarification needed.",
            ideal_length="15-20 seconds",
        ),
        QuestionIntentCard(
            prompt="What motivates you?",
            hidden_assessment="Whether Christian's motivation aligns with the actual work, not just with compensation or vague purpose language.",
            bad_answer_trap="Lifestyle-only answers, generic passion language, or motivations that the role cannot realistically satisfy.",
            story_angle=f"Anchor motivation to the work pattern in {role_title}: clearer workflows, stronger adoption, better reporting, or more useful customer outcomes.",
            clarify_when="Ask whether they mean in the role or more generally only if the question is explicitly ambiguous.",
            ideal_length="15-20 seconds",
        ),
        QuestionIntentCard(
            prompt="Do you prefer working on a team or independently?",
            hidden_assessment="Balance between autonomy and collaboration, plus whether Christian can adapt to the team's working rhythm.",
            bad_answer_trap="An absolute answer that ignores the role's actual collaboration model.",
            story_angle=question_story_angle(
                "Do you prefer working on a team or independently?",
                "Use a teamwork story that still shows independent ownership.",
                stories,
                "Use a story where Christian owned his part while making the team stronger.",
            ),
            clarify_when="No clarification needed unless the team structure is still unknown.",
            ideal_length="20-30 seconds",
        ),
        QuestionIntentCard(
            prompt="Describe your ideal boss.",
            hidden_assessment="Management fit, autonomy needs, and whether Christian can describe a productive environment without sounding negative or needy.",
            bad_answer_trap="Listing what Christian hated about past bosses or describing a fantasy manager instead of a workable operating style.",
            story_angle=f"Describe the kind of manager who makes {candidate_problem_phrase(profile)} easier to solve: clear priorities, trust, coaching, decision access, and useful feedback loops.",
            clarify_when="No clarification needed.",
            ideal_length="20-30 seconds",
        ),
    ]


def question_intent_grid_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    stories: list[StoryCard],
) -> list[str]:
    lines: list[str] = []
    for card in question_intent_framework(profile, company_name, role_title, job_description, stories):
        lines.append(
            f"{card.prompt} Hidden test: {card.hidden_assessment} Trap: {card.bad_answer_trap} Best angle: {card.story_angle} Clarify if: {card.clarify_when} Ideal length: {card.ideal_length}."
        )
    return lines


def behavioral_answer_scripts(
    profile: build_resume.JobProblemProfile,
    stories: list[StoryCard],
    job_description: str = "",
    company_name: str = "",
    role_title: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> list[PreparedAnswer]:
    if not stories:
        return []
    achievement = story_for_type(stories, "Individual Achievement") or stories[0]
    leadership = story_for_type(stories, "Managing and Leading") or achievement
    persuasion = story_for_type(stories, "Persuasion") or achievement
    analysis = story_for_type(stories, "Analysis and Decision") or achievement
    failure = story_for_type(stories, "Challenge and Failure") or achievement
    teamwork = story_for_type(stories, "Teamwork") or leadership
    rapid = story_for_type(stories, "Rapid Learning") or achievement
    role_bridge = lane_interview_bridge(profile)
    proof_line = lane_interview_proof_line(profile)
    tell_me_answer = (
        sixty_second_pitch(profile, company_name, role_title, job_description, resume_text, notes_text)
        if company_name and role_title
        else join_answer_sentences(
            "I help teams turn complex system, customer, and adoption problems into practical outcomes",
            proof_line,
            f"That is the pattern I would bring to this role by continuing to {role_bridge}",
        )
    )
    return [
        PreparedAnswer(
            "Delivery reminder before every answer",
            "Use the word 'because' explicitly when answering why-questions - it triggers a concrete reason and doubles as verbal keyword alignment. Keep short answers to 15-20 seconds. Eliminate fillers: 'as well,' 'you know,' 'also.' Sentences ending on a lower pitch sound confident; ending higher sounds uncertain. Pause before answering - speed is not competence.",
        ),
        PreparedAnswer(
            "Story prep method: rehearsed anchors",
            "Build the answer out loud until the structure is automatic. Rehearse five anchors per story: hook, discovery, what I noticed, action, and result/bridge. The goal is a practiced answer that can flex naturally when the interviewer redirects or drills deeper.",
        ),
        PreparedAnswer(
            "Tell me about yourself",
            tell_me_answer,
        ),
        PreparedAnswer(
            "Why should we hire you?",
            join_answer_sentences(
                f"Hire me if the role needs someone who can {role_bridge}",
                story_natural_reference(achievement),
                "I have delivered the clearest results in roles where I own the problem end to end and can connect the analysis directly to an action",
            ),
        ),
        PreparedAnswer(
            "What challenge do you expect in this role?",
            join_answer_sentences(
                f"The biggest challenge will likely be turning {candidate_problem_phrase(profile)} into decisions, adoption, and measurable progress without losing stakeholder trust",
                "My first move would be to clarify the outcome, validate where the facts are thin, and make the next decision and owner explicit",
                "From there I would create a simple operating rhythm that keeps risk visible and proves an early win",
            ),
        ),
        PreparedAnswer(
            "Ambiguous problem",
            join_answer_sentences(
                "I break ambiguity into facts, assumptions, risks, and owners",
                story_natural_reference(analysis),
                "This approach keeps messy work moving and has delivered cleaner decisions and faster progress than trying to resolve the ambiguity before acting",
            ),
        ),
        PreparedAnswer(
            "Rapid learning",
            join_answer_sentences(
                "I learn fastest by tying product details to a real workflow and customer problem",
                story_natural_reference(rapid),
                "This gives me a repeatable way to get useful quickly and has led to clearer contributions earlier in every new engagement",
            ),
        ),
        PreparedAnswer(
            "Opposing stakeholder views",
            join_answer_sentences(
                "I slow the conversation down enough to understand the constraint behind each position",
                story_natural_reference(persuasion),
                "The key move was making one usable path forward visible to everyone involved",
            ),
        ),
        PreparedAnswer(
            "Difficult coworker or client",
            join_answer_sentences(
                "When a coworker or client gets difficult, I first figure out whether the real issue is the problem itself or the fact that nobody owns it yet",
                story_natural_reference(persuasion),
                "What changed was the working rhythm, which improved ownership clarity and reduced the friction that was keeping the work stuck",
            ),
        ),
        PreparedAnswer(
            "Failure",
            join_answer_sentences(
                "The lesson was to strengthen validation before risk reaches users",
                story_natural_reference(failure),
                "Since then I build validation earlier, which has reduced implementation risk in every project that followed",
            ),
        ),
        PreparedAnswer(
            "Different backgrounds",
            join_answer_sentences(
                "I work well across different backgrounds by listening for what each group is optimizing for",
                story_natural_reference(teamwork),
                "The goal is to make the next step usable for everyone involved, not just win the argument",
            ),
        ),
        PreparedAnswer(
            "Strengths",
            join_answer_sentences(
                "My strongest strengths for this kind of role are structured discovery, stakeholder alignment, and turning messy work into a usable next step",
                proof_line,
                f"I add value fastest when a team needs someone who can {role_bridge}",
            ),
        ),
        PreparedAnswer(
            "Weaknesses",
            join_answer_sentences(
                "One weakness I manage actively is that if the downstream impact is unclear, I can spend too long pressure-testing the risk before moving",
                "The improvement has been to put structure around that instinct: define the decision, run a short validation step, and move forward with checkpoints instead of waiting for perfect certainty",
                "The structure keeps the caution useful without letting it slow the work down",
            ),
        ),
        PreparedAnswer(
            "Influence without authority",
            join_answer_sentences(
                "I influence without authority by making the business tradeoff clear and giving each stakeholder a path to say yes",
                story_natural_reference(persuasion),
                "It has led to better decisions and enabled progress in situations where formal authority was not available",
            ),
        ),
        PreparedAnswer(
            "Process improvement: inefficient process",
            join_answer_sentences(
                "There was a high-volume inventory adjustment process that used to be mostly manual, and fixing it reduced manual work by 78% and discrepancies by 22%",
                "I started by mapping the current state, identifying the real bottleneck, and validating the fix with the people who used the workflow every day",
                (
                    "What I would bring forward from that example is the same discipline: measure before and after, validate with SMEs, and make the new standard easy to repeat through practical documentation"
                    if not build_resume.jd_mentions(job_description, "state farm", "claims process engineering", "claims process")
                    else "What I would bring forward to claims process engineering is the same discipline: measure before and after, validate with SMEs, and make the new standard easy to repeat through SOP-style documentation"
                ),
            ),
        ),
        PreparedAnswer(
            "Continuous improvement mindset",
            "Use the 78% / 22% process improvement story, then add the operating habit: capture lessons learned while the work is happening, apply feedback immediately, pilot before scaling, and use retrospectives or 30/60/90 checks so improvement does not stop at rollout.",
        ),
        PreparedAnswer(
            "Project management in process engineering",
            join_answer_sentences(
                "I handle project management like an integrator, which means I balance scope, time, cost, quality, risk, and stakeholder adoption instead of optimizing only one constraint",
                "I review the current plan, investigate before acting, and keep a live risk and issue log.",
                "It also makes tradeoffs visible before workload or scope becomes unsustainable.",
            ),
        ),
        PreparedAnswer(
            "Data-driven process decision",
            join_answer_sentences(
                "My approach is to start with the operating decision that needs to change, not the report itself",
                "The clearest example from my background is 200+ dashboards and decision visibility, where I clarified the business question, pulled the right data, and turned it into a recommendation leaders could act on",
                "The point is that the analysis changed the workflow or decision, not just the formatting of the report",
            ),
        ),
        PreparedAnswer(
            "Stakeholder resistance to process change",
            join_answer_sentences(
                "When people resist a process change, I start by understanding what they are protecting before I try to persuade them",
                "The best example is Operations versus finance alignment, where the move was to acknowledge the concern, ask clarifying questions, surface the constraint, and then propose a pilot with clear rationale",
                "It usually works better than pushing harder because it gives people a credible path to yes",
            ),
        ),
        PreparedAnswer(
            "What did you dislike about your last boss or job?",
            join_answer_sentences(
                "What I wanted more of was direct customer contact, because I do my best work when I can hear the workflow problem firsthand and turn it into a clearer next step",
                "It is one reason this role stands out to me",
            ),
        ),
        PreparedAnswer(
            "Why should we not hire you?",
            join_answer_sentences(
                "If you need someone with exact day-one depth in every specialty detail, that would be the learning curve in my background",
                "If you need someone who can learn the workflow quickly, make risk visible early, and keep ownership clear under pressure, that is where I am strongest",
            ),
        ),
        PreparedAnswer(
            "How you structure a messy problem",
            join_answer_sentences(
                "I start by defining the business outcome and the first decision that has to be made",
                "Then I separate facts from assumptions, map the risks, and validate the path with the people affected by it",
                "This approach keeps the answer practical instead of letting the problem stay abstract for too long",
            ),
        ),
    ]
    story_diversity_warning(
        question_labels=[answer.prompt for answer in answers],
        assigned_stories=[achievement, leadership, persuasion, analysis, failure, teamwork, rapid],
    )
    return answers


def story_diversity_warning(
    question_labels: list[str],
    assigned_stories: list[StoryCard],
) -> None:
    story_title_counts = Counter(card.title for card in assigned_stories if card.title)
    repeated_titles = [title for title, count in story_title_counts.items() if count >= 3]
    if repeated_titles:
        for title in repeated_titles:
            print(
                f"STORY DIVERSITY WARNING: '{title}' is assigned to "
                f"{story_title_counts[title]} behavioral questions. "
                f"Add more story types to the source or accept repeated examples."
            )


def validate_behavioral_answer_bank(answer_bank: Sequence[PreparedAnswer]) -> None:
    answers = {item.prompt: item.answer for item in answer_bank}
    for prompt, minimum in (
        ("Tell me about yourself", 30),
        ("Why should we hire you?", 30),
        ("Ambiguous problem", 28),
        ("Rapid learning", 28),
        ("Difficult coworker or client", 30),
        ("Failure", 28),
        ("Influence without authority", 28),
    ):
        answer = answers.get(prompt, "")
        if not answer:
            fail(f"behavioral answer bank is missing the '{prompt}' answer")
        assert_full_spoken_answer(prompt, answer, min_words=minimum)
        assert_claim_then_prove_answer(prompt, answer, min_words=minimum)


def answer_do_dont(job_description: str = "") -> list[str]:
    lines = [
        "Interview answers must show Level 3 traits. Do not merely claim a trait or name a trait. Show what Christian noticed, what he said or did, how he adjusted, and what result it produced.",
        "For operational answers, rehearse three layers: name the process area, state Christian's exact role, and finish with the business impact.",
        "Rehearse story anchors until they feel automatic: hook, discovery, what I noticed, action, result, bridge.",
        "Lead with tension or result before context. Context should explain the stakes, not warm up the story.",
        "Add one 'what I noticed' line to every major story. This is the difference between reporting actions and showing judgment.",
        "Use before/after contrast for impact: before the change, what was broken; after the change, what became faster, cleaner, safer, or more measurable.",
        "If you lose your place, recover with: 'Let me come back to that from a slightly different angle,' then state the result and offer to walk through how.",
        "Lead top-down: give the answer first, then give three takeaways, then tell the story.",
        "Tell me about yourself should be results-based, not chronological. Lead with wins and the role-fit pattern.",
        "Balance I and team: name the team context, then state exactly what Christian personally did to influence the result.",
        "Keep most answers under two minutes unless the interviewer asks for depth.",
        "Make interest explicit at least once. Say the role is interesting, name the specific reason, and connect it to the team's needs.",
        "Do start with the problem, then the action, then the business result.",
        "Do keep each answer anchored to one example instead of listing the whole resume.",
        "Do use supported scope markers like 80+ clients, 60+ workshops, 200+ dashboards, five sites, 150+ users, and $1M+ revenue risk when relevant.",
        "Do use the strongest truthful bridge available: lead with the supported operating pattern so the interviewer does not have to infer the fit.",
        "Use the word 'because' in all why-answers. It forces a concrete reason and naturally inserts keywords.",
        "End sentences with a downward tone - confident and declarative. Upward pitch sounds uncertain.",
        "15-20 seconds for short answers. Let the interviewer pull for more detail.",
        "Do not claim direct ownership of unsupported leadership, engineering, or specialty domains.",
        "Do not open with disclaimers when adjacent evidence can be framed through a clear, supported bridge statement.",
        "For continuous improvement answers, include a feedback loop: lesson captured, pilot or test, stakeholder validation, and what changed afterward.",
    ]
    if build_resume.jd_mentions(
        job_description,
        "process engineer",
        "lean six sigma",
        "spc",
        "pid",
        "manufacturing process",
        "aspen",
        "matlab",
        "autocad",
    ):
        lines.insert(
            -1,
            "For process engineering terms like SPC, PFDs, P&IDs, ISO, GMP, Aspen, HYSYS, MATLAB, or AutoCAD: show awareness only unless you have direct supported experience. Use them as smart questions, not claims.",
        )
    return lines


def recording_audit_lines(profile: build_resume.JobProblemProfile, role_title: str) -> list[str]:
    lines = [
        "Record one answer and review it three ways: full audio/video for structure, audio-only for filler and hedge words, and silent-watch for posture and drift.",
        "Watch the first ten seconds only. The opening should sound like Christian already knows the point, not like he is still looking for it.",
        "Trim any answer that looks or sounds finished before the last sentence ends. That is usually where rambling starts.",
        "If a result line lands weakly, slow down slightly on the proof instead of adding more setup.",
    ]
    if profile.primary_lane == "implementation_delivery":
        lines.append(f"For the {role_title} lane, recordings should prove calm control of scope, risk, handoffs, and adoption.")
    elif profile.primary_lane == "presales_solution":
        lines.append(f"For the {role_title} lane, recordings should sound buyer-aware and commercially useful, not like product narration.")
    elif profile.primary_lane == "customer_success":
        lines.append(f"For the {role_title} lane, recordings should make value, risk, and trust visible after the sale.")
    elif profile.primary_lane == "change_enablement":
        lines.append(f"For the {role_title} lane, recordings should prove behavior change and stakeholder clarity, not just communications activity.")
    elif profile.primary_lane == "analytics_operations":
        lines.append(f"For the {role_title} lane, recordings should connect data work to a decision, not just a dashboard.")
    else:
        lines.append(f"For the {role_title} lane, recordings should sound like Christian has already handled similar complexity, not like adjacent experience narrated from a distance.")
    return lines


def search_progress_question(job_description: str = "", context_text: str = "") -> str | None:
    combined = f"{job_description}\n{context_text}"
    later_stage_signals = (
        "final round",
        "panel interview",
        "references",
        "offer",
        "background check",
    )
    long_open_signals = (
        "open until filled",
        "backfill",
        "urgent hire",
        "hiring immediately",
        "reopened role",
        "extended search",
        "long search",
    )
    if mentions(context_text, *later_stage_signals) or mentions(combined, *long_open_signals):
        return "Situational, later-stage only: What seems to be missing in other candidates so far, and where would you most want the next person to feel stronger?"
    return None


def questions_to_ask(
    company_name: str,
    profile: build_resume.JobProblemProfile,
    job_description: str = "",
    context_text: str = "",
) -> list[str]:
    base_questions = {
        "presales_solution": [
            f"What separates a strong {profile.lane_label} from an average one at this stage of {company_name}'s growth?",
            "Where do enterprise buyers most often need help connecting product capabilities to business outcomes?",
            "How do Solutions Engineers partner with Account Executives and Implementation after a technical win?",
            "What does the evaluation process look like for this role, and what would a strong first 90 days signal?",
        ],
        "customer_success": [
            "What customer signals tell the team that adoption is healthy or at risk before a renewal conversation starts?",
            "How does the CSM team balance product expertise, renewal motion, and scalable customer education?",
            f"What would make you confident after 90 days that the person in this role is the right hire at {company_name}?",
            "Where does the handoff from implementation to customer success tend to create the most friction right now?",
        ],
        "implementation_delivery": [
            "What tends to make implementations successful or risky for this customer segment?",
            "How does the team define readiness for go-live and handoff?",
            f"What would success look like in the first 90 days for this role at {company_name}?",
            "Where do scope, data quality, or stakeholder alignment tend to create the most downstream problems?",
        ],
        "change_enablement": [
            "Where does resistance tend to show up most predictably in this kind of change program?",
            "How does the team measure adoption, and at what point does a metric become a leading indicator of risk?",
            f"What would make you confident after 90 days that the person in this role is changing behavior, not just running training?",
            "Which stakeholder groups tend to be hardest to bring along, and why?",
        ],
        "corporate_strategy": [
            f"What separates a strong {profile.lane_label} from an average one on this team?",
            "Where do case teams most often need better structure: problem framing, analysis quality, stakeholder alignment, or implementation follow-through?",
            f"What would make you confident after 90 days that the person in this role is adding value at {company_name}?",
            "How does the team balance sharp analysis with recommendations that clients can actually execute?",
        ],
        "analytics_operations": [
            "Which operating metrics do leaders currently trust, and which ones do they question?",
            "Where does manual work still sit inside what should be automated or dashboard-driven?",
            f"What decision does {company_name} most need faster visibility into right now?",
            "How does the team validate that a report is both accurate and actually used by the people it was built for?",
        ],
        "process_improvement": [
            "What process problem would you most want this person to improve in the first 12 months?",
            "How does the team currently decide which process improvements are worth prioritizing?",
            "Which operational metrics are trusted today, and where does the team still need better measurement?",
            "How mature is the process improvement culture here: bottom-up, top-down, or both?",
            "How does the team capture lessons learned and feedback loops during the project, not just after close?",
            "Where do SOPs, standard work, or quality-control checkpoints most need improvement today?",
        ],
    }
    questions = list(base_questions.get(profile.primary_lane, base_questions["implementation_delivery"]))
    questions.insert(0, biggest_gap_question())
    questions.extend([
        "What projects or customer problems most need help right now?",
        "What is the biggest business goal this role helps the team hit?",
        "What do your best people in this kind of role have in common?",
        "What shows someone is doing great in this role after six months?",
        "What is one problem you would love to see solved better?",
        "What recent team win best reflects how the group operates when the work is going well?",
    ])
    if profile.primary_lane == "process_improvement":
        questions.extend([
            "What does success look like after the data exercise and case-study stage for the person you ultimately hire?",
            "Where do process changes most often lose momentum today: data quality, stakeholder adoption, communication, or execution follow-through?",
            "How do you balance efficiency, quality, compliance, and customer experience when those goals pull against each other?",
        ])
    if build_resume.jd_mentions(job_description, "ai", "automation", "chatbot"):
        questions.append("How is the team currently thinking about AI-assisted workflows, and where is the biggest opportunity to reduce manual work?")
    elif build_resume.jd_mentions(job_description, "data migration", "etl", "cutover"):
        questions.append("What data quality or validation issues tend to surface most often during migration, and how does the team currently manage them?")
    elif build_resume.jd_mentions(job_description, "executive", "c-suite", "board"):
        questions.append("How often does this role interact directly with executive or C-suite stakeholders, and what format do those conversations usually take?")
    first_priority = next(iter(build_resume.jd_priority_phrases(job_description)), "")
    if first_priority:
        questions.append(f"The posting emphasizes '{first_priority}'. Where is that most difficult in practice today?")
    progress_question = search_progress_question(job_description, context_text)
    if progress_question:
        questions.append(progress_question)
    final_questions = questions[:6]
    if progress_question and progress_question not in final_questions:
        final_questions = questions[:5] + [progress_question]
    return final_questions


def company_profile_interview_callout(company_name: str, job_description: str, profile: build_resume.JobProblemProfile) -> list[str]:
    # stub path: never fires; see detect_company_profile docstring
    firm = build_resume.detect_company_profile(company_name, job_description)
    if not firm:
        if re.search(r"state\s*farm", company_name, re.I):
            return [
                "State Farm-specific note: Claims Process Engineering role detected.",
                "Interview emphasis: initiative, Lean Six Sigma-style process improvement, root-cause analysis, data exercise readiness, quality controls, communication materials, and stakeholder adoption.",
                "Proof to emphasize: 78% manual-work reduction, 22% discrepancy reduction, 200+ reporting tools, role-based training, cross-functional stakeholder alignment, and validation checkpoints.",
                "Avoid in this interview: sounding like a pure ERP/systems candidate. Translate every systems example into process, service, quality, customer experience, or operational efficiency.",
            ]
        return []
    display_name = str(firm["key"]).replace("_", " ").title()
    return [
        f"Firm-specific note: {display_name} profile detected.",
        f"Interview emphasis: {firm['interview_emphasis']}",
        f"Proof to emphasize: {firm['proof_emphasis']}",
        f"Avoid in this interview: {firm['avoid']}",
    ]


def closing_desire_focus(profile: build_resume.JobProblemProfile, job_description: str, supplied_context: str = "") -> str:
    supplied_lines = supplied_company_background_lines("", supplied_context)
    if supplied_lines:
        selected = re.sub(r"\s+", " ", supplied_lines[0]).strip()
        if len(selected) > 120:
            selected = selected[:117].rstrip(" ,.;:") + "..."
        return f"what I learned from the supplied company context: {selected}" if selected else candidate_problem_phrase(profile)

    first_priority = next(iter(build_resume.jd_priority_phrases(job_description)), "")
    if first_priority:
        return f"the chance to help with {first_priority}"

    specialty = build_resume.role_specialty_phrase(job_description, candidate_problem_phrase(profile))
    return f"the chance to work on {specialty}"


def closing_mechanics(
    company_name: str,
    role_title: str,
    profile: build_resume.JobProblemProfile,
    job_description: str = "",
    supplied_context: str = "",
) -> tuple[tuple[str, tuple[str, ...]], ...]:
    desire_focus = closing_desire_focus(profile, job_description, supplied_context)
    return (
        (
            "SECTION 1 - THE COMMITMENT LOCK",
            (
                "Run this before leaving every interview: \"Before I let you go, could you walk me through what the next steps look like and when I should expect to hear back?\"",
                "Follow-up: \"And if I have not heard back after the timeline you shared, is it okay to give you a polite nudge?\"",
                "Why this works: once the interviewer states the timeline out loud, they are more likely to follow through.",
            ),
        ),
        (
            "SECTION 2 - THE HESITANCY CLOSE",
            (
                "Ask this before the commitment lock: \"Based on our conversation today, is there anything about my background or what I've shared that gives you pause? I'd rather address it directly.\"",
                f"Alternate diagnostic close: \"{biggest_gap_question()} I would rather understand it now so I can address it directly.\"",
                "Why this works: it surfaces objections while Christian can still respond, and secure candidates are the only ones who ask it, which signals confidence.",
            ),
        ),
        (
            "SECTION 3 - THE DESIRE SIGNAL",
            (
                f"Say this before leaving: \"I want to be direct: I'm genuinely excited about the {role_title} role at {company_name}, and specifically about {desire_focus}. I look forward to the next steps.\"",
            ),
        ),
    )


def closing_fit_summary(profile: build_resume.JobProblemProfile) -> tuple[str, ...]:
    return (
        f"What I am hearing is that this role needs someone who can bring structure to {candidate_problem_phrase(profile)} and keep the work tied to business outcomes.",
        "The reason my background fits directly is that I have done that through enterprise system ownership, client implementation, executive alignment, reporting, training, and account recovery.",
        "Is that how you see me contributing in this role, or is there another challenge you would want me to lean into first?",
    )


def communication_audit_reference(
    job_description: str = "",
    interview_notes: str = "",
) -> tuple[tuple[str, str, str], ...]:
    combined = f"{job_description}\n{interview_notes}"
    video_context = mentions(
        combined,
        "video",
        "virtual",
        "zoom",
        "teams",
        "remote interview",
        "on-demand video",
        "hirevue",
    )
    on_demand_context = mentions(
        combined,
        "on-demand video",
        "hirevue",
        "recorded response",
        "asynchronous interview",
    )
    answer_length_description = (
        "On-demand video responses are typically cut off at 90 seconds. Target 45 to 60 seconds per response. "
        "Practice starting with the result, stating the method in one sentence, and bridging to the role before the time limit. "
        "There is no interviewer to pull you for more detail."
        if on_demand_context
        else "Behavioral questions should usually land in 60 to 90 seconds. \"Tell Me About Yourself\" should be 60 to 90 seconds. Simple questions should be 15 to 30 seconds."
    )
    eye_contact_description = (
        "Look directly at the camera lens, not at the interviewer's face on screen. Practice pointing at the camera dot before the interview so it becomes automatic. "
        "Pan your gaze across the screen between thoughts but return to the lens when making a key point."
        if video_context
        else "In video, look at the camera lens, not the screen. In panels, distribute eye contact."
    )
    eye_contact_fix = (
        "Fix: practice camera-lens delivery for key points, then glance back to the screen between thoughts."
        if video_context
        else "Fix: start with the questioner for 20 to 30 seconds, then pan the room."
    )
    return (
        (
            "Filler words",
            "Count \"um,\" \"like,\" \"you know,\" and \"sort of\" per minute. Target is under two per minute.",
            "Fix: replace every filler with a deliberate pause.",
        ),
        (
            "Upward inflection",
            "Ending declarative statements with a rising pitch makes them sound like questions.",
            "Fix: lower your pitch on key statements, especially when stating metrics.",
        ),
        (
            "Answer length",
            answer_length_description,
            "Fix: time your practice.",
        ),
        (
            "Vague language",
            "Every adjective such as large, significant, complex, or difficult should be replaced with a specific number, proper noun, or scale marker.",
            "Fix: swap adjectives for measurable scope, named systems, named audiences, or concrete stakes.",
        ),
        (
            "Ownership language",
            "\"We\" describes team context. \"I\" is required for every decision and action you personally drove.",
            "Fix: audit each story for missing \"I\" statements.",
        ),
        (
            "Vocal variety",
            "Slow down for key points. Speed up for background context.",
            "Fix: pause after major metrics for two seconds before continuing.",
        ),
        (
            "Eye contact",
            eye_contact_description,
            eye_contact_fix,
        ),
        (
            "Recovery",
            "If you stumble, do not over-apologize.",
            "Fix: say \"Let me come back to that from a slightly different angle,\" land the result, and move forward.",
        ),
        (
            "Closing execution",
            "Check whether you used the commitment lock, asked the hesitancy close, and left with confirmed next steps.",
            "Fix: rehearse the close out loud before the interview so it is available under pressure.",
        ),
        (
            "Memorability",
            "If the interviewer retells your interview to a colleague, what is the one thing they would say?",
            "Fix: name that moment deliberately before you go in.",
        ),
    )


def interview_scorecard_template(company_name: str, role_title: str) -> tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...]]:
    dimensions = (
        "PREPARATION: Research depth, JD analysis, story readiness. Score: ___ / 10",
        "OPENING IMPRESSION: First 60 seconds, energy, smile, halo effect. Score: ___ / 10",
        "TELL ME ABOUT YOURSELF: 3-part formula, 90 seconds, role-specific, compelling. Score: ___ / 10",
        "BEHAVIORAL ANSWERS: CAR or CART structure, action-heavy, quantified results, personally owned. Score: ___ / 10",
        "EMPLOYER-CENTRICITY: Answers tied to their problem, future-pacing used, not self-focused. Score: ___ / 10",
        "COMMUNICATION QUALITY: Clarity, conciseness, no rambling, filler word control. Score: ___ / 10",
        "NON-VERBAL PRESENCE: Eye contact, posture, voice, energy, gestures. Score: ___ / 10",
        "QUESTIONS ASKED: Quality, specificity, hesitancy close used, not generic. Score: ___ / 10",
        "CLOSING EXECUTION: Commitment lock used, next steps confirmed, desire signal stated. Score: ___ / 10",
        "OVERALL MEMORABILITY: Likability, authority, desire demonstrated, at least one memorable moment created. Score: ___ / 10",
        "TOTAL: ___ / 100",
    )
    interpretation = (
        "90 to 100: Offer-ready. Focus on negotiation.",
        "80 to 89: Strong candidate. One or two dimensions pulling you down. Identify and drill them.",
        "70 to 79: Competitive but forgettable. Review the two lowest scores before next round.",
        "60 to 69: Systemic gaps. Review behavioral and communication modules immediately.",
        "Below 60: Fundamental preparation issues. Do not advance to next application before addressing.",
    )
    prompts = (
        "Write down every question asked within one hour.",
        "Identify the two weakest answers. Reconstruct them using the appropriate formula in this guide.",
        "What is the one memorable moment you created? If you cannot name one, that is the priority to fix.",
        "Did you use the commitment lock and hesitancy close? If not, add them to the pre-interview checklist for next time.",
    )
    return dimensions, interpretation, prompts


def negotiation_preparation_section(
    company_name: str,
    role_title: str,
    job_description: str,
    lane_key: str = "",
) -> tuple[tuple[str, tuple[str, ...]], ...]:
    levers = [
        ("Base Salary", "Ask first. This is hardest to move but highest impact."),
        ("Signing Bonus", "Easier to get than base. Justify with unvested equity or bonus you are leaving behind."),
        ("Equity Grant", "For roles with equity, negotiate grant size, vesting schedule, and cliff."),
        ("Performance Bonus", "Clarify target percentage and range between threshold and maximum."),
        ("Early Performance Review", "Request a 6-month review with compensation conversation built in."),
        ("Title", "Sometimes easier to move than compensation and helps in future negotiations."),
        ("Start Date", "Use to vest additional current equity or to give yourself preparation time."),
        ("Remote or Hybrid Flexibility", "Negotiate after offer, not before. Frame as efficiency, not preference."),
        ("Professional Development Budget", "Annual amount for courses, conferences, and coaching."),
        ("Additional PTO", "Especially if you have tenure-based PTO at your current employer."),
    ]
    if lane_key in {"presales_solution", "customer_success"}:
        priority = {"Base Salary": 0, "Performance Bonus": 1, "Equity Grant": 2}
        levers = sorted(levers, key=lambda item: priority.get(item[0], 10 + levers.index(item)))
    elif lane_key in {"implementation_delivery", "analytics_operations"}:
        priority = {"Base Salary": 0, "Signing Bonus": 1}
        levers = sorted(levers, key=lambda item: priority.get(item[0], 10 + levers.index(item)))
    lever_lines = tuple(f"{index}. {name}: {description}" for index, (name, description) in enumerate(levers, start=1))
    return (
        (
            "Salary Deflection Sequence",
            (
                "Round 1 (first ask): Deflect. Script: \"I'm really focused on finding the right fit first. I'm confident that if we get to that stage, we can find something that works. What is the range budgeted for this role?\"",
                "Round 2 (if pushed): Give a researched range. Script: \"Based on my research for this level in this market, I would expect something in the range I prepared ahead of time. But I'm genuinely more interested in the full picture.\"",
                "Round 3 (if they need a number): Give one number and frame it. Script: \"I would share the target number I prepared in advance. That said, I am not going to let a number get in the way of the right opportunity.\"",
            ),
        ),
        (
            "The 10 Negotiation Levers",
            lever_lines,
        ),
        (
            "The Counter Script",
            (
                "\"I am genuinely excited. This is the opportunity I have been most interested in throughout this process. I want to be transparent: based on my research and the strongest differentiator in my candidacy, I was expecting something closer to the target number I prepared in advance. Is there flexibility to get there?\"",
            ),
        ),
        (
            "The Final Offer Close",
            (
                "\"I understand. Before I give you my final answer: is there truly no flexibility anywhere in the package? I want to say yes today if we can find something that works.\"",
                "Never accept the first offer immediately. Express enthusiasm, ask for two to three days, research the market, and schedule a separate negotiation call. The offer does not disappear because you asked politely.",
            ),
        ),
    )


def thank_you_email_template_lines(company_name: str, role_title: str) -> list[str]:
    return [
        f"Subject: Thank you - {role_title} conversation",
        f"Hi interviewer name, thank you for taking time to discuss the {role_title} role at {company_name}.",
        "A moment that stood out was the discussion about the team's current priority; connect that moment to one resume-supported proof point from the conversation.",
        "The strongest proof to reinforce is one quantified story, such as the 78% manual-work reduction, 22% discrepancy reduction, or $1M+ at-risk account recovery.",
        f"Close by saying the conversation increased Christian's interest in {company_name} and that he would welcome the next discussion.",
    ]


def _latest_round_review(round_records: Sequence[Mapping[str, object]]) -> dict[str, object]:
    if not round_records:
        return {}
    review = round_records[0].get("performance_review", {})
    return review if isinstance(review, dict) else {}


def _latest_coaching_signals(round_records: Sequence[Mapping[str, object]]) -> list[dict[str, str]]:
    review = _latest_round_review(round_records)
    signals = review.get("coaching_signals", []) if isinstance(review, dict) else []
    results: list[dict[str, str]] = []
    for item in signals if isinstance(signals, Sequence) else []:
        if isinstance(item, dict):
            results.append(
                {
                    "key": str(item.get("key", "")).strip(),
                    "label": str(item.get("label", "")).strip(),
                    "detail": str(item.get("detail", "")).strip(),
                    "evidence": str(item.get("evidence", "")).strip(),
                }
            )
    return results


def _selected_coaching_records(
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> tuple[Mapping[str, object], ...]:
    if _latest_coaching_signals(round_records):
        return tuple(round_records)
    return tuple(global_round_records) if _latest_coaching_signals(global_round_records) else tuple(round_records)


def coaching_signal_keys(
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> set[str]:
    keys: set[str] = set()
    for item in _latest_coaching_signals(_selected_coaching_records(round_records, global_round_records)):
        key = str(item.get("key", "")).strip()
        if key:
            keys.add(key)
    return keys


def rehearsal_foundation_lines(
    role_title: str,
    round_records: Sequence[Mapping[str, object]] = (),
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    keys = coaching_signal_keys(round_records, global_round_records)
    lines = [
        f"Build a rehearsed foundation for the {role_title} interview: repeat the opening, three core stories, and the close out loud until the facts and order feel automatic.",
        "The goal is not to sound spontaneous from scratch. The goal is to sound natural because the structure has already been practiced enough times to survive redirects and follow-ups.",
        "Rehearse the transitions, not just the middle of the answer: answer first, one proof line, one business-value close.",
    ]
    if "rambling" in keys or "filler_restarts" in keys:
        lines.append("Practice the short version first. If the short version lands cleanly, then rehearse the longer follow-up version instead of starting long.")
    if "wrong_example_first" in keys:
        lines.append("Rehearse the first ten seconds of each answer with the best-fit example already chosen so the answer does not warm up with the wrong story.")
    return lines[:4]


def top_answer_risk_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    selected_records = _selected_coaching_records(round_records, global_round_records)
    review = _latest_round_review(selected_records)
    signal_map = {item["key"]: item for item in _latest_coaching_signals(selected_records) if item.get("key")}
    lines: list[str] = []
    summary = str(review.get("summary", "")).strip()
    if summary:
        lines.append(f"Latest interview read: {summary}")
    if "delayed_answer" in signal_map:
        lines.append("Lead with the answer in sentence one. If the interviewer still does not know the point by sentence two, restart shorter.")
    if "rambling" in signal_map:
        lines.append("Length control: cut the answer by roughly one third, land one proof point, and stop before the explanation loops back.")
    if "wrong_example_first" in signal_map:
        lines.append(f"Example control for {role_title}: choose the closest direct example first instead of warming up with adjacent background.")
    if "filler_restarts" in signal_map:
        lines.append("Filler control: replace setup phrases and repeated restarts with a short pause, then a direct ownership sentence.")
    if "consultative_questions" in signal_map:
        lines.append(f"Question strategy at {company_name}: ask about success, ramp, risk, and client trust before tactical tool questions.")
    if "executive_presence" in signal_map:
        lines.append("Executive presence: use shorter declarative sentences, one metric or operating detail, and a business-value closer.")
    next_round_risks = review.get("next_round_risks", []) if isinstance(review, dict) else []
    for risk in interview_context._split_lines(next_round_risks)[:3]:
        lines.append(f"Drill next: {risk}")
    if not lines:
        lines.append(f"Default risk read for {role_title}: answer the question, give one example, state the result, and stop when the value is clear.")
    return lines[:6]


def best_example_to_use_first_lines(
    profile: build_resume.JobProblemProfile,
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    selected_records = _selected_coaching_records(round_records, global_round_records)
    signal_keys = {item.get("key") for item in _latest_coaching_signals(selected_records)}
    review = _latest_round_review(selected_records)
    risk_text = " ".join(interview_context._split_lines(review.get("next_round_risks", []))).lower()
    if profile.primary_lane == "implementation_delivery" and ("wrong_example_first" in signal_keys or "aptian" in risk_text or "east west" in risk_text):
        return [
            f"Lifecycle questions: lead with {COMPANY_APTEAN} or other vendor-side implementation work first when the interviewer wants discovery through go-live.",
            f"Reserve {COMPANY_EAST_WEST} for customizations, upgrades, and already-live environment change work so the example fit stays crisp.",
        ]
    return [
        "Choose the closest direct proof story first. If the example needs a long preface to sound relevant, pick a different one.",
        "Use adjacent background only after the interviewer already trusts that you have answered the main question.",
    ]


def latest_positioning_diagnosis_lines(
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    selected_records = _selected_coaching_records(round_records, global_round_records)
    analysis = interview_context.latest_review_analysis(selected_records)
    diagnosis = analysis.get("positioning_diagnosis", {}) if isinstance(analysis, Mapping) else {}
    if not isinstance(diagnosis, Mapping):
        return []
    lines: list[str] = []
    headline = str(diagnosis.get("headline", "")).strip()
    if headline:
        lines.append(headline)
    for reason in interview_context._split_lines(diagnosis.get("reasons", []))[:3]:
        lines.append(f"Watch for: {reason}")
    hard_truth = str(diagnosis.get("hard_truth", "")).strip()
    if hard_truth:
        lines.append("Hard truth: " + hard_truth)
    return lines[:5]


def ownership_language_rewrite_lines(
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    selected_records = _selected_coaching_records(round_records, global_round_records)
    analysis = interview_context.latest_review_analysis(selected_records)
    rewrites = analysis.get("language_rewrites", {}) if isinstance(analysis, Mapping) else {}
    if not isinstance(rewrites, Mapping):
        return []
    avoid = interview_context._split_lines(rewrites.get("avoid", []))[:3]
    prefer = interview_context._split_lines(rewrites.get("prefer", []))[:4]
    lines: list[str] = []
    if avoid:
        lines.append("Avoid soft openers like: " + "; ".join(avoid))
    if prefer:
        lines.append("Prefer ownership-forward phrasing like: " + "; ".join(prefer))
    return lines[:3]


def answer_mode_lines() -> list[str]:
    return [
        "15 seconds: Point -> one reason -> stop. Use this for yes/no, fit, and quick-clarification questions.",
        "30 seconds: Point -> one proof line -> why it matters here. This is the default recruiter-screen mode.",
        "60 to 90 seconds: Problem -> your role -> what you changed -> business result -> bridge back to the role.",
    ]


def executive_presence_correction_lines(
    company_name: str,
    role_title: str,
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    signal_keys = coaching_signal_keys(round_records, global_round_records)
    lines = [
        f"For the {role_title} conversation at {company_name}, executive presence means naming the point before the background.",
        "Use one concrete stake, one ownership line, and one business consequence instead of narrating the whole path that got there.",
        "If the answer sounds finished before the mouth stops moving, close earlier and let the result line breathe.",
    ]
    if "executive_presence" in signal_keys or "rambling" in signal_keys:
        lines.append("Slow down on the proof line, lower the pitch on declarative statements, and end on the decision, risk, or customer value.")
    return lines[:4]


def pitch_delivery_note_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    round_records: Sequence[Mapping[str, object]] = (),
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    signal_keys = coaching_signal_keys(round_records, global_round_records)
    lines = [
        f"Open the {role_title} pitch with the 15-second claim, then one proof line, then the role bridge back to {company_name}.",
        "Use one idea per sentence. If the point already landed, cut the answer by roughly one third instead of adding more setup.",
        "Pick the best example immediately. Do not warm up with adjacent background if a direct story is available.",
        "Lower the pitch on the proof line and end on the decision, delivery risk, or customer value.",
    ]
    if "rambling" in signal_keys:
        lines.insert(2, "Rehearse the short version until it sounds complete on its own. The longer version should feel like an extension, not the default.")
    if "wrong_example_first" in signal_keys:
        lines.append(f"For lifecycle questions, default to {COMPANY_APTEAN} first and save {COMPANY_EAST_WEST} for already-live environment change work.")
    return lines[:5]


def pre_interview_routine_lines(
    role_title: str,
    round_records: Sequence[Mapping[str, object]] = (),
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    signal_keys = coaching_signal_keys(round_records, global_round_records)
    lines = [
        "Two minutes before the interview, get both feet flat on the floor and drop the shoulders before the first greeting.",
        "Open with a calm smile, open posture, and relaxed jaw so the first answer starts from steadiness instead of tension.",
        "Take three slow breaths with longer exhales than inhales so the voice settles before the first answer.",
        "Use one grounding sentence: 'My job is to answer clearly, not perform perfectly.'",
        "On hard questions, let one short pause happen before answering. A composed pause reads as judgment, not uncertainty.",
        "Reframe the adrenaline on purpose: this is your body preparing to focus, not proof that the answer is going badly.",
        f"First-answer routine for the {role_title} conversation: answer first, choose one example, state the result, and stop when the value is clear.",
    ]
    if "filler_restarts" in signal_keys:
        lines.append("If you feel the urge to restart, pause once instead. A clean pause sounds stronger than a verbal reset.")
    return lines[:6]


def story_anchor_fact_line(card: StoryCard, profile: build_resume.JobProblemProfile) -> str:
    anchors: list[str] = [card.title]
    if card.evidence_terms:
        anchors.append(build_resume.comma_series(card.evidence_terms[:3]))
    if card.result:
        anchors.append(card.result.rstrip("."))
    anchors.append(story_specific_bridge(card, profile).replace("Bridge: ", "").rstrip("."))
    return "Anchor facts to keep fixed: " + "; ".join(item for item in anchors if item)


def story_adaptation_drill_lines(
    card: StoryCard,
    profile: build_resume.JobProblemProfile,
    company_name: str = "",
    role_title: str = "",
    job_description: str = "",
) -> list[str]:
    role_fit = role_title.lower() if role_title else "the role"
    return [
        f"If they ask what made {card.title} hard, answer with the constraint first: {card.hook}",
        f"If they ask what you personally owned, answer with your role first: {card.evidence}",
        f"If they ask why this matters for {role_fit}, bridge with: {story_specific_bridge(card, profile).replace('Bridge: ', '')}",
    ]


def story_delivery_note_lines(
    card: StoryCard,
    profile: build_resume.JobProblemProfile,
    round_records: Sequence[Mapping[str, object]] = (),
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    signal_keys = coaching_signal_keys(round_records, global_round_records)
    lines = [
        "Answer the question first, then move into this example.",
        f"Land the result early: {card.result}",
        "Stop after the result and role bridge unless the interviewer asks for another layer.",
    ]
    if "executive_presence" in signal_keys or "rambling" in signal_keys:
        lines.append("Keep the answer declarative: one ownership line, one proof line, one business-value close.")
    if "wrong_example_first" in signal_keys:
        lines.append("Use this story only when it is the closest fit. Do not preface it with a weaker example.")
    return lines[:5]


def consultative_question_drill_lines(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
) -> list[str]:
    if profile.primary_lane == "implementation_delivery":
        return [
            biggest_gap_question(),
            f"What separates strong {role_title.lower()} hires from average ones in the first 6 to 12 months at {company_name}?",
            "Where do new hires usually struggle most after onboarding, and how do the best ones recover quickly?",
            "How much of the role is client-facing coordination versus internal execution during the first few months?",
            "What would make someone feel obviously successful by the 90-day mark?",
        ]
    if profile.primary_lane == "presales_solution":
        return [
            biggest_gap_question(),
            "What makes discovery especially credible with your buyers here?",
            "Where do deals usually stall when the pre-sale and delivery handoff is not tight enough?",
            "What separates strong solution consultants here from product narrators?",
        ]
    if profile.primary_lane == "customer_success":
        return [
            biggest_gap_question(),
            "Where does the team most need stronger customer judgment right now: adoption depth, renewal risk, or executive alignment?",
            "What makes a customer conversation especially high stakes in this role?",
            "What would make someone clearly valuable in the first 90 days?",
        ]
    return [
        biggest_gap_question(),
        f"What tends to separate strong performers in the {role_title} role from average ones in the first few months?",
        "Where does the team most need clearer ownership, judgment, or communication right now?",
        "What would make someone feel obviously useful by the 90-day mark?",
    ]


def phone_screen_script_answers(
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    stories: list[StoryCard],
    selected_stories: list[StoryCard],
    pitch_versions: Mapping[str, str],
    supplied_context: str = "",
    resume_text: str = "",
    notes_text: str = "",
) -> list[PreparedAnswer]:
    answer_map = {
        item.prompt: item.answer
        for item in common_interview_answers(
            profile,
            company_name,
            role_title,
            job_description,
            stories,
            panel_context=False,
            supplied_context=supplied_context,
            resume_text=resume_text,
            notes_text=notes_text,
        )
    }
    closest_story = (
        story_for_type(selected_stories, "Individual Achievement")
        or story_for_type(selected_stories, "Analysis and Decision")
        or (selected_stories[0] if selected_stories else None)
    )
    role_focus = build_resume.role_specialty_phrase(job_description, candidate_problem_phrase(profile))
    next_step_phrase = {
        "presales_solution": "discovery, buyer-facing solution work, and implementation-aware customer conversations",
        "customer_success": "adoption, customer health, and measurable value after go-live",
        "analytics_operations": "workflow visibility, reporting clarity, and decision support",
        "change_enablement": "adoption, stakeholder translation, and behavior change work",
        "process_improvement": "process diagnosis, measurable workflow improvement, and stronger operating controls",
    }.get(
        profile.primary_lane,
        f"direct ownership around {candidate_problem_phrase(profile)}",
    )
    similar_work_answer = (
        spoken_story_answer(
            closest_story,
            profile,
            company_name,
            role_title,
            job_description,
        )
        if closest_story
        else answer_map.get("Why should we hire you?", "")
    )
    why_company_answer = interview_join(
        f"I was interested in {company_name} because the posting points to {role_focus}",
        "The work reads like project-based ERP delivery where configuration choices, client trust, and follow-through all matter",
        "This is the kind of implementation work where my background is strongest",
    )
    why_hire_answer = interview_join(
        "You should hire me if this role needs someone who can learn the workflow quickly, make risk visible early, and turn the work into a usable next step",
        story_natural_reference(closest_story) if closest_story else "",
        "I bring structured execution without losing sight of the business outcome",
    )
    answers = [
        PreparedAnswer("Tell me about yourself", pitch_versions["30_second"]),
        PreparedAnswer(f"Why {company_name}?", why_company_answer),
        PreparedAnswer(f"Why this {role_title} role?", answer_map.get(f"Why this {role_title} role?", "")),
        PreparedAnswer("Why should we hire you?", why_hire_answer),
        PreparedAnswer("What is the most similar work you have done before?", similar_work_answer),
        PreparedAnswer(
            "What are you looking for next?",
            f"I am looking for more {next_step_phrase}, because that is the part of my background I want to keep building.",
        ),
        PreparedAnswer(
            "Why are you leaving?",
            "I have learned a lot in my current lane, and now I want a role with more direct ownership of the operating problem itself. This opportunity stands out because it sits closer to that work every day.",
        ),
        PreparedAnswer("What is a gap we should know about?", answer_map.get("What is a gap we should know about?", "")),
    ]
    for prompt in (
        "Tell me about your education and credentials.",
        "Walk me through your career.",
        "How would you learn a new product quickly?",
        "What are your compensation expectations and practical constraints?",
    ):
        if answer_map.get(prompt):
            answers.append(PreparedAnswer(prompt, answer_map[prompt]))
    conditional_answers = keyword_ready_answers(
        profile,
        job_description,
        stories,
        resume_text=resume_text,
        supplied_context=supplied_context,
    )
    answers.extend(
        item
        for item in conditional_answers
        if item.prompt.startswith(("Manufacturing Experience:", "Technical Depth Boundary:"))
    )
    return answers


def story_scale_anchor(card: StoryCard) -> tuple[str, bool]:
    """Returns (anchor_text, is_numeric). When no real number or scale figure exists in the
    story's evidence, falls back to a short qualitative anchor distinct from the Result field
    instead of silently repeating it."""
    candidates = [*card.evidence_terms, card.result, card.evidence]
    for candidate in candidates:
        cleaned = re.sub(r"\s+", " ", candidate).strip().rstrip(".")
        if not cleaned:
            continue
        if re.search(r"\d|\$|one million|two million|three million|four million|five|six|seven|eight|nine|ten", cleaned, re.I):
            return (cleaned if len(cleaned) <= 110 else cleaned[:107].rstrip(" ,.;:") + "...", True)
    fallback = card.takeaways[0] if card.takeaways else card.title
    return (re.sub(r"\s+", " ", fallback).strip().rstrip("."), False)


def story_scale_anchor_label(card: StoryCard) -> str:
    """Renders the anchor with an honest label, never claiming a qualitative anchor is a number."""
    anchor_text, is_numeric = story_scale_anchor(card)
    label = "Number or scale anchor" if is_numeric else "Scale anchor (qualitative)"
    return f"{label}: {anchor_text}"


def four_step_close_lines(
    company_name: str,
    role_title: str,
    profile: build_resume.JobProblemProfile,
    job_description: str = "",
    supplied_context: str = "",
) -> list[str]:
    desire_focus = closing_desire_focus(profile, job_description, supplied_context)
    return [
        f"Step 1 - Fit summary: What I am hearing is that this role needs someone who can bring structure to {candidate_problem_phrase(profile)} and keep the work tied to business outcomes.",
        "Step 2 - Match statement: My background fits because I have done that through systems ownership, client delivery, executive communication, reporting, training, and follow-through.",
        f"Step 3 - Interest signal: I am genuinely excited about the {role_title} role at {company_name}, especially because of {desire_focus}.",
        f"Step 4 - Next steps: \"Before we wrap, could you walk me through next steps and timing? {biggest_gap_question()} I would rather address it directly.\"",
    ]


def red_flag_lines(
    profile: build_resume.JobProblemProfile,
    role_title: str,
    job_description: str,
    round_records: Sequence[Mapping[str, object]],
    global_round_records: Sequence[Mapping[str, object]] = (),
) -> list[str]:
    lines = [
        "Do not memorize full paragraphs. Memorize the opening line, the proof number, and the result so the answer can still sound spoken.",
        "Do not start with disclaimers when a supported bridge is available. Lead with the closest proof first.",
        "Do not answer with two examples unless the interviewer asks. One clean story beats two partial stories.",
        "Do not restart out loud. Pause once, state the point directly, and keep moving.",
        "Do not let a good result get buried under setup. For this role, the first sentence should already sound useful.",
    ]
    signal_keys = coaching_signal_keys(round_records, global_round_records)
    if "wrong_example_first" in signal_keys:
        lines.insert(2, f"Do not warm up with adjacent background. For the {role_title} interview, pick the closest-fit example first.")
    if "rambling" in signal_keys:
        lines.insert(3, "Do not keep talking after the result lands. End the answer and let the interviewer pull for more.")
    if profile.primary_lane == "process_improvement":
        lines.append("Do not drift back into generic systems language. Keep the answer tied to process, control, service, or measurable workflow improvement.")
    return lines[:6]


def add_recent_interview_question_prep_section(
    document: Document,
    job_description: str,
    company_name: str,
    role_title: str,
    *,
    jobs_dir: Path = PROJECT_ROOT / "jobs",
    profile: build_resume.JobProblemProfile | None = None,
    stories: Sequence[StoryCard] | None = None,
    resume_text: str = "",
) -> bool:
    items = question_prep.recent_interviewer_question_prep_items(
        job_description,
        company_name,
        role_title,
        jobs_dir=jobs_dir,
    )
    if not items:
        return False
    if not resume_text:
        _, _, resume_text = question_prep.selected_resume_snapshot(job_description)
    if profile is None:
        profile = adjusted_profile_for_role(
            build_resume.job_problem_profile(job_description, resume_text), role_title, job_description
        )
    if stories is None:
        stories = supported_story_bank(resume_text)
    add_section(document, "Recent Interview Questions To Be Ready For")
    used_titles: set[str] = set()
    for item in items[:5]:
        if item.category in question_prep.NON_STORY_INTERVIEWER_CATEGORIES:
            spoken = prose_engine.spoken_register(
                question_prep.interviewer_question_factual_script(item.prompt, job_description, resume_text)
            ).text
        else:
            story = likely_question_story(InterviewQuestion(item.prompt, item.answer_angle), list(stories), used_titles)
            used_titles.add(story.title)
            spoken = spoken_story_answer(story, profile, company_name, role_title, job_description)
        add_bullet(document, f"{item.prompt} Say this: {spoken}")
    return True


def strip_external_question_prompts_for_validation(
    text: str,
    company_name: str = "",
    role_title: str = "",
    *,
    jobs_dir: Path = PROJECT_ROOT / "jobs",
) -> str:
    prompts = list(question_prep.load_active_application_prompts())
    if company_name and role_title:
        prompts.extend(interview_context.recent_interviewer_questions(jobs_dir, company_name, role_title, limit=10))
    cleaned = text
    for prompt in prompts:
        prompt_text = question_prep.normalize_spaces(prompt)
        if not prompt_text:
            continue
        cleaned = re.sub(re.escape(prompt_text) + r"\s*:\s*", "", cleaned)
        cleaned = re.sub(re.escape(prompt_text), "", cleaned)
    return cleaned


def build_document(company_name: str, role_title: str, job_description: str, resume_docx: Path, output_docx: Path) -> None:
    resume_text = "\n".join(paragraph_texts(resume_docx))
    jobs_dir = PROJECT_ROOT / "jobs"
    context_bundle = interview_context.load_company_context(
        jobs_dir,
        company_name,
        role_title,
        company_research_path=COMPANY_RESEARCH,
        global_interview_notes_path=INTERVIEW_NOTES,
        mode="compact",
    )
    global_round_records = interview_context.global_coaching_fallback_records(jobs_dir, company_name, role_title)
    interview_notes = context_bundle.interview_notes
    supplied_context = context_bundle.supplied_context
    profile = adjusted_profile_for_role(
        build_resume.job_problem_profile(job_description, resume_text),
        role_title,
        job_description,
    )
    selected_stories = hero_stories(profile, job_description, resume_text)
    all_stories = supported_story_bank(resume_text)
    debrief_summary = (
        debrief_analysis.analyze_entries(context_bundle.round_records, company_name)
        if context_bundle.round_records
        else debrief_analysis.DebriefPatternSummary(
            company_name=company_name,
            entry_count=0,
            most_common_question="",
            recurring_question_count=0,
            top_story_title="",
            top_story_count=0,
            repeated_role_language=(),
            top_coaching_signals=(),
        )
    )
    selected_stories = debrief_analysis.reorder_story_cards(selected_stories, debrief_summary)
    pitch_versions = pitch_variants(
        profile,
        company_name=company_name,
        role_title=role_title,
        job_description=job_description,
        resume_text=resume_text,
        notes_text=interview_notes,
    )
    for label, pitch_text in pitch_versions.items():
        enforce_prose_quality(
            pitch_text,
            "interview_pitch",
            label=f"Interview pitch ({label})",
            mode="warn",
            check_template_leakage=False,
        )
    if len(all_stories) < 6:
        fail("not enough resume-supported stories available for a complete interview guide")
    if len(selected_stories) < 3:
        fail("not enough resume-supported HERO stories available for a complete interview guide")
    screen_answers = phone_screen_script_answers(
        profile,
        company_name,
        role_title,
        job_description,
        all_stories,
        selected_stories,
        pitch_versions,
        supplied_context=supplied_context,
        resume_text=resume_text,
        notes_text=interview_notes,
    )
    screen_answers.extend(
        PreparedAnswer(response.prompt, prose_engine.spoken_register(response.answer).text)
        for response in question_prep.element_probe_responses(job_description, resume_text)
    )
    for item in screen_answers:
        enforce_prose_quality(
            item.answer,
            "interview_story_answer",
            label=f"Phone-screen script ({item.prompt[:50]})",
            mode="warn",
            check_template_leakage=False,
        )
    supplied_questions, supplied_changes = reframe_questions_to_positive(supplied_smart_questions(supplied_context))
    if supplied_changes:
        log_reframed_questions("supplied questions", supplied_changes)
    generated_questions, generated_changes = reframe_questions_to_positive(
        questions_to_ask(company_name, profile, job_description, supplied_context)
    )
    if generated_changes:
        log_reframed_questions("generated questions", generated_changes)
    final_questions: list[str] = []
    seen_questions = {question.lower() for question in supplied_questions}
    for question in generated_questions:
        if question.lower() in seen_questions:
            continue
        final_questions.append(question)
    if supplied_questions:
        final_questions = list(supplied_questions[:2]) + final_questions
    final_questions = final_questions[:6]

    document = Document()
    set_default_style(document)
    add_title(document, company_name, role_title)

    add_section(document, "Role Focus For This Call")
    for line in company_background_lines(job_description, profile, company_name, role_title):
        add_bullet(document, line)
    for line in top_answer_risk_lines(profile, company_name, role_title, context_bundle.round_records, global_round_records)[:3]:
        add_bullet(document, line)
    add_section(document, "Pre-Call Routine")
    for line in pre_interview_routine_lines(role_title, context_bundle.round_records, global_round_records):
        add_bullet(document, line)
    add_section(document, "Call Pacing")
    for line in phone_screen_first_round_playbook(profile, company_name, role_title, job_description, resume_text, interview_notes)[:6]:
        add_bullet(document, line)
    add_section(document, "Answer-Length Guide")
    for line in answer_mode_lines():
        add_bullet(document, line)
    add_section(document, "Anti-Filler Guide")
    for title, description, fix in communication_audit_reference(job_description, interview_notes)[:3]:
        add_bullet(document, f"{title}: {description} {fix}")
    add_section(document, "Rehearsal Method")
    add_bullet(document, "Use the 3-point spine for every story: hook, number or scale anchor, result.")
    add_bullet(document, "Memorize the first line and the last line of each story. Let the middle stay conversational.")
    for line in rehearsal_foundation_lines(role_title, context_bundle.round_records, global_round_records)[:2]:
        add_bullet(document, line)
    if debrief_summary.top_coaching_signals:
        add_section(document, "Recurring Delivery Habits")
        for item in debrief_summary.top_coaching_signals[:4]:
            add_bullet(document, item)

    add_page_break(document)

    add_section(document, "Likely First-Round Questions")
    for item in screen_answers:
        add_bullet(document, f"{item.prompt} Model answer: {item.answer}")
    add_recent_interview_question_prep_section(
        document,
        job_description,
        company_name,
        role_title,
        jobs_dir=jobs_dir,
        profile=profile,
        stories=all_stories,
        resume_text=resume_text,
    )

    add_section(document, "Memory System: Three-Point Spine")
    add_bullet(document, "For each story, lock in three things: the hook, the number or scale anchor, and the result.")
    add_bullet(document, "If you lose the middle, go back to the anchor phrase, then state the result and bridge the answer back to the role.")
    add_bullet(document, "Use the first sentence to name the situation and the last sentence to name why it matters here.")

    add_section(document, "Five Story Cards")
    for index, card in enumerate(selected_stories[:5], start=1):
        add_subsection(document, f"Story {index}: {card.title}")
        add_bullet(document, f"Memory anchor: {card.title}")
        add_bullet(document, f"Hook: {card.hook}")
        add_bullet(document, story_scale_anchor_label(card))
        add_bullet(document, f"Result: {card.result}")
        add_bullet(document, f"Role bridge: {story_specific_bridge(card, profile).replace('Bridge: ', '')}")
        add_bullet(document, f"Use for: {', '.join(card.story_types)}")

    add_section(document, "Questions To Ask")
    for question in final_questions:
        add_bullet(document, question)
    warnings = interview_question_quality_warnings(final_questions)
    if warnings:
        add_section(document, "Question Quality Audit")
        for warning in warnings:
            add_bullet(document, warning)

    add_section(document, "Four-Step Close")
    for line in four_step_close_lines(company_name, role_title, profile, job_description, supplied_context):
        add_bullet(document, line)

    add_section(document, "Thank-You Email Template")
    for line in thank_you_email_template_lines(company_name, role_title):
        add_bullet(document, line)

    add_section(document, "Red Flags To Avoid")
    for line in red_flag_lines(profile, role_title, job_description, context_bundle.round_records, global_round_records):
        add_bullet(document, line)

    add_section(document, "Post-Interview Self-Scorecard")
    add_body(document, f"Complete this within one hour of the {role_title} interview with {company_name}.")
    scorecard_dimensions, scorecard_interpretation, scorecard_prompts = interview_scorecard_template(company_name, role_title)
    for line in scorecard_dimensions:
        add_bullet(document, line)
    add_subsection(document, "Interpretation Guide")
    for line in scorecard_interpretation:
        add_bullet(document, line)
    add_subsection(document, "Debrief Prompts")
    for line in scorecard_prompts[:2]:
        add_bullet(document, line)

    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    scrub_document_for_job_language(document, job_description)
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    build_resume.assert_no_erp_language_for_non_erp_role(body, job_description, "interview cheat sheet")
    validate_text(body, company_name=company_name, role_title=role_title)
    assert_cheat_sheet_qc(body)
    output_docx.parent.mkdir(exist_ok=True)
    document.save(str(output_docx))


def assert_cheat_sheet_qc(text: str) -> None:
    """Enforce Master Appendix I interview cheat sheet quality controls before saving."""
    lowered = text.lower()

    tell_me_match = re.search(r"tell me about yourself\.?.{0,700}model answer", text, re.I | re.S)
    if not tell_me_match:
        fail("interview cheat sheet QC failed: Tell Me About Yourself script is missing")

    quantified_story = re.search(
        r"\b(?:story|result|proof)\b.{0,220}(?:\d+%|\$\d[\d,]*(?:\+|\s*m|\s*million)?|\b\d+\+?\s+(?:clients|users|sites|tools|reports|dashboards|accounts)\b)",
        text,
        re.I | re.S,
    )
    if not quantified_story:
        fail("interview cheat sheet QC failed: no story with a quantified result was found")

    question_section_match = re.search(
        r"(?:five power questions|questions to ask|closing questions|questions and closing)(.*)",
        text,
        re.I | re.S,
    )
    if not question_section_match or question_section_match.group(1).count("?") < 5:
        fail("interview cheat sheet QC failed: five power questions or equivalent closing questions section is missing")

    if not (
        re.search(r"\bthank[- ]you email template\b", text, re.I)
        or ("thank-you email" in lowered and "template" in lowered)
        or ("thank you email" in lowered and "template" in lowered)
    ):
        fail("interview cheat sheet QC failed: thank-you email template is missing")

    if not (
        "seven deadly mistakes" in lowered
        or "likely gaps and how to answer" in lowered
        or "red flags to avoid" in lowered
        or "gap to manage" in lowered
    ):
        fail("interview cheat sheet QC failed: seven deadly mistakes or equivalent gap section is missing")


def validate_text(
    text: str,
    *,
    company_name: str = "",
    role_title: str = "",
    jobs_dir: Path = PROJECT_ROOT / "jobs",
) -> None:
    validation_text = strip_external_question_prompts_for_validation(
        text,
        company_name,
        role_title,
        jobs_dir=jobs_dir,
    )
    if "--" in validation_text:
        fail("interview cheat sheet contains double dashes")
    assert_no_template_leakage(validation_text)
    placeholder_scan_text = remove_approved_bracketed_metadata(validation_text)
    for pattern in PLACEHOLDER_PATTERNS:
        if re.search(pattern, placeholder_scan_text, re.I):
            fail(f"placeholder text detected in interview cheat sheet: {pattern}")
    if build_resume.contains_cliche(validation_text):
        fail("interview cheat sheet contains generic cliché language")
    if build_resume.contains_ai_writing_word(validation_text):
        fail("interview cheat sheet contains banned AI-writing words")
    forbidden_fragments = (
        "The reason I think my background fits",
        "That matters here because this role requires",
        "process waste, operational variation, data-backed root cause analysis, and measurable efficiency improvement",
    )
    for fragment in forbidden_fragments:
        if fragment.lower() in validation_text.lower():
            fail(f"interview cheat sheet contains stale template language: {fragment}")


def build_interview_cheat_sheet() -> CheatSheetResult:
    build_resume.require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    build_resume.require_file(JOB_DESCRIPTION, "job description")

    job_description = read_text(JOB_DESCRIPTION)
    if not job_description:
        fail("jobs/job_description.txt is empty; refusing to create a placeholder cheat sheet")

    company_name = build_resume.extract_output_name(job_description)
    output_target_name = build_resume.extract_output_target_name(job_description)
    role_title = build_cover_letter.extract_role_title(job_description) or "Role"
    resume_docx = build_cover_letter.find_resume_output(job_description)
    resume_audit_state = resume_analysis.output_audit_state(resume_docx)
    output_name = f"Christian Estrada - {output_target_name} Interview Cheat Sheet.docx"
    if resume_audit_state == "FAIL":
        output_name = f"Christian Estrada - {output_target_name} FAIL Interview Cheat Sheet.docx"
    elif resume_audit_state == "POOR":
        output_name = f"Christian Estrada - {output_target_name} POOR Interview Cheat Sheet.docx"
    elif resume_audit_state == "BRIDGE":
        output_name = f"Christian Estrada - {output_target_name} BRIDGE Interview Cheat Sheet.docx"
    output_docx = OUTPUT_DIR / output_name
    return build_interview_cheat_sheet_for_inputs(
        job_description=job_description,
        resume_docx=resume_docx,
        output_docx=output_docx,
        company_name=company_name,
        role_title=role_title,
    )


def build_interview_cheat_sheet_for_inputs(
    *,
    job_description: str,
    resume_docx: Path,
    output_docx: Path,
    company_name: str,
    role_title: str,
) -> CheatSheetResult:
    build_resume.require_file(PROJECT_ROOT / "AGENTS.md", "AGENTS.md")
    if not job_description.strip():
        fail("job description is empty; refusing to create a placeholder cheat sheet")
    if not role_title or role_title == "Role":
        fail("could not determine role title; refusing to create a placeholder cheat sheet")
    assert_company_name_in_source(company_name, job_description, label="interview cheat sheet")
    prompt_state = question_prep.load_application_prompt_state()
    question_issues = question_prep.application_question_context_issues(job_description, prompt_state, workflow="commercial")
    actual_output = question_prep.application_question_draft_path(output_docx) if question_issues else output_docx
    with prose_engine.collect_spoken_repair_issues() as spoken_issues:
        build_document(company_name, role_title, job_description, resume_docx, actual_output)
    review_issues = tuple(dict.fromkeys((*question_issues, *spoken_issues)))
    if spoken_issues and not question_issues:
        draft_output = question_prep.application_question_draft_path(output_docx)
        actual_output.replace(draft_output)
        actual_output = draft_output
    if review_issues:
        question_prep.mark_docx_as_draft(actual_output, review_issues)
    render_checks.render_docx(actual_output)
    return CheatSheetResult(company_name, role_title, resume_docx, actual_output)


def main() -> None:
    result = build_interview_cheat_sheet()
    print(f"Company: {result.company_name}")
    print(f"Role: {result.role_title}")
    print(f"Resume source: {result.resume_docx}")
    print(f"Output DOCX: {result.output_docx}")


if __name__ == "__main__":
    main()
