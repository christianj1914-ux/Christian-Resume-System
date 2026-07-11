#!/usr/bin/env python3
"""Generate a post-interview follow-up email as a Word document."""

from __future__ import annotations

import re
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt

import build_resume
import build_thank_you
import interview_context
import job_search_guidance as guidance
import proof_text
import resume_analysis
from utils import assert_no_template_leakage, discussion_topic_sentence, enforce_prose_quality


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOBS_DIR = PROJECT_ROOT / "jobs"
OUTPUT_DIR = PROJECT_ROOT / "output"
JOB_DESCRIPTION = JOBS_DIR / "job_description.txt"
DEBRIEF_HISTORY = JOBS_DIR / "debrief_history.txt"
COMPANY_RESEARCH = JOBS_DIR / "company_research.txt"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig") if path.exists() else ""


def field_value(entry: str, label: str) -> str:
    match = re.search(rf"(?im)^\s*{re.escape(label)}:\s*(.*?)\s*$", entry)
    return match.group(1).strip() if match else ""


def suggest_followup_window(debrief_entry: str) -> str:
    interview_date_raw = field_value(debrief_entry, "Interview date")
    feedback = build_thank_you.section_value(debrief_entry, "Feedback received").lower()
    try:
        interview_date = datetime.strptime(interview_date_raw, "%m/%d/%Y").date()
    except ValueError:
        return "Suggested send window: about 5 to 7 business days after the interview, unless the team gave a different timeline."

    if "1 to 2 weeks" in feedback or "1-2 weeks" in feedback or "one to 2 weeks" in feedback:
        start = interview_date + timedelta(days=7)
        end = interview_date + timedelta(days=14)
        return f"Suggested send window: {start.isoformat()} to {end.isoformat()} based on the timeline mentioned in the interview."
    send_date = interview_date + timedelta(days=5)
    return f"Suggested send window: on or after {send_date.isoformat()} if no update has arrived."


def interview_followup_body(
    company_name: str,
    role_title: str,
    debrief_entry: str,
    proof_points: list[str],
    news_line: str = "",
) -> tuple[str, str, list[str]]:
    job_description = read_text(JOB_DESCRIPTION)
    primary_lane = build_resume.job_problem_profile(job_description, "").primary_lane if job_description else "implementation_delivery"
    role_language = build_thank_you.section_value(debrief_entry, "Specific interviewer language about the role")
    if not role_language:
        role_language = build_thank_you.section_value(debrief_entry, "Specific language the interviewer used about the role")
    followups = build_thank_you.section_value(debrief_entry, "Stories that generated follow-up questions")

    subject = f"Following up on the {role_title} interview process"
    discussion_sentence = discussion_topic_sentence(
        role_language,
        intro="I especially appreciated the detail around",
    )
    proof_sentence = (
        f"One proof point from my background kept coming to mind: {proof_text.sanitize_proof_sentence(proof_points[0], max_words=22, artifact='interview_followup_body')}"
        if proof_points
        else "That discussion reinforced how closely my background lines up with implementation discipline, stakeholder coordination, and practical follow-through."
    )
    bridge_sentence = discussion_topic_sentence(
        followups,
        intro="I also kept thinking about the follow-up around",
    )
    if bridge_sentence:
        bridge_sentence = bridge_sentence.rstrip(".") + "; it felt like a useful window into where the team most needs confidence."
    else:
        bridge_sentence = build_thank_you.lane_followup_fallback_sentence(primary_lane)
    paragraphs = [
        f"Hello [Name],\n\nI wanted to follow up on the {role_title} interview process and reiterate my interest in the opportunity. The conversation made the role feel even more aligned with the kind of work where I can be most useful.",
        discussion_sentence or "",
        f"I also noticed {news_line}, which added a helpful layer of context around where the team is headed." if news_line else "",
        f"{proof_sentence} {bridge_sentence}",
        "If the team needs any additional detail as the process continues, I can send it right away. Thank you again for your time and consideration.\n\nBest,\nChristian Estrada",
    ]
    notes: list[str] = []
    if not debrief_entry:
        notes.append("No matching debrief entry was found; this follow-up uses active job context only.")
    if not proof_points:
        notes.append("No matching generated resume proof point was found; add one concrete result before sending if needed.")
    body = "\n\n".join(paragraphs)
    if "concrete implementation, analytics, and customer-adoption" in body:
        print("THANK-YOU WARNING: hardcoded fallback phrase detected. Check lane detection.")
    return subject, body, notes


def build_interview_followup() -> Path:
    job_description = read_text(JOB_DESCRIPTION).strip()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")

    company_name = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role_title = resume_analysis.extract_job_title(job_description) or "the role"
    debrief_entry = build_thank_you.find_latest_debrief_for_company(company_name, DEBRIEF_HISTORY)
    if not debrief_entry:
        raise SystemExit(f"No debrief found for {company_name}. Run python tasks.py debrief first.")

    resume_docx = build_thank_you.latest_matching_resume(job_description)
    proof_points = build_thank_you.proof_points_from_resume(resume_docx)
    company_context = interview_context.company_research_context(
        JOBS_DIR,
        company_name,
        role_title,
        company_research_path=COMPANY_RESEARCH,
    )
    news_line = interview_context.recent_company_news_line(company_context, company_name)
    subject, body, notes = interview_followup_body(company_name, role_title, debrief_entry, proof_points, news_line)
    timing = suggest_followup_window(debrief_entry)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {build_thank_you.clean_filename_piece(output_target_name)} Interview Follow-Up.docx"

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.add_run(f"Interview Follow-Up - {company_name}").bold = True
    doc.add_paragraph(f"Role: {role_title}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(timing)
    doc.add_paragraph(f"Subject: {subject}")
    doc.add_paragraph("Email rules: " + " ".join(guidance.concise_email_rules()))
    for paragraph in body.split("\n\n"):
        doc.add_paragraph(paragraph)
    if news_line:
        doc.add_paragraph(f"Verified company update used: {news_line}")
    if notes:
        doc.add_paragraph("Review Notes").runs[0].bold = True
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")

    visible_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    assert_no_template_leakage(visible_text)
    enforce_prose_quality(body, "interview_followup_body", label="Interview follow-up body")
    doc.save(output)
    print(f"Interview follow-up created: {output}")
    return output


def main() -> None:
    build_interview_followup()


if __name__ == "__main__":
    main()
