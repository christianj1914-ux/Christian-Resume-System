#!/usr/bin/env python3
"""Build a compact interview review document from structured debrief data."""

from __future__ import annotations

import _bootstrap

_bootstrap.ensure_script_path()

import re
from pathlib import Path
from typing import Mapping, Sequence

from docx import Document
from docx.shared import Pt

import build_resume
import interview_context
from config.paths import JOB_DESCRIPTION, OUTPUT_DIR, PROJECT_ROOT
from utils import fail, optional_text


JOBS_DIR = PROJECT_ROOT / "jobs"


def active_target() -> tuple[str, str]:
    job_description = optional_text(JOB_DESCRIPTION)
    if not job_description:
        return "", ""
    company_name = build_resume.extract_output_name(job_description)
    role_title = build_resume.extract_job_title(job_description) or ""
    return company_name, role_title


def preferred_round_record() -> dict[str, object]:
    company_name, role_title = active_target()
    if company_name:
        scoped = interview_context.load_round_records(JOBS_DIR, company_name, role_title)
        if scoped:
            return scoped[0]
    records = interview_context.load_round_records(JOBS_DIR)
    if records:
        return records[0]
    fail("no structured interview debriefs were found")


def _clean_lines(value: object, limit: int = 8) -> list[str]:
    return interview_context._split_lines(value)[:limit]


def interview_review_sections(record: Mapping[str, object]) -> list[tuple[str, list[str]]]:
    normalized = interview_context.normalize_round_record(record)
    review = interview_context._serialize_performance_review(
        normalized.get("performance_review") if isinstance(normalized.get("performance_review"), Mapping) else {}
    )
    analysis = interview_context.review_analysis_from_record(normalized)
    decision_signal = analysis.get("decision_signal", {}) if isinstance(analysis, Mapping) else {}
    diagnosis = analysis.get("positioning_diagnosis", {}) if isinstance(analysis, Mapping) else {}
    rewrites = analysis.get("language_rewrites", {}) if isinstance(analysis, Mapping) else {}
    strategy = analysis.get("answer_strategy", {}) if isinstance(analysis, Mapping) else {}
    assets = analysis.get("answer_assets", {}) if isinstance(analysis, Mapping) else {}
    career = analysis.get("career_targeting", {}) if isinstance(analysis, Mapping) else {}

    sections: list[tuple[str, list[str]]] = []
    sections.append(
        (
            "Round Facts",
            [
                f"Company: {normalized.get('company_name', '')}",
                f"Role: {normalized.get('role_title', '') or 'None supplied.'}",
                f"Interview date: {normalized.get('interview_date', '')}",
                f"Round: {normalized.get('round_number', '') or 'general'}",
                f"Outcome: {normalized.get('outcome', '') or 'unknown'}",
            ],
        )
    )
    sections.append(
        (
            "Decision Signal",
            [
                str(decision_signal.get("headline", "")).strip(),
                *_clean_lines(decision_signal.get("translation", []), limit=4),
            ],
        )
    )
    sections.append(
        (
            "Positioning Diagnosis",
            [
                *_clean_lines(diagnosis.get("reasons", []), limit=5),
                *([f"Good news: {line}" for line in _clean_lines(diagnosis.get("good_news", []), limit=3)]),
                *([f"Hard truth: {str(diagnosis.get('hard_truth', '')).strip()}"] if str(diagnosis.get("hard_truth", "")).strip() else []),
            ],
        )
    )
    sections.append(
        (
            "Language Rewrites",
            [
                "Ownership ladder: " + " -> ".join(_clean_lines(rewrites.get("ownership_ladder", []), limit=5)),
                *([f"Avoid: {line}" for line in _clean_lines(rewrites.get("avoid", []), limit=4)]),
                *([f"Prefer: {line}" for line in _clean_lines(rewrites.get("prefer", []), limit=5)]),
                *([f"Consultative phrase: {line}" for line in _clean_lines(rewrites.get("consultative_phrases", []), limit=4)]),
            ],
        )
    )
    sections.append(
        (
            "Answer Strategy",
            [
                "Default structure: " + " -> ".join(_clean_lines(strategy.get("default_structure", []), limit=4)),
                *(_clean_lines(strategy.get("delivery_shifts", []), limit=6)),
            ],
        )
    )
    sections.append(
        (
            "Answer Assets",
            [
                *([f"Best takeaway: {str(assets.get('best_takeaway_statement', '')).strip()}"] if str(assets.get("best_takeaway_statement", "")).strip() else []),
                *([f"Interviewer question: {line}" for line in _clean_lines(assets.get("interviewer_questions", []), limit=6)]),
                *([f"Role language: {line}" for line in _clean_lines(assets.get("role_language_lines", []), limit=5)]),
                *([f"Company signal: {line}" for line in _clean_lines(assets.get("company_signal_lines", []), limit=4)]),
            ],
        )
    )
    sections.append(
        (
            "Career Targeting",
            [
                str(career.get("fit_read", "")).strip(),
                *(_clean_lines(career.get("next_steps", []), limit=5)),
            ],
        )
    )
    sections.append(
        (
            "Performance Summary",
            [
                str(review.get("summary", "")).strip(),
                *([f"Coaching signal: {item.get('label', '')} - {item.get('detail', '')}".strip(" -") for item in review.get("coaching_signals", []) if isinstance(item, Mapping)]),
                *([f"Next-round risk: {line}" for line in _clean_lines(review.get("next_round_risks", []), limit=5)]),
            ],
        )
    )
    appendix_lines = _clean_lines(normalized.get("imported_artifacts", []), limit=4)
    review_appendix = str(normalized.get("review_appendix_path", "")).strip()
    if review_appendix:
        appendix_lines.append("Review appendix: " + review_appendix)
    if appendix_lines:
        sections.append(("Appendix References", appendix_lines))
    return [(title, [line for line in lines if line.strip()]) for title, lines in sections if any(line.strip() for line in lines)]


def _set_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def _add_title(document: Document, record: Mapping[str, object]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        f"Christian Estrada - Interview Review - {record.get('company_name', '')} - {record.get('role_title', '') or 'General'}"
    )
    run.bold = True
    run.font.size = Pt(16)


def _add_section(document: Document, title: str, lines: Sequence[str]) -> None:
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    for line in lines:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(line)


def output_path_for_record(record: Mapping[str, object]) -> Path:
    company = re.sub(r'[\\/:*?"<>|]+', "", str(record.get("company_name", "")).strip()) or "Unknown Company"
    role = re.sub(r'[\\/:*?"<>|]+', "", str(record.get("role_title", "")).strip()) or "General"
    filename = f"Christian Estrada - Interview Review - {company} - {role}.docx"
    return OUTPUT_DIR / filename


def build_document(record: Mapping[str, object], output_path: Path) -> None:
    document = Document()
    _set_default_style(document)
    normalized = interview_context.normalize_round_record(record)
    _add_title(document, normalized)
    for title, lines in interview_review_sections(normalized):
        _add_section(document, title, lines)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    record = preferred_round_record()
    output_path = output_path_for_record(record)
    build_document(record, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
