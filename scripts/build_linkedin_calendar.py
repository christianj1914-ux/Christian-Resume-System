#!/usr/bin/env python3
"""Generate a 30-day LinkedIn content calendar for the active search focus."""

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from docx import Document

import job_search_guidance as guidance
import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"


CONTENT_TYPES = ("Proof Point", "Lesson Learned", "Perspective", "Process", "Reflection")


def read_job() -> str:
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip() if JOB_DESCRIPTION.exists() else ""


def build_calendar() -> Path:
    job_description = read_job()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")
    profile = resume_analysis.job_problem_profile(job_description)
    specialty = resume_analysis.role_specialty_phrase(job_description, profile.core_problem)
    start = date.today()

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - LinkedIn Content Calendar {start.strftime('%B %Y')}.docx"
    doc = Document()
    doc.add_heading("30-Day LinkedIn Content Calendar", level=1)
    doc.add_paragraph(f"Search focus: {profile.lane_label} / {specialty}")

    doc.add_heading("Posting Strategy", level=2)
    doc.add_paragraph(
        "Balance proof, lessons, perspective, process, and reflection. Keep claims tied to verified resume evidence: 80+ client engagements, 200+ reporting tools, 60+ workshops/QBRs, five-site systems ownership, and $1M+ account-risk stabilization."
    )
    doc.add_paragraph("Use operator language, not influencer language. Each post should teach something practical about customer problems, workflow clarity, reporting, implementation, or adoption.")

    doc.add_heading("Theme Bank", level=2)
    for item in guidance.safe_thought_leadership_themes(profile, job_description):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Comment Strategy", level=2)
    for item in guidance.linkedin_comment_strategy_lines(profile, job_description):
        doc.add_paragraph(item, style="List Bullet")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Date", "Type", "Topic", "Draft Angle"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for offset in range(30):
        current = start + timedelta(days=offset)
        content_type = CONTENT_TYPES[offset % len(CONTENT_TYPES)]
        topic = {
            "Proof Point": f"A measurable {specialty} outcome",
            "Lesson Learned": "What implementation work teaches about adoption",
            "Perspective": f"Why {profile.core_problem} needs practical operating rhythm",
            "Process": "A simple way to make risks and owners visible",
            "Reflection": "How customer trust is rebuilt through consistent follow-through",
        }[content_type]
        angle = {
            "Proof Point": "Use one metric, explain the business problem, and avoid confidential details.",
            "Lesson Learned": "Share the insight without naming an employer or implying unsupported ownership.",
            "Perspective": "Comment on the work pattern, not a hot take for attention.",
            "Process": "Describe a reusable checklist or diagnostic question.",
            "Reflection": "Keep it concise and tied to customer, team, or workflow outcomes.",
        }[content_type]
        row = table.add_row().cells
        row[0].text = current.isoformat()
        row[1].text = content_type
        row[2].text = topic
        row[3].text = angle

    doc.save(output)
    print(f"LinkedIn content calendar created: {output}")
    return output


def main() -> None:
    build_calendar()


if __name__ == "__main__":
    main()
