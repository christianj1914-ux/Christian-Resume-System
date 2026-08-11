#!/usr/bin/env python3
"""Generate tailored LinkedIn profile update suggestions from verified resume/JD evidence."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

import job_search_guidance as guidance
import resume_analysis
from utils import docx_visible_text, enforce_prose_quality, optional_text


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"
def headline_options(profile: resume_analysis.JobProblemProfile, job_description: str) -> list[str]:
    options_by_lane = {
        "presales_solution": [
            "Solutions Consultant | 80+ Enterprise Clients | Discovery to Go-Live",
            "Pre-Sales Consultant | Executive Workshops | Practical Solution Design",
            "Solution Consulting Leader | Customer Discovery | Enterprise Software",
        ],
        "implementation_delivery": [
            "Implementation Consultant | 80+ Client Engagements | Go-Live Readiness",
            "Enterprise Software Consultant | ERP to Adoption | Customer Delivery",
            "Implementation Leader | Data Migration | Testing, Training, Go-Live",
        ],
        "customer_success": [
            "Customer Success Consultant | $1M+ Risk Stabilized | Adoption to Renewal",
            "Technical CSM | Executive QBRs | Enterprise Software Adoption",
            "Customer Outcomes Leader | Account Health | Value Realization",
        ],
        "analytics_operations": [
            "Analytics Consultant | 200+ KPI Tools Built | Workflow to Decisions",
            "Operations Analytics Leader | SQL Validation | Executive Reporting",
            "Business Systems Analyst | Dashboards to Action | Process Visibility",
        ],
        "change_enablement": [
            "Change Adoption Consultant | Training to Workflow | Stakeholder Alignment",
            "Enterprise Systems Consultant | Adoption, Training, Reporting | 150+ Users",
            "Change Enablement Leader | Workshops to Durable Process Change",
        ],
        "process_improvement": [
            "Process Improvement Consultant | 78% Manual Work Reduction | Workflow Design",
            "Operations Improvement Leader | Root Cause to Adoption | KPI Visibility",
            "Business Process Consultant | 22% Fewer Discrepancies | Practical Controls",
        ],
    }
    options = options_by_lane.get(profile.primary_lane, options_by_lane["implementation_delivery"])
    return [option[:119] for option in options]


def about_components(profile: resume_analysis.JobProblemProfile, job_description: str) -> tuple[str, str, tuple[str, ...], str]:
    specialty = resume_analysis.role_specialty_phrase(job_description, profile.core_problem)
    credential = (
        "Enterprise software, implementation, customer-success, and operations professional with more than ten years of experience "
        "turning customer and system work into reliable operating results."
    )
    capability = (
        f"I work where {specialty} requires structured discovery, practical workflow design, stakeholder alignment, adoption, "
        "and decision-ready reporting."
    )
    proof_bullets = (
        "Supported 80+ client engagements across enterprise software implementation and customer delivery.",
        "Owned mission-critical systems across five sites and 150+ users.",
        "Built 200+ reporting tools to strengthen decision visibility.",
        "Facilitated 60+ workshops and QBRs for users and executive stakeholders.",
        "Stabilized high-risk customer accounts representing $1M+ in annual revenue.",
    )
    cta = f"Open to conversations about roles where {profile.core_problem} needs to become practical, sustained operating change."
    return credential, capability, proof_bullets, cta


def about_section(profile: resume_analysis.JobProblemProfile, job_description: str) -> str:
    """Plain-text representation for review; the Word builder renders proof as real bullets."""
    credential, capability, proof_bullets, cta = about_components(profile, job_description)
    return "\n\n".join((credential, capability, "\n".join(proof_bullets), cta))


def featured_proof_points(resume_text: str) -> list[str]:
    patterns = (r"80\+", r"150\+", r"200\+", r"60\+", r"\$1M\+", r"78%", r"22%")
    points = []
    for line in resume_text.splitlines():
        clean = re.sub(r"\s+", " ", line).strip()
        if any(re.search(pattern, clean) for pattern in patterns) and clean not in points:
            points.append(clean)
        if len(points) >= 5:
            break
    return points or [
        "80+ client engagements across implementation, customer success, and adoption.",
        "200+ reporting tools built to improve decision visibility.",
        "$1M+ in at-risk annual revenue stabilized through account recovery work.",
    ]


def skill_suggestions(profile: resume_analysis.JobProblemProfile, job_description: str) -> list[str]:
    base = [
        "Enterprise Software Implementation",
        "Customer Success",
        "Requirements Gathering",
        "Data Migration",
        "Executive Workshops",
        "Adoption Planning",
        "SQL Validation",
        "Business Intelligence",
    ]
    base.extend(profile.safe_terms[:6])
    base.extend(resume_analysis.visible_role_specialties(job_description))
    deduped = []
    for item in base:
        clean = item.strip()
        if clean and clean.lower() not in {existing.lower() for existing in deduped}:
            deduped.append(clean)
    return deduped[:15]


def recruiter_keywords(profile: resume_analysis.JobProblemProfile, job_description: str) -> list[str]:
    return guidance.recruiter_facing_keywords(profile, job_description)


def thought_leadership_themes(profile: resume_analysis.JobProblemProfile, job_description: str) -> list[str]:
    return guidance.safe_thought_leadership_themes(profile, job_description)


def comment_strategy(profile: resume_analysis.JobProblemProfile, job_description: str) -> list[str]:
    return guidance.linkedin_comment_strategy_lines(profile, job_description)


def build_linkedin_update() -> Path:
    job_description = optional_text(JOB_DESCRIPTION)
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")
    selected_resume = resume_analysis.choose_resume(job_description)
    resume_text = docx_visible_text(selected_resume)
    profile = resume_analysis.job_problem_profile(job_description, resume_text)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / "Christian Estrada - LinkedIn Update Guide.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Christian Estrada - LinkedIn Update Guide", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Target lane: {profile.lane_label}")

    doc.add_heading("Headline Options", level=2)
    for option in headline_options(profile, job_description):
        doc.add_paragraph(option, style="List Bullet")

    doc.add_heading("About Section Draft", level=2)
    credential, capability, proof_bullets, cta = about_components(profile, job_description)
    composed_about = " ".join((credential, capability, *proof_bullets, cta))
    enforce_prose_quality(
        composed_about,
        "generic",
        label="LinkedIn About section",
        mode="warn",
        check_template_leakage=False,
    )
    doc.add_paragraph(credential)
    doc.add_paragraph(capability)
    for proof_bullet in proof_bullets:
        doc.add_paragraph(proof_bullet, style="List Bullet")
    doc.add_paragraph(cta)

    doc.add_heading("Featured Proof Points", level=2)
    for point in featured_proof_points(resume_text):
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Featured Proof Point Review", level=2)
    doc.add_paragraph("Self-check only: review scale already supported by the approved source resumes; do not turn blank cells into new claims.")
    review = doc.add_table(rows=1, cols=5)
    review.style = "Table Grid"
    for cell, header in zip(review.rows[0].cells, ("Proof point", "Time", "Money", "Team", "Scope")):
        cell.text = header
    review_rows = (
        ("80+ client engagements", "", "", "", "Client-delivery breadth"),
        ("Five sites / 150+ users", "", "", "150+ users", "Multi-site systems ownership"),
        ("200+ reporting tools", "", "", "", "Decision visibility"),
        ("60+ workshops and QBRs", "", "", "Users and executives", "Enablement and executive communication"),
        ("$1M+ stabilized annual revenue", "", "$1M+ annual revenue", "Cross-functional recovery", "At-risk account stabilization"),
    )
    for row_values in review_rows:
        row = review.add_row().cells
        for cell, value in zip(row, row_values):
            cell.text = value

    doc.add_heading("Skills to Consider", level=2)
    doc.add_paragraph("; ".join(skill_suggestions(profile, job_description)))

    doc.add_heading("Recruiter-Facing Keywords To Keep Visible", level=2)
    doc.add_paragraph("; ".join(recruiter_keywords(profile, job_description)))

    doc.add_heading("Safe Thought-Leadership Themes", level=2)
    for item in thought_leadership_themes(profile, job_description):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Comment Strategy", level=2)
    for item in comment_strategy(profile, job_description):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Informational Interview Positioning", level=2)
    for item in guidance.informational_interview_lines("", ""):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Profile Guardrails", level=2)
    guardrails = [
        "Do not use LinkedIn page content as source material for resumes.",
        "Do not claim direct quota, NRR, GRR, or closed expansion ownership unless Christian adds verified support.",
        "Keep ERP ownership tied to East West Manufacturing; keep Aptean Encompix tied to Aptean implementation/customer success work.",
        "Treat this as suggested profile language, not a resume-source database.",
        "Avoid influencer-style posting. Proof, observation, and practical operator language are stronger than generic personal-brand talk.",
    ]
    for guardrail in guardrails:
        doc.add_paragraph(guardrail, style="List Bullet")

    doc.save(output)
    print(f"LinkedIn update guide created: {output}")
    return output


def main() -> None:
    build_linkedin_update()


if __name__ == "__main__":
    main()
