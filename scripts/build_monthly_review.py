#!/usr/bin/env python3
"""Generate a monthly job-search system review report."""

from __future__ import annotations

import csv
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
APPLICATIONS_CSV = PROJECT_ROOT / "scratch" / "applications.csv"
DEBRIEF_HISTORY = PROJECT_ROOT / "jobs" / "debrief_history.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"
ADVANCED_STATUSES = {"phone_screen", "interview", "final_round", "offer"}


def rows() -> list[dict[str, str]]:
    if not APPLICATIONS_CSV.exists():
        return []
    with APPLICATIONS_CSV.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def parse_date(value: str) -> datetime | None:
    try:
        return datetime.strptime(value, "%Y-%m-%d")
    except (TypeError, ValueError):
        return None


def current_month_rows(all_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    now = datetime.now()
    result = []
    for row in all_rows:
        added = parse_date(row.get("date_added", ""))
        if added and added.year == now.year and added.month == now.month:
            result.append(row)
    return result


def output_builds_this_month() -> int:
    now = datetime.now()
    count = 0
    if not OUTPUT_DIR.exists():
        return 0
    for path in OUTPUT_DIR.glob("*.docx"):
        modified = datetime.fromtimestamp(path.stat().st_mtime)
        if modified.year == now.year and modified.month == now.month:
            count += 1
    return count


def debrief_count() -> int:
    if not DEBRIEF_HISTORY.exists():
        return 0
    return DEBRIEF_HISTORY.read_text(encoding="utf-8-sig").count("POST-INTERVIEW DEBRIEF CAPTURED")


def build_monthly_review() -> Path:
    all_rows = rows()
    month_rows = current_month_rows(all_rows)
    status_counts = Counter(row.get("current_status", "draft") or "draft" for row in month_rows)
    advanced = sum(1 for row in month_rows if row.get("current_status") in ADVANCED_STATUSES)
    now = datetime.now()

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - Monthly Review {now.strftime('%B %Y')}.docx"
    doc = Document()
    doc.add_heading(f"Monthly Review - {now.strftime('%B %Y')}", level=1)

    doc.add_heading("Volume and Activity", level=2)
    doc.add_paragraph(f"Applications added this month: {len(month_rows)}")
    doc.add_paragraph(f"Total tracked applications: {len(all_rows)}")
    doc.add_paragraph(f"Debrief entries captured overall: {debrief_count()}")
    doc.add_paragraph(f"DOCX outputs modified this month: {output_builds_this_month()}")
    if not month_rows:
        doc.add_paragraph("Activity gap: no applications were added to the tracker this month.")

    doc.add_heading("Advancement Rates", level=2)
    doc.add_paragraph(f"Advanced this month: {advanced}")
    rate = "0%" if not month_rows else f"{round((advanced / len(month_rows)) * 100)}%"
    doc.add_paragraph(f"Advancement rate: {rate}")

    doc.add_heading("Pipeline Status", level=2)
    if status_counts:
        for status, count in sorted(status_counts.items()):
            doc.add_paragraph(f"{status}: {count}", style="List Bullet")
    else:
        doc.add_paragraph("No monthly tracker rows available.")

    doc.add_heading("Next Month Focus", level=2)
    if len(month_rows) < 10:
        doc.add_paragraph("Increase application volume or improve tracker discipline before drawing strong conclusions.")
    elif advanced == 0:
        doc.add_paragraph("Review target fit, resume top-third proof, and outreach strategy before increasing volume.")
    else:
        doc.add_paragraph("Study the applications that advanced and bias next month's targeting toward similar lanes and company types.")

    doc.save(output)
    print(f"Monthly review created: {output}")
    return output


def main() -> None:
    build_monthly_review()


if __name__ == "__main__":
    main()
