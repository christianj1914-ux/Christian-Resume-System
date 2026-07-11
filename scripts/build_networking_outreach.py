#!/usr/bin/env python3
"""Generate networking outreach templates for the active target company."""

from __future__ import annotations

from pathlib import Path

from docx import Document

import job_search_guidance as guidance
import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"


def read_job() -> str:
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip() if JOB_DESCRIPTION.exists() else ""


def build_outreach() -> Path:
    job_description = read_job()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")
    company = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role = resume_analysis.extract_job_title(job_description) or "the role"
    profile = resume_analysis.job_problem_profile(job_description)
    specialty = resume_analysis.role_specialty_phrase(job_description, profile.core_problem)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {output_target_name} Networking Outreach.docx"
    doc = Document()
    doc.add_heading(f"Networking Outreach - {company}", level=1)
    doc.add_paragraph(f"Role: {role}")

    linkedin = (
        f"Hi [Name], I am exploring the {role} role at {company}. "
        f"Your work in {specialty} stood out, and I would value any quick advice on the team."
    )
    linkedin = linkedin[:299]
    email = (
        f"Subject: Quick question about {company}\n\n"
        f"Hi [Name],\n\nI am exploring the {role} opportunity at {company} and noticed your background connects closely to {specialty}. "
        "I am coming from enterprise software implementation, customer adoption, analytics, and workflow improvement work, and I am trying to understand what makes someone successful on this team.\n\n"
        "If you are open to it, I would appreciate one or two pieces of advice about the role, the team, or the customer problems the group is solving.\n\n"
        "Best,\nChristian Estrada"
    )
    informational = (
        f"Subject: Informational interview request - {company}\n\n"
        f"Hi [Name],\n\nI am exploring the {role} role at {company} and your path stood out because it sits close to {specialty}. "
        "If you would be open to a 20-minute conversation, I would appreciate the chance to learn what the team values most and where new hires become useful fastest.\n\n"
        "Best,\nChristian Estrada"
    )
    referral = (
        f"Subject: Possible referral question - {role}\n\n"
        f"Hi [Name],\n\nI am applying for the {role} role at {company}. The role appears to focus on {profile.core_problem}, which connects strongly to my background supporting 80+ client engagements, building 200+ reporting tools, and driving implementation/adoption work across complex customer environments.\n\n"
        "If you feel comfortable after reviewing my background, would you be open to referring me or pointing me toward the best person to contact? No pressure either way; I appreciate any guidance you can share.\n\n"
        "Best,\nChristian Estrada"
    )

    for title, body in (
        ("LinkedIn Connection Request", linkedin),
        ("Informational Interview Ask", informational),
        ("First-Degree or Warm Email", email),
        ("Referral Ask", referral),
    ):
        doc.add_heading(title, level=2)
        for paragraph in body.split("\n\n"):
            doc.add_paragraph(paragraph)

    doc.add_heading("Informational Interview Rules", level=2)
    for line in guidance.informational_interview_lines(company, role):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Follow-Up Sequence", level=2)
    for line in guidance.follow_up_sequence_lines(company, role):
        doc.add_paragraph(line, style="List Bullet")

    doc.save(output)
    print(f"Networking outreach document created: {output}")
    return output


def main() -> None:
    build_outreach()


if __name__ == "__main__":
    main()
