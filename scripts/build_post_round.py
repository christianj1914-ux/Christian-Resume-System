#!/usr/bin/env python3
"""Generate post-round follow-up and next-round prep from the latest debrief."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

import build_thank_you
import business_context
import interview_context
import resume_analysis
from utils import assert_no_template_leakage, discussion_topic_sentence, enforce_prose_quality


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOBS_DIR = PROJECT_ROOT / "jobs"
OUTPUT_DIR = PROJECT_ROOT / "output"
JOB_DESCRIPTION = JOBS_DIR / "job_description.txt"
DEBRIEF_HISTORY = JOBS_DIR / "debrief_history.txt"
COMPANY_RESEARCH = JOBS_DIR / "company_research.txt"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig") if path.exists() else ""


def section_value(entry: str, label: str) -> str:
    match = re.search(rf"(?ims)^{re.escape(label)}:\s*(.*?)(?=\n[A-Z][A-Za-z ]+:\s|\Z)", entry)
    return re.sub(r"\s+", " ", match.group(1)).strip() if match else ""


def latest_debrief(company: str) -> str:
    return build_thank_you.find_latest_debrief_for_company(company, DEBRIEF_HISTORY)


def bullets_from_signals(entry: str) -> list[str]:
    text = entry.lower()
    bullets = []
    if "excel" in text:
        bullets.append("Excel comfort: answer with reporting, exports, validation, analysis, and the 200+ dashboards/reporting-tools proof point.")
    if "training" in text or "train others" in text:
        bullets.append("Training others: use 60+ workshops/QBRs, role-based customer enablement, workflow clarity, and adoption outcomes.")
    if "technical project" in text or "technical projects" in text:
        bullets.append("Technical projects: use implementation lifecycle, data migration, SQL validation, issue tracking, integrations, and technical handoffs.")
    if "spectra path" in text:
        bullets.append("Named product line: ask what implementation complexity, customer workflow, and ticket patterns are most common for that product.")
    if "jira" in text:
        bullets.append("Jira: speak to issue ownership, priority, dependencies, business impact, and customer update rhythm.")
    if "azure" in text:
        bullets.append("Azure: bridge carefully to cloud-adjacent collaboration and documentation, not unsupported Azure administration.")
    if "hipaa" in text or "phi" in text:
        bullets.append("HIPAA/PHI: do not overclaim compliance ownership; emphasize documentation, validation, escalation, access awareness, and regulated-workflow care.")
    if not bullets:
        bullets.append("Review the latest debrief and prepare answers for any topic that created follow-up questions or hesitation.")
    return bullets


def post_round_lane_phrase(role_title: str, job_description: str = "") -> str:
    lane_phrase_by_lane = {
        "implementation_delivery": "connect implementation discipline, customer communication, and technical issue ownership",
        "customer_success": "connect customer communication, adoption signals, and issue ownership",
        "change_enablement": "connect stakeholder communication, training, and adoption follow-through",
        "analytics_operations": "connect reporting, workflow clarity, and customer-facing follow-through",
        "presales_solution": "connect discovery, stakeholder communication, and solution credibility",
        "corporate_strategy": "connect structured analysis, stakeholder communication, and decision follow-through",
    }
    role_lower = role_title.lower()
    if "analyst" in role_lower or "analytics" in role_lower:
        return lane_phrase_by_lane["analytics_operations"]
    if "change" in role_lower or "adoption" in role_lower:
        return lane_phrase_by_lane["change_enablement"]
    if job_description.strip():
        profile = resume_analysis.job_problem_profile(job_description)
        return lane_phrase_by_lane.get(profile.primary_lane, "connect customer communication, ownership, and practical follow-through")
    return lane_phrase_by_lane["implementation_delivery"]


def post_round_followup_email(role_title: str, company: str, entry: str, news_line: str = "", job_description: str = "") -> str:
    role_language = section_value(entry, "Specific interviewer language about the role")
    role_topic_sentence = discussion_topic_sentence(
        role_language,
        intro="I appreciated the detail around",
    ) or "I appreciated the chance to learn more about the role and the team."
    lane_phrase = post_round_lane_phrase(role_title, job_description)
    return (
        f"Subject: Thank you - {role_title}\n\n"
        f"Hi [Name],\n\nThank you again for speaking with me about the {role_title} role. "
        f"{role_topic_sentence} "
        f"The conversation reinforced my interest in {company}, especially the opportunity to {lane_phrase}."
        f"{' I also noticed ' + news_line + '.' if news_line else ''}\n\n"
        "If helpful, I can answer any additional questions as the process continues.\n\n"
        "Best,\nChristian Estrada"
    )


def build_post_round() -> Path:
    job_description = read_text(JOB_DESCRIPTION).strip()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")
    company = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role_title = resume_analysis.extract_job_title(job_description) or "Target Role"
    entry = latest_debrief(company)
    if not entry:
        raise SystemExit(f"No debrief found for {company}. Run python tasks.py debrief first.")
    news_line = interview_context.recent_company_news_line(
        interview_context.company_research_context(
            JOBS_DIR,
            company,
            role_title,
            company_research_path=COMPANY_RESEARCH,
        ),
        company,
    )

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {output_target_name} Post-Round Follow-Up and Round 3 Prep.docx"
    doc = Document()
    doc.add_heading(f"Post-Round Follow-Up and Next-Round Prep - {company}", level=1)
    doc.add_paragraph(f"Role: {role_title}")
    doc.add_heading("Follow-Up Email", level=2)
    doc.add_paragraph(post_round_followup_email(role_title, company, entry, news_line, job_description))
    doc.add_heading("What To Address Before The Next Round", level=2)
    for bullet in bullets_from_signals(entry):
        doc.add_paragraph(bullet, style="List Bullet")
    questions = business_context.post_round_business_questions(entry)
    if questions:
        doc.add_heading("Business-Facing Questions To Prepare", level=2)
        for item in questions:
            doc.add_paragraph(f"{item.question} Answer angle: {item.answer_angle} Ask back: {item.ask_back}", style="List Bullet")
    doc.add_heading("Closing Question For Next Round", level=2)
    doc.add_paragraph("What is the biggest gap between the experience you are seeing and what you actually need for this role? I would rather understand it now so I can address it directly.")
    visible_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    assert_no_template_leakage(visible_text)
    enforce_prose_quality(
        post_round_followup_email(role_title, company, entry, news_line, job_description),
        "post_round_email_body",
        label="Post-round follow-up body",
    )
    doc.save(output)
    print(f"Post-round follow-up and prep created: {output}")
    return output


def main() -> None:
    build_post_round()


if __name__ == "__main__":
    main()
