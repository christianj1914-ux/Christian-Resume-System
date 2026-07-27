#!/usr/bin/env python3
"""Build a read-only Word audit of question bank coverage and redundancy."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import _bootstrap

_bootstrap.ensure_script_path()

from docx import Document

import question_bank_audit
from config.paths import JOB_DESCRIPTION, OUTPUT_DIR
from utils import optional_text


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_table_rows(document: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_document(audit: question_bank_audit.QuestionBankAudit) -> Document:
    document = Document()
    document.add_heading("Question Bank Audit", level=1)
    document.add_paragraph(f"Date: {date.today().isoformat()}")
    document.add_paragraph("Read-only audit. This report groups and warns; it does not edit source question banks.")

    document.add_heading("Application Bank Health", level=2)
    add_bullet(document, f"Application prompts: {len(audit.rows)}")
    add_bullet(document, f"Unmapped generic_bridge prompts: {len(audit.unmapped_prompts)}")
    add_bullet(document, f"Category collisions: {len(audit.category_collisions)}")
    add_bullet(document, f"Internal duplicate defects: {len(audit.exact_duplicate_groups)}")
    add_bullet(document, f"Application near-duplicate hints: {len(audit.application_near_duplicates)}")

    document.add_heading("Category Collisions", level=3)
    if audit.category_collisions:
        for category, prompts in audit.category_collisions.items():
            document.add_heading(category, level=4)
            for prompt in prompts:
                add_bullet(document, prompt)
    else:
        document.add_paragraph("None.")

    document.add_heading("Unmapped Prompts", level=3)
    if audit.unmapped_prompts:
        for prompt in audit.unmapped_prompts:
            add_bullet(document, prompt)
    else:
        document.add_paragraph("None.")

    document.add_heading("Internal Duplicate Defects", level=3)
    if audit.exact_duplicate_groups:
        for group in audit.exact_duplicate_groups:
            add_bullet(document, " / ".join(group))
    else:
        document.add_paragraph("None.")

    document.add_heading("Application Near Duplicate Hints", level=3)
    if audit.application_near_duplicates:
        for first, second, score in audit.application_near_duplicates:
            add_bullet(document, f"{score:.2f}: {first} / {second}")
    else:
        document.add_paragraph("None.")

    document.add_heading("Application Bank Prompt Table", level=2)
    add_table_rows(
        document,
        ("Category", "Prompt", "Sources", "Study refs"),
        [
            (
                row.category,
                row.prompt,
                "; ".join(row.sources),
                "; ".join(row.theme_track_refs),
            )
            for row in audit.rows
        ],
    )

    document.add_heading("Interview Question Corpus", level=2)
    document.add_paragraph(
        "Reference material only. generic_bridge is expected here and is not an application-bank defect."
    )
    add_bullet(document, f"Interview corpus questions: {len(audit.interview_corpus_rows)}")
    add_bullet(document, f"Interview near-duplicate pairs: {len(audit.interview_near_duplicates)}")
    if audit.interview_near_duplicates:
        for first, second, score in audit.interview_near_duplicates:
            add_bullet(document, f"{score:.2f}: {first} / {second}")
    else:
        document.add_paragraph("No near-duplicate hints.")
    return document


def main() -> None:
    job_description = optional_text(JOB_DESCRIPTION)
    audit = question_bank_audit.build_audit(job_description=job_description)
    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Question_Bank_Audit_{date.today().isoformat()}.docx"
    build_document(audit).save(output)
    print(f"Output DOCX: {output}")
    print(
        "Application bank: "
        f"{len(audit.rows)} prompts, "
        f"{len(audit.unmapped_prompts)} unmapped, "
        f"{len(audit.category_collisions)} category collisions, "
        f"{len(audit.exact_duplicate_groups)} internal duplicate defects."
    )
    print(
        "Interview question corpus (reference): "
        f"{len(audit.interview_corpus_rows)} questions, "
        f"{len(audit.interview_near_duplicates)} near-duplicate pairs."
    )


if __name__ == "__main__":
    main()
