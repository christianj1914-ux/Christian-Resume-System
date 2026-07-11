#!/usr/bin/env python3
"""Build a salary negotiation preparation guide for the active target role."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document

import job_search_guidance as guidance
import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"


def read_job() -> str:
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip() if JOB_DESCRIPTION.exists() else ""


def bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def compensation_scope(job_description: str) -> list[str]:
    scope = []
    if "$" in job_description or "salary" in job_description.lower() or "compensation" in job_description.lower():
        scope.append("Capture the posted salary or compensation range exactly as written before any negotiation call.")
    if "remote" in job_description.lower():
        scope.append("Compare remote national ranges against Atlanta/local ranges so the anchor is not undercut by geography.")
    if any(term in job_description.lower() for term in ("bonus", "commission", "equity", "stock", "incentive")):
        scope.append("Separate base salary, variable pay, equity, and bonus opportunity before judging the full offer.")
    if not scope:
        scope.append("Ask for the role's approved compensation range early enough to avoid negotiating in the dark.")
    return scope


def evidence_inventory(profile: resume_analysis.JobProblemProfile) -> list[str]:
    evidence = [
        "80+ client engagements across implementation, customer success, reporting, and operational improvement work.",
        "200+ reporting tools and dashboards built to improve decision quality, visibility, and operating discipline.",
        "Enterprise software implementation, configuration, testing, go-live, post-go-live support, and adoption experience.",
        "ERP ownership across five sites and 150+ users, including data validation, migration support, and cutover coordination.",
        "High-risk account stabilization and customer-facing recovery work tied to renewal risk, adoption, and trust rebuilding.",
    ]
    if profile.primary_lane == "presales_solution":
        evidence.append("Pre-sales discovery, requirements definition, SOW support, and buyer-facing solution framing.")
    if profile.primary_lane == "analytics_operations":
        evidence.append("Analytics and operations background that connects reporting, workflow clarity, and business decisions.")
    if profile.primary_lane == "customer_success":
        evidence.append("Customer success positioning around account health, value realization, renewal risk, and expansion discovery.")
    return evidence


def walkaway_criteria_lines(job_description: str) -> list[str]:
    lines = [
        "Set the walkaway point before the process gets emotional. Late-stage momentum should not redefine the floor.",
        "Treat pay, title level, reporting line, location policy, travel load, and scope clarity as one decision package rather than judging salary in isolation.",
        "If the company cannot move on cash, decide in advance which non-cash items would still make the role worth taking: PTO, flexibility, review timing, title, or development budget.",
    ]
    if "travel" in job_description.lower():
        lines.append("Travel load belongs in the walkaway test. A strong salary is not a clean trade if the travel rhythm is unsustainable.")
    if "remote" in job_description.lower() or "hybrid" in job_description.lower():
        lines.append("Work-location policy belongs in the walkaway test. Confirm state restrictions, onsite rhythm, and schedule expectations before the offer stage.")
    return lines


def build_salary_guide() -> Path:
    job_description = read_job()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")

    company = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role = resume_analysis.extract_job_title(job_description) or "Target Role"
    profile = resume_analysis.job_problem_profile(job_description)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {output_target_name} Salary Negotiation Guide.docx"
    doc = Document()
    doc.add_heading(f"Salary Negotiation Guide - {company}", level=1)
    doc.add_paragraph(f"Role: {role}")
    doc.add_paragraph(f"Prepared: {date.today().isoformat()}")
    doc.add_paragraph(f"Positioning lane: {profile.lane_label}")

    doc.add_heading("Market Research Framework", level=2)
    bullet_list(
        doc,
        [
            "Check posted pay transparency first: the job posting, company careers page, recruiter email, and any state-specific range language.",
            "Compare at least three public sources: Levels.fyi, Glassdoor, LinkedIn salary insights, RepVue for revenue-adjacent roles, and Built In or Wellfound when relevant.",
            "Use Atlanta, remote US, and role-specific market cuts separately instead of averaging them into one vague number.",
            "Use the closest Bureau of Labor Statistics OEWS occupation only as a baseline sanity check, not as the final target.",
            "Log source, date checked, title match, geography, base range, bonus/variable pay, and confidence level.",
        ],
    )

    doc.add_heading("Compensation Scope", level=2)
    bullet_list(doc, compensation_scope(job_description))
    bullet_list(
        doc,
        [
            "Define three numbers before the recruiter call: target, acceptable floor, and walkaway point.",
            "Keep base salary separate from benefits, bonus, commission, equity, PTO, remote flexibility, and professional-development budget.",
            "Treat title, scope, reporting line, and promotion path as part of the offer if cash flexibility is limited.",
        ],
    )

    doc.add_heading("Recruiter / Range-First Framing", level=2)
    bullet_list(
        doc,
        [
            "Ask for the approved range early enough to avoid negotiating against an unknown budget.",
            "Use the recruiter call to confirm level, location policy, and bonus or equity structure before naming a final number.",
            "If asked first, keep the answer collaborative: ask for the range, reference market research, and avoid anchoring against yourself too early.",
        ],
    )

    doc.add_heading("Walkaway Criteria", level=2)
    bullet_list(doc, walkaway_criteria_lines(job_description))

    doc.add_heading("Evidence Inventory", level=2)
    bullet_list(doc, evidence_inventory(profile))

    doc.add_heading("Opening Script", level=2)
    doc.add_paragraph(
        f"I am excited about the {role} opportunity because it connects directly to the kind of implementation, customer adoption, and workflow-improvement work I have already done. "
        "Before we get too far, could you share the approved compensation range for the role, including base salary and any bonus or variable component?"
    )

    doc.add_heading("If Asked for Expectations First", level=2)
    doc.add_paragraph(
        "I would like to understand the full scope and the company's range before naming a final number. Based on the role responsibilities, market research, and my background in implementation, customer-facing recovery, analytics, and ERP/process ownership, I would expect the offer to be competitive for this level. What range has been approved?"
    )

    doc.add_heading("Counteroffer Script", level=2)
    doc.add_paragraph(
        "Thank you for the offer. I am genuinely interested in the role and the problems the team is solving. Based on the scope, the market data I gathered, and the direct experience I bring in implementation delivery, customer adoption, analytics, and cross-functional operating improvement, I was hoping to be closer to [target number]. Is there room to improve the base salary or total package?"
    )

    doc.add_heading("Counteroffer Evaluation Checklist", level=2)
    bullet_list(doc, guidance.counteroffer_evaluation_lines())

    doc.add_heading("Risk Controls", level=2)
    bullet_list(
        doc,
        [
            "Do not apologize for negotiating; keep the tone collaborative and evidence-based.",
            "Do not cite unsupported salary numbers. Keep a research log so every anchor can be explained.",
            "Do not inflate direct quota, NRR, GRR, or closed expansion ownership; use supported commercial-adjacent language instead.",
            "Do not negotiate only base pay if the company has fixed bands; ask about bonus, sign-on, title, review timing, PTO, remote flexibility, and development budget.",
            "Ask for written terms before making a final decision.",
        ],
    )

    doc.add_heading("Decision Worksheet", level=2)
    for label in (
        "Target compensation:",
        "Acceptable floor:",
        "Walkaway point:",
        "Strongest evidence to cite:",
        "Non-salary items to request:",
        "Questions still unanswered:",
    ):
        doc.add_paragraph(label)

    doc.save(output)
    print(f"Salary negotiation guide created: {output}")
    return output


def main() -> None:
    build_salary_guide()


if __name__ == "__main__":
    main()
