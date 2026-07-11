#!/usr/bin/env python3
"""Generate a job-search analytics report from the application tracker."""

from __future__ import annotations

import csv
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
import track_applications


PROJECT_ROOT = Path(__file__).resolve().parents[1]
APPLICATIONS_CSV = PROJECT_ROOT / "scratch" / "applications.csv"
OUTPUT_DIR = PROJECT_ROOT / "output"
ADVANCED_STATUSES = {"phone_screen", "interview", "final_round", "offer"}


def read_rows() -> list[dict[str, str]]:
    if not APPLICATIONS_CSV.exists():
        return []
    with APPLICATIONS_CSV.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def percent(numerator: int, denominator: int) -> str:
    if denominator == 0:
        return "0%"
    return f"{round((numerator / denominator) * 100)}%"


def build_report() -> Path:
    rows = read_rows()
    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - Job Search Analytics {datetime.now().strftime('%Y-%m-%d')}.docx"
    doc = Document()
    doc.add_heading("Christian Estrada - Job Search Analytics", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")

    status_counts = Counter(row.get("current_status", "draft") or "draft" for row in rows)
    fit_counts = Counter(
        track_applications.tracker_fit_status(row) or "Not classified"
        for row in rows
    )
    audit_counts = Counter(
        track_applications.tracker_audit_flag(row)
        for row in rows
        if track_applications.tracker_audit_flag(row)
    )
    lane_counts = Counter(
        track_applications.tracker_lane_label(row) or "Unknown"
        for row in rows
    )

    doc.add_heading("Pipeline Overview", level=2)
    doc.add_paragraph(f"Total applications: {len(rows)}")
    doc.add_paragraph("Status breakdown: " + (", ".join(f"{key}: {value}" for key, value in sorted(status_counts.items())) or "No applications tracked."))
    doc.add_paragraph("Fit breakdown: " + (", ".join(f"{key}: {value}" for key, value in sorted(fit_counts.items())) or "No fit classifications tracked."))
    if audit_counts:
        doc.add_paragraph("Flagged outputs: " + ", ".join(f"{key}: {value}" for key, value in sorted(audit_counts.items())))
    doc.add_paragraph("Lane breakdown: " + (", ".join(f"{key}: {value}" for key, value in sorted(lane_counts.items())) or "No lane data tracked."))
    advanced = sum(1 for row in rows if row.get("current_status") in ADVANCED_STATUSES)
    doc.add_paragraph(f"Advanced past initial application: {advanced} ({percent(advanced, len(rows))})")

    doc.add_heading("Lane Performance", level=2)
    lanes: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        lanes[track_applications.tracker_lane_label(row) or "Unknown"].append(row)
    if not lanes:
        doc.add_paragraph("No lane data yet. Add applications with python tasks.py track.")
    for lane, lane_rows in sorted(lanes.items()):
        lane_advanced = sum(1 for row in lane_rows if row.get("current_status") in ADVANCED_STATUSES)
        doc.add_paragraph(f"{lane}: {len(lane_rows)} application(s), {lane_advanced} advanced, {percent(lane_advanced, len(lane_rows))} advancement rate.")

    doc.add_heading("Current Signals", level=2)
    if not rows:
        doc.add_paragraph("The tracker is empty, so no search pattern is available yet.")
    elif advanced == 0 and len(rows) >= 5:
        doc.add_paragraph("No tracked application has advanced yet. Review target lane, seniority, and top-third resume proof before increasing volume.")
    elif advanced:
        doc.add_paragraph("At least one tracked application has advanced. Compare that target's lane, company type, and resume proof against applications that stayed in applied status.")
    else:
        doc.add_paragraph("Keep tracking applications until there is enough volume to identify reliable patterns.")

    doc.save(output)
    print(f"Job search analytics report created: {output}")
    return output


def main() -> None:
    build_report()


if __name__ == "__main__":
    main()
