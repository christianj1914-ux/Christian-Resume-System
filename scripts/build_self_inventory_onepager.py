#!/usr/bin/env python3
"""Build Christian's provisional self-inventory one-pager."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

import _bootstrap

_bootstrap.ensure_script_path()
_bootstrap.configure_fresh_pycache()

import interview_intelligence


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def build_self_inventory_onepager() -> Path:
    inventory = interview_intelligence.load_self_inventory()
    content = interview_intelligence.build_self_inventory_onepager_content(inventory)
    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - Self-Inventory One-Pager {datetime.now().strftime('%Y-%m-%d')}.docx"

    document = Document()
    document.add_heading(content["title"], level=1)
    document.add_paragraph(content["status"])

    document.add_heading("Three Strengths", level=2)
    for strength in content["strengths"]:
        add_bullet(document, f"{strength['name']}: {strength['line']} Proof: {strength['proof']}.")

    document.add_heading("Three Development Areas", level=2)
    for weakness in content["weaknesses"]:
        add_bullet(document, f"{weakness['line']} Improvement: {weakness['improvement']} Status: {weakness['status']}.")

    document.add_heading("Motivation", level=2)
    document.add_paragraph(content["motivation"])

    document.add_heading("Five Signature Stories", level=2)
    for story in content["signature_stories"]:
        add_bullet(document, f"{story['name']}: {story['summary']}")

    document.add_heading("Spoken Strengths Answer", level=2)
    document.add_paragraph(content["strengths_answer"])

    document.add_heading("Spoken Weaknesses Answer", level=2)
    document.add_paragraph(content["weaknesses_answer"])

    rendered_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    interview_intelligence.assert_safe_generated_text(rendered_text, inventory)
    document.save(output)
    print(f"Self-inventory one-pager created: {output}")
    print()
    print("3 strengths answer:")
    print(content["strengths_answer"])
    print()
    print("3 weaknesses answer:")
    print(content["weaknesses_answer"])
    return output


def main() -> None:
    build_self_inventory_onepager()


if __name__ == "__main__":
    main()
