#!/usr/bin/env python3
"""Analyze recurring target-job skills not strongly visible in source resumes."""

from __future__ import annotations

from collections import Counter
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document

import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
TARGET_JDS = PROJECT_ROOT / "scratch" / "target_jds"
OUTPUT_DIR = PROJECT_ROOT / "output"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WORD_NS}}}"


def ensure_target_dir() -> None:
    TARGET_JDS.mkdir(parents=True, exist_ok=True)
    readme = TARGET_JDS / "README.txt"
    if not readme.exists():
        readme.write_text(
            "Place additional target job descriptions as .txt files in this folder for broader skills-gap analysis.\n",
            encoding="utf-8",
        )


def docx_visible_text(path: Path) -> str:
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    return "\n".join(
        "".join(node.text or "" for node in paragraph.findall(f".//{W}t"))
        for paragraph in root.findall(f".//{W}p")
    )


def job_descriptions() -> list[tuple[str, str]]:
    ensure_target_dir()
    items: list[tuple[str, str]] = []
    if JOB_DESCRIPTION.exists() and JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip():
        items.append(("current", JOB_DESCRIPTION.read_text(encoding="utf-8-sig")))
    for path in sorted(TARGET_JDS.glob("*.txt")):
        if path.name.lower() == "readme.txt":
            continue
        text = path.read_text(encoding="utf-8-sig").strip()
        if text:
            items.append((path.stem, text))
    return items


def build_skills_gap() -> Path:
    descriptions = job_descriptions()
    if not descriptions:
        raise SystemExit("No job descriptions found. Add jobs/job_description.txt or .txt files in scratch/target_jds.")
    resume_texts = "\n".join(docx_visible_text(path) for path in (resume_analysis.IMPLEMENTATION_RESUME, resume_analysis.PRESALES_CSM_RESUME))
    keyword_counts: Counter[str] = Counter()
    for _, text in descriptions:
        keyword_counts.update(resume_analysis.audit_keywords(text))

    gaps = []
    covered = []
    for keyword, count in keyword_counts.most_common():
        if resume_analysis.text_mentions(resume_texts, keyword):
            covered.append((keyword, count))
        else:
            gaps.append((keyword, count))

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - Skills Gap Analysis {datetime.now().strftime('%Y-%m-%d')}.docx"
    doc = Document()
    doc.add_heading("Future-Proofing Skills Gap Analysis", level=1)
    doc.add_paragraph(f"Job descriptions analyzed: {len(descriptions)}")

    doc.add_heading("Potential Gaps", level=2)
    for keyword, count in gaps[:25]:
        doc.add_paragraph(f"{keyword} ({count} posting signal(s))", style="List Bullet")
    if not gaps:
        doc.add_paragraph("No major keyword gaps detected against the current source resumes.")

    doc.add_heading("Already Supported Signals", level=2)
    for keyword, count in covered[:25]:
        doc.add_paragraph(f"{keyword} ({count} posting signal(s))", style="List Bullet")

    doc.add_heading("How to Use This", level=2)
    doc.add_paragraph("Treat gaps as professional-development prompts, not resume claims. Add a term to resumes only after Christian has verified support from source experience, current professional use, or a clearly defensible short-learning-curve bridge.")
    doc.save(output)
    print(f"Skills gap analysis created: {output}")
    return output


def main() -> None:
    build_skills_gap()


if __name__ == "__main__":
    main()
