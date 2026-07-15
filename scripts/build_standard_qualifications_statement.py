#!/usr/bin/env python3
"""Build a commercial qualifications statement for job-specific application questions."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

import _bootstrap

_bootstrap.ensure_script_path()

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import build_interview_cheat_sheet
import build_resume
import question_prep
import prose_engine
import render_checks
import resume_analysis
from build_skills_database import ResumeSnapshot
from config.paths import APPLICATION_QUESTIONS, JOB_DESCRIPTION, OUTPUT_DIR
from utils import read_text


BODY_FONT_SIZE = 11
QUESTION_FONT_SIZE = 13
NAME_FONT_SIZE = 14
TITLE_FONT_SIZE = 16
TITLE_BLUE = RGBColor(46, 116, 181)
QualificationsResponse = question_prep.QualificationsResponse
normalize_spaces = question_prep.normalize_spaces
parse_question_blocks = question_prep.parse_question_blocks
selected_resume_snapshot = question_prep.selected_resume_snapshot
build_question_responses = question_prep.build_question_responses
MONTH_NAME_TO_NUMBER = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}


@dataclass(frozen=True)
class QualificationsBuildResult:
    output_docx: Path
    company_name: str
    question_count: int
    used_custom_questions: bool


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Christian Estrada's commercial qualifications statement.")
    parser.add_argument("--no-pdf", action="store_true", help="Accepted for clarity; PDFs are never created.")
    return parser.parse_args()


def parse_month_year(text: str) -> tuple[int, int]:
    normalized = normalize_spaces(text)
    if normalized.lower() == "present":
        today = date.today()
        return today.year, today.month
    month_name, year_text = normalized.rsplit(" ", 1)
    month = MONTH_NAME_TO_NUMBER[month_name.lower()]
    return int(year_text), month


def role_month_count(snapshot: ResumeSnapshot) -> int:
    covered_months: set[int] = set()
    for role in snapshot.roles:
        if " to " not in role.dates:
            continue
        start_text, end_text = role.dates.split(" to ", 1)
        start_year, start_month = parse_month_year(start_text)
        end_year, end_month = parse_month_year(end_text)
        start_index = start_year * 12 + start_month - 1
        end_index = end_year * 12 + end_month - 1
        for month_index in range(start_index, end_index + 1):
            covered_months.add(month_index)
    return len(covered_months)


def direct_experience_years(snapshot: ResumeSnapshot) -> int:
    return max(1, role_month_count(snapshot) // 12)


def direct_experience_years_label(snapshot: ResumeSnapshot) -> str:
    months = role_month_count(snapshot)
    years = max(1, months // 12)
    return f"{years}+ years" if months % 12 else f"{years} years"


def education_years_answer(snapshot: ResumeSnapshot) -> str:
    education_lines = " ".join(snapshot.education).lower()
    estimated_years = 0
    if "bachelor" in education_lines:
        estimated_years += 4
    if "master" in education_lines:
        estimated_years += 2
    if estimated_years >= 6:
        years_text = "6+ years"
    elif estimated_years >= 4:
        years_text = "4+ years"
    else:
        years_text = "Related post-secondary education"
    development = ", ".join(snapshot.professional_development)
    return normalize_spaces(
        f"{years_text} of related post-secondary education, supported by bachelor's- and master's-level study in information systems, plus ongoing professional development through {development}."
    )


def direct_experience_years_answer(snapshot: ResumeSnapshot) -> str:
    return f"Approximately {direct_experience_years_label(snapshot)} of direct relevant experience."


def relevant_experience_answer(job_description: str, snapshot: ResumeSnapshot, resume_text: str) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    aptean_sentence = (
        "At Aptean, I managed 80+ international client engagements, led discovery, requirements definition, data migration, testing, training, and post-go-live support, and translated complex implementations into Statements of Work, Functional Requirements Documents, executive workshops, and measurable adoption plans."
    )
    if profile.primary_lane == "implementation_delivery":
        aptean_sentence = (
            "At Aptean, I delivered 12 full-lifecycle ERP implementations and managed up to four at a time across 80+ international client engagements, leading discovery, requirements definition, data migration, testing, training, and post-go-live support."
        )
    parts = [
        f"I bring approximately {direct_experience_years_label(snapshot)} of related experience across enterprise systems, analytics, customer-facing solution delivery, and cross-functional operational execution.",
        "At East West Manufacturing, I owned a mission-critical enterprise platform across five sites and 150+ users, supported migration planning and cutover work, built 200+ SQL-based KPI dashboards and reporting tools, and led training, testing, and go-live readiness across operations, finance, and engineering.",
        aptean_sentence,
    ]
    if re.search(r"\b(?:crm|salesforce)\b", job_description, re.I):
        parts.append("That experience is reinforced by Salesforce-backed reporting, case visibility, and backlog-ready requirements work at both Aptean and The Home Depot.")
    if re.search(r"\b(?:competitor|competitive)\b", job_description, re.I):
        parts.append("I also bring adjacent competitive-positioning experience through pre-sales discovery, product demonstrations, solution-fit conversations, and buyer-facing tradeoff discussions.")
    if re.search(r"\b(?:pilot|prototype|innovation|business case|road map|roadmap)\b", job_description, re.I):
        parts.append("The strongest bridge to innovation work is that I have helped move ambiguous ideas into usable pilots and operational rollouts, including a zero-to-one internal SMS support channel and a new warehouse and robotics ERP launch.")
    parts.append(
        f"Overall, the strongest match to this role is experience in {profile.core_problem}, backed by reporting, solution development, stakeholder facilitation, and practical follow-through."
    )
    return normalize_spaces(" ".join(parts))


def public_agency_experience_answer() -> str:
    return normalize_spaces(
        "0 years of direct public-agency or cooperative experience. My background is primarily private-sector, but it has centered on structured project delivery, stakeholder coordination, reporting, training, and solution development in complex multi-stakeholder environments."
    )


def unique_qualifications_answer(job_description: str, resume_text: str) -> str:
    bridges = [
        "80+ international client engagements, 60+ executive workshops and QBRs, five-site systems ownership for 150+ users, and 200+ reporting tools",
        "measurable process improvement such as a 78% reduction in manual inventory work and $1M+ in at-risk account recovery",
    ]
    if re.search(r"\b(?:pilot|prototype|innovation|prove|scale)\b", job_description, re.I):
        bridges.append(
            "experience moving ambiguous work from early concept into usable rollout, including a zero-to-one internal SMS support pilot and a new warehouse and robotics ERP launch"
        )
    if re.search(r"\b(?:competitor|competitive)\b", job_description, re.I):
        bridges.append("pre-sales discovery, product demonstrations, and solution-fit conversations that helped shape competitive buying decisions")
    if re.search(r"\b(?:salesforce|crm|power bi|excel)\b", job_description, re.I):
        bridges.append("hands-on reporting and CRM depth across Salesforce, Excel Power Query, Power BI, SQL, and SAP Crystal Reports")
    return normalize_spaces(
        "What may uniquely qualify me is the combination of solution consulting, operational analytics, and execution discipline. "
        + "That includes "
        + ", ".join(bridges[:-1])
        + f", and {bridges[-1]}."
    )


def certifications_answer(snapshot: ResumeSnapshot) -> str:
    development_text = ", ".join(snapshot.professional_development)
    if re.search(r"\bitil 4 foundation\b", development_text, re.I):
        return normalize_spaces(
            "Relevant professional development includes ITIL 4 Foundation, ServiceNow NextGen training, and McKinsey Forward Learners."
        )
    return normalize_spaces(
        f"Relevant professional development listed in my current materials includes {development_text}."
    )


def generic_bridge_answer(job_description: str, resume_text: str) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    return normalize_spaces(
        f"My strongest supported bridge for this application is experience in {profile.core_problem}, backed by 80+ client engagements, 200+ reporting tools, 60+ executive workshops and QBRs, and multi-site enterprise systems ownership. I would answer this item with direct proof where available and honest bridge language where the requirement is adjacent rather than identical."
    )


def latest_resume_state(job_description: str) -> str:
    matches = build_resume.matching_output_files(OUTPUT_DIR, job_description, "Resume.docx")
    return resume_analysis.output_audit_state(matches[0] if matches else None)


def set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(BODY_FONT_SIZE)


def add_paragraph(
    document: Document,
    text: str,
    *,
    size: float = BODY_FONT_SIZE,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    centered: bool = False,
    space_after: float = 6,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def build_recent_interviewer_scripts(
    items: tuple[question_prep.InterviewQuestionPrep, ...],
    job_description: str,
    resume_text: str,
    company_name: str,
    role_title: str,
) -> tuple[tuple[str, str], ...]:
    """Render real, spoken-ready answers for the Recent Interview Questions section.

    item.answer_angle is coaching guidance (how to construct an answer), not an
    answer itself. The interview cheat sheet and detailed guide builders already
    resolve each item to an actual spoken script via interviewer_question_factual_script()
    for non-story categories, or a story-bank answer otherwise. This mirrors that
    logic so the qualifications statement shows the same ready-to-say text instead
    of bare instructions.
    """
    if not items:
        return ()
    profile = build_interview_cheat_sheet.adjusted_profile_for_role(
        build_resume.job_problem_profile(job_description, resume_text), role_title, job_description
    )
    stories = build_interview_cheat_sheet.supported_story_bank(resume_text)
    used_titles: set[str] = set()
    scripts: list[tuple[str, str]] = []
    for item in items:
        if item.category in question_prep.NON_STORY_INTERVIEWER_CATEGORIES:
            spoken = prose_engine.spoken_register(
                question_prep.interviewer_question_factual_script(item.prompt, job_description, resume_text)
            ).text
        else:
            story = build_interview_cheat_sheet.likely_question_story(
                build_interview_cheat_sheet.InterviewQuestion(item.prompt, item.answer_angle),
                list(stories),
                used_titles,
            )
            used_titles.add(story.title)
            spoken = build_interview_cheat_sheet.spoken_story_answer(
                story, profile, company_name, role_title, job_description
            )
        if "Fill in" in spoken:
            print(
                f"[WARN] federal recent-interviewer: placeholder detected for '{item.prompt[:60]}', "
                "falling back to coaching angle"
            )
            spoken = item.answer_angle
        scripts.append((item.prompt, spoken))
    return tuple(scripts)


def build_document(
    company_name: str,
    role_title: str,
    responses: tuple[question_prep.QualificationsResponse, ...],
    *,
    recent_interviewer_scripts: tuple[tuple[str, str], ...] = (),
    used_custom_questions: bool,
) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    set_normal_style(document)

    add_paragraph(document, "Christian Estrada", size=NAME_FONT_SIZE, bold=True, centered=True, space_after=2)
    add_paragraph(document, "Qualifications Statement", size=TITLE_FONT_SIZE, bold=True, centered=True, color=TITLE_BLUE, space_after=4)
    add_paragraph(document, f"Target Position: {company_name} - {role_title}", italic=True, centered=True, space_after=10)
    if used_custom_questions:
        add_paragraph(document, "Customized to the current application questions.", size=10, italic=True, centered=True, space_after=10)

    for response in responses:
        add_paragraph(document, response.prompt, size=QUESTION_FONT_SIZE, bold=True, color=TITLE_BLUE, space_after=4)
        for paragraph in re.split(r"\n{2,}", response.answer):
            cleaned = normalize_spaces(paragraph)
            if cleaned:
                add_paragraph(document, cleaned)
    if recent_interviewer_scripts:
        add_paragraph(
            document,
            "Recent Interview Questions To Be Ready For",
            size=QUESTION_FONT_SIZE,
            bold=True,
            color=TITLE_BLUE,
            space_after=4,
        )
        for prompt, spoken in recent_interviewer_scripts:
            add_paragraph(document, prompt, size=BODY_FONT_SIZE, bold=True, space_after=3)
            add_paragraph(document, spoken)
    return document


def build_standard_qualifications_statement() -> QualificationsBuildResult:
    job_description = build_resume.validate_inputs(read_text(JOB_DESCRIPTION))
    company_name = build_resume.extract_output_name(job_description)
    output_target_name = build_resume.extract_output_target_name(job_description)
    role_title = build_resume.extract_job_title(job_description) or "Target Role"
    _, snapshot, resume_text = question_prep.selected_resume_snapshot(job_description)

    prompt_state = question_prep.load_application_prompt_state(APPLICATION_QUESTIONS)
    prompts = prompt_state.effective_prompts
    used_custom_questions = not prompt_state.uses_default_questions
    responses = question_prep.build_question_responses(prompts, job_description, snapshot, resume_text)
    recent_interviewer_items = question_prep.recent_interviewer_question_prep_items(
        job_description,
        company_name,
        role_title,
    )
    with prose_engine.collect_spoken_repair_issues() as spoken_issues:
        recent_interviewer_scripts = build_recent_interviewer_scripts(
            recent_interviewer_items,
            job_description,
            resume_text,
            company_name,
            role_title,
        )
    for response in responses:
        if response.warning:
            print(response.warning)

    audit_state = latest_resume_state(job_description)
    audit_suffix = resume_analysis.output_audit_suffix(audit_state)
    question_issues = question_prep.application_question_context_issues(job_description, prompt_state, workflow="commercial")
    review_issues = tuple(dict.fromkeys((*question_issues, *spoken_issues)))
    draft_suffix = " DRAFT" if review_issues else ""
    output_docx = OUTPUT_DIR / f"Christian Estrada - {output_target_name}{audit_suffix}{draft_suffix} Qualifications Statement.docx"
    OUTPUT_DIR.mkdir(exist_ok=True)
    document = build_document(
        company_name,
        role_title,
        responses,
        recent_interviewer_scripts=recent_interviewer_scripts,
        used_custom_questions=used_custom_questions,
    )
    document.save(str(output_docx))
    if review_issues:
        question_prep.mark_docx_as_draft(output_docx, review_issues)
    render_checks.render_docx(output_docx)
    return QualificationsBuildResult(
        output_docx=output_docx,
        company_name=company_name,
        question_count=len(responses),
        used_custom_questions=used_custom_questions,
    )


def main() -> None:
    parse_args()
    result = build_standard_qualifications_statement()
    print(f"Company: {result.company_name}")
    print(f"Output DOCX: {result.output_docx}")
    print(f"Questions answered: {result.question_count}")
    print(
        "Question source: custom application questions"
        if result.used_custom_questions
        else "Question source: default role-interest question"
    )


if __name__ == "__main__":
    main()
