#!/usr/bin/env python3
"""Audit an interview story and suggest stronger framing."""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document

from utils import has_great_eight_signal


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Audit a candidate-provided interview story.")
    parser.add_argument("--sample", action="store_true", help="Use a sample story for testing.")
    return parser.parse_args()


def prompt_multiline(label: str) -> str:
    print(label)
    print("Paste the story. Leave a blank line when finished.")
    lines = []
    while True:
        try:
            line = input("> ")
        except (KeyboardInterrupt, EOFError):
            raise SystemExit("Story audit canceled.")
        if not line.strip():
            break
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


def sample_story() -> str:
    return (
        "A customer was frustrated after go-live because support ownership was fragmented. "
        "I noticed the issue was not only technical; the customer did not know who owned each next step. "
        "I created a single recovery path, aligned internal owners, and gave the customer clearer updates. "
        "The account stabilized and the team rebuilt trust."
    )


def present(label: str, text: str, patterns: tuple[str, ...]) -> str:
    return "Present" if any(re.search(pattern, text, re.I) for pattern in patterns) else "Needs work"


def audit_story(story: str) -> dict[str, str]:
    return {
        "Hook": present("Hook", story, (r"\b(problem|challenge|frustrated|risk|issue|unclear|stalled)\b",)),
        "Noticing": present("Noticing", story, (r"\b(?:noticed|realized|saw|learned|recognized|understood)\b",)),
        "Action": present("Action", story, (r"\b(?:created|built|aligned|led|changed|mapped|validated|coordinated)\b",)),
        "Result": "Present" if has_great_eight_signal(story) or re.search(r"\b(stabilized|reduced|improved|increased|saved|rebuilt|launched)\b", story, re.I) else "Needs work",
        "Bridge": present("Bridge", story, (r"\b(?:this matters|same pattern|that is why|connects to|for this role)\b",)),
        "Internal Monologue": present("Internal Monologue", story, (r"\b(?:noticed|decided|realized|because|so I|I chose)\b",)),
        "Show Don't Tell": "Needs work" if re.search(r"\b(strong|great|excellent|proactive|hardworking|strategic)\b", story, re.I) else "Present",
    }


def improvement_lines(audit: dict[str, str]) -> list[str]:
    lines = []
    for area, status in audit.items():
        if status == "Needs work":
            lines.append(f"Strengthen {area}: add a concrete sentence that shows what happened, what Christian noticed, what action he took, or what changed.")
    return lines or ["Story has the core structure. Tighten wording, add role-specific bridge language, and keep the result measurable."]


def build_story_audit(sample: bool = False) -> Path:
    story = sample_story() if sample else prompt_multiline("Story to audit")
    if not story:
        raise SystemExit("No story supplied.")
    audit = audit_story(story)
    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - Story Audit {datetime.now().strftime('%Y-%m-%d')}.docx"
    doc = Document()
    doc.add_heading("Story Quality Audit", level=1)
    doc.add_heading("Original Story", level=2)
    doc.add_paragraph(story)
    doc.add_heading("Assessment", level=2)
    for area, status in audit.items():
        doc.add_paragraph(f"{area}: {status}", style="List Bullet")
    doc.add_heading("Upgrade Suggestions", level=2)
    for line in improvement_lines(audit):
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Show-Don't-Tell Rewrite Direction", level=2)
    doc.add_paragraph("Replace adjective claims with observed behavior: what the customer said, what Christian noticed, what action he took, and what changed afterward.")
    doc.save(output)
    print(f"Story audit created: {output}")
    return output


def main() -> None:
    args = parse_args()
    build_story_audit(args.sample)


if __name__ == "__main__":
    main()
