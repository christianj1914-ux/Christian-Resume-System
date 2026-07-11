#!/usr/bin/env python3
"""Generate a tailored post-interview thank-you note as a Word document."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

import build_resume
import interview_context
import resume_analysis
import business_context
import proof_text
from utils import (
    assert_no_template_leakage,
    enforce_prose_quality,
    extract_single_discussion_topic,
    normalize_generated_value,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
JOBS_DIR = PROJECT_ROOT / "jobs"
OUTPUT_DIR = PROJECT_ROOT / "output"
JOB_DESCRIPTION = JOBS_DIR / "job_description.txt"
DEBRIEF_HISTORY = JOBS_DIR / "debrief_history.txt"
DEBRIEF_DELIMITER = "POST-INTERVIEW DEBRIEF CAPTURED"


def read_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8-sig")


def clean_filename_piece(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", "", value).strip()
    return re.sub(r"\s+", " ", cleaned)[:120] or "Company"


def find_latest_debrief_for_company(company_name: str, debrief_path: Path = DEBRIEF_HISTORY) -> str:
    if debrief_path == DEBRIEF_HISTORY:
        structured = interview_context.latest_debrief_entry(PROJECT_ROOT / "jobs", company_name)
        if structured:
            return structured
    text = read_text(debrief_path)
    if not text.strip() or not company_name.strip():
        return ""
    entries = [
        f"{DEBRIEF_DELIMITER}{entry}".strip()
        for entry in text.split(DEBRIEF_DELIMITER)
        if entry.strip()
    ]
    company_lower = company_name.lower()
    matches = [entry for entry in entries if company_lower in entry.lower()]
    return matches[-1] if matches else ""


def section_value(entry: str, label: str) -> str:
    match = re.search(rf"(?ims)^{re.escape(label)}:\s*(.*?)(?=\n[A-Z][A-Za-z ]+:\s|\Z)", entry)
    return re.sub(r"\s+", " ", match.group(1)).strip() if match else ""


def latest_matching_resume(job_description: str) -> Path | None:
    if not OUTPUT_DIR.exists():
        return None
    candidates = resume_analysis.matching_output_files(OUTPUT_DIR, job_description, "Resume.docx")
    return candidates[0] if candidates else None


def docx_paragraph_texts(path: Path) -> list[str]:
    document = Document(str(path))
    return [re.sub(r"\s+", " ", paragraph.text).strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def is_resume_contact_line(text: str) -> bool:
    lowered = text.lower()
    return bool(
        ("|" in text and ("@" in text or "linkedin.com" in lowered or re.search(r"\b\d{3}[-.) ]+\d{3}[-. ]+\d{4}\b", text)))
        or lowered.startswith(("atlanta, ga", "christian estrada"))
        or "linkedin.com/in/" in lowered
        or re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, re.I)
    )


def compact_proof_point(text: str, *, max_words: int = 22) -> str:
    return proof_text.sanitize_proof_sentence(
        text,
        max_words=max_words,
        artifact="thank_you_body",
    )


def proof_points_from_resume(resume_docx: Path | None) -> list[str]:
    if resume_docx is None:
        return []
    action_pattern = re.compile(
        r"\b(?:reduced|improved|built|created|stabilized|managed|owned|delivered|launched|enabled|led|converted|protected)\b",
        re.I,
    )
    date_range_pattern = re.compile(
        r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*-\s*"
        r"(?:Present|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b",
        re.I,
    )
    points: list[str] = []
    in_experience = False
    for text in docx_paragraph_texts(resume_docx):
        lowered = text.lower()
        if lowered == "professional experience":
            in_experience = True
            continue
        if not in_experience:
            continue
        if lowered in {"education", "professional development"}:
            break
        if is_resume_contact_line(text):
            continue
        if len(text.split()) < 8 or len(text.split()) > 35 or lowered in {"professional summary", "professional experience"}:
            continue
        if date_range_pattern.search(text):
            continue
        if action_pattern.search(text):
            points.append(text)
        if len(points) >= 8:
            break
    return proof_text.best_proof_sentences(
        points,
        max_sentences=3,
        max_words=22,
        artifact="thank_you_body",
    )


def lane_followup_fallback_sentence(primary_lane: str) -> str:
    sentences = {
        "presales_solution": (
            "The conversation also clarified how closely the role connects to discovery, solution positioning, and the kind of partner or customer adoption that follows a well-scoped engagement."
        ),
        "customer_success": (
            "The conversation also clarified how closely the role connects to adoption, retention, and the kind of customer confidence that drives renewal and expansion."
        ),
        "analytics_operations": (
            "The conversation also clarified how closely the role connects to operational reporting, data visibility, and the decisions those outputs enable for teams that need faster answers."
        ),
        "change_enablement": (
            "The conversation also clarified how closely the role connects to stakeholder alignment, behavioral adoption, and the discipline of making change sustainable rather than just announced."
        ),
        "corporate_strategy": (
            "The conversation also clarified how closely the role connects to problem framing, cross-functional coordination, and the quality of recommendations that teams can actually implement."
        ),
        "implementation_delivery": (
            "The conversation also clarified how closely the role connects to implementation discipline, stakeholder coordination, and the kind of delivery follow-through that protects go-live timelines."
        ),
        "process_improvement": (
            "The conversation also clarified how closely the role connects to operational rigor, process design, and the measurement discipline that separates sustained improvement from a one-time fix."
        ),
    }
    return sentences.get(
        primary_lane,
        "The conversation also clarified how closely the role connects to the kind of structured problem-solving and cross-functional delivery I have found most rewarding."
    )


def thank_you_body(company_name: str, role_title: str, debrief_entry: str, proof_points: list[str]) -> tuple[str, list[str]]:
    def lower_lead(text: str) -> str:
        if not text:
            return ""
        return text[0].lower() + text[1:] if len(text) > 1 and text[0].isupper() else text

    interviewer_language = section_value(debrief_entry, "Specific language the interviewer used about the role")
    if not interviewer_language:
        interviewer_language = section_value(debrief_entry, "Specific interviewer language about the role")
    followups = section_value(debrief_entry, "Stories that generated follow-up questions")
    intelligence = section_value(debrief_entry, "Insider company intelligence learned")
    job_description = read_text(JOB_DESCRIPTION)
    primary_lane = build_resume.job_problem_profile(job_description, "").primary_lane if job_description else "implementation_delivery"
    combined_context = f"{job_description}\n{debrief_entry}"
    context = business_context.extract_business_context(combined_context)
    concrete_detail = ""
    if context.technical_stack:
        concrete_detail = f"how the team uses {', '.join(context.technical_stack[:3])} to keep implementation and support work visible"
    elif context.product_or_service or context.customer_type:
        context_items = [
            item
            for item in (context.product_or_service, context.customer_type, context.industry)
            if normalize_generated_value(item)
            and normalize_generated_value(item).lower() not in {"professional services", "b2b customers or client accounts", "software platform"}
        ]
        if context_items:
            concrete_detail = "the business context around " + ", ".join(context_items)

    primary_topic = extract_single_discussion_topic(interviewer_language or intelligence)
    followup_topic = extract_single_discussion_topic(followups)
    proof_point = proof_points[0].strip() if proof_points else ""
    if proof_point and proof_point[-1] not in ".!?":
        proof_point += "."
    role_topic_sentence = (
        f"After our conversation, I kept thinking about {lower_lead(primary_topic)}."
        if primary_topic
        else ""
    )
    proof_sentence = (
        f"I left the conversation clearer on where I could create value quickly: {proof_point}"
        if proof_point
        else f"I left the conversation clearer on where my implementation, customer-facing delivery, and workflow-improvement background would transfer fastest for {company_name}."
    )
    fit_detail_sentence = (
        f"The detail around {concrete_detail} made the fit more concrete because it points to the kind of execution discipline the role needs day to day."
        if concrete_detail
        else ""
    )
    followup_sentence = (
        f"The follow-up around {lower_lead(followup_topic)} also stood out because it is the kind of detail that shapes adoption, service quality, and client confidence."
        if followup_topic
        else lane_followup_fallback_sentence(primary_lane)
    )

    paragraphs = [
        f"Thank you again for the conversation about the {role_title} role.",
        role_topic_sentence or "The conversation gave me a clearer picture of where the role creates value.",
        fit_detail_sentence,
        proof_sentence,
        followup_sentence,
        "I remain very interested in the opportunity and look forward to the next conversation.",
        "Best,\nChristian Estrada",
    ]
    notes = []
    if not debrief_entry:
        notes.append("No matching debrief entry was found; the note uses job-description and resume context only.")
    if not proof_points:
        notes.append("No matching generated resume was found; add a specific proof point before sending if needed.")
    body = "\n\n".join(paragraph for paragraph in paragraphs if paragraph)
    if "concrete implementation, analytics, and customer-adoption" in body:
        print("THANK-YOU WARNING: hardcoded fallback phrase detected. Check lane detection.")
    return body, notes


def build_thank_you() -> Path:
    job_description = read_text(JOB_DESCRIPTION).strip()
    if not job_description:
        raise SystemExit("jobs/job_description.txt is empty. Add the active job description first.")

    company_name = resume_analysis.extract_output_name(job_description)
    output_target_name = resume_analysis.extract_output_target_name(job_description)
    role_title = resume_analysis.extract_job_title(job_description) or "the role"
    debrief_entry = find_latest_debrief_for_company(company_name)
    resume_docx = latest_matching_resume(job_description)
    proof_points = proof_points_from_resume(resume_docx)
    body, notes = thank_you_body(company_name, role_title, debrief_entry, proof_points)

    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - {clean_filename_piece(output_target_name)} Thank-You Note.docx"

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.add_run(f"Thank-You Note - {company_name}").bold = True
    doc.add_paragraph(f"Role: {role_title}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Subject: Thank you - {role_title}")
    for paragraph in body.split("\n\n"):
        doc.add_paragraph(paragraph)
    if notes:
        doc.add_paragraph("Review Notes").runs[0].bold = True
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")

    visible_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    assert_no_template_leakage(visible_text)
    enforce_prose_quality(body, "thank_you_body", label="Thank-you note body")
    doc.save(output)
    print(f"Thank-you note created: {output}")
    return output


def main() -> None:
    build_thank_you()


if __name__ == "__main__":
    main()
