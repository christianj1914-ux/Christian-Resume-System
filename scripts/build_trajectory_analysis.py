#!/usr/bin/env python3
"""Capture and generate an Up-or-Out Trajectory Analysis."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SCRATCH_DIR = PROJECT_ROOT / "scratch"
HISTORY = SCRATCH_DIR / "trajectory_history.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"

FIELDS = (
    "target role in two years",
    "target role in five years",
    "preferred company size",
    "preferred working style",
    "skill to be known for",
    "career risk currently being taken",
    "what would cause leaving within 12 months",
    "work that energizes most",
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build an Up-or-Out Trajectory Analysis.")
    parser.add_argument("--sample", action="store_true", help="Generate a sample/template without interactive input.")
    return parser.parse_args()


def safe_input(prompt: str) -> str:
    try:
        return input(prompt)
    except (KeyboardInterrupt, EOFError):
        raise SystemExit("Trajectory capture canceled.")


def prompt_required(label: str) -> str:
    while True:
        value = safe_input(f"{label}: ").strip()
        if value:
            return value
        print(f"{label} is required.")


def collect(sample: bool = False) -> dict[str, str]:
    if sample:
        return {field: f"[{field}]" for field in FIELDS}
    return {field: prompt_required(field.title()) for field in FIELDS}


def append_history(data: dict[str, str]) -> None:
    SCRATCH_DIR.mkdir(exist_ok=True)
    lines = [f"TRAJECTORY CAPTURED {datetime.now().isoformat(timespec='seconds')}"]
    lines.extend(f"{field}: {value}" for field, value in data.items())
    with HISTORY.open("a", encoding="utf-8") as handle:
        if HISTORY.exists() and HISTORY.stat().st_size:
            handle.write("\n")
        handle.write("\n".join(lines) + "\n")


def build_trajectory(sample: bool = False) -> Path:
    data = collect(sample)
    append_history(data)
    OUTPUT_DIR.mkdir(exist_ok=True)
    suffix = " TEMPLATE" if sample else ""
    output = OUTPUT_DIR / f"Christian Estrada - Trajectory Analysis{suffix} {datetime.now().strftime('%Y-%m-%d')}.docx"
    doc = Document()
    doc.add_heading("Up-or-Out Trajectory Analysis", level=1)
    for field, value in data.items():
        doc.add_heading(field.title(), level=2)
        doc.add_paragraph(value)
    doc.add_heading("Positioning Implication", level=2)
    doc.add_paragraph("Use this analysis to decide whether each target role builds the desired next capability, increases future option value, and reduces the risk of drifting into work Christian does not want to be known for.")
    doc.save(output)
    print(f"Trajectory analysis created: {output}")
    return output


def main() -> None:
    args = parse_args()
    build_trajectory(args.sample)


if __name__ == "__main__":
    main()
