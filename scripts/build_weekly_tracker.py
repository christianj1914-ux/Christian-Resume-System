#!/usr/bin/env python3
"""Generate a weekly job-search cadence plan."""

from __future__ import annotations

import csv
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document

import resume_analysis


PROJECT_ROOT = Path(__file__).resolve().parents[1]
APPLICATIONS_CSV = PROJECT_ROOT / "scratch" / "applications.csv"
JOB_DESCRIPTION = PROJECT_ROOT / "jobs" / "job_description.txt"
OUTPUT_DIR = PROJECT_ROOT / "output"


def read_rows() -> list[dict[str, str]]:
    if not APPLICATIONS_CSV.exists():
        return []
    with APPLICATIONS_CSV.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def monday_of_week(today: date | None = None) -> date:
    today = today or date.today()
    return today - timedelta(days=today.weekday())


def read_job() -> str:
    return JOB_DESCRIPTION.read_text(encoding="utf-8-sig").strip() if JOB_DESCRIPTION.exists() else ""


def recommended_actions(rows: list[dict[str, str]], job_description: str) -> list[str]:
    actions: list[str] = []
    if len(rows) < 10:
        actions.append(f"Add {10 - len(rows)} more tracked applications to reach the 10-application baseline.")
    if any(row.get("current_status") in {"phone_screen", "interview", "final_round"} for row in rows):
        actions.append("Run tasks.py guide for the most advanced active interview opportunity.")
    silent = []
    today = date.today()
    for row in rows:
        if row.get("current_status") == "applied" and row.get("applied_date"):
            try:
                applied = datetime.strptime(row["applied_date"], "%Y-%m-%d").date()
            except ValueError:
                continue
            if today - applied >= timedelta(days=14):
                silent.append(row.get("company", "Unknown company"))
    if silent:
        actions.append(f"Run tasks.py followup for silent applications: {', '.join(silent[:3])}.")
    if job_description:
        company = resume_analysis.extract_output_name(job_description)
        role = resume_analysis.extract_job_title(job_description) or "active role"
        actions.append(f"Complete or update the application package for {company} - {role}.")
    return actions[:4] or ["Add an active job description, then run tasks.py morning to choose the next action."]


def build_weekly_tracker() -> Path:
    rows = read_rows()
    job_description = read_job()
    monday = monday_of_week()
    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Christian Estrada - Week of {monday.isoformat()} Job Search Plan.docx"
    doc = Document()
    doc.add_heading(f"Week of {monday.strftime('%B %d, %Y')} Job Search Plan", level=1)

    doc.add_heading("This Week's Targets", level=2)
    for action in recommended_actions(rows, job_description):
        doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("Daily Action Plan", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Day", "Primary Focus", "Done"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    plans = [
        ("Monday", "Set weekly targets, refresh tracker, choose priority applications.", ""),
        ("Tuesday", "Apply to new roles and tailor first-pass materials.", ""),
        ("Wednesday", "Network or follow up on silent applications.", ""),
        ("Thursday", "Prepare interview stories, guides, or company research.", ""),
        ("Friday", "Review tracker, capture debriefs, and set next week's carryovers.", ""),
    ]
    for plan in plans:
        row = table.add_row().cells
        for index, value in enumerate(plan):
            row[index].text = value

    doc.add_heading("Pipeline Snapshot", level=2)
    doc.add_paragraph(f"Tracked applications: {len(rows)}")
    active = [row for row in rows if row.get("current_status") not in {"rejected", "withdrawn"}]
    doc.add_paragraph(f"Active tracked opportunities: {len(active)}")

    doc.save(output)
    print(f"Weekly tracker created: {output}")
    return output


def main() -> None:
    build_weekly_tracker()


if __name__ == "__main__":
    main()
