#!/usr/bin/env python3
"""Compare two tailored resume DOCX files section by section."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document

from resume_format import REQUIRED_SECTIONS, W, normalize_required_section_name


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Compare two generated resume DOCX files.")
    parser.add_argument("before", type=Path, help="Earlier or baseline resume DOCX.")
    parser.add_argument("after", type=Path, help="Later or comparison resume DOCX.")
    parser.add_argument("--format", choices=("text", "docx"), default="text", help="Output format. Default: text.")
    return parser.parse_args()


def paragraph_texts(docx_path: Path) -> list[str]:
    with ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    paragraphs = []
    for paragraph in root.findall(f".//{W}p"):
        text = "".join(node.text or "" for node in paragraph.findall(f".//{W}t"))
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def section_map(paragraphs: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {section: [] for section in REQUIRED_SECTIONS}
    current = ""
    for paragraph in paragraphs:
        normalized_section = normalize_required_section_name(paragraph)
        if normalized_section:
            current = normalized_section
            continue
        if current:
            sections[current].append(paragraph)
    return sections


def changed_lines(before: list[str], after: list[str]) -> tuple[list[str], list[str]]:
    before_set = set(before)
    after_set = set(after)
    removed = [line for line in before if line not in after_set]
    added = [line for line in after if line not in before_set]
    return removed, added


def build_comparison(before_path: Path, after_path: Path) -> dict[str, tuple[list[str], list[str]]]:
    before_sections = section_map(paragraph_texts(before_path))
    after_sections = section_map(paragraph_texts(after_path))
    return {
        section: changed_lines(before_sections.get(section, []), after_sections.get(section, []))
        for section in REQUIRED_SECTIONS
    }


def format_text(before_path: Path, after_path: Path, comparison: dict[str, tuple[list[str], list[str]]]) -> str:
    lines = [
        "Resume Comparison",
        f"Before: {before_path}",
        f"After: {after_path}",
        "",
    ]
    for section, (removed, added) in comparison.items():
        lines.append(section)
        lines.append("-" * len(section))
        if not removed and not added:
            lines.append("No visible text changes detected.")
        if removed:
            lines.append("Removed:")
            lines.extend(f"- {line}" for line in removed)
        if added:
            lines.append("Added:")
            lines.extend(f"+ {line}" for line in added)
        lines.append("")
    return "\n".join(lines).rstrip()


def write_docx(before_path: Path, after_path: Path, comparison: dict[str, tuple[list[str], list[str]]]) -> Path:
    OUTPUT_DIR.mkdir(exist_ok=True)
    output = OUTPUT_DIR / f"Resume Comparison - {before_path.stem} vs {after_path.stem}.docx"
    doc = Document()
    doc.add_heading("Resume Comparison", level=1)
    doc.add_paragraph(f"Before: {before_path}")
    doc.add_paragraph(f"After: {after_path}")
    for section, (removed, added) in comparison.items():
        doc.add_heading(section, level=2)
        if not removed and not added:
            doc.add_paragraph("No visible text changes detected.")
            continue
        if removed:
            doc.add_paragraph("Removed").runs[0].bold = True
            for line in removed:
                doc.add_paragraph(line, style="List Bullet")
        if added:
            doc.add_paragraph("Added").runs[0].bold = True
            for line in added:
                doc.add_paragraph(line, style="List Bullet")
    doc.save(output)
    return output


def main() -> None:
    args = parse_args()
    if not args.before.exists() or not args.after.exists():
        raise SystemExit("Both DOCX paths must exist.")
    comparison = build_comparison(args.before, args.after)
    if args.format == "docx":
        output = write_docx(args.before, args.after, comparison)
        print(f"Comparison document created: {output}")
    else:
        print(format_text(args.before, args.after, comparison))


if __name__ == "__main__":
    main()
