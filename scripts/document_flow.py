#!/usr/bin/env python3
"""Shared document-flow and visible submission-status helpers."""

from __future__ import annotations

import re
from collections.abc import Iterable
from xml.etree import ElementTree as ET

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WORD_NS}}}"
_SECTION_HEADINGS = {
    "professional summary",
    "professional experience",
    "education",
    "skills",
    "core competencies",
    "professional development",
}


def status_message(audit_status: str, reasons: Iterable[str] = ()) -> tuple[str, str, RGBColor]:
    """Return the visible label, detail, and color for a linked resume state."""
    normalized = (audit_status or "PASS").upper()
    detail = "; ".join(item.strip() for item in reasons if item and item.strip())
    if normalized in {"FAIL", "POOR"}:
        return (
            "NOT READY FOR SUBMISSION",
            detail or "The linked resume did not pass the evidence or sendability checks.",
            RGBColor(192, 0, 0),
        )
    if normalized in {"BRIDGE", "DRAFT"}:
        return (
            "REVIEW REQUIRED BEFORE SUBMISSION",
            detail or "The linked resume needs explicit, evidence-safe bridge review before submission.",
            RGBColor(192, 0, 0),
        )
    return ("", "", RGBColor(0, 0, 0))


def submission_readiness(audit_status: str, *, document_is_draft: bool = False) -> str:
    normalized = (audit_status or "PASS").upper()
    if normalized in {"FAIL", "POOR"}:
        return "blocked"
    if normalized in {"BRIDGE", "DRAFT"} or document_is_draft:
        return "review-required"
    return "ready"


def add_status_banner(document, audit_status: str, reasons: Iterable[str] = ()) -> bool:
    """Insert a body-visible warning before all candidate content when required."""
    label, detail, color = status_message(audit_status, reasons)
    if not label:
        return False
    banner = document.add_paragraph()
    if document.paragraphs:
        document._body._element.remove(banner._p)
        document._body._element.insert(0, banner._p)
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner.paragraph_format.space_before = Pt(0)
    banner.paragraph_format.space_after = Pt(2)
    banner.paragraph_format.keep_with_next = True
    headline = banner.add_run(label)
    headline.bold = True
    headline.font.color.rgb = color
    headline.font.size = Pt(11)
    detail_paragraph = document.add_paragraph()
    document._body._element.remove(detail_paragraph._p)
    document._body._element.insert(1, detail_paragraph._p)
    detail_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail_paragraph.paragraph_format.space_before = Pt(0)
    detail_paragraph.paragraph_format.space_after = Pt(6)
    detail_paragraph.paragraph_format.keep_with_next = True
    detail_run = detail_paragraph.add_run(detail)
    detail_run.italic = True
    detail_run.font.color.rgb = color
    detail_run.font.size = Pt(9)
    return True


def _is_docx_heading(paragraph) -> bool:
    text = re.sub(r"\s+", " ", paragraph.text).strip().lower()
    if text in _SECTION_HEADINGS:
        return True
    style_name = getattr(getattr(paragraph, "style", None), "name", "").lower()
    if "heading" in style_name:
        return True
    return bool(paragraph.runs and all(run.bold for run in paragraph.runs if run.text.strip()) and len(text) < 180)


def apply_docx_flow_controls(document) -> None:
    """Apply safe paragraph-level flow controls without locking long blocks together."""
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    for index, paragraph in enumerate(paragraphs):
        paragraph.paragraph_format.widow_control = True
        if _is_docx_heading(paragraph) and index + 1 < len(paragraphs):
            paragraph.paragraph_format.keep_with_next = True


def _ensure_ppr(paragraph: ET.Element) -> ET.Element:
    p_pr = paragraph.find(f"{W}pPr")
    if p_pr is None:
        p_pr = ET.Element(f"{W}pPr")
        paragraph.insert(0, p_pr)
    return p_pr


def _paragraph_text(paragraph: ET.Element) -> str:
    return re.sub(r"\s+", " ", "".join(paragraph.itertext())).strip()


def apply_resume_xml_flow_controls(root: ET.Element) -> None:
    """Apply the same safe flow rules to the direct-OOXML resume formatter."""
    paragraphs = root.findall(f".//{W}p")
    for index, paragraph in enumerate(paragraphs):
        p_pr = _ensure_ppr(paragraph)
        if p_pr.find(f"{W}widowControl") is None:
            ET.SubElement(p_pr, f"{W}widowControl")
        text = _paragraph_text(paragraph).lower()
        if text in _SECTION_HEADINGS and index + 1 < len(paragraphs):
            if p_pr.find(f"{W}keepNext") is None:
                ET.SubElement(p_pr, f"{W}keepNext")
