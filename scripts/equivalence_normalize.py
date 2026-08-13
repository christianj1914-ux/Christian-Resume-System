"""Canonical capture helpers for Release A output equivalence records."""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CORE_NS = {
    "dcterms": "http://purl.org/dc/terms/",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
}
VOLATILE_CORE_TAGS = {
    f"{{{CORE_NS['dcterms']}}}created",
    f"{{{CORE_NS['dcterms']}}}modified",
    f"{{{CORE_NS['cp']}}}revision",
}


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_text(value: str) -> str:
    return sha256_bytes(value.encode("utf-8"))


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def canonical_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def _strip_volatile_xml(root: ET.Element) -> None:
    for parent in root.iter():
        for child in list(parent):
            if child.tag in VOLATILE_CORE_TAGS:
                parent.remove(child)
        for attribute in list(parent.attrib):
            if attribute.startswith(f"{{{W_NS}}}rsid"):
                del parent.attrib[attribute]


def _relationship_map(archive: zipfile.ZipFile, part_name: str) -> dict[str, tuple[str, str, str]]:
    part_path = Path(part_name)
    rel_name = str(part_path.parent / "_rels" / f"{part_path.name}.rels").replace("\\", "/")
    if rel_name not in archive.namelist():
        return {}
    root = ET.fromstring(archive.read(rel_name))
    mapping: dict[str, tuple[str, str, str]] = {}
    for relationship in root.findall(f"{{{REL_NS}}}Relationship"):
        relationship_id = relationship.attrib.get("Id", "")
        mapping[relationship_id] = (
            relationship.attrib.get("Type", ""),
            relationship.attrib.get("Target", ""),
            relationship.attrib.get("TargetMode", "Internal"),
        )
    return mapping


def canonical_xml_part(archive: zipfile.ZipFile, part_name: str) -> str:
    if part_name not in archive.namelist():
        return ""
    root = ET.fromstring(archive.read(part_name))
    _strip_volatile_xml(root)
    relationships = _relationship_map(archive, part_name)
    for element in root.iter():
        for attribute, value in list(element.attrib.items()):
            if attribute == f"{{{R_NS}}}id" and value in relationships:
                element.attrib[attribute] = "rel:" + sha256_text(canonical_json(relationships[value]))[:16]
        if element.attrib:
            ordered = sorted(element.attrib.items())
            element.attrib.clear()
            element.attrib.update(ordered)
    return ET.tostring(root, encoding="unicode", short_empty_elements=True)


def relationship_records(archive: zipfile.ZipFile) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    for name in sorted(item for item in archive.namelist() if item.endswith(".rels")):
        root = ET.fromstring(archive.read(name))
        for relationship in root.findall(f"{{{REL_NS}}}Relationship"):
            records.append(
                {
                    "part": name,
                    "type": relationship.attrib.get("Type", ""),
                    "target": relationship.attrib.get("Target", ""),
                    "target_mode": relationship.attrib.get("TargetMode", "Internal"),
                }
            )
    return sorted(records, key=lambda item: canonical_json(item))


def _iter_block_text(document: Document) -> list[str]:
    body = document.element.body
    paragraphs_by_element = {paragraph._p: paragraph for paragraph in document.paragraphs}
    tables_by_element = {table._tbl: table for table in document.tables}
    lines: list[str] = []
    for child in body.iterchildren():
        paragraph = paragraphs_by_element.get(child)
        if paragraph is not None:
            lines.append(paragraph.text)
            continue
        table = tables_by_element.get(child)
        if table is not None:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
    return lines


def visible_text(docx_path: Path) -> str:
    return "\n".join(_iter_block_text(Document(str(docx_path))))


def _length(value: Any) -> int | None:
    return int(value) if value is not None else None


def formatting_snapshot(document: Document) -> dict[str, Any]:
    sections = []
    for section in document.sections:
        sections.append(
            {
                "page_width": _length(section.page_width),
                "page_height": _length(section.page_height),
                "top_margin": _length(section.top_margin),
                "right_margin": _length(section.right_margin),
                "bottom_margin": _length(section.bottom_margin),
                "left_margin": _length(section.left_margin),
                "header_distance": _length(section.header_distance),
                "footer_distance": _length(section.footer_distance),
                "gutter": _length(section.gutter),
                "orientation": str(section.orientation),
            }
        )
    paragraphs = []
    for paragraph in document.paragraphs:
        fmt = paragraph.paragraph_format
        runs = []
        for run in paragraph.runs:
            color = run.font.color.rgb
            runs.append(
                {
                    "text": run.text,
                    "name": run.font.name,
                    "size": _length(run.font.size),
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": str(run.underline),
                    "color": str(color) if color is not None else None,
                }
            )
        paragraphs.append(
            {
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": str(paragraph.alignment),
                "left_indent": _length(fmt.left_indent),
                "right_indent": _length(fmt.right_indent),
                "first_line_indent": _length(fmt.first_line_indent),
                "space_before": _length(fmt.space_before),
                "space_after": _length(fmt.space_after),
                "line_spacing": str(fmt.line_spacing),
                "keep_together": fmt.keep_together,
                "keep_with_next": fmt.keep_with_next,
                "page_break_before": fmt.page_break_before,
                "widow_control": fmt.widow_control,
                "runs": runs,
            }
        )
    tables = []
    for table in document.tables:
        tables.append(
            {
                "style": table.style.name if table.style else None,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": [[cell.text for cell in row.cells] for row in table.rows],
            }
        )
    return {"sections": sections, "paragraphs": paragraphs, "tables": tables}


def document_record(docx_path: Path) -> dict[str, Any]:
    docx_path = Path(docx_path)
    text = visible_text(docx_path)
    document = Document(str(docx_path))
    with zipfile.ZipFile(docx_path) as archive:
        parts = {
            part: canonical_xml_part(archive, part)
            for part in (
                "word/document.xml",
                "word/styles.xml",
                "word/settings.xml",
                "word/numbering.xml",
            )
        }
        relationships = relationship_records(archive)
    external_targets = sorted(
        item["target"] for item in relationships if item["target_mode"].lower() == "external"
    )
    return {
        "filename": docx_path.name,
        "raw_docx_sha256": file_sha256(docx_path),
        "visible_text": text,
        "visible_text_sha256": sha256_text(text),
        "xml": parts,
        "xml_sha256": {name: sha256_text(value) for name, value in parts.items()},
        "relationships": relationships,
        "external_hyperlink_targets": external_targets,
        "formatting": formatting_snapshot(document),
    }


def canonical_volatile_text(text: str, roots: tuple[Path, ...] = ()) -> str:
    """Normalize only approved transient text and remain idempotent."""
    value = text.replace("\r\n", "\n").replace("\\", "/")
    for root in sorted((path.resolve().as_posix() for path in roots), key=len, reverse=True):
        value = re.sub(re.escape(root), "<WORKSPACE>", value, flags=re.I)
    value = re.sub(r"(?i)^[A-Z]:/.*/scratch/e/r/[0-9a-f]+", "<WORKSPACE>", value)
    value = re.sub(
        r"\b20\d{2}-\d{2}-\d{2}(?:[T ]\d{2}:\d{2}:\d{2}(?:\.\d+)?(?:Z|[+-]\d{2}:\d{2})?)?\b",
        "<TIMESTAMP>",
        value,
    )
    value = re.sub(r"20\d{6}[_T]\d{6}(?:_\d{6})?", "<TIMESTAMP>", value)
    value = re.sub(
        r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December) "
        r"(?:[1-9]|[12]\d|3[01]), \d{4}\b",
        "<BUILD_DATE>",
        value,
    )
    value = re.sub(r"\b\d+(?:\.\d+)?s\b", "<DURATION>", value)
    value = re.sub(r"(?<![0-9a-f])[0-9a-f]{32}(?![0-9a-f])", "<RUN_ID>", value, flags=re.I)
    value = re.sub(r"(?<=_)[0-9a-f]{8}(?=(?:\s|/|\.|$))", "<RUN_ID>", value, flags=re.I)
    return value


def normalized_console(text: str, roots: tuple[Path, ...]) -> str:
    return canonical_volatile_text(text, roots)


def normalized_json_value(value: object, roots: tuple[Path, ...]) -> object:
    """Normalize decoded JSON values without corrupting JSON escape syntax."""
    if isinstance(value, dict):
        return {key: normalized_json_value(item, roots) for key, item in value.items()}
    if isinstance(value, list):
        return [normalized_json_value(item, roots) for item in value]
    if isinstance(value, str):
        return canonical_volatile_text(value, roots)
    return value
