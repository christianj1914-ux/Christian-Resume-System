"""State Farm Process Engineer interview guide playbook."""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

import build_interview_cheat_sheet as cheat
import build_resume
from utils import fail


FONT = "Carlito"
SECTION_BLUE = RGBColor(31, 78, 121)
SUBSECTION_BLUE = RGBColor(47, 84, 150)
SECTION_SIZE = 15
BODY_SIZE = 12
SMALL_SIZE = 10.5


@dataclass(frozen=True)
class PrepInsights:
    situation_read: tuple[str, ...]
    interviewer_read: tuple[str, ...]
    likely_scorecard: tuple[str, ...]
    pushbacks: tuple[tuple[str, str], ...]
    selling_angles: tuple[str, ...]
    anticipated_questions: tuple[tuple[str, str], ...]
    smart_questions: tuple[str, ...]
    brevity_rules: tuple[str, ...]


def text_has(text: str, *terms: str) -> bool:
    lowered = text.lower()
    return any(term.lower() in lowered for term in terms)


def add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(SECTION_SIZE)
    run.font.color.rgb = SECTION_BLUE


def add_subsection(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(BODY_SIZE)
    run.font.color.rgb = SUBSECTION_BLUE


def add_body(document: Document, text: str, *, small: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SMALL_SIZE if small else BODY_SIZE)


def add_bullet(document: Document, text: str, *, small: bool = False) -> None:
    paragraph = document.add_paragraph(style=None)
    paragraph.style = document.styles["List Bullet"]
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SMALL_SIZE if small else BODY_SIZE)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def is_state_farm_process_engineer_context(
    company_name: str,
    role_title: str,
    job_description: str,
    company_research: str = "",
    interview_notes: str = "",
) -> bool:
    combined = f"{company_name}\n{role_title}\n{job_description}\n{company_research}\n{interview_notes}".lower()
    return "state farm" in combined and text_has(
        combined,
        "process engineer",
        "claims process engineering",
        "lean six sigma",
        "process improvement",
        "process design",
    )

def state_farm_story_key(title: str) -> str:
    lowered = title.lower()
    if "inventory" in lowered:
        return "inventory"
    if "account" in lowered or "$1m" in lowered:
        return "account"
    if "dashboard" in lowered or "decision visibility" in lowered:
        return "dashboards"
    if "rapid" in lowered or "product learning" in lowered:
        return "learning"
    if "operations versus finance" in lowered or "finance alignment" in lowered:
        return "ops_finance"
    if "failure" in lowered or "validation" in lowered:
        return "failure"
    if "workshop" in lowered or "qbr" in lowered:
        return "workshops"
    return "default"

def state_farm_story_noticing(card: cheat.StoryCard) -> str:
    key = state_farm_story_key(card.title)
    return {
        "inventory": "What I noticed was that this was not a people problem. The errors were concentrated around a structural process gap: there was no reliable checkpoint between input and posting.",
        "account": "What I noticed was that the customer did not need another status update. They needed one accountable owner and a recovery path they could actually trust.",
        "dashboards": "What I noticed was that nobody was asking the question before requesting the report: what decision needs to be made, and what would change the answer? That missing question was why technically accurate reports were still operationally weak.",
        "learning": "What I noticed was that the fastest way to learn the domain was not documentation alone. It was treating each client problem as a live map of where the workflow broke down.",
        "ops_finance": "What I noticed was that both sides were protecting against different risks. Operations was protecting speed and flow; finance was protecting control and auditability. Naming that tradeoff changed the conversation.",
        "failure": "What I noticed afterward was that validation had happened too late in the sequence. The fix was not more effort; it was earlier SME input, clearer criteria, and checkpoints before the decision became expensive to reverse.",
        "workshops": "What I noticed was that executives, operators, and delivery teams did not need the same meeting. Each group needed the same goal translated into the decision, workflow, or risk they owned.",
    }.get(key, card.level3_trait)

def state_farm_story_bridge(card: cheat.StoryCard) -> str:
    key = state_farm_story_key(card.title)
    return {
        "inventory": "That pattern maps directly to claims: map what is actually happening rather than what the documentation says is happening, find the structural gap, validate the fix with the people doing the work, pilot it, and measure whether anything actually changed.",
        "account": "In claims, the same dynamic plays out at scale: when a customer's experience is breaking down, what they need first is one person who owns it and a credible path forward, not a better status cadence.",
        "dashboards": "For claims and task metrics, the discipline is identical: define the decision before touching the data. The question is not what cycle time shows; it is what decision this data needs to support and whether the segmentation makes that obvious.",
        "learning": "That is exactly how I would close the claims knowledge gap: learn through the workflow by being inside it, not through documentation. Every conversation with an adjuster is a case study in how the process actually operates versus how it is supposed to operate.",
        "ops_finance": "That same tension exists in claims process engineering every time a workflow touches speed, quality, compliance, and customer experience simultaneously. The job is to surface what each group is actually protecting before designing something that has to satisfy all of them.",
        "failure": "In claims, the stakes of that lesson are higher than in manufacturing: a process error can become a customer experience failure, a compliance exposure, or a financial risk. Quality control has to be designed into the process, not inspected in at the end.",
        "workshops": "For State Farm, that means translating the same process goal differently for claims leaders, execution partners, analysts, and frontline users so the change can actually stick.",
    }.get(key, "For State Farm, I would connect the result to the specific process, data, stakeholder, or customer outcome the team is trying to improve.")

def state_farm_calibration_question(card: cheat.StoryCard) -> str:
    key = state_farm_story_key(card.title)
    return {
        "inventory": "Is finding the structural root cause, rather than building better error detection, the kind of thinking this team is looking for?",
        "account": "Is that kind of cross-functional coordination without formal authority a regular part of how this role operates?",
        "dashboards": "Is defining the business decision before touching the data a discipline the team is still building, or is that infrastructure mature?",
        "learning": "Is a 30-to-60 day ramp where I learn the claims environment before recommending anything the right expectation for how someone starts in this role?",
        "ops_finance": "Is navigating that kind of competing-priority tension, where process and compliance teams each define success differently, a regular part of what this role manages?",
        "failure": "Is building SME validation into every process change before go-live a strong discipline here already, or an area to strengthen?",
        "workshops": "How often does this role need to translate the same process change differently for leaders, analysts, and frontline teams?",
    }.get(key, "Is that the right way to think about how this experience maps to the role?")

def is_state_farm_active(
    company_name: str,
    role_title: str,
    job_description: str,
    company_research: str = "",
    interview_notes: str = "",
) -> bool:
    return is_state_farm_process_engineer_context(company_name, role_title, job_description, company_research, interview_notes)


def state_farm_prep_insights() -> PrepInsights:
    return PrepInsights(
        situation_read=(
            "This is not a generic analyst interview. State Farm is testing whether Christian can improve claims operations at scale: identify waste, prove root cause with data, redesign the process, and bring claim/process partners along without sacrificing service, quality, or compliance.",
            "The process likely moves through on-demand video, a data exercise, and a live case-study conversation. Each step rewards the same behavior: lead with the answer, show the structure behind the recommendation, and keep the improvement tied to customer experience and operational reliability.",
            "Because the posting names Claims Process Engineering, Auto, Fire, Shared Services, tollgate materials, cost/benefit analysis, agile prioritization, and Lean Six Sigma, the winning narrative should sound like process engineering first and systems background second.",
            "The newer notes broaden the lens: process engineers are bridges between engineering principles and daily operations. For State Farm, translate that from plant operations into claims operations: workflow, bottlenecks, standard work, quality controls, risk, data signals, and customer impact.",
        ),
        interviewer_read=(),
        likely_scorecard=(
            "Can map current-state and future-state workflows, identify waste or variation, and turn findings into standard work that people can follow.",
            "Can use operational metrics, claims/task/phone data, Excel, PowerPoint, Visio-style process visuals, and cost/benefit logic to support decisions.",
            "Can influence stakeholders who do not report to him, especially when the process change creates discomfort or changes local routines.",
            "Can communicate clearly in tollgate updates, decision documents, case presentations, and short video responses.",
            "Can protect quality and accuracy through SME validation, pilots, defined metrics, and post-launch feedback loops.",
            "Can show continuous improvement habits: capture lessons learned during the project, apply feedback immediately, run retrospectives, and reinforce the business vision so teams do not lose the why.",
            "Can think like a project integrator: balance scope, time, cost, quality, risk, resources, and stakeholder adoption instead of optimizing only one metric.",
        ),
        pushbacks=(
            (
                "You do not have direct insurance claims experience.",
                "That is true, and I would not overstate it. The transferable pattern is the part I have done repeatedly: map the current process, find the root cause with data, validate the change with the people doing the work, and measure whether the new process actually improves quality, speed, or service. I would learn State Farm's claims language quickly by shadowing the workflow and grounding my recommendations in claim/task/phone metrics rather than assumptions.",
            ),
            (
                "Your background may look too systems-heavy for a Process Engineer role.",
                "The systems were usually the vehicle, not the point. The business work was process improvement: reducing manual work, improving data quality, training users, documenting standard work, and helping teams adopt a better way of operating. For State Farm, I would translate that into claims process design, operational metrics, and improvement routines.",
            ),
            (
                "Do you have formal Lean Six Sigma certification?",
                "I would answer this accurately and not overclaim. If asked, the strongest bridge is: I have applied Lean Six Sigma-style thinking through current-state mapping, root-cause analysis, waste reduction, measurable before/after metrics, pilot validation, and stakeholder adoption. I am also comfortable learning and applying State Farm's preferred methodology and terminology.",
            ),
            (
                "Can you handle ambiguity and multiple improvement priorities?",
                "Yes. My approach is to separate urgent noise from business impact: risk/customer impact, effort-to-value, stakeholder dependencies, and real deadlines. If a commitment is at risk, I surface it early with options rather than letting trust erode quietly.",
            ),
            (
                "You are not a traditional industrial process engineer with plant-floor P&ID or simulation-tool ownership.",
                "Correct. I would not claim direct ownership of PFDs, P&IDs, Aspen, HYSYS, MATLAB, or formal plant commissioning unless the work required it and I had done it. My fit is process engineering in an operations environment: mapping workflows, using data to find bottlenecks and root cause, improving standard work, building documentation, and helping stakeholders adopt the change. If State Farm uses specific process-engineering tools or templates, I would learn them quickly and anchor them to the claims workflow.",
            ),
        ),
        selling_angles=(
            "Lead with the 78% manual-work reduction and 22% discrepancy reduction story as the closest process-improvement proof: broken workflow, root cause, redesigned control, measurable result.",
            "Use the 200+ dashboard/reporting-tools story as data-to-decision proof: he does not just build reports; he starts with the operational decision and validates the source data.",
            "Use 60+ executive workshops and QBRs as stakeholder-influence proof: he can explain complex operational issues clearly and get decision-makers aligned.",
            "Use five-site and 150+ user ownership carefully: frame it as process adoption, training, documentation, and operating discipline, not as pure systems administration.",
            "Use the career CI story as a mindset proof: you are someone who seeks feedback early, documents what changed, and treats learning as part of the operating system rather than a once-a-year review event.",
        ),
        anticipated_questions=(
            (
                "Tell me about a time you identified and improved an inefficient process.",
                "The strongest example is an inventory adjustment workflow where manual work created delay and discrepancy risk. I would lead with the result: I reduced manual effort by 78% and discrepancies by 22%. Then I would explain the method: I mapped the current process, found where repeated manual touches were creating errors, built an automated and auditable workflow, validated it with the users, and measured before/after results. The lesson for State Farm is the same: process improvement has to start with current-state truth, not assumptions.",
            ),
            (
                "How do you prioritize multiple process improvement projects?",
                "I use a simple four-part filter: business impact, risk or compliance urgency, effort-to-value, and stakeholder dependencies. If two items look equal, I ask which one is blocking another team or creating customer impact. I also communicate tradeoffs early. The point is not to look busy; it is to protect the work that most improves service, quality, or operational reliability.",
            ),
            (
                "Tell me about a time you used data to drive a process decision.",
                "I would use the dashboard and KPI story. I built 200+ reporting tools, but the important part is that I started with the decision leaders needed to make, validated the source data, segmented the problem, and turned raw exports into clearer operating signals. For State Farm, I would apply that same approach to claim, task, phone, or workflow metrics: define the question, test the data, identify the lever, recommend the change, then monitor the KPI after rollout.",
            ),
            (
                "How do you get buy-in from stakeholders resistant to a process change?",
                "I do not start by pushing harder. I acknowledge the concern, ask clarifying questions, listen for the operational fear underneath the objection, share information they may not have, and then propose an alternative with a clear rationale. If possible, I pilot the change with an exit ramp. That gives people evidence and control instead of making the change feel imposed on them.",
            ),
            (
                "How do you ensure accuracy and quality when redesigning or documenting a process?",
                "My system is current-state documentation, SME validation, pilot testing, defined success metrics, and a feedback loop after rollout. I do not trust a process map just because it looks neat. I validate it with the people doing the work and build checkpoints so the improvement does not create a quality problem downstream.",
            ),
            (
                "Why State Farm and this Process Engineer role?",
                "Use the 50/50 rule. First: State Farm operates at a scale where process quality matters to real customers, not just internal efficiency. Claims processes affect speed, trust, service, quality, and the good-neighbor experience. Second: my fit is the operating pattern I bring: process diagnosis, data-backed recommendations, stakeholder alignment, documentation, training, and measurable adoption.",
            ),
            (
                "Tell me about continuous improvement in your own work.",
                "Use the proactive feedback loop: I do not wait until the formal review or project close to learn. I listen for feedback signals, document what I am hearing, test improvements quickly, and then check back with the stakeholder or manager to confirm whether the change is working. That is the same habit I use in process work: lessons learned during the project, not only after it ends.",
            ),
            (
                "How do project management principles show up in your process work?",
                "I think of the role as an integrator. A process change affects scope, time, cost, quality, resources, risk, and stakeholder adoption at the same time. Before acting, I review the plan or current standard, investigate the root cause, consult the people closest to the work, and then communicate the tradeoff and recommendation clearly.",
            ),
        ),
        smart_questions=(
            "What does this person need to accomplish in the first 12 months for you to consider the hire a success?",
            "What is the most challenging claims process problem your team is working on right now?",
            "Which process metrics does the team trust most today, and which ones still require better visibility or validation?",
            "How mature is the process improvement culture here: more bottom-up, top-down, or a mix of both?",
            "What tools and methodologies does the team use most often for process analysis, tollgate communication, and case-study recommendations?",
            "What behaviors separate the most effective people on this Claims Process Engineering team?",
            "Do you have any concerns about my ability to do this job that I can address directly?",
            "How does the team use lessons learned, retrospectives, or feedback loops while a process improvement is still live?",
            "Where do standard work, SOP-style documentation, or quality checkpoints most need reinforcement today?",
            "If the data exercise surfaces multiple possible root causes, how does the team expect candidates to frame uncertainty and next-step validation?",
        ),
        brevity_rules=(
            "Use the working story sequence for behavioral answers: Hook, Noticing, Action, Result, Bridge, Calibration. Lead with the result when the question is direct, then explain the method and State Farm bridge.",
            "For on-demand video, keep most answers to 60-90 seconds and avoid sounding scripted. Smile, pause, then answer like a calm operator.",
            "Use the Rule of Three when structuring any answer: business problem, action, result; or current state, root cause, recommendation.",
            "For the data exercise, narrate assumptions and tradeoffs. They are evaluating judgment, not just math.",
            "For the case-study presentation, make the slide logic executive-friendly: answer first, evidence second, recommendation third, risks and rollout last.",
        ),
    )


def add_state_farm_process_engineer_playbook(document: Document, profile: build_resume.JobProblemProfile) -> None:
    add_section(document, "State Farm Process Engineer Playbook")
    add_subsection(document, "What This Role Is Really About")
    add_body(
        document,
        "A Process Engineer in Claims Process Engineering is being evaluated as an operator: can Christian spot broken systems, prove the root cause, redesign the process, and bring people along so the change improves speed, service, quality, and reliability? The safest positioning is process-first, data-backed, and customer-aware."
    )
    add_subsection(document, "Five Qualities To Prove")
    qualities = (
        "Leadership / Initiative: he found the process issue instead of waiting for someone to define it for him.",
        "Resilience: he can recover when rollout, adoption, or data validation does not go perfectly.",
        "Teamwork: he can work across Claims, operations, analytics, leadership, and frontline users without direct authority.",
        "Influence / Persuasion: he can turn resistance into a pilot, evidence, and adoption rather than a debate.",
        "Ethics / Integrity: he will surface quality, compliance, or data-risk concerns even when doing so slows the timeline."
    )
    for quality in qualities:
        add_bullet(document, quality)
    add_subsection(document, "Interview Process Strategy")
    for line in (
        "On-demand video: answer in 60-90 seconds. Lead with the result, then give the structure. Do not sound like you are reading.",
        "Data exercise: state assumptions, define the metric, check data quality, segment the problem, identify likely root cause, and recommend a pilot with success measures.",
        "Live panel / case study: use Pyramid Principle. Recommendation first, evidence second, rollout and risks third. Make the slides executive-friendly.",
        "Across all stages: tie process improvement back to customer experience. For State Farm, operational excellence is part of the good-neighbor experience."
    ):
        add_bullet(document, line)

def add_state_farm_data_exercise_strategy(document: Document) -> None:
    add_section(document, "Data Exercise And Case Study Strategy")
    add_body(document, "Use this structure if the interview includes claims/task/phone/operational metrics or a case-study presentation.")
    steps = (
        "1. Clarify the business question: speed, quality, waste, cost, service, workload, or customer experience.",
        "2. Define the unit of analysis: claim, task, call, adjuster/team, channel, process step, or time period.",
        "3. Check data quality: missing values, duplicates, outliers, inconsistent definitions, and whether the metric is leading or lagging.",
        "4. Segment before concluding: compare by claim type, channel, geography, team, complexity, age, volume, or handoff point.",
        "5. Identify controllable levers: rework, queue time, duplicate touches, unclear ownership, training gaps, routing logic, or documentation gaps.",
        "6. Prioritize with impact/effort/risk: show why this improvement is worth doing before another one.",
        "7. Recommend a pilot: scope, owner, timeline, control group if possible, success metric, risk, and rollback/adjustment path.",
        "8. Close with what leadership should decide: approve pilot, collect more data, change standard work, or escalate a dependency."
    )
    for step in steps:
        add_bullet(document, step)
    add_subsection(document, "MECE Segmentation Check")
    add_body(document, "Before presenting segments, verify they are mutually exclusive and collectively exhaustive. A simple way to say it: 'I segmented by claim type - Auto simple, Auto multi-vehicle, and Disputed Liability - so no claim belongs to two categories at once and the full dataset is covered. I ruled out time-period segmentation because the pattern was consistent across all 13 weeks, which suggests a structural issue rather than seasonality.'")
    add_subsection(document, "What / So What / Now What")
    add_body(document, "Use this to lead each slide: What is the finding, so what does it mean for State Farm, now what should the team do next?")
    add_bullet(document, "What: multi-vehicle claims show a consistent 4-5 day routing delay.")
    add_bullet(document, "So What: that delay adds avoidable cycle time to hundreds of claims and puts the customer's first contact experience at risk.")
    add_bullet(document, "Now What: recommend a six-week routing pilot with a 24-hour service-level target and quality guardrails.")
    add_subsection(document, "Case Presentation Slide Logic")
    for line in (
        "Slide 1: answer first - the process problem, recommended action, and expected business impact.",
        "Slide 2: evidence - the data pattern, segmentation, root-cause hypothesis, and what you ruled out.",
        "Slide 3: implementation - pilot scope, stakeholder plan, communication/tollgate rhythm, risks, and success metrics."
    ):
        add_bullet(document, line)

def add_state_farm_answer_bank(document: Document) -> None:
    add_section(document, "State Farm Master Q&A Bank")
    add_subsection(document, "Start Here: Two Question Types")
    add_body(document, "Every question you will hear is one of two types: a past experience question ('Tell me about a time...') or a future approach question ('How would you handle...'). Past experience questions get the story sequence. Future approach questions get a position statement with three supporting points. Classify before answering.")
    answer_bank = (
        (
            "Tell me about a time you identified and improved an inefficient process.",
            "What they are assessing: initiative, process thinking, and measurable impact. Lead with the result: 'I reduced manual work by 78% and discrepancies by 22%.' Then explain that you mapped the inventory adjustment workflow, found repeated manual touchpoints, built an automated/auditable process, validated with users, and measured before/after. Bridge: 'I would use the same pattern in Claims Process Engineering: current state, root cause, pilot, metric, adoption.'"
        ),
        (
            "How do you prioritize competing process improvement projects?",
            "Use a calm, confident system: business/customer impact, risk or compliance urgency, effort-to-value, stakeholder dependencies, and real deadlines. Say that fast-moving environments energize you, but you do not rely on energy alone; you use a visible prioritization method and surface tradeoffs early."
        ),
        (
            "Tell me about stakeholder resistance.",
            "Use the five-step influence process: acknowledge their perspective, ask clarifying questions, listen to understand, share missing information, and propose an alternative with rationale. The best ending is not 'I won'; it is 'they trusted the pilot, adopted the change, and we could work together again.'"
        ),
        (
            "Influence without authority - $1M account version",
            "Emphasize method, not drama: I did not own every function involved, so I created single-point ownership, structured the sessions, named owners, and gave stakeholders a clean decision rhythm. What I noticed was that the account did not need more status updates; it needed accountable coordination. Result: $1M+ in account risk stabilized. Calibration: Is that kind of cross-functional coordination without formal authority a regular part of how this role operates?"
        ),
        (
            "Opposing stakeholder views - $1M account version",
            "Emphasize discovery: the stated disagreement was about tickets and timelines, but what I noticed was that each group was protecting a different risk. One side wanted speed and customer relief; another wanted control, feasibility, or quality. I made that tradeoff visible so the room could decide instead of arguing past each other. Calibration: Is navigating that kind of competing-priority tension common in claims process work here?"
        ),
        (
            "Difficult customer or client - $1M account version",
            "Emphasize relationship: the client was frustrated because they felt passed around, not because they wanted a perfect explanation. What I noticed was emotional first: they needed to believe one person owned the recovery. I listened, clarified what 'fixed' meant to them, consolidated ownership, and kept the communication reliable until confidence returned. Calibration: When customer experience is breaking down, does this team expect the Process Engineer to help create ownership across the workflow?"
        ),
        (
            "Describe a time you used data to drive a process decision.",
            "Use the dashboard/KPI story. Start with the business question, not the tool. Explain how you validated source data, segmented the issue, built the reporting view, and helped leaders make decisions from operational signals instead of raw exports or verbal status updates."
        ),
        (
            "How do you ensure quality and accuracy when documenting or redesigning a process?",
            "Answer with a repeatable system: observe the actual current state, document it, validate with SMEs, pilot the redesign, define success metrics, version-control the process documentation, and run 30/60/90-day checks. This is especially important in insurance because process errors can become customer, compliance, or financial risk."
        ),
        (
            "Tell me about a process improvement that did not go as planned.",
            "Keep the failure brief and spend most of the answer on diagnosis and the safeguard you built afterward. A strong answer sounds like: 'What I underestimated was X. Once the issue surfaced, I did Y. What I changed afterward was Z.' Avoid a story where the failure was negligence."
        ),
        (
            "How would you explain a complex process change to a nontechnical audience?",
            "Use SCR: Situation, Complication, Resolution. Start with what the audience already cares about, avoid jargon, use an analogy if helpful, and check understanding by asking them to explain how the change affects their team. Signal that you enjoy translation work rather than treating it as a burden."
        ),
        (
            "Why State Farm?",
            "Use the 50/50 rule and avoid the uniqueness trap. First half: one authentic State Farm reason with specificity - scale, good-neighbor customer impact, tuition reimbursement/learning, workplace reputation, or community investment. Second half: Christian brings process diagnosis, data-backed recommendation, stakeholder alignment, documentation, and measurable adoption. Use because to force a concrete reason."
        ),
        (
            "Leadership style / How do you lead?",
            "My leadership philosophy is to build capability rather than create dependency. In practice, that means three things: understand how each person communicates and what motivates them, set clear goals with individual accountability, and reflect inward first when something goes wrong: was my direction clear, was the environment supportive, and were the expectations realistic? In claims process engineering, that translates directly to adoption: I am not just changing a process; I am helping the people doing the work use the new standard confidently."
        ),
        (
            "Future approach: data and SMEs disagree",
            "I would treat the gap as the most interesting signal in the room, not a conflict to resolve by picking a side. First I would revisit the data with the SME's perspective as a hypothesis. Second I would define what we should expect to see if the SME is right. Third I would validate the revised conclusion with at least two frontline people before presenting anything. For claims, where adjuster pattern recognition can be deep, I would respect the SME read until the data gives me a specific reason not to."
        ),
    )
    for question, answer in answer_bank:
        add_subsection(document, question)
        add_body(document, answer)

def add_state_farm_process_engineering_lens(document: Document) -> None:
    add_section(document, "Process Engineering Lens For State Farm")
    add_body(document, "Use this section to sound like you understand process engineering without pretending to be a chemical, plant, or safety engineer if the interviewer goes there. Translate the industrial language into claims operations language.")
    rows = (
        ("Bridge role", "Process engineers bridge technical principles and daily operations. In Claims Process Engineering, that means connecting process design, operational data, frontline workflow, leadership decisions, and customer experience."),
        ("Process analysis", "Industrial examples use yield, cycle time, and waste. State Farm examples may use claim cycle time, task aging, call volume, handoff delays, rework, quality, service, and customer experience."),
        ("Documentation", "PFDs, P&IDs, and SOPs are blueprints for consistent operations. Your bridge is workflow maps, standard work, communication plans, decision documents, tollgate slides, training materials, and repeatable operating instructions."),
        ("Quality control", "ISO/GMP/Six Sigma language maps to auditability, root-cause analysis, corrective actions, validation checkpoints, and making sure a faster process does not create downstream quality or compliance risk."),
        ("Safety and risk", "Do not claim plant safety ownership. Translate safety into risk controls: protect customers, data quality, compliance, employee workload, and service reliability when a process changes."),
        ("Data and SPC", "If asked about SPC or statistical testing, describe the transferable method: define the metric, check variation, segment the process, identify a likely root cause, test the change, and monitor before/after results."),
        ("Tools boundary", "Aspen Plus, HYSYS, MATLAB, AutoCAD, PFDs, and P&IDs are awareness terms unless directly supported. Use them as smart questions about State Farm's preferred process-mapping and analysis tools, not as experience claims."),
    )
    for title, body in rows:
        add_subsection(document, title)
        add_body(document, body)

def add_state_farm_continuous_improvement_pm_section(document: Document) -> None:
    add_section(document, "Continuous Improvement And PM Operating System")
    add_subsection(document, "Continuous Improvement Habits To Show")
    for line in (
        "Capture lessons learned throughout the project, not only at close. Say: 'If I learn something in week two, I want the project to benefit in week three.'",
        "Apply feedback loops immediately. Do not wait for the next phase if the evidence is already clear.",
        "Run retrospectives even when the project went well, because good outcomes can still hide preventable waste.",
        "Repeat the project vision. Teams lose the big picture when they are buried in tasks; reconnect the process change to service, quality, and customer impact.",
        "Pilot before full rollout. A pilot gives skeptics evidence and protects the team from scaling an unvalidated change.",
        "Use the proactive feedback loop personally: anticipate feedback, improve, check back, and get confirmation that the change worked.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Project Management Principles To Weave Into Answers")
    for line in (
        "Stakeholders: evaluate requests before dismissing them; understand the business value and communicate differently by audience.",
        "Change control: every change ripples into scope, time, cost, quality, resources, risk, and adoption. Make the tradeoff visible before acting.",
        "Decision-making: investigate before fixing. Review existing plans, current standards, and known risks before proposing a solution.",
        "Leadership: act as a servant leader and integrator. Consult the team, protect people from unsustainable workload, and avoid public embarrassment or blame.",
        "Risk: treat the risk register as living. When an issue appears, ask whether it was already a known risk and whether a response plan exists.",
        "Agile: if the team works agile, respect backlog ownership and use pilots, sprint reviews, and customer feedback to adjust."
    ):
        add_bullet(document, line)
    add_subsection(document, "Power Phrase")
    add_body(document, "'My process-improvement style is to avoid optimizing one metric in isolation. I try to balance efficiency, quality, risk, stakeholder adoption, and customer experience, then use a pilot and feedback loop to prove the change works before scaling it.'")

def add_state_farm_interview_delivery_addendum(document: Document) -> None:
    add_section(document, "Interview Delivery Addendum From New Notes")
    add_subsection(document, "Behavioral Story Structure")
    for line in (
        "Use one working sequence for this State Farm interview: Hook -> Noticing -> Action -> Result -> Bridge -> Calibration. The label matters less than the order.",
        "Past-experience questions get the full story sequence. Future-approach questions get a direct position, three steps, one proof point, and a State Farm bridge.",
        "Story selection rule: relevancy beats recency. A 10-year-old story that mirrors their process challenge is better than a recent story that only sounds impressive.",
        "Best challenge-story arc: the situation was not your fault, you went above and beyond gladly, stakeholders moved from frustrated or skeptical to confident, and you can name what you learned.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Conviction And Pacing Mechanics")
    for line in (
        "Emphasize key words inside each sentence, not only the final word. Say: 'I reduced manual work by seventy-eight percent' with emphasis on reduced and seventy-eight percent.",
        "After stating each major metric, hold eye contact with the camera lens for two full seconds before continuing. Holding eye contact after the claim signals you own it.",
        "When you notice your pace accelerating, use the silent count: count one, two, three internally before the next sentence. Use it before your opening answer and before stating your result.",
        "Tell them the time, do not build them the clock. Give the result, method, and bridge; let them ask for the mechanism if they want it.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Video Interview Details")
    for line in (
        "Quiet environment: no distracting sounds or movement; test audio levels before starting.",
        "Clean background: simple and professional. Avoid fake or blurred backgrounds if possible.",
        "Camera at eye level: prop the laptop on books if needed; keep torso and hand gestures visible.",
        "Lighting: face a window during the day or use a lamp behind the camera at night.",
        "Business casual attire: no t-shirts, hats, tank tops, or pajamas.",
        "Practice method: record one answer, watch with sound on, listen sound-only for filler words, then watch sound-off for body language.",
        "Filler control: use the silent count two ways. Replace 'um,' 'you know,' and 'as well' with a one-two-three pause, and also use that count between major sentences whenever pace starts accelerating."
    ):
        add_bullet(document, line)

def add_state_farm_video_interview_section(document: Document) -> None:
    add_section(document, "Video Interview Setup And Delivery")
    for line in (
        "Camera at eye level with torso visible. Keep gestures in the truth plane, roughly waist to shoulders.",
        "Use light behind the camera facing you. Avoid fake or blurred backgrounds if possible; clean and real reads better.",
        "Audio matters more than camera quality. Test microphone, connection, and recording environment before the interview.",
        "Look through the camera lens, not at the screen. Place the interviewer's window close beneath the camera.",
        "Practice by recording one answer, then review three times: sound on for content, sound-only for filler words, and sound off for posture and energy.",
        "Nod slowly while listening and paraphrase complex questions before answering. It buys time and shows you heard the actual question."
    ):
        add_bullet(document, line)

def add_state_farm_response_calibration_section(document: Document) -> None:
    add_section(document, "Response Calibration And Short Answer Frameworks")
    add_subsection(document, "Only Two Question Types")
    add_body(document, "Every question is either a past-experience question or a future-approach question. Classify it before answering. That one-second classification prevents using a long story when the interviewer only asked for a short point, or giving theory when they asked for proof.")
    for line in (
        "Past experience: prompts like 'Tell me about a time...' get Hook -> Noticing -> Action -> Result -> Bridge -> Calibration.",
        "Future approach: prompts like 'How would you handle...' get a position statement, three supporting steps, one past proof point, and a State Farm bridge.",
    ):
        add_bullet(document, line)
    add_subsection(document, "PREP For Short Direct Answers")
    add_body(document, "Use PREP for 15-20 second answers: Point, Reason, Example, Point restatement.")
    add_body(document, "Example: 'Yes, I thrive in ambiguous environments because ambiguity is where structure creates the most value. The inventory adjustment project started with no defined root cause or fix, and my job was to build both. That is the kind of environment where I do my best work.'")
    add_subsection(document, "Executive Presence Q&A Calibration")
    for line in (
        "True/False, yes/no: one sentence answer plus one evidence sentence. 15 seconds max. 'Yes, and here is the clearest proof...'",
        "Multiple choice: name the choice, explain why, eliminate the alternatives briefly. About 30 seconds.",
        "Fill-in-the-blank, such as 'My approach to X is...': use PREP. 20-30 seconds.",
        "Essay, such as 'Tell me about a time...': use the full story sequence. 60-90 seconds.",
        "Most candidates answer every question as an essay. That is why they run too long on simple questions and sound less concise than they are.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Future Approach Template")
    for line in (
        "Step 1: state your position clearly in one sentence; do not hedge.",
        "Step 2: give three specific steps or principles you would apply.",
        "Step 3: ground at least one step in a past experience as validation.",
        "Step 4: connect the approach to State Farm's claims/process environment.",
        "Example: 'If the data and SMEs disagree, I would treat that gap as the most interesting signal in the room. First I would revisit the data with the SME's perspective as a hypothesis. Second I would ask what we should see if they are right. Third I would validate the revised conclusion with two frontline people before presenting anything.'",
    ):
        add_bullet(document, line)
    add_subsection(document, "Tell Them The Time")
    add_body(document, "Self-monitoring rule: tell them the time, do not build them the clock. State the result, the key method, and the bridge. If they want the mechanism behind the clock, they will ask. If you find yourself in sentence five of an answer that started as yes/no, stop, land the result, and ask a calibration question.")

def add_state_farm_uniqueness_trap_warning(document: Document) -> None:
    add_section(document, "Why State Farm - Authenticity Warning")
    add_body(document, "Avoid the uniqueness trap. Every State Farm interviewer has heard 'I am passionate about insurance' and 'I admire your community investment.' You do not need an original reason; you need an authentic one.")
    add_body(document, "The difference is specificity: 'I value companies that invest in their people' is generic. 'The tuition reimbursement signal matters to me because I am actively studying toward project-management growth and I want a company that treats learning as part of the job, not something you do in spite of it' is authentic. Same motivation, completely different impact.")

def add_state_farm_drill_down_layers(document: Document) -> None:
    add_section(document, "5-7 Layers Of Drill-Down Readiness")
    add_body(document, "Interviewers can spend 15-45 minutes on a single story. The skeleton gets the story started; these layers keep Christian ready when they ask 'tell me more,' 'walk me through exactly how,' or 'what happened next?' ")
    drilldowns = (
        ("Inventory Adjustment", (
            "L1: what you did and the result: 78% less manual work, 22% fewer discrepancies.",
            "L2: how you mapped the workflow: steps, handoffs, manual corrections, users involved, and where data entered the process.",
            "L3: what was unexpected: the issue was structural, not a people-performance problem.",
            "L4: how you validated: bring the warehouse team into the workflow review before touching the live process.",
            "L5: resistance: address fear of disruption by piloting and showing how the checkpoint reduced rework.",
            "L6: measurement: define baseline manual effort/discrepancies, compare after rollout, and watch for unintended effects.",
            "L7: what you would do differently: document the standard sooner and create a feedback loop immediately after launch.",
        )),
        ("$1M Account Stabilization", (
            "L1: account was at risk; structured recovery stabilized $1M+ in annual revenue.",
            "L2: first diagnosis: customer had no single accountable owner across open issues.",
            "L3: emotional read: frustration was about trust and follow-through, not just defects.",
            "L4: operating rhythm: weekly working sessions, named owners, decision log, product/development in room.",
            "L5: resistance: internal teams had competing priorities; the agenda made tradeoffs visible.",
            "L6: measurement: movement from escalation calls to renewal conversations and retained revenue risk.",
            "L7: lesson: influence without authority starts with accountable structure, not louder persuasion.",
        )),
        ("200+ Dashboards", (
            "L1: built 200+ tools that improved decision visibility.",
            "L2: first question: what decision does this report need to support?",
            "L3: data validation: check source definitions, variance, gaps, and whether the metric is trusted.",
            "L4: what changed: leaders moved from raw exports/verbal updates to shared operating signals.",
            "L5: resistance: translate reporting into the stakeholder's decision language, not the builder's tool language.",
            "L6: measurement: adoption by leaders, recurring use, and improved quality of operating conversations.",
            "L7: State Farm transfer: claim/task/phone metrics should answer a decision, not just describe activity.",
        )),
        ("Rapid Learning", (
            "L1: became credible across 80+ international clients in a complex domain.",
            "L2: learning method: use live problems as case studies of the workflow.",
            "L3: SME validation: confirm understanding with people closest to the work before recommending changes.",
            "L4: first 30-60 days: shadow, document, ask 'newbie' questions, and learn the trusted metrics.",
            "L5: resistance: avoid pretending expertise; earn credibility by being useful quickly.",
            "L6: proof: discovery through go-live and post-go-live support across international clients.",
            "L7: State Farm transfer: learn claims through workflow, data, and frontline validation.",
        )),
        ("Operations Versus Finance", (
            "L1: competing stakeholders were optimizing different constraints.",
            "L2: what each side protected: operations protected speed; finance protected control and auditability.",
            "L3: room dynamic: name the tradeoff without making either side wrong.",
            "L4: action: build options with cost, timeline, risk, and adoption impact visible.",
            "L5: resistance: acknowledge why each group is rational from its own operating reality.",
            "L6: measurement: decision quality improved because tradeoffs were explicit.",
            "L7: State Farm transfer: claims process work will often balance speed, quality, compliance, and customer experience.",
        )),
    )
    for title, layers in drilldowns:
        add_subsection(document, title)
        for layer in layers:
            add_bullet(document, layer)

def add_state_farm_coachability_protocol(document: Document) -> None:
    add_section(document, "Coachability Signals In The Panel Round")
    add_body(document, "Coachability is demonstrated in the moment, especially during a data exercise or case presentation. The panel is watching whether Christian treats pushback as useful information or as an attack.")
    rows = (
        ("When they redirect mid-presentation", "Do not finish your slide. Go directly to what they asked: 'Good direction - let me go there.' That shows you are tracking their priority, not protecting your script."),
        ("When they challenge a finding", "Do not defend first. Say: 'That is a fair challenge. If that is true, my recommendation changes in this direction.' Then explain how."),
        ("When you do not know", "Say: 'I do not have enough information to be confident there. My assumption is X; if that is wrong, here is what I would need to validate.'"),
    )
    for title, body in rows:
        add_subsection(document, title)
        add_body(document, body)

def add_state_farm_skeleton_study_system(document: Document) -> None:
    add_section(document, "State Farm 2-Day Study System")
    add_subsection(document, "Core Rule: Skeletons, Not Scripts")
    add_body(document, "Do not memorize full answers. Memorize the anchors that let you rebuild the answer naturally under pressure. The full versions in this guide are practice models, not scripts to recite.")
    for line in (
        "Memorize five anchors per story: hook, discovery, what I noticed, action, result/bridge.",
        "Lock in only the few exact phrases that matter: 30-second pitch, 78% / 22%, 200+ tools, 60+ workshops, Why State Farm angle, and the closing concern question.",
        "Practice by speaking out loud, not by re-reading. The gap between head and mouth is where interviews usually break down.",
        "Record once with sound on for content, once sound-only for fillers and pacing, and once muted for posture, gestures, and energy.",
        "Stop heavy practice the night before. Review skeletons lightly, then let the brain consolidate instead of turning preparation into stiffness.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Two-Day Ramp")
    for line in (
        "Day 1 - Build day, 4 hours total: Hour 1 skeleton cards for all primary stories, bullet points only. Hour 2 record Tell Me About Yourself, watch twice, note one fix, record once more. Hour 3 practice Why State Farm and the closing sequence out loud. Hour 4 complete one data-exercise walkthrough out loud, narrating each step. Hard stop.",
        "Day 2 - Pressure test day, 3 hours total: Hour 1 reconstruct all story skeletons from memory without notes. Hour 2 run a five-question simulated interview and record it. Hour 3 review the recording, fix one thing only, then stop. Evening is light skeleton review and one final recording, with a hard stop by 9pm.",
        "Morning of: no new prep. Equipment check, breathing sequence, one skeleton read, and conversation mode.",
    ):
        add_bullet(document, line)

def add_natural_storytelling_system(document: Document) -> None:
    add_section(document, "Natural Storytelling System")
    add_body(document, "The goal is not to sound polished. The goal is to sound like a credible person telling a real professional memory with judgment, evidence, and a clean bridge to State Farm.")
    rows = (
        ("Lead with tension", "Open with what was broken or at risk, not with title, company, and chronology. Example: 'We had a recurring warehouse accuracy problem that kept coming back.'"),
        ("Use the noticing line", "Between discovery and action, add: 'What I noticed was...' This shows judgment and separates Christian from candidates who only list actions."),
        ("Use before / after", "State what changed because Christian was there. Before the fix, what was slow, messy, manual, or risky? After the fix, what became faster, cleaner, safer, or measurable?"),
        ("Use And then / But / So", "Use simple connectors to keep stories human: 'and then' creates movement, 'but' creates tension, and 'so' shows the decision."),
        ("Shift into memory mode", "During the action phase, it is okay to sound conversational: 'So I map the workflow and start seeing...' Return to past tense for the result."),
        ("Recover cleanly", "If you lose your place, say: 'Let me come back to that from a slightly different angle.' Then land the result and offer to explain the path."),
        ("Wrong question recovery", "If you realize mid-answer that you may be answering the wrong question, stop and say: 'Actually, let me make sure I am addressing what you are really asking. Are you looking more for X or Y?'"),
        ("Pushback recovery", "When pushback lands, say: 'That is a fair challenge.' Pause for two seconds, then explain how the recommendation or answer changes if their challenge is true."),
        ("Silent count", "Use the silent one-two-three in two ways: replace fillers when an um or you know wants to come out, and control rhythm between major sentences when anxiety speeds up the answer."),
    )
    for title, body in rows:
        add_subsection(document, title)
        add_body(document, body)
    add_subsection(document, "Bridge Phrase Library")
    for line in (
        "From situation to discovery: 'The thing nobody had looked at yet was...'",
        "From discovery to action: 'Once I understood that, the path forward was clear...'",
        "From action to result: 'When we checked the before-and-after, here's what the data showed...'",
        "From result to State Farm: 'The reason that matters in Claims Process Engineering is...'",
    ):
        add_bullet(document, line)

def add_state_farm_natural_story_models(document: Document) -> None:
    add_section(document, "State Farm Story Skeletons And Spoken Models")
    add_body(document, "Use these as practice models. Read the model once or twice, then cover it and reconstruct from the skeleton. The wording should change each time; the anchors should not.")
    stories = (
        (
            "Inventory Adjustment System - Primary Story",
            "Process improvement, data-driven decision, initiative, ambiguous problem.",
            (
                "Hook: recurring warehouse discrepancy problem with repeated cleanup and no root-cause fix.",
                "Discovery: mapped workflow and found the same manual step appearing in multiple places.",
                "What I noticed: not a people problem; a process gap with no checkpoint between input and posting.",
                "Action: built automated, auditable workflow; validated with warehouse users; piloted before scaling.",
                "Result / bridge: 78% less manual work and 22% fewer discrepancies; same pattern for claims: map, root cause, pilot, measure.",
            ),
            "Practice model: We had a warehouse accuracy problem that kept coming back. Every month, the same discrepancies, and every month someone was manually correcting them. So I mapped the workflow and saw the same manual step showing up in multiple places. What changed the approach was realizing this was not a people problem. It was a process gap. We built an automated checkpoint, validated it with the team, piloted it, and measured before and after. Manual work dropped 78%, discrepancies dropped 22%, and the team stopped fighting the same fire every month. That is the exact pattern I would bring to claims process engineering.",
        ),
        (
            "$1M+ Account Stabilization",
            "Influence without authority, stakeholder resistance, difficult customer, persuasion.",
            (
                "Hook: account escalating for months; customer losing trust.",
                "Discovery: issue was not only technical; no one owned the full recovery path.",
                "What I noticed: the customer needed one accountable owner, not another status email.",
                "Action: consolidated cases, escalation path, weekly agenda, and product/development coordination.",
                "Result / bridge: $1M+ in at-risk revenue stabilized; influence means giving people one credible path forward.",
            ),
            "Practice model: I was brought into an account that had been escalating for months. The technical issues were real, but what I noticed in the first working session was that the customer had no single owner. Three contact paths, multiple open items, no credible recovery rhythm. So I consolidated ownership, stopped relying on status emails, and ran structured working sessions with product, development, and the customer in the room. Over three months we moved from escalation calls to renewal conversations, stabilizing more than $1M in at-risk revenue.",
        ),
        (
            "200+ Dashboards And Decision Visibility",
            "Data-driven process decisions, analytical thinking, operational metrics.",
            (
                "Hook: leaders using raw exports and verbal updates; no shared operating signal.",
                "Discovery: report requests were really decision requests.",
                "What I noticed: the business question had to be defined before the report could be trusted.",
                "Action: validated source data, segmented the problem, built SQL/Power BI/Crystal decision tools.",
                "Result / bridge: 200+ reporting tools; same discipline for claims metrics and data exercises.",
            ),
            "Practice model: Leaders were making decisions from raw exports and verbal updates, so every meeting had a slightly different version of reality. I started asking a different question: what decision needs to be made, and what would change the answer? From there I validated source data, built the right KPI views, and created 200+ tools that turned operating noise into signals. For State Farm, I would apply the same discipline to claim, task, phone, and workflow metrics.",
        ),
        (
            "Rapid Product And Domain Learning",
            "No direct claims background, rapid learning, new domain credibility.",
            (
                "Hook: joined complex manufacturing software environment without the domain background.",
                "Discovery: fastest learning came from client problems, not static documentation.",
                "What I noticed: every support or implementation issue was a case study in how the workflow actually operated.",
                "Action: learned through discovery, requirements, configuration, data migration, go-live, and post-go-live support.",
                "Result / bridge: credible across 80+ international clients; same method for learning claims operations quickly.",
            ),
            "Practice model: I did not enter Aptean already fluent in every manufacturing environment. I became credible by learning through the work: client questions, workflow breakdowns, data issues, and go-live problems. Every ticket became a case study in how the business actually operated. That is how I would close the claims knowledge gap at State Farm: shadow the workflow, learn the metrics people already trust, validate with SMEs, and earn credibility through useful analysis rather than claiming domain expertise I have not yet earned.",
        ),
        (
            "Operations Versus Finance Alignment",
            "Competing priorities, stakeholder conflict, tradeoff communication.",
            (
                "Hook: operations and finance used the same words but meant different risks.",
                "Discovery: one group optimized for speed; the other optimized for control and auditability.",
                "What I noticed: the disagreement was not about the system; it was about whose risk mattered more.",
                "Action: named the tradeoff, built options with cost/timeline/risk visible, and created a decision path.",
                "Result / bridge: stakeholders could choose consciously instead of arguing past each other.",
            ),
            "Practice model: Operations and finance were not really fighting about the system. They were using the same word, accuracy, to mean two different things. Operations meant speed and flow. Finance meant control and auditability. Once I named that tradeoff, the room changed. We could stop debating who was right and start comparing options: what each path cost, what risk it reduced, and what it delayed. That is the same habit I would bring to claims process work when speed, quality, compliance, and customer experience pull in different directions.",
        ),
        (
            "Failure And Stronger Validation",
            "Failure, growth mindset, quality control, prevention habit.",
            (
                "Hook: an implementation moved faster than the validation rhythm could support.",
                "Discovery: the issue was not effort; it was that a stakeholder checkpoint was too late.",
                "What I noticed: quality risk was entering after the design decision, when correction was more expensive.",
                "Action: added earlier SME validation, clearer acceptance criteria, and stronger post-launch checks.",
                "Result / bridge: stronger validation habit for every later project; useful in regulated claims operations.",
            ),
            "Practice model: The failure I would choose is not negligence; it was a validation miss in a complex rollout. What I underestimated was how late a key stakeholder checkpoint was happening. By the time feedback surfaced, correcting it created more friction than it should have. What I changed afterward was specific: earlier SME validation, clearer acceptance criteria, and checkpoints before the work became expensive to unwind. The lesson for State Farm is that quality control has to be designed into the process, not inspected in at the end.",
        ),
    )
    calibration_by_key = {
        "inventory": "Is finding the structural root cause, rather than building better error detection, the kind of thinking this team is looking for?",
        "account": "Is cross-functional coordination without formal authority a regular part of how this role operates?",
        "dashboards": "Is defining the business decision before touching the data a discipline the team is still building, or is that infrastructure mature?",
        "learning": "Is a 30-to-60 day ramp where I learn the claims environment before recommending anything the right expectation for how someone starts in this role?",
        "ops_finance": "Is navigating that kind of competing-priority tension, where process and compliance teams each define success differently, a regular part of what this role manages?",
        "failure": "Is building SME validation into every process change before go-live a strong discipline here already, or an area to strengthen?",
    }
    for title, use_for, skeleton, model in stories:
        add_subsection(document, title)
        add_body(document, f"Use for: {use_for}")
        for line in skeleton:
            add_bullet(document, line)
        add_body(document, model)
        add_body(document, f"Calibration question: {calibration_by_key.get(state_farm_story_key(title), 'Is that the right way to think about how this experience maps to the role?')}")

def add_state_farm_on_demand_video_deep_dive(document: Document) -> None:
    add_section(document, "On-Demand Video Round Deep Dive")
    add_body(document, "Treat the on-demand video as a recorded recommendation, not a live conversation. There is no interviewer to rescue a wandering answer, so every response needs a clean opening, visible structure, and a strong stop.")
    for line in (
        "Read the question twice. Identify whether it is asking for priority, process, customer, failure, or why-fit.",
        "Take one breath, smile slightly, and start with the result or answer. Two seconds of silence reads as thoughtfulness, not freezing.",
        "Use Result -> Method -> Bridge: what changed, how you got there, why that same method matters for State Farm.",
        "Speak 10-15% louder than normal. Energy leaks through screens, especially in recorded formats.",
        "Pause between sections. The pause creates structure when there is no interviewer nodding along.",
        "End with a short bridge or calibration-style close; do not trail off with 'so yeah.'",
    ):
        add_bullet(document, line)
    add_subsection(document, "Recovery Phrases")
    for line in (
        "Need time: 'That's a good question. Let me think about that for a second.'",
        "Lost the thread: 'Let me come back to that from a slightly different angle.'",
        "Too long: 'The short version is...' then land the metric and stop.",
        "Gap question: 'That's accurate, and I would not overclaim that experience. The transferable pattern is...'",
    ):
        add_bullet(document, line)

def add_state_farm_worked_data_exercise(document: Document) -> None:
    add_section(document, "Worked Data Exercise Example")
    add_body(document, "Use this as a model for how to narrate analysis. The specific numbers are practice assumptions, not claims about State Farm data. The method is the point: define the question, test the data, segment, hypothesize, pilot, and measure.")
    add_subsection(document, "Scenario")
    add_body(document, "You receive claims data showing multi-vehicle auto claims have longer assignment-to-first-contact time and higher rework than simple auto claims. Leadership wants to know where to intervene first.")
    add_subsection(document, "How To Think Out Loud")
    for line in (
        "Business question: Are we trying to improve speed, quality, customer experience, workload balance, or all four? I would clarify the decision before calculating anything.",
        "Data quality: check date definitions, missing timestamps, duplicate claims, outliers, claim-type coding, and whether rework is measured consistently.",
        "Segmentation: compare by claim type, complexity, hub, adjuster capacity, day of week, channel, assignment path, and handoff count.",
        "Hypothesis: if delay clusters before assignment, routing or ownership may be the lever; if delay appears throughout processing, training, documentation, or decision authority may be the lever.",
        "Recommendation: choose one structurally discrete pilot first so the result can be attributed to the intervention.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Three-Slide Case Logic")
    for line in (
        "Slide 1 - Recommendation: pilot a routing or ownership change for the highest-delay segment; expected impact is faster first contact without quality deterioration.",
        "Slide 2 - Evidence: show segmentation, what you ruled out, and why the selected lever is the cleanest first test.",
        "Slide 3 - Implementation: six-week pilot, control group if available, primary metric, guardrail metrics, stakeholder plan, rollback criteria, and decision ask.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Panel Pushbacks To Expect")
    for line in (
        "If they ask why not solve every issue at once: explain that multiple simultaneous interventions make it impossible to know what worked.",
        "If they ask whether it could be training instead of routing: agree it is plausible, then explain which timestamp or quality pattern would confirm or disprove it.",
        "If they challenge the timeline: explain the tradeoff between speed and enough volume to distinguish real improvement from random variation.",
        "If they ask what happens after the pilot: scale only if the primary metric improves and guardrails hold; otherwise revise the root-cause hypothesis based on pilot data.",
    ):
        add_bullet(document, line)

def add_state_farm_objection_handling_deep_dive(document: Document) -> None:
    add_section(document, "Pushback And Objection Handling")
    add_body(document, "The goal is to answer doubts directly without sounding defensive. Use the formula: acknowledge, bridge to transferable pattern, give evidence, close with confidence.")
    rows = (
        ("No direct insurance or claims background", "That is accurate; I would not claim direct claims experience. What I have is a repeatable method for learning operational environments quickly: shadow the workflow, validate the metrics people already trust, find where the process breaks, and test changes with SMEs before scaling them. I used that approach to become credible with 80+ international manufacturing clients, and I would apply it to claims operations from day one."),
        ("Too systems-heavy", "The systems were the vehicle, not the point. The business work was process improvement: reducing manual work, improving data quality, training users, documenting standard work, and helping teams adopt a better operating rhythm."),
        ("Limited formal Lean Six Sigma certification", "I would be precise: I have applied Lean Six Sigma-style thinking through current-state mapping, waste reduction, root-cause analysis, pilot validation, before/after measurement, and stakeholder adoption. I would also learn and use State Farm's preferred terminology and templates."),
        ("Not a traditional industrial process engineer", "Correct. My fit is process engineering in an operations environment: workflow mapping, operational metrics, quality controls, documentation, stakeholder adoption, and measurable improvement. I would use industrial terms as awareness, not as unsupported claims."),
        ("Concern about moving too slowly", "My natural caution shows up when downstream impact is unclear. I manage that with pilot-first design: small test, clear metric, rollback path, and decision checkpoint so the team can move without pretending the risk is zero."),
    )
    for concern, answer in rows:
        add_subsection(document, concern)
        add_body(document, answer)

def add_state_farm_first_90_days_deep_dive(document: Document) -> None:
    add_section(document, "First 90 Days - Process Engineer Plan")
    add_body(document, "If asked, answer in three phases. The theme is learn before prescribing, validate before scaling, and produce something useful by day 90.")
    phases = (
        ("Days 1-30: Learn the operating system", "Shadow claims workflows, learn the team's language, map stakeholder groups, understand trusted metrics, review active tollgate materials, and identify where process work currently slows down."),
        ("Days 31-60: Diagnose and prioritize", "Build a short list of friction points, separate symptoms from likely root causes, validate with SMEs, quantify impact/effort/risk, and recommend one pilot candidate with success and guardrail metrics."),
        ("Days 61-90: Pilot and package the operating rhythm", "Run or support a focused pilot, create communication materials, document the standard work, monitor before/after metrics, capture lessons learned, and present the next decision: scale, adjust, or stop."),
    )
    for title, body in phases:
        add_subsection(document, title)
        add_body(document, body)
    add_subsection(document, "90-Day Deliverables")
    for line in (
        "A current-state map of at least one priority workflow.",
        "A metric definition and data-quality check for the selected process problem.",
        "A stakeholder map with sponsors, SMEs, blockers, and adoption risks.",
        "A pilot recommendation with business case, success metric, guardrail metric, and rollback path.",
        "A lessons-learned loop that turns pilot feedback into the next iteration.",
    ):
        add_bullet(document, line)

def add_state_farm_day_of_protocol(document: Document) -> None:
    add_section(document, "Day-Of Protocol")
    add_subsection(document, "Night Before")
    for line in (
        "Do one light review of story skeletons, questions to ask, and the closing sequence. Do not drill until exhausted.",
        "Record Tell Me About Yourself once, note one improvement, then stop.",
        "Hard stop by evening. Rest is part of preparation because recall and verbal fluency depend on recovery.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Morning Of")
    for line in (
        "Eat real food. Anxiety on an empty stomach is louder than anxiety with fuel.",
        "Dress fully in interview attire before final setup; it helps shift into professional mode.",
        "Do three breathing cycles: inhale 4, hold 4, exhale 6.",
        "Say the frame out loud: 'I have evidence for everything I say. I am here to explore whether this is the right environment for my skills. I am ready.'",
        "Join early, check camera and audio, then stop reviewing notes five minutes before start.",
    ):
        add_bullet(document, line)
    add_subsection(document, "During")
    for line in (
        "If anxiety spikes, acknowledge it silently and let it pass. Do not fight it or narrate it.",
        "If the interviewer is flat, do not read it as negative feedback. Keep energy steady.",
        "If an answer runs long, land the result and stop. Do not keep repairing aloud.",
        "If they correct your framing, treat it as useful context and adjust: 'That changes how I would frame the most relevant part of my background.'",
    ):
        add_bullet(document, line)

def add_state_farm_post_interview_strategy_deep(document: Document) -> None:
    add_section(document, "Post-Interview Strategy")
    add_subsection(document, "First 30 Minutes")
    for line in (
        "Capture names, roles, specific points each person made, concerns raised, weak answers, and exact language they used for the role's biggest challenge.",
        "Give yourself a 15-minute emotional reset before writing the thank-you note so it does not sound anxious or over-effusive.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Thank-You Note Scenarios")
    for line in (
        "If it went well: thank them, reference one specific thing this interviewer said, connect that point to one specific proof from Christian's background, add one useful thought that occurred after the conversation, and close forward-looking.",
        "If a concern was raised: add one sentence acknowledging it and reframing it with evidence. Do not write a paragraph defending yourself.",
        "If there were multiple interviewers: send separate notes with one specific point from each person's conversation.",
        "If there was a data exercise: add one thoughtful follow-up insight only if it improves the analysis without sounding like a correction spiral.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Required Four-Part Template")
    for line in (
        "1. Reference one specific thing this interviewer said, not a general topic.",
        "2. Connect that point to one specific proof from the resume or story bank.",
        "3. Add one new thought after the conversation, framed as an addition rather than a correction.",
        "4. Close with specific interest in the work they described.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Round Two Intelligence")
    for line in (
        "Use the exact language they used in round one to describe the challenge.",
        "Prepare deeper answers for any question you only partially answered.",
        "Research any tool, process, or methodology they mentioned that you did not know deeply.",
        "Shift from 'can I do the work?' to 'am I the safest, clearest, most useful person to hire?'",
    ):
        add_bullet(document, line)

def add_state_farm_questions_and_rapport_section(document: Document) -> None:
    add_section(document, "Questions, Rapport, And Tell Me More")
    add_body(document, "Strong questions are not only for the end. Use them during the interview to gather the language you will mirror in your close and thank-you note.")
    add_subsection(document, "Tell Me More")
    add_body(document, "Use 'Tell me more about that' when an interviewer says something interesting about the team's biggest challenge, what success looks like, or what is not working today. This is not a delay tactic; it signals genuine curiosity and gives you raw material for your closing summary.")
    add_subsection(document, "Priority Questions")
    for line in (
        "What does success look like in the first 12 months?",
        "What is the most challenging process problem your team is working on right now?",
        "Which operating metrics does the team trust today, and which ones still need better definition?",
        "What behaviors separate the most effective people on this Claims Process Engineering team?",
        "Do you have any questions or concerns about my ability to do the job?",
    ):
        add_bullet(document, line)
    add_subsection(document, "Mirror Their Answer")
    add_body(document, "When they answer, mirror one phrase back before pitching yourself. Example: 'That point about data quality creating downstream rework is helpful context. The closest pattern in my background is...' This keeps the interview conversational rather than rehearsed.")

def add_state_farm_master_checklist(document: Document) -> None:
    add_section(document, "Final Master Checklist")
    categories = (
        ("Mindset", ("I have evidence for everything I say", "I am showing what it is like to work with me", "Resistance is information, not an attack", "The interview is a real conversation about a real problem", "I am evaluating them as much as they are evaluating me")),
        ("Stories", ("Inventory story includes what-I-noticed moment", "$1M account story includes ownership/accountability insight", "Dashboard story starts with decision before data", "Rapid learning story bridges no-claims gap honestly", "Failure story ends with a specific validation habit")),
        ("Core Answers", ("Tell Me About Yourself under 90 seconds", "Why State Farm uses one researched angle and because bridge", "Why Hire You names three role-specific reasons", "First 90 Days has three phases and deliverables", "Greatest Weakness is away from the core job and resolved with pilot discipline")),
        ("Delivery", ("Result first", "I not we", "Downward tone on metrics", "2-second pause feels normal", "Recovery phrase ready", "No trailing 'does that make sense?'")),
        ("Questions", ("12-month success", "Most challenging current process problem", "Trusted metrics and data gaps", "Best team member behaviors", "Any concerns about my ability to do the job")),
        ("Setup", ("Camera at eye level", "Torso and hands visible", "Light behind camera", "Real background", "Mic tested", "Phone silenced", "Water and notes off camera")),
    )
    for title, items in categories:
        add_subsection(document, title)
        for item in items:
            add_bullet(document, item)

def add_state_farm_workbook_cover_note(document: Document) -> None:
    add_section(document, "Workbook Promise")
    add_body(document, "This is not a cheat sheet. It is a rehearsal workbook for the State Farm Process Engineer interview process: on-demand video, data exercise, case presentation, and live panel. The goal is not to memorize paragraphs. The goal is to build enough depth that Christian can answer the first question cleanly and survive five to seven follow-up layers without sounding scripted.")
    add_body(document, "Use the short interview cheat sheet the day of the interview. Use this workbook before then: mark it up, practice the drills, record answers, and pressure-test the stories. If a section feels too long, that is intentional. The guide is meant to make the actual interview feel smaller than the preparation.")
    for line in (
        "Primary interview thesis: Christian improves messy operations by finding the real process gap, proving it with data, building a practical fix, and helping people adopt it.",
        "State Farm translation: claims operations need speed, quality, service, cost discipline, and customer experience to improve together, not one metric at the expense of the others.",
        "Main risk to manage: sounding like a systems candidate. Every technical example must land as process improvement, quality control, operational measurement, stakeholder adoption, or customer impact.",
        "Primary story sequence: Hook -> Noticing -> Action -> Result -> Bridge -> Calibration.",
    ):
        add_bullet(document, line)

def add_state_farm_role_deconstruction_workbook(document: Document, job_description: str) -> None:
    add_page_break(document)
    add_section(document, "State Farm Role Deconstruction")
    add_body(document, "The job description is not asking for a generic analyst. It is asking for someone who can improve claims processes at scale. The repeated signals are process design, Lean Six Sigma, measurable efficiency, data analysis, communication materials, agile prioritization, and positive change across complex business partners.")
    rows = (
        ("Conduct research and analysis for end-to-end process design", "Show current-state mapping, workflow observation, root-cause analysis, and process redesign. Use Inventory, Dashboards, and Operations/Finance stories."),
        ("Utilize Lean Six Sigma methodologies", "Use practical Lean language without overclaiming certification: waste, variation, root cause, pilot, before/after measurement, control checks, feedback loops."),
        ("Execute implementation of process improvements", "Do not stop at analysis. Show rollout, training, documentation, adoption, and measurable results."),
        ("Drive collaboration and positive change", "Show influence without authority, stakeholder resistance, executive workshops, and cross-functional alignment."),
        ("Understand structure, standard work, segmentation, detailed processes", "Use language around actual process versus documented process, standard work, segmentation, and operating rhythm."),
        ("Analyze claim, task, phone, and operational metrics", "Bridge SQL/KPI/dashboard proof to claims metrics. The key is decision-first reporting."),
        ("Create tollgate slides and decision documents", "Show executive communication: answer first, evidence second, recommendation third, risks and rollout last."),
        ("Measure cost/benefit and prioritize projects", "Show impact/effort/risk tradeoffs, quick wins versus structural fixes, and visible decision criteria."),
        ("Manage multiple objectives in agile environment", "Show calm prioritization, transparency, stakeholder dependency management, and weekly re-planning."),
    )
    for jd_signal, christian_proof in rows:
        add_subsection(document, jd_signal)
        add_body(document, christian_proof)
    add_subsection(document, "Words To Use Naturally")
    for line in (
        "process design, process improvement, current state, future state, root cause, waste, variation, standard work, segmentation, measurable results",
        "operational efficiency, service, quality, customer experience, cost/benefit, prioritization, tollgate materials, communication plan",
        "claim metrics, task metrics, phone metrics, approved operational metrics, data quality, trend identification, statistical testing",
        "agile prioritization, multiple work objectives, ambiguity, self-starter, resourceful, team environment, shared responsibilities",
    ):
        add_bullet(document, line)
    add_subsection(document, "Words To Use Carefully")
    for line in (
        "ERP, implementation, SQL, dashboards, and systems language are useful only when immediately translated into process, data, or adoption outcomes.",
        "Industrial process engineering terms like PFDs, P&IDs, GMP, ISO, Aspen, HYSYS, MATLAB, and AutoCAD should be awareness terms only unless asked. Do not claim direct plant engineering ownership.",
        "Insurance and claims domain knowledge should be framed as a learning plan, not a fake claim of experience.",
    ):
        add_bullet(document, line)

def add_state_farm_interview_process_map(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Interview Process Map")
    add_body(document, "State Farm describes a three-step process: on-demand video, data exercise, and live hiring team interview with case-study presentation. Treat each round as a different test of the same operating capability.")
    stages = (
        ("Round 1: On-Demand Video", "They are testing clarity, confidence, concise structure, customer/process orientation, and whether Christian sounds calm under time pressure. Use Result -> Method -> State Farm bridge. Do not ramble because no interviewer will interrupt."),
        ("Round 2: Data Exercise", "They are testing how Christian thinks: data quality, segmentation, root-cause hypotheses, controllable levers, cost/benefit, and recommendation discipline. They are not only testing math."),
        ("Round 3: Live Panel And Case Study", "They are testing executive presence, coachability, stakeholder judgment, recommendation defense, and how Christian responds when challenged."),
    )
    for title, body in stages:
        add_subsection(document, title)
        add_body(document, body)
    add_subsection(document, "One Skill They Test In Every Round")
    add_body(document, "Can Christian turn ambiguity into a structured, measurable, human-adoptable improvement? If an answer does not connect to that, it is probably too far from the role.")

def add_state_farm_answer_operating_system(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Answer Operating System")
    add_subsection(document, "Only Two Question Types")
    add_body(document, "Every question is one of two types: past experience or future approach. Past-experience questions start with 'Tell me about a time...' and get a story. Future-approach questions start with 'How would you...' and get a position statement, three steps, one proof anchor, and a State Farm bridge. Classify the question first. It takes one second and determines the entire answer structure.")
    add_subsection(document, "Past Experience Sequence")
    for line in (
        "Hook: name the stakes before chronology. What was broken, risky, slow, manual, unclear, or misaligned?",
        "Noticing: say what Christian saw that others were missing. This is the judgment signal.",
        "Action: explain what Christian personally did. Use I, not vague team language.",
        "Result: quantify what changed. Say the metric cleanly and hold eye contact after it.",
        "Bridge: connect this exact story to claims process engineering. No generic repeated bridge sentence.",
        "Calibration: ask a story-specific question that invites the interviewer to connect the dots or clarify the role.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Future Approach Sequence")
    for line in (
        "Step 1: state your position clearly. Do not hedge.",
        "Step 2: give three specific steps or principles you would apply.",
        "Step 3: ground one step in a past experience as evidence.",
        "Step 4: bridge to State Farm's claims/process environment.",
        "Example: 'I would observe before analyzing, and analyze before recommending. First, I would watch the real workflow. Second, I would label my assumptions before pulling metrics. Third, I would validate the first hypothesis with two SMEs before leadership sees it. At Aptean, I became credible in new client environments by learning through the workflow problems, and I would use that same method with claims.'",
    ):
        add_bullet(document, line)
    add_subsection(document, "PREP For Short Direct Answers")
    add_body(document, "For 15-20 second questions, use PREP: Point, Reason, Example, Point restatement. Do not use a full behavioral story for a yes/no question.")
    add_body(document, "Example: 'Yes, I thrive in ambiguous environments because ambiguity is where structure creates the most value. The inventory adjustment project started with no defined root cause, just a recurring problem, and my job was to build the structure around it. That is actually where I do my best work.'")
    add_subsection(document, "Executive Presence Response Calibration")
    for line in (
        "True/False or yes/no: one direct sentence plus one evidence sentence. Fifteen seconds.",
        "Multiple choice: name the choice, explain why, eliminate alternatives briefly. About thirty seconds.",
        "Fill-in-the-blank, such as 'My approach to X is...': use PREP. Twenty to thirty seconds.",
        "Essay, such as 'Tell me about a time...': use the full story sequence. Sixty to ninety seconds.",
        "Most answer failures come from treating a true/false question as an essay. If the question has a direct answer, give it directly and let them pull for more.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Tell Them The Time")
    add_body(document, "Self-monitoring rule: tell them the time, do not build them the clock. State the result, the key method, and the bridge. If they want the mechanism behind the clock, they will ask. If you are in sentence five of an answer that started as a yes/no question, stop, land the result, and ask a calibration question.")

def add_state_farm_core_positioning_lab(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Core Positioning Answer Lab")
    add_subsection(document, "Tell Me About Yourself - Career Pattern")
    add_body(document, "Do not present three separate jobs. Present one operating pattern across three environments: Christian finds process gaps other people work around and builds something measurable that makes the problem stop recurring.")
    for line in (
        "Home Depot: leaders were making decisions from manual data and fragmented signals, so Christian built reporting structure and visibility.",
        "Aptean: Christian learned operational environments through the problems in them, not through static documentation, then translated that into implementation and customer outcomes.",
        "East West: Christian found a recurring process failure nobody had root-caused, automated the workflow, reduced manual work, and lowered discrepancies.",
        "State Farm: the same operating pattern applies to claims processes with customer stakes, service quality, and measurable efficiency.",
    ):
        add_bullet(document, line)
    add_body(document, "Career-pattern sentence: 'I have usually been the person who finds the process gap others work around and builds something that makes it stop happening.'")
    variants = (
        ("30-second version", "My background sits at the intersection of process improvement, operational data, and stakeholder adoption. At Home Depot, Aptean, and East West, the pattern was consistent: find the workflow gap, prove it with data, build the fix, and help people use it. That is why State Farm's Claims Process Engineering role makes sense to me. It is the same operating pattern, applied to claims processes where speed, quality, and customer experience matter at scale."),
        ("60-second version", "I have spent my career around operational systems, but the work I keep returning to is process improvement. At Home Depot, I built reporting visibility when leaders were working from raw or manual data. At Aptean, I learned customer workflows by working through real implementation problems, requirements, data issues, and go-lives. At East West, I found a recurring inventory discrepancy problem, mapped the workflow, automated the adjustment process, and reduced manual work by 78% while lowering discrepancies by 22%. So I do not see State Farm as a random industry switch. I see it as the same work: understand the real process, identify waste or variation, use data to prove the root cause, and implement a change people can actually adopt."),
        ("Two-minute version", "The easiest way to explain my background is that I have repeatedly worked where operations, data, systems, and people intersect. At Home Depot, that looked like turning customer and workflow data into reporting that leaders could use. At Aptean, it meant supporting 80+ manufacturing clients through discovery, requirements, data migration, configuration, go-live, and executive conversations. At East West, it became internal process ownership: five sites, 150+ users, ERP operations, training, dashboards, and a large inventory adjustment improvement that reduced manual work by 78% and discrepancies by 22%. The common pattern is not the software. It is the process discipline: map what is happening, find the structural gap, validate it with users and data, implement the fix, and measure adoption. Claims Process Engineering is compelling because State Farm is asking for exactly that pattern in a customer-facing operational environment: reduce waste, improve service and quality, analyze approved metrics, create clear communication materials, and help complex teams change how work gets done."),
    )
    for title, body in variants:
        add_subsection(document, title)
        add_body(document, body)
    add_subsection(document, "Why State Farm - Authenticity Warning")
    add_body(document, "Do not manufacture a completely original reason. Most motivations have been heard before. That is fine. What matters is specificity, not originality. 'I admire your community investment' is generic. 'I value companies that invest in people because I am actively building my PMP and process improvement skill set, and I want learning to be part of the job rather than something done in spite of it' is credible.")
    add_subsection(document, "Why State Farm - Answer Model")
    add_body(document, "State Farm stands out to me because claims processes affect real people at moments when speed, clarity, and trust matter. This is not process improvement for a dashboard alone. It connects directly to customer experience, service quality, and operational reliability. My fit is that I have repeatedly improved messy workflows by combining process mapping, data analysis, stakeholder alignment, and adoption. I would be excited to bring that pattern into a Claims Process Engineering team where the work has scale and real customer stakes.")
    add_subsection(document, "Why This Role")
    add_body(document, "The Process Engineer role is compelling because it is not only analysis. The posting asks for research, process design, Lean Six Sigma-style improvement, operational metrics, communication materials, cost/benefit thinking, agile prioritization, and implementation. That is the kind of work I enjoy most: finding the root cause, making the tradeoffs visible, and turning the recommendation into something teams can actually use.")

def state_farm_story_workbook_data(card: cheat.StoryCard) -> dict[str, object]:
    key = state_farm_story_key(card.title)
    base = {
        "inventory": {
            "label": "Inventory Adjustment Process Improvement",
            "use_for": "process improvement, root cause, ambiguity, measurable efficiency, Lean-style waste reduction",
            "noticing": "What I noticed was that the team was treating the recurring discrepancy as an execution issue, but the real problem was structural: the workflow depended on too many manual touches and did not make the error path visible.",
            "bridge": "That pattern translates directly to claims: map what is actually happening, find the structural gap, validate with the people doing the work, pilot it, and measure whether anything changed, not just whether the process looks better on paper.",
            "calibration": "Is finding the structural root cause, rather than building better error detection, the kind of thinking this team is looking for?",
            "drills": (
                ("L1 - skeleton", "What did you do and what changed?", "I mapped a recurring inventory adjustment problem, built an automated and auditable workflow, reduced manual work by 78%, and lowered discrepancies by 22%."),
                ("L2 - mapping", "How exactly did you map the workflow?", "I started with the current path instead of the documented path: who touched the adjustment, where the data came from, where manual re-entry happened, and where exceptions were handled outside the normal process. I treated each handoff as a possible failure point."),
                ("L3 - discovery", "What did you find that others had missed?", "The issue was not that people were careless. The process made the same failure easy to repeat because the manual adjustment path was doing work that should have been controlled, visible, and repeatable."),
                ("L4 - validation", "How did you validate it with the warehouse team?", "I walked the actual workflow back with the users who touched it, asked them where the workaround happened on busy days, and checked whether the proposed process would fit their real constraints rather than an ideal version of the work."),
                ("L5 - resistance", "What pushback did you get?", "The natural concern was that a new workflow would slow people down or add oversight. I addressed it by showing the before/after burden: less manual effort, cleaner auditability, and fewer recurring corrections."),
                ("L6 - measurement", "How did you measure the before and after?", "I anchored on manual effort and discrepancy rate because those were the pain points the operation could feel. The result was a 78% reduction in manual work and a 22% reduction in discrepancies."),
                ("L7 - learning", "What would you do differently now?", "I would define the control metric earlier and set a formal check-in after rollout so the process owner can catch drift before the old workaround returns."),
            ),
        },
        "account": {
            "label": "$1M+ Account Stabilization",
            "use_for": "high-risk stabilization, influence without authority, difficult customer, stakeholder coordination, recovery under pressure",
            "noticing": "What I noticed was that the customer was not only frustrated with the issue. They were frustrated because nobody appeared to own the full recovery path.",
            "bridge": "In claims operations, the same dynamic plays out at scale: when a customer's experience is breaking down, what they need first is one accountable owner and a credible path forward, not a better status cadence.",
            "calibration": "Is cross-functional coordination without formal authority a regular part of how this role operates?",
            "drills": (
                ("L1 - skeleton", "What was the risk and result?", "The account had unresolved workflow and integration issues that put more than $1M in annual revenue at risk. I consolidated ownership, coordinated product and development partners, and stabilized the account."),
                ("L2 - emotional read", "What did you notice about the customer?", "The customer was tired of being passed between explanations. The emotional issue was trust: they needed to believe one person knew the recovery path and would stay with it."),
                ("L3 - influence", "How did you influence without authority?", "I did not have formal authority over every group, so I created structure: named owners, working sessions, open issue tracking, escalation paths, and a cadence that made progress visible."),
                ("L4 - conflicting groups", "What did each group want?", "The customer wanted relief and ownership. Product and development needed feasible scope and technical clarity. Account stakeholders needed risk reduced quickly. I translated those needs into a recovery plan each group could act on."),
                ("L5 - difficult client version", "How would you tell this as a difficult client story?", "I would lead with the relationship read: the customer was not angry about software alone; they were angry because nobody had owned the problem. My first job was to make them feel heard and then make the recovery path credible."),
                ("L6 - opposing views version", "How would you tell this as opposing stakeholder views?", "I would lead with the discovery that each group was protecting a different risk. Once I named that, we could stop treating disagreement as resistance and start building a shared recovery plan."),
                ("L7 - claims bridge", "How does this map to State Farm?", "Claims workflows can break customer trust when ownership is unclear. The Process Engineer may not own every team, but can create visibility, accountability, and a better operating rhythm."),
            ),
        },
        "dashboards": {
            "label": "200+ Dashboards And Decision Visibility",
            "use_for": "data analysis, metrics, reporting, decision support, executive communication, data exercise",
            "noticing": "What I noticed was that nobody was asking the first question before requesting a report: what decision needs to be made, and what would change the answer?",
            "bridge": "For claims and task metrics, the discipline is identical: define the decision before touching the data. The question is not only what cycle time shows. It is what decision the cycle time data needs to support and whether the segmentation makes that obvious.",
            "calibration": "Is defining the business decision before touching the data a discipline the team is still building, or is that infrastructure mature?",
            "drills": (
                ("L1 - skeleton", "What did you build?", "I built 200+ dashboards, KPI reports, and analytics tools that improved visibility into operational performance, customer experience metrics, and process gaps."),
                ("L2 - decision-first", "How did you decide what to build?", "I started with the operating decision, not the tool. I asked what leaders needed to decide, which metric would change the answer, and whether the source data was trustworthy enough to support that decision."),
                ("L3 - data quality", "How did you validate data?", "I checked source definitions, reconciliation points, duplicate or missing records, and whether the metric meant the same thing to different teams."),
                ("L4 - segmentation", "How did segmentation matter?", "A single average can hide the process problem. I looked for cuts by workflow, customer type, site, owner, timing, or exception path so the recommendation pointed to a controllable lever."),
                ("L5 - executive communication", "How did you present it?", "I kept the view decision-oriented: what changed, why it matters, what action it supports, and what risk remains if the team does nothing."),
                ("L6 - data exercise", "How would this show up in State Farm's data exercise?", "I would define the business question, validate claim/task/phone metric definitions, segment before concluding, state assumptions, then recommend a pilot with primary and guardrail metrics."),
                ("L7 - learning", "What would you do differently?", "I would document metric definitions even more explicitly up front because a trusted dashboard is as much a shared language artifact as a technical product."),
            ),
        },
        "learning": {
            "label": "Rapid Product And Domain Learning",
            "use_for": "no direct claims background, ramping in new domain, humility, learning process, domain credibility",
            "noticing": "What I noticed was that every support or implementation issue was a case study in how the customer's workflow actually operated.",
            "bridge": "That is exactly how I would close the claims knowledge gap: learn through the workflow by being inside it, not through documentation alone. Every conversation with an adjuster is a case study in how the process actually operates versus how it is supposed to operate.",
            "calibration": "Is a 30-to-60 day ramp where I learn the claims environment before recommending anything the right expectation for how someone starts in this role?",
            "drills": (
                ("L1 - skeleton", "What is the story?", "I became credible in complex manufacturing software environments by learning through client problems, not by pretending I already knew every domain."),
                ("L2 - first moves", "What did you do first in a new domain?", "I asked workflow questions, traced real examples, listened for repeated pain points, and translated issues into requirements or process decisions."),
                ("L3 - humility", "How do you avoid overclaiming?", "I separate what I know from what I need to validate. I can own the learning process without pretending to know claims on day one."),
                ("L4 - SME learning", "How would you learn claims?", "Shadow the workflow, ask adjusters to walk through normal and exception cases, learn the metrics they trust, and validate hypotheses before recommending changes."),
                ("L5 - credibility", "How do you build credibility quickly?", "By producing something useful early: a clean process map, clarified metric definition, issue log, or decision document that makes the team's work easier."),
                ("L6 - risk", "What is the risk in a new domain?", "The risk is recommending from a shallow understanding. I manage that by labeling assumptions, validating with SMEs, and piloting before scaling."),
                ("L7 - State Farm bridge", "How do you say this in the interview?", "I would not claim claims expertise. I would claim a disciplined way to learn claims operations and turn that learning into process improvement."),
            ),
        },
        "ops_finance": {
            "label": "Operations Versus Finance Alignment",
            "use_for": "opposing views, competing priorities, compliance versus speed, stakeholder tension, tradeoff communication",
            "noticing": "What I noticed was that both sides were defending against different risks. Operations was protecting speed and flow. Finance was protecting control and auditability.",
            "bridge": "That same tension exists in claims process engineering whenever a workflow affects speed, quality, compliance, and customer experience simultaneously. The job is to surface what each group is protecting before designing something that satisfies all of them.",
            "calibration": "Is navigating competing-priority tension, where process and compliance teams define success differently, a regular part of what this role manages?",
            "drills": (
                ("L1 - skeleton", "What was the conflict?", "Operations and finance had competing priorities around cost, timeline, workflow impact, and control. I clarified the tradeoffs and helped stakeholders compare options instead of arguing past each other."),
                ("L2 - noticing", "What did you notice beneath the disagreement?", "They were not disagreeing about effort. They were protecting different business risks, and neither side trusted that the other risk was being taken seriously."),
                ("L3 - method", "How did you make progress?", "I named the tradeoff, separated facts from assumptions, documented options, and made cost, timeline, risk, and operational impact visible."),
                ("L4 - language", "What did you say in the room?", "I would frame it like: I do not think we have one right answer yet. I think we have competing risks. Let's name what each option protects and what each option exposes."),
                ("L5 - resistance", "What if one side still resisted?", "I would ask what condition would make the option acceptable. That turns a hard no into design criteria."),
                ("L6 - claims example", "How does this map to claims?", "A claims process change may improve speed but create quality risk, workload imbalance, compliance exposure, or customer confusion. The role has to make that visible."),
                ("L7 - learning", "What did this teach you?", "Stakeholder disagreement is often useful data. It tells you which risk the process design has not addressed yet."),
            ),
        },
        "failure": {
            "label": "Failure, Validation, And Quality Control",
            "use_for": "failure question, growth mindset, quality, validation, risk control, lessons learned",
            "noticing": "What I noticed was that quality risk was entering after the design decision, when correction was more expensive and adoption friction was higher.",
            "bridge": "In claims, the stakes of that lesson are higher than in manufacturing: a process error can become a customer experience failure, compliance exposure, or financial risk. That is why validation has to be built into the process, not inspected in at the end.",
            "calibration": "Is building SME validation into every process change before go-live a strong discipline here already, or an area to strengthen?",
            "drills": (
                ("L1 - skeleton", "What failed?", "A rollout exposed that a validation checkpoint was too late. The issue was recoverable, but it created avoidable friction."),
                ("L2 - ownership", "What do you own in the failure?", "I own that I should have moved SME validation earlier and clarified acceptance criteria before the decision became expensive to unwind."),
                ("L3 - internal moment", "What made the lesson land?", "The hard moment was not fixing the issue. It was realizing, after hearing the operations lead walk through what did not work, that the people who would have caught the issue were not the ones who had run the test."),
                ("L4 - response", "What did you do when it surfaced?", "I helped triage the issue, gathered the right stakeholders, clarified what had to be corrected, and adjusted the validation rhythm for future work."),
                ("L5 - learning", "What changed afterward?", "Earlier SME reviews, clearer acceptance criteria, stronger sandbox or pilot validation, and post-launch checks became part of my operating method."),
                ("L6 - growth mindset", "How do you avoid sounding defensive?", "Keep the failure brief and spend more time on the changed behavior. Say what you learned and what you now do differently."),
                ("L7 - State Farm bridge", "How does this apply to claims?", "Claims processes need guardrail metrics and validation before full rollout because speed gains cannot come at the expense of quality or customer trust."),
            ),
        },
        "workshops": {
            "label": "Executive Workshops And QBR Alignment",
            "use_for": "executive communication, facilitation, leadership updates, tollgate materials, alignment",
            "noticing": "What I noticed was that each audience needed a different version of the same truth: leaders needed decisions, delivery teams needed next actions, and users needed how the change affected their work.",
            "bridge": "For State Farm, that maps to tollgate materials and execution leadership communication: make the finding clear, show the evidence, name the decision, and make the next action easy to approve.",
            "calibration": "How often does this role need to translate the same process change differently for leaders, analysts, and frontline teams?",
            "drills": (
                ("L1 - skeleton", "What is the story?", "I facilitated 60+ executive workshops and QBRs to align stakeholders around roadmap, adoption, business outcomes, and delivery risk."),
                ("L2 - audience", "How did you adapt by audience?", "I changed the framing: executives got decision and risk, delivery teams got scope and sequence, users got workflow impact and support."),
                ("L3 - conflict", "What if stakeholders disagreed?", "I surfaced the decision criteria, named tradeoffs, and kept the conversation tied to the business outcome rather than preferences."),
                ("L4 - materials", "What materials did you create?", "Agendas, roadmap views, status summaries, decision notes, adoption materials, and follow-up actions with owners."),
                ("L5 - State Farm bridge", "How does this map to tollgate slides?", "A tollgate deck should answer: what did we find, why does it matter, what do we recommend, what changes operationally, and how will we measure it?"),
                ("L6 - executive presence", "How do you sound concise?", "Lead with the recommendation, then provide evidence only as needed. Do not build the clock."),
                ("L7 - improvement", "What would you improve now?", "I would make decision asks even more explicit: approve, reject, revise, or request more data."),
            ),
        },
    }
    fallback = {
        "label": card.title,
        "use_for": ", ".join(card.story_types),
        "noticing": state_farm_story_noticing(card),
        "bridge": state_farm_story_bridge(card),
        "calibration": state_farm_calibration_question(card),
        "drills": tuple((f"L{i}", "What detail should you prepare?", "Prepare one concrete detail about the action, decision, stakeholder, metric, resistance, or lesson.") for i in range(1, 8)),
    }
    data = base.get(key, fallback).copy()
    data["key"] = key
    data["card"] = card
    return data

def add_state_farm_story_deep_dive(document: Document, data: dict[str, object]) -> None:
    card: cheat.StoryCard = data["card"]  # type: ignore[assignment]
    add_page_break(document)
    add_section(document, f"Story Deep Dive - {data['label']}")
    add_subsection(document, "When To Use This Story")
    add_body(document, str(data["use_for"]))
    add_subsection(document, "Story Card")
    for line in (
        f"Hook: {card.hook}",
        f"Noticing: {data['noticing']}",
        f"Action: {card.evidence}",
        f"Result: {card.result}",
        f"Bridge: {data['bridge']}",
        f"Calibration: {data['calibration']}",
    ):
        add_bullet(document, line)
    add_subsection(document, "60-Second Spoken Model")
    add_body(document, f"{card.hook} {data['noticing']} {card.evidence} The result was {card.result.lower()} {data['bridge']}")
    add_subsection(document, "90-Second Spoken Model")
    add_body(document, f"The way I would tell this story is to start with the process tension: {card.hook} The important part is what I noticed: {data['noticing']} From there, I focused on the controllable path: {card.evidence} That produced a concrete result: {card.result} For State Farm, the reason I would use this story is specific: {data['bridge']} {data['calibration']}")
    add_subsection(document, "What Not To Say")
    for line in (
        "Do not turn this into a software story. Translate every technical action into process, data, quality, adoption, or customer impact.",
        "Do not use a generic bridge sentence. The bridge must match this story's specific lesson.",
        "Do not over-answer the first question. Give the clean story first, then use the drill-down layers only when asked.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Drill-Down Questions And Answers")
    for layer, question, answer in data["drills"]:  # type: ignore[index]
        add_subsection(document, f"{layer}: {question}")
        add_body(document, answer)
    add_subsection(document, "Likely Follow-Up Pushbacks")
    for line in (
        "If they ask how you know the result was caused by your change: explain the baseline, the changed workflow, and the before/after metric.",
        "If they ask what others did: give credit briefly, then clarify Christian's specific ownership.",
        "If they ask how it applies to claims: translate into current-state mapping, approved metrics, process segmentation, service/quality/customer impact, and pilot measurement.",
        "If they ask what you learned: name one behavior you changed afterward, not a vague lesson.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Practice Reps")
    for line in (
        "Rep 1: answer in 45 seconds with only hook, noticing, result, and bridge.",
        "Rep 2: answer in 90 seconds with one specific action detail.",
        "Rep 3: answer the deepest drill-down question without adding unrelated resume history.",
        "Rep 4: record yourself and check whether the noticing line sounds like judgment rather than narration.",
        "Rep 5: state the result, hold eye contact for two seconds, then stop talking.",
    ):
        add_bullet(document, line)

def add_state_farm_story_workbook(document: Document, hero_stories: list[cheat.StoryCard], resume_text: str) -> list[dict[str, object]]:
    add_page_break(document)
    add_section(document, "Primary Story Workbook")
    add_body(document, "This is the heart of the workbook. Each story has a surface answer and seven prepared drill-down layers. The panel can stay on one story for a long time. These layers prevent Christian from running out of detail or inventing detail under pressure.")
    story_cards = cheat.supported_story_bank(resume_text)
    priority_keys = ["inventory", "account", "dashboards", "learning", "ops_finance", "failure", "workshops"]
    selected: list[cheat.StoryCard] = []
    for key in priority_keys:
        match = next((card for card in story_cards if state_farm_story_key(card.title) == key), None)
        if match:
            selected.append(match)
    for card in hero_stories:
        if card not in selected and len(selected) < 8:
            selected.append(card)
    data_items = [state_farm_story_workbook_data(card) for card in selected]
    add_subsection(document, "Story Selection Rule")
    add_body(document, "Relevancy beats recency. Use the story that mirrors the State Farm problem, not the story that happened most recently. The best story is the one that proves the exact behavior they are testing.")
    for data in data_items:
        add_bullet(document, f"{data['label']}: {data['use_for']}")
    for data in data_items:
        add_state_farm_story_deep_dive(document, data)
    return data_items

def add_state_farm_master_qa_workbook(document: Document, story_data: list[dict[str, object]]) -> None:
    add_page_break(document)
    add_section(document, "Master Q&A Workbook")
    add_body(document, "This replaces overlapping question sections. Each entry shows what the interviewer is really testing, which story to use, how to answer, what follow-up may come next, and how to practice it.")
    by_key = {str(item["key"]): item for item in story_data}

    def story_label(key: str) -> str:
        return str(by_key.get(key, {}).get("label", "Use the strongest matching story"))

    categories = (
        ("Opening And Positioning", (
            ("Tell me about yourself.", "Can Christian explain the career story without sounding scattered?", "career pattern", "Use the career pattern: process gaps others worked around. Home Depot for reporting visibility, Aptean for learning workflows through problems, East West for root-cause improvement, State Farm for claims operations at customer scale.", "Ask whether the role's biggest need is more diagnosis, implementation, or stakeholder adoption."),
            ("Why State Farm?", "Is the interest authentic or manufactured?", "authenticity", "Use the 50/50 rule: State Farm scale and customer stakes first, then Christian's process/data/adoption fit. Avoid trying to sound uniquely passionate about insurance.", "Ask what success looks like in the first 12 months."),
            ("Why this Process Engineer role?", "Does Christian understand the actual job?", "role", "Name the JD signals: process design, Lean Six Sigma-style improvement, approved operational metrics, tollgate materials, cost/benefit, agile priorities, and collaboration.", "Ask which part of the role is most urgent right now."),
            ("What is your biggest transferable strength?", "Can Christian simplify his value?", "inventory", "My strongest transferable strength is turning recurring operational problems into structured, measurable workflows that people can keep using.", "Ask whether the team needs more help with diagnosis or implementation."),
        )),
        ("Process Engineering Core", (
            ("Tell me about a time you improved an inefficient process.", "Initiative, root cause, measurable improvement.", "inventory", "Lead with 78% manual-work reduction and 22% discrepancy reduction. Explain current-state mapping, structural gap, validation with users, workflow improvement, and before/after measurement.", "Expect: how did you map it, how did you measure it, what resistance did you see?"),
            ("How do you identify waste in a process?", "Does Christian understand Lean-style thinking?", "inventory", "Start by observing the actual workflow, not the documented one. Look for rework, waiting, duplicate entry, unclear ownership, avoidable handoffs, exception paths, and metrics that show variation.", "Expect: what if the data and SME feedback disagree?"),
            ("How do you make process improvements stick?", "Implementation and adoption, not just analysis.", "failure", "A process sticks when the people doing the work understand it, the metric proves it, the owner can maintain it, and a feedback loop catches drift. Use validation and control checks.", "Expect: what did you do when adoption was weak?"),
            ("How do you balance efficiency, quality, and customer experience?", "Tradeoff judgment.", "ops_finance", "I do not optimize one metric in isolation. I name the tradeoff, define primary and guardrail metrics, pilot the change, and validate whether speed is improving without quality or customer experience deteriorating.", "Expect: how would this apply to claims?"),
            ("How would you approach a process you do not know yet?", "Humility, learning method, claims ramp.", "learning", "Observe before analyzing, analyze before recommending. Shadow real cases, label assumptions, learn trusted metrics, validate with SMEs, and produce a useful map or issue log early.", "Expect: how quickly can you become productive?"),
        )),
        ("Data And Analysis", (
            ("Describe a time you used data to drive a process decision.", "Decision-first data thinking.", "dashboards", "Use the dashboard story. The noticing line is that report requests were really decision requests. Explain source validation, segmentation, KPI view, and leadership decision support.", "Expect: how did you know the data was accurate?"),
            ("How would you handle a data exercise with limited context?", "Structured ambiguity.", "dashboards", "Clarify the business question, check data quality, define the unit of analysis, segment using MECE categories, identify controllable levers, recommend a pilot, and state assumptions.", "Expect: what would you do if your hypothesis is wrong?"),
            ("What metrics would matter in claims process work?", "Can Christian translate metrics to claims?", "dashboards", "Cycle time, task aging, first contact time, rework, handoff count, claim type, queue time, quality, service level, customer experience, cost/benefit, and guardrail metrics.", "Expect: how would you avoid optimizing the wrong metric?"),
            ("How do you present findings to leadership?", "Executive communication.", "workshops", "Use What/So What/Now What. What did you find, why does it matter for service, quality, cost, or customer experience, and what decision should leadership make next?", "Expect: what if they challenge your recommendation?"),
        )),
        ("Stakeholder And Influence", (
            ("Tell me about influencing without authority.", "Can Christian move people he does not manage?", "account", "Use the $1M account method version: consolidate ownership, structure sessions, name owners, and create a decision rhythm. The noticing line is one accountable owner, not more status.", "Expect: what authority did you actually have?"),
            ("Tell me about opposing stakeholder views.", "Can Christian discover the real tension?", "ops_finance", "Use Operations/Finance or the $1M opposing-views version. The noticing line is that people were protecting different risks, not simply disagreeing.", "Expect: what did you say in the room?"),
            ("Tell me about a difficult client or customer.", "Relationship read and recovery.", "account", "Use the difficult-client version: the customer was not angry about software alone; they were angry because nobody owned the problem. Listen, define what fixed means, create ownership, and communicate reliably.", "Expect: how did you avoid taking it personally?"),
            ("How do you build cross-functional trust?", "Collaboration habits.", "workshops", "Learn what each group is measured on, communicate in their language, make tradeoffs visible, and follow through on decisions and owners.", "Expect: how do you handle someone who still resists?"),
        )),
        ("Failure, Ambiguity, And Growth", (
            ("Tell me about a process improvement that did not go as planned.", "Coachability and learning.", "failure", "Keep the failure brief. Own the validation miss. Spend most of the answer on what changed: earlier SME validation, acceptance criteria, guardrail metrics, and post-launch checks.", "Expect: what would you do differently?"),
            ("Are you comfortable leading through ambiguity?", "Composure and structure.", "inventory", "PREP answer: yes, because ambiguity is where structure creates value. The inventory project started as a recurring problem without a clean root cause, and the value was building the structure around it.", "Expect: how do you prioritize when everything is ambiguous?"),
            ("How do you handle feedback or pushback?", "Coachability.", "failure", "Treat pushback as data. Pause, restate the challenge, say how the recommendation would change if the challenge is true, and identify what you would validate next.", "Expect: tell me about a time feedback changed your approach."),
            ("What is a weakness or gap for this role?", "Honesty without self-sabotage.", "learning", "I do not have direct claims experience, and I would not overclaim it. My plan is to close that gap through workflow shadowing, SME validation, metric learning, and useful early deliverables.", "Expect: why should we take that risk?"),
        )),
    )
    for category, questions in categories:
        add_page_break(document)
        add_section(document, category)
        for q_index, (question, hidden, story_key, answer, followup) in enumerate(questions):
            if q_index:
                add_page_break(document)
            add_section(document, question)
            add_body(document, f"Hidden assessment: {hidden}")
            if story_key in by_key:
                add_body(document, f"Primary story: {story_label(story_key)}")
                add_body(document, f"Noticing line to use: {by_key[story_key]['noticing']}")
                add_body(document, f"State Farm bridge: {by_key[story_key]['bridge']}")
            else:
                add_body(document, f"Primary angle: {story_key}")
            add_body(document, f"Model answer direction: {answer}")
            add_body(document, f"Likely follow-up: {followup}")
            add_subsection(document, "Workbook Practice")
            for practice_line in (
                "Rep 1: answer in 20-30 seconds. Use only the direct answer, one proof point, and one State Farm bridge.",
                "Rep 2: answer in 60-90 seconds. Add the noticing line and one action detail, then stop.",
                "Rep 3: answer the likely follow-up without restarting the original answer.",
                "Rep 4: write the one sentence you must not forget. This is the anchor if anxiety spikes.",
                "Rep 5: say the final metric or claim, hold camera eye contact for two seconds, and stop talking.",
            ):
                add_bullet(document, practice_line)
            add_subsection(document, "Score This Answer")
            for score_line in (
                "Did I classify the question type before answering?",
                "Did I lead with the point instead of context?",
                "Did I use one story rather than stacking examples?",
                "Did I include a role-specific State Farm bridge?",
                "Did I avoid unsupported claims about direct claims experience?",
            ):
                add_bullet(document, score_line)

def add_state_farm_data_exercise_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Data Exercise Workbook")
    add_body(document, "The data exercise is a thinking test. They want to see whether Christian can define the business question, cleanly segment the data, avoid overclaiming, and recommend a practical pilot with measurable results.")
    add_subsection(document, "Eight-Step Analysis Flow")
    for line in (
        "1. Clarify the business question: speed, waste, cost, service, quality, workload, or customer experience.",
        "2. Define the unit of analysis: claim, task, call, adjuster, team, segment, process step, channel, or time period.",
        "3. Check data quality: missing values, duplicate records, outliers, inconsistent definitions, timestamp logic, and whether the metric is leading or lagging.",
        "4. Segment before concluding: claim type, complexity, channel, team, geography, age, volume, handoff point, queue stage, or exception path.",
        "5. Run the MECE test: categories should not overlap and should cover the full dataset.",
        "6. Identify controllable levers: rework, queue time, duplicate touches, unclear ownership, training gaps, routing logic, documentation gaps, approval delays.",
        "7. Recommend a pilot: scope, owner, timeline, control group if possible, success metric, guardrail metric, rollback path.",
        "8. Close with the leadership decision: approve pilot, collect more data, revise process standard, or escalate a dependency.",
    ):
        add_bullet(document, line)
    add_subsection(document, "MECE Segmentation Script")
    add_body(document, "'I segmented by claim type: simple auto, multi-vehicle, and disputed liability. These are mutually exclusive by definition and collectively cover the full dataset. I ruled out time-period segmentation because the pattern was consistent across all 13 weeks, which suggests a structural issue rather than seasonality.'")
    add_subsection(document, "What / So What / Now What Slide Logic")
    for line in (
        "What: multi-vehicle claims show a consistent 4-5 day routing delay.",
        "So What: that delay adds avoidable cycle time to hundreds of claims and puts the customer's first contact experience at risk.",
        "Now What: recommend a six-week routing pilot with a 24-hour service-level target and quality guardrails.",
    ):
        add_bullet(document, line)
    scenarios = (
        ("Scenario 1 - Routing Delay", "Multi-vehicle claims are assigned later than simple claims, and first contact is delayed.", "Hypothesis: routing complexity and unclear ownership are creating queue time before first action.", "Pilot: define routing rules for the segment, create an owner queue, measure assignment-to-first-contact, and track quality guardrails."),
        ("Scenario 2 - Rework Spike", "A claim segment has normal cycle time but unusually high rework.", "Hypothesis: the process is fast but not clean; documentation, training, or decision authority may be weak.", "Pilot: add a validation checklist or SME review at the highest-error step, then measure rework rate and cycle-time impact."),
        ("Scenario 3 - Phone Volume Pattern", "Call volume spikes after a claim-status handoff.", "Hypothesis: customers are not receiving clear status or next-step expectations at the handoff point.", "Pilot: revise handoff communication, create a status trigger, measure repeat calls, customer experience, and task aging."),
        ("Scenario 4 - Workload Imbalance", "One team has slower task closure despite similar claim volume.", "Hypothesis: complexity mix, assignment logic, training, or exception handling differs by team.", "Pilot: rebalance by complexity or create a support path for exception tasks. Guardrail quality and employee workload."),
    )
    for title, fact_pattern, hypothesis, pilot in scenarios:
        add_page_break(document)
        add_section(document, title)
        add_body(document, f"Fact pattern: {fact_pattern}")
        add_body(document, f"Likely hypothesis: {hypothesis}")
        add_body(document, f"Recommended pilot: {pilot}")
        add_subsection(document, "Panel Pushbacks And Answers")
        for push, answer in (
            ("Why not fix everything at once?", "Multiple simultaneous interventions make attribution impossible. I would start with the cleanest controllable lever, measure it, then scale or revise."),
            ("What if this is just seasonality?", "I would compare across weeks and segments. If the pattern persists across the period, I would treat it as structural; if it spikes by time, I would investigate capacity or seasonal drivers."),
            ("How do you know the metric is trustworthy?", "I would validate definitions, missing timestamps, duplicate records, and whether each team records the event the same way."),
            ("What if the pilot improves speed but hurts quality?", "That is why I would define guardrail metrics before launch. A faster process that creates rework or customer dissatisfaction is not a successful process improvement."),
        ):
            add_body(document, f"Pushback: {push}")
            add_body(document, f"Answer: {answer}")
        add_subsection(document, "Practice Drill")
        add_body(document, "Present this scenario in three minutes: recommendation first, evidence second, pilot third, risk/guardrail fourth. Then answer one pushback without defending emotionally.")

def add_state_farm_panel_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Live Panel And Case Presentation Workbook")
    add_subsection(document, "Coachability Protocol")
    for line in (
        "When they redirect mid-presentation: stop your planned sequence and say, 'Good direction. Let me go there.' Then answer what they asked.",
        "When they challenge a finding: say, 'That is a fair challenge. If that is true, my recommendation changes in this direction.' Then explain how.",
        "When you do not know: say, 'I do not have enough information to be confident there. My assumption is X. If that is wrong, I would validate Y.'",
        "When they push for more detail: answer the exact layer they requested. Do not restart the whole story.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Executive Presence Mechanics")
    for line in (
        "Emphasize key words inside sentences, not only the last word. Say 'I reduced manual work by seventy-eight percent' with emphasis on reduced and seventy-eight percent.",
        "After each metric, hold eye contact with the camera for two seconds before continuing.",
        "Use the silent count when pace accelerates: count one, two, three internally before the next sentence.",
        "End recommendation sentences with a downward tone. Do not lift your pitch at the end of a claim.",
        "Pause after the answer. The pause lets the interviewer pull for more and prevents over-repairing.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Case Presentation Skeleton")
    for line in (
        "Opening: My recommendation is X because Y is the highest-impact controllable lever.",
        "Evidence: The pattern appears in this segment, not evenly across the whole process.",
        "Ruling out: I checked whether time period, volume, or data quality explained it, and those did not fully account for the pattern.",
        "Pilot: I recommend a focused pilot with primary metric, guardrail metric, owner, timeline, and rollback path.",
        "Decision ask: approve the pilot, request one additional validation step, or choose a different lever based on the tradeoff.",
    ):
        add_bullet(document, line)

def add_state_farm_video_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "On-Demand Video Workbook")
    add_body(document, "Recorded answers punish wandering. There is no interviewer to interrupt, redirect, or rescue the answer. Use answer-first structure and a clean stop.")
    prompts = (
        ("Why State Farm?", "State Farm matters to me because claims processes affect real people at high-stakes moments. My fit is process improvement: map the work, find waste or variation, use data to prove the root cause, and implement changes people can adopt."),
        ("Describe a process improvement.", "I reduced manual inventory adjustment work by 78% and discrepancies by 22% by mapping the actual process, finding the structural gap, automating the workflow, and validating the change with users."),
        ("How do you prioritize?", "I prioritize by business/customer impact, risk, effort-to-value, stakeholder dependencies, and real deadlines. If something will slip, I surface it early so the team can adjust without losing trust."),
        ("How do you handle ambiguity?", "I handle ambiguity by creating structure: clarify the question, label assumptions, map the process, validate with SMEs, and choose a small measurable next step."),
        ("How do you work with teams?", "I start by learning what each group is protecting: speed, quality, compliance, cost, workload, or customer experience. Once the tradeoff is visible, teams can make better decisions."),
    )
    for prompt, model in prompts:
        add_subsection(document, prompt)
        add_body(document, f"Model: {model}")
        add_body(document, "Practice: record once at 45 seconds, once at 75 seconds. Watch sound-only for fillers, then sound-off for posture and eye contact.")
    add_subsection(document, "Video Setup Checklist")
    for line in (
        "Camera at eye level, torso visible, hands in the truth plane from waist to shoulders.",
        "Light source behind the camera facing you. Avoid fake or blurred backgrounds when possible.",
        "Audio matters more than video. Test the microphone and connection before the interview.",
        "Look through the camera lens, not at the screen. Put the interviewer window just below the camera if live.",
        "After stating a metric or major claim, hold eye contact with the camera lens for two full seconds before continuing.",
        "Use the silent one-two-three count both to replace fillers and to slow rhythm between major sentences.",
        "Business casual, no distracting background movement, and no last-minute setup changes.",
    ):
        add_bullet(document, line)

def add_state_farm_mock_dialogue_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Mock Interview Dialogue Workbook")
    add_body(document, "This section is intentionally dialogue-heavy. Practice it out loud. The goal is not to memorize every answer, but to hear how a strong answer sounds when the interviewer asks follow-ups, challenges the framing, or redirects the conversation.")
    dialogues = (
        (
            "Opening: Tell Me About Yourself",
            "Interviewer: Tell me about yourself.",
            "Christian: My background sits at the intersection of process improvement, operational data, and stakeholder adoption. The common pattern is that I find process gaps people have learned to work around and build something measurable that makes the issue stop recurring. At Home Depot that looked like reporting visibility. At Aptean it looked like learning client workflows through implementation problems. At East West it looked like reducing manual inventory adjustment work by 78% and discrepancies by 22%. That is why this State Farm Process Engineer role makes sense to me: it is the same operating pattern applied to claims processes where speed, quality, and customer experience matter at scale.",
            ("Interviewer: Why not just stay in systems work?", "Christian: The systems work was always the vehicle. The part I care most about is the process outcome: cleaner workflow, better data, less rework, stronger adoption, and clearer decisions."),
        ),
        (
            "Why State Farm",
            "Interviewer: Why State Farm?",
            "Christian: State Farm stands out because claims processes affect real people at moments when speed, clarity, and trust matter. I am not trying to manufacture a unique reason; what is authentic for me is that this is process improvement with real customer stakes. My background fits because I have repeatedly mapped broken workflows, used data to prove the issue, built practical fixes, and helped people adopt the change.",
            ("Interviewer: What specifically about our role interests you?", "Christian: The posting combines process design, Lean Six Sigma-style improvement, operational metrics, communication materials, cost/benefit thinking, and implementation. That is the mix I enjoy: not just analysis, but making the recommendation usable."),
        ),
        (
            "Process Improvement Story",
            "Interviewer: Tell me about a time you identified and improved an inefficient process.",
            "Christian: I reduced manual inventory adjustment work by 78% and lowered discrepancies by 22%. The issue started as a recurring warehouse accuracy problem, but what I noticed was that the team was treating it like an execution issue when the real problem was structural. The workflow depended on too many manual touches and did not make the error path visible. I mapped the current process, validated the actual work with users, built an automated and auditable workflow, and measured the before and after. For claims, I would use the same pattern: map what is actually happening, find the structural gap, validate it with the people doing the work, pilot the fix, and measure whether anything changed.",
            ("Interviewer: How did you know the result came from your change?", "Christian: I anchored the measurement to the pain points the operation already felt: manual effort and discrepancy rate. The workflow changed at the control point creating the issue, and the before/after movement was visible in those two metrics."),
        ),
        (
            "Data Decision Story",
            "Interviewer: Describe a time you used data to drive a process decision.",
            "Christian: I built 200+ dashboards and KPI tools, but the important part is how I approached them. What I noticed was that many report requests were really decision requests in disguise. Leaders were asking for data, but the missing question was: what decision needs to be made and what would change the answer? I validated source data, clarified metric definitions, and built views around the operating decision instead of the tool. For State Farm, I would apply that same discipline to claim, task, phone, and operational metrics: define the decision first, validate the data, then build the view that makes the right action obvious.",
            ("Interviewer: What if the data is incomplete?", "Christian: I would say that directly. I would separate what the data can support from what it cannot, state the assumption, and recommend the next validation step before making a strong claim."),
        ),
        (
            "Influence Without Authority",
            "Interviewer: Tell me about influencing without authority.",
            "Christian: I would use a $1M+ account stabilization example. The account had unresolved workflow and integration issues, and multiple groups were involved. What I noticed was that the customer did not need another status email; they needed one accountable recovery path. I did not formally manage every team involved, so I created structure: single-point ownership, working sessions, named owners, open issues, and an escalation rhythm. That stabilized more than $1M in at-risk revenue. In claims, I think the same dynamic can happen when ownership is unclear across teams. A Process Engineer may not own every group, but can create the operating structure that lets the work move.",
            ("Interviewer: What authority did you actually have?", "Christian: Mostly influence and structure. I had credibility because I understood the issue and stayed close to the recovery path, but the power came from making ownership visible and giving each group a practical next action."),
        ),
        (
            "Difficult Customer",
            "Interviewer: Tell me about a difficult customer or client.",
            "Christian: The customer was frustrated, but the key was not to treat it as a personality issue. What I noticed was that they were angry because nobody had owned the problem end to end. I listened first, clarified what fixed actually meant to them, consolidated the open items, and gave them a more reliable communication rhythm. The relationship improved because they could see ownership and progress, not just explanations. For State Farm, that matters because when a customer's claims experience is breaking down, the process needs ownership just as much as it needs information.",
            ("Interviewer: How did you avoid taking it personally?", "Christian: I treated frustration as information. If someone is that frustrated, the process has probably made them repeat themselves or wait too long. My job is to find the friction and create a credible path forward."),
        ),
        (
            "Opposing Stakeholder Views",
            "Interviewer: Tell me about a time stakeholders disagreed.",
            "Christian: In one situation, operations and finance were using the same words but protecting different risks. Operations cared about speed and flow. Finance cared about control and auditability. What I noticed was that the disagreement was not really about effort; it was about which risk was being prioritized. I named the tradeoff, separated facts from assumptions, and helped the group compare options by cost, timeline, risk, and operational impact. For claims process engineering, I would expect similar tension between speed, quality, compliance, workload, and customer experience.",
            ("Interviewer: What did you say to lower the temperature?", "Christian: I would say something like, 'I do not think we have one right answer yet. I think we have competing risks. Let's name what each option protects and what each option exposes.'"),
        ),
        (
            "Failure And Learning",
            "Interviewer: Tell me about a process improvement that did not go as planned.",
            "Christian: I would choose a validation miss rather than a negligence story. In a complex rollout, what I underestimated was how late a key stakeholder checkpoint was happening. The issue was recoverable, but it created avoidable friction because feedback came after the design decision was already harder to unwind. What I changed afterward was specific: earlier SME validation, clearer acceptance criteria, and stronger post-launch checks. In claims, that lesson matters because a process error can become a customer experience failure, compliance exposure, or financial risk.",
            ("Interviewer: What did you personally own?", "Christian: I own that I should have moved validation earlier. The lesson was not 'work harder.' It was design the validation rhythm so risk surfaces before the expensive decision point."),
        ),
        (
            "No Claims Experience",
            "Interviewer: You do not have direct claims experience. Why should we take that risk?",
            "Christian: That is fair, and I would not overclaim claims experience. What I do have is a repeatable method for learning operational environments quickly: shadow the workflow, learn the metrics people trust, validate with SMEs, and turn that learning into useful process documentation or analysis. At Aptean, I became credible with manufacturing clients by learning through their problems rather than through documentation alone. I would bring that same discipline to claims.",
            ("Interviewer: How would you ramp in the first 30 days?", "Christian: I would shadow normal and exception cases, learn claim/task/phone metric definitions, map one priority workflow, and validate what I am seeing with frontline SMEs before recommending changes."),
        ),
        (
            "Lean Six Sigma Gap",
            "Interviewer: What is your Lean Six Sigma experience?",
            "Christian: I would be precise. I have applied Lean Six Sigma-style methods through current-state mapping, waste reduction, root-cause analysis, pilot validation, before/after measurement, and adoption work. I would not claim a belt I do not hold, but the operating method is very familiar: identify waste or variation, find the controllable lever, test the improvement, and measure results.",
            ("Interviewer: What Lean concept do you use most?", "Christian: Current-state mapping and root cause. Most weak fixes come from solving the documented process instead of the actual process."),
        ),
        (
            "Prioritization",
            "Interviewer: How do you prioritize multiple process improvement projects?",
            "Christian: I use a simple filter: business or customer impact, risk or compliance urgency, effort-to-value, stakeholder dependencies, and real deadlines. I like fast-moving environments, but I do not rely on energy alone. I make the tradeoffs visible, review priorities consistently, and if something will slip, I surface it early to the impacted stakeholder.",
            ("Interviewer: What if everything is urgent?", "Christian: Then I separate urgency from consequence. What creates customer, compliance, financial, or downstream team risk comes first; everything else gets sequenced by effort-to-value and dependency."),
        ),
        (
            "Ambiguity",
            "Interviewer: Are you comfortable leading through ambiguity?",
            "Christian: Yes. I actually do well in ambiguity because that is where structure creates value. The inventory adjustment project did not start with a clean root cause or defined fix; it started as a recurring problem. My job was to map the work, find the structural issue, validate it, and build the process around it. That is the kind of environment where I do strong work.",
            ("Interviewer: How do you avoid moving too slowly?", "Christian: I use pilots. A small test with a clear metric and guardrail lets the team move without pretending uncertainty is zero."),
        ),
        (
            "Leadership Style",
            "Interviewer: How do you lead?",
            "Christian: My approach is to build capability rather than dependency. I want the people I work with to be stronger after the engagement. In practice, I start by understanding how each person communicates and what motivates them, set clear goals with individual accountability, and reflect inward first when something goes wrong. In claims process engineering, that means I am not just managing the process change; I am helping the people who use the process understand and execute the new standard confidently.",
            ("Interviewer: How do you lead without a title?", "Christian: By creating clarity: shared goal, named owners, visible risks, and a practical rhythm for decisions."),
        ),
        (
            "Coachability",
            "Interviewer: What would you do if the panel challenged your recommendation?",
            "Christian: I would treat the challenge as useful information, not as an attack. I would say, 'That is a fair challenge. If that is true, my recommendation changes in this direction.' Then I would explain what I would validate next. In a data exercise, credibility comes from showing how your thinking changes when new information matters.",
            ("Interviewer: What if you are wrong?", "Christian: Then I want to find out quickly. That is why I prefer pilots, guardrail metrics, and explicit assumptions."),
        ),
        (
            "Data Exercise Start",
            "Interviewer: Here is a dataset. How would you start?",
            "Christian: First I would clarify the business question. Are we trying to improve speed, quality, service, cost, workload, or customer experience? Then I would define the unit of analysis, check data quality, segment before concluding, and look for controllable levers. I would avoid jumping straight to a recommendation before I know whether the metric is trustworthy and what decision the analysis needs to support.",
            ("Interviewer: Why not start calculating immediately?", "Christian: Because calculations can be accurate and still irrelevant if they answer the wrong business question."),
        ),
        (
            "MECE Segmentation",
            "Interviewer: How would you segment claims data?",
            "Christian: I would choose categories that are mutually exclusive and collectively exhaustive. For example, if claim type is the driver, I might segment simple auto, multi-vehicle, and disputed liability so no claim sits in two groups and the dataset is covered. Then I would test whether the pattern holds across time, team, volume, and complexity before calling it structural.",
            ("Interviewer: What if categories overlap?", "Christian: Then the segmentation is not clean enough to support a recommendation. I would redefine the unit of analysis or choose a different cut."),
        ),
        (
            "Presenting Findings",
            "Interviewer: Present your recommendation to the panel.",
            "Christian: My recommendation is a focused routing pilot for the highest-delay claim segment. What I found is that multi-vehicle claims show a consistent routing delay. So what that means is avoidable cycle time before first contact, which can hurt customer experience. Now what I recommend is a six-week pilot with a 24-hour service-level target, quality guardrails, and a rollback path if rework increases.",
            ("Interviewer: Why this pilot first?", "Christian: It is the cleanest controllable lever. If we change too many things at once, we may improve the metric but not know why."),
        ),
        (
            "Accuracy And Documentation",
            "Interviewer: How do you ensure accuracy when redesigning a process?",
            "Christian: I do not redesign from memory. I document the current state by observing the real process, validate it with SMEs, identify where the process differs from the documented standard, and define success and guardrail metrics before rollout. After implementation, I want a 30/60/90-day check to catch drift or unintended consequences.",
            ("Interviewer: Why does that matter in insurance?", "Christian: Because process errors can become customer experience issues, compliance exposure, or financial risk."),
        ),
        (
            "Explaining Complexity",
            "Interviewer: How would you explain a complex process change to a nontechnical audience?",
            "Christian: I would start with what they already care about, then explain the change in terms of their workflow. I avoid jargon and use SCR: situation, complication, resolution. I also check understanding by asking them to walk back how the change affects their team, because 'does that make sense?' often gets a polite yes even when the change did not land.",
            ("Interviewer: What if they still do not understand?", "Christian: I change the explanation, not the audience. I would use a simpler workflow example or show the before and after visually."),
        ),
        (
            "Team Culture",
            "Interviewer: How do you contribute to a collaborative culture?",
            "Christian: I like being part of high-performing teams, especially when the work is complex. I try to contribute by being reliable, helping teammates when they hit a wall, documenting what I learn so others can use it, and making disagreements easier to resolve by naming the tradeoff instead of personalizing it.",
            ("Interviewer: What if a teammate is struggling?", "Christian: I would first understand whether the issue is clarity, workload, skill, or motivation. The response depends on the actual constraint."),
        ),
        (
            "Customer Service Lens",
            "Interviewer: What customer service skills matter in this role?",
            "Christian: Even though this is a process role, customer service still matters because claims processes shape customer experience. The core skill is meeting the customer where they are: listen, understand the friction, solve what can be solved, and design the process so the same issue is less likely to happen again.",
            ("Interviewer: How does that connect to process engineering?", "Christian: A good process protects the customer from having to chase, repeat, wait, or interpret confusing handoffs."),
        ),
        (
            "Closing Ask",
            "Interviewer: What questions do you have for us?",
            "Christian: I have a few, but the one I care about most is: what does this person need to accomplish in the first 12 months for you to consider the hire a success?",
            ("Interviewer: Anything else?", "Christian: Yes. What is the most challenging claims process problem your team is working on right now? And do you have any concerns about my ability to do this job that I can address directly?"),
        ),
    )
    for dialogue_index, (title, prompt, answer, follow) in enumerate(dialogues):
        if dialogue_index:
            add_page_break(document)
        add_section(document, title)
        add_body(document, prompt)
        add_body(document, f"Model answer: {answer}")
        add_subsection(document, "Follow-Up Pressure")
        add_body(document, follow[0])
        add_body(document, f"Model response: {follow[1]}")
        add_subsection(document, "Practice Drill")
        for line in (
            "Say the answer once naturally without looking.",
            "Repeat it in half the time without losing the point.",
            "Answer the follow-up directly without retelling the whole story.",
            "Identify the one phrase that must stay in the answer.",
        ):
            add_bullet(document, line)

def add_state_farm_questions_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Questions, Rapport, And Closing Workbook")
    add_subsection(document, "Tell Me More As A Strategic Tool")
    add_body(document, "When an interviewer says something revealing, do not rush to the next prepared question. Use the moment. Say 'Tell me more about that,' 'What specifically makes that challenging?', or 'How long has that been a friction point for the team?' This gives you the raw material for your close and thank-you note.")
    add_subsection(document, "Primary Questions To Ask")
    for line in (
        "What does this person need to accomplish in the first 12 months for you to consider the hire a success?",
        "What is the most challenging claims process problem your team is working on right now?",
        "Which operational metrics does the team trust most today, and which still need better visibility or validation?",
        "How mature is the process improvement culture here: bottom-up, top-down, or both?",
        "Where do process, execution, and business partners most often define success differently?",
        "What tools, templates, or methodologies does the team use for tollgate materials and process analysis?",
        "What behaviors separate the most effective people on this Claims Process Engineering team?",
        "Do you have any concerns about my ability to do this job that I can address directly?",
    ):
        add_bullet(document, line)
    add_subsection(document, "Closing Summary")
    add_body(document, "What I am taking away is that this role needs someone who can improve claims processes with structure, data, and adoption discipline. That lines up with the work I have done: finding recurring process gaps, proving them with metrics, building practical fixes, and helping stakeholders use the new process. I am very interested in the role, and I would welcome the chance to bring that operating style to State Farm.")

def add_state_farm_deep_questions_to_ask_bank(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Deep Questions To Ask Bank")
    add_body(document, "Do not ask all of these. Pick six to eight based on the conversation. The purpose is to learn what the role actually needs, signal operating maturity, and gather language for the closing summary and thank-you note.")
    banks = (
        ("Success And Expectations", (
            "What does this person need to accomplish in the first 12 months for you to consider the hire a success?",
            "What would make you confident after 90 days that the person in this role is ramping the right way?",
            "Which outcome matters most in the first year: efficiency, service, quality, cost, customer experience, or adoption?",
            "Where have previous process improvement efforts succeeded or struggled?",
            "What would be a visible early win for this role without creating unnecessary risk?",
        )),
        ("Current Process Pain", (
            "What is the most challenging claims process problem your team is working on right now?",
            "Where does waste show up most visibly today: waiting, rework, duplicate touches, handoffs, unclear ownership, or inconsistent standards?",
            "Which processes are most dependent on tribal knowledge or informal workarounds?",
            "Where do process issues most directly affect customer experience?",
            "If you could fix one recurring process friction this quarter, what would it be?",
        )),
        ("Data And Metrics", (
            "Which claim, task, phone, or operational metrics does the team trust most today?",
            "Which metrics are debated because definitions or data quality are inconsistent?",
            "How does the team decide whether a metric is a leading indicator or a lagging indicator?",
            "Where does the team need better segmentation to understand root cause?",
            "What tools are currently used for process analysis, reporting, and tollgate materials?",
        )),
        ("Stakeholders And Change", (
            "Which partner groups are usually involved in process improvement work?",
            "Where do process, execution, and business partners most often define success differently?",
            "How does the team handle resistance when a process change affects frontline workflow?",
            "What communication style works best with execution leadership here?",
            "How does the team keep adoption from fading after rollout?",
        )),
        ("Methods And Operating Rhythm", (
            "How formally does the team use Lean Six Sigma methods versus adapting the tools pragmatically?",
            "What does a strong tollgate presentation usually include here?",
            "How are cost/benefit decisions made when there are multiple improvement opportunities?",
            "How does agile prioritization work for the Claims Process Engineering team?",
            "How often does the team run retrospectives or lessons learned during a live project?",
        )),
        ("Team Culture And Fit", (
            "What behaviors separate the most effective people on this team?",
            "What attracted you to State Farm, and what keeps you here?",
            "How would you describe the process improvement culture: more bottom-up, top-down, or both?",
            "What does strong collaboration look like on this team when priorities conflict?",
            "Do you have any concerns about my ability to do this job that I can address directly?",
        )),
    )
    for title, questions in banks:
        add_page_break(document)
        add_section(document, title)
        for question in questions:
            add_bullet(document, question)
        add_subsection(document, "Follow-Up Move")
        add_body(document, "If the interviewer gives a revealing answer, say: 'Tell me more about that.' Then ask what makes it difficult, how long it has been a friction point, or what success would look like if it improved.")

def add_state_farm_practice_plan_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Practice Plan")
    add_subsection(document, "Two-Day Sprint If The Interview Is Imminent")
    for line in (
        "Day 1, 4 hours total: Hour 1 write story skeleton cards in bullets only. Hour 2 record Tell Me About Yourself, watch twice, note one fix, record once more. Hour 3 practice Why State Farm and the closing sequence out loud. Hour 4 do one full data exercise walkthrough out loud. Then stop.",
        "Day 2, 3 hours total: Hour 1 reconstruct all story skeletons from memory. Hour 2 record a five-question mock interview. Hour 3 review the recording and fix one thing only. Evening is light review and one final recording, with a hard stop by 9pm.",
        "Morning of: no new content. Equipment check, breathing sequence, one skeleton read, and conversation mode.",
    ):
        add_bullet(document, line)
    add_subsection(document, "Five-Day Plan If There Is More Time")
    days = (
        ("Day 1 - Role And Story Foundation", ("Read role deconstruction", "Practice Tell Me About Yourself in 30/60/120 seconds", "Record Inventory story once", "Mark weak drill-down layers")),
        ("Day 2 - Process And Data", ("Practice data exercise flow", "Do two What/So What/Now What reps", "Practice MECE segmentation script", "Answer dashboard drill-down questions")),
        ("Day 3 - Stakeholders And Influence", ("Practice $1M story in three variants", "Practice Operations/Finance story", "Record one difficult-client answer", "Practice Tell me more responses")),
        ("Day 4 - Failure And Coachability", ("Practice failure story", "Answer three panel pushbacks", "Practice saying 'fair challenge' without defending", "Review gap answer for no claims experience")),
        ("Day 5 - Video And Final Polish", ("Record five on-demand prompts", "Watch sound-only for filler words", "Watch sound-off for gestures and eye contact", "Do final question list and closing ask")),
    )
    for day, tasks in days:
        add_subsection(document, day)
        for task in tasks:
            add_bullet(document, task)
    add_subsection(document, "Scorecard For Practice Answers")
    for line in (
        "Did the answer start with the answer or stakes within the first sentence?",
        "Did it include a noticing line that shows judgment?",
        "Did it use one story instead of stacking the whole resume?",
        "Did it include a metric or concrete result?",
        "Did the bridge mention claims, process, data, service, quality, customer experience, or stakeholder adoption?",
        "Did Christian stop cleanly instead of over-explaining?",
    ):
        add_bullet(document, line)

def add_state_farm_leadership_and_pushback_workbook(document: Document) -> None:
    add_page_break(document)
    add_section(document, "Leadership, Prioritization, And Pushback Workbook")
    add_body(document, "This section prepares for panel questions that are not pure stories: leadership style, priorities, coachability, gaps, and pressure. Each answer should be direct, calm, and specific.")
    entries = (
        ("Leadership style / How do you lead?", "My approach is to build capability rather than create dependency. I start by understanding how each person communicates and what motivates them, set clear goals with individual accountability, and reflect inward first when something goes wrong: was my direction clear, was the environment supportive, and were expectations realistic? In claims process engineering, that means adoption work: I am not just changing a process; I am helping the people who use the process execute it confidently.", "Practice: answer once as a 30-second philosophy and once as a 90-second answer with the workshops or Operations/Finance story."),
        ("How do you prioritize multiple process improvements?", "I prioritize by business/customer impact, risk or compliance urgency, effort-to-value, stakeholder dependencies, and real deadlines. I like fast-moving environments, but I do not rely on energy alone. I use a visible prioritization method and surface tradeoffs early if something is at risk.", "Practice: give one quick-win example and one structural-fix example."),
        ("No direct claims background", "That is accurate, and I would not overclaim it. The transferable pattern is that I learn operational environments through real workflows, SMEs, metrics, and exceptions. My first goal would be to understand the claims process deeply enough to make useful recommendations, not to pretend expertise on day one.", "Practice: say this without apologizing and then bridge to the Rapid Learning story."),
        ("Limited formal Lean Six Sigma certification", "I would be precise: I have applied Lean Six Sigma-style habits through current-state mapping, waste reduction, root-cause analysis, pilot validation, before/after measurement, and stakeholder adoption. I would also learn and use State Farm's preferred terminology and templates.", "Practice: name three Lean-style behaviors without claiming a belt you do not hold."),
        ("Too systems-heavy", "The systems were the vehicle. The business work was process improvement: reducing manual work, improving data quality, training users, documenting standard work, and helping teams adopt a better operating rhythm.", "Practice: translate one ERP bullet into process language."),
        ("Concern about ambiguity", "Ambiguity does not bother me because it gives structure room to create value. My method is to clarify the business question, label assumptions, map the actual process, validate with SMEs, and choose a measurable next step.", "Practice: answer this with PREP in under 20 seconds."),
        ("Concern about speed", "I move fastest when the next step is clear. If risk is unclear, I use a pilot-first approach: small test, clear metric, guardrail metric, and decision checkpoint. That lets the team move without pretending uncertainty is zero.", "Practice: connect this to data exercise pilot logic."),
        ("What if your recommendation is challenged?", "I would treat that as useful information. I would say: that is a fair challenge. If that is true, my recommendation changes this way. Then I would identify what to validate next.", "Practice: say this slowly enough that it sounds receptive, not defensive."),
        ("What did you dislike about your last role?", "Use the omission rule. Do not name something negative. Say what you wanted more of: direct customer contact, clearer metrics, more ownership of process outcomes, or more exposure to the business problem.", "Practice: keep this under 20 seconds."),
        ("Why should we not hire you?", "Treat it like why should we hire you, just inverted. Do not list a real weakness. Say: if you need someone with direct claims experience on day one, that is not my background. If you need someone who can learn the workflow, structure ambiguity, analyze metrics, and implement process improvements, that is exactly where my background fits.", "Practice: answer without sounding cute or evasive."),
    )
    for entry_index, (question, model, practice) in enumerate(entries):
        if entry_index:
            add_page_break(document)
        add_section(document, question)
        add_body(document, f"Model answer: {model}")
        add_body(document, practice)
        add_subsection(document, "Likely Follow-Up")
        add_body(document, "If they push deeper, give one concrete example and stop. Do not reopen the entire resume.")
        add_subsection(document, "Self-Check")
        for line in (
            "Did I answer the question directly?",
            "Did I avoid sounding defensive?",
            "Did I connect the answer to claims process engineering?",
            "Did I stop before over-explaining?",
        ):
            add_bullet(document, line)

def add_state_farm_full_workbook(
    document: Document,
    profile: build_resume.JobProblemProfile,
    company_name: str,
    role_title: str,
    job_description: str,
    resume_text: str,
    hero_stories: list[cheat.StoryCard],
    add_keyword_question_bank_callback: Callable[..., None],
) -> None:
    add_state_farm_workbook_cover_note(document)
    add_state_farm_role_deconstruction_workbook(document, job_description)
    add_state_farm_interview_process_map(document)
    add_state_farm_answer_operating_system(document)
    add_keyword_question_bank_callback(document, profile, job_description, resume_text, hero_stories, company_name, role_title)
    add_natural_storytelling_system(document)
    add_state_farm_core_positioning_lab(document)
    story_data = add_state_farm_story_workbook(document, hero_stories, resume_text)
    add_state_farm_master_qa_workbook(document, story_data)
    add_state_farm_mock_dialogue_workbook(document)
    add_state_farm_data_exercise_workbook(document)
    add_state_farm_panel_workbook(document)
    add_state_farm_leadership_and_pushback_workbook(document)
    add_state_farm_video_workbook(document)
    add_state_farm_questions_workbook(document)
    add_state_farm_deep_questions_to_ask_bank(document)
    add_state_farm_post_interview_strategy_deep(document)
    add_state_farm_practice_plan_workbook(document)
    add_page_break(document)
    add_state_farm_master_checklist(document)

def validate_state_farm_workbook_text(text: str) -> None:
    # Coverage surface includes state_farm_prep_insights() plus the full workbook sections below.
    forbidden = (
        "Jacob", "Anticipated Questions From Notes",
        "Default answer framework for this guide: CAAR", "SAR Quality Audit",
    )
    for term in forbidden:
        if term.lower() in text.lower():
            fail(f"State Farm workbook contains stale or duplicate guide language: {term}")
    repeated_bridge = "process waste, operational variation, data-backed root cause analysis, and measurable efficiency improvement"
    if repeated_bridge.lower() in text.lower():
        fail("State Farm workbook contains the old generic repeated story bridge")
    required = (
        "Workbook Promise", "Drill-Down Questions And Answers", "Master Q&A Workbook",
        "Data Exercise Workbook", "MECE", "What / So What / Now What", "PREP", "Tell Them The Time",
        "Tell Me More As A Strategic Tool", "Coachability Protocol", "Practice Plan",
    )
    for term in required:
        if term.lower() not in text.lower():
            fail(f"State Farm workbook missing required deep-guide section: {term}")
