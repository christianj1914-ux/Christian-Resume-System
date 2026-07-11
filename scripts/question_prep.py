"""Shared helpers for active application questions and recent interview prep."""

from __future__ import annotations

import re
import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path

import build_resume
import evidence_engine
import interview_context
from build_skills_database import ResumeSnapshot, SOURCE_FILES, parse_resume
from config.language_rules import MANDATORY_REORG_SENTENCE
from config.paths import APPLICATION_QUESTIONS, COMPANY_RESEARCH, FEDERAL_RESUME_SOURCE, GLOBAL_NOTES, INTERVIEW_NOTES, JOBS_DIR
from utils import read_text
from text_safety import collision_safe_substitute, fix_indefinite_articles, sentence_splice_issues


MONTH_NAME_TO_NUMBER = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}
SOFTWARE_ITEMS = (
    ("Microsoft Excel (including Power Query)", (r"\bexcel power query\b",)),
    ("Power BI", (r"\bpower bi\b",)),
    ("SQL", (r"\bsql\b",)),
    ("SAP Crystal Reports", (r"\bsap crystal reports\b|\bcrystal reports\b",)),
    ("Tableau", (r"\btableau\b",)),
    ("Salesforce CRM", (r"\bsalesforce\b",)),
    ("Salesforce Service Cloud", (r"\bservice cloud\b",)),
    ("Salesforce Marketing Cloud", (r"\bmarketing cloud\b",)),
    ("Salesforce AppExchange", (r"\bappexchange\b",)),
    ("ServiceNow", (r"\bservicenow\b",)),
    ("Jira", (r"\bjira\b",)),
    ("Aptean Intuitive", (r"\baptean intuitive\b",)),
    ("Aptean Encompix", (r"\baptean encompix\b|\bencompix\b",)),
    ("Epicor Kinetic", (r"\bepicor kinetic\b",)),
    ("Microsoft Dynamics 365", (r"\bmicrosoft dynamics 365\b",)),
    ("LivePerson LiveEngage", (r"\bliveperson liveengage\b|\bliveengage\b",)),
    ("Claude", (r"\bclaude\b",)),
)
PROFESSIONAL_DEVELOPMENT_NAMES = tuple(
    json.loads(FEDERAL_RESUME_SOURCE.read_text(encoding="utf-8")).get("professional_development", ())
)


@dataclass(frozen=True)
class QualificationsResponse:
    prompt: str
    answer: str
    warning: str = ""


@dataclass(frozen=True)
class InterviewQuestionPrep:
    prompt: str
    answer_angle: str
    category: str = ""


@dataclass(frozen=True)
class PositioningBrief:
    company_name: str
    role_title: str
    primary_lane: str
    employer_type: str
    mission_or_context: str
    role_core_problem: str
    role_problem_phrase: str
    personal_reason_to_care: str
    personal_reason_source: str
    strongest_direct_proofs: list[str]
    strongest_bridge_theme: str
    top_proof_anchors: list[str]
    company_specific_fact: str
    gap_honesty_boundary: str
    selected_proof_sentences: list[str]


APPLICATION_BANNED_PHRASES = (
    "the role sits close to",
    "it is a strong match with how i have worked",
    "my closest match is",
    "i would ramp honestly",
    "note:",
    "relevant to this role",
    "tailored to this role",
)

MISSION_LED_EMPLOYER_PATTERN = re.compile(
    r"\b(?:nonprofit|mission[- ]led|social impact|economic mobility|students?|young adults?|education access|community impact)\b",
    re.I,
)
ANCHOR_NUMBER_PATTERN = re.compile(r"\$[\d,]+(?:\+)?(?:\s*[MK])?|\b\d+\+?%?\b", re.I)
PROPER_NOUN_PATTERN = re.compile(r"\b[A-Z][a-z0-9]+(?:\s+[A-Z][a-z0-9]+)+\b")
ACTION_VERB_PATTERN = re.compile(
    r"\b(?:led|owned|built|launched|reduced|protected|improved|increased|stabilized|designed|translated|"
    r"coordinated|delivered|managed|supported|created|developed|drove|streamlined|implemented|validated|"
    r"facilitated|automated)\b",
    re.I,
)
JOB_LABEL_PATTERN = re.compile(
    r"^(?:company|job title|role|position summary|about us|about the role|responsibilities|qualifications)\s*:",
    re.I,
)
DATE_RANGE_PATTERN = re.compile(
    r"\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\b.*\b20\d{2}\b",
    re.I,
)


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def normalize_question(text: str) -> str:
    normalized = normalize_spaces(text).lower()
    return normalized.rstrip(":")


def approved_source_resume_text() -> str:
    return "\n".join(build_resume.docx_visible_text_from_path(path) for path in SOURCE_FILES)


def split_into_sentences(text: str) -> list[str]:
    if not text or not text.strip():
        return []
    sentences: list[str] = []
    # Split on newlines before collapsing whitespace. Job descriptions commonly
    # use a "Company: X\n\nJob Title: Y\n\nAbout the job\n..." header with no
    # terminal punctuation, so normalizing straight to a single line first would
    # glue that header onto the next real sentence. looks_like_resume_noise_line()
    # then discards the whole merged blob, including the real content. Treating
    # each line as its own sentence boundary keeps the noise filter scoped to the
    # actual noise line and leaves the real sentence available for extraction.
    for line in text.splitlines():
        normalized_line = normalize_spaces(line)
        if not normalized_line:
            continue
        parts = re.split(r"(?<=[.!?])\s+", normalized_line)
        sentences.extend(part.strip() for part in parts if part.strip())
    return sentences


def sentence_has_named_context(sentence: str, company_name: str = "") -> bool:
    if company_name and re.search(re.escape(company_name), sentence, re.I):
        return True
    return bool(PROPER_NOUN_PATTERN.search(sentence))


def clean_answer_sentence(text: str) -> str:
    cleaned = normalize_spaces(text).strip().strip('"')
    if not cleaned:
        return ""
    if cleaned[-1] not in ".!?":
        cleaned += "."
    return cleaned


def clean_selected_proof_sentence(text: str) -> str:
    sentences = split_into_sentences(text)
    if (
        sentences
        and normalize_spaces(sentences[-1]).rstrip(".!?").lower()
        == normalize_spaces(MANDATORY_REORG_SENTENCE).rstrip(".!?").lower()
    ):
        sentences = sentences[:-1]
    if not sentences:
        return ""
    return clean_answer_sentence(" ".join(sentences))


def looks_like_resume_noise_line(text: str) -> bool:
    normalized = normalize_spaces(text)
    lowered = normalized.lower()
    if not normalized:
        return True
    if re.match(r"^(POST-INTERVIEW NOTE|INTERVIEW NOTES FILE:)", normalized, re.I):
        return True
    if JOB_LABEL_PATTERN.match(normalized):
        return True
    if "linkedin.com/" in lowered or re.search(r"\b[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}\b", lowered):
        return True
    if normalized.startswith("Atlanta, GA"):
        return True
    if "|" in normalized and not ACTION_VERB_PATTERN.search(lowered):
        return True
    if DATE_RANGE_PATTERN.search(lowered) and not ACTION_VERB_PATTERN.search(lowered):
        return True
    if normalized.count(":") >= 2 and not ACTION_VERB_PATTERN.search(lowered):
        return True
    if "proof includes" in lowered or "strong fit for roles" in lowered:
        return True
    if re.match(r"^[A-Z][A-Za-z& .'-]+\s+is\s+(?:a|an|the)\b", normalized) and not ACTION_VERB_PATTERN.search(lowered):
        return True
    if (
        re.search(r"\b(?:education|professional development|core competencies|technical skills|professional summary)\b", lowered)
        and len(normalized.split()) <= 8
    ):
        return True
    return False


def compact_anchor_phrase(text: str) -> str:
    patterns = (
        r"\$[\d,]+(?:\+)?(?:\s*[MK])?(?:\s+in\s+[^.,;]{0,45})?",
        r"\b\d+\+?\s+(?:\w+\s+){0,2}(?:client engagements?|clients?|users?|sites?|dashboards?|kpi tools|reporting tools|workshops?|qbrs?|offices|accounts?)\b[^.,;]{0,45}",
        r"\b\d+%\s+(?:reduction|lower|fewer|improvement|improved|cut)\b[^.,;]{0,55}",
    )
    candidates: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, re.I):
            snippet = normalize_spaces(match.group(0)).rstrip(".,;:")
            snippet = re.sub(r"\s+by\s+\w+ing$", "", snippet, flags=re.I)
            end_pos = match.end()
            if end_pos < len(text) and text[end_pos : end_pos + 1].isalpha():
                snippet = re.sub(r"\s+\S+$", "", snippet)
            snippet = re.sub(r"\s+by\s+\w+ing$", "", snippet, flags=re.I)
            if snippet:
                candidates.append(snippet)
    if candidates:
        return max(candidates, key=lambda item: (proof_anchor_value(item), len(item)))
    return normalize_spaces(text).rstrip(".,;:")


def sentence_word_count(text: str) -> int:
    return len(re.findall(r"\b[\w+.#'-]+\b", text))


def lower_lead(text: str) -> str:
    cleaned = normalize_spaces(text).strip()
    if not cleaned:
        return ""
    if len(cleaned) > 1 and cleaned[0].isupper() and cleaned[1].islower():
        return cleaned[0].lower() + cleaned[1:]
    return cleaned


def ensure_company_named(sentence: str, company_name: str) -> str:
    cleaned = clean_answer_sentence(sentence)
    if not cleaned:
        return ""
    if company_name and not re.search(re.escape(company_name), cleaned, re.I):
        return clean_answer_sentence(f"{company_name} stands out to me because {lower_lead(cleaned)}")
    return cleaned


def assert_no_application_banned_phrases(text: str) -> None:
    lowered = text.lower()
    for phrase in APPLICATION_BANNED_PHRASES:
        if phrase in lowered:
            raise ValueError(f"Banned phrase detected in generated answer: {phrase}")
    if "--" in text:
        raise ValueError("Generated answer contains a double-dash.")
    if re.search(r"(^|[.!?]\s+)That\b", text):
        raise ValueError("Generated answer contains a sentence starting with 'That'.")


def validate_answer_word_count(text: str, minimum: int, maximum: int, label: str) -> None:
    words = sentence_word_count(text)
    if words < minimum or words > maximum:
        raise ValueError(f"{label} must be between {minimum} and {maximum} words; got {words}.")


def note_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        line = re.sub(r"^[\-\*•#\d.\s]+", "", normalize_spaces(raw)).strip().strip('"')
        if len(line) >= 8:
            lines.append(line)
    return lines


def motivation_note_line(notes_text: str) -> str:
    motivation_patterns = (
        r"\bdrawn to\b",
        r"\bcare about\b",
        r"\bmotivates me\b",
        r"\blove about\b",
        r"\bwhat keeps me\b",
        r"\bwhy i\b",
        r"\bwhat matters to me\b",
        r"\bmeaningful to me\b",
        r"\bpersonally\b",
        r"\bgiving back\b",
        r"\bpassion\b",
    )
    for line in note_lines(notes_text):
        if not re.search(r"\b(i|me|my)\b", line, re.I):
            continue
        if any(re.search(pattern, line, re.I) for pattern in motivation_patterns):
            cleaned = line.strip().rstrip(".")
            if len(cleaned) >= 20:
                return f"{cleaned}."
    return ""


def lane_default_human_motivation_sentence(profile: build_resume.JobProblemProfile) -> str:
    lane_sentences = {
        "analytics_operations": (
            "What keeps me engaged is the moment when a business question stops being guesswork and becomes clear enough for someone to act on with confidence."
        ),
        "implementation_delivery": (
            "The part I keep coming back to is when a system stops being a project plan and starts working for the people who have to live in the workflow every day."
        ),
        "customer_success": (
            "What matters to me in customer work is the shift from a customer managing issues to a customer feeling confident enough to get real value from the platform."
        ),
        "presales_solution": (
            "I like the discovery side of this work because the best solution conversations start by getting honest about the workflow before anyone jumps to features."
        ),
        "change_enablement": (
            "What I have learned is that resistance is usually not about the tool itself. It is about whether people feel clear, prepared, and heard."
        ),
        "process_improvement": (
            "What keeps me engaged is finding the friction people have learned to tolerate and replacing it with a process that actually makes the day easier."
        ),
        "corporate_strategy": (
            "What keeps me engaged is making an ambiguous problem concrete enough that a team can act on it instead of continuing to debate it."
        ),
    }
    return lane_sentences.get(
        profile.primary_lane,
        "What keeps me engaged is working on problems where the technical answer is only half the solution and the human follow-through determines whether the work sticks."
    )


def candidate_problem_phrase(profile: build_resume.JobProblemProfile) -> str:
    return {
        "implementation_delivery": "complex ERP implementations moving cleanly from discovery through go-live and adoption",
        "customer_success": "customer adoption, risk, and renewal work staying steady under pressure",
        "presales_solution": "buyer questions turning into workable solution decisions",
        "analytics_operations": "systems, data, and workflow questions turning into decisions people can use",
        "change_enablement": "systems and process changes turning into real adoption",
        "process_improvement": "operational friction turning into measurable process improvement",
        "corporate_strategy": "complex cross-functional questions turning into decisions people can act on",
    }.get(profile.primary_lane, build_resume.natural_problem_phrase(profile))


def candidate_problem_phrase_from_brief(brief: PositioningBrief) -> str:
    return {
        "implementation_delivery": "complex ERP implementations moving cleanly from discovery through go-live and adoption",
        "customer_success": "customer adoption, risk, and renewal work staying steady under pressure",
        "presales_solution": "buyer questions turning into workable solution decisions",
        "analytics_operations": "systems, data, and workflow questions turning into decisions people can use",
        "change_enablement": "systems and process changes turning into real adoption",
        "process_improvement": "operational friction turning into measurable process improvement",
        "corporate_strategy": "complex cross-functional questions turning into decisions people can act on",
    }.get(brief.primary_lane, brief.role_problem_phrase)


def active_positioning_context(job_description: str) -> tuple[str, str]:
    company_name = build_resume.extract_output_name(job_description)
    role_title = build_resume.extract_job_title(job_description) or ""
    bundle = interview_context.load_company_context(
        JOBS_DIR,
        company_name,
        role_title,
        company_research_path=COMPANY_RESEARCH,
        global_interview_notes_path=INTERVIEW_NOTES,
        mode="compact",
    )
    notes_parts = [
        part
        for part in (
            bundle.interview_notes,
            bundle.coaching_notes,
            bundle.supplied_context,
            read_text(INTERVIEW_NOTES),
            read_text(GLOBAL_NOTES) if GLOBAL_NOTES.exists() else "",
        )
        if normalize_spaces(part)
    ]
    notes_text = "\n".join(dict.fromkeys(note.strip() for note in notes_parts if note.strip()))
    research_text = bundle.company_research or interview_context.company_research_context(
        JOBS_DIR,
        company_name,
        role_title,
        company_research_path=COMPANY_RESEARCH,
    )
    return notes_text, research_text


def role_keyword_terms(job_description: str, profile: build_resume.JobProblemProfile) -> tuple[str, ...]:
    keywords = list(build_resume.audit_keywords(job_description))
    keywords.extend(profile.direct_matches)
    keywords.extend(profile.adjacent_matches)
    keywords.extend(re.findall(r"[A-Za-z][A-Za-z0-9/+.-]+", profile.core_problem))
    seen: list[str] = []
    for keyword in keywords:
        cleaned = normalize_spaces(keyword).lower()
        if len(cleaned) < 3 or cleaned in seen:
            continue
        seen.append(cleaned)
    return tuple(seen)


def canonical_evidence_supported(
    item: evidence_engine.CanonicalEvidence,
    resume_text: str,
) -> bool:
    normalized_resume = normalize_spaces(resume_text).lower()
    return all(
        normalize_spaces(term).lower() in normalized_resume
        for term in evidence_engine.canonical_evidence_support_terms(item)
    )


def canonical_lane_terms(profile: build_resume.JobProblemProfile) -> tuple[str, ...]:
    return {
        "implementation_delivery": (
            "implementation",
            "erp",
            "go-live",
            "configuration",
            "migration",
            "requirements",
            "training",
            "client side",
            "finance",
        ),
        "customer_success": (
            "customer",
            "retention",
            "renewal",
            "revenue",
            "qbr",
            "risk",
        ),
        "presales_solution": (
            "discovery",
            "solution",
            "demo",
            "buyer",
            "requirements",
        ),
        "analytics_operations": (
            "analytics",
            "reporting",
            "dashboard",
            "sql",
            "workflow",
        ),
        "change_enablement": (
            "change",
            "adoption",
            "training",
            "enablement",
            "stakeholder",
        ),
        "corporate_strategy": (
            "decision",
            "analysis",
            "recommendation",
            "risk",
            "stakeholder",
        ),
        "process_improvement": (
            "workflow",
            "automation",
            "root cause",
            "improvement",
            "auditable",
        ),
    }.get(profile.primary_lane, ())


def canonical_evidence_score(
    item: evidence_engine.CanonicalEvidence,
    job_description: str,
    profile: build_resume.JobProblemProfile,
) -> int:
    jd_lower = job_description.lower()
    item_text = f"{item.claim} {' '.join(item.lane_signals)}".lower()
    lane_terms = canonical_lane_terms(profile)
    score = sum(8 for term in lane_terms if term in item_text)
    score += sum(5 for signal in item.lane_signals if signal.lower() in jd_lower)
    if profile.primary_lane == "implementation_delivery":
        if re.search(r"\b(?:client-side|cfo|controller|finance)\b", item_text):
            score += 18
        if re.search(r"\b(?:erp|go-live|configuration|migration|cutover|training|requirements)\b", item_text):
            score += 16
        if re.search(r"\b(?:amazon robotics|acumatica|warehouse)\b", item_text):
            score += 10
    score += min(40, proof_anchor_value(item.claim))
    return score


def supported_canonical_evidence(
    resume_text: str,
    job_description: str,
    profile: build_resume.JobProblemProfile,
) -> list[evidence_engine.CanonicalEvidence]:
    supported = [
        item
        for item in evidence_engine.commercial_canonical_evidence()
        if canonical_evidence_supported(item, resume_text)
    ]
    supported.sort(
        key=lambda item: canonical_evidence_score(item, job_description, profile),
        reverse=True,
    )
    return supported


def canonical_claim_sentence(item: evidence_engine.CanonicalEvidence) -> str:
    claim = normalize_spaces(item.claim)
    if not claim:
        return ""
    claim = claim[0].upper() + claim[1:] if claim[0].islower() else claim
    return clean_answer_sentence(claim)


def canonical_anchor_phrase(item: evidence_engine.CanonicalEvidence) -> str:
    anchored_patterns = (
        r"\b\d+\+?\s+(?:users?|clients?|sites?|workshops?|accounts?|offices)\b(?:\s+across\s+[^.,;]+)?",
        r"\b[A-Z][A-Za-z0-9 ]+\s+certification before go-live\b",
        r"\b\d+-month\b[^.,;]{0,20}",
        r"\$[\d,]+(?:\+)?(?:\s*[MK])?[^.,;]{0,20}",
    )
    for source in (item.claim, item.source_text):
        for pattern in anchored_patterns:
            match = re.search(pattern, source, re.I)
            if match:
                return normalize_spaces(match.group(0)).rstrip(".,;:")
        compact = compact_anchor_phrase(source)
        if compact and compact != normalize_spaces(source).rstrip(".,;:") and sentence_word_count(compact) <= 12:
            return compact.rstrip(".,;:")
    for metric in item.metrics:
        cleaned = normalize_spaces(metric).rstrip(".,;:")
        if cleaned:
            return cleaned
    return ""


def proof_anchor_value(text: str) -> int:
    score = 0
    for match in ANCHOR_NUMBER_PATTERN.findall(text):
        digits = re.sub(r"[^\d]", "", match)
        if digits:
            score += int(digits[:6])
    if "$" in text:
        score += 500
    if "%" in text:
        score += 200
    if "+" in text:
        score += 100
    return score


def extract_top_proof_anchors(resume_text: str, job_description: str, profile: build_resume.JobProblemProfile) -> list[str]:
    canonical_anchors: list[str] = []
    for item in supported_canonical_evidence(resume_text, job_description, profile):
        anchor = canonical_anchor_phrase(item)
        if anchor and anchor not in canonical_anchors:
            canonical_anchors.append(anchor)
        if len(canonical_anchors) >= 2:
            return canonical_anchors

    role_terms = role_keyword_terms(job_description, profile)
    candidates: list[tuple[int, str]] = []
    for line in resume_text.splitlines():
        normalized = normalize_spaces(line)
        lowered = normalized.lower()
        if looks_like_resume_noise_line(normalized) or not ANCHOR_NUMBER_PATTERN.search(normalized):
            continue
        if role_terms and not any(term in lowered for term in role_terms):
            scope_terms = ("users", "clients", "sites", "dashboards", "revenue", "workshops", "qbr", "accounts")
            if not any(term in lowered for term in scope_terms):
                continue
        compact = compact_anchor_phrase(normalized)
        candidates.append((proof_anchor_value(compact), compact))
    candidates.sort(key=lambda item: item[0], reverse=True)
    anchors: list[str] = []
    for _score, line in candidates:
        if line not in anchors:
            anchors.append(line)
        if len(anchors) >= 2:
            break
    return anchors


def extract_selected_proof_sentences(resume_text: str, job_description: str, profile: build_resume.JobProblemProfile) -> list[str]:
    selected: list[str] = []
    for item in supported_canonical_evidence(resume_text, job_description, profile):
        sentence = canonical_claim_sentence(item)
        if sentence and sentence not in selected:
            selected.append(sentence)
        if len(selected) >= 3:
            return selected

    role_terms = role_keyword_terms(job_description, profile)
    candidates: list[tuple[int, str]] = []
    for line in resume_text.splitlines():
        normalized = clean_selected_proof_sentence(line)
        lowered = normalized.lower()
        if not normalized or looks_like_resume_noise_line(normalized):
            continue
        if not ACTION_VERB_PATTERN.search(lowered):
            continue
        if re.match(r"^(?:implementation consultant|customer success consultant|enterprise systems manager)\b", lowered):
            continue
        if re.search(r"\bbrings the most value when\b", lowered):
            continue
        score = 0
        if any(term in lowered for term in role_terms):
            score += 5
        if ANCHOR_NUMBER_PATTERN.search(normalized):
            score += 6
        score += 3
        if score <= 0:
            continue
        candidates.append((score + proof_anchor_value(normalized), normalized))
    candidates.sort(key=lambda item: item[0], reverse=True)
    for _score, sentence in candidates:
        if sentence not in selected:
            selected.append(sentence)
        if len(selected) >= 3:
            break
    return selected


def mission_or_context_sentence(company_name: str, company_research_text: str, job_description: str) -> str:
    for sentence in split_into_sentences(company_research_text):
        lowered = sentence.lower()
        if looks_like_resume_noise_line(sentence):
            continue
        if company_name and company_name.lower() in lowered and re.search(r"\b(?:helps|provides|builds|serves|prepares|supports|focuses on|connects)\b", lowered):
            return clean_answer_sentence(sentence)
    for sentence in split_into_sentences(job_description[:400]):
        lowered = sentence.lower()
        if looks_like_resume_noise_line(sentence):
            continue
        if re.match(r"^(?:company|job title|role|position(?: title)?)\s*:", sentence, re.I):
            continue
        if re.search(r"\b(?:mission|purpose|focuses|prepare|supports|serves|helps|builds|empowers)\b", lowered):
            return ensure_company_named(sentence, company_name)
    return ""


def company_specific_fact_sentence(company_name: str, company_research_text: str) -> str:
    generic_only = (
        "the company",
        "the organization",
        "we believe",
        "our mission is to",
    )
    for sentence in split_into_sentences(company_research_text):
        lowered = sentence.lower()
        if looks_like_resume_noise_line(sentence):
            continue
        if any(fragment in lowered for fragment in generic_only) and not sentence_has_named_context(sentence, company_name):
            continue
        if sentence_has_named_context(sentence, company_name):
            return clean_answer_sentence(sentence)
    return ""


def derive_employer_type(job_description: str, profile: build_resume.JobProblemProfile, company_research_text: str = "") -> str:
    lowered = f"{job_description}\n{company_research_text}".lower()
    if MISSION_LED_EMPLOYER_PATTERN.search(lowered):
        return "mission-led nonprofit"
    context = build_resume.primary_employer_context(job_description)
    if context:
        key = str(context.get("key", "")).strip()
        if key == "saas":
            return "B2B SaaS growth stage" if re.search(r"\b(series|growth stage|startup|scale)\b", lowered) else "B2B SaaS"
        if key == "consulting":
            return "consulting firm"
        if key == "financial_services":
            return "financial services employer"
        if key == "education_assessment":
            return "education-focused employer"
        if key == "customer_experience":
            return "customer experience platform"
        return key.replace("_", " ")
    return profile.lane_label


def build_positioning_brief(
    job_description: str,
    resume_text: str,
    notes_text: str = "",
    company_research_text: str = "",
) -> PositioningBrief:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    company_name = build_resume.extract_output_name(job_description)
    role_title = build_resume.extract_job_title(job_description) or "the role"
    personal_reason = motivation_note_line(notes_text)
    personal_reason_source = "notes" if personal_reason else "lane_default"
    if not personal_reason:
        personal_reason = lane_default_human_motivation_sentence(profile)
    direct_proofs = list(profile.direct_matches[:3])
    if not direct_proofs:
        direct_proofs = list(profile.adjacent_matches[:2]) or [build_resume.natural_problem_phrase(profile)]
    if direct_proofs:
        job_lower = job_description.lower()

        def direct_proof_score(proof: str) -> int:
            lowered = proof.lower()
            score = 0
            if profile.primary_lane == "implementation_delivery":
                if any(term in lowered for term in ("erp", "client-side", "finance", "implementation", "delivery", "migration")):
                    score += 8
                if "acumatica" in job_lower and "erp" in lowered:
                    score += 4
                if any(term in job_lower for term in ("controller", "cfo", "finance")) and any(term in lowered for term in ("client-side", "finance")):
                    score += 4
            if any(term in lowered for term in ("change", "adoption", "training", "enablement")) and any(term in job_lower for term in ("training", "adoption", "change")):
                score += 2
            return score

        direct_proofs.sort(key=direct_proof_score, reverse=True)
    strongest_bridge_theme = profile.adjacent_matches[0] if profile.adjacent_matches else ""
    brief = PositioningBrief(
        company_name=company_name,
        role_title=role_title,
        primary_lane=profile.primary_lane,
        employer_type=derive_employer_type(job_description, profile, company_research_text),
        mission_or_context=mission_or_context_sentence(company_name, company_research_text, job_description),
        role_core_problem=profile.core_problem,
        role_problem_phrase=build_resume.natural_problem_phrase(profile),
        personal_reason_to_care=personal_reason,
        personal_reason_source=personal_reason_source,
        strongest_direct_proofs=direct_proofs,
        strongest_bridge_theme=strongest_bridge_theme,
        top_proof_anchors=extract_top_proof_anchors(resume_text, job_description, profile),
        company_specific_fact=company_specific_fact_sentence(company_name, company_research_text),
        gap_honesty_boundary=(
            f"The area where I am building rather than established is {profile.unsupported_requirements[0]}."
            if profile.unsupported_requirements else ""
        ),
        selected_proof_sentences=extract_selected_proof_sentences(resume_text, job_description, profile),
    )
    validate_positioning_brief(brief)
    return brief


def validate_positioning_brief(brief: PositioningBrief) -> None:
    if not brief.role_core_problem.strip():
        raise ValueError("PositioningBrief role_core_problem must be non-empty.")
    if not brief.role_problem_phrase.strip():
        raise ValueError("PositioningBrief role_problem_phrase must be non-empty.")
    if not brief.strongest_direct_proofs:
        raise ValueError("PositioningBrief strongest_direct_proofs must include at least one supported proof theme.")
    if not brief.personal_reason_to_care.strip():
        raise ValueError("PositioningBrief personal_reason_to_care must be non-empty.")


def finalize_candidate_answer(
    label: str,
    text: str,
    *,
    claim_first: bool = False,
    min_words: int = 18,
) -> str:
    cleaned = normalize_spaces(fix_indefinite_articles(text))
    issues = sentence_splice_issues(cleaned)
    if issues:
        raise ValueError(f"{label} produced unresolved grammar fragments: {', '.join(issues)}")
    if claim_first:
        import build_interview_cheat_sheet

        build_interview_cheat_sheet.assert_claim_then_prove_answer(label, cleaned, min_words=min_words)
    return cleaned


def active_positioning_brief(job_description: str, resume_text: str) -> PositioningBrief:
    notes_text, company_research_text = active_positioning_context(job_description)
    return build_positioning_brief(
        job_description=job_description,
        resume_text=resume_text,
        notes_text=notes_text,
        company_research_text=company_research_text,
    )

def parse_question_blocks(text: str) -> tuple[str, ...]:
    prompts: list[str] = []
    current: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if current:
                prompts.append(normalize_spaces(" ".join(current)))
                current = []
            continue
        if re.fullmatch(r"(application )?(supplemental )?(qualifications )?questions?", line, re.I):
            continue
        current.append(line)
    if current:
        prompts.append(normalize_spaces(" ".join(current)))
    return tuple(dict.fromkeys(prompt for prompt in prompts if prompt))


DEFAULT_APPLICATION_QUESTIONS: tuple[str, ...] = ("Why are you interested in this role?",)


@dataclass(frozen=True)
class ApplicationPromptState:
    explicit_prompts: tuple[str, ...]
    effective_prompts: tuple[str, ...]
    uses_default_questions: bool


def load_application_prompt_state(question_path: Path = APPLICATION_QUESTIONS) -> ApplicationPromptState:
    """Canonical loader. Returns provenance-aware state; never blocks on an empty file."""
    if not question_path.exists():
        return ApplicationPromptState(
            explicit_prompts=(),
            effective_prompts=DEFAULT_APPLICATION_QUESTIONS,
            uses_default_questions=True,
        )
    text = question_path.read_text(encoding="utf-8-sig")
    parsed = parse_question_blocks(text)
    if not parsed or set(parsed) == set(DEFAULT_APPLICATION_QUESTIONS):
        return ApplicationPromptState(
            explicit_prompts=parsed,
            effective_prompts=DEFAULT_APPLICATION_QUESTIONS,
            uses_default_questions=True,
        )
    return ApplicationPromptState(
        explicit_prompts=parsed,
        effective_prompts=parsed,
        uses_default_questions=False,
    )


def load_active_application_prompts(question_path: Path = APPLICATION_QUESTIONS) -> tuple[str, ...]:
    return load_application_prompt_state(question_path).effective_prompts


def require_active_application_prompts(
    question_path: Path = APPLICATION_QUESTIONS,
    *,
    workflow_name: str = "the standard commercial workflow",
) -> tuple[str, ...]:
    return load_application_prompt_state(question_path).effective_prompts


def application_question_context_issues(
    job_description: str,
    prompts: tuple[str, ...] | ApplicationPromptState,
    *,
    workflow: str = "commercial",
) -> tuple[str, ...]:
    """Identify strongly stale prompts without silently deleting user input.

    Accepts either a raw tuple (legacy callers) or ApplicationPromptState.
    Default-fallback builds skip archive-pairing and heuristic checks entirely.
    """

    if isinstance(prompts, ApplicationPromptState):
        effective = prompts.effective_prompts
        uses_default = prompts.uses_default_questions
    else:
        effective = prompts
        uses_default = False

    jd = normalize_question(job_description)
    issues: list[str] = []
    if workflow == "commercial" and effective and not uses_default:
        try:
            import job_context_archive

            pairing_issue = job_context_archive.application_question_pairing_issue(
                job_description,
                "\n\n".join(effective),
            )
        except (OSError, ValueError):
            pairing_issue = ""
        if pairing_issue:
            issues.append(f"{pairing_issue} First active prompt: {effective[0]}")
    if not uses_default:
        for prompt in effective:
            normalized = normalize_question(prompt)
            if workflow == "commercial":
                if "public agencies or cooperatives" in normalized and not re.search(r"\b(?:public agenc|cooperative)\b", jd):
                    issues.append(prompt)
                elif re.search(r"\bai product or service\b", normalized) and not re.search(r"\b(?:ai|artificial intelligence)\b", jd):
                    issues.append(prompt)
            else:
                if re.search(r"\b(?:joining this company|saas or technology company|ai product or service)\b", normalized):
                    if not any(term in jd for term in ("saas", "artificial intelligence", " ai ")):
                        issues.append(prompt)
    return tuple(dict.fromkeys(issues))


def element_probe_responses(job_description: str, resume_text: str, *, limit: int = 3) -> tuple[QualificationsResponse, ...]:
    """Prepare bridge answers for named requirements not explicit in source text."""

    from requirement_engine import parse_commercial_requirements

    resume_normalized = normalize_question(resume_text)
    responses: list[QualificationsResponse] = []
    seen_terms: set[str] = set()
    for element in parse_commercial_requirements(job_description):
        missing = [term for term in element.canonical_terms if normalize_question(term) not in resume_normalized]
        if not element.required or not missing:
            continue
        term = missing[0]
        if term in seen_terms:
            continue
        seen_terms.add(term)
        prompt = f"How would your experience transfer to work involving {term}?"
        answer = normalize_spaces(
            f"My experience is adjacent rather than identical to {term}. "
            f"I would bridge it honestly this way: {generic_bridge_answer(job_description, resume_text)}"
        )
        responses.append(QualificationsResponse(prompt, answer))
        if len(responses) >= limit:
            break
    return tuple(responses)


def application_question_draft_path(output_docx: Path) -> Path:
    if " DRAFT" in output_docx.stem.upper():
        return output_docx
    return output_docx.with_name(output_docx.stem + " DRAFT" + output_docx.suffix)


def mark_docx_as_draft(output_docx: Path, issues: tuple[str, ...]) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document(str(output_docx))
    title = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip()), None)
    if title is None:
        title = document.add_paragraph()
    if not title.runs:
        title.add_run()
    if not title.text.upper().startswith("DRAFT"):
        first_run = title.runs[0]
        first_run.text = "DRAFT - " + first_run.text
        first_run.bold = True
        first_run.font.color.rgb = RGBColor(192, 0, 0)

    # Put the diagnostic at the end so marking an already paginated document does
    # not push a manual page break onto an otherwise blank page.
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    spoken_issue = any(issue.startswith("Spoken prose repair") for issue in issues)
    heading_text = (
        "DRAFT - REQUIRES HUMAN REVIEW\n"
        if spoken_issue
        else "DRAFT - APPLICATION QUESTIONS REQUIRE REVIEW\n"
    )
    detail_label = "Review issue(s): " if spoken_issue else "Suspect prompt(s): "
    heading = note.add_run(heading_text)
    heading.bold = True
    heading.font.color.rgb = RGBColor(192, 0, 0)
    detail = note.add_run(detail_label + " | ".join(issues))
    detail.font.size = Pt(8)
    document.save(str(output_docx))


def selected_resume_snapshot(job_description: str) -> tuple[Path, ResumeSnapshot, str]:
    resume_path = build_resume.choose_resume(job_description)
    snapshot = parse_resume(resume_path)
    resume_text = approved_source_resume_text()
    return resume_path, snapshot, resume_text


def parse_month_year(text: str) -> tuple[int, int]:
    normalized = normalize_spaces(text)
    if normalized.lower() == "present":
        today = date.today()
        return today.year, today.month
    month_name, year_text = normalized.rsplit(" ", 1)
    month = MONTH_NAME_TO_NUMBER[month_name.lower()]
    return int(year_text), month


def role_month_count(snapshot: ResumeSnapshot) -> int:
    covered_months: set[int] = set()
    for role in snapshot.roles:
        if " to " not in role.dates:
            continue
        start_text, end_text = role.dates.split(" to ", 1)
        start_year, start_month = parse_month_year(start_text)
        end_year, end_month = parse_month_year(end_text)
        start_index = start_year * 12 + start_month - 1
        end_index = end_year * 12 + end_month - 1
        for month_index in range(start_index, end_index + 1):
            covered_months.add(month_index)
    return len(covered_months)


def direct_experience_years_label(snapshot: ResumeSnapshot) -> str:
    months = role_month_count(snapshot)
    years = max(1, months // 12)
    return f"{years}+ years" if months % 12 else f"{years} years"


def education_years_answer(snapshot: ResumeSnapshot) -> str:
    education_lines = " ".join(snapshot.education).lower()
    estimated_years = 0
    if "bachelor" in education_lines:
        estimated_years += 4
    if "master" in education_lines:
        estimated_years += 2
    if estimated_years >= 6:
        years_text = "6+ years"
    elif estimated_years >= 4:
        years_text = "4+ years"
    else:
        years_text = "Related post-secondary education"
    development = ", ".join(snapshot.professional_development)
    return normalize_spaces(
        f"{years_text} of related post-secondary education, supported by bachelor's- and master's-level study in information systems, plus ongoing professional development through {development}."
    )


def direct_experience_years_answer(snapshot: ResumeSnapshot) -> str:
    label = direct_experience_years_label(snapshot)
    plus_match = re.fullmatch(r"(\d+)\+ years", label)
    if plus_match:
        return f"More than {plus_match.group(1)} years of direct relevant experience."
    return f"Approximately {label} of direct relevant experience."


def select_within_word_budget(
    required_lead: str,
    optional_parts: list[str],
    required_trailer: str,
    maximum: int,
) -> list[str]:
    """Greedily keep optional_parts, in order, so the assembled answer stays within maximum words.

    Always keeps at least the first non-empty optional part (even if that alone would
    exceed the budget) so a single long sentence does not silently collapse the answer to
    nothing; validate_answer_word_count remains the final safety check after assembly.
    """
    selected: list[str] = []
    for part in optional_parts:
        if not part:
            continue
        candidate = " ".join(p for p in (required_lead, *selected, part, required_trailer) if p)
        if sentence_word_count(candidate) <= maximum or not selected:
            selected.append(part)
        else:
            break
    return selected


def build_why_company_answer(brief: PositioningBrief) -> str:
    problem_phrase = candidate_problem_phrase_from_brief(brief)
    if "mission" in brief.employer_type.lower() or "nonprofit" in brief.employer_type.lower():
        if brief.personal_reason_source != "notes":
            raise ValueError(
                "Mission-led employer detected but personal_reason_to_care is empty. "
                "Add motivation context to jobs/interview_notes.txt before generating."
            )
    opening = (
        ensure_company_named(brief.company_specific_fact, brief.company_name)
        or ensure_company_named(brief.mission_or_context, brief.company_name)
        or clean_answer_sentence(
            f"{brief.company_name} is the kind of {brief.employer_type} environment where {problem_phrase} directly shapes the work."
        )
    )
    motivation = clean_answer_sentence(
        f"I am interested in joining {brief.company_name} because {lower_lead(brief.personal_reason_to_care)}"
    )
    proof_sentence = clean_answer_sentence(
        brief.selected_proof_sentences[1]
        if len(brief.selected_proof_sentences) > 1
        else brief.selected_proof_sentences[0]
        if brief.selected_proof_sentences
        else f"The background already includes {brief.strongest_direct_proofs[0]} in complex operating environments."
    )
    role_fit = clean_answer_sentence(
        f"What stands out about the {brief.role_title} role is the need to solve {problem_phrase}, and that is already visible in the background."
    )
    anchor = clean_answer_sentence(
        f"A concrete proof point is {brief.top_proof_anchors[0]}."
        if brief.top_proof_anchors
        else f"A concrete proof point is {brief.strongest_direct_proofs[0]}."
    )
    lead = " ".join(part for part in (opening, motivation) if part)
    optional_parts = select_within_word_budget(lead, [role_fit, proof_sentence, anchor], "", 165)
    answer = " ".join(part for part in (lead, *optional_parts) if part)
    assert_no_application_banned_phrases(answer)
    validate_answer_word_count(answer, 80, 165, "Why-company answer")
    return answer


def build_relevant_experience_answer(brief: PositioningBrief) -> str:
    problem_phrase = candidate_problem_phrase_from_brief(brief)
    lead_theme = brief.strongest_direct_proofs[0]
    sentence_one = clean_answer_sentence(
        f"{lead_theme.capitalize()} has been a consistent part of the work, especially in environments where {problem_phrase} mattered."
    )
    sentence_two = clean_answer_sentence(
        brief.selected_proof_sentences[0]
        if brief.selected_proof_sentences
        else (
            f"The work includes {brief.top_proof_anchors[0]}."
            if brief.top_proof_anchors
            else f"The work includes direct proof in {brief.strongest_direct_proofs[0]}."
        )
    )
    supporting_sentences: list[str] = []
    if brief.primary_lane == "implementation_delivery":
        if re.search(r"\b80\+", sentence_two):
            impl_count_text = "At Aptean, I delivered 12 full-lifecycle ERP implementations and managed up to four at a time."
        else:
            impl_count_text = (
                "At Aptean, I delivered 12 full-lifecycle ERP implementations and managed up to four at a time "
                "across 80+ international client engagements."
            )
        implementation_count_sentence = clean_answer_sentence(impl_count_text)
        if implementation_count_sentence.lower() not in sentence_two.lower():
            supporting_sentences.append(implementation_count_sentence)
    if brief.top_proof_anchors:
        anchor_sentence = clean_answer_sentence(f"Scope-wise, the work included {brief.top_proof_anchors[0]}.")
        if brief.top_proof_anchors[0].lower() not in sentence_two.lower():
            supporting_sentences.append(anchor_sentence)
    for proof_theme in brief.strongest_direct_proofs[1:3]:
        candidate = clean_answer_sentence(
            f"Additional direct proof shows up in {proof_theme}, which kept the work tied to visible owners, "
            "clearer decisions, and measurable follow-through."
        )
        if candidate.lower() not in {item.lower() for item in supporting_sentences}:
            supporting_sentences.append(candidate)
            break
    if brief.strongest_bridge_theme:
        bridge_theme_sentence = clean_answer_sentence(
            f"A useful adjacent strength is {brief.strongest_bridge_theme}, which helps when priorities and stakeholders have to stay aligned."
        )
        if bridge_theme_sentence.lower() not in {item.lower() for item in supporting_sentences}:
            supporting_sentences.append(bridge_theme_sentence)
    bridge = clean_answer_sentence(
        f"Taken together, the combination supports the role because it keeps {problem_phrase} grounded in execution, visibility, and follow-through."
    )
    lead = " ".join(part for part in (sentence_one, sentence_two) if part)
    deduped_support = [
        sentence for sentence in supporting_sentences
        if sentence and sentence.lower() != lead.lower()
    ]
    selected_support = select_within_word_budget(lead, deduped_support, bridge, 195)
    answer = " ".join(part for part in (lead, *selected_support, bridge) if part)
    if sentence_word_count(answer) < 95:
        fallback_sentences = [
            clean_answer_sentence(sentence)
            for sentence in brief.selected_proof_sentences[1:3]
            if clean_answer_sentence(sentence)
        ]
        for fallback in fallback_sentences:
            if fallback.lower() in {item.lower() for item in selected_support} or fallback.lower() in lead.lower():
                continue
            candidate_answer = " ".join(part for part in (lead, *selected_support, fallback, bridge) if part)
            if sentence_word_count(candidate_answer) <= 195:
                selected_support.append(fallback)
                answer = candidate_answer
            if sentence_word_count(answer) >= 95:
                break
    assert_no_application_banned_phrases(answer)
    validate_answer_word_count(answer, 95, 195, "Relevant-experience answer")
    return answer


def build_unique_qualifications_answer(brief: PositioningBrief) -> str:
    problem_phrase = candidate_problem_phrase_from_brief(brief)
    proof_a = brief.strongest_direct_proofs[0]
    proof_b = brief.strongest_direct_proofs[1] if len(brief.strongest_direct_proofs) > 1 else brief.strongest_bridge_theme or brief.role_problem_phrase
    sentence_one = clean_answer_sentence(
        f"The differentiator in this background is the combination of {proof_a} and {proof_b}, which is not often packaged together in one operator."
    )
    evidence_lines = [clean_answer_sentence(line) for line in brief.selected_proof_sentences[1:3] if clean_answer_sentence(line)]
    if brief.top_proof_anchors:
        evidence_lines.append(clean_answer_sentence(f"Scope-wise, the work reached {brief.top_proof_anchors[0]}."))
        if len(brief.top_proof_anchors) > 1:
            evidence_lines.append(clean_answer_sentence(f"A second scope marker is {brief.top_proof_anchors[1]}."))
    if len(brief.strongest_direct_proofs) > 2:
        evidence_lines.append(clean_answer_sentence(f"Another direct proof area is {brief.strongest_direct_proofs[2]}."))
    if brief.strongest_bridge_theme:
        evidence_lines.append(
            clean_answer_sentence(
                f"A useful adjacent strength is {brief.strongest_bridge_theme}, which helps when the work spans systems, stakeholders, and operating follow-through."
            )
        )
    tech_line = ""
    if re.search(r"\b(ai|automation|system|erp|data|analytics)\b", brief.role_core_problem, re.I):
        tech_sentence = next(
            (
                sentence for sentence in brief.selected_proof_sentences
                if re.search(r"\b(claude|sql|salesforce|power bi|dashboard|erp|automation|liveperson|liveengage)\b", sentence, re.I)
            ),
            "",
        )
        tech_line = clean_answer_sentence(tech_sentence) if tech_sentence else ""
    closing = clean_answer_sentence(
        f"The combination is most useful when {problem_phrase} needs structured execution plus business-facing translation."
    )
    ordered_evidence: list[str] = []
    for line in evidence_lines:
        if line and line.lower() not in {item.lower() for item in ordered_evidence}:
            ordered_evidence.append(line)
    candidate_extras = [*ordered_evidence[:3], tech_line]
    selected_extras = select_within_word_budget(sentence_one, candidate_extras, closing, 155)
    answer_parts = [sentence_one, *selected_extras, closing]
    answer = " ".join(part for part in answer_parts if part)
    assert_no_application_banned_phrases(answer)
    validate_answer_word_count(answer, 70, 155, "Unique-qualifications answer")
    return answer


def communication_question_topic(prompt: str) -> str:
    lowered = normalize_question(prompt)
    if "implementation" in lowered:
        return "implementation methodology"
    if re.search(r"\bartificial intelligence\b|\bai\b", lowered):
        return "AI tool usage"
    if "stakeholder" in lowered or "communicat" in lowered or "presentation" in lowered:
        return "stakeholder communication"
    return "cross-functional coordination"


def build_communication_or_implementation_answer(
    brief: PositioningBrief,
    question_topic: str,
) -> str:
    problem_phrase = candidate_problem_phrase_from_brief(brief)
    topic = normalize_spaces(question_topic or "cross-functional coordination")
    sentence_one = clean_answer_sentence(
        f"My approach to {topic} is to make the workflow, owner, risk, and next decision visible early enough for people to act."
    )
    relevant_lines = [
        sentence for sentence in brief.selected_proof_sentences
        if any(term in sentence.lower() for term in topic.lower().split())
    ]
    if not relevant_lines:
        relevant_lines = brief.selected_proof_sentences[2:4]
    if not relevant_lines:
        relevant_lines = [clean_answer_sentence(f"The background includes {brief.strongest_direct_proofs[0]}.")]
    if brief.top_proof_anchors:
        anchor_line = clean_answer_sentence(f"A representative scope marker is {brief.top_proof_anchors[0]}.")
        if anchor_line.lower() not in {clean_answer_sentence(line).lower() for line in relevant_lines}:
            relevant_lines = [*relevant_lines[:1], anchor_line, *relevant_lines[1:]]
    bridge = clean_answer_sentence(
        f"This discipline matters here because the role depends on {problem_phrase} without losing momentum across teams."
    )
    support_parts = [clean_answer_sentence(line) for line in relevant_lines[:2] if clean_answer_sentence(line)]
    if len(support_parts) < 2 and len(brief.strongest_direct_proofs) > 1:
        support_parts.append(clean_answer_sentence(f"A second proof area is {brief.strongest_direct_proofs[1]}."))
    if len(support_parts) < 2 and brief.strongest_bridge_theme:
        support_parts.append(clean_answer_sentence(f"A transferable strength here is {brief.strongest_bridge_theme}."))
    if any(term in topic.lower() for term in ("communication", "presentation", "facilitation", "training", "stakeholder")):
        public_speaking_line = (
            "The strongest public-speaking proof includes 60+ executive workshops and QBRs plus pre-sales discovery and product demonstrations."
            if brief.primary_lane == "presales_solution"
            else "The strongest communication proof includes 60+ executive workshops and QBRs plus role-based training and executive-facing facilitation."
        )
        support_parts = [clean_answer_sentence(public_speaking_line), *support_parts]
    deduped_support = []
    for line in support_parts:
        if line and line.lower() not in {item.lower() for item in deduped_support}:
            deduped_support.append(line)
    selected_support = select_within_word_budget(sentence_one, deduped_support, bridge, 155)
    answer = " ".join([sentence_one, *selected_support, bridge])
    assert_no_application_banned_phrases(answer)
    validate_answer_word_count(answer, 80, 155, "Communication or implementation answer")
    return answer


def relevant_experience_answer(job_description: str, snapshot: ResumeSnapshot, resume_text: str) -> str:
    return build_relevant_experience_answer(active_positioning_brief(job_description, resume_text))


def public_agency_experience_answer() -> str:
    return normalize_spaces(
        "0 years of direct public-agency or cooperative experience. My background is primarily private-sector, but it has centered on structured project delivery, stakeholder coordination, reporting, training, and solution development in complex multi-stakeholder environments."
    )


def unique_qualifications_answer(job_description: str, resume_text: str) -> str:
    return build_unique_qualifications_answer(active_positioning_brief(job_description, resume_text))


def certifications_answer(snapshot: ResumeSnapshot) -> str:
    development_text = ", ".join(snapshot.professional_development)
    if re.search(r"\bitil 4 foundation\b", development_text, re.I):
        return normalize_spaces(
            f"Relevant professional development includes {', '.join(PROFESSIONAL_DEVELOPMENT_NAMES)}."
        )
    return normalize_spaces(
        f"Relevant professional development listed in my current materials includes {development_text}."
    )


def software_inventory_answer(resume_text: str) -> str:
    found: list[str] = []
    for label, patterns in SOFTWARE_ITEMS:
        if any(re.search(pattern, resume_text, re.I) for pattern in patterns):
            found.append(label)
    if not found:
        return "Intermediate or higher experience with enterprise systems, reporting tools, CRM workflows, and AI-assisted documentation tools."
    if len(found) == 1:
        return f"Intermediate or higher experience with {found[0]}."
    return normalize_spaces(
        "Intermediate or higher experience with "
        + ", ".join(found[:-1])
        + f", and {found[-1]}."
    )


def software_inventory_answer_for_job(job_description: str, resume_text: str) -> str:
    answer = software_inventory_answer(resume_text)
    if build_resume.jd_explicitly_requires_erp(job_description):
        return answer
    named_erp = ("Aptean Intuitive", "Aptean Encompix", "Epicor Kinetic", "Microsoft Dynamics 365", "Oracle ERP")
    for label in named_erp:
        answer = re.sub(rf"(?:,\s*|\s+and\s+)?{re.escape(label)}", "", answer, flags=re.I)
    answer = re.sub(r",\s*,", ",", answer)
    answer = re.sub(r"\s+,", ",", answer)
    answer = re.sub(r",\s+and\s+\.", ".", answer)
    return normalize_spaces(answer)


def communication_answer() -> str:
    job_description = read_text(JOBS_DIR / "job_description.txt") if (JOBS_DIR / "job_description.txt").exists() else ""
    if not job_description.strip():
        job_description = (
            "Company: Target Company\n"
            "Role: Cross-Functional Implementation Role\n"
            "This role requires stakeholder communication, implementation coordination, and practical follow-through."
        )
    resume_text = approved_source_resume_text()
    brief = active_positioning_brief(job_description, resume_text)
    return build_communication_or_implementation_answer(brief, "stakeholder communication")


def company_interest_answer(job_description: str, resume_text: str) -> str:
    return build_why_company_answer(active_positioning_brief(job_description, resume_text))


def customer_profile_answer() -> str:
    return normalize_spaces(
        "The typical customers I have supported were B2B manufacturing and enterprise software environments, including manufacturing, supply chain, distribution, and adjacent operational workflows. Exact employee counts were not always disclosed, so I would answer that honestly, but the environments ranged from multi-site mid-market operations to larger enterprise organizations, including a five-site 150+ user ERP environment and an international portfolio of manufacturing clients. The main stakeholder roles typically included operations leaders, finance teams, engineering or IT partners, executive sponsors, and business users responsible for adoption, reporting, and day-to-day workflow execution."
    )


def implementation_volume_answer() -> str:
    return normalize_spaces(
        "My work has involved multiple concurrent implementations or workstreams at the same time rather than one isolated project. On the client-delivery side, I supported a portfolio of 80+ international manufacturing client engagements, and on the internal ERP side I coordinated migration, validation, training, and go-live activities across concurrent operations, finance, and engineering tracks. Exact project counts and a single average duration were not always the way those roles were measured, so I would answer this honestly as multiple active implementations or workstreams at once, with duration ranging from shorter scoped deployment phases to longer multi-phase enterprise programs that continued from discovery through post-go-live support."
    )


def implementation_success_answer() -> str:
    return normalize_spaces(
        "I measure implementation performance and success across scope clarity, testing and data accuracy, go-live readiness, user adoption, and post-launch business visibility. In practice, that means confirming the requirements are actually met, validating that data and workflows behave correctly in UAT before release, making sure users are trained and can operate in the new process, tracking whether issues are resolved quickly after go-live, and using reporting to see whether the implementation is improving decision-making or workflow performance. This approach is grounded in hands-on work leading training, testing, and go-live readiness across a five-site 150+ user ERP environment, supporting 80+ client engagements, and building 200+ KPI dashboards and reporting tools that made performance easier to see."
    )


def ai_passion_answer() -> str:
    return normalize_spaces(
        "Professionally, I have put AI to work in practical ways rather than treating it as a buzzword. It includes using Claude and AI-assisted tools to accelerate documentation, reporting, SQL troubleshooting, and operational analysis across recurring systems and implementation processes, and helping build a zero-to-one internal SMS support channel that included LivePerson LiveEngage workflows, automated greetings and closings, and AI-assisted chatbot logic. What makes AI compelling to me is when it reduces manual effort, speeds usable insight, and helps teams turn ambiguous workflow problems into something more scalable and supportable."
    )


def saas_ai_company_experience_answer() -> str:
    return normalize_spaces(
        "Yes. My background includes SaaS and technology-company experience in enterprise software environments such as Aptean and Aderant, along with hands-on work using AI-enabled tools and customer-service workflows. At Aptean, I supported industry-specific software delivery across 80+ manufacturing client engagements, and at Aderant I supported enterprise legal technology software used across 600+ law firm offices. On the AI side, I have used Claude and AI-assisted tools to accelerate documentation, reporting, SQL troubleshooting, and operational analysis, and I helped build a zero-to-one internal SMS support channel using LivePerson LiveEngage workflows, automated greetings and closings, and AI-assisted chatbot logic. I would not position that as experience at a purely AI-first employer, but it is credible SaaS and technology-company experience that included AI-enabled service and workflow exposure."
    )


def ambiguity_delivery_answer() -> str:
    return finalize_candidate_answer(
        "Ambiguity delivery answer",
        "One strong example is the new warehouse operation and Amazon Robotics program I helped launch at East West Manufacturing. The business outcome I was accountable for was a production-ready system setup that could support go-live across operations, finance, and engineering without losing control of the work across a five-site environment. What was ambiguous at the start was that there was no mature internal playbook for how the new operation, master data, cross-functional dependencies, training needs, and launch sequence should come together, so multiple workstreams had to be defined in parallel. In the first few weeks, I focused on turning that ambiguity into sequence by clarifying scope, identifying the highest-risk decisions, and separating the work into concurrent operations, finance, engineering, and system-readiness tracks. The structure I created centered on documented scope, requirements, milestone schedules, risk assumptions, role-based training materials, and change communications so the team had a practical operating rhythm rather than informal updates. I had to influence VP- and director-level stakeholders, vendors, business users, and technical teams without direct supervisory authority, mainly by making dependencies and tradeoffs visible early enough for decisions to happen before go-live. The final outcome was a production-ready launch with clearer ownership, cross-site readiness, and a more controlled path from initial requirements through training and go-live.",
        claim_first=True,
    )


def complex_project_leadership_answer() -> str:
    return finalize_candidate_answer(
        "Complex project leadership answer",
        "Yes. My project leadership experience on complex initiatives has centered on enterprise platform migration and new-program setup where multiple workstreams had to move together without a simple template. At East West Manufacturing, I led the IT program management lifecycle for an enterprise platform migration and also helped launch a production-ready system setup for a new warehouse operation and Amazon Robotics program. The biggest challenges were cross-functional dependencies across operations, finance, and engineering, migration and go-live risk, data and cutover readiness, vendor coordination, and the need to drive alignment without direct supervisory authority. I addressed that by creating structure early through scope baselines, requirements, milestone schedules, risk assumptions, testing and validation checkpoints, role-based training, and a stakeholder rhythm that kept decisions moving before issues became launch blockers. The outcome was clearer ownership across concurrent workstreams and a stronger path to migration readiness and production go-live.",
        claim_first=True,
    )


def expectation_gap_answer() -> str:
    return finalize_candidate_answer(
        "Expectation gap answer",
        "I have handled expectation gaps by naming early when a requested customization, integration, or workflow outcome would not fit the platform cleanly within the available scope or timeline. My first step was to diagnose the gap between expectation and reality by walking backward from the desired outcome into the actual workflow, then reviewing requirements, downstream reporting impacts, and technical feasibility with business stakeholders, development teams, and product or vendor partners. Once the reality was clear, I communicated the tradeoffs directly: what the customer wanted, what the platform could support through configuration, what would require higher-risk customization, and what the cost, timeline, and support implications would be. I protected trust by surfacing the constraint early, documenting scope and risk in Statements of Work and Functional Requirements Documents, and giving stakeholders a workable path forward instead of a flat no. This approach helped steer customers away from high-risk or low-value customizations, improved delivery predictability, and protected both customer confidence and long-term account health because expectations were reset before the work drifted into avoidable frustration. The main lesson for me was that bad news is usually best received when it is specific, early, and paired with an honest alternative.",
        claim_first=True,
    )


def training_strategy_leadership_answer() -> str:
    return finalize_candidate_answer(
        "Training strategy leadership answer",
        "Yes. I have led training-strategy work for large-scale technology and enterprise system implementations, especially in ERP, migration, and post-go-live environments. At East West Manufacturing, that included cross-site training, onboarding materials, and change communications tied to enterprise platform migration, a five-site user base, and the warehouse and robotics launch. At Aptean, it included designing and delivering training programs, adoption guides, and client-facing enablement resources for complex enterprise software implementations across 80+ international client engagements. My role was to turn the implementation plan into role-based readiness by defining who needed what training, creating the materials, coordinating with business and technical stakeholders, and reinforcing adoption through post-go-live support and communication. The goal was not just to train users once, but to improve go-live readiness, support adoption, and reduce reactive post-launch support volume.",
        claim_first=True,
    )


def generic_bridge_answer(job_description: str, resume_text: str) -> str:
    profile = build_resume.job_problem_profile(job_description, resume_text)
    return normalize_spaces(
        f"My strongest supported bridge for this application is experience in {build_resume.natural_problem_phrase(profile)}, backed by 80+ client engagements, 200+ reporting tools, 60+ executive workshops and QBRs, and multi-site enterprise systems ownership. I would answer this item with direct proof where available and honest bridge language where the requirement is adjacent rather than identical."
    )


def prioritization_answer() -> str:
    return finalize_candidate_answer(
        "Prioritization answer",
        "I sequence on risk and dependencies first, the work that blocks go-live or breaks if it is wrong, then use a quick impact-versus-effort read for the rest. With finance-led clients I protect the financial calendar so a cutover never threatens month-end or payments, and I do not commit a timeline until I have validated the real effort with the client. On my projects I ran that sequencing in Microsoft Project, and Smartsheet is the same discipline.",
        claim_first=True,
    )


def platform_ramp_answer() -> str:
    return finalize_candidate_answer(
        "Platform ramp answer",
        "I have not worked in Acumatica directly, but ramping fast on a new ERP is something I have done under pressure. I owned the migration from Aptean Intuitive to Epicor Kinetic and was productive in the new platform within weeks while running live cutover. I learn a system by its core workflows first, then how they map to the business, then the common client scenarios, so I would be contributing on Acumatica quickly.",
        claim_first=True,
    )


def cover_lane_reason(profile: build_resume.JobProblemProfile) -> str:
    return {
        "presales_solution": "buyer-facing discovery, solution design, and implementation handoff",
        "implementation_delivery": "complex implementations, stakeholder alignment, and go-live readiness",
        "customer_success": "adoption, renewal risk, and executive communication",
        "analytics_operations": "workflow visibility, reporting clarity, and decision support",
        "change_enablement": "training, stakeholder alignment, and adoption follow-through",
        "process_improvement": "workflow improvement, measurement, and operational follow-through",
        "corporate_strategy": "structured analysis, recommendations, and implementation follow-through",
    }.get(profile.primary_lane, build_resume.natural_problem_phrase(profile))


def scrub_cover_answer_for_job(job_description: str, text: str) -> str:
    if build_resume.jd_explicitly_requires_erp(job_description):
        return text
    updated = build_resume.scrub_erp_language_for_non_erp_text(text, job_description)
    replacements = (
        (r"\bAptean Intuitive\b", "enterprise platform"),
        (r"\bAptean Encompix\b", "software platform"),
        (r"\bEpicor Kinetic\b", "successor platform"),
        (r"\bMicrosoft Dynamics 365\b", "enterprise platform"),
        (r"\bOracle ERP\b", "enterprise platform"),
    )
    return collision_safe_substitute(updated, replacements, list_mode=True)


def cover_company_interest_answer(
    job_description: str,
    resume_text: str,
    *,
    company_name: str = "",
    role_title: str = "",
) -> str:
    return scrub_cover_answer_for_job(
        job_description,
        build_why_company_answer(active_positioning_brief(job_description, resume_text)),
    )


def cover_relevant_experience_answer(job_description: str, snapshot: ResumeSnapshot, resume_text: str) -> str:
    return scrub_cover_answer_for_job(
        job_description,
        build_relevant_experience_answer(active_positioning_brief(job_description, resume_text)),
    )


def cover_unique_qualifications_answer(job_description: str, resume_text: str) -> str:
    return scrub_cover_answer_for_job(
        job_description,
        build_unique_qualifications_answer(active_positioning_brief(job_description, resume_text)),
    )


def cover_communication_answer() -> str:
    return scrub_cover_answer_for_job(
        read_text(JOBS_DIR / "job_description.txt") if (JOBS_DIR / "job_description.txt").exists() else "",
        communication_answer(),
    )


def question_category(prompt: str) -> str:
    normalized = normalize_question(prompt)
    if "post-secondary" in normalized and "education" in normalized:
        return "education_years"
    if "direct relevant experience" in normalized and "job duties" in normalized:
        return "direct_experience_years"
    if normalized.startswith("briefly describe your relevant experience"):
        return "relevant_experience"
    if "public agencies or cooperatives" in normalized:
        return "public_agency_experience"
    if ("anything else" in normalized or "what else" in normalized) and "qualify" in normalized:
        return "unique_qualifications"
    if "certifications and licenses" in normalized:
        return "certifications"
    if "software packages" in normalized or "systems and programs" in normalized:
        return "software_inventory"
    if "public speaking" in normalized or "group facilitation" in normalized or "presentation experience" in normalized:
        return "communication"
    if "interested in joining this company" in normalized or (
        "why are you interested" in normalized and ("company" in normalized or "role" in normalized)
    ):
        return "company_interest"
    if "typical customer" in normalized and "stakeholder roles" in normalized:
        return "customer_profile"
    if "concurrent implementations" in normalized or "average length of an implementation process" in normalized:
        return "implementation_volume"
    if "implementation" in normalized and "success" in normalized and "measure" in normalized:
        return "implementation_success"
    if "passionate about ai" in normalized or "passionate about artificial intelligence" in normalized:
        return "ai_passion"
    if "put it to work for you" in normalized and ("ai" in normalized or "artificial intelligence" in normalized):
        return "ai_passion"
    if "saas or technology company" in normalized and "ai product or service" in normalized:
        return "saas_ai_company_experience"
    if "no clear delivery playbook" in normalized or "no clear playbook" in normalized:
        return "ambiguity_delivery"
    if "project leadership roles for complex initiatives" in normalized:
        return "complex_project_leadership"
    if "complex initiatives" in normalized and "challenges you faced" in normalized and "how did you address them" in normalized:
        return "complex_project_leadership"
    if "operating rhythm" in normalized and "artifacts" in normalized and "move forward" in normalized:
        return "ambiguity_delivery"
    if "established way of working" in normalized and "ambiguous at the start" in normalized:
        return "ambiguity_delivery"
    if "could not actually deliver as promised" in normalized:
        return "expectation_gap"
    if "leading training strategies" in normalized and "enterprise system implementations" in normalized:
        return "training_strategy_leadership"
    if "large-scale technology" in normalized and "what types of projects have you worked on" in normalized and "what was your role" in normalized:
        return "training_strategy_leadership"
    if "gap between expectation and reality" in normalized and ("tradeoffs" in normalized or "bad news" in normalized):
        return "expectation_gap"
    if "protect trust" in normalized and "path forward" in normalized and "tradeoffs" in normalized:
        return "expectation_gap"
    if "priorit" in normalized or "competing deadlines" in normalized or "multiple deadlines" in normalized:
        return "prioritization"
    if "acumatica" in normalized or ("platform experience" in normalized and "erp" in normalized):
        return "platform_ramp"
    return "generic_bridge"


def answer_prompt(prompt: str, job_description: str, snapshot: ResumeSnapshot, resume_text: str) -> QualificationsResponse:
    category = question_category(prompt)
    if category == "education_years":
        response = QualificationsResponse(prompt, education_years_answer(snapshot))
    elif category == "direct_experience_years":
        response = QualificationsResponse(prompt, direct_experience_years_answer(snapshot))
    elif category == "relevant_experience":
        response = QualificationsResponse(prompt, relevant_experience_answer(job_description, snapshot, resume_text))
    elif category == "public_agency_experience":
        response = QualificationsResponse(prompt, public_agency_experience_answer())
    elif category == "unique_qualifications":
        response = QualificationsResponse(prompt, unique_qualifications_answer(job_description, resume_text))
    elif category == "certifications":
        response = QualificationsResponse(prompt, certifications_answer(snapshot))
    elif category == "software_inventory":
        response = QualificationsResponse(prompt, software_inventory_answer_for_job(job_description, resume_text))
    elif category == "communication":
        response = QualificationsResponse(
            prompt,
            build_communication_or_implementation_answer(
                active_positioning_brief(job_description, resume_text),
                communication_question_topic(prompt),
            ),
        )
    elif category == "company_interest":
        response = QualificationsResponse(prompt, company_interest_answer(job_description, resume_text))
    elif category == "customer_profile":
        response = QualificationsResponse(prompt, customer_profile_answer())
    elif category == "implementation_volume":
        response = QualificationsResponse(prompt, implementation_volume_answer())
    elif category == "implementation_success":
        response = QualificationsResponse(prompt, implementation_success_answer())
    elif category == "ai_passion":
        response = QualificationsResponse(prompt, ai_passion_answer())
    elif category == "saas_ai_company_experience":
        response = QualificationsResponse(prompt, saas_ai_company_experience_answer())
    elif category == "ambiguity_delivery":
        response = QualificationsResponse(prompt, ambiguity_delivery_answer())
    elif category == "complex_project_leadership":
        response = QualificationsResponse(prompt, complex_project_leadership_answer())
    elif category == "expectation_gap":
        response = QualificationsResponse(prompt, expectation_gap_answer())
    elif category == "training_strategy_leadership":
        response = QualificationsResponse(prompt, training_strategy_leadership_answer())
    elif category == "prioritization":
        response = QualificationsResponse(prompt, prioritization_answer())
    elif category == "platform_ramp":
        response = QualificationsResponse(prompt, platform_ramp_answer())
    else:
        response = QualificationsResponse(
            prompt,
            generic_bridge_answer(job_description, resume_text),
            warning=f"QUALIFICATIONS WARNING: used generic bridge answer for unrecognized question: {prompt}",
        )
    claim_first_categories = {
        "ambiguity_delivery",
        "complex_project_leadership",
        "expectation_gap",
        "training_strategy_leadership",
        "prioritization",
        "platform_ramp",
    }
    finalized = finalize_candidate_answer(
        prompt,
        response.answer,
        claim_first=category in claim_first_categories,
        min_words=18 if category in claim_first_categories else 8,
    )
    return QualificationsResponse(response.prompt, finalized, response.warning)


def build_question_responses(
    prompts: tuple[str, ...],
    job_description: str,
    snapshot: ResumeSnapshot,
    resume_text: str,
) -> tuple[QualificationsResponse, ...]:
    return tuple(answer_prompt(prompt, job_description, snapshot, resume_text) for prompt in prompts)


def default_responses(job_description: str, snapshot: ResumeSnapshot, resume_text: str) -> tuple[QualificationsResponse, ...]:
    return (
        QualificationsResponse("Relevant experience summary", relevant_experience_answer(job_description, snapshot, resume_text)),
        QualificationsResponse("Additional qualifications", unique_qualifications_answer(job_description, resume_text)),
        QualificationsResponse("Relevant systems and tools", software_inventory_answer_for_job(job_description, resume_text)),
        QualificationsResponse("Communication, facilitation, and training experience", communication_answer()),
    )


def active_application_question_responses(
    job_description: str,
    question_path: Path = APPLICATION_QUESTIONS,
    *,
    require_prompts: bool = False,
) -> tuple[QualificationsResponse, ...]:
    prompts = (
        require_active_application_prompts(question_path)
        if require_prompts
        else load_active_application_prompts(question_path)
    )
    if not prompts:
        return ()
    _, snapshot, resume_text = selected_resume_snapshot(job_description)
    return build_question_responses(prompts, job_description, snapshot, resume_text)


def dedupe_question_responses(
    responses: tuple[QualificationsResponse, ...],
    seen_prompts: tuple[str, ...] = (),
) -> tuple[QualificationsResponse, ...]:
    seen = {normalize_question(prompt) for prompt in seen_prompts if prompt}
    filtered: list[QualificationsResponse] = []
    for response in responses:
        key = normalize_question(response.prompt)
        if key in seen:
            continue
        filtered.append(response)
        seen.add(key)
    return tuple(filtered)


# Categories where the system has no real fact to draw on (the true reason a
# role ended is Christian's alone, and a specific interviewer's product is
# unknowable ahead of the call) or where the honest answer is short, fixed
# boilerplate (the AI/SaaS exposure claim, already verified elsewhere in this
# codebase). These get a concrete fill-in script instead of a story match, so
# the build never has to invent a fact or force an unrelated story onto a
# question a story can't honestly answer.
NON_STORY_INTERVIEWER_CATEGORIES = frozenset(
    {"role_end", "product_explainer", "ai_saas_experience", "comprehension_check"}
)

_INTERVIEWER_QUESTION_ANGLES: dict[str, str] = {
    "role_end": "Lead with the clean factual answer in sentence one, then add one scope line that reinforces fit.",
    "implementation_ownership": "Lead with the closest full-lifecycle implementation example first, then separate already-live enhancement work from discovery-through-go-live ownership.",
    "product_explainer": "Translate the product into the closest supported workflow or system category first, then explain the user, the business problem, and the value in plain language.",
    "comprehension_check": "Confirm understanding directly and invite the interviewer to go deeper on anything that was unclear, without restating unrelated product or background material.",
    "functional_ownership": "Answer from the functional-ownership side: test scenarios, data accuracy, cross-team coordination, and what had to be true before release or go-live.",
    "go_live_sequence": "Frame the answer as a sequence: launch readiness, transition support, issue resolution, and post-go-live adoption continuity.",
    "ai_saas_experience": "Lead with yes in sentence one, separate SaaS and technology-company context from AI-enabled workflow exposure, and keep the AI claim bounded to supported tools, messaging workflows, and service use cases rather than implying an AI-first employer.",
    "from_scratch_delivery": "Lead with the warehouse and robotics or strongest from-scratch example first, then walk through the first weeks, the structure you created, the stakeholders you influenced, and the business outcome.",
    "complex_project_leadership": "Lead with the enterprise migration or warehouse and robotics example, then cover the business outcome, main risks, the structure you created, and how you kept cross-functional decisions moving.",
    "expectation_gap": "Lead with a scope or customization expectation-reset example, show how you diagnosed feasibility and downstream risk, then close with the tradeoff conversation and the healthier path you drove.",
    "training_strategy_leadership": "Lead with the cross-site ERP or client-implementation training example first, then explain audience segmentation, training artifacts, adoption support, and the readiness or support-volume outcome.",
    "rapid_ramp": "Acknowledge the domain ramp directly, then bridge to implementation discipline and how you learn the business language quickly in context.",
    "stakeholder_tension": "Lead with the tension first, then show how you created clarity, surfaced tradeoffs, and kept the work moving without overstating authority.",
}


def interviewer_question_category(prompt: str) -> str:
    """Classify a recent-interviewer prompt into a stable category key so the
    guide builders can decide, per question, whether a real story can answer
    it (most categories) or whether it needs a fact-bound fill-in script
    instead (NON_STORY_INTERVIEWER_CATEGORIES)."""
    normalized = normalize_question(prompt)
    if "why did your most recent role end" in normalized or "why are you leaving" in normalized or "role end" in normalized:
        return "role_end"
    if (
        "implementation you owned" in normalized
        or "implementation flow" in normalized
        or "full implementation" in normalized
        or "discovery, configuration, data migration" in normalized
        or ("implementation" in normalized and "cycle" in normalized)
    ):
        return "implementation_ownership"
    if "does that make sense" in normalized:
        return "comprehension_check"
    if "target audience" in normalized or "business value" in normalized or "product offerings" in normalized:
        return "product_explainer"
    if "uat" in normalized or "data validation" in normalized or "functional testing" in normalized or "testing" in normalized:
        return "functional_ownership"
    if "go-live" in normalized:
        return "go_live_sequence"
    if "saas or technology company" in normalized and "ai product or service" in normalized:
        return "ai_saas_experience"
    if "no clear delivery playbook" in normalized or "operating rhythm" in normalized or "established way of working" in normalized:
        return "from_scratch_delivery"
    if "project leadership roles for complex initiatives" in normalized or ("complex initiatives" in normalized and "how did you address them" in normalized):
        return "complex_project_leadership"
    if "could not actually deliver as promised" in normalized or "gap between expectation and reality" in normalized or "tradeoffs or bad news" in normalized:
        return "expectation_gap"
    if "leading training strategies" in normalized or "enterprise system implementations" in normalized and "what was your role" in normalized:
        return "training_strategy_leadership"
    if "ramp" in normalized or "new industry" in normalized or "learning curve" in normalized:
        return "rapid_ramp"
    if "stakeholder" in normalized or "difficult situation" in normalized or "convince" in normalized:
        return "stakeholder_tension"
    return "generic_bridge"


def interviewer_question_answer_angle(prompt: str, job_description: str, resume_text: str) -> str:
    category = interviewer_question_category(prompt)
    angle = _INTERVIEWER_QUESTION_ANGLES.get(category)
    if angle:
        return angle
    profile = build_resume.job_problem_profile(job_description, resume_text)
    return normalize_spaces(
        f"Lead with a direct answer in sentence one, use one strongest proof example tied to {candidate_problem_phrase(profile)}, and close with why that experience transfers cleanly to this role."
    )


def interviewer_question_factual_script(prompt: str, job_description: str, resume_text: str) -> str:
    """Real, spoken-ready scripts for the NON_STORY_INTERVIEWER_CATEGORIES.

    These categories depend on something a story cannot always supply. The
    role-ending answer uses a verified reorganization fact when the resume
    supports it and otherwise stays factual without inventing a reason. The
    product answer describes the supported workflow specialization instead of
    guessing at an unnamed offering. The AI/SaaS exposure category restates a claim
    already verified elsewhere in this codebase (the federal resume's
    "AI-enabled workflows" framing), so it gets a complete, ready-to-say
    script rather than a fill-in. The comprehension-check category ("does
    that make sense?") is a plain acknowledgment, not a product question or a
    story prompt, so it gets its own short, complete script instead of being
    folded into the product placeholder.
    """
    category = interviewer_question_category(prompt)
    profile = build_resume.job_problem_profile(job_description, resume_text)
    if category == "role_end":
        if re.search(r"\b(?:position|role) impacted by (?:company )?reorgani[sz]ation\b", resume_text, re.I):
            return (
                "My most recent position was impacted by company reorganization. "
                f"I am now looking for a role where I can apply {candidate_problem_phrase(profile)} more consistently, "
                "and this role stood out for that reason."
            )
        return (
            "I keep the reason a role ended factual and brief, without adding details the record does not support. "
            f"The important next step for me is finding a role where I can apply {candidate_problem_phrase(profile)} more consistently, "
            "and this role stood out for that reason."
        )
    if category == "product_explainer":
        return (
            f"My specialization has been in systems and workflows tied to {candidate_problem_phrase(profile)}. "
            "When I explain a product, I start with who uses it, the business problem it solves, and the value in plain language."
        )
    if category == "ai_saas_experience":
        return (
            "Yes, I've worked in SaaS and technology-company environments, and that included AI-enabled workflows in "
            "supported tools, messaging, and service use cases. I haven't built AI products myself, but I've evaluated "
            "AI-enabled workflows, explained them to stakeholders, and applied them where they actually fit."
        )
    if category == "comprehension_check":
        return (
            "Yes, that makes sense. If any part of what I said was unclear, happy to go deeper on it."
        )
    return ""


def recent_interviewer_question_prep_items(
    job_description: str,
    company_name: str,
    role_title: str,
    *,
    jobs_dir: Path = JOBS_DIR,
    limit: int = 10,
) -> tuple[InterviewQuestionPrep, ...]:
    prompts = interview_context.recent_interviewer_questions(jobs_dir, company_name, role_title, limit=limit)
    if not prompts:
        return ()
    _, _, resume_text = selected_resume_snapshot(job_description)
    items = [
        InterviewQuestionPrep(
            prompt=prompt,
            answer_angle=interviewer_question_answer_angle(prompt, job_description, resume_text),
            category=interviewer_question_category(prompt),
        )
        for prompt in prompts
    ]
    return tuple(item for item in items if item.category != "comprehension_check")
