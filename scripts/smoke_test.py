#!/usr/bin/env python3
"""Smoke tests for the resume automation scripts."""

from __future__ import annotations

import contextlib
import importlib
import io
import json
import os
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime, timedelta
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from xml.sax.saxutils import escape


SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

MAJOR_SCRIPTS = (
    "business_context",
    "claude_review_bundle",
    "build_claude_prompt",
    "build_claude_review_packet",
    "build_federal_resume",
    "build_federal_cover_letter",
    "build_federal_interview_cheat_sheet",
    "build_federal_detailed_interview_guide",
    "build_followup_email",
    "build_resume",
    "build_cover_letter",
    "commercial_resume_model",
    "build_standard_qualifications_statement",
    "build_interview_cheat_sheet",
    "build_interview_companions",
    "build_first_90_days",
    "build_detailed_interview_guide",
    "build_interview_review",
    "build_debrief_analysis",
    "build_interview_followup",
    "build_linkedin_calendar",
    "build_linkedin_update",
    "build_networking_outreach",
    "build_post_round",
    "build_salary_guide",
    "build_skills_database",
    "build_application_checklist",
    "build_thank_you",
    "build_general_advice",
    "cleanup_output",
    "cleanup_render_checks",
    "extract_writing_examples",
    "federal_supporting_docs",
    "interview_stage",
    "interview_context",
    "job_search_guidance",
    "job_context_archive",
    "post_interview_debrief",
    "prose_engine",
    "render_checks",
    "run_federal_resume_workflow",
    "run_resume_workflow",
    "reset_jobs",
    "refresh_claude_review_bundle",
    "utils",
    "writing_eval",
)
CONFIG_MODULES = (
    "config.language_rules",
    "config.job_profiles",
)

DUMMY_JOB_DESCRIPTION = """
Company: Smoke Test Systems
Role: Implementation Consultant
This role owns implementation, configuration, requirements, scope, data migration,
integration, go-live readiness, hypercare, stakeholder communication, adoption,
reporting, and customer onboarding for enterprise software customers.
"""

MICRO1_SMOKE_JOB_DESCRIPTION = """
Company: micro1
Job Title: Professional Services Consultant
micro1 is an AI data lab for training frontier models and evaluating AI agents.
This role turns ambiguous client problems into structured decisions and clear written guidance.
Required skills include analytical thinking, stakeholder communication, slide-deck preparation,
 contract review, distributed-team collaboration, and synthesizing complex concepts.
"""

VENSURE_STYLE_JOB_DESCRIPTION = """
Company: VensureHR
Job Title: Senior Implementation Consultant

About Us
Vensure Employer Solutions is the largest privately held organization in the HR technology and service sector.

Position Summary
The Senior Implementation Consultant serves as a highly experienced subject matter expert responsible for leading complex implementations and ensuring an exceptional client experience.

Essential Duties and Responsibilities
Lead the implementation process for new clients from project initiation through completion.
Provide expert guidance on configuration, system setup, carrier connections, and implementation best practices.
Manage multiple implementations simultaneously while prioritizing workload, meeting deadlines, and maintaining accuracy and organization.

Knowledge, Skills, and Abilities
Strong understanding of implementation and administration processes, including system configuration, eligibility rules, carrier connections, and workflow best practices.

Education & Experience
Minimum of 5 years of related implementation experience.

Required Licenses And/Or Certifications
Industry certification preferred.
"""

LANE_JOB_DESCRIPTIONS = {
    "change_enablement": """
Role: Change Enablement Lead
The role owns change enablement, change management, organizational effectiveness,
enterprise transformation, operating model adoption, resistance planning, future of
work, team member experience, stakeholder alignment, manager readiness, communications,
training, and measurable adoption outcomes.
""",
    "presales_solution": """
Role: Pre-Sales Solution Consultant
The role requires pre-sales, presales, solution consulting, solution consultant,
sales engineer support, demos, demo delivery, discovery, value selling, ROI framing,
executive buyers, technical discovery, product fit, and sales partnership.
""",
    "customer_success": """
Role: Customer Success Manager
The role manages customer success, renewal, retention, expansion, book of business,
ARR, NRR, GRR, churn risk, QBRs, account health, value realization, adoption,
executive relationships, and customer lifecycle outcomes.
""",
    "implementation_delivery": """
Role: Implementation Project Manager
The role owns implementation, go-live, configuration, data migration, integration,
hypercare, project management, scope, deliverables, timeline, requirements, UAT,
customer onboarding, and launch readiness.
""",
    "analytics_operations": """
Role: Analytics Operations Analyst
The role focuses on analytics, data analysis, dashboards, KPI reporting, forecast
quality, operational efficiency, process improvement, business intelligence, workflow
measurement, performance reporting, and decision support.
""",
}

OLLIE_ANALYTICS_JOB_DESCRIPTION = """
Company: Ollie
Job Title: Data Analyst, Retention

Ollie makes human-grade pet food tailored to each dog's nutritional needs and delivers it directly to customers' doors.

Ollie is looking for a Data Analyst to join the cross functional team responsible for building subscription retention and member loyalty.
This role is ideal for someone who is highly analytical, comfortable working with imperfect business questions,
and excited to translate data into decisions.

What You'll Do:
Drive a data-driven approach to optimizing the Ollie member experience alongside peers from product, marketing, CX, and research.
Lead analysis across Ollie's lifecycle comms, mobile apps, and new member services.
Support the team through ad hoc analysis, hypothesis valuation, test design and measurement, and metric definition.

Who You Are:
You have experience in data analytics, growth analytics, lifecycle analytics, or retention analytics in a subscription business.
You balance SQL, Excel, test design and measurement, and subscription-based metrics including retention, churn, LTV, reactivation, and cohort analysis.
You are comfortable sourcing and joining information from curated reporting tools, production databases, Segment, MixPanel, Looker, and Snowflake.
"""

SOURCEWELL_QUALIFICATIONS_JOB_DESCRIPTION = """
Company: Sourcewell
Job Title: Senior Solution Consultant

The Senior Solution Consultant role supports innovation work from Discovery to Pilot to Prove to Scale. The role develops project plans, coordinates cross-functional teams, contributes research and solution ideas, shapes business cases, works with CRM systems such as Salesforce, uses Excel and Power BI reporting, and requires public speaking or sales presentation experience.
"""

OLLIE_RESUME_TEXT = """
Professional Summary
Enabled real-time executive decision-making by building 200+ SQL-based BI and reporting tools that replaced raw exports and verbal status updates.
Facilitated 60+ executive workshops and QBRs that aligned technology decisions to business objectives and measurable outcomes.
Built more than 200 dashboards and reporting tools that converted operational data into clearer executive and operational decisions.
"""

FEDERAL_DUMMY_JOB_DESCRIPTION = """
Agency: Department of Veterans Affairs
Role: IT Program Manager
This federal role leads enterprise implementation, data migration, acquisition support,
program governance, executive stakeholder briefings, reporting, change management,
and cross-functional delivery across a multi-site environment.
"""

FEDERAL_AI_JOB_DESCRIPTION = """
Agency: Department of the Treasury
Role: IT Specialist (AI)
Selective Factor: Demonstrated experience implementing AI solutions in production or test environments.
GS-15: Directing collaborative projects, leading enterprise cloud-native modernization efforts,
and developing production-grade AI-enabled systems.
GS-14: Collaborating with Senior Executives to align emerging technology adoption with Administration priorities
and conducting cybersecurity risk assessments for AI-enabled systems.
GS-13: Applying cybersecurity practices in AI or cloud environments and participating in Agile development teams
during sprint planning or execution.
GS-12: Assisting in designing, developing, testing, or deploying AI models and prototypes.
"""

FEDERAL_STANDALONE_AGENCY_DESCRIPTION = """
Position: IT Specialist (Customer Support Systems Analysis)
Department of Veterans Affairs
Deputy Assistant Secretary for Information and Technology

Qualifications
Attention to Detail - Is thorough when performing work and conscientious about attending to detail.
Customer Service - Works with clients and customers.
"""

PROCORE_JOB_DESCRIPTION = """
Company: Procore
Job Title: Senior Solutions Architect

We’re looking for a Senior Solutions Architect to join Procore’s Technical Services Team. In this role, you’ll lead technical scoping and requirements gathering to design solutions around integrations and data migrations to ultimately bring success to our customers through our cloud based software.

As a Senior Solutions Architect, you’ll partner with clients, prospects, partners, and our engineering team to analyze requirements, design solutions, and lead technical integration and data-migration efforts. Use your technical consulting, project management, and customer engagement skills to deliver high value professional services and successful integrations.

What you’ll do:
Lead technical scoping and requirements gathering with clients, prospects, and partners.
Design and implement custom integrations and data migrations leveraging Procore's APIs.
Draft and manage Statements of Work (SOW), including requirements and pricing.
Project manage technical development efforts including requirements, design, testing, and delivery.
Collaborate with Sales, Customer Success, and technical teams to align solutions with business outcomes.
Serve as a subject-matter expert for integrations with ERP/accounting systems.
"""

PROCORE_RESUME_TEXT = """
Managed the full ERP lifecycle for enterprise manufacturing clients across discovery, requirements definition, configuration, data migration, integration, testing, go-live, and post-go-live support.
Protected migration stability by leading implementation readiness, scope alignment, sandbox testing, UAT validation, and targeted training across concurrent program tracks.
Converted complex implementation, data migration, integration, and customization needs into statements of work and functional requirements that clarified scope, milestones, and cost baselines before build work began.
Turned ambiguous operations, finance, and engineering needs into scoped system recommendations, vendor tradeoffs, cost and timeline options, and risk-aware implementation plans for directors and VPs before build work began.
Facilitated 60+ executive workshops and QBRs that aligned technology decisions to business objectives and measurable outcomes.
Built 200+ dashboards and KPI reporting tools that gave leaders clearer visibility into operational performance and decision-making.
"""

PEARSON_MEASUREMENT_JOB_DESCRIPTION = """
Company: Pearson
Job Title: Lead Specialist, Measurement

This role aligns to industry level title of Lead Specialist, Applied AI Measurement & Learning Systems

Measurement & Learning | School Assessment

Role Overview
We are seeking a Lead Specialist, Measurement to support the design, validation, and continuous improvement of intelligent measurement systems and AI-assisted assessment and learning workflows.

This role sits within the Measurement & Learning team and operates at the intersection of Measurement, Content Development, Learning Science, Technology, and AI Science.

Validation, Monitoring & Quality Frameworks
Design and oversee validation studies evaluating the quality, appropriateness, effectiveness, interpretability, and educational defensibility of AI-assisted systems, outputs, and workflows.
Develop frameworks for ongoing monitoring, drift detection, and continuous improvement of AI-assisted systems over time.
Ensure AI-assisted outputs meet expectations for validity, accuracy, fairness, accessibility, and instructional appropriateness.

Cross-Functional Collaboration
Partner with Technology teams implementing AI infrastructure and intelligent systems.
Partner with content teams authoring assessment and learning content.
"""

GUIDEHOUSE_ERP_ASSESSMENT_JOB_DESCRIPTION = """
Company: Guidehouse
Job Title: ERP Assessment Leader

About Us
Guidehouse is a global AI-led professional services firm delivering advisory, technology, and managed services to the commercial and government sectors.

What You Will Do
Guidehouse is seeking a highly skilled ERP Assessment Lead to join a complex financial modernization project.
Serve as the state's lead responsible for translating executive objectives into an integrated strategy, roadmap, and product-level decisions while ensuring alignment across oversight, delivery, and stakeholder workstreams.
Will be part of a larger program team accountable for delivering a successful program turnaround with budget, timing, and quality metrics.
Ability to maintain complete line of sight across program activities and speak to the status of each initiative.

What You Will Need
Experience in consulting, including large public-sector ERP or transformation programs.
Experience leading planning efforts, roadmaps, or operating model design.
Strong executive-facing presence with ability to drive decisions under ambiguity.
"""


class SmokeFailure(AssertionError):
    pass


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise SmokeFailure(message)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w+.#'-]+\b", text))


def import_major_scripts() -> dict[str, object]:
    modules: dict[str, object] = {}
    failures: list[tuple[str, Exception]] = []
    
    for module_name in MAJOR_SCRIPTS:
        try:
            modules[module_name] = importlib.import_module(module_name)
        except Exception as error:
            failures.append((module_name, error))
    
    if failures:
        error_lines = [f"{name}: {type(error).__name__}: {error}" for name, error in failures]
        raise SmokeFailure(
            f"Failed to import {len(failures)} major script(s):\n  " + "\n  ".join(error_lines)
        )
    
    return modules


def import_config_files() -> None:
    for module_name in CONFIG_MODULES:
        try:
            importlib.import_module(module_name)
        except AssertionError as error:
            raise SmokeFailure(f"{module_name} failed config validation: {error}") from error
    print("Config validation import PASSED: language_rules.py and job_profiles.py loaded cleanly.")


def test_validate_inputs(build_resume: object) -> None:
    result = build_resume.validate_inputs(DUMMY_JOB_DESCRIPTION)
    assert_true(isinstance(result, str), "validate_inputs() did not return a string")
    assert_true("Smoke Test Systems" in result, "validate_inputs() did not return the supplied dummy job text")


def test_validate_inputs_strips_linkedin_boilerplate(build_resume: object) -> None:
    linkedin_snippet = """
Company: Storm3
Role: Technical Implementation Manager
Responsibilities
Own delivery of workflows to onboard customers
Work closely with customers to understand their clinical data needs
Interested in applying? Please click on the Easy Apply button
Storm3 is a HealthTech recruitment firm with clients across major Tech hubs in North America. Follow the Storm3 LinkedIn page for the latest jobs and intel
"""
    result = build_resume.validate_inputs(linkedin_snippet)
    assert_true("Own delivery of workflows" in result, "validate_inputs() should preserve the usable job content")
    assert_true("Easy Apply" not in result, "validate_inputs() should strip Easy Apply boilerplate from LinkedIn snippets")
    assert_true("LinkedIn page" not in result, "validate_inputs() should strip recruiter/footer boilerplate from LinkedIn snippets")


def test_role_requirement_text_resumes_after_about_us_sections(build_resume: object) -> None:
    scoped = build_resume.role_requirement_text(VENSURE_STYLE_JOB_DESCRIPTION)
    assert_true(
        "Lead the implementation process for new clients from project initiation through completion." in scoped,
        "role_requirement_text() should preserve duties after an About Us section when the JD uses Position Summary and Essential Duties headers",
    )
    assert_true(
        "Strong understanding of implementation and administration processes" in scoped,
        "role_requirement_text() should preserve Knowledge, Skills, and Abilities content after boilerplate sections",
    )
    assert_true(
        "largest privately held organization" not in scoped,
        "role_requirement_text() should still strip the About Us company overview boilerplate",
    )

    resume_text = (
        "Professional Summary\n"
        "Senior implementation consultant with experience in go-live, configuration, data migration, integration, "
        "and requirements definition.\n"
        "Professional Experience\n"
        "Implementation Consultant\n"
        "Led client implementations from kickoff through go-live.\n"
    )
    profile = build_resume.job_problem_profile(VENSURE_STYLE_JOB_DESCRIPTION, resume_text)
    assert_true(
        "Implementation Delivery" in profile.direct_matches,
        "job_problem_profile() should detect direct implementation evidence after role_requirement_text() resumes on Vensure-style section headers",
    )
    assert_true(
        "Benefits Administration Domain" in profile.specialty_gaps,
        "job_problem_profile() should surface benefits-domain specialty gaps separately from implementation fit on Vensure-style roles",
    )


def test_output_target_name_prefers_company_and_role(build_resume: object) -> None:
    job_description = """
Company: State Farm
Role: Digital Experience Analytics Analyst
This role supports digital analytics, Adobe Analytics, journey analysis, and customer experience reporting.
"""
    target_name = build_resume.extract_output_target_name(job_description)
    assert_true(
        target_name == "State Farm - Digital Experience Analytics Analyst",
        f"extract_output_target_name() should include company and role; got {target_name!r}",
    )
    candidates = build_resume.output_name_candidates(job_description)
    assert_true(candidates[0] == target_name, "output_name_candidates() should prefer the exact company-role name first")
    assert_true("State Farm" in candidates, "output_name_candidates() should keep the company-only fallback for older files")


def test_choose_resume(build_resume: object) -> None:
    presales_job = "Pre-sales solution consulting role with demos, discovery, ROI, and executive buyers."
    implementation_job = "Implementation delivery role with configuration, data migration, go-live, scope, and requirements."

    assert_true(
        build_resume.choose_resume(presales_job) == build_resume.PRESALES_CSM_RESUME,
        "choose_resume() did not select the Pre-Sales/CSM resume for pre-sales signals",
    )
    assert_true(
        build_resume.choose_resume(implementation_job) == build_resume.IMPLEMENTATION_RESUME,
        "choose_resume() did not select the Implementation resume for implementation signals",
    )


def test_lane_profiles_and_summaries(build_resume: object) -> None:
    for expected_lane, job_description in LANE_JOB_DESCRIPTIONS.items():
        profile = build_resume.job_problem_profile(job_description)
        assert_true(
            profile.primary_lane == expected_lane,
            f"job_problem_profile() returned {profile.primary_lane!r} for {expected_lane!r}",
        )

        summary = build_resume.build_problem_first_summary(job_description)
        summary_words = word_count(summary)
        assert_true(
            build_resume.PROFESSIONAL_SUMMARY_MIN_WORDS <= summary_words <= build_resume.PROFESSIONAL_SUMMARY_MAX_WORDS,
            (
                f"summary for {expected_lane!r} has {summary_words} words; expected "
                f"{build_resume.PROFESSIONAL_SUMMARY_MIN_WORDS}-{build_resume.PROFESSIONAL_SUMMARY_MAX_WORDS}"
            ),
        )
        assert_true("--" not in summary, f"summary for {expected_lane!r} contains double dashes")
        assert_true(
            not build_resume.contains_first_person(summary),
            f"summary for {expected_lane!r} contains first-person pronouns",
        )
        assert_true(
            not build_resume.contains_ai_writing_word(summary),
            f"summary for {expected_lane!r} contains banned AI-writing words",
        )
        assert_true(
            not any(sentence.startswith("That") for sentence in build_resume.summary_sentences(summary)),
            f"summary for {expected_lane!r} contains a sentence that starts with 'That'",
        )
        assert_true(
            not re.search(r"\b[Tt]hat\b", summary),
            f"summary for {expected_lane!r} still uses the word 'that': {summary}",
        )
        assert_true(
            "Teams get" not in summary and "Leaders get" not in summary,
            f"summary for {expected_lane!r} should avoid stock summary closers: {summary}",
        )
        sentences = build_resume.summary_sentences(summary)
        assert_true(
            len(sentences) == 3,
            f"summary for {expected_lane!r} must have 3 sentences; found {len(sentences)}: {summary}",
        )
        assert_true(
            summary.count(";") <= 1,
            f"summary for {expected_lane!r} has too many semicolons: {summary}",
        )
        assert_true(
            not build_resume.contains_prompt_leak(summary),
            f"summary for {expected_lane!r} contains prompt-like phrasing: {summary}",
        )
        opening = sentences[0]
        comma_items = [item.strip() for item in opening.split(",") if item.strip()]
        assert_true(
            len(comma_items) <= 4,
            f"summary opening for {expected_lane!r} contains more than four comma-separated items: {opening}",
        )


def test_first_person_detector_ignores_role_level_i(build_resume: object) -> None:
    assert_true(
        not build_resume.contains_first_person("Account Manager I, Fulfillment"),
        "contains_first_person() should not treat role-level Roman numeral I as a pronoun",
    )
    assert_true(
        not build_resume.contains_first_person("One of the largest eCommerce operations in US retail."),
        "contains_first_person() should not treat uppercase US as the pronoun 'us'",
    )
    assert_true(
        build_resume.contains_first_person("I led customer adoption and renewal-risk recovery work."),
        "contains_first_person() should still catch real first-person pronoun usage",
    )


def test_federal_agency_extraction(build_resume: object) -> None:
    detected = build_resume.extract_company_name(FEDERAL_DUMMY_JOB_DESCRIPTION)
    assert_true(
        detected == "Department of Veterans Affairs",
        f"extract_company_name() should detect agency lines for federal postings; got {detected!r}",
    )


def test_federal_source_load(build_federal_resume: object) -> None:
    source = build_federal_resume.load_federal_source()
    assert_true(len(source.roles) >= 4, "Federal source should load at least four roles")
    assert_true(
        all(role.company_summary and role.job_summary for role in source.roles),
        "Federal source roles should include both company and job summaries",
    )
    assert_true(
        any("Claude" == skill for skill in source.technical_skills),
        "Federal source should preserve approved Claude workflow evidence in technical skills",
    )
    assert_true(
        any("Codex" == skill for skill in source.technical_skills),
        "Federal source should preserve approved Codex workflow evidence in technical skills",
    )
    blocks = {role.company + "|" + role.title: "\n".join(role.bullets) for role in source.roles}
    assert_true(
        all(term in blocks["East West Manufacturing|Enterprise Systems Manager"] for term in ("views", "stored procedures", "reporting data models")),
        "Confirmed SQL-object evidence must load under East West.",
    )
    assert_true(
        all(term in blocks["Aptean|Customer Success Consultant"] for term in ("Azure DevOps/TFS", "Git/GitHub"))
        and not any("Azure DevOps/TFS" in text or "Git/GitHub" in text for key, text in blocks.items() if not key.startswith("Aptean|")),
        "Confirmed version-control evidence must load under Aptean only.",
    )
    assert_true(
        all(term in blocks["The Home Depot|Support Operations Analyst"] for term in ("contact center", "service level", "customer interaction data")),
        "Confirmed contact-center evidence must load under The Home Depot.",
    )
    assert_true(
        "recovery testing" in blocks["Aderant|Interim Systems Administrator"]
        and not re.search(r"\b(?:owned|administered|led)\b[^.]{0,80}\b(?:backup|restore|recovery)\b", blocks["Aderant|Interim Systems Administrator"], re.I),
        "Aderant recovery evidence must stay capped at supported/performed language.",
    )


def test_federal_standard_essay_responses(build_federal_resume: object) -> None:
    prompts = build_federal_resume.load_federal_standard_essays()
    assert_true(
        len(prompts) >= 4,
        "Federal standard essay source should include the full question bank",
    )
    source = build_federal_resume.load_federal_source()
    audit = build_federal_resume.federal_requirement_audit(source, FEDERAL_AI_JOB_DESCRIPTION)
    document = build_federal_resume.build_qualifications_document(
        source,
        FEDERAL_AI_JOB_DESCRIPTION,
        build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS[0],
        audit,
    )
    text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    assert_true(
        "STANDARD FEDERAL ESSAY RESPONSES" in text,
        "Federal qualifications documents should include the standard essay-response section",
    )
    assert_true(
        "commitment to the Constitution" in text,
        "Federal essay section should preserve the Constitution-service prompt",
    )
    assert_true(
        "Executive Order 13985" in text and "Executive Order 14028" in text,
        "Federal essay section should carry forward the stored Executive Order response",
    )
    assert_true(
        "AI-enabled workflows" in text or "secure modernization" in text,
        "Federal essay responses should add role-specific federal tailoring language",
    )


def test_federal_qualifications_append_additional_questions_and_recent_interview_prep(
    build_federal_resume: object,
) -> None:
    source = build_federal_resume.load_federal_source()
    audit = build_federal_resume.federal_requirement_audit(source, FEDERAL_AI_JOB_DESCRIPTION)
    profile = build_federal_resume.job_problem_profile(
        FEDERAL_AI_JOB_DESCRIPTION,
        build_federal_resume.source_visible_text(source),
    )
    questions = build_federal_resume.federal_application_questions(FEDERAL_AI_JOB_DESCRIPTION, profile.primary_lane)
    duplicate_prompt = questions.competencies[0].name if questions.competencies else "Oral Communication"
    original_active = build_federal_resume.question_prep.active_application_question_responses
    original_recent = build_federal_resume.question_prep.recent_interviewer_question_prep_items
    try:
        build_federal_resume.question_prep.active_application_question_responses = lambda _job_description: (
            build_federal_resume.question_prep.QualificationsResponse(duplicate_prompt, "Duplicate should not append twice."),
            build_federal_resume.question_prep.QualificationsResponse(
                "Please list all software packages, systems, and programs for which you rate your skills at an intermediate or higher level.",
                "Intermediate or higher experience with Power BI, SQL, and ServiceNow.",
            ),
        )
        build_federal_resume.question_prep.recent_interviewer_question_prep_items = lambda *_args, **_kwargs: (
            build_federal_resume.question_prep.InterviewQuestionPrep(
                "Why did your most recent role end?",
                "Lead with the clean factual answer in sentence one, then add one scope line that reinforces fit.",
                category="role_end",
            ),
        )
        document = build_federal_resume.build_qualifications_document(
            source,
            FEDERAL_AI_JOB_DESCRIPTION,
            build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS[0],
            audit,
        )
    finally:
        build_federal_resume.question_prep.active_application_question_responses = original_active
        build_federal_resume.question_prep.recent_interviewer_question_prep_items = original_recent

    text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    assert_true(
        ("QUESTION 1 - SPECIALIZED EXPERIENCE" in text or "CORE EXPERIENCE" in text)
        and "STANDARD FEDERAL ESSAY RESPONSES" in text
        and "ADDITIONAL APPLICATION QUESTIONS" in text
        and "RECENT INTERVIEW QUESTIONS TO BE READY FOR" in text,
        "Federal qualifications documents should keep posting-derived content first, then append additional application questions and recent interviewer prep.",
    )
    assert_true(
        text.count(duplicate_prompt) == 1,
        f"Federal qualifications documents should not duplicate prompts already covered by posting-derived sections; got {text.count(duplicate_prompt)} instances of {duplicate_prompt!r}",
    )
    assert_true(
        "Why did your most recent role end?" in text
        and "position was impacted by company reorganization" in text
        and "Fill in" not in text,
        "Federal qualifications documents should render the verified reorganization fact without a placeholder; "
        f"got {text!r}",
    )


def test_federal_summary_structure(build_federal_resume: object) -> None:
    source = build_federal_resume.load_federal_source()
    summary = build_federal_resume.build_gs14_summary(source, FEDERAL_DUMMY_JOB_DESCRIPTION)
    summary_words = word_count(summary)
    sentences = [part.strip() for part in re.split(r"(?<=[.])\s+", summary) if part.strip()]
    assert_true(
        build_federal_resume.MIN_SUMMARY_WORDS <= summary_words <= build_federal_resume.MAX_SUMMARY_WORDS,
        (
            f"Federal summary has {summary_words} words; expected "
            f"{build_federal_resume.MIN_SUMMARY_WORDS}-{build_federal_resume.MAX_SUMMARY_WORDS}"
        ),
    )
    assert_true(len(sentences) == 3, f"Federal summary should have 3 sentences; found {len(sentences)}: {summary}")
    assert_true("VP-" in summary or "director-level" in summary, "Federal summary should surface senior-scope evidence")
    assert_true("--" not in summary, "Federal summary should not contain double dashes")


def test_federal_ai_summary_and_selection(build_federal_resume: object) -> None:
    source = build_federal_resume.load_federal_source()
    summary = build_federal_resume.build_gs14_summary(source, FEDERAL_AI_JOB_DESCRIPTION)
    assert_true(
        any(term in summary for term in ("AI workflow", "Codex and Claude", "Claude-assisted", "AI-enabled systems")),
        f"Federal AI summary should surface supported AI evidence; got {summary}",
    )
    profile = build_federal_resume.job_problem_profile(
        FEDERAL_AI_JOB_DESCRIPTION,
        build_federal_resume.source_visible_text(source),
    )
    bullet_groups = build_federal_resume.selected_bullets_by_role(
        source,
        FEDERAL_AI_JOB_DESCRIPTION,
        profile.primary_lane,
        build_federal_resume.keyword_set(FEDERAL_AI_JOB_DESCRIPTION),
        build_federal_resume.FEDERAL_LAYOUT_PROFILES[0],
    )
    selected_text = "\n".join(bullet for group in bullet_groups for bullet in group)
    assert_true(
        any(term in selected_text for term in ("Codex", "Claude", "AI-assisted chatbot", "LivePerson")),
        "Federal AI bullet selection should keep at least one AI-supporting proof point visible",
    )


def test_federal_requirement_audit_and_keywords(build_federal_resume: object) -> None:
    source = build_federal_resume.load_federal_source()
    audit = build_federal_resume.federal_requirement_audit(source, FEDERAL_AI_JOB_DESCRIPTION)
    labels = {coverage.bucket.label for coverage in audit.buckets}
    assert_true("Selective Factor" in labels, f"Federal audit should parse Selective Factor buckets; got {labels}")
    assert_true("GS-15" in labels and "GS-12" in labels, f"Federal audit should parse GS-level buckets; got {labels}")
    keyword_text = " | ".join(audit.keyword_targets)
    assert_true(
        any(term in keyword_text for term in ("AI-enabled systems", "workflow automation", "risk management")),
        f"Federal keyword targets should include job-description and government-style qualification language; got {keyword_text}",
    )


def test_federal_layouts_stay_at_ten_point(build_federal_resume: object) -> None:
    assert_true(
        all(layout.font_size >= 10.0 for layout in build_federal_resume.FEDERAL_LAYOUT_PROFILES),
        "Federal resume layouts must not shrink body text below 10pt",
    )
    assert_true(
        all(layout.font_size >= 10.0 for layout in build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS),
        "Federal qualifications layouts must not shrink body text below 10pt",
    )


def test_federal_visibility_report_tracks_selected_requirements(build_federal_resume: object) -> None:
    source = build_federal_resume.load_federal_source()
    audit = build_federal_resume.federal_requirement_audit(source, FEDERAL_AI_JOB_DESCRIPTION)
    profile = build_federal_resume.job_problem_profile(
        FEDERAL_AI_JOB_DESCRIPTION,
        build_federal_resume.source_visible_text(source),
    )
    bullet_groups = build_federal_resume.selected_bullet_candidates_by_role(
        source,
        FEDERAL_AI_JOB_DESCRIPTION,
        profile.primary_lane,
        build_federal_resume.keyword_set(FEDERAL_AI_JOB_DESCRIPTION),
        build_federal_resume.FEDERAL_LAYOUT_PROFILES[0],
        audit,
    )
    visible_audit = build_federal_resume.apply_selection_visibility(source, audit, bullet_groups)
    reference_lines = "\n".join(build_federal_resume.selected_bullet_reference_lines(source, bullet_groups))
    assert_true(
        any(coverage.selected_refs for coverage in visible_audit.buckets if coverage.bucket.label == "Selective Factor"),
        "Federal visibility audit should mark at least one Selective Factor evidence point as visible in the selected resume",
    )
    assert_true(
        any(token in reference_lines for token in ("East West Manufacturing [8]", "The Home Depot [1]")),
        f"Federal bullet reference report should surface the strongest AI-adjacent bullets; got {reference_lines}",
    )


def test_federal_output_name_with_standalone_agency(build_federal_resume: object) -> None:
    agency = build_federal_resume.extract_federal_agency_name(FEDERAL_STANDALONE_AGENCY_DESCRIPTION)
    role = build_federal_resume.extract_federal_role_title(FEDERAL_STANDALONE_AGENCY_DESCRIPTION)
    output_name = build_federal_resume.extract_federal_output_name(FEDERAL_STANDALONE_AGENCY_DESCRIPTION)
    assert_true(agency == "Department of Veterans Affairs", f"Federal agency extraction should use standalone department lines; got {agency!r}")
    assert_true(role == "IT Specialist (Customer Support Systems Analysis)", f"Federal role extraction should prefer Position: lines; got {role!r}")
    assert_true(
        output_name == "Department of Veterans Affairs - IT Specialist (Customer Support Systems Analysis)",
        f"Federal output name should combine agency and role; got {output_name!r}",
    )


def test_federal_supporting_doc_resolution(build_federal_resume: object, federal_supporting_docs: object) -> None:
    import os

    validated = federal_supporting_docs.read_validated_federal_job_description(FEDERAL_DUMMY_JOB_DESCRIPTION)
    output_target_name = build_federal_resume.extract_federal_output_name(validated)
    expected_cover_name = f"Christian Estrada - {output_target_name} Federal Cover Letter.docx"

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        output_dir = Path(temp_name)
        older_resume = output_dir / f"Christian Estrada - {output_target_name} Federal Resume.docx"
        newer_resume = output_dir / f"Christian Estrada - {output_target_name} v2 Federal Resume.docx"
        distractor = output_dir / "Christian Estrada - Another Agency - Analyst Federal Resume.docx"
        older_resume.write_text("older federal resume", encoding="utf-8")
        newer_resume.write_text("newer federal resume", encoding="utf-8")
        distractor.write_text("distractor", encoding="utf-8")
        now = datetime.now().timestamp()
        os.utime(older_resume, (now - 120, now - 120))
        os.utime(newer_resume, (now - 30, now - 30))

        original_output_dir = federal_supporting_docs.OUTPUT_DIR
        try:
            federal_supporting_docs.OUTPUT_DIR = output_dir
            matches = federal_supporting_docs.matching_federal_resume_outputs(output_target_name)
            assert_true(
                matches[:2] == [newer_resume, older_resume],
                f"Federal supporting-doc lookup should return newest matching resumes first; got {matches}",
            )
            selected = federal_supporting_docs.find_federal_resume_output(validated)
            assert_true(
                selected == newer_resume,
                f"find_federal_resume_output() should pick the latest matching federal resume; got {selected}",
            )
            cover_path = federal_supporting_docs.supporting_output_path(output_target_name, "Cover Letter")
            assert_true(
                cover_path.name == expected_cover_name,
                f"Federal supporting doc output should use the expected naming pattern; got {cover_path.name!r}",
            )
        finally:
            federal_supporting_docs.OUTPUT_DIR = original_output_dir


def test_federal_plain_text_validation_splits_hours_and_salary_warnings(build_federal_resume: object) -> None:
    from docx import Document

    with TemporaryDirectory(prefix="federal_plain_text_") as temp_name:
        docx_path = Path(temp_name) / "federal_resume.docx"
        document = Document()
        for section in build_federal_resume.FEDERAL_REQUIRED_SECTIONS:
            document.add_paragraph(section)
        document.add_paragraph("christian@example.com")
        document.add_paragraph("Known Agency")
        document.add_paragraph("Implementation Program Manager")
        document.add_paragraph("January 2020 - Present")
        document.add_paragraph("Supervisor: Jane Doe")
        document.add_paragraph(
            "Supported implementation governance, reporting, and adoption work across a multi-site program."
        )
        document.save(str(docx_path))
        report = build_federal_resume.federal_plain_text_validation(docx_path)

    warnings = report["warnings"]
    assert_true(
        any("Hours Per Week" in warning for warning in warnings),
        f"federal_plain_text_validation() should warn separately when Hours Per Week is missing; got {warnings}",
    )
    assert_true(
        any("Salary details" in warning for warning in warnings),
        f"federal_plain_text_validation() should warn separately when salary data is missing; got {warnings}",
    )
    assert_true(
        not any("Hours Per Week" in warning and "Salary" in warning for warning in warnings),
        f"federal_plain_text_validation() should split the Hours Per Week and salary warnings instead of merging them; got {warnings}",
    )


def test_federal_resume_plan_warns_when_page_count_is_unverified(build_federal_resume: object) -> None:
    source = build_federal_resume.load_federal_source()
    original_page_count_for_docx = build_federal_resume.page_count_for_docx
    original_resume_layouts = build_federal_resume.FEDERAL_LAYOUT_PROFILES
    original_qualification_layouts = build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS
    try:
        build_federal_resume.page_count_for_docx = lambda *_args, **_kwargs: None
        build_federal_resume.FEDERAL_LAYOUT_PROFILES = build_federal_resume.FEDERAL_LAYOUT_PROFILES[:1]
        build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS = build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS[:1]
        with TemporaryDirectory(prefix="federal_resume_plan_") as temp_name:
            plan = build_federal_resume.resume_plan(Path(temp_name), source, FEDERAL_DUMMY_JOB_DESCRIPTION)
    finally:
        build_federal_resume.page_count_for_docx = original_page_count_for_docx
        build_federal_resume.FEDERAL_LAYOUT_PROFILES = original_resume_layouts
        build_federal_resume.FEDERAL_QUALIFICATION_LAYOUTS = original_qualification_layouts

    warning_text = " ".join(plan.audit.warnings)
    assert_true(
        plan.resume_page_count is None and plan.qualifications_page_count is None,
        f"resume_plan() should preserve None page counts when the renderer is unavailable; got resume={plan.resume_page_count}, qualifications={plan.qualifications_page_count}",
    )
    assert_true(
        "resume candidate with an unverified page count" in warning_text,
        f"resume_plan() should warn clearly when it selects an unverified federal resume layout; got {plan.audit.warnings}",
    )
    assert_true(
        "qualifications layout with an unverified page count" in warning_text,
        f"resume_plan() should warn clearly when it selects an unverified federal qualifications layout; got {plan.audit.warnings}",
    )


def test_commercial_builder_entrypoints_delegate_to_input_helpers(
    build_cover_letter: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        temp_root = Path(temp_name)
        job_path = temp_root / "job_description.txt"
        job_path.write_text(DUMMY_JOB_DESCRIPTION.strip(), encoding="utf-8")
        output_dir = temp_root / "output"
        output_dir.mkdir()
        resume_docx = temp_root / "Commercial Resume.docx"
        resume_docx.write_text("resume placeholder", encoding="utf-8")

        original_cover_job = build_cover_letter.JOB_DESCRIPTION
        original_cover_output = build_cover_letter.OUTPUT_DIR
        original_cover_find_resume = build_cover_letter.find_resume_output
        original_cover_extract_role = build_cover_letter.extract_role_title
        original_cover_helper = build_cover_letter.build_cover_letter_for_inputs
        original_interview_job = build_interview_cheat_sheet.JOB_DESCRIPTION
        original_interview_output = build_interview_cheat_sheet.OUTPUT_DIR
        original_interview_helper = build_interview_cheat_sheet.build_interview_cheat_sheet_for_inputs
        original_guide_job = build_detailed_interview_guide.JOB_DESCRIPTION
        original_guide_output = build_detailed_interview_guide.OUTPUT_DIR
        original_guide_helper = build_detailed_interview_guide.build_detailed_interview_guide_for_inputs

        cover_calls: dict[str, object] = {}
        interview_calls: dict[str, object] = {}
        guide_calls: dict[str, object] = {}

        try:
            build_cover_letter.JOB_DESCRIPTION = job_path
            build_cover_letter.OUTPUT_DIR = output_dir
            build_cover_letter.find_resume_output = lambda _job_description: resume_docx
            build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_cover_letter.build_cover_letter_for_inputs = lambda **kwargs: cover_calls.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
                bullets_used=3,
                audit_status="PASS",
                specificity_warnings=[],
                mode=kwargs["mode"],
            )

            build_interview_cheat_sheet.JOB_DESCRIPTION = job_path
            build_interview_cheat_sheet.OUTPUT_DIR = output_dir
            build_interview_cheat_sheet.build_cover_letter.find_resume_output = lambda _job_description: resume_docx
            build_interview_cheat_sheet.build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_interview_cheat_sheet.build_interview_cheat_sheet_for_inputs = lambda **kwargs: interview_calls.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
            )

            build_detailed_interview_guide.JOB_DESCRIPTION = job_path
            build_detailed_interview_guide.OUTPUT_DIR = output_dir
            build_detailed_interview_guide.build_cover_letter.find_resume_output = lambda _job_description: resume_docx
            build_detailed_interview_guide.build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_detailed_interview_guide.build_detailed_interview_guide_for_inputs = lambda **kwargs: guide_calls.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
            )

            build_cover_letter.build_cover_letter()
            build_interview_cheat_sheet.build_interview_cheat_sheet()
            build_detailed_interview_guide.build_detailed_interview_guide()
        finally:
            build_cover_letter.JOB_DESCRIPTION = original_cover_job
            build_cover_letter.OUTPUT_DIR = original_cover_output
            build_cover_letter.find_resume_output = original_cover_find_resume
            build_cover_letter.extract_role_title = original_cover_extract_role
            build_cover_letter.build_cover_letter_for_inputs = original_cover_helper
            build_interview_cheat_sheet.JOB_DESCRIPTION = original_interview_job
            build_interview_cheat_sheet.OUTPUT_DIR = original_interview_output
            build_interview_cheat_sheet.build_interview_cheat_sheet_for_inputs = original_interview_helper
            build_detailed_interview_guide.JOB_DESCRIPTION = original_guide_job
            build_detailed_interview_guide.OUTPUT_DIR = original_guide_output
            build_detailed_interview_guide.build_detailed_interview_guide_for_inputs = original_guide_helper

        expected_output_target = "Smoke Test Systems - Implementation Consultant"
        assert_true(
            cover_calls["job_description"] == DUMMY_JOB_DESCRIPTION.strip(),
            "build_cover_letter() should pass the commercial job description text into the shared helper",
        )
        assert_true(
            cover_calls["output_docx"].name == f"Christian Estrada - {expected_output_target} Cover Letter.docx",
            f"build_cover_letter() should preserve the commercial cover-letter naming pattern; got {cover_calls['output_docx']}",
        )
        assert_true(
            interview_calls["output_docx"].name == f"Christian Estrada - {expected_output_target} Interview Cheat Sheet.docx",
            f"build_interview_cheat_sheet() should preserve the commercial cheat-sheet naming pattern; got {interview_calls['output_docx']}",
        )
        assert_true(
            guide_calls["output_docx"].name == f"Christian Estrada - {expected_output_target} Detailed Interview Guide.docx",
            f"build_detailed_interview_guide() should preserve the commercial detailed-guide naming pattern; got {guide_calls['output_docx']}",
        )
        assert_true(
            guide_calls.get("stage") == "all",
            f"build_detailed_interview_guide() should default to the all-stages guide when no stage flag is supplied; got {guide_calls}",
        )


def test_standard_qualifications_default_question_when_file_empty(build_standard_qualifications_statement: object, question_prep: object) -> None:
    # An empty application_questions.txt should no longer block the build.
    # load_active_application_prompts() must return the default question instead.
    import tempfile as _tempfile
    with _tempfile.TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        temp_root = Path(temp_name)
        empty_path = temp_root / "application_questions.txt"
        empty_path.write_text("", encoding="utf-8")
        prompts = question_prep.load_active_application_prompts(empty_path)
        assert_true(
            len(prompts) > 0,
            f"load_active_application_prompts() should return the default question when file is empty; got {prompts!r}",
        )
        assert_true(
            any("interested" in p.lower() for p in prompts),
            f"Default question should be a company/role-interest question; got {prompts!r}",
        )
        missing_path = temp_root / "nonexistent_questions.txt"
        prompts_missing = question_prep.load_active_application_prompts(missing_path)
        assert_true(
            len(prompts_missing) > 0,
            f"load_active_application_prompts() should return the default question when file is absent; got {prompts_missing!r}",
        )


def test_dry_run_reports_default_question_usage(question_prep: object) -> None:
    import tempfile as _tempfile
    with _tempfile.TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        empty_path = Path(temp_name) / "application_questions.txt"
        empty_path.write_text("", encoding="utf-8")
        state = question_prep.load_application_prompt_state(empty_path)
        assert_true(
            state.uses_default_questions,
            f"Empty file should produce uses_default_questions=True; got {state!r}",
        )
        assert_true(
            not state.explicit_prompts,
            f"Empty file should have empty explicit_prompts; got {state.explicit_prompts!r}",
        )
        assert_true(
            bool(state.effective_prompts),
            f"Empty file should still have non-empty effective_prompts; got {state.effective_prompts!r}",
        )


def test_default_questions_skip_stale_pairing_check(question_prep: object) -> None:
    import tempfile as _tempfile
    with _tempfile.TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        empty_path = Path(temp_name) / "application_questions.txt"
        empty_path.write_text("", encoding="utf-8")
        state = question_prep.load_application_prompt_state(empty_path)
        issues = question_prep.application_question_context_issues(
            DUMMY_JOB_DESCRIPTION,
            state,
            workflow="commercial",
        )
        assert_true(
            not issues,
            f"Default-only build should produce no context issues; got {issues!r}",
        )


def test_explicit_stale_questions_still_flagged(question_prep: object) -> None:
    import tempfile as _tempfile
    with _tempfile.TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        stale_path = Path(temp_name) / "application_questions.txt"
        stale_path.write_text("Have you worked with public agencies or cooperatives?", encoding="utf-8")
        state = question_prep.load_application_prompt_state(stale_path)
        assert_true(
            not state.uses_default_questions,
            f"Explicit non-default questions should produce uses_default_questions=False; got {state!r}",
        )
        issues = question_prep.application_question_context_issues(
            DUMMY_JOB_DESCRIPTION,
            state,
            workflow="commercial",
        )
        assert_true(
            issues,
            f"Stale explicit question should produce context issues; got {issues!r}",
        )


def test_qualifications_builder_uses_question_prep_response_engine() -> None:
    source = Path("scripts/build_standard_qualifications_statement.py").read_text(encoding="utf-8")
    assert_true(
        "responses = question_prep.build_question_responses(" in source
        and "question_prep.selected_resume_snapshot(job_description)" in source,
        "The qualifications builder should call the shared question_prep response engine directly so answer changes propagate without a shadow-copy seam.",
    )


def test_qualifications_builder_removes_local_shadow_answer_helpers() -> None:
    source = Path("scripts/build_standard_qualifications_statement.py").read_text(encoding="utf-8")
    assert_true(
        all(
            marker not in source
            for marker in (
                "def software_inventory_answer(",
                "def communication_answer(",
                "def answer_prompt(",
                "def default_responses(",
            )
        ),
        "The qualifications builder should not keep a second local answer engine once the shared question_prep path is active.",
    )


def test_interview_outputs_inherit_bridge_resume_name(
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        temp_root = Path(temp_name)
        job_path = temp_root / "job_description.txt"
        job_path.write_text(DUMMY_JOB_DESCRIPTION.strip(), encoding="utf-8")
        output_dir = temp_root / "output"
        output_dir.mkdir()
        bridge_resume = temp_root / "Christian Estrada - Smoke Test Systems - Implementation Consultant BRIDGE Resume.docx"
        bridge_resume.write_text("resume placeholder", encoding="utf-8")

        original_interview_job = build_interview_cheat_sheet.JOB_DESCRIPTION
        original_interview_output = build_interview_cheat_sheet.OUTPUT_DIR
        original_interview_helper = build_interview_cheat_sheet.build_interview_cheat_sheet_for_inputs
        original_guide_job = build_detailed_interview_guide.JOB_DESCRIPTION
        original_guide_output = build_detailed_interview_guide.OUTPUT_DIR
        original_guide_helper = build_detailed_interview_guide.build_detailed_interview_guide_for_inputs
        original_interview_find_resume = build_interview_cheat_sheet.build_cover_letter.find_resume_output
        original_interview_extract_role = build_interview_cheat_sheet.build_cover_letter.extract_role_title
        original_guide_find_resume = build_detailed_interview_guide.build_cover_letter.find_resume_output
        original_guide_extract_role = build_detailed_interview_guide.build_cover_letter.extract_role_title

        interview_calls: dict[str, object] = {}
        guide_calls: dict[str, object] = {}

        try:
            build_interview_cheat_sheet.JOB_DESCRIPTION = job_path
            build_interview_cheat_sheet.OUTPUT_DIR = output_dir
            build_interview_cheat_sheet.build_cover_letter.find_resume_output = lambda _job_description: bridge_resume
            build_interview_cheat_sheet.build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_interview_cheat_sheet.build_interview_cheat_sheet_for_inputs = lambda **kwargs: interview_calls.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
            )

            build_detailed_interview_guide.JOB_DESCRIPTION = job_path
            build_detailed_interview_guide.OUTPUT_DIR = output_dir
            build_detailed_interview_guide.build_cover_letter.find_resume_output = lambda _job_description: bridge_resume
            build_detailed_interview_guide.build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_detailed_interview_guide.build_detailed_interview_guide_for_inputs = lambda **kwargs: guide_calls.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
            )

            build_interview_cheat_sheet.build_interview_cheat_sheet()
            build_detailed_interview_guide.build_detailed_interview_guide()
        finally:
            build_interview_cheat_sheet.JOB_DESCRIPTION = original_interview_job
            build_interview_cheat_sheet.OUTPUT_DIR = original_interview_output
            build_interview_cheat_sheet.build_interview_cheat_sheet_for_inputs = original_interview_helper
            build_detailed_interview_guide.JOB_DESCRIPTION = original_guide_job
            build_detailed_interview_guide.OUTPUT_DIR = original_guide_output
            build_detailed_interview_guide.build_detailed_interview_guide_for_inputs = original_guide_helper
            build_interview_cheat_sheet.build_cover_letter.find_resume_output = original_interview_find_resume
            build_interview_cheat_sheet.build_cover_letter.extract_role_title = original_interview_extract_role
            build_detailed_interview_guide.build_cover_letter.find_resume_output = original_guide_find_resume
            build_detailed_interview_guide.build_cover_letter.extract_role_title = original_guide_extract_role

        assert_true(
            " BRIDGE Interview Cheat Sheet.docx" in interview_calls["output_docx"].name,
            f"build_interview_cheat_sheet() should propagate BRIDGE naming from the matched resume; got {interview_calls['output_docx']}",
        )
        assert_true(
            " BRIDGE Detailed Interview Guide.docx" in guide_calls["output_docx"].name,
            f"build_detailed_interview_guide() should propagate BRIDGE naming from the matched resume; got {guide_calls['output_docx']}",
        )


def test_interview_stage_resolution_and_context_parsing(
    interview_stage: object,
) -> None:
    labeled = interview_stage.parse_interviewer_context(
        "\n".join(
            [
                "Name: Jordan Lee",
                "Title: Director of Technical Delivery",
                "Stage: panel",
                "Recruiter Feedback: Keep answers tighter; lead with outcomes",
                "Emphasize: automation, ERP, stakeholder alignment",
            ]
        )
    )
    assert_true(
        labeled.name == "Jordan Lee"
        and labeled.title == "Director of Technical Delivery"
        and labeled.stage_hint == "panel"
        and labeled.recruiter_feedback == ("Keep answers tighter", "lead with outcomes")
        and labeled.emphasized_terms == ("automation", "ERP", "stakeholder alignment"),
        f"parse_interviewer_context() should parse labeled interviewer notes; got {labeled}",
    )
    bare = interview_stage.parse_interviewer_context("Sam Rivera, Hiring Manager")
    assert_true(
        bare.name == "Sam Rivera" and bare.title == "Hiring Manager",
        f"parse_interviewer_context() should parse a bare 'Name, Title' line; got {bare}",
    )
    empty = interview_stage.parse_interviewer_context("")
    assert_true(
        not empty.name and not empty.title and not empty.stage_hint,
        f"parse_interviewer_context() should return an empty context object for empty input; got {empty}",
    )
    resolved = interview_stage.resolve_stage("technical", labeled, None)
    assert_true(
        resolved.key == "technical",
        f"resolve_stage() should prefer the explicit CLI stage over interviewer-context stage hints; got {resolved}",
    )
    resolved_from_context = interview_stage.resolve_stage("", labeled, None)
    assert_true(
        resolved_from_context.key == "panel",
        f"resolve_stage() should use the interviewer-context Stage: hint when no CLI stage is supplied; got {resolved_from_context}",
    )
    try:
        interview_stage.resolve_stage("mystery_round", empty, None)
    except ValueError:
        pass
    else:
        raise AssertionError("resolve_stage() should fail on unknown stage keys")


def test_stage_filename_suffix_composes_with_existing_detailed_guide_names(
    interview_stage: object,
    build_detailed_interview_guide: object,
    question_prep: object,
) -> None:
    suffixes = {
        key: interview_stage.stage_filename_suffix(profile)
        for key, profile in interview_stage.STAGE_PROFILES.items()
    }
    non_empty = [value for key, value in suffixes.items() if key != "all"]
    assert_true(
        suffixes["all"] == "" and len(non_empty) == len(set(non_empty)),
        f"stage_filename_suffix() should leave all unsuffixed and keep every other stage suffix unique; got {suffixes}",
    )
    base = Path("C:/tmp/Christian Estrada - Acme FAIL Detailed Interview Guide.docx")
    staged = build_detailed_interview_guide.output_path_for_stage(base, interview_stage.STAGE_PROFILES["hr_screen"])
    assert_true(
        staged.name == "Christian Estrada - Acme FAIL Detailed Interview Guide (HR Screen).docx",
        f"output_path_for_stage() should insert the stage suffix after the Detailed Interview Guide stem; got {staged.name}",
    )
    drafted = question_prep.application_question_draft_path(staged)
    assert_true(
        "(HR Screen)" in drafted.name,
        f"application_question_draft_path() should preserve the stage suffix on staged guide names; got {drafted.name}",
    )


def test_federal_detailed_guide_wrapper_keeps_stage_params_optional(
    build_federal_detailed_interview_guide: object,
    build_detailed_interview_guide: object,
    federal_supporting_docs: object,
) -> None:
    captured: dict[str, object] = {}
    original_context = federal_supporting_docs.resolve_federal_context
    original_output_path = federal_supporting_docs.supporting_output_path
    original_helper = build_detailed_interview_guide.build_detailed_interview_guide_for_inputs
    try:
        federal_supporting_docs.resolve_federal_context = lambda: SimpleNamespace(
            job_description="Agency: Example\nJob Title: Program Analyst\n",
            resume_docx=Path("federal_resume.docx"),
            output_target_name="Example Agency - Program Analyst",
            company_name="Example Agency",
            role_title="Program Analyst",
        )
        federal_supporting_docs.supporting_output_path = lambda target_name, label: Path(f"{target_name} {label}.docx")
        build_detailed_interview_guide.build_detailed_interview_guide_for_inputs = (
            lambda **kwargs: captured.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
            )
        )
        build_federal_detailed_interview_guide.main()
    finally:
        federal_supporting_docs.resolve_federal_context = original_context
        federal_supporting_docs.supporting_output_path = original_output_path
        build_detailed_interview_guide.build_detailed_interview_guide_for_inputs = original_helper
    assert_true(
        "stage" not in captured and "interviewer_context" not in captured,
        f"The federal detailed-guide wrapper should still call the helper without the new stage params because they remain optional; got {captured}",
    )


def test_summary_three_sentence_structure(build_resume: object) -> None:
    job_description = """
    Company: Acme HealthTech
    Role: Implementation Consultant
  Healthcare implementation role supporting clinical teams, patient workflows, claims reporting, adoption, and go-live readiness.
    """
    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(len(build_resume.summary_sentences(summary)) == 3, f"expected 3 sentences: {summary}")
    assert_true("Target context:" not in summary, f"summary must not use Target context label: {summary}")
    assert_true("Strong fit for roles" not in summary, f"summary must not use Strong fit closer: {summary}")
    assert_true("Experience spans" not in summary, f"summary must not use Experience spans: {summary}")


def test_supply_chain_summary_stays_in_lane_context(build_resume: object) -> None:
    job_description = """
    Company: The Clorox Company
    Role: Supply Chain Network Optimization Analyst
    The role supports network optimization, cost modeling, supply chain analytics, ETL workflows, large operational datasets,
    executive reporting, Power BI, and cross-functional decision support across a multi-site environment.
    """
    summary = build_resume.build_problem_first_summary(job_description)
    lowered = summary.lower()
    assert_true(
        "cloud integration and data management buyers" not in lowered,
        f"supply-chain analytics summary should not drift into integration-buyer context: {summary}",
    )
    assert_true(
        "supply chain" in lowered or "manufacturing" in lowered or "operating environments" in lowered,
        f"supply-chain analytics summary should keep an operations context visible: {summary}",
    )
    assert_true(
        word_count(summary) >= build_resume.PROFESSIONAL_SUMMARY_MIN_WORDS,
        (
            f"supply-chain analytics summary should stay above the "
            f"{build_resume.PROFESSIONAL_SUMMARY_MIN_WORDS}-word minimum: {summary}"
        ),
    )


def test_retention_analytics_summary_meets_minimum_word_count(build_resume: object) -> None:
    job_description = """
    Company: Ollie
    Role: Data Analyst, Retention
    This role focuses on retention analytics, curated reporting, data-driven decision support,
    customer behavior trends, workflow quality, and measurable process improvement.
    """
    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(len(build_resume.summary_sentences(summary)) == 3, f"retention analytics summary should stay at 3 sentences: {summary}")
    assert_true(
        word_count(summary) >= build_resume.PROFESSIONAL_SUMMARY_MIN_WORDS,
        (
            f"retention analytics summary should stay above the "
            f"{build_resume.PROFESSIONAL_SUMMARY_MIN_WORDS}-word minimum: {summary}"
        ),
    )
    opening = build_resume.summary_sentences(summary)[0]
    comma_items = [item.strip() for item in opening.split(",") if item.strip()]
    assert_true(
        len(comma_items) <= 4,
        f"retention analytics summary opening should stay at four or fewer comma-separated segments: {opening}",
    )
    assert_true(
        "kpi reporting" in summary.lower() or "data analysis" in summary.lower(),
        f"retention analytics summary should surface supported analytics checklist language when needed: {summary}",
    )
    assert_true(
        "data analyst" in summary.lower() and "retention" in summary.lower() and "curated reporting" in summary.lower(),
        f"retention analytics summary should weave top role-language signals into the summary when supported: {summary}",
    )


def test_title_phrase_candidates_do_not_cross_comma_title_segments(build_resume: object) -> None:
    job_description = """
    Company: Ollie
    Job Title: Data Analyst, Retention
    """
    phrases = build_resume.title_phrase_candidates(job_description)
    assert_true(
        "data analyst" in phrases,
        f"title_phrase_candidates() should keep the primary title segment; got {phrases}",
    )
    assert_true(
        "data analyst retention" not in phrases and "analyst retention" not in phrases,
        f"title_phrase_candidates() should not create cross-segment comma title phrases; got {phrases}",
    )


def test_clorox_style_job_title_and_specialties(build_resume: object) -> None:
    job_description = """
    Company: The Clorox Company

    Job Title:

    Your role at Clorox: Supply Chain Network Optimization Analyst
    This role supports supply chain analytics, cost modeling, KPI interpretation, Data Guru workflows, and Power BI reporting.
    #LI-Hybrid
    """
    extracted_title = build_resume.extract_job_title(job_description)
    assert_true(
        extracted_title == "Supply Chain Network Optimization Analyst",
        f"extract_job_title() should prefer the named role over job-board tags; got {extracted_title!r}",
    )
    specialties = build_resume.visible_role_specialties(job_description)
    assert_true(
        "ERP implementation" not in specialties,
        f"visible_role_specialties() should not infer ERP from unrelated words like interpretation; got {specialties}",
    )
    assert_true(
        "Power BI" in specialties,
        f"visible_role_specialties() should keep real analytical specialty signals; got {specialties}",
    )


def test_multiline_job_title_extraction_with_bom(build_resume: object) -> None:
    job_description = "\ufeffCompany: Clorox\n\nJob Title: Supply Chain Network Optimization Analyst\n\nYour role at Clorox:\n\nThis role supports the Network Optimization team by executing cost modeling and network optimization analyses."
    assert_true(
        build_resume.extract_company_name(job_description) == "Clorox",
        f"extract_company_name() should ignore a UTF-8 BOM on the Company line; got {build_resume.extract_company_name(job_description)!r}",
    )
    assert_true(
        build_resume.extract_job_title(job_description) == "Supply Chain Network Optimization Analyst",
        f"extract_job_title() should prefer the labeled title over body copy; got {build_resume.extract_job_title(job_description)!r}",
    )
    assert_true(
        build_resume.extract_output_target_name(job_description) == "Clorox - Supply Chain Network Optimization Analyst",
        f"extract_output_target_name() should keep clean company/title naming for multiline Clorox-style posts; got {build_resume.extract_output_target_name(job_description)!r}",
    )


def test_supply_chain_analytics_summary_promotes_supported_delivery_terms(build_resume: object) -> None:
    job_description = """
    Company: Clorox
    Job Title: Supply Chain Network Optimization Analyst
    This role supports network optimization, cost modeling, data management, project delivery, model quality,
    Power BI reporting, ETL workflows, and executive-level reporting across multi-site supply chain operations.
    """
    summary = build_resume.build_problem_first_summary(job_description)
    lowered = summary.lower()
    assert_true("data management" in lowered, f"supply-chain analytics summary should surface data management when supported: {summary}")
    assert_true("project delivery" in lowered, f"supply-chain analytics summary should surface project delivery when supported: {summary}")
    assert_true("model quality" in lowered, f"supply-chain analytics summary should surface model quality when supported: {summary}")


def test_startup_operator_summary_structure(build_resume: object) -> None:
    job_description = """
    Company: Growth Health
    Role: Technical Implementation Manager
    High-growth healthcare AI startup seeking a technical implementation manager to own customer onboarding,
    workflow automation, implementation delivery, data migration, and adoption readiness.
    """
    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(len(build_resume.summary_sentences(summary)) == 3, f"startup operator summary should stay at 3 sentences: {summary}")
    assert_true(
        not any(sentence.startswith("That") for sentence in build_resume.summary_sentences(summary)),
        f"startup operator summary should avoid sentence openers with 'That': {summary}",
    )
    assert_true(
        not re.search(r"\b[Tt]hat\b", summary),
        f"startup operator summary should avoid the word 'that' when cleaner phrasing is available: {summary}",
    )


def test_summary_detectors_avoid_member_and_operator_false_positives(build_resume: object) -> None:
    assert_true(
        not build_resume.is_startup_or_broad_operator_role(VENSURE_STYLE_JOB_DESCRIPTION),
        "is_startup_or_broad_operator_role() should not trigger on enterprise implementation JDs with only one broad-operator signal",
    )
    summary = build_resume.build_problem_first_summary(VENSURE_STYLE_JOB_DESCRIPTION)
    lowered = summary.lower()
    assert_true(
        "member-facing" not in lowered and "subscription" not in lowered,
        f"build_problem_first_summary() should not infer subscription/member context from phrases like team members: {summary}",
    )


def test_implementation_summary_surfaces_quality_and_process_language(build_resume: object) -> None:
    job_description = """
    Company: Quality Systems
    Role: Senior Implementation Consultant
    This role leads implementation, workflow best practices, process improvement initiatives,
    quality control, setup accuracy, and go-live readiness across enterprise software clients.
    """
    summary = build_resume.build_problem_first_summary(job_description)
    lowered = summary.lower()
    assert_true(
        "quality" in lowered,
        f"implementation summaries should surface quality language when the JD stresses quality control and accuracy: {summary}",
    )
    assert_true(
        "process" in lowered,
        f"implementation summaries should surface process language when the JD stresses process improvement and workflow best practices: {summary}",
    )


def test_natural_top_bullet_meta_penalty(build_resume: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    keywords = {"adoption", "launch", "dashboard", "sql"}
    clean_bullet = (
        "Built 200+ dashboards and reporting tools that improved launch readiness and adoption visibility across complex implementations."
    )
    meta_bullet = (
        "Built 200+ dashboards and reporting tools that improved launch readiness and adoption visibility across complex implementations, "
        "a pattern that supports steadier launches."
    )
    clean_score = build_resume.natural_top_bullet_score(clean_bullet, keywords, profile)
    meta_score = build_resume.natural_top_bullet_score(meta_bullet, keywords, profile)
    assert_true(
        clean_score > meta_score,
        f"natural_top_bullet_score() should penalize forward-looking meta-commentary; clean={clean_score}, meta={meta_score}",
    )


def test_boomi_presales_summary(build_resume: object) -> None:
    job_description = """
    Company: Boomi
    Role: Senior Presales Solutions Engineer
    Boomi provides an intelligent integration and automation platform for B2B customers.
    The role requires pre-sales, technical discovery, proof of concept, data management,
    master data management, data quality, data governance, iPaaS, API-led connectivity,
    executive presentations, and product demonstrations for North American clients.
    """
    profile = build_resume.job_problem_profile(job_description)
    assert_true(
        profile.primary_lane == "presales_solution",
        f"expected presales_solution lane, got {profile.primary_lane!r}",
    )
    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(len(build_resume.summary_sentences(summary)) == 3, f"expected 3 sentences: {summary}")
    lowered = summary.lower()
    assert_true(
        not re.search(r"\berp\b", lowered),
        f"pre-sales summary should not mention ERP without explicit ERP requirement: {summary}",
    )
    assert_true(
        "discovery" in lowered or "demonstration" in lowered,
        f"pre-sales summary should name discovery or demonstrations: {summary}",
    )


def test_customer_success_summary_capitalization(build_resume: object) -> None:
    job_description = LANE_JOB_DESCRIPTIONS["customer_success"]
    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(
        "Revenue-focused customer success and enterprise software consultant" in summary,
        f"customer-success summary should keep the role label lowercase inside the sentence; got {summary}",
    )


def test_customer_success_summary_clarifies_portfolio_scope(build_resume: object) -> None:
    job_description = LANE_JOB_DESCRIPTIONS["customer_success"]
    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(
        "within a $6M+ book of business" in summary,
        f"customer-success summary should surface the broader book of business; got {summary}",
    )
    assert_true(
        "$1M+ in at-risk annual revenue" in summary,
        f"customer-success summary should keep the at-risk revenue detail explicit; got {summary}",
    )
    assert_true(
        "high-risk account recovery tied to $1M+ in annual revenue" not in summary,
        f"customer-success summary should not read like $1M+ was the whole portfolio; got {summary}",
    )


def test_general_consulting_lane_isolated_from_other_consultant_roles(build_resume: object) -> None:
    consulting_job = """
    Company: Bain
    Role: Consultant
    Management consulting role working on case teams to help executives solve difficult client problems,
    lead analyses, build relationships, and turn recommendations into measurable outcomes.
    """
    consulting_profile = build_resume.job_problem_profile(consulting_job)
    assert_true(
        consulting_profile.primary_lane == "corporate_strategy",
        f"general consulting posting should map to corporate_strategy; got {consulting_profile.primary_lane!r}",
    )
    consulting_summary = build_resume.build_problem_first_summary(consulting_job)
    assert_true(
        "consulting" in consulting_summary.lower(),
        f"general consulting summary should make consulting visible: {consulting_summary}",
    )
    assert_true(
        not re.search(r"\b[Tt]hat\b", consulting_summary),
        f"general consulting summary should avoid the word 'that' when cleaner phrasing is available: {consulting_summary}",
    )

    implementation_job = """
    Company: Acme Software
    Role: Solutions Implementation Consultant
    Customer-facing implementation role supporting software configuration, data migration, integration,
    user acceptance testing, go-live readiness, and stakeholder communication for enterprise SaaS clients.
    """
    implementation_profile = build_resume.job_problem_profile(implementation_job)
    assert_true(
        implementation_profile.primary_lane == "implementation_delivery",
        f"implementation consultant should stay on implementation_delivery; got {implementation_profile.primary_lane!r}",
    )
    implementation_summary = build_resume.build_problem_first_summary(implementation_job)
    assert_true(
        "consulting-style" not in implementation_summary.lower(),
        f"implementation consultant summary should not inherit the general consulting summary: {implementation_summary}",
    )


def test_strategy_transformation_consultant_prefers_corporate_strategy(build_resume: object) -> None:
    strategy_job = """
    Company: Guidehouse
    Job Title: Strategy & Transformation Senior Consultant
    Guidehouse is a global consulting firm serving public-sector and commercial clients.
    This senior consultant role helps clients solve complex business challenges from strategy through execution.
    Responsibilities include requirements gathering and strategy sessions with senior-level clients, root cause and gap analysis,
    findings and recommendation reports, process flows, reporting dashboards, stakeholder communication, and change management support.
    """
    strategy_profile = build_resume.job_problem_profile(strategy_job)
    assert_true(
        strategy_profile.primary_lane == "corporate_strategy",
        f"strategy/transformation consulting titles should map to corporate_strategy even when the JD also includes implementation-style execution language; got {strategy_profile.primary_lane!r}",
    )
    strategy_summary = build_resume.build_problem_first_summary(strategy_job)
    assert_true(
        any(term in strategy_summary.lower() for term in ("strategy", "consult", "recommendation", "client")),
        f"strategy/transformation consulting summaries should stay in strategy/consulting language; got {strategy_summary}",
    )


def test_consulting_audit_keywords_filter_recruiting_fluff(build_resume: object) -> None:
    job_description = """
    Company: Bain
    Job Title: Consultant
    An opportunity to build your own career and learn a unique problem-solving toolkit.
    Management consulting role working on case teams to help executives solve difficult client problems,
    lead analyses, and turn recommendations into measurable outcomes.
    """
    keywords = build_resume.audit_keywords(job_description)
    assert_true("consulting" in keywords, f"consulting should remain a visible audit keyword: {keywords}")
    assert_true("bain" not in keywords, f"company names should not become audit keywords: {keywords}")
    assert_true("opportunity" not in keywords, f"recruiting fluff should not become audit keywords: {keywords}")


def test_audit_keywords_filter_company_affiliates_and_weak_fragments(build_resume: object) -> None:
    job_description = """
    Company: Anord Mardix
    Job Title: Field Service Technical Support
    Anord Mardix, a Flex company, is a global leader in critical power solutions supporting data centers.
    If this sounds interesting, we'd like to meet you.
    Reviews customer requirements and provides recommendations that will meet both the capabilities Flex and the specifications of the customer.
    Coordinates product integration efforts between component suppliers, Flex, and the customer.
    Provides complex technical assistance related to manufacturing capabilities including design, fabrication, and assembly abilities.
    Typically requires a bachelor's degree in related field or equivalent experience.
    Demonstrates expert knowledge of the function and a thorough understanding of Flex and related business.
    Executes functional strategic plans and objectives at the site.
    """
    keywords = build_resume.audit_keywords(job_description)
    assert_true("flex" not in keywords, f"affiliate company names should not become audit keywords: {keywords}")
    assert_true("meet" not in keywords, f"generic recruiting verbs should not become audit keywords: {keywords}")
    assert_true("related" not in keywords, f"generic JD fragments should not become audit keywords: {keywords}")
    assert_true("functional" not in keywords and "defining" not in keywords, f"generic JD scaffolding should not become audit keywords: {keywords}")
    assert_true(
        "field service" in keywords and "technical support" in keywords and "product integration" in keywords,
        f"title-aligned and domain phrases should survive audit cleanup: {keywords}",
    )
    assert_true(
        "coordinates product integration effort" not in keywords and "executes functional strategic plan" not in keywords,
        f"action-led JD fragments should not become audit keywords: {keywords}",
    )
    assert_true(
        any("data center" in keyword for keyword in keywords) and "integration" in keywords and "technical" in keywords and "customer" in keywords,
        f"role-relevant domain terms should survive audit cleanup: {keywords}",
    )


def test_audit_keywords_filter_boilerplate_adjectives_and_normalize_es_plurals(build_resume: object) -> None:
    keywords = build_resume.audit_keywords(VENSURE_STYLE_JOB_DESCRIPTION)
    assert_true("process" in keywords, f"audit_keywords() should preserve the canonical singular form of processes: {keywords}")
    assert_true("processe" not in keywords, f"audit_keywords() should not create broken singular forms like 'processe': {keywords}")
    assert_true("strong" not in keywords and "high" not in keywords, f"boilerplate adjectives should not become audit keywords: {keywords}")
    assert_true("guidance" not in keywords and "serve" not in keywords, f"generic serve/guidance phrasing should not become audit keywords: {keywords}")


def test_keyword_placement_prefers_role_phrases(build_resume: object) -> None:
    job_description = """
    Company: Anord Mardix
    Job Title: Field Service Technical Support
    Coordinates product integration efforts between component suppliers and the customer.
    Supports data centers and manufacturing environments with technical issue ownership.
    """
    resume_text = "\n".join(
        [
            "Professional Summary",
            "Implementation consultant supporting customer-facing delivery, reporting clarity, and stakeholder coordination.",
            "Professional Experience",
            "Some Role    March 2023 - Present",
            "Known Company | Knoxville, TN",
            "Improved adoption and issue visibility across complex implementations.",
            "Education",
            "Bachelor of Arts",
            "Skills",
            "Implementation and Delivery:  SQL  |  Reporting  |  Dashboard Design",
            "Professional Development",
            "Certified Scrum Product Owner",
        ]
    )
    report = build_resume.keyword_placement_audit(job_description, resume_text, limit=5)
    gap_keywords = [str(gap.get("keyword", "")) for gap in report.get("gaps", []) if isinstance(gap, dict)]
    assert_true(
        any(keyword in gap_keywords[:3] for keyword in ("field service", "technical support", "product integration")),
        f"keyword_placement_audit() should surface role phrases before generic single words when they are missing; got {gap_keywords}",
    )


def test_ai_evidence_area_config(build_resume: object) -> None:
    job_description = """
    Company: AI Workflow Systems
    Role: AI Integration Consultant
    The role supports AI strategy, AI integration, AI-driven workflow improvement,
    chatbot adoption, natural language processing, and NLP-based customer messaging.
    """
    resume_text = (
        "AI tools, AI-Assisted Documentation, Generative AI, AI-assisted workflow support, "
        "LivePerson, LiveEngage, Conversational AI, and NLP-based Messaging."
    )
    profile = build_resume.job_problem_profile(job_description, resume_text)
    assert_true(
        "AI-Assisted Workflow Improvement" in profile.direct_matches,
        "job_problem_profile() should detect AI-Assisted Workflow Improvement through BRIDGE_EVIDENCE_AREAS",
    )
    assert_true(
        "AI-assisted workflow improvement" in profile.safe_terms,
        "AI safe terms should come from the config-driven bridge evidence entry",
    )


def professional_summary_xml(summary: str) -> str:
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
<w:p><w:r><w:t>{summary}</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
</w:body></w:document>'''


def resume_with_competencies_xml(
    summary: str,
    competency_items: list[str],
    bullets: list[str] | None = None,
    *,
    section_heading: str = "Skills",
) -> str:
    role_bullets = bullets or [
        "Reduced onboarding delays by 22% across five sites for 150+ users.",
        "Built SQL reporting that improved go-live readiness and stakeholder decisions.",
    ]
    bullet_xml = "".join(f"<w:p><w:r><w:t>{bullet}</w:t></w:r></w:p>" for bullet in role_bullets)
    competency_text = "Implementation and Delivery:  " + "  |  ".join(competency_items)
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
<w:p><w:r><w:t>{summary}</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
<w:p><w:r><w:t>ERP Systems Manager    March 2023 - Present</w:t></w:r></w:p>
<w:p><w:r><w:t>Known Company | Knoxville, TN</w:t></w:r></w:p>
{bullet_xml}
<w:p><w:r><w:t>{section_heading}</w:t></w:r></w:p>
<w:p><w:r><w:t>{competency_text}</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Development</w:t></w:r></w:p>
<w:p><w:r><w:t>Certified Scrum Product Owner</w:t></w:r></w:p>
</w:body></w:document>'''


def write_docx_with_lines(path: Path, lines: list[str]) -> None:
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>'
        + "".join(f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>" for line in lines)
        + "</w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("word/document.xml", document_xml)


def simple_document_xml(lines: list[str]) -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>'
        + "".join(f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>" for line in lines)
        + "</w:body></w:document>"
    )


def rebalance_summary_text(build_resume: object, summary: str) -> tuple[int, str]:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        document_xml = Path(temp_name) / "document.xml"
        document_xml.write_text(professional_summary_xml(summary), encoding="utf-8")
        changed = build_resume.rebalance_professional_summary_erp_mentions(document_xml)
        return changed, build_resume.paragraph_infos(document_xml)[1].text


def condense_summary_text(build_resume: object, summary: str) -> tuple[int, str]:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        document_xml = Path(temp_name) / "document.xml"
        document_xml.write_text(professional_summary_xml(summary), encoding="utf-8")
        changed = build_resume.condense_professional_summary(document_xml)
        return changed, build_resume.paragraph_infos(document_xml)[1].text


def test_summary_condense_guard(build_resume: object) -> None:
    summary = (
        "Implementation consultant with more than ten years of experience turning business requirements, data issues, "
        "reporting needs, and workflow gaps into clearer delivery plans for enterprise software teams. Background includes "
        "stakeholder alignment, go-live readiness, customer adoption, and measurable process improvement across complex "
        "systems environments."
    )
    changed, updated = condense_summary_text(build_resume, summary)
    assert_true(changed == 0, "condense_professional_summary() should not change summaries at or below 130 words")
    assert_true("more than ten years" in updated, "condense_professional_summary() should preserve longer phrasing when the summary is already short enough")


def test_erp_summary_rebalance(build_resume: object) -> None:
    three_mentions = "ERP reporting improved ERP adoption while erp governance stayed visible."
    changed, updated = rebalance_summary_text(build_resume, three_mentions)
    assert_true(changed == 1, "ERP rebalance should report one changed summary paragraph for three ERP mentions")
    assert_true(
        len(re.findall(r"\berp\b", updated, flags=re.I)) == 2,
        f"ERP rebalance should leave exactly two ERP mentions; got: {updated}",
    )
    assert_true(
        "enterprise systems" in updated,
        "ERP rebalance should replace the third ERP mention with enterprise systems",
    )

    one_mention = "ERP delivery improved reporting and adoption across users."
    changed, updated = rebalance_summary_text(build_resume, one_mention)
    assert_true(changed == 0, "ERP rebalance should not change a summary with one ERP mention")
    assert_true(one_mention in updated, "ERP rebalance should leave a one-mention summary unchanged")


def test_cover_letter_colon_smoothing(build_cover_letter: object) -> None:
    text = "Interview starts at 9:00 AM. Result: stronger customer trust."
    smoothed = build_cover_letter.smooth_cover_letter_text(text)
    assert_true(
        "9:00 AM" in smoothed,
        f"smooth_cover_letter_text() corrupted a time expression: {smoothed}",
    )
    startup_text = (
        "Tabs is doing something genuinely difficult: moving from early traction to a repeatable operating rhythm "
        "without losing speed. The role is designed to solve that kind of operating problem, which is why my "
        "background fits here: ambiguous requirements and multiple stakeholder groups. If the role needs customer "
        "success with commercial judgment, the strongest bridge in my background is value-led growth: using "
        "discovery, executive reviews, and risk recovery to protect revenue."
    )
    startup_smoothed = build_cover_letter.smooth_cover_letter_text(startup_text)
    assert_true(
        startup_smoothed.count(":") <= 2,
        f"smooth_cover_letter_text() should trim colon-heavy startup phrasing; got {startup_smoothed}",
    )


def test_cover_sentence_score_prioritizes_signal_density_over_length(build_cover_letter: object) -> None:
    company_name = "Acme Corp"
    role_title = "Implementation Specialist"
    job_description = "Acme Corp is hiring an Implementation Specialist to lead enterprise rollouts."
    dense_sentence = (
        "Owned a mission-critical enterprise platform across five sites and 150+ users, supported 80+ "
        "manufacturing clients across the Americas, Europe, and Asia within a $6M+ client book of business, "
        "built 200+ dashboards and KPI reporting tools, facilitated 60+ executive workshops and QBRs, and "
        "helped stabilize $1M+ in at-risk annual revenue."
    )
    thin_sentence = (
        "Resolved enterprise application issues across more than 600 law firm offices, spanning Active "
        "Directory, SQL Server, third-party integrations, and Windows services during a support transition."
    )
    dense_score = build_cover_letter.cover_sentence_score(dense_sentence, company_name, role_title, job_description)
    thin_score = build_cover_letter.cover_sentence_score(thin_sentence, company_name, role_title, job_description)
    assert_true(
        dense_score > thin_score,
        "cover_sentence_score() should rate a sentence carrying many concrete proof signals (scale, revenue, "
        f"named scope) above a thinner one even though the thin sentence is shorter; got dense={dense_score} "
        f"thin={thin_score}",
    )

    dense_priority = build_cover_letter.standard_sentence_trim_priority(
        dense_sentence, company_name, role_title, job_description
    )
    thin_priority = build_cover_letter.standard_sentence_trim_priority(
        thin_sentence, company_name, role_title, job_description
    )
    assert_true(
        thin_priority < dense_priority,
        "standard_sentence_trim_priority() should rank the thinner sentence for removal first when a cover "
        "letter must be trimmed to fit the word budget, not the proof-dense sentence just because it is "
        f"longer; got dense={dense_priority} thin={thin_priority}",
    )


def test_mission_or_context_sentence_survives_job_label_header(build_cover_letter: object) -> None:
    question_prep = build_cover_letter.question_prep
    job_description = (
        "Company: Techolution\n"
        "\n"
        "\n"
        "Job Title: AI Adoption & Transformation Manager\n"
        "\n"
        "\n"
        "About the job\n"
        "As a Google Premium Partner, Techolution helps leading enterprise organizations transform the way "
        "they work with Google Gemini Enterprise. We are looking for an AI Enablement & Adoption Consultant "
        "to work directly with enterprise customers across the US."
    )
    sentences = question_prep.split_into_sentences(job_description[:400])
    assert_true(
        any("helps leading enterprise organizations" in sentence for sentence in sentences),
        "split_into_sentences() should not glue the 'Company:'/'Job Title:' header lines onto the real "
        f"opening sentence and lose it; got {sentences}",
    )

    mission = question_prep.mission_or_context_sentence("Techolution", "", job_description)
    assert_true(
        "helps leading enterprise organizations" in mission,
        "mission_or_context_sentence() should fall back to the JD's own 'About the job' sentence instead of "
        f"returning empty when company_research_text is blank; got {mission!r}",
    )
    foundant_job_description = (
        "Company: Foundant\n\n"
        "Job Title: Associate Implementation Consultant - Strategic Advancement Focus\n\n"
        "Associate Implementation Consultant - Strategic Advancement Focus\n"
        "About Foundant:\n\n"
        "At Foundant, we empower mission-driven organizations to manage their data, workflows, and impact "
        "with our comprehensive software solutions."
    )
    foundant_mission = question_prep.mission_or_context_sentence(
        "Foundant",
        "",
        foundant_job_description,
    )
    assert_true(
        "empower mission-driven organizations" in foundant_mission.lower(),
        "mission_or_context_sentence() should ignore 'Focus' in a role title and select the actual company "
        f"context sentence; got {foundant_mission!r}",
    )


def test_proof_first_opening_avoids_list_density_with_comma_heavy_core_problem(
    build_cover_letter: object,
    writing_eval: object,
) -> None:
    brief = build_cover_letter.question_prep.PositioningBrief(
        company_name="QTS Data Centers",
        role_title="Legal Operations Specialist",
        primary_lane="implementation_delivery",
        employer_type="legal operations",
        mission_or_context="",
        role_core_problem="implementation scope, technical complexity, timeline risk, and adoption readiness",
        role_problem_phrase="implementation scope becoming real operating progress",
        personal_reason_to_care="The work fits the way I like to turn system ambiguity into usable operating progress.",
        personal_reason_source="notes",
        strongest_direct_proofs=[
            "implementation delivery",
            "stakeholder communication",
            "operational improvement",
        ],
        strongest_bridge_theme="stakeholder communication",
        top_proof_anchors=[],
        company_specific_fact="",
        gap_honesty_boundary="",
        selected_proof_sentences=[],
    )
    opening = build_cover_letter.proof_first_opening_paragraph(
        brief,
        "QTS Data Centers",
        "Legal Operations Specialist",
    )
    report = build_cover_letter.prose_quality_report(opening, "cover_letter_full")
    result = writing_eval.evaluate_text("cover_letter_full", opening, sample_id="proof_first_opening_dense_list_guard")
    issue_codes = {issue.code for issue in result.issues}

    assert_true(
        "implementation scope" in opening and "technical complexity" not in opening,
        "proof_first_opening_paragraph() should truncate a comma-heavy role_core_problem inside literal proof-first "
        f"sentence templates; got {opening!r}",
    )
    assert_true(
        "list_density_overload" not in issue_codes,
        "proof_first_opening_paragraph() should avoid list-density failures when role_core_problem is a long "
        f"comma-heavy list; got issues={sorted(issue_codes)} opening={opening!r}",
    )
    assert_true(
        all("stacked list" not in failure.lower() for failure in report["failures"] + report["warnings"]),
        f"prose_quality_report() should not flag the proof-first opening as a stacked list after truncation; got {report}",
    )


def test_cover_letter_prose_check_text_strips_header_before_quality_eval(
    build_cover_letter: object,
    writing_eval: object,
) -> None:
    body_text = "\n".join(
        [
            "Christian Estrada",
            "Atlanta, GA",
            "https://www.linkedin.com/in/cjne/",
            "July 1, 2026",
            "",
            "Dear QTS Data Centers Team,",
            "",
            "The Legal Operations Specialist role centers on implementation scope becoming real operating progress.",
            "That is the kind of work where I have been strongest.",
        ]
    )
    stripped = build_cover_letter.cover_letter_prose_check_text(body_text)
    report = build_cover_letter.prose_quality_report(stripped, "cover_letter_full")
    result = writing_eval.evaluate_text("cover_letter_full", stripped, sample_id="cover_letter_header_strip")
    issue_codes = {issue.code for issue in result.issues}

    assert_true(
        stripped.startswith("The Legal Operations Specialist role centers on implementation scope becoming real operating progress."),
        f"cover_letter_prose_check_text() should strip the header and salutation before prose checking; got {stripped!r}",
    )
    assert_true(
        "list_density_overload" not in issue_codes,
        "cover_letter_prose_check_text() should prevent header commas from creating a false dense-list prose "
        f"failure; got issues={sorted(issue_codes)} stripped={stripped!r}",
    )
    assert_true(
        all("stacked list" not in failure.lower() for failure in report["failures"] + report["warnings"]),
        f"Header stripping should keep prose_quality_report() from flagging a false stacked-list sentence; got {report}",
    )


def test_word_budget_trims_opening_filler_before_dense_proof(build_cover_letter: object) -> None:
    company_name = "Techolution"
    role_title = "AI Adoption & Transformation Manager"
    job_description = (
        "Company: Techolution\n\nJob Title: AI Adoption & Transformation Manager\n\nAbout the job\n"
        "As a Google Premium Partner, Techolution helps leading enterprise organizations transform the way "
        "they work with Google Gemini Enterprise."
    )
    opening = (
        "As a Google Premium Partner, Techolution helps leading enterprise organizations transform the way "
        "they work with Google Gemini Enterprise. The AI Adoption & Transformation Manager role centers on "
        "the point where ambiguous client problems that need structured analysis, executive alignment, and "
        "practical recommendations that hold up in execution has to become real operating progress. What "
        "keeps me engaged is making an ambiguous problem concrete enough that a team can act on it instead "
        "of continuing to debate it."
    )
    proof = (
        "Client-facing consultant with 10+ years bringing consulting-style discovery, analysis, and "
        "stakeholder alignment to complex client delivery and transformation programs. Owned a "
        "mission-critical enterprise platform across five sites and 150+ users, supported 80+ manufacturing "
        "clients across the Americas, Europe, and Asia within a $6M+ client book of business, built 200+ "
        "dashboards and KPI reporting tools, facilitated 60+ executive workshops and QBRs, and helped "
        "stabilize $1M+ in at-risk annual revenue. Those engagements needed structured problem solving, "
        "data-backed recommendations, stakeholder alignment, and execution judgment to hold up through "
        "delivery. The scope of that work included Resolved enterprise application issues across more than "
        "600 law firm offices, spanning Active Directory, SQL Server, third-party integrations, and Windows "
        "services during a support transition."
    )
    close = (
        "The combination of Consulting and Structured Problem Solving and ambiguous client problems that "
        "need structured analysis, executive alignment, and practical recommendations that hold up in "
        "execution is why this AI Adoption & Transformation Manager opportunity feels concrete rather than "
        "adjacent. I would welcome the chance to discuss where that background could help Techolution "
        "move faster and with more confidence."
    )
    trimmed = build_cover_letter.fit_standard_cover_word_budget(
        "Dear Techolution Team,",
        (opening, proof, close),
        company_name,
        role_title,
        job_description,
        mode=build_cover_letter.STANDARD_COVER_MODE,
    )
    proof_paragraph = trimmed[1]
    assert_true(
        "$6M" in proof_paragraph and "$1M" in proof_paragraph,
        "fit_standard_cover_word_budget() should keep the dense proof sentence (scale, revenue) over generic "
        f"opening/closing filler when trimming to the word cap; got body_paragraphs={trimmed}",
    )
    assert_true(
        "What keeps me engaged" not in trimmed[0],
        "fit_standard_cover_word_budget() should be able to trim generic motivation filler out of the "
        f"opening paragraph before stripping concrete proof content elsewhere; got opening={trimmed[0]!r}",
    )
    assert_true(
        role_title in trimmed[0],
        "fit_standard_cover_word_budget() must never trim the opening paragraph's role-naming sentence, "
        "since opening_quality_problem() requires role_title to appear in the first paragraph; got "
        f"opening={trimmed[0]!r}",
    )


def test_standard_cover_trim_enforces_body_sentence_cap_even_within_word_budget(build_cover_letter: object) -> None:
    company_name = "Atomicwork"
    role_title = "Enterprise Solution Consultant"
    job_description = (
        "Company: Atomicwork\n\nJob Title: Enterprise Solution Consultant\n\n"
        "Atomicwork is on a mission to transform the digital workplace experience by uniting people, "
        "processes, and platforms through AI workforce."
    )
    opening = (
        "Atomicwork is on a mission to transform the digital workplace experience by uniting people, "
        "processes, and platforms through AI workforce. The Enterprise Solution Consultant role centers on "
        "the point where client discovery has to become real operating progress. I like the discovery side "
        "of this work because the best solution conversations start by getting honest about the workflow "
        "before anyone jumps to features."
    )
    proof = (
        "Solutions consultant with 10+ years leading technical discovery, executive conversations, and "
        "customer-value positioning for enterprise software decisions built to hold up in delivery. Managed "
        "80+ manufacturing client engagements within a $6M+ book of business, led 60+ executive workshops "
        "and QBRs, and helped protect $1M+ in at-risk annual revenue through recovery and adoption work. "
        "Best used where discovery, executive communication, and implementation judgment all shape revenue, "
        "adoption, and expansion conversations."
    )
    close = (
        "I would welcome the chance to discuss where that background could help Atomicwork move faster "
        "and with more confidence."
    )
    trimmed = build_cover_letter.fit_standard_cover_word_budget(
        "Dear Atomicwork Team,",
        (opening, proof, close),
        company_name,
        role_title,
        job_description,
        mode=build_cover_letter.STANDARD_COVER_MODE,
    )
    sentence_total = sum(len(build_cover_letter.sentence_list(paragraph)) for paragraph in trimmed)

    assert_true(
        sentence_total == build_cover_letter.MAX_STANDARD_BODY_SENTENCES,
        "fit_standard_cover_word_budget() should trim to the standard cover-letter sentence cap even when the "
        f"draft is already within the word budget; got sentence_total={sentence_total} trimmed={trimmed}",
    )
    assert_true(
        "I like the discovery side of this work" not in trimmed[0],
        "fit_standard_cover_word_budget() should trim the Atomicwork opening motivation sentence before any "
        f"concrete proof or role sentence; got opening={trimmed[0]!r}",
    )
    assert_true(
        "$6M+" in trimmed[1] and "$1M+" in trimmed[1],
        "fit_standard_cover_word_budget() should preserve the dense proof paragraph while fixing a "
        f"sentence-count-only violation; got proof={trimmed[1]!r}",
    )


def test_consulting_story_summary_avoids_list_density_overload(
    writing_eval: object, build_resume: object
) -> None:
    import resume_content

    summary = resume_content.consulting_story_summary(
        "Looking for a transformation consultant to lead modernization work."
    )
    sentences = writing_eval.split_sentences(summary)
    assert_true(
        len(sentences) == 3,
        "consulting_story_summary() feeds the Professional Summary, which build_resume.py's "
        "assert_professional_summary_structure() requires to be exactly 3 sentences; got "
        f"{len(sentences)}: {sentences}",
    )
    assert_true(
        summary.count(";") <= 1,
        f"consulting_story_summary() must use at most one semicolon; got {summary.count(';')} in {summary!r}",
    )
    proof_and_close = sentences[1:]
    semicolon_sentences = [sentence for sentence in proof_and_close if ";" in sentence]
    assert_true(
        not semicolon_sentences,
        "build_resume.py's assert_professional_summary_structure() bans semicolons in every "
        "Professional Summary sentence except the opening one; got semicolons in "
        f"{semicolon_sentences}",
    )
    opening_comma_items = [item.strip() for item in sentences[0].split(",") if item.strip()]
    assert_true(
        len(opening_comma_items) <= 4,
        "build_resume.py's assert_professional_summary_structure() caps the opening sentence at four "
        f"comma-separated segments; got {len(opening_comma_items)} in {sentences[0]!r}",
    )
    overloaded = [
        sentence for sentence in sentences if writing_eval.list_density_issue(sentence) is not None
    ]
    assert_true(
        not overloaded,
        "consulting_story_summary() should not contain any single sentence that reads as a stacked list "
        f"(>=18 words, >=3 commas, >=3 connectors); got overloaded sentences={overloaded}",
    )
    soft_hits = sum(
        1
        for pattern, _label in build_resume.TOP_THIRD_OWNERSHIP_SOFTENERS
        if re.search(pattern, summary, re.I)
    )
    strong_hits = len(build_resume.TOP_THIRD_STRONG_OWNERSHIP_RE.findall(summary))
    assert_true(
        not (soft_hits >= 2 and strong_hits <= soft_hits),
        "build_resume.py's top_third_ownership_issues() flags wording as too support-oriented when "
        "soft verbs (supported/helped/assisted/...) outnumber or match strong ownership verbs "
        f"(owned/led/built/stabilized/...); got soft_hits={soft_hits}, strong_hits={strong_hits} in {summary!r}",
    )


def test_cover_letter_compaction(build_cover_letter: object) -> None:
    job_description = (
        "Company: 4flow\n"
        "Job Title: PreSales Consultant\n"
        "Help customers optimize logistics operations and supply chain networks.\n"
        "Design tailored demos, support proofs of concept, and connect customer needs to product innovation.\n"
    )
    opening = build_cover_letter.opening_method_paragraph(
        "4flow",
        "PreSales Consultant",
        job_description,
        OLLIE_RESUME_TEXT,
        mode=build_cover_letter.LONG_COVER_MODE,
    )
    assert_true(
        "My background fits directly:" not in opening and opening.count(":") == 0,
        f"opening_method_paragraph() should avoid colon-led list phrasing for presales openings; got {opening}",
    )

    bridge = build_cover_letter.anticipated_fit_bridge(
        "4flow",
        "PreSales Consultant",
        job_description,
        OLLIE_RESUME_TEXT,
    )
    assert_true(
        not bridge.lower().startswith("the bridge from implementation into solution consulting"),
        f"anticipated_fit_bridge() should use direct language instead of template framing; got {bridge}",
    )

    try:
        with contextlib.redirect_stderr(io.StringIO()):
            build_cover_letter.validate_cover_letter_shape(
                "\n".join(
                    [
                        "Dear Hiring Manager,",
                        "4flow needs a presales consultant who can connect customer requirements to a credible solution path.",
                        "I bring discovery depth; implementation judgment; and practical solution design.",
                        "Thank you for your time and consideration,",
                        "Christian Estrada",
                    ]
                )
            )
        raise SmokeFailure("validate_cover_letter_shape() should fail when cover letter prose uses too many semicolons")
    except BaseException as error:
        assert_true(
            isinstance(error, SystemExit),
            f"validate_cover_letter_shape() should hard-fail on semicolon-heavy body text; got {type(error).__name__}: {error}",
        )


def test_cover_letter_unicode_company_pronoun_rewrite(build_cover_letter: object) -> None:
    rewritten = build_cover_letter.rewrite_company_pronouns(
        "We’re more than just a software company.",
        "AppFolio",
    )
    assert_true(
        rewritten == "AppFolio is more than just a software company.",
        f"rewrite_company_pronouns() should normalize curly-apostrophe contractions before generic company substitution; got {rewritten!r}",
    )


def test_naturalness_score_and_adverb_cleanup(build_resume: object) -> None:
    high_signal = (
        "Seamlessly spearheaded a robust, best-in-class initiative that moved the needle "
        "through innovative solutions and forward-thinking execution."
    )
    high_score = build_resume.naturalness_score(high_signal)
    assert_true(
        int(high_score["score"]) > 10,
        f"naturalness_score() should flag high-signal AI phrasing; got {high_score}",
    )

    plain_bullet = "Reduced manual inventory work by 78% across five sites for 150+ users."
    plain_score = build_resume.naturalness_score(plain_bullet)
    assert_true(
        int(plain_score["score"]) < 5,
        f"naturalness_score() should stay low for factual proof-heavy text; got {plain_score}",
    )

    cleaned = build_resume.strengthen_outcome_framing("Successfully reduced manual inventory work by 78%.")
    assert_true(
        cleaned == "Reduced manual inventory work by 78%.",
        f"strengthen_outcome_framing() should remove adverb openers cleanly; got {cleaned!r}",
    )


def test_ownership_language_rewrites(build_resume: object, build_detailed_interview_guide: object, build_interview_cheat_sheet: object, build_cover_letter: object) -> None:
    rewritten = build_resume.strengthen_outcome_framing("Collaborated with Product to deliver a 22% reduction in support backlog.")
    assert_true(
        rewritten.startswith("Partnered across teams to deliver"),
        f"strengthen_outcome_framing() should rewrite collective ownership openers with proof; got {rewritten!r}",
    )

    strong_action = "Reduced manual inventory work by 22% across five sites."
    assert_true(
        build_resume.strengthen_outcome_framing(strong_action) == strong_action,
        "strengthen_outcome_framing() should not alter bullets that already start with a strong action verb",
    )

    reminders = build_interview_cheat_sheet.recruiter_reminders()
    assert_true(
        any("Ownership audit:" in reminder for reminder in reminders),
        "recruiter_reminders() should include an ownership audit reminder",
    )

    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    card = build_interview_cheat_sheet.StoryCard(
        title="Team-Owned Delivery Story",
        story_types=("Teamwork",),
        hook="A customer launch needed tighter coordination.",
        takeaways=("Ownership",),
        evidence="We mapped the issue, we aligned the team, and I handled the final escalation path.",
        level3_trait="Clarified ownership under pressure.",
        result="Protected the launch timeline.",
        outcome="The rollout stayed on track.",
        evidence_terms=("launch",),
        signals=("teamwork",),
    )
    audit_lines = build_detailed_interview_guide.story_quality_audit(card, profile)
    assert_true(
        any("OWNERSHIP check:" in line for line in audit_lines),
        "story_quality_audit() should flag evidence that uses 'we' more than 'I'",
    )

    cover_issues, cover_warnings = build_cover_letter.validate_cover_letter_text(
        "\n".join(
            [
                "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
                "Dear Hiring Manager,",
                "The role matters because it sits close to customer implementation and adoption.",
                "We have improved delivery quality. Our team kept customers aligned. We delivered the rollout and we built cleaner reporting.",
                "I would welcome a conversation about the role.",
            ]
        ),
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
        "Acme Systems",
    )
    assert_true(
        any("collective 'we' language" in issue for issue in cover_issues),
        "validate_cover_letter_text() should warn when cover letters overuse collective 'we' language",
    )
    assert_true(
        cover_warnings == [],
        f"Collective-language validation should stay in the specificity bucket, not the cover-warning bucket; got {cover_warnings}",
    )
    ownership_issues = build_resume.top_third_ownership_issues(
        "Implementation consultant supporting go-live readiness and cross-functional work.",
        ["Supported testing and rollout across multiple sites."],
    )
    assert_true(
        any("ownership line" in issue.lower() or "support-oriented" in issue.lower() for issue in ownership_issues),
        f"top_third_ownership_issues() should flag overly support-oriented top-third language; got {ownership_issues}",
    )
    strong_score = build_cover_letter.proof_selection_score(
        "I owned five-site operations and built 200+ reporting tools that improved workflow visibility.",
        "implementation_delivery",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    soft_score = build_cover_letter.proof_selection_score(
        "I supported multi-stakeholder customer environments where teams needed a clear path through go-live.",
        "implementation_delivery",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    assert_true(
        strong_score > soft_score,
        f"Cover proof selection should prefer ownership-forward evidence over support-style bridge language; got strong={strong_score}, soft={soft_score}",
    )

    try:
        with contextlib.redirect_stderr(io.StringIO()):
            build_cover_letter.validate_cover_letter_text(
                "\n".join(
                    [
                        "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
                        "Dear Hiring Manager,",
                        "The business context is software platform, healthcare, cross-functional delivery.",
                        "I would welcome a conversation about the role.",
                    ]
                ),
                LANE_JOB_DESCRIPTIONS["implementation_delivery"],
                "Acme Systems",
            )
        raise SmokeFailure("validate_cover_letter_text() should fail when business-context debug text leaks into the letter")
    except BaseException as error:
        assert_true(
            isinstance(error, SystemExit),
            f"validate_cover_letter_text() should hard-fail when business-context debug text leaks into the letter; got {type(error).__name__}: {error}",
        )


def test_cover_letter_no_jd_artifacts(build_resume: object, build_cover_letter: object) -> None:
    dummy_jd = """
Company: aCCELA
Job Title: Consultant

RESPONSIBILITIES:
• Manage ticket intake and triage from customer contacts
• Monitor and oversee MAS support

What you'll do:
• Support testing across mobile, email, and app workflows
• Improve reporting for customer support operations

QUALIFICATIONS:
Strong communication, workflow analysis, reporting, and testing judgment.

OVERVIEW:
aCCELA needs cleaner support workflows, reporting visibility, and testing follow-through.
"""
    draft = build_cover_letter.compose_cover_letter_draft(
        "Accela",
        "Consultant",
        dummy_jd,
        OLLIE_RESUME_TEXT,
        mode=build_cover_letter.STANDARD_COVER_MODE,
    )
    combined = " ".join(
        [
            *draft.body_paragraphs,
            build_cover_letter.anticipated_fit_bridge("Accela", "Consultant", dummy_jd, OLLIE_RESUME_TEXT),
        ]
    )
    assert_true(
        not re.search(r"[•\-\*]\s+\w", combined),
        f"Cover-letter draft text should remove bullet artifacts from synthetic JDs; got {combined!r}",
    )
    assert_true(
        not re.search(r"\b[a-z][A-Z]+\w*\b", combined),
        f"Cover-letter draft text should remove mixed-case JD artifacts; got {combined!r}",
    )
    assert_true(
        all(signal not in combined.lower() for signal in build_cover_letter.WEAK_OPENER_SIGNALS),
        f"Cover-letter draft text should avoid weak opener signals; got {combined!r}",
    )
    assert_true(
        build_resume.keyword_hits(combined, build_resume.audit_keywords(dummy_jd)) >= 3,
        f"Cover-letter draft text should keep direct JD keyword coverage after cleanup; got {combined!r}",
    )


def test_has_bullet_artifact_ignores_role_title_dash(build_cover_letter: object) -> None:
    # A mid-sentence dash in a role title (e.g. "Consultant - Strategic Focus") must NOT
    # be treated as a bullet-list artifact. Only line-initial dashes count.
    mid_sentence = "the Associate Implementation Consultant - Strategic Advancement Focus role centers on scope."
    assert_true(
        not build_cover_letter.has_bullet_artifact(mid_sentence),
        f"has_bullet_artifact() should not fire on a mid-sentence role-title dash; got True for {mid_sentence!r}",
    )
    # A line-initial dash (real bullet artifact) MUST still fire.
    line_initial = "Dear Hiring Manager,\n- Strategic Advancement Focus: manage implementations."
    assert_true(
        build_cover_letter.has_bullet_artifact(line_initial),
        f"has_bullet_artifact() should fire on a line-initial dash bullet; got False for {line_initial!r}",
    )
    role_title = "Associate Implementation Consultant - Strategic Advancement Focus"
    prose_letter = "\n".join(
        [
            "Dear Foundant Team,",
            f"Foundant's {role_title} role centers on turning implementation scope into operating progress.",
            "I led complex customer implementations, stakeholder alignment, training, and go-live readiness across enterprise environments.",
            "I would welcome a conversation about supporting Foundant's client delivery work.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    prose_warnings = build_cover_letter.cover_letter_preflight(
        prose_letter,
        "Foundant",
        role_title,
        "Foundant needs implementation, stakeholder communication, training, client delivery, and go-live readiness.",
    )
    assert_true(
        "Bullet characters found in cover letter body." not in prose_warnings,
        f"cover-letter preflight should not flag a mid-sentence role-title dash; got {prose_warnings!r}",
    )
    bullet_letter = prose_letter.replace(
        "I led complex customer implementations",
        "- I led complex customer implementations",
    )
    bullet_warnings = build_cover_letter.cover_letter_preflight(
        bullet_letter,
        "Foundant",
        role_title,
        "Foundant needs implementation, stakeholder communication, training, client delivery, and go-live readiness.",
    )
    assert_true(
        "Bullet characters found in cover letter body." in bullet_warnings,
        f"cover-letter preflight should still flag a real line-initial bullet; got {bullet_warnings!r}",
    )


def test_cover_letter_validator_blocks_jd_artifacts_and_warns_on_switch(build_cover_letter: object) -> None:
    raw_artifact_text = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Hiring Manager,",
            "RESPONSIBILITIES:",
            "Manage ticket intake and triage from customer contacts.",
            "Built 200+ dashboards that improved reporting quality across customer support workflows.",
            "I would welcome a conversation about how Accela needs faster early wins across support workflows.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    stderr_buffer = io.StringIO()
    try:
        with contextlib.redirect_stderr(stderr_buffer):
            build_cover_letter.validate_cover_letter_text(
                raw_artifact_text,
                DUMMY_JOB_DESCRIPTION,
                "Accela",
            )
        raise SmokeFailure("validate_cover_letter_text() should hard-fail when raw JD artifacts leak into the letter")
    except BaseException as error:
        assert_true(
            isinstance(error, SystemExit),
            f"validate_cover_letter_text() should hard-fail on raw JD artifacts; got {type(error).__name__}: {error}",
        )
        assert_true(
            "raw jd artifacts" in stderr_buffer.getvalue().lower(),
            f"JD-artifact failures should explain the root cause clearly; got {stderr_buffer.getvalue()!r}",
        )

    warned_text = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Hiring Manager,",
            "At Accela, the Consultant role protects customer support workflows and reporting quality when ticket triage, testing, and cross-functional handoffs all need to stay aligned.",
            "Built 200+ dashboards that improved reporting quality across customer support workflows. I validate the workflow, document risks, and keep product, CX, and support aligned.",
            "I would welcome a conversation about how Accela needs faster early wins across support workflows.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    specificity_warnings, cover_warnings = build_cover_letter.validate_cover_letter_text(
        warned_text,
        DUMMY_JOB_DESCRIPTION,
        "Accela",
    )
    assert_true(
        any("abrupt first-person switch" in warning for warning in cover_warnings),
        f"validate_cover_letter_text() should emit a cover warning for abrupt first-person switches; got {cover_warnings}",
    )
    assert_true(
        isinstance(specificity_warnings, list),
        "validate_cover_letter_text() should still return specificity warnings alongside cover warnings",
    )


def test_cover_letter_validator_allows_requirements_in_normal_prose(build_cover_letter: object) -> None:
    valid_text = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Guidehouse Team,",
            "I am interested in Guidehouse because this role sits at the intersection of strategy, transformation, and implementation work leaders need to execute clearly. I have turned strategy sessions and functional requirements into statements of work, milestone clarity, and more credible stakeholder decisions across 80+ international client engagements.",
            "",
            "I have also supported implementation planning, stakeholder communication, and reporting work that helped teams move from ambiguity into clearer action. I would welcome a conversation about how that background could help Guidehouse create early wins for clients.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    specificity_warnings, cover_warnings = build_cover_letter.validate_cover_letter_text(
        valid_text,
        DUMMY_JOB_DESCRIPTION,
        "Guidehouse",
    )
    assert_true(
        isinstance(specificity_warnings, list) and isinstance(cover_warnings, list),
        "validate_cover_letter_text() should allow normal prose that uses requirement language without treating it as raw JD leakage",
    )


def test_cover_letter_sections_recognize_extended_headers(build_cover_letter: object) -> None:
    jd = """
Company: Guidehouse
Job Title: Requirements/Documentation Analyst

What You Will Do:
Be responsible for the requirements documentation and analysis project workstream

What You Will Need:
Minimum of four years of experience distilling, documenting, and tracking requirements

What Would Be Nice To Have:
Experience using Jira and Confluence
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    assert_true(
        "What You Will Need" not in sections["what_youll_do"] and "tracking requirements" in sections["who_you_are"],
        f"cover_letter_jd_sections() should route 'What You Will Need' content into the qualifications bucket; got {sections}",
    )
    assert_true(
        "What Would Be Nice To Have" not in sections["what_youll_do"] and "Jira and Confluence" in sections["who_you_are"],
        f"cover_letter_jd_sections() should route nice-to-have content into the qualifications bucket; got {sections}",
    )
    cleaned = build_cover_letter.sanitize_for_spoken_text(
        "The role requires documenting requirements and communicating acceptance criteria clearly."
    )
    assert_true(
        "requirements" in cleaned.lower(),
        f"sanitize_for_spoken_text() should preserve normal requirement language instead of stripping the word globally; got {cleaned!r}",
    )


def test_cover_letter_sections_split_typical_day_headers(build_cover_letter: object) -> None:
    jd = """
Company: Infor
Job Title: CPQ Implementation Consultant, Senior

As a Functional Consultant supporting Infor's CPQ solutions, you will guide customers through end-to-end implementations of Configure, Price, and Quote capabilities.

A Typical Day in the Life Includes:
Participating in full-cycle implementations, including Kickoff, Requirements Gathering, Functional Design and Documentation, System Configuration, Testing, and Training.

Basic Qualifications:
Experience with configure, price, quote software.

About Infor
Infor is a global leader in business cloud software products for companies in industry specific markets.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    assert_true(
        "full-cycle implementations" in sections["what_youll_do"],
        f"cover_letter_jd_sections() should route 'A Typical Day in the Life Includes' content into duties; got {sections}",
    )
    assert_true(
        "configure, price, quote software" in sections["who_you_are"],
        f"cover_letter_jd_sections() should route 'Basic Qualifications' content into qualifications; got {sections}",
    )
    assert_true(
        "global leader in business cloud software" in sections["intro"].lower(),
        f"cover_letter_jd_sections() should keep clean company-background lines available for mission extraction; got {sections}",
    )


def test_cover_letter_plan_normalizes_slash_role_titles(build_cover_letter: object) -> None:
    jd = """
Company: Guidehouse
Job Title: Requirements/Documentation Analyst

What You Will Do:
Be responsible for the requirements documentation and analysis project workstream
Support the Project Manager in translating contractual requirements into user stories

What You Will Need:
Experience distilling, documenting, and tracking requirements for multi-workstream projects
"""
    plan = build_cover_letter.build_cover_letter_plan(
        "Guidehouse",
        "Requirements/Documentation Analyst",
        jd,
        PROCORE_RESUME_TEXT,
        mode=build_cover_letter.STANDARD_COVER_MODE,
    )
    draft = build_cover_letter.compose_cover_letter_from_plan(plan)
    assert_true(
        "Requirements and Documentation Analyst" in draft.body_paragraphs[0],
        f"Cover-letter planning should normalize slash-heavy role titles for readable prose; got {draft.body_paragraphs[0]!r}",
    )
    full_text = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            draft.salutation,
            *draft.body_paragraphs,
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    build_cover_letter.assert_cover_letter_qc(
        full_text,
        "Guidehouse",
        plan.prose_role_title,
        jd,
        mode=build_cover_letter.STANDARD_COVER_MODE,
    )


def test_cover_letter_blocks_unsupported_target_context(build_cover_letter: object) -> None:
    non_education_jd = """
Company: M-Tech Systems
Job Title: Implementation Specialist (Project Consultant)
This role contributes to SaaS implementation, training, user adoption, and go-live support for global clients.
"""
    invalid_text = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Hiring Manager,",
            "M-Tech Systems supports assessment and learning systems for schools, educators, and learners.",
            "The Implementation Specialist role turns that value into day-to-day execution through SaaS implementation work.",
            "I would welcome a conversation about where M-Tech Systems needs faster early wins across training and user adoption.",
        ]
    )
    try:
        with contextlib.redirect_stderr(io.StringIO()):
            build_cover_letter.validate_cover_letter_text(
                invalid_text,
                non_education_jd,
                "M-Tech Systems",
            )
        raise SmokeFailure("validate_cover_letter_text() should fail when the opening drifts into an unsupported target-company context")
    except BaseException as error:
        assert_true(
            isinstance(error, SystemExit),
            f"validate_cover_letter_text() should hard-fail on unsupported target-company context; got {type(error).__name__}: {error}",
        )


def test_extract_company_mission_prefers_clean_intro_line(build_cover_letter: object) -> None:
    jd = """
Company: M-Tech Systems
Job Title: Implementation Specialist (Project Consultant)
Implementation Specialist (Project Consultant)
Reports To: Project Team Lead Department: Projects
Location: Atlanta, GA (Hybrid - 3 days onsite)
At M-Tech Systems, our company mission is to increase yield in protein production to help feed the growing world population without compromising animal welfare or damaging the planet.
We aim to create software that delivers real-time data to the entire supply chain.
Job Summary
This is a hands-on functional implementation role.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    mission = build_cover_letter.extract_company_mission("M-Tech Systems", sections, jd)
    assert_true(
        "increase yield in protein production" in mission.lower(),
        f"extract_company_mission() should prefer the actual mission line before a fallback summary; got {mission!r}",
    )
    normalized = build_cover_letter.normalized_company_mission("M-Tech Systems", mission)
    assert_true(
        "its company mission" not in normalized.lower(),
        f"normalized_company_mission() should rewrite awkward 'its company mission' phrasing; got {normalized!r}",
    )
    assert_true(
        not normalized.lower().endswith("to help."),
        f"normalized_company_mission() should trim long mission clauses at a clean boundary; got {normalized!r}",
    )


def test_extract_company_mission_prefers_company_descriptor_over_role_intro(build_cover_letter: object) -> None:
    jd = """
Company: Infor
Job Title: CPQ Implementation Consultant, Senior

As a Functional Consultant supporting Infor's CPQ solutions, you will guide customers through end-to-end implementations of Configure, Price, and Quote capabilities.
A Typical Day in the Life Includes:
Participating in full-cycle implementations, including Kickoff, Requirements Gathering, Functional Design and Documentation, System Configuration, Testing, and Training.
Basic Qualifications:
Experience with configure, price, quote software.

About Infor
Infor is a global leader in business cloud software products for companies in industry specific markets.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    mission = build_cover_letter.extract_company_mission("Infor", sections, jd)
    assert_true(
        "global leader in business cloud software" in mission.lower(),
        f"extract_company_mission() should prefer company background over role-description lines; got {mission!r}",
    )


def test_extract_company_mission_filters_recruiting_slogans(build_cover_letter: object) -> None:
    jd = """
Company: JFrog
Job Title: Senior Technical Project Manager

At JFrog, we're reinventing DevOps to help the world's greatest companies innovate -- and we want you along for the ride.
This is a special place with a unique combination of brilliance, spirit, and just all-around great people.
Thousands of customers, including the majority of the Fortune 100, trust JFrog to manage, accelerate, and secure their software delivery from code to production -- a concept we call liquid software.

We are looking for a Senior Technical Project Manager who will manage high-complexity projects to successful completion.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    mission = build_cover_letter.extract_company_mission("JFrog", sections, jd)
    normalized = build_cover_letter.normalized_company_mission("JFrog", mission)
    assert_true(
        "--" not in normalized and "along for the ride" not in normalized.lower() and "special place" not in normalized.lower(),
        f"company mission extraction should strip recruiting slogans and double dashes; got {normalized!r}",
    )
    assert_true(
        "jfrog" in normalized.lower(),
        f"company mission extraction should still produce company-specific context after filtering slogans; got {normalized!r}",
    )


def test_extract_company_mission_prefers_business_descriptor_over_culture_line(build_cover_letter: object) -> None:
    jd = """
Company: BlackLine
Job Title: Principal Presales Consultant

At BlackLine, we're committed to bringing passion and customer focus to the business of enterprise applications.
Since being founded in 2001, BlackLine has become a leading provider of cloud software that automates and controls the entire financial close process.

Thrive at BlackLine Because You Are Joining:
A technology-based company with a sense of adventure and a vision for the future. Every door at BlackLine is open.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    mission = build_cover_letter.extract_company_mission("BlackLine", sections, jd)
    assert_true(
        "leading provider of cloud software" in mission.lower() and "is open" not in mission.lower(),
        f"extract_company_mission() should prefer the business descriptor over culture copy; got {mission!r}",
    )


def test_extract_company_mission_ignores_federal_headers(build_cover_letter: object) -> None:
    jd = """
Position Title: Logistics Management Specialist

Department of the Air Force
Air Force Materiel Command

1. Works with the senior specialist and/or project/program manager to determine the nature and scope of acquisition efforts.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    mission = build_cover_letter.extract_company_mission("Department of the Air Force", sections, jd)
    assert_true(
        not mission or "position title" not in mission.lower(),
        f"extract_company_mission() should ignore federal header text instead of turning it into a company mission; got {mission!r}",
    )


def test_federal_cover_fallback_does_not_invent_education_mission(build_cover_letter: object) -> None:
    jd = """
Position Title: Logistics Management Specialist

Department of the Air Force
Air Force Materiel Command

1. Participates in the development and coordination of assessment criteria, indicators and procedures for measuring system performance.

KNOWLEDGE, SKILLS AND ABILITIES (KSAs):
1. Ability to communicate effectively, both orally and in writing.
"""
    sections = build_cover_letter.cover_letter_jd_sections(jd)
    mission = build_cover_letter.extract_company_mission("Department of the Air Force", sections, jd)
    normalized = build_cover_letter.normalized_company_mission("Department of the Air Force", mission)
    assert_true(
        not normalized,
        f"federal cover fallback should stay blank instead of inventing an education-style mission; got {normalized!r}",
    )


def test_cover_letter_smoothing_preserves_job_acronyms(build_cover_letter: object) -> None:
    allowed = build_cover_letter.cover_allowed_acronyms(
        role_title="CPQ Implementation Consultant, Senior",
        job_description="CPQ implementations require CPQ configuration, testing, and customer-facing consulting.",
        company_name="Infor",
    )
    smoothed = build_cover_letter.smooth_cover_letter_text(
        "The CPQ Implementation Consultant, Senior role supports CPQ delivery and testing.",
        allowed_acronyms=allowed,
    )
    assert_true(
        "CPQ Implementation Consultant, Senior" in smoothed and "CPQ delivery" in smoothed,
        f"smooth_cover_letter_text() should preserve job-specific acronyms that the JD and role title rely on; got {smoothed!r}",
    )


def test_cover_letter_smoothing_normalizes_double_dashes(build_cover_letter: object) -> None:
    smoothed = build_cover_letter.smooth_cover_letter_text(
        "JFrog is reinventing DevOps -- and we want you along for the ride."
    )
    assert_true(
        "--" not in smoothed,
        f"smooth_cover_letter_text() should normalize double-dash punctuation before validation; got {smoothed!r}",
    )


def test_expand_short_cover_opening_reaches_qc_minimum(build_cover_letter: object) -> None:
    signals = build_cover_letter.CoverLetterSignals(
        company_mission="",
        role_core_function="turning buyer questions into credible solution decisions",
        top_accomplishment="Led discovery and solution framing across complex enterprise buying cycles.",
        fit_bridge="",
        jd_skill_terms=("solution", "buyer"),
        ambiguity_process="trace the buyer workflow before turning it into a recommendation path",
        jd_test_environments=(),
        communication_metric="Led 60+ executive workshops and QBRs across the Americas, Europe, and Asia.",
        partner_functions=("sales", "product"),
        jd_pain_area="buyer questions and solution decisions",
    )
    expanded = build_cover_letter.expand_short_cover_opening(
        "The Principal Presales Consultant role turns buyer questions into credible solution decisions.",
        signals,
    )
    assert_true(
        build_cover_letter.word_count(expanded) >= 30,
        f"expand_short_cover_opening() should pad short openers past the QC minimum; got {expanded!r}",
    )


def test_cover_communication_metric_stays_lane_relevant(build_cover_letter: object) -> None:
    job_description = """
Company: Infor
Job Title: CPQ Implementation Consultant, Senior

This implementation role partners with Product Management, Product Development, and Customer Support to drive successful implementations.
Responsibilities include requirements gathering, testing, and customer-facing presentations.
"""
    resume_text = "\n".join(
        [
            "Translated LivePerson and Salesforce interaction data into support insights that surfaced service gaps and recurring customer questions.",
            "Led 60+ executive workshops and QBRs across the Americas, Europe, and Asia.",
            "Converted complex implementation, data migration, integration, and customization needs into statements of work and functional requirements across 80+ international manufacturing client engagements.",
        ]
    )
    metric = build_cover_letter.extract_communication_metric(
        job_description,
        resume_text,
        lane_key="implementation_delivery",
        role_title="CPQ Implementation Consultant, Senior",
    )
    assert_true(
        "LivePerson" not in metric and "workshops" in metric,
        f"extract_communication_metric() should avoid customer-support examples for implementation roles that only mention support as a partner team; got {metric!r}",
    )


def test_short_proof_paragraph_pulls_supporting_metric(build_cover_letter: object) -> None:
    signals = build_cover_letter.CoverLetterSignals(
        company_mission="M-Tech Systems delivers enterprise software for supply-chain operators.",
        role_core_function="contributing to SaaS implementation and user adoption",
        top_accomplishment="Protected migration stability by leading implementation readiness, scope alignment, sandbox testing, validation, and targeted training across concurrent program tracks.",
        fit_bridge="",
        jd_skill_terms=("reporting", "testing"),
        ambiguity_process="validate the workflow, dependencies, and testing path early so go-live risk does not compound",
        jd_test_environments=(),
        communication_metric="Led 60+ executive workshops and QBRs across the Americas, Europe, and Asia.",
        partner_functions=("product", "engineering"),
        jd_pain_area="technical scoping, integrations, data migration, and delivery risk",
    )
    paragraph, _ = build_cover_letter.proof_reframing_paragraph("M-Tech Systems", signals)
    assert_true(
        "Led 60+ executive workshops and QBRs" in paragraph,
        f"proof_reframing_paragraph() should add a second proof sentence when the proof paragraph is too short; got {paragraph!r}",
    )


def test_compact_cover_sentence_avoids_incomplete_clause_fragments(build_cover_letter: object) -> None:
    sentence = build_cover_letter.compact_cover_sentence(
        "stabilized high-risk accounts across a $6M+ book of business, including more than one million dollars in at-risk annual revenue, by diagnosing root causes, consolidating ownership, and driving resolution",
        max_words=22,
    )
    assert_true(
        "diagnosing root." not in sentence.lower() and sentence.endswith("."),
        f"compact_cover_sentence() should not leave incomplete clause fragments after compaction; got {sentence!r}",
    )


def test_change_consulting_cover_letter_stays_in_change_lane(build_cover_letter: object) -> None:
    jd = """
Company: Guidehouse
Job Title: Organization Development Consultant

Guidehouse is seeking an Organization Development Consultant to support organizational transformation and change initiatives. This role contributes to organizational design, leadership development, and team effectiveness efforts to help clients achieve mission outcomes.

Responsibilities include:
Support organizational design and restructuring efforts, including contributing to current and future state analyses and transition planning
Assist in the development and implementation of change management and organization development strategies aligned with client mission and workforce needs
Contribute to the design and delivery of leadership development activities, including training materials, workshops, and coaching support
Help facilitate team effectiveness sessions, such as strategic planning workshops, team charters, and working sessions
Participate in organizational assessments, including data collection, analysis, and synthesis of findings into actionable insights
"""
    resume_text = "\n".join(
        [
            "Facilitated more than 60 executive workshops and QBRs that aligned technology decisions to business objectives and measurable outcomes.",
            "Built 200+ dashboards and KPI reporting tools that gave leaders clearer visibility into operational performance and decision-making.",
            "Owned a mission-critical enterprise system across five sites and 150+ users while supporting data migration, validation, training, and adoption.",
            "Converted complex implementation, data migration, integration, and customization needs into statements of work and functional requirements across 80+ international manufacturing client engagements, clarifying scope, milestones, and cost baselines before build work began.",
        ]
    )
    profile = build_cover_letter.build_resume.job_problem_profile(jd, resume_text)
    lane_key = build_cover_letter.effective_lane_key("Organization Development Consultant", jd, profile)
    assert_true(
        lane_key == "change_enablement",
        f"effective_lane_key() should keep organization-development consulting roles in change_enablement; got {lane_key!r}",
    )
    draft = build_cover_letter.compose_cover_letter_draft(
        "Guidehouse",
        "Organization Development Consultant",
        jd,
        resume_text,
    )
    combined = " ".join(draft.body_paragraphs)
    assert_true(
        "b2b customers or client accounts" not in combined.lower(),
        f"compose_cover_letter_draft() should drop generic consulting-fallback mission text; got {combined!r}",
    )
    assert_true(
        any(term in combined.lower() for term in ("change management", "organizational design", "leadership development", "team effectiveness")),
        f"compose_cover_letter_draft() should surface change-specific specialty language for organization-development roles; got {combined!r}",
    )
    letter_text = "\n".join(
        [
            f"Christian Estrada | christianj1914@gmail.com | {build_cover_letter.build_resume.LINKEDIN_URL}",
            draft.salutation,
            *draft.body_paragraphs,
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    specificity_warnings, cover_warnings = build_cover_letter.validate_cover_letter_text(
        letter_text,
        jd,
        "Guidehouse",
    )
    assert_true(
        not any("specialty areas" in warning for warning in specificity_warnings),
        f"validate_cover_letter_text() should not emit specialty warnings when the draft uses change-specific language; got {specificity_warnings}",
    )
    assert_true(
        not any("abrupt first-person switch" in warning for warning in cover_warnings),
        f"validate_cover_letter_text() should avoid abrupt first-person-switch warnings for the generated draft; got {cover_warnings}",
    )


def test_finalize_communication_metric_candidate_rejects_fragment(build_cover_letter: object) -> None:
    fragment = (
        "Led change adoption and continuous improvement for mission-critical enterprise system "
        "as the primary mission-critical enterprise system across a."
    )
    assert_true(
        build_cover_letter.finalize_communication_metric_candidate(fragment) == "",
        "finalize_communication_metric_candidate() should reject compacted fragments that end mid-thought",
    )


def test_cover_prompt_leak_patterns_catch_maps_directly_to() -> None:
    import config.language_rules as language_rules

    assert_true(
        any(re.search(pattern, "maps directly to stakeholder communication work", re.I) for pattern in language_rules.PROMPT_LEAK_PATTERNS),
        "PROMPT_LEAK_PATTERNS should catch the phrase 'maps directly to'",
    )
    assert_true(
        any(re.search(pattern, "maps to this role", re.I) for pattern in language_rules.PROMPT_LEAK_PATTERNS),
        "PROMPT_LEAK_PATTERNS should catch the phrase 'maps to this role'",
    )


def test_positive_question_framing(build_interview_cheat_sheet: object) -> None:
    needs_reframe, reframe = build_interview_cheat_sheet.audit_question_framing("What are the biggest challenges?")
    assert_true(
        needs_reframe and bool(reframe),
        "audit_question_framing() should detect and reframe negatively framed questions",
    )

    needs_reframe, reframe = build_interview_cheat_sheet.audit_question_framing("What does success look like?")
    assert_true(
        not needs_reframe and reframe == "",
        "audit_question_framing() should leave positive-framed questions alone",
    )

    cleaned, changes = build_interview_cheat_sheet.reframe_questions_to_positive(
        [
            "What are the biggest challenges?",
            "What does success look like?",
            "Why is this role open right now?",
        ]
    )
    assert_true(
        len(cleaned) == 3 and len(changes) == 2,
        f"reframe_questions_to_positive() should reframe only the negative items; got cleaned={cleaned}, changes={changes}",
    )
    warnings = build_interview_cheat_sheet.interview_question_quality_warnings(
        ["What are the biggest challenges?"]
    )
    assert_true(
        any("Negative framing:" in warning for warning in warnings),
        "interview_question_quality_warnings() should surface negative-framing warnings",
    )


def test_alignment_score_report(build_resume: object) -> None:
    resume_text = "\n".join(
        [
            "Professional Summary",
            "Implementation consultant with 10+ years improving go-live readiness, data migration quality, and adoption outcomes for enterprise software teams.",
            "Professional Experience",
            "Implementation Project Manager    March 2023 - Present",
            "Known Company | Atlanta, GA",
            "Reduced onboarding delays by 22% across five sites for 150+ users.",
            "Built SQL-based reporting that improved launch readiness and stakeholder decision-making.",
            "Education",
        ]
    )
    report = build_resume.alignment_score_report(DUMMY_JOB_DESCRIPTION, resume_text)
    assert_true(
        isinstance(report, dict),
        f"alignment_score_report() should return a dict; got {type(report)}",
    )
    for key in (
        "total_score",
        "keyword_coverage",
        "requirement_coverage",
        "lane_fit",
        "specialty_fit",
        "business_context",
        "outcome_density",
        "grade",
        "score_scale_max",
        "minimum_pass_score",
        "preferred_target_score",
        "minimum_pass_met",
        "preferred_target_met",
    ):
        assert_true(key in report, f"alignment_score_report() is missing key: {key}")
    assert_true(
        0 <= int(report["total_score"]) <= int(report["score_scale_max"]),
        f"alignment_score_report() should stay within 0-{report['score_scale_max']}; got {report['total_score']}",
    )
    expected_total = (
        int(report["keyword_coverage"]["score"])
        + int(report["requirement_coverage"]["score"])
        + int(report["lane_fit"]["score"])
        + int(report["specialty_fit"]["score"])
        + int(report["business_context"]["score"])
        + int(report["outcome_density"]["score"])
    )
    assert_true(
        int(report["total_score"]) == expected_total,
        f"alignment_score_report() should include requirement and specialty scores in the total; expected {expected_total}, got {report['total_score']}",
    )
    assert_true(
        report["grade"] in {"Strong Fit", "Adjacent Fit", "Stretch Fit", "Poor Fit"},
        f"alignment_score_report() returned an invalid grade: {report['grade']}",
    )


def test_alignment_score_report_distinguishes_lane_and_domain_fit(build_resume: object) -> None:
    resume_text = "\n".join(
        [
            "Professional Summary",
            "Implementation consultant with 10+ years improving go-live readiness, data migration quality, customer-facing delivery, risk mitigation, and reporting outcomes for enterprise software teams.",
            "Professional Experience",
            "Implementation Project Manager    March 2023 - Present",
            "Known Company | Atlanta, GA",
            "Led cross-functional client implementations from requirements definition through configuration, testing, go-live, and post-go-live support.",
            "Facilitated executive workshops, managed issue triage, and built SQL-based reporting that improved launch readiness and stakeholder decision-making.",
            "Education",
        ]
    )
    report = build_resume.alignment_score_report(VENSURE_STYLE_JOB_DESCRIPTION, resume_text)
    specialty_fit = report.get("specialty_fit", {})
    assert_true(
        int(report.get("lane_fit", {}).get("score", 0)) > int(specialty_fit.get("score", 0)),
        f"alignment_score_report() should still recognize materially stronger implementation lane fit than domain-specialty fit for Vensure-style roles; got {report}",
    )
    assert_true(
        isinstance(specialty_fit, dict) and int(specialty_fit.get("required", 0)) >= 1,
        f"alignment_score_report() should surface a dedicated specialty-fit report for HR/benefits/payroll roles; got {report}",
    )
    gap_labels = tuple(str(label) for label in specialty_fit.get("gap_labels", ()))
    assert_true(
        "Benefits Administration Domain" in gap_labels,
        f"specialty_fit should list the domain gap instead of flattening it into lane fit: {report}",
    )


def test_build_resume_uses_selected_resume_text_for_profile_and_alignment_report(build_resume: object) -> None:
    class StopBuild(Exception):
        pass

    selected_resume_text = "selected resume text sentinel"
    profile_calls: list[tuple[str, str]] = []
    captured_alignment: tuple[str, str] | None = None
    dummy_profile = build_resume.JobProblemProfile(
        primary_lane="implementation_delivery",
        lane_label="Implementation Delivery",
        core_problem="turning implementation work into usable operating outcomes",
        audience="customers and stakeholders",
        outcomes=("adoption", "delivery", "visibility"),
        direct_matches=("Implementation Delivery",),
        adjacent_matches=("Analytics",),
        unsupported_requirements=(),
        safe_terms=("implementation",),
    )

    original_validate_config_integrity = build_resume.validate_config_integrity
    original_validate_inputs = build_resume.validate_inputs
    original_extract_output_name = build_resume.extract_output_name
    original_extract_output_target_name = build_resume.extract_output_target_name
    original_choose_resume = build_resume.choose_resume
    original_docx_visible_text_from_path = build_resume.docx_visible_text_from_path
    original_job_problem_profile = build_resume.job_problem_profile
    original_alignment_score_report = build_resume.alignment_score_report
    try:
        build_resume.validate_config_integrity = lambda: None
        build_resume.validate_inputs = lambda: DUMMY_JOB_DESCRIPTION
        build_resume.extract_output_name = lambda _jd: "Smoke Test Systems"
        build_resume.extract_output_target_name = lambda _jd: "Smoke Test Systems - Implementation Consultant"
        build_resume.choose_resume = lambda _jd: Path("selected_resume.docx")
        build_resume.docx_visible_text_from_path = lambda _path: selected_resume_text

        def fake_profile(job_description: str, resume_text: str = ""):
            profile_calls.append((job_description, resume_text))
            return dummy_profile

        def fake_alignment(job_description: str, resume_text: str):
            nonlocal captured_alignment
            captured_alignment = (job_description, resume_text)
            raise StopBuild()

        build_resume.job_problem_profile = fake_profile
        build_resume.alignment_score_report = fake_alignment

        try:
            build_resume.build_resume()
            raise SmokeFailure("build_resume() should have been interrupted after alignment scoring")
        except StopBuild:
            pass
    finally:
        build_resume.validate_config_integrity = original_validate_config_integrity
        build_resume.validate_inputs = original_validate_inputs
        build_resume.extract_output_name = original_extract_output_name
        build_resume.extract_output_target_name = original_extract_output_target_name
        build_resume.choose_resume = original_choose_resume
        build_resume.docx_visible_text_from_path = original_docx_visible_text_from_path
        build_resume.job_problem_profile = original_job_problem_profile
        build_resume.alignment_score_report = original_alignment_score_report

    assert_true(
        bool(profile_calls) and profile_calls[0] == (DUMMY_JOB_DESCRIPTION, selected_resume_text),
        f"build_resume() should build its primary JobProblemProfile with the selected resume text before downstream helpers run; got {profile_calls}",
    )
    assert_true(
        captured_alignment == (DUMMY_JOB_DESCRIPTION, selected_resume_text),
        f"build_resume() should pass the same selected resume text into alignment_score_report(); got {captured_alignment}",
    )


def test_alignment_gate_decision(build_resume: object) -> None:
    strong_decision, strong_actions = build_resume.alignment_gate_decision(
        build_resume.ALIGNMENT_TARGET_SCORE + 2,
        {
            "keyword_coverage": {"covered": 10, "total_keywords": 12},
            "lane_fit": {"unsupported": 0},
            "specialty_fit": {"gap_labels": ()},
        },
        DUMMY_JOB_DESCRIPTION,
        "Acme Systems",
    )
    assert_true(
        strong_decision == "STRONG FIT - apply now" and 1 <= len(strong_actions) <= 3,
        f"alignment_gate_decision() should keep STRONG FIT concise; got {strong_decision}, {strong_actions}",
    )

    adjacent_decision, adjacent_actions = build_resume.alignment_gate_decision(
        build_resume.ALIGNMENT_FAIL_FLOOR + 2,
        {
            "keyword_coverage": {"covered": 6, "total_keywords": 10},
            "lane_fit": {"unsupported": 2},
            "specialty_fit": {"gap_labels": ("Benefits Administration Domain", "Payroll and HR Platform Domain")},
        },
        DUMMY_JOB_DESCRIPTION,
        "Acme Systems",
    )
    assert_true(
        adjacent_decision == "ADJACENT FIT - make targeted edits first" and bool(adjacent_actions),
        f"alignment_gate_decision() should return actions for ADJACENT FIT; got {adjacent_decision}, {adjacent_actions}",
    )
    assert_true(
        any("domain-specialty fit" in action for action in adjacent_actions),
        f"alignment_gate_decision() should mention specialty gaps when implementation fit is stronger than domain fit; got {adjacent_actions}",
    )

    stretch_decision, stretch_actions = build_resume.alignment_gate_decision(
        build_resume.ALIGNMENT_FAIL_FLOOR - 10,
        {
            "keyword_coverage": {"covered": 3, "total_keywords": 10},
            "lane_fit": {"unsupported": 3},
            "specialty_fit": {"gap_labels": ()},
        },
        DUMMY_JOB_DESCRIPTION,
        "Acme Systems",
    )
    assert_true(
        stretch_decision == "STRETCH FIT - evaluate before applying" and len(stretch_actions) == 5,
        f"alignment_gate_decision() should return five actions for STRETCH FIT; got {stretch_decision}, {stretch_actions}",
    )


def test_dynamic_header_title_line(build_resume: object) -> None:
    lane_jobs = list(LANE_JOB_DESCRIPTIONS.values())[:5]
    for job_description in lane_jobs:
        line = build_resume.dynamic_header_title_line(job_description)
        assert_true(
            isinstance(line, str) and 0 < len(line) <= 110,
            f"dynamic_header_title_line() should return a non-empty string under 110 characters; got {line!r}",
        )

    storm3_like_job = """
    Company: Storm3
    Job Title: Technical Implementation Manager
    Healthcare AI
    Responsibilities
    Own delivery of workflows to onboard customers
    Work closely with customers to understand their clinical data needs
    Translate clinical requirements to technical specifications
    Requirements
    Experience working with healthcare data
    Strong understanding of API-led integrations
    Why apply?
    Platinum Health Insurance plan
    """
    storm3_line = build_resume.dynamic_header_title_line(storm3_like_job)
    assert_true(
        storm3_line.startswith("Technical Implementation Manager"),
        f"dynamic_header_title_line() should lead with the cleanest job title; got {storm3_line!r}",
    )
    assert_true(
        "Financial Services" not in storm3_line and "Insurance" not in storm3_line,
        f"dynamic_header_title_line() should not pull benefit language into the header; got {storm3_line!r}",
    )

    state_farm_like_job = """
    Company: State Farm
    Job Title: Digital Experience Analytics Analyst
    Qualifications
    Strong business acumen in insurance/financial services, including understanding of digital self-service capabilities.
    Proven collaboration skills to build strong working relationships with business partners and State Farm associates.
    """
    state_farm_line = build_resume.dynamic_header_title_line(state_farm_like_job)
    assert_true(
        "Healthcare Technology" not in state_farm_line,
        f"dynamic_header_title_line() should not misclassify collaboration text as healthcare; got {state_farm_line!r}",
    )

    long_title_job = """
    Company: Acme Enterprise
    Job Title: Enterprise Program Office - Senior Change Delivery Lead
    Responsibilities
    Lead enterprise change delivery, adoption planning, and cross-functional program execution
    Requirements
    Experience with transformation programs and stakeholder communication
    """
    long_title_line = build_resume.dynamic_header_title_line(long_title_job)
    assert_true(
        long_title_line.startswith("Senior Change Delivery Lead"),
        f"dynamic_header_title_line() should put the cleanest title segment first; got {long_title_line!r}",
    )
    assert_true(
        "Enterprise Program" in long_title_line,
        f"dynamic_header_title_line() should preserve relevant program context from long titles; got {long_title_line!r}",
    )

    original_extract_job_title = build_resume.extract_job_title
    try:
        build_resume.extract_job_title = lambda _job_description: None
        fallback = build_resume.dynamic_header_title_line(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    finally:
        build_resume.extract_job_title = original_extract_job_title
    assert_true(
        fallback == build_resume.IMPLEMENTATION_MASTER_TITLE,
        f"dynamic_header_title_line() should fall back to the master title when no job title is detected; got {fallback!r}",
    )


def test_header_dedupe_avoids_near_duplicate_consultant_titles(build_resume: object) -> None:
    job_description = """
Company: Guidehouse
Job Title: Strategy & Transformation Senior Consultant | Strategy & Transformation Consultant | Power BI
This consulting role focuses on strategy, transformation, analytics, executive stakeholders, and Power BI reporting.
"""
    line = build_resume.dynamic_header_title_line(job_description)
    assert_true(
        "Strategy & Transformation Senior Consultant  |  Strategy & Transformation Consultant" not in line,
        f"dynamic_header_title_line() should dedupe near-identical consultant title variants; got {line!r}",
    )
    assert_true(
        "Power BI" in line,
        f"dynamic_header_title_line() should keep the distinct specialty fragment after deduping titles; got {line!r}",
    )


def test_scope_marker_injection(build_resume: object) -> None:
    healthcare_jd = """
    Company: Acme Health
    Role: Implementation Consultant
    Healthcare implementation role supporting clinical teams, patient workflows, adoption, reporting, and go-live readiness.
    """
    injected = build_resume.strengthen_outcome_framing(
        "Reduced onboarding delays by 22%.",
        healthcare_jd,
    )
    assert_true(
        "in a healthcare workflow environment." in injected,
        f"strengthen_outcome_framing() should append a healthcare scope suffix when metric proof lacks visible scale; got {injected!r}",
    )

    already_scoped = build_resume.strengthen_outcome_framing(
        "Reduced onboarding delays by 22% across five sites.",
        healthcare_jd,
    )
    assert_true(
        already_scoped == "Reduced onboarding delays by 22% across five sites.",
        "strengthen_outcome_framing() should not append a second scope suffix when scope is already visible",
    )


def test_competency_relevance_and_page_guards(build_resume: object) -> None:
    import resume_content

    sql_score = build_resume.skill_relevance_score("SQL", "Role requires SQL, dashboards, and analytics.")
    absent_score = build_resume.skill_relevance_score("Customer Retention", "Role requires SQL, dashboards, and analytics.")
    assert_true(
        sql_score >= 8 and absent_score == 0,
        f"skill_relevance_score() should reward direct matches and ignore absent skills; got sql={sql_score}, absent={absent_score}",
    )
    assert_true(
        build_resume.title_case_skill_phrase("llm-based data cleaning") == "LLM-based Data Cleaning",
        "title_case_skill_phrase() should preserve LLM capitalization inside hyphenated competency phrases",
    )
    assert_true(
        build_resume.title_case_skill_phrase("CPQ implementation consultant") == "CPQ Implementation Consultant",
        "build_resume.title_case_skill_phrase() should preserve short all-caps acronyms instead of title-casing them",
    )
    assert_true(
        resume_content.title_case_skill_phrase("CPQ implementation consultant") == "CPQ Implementation Consultant",
        "resume_content.title_case_skill_phrase() should preserve short all-caps acronyms instead of title-casing them",
    )
    assert_true(
        build_resume.title_case_skill_phrase("cpq implementation consultant") == "CPQ Implementation Consultant",
        "build_resume.title_case_skill_phrase() should restore lowercase acronyms such as cpq to their canonical uppercase form",
    )
    assert_true(
        resume_content.title_case_skill_phrase("cpq implementation consultant") == "CPQ Implementation Consultant",
        "resume_content.title_case_skill_phrase() should restore lowercase acronyms such as cpq to their canonical uppercase form",
    )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        low_count_xml = Path(temp_name) / "low_count.xml"
        low_count_xml.write_text(
            resume_with_competencies_xml(
                "Implementation consultant with 10+ years improving adoption, reporting, and go-live execution for enterprise software teams. Reduced delays by 22% across five sites for 150+ users. Built SQL reporting that improved launch decisions and workflow quality.",
                ["SQL", "Reporting", "Data Migration"],
            ),
            encoding="utf-8",
        )
        _status, low_notes = build_resume.final_fit_audit(low_count_xml, DUMMY_JOB_DESCRIPTION)
        assert_true(
            any("Optimal range is 15-25" in note for note in low_notes),
            f"final_fit_audit() should flag thin competency sections; got {low_notes}",
        )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        high_count_xml = Path(temp_name) / "high_count.xml"
        high_count_xml.write_text(
            resume_with_competencies_xml(
                "Implementation consultant with 10+ years improving adoption, reporting, and go-live execution for enterprise software teams. Reduced delays by 22% across five sites for 150+ users. Built SQL reporting that improved launch decisions and workflow quality.",
                [f"Skill {index}" for index in range(1, 27)],
            ),
            encoding="utf-8",
        )
        _status, high_notes = build_resume.final_fit_audit(high_count_xml, DUMMY_JOB_DESCRIPTION)
        assert_true(
            any("Over 25 items can dilute signal" in note for note in high_notes),
            f"final_fit_audit() should flag oversized competency sections; got {high_notes}",
        )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        trim_xml = Path(temp_name) / "trim_count.xml"
        trim_xml.write_text(
            resume_with_competencies_xml(
                "Professional Summary: Delivery context with measurable outcomes and clear stakeholder communication.",
                [f"Skill {index}" for index in range(1, 28)],
                section_heading="Core Competencies",
            ),
            encoding="utf-8",
        )
        removed = build_resume.remove_irrelevant_core_competencies(trim_xml, DUMMY_JOB_DESCRIPTION)
        remaining = build_resume.resume_snapshot(trim_xml).competency_items
        summary_text = build_resume.paragraph_infos(trim_xml)[1].text
        assert_true(
            removed == 4 and len(remaining) == 23,
            f"remove_irrelevant_core_competencies() should trim overflow to 23 items; removed={removed}, remaining={len(remaining)}",
            )
        assert_true(
            summary_text == "Professional Summary: Delivery context with measurable outcomes and clear stakeholder communication.",
            f"remove_irrelevant_core_competencies() should only edit the skills section; got summary={summary_text!r}",
        )


def test_final_fit_audit_accepts_presales_style_top_role_heading(build_resume: object) -> None:
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
<w:p><w:r><w:t>Implementation consultant leading implementation, data migration, onboarding, reporting, and go-live readiness across enterprise software workflows for multi-site customers.</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
<w:p><w:r><w:t>Solutions Consultant    March 2023 - Present</w:t></w:r></w:p>
<w:p><w:r><w:t>Known Company | Knoxville, TN</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Led implementation discovery, data migration validation, and go-live reporting across five sites for 150+ users.</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Improved onboarding pace and stakeholder visibility by building KPI reporting and workflow checkpoints.</w:t></w:r></w:p>
<w:p><w:r><w:t>Education</w:t></w:r></w:p>
</w:body></w:document>"""
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        document_xml = Path(temp_name) / "presales_top_role.xml"
        document_xml.write_text(xml, encoding="utf-8")
        _status, notes = build_resume.final_fit_audit(document_xml, DUMMY_JOB_DESCRIPTION)

    assert_true(
        not any("Could not locate top role bullets" in note for note in notes),
        f"final_fit_audit() should recognize pre-sales/CSM top-role variants instead of failing structurally; got {notes}",
    )


def test_final_fit_audit_promotes_bridge_for_guidehouse_fixture(build_resume: object) -> None:
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
<w:p><w:r><w:t>ERP Assessment Leader and client-facing strategy and transformation consultant with 10+ years leading executive analysis, delivery oversight, and stakeholder decisions across complex modernization programs. Owned a mission-critical enterprise platform across five sites and 150+ users, supported 80+ manufacturing clients across the Americas, Europe, and Asia, built 200+ dashboards and KPI reporting tools, and facilitated 60+ executive workshops and QBRs. Brings consulting, roadmap, reporting, delivery, transformation, stakeholder alignment, and status visibility experience for public-sector modernization work that needs structured recommendations and measurable follow-through.</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
<w:p><w:r><w:t>Enterprise Systems Manager    March 2023 - Present</w:t></w:r></w:p>
<w:p><w:r><w:t>Known Company | Knoxville, TN</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Led ERP assessment workshops, roadmap planning, and status reporting across a financial modernization program, translating executive objectives into measurable delivery and quality checkpoints for 150+ users.</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Built KPI dashboards and executive-facing program reporting that improved stakeholder communication, delivery oversight, and program decisions across enterprise system migration work.</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Coordinated project plans, risk tracking, and quality reviews that kept strategy, delivery, and stakeholder workstreams aligned through turnaround work.</w:t></w:r></w:p>
<w:p><w:r><w:t>Core Competencies</w:t></w:r></w:p>
<w:p><w:r><w:t>Consulting and Delivery: Consulting | Program Management | Roadmap Planning | Status Reporting | KPI Tracking | Stakeholder Alignment | Executive Communication | Risk Management | Quality Management | Project Delivery | Transformation | Financial Modernization | Analytics | Implementation | Workflow Analysis</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Development</w:t></w:r></w:p>
<w:p><w:r><w:t>Certified Scrum Product Owner</w:t></w:r></w:p>
</w:body></w:document>"""
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        document_xml = Path(temp_name) / "guidehouse_bridge.xml"
        document_xml.write_text(xml, encoding="utf-8")
        status, notes = build_resume.final_fit_audit(document_xml, GUIDEHOUSE_ERP_ASSESSMENT_JOB_DESCRIPTION)
    assert_true(
        status == "BRIDGE",
        f"final_fit_audit() should classify strong-near-fit Guidehouse-style resumes as BRIDGE when one unsupported requirement remains; got status={status!r}, notes={notes}",
    )
    assert_true(
        any("unsupported or human-review requirements" in note.lower() for note in notes),
        f"final_fit_audit() should still preserve the honest bridge-gap note; got {notes}",
    )


def test_add_simple_core_competencies_respects_cap(build_resume: object) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        cap_xml = Path(temp_name) / "competency_cap.xml"
        existing_items = [f"Implementation Skill {index}" for index in range(1, 23)]
        cap_xml.write_text(
            resume_with_competencies_xml(
                "Implementation consultant with 10+ years improving adoption, reporting, and go-live execution for enterprise software teams. Reduced delays by 22% across five sites for 150+ users. Built SQL reporting that improved launch decisions and workflow quality.",
                existing_items,
                section_heading="Core Competencies",
            ),
            encoding="utf-8",
        )
        added = build_resume.add_simple_core_competencies(
            cap_xml,
            "Role: Implementation Consultant. Requires Agile project management, dashboards, analytics, reporting, and stakeholder communication.",
        )
        remaining = build_resume.resume_snapshot(cap_xml).competency_items
        assert_true(
            added == 1 and len(remaining) == 23,
            f"add_simple_core_competencies() should stop at 23 items; added={added}, remaining={len(remaining)}",
        )


def test_xml_page_estimate_uses_word_guard(build_resume: object) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        estimate_xml = Path(temp_name) / "estimate.xml"
        line = " ".join(["implementation"] * 20)
        competency_text = "Implementation and Delivery:  " + "  |  ".join(
            [f"Skill {index}" for index in range(1, 31)]
        )
        estimate_xml.write_text(
            simple_document_xml(
                [line for _ in range(52)]
                + ["Core Competencies", competency_text]
            ),
            encoding="utf-8",
        )
        estimate = build_resume.estimate_page_count_from_xml(estimate_xml)
        assert_true(
            estimate == 3,
            f"estimate_page_count_from_xml() should use the conservative word guard for dense text and oversized competency rows; got {estimate}",
        )


def test_xml_page_estimate_shrinks_with_compact_separator_font(build_resume: object) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        estimate_xml = Path(temp_name) / "separator_estimate.xml"
        lines = [f"Line {index}" for index in range(40)]
        lines_with_separators: list[str] = []
        for index, line in enumerate(lines, start=1):
            lines_with_separators.append(line)
            if index % 4 == 0:
                lines_with_separators.append("")
        estimate_xml.write_text(simple_document_xml(lines_with_separators), encoding="utf-8")
        baseline = build_resume.estimate_page_count_from_xml(
            estimate_xml,
            body_size_hp=20.0,
            separator_size_hp=20.0,
        )
        compact_separator = build_resume.estimate_page_count_from_xml(
            estimate_xml,
            body_size_hp=20.0,
            separator_size_hp=12.0,
        )
        assert_true(
            compact_separator < baseline,
            "estimate_page_count_from_xml() should lower the fallback page estimate when the same XML uses more "
            f"compact separator paragraphs; got baseline={baseline}, compact={compact_separator}",
        )


def test_resume_integrity_allows_planned_competency_trim(build_resume: object) -> None:
    source_items = [
        "Account Recovery",
        "Adoption Strategy",
        "AI Assisted Analysis",
        "AI Assisted Documentation",
        "Aptean Encompix",
        "Aptean Intuitive",
        "Claude",
    ] + [f"Implementation Skill {index}" for index in range(1, 26)]
    source_competencies = {build_resume.normalize_compare(item) for item in source_items}
    retained = build_resume.retained_competency_items(DUMMY_JOB_DESCRIPTION, source_competencies)
    final_competencies = {item for item in source_competencies if item in retained}

    required_roles = [
        build_resume.RoleInfo(
            title="ERP Systems Manager",
            company="East West Manufacturing",
            company_context="",
            bullet_count=5,
            block_text=build_resume.MANDATORY_REORG_SENTENCE,
        ),
        build_resume.RoleInfo(
            title="Customer Success Manager",
            company="Aptean",
            company_context="",
            bullet_count=5,
            block_text=build_resume.MANDATORY_REORG_SENTENCE,
        ),
    ]
    source_snapshot = build_resume.ResumeSnapshot(
        full_text="Source snapshot",
        sections=set(build_resume.REQUIRED_SECTIONS),
        roles=required_roles,
        competency_labels={"implementation and delivery"},
        competency_items=source_competencies,
        professional_development_items={"certified scrum product owner"},
    )
    final_snapshot = build_resume.ResumeSnapshot(
        full_text="Final snapshot",
        sections=set(build_resume.REQUIRED_SECTIONS),
        roles=required_roles,
        competency_labels={"implementation and delivery"},
        competency_items=final_competencies,
        professional_development_items={"certified scrum product owner"},
    )

    build_resume.validate_resume_integrity(source_snapshot, final_snapshot, DUMMY_JOB_DESCRIPTION)


def test_resume_non_erp_audit_ignores_company_context(build_resume: object) -> None:
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
<w:p><w:r><w:t>Implementation leader improving onboarding quality and adoption outcomes.</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
<w:p><w:r><w:t>Implementation Manager    March 2023 - Present</w:t></w:r></w:p>
<w:p><w:r><w:t>Aptean | Alpharetta, GA</w:t></w:r></w:p>
<w:p><w:r><w:t>Aptean is a global provider of mission-critical, industry-specific software solutions for manufacturers, distributors, and other specialized organizations. More than 10,000 organizations across 20-plus industries and 80 countries rely on Aptean's purpose-built ERP, supply chain, and compliance platforms to streamline daily operations.</w:t></w:r></w:p>
<w:p><w:r><w:t>Owned Aptean Intuitive administration and Epicor Kinetic transition support while leading onboarding, requirements, and adoption planning.</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Reduced onboarding delays by 22%.</w:t></w:r></w:p>
<w:p><w:r><w:t>Core Competencies</w:t></w:r></w:p>
<w:p><w:r><w:t>Implementation and Delivery:  Customer Onboarding  |  Adoption Planning</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Development</w:t></w:r></w:p>
<w:p><w:r><w:t>Certified Scrum Product Owner</w:t></w:r></w:p>
<w:p><w:r><w:t>Education</w:t></w:r></w:p>
</w:body></w:document>"""
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        document_xml = Path(temp_name) / "document.xml"
        document_xml.write_text(xml, encoding="utf-8")
        audit_text = build_resume.resume_text_for_non_erp_audit(document_xml)
    assert_true(
        "Aptean Intuitive" not in audit_text and "Epicor Kinetic" not in audit_text and "ERP" not in audit_text,
        f"resume_text_for_non_erp_audit() should focus on summary and skills, not role proof or company context: {audit_text}",
    )


def test_keyword_placement_audit(build_resume: object) -> None:
    resume_text = "\n".join(
        [
            "Professional Summary",
            "Implementation consultant improving data migration quality and enterprise software adoption for customer-facing teams.",
            "Professional Experience",
            "ERP Systems Manager    March 2023 - Present",
            "Known Company | Knoxville, TN",
            "Reduced onboarding delays by 22% across five sites for 150+ users.",
            "Education",
            "Bachelor of Arts",
            "Skills",
            "Implementation and Delivery:  SQL  |  Reporting  |  Dashboard Design",
            "Professional Development",
            "Certified Scrum Product Owner",
        ]
    )
    original_audit_keywords = build_resume.audit_keywords
    try:
        build_resume.audit_keywords = lambda _job_description: {"SQL", "go-live readiness", "data migration"}
        report = build_resume.keyword_placement_audit(
            "SQL SQL go-live readiness data migration",
            resume_text,
        )
    finally:
        build_resume.audit_keywords = original_audit_keywords

    gaps = report.get("gaps", [])
    assert_true(
        any(isinstance(gap, dict) and gap.get("keyword") == "SQL" and "Skills only" in str(gap.get("issue")) for gap in gaps),
        f"keyword_placement_audit() should flag keywords that only appear in Skills; got {gaps}",
    )
    assert_true(
        any(isinstance(gap, dict) and gap.get("keyword") == "go-live readiness" and "missing from the resume" in str(gap.get("issue")) for gap in gaps),
        f"keyword_placement_audit() should flag missing priority keywords; got {gaps}",
    )
    assert_true(
        not any(isinstance(gap, dict) and gap.get("keyword") == "data migration" for gap in gaps),
        f"keyword_placement_audit() should not flag keywords already visible in the summary; got {gaps}",
    )


def test_obvious_choice_positioning(build_resume: object, build_cover_letter: object, build_interview_cheat_sheet: object) -> None:
    job_description = LANE_JOB_DESCRIPTIONS["change_enablement"]
    profile = build_resume.job_problem_profile(job_description)
    positioning = build_resume.obvious_choice_positioning(profile, job_description)
    assert_true(
        "obvious choice" in str(positioning.get("sentence", "")).lower() and str(positioning.get("short_line", "")).startswith("Obvious-choice frame:"),
        f"obvious_choice_positioning() should return both sentence and short-line variants; got {positioning}",
    )

    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(
        "obvious choice" not in summary.lower(),
        f"build_problem_first_summary() should keep summary language direct instead of leaking coaching phrasing; got {summary}",
    )

    original_select_opening_pattern = build_cover_letter._select_opening_pattern
    try:
        build_cover_letter._select_opening_pattern = (
            lambda company_name, role_title, _job_description: f"The {role_title} role at {company_name} is an exciting opportunity."
        )
        opening = build_cover_letter.opening_method_paragraph(
            "Acme Systems",
            "Implementation Consultant",
            job_description,
            OLLIE_RESUME_TEXT,
            mode=build_cover_letter.LONG_COVER_MODE,
        )
    finally:
        build_cover_letter._select_opening_pattern = original_select_opening_pattern
    assert_true(
        "obvious choice" not in opening.lower(),
        f"opening_method_paragraph() should keep cover letter openings free of obvious-choice coaching language; got {opening}",
    )

    memorable_lines = build_interview_cheat_sheet.memorable_candidate_lines(profile, "Acme Systems")
    assert_true(
        memorable_lines and memorable_lines[0].startswith("Obvious-choice frame:"),
        f"memorable_candidate_lines() should lead with the obvious-choice coaching line; got {memorable_lines}",
    )


def test_positioning_statement_output(build_resume: object, build_interview_cheat_sheet: object) -> None:
    job_description = LANE_JOB_DESCRIPTIONS["customer_success"]
    profile = build_resume.job_problem_profile(job_description)
    statement = build_resume.generate_positioning_statement(profile, job_description)
    assert_true(
        statement.startswith("Christian is strongest"),
        f"generate_positioning_statement() should return a concise positioning statement; got {statement!r}",
    )

    variants = build_interview_cheat_sheet.spoken_positioning_variants(profile, job_description)
    assert_true(
        len(variants) == 3 and variants[0].startswith("Direct:") and variants[1].startswith("Conversational:") and variants[2].startswith("Executive:"),
        f"spoken_positioning_variants() should return three short spoken variants; got {variants}",
    )


def test_future_bridge_summary_and_bullet_clause(build_resume: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    bridge = build_resume.summary_future_bridge(profile, DUMMY_JOB_DESCRIPTION)
    assert_true(
        "the same pattern fits" not in bridge.lower() and bridge.endswith("."),
        f"summary_future_bridge() should return a short future-facing bridge sentence; got {bridge!r}",
    )

    summary = (
        "Implementation consultant with 10+ years leading enterprise software delivery. "
        "Reduced onboarding delays by 22% across five sites for 150+ users. "
        "Helps teams move complex system work into stable adoption and clearer decisions."
    )
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        summary_xml = Path(temp_name) / "summary_bridge.xml"
        summary_xml.write_text(professional_summary_xml(summary), encoding="utf-8")
        changed = build_resume.append_summary_future_bridge(summary_xml, DUMMY_JOB_DESCRIPTION)
        updated_summary = build_resume.paragraph_infos(summary_xml)[1].text
        assert_true(
            changed == 0 and len(build_resume.summary_sentences(updated_summary)) == 3,
            f"append_summary_future_bridge() should stay disabled so summaries remain three sentences; got changed={changed}, summary={updated_summary!r}",
        )
        build_resume.assert_professional_summary_structure(summary_xml)

    w_ns = build_resume.W.strip("{}")
    bullet_doc = f"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="{w_ns}">
  <w:body>
    <w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
    <w:p><w:r><w:t>{summary}</w:t></w:r></w:p>
    <w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
    <w:p><w:r><w:t>ERP Systems Manager    March 2023 - Present</w:t></w:r></w:p>
    <w:p><w:r><w:t>Known Company | Knoxville, TN</w:t></w:r></w:p>
    <w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>Coordinated testing and launch support across teams.</w:t></w:r></w:p>
    <w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>Reduced onboarding delays by 22% across five sites for 150+ users.</w:t></w:r></w:p>
    <w:p><w:r><w:t>Education</w:t></w:r></w:p>
  </w:body>
</w:document>
"""
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        bullet_xml = Path(temp_name) / "bullet_bridge.xml"
        bullet_xml.write_text(bullet_doc, encoding="utf-8")
        changed = build_resume.reorder_bullets(
            bullet_xml,
            {"adoption", "launch"},
            profile,
        )
        first_bullet = build_resume.experience_bullet_texts(bullet_xml)[0]
        assert_true(
            changed >= 1 and "a pattern that" not in first_bullet.lower(),
            f"reorder_bullets() should surface the strongest bullet without appending forward-looking meta-commentary; got {first_bullet!r}",
        )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        revert_xml = Path(temp_name) / "bullet_bridge_revert.xml"
        revert_xml.write_text(bullet_doc, encoding="utf-8")
        original_estimate = build_resume.estimate_page_count_from_xml
        call_count = {"value": 0}

        def fake_estimate(_path: object) -> int:
            call_count["value"] += 1
            return 2 if call_count["value"] == 1 else 3

        try:
            build_resume.estimate_page_count_from_xml = fake_estimate
            build_resume.reorder_bullets(revert_xml, {"adoption", "launch"}, profile)
        finally:
            build_resume.estimate_page_count_from_xml = original_estimate
        reverted_first_bullet = build_resume.experience_bullet_texts(revert_xml)[0]
        assert_true(
            "a pattern that" not in reverted_first_bullet.lower(),
            f"reorder_bullets() should never append forward-looking meta-commentary to the first bullet; got {reverted_first_bullet!r}",
        )


def test_offer_blocker_logic(build_resume: object, build_cover_letter: object, build_detailed_interview_guide: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    blockers = build_detailed_interview_guide.six_offer_blocker_lines(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        len(blockers) == 6 and blockers[0].startswith("Offer blocker 1"),
        f"six_offer_blocker_lines() should return six concrete interview blockers; got {blockers}",
    )

    skim_issues = build_resume.hiring_manager_skim_issues(
        "Implementation consultant with 10+ years helping teams through complex rollout work. Strong communicator who partners well across teams. Helps teams move complex system work into stable adoption and clearer decisions.",
        ["Coordinated training and launch support across teams."],
        profile,
    )
    assert_true(
        any("surface proof fast enough" in issue for issue in skim_issues),
        f"hiring_manager_skim_issues() should flag proof-light top thirds more aggressively; got {skim_issues}",
    )

    specificity_issues = build_cover_letter.validate_cover_letter_specificity(
        "\n".join(
            [
                "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
                "Dear Hiring Manager,",
                "The Implementation Consultant role at Acme Systems is an exciting opportunity for me.",
                "I help teams turn implementation, data migration, and adoption work into practical outcomes because the customer impact is real.",
                "Thank you for your consideration.",
            ]
        ),
        "Acme Systems",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        any("Offer blocker: generic opener phrase" in issue for issue in specificity_issues),
        f"validate_cover_letter_specificity() should flag generic opener phrases as offer blockers; got {specificity_issues}",
    )


def test_ats_plain_text_validation(build_resume: object) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        valid_docx = Path(temp_name) / "valid_resume.docx"
        write_docx_with_lines(
            valid_docx,
            [
                "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
                "Professional Summary",
                "Implementation consultant with 10+ years improving adoption, reporting, and go-live readiness for enterprise software teams.",
                "Professional Experience",
                "ERP Systems Manager    March 2023 - Present",
                "Known Company | Knoxville, TN",
                "Reduced onboarding delays by 22% across five sites for 150+ users.",
                "Education",
                "Bachelor of Arts",
                "Skills",
                "Implementation and Delivery:  SQL  |  Reporting  |  Data Migration",
                "Professional Development",
                "Certified Scrum Product Owner",
            ],
        )
        valid_report = build_resume.ats_plain_text_validation(valid_docx)
        assert_true(
            not valid_report["blockers"],
            f"ats_plain_text_validation() should pass a resume with sections, email, and date ranges; got {valid_report}",
        )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        legacy_docx = Path(temp_name) / "legacy_resume.docx"
        write_docx_with_lines(
            legacy_docx,
            [
                "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
                "Professional Summary",
                "Implementation consultant with 10+ years improving adoption, reporting, and go-live readiness for enterprise software teams.",
                "Professional Experience",
                "ERP Systems Manager    March 2023 - Present",
                "Known Company | Knoxville, TN",
                "Reduced onboarding delays by 22% across five sites for 150+ users.",
                "Education",
                "Bachelor of Arts",
                "Core Competencies",
                "Implementation and Delivery:  SQL  |  Reporting  |  Data Migration",
                "Professional Development",
                "Certified Scrum Product Owner",
            ],
        )
        legacy_report = build_resume.ats_plain_text_validation(legacy_docx)
        assert_true(
            not legacy_report["blockers"],
            f"ats_plain_text_validation() should still accept legacy Core Competencies headings; got {legacy_report}",
        )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        invalid_docx = Path(temp_name) / "invalid_resume.docx"
        write_docx_with_lines(
            invalid_docx,
            [
                "Christian Estrada | linkedin.com/in/cjne",
                "Professional Summary",
                "Implementation consultant improving adoption and reporting.",
                "Professional Experience",
                "ERP Systems Manager",
                "Known Company | Knoxville, TN",
                "Supported implementation work for customers.",
                "Education",
                "Bachelor of Arts",
                "Skills",
                "Implementation and Delivery:  SQL  |  Reporting",
                "Professional Development",
                "Certified Scrum Product Owner",
            ],
        )
        invalid_report = build_resume.ats_plain_text_validation(invalid_docx)
        assert_true(
            any("email address" in blocker for blocker in invalid_report["blockers"]) and any("date ranges" in blocker for blocker in invalid_report["blockers"]),
            f"ats_plain_text_validation() should treat missing email and missing date ranges as blocking; got {invalid_report}",
        )

    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        joined_docx = Path(temp_name) / "joined_resume.docx"
        write_docx_with_lines(
            joined_docx,
            [
                "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
                "Professional Summary",
                "Implementation consultant with 10+ years improving adoption, reporting, and go-live readiness for enterprise software teams.",
                "Professional Experience",
                "ERP Systems ManagerMarch 2023 - Present",
                "Known Company | Knoxville, TN",
                "Reduced onboarding delays by 22% across five sites for 150+ users.",
                "Education",
                "Bachelor of Arts",
                "Skills",
                "Implementation and Delivery:  SQL  |  Reporting  |  Data Migration",
                "Professional Development",
                "Certified Scrum Product Owner",
            ],
        )
        joined_report = build_resume.ats_plain_text_validation(joined_docx)
        assert_true(
            not joined_report["blockers"],
            f"ats_plain_text_validation() should tolerate title/date lines where the month is attached to the title; got {joined_report}",
        )


def test_moment_in_time_context(build_resume: object, build_cover_letter: object) -> None:
    job_description = """
    Company: Acme Systems
    Role: Implementation Consultant
    The team is scaling after an acquisition and major platform migration, with heavy focus on integration, adoption, and go-live execution.
    """
    moment_context = build_cover_letter.moment_in_time_context(job_description)
    assert_true(
        "role sits where" in moment_context.lower(),
        f"moment_in_time_context() should detect time-sensitive operating context; got {moment_context!r}",
    )

    summary = build_resume.build_problem_first_summary(job_description)
    assert_true(
        "at a point where" not in summary.lower(),
        f"build_problem_first_summary() should keep the summary opener direct instead of injecting moment-in-time framing; got {summary}",
    )

    original_select_opening_pattern = build_cover_letter._select_opening_pattern
    try:
        build_cover_letter._select_opening_pattern = (
            lambda company_name, role_title, _job_description: f"The {role_title} role at {company_name} needs strong implementation judgment."
        )
        opening = build_cover_letter.opening_method_paragraph(
            "Acme Systems",
            "Implementation Consultant",
            job_description,
            OLLIE_RESUME_TEXT,
            mode=build_cover_letter.LONG_COVER_MODE,
        )
    finally:
        build_cover_letter._select_opening_pattern = original_select_opening_pattern
    assert_true(
        not opening.lower().startswith("the team appears to be at a point where"),
        f"opening_method_paragraph() should keep the opening specific instead of prepending generic moment-in-time framing; got {opening}",
    )


def test_customer_success_opening_surfaces_concrete_context(build_resume: object, build_cover_letter: object) -> None:
    job_description = """
    Company: HubSpot
    Job Title: Senior Customer Success Manager
    HubSpot is an AI-powered customer platform.
    This role owns a complex book of business, retention forecasting, mitigation planning, customer renewals within six months,
    and growth across HubSpot's product suite.
    """
    opening = build_cover_letter._tension_opening(
        "HubSpot",
        "Senior Customer Success Manager",
        job_description,
    )
    lowered = opening.lower()
    assert_true(
        "platform" in lowered and "book of business" in lowered,
        f"customer-success opening should name concrete role context like platform and book of business; got {opening}",
    )
    assert_true(
        "renewal" in lowered or "mitigation" in lowered,
        f"customer-success opening should surface renewal or risk pressure early; got {opening}",
    )
    assert_true(
        opening.count(":") == 0,
        f"customer-success opening should avoid colon-led framing; got {opening}",
    )
    assert_true(
        build_resume.objective_context_signal_count(opening) >= 1,
        f"customer-success opening should contribute objective business context before the proof paragraph; got {opening}",
    )


def test_customer_success_support_paragraph_stays_separate(build_cover_letter: object) -> None:
    paragraphs = build_cover_letter.finalize_body_paragraphs(
        "Opening paragraph about HubSpot and the role.",
        "Proof paragraph with revenue, dashboards, and workshops.",
        [
            "Support paragraph about commercial judgment and account recovery. In the first 90 days, I would focus on account health signals."
        ],
        "",
        "Closing paragraph asking for a conversation.",
    )
    assert_true(
        paragraphs == [
            "Opening paragraph about HubSpot and the role.",
            "Proof paragraph with revenue, dashboards, and workshops.",
            "Support paragraph about commercial judgment and account recovery. In the first 90 days, I would focus on account health signals.",
            "Closing paragraph asking for a conversation.",
        ],
        f"finalize_body_paragraphs() should keep a support paragraph separate instead of merging it into the proof paragraph; got {paragraphs}",
    )


def test_cover_letter_uses_company_specific_context(build_cover_letter: object) -> None:
    job_description = """
    Company: BioTouch
    Role: Solutions Implementation Consultant
    Support customer implementations, issue tracking, workflow documentation, and stakeholder communication.
    """
    company_context = (
        "POST-INTERVIEW NOTE [05/26/2026]: BioTouch, round 2, outcome pending.\n"
        "We would be supporting Spectra Path, internally using Jira and Azure to track tickets and document issues"
    )
    original_select_opening_pattern = build_cover_letter._select_opening_pattern
    try:
        build_cover_letter._select_opening_pattern = (
            lambda company_name, role_title, _job_description: f"The {role_title} role at {company_name} needs steady implementation follow-through."
        )
        opening = build_cover_letter.opening_method_paragraph(
            "BioTouch",
            "Solutions Implementation Consultant",
            job_description,
            OLLIE_RESUME_TEXT,
            company_context,
            mode=build_cover_letter.LONG_COVER_MODE,
        )
    finally:
        build_cover_letter._select_opening_pattern = original_select_opening_pattern
    assert_true(
        all(term in opening for term in ("Spectra Path", "Jira", "Azure")),
        f"opening_method_paragraph() should surface concrete company context when supplied; got {opening}",
    )
    assert_true(
        opening.startswith("The Solutions Implementation Consultant role at BioTouch needs steady implementation follow-through. The detail that the team would support"),
        f"opening_method_paragraph() should turn company notes into a direct second sentence; got {opening}",
    )


def test_gap_address_paragraph(build_cover_letter: object) -> None:
    original_poor_fit_requirements = build_cover_letter.build_resume.poor_fit_requirements
    try:
        build_cover_letter.build_resume.poor_fit_requirements = lambda _job_description, _resume_text: ["direct people management"]
        paragraph = build_cover_letter.gap_address_paragraph(
            "Acme Systems",
            "Implementation Consultant",
            DUMMY_JOB_DESCRIPTION,
            OLLIE_RESUME_TEXT,
        )
    finally:
        build_cover_letter.build_resume.poor_fit_requirements = original_poor_fit_requirements

    sentences = [part.strip() for part in re.split(r"(?<=[.!?])\s+", paragraph) if part.strip()]
    assert_true(
        len(sentences) == 3 and "direct people management" in paragraph.lower(),
        f"gap_address_paragraph() should return a three-sentence gap bridge tied to the poor-fit signal; got {paragraph!r}",
    )
    assert_true(
        not paragraph.lower().startswith(("while i have not", "although i do not", "i know i lack", "i may not have")),
        f"gap_address_paragraph() should avoid defensive opener language; got {paragraph!r}",
    )


def test_first_90_days_cover_sentence(build_cover_letter: object) -> None:
    senior_sentence = build_cover_letter.first_90_days_cover_sentence(
        "Senior Implementation Manager",
        DUMMY_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
    )
    non_senior_sentence = build_cover_letter.first_90_days_cover_sentence(
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
    )
    assert_true(
        senior_sentence.startswith("In the first 90 days") and non_senior_sentence == "",
        f"first_90_days_cover_sentence() should trigger only for senior-role titles; got senior={senior_sentence!r}, non_senior={non_senior_sentence!r}",
    )


def test_diagnose_before_selling_framework(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    framework = build_interview_cheat_sheet.diagnose_before_selling_framework(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        len(framework) >= 4 and "Use the pivot question" in framework[2],
        f"diagnose_before_selling_framework() should return the early interview framework and pivot prompt; got {framework}",
    )

    story = build_interview_cheat_sheet.StoryCard(
        title="Implementation Recovery",
        story_types=("Individual Achievement",),
        hook="A customer rollout was slipping.",
        takeaways=("Discovery", "Ownership", "Adoption"),
        evidence="Mapped the issue, aligned the owners, and reset the operating rhythm.",
        level3_trait="Separated symptoms from the root cause quickly.",
        result="Protected the timeline and improved customer trust.",
        outcome="The launch stayed on track.",
        evidence_terms=("implementation",),
        signals=("delivery",),
    )
    answers = build_interview_cheat_sheet.common_interview_answers(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        [story],
    )
    assert_true(
        answers and answers[0].prompt == "Diagnose-before-selling pivot question",
        f"common_interview_answers() should place the pivot question first; got {[answer.prompt for answer in answers[:3]]}",
    )


def test_bold_diagnostic_questions(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    questions = build_interview_cheat_sheet.bold_diagnostic_questions(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        len(questions) >= 4 and questions[0].startswith("When implementations slow down"),
        f"bold_diagnostic_questions() should lead with the lane-specific structural question; got {questions}",
    )


def test_preloaded_questions(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    questions = build_interview_cheat_sheet.preloaded_questions(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        len(questions) == 3 and questions[0].startswith("When implementations slow down"),
        f"preloaded_questions() should return exactly three questions with the structural question first; got {questions}",
    )


def test_slot_based_summary_and_interview_answers(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    analytics_slot_jd = (
        OLLIE_ANALYTICS_JOB_DESCRIPTION
        + "\nThis analytics role handles data analysis, reporting, SQL, cohort analysis, retention analysis, "
        + "ad hoc analysis, metric definition, customer behavior, and imperfect business questions."
    )
    summary = build_resume.build_problem_first_summary(analytics_slot_jd, OLLIE_RESUME_TEXT)
    assert_true(
        "turning unclear business questions into retention analysis" in summary,
        f"build_problem_first_summary() should reuse the compact ambiguity-to-outcome slot for analytics summaries; got {summary}",
    )

    profile = build_resume.job_problem_profile(analytics_slot_jd, OLLIE_RESUME_TEXT)
    story = build_interview_cheat_sheet.StoryCard(
        title="Analytics Visibility",
        story_types=("Analysis and Decision", "Individual Achievement", "Rapid Learning"),
        hook="A reporting gap was slowing decisions.",
        takeaways=("Mapped the question", "Made the data usable", "Kept decisions moving"),
        evidence="Built the reporting path and aligned owners on the next decision.",
        level3_trait="Clarified the decision before diving into the data.",
        result="Built 200+ dashboards and reporting tools that made operating decisions clearer.",
        outcome="Leaders had cleaner visibility and faster follow-through.",
        evidence_terms=("200+ dashboards",),
        signals=("analytics", "reporting", "decision"),
    )
    answers = {
        item.prompt: item.answer
        for item in build_interview_cheat_sheet.behavioral_answer_scripts(profile, [story], analytics_slot_jd)
    }
    tell_me = answers["Tell me about yourself"]
    ambiguous = answers["Ambiguous problem"]
    assert_true(
        "Three takeaways:" not in tell_me and "200+" in tell_me,
        f"behavioral_answer_scripts() should use a tighter point-proof-bridge answer for tell me about yourself; got {tell_me}",
    )
    assert_true(
        "Three takeaways:" not in ambiguous and "facts, assumptions, risks, and owners" in ambiguous,
        f"behavioral_answer_scripts() should use the direct ambiguity workflow instead of the older takeaways list; got {ambiguous}",
    )

    challenge = build_detailed_interview_guide.role_challenge_answer(
        "Ollie",
        "Data Analyst, Retention",
        profile,
        analytics_slot_jd,
    )
    assert_true(
        "clarify the outcome" in challenge and "decision rhythm" in challenge,
        f"role_challenge_answer() should stay in the compact challenge-approach-early-win structure; got {challenge}",
    )


def test_value_compression_opening(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    opening = build_interview_cheat_sheet.value_compression_opening(profile, DUMMY_JOB_DESCRIPTION)
    reminders = build_interview_cheat_sheet.recruiter_reminders()
    assert_true(
        opening.startswith("I help teams move from messy implementation work"),
        f"value_compression_opening() should produce a short value-first opener for the lane; got {opening!r}",
    )
    assert_true(
        reminders[0] == "Value first, proof second, chronology never.",
        f"recruiter_reminders() should surface the value-first reminder early; got {reminders[:3]}",
    )


def test_pitch_variants_reuse_cover_letter_logic(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(OLLIE_ANALYTICS_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    variants = build_interview_cheat_sheet.pitch_variants(
        profile,
        company_name="Ollie",
        role_title="Data Analyst, Retention",
        job_description=OLLIE_ANALYTICS_JOB_DESCRIPTION,
        resume_text=OLLIE_RESUME_TEXT,
    )
    assert_true(
        set(variants) >= {"30_second", "60_second", "90_second"},
        f"pitch_variants() should return the expected keys; got {tuple(variants)}",
    )
    for key in ("30_second", "60_second", "90_second"):
        assert_true(
            len(variants[key].strip()) >= 40,
            f"pitch_variants()['{key}'] should be a populated spoken answer; got {variants[key]!r}",
        )
    assert_true(
        "Ollie" in variants["60_second"] or "Data Analyst, Retention" in variants["60_second"],
        f"pitch_variants()['60_second'] should pull in company or role-specific framing; got {variants['60_second']}",
    )
    assert_true(
        "What keeps me engaged" in variants["90_second"] or "What matters to me" in variants["90_second"] or "I like the discovery side" in variants["90_second"],
        f"pitch_variants()['90_second'] should include the human-motivation layer; got {variants['90_second']}",
    )


def test_human_motivation_sentence_has_lane_coverage(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    for lane_name, lane_jd in LANE_JOB_DESCRIPTIONS.items():
        profile = build_resume.job_problem_profile(lane_jd, OLLIE_RESUME_TEXT)
        sentence = build_interview_cheat_sheet.human_motivation_sentence(profile)
        assert_true(
            len(sentence.strip()) >= 25,
            f"human_motivation_sentence() should return a substantive sentence for {lane_name}; got {sentence!r}",
        )


def test_adjusted_profile_for_role_preserves_non_lane_fields(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    original_effective_lane_key = build_resume.effective_lane_key
    try:
        build_resume.effective_lane_key = lambda *args, **kwargs: "presales_solution"
        adjusted = build_interview_cheat_sheet.adjusted_profile_for_role(
            profile,
            "Senior Solutions Architect",
            PROCORE_JOB_DESCRIPTION,
        )
    finally:
        build_resume.effective_lane_key = original_effective_lane_key

    assert_true(
        adjusted.primary_lane == "presales_solution",
        f"adjusted_profile_for_role() should update the lane when the effective lane shifts; got {adjusted.primary_lane!r}",
    )
    assert_true(
        adjusted.direct_matches == profile.direct_matches
        and adjusted.adjacent_matches == profile.adjacent_matches
        and adjusted.unsupported_requirements == profile.unsupported_requirements
        and adjusted.safe_terms == profile.safe_terms,
        "adjusted_profile_for_role() should preserve all non-lane JobProblemProfile fields when the lane changes",
    )


def test_story_answer_parts_preserve_spoken_alias(build_detailed_interview_guide: object) -> None:
    answer = build_detailed_interview_guide.StoryAnswerParts(
        full="I led the cleanup and stabilized the rollout.",
        meat_first="I led the cleanup.",
        stretch_modules=(("Proof", "I rebuilt the validation path."),),
        alternate="The point is that I made the risk visible early.",
        pushback_branches=(("What was yours?", "I owned the structure and validation path."),),
        coaching_note="Coaching note: keep the close direct.",
    )
    assert_true(
        answer.spoken == answer.full,
        f"StoryAnswerParts.spoken should remain a read-only alias of full; got spoken={answer.spoken!r}, full={answer.full!r}",
    )


def test_story_answer_constructors_use_full_keyword() -> None:
    source = Path("scripts/build_detailed_interview_guide.py").read_text(encoding="utf-8")
    assert_true(
        "StoryAnswerParts(spoken=" not in source,
        "build_detailed_interview_guide.py should migrate every StoryAnswerParts constructor to full=",
    )


def test_extended_tmay_sections_build_time_ladder(
    build_resume: object,
    build_detailed_interview_guide: object,
) -> None:
    profile = build_resume.job_problem_profile(OLLIE_ANALYTICS_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    sections = build_detailed_interview_guide.build_extended_tmay_sections(
        profile,
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
    )
    labels = [item.label for item in sections]
    assert_true(
        labels == ["20 to 30 seconds", "60 to 90 seconds", "2 minutes", "4 minutes"],
        f"build_extended_tmay_sections() should return the four-rung time ladder; got {labels}",
    )
    assert_true(
        sections[0].script.startswith("My background is"),
        f"The anchor TMAY script should start with the direct claim sentence; got {sections[0].script!r}",
    )
    two_minute_modules = {label for label, _text in sections[2].modules}
    assert_true(
        "Why this work" in two_minute_modules and "Why this role" in two_minute_modules,
        f"The longer TMAY rungs should add human and role-bridge modules; got {sections[2].modules}",
    )


def test_extended_tmay_sections_use_module_level_superset(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    sparse_story = build_interview_cheat_sheet.StoryCard(
        title="Sparse Story",
        story_types=("Analysis and Decision",),
        hook="an implementation was at risk of missing key validation steps",
        takeaways=("validate early", "keep owners aligned", "protect the timeline"),
        evidence="I built the checkpoint plan and clarified ownership before launch",
        level3_trait="",
        result="",
        outcome="clearer validation rhythm",
        evidence_terms=("implementation", "validation"),
        signals=("implementation", "validation"),
    )
    original_hero_stories = build_detailed_interview_guide.cheat.hero_stories
    try:
        build_detailed_interview_guide.cheat.hero_stories = lambda *args, **kwargs: [sparse_story]
        sections = build_detailed_interview_guide.build_extended_tmay_sections(
            profile,
            "Smoke Test Systems",
            "Implementation Consultant",
            DUMMY_JOB_DESCRIPTION,
            OLLIE_RESUME_TEXT,
        )
    finally:
        build_detailed_interview_guide.cheat.hero_stories = original_hero_stories

    module_sets = [set(label for label, _text in section.modules) for section in sections]
    assert_true(
        module_sets[0] < module_sets[1] < module_sets[2] < module_sets[3],
        f"Each longer TMAY rung should be a strict module-level superset of the shorter rung; got {module_sets}",
    )
    proof_text = next(text for label, text in sections[0].modules if label == "Proof")
    assert_true(
        "What I noticed early was" not in proof_text and "The result was" not in proof_text,
        f"Sparse TMAY proof modules should skip empty fragments instead of inserting placeholders; got {proof_text!r}",
    )
    assert_true(
        build_interview_cheat_sheet.lower_clause(sparse_story.hook) in proof_text and "I built the checkpoint plan" in proof_text,
        f"TMAY proof should keep the populated story fragments; got {proof_text!r}",
    )


def test_behavioral_answer_scripts_empty_story_guard(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    answers = build_interview_cheat_sheet.behavioral_answer_scripts(profile, [], DUMMY_JOB_DESCRIPTION)
    assert_true(answers == [], f"behavioral_answer_scripts() should return [] when no stories are available; got {answers!r}")


def test_story_sample_answer_separates_coaching_note(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    card = build_interview_cheat_sheet.StoryCard(
        title="Cross-Functional Rescue",
        story_types=("Challenge and Failure", "Teamwork"),
        hook="a launch path had unclear ownership across teams",
        takeaways=("clarify the owner", "stabilize the workflow", "protect adoption"),
        evidence="I aligned the owners, rebuilt the validation path, and gave leaders a clearer checkpoint plan",
        level3_trait="the risk was less technical than it first looked because ownership was fragmented",
        result="the launch stabilized and the team had a repeatable checkpoint structure",
        outcome="stabilized delivery",
        evidence_terms=("implementation", "stakeholder alignment"),
        signals=("implementation", "alignment"),
    )
    answer = build_detailed_interview_guide.story_sample_answer(
        card,
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        "",
        OLLIE_RESUME_TEXT,
    )
    assert_true(
        "Human layer:" not in answer.spoken and "The role-specific bridge is:" not in answer.spoken,
        f"story_sample_answer().spoken should not include inline coaching labels; got {answer.spoken!r}",
    )
    assert_true(
        answer.coaching_note.startswith("Coaching note: ")
        and "Human layer to weave in naturally:" in answer.coaching_note
        and "Role bridge:" in answer.coaching_note,
        f"story_sample_answer() should move coaching guidance into a separate note; got {answer.coaching_note!r}",
    )


def test_story_sample_answer_reuses_claim_sentence_in_full(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    card = build_interview_cheat_sheet.StoryCard(
        title="Delivery Recovery",
        story_types=("Challenge and Failure", "Individual Achievement"),
        hook="a launch path had unclear ownership and too many hidden dependencies",
        takeaways=("clarify the owner", "surface the hidden dependency", "protect adoption"),
        evidence="I rebuilt the validation path, reset owner checkpoints, and kept stakeholders aligned through release",
        level3_trait="the first warning sign was that people could describe the issue but not the actual owner or checkpoint",
        result="the launch stabilized and the team had a repeatable checkpoint structure",
        outcome="stabilized delivery",
        evidence_terms=("implementation", "stakeholder alignment"),
        signals=("implementation", "alignment"),
    )
    answer = build_detailed_interview_guide.story_sample_answer(
        card,
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        "",
        OLLIE_RESUME_TEXT,
    )
    assert_true(
        build_detailed_interview_guide.first_sentence(answer.full) == build_detailed_interview_guide.first_sentence(answer.meat_first),
        f"story_sample_answer().full should begin with the exact same claim sentence used in meat_first; got meat_first={answer.meat_first!r}, full={answer.full!r}",
    )
    assert_true(
        answer.stretch_modules and answer.pushback_branches and answer.alternate,
        f"story_sample_answer() should populate the layered script fields; got {answer}",
    )


def test_delivery_validator_only_scans_scripted_strings(build_detailed_interview_guide: object) -> None:
    clean_answer = build_detailed_interview_guide.StoryAnswerParts(
        full="I led the rollout cleanup and stabilized the launch.",
        meat_first="I led the rollout cleanup.",
        stretch_modules=(("Proof", "I rebuilt the validation path."),),
        alternate="The key was making ownership visible early.",
        pushback_branches=(("What was yours?", "I owned the structure and the checkpoint rhythm."),),
    )
    build_detailed_interview_guide.validate_delivery_principles(clean_answer, label="clean scripted answer")
    hedged_answer = build_detailed_interview_guide.StoryAnswerParts(
        full="I guess the closest example is a rollout cleanup.",
        meat_first="I guess the closest example is a rollout cleanup.",
    )
    try:
        build_detailed_interview_guide.validate_delivery_principles(hedged_answer, label="hedged scripted answer")
    except SystemExit:
        pass
    else:
        raise AssertionError("validate_delivery_principles() should fail on banned hedge language inside scripted answer text")


def test_scripted_answer_validator_rejects_unsupported_metrics(build_detailed_interview_guide: object) -> None:
    answer = build_detailed_interview_guide.StoryAnswerParts(
        full="I rebuilt the workflow and reduced manual work by 99%.",
        meat_first="I rebuilt the workflow.",
    )
    try:
        build_detailed_interview_guide.validate_scripted_answer(
            answer,
            label="unsupported metric story",
            allowed_text="I rebuilt the workflow and reduced manual work by 22%.",
        )
    except SystemExit:
        pass
    else:
        raise AssertionError("validate_scripted_answer() should fail when a script introduces an unsupported metric token")


def test_ai_customer_work_answer_uses_confirmed_qualitative_story(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    answer = build_detailed_interview_guide.ai_customer_work_answer()
    combined = " ".join(
        [
            answer.full,
            answer.alternate,
            *[f"{label} {text}" for label, text in answer.stretch_modules],
            *[f"{concern} {text}" for concern, text in answer.pushback_branches],
        ]
    )
    assert_true(
        "ChatGPT Codex" in combined
        and "Claude" in combined
        and "Gemini" in combined
        and "Copilot" in combined
        and "Aptean Import Wizard" in combined
        and "GL-account" in combined
        and "sandbox" in combined
        and "that sounds like scripting more than ai" in combined.lower(),
        f"ai_customer_work_answer() should use the confirmed qualitative workflow story and pushback branch; got {combined!r}",
    )
    keyword_answers = build_interview_cheat_sheet.keyword_ready_answers(
        build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT),
        "AI workflow, automation, and customer enablement role.",
        build_interview_cheat_sheet.supported_story_bank(OLLIE_RESUME_TEXT),
    )
    ai_answer = next((item.answer for item in keyword_answers if item.prompt.startswith("AI-Enabled Customer Work:")), "")
    assert_true(
        "ChatGPT Codex" in ai_answer and "Aptean Import Wizard" in ai_answer,
        f"keyword_ready_answers() should reuse the confirmed AI workflow story; got {ai_answer!r}",
    )


def test_story_sample_answer_does_not_call_dead_star_selection(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    card = build_interview_cheat_sheet.StoryCard(
        title="Implementation Recovery",
        story_types=("Challenge and Failure",),
        hook="an implementation was at risk of slipping",
        takeaways=("clarified the real blocker", "reset ownership", "protected the launch"),
        evidence="I rebuilt the validation path and kept stakeholders aligned through release",
        level3_trait="the issue pattern was visible before the launch risk fully surfaced",
        result="protected the launch timeline",
        outcome="kept the rollout on track",
        evidence_terms=("implementation", "validation"),
        signals=("implementation", "stakeholder"),
    )
    original_selection = build_detailed_interview_guide.cheat.answer_framework_selection
    try:
        build_detailed_interview_guide.cheat.answer_framework_selection = (
            lambda *_args, **_kwargs: (_ for _ in ()).throw(SmokeFailure("dead STAR selection branch should not run"))
        )
        answer = build_detailed_interview_guide.story_sample_answer(
            card,
            profile,
            "Smoke Test Systems",
            "Implementation Consultant",
            DUMMY_JOB_DESCRIPTION,
            "",
            OLLIE_RESUME_TEXT,
        )
    finally:
        build_detailed_interview_guide.cheat.answer_framework_selection = original_selection
    assert_true(
        build_interview_cheat_sheet.lower_clause(card.hook) in answer.spoken,
        f"story_sample_answer() should still return a spoken answer after the dead STAR branch is removed; got {answer}",
    )


def test_pitch_variants_add_fifteen_second_claim_and_validate_opening(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(PROCORE_JOB_DESCRIPTION, PROCORE_RESUME_TEXT)
    variants = build_interview_cheat_sheet.pitch_variants(
        profile,
        company_name="Procore",
        role_title="Senior Solutions Architect",
        job_description=PROCORE_JOB_DESCRIPTION,
        resume_text=PROCORE_RESUME_TEXT,
    )
    assert_true(
        "15_second" in variants and variants["15_second"].startswith("My background is"),
        f"pitch_variants() should use the natural identity-first opening; got {variants}",
    )
    assert_true(
        build_interview_cheat_sheet.claim_then_prove_issues(variants["60_second"]) == [],
        f"pitch_variants()['60_second'] should pass the claim-then-prove validator; got {variants['60_second']!r}",
    )


def test_claim_then_prove_validator_catches_delayed_openings(build_interview_cheat_sheet: object) -> None:
    issues = build_interview_cheat_sheet.claim_then_prove_issues(
        "Let me give a little background first before I answer directly. I led the reporting rebuild across 200+ dashboards.",
        require_proof=False,
    )
    assert_true(
        any("delays the point" in issue for issue in issues),
        f"claim_then_prove_issues() should catch delayed openings; got {issues}",
    )


def test_natural_voice_validation_and_answer_budgets(build_interview_cheat_sheet: object) -> None:
    identity_answer = (
        "My background is in systems and implementation work. "
        "I supported 80+ manufacturing clients through requirements, configuration, testing, and go-live."
    )
    assert_true(
        build_interview_cheat_sheet.claim_then_prove_issues(identity_answer) == [],
        "identity-first answers should pass claim-then-prove validation when supported proof follows",
    )
    filler_issues = build_interview_cheat_sheet.claim_then_prove_issues(
        "Sure. I supported 80+ manufacturing clients through implementation.",
    )
    assert_true(
        any("filler" in issue for issue in filler_issues),
        f"standalone conversational filler should be rejected; got {filler_issues}",
    )
    product = build_interview_cheat_sheet.product_learning_answer()
    metrics = build_interview_cheat_sheet.spoken_answer_metrics(product)
    assert_true(
        70 <= metrics["words"] <= 105 and 15 <= metrics["average_sentence_words"] <= 22,
        f"product-learning answer should meet its spoken budget; got {metrics}",
    )


def test_natural_voice_question_coverage(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    stories = build_interview_cheat_sheet.expanded_story_bank()
    answers = build_interview_cheat_sheet.common_interview_answers(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        stories,
        resume_text=OLLIE_RESUME_TEXT,
    )
    answer_map = {item.prompt: item.answer for item in answers}
    expected = {
        "Tell me about your education and credentials.",
        "Walk me through your career.",
        "How would you learn a new product quickly?",
        "What are your compensation expectations and practical constraints?",
    }
    assert_true(expected <= set(answer_map), f"common answer bank is missing natural-voice questions: {expected - set(answer_map)}")
    walkthrough = answer_map["Walk me through your career."]
    assert_true(
        "From there" in walkthrough and "Most recently" in walkthrough and "Tell me about yourself" not in walkthrough,
        f"career walkthrough should be chronological and distinct from TMAY; got {walkthrough}",
    )
    story_titles = {story.title for story in stories}
    assert_true(
        not any(title in answer_map["What is a gap we should know about?"] for title in story_titles),
        "spoken gap answer should not say an internal story title aloud",
    )
    modernization_story = next(story for story in stories if story.title == "13-month modernization complexity")
    modernization_answer = build_interview_cheat_sheet.spoken_caar_answer(modernization_story, profile)
    assert_true(
        "show what was done" not in modernization_answer.lower()
        and "show what was noticed" not in modernization_answer.lower(),
        f"spoken story answers should not expose rehearsal labels; got {modernization_answer}",
    )


def test_keyword_reference_uses_natural_story_intro(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    story = build_interview_cheat_sheet.expanded_story_bank()[0]
    answer = build_detailed_interview_guide.keyword_sample_answer(
        "operations",
        profile,
        "Acme Systems",
        "Implementation Consultant",
        story,
    )
    assert_true(
        f"my best example is {story.title}".lower() not in answer.lower()
        and "my best example" not in answer.lower(),
        f"keyword sample answers should use the shared spoken foundation without announcing internal story titles; got {answer}",
    )
    assert_true(
        ". I " in answer or ". The " in answer,
        f"keyword sample answers should preserve sentence boundaries between the observation and action; got {answer}",
    )


def test_industry_depth_and_company_scoped_logistics(build_resume: object, build_interview_cheat_sheet: object) -> None:
    jd = DUMMY_JOB_DESCRIPTION + " Manufacturing controls, electrical schematics, PLC, SCADA, shop floor, and BOM workflows."
    profile = build_resume.job_problem_profile(jd, OLLIE_RESUME_TEXT)
    answers = build_interview_cheat_sheet.keyword_ready_answers(
        profile,
        jd,
        build_interview_cheat_sheet.expanded_story_bank(),
        resume_text=OLLIE_RESUME_TEXT,
    )
    answer_map = {item.prompt: item.answer for item in answers}
    manufacturing = next((answer for prompt, answer in answer_map.items() if prompt.startswith("Manufacturing Experience:")), "")
    technical = next((answer for prompt, answer in answer_map.items() if prompt.startswith("Technical Depth Boundary:")), "")
    assert_true(
        "inventory accuracy" in manufacturing and "plant engineer" in manufacturing,
        f"manufacturing answer should connect process, role, impact, and an honest boundary; got {manufacturing}",
    )
    assert_true(
        "business analytics and ERP configuration" in technical
        and "not PLC programming or code writing" in technical
        and "barcode scanners tied into the software layer" in technical,
        f"technical-depth answer should state the boundary and partnership model; got {technical}",
    )
    supplied = (
        "Acme Systems interview notes\n"
        "I am targeting $100K to $110K, and I am flexible based on the full package.\n"
        "I am comfortable with about 2 days per month of travel.\n"
        "I can move quickly if the process is a mutual fit."
    )
    scoped = build_interview_cheat_sheet.compensation_logistics_answer(supplied)
    generic = build_interview_cheat_sheet.compensation_logistics_answer("")
    assert_true(
        "$100K to $110K" in scoped and "2 days per month" in scoped and "move quickly" in scoped,
        f"company-scoped notes should supply supported logistics details; got {scoped}",
    )
    assert_true(
        "$100K" not in generic and "approved range" in generic and "[" not in generic,
        f"missing notes should produce a complete nonnumeric answer without placeholders; got {generic}",
    )


def test_first_90_day_plan_reuses_shared_stage_source(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    plan = build_interview_cheat_sheet.first_90_day_plan(profile)
    ladder = build_interview_cheat_sheet.first_90_day_approach(profile)
    assert_true(
        len(plan.stages) == 3 and len(ladder) == 3,
        f"first_90_day_plan() should expose three stage blocks and first_90_day_approach() should render the same three; got plan={plan.stages} ladder={ladder}",
    )
    assert_true(
        plan.answer.startswith("In the first 30 days") and plan.validation_question == "Does that match what the team expects at this level?",
        f"first_90_day_plan() should build the spoken answer plus the answer-first validation question; got {plan}",
    )
    for (label, text), rendered in zip(plan.stages, ladder):
        assert_true(
            label in rendered and text in rendered,
            f"first_90_day_approach() should render the same stage lines exposed by first_90_day_plan(); got stage={(label, text)} rendered={rendered!r}",
        )


def test_application_answers_use_written_confirmation_and_automation_boundary(question_prep: object) -> None:
    jd = (
        "Company: Acme Automation\n"
        "Role: Solution Consultant\n"
        "This warehouse automation role includes PLC, controls, conveyors, WCS, and stakeholder communication."
    )
    software = question_prep.software_inventory_answer_for_job(jd, OLLIE_RESUME_TEXT)
    communication = question_prep.communication_answer(
        "stakeholder communication",
        job_description=jd,
        resume_text=OLLIE_RESUME_TEXT,
    )
    assert_true(
        "business analytics and ERP configuration" in software
        and "not PLC programming or code writing" in software
        and "least-privilege access controls" in software
        and "barcode scanners tied into the software layer" in software
        and "access-control layer" in software,
        f"software_inventory_answer_for_job() should self-categorize technical depth for automation-heavy roles; got {software!r}",
    )
    assert_true(
        "confirm it in writing" in communication and "functional requirements" in communication,
        f"communication_answer() should lead with written-confirmation discipline for communication prompts; got {communication!r}",
    )


def test_detailed_guide_stage_patterns_and_debrief_overlay(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    from docx import Document

    jd = (
        "Company: Acme Automation\n"
        "Role: Solution Consultant\n"
        "This warehouse automation role includes PLC, controls, conveyors, WCS, and stakeholder communication."
    )
    profile = build_resume.job_problem_profile(jd, OLLIE_RESUME_TEXT)
    stories = build_interview_cheat_sheet.expanded_story_bank()

    hr_document = Document()
    build_detailed_interview_guide.add_hr_screen_prep_section(
        hr_document,
        profile,
        "Acme Automation",
        "Solution Consultant",
        jd,
        OLLIE_RESUME_TEXT,
        "",
        "",
        stories,
    )
    hr_text = build_detailed_interview_guide.document_text(hr_document)
    assert_true(
        "Company snapshot" in hr_text
        and "Recruiter checklist" in hr_text
        and "What would you do in the first 90 days?" in hr_text,
        f"HR screen prep should surface the reusable company snapshot, recruiter checklist, and staged recruiter answer pairs; got {hr_text!r}",
    )

    hiring_document = Document()
    build_detailed_interview_guide.add_hiring_manager_prep_section(
        hiring_document,
        profile,
        "Acme Automation",
        "Solution Consultant",
        jd,
        stories,
    )
    hiring_text = build_detailed_interview_guide.document_text(hiring_document)
    assert_true(
        "First 90-day glance" in hiring_text
        and "Software-boundary answer" in hiring_text
        and "Hardware-gap bridge" in hiring_text
        and "Four anchor stories for this round" in hiring_text,
        f"Hiring-manager prep should surface the shared 90-day answer, the software and hardware boundary scripts, and the story anchor table; got {hiring_text!r}",
    )

    overlay_document = Document()
    rendered = build_detailed_interview_guide.add_debrief_overlay_section(
        overlay_document,
        (
            {
                "unexpected_questions": ["What would you do in the first 90 days?"],
                "feedback_received": ["Keep answers tighter."],
                "review_analysis": {
                    "positioning_diagnosis": {
                        "headline": "Qualified but too long.",
                        "reasons": ["Lead with the answer sooner."],
                    },
                    "answer_strategy": {
                        "delivery_shifts": ["Lead with the answer in sentence one."],
                    },
                    "answer_assets": {
                        "role_language_lines": ["Translate customer inputs into recommendations."],
                        "company_signal_lines": ["Discovery quality matters early."],
                    },
                },
            },
        ),
    )
    overlay_text = build_detailed_interview_guide.document_text(overlay_document)
    assert_true(rendered, "add_debrief_overlay_section() should render when a structured round record exists.")
    overlay_lower = overlay_text.lower()
    assert_true(
        "debrief-to-prep overlay" in overlay_lower
        and "interviewer language to reuse" in overlay_lower
        and "unexpected questions to drill" in overlay_lower
        and "safe repeatable company terms" in overlay_lower,
        f"Debrief overlay should render the reusable prep surfaces from structured records; got {overlay_text!r}",
    )


def test_interview_companion_documents_reuse_shared_stage_sources(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    jd = (
        "Company: Acme Automation\n"
        "Role: Solution Consultant\n"
        "This warehouse automation role includes PLC, controls, conveyors, WCS, and stakeholder communication."
    )
    profile = build_interview_cheat_sheet.adjusted_profile_for_role(
        build_resume.job_problem_profile(jd, OLLIE_RESUME_TEXT),
        "Solution Consultant",
        jd,
    )
    stories = build_interview_cheat_sheet.expanded_story_bank()

    recruiter_document = build_detailed_interview_guide.build_recruiter_screen_companion_document(
        profile,
        "Acme Automation",
        "Solution Consultant",
        jd,
        OLLIE_RESUME_TEXT,
        "",
        "",
        stories,
    )
    recruiter_text = build_detailed_interview_guide.document_text(recruiter_document)
    recruiter_upper = recruiter_text.upper()
    assert_true(
        "RECRUITER SCREEN PREP" in recruiter_upper
        and "KNOW THE COMPANY" in recruiter_upper
        and "RECRUITER CHECKLIST" in recruiter_upper
        and "What would you do in the first 90 days?" in recruiter_text,
        f"Recruiter-screen companion should stay tied to the shared recruiter bank and checklist content; got {recruiter_text!r}",
    )

    plan_document = build_detailed_interview_guide.build_first_90_day_one_pager_document(
        profile,
        "Acme Automation",
        "Solution Consultant",
        jd,
    )
    plan_text = build_detailed_interview_guide.document_text(plan_document)
    plan = build_interview_cheat_sheet.first_90_day_plan(profile)
    assert_true(
        plan.validation_question in plan_text
        and all(text in plan_text for _label, text in plan.stages),
        f"90-day companion should reuse the shared stage plan and validation question; got {plan_text!r}",
    )

    addendum_document = build_detailed_interview_guide.build_debrief_addendum_document(
        "Acme Automation",
        "Solution Consultant",
        (
            {
                "unexpected_questions": ["What would you do in the first 90 days?"],
                "feedback_received": ["Keep answers tighter."],
                "review_analysis": {
                    "positioning_diagnosis": {
                        "headline": "Qualified but too long.",
                        "reasons": ["Lead with the answer sooner."],
                    },
                    "answer_strategy": {
                        "delivery_shifts": ["Lead with the answer in sentence one."],
                    },
                    "answer_assets": {
                        "role_language_lines": ["Translate customer inputs into recommendations."],
                        "company_signal_lines": ["Discovery quality matters early."],
                    },
                },
            },
        ),
    )
    assert_true(addendum_document is not None, "Debrief companion should render when structured debrief signals exist.")
    addendum_text = build_detailed_interview_guide.document_text(addendum_document)
    addendum_upper = addendum_text.upper()
    assert_true(
        "DEBRIEF PREP ADDENDUM" in addendum_upper
        and "DEBRIEF-TO-PREP OVERLAY" in addendum_upper
        and "FIX CARDS" in addendum_upper,
        f"Debrief companion should stay tied to the shared overlay sections; got {addendum_text!r}",
    )


def test_behavioral_answer_scripts_use_spoken_answers_for_core_prompts(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    stories = build_interview_cheat_sheet.expanded_story_bank()
    answers = {
        item.prompt: item.answer
        for item in build_interview_cheat_sheet.behavioral_answer_scripts(
            profile,
            stories,
            DUMMY_JOB_DESCRIPTION,
            "Smoke Test Systems",
            "Implementation Consultant",
            OLLIE_RESUME_TEXT,
        )
    }
    assert_true(
        not answers["Difficult coworker or client"].startswith("Use ")
        and not answers["Process improvement: inefficient process"].startswith("Use ")
        and not answers["Project management in process engineering"].startswith("Answer like"),
        f"behavioral_answer_scripts() should return spoken answers instead of instruction fragments; got {answers}",
    )
    assert_true(
        answers["Why should we not hire you?"].startswith("If you need someone"),
        f"behavioral_answer_scripts() should use an honest bridge answer for inverted-fit questions; got {answers['Why should we not hire you?']!r}",
    )


def test_story_adaptation_and_pre_interview_routine_helpers(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    card = build_interview_cheat_sheet.expanded_story_bank()[0]
    routine = build_interview_cheat_sheet.pre_interview_routine_lines(
        "Implementation Consultant",
        (
            {
                "performance_review": {
                    "coaching_signals": [
                        {"key": "rambling", "label": "Rambling", "detail": "Shorten answers."},
                        {"key": "filler_restarts", "label": "Filler", "detail": "Cut restarts."},
                    ]
                }
            },
        ),
    )
    adaptation = build_interview_cheat_sheet.story_adaptation_drill_lines(
        card,
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    delivery = build_interview_cheat_sheet.story_delivery_note_lines(
        card,
        profile,
        (
            {
                "performance_review": {
                    "coaching_signals": [
                        {"key": "executive_presence", "label": "Executive presence", "detail": "Be more declarative."},
                    ]
                }
            },
        ),
    )
    anchor = build_interview_cheat_sheet.story_anchor_fact_line(card, profile)
    spoken = build_interview_cheat_sheet.spoken_story_answer(
        card,
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        any("longer exhales" in line for line in routine) and any("adrenaline" in line.lower() for line in routine),
        f"pre_interview_routine_lines() should include breathing and arousal reframing; got {routine}",
    )
    assert_true(
        any("personally owned" in line.lower() or "your role first" in line.lower() for line in adaptation)
        and any("role bridge" in line.lower() or "why this matters" in line.lower() for line in adaptation),
        f"story_adaptation_drill_lines() should create follow-up prompts for ownership and role fit; got {adaptation}",
    )
    assert_true(
        anchor.startswith("Anchor facts to keep fixed:")
        and any("declarative" in line.lower() or "one ownership line" in line.lower() for line in delivery),
        f"story anchor and delivery helpers should surface fixed facts and delivery cues; got anchor={anchor!r}, delivery={delivery}",
    )
    assert_true(
        "Hook:" not in spoken
        and "CONTEXT" not in spoken
        and build_interview_cheat_sheet.lower_clause(card.hook) in spoken,
        f"spoken_story_answer() should return a rehearsable spoken answer instead of framework labels; got {spoken!r}",
    )


def test_spoken_sentence_split_preserves_leading_word_characters() -> None:
    import prose_engine

    source = (
        "I led a modernization engagement where requirements revealed significant hardware upgrades, "
        "costing tens of thousands of dollars on their side, before implementation could begin and the team could proceed safely."
    )
    repaired = prose_engine.spoken_register(source).text
    assert_true(
        "dollars" in repaired.lower()
        and "ollars" not in repaired.replace("dollars", "").lower()
        and "of. dollars" not in repaired.lower(),
        f"spoken sentence splitting must preserve the first character of ordinary words; got {repaired!r}",
    )


def test_nonconverged_spoken_repairs_are_collected() -> None:
    import prose_engine

    nested = "I managed scope, risk, testing, deployment, including data, security, or training and support."
    with prose_engine.collect_spoken_repair_issues() as issues:
        outcome = prose_engine.spoken_register(nested)
    assert_true(
        not outcome.converged
        and any("PROSE_NESTED_LIST" in issue for issue in issues),
        f"Non-converged spoken repairs must expose rule IDs to the document DRAFT gate; got {outcome}, {issues}",
    )


def test_federal_version_control_scope_gate_is_unconditional(build_federal_resume: object) -> None:
    aptean = "Used Azure DevOps/TFS and Git/GitHub for repositories and release tracking."
    build_federal_resume.assert_version_control_employer_scope(aptean, "No version tools are claimed here.")
    for leaked_term in ("Git", "GitHub", "TFS", "Azure DevOps", "version control"):
        try:
            build_federal_resume.assert_version_control_employer_scope(
                aptean,
                f"East West Manufacturing used {leaked_term} for delivery work.",
            )
        except SystemExit:
            continue
        raise AssertionError(f"Bare non-Aptean version-control term should fail unconditionally: {leaked_term}")


def test_foundant_summary_uses_human_close(build_resume: object) -> None:
    foundant_job_description = """
Company: Foundant
Role: Associate Implementation Consultant
Lead implementation discovery, data migration, testing, role-based training, quality control, go-live readiness, and customer adoption.
"""
    summary = build_resume.build_problem_first_summary(
        foundant_job_description,
        OLLIE_RESUME_TEXT,
        variant_index=0,
    )
    assert_true(
        "Best in implementations needing" not in summary
        and "keeping operating continuity as visible" not in summary
        and "Brings the most value when" in summary,
        f"Commercial implementation summary should use the human-register closing sentence; got {summary!r}",
    )


def test_commercial_resume_model_provenance_and_render(
    commercial_resume_model: object,
    build_resume: object,
) -> None:
    with TemporaryDirectory(prefix="commercial_model_") as temp_name:
        work_dir = Path(temp_name) / "work"
        build_resume.unpack_docx(build_resume.IMPLEMENTATION_RESUME, work_dir)
        document_xml = work_dir / "word" / "document.xml"
        before = commercial_resume_model.build_content_model(
            build_resume.IMPLEMENTATION_RESUME,
            document_xml,
            document_xml,
        )
        commercial_resume_model.render_content_model(document_xml, before)
        after = commercial_resume_model.build_content_model(
            build_resume.IMPLEMENTATION_RESUME,
            document_xml,
            document_xml,
        )
        manifest = commercial_resume_model.write_manifest(before, Path(temp_name) / "model.json")
    assert_true(
        before.content_hash == after.content_hash and manifest.name == "model.json",
        "Commercial model render should preserve the authoritative modeled content hash.",
    )
    assert_true(
        bool(before.summary)
        and bool(before.roles)
        and all(item.provenance for role in before.roles for item in (*role.summaries, *role.bullets))
        and all(
            ref.employer == role.employer
            for role in before.roles
            for item in (*role.summaries, *role.bullets)
            for ref in item.provenance
        ),
        "Every modeled summary/bullet must carry same-employer approved-source provenance.",
    )


def test_application_question_pairing_detects_jd_swap() -> None:
    import job_context_archive

    original_archive_root = job_context_archive.SCRATCH_JD_LIBRARY
    original_index_path = job_context_archive.INDEX_PATH
    original_sync_complete = job_context_archive._SYNC_COMPLETE
    try:
        with TemporaryDirectory(prefix="question_pairing_") as temp_name:
            archive_root = Path(temp_name) / "jd_library"
            job_context_archive.SCRATCH_JD_LIBRARY = archive_root
            job_context_archive.INDEX_PATH = archive_root / "index.csv"
            job_context_archive._SYNC_COMPLETE = True
            questions = "Why do you want to join Acme?\n\nDescribe your implementation experience."
            job_context_archive.archive_texts(
                job_description_text="Company: Acme\nRole: Implementation Consultant\nLead customer implementations.",
                application_questions_text=questions,
                source_command="smoke",
                sync_legacy=False,
            )
            issue = job_context_archive.application_question_pairing_issue(
                "Company: NewCo\nRole: Customer Success Manager\nOwn adoption and retention.",
                questions,
            )
    finally:
        job_context_archive.SCRATCH_JD_LIBRARY = original_archive_root
        job_context_archive.INDEX_PATH = original_index_path
        job_context_archive._SYNC_COMPLETE = original_sync_complete
    assert_true(
        "previously archived with Acme" in issue and "different content hash" in issue,
        f"A reused question hash with a changed JD hash should trigger the swap guard; got {issue!r}",
    )
    workflow_source = Path("scripts/run_resume_workflow.py").read_text(encoding="utf-8")
    assert_true(
        "application_question_pairing_issue" in workflow_source and "Application-question/JD swap check" in workflow_source,
        "The commercial workflow runner must stop at the JD-swap boundary before building outputs.",
    )


def test_repo_guidance_prefers_rehearsed_foundation() -> None:
    agents_text = Path("AGENTS.md").read_text(encoding="utf-8")
    rules_text = Path(".context/RULES_FOR_CLAUDE.md").read_text(encoding="utf-8")
    claude_text = Path("CLAUDE.md").read_text(encoding="utf-8")
    assert_true(
        "rehearsed-foundation study method" in agents_text and "scripts to memorize" not in agents_text,
        "AGENTS.md should prefer rehearsed-foundation language over memorize-the-script guidance",
    )
    assert_true(
        "Build a rehearsed foundation" in rules_text,
        ".context/RULES_FOR_CLAUDE.md should describe the rehearsed interview-prep philosophy",
    )
    assert_true(
        "rehearsable TMAY" in claude_text,
        "CLAUDE.md should describe the interview packet in rehearsable-answer terms",
    )


def test_pitch_helpers_handle_missing_cover_letter_pitch_parts(
    build_resume: object,
    build_interview_cheat_sheet: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    card = build_interview_cheat_sheet.StoryCard(
        title="Fallback Story",
        story_types=("Analysis and Decision",),
        hook="leaders needed clearer reporting before launch decisions",
        takeaways=("clarify the question", "build trust in the data", "make the next step obvious"),
        evidence="I built a cleaner reporting view and aligned the decision path with stakeholders",
        level3_trait="leaders were debating inputs instead of the decision",
        result="the team moved faster with clearer visibility",
        outcome="clearer decisions",
        evidence_terms=("reporting", "analytics"),
        signals=("reporting", "analytics"),
    )
    variants = build_interview_cheat_sheet.pitch_variants(
        profile,
        company_name="Smoke Test Systems",
        role_title="Implementation Consultant",
        job_description=DUMMY_JOB_DESCRIPTION,
        resume_text=OLLIE_RESUME_TEXT,
    )
    reference_lines = build_interview_cheat_sheet.human_layer_reference_lines(
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
    )
    story_line = build_interview_cheat_sheet.story_human_connection_line(
        card,
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
    )

    assert_true(
        all(variants[key].strip() for key in ("30_second", "60_second", "90_second")),
        f"pitch_variants() should return populated interview-native answers without relying on cover-letter helpers; got {variants}",
    )
    assert_true(
        reference_lines
        and reference_lines[0].startswith("Human line to keep in reserve:")
        and len(reference_lines) >= 3,
        f"human_layer_reference_lines() should stay usable with the interview-native pitch parts; got {reference_lines}",
    )
    assert_true(
        "human side" in story_line.lower() or "what keeps me engaged" in story_line.lower(),
        f"story_human_connection_line() should still return a usable human-connection line after the cover-letter decoupling; got {story_line!r}",
    )


def test_read_the_room_opening(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    lines = build_interview_cheat_sheet.read_the_room_opening(
        profile,
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        len(lines) >= 4 and lines[0].startswith("Read the room before you sell"),
        f"read_the_room_opening() should return opening-five-minutes guidance; got {lines}",
    )


def test_interview_question_prep_sections_render(
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    from docx import Document

    cheat_recent = build_interview_cheat_sheet.question_prep.recent_interviewer_question_prep_items
    detailed_active = build_detailed_interview_guide.question_prep.active_application_question_responses
    detailed_recent = build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items
    cheat_load_prompts = build_interview_cheat_sheet.question_prep.load_active_application_prompts
    recent_lookup = build_interview_cheat_sheet.interview_context.recent_interviewer_questions
    try:
        supplied_response = build_interview_cheat_sheet.question_prep.QualificationsResponse(
            "We are looking for people who are personally and/or professionally passionate about AI. Please briefly explain how you have put it to work for you in either or both areas of your life.",
            "Professionally, I have used AI in reporting, documentation, and workflow support.",
        )
        supplied_recent = build_interview_cheat_sheet.question_prep.InterviewQuestionPrep(
            "Why are you passionate about AI?",
            "Lead with practical business use cases and concrete workflow examples.",
        )
        # strip_external_question_prompts_for_validation reads load_active_application_prompts
        # to know which prompts to remove before the AI-writing check. Patch it so the
        # injected question prompt is included, preventing a false-positive "passionate about" hit.
        build_interview_cheat_sheet.question_prep.load_active_application_prompts = lambda *_args, **_kwargs: (supplied_response.prompt,)
        build_interview_cheat_sheet.question_prep.recent_interviewer_question_prep_items = lambda *_args, **_kwargs: (supplied_recent,)
        build_detailed_interview_guide.question_prep.active_application_question_responses = lambda _job_description: (supplied_response,)
        build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items = lambda *_args, **_kwargs: (supplied_recent,)
        build_interview_cheat_sheet.interview_context.recent_interviewer_questions = (
            lambda *_args, **_kwargs: (supplied_recent.prompt,)
        )

        cheat_document = Document()
        detailed_document = Document()
        cheat_recent_rendered = build_interview_cheat_sheet.add_recent_interview_question_prep_section(
            cheat_document,
            DUMMY_JOB_DESCRIPTION,
            "Acme Systems",
            "Implementation Consultant",
        )
        detailed_rendered = build_detailed_interview_guide.add_application_question_prep_section(detailed_document, DUMMY_JOB_DESCRIPTION)
        detailed_recent_rendered = build_detailed_interview_guide.add_recent_interview_question_prep_section(
            detailed_document,
            DUMMY_JOB_DESCRIPTION,
            "Acme Systems",
            "Implementation Consultant",
        )
        cheat_text = "\n".join(paragraph.text.strip() for paragraph in cheat_document.paragraphs if paragraph.text.strip())
        # The detailed interview guide renders section/subsection headers and
        # Q&A answers as styled tables (see iter_all_paragraphs() in
        # build_detailed_interview_guide.py), so a plain document.paragraphs
        # walk would miss that text entirely.
        detailed_text = build_detailed_interview_guide.document_text(detailed_document)
        build_interview_cheat_sheet.validate_text(
            cheat_text,
            company_name="Acme Systems",
            role_title="Implementation Consultant",
        )
        build_detailed_interview_guide.validate_text(
            detailed_text,
            company_name="Acme Systems",
            role_title="Implementation Consultant",
        )

        build_interview_cheat_sheet.question_prep.recent_interviewer_question_prep_items = lambda *_args, **_kwargs: ()
        build_detailed_interview_guide.question_prep.active_application_question_responses = lambda _job_description: ()
        build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items = lambda *_args, **_kwargs: ()
        empty_cheat_document = Document()
        empty_detailed_document = Document()
        cheat_recent_empty = build_interview_cheat_sheet.add_recent_interview_question_prep_section(
            empty_cheat_document,
            DUMMY_JOB_DESCRIPTION,
            "Acme Systems",
            "Implementation Consultant",
        )
        detailed_empty = build_detailed_interview_guide.add_application_question_prep_section(empty_detailed_document, DUMMY_JOB_DESCRIPTION)
        detailed_recent_empty = build_detailed_interview_guide.add_recent_interview_question_prep_section(
            empty_detailed_document,
            DUMMY_JOB_DESCRIPTION,
            "Acme Systems",
            "Implementation Consultant",
        )
    finally:
        build_interview_cheat_sheet.question_prep.load_active_application_prompts = cheat_load_prompts
        build_interview_cheat_sheet.question_prep.recent_interviewer_question_prep_items = cheat_recent
        build_detailed_interview_guide.question_prep.active_application_question_responses = detailed_active
        build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items = detailed_recent
        build_interview_cheat_sheet.interview_context.recent_interviewer_questions = recent_lookup

    assert_true(
        cheat_recent_rendered
        and "RECENT INTERVIEW QUESTIONS TO BE READY FOR" in cheat_text
        and "APPLICATION / SUPPLEMENTAL QUESTIONS TO BE READY FOR" not in cheat_text,
        "Interview cheat sheet should stay focused on live-call prep by rendering recent interviewer questions without pulling in the application-question bank.",
    )
    assert_true(
        detailed_rendered
        and detailed_recent_rendered
        and "APPLICATION / SUPPLEMENTAL QUESTIONS TO BE READY FOR" in detailed_text
        and "RECENT INTERVIEW QUESTIONS TO BE READY FOR" in detailed_text,
        "Detailed interview guide helpers should render both new question-prep sections when shared data is available.",
    )
    assert_true(
        not cheat_recent_empty and not empty_cheat_document.paragraphs,
        "Interview cheat sheet should omit the recent-question section cleanly when no shared interview-prep data exists.",
    )
    assert_true(
        not detailed_empty and not detailed_recent_empty and not empty_detailed_document.paragraphs,
        "Detailed interview guide helpers should omit the new sections cleanly when no shared question-prep data exists.",
    )


def test_recent_interview_question_classification_and_factual_scripts(
    build_interview_cheat_sheet: object,
) -> None:
    question_prep = build_interview_cheat_sheet.question_prep
    role_end_prompt = "Why did your most recent role end?"
    product_prompt = "Can you walk me through the target audience and business value of our product offerings?"
    ai_prompt = "Have you worked at a SaaS or technology company that builds an AI product or service?"
    training_prompt = "Tell me about leading training strategies for enterprise system implementations. What was your role?"
    comprehension_prompt = "Does that make sense?"

    assert_true(
        question_prep.interviewer_question_category(role_end_prompt) == "role_end",
        "interviewer_question_category() should classify role-ending prompts as role_end",
    )
    assert_true(
        question_prep.interviewer_question_category(product_prompt) == "product_explainer",
        "interviewer_question_category() should classify product-explainer prompts correctly",
    )
    assert_true(
        question_prep.interviewer_question_category(ai_prompt) == "ai_saas_experience",
        "interviewer_question_category() should classify AI/SaaS exposure prompts correctly",
    )
    assert_true(
        question_prep.interviewer_question_category(training_prompt) == "training_strategy_leadership",
        "interviewer_question_category() should keep story-answerable prompts out of the non-story set",
    )
    assert_true(
        "training_strategy_leadership" not in question_prep.NON_STORY_INTERVIEWER_CATEGORIES,
        "Story-answerable categories must stay out of NON_STORY_INTERVIEWER_CATEGORIES",
    )
    assert_true(
        question_prep.interviewer_question_category(comprehension_prompt) == "comprehension_check",
        "interviewer_question_category() should classify a plain 'does that make sense?' check as its own "
        "comprehension_check category, not fold it into product_explainer",
    )

    role_end_script = question_prep.interviewer_question_factual_script(role_end_prompt, DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    product_script = question_prep.interviewer_question_factual_script(product_prompt, DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    ai_script = question_prep.interviewer_question_factual_script(ai_prompt, DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    comprehension_script = question_prep.interviewer_question_factual_script(comprehension_prompt, DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)

    assert_true(
        "(Fill in:" not in role_end_script and "factual and brief" in role_end_script,
        f"role_end factual script should avoid placeholders and unsupported reasons; got {role_end_script!r}",
    )
    assert_true(
        "(Fill in:" not in product_script and "My specialization has been" in product_script,
        f"product_explainer factual script should use the supported workflow category without placeholders; got {product_script!r}",
    )
    assert_true(
        "(Fill in:" not in ai_script and "AI-enabled workflows" in ai_script,
        f"ai_saas_experience factual script should reuse the verified AI-enabled-workflows claim instead of a fill-in; got {ai_script!r}",
    )
    assert_true(
        "(Fill in:" not in comprehension_script and comprehension_script != product_script,
        "comprehension_check should get its own complete script, distinct from the product placeholder, since "
        f"'does that make sense?' is not a product question; got {comprehension_script!r}",
    )


def test_compact_anchor_phrase_handles_intermediate_scope_words() -> None:
    import question_prep

    law_firm_scope = question_prep.compact_anchor_phrase(
        "Supported enterprise applications across more than 600 law firm offices, resolving technical issues "
        "spanning Active Directory, SQL Server, third-party integrations, and Windows services."
    )
    manufacturing_scope = question_prep.compact_anchor_phrase(
        "Protected delivery quality across 80+ international manufacturing clients, keeping implementation pace and "
        "customer trust steady during rollout."
    )
    clipped_tail_scope = question_prep.compact_anchor_phrase(
        "Managed 80+ international manufacturing clients in a $6M+ book of business by running full-lifecycle "
        "delivery across ERP implementations."
    )
    gerund_tail_scope = question_prep.compact_anchor_phrase(
        "Managed 80+ international manufacturing clients by running full-lifecycle delivery across ERP implementations."
    )
    assert_true(
        law_firm_scope == "600 law firm offices",
        f"compact_anchor_phrase() should extract the compact noun phrase for multi-word office scope; got {law_firm_scope!r}",
    )
    assert_true(
        manufacturing_scope == "80+ international manufacturing clients",
        "compact_anchor_phrase() should keep up to two intermediate scope words before the noun instead of "
        f"falling back to the full sentence; got {manufacturing_scope!r}",
    )
    assert_true(
        not re.search(r"-[A-Za-z]?$", clipped_tail_scope) and "by running" not in clipped_tail_scope.lower(),
        "compact_anchor_phrase() should strip a clipped trailing token when the regex ends mid-word instead of "
        f"returning a truncated hyphen tail; got {clipped_tail_scope!r}",
    )
    assert_true(
        "by running" not in gerund_tail_scope.lower(),
        f"compact_anchor_phrase() should not leave a dangling gerund tail in a noun-phrase anchor; got {gerund_tail_scope!r}",
    )


def test_extract_selected_proof_sentences_strips_trailing_reorg_annotation() -> None:
    import build_resume
    import question_prep

    job_description = (
        "Implementation Consultant role focused on ERP implementation, data migration, training, stakeholder "
        "alignment, and post-go-live adoption."
    )
    resume_text = (
        "Protected delivery quality and customer trust across 80+ international manufacturing clients in a $6M+ "
        "book of business by running full-lifecycle ERP delivery from discovery and configuration through data "
        "migration, testing, training, and post-go-live adoption. Position impacted by company reorganization."
    )
    profile = build_resume.job_problem_profile(job_description, resume_text)
    selected = question_prep.extract_selected_proof_sentences(resume_text, job_description, profile)

    assert_true(
        bool(selected),
        "extract_selected_proof_sentences() should keep a valid proof sentence even when a trailing reorg "
        "annotation is present.",
    )
    assert_true(
        "Position impacted by company reorganization." not in selected[0]
        and "Protected delivery quality and customer trust" in selected[0],
        "extract_selected_proof_sentences() should strip the trailing reorg annotation while preserving the main "
        f"proof sentence; got {selected[0]!r}",
    )


def test_application_answers_rotate_selected_proof_sentences() -> None:
    from dataclasses import replace

    import question_prep

    brief = question_prep.PositioningBrief(
        company_name="Revive ERP",
        role_title="Implementation Specialist",
        primary_lane="implementation_delivery",
        employer_type="manufacturing",
        mission_or_context=(
            "Revive ERP is the kind of client-facing ERP environment where implementation scope, technical "
            "complexity, timeline risk, and adoption readiness shape the work every day."
        ),
        role_core_problem="implementation scope, technical complexity, timeline risk, and adoption readiness",
        role_problem_phrase="implementation scope becoming usable delivery progress",
        personal_reason_to_care=(
            "the client-side empathy and system-change discipline in the role fit the way I like to work with "
            "teams navigating an ERP rollout"
        ),
        personal_reason_source="notes",
        strongest_direct_proofs=[
            "change adoption and enablement",
            "implementation delivery",
            "customer and revenue outcomes",
        ],
        strongest_bridge_theme="stakeholder communication",
        top_proof_anchors=[
            "600 law firm offices",
            "80+ international manufacturing clients",
        ],
        company_specific_fact=(
            "Revive ERP stands out because the team wants implementers who understand the client side of ERP change "
            "and can keep training, delivery, and adoption practical."
        ),
        gap_honesty_boundary="",
        selected_proof_sentences=[
            (
                "Protected delivery quality and customer trust across 80+ international manufacturing clients in a "
                "$6M+ book of business by running full-lifecycle ERP delivery from discovery and configuration "
                "through data migration, testing, training, and post-go-live adoption."
            ),
            (
                "Protected post-sale revenue and customer trust across 80+ international manufacturing accounts in "
                "a $6M+ book of business by guiding onboarding, adoption, renewal risk, and expansion "
                "conversations around enterprise software."
            ),
            (
                "Built 200+ dashboards and KPI reporting tools for operations and finance leaders while leading "
                "60+ executive workshops and role-based training to keep cross-functional decisions visible."
            ),
        ],
    )

    why_company = question_prep.build_why_company_answer(brief)
    relevant_experience = question_prep.build_relevant_experience_answer(brief)
    unique_qualifications = question_prep.build_unique_qualifications_answer(brief)
    communication = question_prep.build_communication_or_implementation_answer(brief, "stakeholder communication")

    answers = (why_company, relevant_experience, unique_qualifications, communication)
    proof_sentences = [question_prep.clean_answer_sentence(sentence) for sentence in brief.selected_proof_sentences]
    proof_counts = {
        sentence: sum(1 for answer in answers if sentence in answer)
        for sentence in proof_sentences
    }

    assert_true(
        proof_sentences[0] in relevant_experience and proof_sentences[0] not in why_company,
        "build_relevant_experience_answer() should keep the strongest proof sentence in the core experience answer "
        f"instead of reusing it in the why-company answer; got why={why_company!r} relevant={relevant_experience!r}",
    )
    assert_true(
        proof_sentences[1] in why_company and proof_counts[proof_sentences[1]] <= 2,
        "The second proof sentence should rotate across why-company and the qualifications answers without "
        f"spilling into more than two answers; got counts={proof_counts}",
    )
    assert_true(
        proof_sentences[2] in communication and proof_counts[proof_sentences[2]] <= 2,
        "The third proof sentence should rotate into the communication answer without appearing in more than two "
        f"answers overall; got counts={proof_counts}",
    )
    assert_true(
        all("included Supported" not in answer and "scope marker is Supported" not in answer for answer in answers),
        f"Application answers should use compact noun anchors instead of full 'Supported ...' sentences; got {answers!r}",
    )
    assert_true(
        "Position impacted by company reorganization." not in relevant_experience,
        "build_relevant_experience_answer() should not leak the trailing company-reorganization annotation into the "
        f"answer; got {relevant_experience!r}",
    )
    assert_true(
        relevant_experience.count("Additional direct proof shows up in") <= 1,
        "build_relevant_experience_answer() should emit at most one additional-proof clause instead of repeating "
        f"the same suffix for multiple proof themes; got {relevant_experience!r}",
    )
    assert_true(
        relevant_experience.count("80+") == 1 and "across 80+ international client engagements" not in relevant_experience,
        "build_relevant_experience_answer() should not repeat the 80+ scope count when the lead proof sentence "
        f"already establishes it; got {relevant_experience!r}",
    )
    assert_true(
        "One anchor is" not in unique_qualifications and "A second anchor is" not in unique_qualifications,
        "build_unique_qualifications_answer() should use scope phrasing instead of metadata-style 'anchor' labels; "
        f"got {unique_qualifications!r}",
    )

    non_implementation_relevant = question_prep.build_relevant_experience_answer(
        replace(brief, primary_lane="presales_solution")
    )
    assert_true(
        "12 full-lifecycle ERP implementations" not in non_implementation_relevant,
        "build_relevant_experience_answer() should keep the implementation-count sentence scoped to the "
        f"implementation-delivery lane; got {non_implementation_relevant!r}",
    )


def test_recent_interview_question_prep_renders_spoken_answers(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    from docx import Document

    question_prep = build_interview_cheat_sheet.question_prep
    cheat_recent = question_prep.recent_interviewer_question_prep_items
    detailed_recent = build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items
    role_end_prompt = "Why did your most recent role end?"
    story_prompt = "Tell me about leading training strategies for enterprise system implementations. What was your role?"
    role_end_item = question_prep.InterviewQuestionPrep(
        role_end_prompt,
        question_prep.interviewer_question_answer_angle(role_end_prompt, DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT),
        category="role_end",
    )
    story_item = question_prep.InterviewQuestionPrep(
        story_prompt,
        question_prep.interviewer_question_answer_angle(story_prompt, DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT),
        category="training_strategy_leadership",
    )
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION, OLLIE_RESUME_TEXT)
    stories = build_interview_cheat_sheet.expanded_story_bank()
    try:
        question_prep.recent_interviewer_question_prep_items = lambda *_args, **_kwargs: (role_end_item, story_item)
        build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items = (
            lambda *_args, **_kwargs: (role_end_item, story_item)
        )

        cheat_document = Document()
        detailed_document = Document()
        build_interview_cheat_sheet.add_recent_interview_question_prep_section(
            cheat_document,
            DUMMY_JOB_DESCRIPTION,
            "Smoke Test Systems",
            "Implementation Consultant",
            profile=profile,
            stories=stories,
            resume_text=OLLIE_RESUME_TEXT,
        )
        build_detailed_interview_guide.add_recent_interview_question_prep_section(
            detailed_document,
            DUMMY_JOB_DESCRIPTION,
            "Smoke Test Systems",
            "Implementation Consultant",
            profile=profile,
            stories=stories,
            resume_text=OLLIE_RESUME_TEXT,
        )
    finally:
        question_prep.recent_interviewer_question_prep_items = cheat_recent
        build_detailed_interview_guide.question_prep.recent_interviewer_question_prep_items = detailed_recent

    cheat_text = "\n".join(paragraph.text for paragraph in cheat_document.paragraphs)
    detailed_text = build_detailed_interview_guide.document_text(detailed_document)

    assert_true(
        "Answer angle:" not in cheat_text and "Answer angle:" not in detailed_text,
        f"Recent-question sections should never print the raw 'Answer angle:' coaching label as the answer; got cheat={cheat_text!r}",
    )
    assert_true(
        "(Fill in:" not in cheat_text and "(Fill in:" not in detailed_text,
        "Recent-question sections should render complete, evidence-safe answers without placeholders",
    )
    assert_true(
        "the clearest example" not in cheat_text.lower()
        and "a good example" not in cheat_text.lower()
        and "the best example" not in cheat_text.lower()
        and "the clearest example" not in detailed_text.lower()
        and "a good example" not in detailed_text.lower()
        and "the best example" not in detailed_text.lower(),
        f"Story-answerable recent questions should not use meta-announcement openers; got cheat={cheat_text!r} detailed={detailed_text!r}",
    )


def test_business_context_question_section_separates_answer_from_coaching(
    build_resume: object,
    build_interview_cheat_sheet: object,
    build_detailed_interview_guide: object,
) -> None:
    from docx import Document

    job_description = (
        "Company: BioTouch\n"
        "Job Title: Solutions Implementation Consultant\n"
        "This cloud-based platform supports B2B healthcare customers with 5-10 implementation projects, "
        "20-30 active customer accounts, APIs, automated file transfers, Jira, Azure, and HIPAA/PHI-sensitive workflows."
    )
    profile = build_resume.job_problem_profile(job_description, OLLIE_RESUME_TEXT)
    stories = build_interview_cheat_sheet.expanded_story_bank()
    document = Document()
    build_detailed_interview_guide.add_business_context_question_section(
        document,
        job_description,
        "",
        stories,
        profile=profile,
        company_name="BioTouch",
        role_title="Solutions Implementation Consultant",
        resume_text=OLLIE_RESUME_TEXT,
    )
    assert_true(len(document.tables) > 0, "Business-context section should render at least one Q&A card")
    for table in document.tables:
        cell_text = table.rows[0].cells[0].text
        if cell_text.startswith("SAY THIS:"):
            assert_true(
                "Hidden business concern:" not in cell_text
                and "Answer angle:" not in cell_text
                and "Best supporting story:" not in cell_text
                and "Ask back:" not in cell_text,
                f"The spoken answer box should not contain raw analyst-note labels; got {cell_text!r}",
            )
    full_text = build_detailed_interview_guide.document_text(document)
    assert_true(
        "Hidden business concern:" in full_text and "Ask back:" in full_text,
        "The hidden-concern and ask-back coaching should still surface, just in the tip footer rather than the spoken answer",
    )


def test_story_natural_reference_avoids_meta_announcement_language(
    build_interview_cheat_sheet: object,
) -> None:
    banned_phrases = ("the clearest example", "a good example", "the best example", "good example is", "best example is")
    for card in build_interview_cheat_sheet.expanded_story_bank():
        reference = build_interview_cheat_sheet.story_natural_reference(card, "Smoke Test Systems")
        lowered = reference.lower()
        assert_true(
            card.title.lower() not in lowered,
            f"story_natural_reference() should never speak the story title out loud; got {reference!r} for {card.title!r}",
        )
        for phrase in banned_phrases:
            assert_true(
                phrase not in lowered,
                f"story_natural_reference() should not use meta-announcement framing ({phrase!r}); got {reference!r}",
            )


def test_companies_refer_to_same() -> None:
    from utils import companies_refer_to_same, normalize_company_key

    assert_true(normalize_company_key("Boomi, Inc.") == "boomi", "normalize_company_key should strip corporate suffixes")
    assert_true(companies_refer_to_same("Boomi", "BOOMI"), "companies_refer_to_same should ignore case")
    assert_true(companies_refer_to_same("Boomi Inc", "Boomi"), "companies_refer_to_same should ignore suffixes")
    assert_true(not companies_refer_to_same("Boomi", "Acme"), "companies_refer_to_same should reject unrelated names")


def test_debrief_active_job_matching() -> None:
    import post_interview_debrief

    job_description = "Company: Boomi\nJob Title: Senior Presales Solutions Engineer\n"
    assert_true(
        post_interview_debrief.debrief_matches_active_job("BOOMI", job_description),
        "debrief_matches_active_job should match case-insensitive company names",
    )
    assert_true(
        not post_interview_debrief.debrief_matches_active_job("Acme", job_description),
        "debrief_matches_active_job should reject unrelated companies",
    )


def test_great_eight_utils() -> None:
    from utils import has_great_eight_signal

    assert_true(
        has_great_eight_signal("Reduced manual work by 78 percent through a cleaner workflow."),
        "has_great_eight_signal() should detect efficiency outcome language",
    )
    assert_true(
        not has_great_eight_signal("Configured the workflow for the team."),
        "has_great_eight_signal() should not pass task-only text without outcome language",
    )


def test_interview_cheat_sheet_static_content_guards(build_interview_cheat_sheet: object) -> None:
    card = build_interview_cheat_sheet.StoryCard(
        title="Inventory Adjustment Workflow",
        story_types=("Individual Achievement",),
        hook="Inventory work was slow and difficult to trust.",
        takeaways=("Root cause", "Workflow design", "Adoption"),
        evidence="Redesigned the adjustment workflow.",
        level3_trait="Separated symptoms from root cause.",
        result="Reduced manual work by 78%.",
        outcome="It proves measurable process improvement.",
        evidence_terms=("inventory",),
        signals=("process",),
    )
    story_lines = build_interview_cheat_sheet.six_story_type_lines([card])
    assert_true(len(story_lines) == 6, "six_story_type_lines() should always return all six story type slots")
    assert_true(
        any("Teamwork: No resume-supported story available for this type." in line for line in story_lines),
        "six_story_type_lines() should show a fallback when a story type is uncovered",
    )

    saas_guidance = build_interview_cheat_sheet.answer_do_dont(
        "SaaS customer success role focused on adoption, renewal, retention, and account health."
    )
    assert_true(
        not any("SPC, PFDs" in line for line in saas_guidance),
        "process engineering advice should not appear for SaaS customer success roles",
    )
    process_guidance = build_interview_cheat_sheet.answer_do_dont(
        "Process engineer role with lean six sigma, SPC, manufacturing process improvement, MATLAB, and AutoCAD exposure."
    )
    assert_true(
        any("SPC, PFDs" in line for line in process_guidance),
        "process engineering advice should appear for process engineering roles",
    )


def test_late_stage_context_detection(build_interview_cheat_sheet: object) -> None:
    assert_true(
        build_interview_cheat_sheet.detect_late_stage_context(
            "",
            "This role includes total compensation details, benefits, and transparent pay expectations.",
        ),
        "detect_late_stage_context() should detect job-description total compensation signals",
    )
    assert_true(
        not build_interview_cheat_sheet.detect_late_stage_context("", ""),
        "detect_late_stage_context() should return False when no late-stage signals exist",
    )


def test_interview_bullet_marker_is_clean(build_interview_cheat_sheet: object) -> None:
    from docx import Document

    document = Document()
    build_interview_cheat_sheet.add_bullet(document, "Example bullet")
    runs = document.paragraphs[0].runs
    assert_true(
        runs[0].text == chr(8226) + " ",
        f"Interview guide bullets should use a clean bullet marker; got {runs[0].text!r}",
    )


def test_why_company_and_negotiation_context(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    no_context_answer = build_interview_cheat_sheet.why_company_three_reason_answer(
        profile,
        "Acme Systems",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    assert_true(
        "research reminder" not in no_context_answer.lower()
        and "specific recent signal" not in no_context_answer.lower(),
        "why_company_three_reason_answer() should not speak internal research coaching when company notes are absent",
    )
    assert_true(
        "a recent company signal I will confirm before the interview" not in no_context_answer,
        "why_company_three_reason_answer() should not emit the old bare placeholder",
    )

    supplied_answer = build_interview_cheat_sheet.why_company_three_reason_answer(
        profile,
        "Acme Systems",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
        "Founded: Acme Systems launched a new analytics workflow for 600+ healthcare customers.",
    )
    assert_true(
        "Founded: Acme Systems launched a new analytics workflow" in supplied_answer,
        "why_company_three_reason_answer() should use supplied company intelligence when available",
    )

    presales_levers = dict(
        build_interview_cheat_sheet.negotiation_preparation_section(
            "Acme Systems",
            "Solution Consultant",
            LANE_JOB_DESCRIPTIONS["presales_solution"],
            "presales_solution",
        )
    )["The 10 Negotiation Levers"]
    assert_true(
        presales_levers[1].startswith("2. Performance Bonus"),
        "pre-sales negotiation levers should prioritize performance bonus second",
    )
    assert_true(
        presales_levers[2].startswith("3. Equity Grant"),
        "pre-sales negotiation levers should prioritize equity grant third",
    )

    implementation_levers = dict(
        build_interview_cheat_sheet.negotiation_preparation_section(
            "Acme Systems",
            "Implementation Consultant",
            LANE_JOB_DESCRIPTIONS["implementation_delivery"],
            "implementation_delivery",
        )
    )["The 10 Negotiation Levers"]
    assert_true(
        implementation_levers[1].startswith("2. Signing Bonus"),
        "implementation negotiation levers should keep signing bonus second",
    )


def test_communication_audit_context(build_interview_cheat_sheet: object) -> None:
    video_lines = build_interview_cheat_sheet.communication_audit_reference(
        "Remote interview over Zoom with an on-demand video recorded response step.",
        "",
    )
    by_title = {title: description for title, description, _fix in video_lines}
    assert_true(
        "camera lens" in by_title["Eye contact"],
        "communication_audit_reference() should use camera-lens guidance for video interviews",
    )
    assert_true(
        "45 to 60 seconds" in by_title["Answer length"],
        "communication_audit_reference() should use shorter timing for on-demand video interviews",
    )

    default_lines = build_interview_cheat_sheet.communication_audit_reference("", "")
    default_by_title = {title: description for title, description, _fix in default_lines}
    assert_true(
        "60 to 90 seconds" in default_by_title["Answer length"],
        "communication_audit_reference() should keep default answers inside the natural spoken target",
    )


def test_story_quality_audit_by_type(build_resume: object, build_interview_cheat_sheet: object, build_detailed_interview_guide: object) -> None:
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    card = build_interview_cheat_sheet.StoryCard(
        title="Recovery And Persuasion Story",
        story_types=("Challenge and Failure", "Persuasion", "Rapid Learning"),
        hook="A launch risk needed a cleaner recovery plan.",
        takeaways=("Recovery", "Empathy", "Learning method"),
        evidence="Rebuilt the workflow with stakeholders.",
        level3_trait="Changed the operating behavior after the issue.",
        result="Reduced risk and restored adoption confidence.",
        outcome="The team had a stronger launch path.",
        evidence_terms=("workflow",),
        signals=("risk", "persuasion", "learning"),
    )
    lines = build_detailed_interview_guide.story_quality_audit(card, profile)
    joined = "\n".join(lines)
    assert_true("RECOVERY check:" in joined, "story_quality_audit() should add failure/recovery checks")
    assert_true("EMPATHY check:" in joined, "story_quality_audit() should add persuasion empathy checks")
    assert_true("RAMP check:" in joined, "story_quality_audit() should add rapid-learning checks")
    assert_true("TONAL LEAD-BURIAL check:" in joined, "story_quality_audit() should warn when the most human line may be buried too late")
    assert_true(len(lines) <= 8, "story_quality_audit() should cap audit lines at eight")


def test_four_trust_questions_audit(build_resume: object, build_detailed_interview_guide: object) -> None:
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    lines = build_detailed_interview_guide.four_trust_questions_audit(
        profile,
        "Acme Systems",
        "Implementation Project Manager",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    joined = "\n".join(lines)
    assert_true(len(lines) == 4, "four_trust_questions_audit() should return four trust questions")
    assert_true("Trust question 1" in joined and "Trust question 4" in joined, "four_trust_questions_audit() should label each executive trust question clearly")
    assert_true("Acme Systems" in joined, "four_trust_questions_audit() should keep the company context visible")


def test_executive_presence_signals(build_resume: object, build_detailed_interview_guide: object, build_cover_letter: object) -> None:
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    lines = build_detailed_interview_guide.executive_presence_signals(
        profile,
        "Acme Systems",
        "Implementation Project Manager",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    joined = "\n".join(lines)
    assert_true(
        "Open with the answer" in joined and "Acme Systems" in joined,
        "executive_presence_signals() should include direct executive-presence guidance tied to the company context",
    )

    light_proof_letter = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Hiring Manager,",
            "Acme Systems needs an Implementation Project Manager who can keep ERP adoption and workflow clarity aligned through go-live.",
            "My background fits environments where teams need steadier communication, usable process structure, and customer confidence during change.",
            "I bring a practical style that keeps stakeholders aligned and moves the work forward without unnecessary noise.",
            "I would welcome a conversation about the role and how I could support Acme Systems.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    warnings = build_cover_letter.validate_cover_letter_specificity(
        light_proof_letter,
        "Acme Systems",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    assert_true(
        any("Executive presence note:" in warning for warning in warnings),
        "validate_cover_letter_specificity() should add a warning-only executive-presence note when proof density is too light",
    )


def test_lane_specific_recording_focus(build_resume: object, build_detailed_interview_guide: object) -> None:
    implementation_profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    implementation_lines = build_detailed_interview_guide.lane_specific_recording_focus(
        implementation_profile,
        "Implementation Project Manager",
    )
    implementation_joined = "\n".join(implementation_lines)
    assert_true(
        "scope, risk, handoffs, and adoption" in implementation_joined,
        "lane_specific_recording_focus() should give implementation-specific recording guidance",
    )

    customer_success_profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["customer_success"])
    customer_success_lines = build_detailed_interview_guide.lane_specific_recording_focus(
        customer_success_profile,
        "Customer Success Manager",
    )
    assert_true(
        any("renewal risk" in line.lower() or "value realization" in line.lower() for line in customer_success_lines),
        "lane_specific_recording_focus() should include customer-success-specific guidance when the lane is customer_success",
    )


def test_debrief_pattern_analysis(build_interview_cheat_sheet: object, build_debrief_analysis: object) -> None:
    with TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "debrief_history.txt"
        path.write_text(
            "\n\n".join(
                (
                    """POST-INTERVIEW DEBRIEF CAPTURED 2026-01-02 09:00:00

Interview date: 01/02/2026
Company name: Acme Systems
Round number: 1
Outcome: advance

Stories that generated follow-up questions:
Inventory workflow story. How did adoption hold?

Unexpected questions:
How did adoption hold?

Specific interviewer language about the role:
They need measurable implementation discipline.""",
                    """POST-INTERVIEW DEBRIEF CAPTURED 2026-01-05 09:00:00

Interview date: 01/05/2026
Company name: Acme Systems
Round number: 2
Outcome: pending

Stories that generated follow-up questions:
Inventory adjustment work and manual-work reduction. How did adoption hold?

Unexpected questions:
How did adoption hold?

Specific interviewer language about the role:
They need measurable implementation discipline.""",
                )
            ),
            encoding="utf-8",
        )
        entries = build_debrief_analysis.find_debrief_entries_for_company("Acme Systems", path)
        summary = build_debrief_analysis.analyze_entries(entries, "Acme Systems")

    assert_true(len(entries) == 2, "find_debrief_entries_for_company() should return company-matched debrief entries")
    assert_true(
        summary.most_common_question == "How did adoption hold?",
        "analyze_entries() should detect the most common recurring debrief question",
    )
    assert_true(
        summary.top_story_title == "Inventory adjustment system",
        "analyze_entries() should identify a repeated top-performing story theme when one is present",
    )
    reordered = build_debrief_analysis.reorder_story_cards(
        build_interview_cheat_sheet.expanded_story_bank()[:5],
        summary,
    )
    assert_true(
        reordered[0].title == "Inventory adjustment system",
        "reorder_story_cards() should move the top-performing story to the front when debrief analysis finds one",
    )


def test_interview_context_review_analysis(interview_context: object) -> None:
    review = interview_context.analyze_performance_review(
        "\n".join(
            [
                "Overall: qualified, but too long and loose.",
                "Where you were weakest",
                "Keeping answers short",
                "Leading with the direct answer",
                "Biggest risks to fix",
                "Cut your answer length by 30-40%.",
                "Using the wrong example first.",
                "Questions at the end were tactical, not strategic.",
                "If they have similar ERP candidates with stronger executive presence, that could matter.",
                "Filler / style habits",
                "repeated restarts",
            ]
        )
    )
    keys = {item["key"] for item in review["coaching_signals"]}
    assert_true("rambling" in keys, "analyze_performance_review() should detect rambling risk")
    assert_true("delayed_answer" in keys, "analyze_performance_review() should detect delayed-answer risk")
    assert_true("wrong_example_first" in keys, "analyze_performance_review() should detect wrong-example-first risk")
    assert_true("consultative_questions" in keys, "analyze_performance_review() should detect weak consultative-question risk")
    assert_true("executive_presence" in keys, "analyze_performance_review() should detect executive-presence risk")
    assert_true("filler_restarts" in keys, "analyze_performance_review() should detect filler/restart risk")


def test_interview_context_extracts_exact_interviewer_questions(interview_context: object) -> None:
    questions = interview_context.extract_interviewer_questions(
        "\n".join(
            [
                '1) "Do you have any questions from our previous conversation?"',
                '2) "Why did your most recent role end in November?"',
                '3) Your question: "What reporting tools are used?"',
                "Scorecard",
                '4) "This should not be collected."',
            ]
        )
    )
    assert_true(
        questions == (
            "Do you have any questions from our previous conversation?",
            "Why did your most recent role end in November?",
        ),
        f"extract_interviewer_questions() should preserve exact interviewer prompts, stop before coaching sections, and skip Christian's own questions; got {questions}",
    )


def test_interview_context_structured_round_storage(interview_context: object) -> None:
    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        jobs_dir.mkdir()
        record = interview_context.normalize_round_record(
            {
                "company_name": "Acme Systems",
                "role_title": "Implementation Consultant",
                "interview_date": "06/19/2026",
                "round_number": "2",
                "outcome": "pending",
                "raw_notes": "Customer asked about implementation pace and stakeholder alignment.",
                "story_followups": ["Implementation lifecycle", "Stakeholder disagreement"],
                "unexpected_questions": ["How would you ramp in a new industry?"],
                "role_language": ["They need someone who can keep adoption and decision quality visible."],
                "feedback_received": ["Needed crisper direct answers."],
                "company_intelligence": ["Round 3 is an on-site final round."],
                "performance_review": {
                    "summary": "Strong substance, but too long.",
                    "next_round_risks": ["Answer first", "Shorten answers by 30%"],
                    "coaching_signals": [
                        {"key": "rambling", "label": "Rambling / answer length", "detail": "Shorten answers by 30%."},
                        {"key": "executive_presence", "label": "Executive presence / polish", "detail": "Sound more declarative."},
                    ],
                },
            }
        )
        interview_context.save_round_record(jobs_dir, record)
        interview_context.rebuild_company_dossiers(jobs_dir)
        interview_context.rebuild_legacy_exports(jobs_dir)
        bundle = interview_context.load_company_context(jobs_dir, "Acme Systems", "Implementation Consultant", mode="full")
        dossier_path = jobs_dir / "company_notes" / "acme_systems.md"
        debrief_path = jobs_dir / "debrief_history.txt"

        assert_true(dossier_path.exists(), "rebuild_company_dossiers() should write a company dossier markdown file")
        assert_true(debrief_path.exists(), "rebuild_legacy_exports() should regenerate debrief_history.txt")
        assert_true("Round 3 is an on-site final round." in bundle.company_research, "load_company_context() should surface structured company intelligence")
        assert_true("Coaching signals:" in bundle.interview_notes, "load_company_context() should include normalized coaching notes in interview context")
        assert_true("Recurring coaching themes:" in bundle.coaching_notes, "load_company_context() should summarize recurring coaching themes")


def test_agents_word_budget() -> None:
    agents_path = PROJECT_ROOT / "AGENTS.md"
    reference_path = PROJECT_ROOT / "SYSTEM_REFERENCE.md"
    word_count = len(re.findall(r"\b\w+\b", agents_path.read_text(encoding="utf-8")))
    assert_true(agents_path.exists(), "AGENTS.md should exist")
    assert_true(reference_path.exists(), "SYSTEM_REFERENCE.md should exist after the compact AGENTS refactor")
    assert_true(word_count <= 3500, f"AGENTS.md should stay under 3500 words after the compact-context refactor; got {word_count}")


def test_interview_context_compact_context_budget_and_diagnosis(interview_context: object) -> None:
    fixture_path = PROJECT_ROOT / "jobs" / "interview_debriefs" / "m_tech_systems" / "2026-06-19__round_1.json"
    assert_true(fixture_path.exists(), f"Expected real M-Tech fixture at {fixture_path}")
    bundle = interview_context.load_company_context(
        PROJECT_ROOT / "jobs",
        "M-Tech Systems",
        "Project Consultant",
        mode="compact",
    )
    supplied_word_count = len(re.findall(r"\b\w+\b", bundle.supplied_context))
    assert_true(
        supplied_word_count <= 800,
        f"Compact supplied_context should stay under 800 words; got {supplied_word_count}",
    )
    assert_true(
        "ownership-forward" in bundle.supplied_context.lower() or "ownership" in bundle.supplied_context.lower(),
        f"Compact M-Tech context should preserve the ownership diagnosis; got {bundle.supplied_context!r}",
    )
    assert_true(
        "consultative" in bundle.supplied_context.lower(),
        f"Compact M-Tech context should preserve the consultative diagnosis; got {bundle.supplied_context!r}",
    )


def test_interview_context_review_appendix_storage(interview_context: object) -> None:
    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        jobs_dir.mkdir()
        record = {
            "company_name": "Acme Systems",
            "role_title": "Implementation Consultant",
            "interview_date": "06/19/2026",
            "round_number": "1",
            "outcome": "pending",
            "role_language": ["Client-facing ownership matters."],
            "imported_review_text": "\n".join(
                [
                    "Overall: qualified, but not ownership-forward enough.",
                    '1) "How would you handle stakeholder disagreement?"',
                    "Where you were weakest",
                    "Leading with the direct answer",
                    "Keeping answers short",
                    "Questions at the end were tactical, not strategic.",
                ]
            ),
        }
        path = interview_context.save_round_record(jobs_dir, record)
        payload = json.loads(path.read_text(encoding="utf-8"))
        appendix_path = Path(payload["review_appendix_path"])
        appendix_exists = appendix_path.exists()
        interview_context.rebuild_company_dossiers(jobs_dir)
        dossier_text = (jobs_dir / "company_notes" / "acme_systems.md").read_text(encoding="utf-8")

    assert_true(payload["imported_review_text"] == "", "Structured round JSON should not inline the full imported review text by default")
    assert_true(appendix_exists, "save_round_record() should externalize the raw imported review into an appendix file")
    assert_true(
        "review_analysis" in payload and payload["review_analysis"]["decision_signal"]["headline"],
        "Structured round JSON should persist parsed review-analysis sections",
    )
    assert_true(
        "### Imported Performance Review Appendix" in dossier_text and "What you said" not in dossier_text,
        "Default company dossiers should reference the raw review appendix instead of embedding the full imported review body",
    )


def test_interview_context_global_coaching_fallback_is_lazy(interview_context: object) -> None:
    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        jobs_dir.mkdir()
        interview_context.save_round_record(
            jobs_dir,
            {
                "company_name": "Acme Systems",
                "role_title": "Implementation Consultant",
                "interview_date": "06/19/2026",
                "round_number": "1",
                "outcome": "pending",
                "performance_review": {
                    "summary": "Strong substance, but too long.",
                    "coaching_signals": [
                        {"key": "rambling", "label": "Rambling / answer length", "detail": "Shorten answers."}
                    ],
                },
                "raw_notes": "This should not be needed for global fallback.",
            },
        )
        interview_context.save_round_record(
            jobs_dir,
            {
                "company_name": "Beta Labs",
                "role_title": "Customer Success Consultant",
                "interview_date": "06/20/2026",
                "round_number": "2",
                "outcome": "pending",
                "raw_notes": "Global fallback should compact this record.",
            },
        )
        no_fallback = interview_context.global_coaching_fallback_records(jobs_dir, "Acme Systems", "Implementation Consultant")
        fallback = interview_context.global_coaching_fallback_records(jobs_dir, "Missing Company", "Missing Role")

    assert_true(
        no_fallback == (),
        f"global_coaching_fallback_records() should stay empty when same-company coaching already exists; got {no_fallback}",
    )
    assert_true(fallback, "global_coaching_fallback_records() should fall back to compact global records when scoped coaching is unavailable")
    assert_true(
        all(not str(record.get("raw_notes", "")).strip() and not str(record.get("imported_review_text", "")).strip() for record in fallback),
        "Global coaching fallback should use compact records instead of carrying raw-note or raw-review text",
    )


def test_interview_context_recent_question_lookup_prefers_same_company(interview_context: object) -> None:
    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        jobs_dir.mkdir()
        interview_context.save_round_record(
            jobs_dir,
            {
                "company_name": "Acme Systems",
                "role_title": "Implementation Consultant",
                "interview_date": "06/15/2026",
                "round_number": "1",
                "outcome": "pending",
                "imported_review_text": '1) "How would you handle stakeholder disagreement?"',
            },
        )
        interview_context.save_round_record(
            jobs_dir,
            {
                "company_name": "Beta Labs",
                "role_title": "Customer Success Consultant",
                "interview_date": "06/20/2026",
                "round_number": "2",
                "outcome": "pending",
                "imported_review_text": '1) "Why should we move you to the next round?"',
            },
        )
        same_company = interview_context.recent_interviewer_questions(
            jobs_dir,
            "Acme Systems",
            "Implementation Consultant",
        )
        fallback_latest = interview_context.recent_interviewer_questions(
            jobs_dir,
            "Missing Company",
            "Missing Role",
        )
    with TemporaryDirectory() as empty_temp_dir:
        empty_jobs_dir = Path(empty_temp_dir) / "jobs"
        empty_jobs_dir.mkdir()
        no_questions = interview_context.recent_interviewer_questions(
            empty_jobs_dir,
            "Missing Company",
            "Missing Role",
        )

    assert_true(
        same_company == ("How would you handle stakeholder disagreement?",),
        f"recent_interviewer_questions() should prefer same-company history first; got {same_company}",
    )
    assert_true(
        fallback_latest == ("Why should we move you to the next round?",),
        f"recent_interviewer_questions() should fall back to the latest interview overall when no same-company record exists; got {fallback_latest}",
    )
    assert_true(
        no_questions == (),
        f"recent_interviewer_questions() should return an empty tuple when no structured question source exists; got {no_questions}",
    )


def test_interview_context_legacy_repair(interview_context: object) -> None:
    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        scoped_dir = jobs_dir / "interview_notes_by_company"
        jobs_dir.mkdir()
        scoped_dir.mkdir()
        (jobs_dir / "debrief_history.txt").write_text(
            "\n".join(
                [
                    "POST-INTERVIEW DEBRIEF CAPTURED 2026-06-19 14:33:21",
                    "",
                    "Interview date: 6/19/2026",
                    "Company name: M-Tech Systems",
                    "Role title: Project Consultant",
                    "Round number: 1",
                    "Outcome: pending",
                    "",
                    "Stories that generated follow-up questions:",
                    "Implementation lifecycle and data validation",
                    "",
                    "Unexpected questions:",
                    "How would you ramp in a new industry?",
                    "",
                    "Specific interviewer language about the role:",
                    "Industry knowledge matters more than raw technical skill.",
                    "",
                    "Feedback received:",
                    "Need stronger direct answers.",
                    "",
                    "Insider company intelligence learned:",
                    "Three rounds with an on-site final round.",
                ]
            ),
            encoding="utf-8",
        )
        (scoped_dir / "m_tech_systems__project_consultant__round_1.txt").write_text(
            "\n".join(
                [
                    "Company: M-Tech Systems",
                    "Role: Project Consultant",
                    "Round: 1",
                    "Interview date: 6/19/2026",
                    "Outcome:",
                    "",
                    "Paste raw interview notes below this line.",
                    "Keep notes factual. Include interviewer language, questions, objections, feedback, and next-step clues.",
                    "",
                    "Jonathan said the team cares about learning the poultry business quickly.",
                    "",
                    "Company: M-Tech Systems",
                    "Role: Project Consultant",
                    "Round: 1",
                    "Interview date: 6/19/2026",
                    "Outcome: pending",
                    "",
                    "Raw notes:",
                    "Duplicated structured block that should be trimmed.",
                ]
            ),
            encoding="utf-8",
        )
        result = interview_context.repair_legacy_interview_data(jobs_dir)
        records = interview_context.load_round_records(jobs_dir, "M-Tech Systems")
        assert_true(result["records_written"] >= 1, "repair_legacy_interview_data() should write structured round records")
        assert_true(len(records) == 1, "repair_legacy_interview_data() should collapse matching legacy artifacts into one round record")
        assert_true(
            "Duplicated structured block" not in str(records[0].get("raw_notes", "")),
            "repair_legacy_interview_data() should trim duplicated structured note content from legacy raw notes",
        )
        assert_true(
            "Keep notes factual" not in str(records[0].get("raw_notes", "")),
            "repair_legacy_interview_data() should remove raw-note instruction boilerplate from repaired notes",
        )
        assert_true(
            (jobs_dir / "company_notes" / "m_tech_systems.md").exists(),
            "repair_legacy_interview_data() should rebuild a company dossier for repaired legacy entries",
        )


def test_interview_context_dedupes_same_round_records(interview_context: object) -> None:
    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        company_dir = jobs_dir / "interview_debriefs" / "biotouch"
        jobs_dir.mkdir()
        company_dir.mkdir(parents=True)
        duplicate_general = {
            "company_name": "BioTouch",
            "role_title": "",
            "interview_date": "05/26/2026",
            "interview_date_iso": "2026-05-26",
            "round_number": "2",
            "outcome": "pending",
            "raw_notes": "",
            "story_followups": ["Excel comfort"],
            "unexpected_questions": [],
            "role_language": [],
            "feedback_received": [],
            "company_intelligence": [],
            "performance_review": {"summary": "", "strongest_answers": [], "weakest_answers": [], "next_round_risks": [], "coaching_signals": []},
            "imported_review_text": "",
            "imported_artifacts": [],
            "legacy_text": "",
            "captured_at": "2026-05-26 13:00:00",
        }
        duplicate_specific = {
            "company_name": "BioTouch",
            "role_title": "Solutions Implementation Consultant",
            "interview_date": "05/26/2026",
            "interview_date_iso": "2026-05-26",
            "round_number": "2",
            "outcome": "pending",
            "raw_notes": "Detailed raw notes from the same round.",
            "story_followups": [],
            "unexpected_questions": [],
            "role_language": ["Consulting, curious to learn"],
            "feedback_received": [],
            "company_intelligence": ["Supports Spectra Path"],
            "performance_review": {"summary": "", "strongest_answers": [], "weakest_answers": [], "next_round_risks": [], "coaching_signals": []},
            "imported_review_text": "",
            "imported_artifacts": [],
            "legacy_text": "",
            "captured_at": "2026-05-26 14:00:00",
        }
        (company_dir / "2026-05-26__round_2__general.json").write_text(json.dumps(duplicate_general), encoding="utf-8")
        (company_dir / "2026-05-26__round_2__solutions_implementation_consultant.json").write_text(json.dumps(duplicate_specific), encoding="utf-8")
        records = interview_context.load_round_records(jobs_dir, "BioTouch")

        assert_true(len(records) == 1, "load_round_records() should merge duplicate same-round JSON entries into one canonical record")
        assert_true(
            records[0].get("role_title") == "Solutions Implementation Consultant",
            "load_round_records() should preserve the richer role title when merging same-round records",
        )
        assert_true(
            "Detailed raw notes from the same round." in str(records[0].get("raw_notes", "")),
            "load_round_records() should preserve the richer raw notes when merging same-round records",
        )


def test_interview_context_scopes_company_context_and_role_variants(interview_context: object) -> None:
    text = "\n\n".join(
        [
            "POST-INTERVIEW NOTE [05/26/2026]: BioTouch - Solutions Implementation Consultant, round 2, outcome pending.\nWe would support Spectra Path.",
            "POST-INTERVIEW NOTE [6/19/2026]: M-Tech Systems - Project Consultant, round 1, outcome pending.\nThree rounds with an on-site final round.",
        ]
    )
    scoped = interview_context.relevant_company_context(text, "M-Tech Systems")
    assert_true("M-Tech Systems" in scoped, "relevant_company_context() should keep the matching company block")
    assert_true("BioTouch" not in scoped, "relevant_company_context() should drop non-matching company blocks from shared context files")

    with TemporaryDirectory() as temp_dir:
        jobs_dir = Path(temp_dir) / "jobs"
        jobs_dir.mkdir()
        interview_context.save_round_record(
            jobs_dir,
            interview_context.normalize_round_record(
                {
                    "company_name": "M-Tech Systems",
                    "role_title": "Project Consultant",
                    "interview_date": "06/19/2026",
                    "round_number": "1",
                    "outcome": "pending",
                    "raw_notes": "Round-specific raw notes.",
                }
            ),
        )
        bundle = interview_context.load_company_context(jobs_dir, "M-Tech Systems", "Implementation Specialist (Project Consultant)")
        assert_true(bundle.round_records, "load_company_context() should match structured rounds even when the active role title includes a parenthetical variant")


def test_discussion_topic_helpers(utils: object) -> None:
    raw_notes = """
- Goal: implementation pace and customer confidence
- Tasks: training others and technical projects
- Biggest challenge: data validation
"""
    topic = utils.extract_single_discussion_topic(raw_notes)
    assert_true(
        topic == "implementation pace and customer confidence",
        f"extract_single_discussion_topic() should pick one concrete discussion topic instead of a raw note dump; got {topic!r}",
    )
    sentence = utils.discussion_topic_sentence(raw_notes, intro="I appreciated the detail around")
    assert_true(
        sentence == "I appreciated the detail around implementation pace and customer confidence.",
        f"discussion_topic_sentence() should render one clean sentence from labeled bullets; got {sentence!r}",
    )
    assert_true(
        "Tasks:" not in sentence and "Biggest challenge:" not in sentence,
        f"discussion_topic_sentence() should strip raw note labels from generated prose; got {sentence!r}",
    )


def test_template_leakage_validator(utils: object) -> None:
    utils.assert_no_template_leakage("I appreciated the detail around implementation pace and customer confidence.")
    utils.assert_no_template_leakage(
        "Protected cross-functional delivery and post-launch data-migration work across customer-facing implementations."
    )
    utils.assert_no_template_leakage(
        "\n".join(
            [
                "- Build a rehearsed foundation for the interview.",
                "- Use one proof line instead of three examples.",
                "- End on the business result once the point is clear.",
            ]
        )
    )

    for label, text in (
        ("None-supplied placeholder", "I appreciated None supplied during the conversation."),
        (
            "dash-prefixed note dump",
            "I appreciated - Goal: implementation pace - Tasks: training others - Biggest challenge: data validation.",
        ),
        (
            "raw note appendix marker",
            "Appendix: Raw Supplied Notes\n- Company research supplied by Christian:\n- POST-INTERVIEW NOTE [6/19/2026]: M-Tech Systems",
        ),
    ):
        error: SystemExit | None = None
        with contextlib.redirect_stderr(io.StringIO()):
            try:
                utils.assert_no_template_leakage(text)
            except SystemExit as exc:
                error = exc
        assert_true(
            isinstance(error, SystemExit),
            f"assert_no_template_leakage() should fail on {label}; got {error!r}",
        )


def test_thank_you_body_normalizes_discussion_topics(build_thank_you: object, utils: object) -> None:
    debrief_entry = """POST-INTERVIEW DEBRIEF CAPTURED 2026-06-19 16:10:00

Interview date: 06/19/2026
Company name: M-Tech Systems
Role title: Implementation Specialist (Project Consultant)
Round number: 1
Outcome: pending

Stories that generated follow-up questions:
- Biggest challenge: data validation
- Tasks: training others and technical projects

Specific interviewer language about the role:
- Goal: implementation pace and customer confidence
- Tasks: training others and technical projects

Feedback received:
They said they would follow up next week.

Insider company intelligence learned:
Spectra Path implementation work.
"""
    body, notes = build_thank_you.thank_you_body(
        "M-Tech Systems",
        "Implementation Specialist (Project Consultant)",
        debrief_entry,
        ["Built 200+ dashboards and reporting tools that improved decision visibility."],
    )
    assert_true(
        "After our conversation, I kept thinking about implementation pace and customer confidence." in body,
        f"thank_you_body() should open with a concrete post-conversation takeaway; got {body!r}",
    )
    assert_true(
        "I left the conversation clearer on where I could create value quickly: Built 200+ dashboards and reporting tools that improved decision visibility." in body,
        f"thank_you_body() should connect the follow-up to a concrete proof point; got {body!r}",
    )
    assert_true(
        "The follow-up around data validation also stood out because it is the kind of detail that shapes adoption, service quality, and client confidence." in body,
        f"thank_you_body() should turn follow-up questions into one clean business-oriented sentence; got {body!r}",
    )
    assert_true(
        "- Goal:" not in body and "- Tasks:" not in body and "None supplied" not in body,
        f"thank_you_body() should not leak raw note bullets or missing-value placeholders; got {body!r}",
    )
    utils.assert_no_template_leakage(body)
    assert_true(notes == [], "thank_you_body() should not add review notes when debrief and proof are present")


def test_interview_followup_body(build_interview_followup: object) -> None:
    debrief_entry = """POST-INTERVIEW DEBRIEF CAPTURED 2026-05-26 13:35:19

Interview date: 05/26/2026
Company name: BioTouch
Role title: Solutions Implementation Consultant
Round number: 2
Outcome: pending

Stories that generated follow-up questions:
training others and technical projects

Specific interviewer language about the role:
consulting, curious to learn

Feedback received:
They said they would follow up in 1 to 2 weeks.

Insider company intelligence learned:
We would support Spectra Path
"""
    subject, body, notes = build_interview_followup.interview_followup_body(
        "BioTouch",
        "Solutions Implementation Consultant",
        debrief_entry,
        ["Built 200+ dashboards and reporting tools that improved decision visibility."],
    )
    assert_true(
        "Following up on the Solutions Implementation Consultant interview process" == subject,
        "interview_followup_body() should generate the interview follow-up subject line",
    )
    assert_true(
        "Built 200+ dashboards" in body and "training others and technical projects" in body,
        "interview_followup_body() should weave in a concrete proof point and the debrief follow-up theme",
    )
    assert_true(
        build_interview_followup.suggest_followup_window(debrief_entry).startswith("Suggested send window: 2026-06-02 to 2026-06-09"),
        "suggest_followup_window() should use the interview timeline from the debrief when one is provided",
    )
    assert_true(notes == [], "interview_followup_body() should not add review notes when debrief and proof are present")


def test_post_round_followup_email_normalizes_discussion_topics(build_post_round: object, utils: object) -> None:
    entry = """POST-INTERVIEW DEBRIEF CAPTURED 2026-06-19 16:10:00

Specific interviewer language about the role:
- Goal: implementation pace and customer confidence
- Tasks: training others and technical projects
"""
    message = build_post_round.post_round_followup_email(
        "Implementation Specialist (Project Consultant)",
        "M-Tech Systems",
        entry,
    )
    assert_true(
        "I appreciated the detail around implementation pace and customer confidence." in message,
        f"post_round_followup_email() should convert labeled raw notes into one clean sentence; got {message!r}",
    )
    assert_true(
        "- Goal:" not in message and "- Tasks:" not in message,
        f"post_round_followup_email() should not leak raw note labels; got {message!r}",
    )
    utils.assert_no_template_leakage(message)


def test_post_round_followup_email_is_lane_aware(build_post_round: object) -> None:
    entry = """POST-INTERVIEW DEBRIEF CAPTURED 2026-06-19 16:10:00

Specific interviewer language about the role:
customer reporting and decision support
"""
    analytics_message = build_post_round.post_round_followup_email(
        "Data Analyst, Retention",
        "Ollie",
        entry,
        job_description=OLLIE_ANALYTICS_JOB_DESCRIPTION,
    )
    change_message = build_post_round.post_round_followup_email(
        "Change Enablement Lead",
        "Acme Change",
        entry,
        job_description=LANE_JOB_DESCRIPTIONS["change_enablement"],
    )
    assert_true(
        "implementation discipline, customer communication, and technical issue ownership" not in analytics_message,
        f"Analytics post-round emails should not reuse implementation-only phrasing; got {analytics_message!r}",
    )
    assert_true(
        "reporting, workflow clarity, and customer-facing follow-through" in analytics_message,
        f"Analytics post-round emails should use analytics language; got {analytics_message!r}",
    )
    assert_true(
        "stakeholder communication, training, and adoption follow-through" in change_message,
        f"Change-enablement post-round emails should use change-focused phrasing; got {change_message!r}",
    )


def test_build_skills_database(build_skills_database: object) -> None:
    with TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "skills_database.md"
        result_path = build_skills_database.build_skills_database(output_path=output_path, force=True)
        content = result_path.read_text(encoding="utf-8")
    assert_true(result_path == output_path, "build_skills_database() should write to the requested output path")
    assert_true(
        "## Unified Core Competencies" in content and "## Role Evidence Map" in content,
        "build_skills_database() should render the expected structured markdown sections",
    )
    assert_true(
        "Estrada_Resume_Implementation.docx" in content and "Estrada_Resume_PreSales_CSM.docx" in content,
        "build_skills_database() should record both source resumes in the markdown output",
    )


def test_standard_qualifications_answers_known_questions(
    build_standard_qualifications_statement: object,
) -> None:
    prompts = (
        "How many years of business-related experience do you have with public agencies or cooperatives?",
        "Please list all certifications and licenses you currently hold relevant to this position.",
        "Please list all software packages, systems, and programs for which you rate your skills at an intermediate or higher level, as it directly relates to position requirements.",
        "Briefly describe your level of demonstrated effective communication, formal public speaking, presentation experience, group facilitation or training skills, as it relates to the duties of this position.",
    )
    _, snapshot, resume_text = build_standard_qualifications_statement.selected_resume_snapshot(
        SOURCEWELL_QUALIFICATIONS_JOB_DESCRIPTION
    )
    responses = build_standard_qualifications_statement.build_question_responses(
        prompts,
        SOURCEWELL_QUALIFICATIONS_JOB_DESCRIPTION,
        snapshot,
        resume_text,
    )
    response_map = {response.prompt: response.answer for response in responses}
    assert_true(
        "0 years of direct public-agency or cooperative experience" in response_map[prompts[0]],
        "Public-agency answers should stay honest about the direct-experience gap while still bridging adjacent experience.",
    )
    assert_true(
        "ITIL 4 Foundation" in response_map[prompts[1]],
        "Certification answers should surface the supported ITIL credential.",
    )
    assert_true(
        "Power BI" in response_map[prompts[2]]
        and "Salesforce CRM" in response_map[prompts[2]]
        and "ServiceNow" in response_map[prompts[2]],
        "Software answers should include the supported reporting, CRM, and systems tools from the source resumes.",
    )
    assert_true(
        "60+ executive workshops and QBRs" in response_map[prompts[3]]
        and "pre-sales discovery and product demonstrations" in response_map[prompts[3]],
        "Communication answers should use the strongest supported public-speaking and facilitation proof points.",
    )


def test_standard_qualifications_answers_added_company_and_implementation_questions(
    build_standard_qualifications_statement: object,
) -> None:
    prompts = (
        "Why are you interested in joining this company?",
        "Describe the typical customer you support, including industry(s), employee count, and main stakeholder roles.",
        "How many concurrent implementations have you worked on at any given time and what was the average length of an implementation process?",
        "How do you measure performance and success during a software implementation?",
        "We are looking for people who are personally and/or professionally passionate about AI. Please briefly explain how you have put it to work for you in either or both areas of your life.",
        "Tell us about a complex project or program you led where there was no clear delivery playbook, template, or established way of working. What was the business/customer outcome you were accountable for, what was ambiguous at the start, and how did you create enough structure for the team to move forward? Please include how you handled the first few weeks, what operating rhythm or artifacts you created, who you needed to influence, and what the final outcome was.",
        "Describe a time when a customer, executive stakeholder, or internal partner expected something that the product/platform/team could not actually deliver as promised. What was the gap between expectation and reality, how did you diagnose the issue, and how did you communicate the tradeoffs or bad news? Please include who was involved, how you protected trust, what decision or path forward you drove, and what you learned from the situation.",
        "Do you have experience working for a SaaS or technology company with an AI product or service? If yes, please briefly describe your experience.",
        "Can you describe your experience in project leadership roles for complex initiatives? What were some of the challenges you faced, and how did you address them?",
        "Could you share details about your experience leading training strategies for large-scale technology or enterprise system implementations? What types of projects have you worked on, and what was your role?",
    )
    _, snapshot, resume_text = build_standard_qualifications_statement.selected_resume_snapshot(
        SOURCEWELL_QUALIFICATIONS_JOB_DESCRIPTION
    )
    responses = build_standard_qualifications_statement.build_question_responses(
        prompts,
        SOURCEWELL_QUALIFICATIONS_JOB_DESCRIPTION,
        snapshot,
        resume_text,
    )
    response_map = {response.prompt: response.answer for response in responses}
    assert_true(
        "interested in joining Sourcewell" in response_map[prompts[0]]
        or "interested in joining" in response_map[prompts[0]].lower(),
        f"Company-interest answers should produce a tailored company-join rationale; got {response_map[prompts[0]]!r}",
    )
    assert_true(
        "employee counts were not always disclosed" in response_map[prompts[1]]
        and "150+ user ERP environment" in response_map[prompts[1]]
        and "operations leaders" in response_map[prompts[1]],
        "Typical-customer answers should stay honest about undisclosed employee counts while naming supported industries and stakeholder groups.",
    )
    assert_true(
        "80+ international manufacturing client engagements" in response_map[prompts[2]]
        and "multiple concurrent implementations or workstreams" in response_map[prompts[2]]
        and "discovery through post-go-live support" in response_map[prompts[2]],
        "Implementation-volume answers should use supported portfolio and concurrent-workstream proof without inventing a fake average duration.",
    )
    assert_true(
        "scope clarity" in response_map[prompts[3]].lower()
        and "UAT" in response_map[prompts[3]]
        and "five-site 150+ user ERP environment" in response_map[prompts[3]]
        and "200+ KPI dashboards and reporting tools" in response_map[prompts[3]],
        "Implementation-success answers should define the success criteria and tie them to supported testing, adoption, and reporting proof.",
    )
    assert_true(
        "Claude and AI-assisted tools" in response_map[prompts[4]]
        and "zero-to-one internal SMS support channel" in response_map[prompts[4]]
        and "AI-assisted chatbot logic" in response_map[prompts[4]],
        "AI-passion answers should stay bridge-hard by citing the supported AI workflow and chatbot evidence from the source resumes.",
    )
    assert_true(
        "warehouse operation and Amazon Robotics program" in response_map[prompts[5]]
        and "operations, finance, engineering, and system-readiness tracks" in response_map[prompts[5]]
        and "role-based training materials, and change communications" in response_map[prompts[5]],
        "Ambiguity-delivery answers should use the strongest supported from-scratch program example and name the structure Christian created.",
    )
    assert_true(
        "customization, integration, or workflow outcome" in response_map[prompts[6]]
        and "Statements of Work and Functional Requirements Documents" in response_map[prompts[6]]
        and "high-risk or low-value customizations" in response_map[prompts[6]],
        "Expectation-gap answers should ground the tradeoff conversation in supported scope, feasibility, and trust-protection evidence.",
    )
    assert_true(
        response_map[prompts[7]].startswith("Yes.")
        and "Aptean and Aderant" in response_map[prompts[7]]
        and "LivePerson LiveEngage workflows" in response_map[prompts[7]]
        and "AI-assisted chatbot logic" in response_map[prompts[7]],
        "SaaS and AI-company answers should use a bounded yes posture with supported SaaS, legal-tech, and AI-workflow evidence.",
    )
    assert_true(
        "enterprise platform migration" in response_map[prompts[8]]
        and "warehouse operation and Amazon Robotics program" in response_map[prompts[8]]
        and "risk assumptions, testing and validation checkpoints, role-based training, and a stakeholder rhythm" in response_map[prompts[8]],
        "Complex-project-leadership answers should use the strongest supported migration and warehouse-launch leadership proof.",
    )
    assert_true(
        response_map[prompts[9]].startswith("Yes.")
        and "cross-site training, onboarding materials, and change communications" in response_map[prompts[9]]
        and "80+ international client engagements" in response_map[prompts[9]]
        and "reduce reactive post-launch support volume" in response_map[prompts[9]],
        "Training-strategy answers should tie enterprise implementation training work to readiness, adoption, and lower support volume.",
    )


def test_standard_application_question_parser_dedupes_blocks(
    build_standard_qualifications_statement: object,
) -> None:
    prompts = build_standard_qualifications_statement.parse_question_blocks(
        "\n".join(
            [
                "Application Questions",
                "",
                "How many years of direct relevant experience do you have related to the job duties of this position?",
                "",
                "Briefly describe your relevant experience",
                "demonstrating how it directly relates to the position.",
                "",
                "How many years of direct relevant experience do you have related to the job duties of this position?",
            ]
        )
    )
    assert_true(
        prompts == (
            "How many years of direct relevant experience do you have related to the job duties of this position?",
            "Briefly describe your relevant experience demonstrating how it directly relates to the position.",
        ),
        f"parse_question_blocks() should skip headings, preserve multiline prompts, and dedupe repeats; got {prompts}",
    )


def test_standard_qualifications_document_renders_recent_interview_questions(
    build_standard_qualifications_statement: object,
) -> None:
    responses = (
        build_standard_qualifications_statement.QualificationsResponse(
            "Relevant experience summary",
            "Supported summary text.",
        ),
    )
    recent_scripts = (
        (
            "What was your role in go-live?",
            "I led the go-live cutover end to end, coordinating readiness checks, training, and stakeholder sign-off.",
        ),
    )
    document = build_standard_qualifications_statement.build_document(
        "Acme Systems",
        "Implementation Consultant",
        responses,
        recent_interviewer_scripts=recent_scripts,
        used_custom_questions=False,
    )
    text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    assert_true(
        "Recent Interview Questions To Be Ready For" in text
        and "What was your role in go-live?" in text
        and "I led the go-live cutover end to end" in text,
        "Commercial qualifications statements should render the resolved spoken script for each recent interviewer-question appendix item.",
    )


def test_standard_qualifications_recent_interviewer_scripts_resolve_factual_script(
    build_standard_qualifications_statement: object,
) -> None:
    items = (
        build_standard_qualifications_statement.question_prep.InterviewQuestionPrep(
            prompt="Why did your most recent role end?",
            answer_angle="Frame the answer around the role's core responsibilities.",
            category="role_end",
        ),
        build_standard_qualifications_statement.question_prep.InterviewQuestionPrep(
            prompt="Are there specific product offerings you specialized in?",
            answer_angle="Translate the product into the closest supported workflow or system category first.",
            category="product_explainer",
        ),
    )
    job_description = (
        "Implementation Specialist supporting enterprise system rollouts, configuration, and client training "
        "for manufacturing customers."
    )
    resume_text = "Led enterprise system implementations across five manufacturing sites, training end users on new workflows."
    scripts = build_standard_qualifications_statement.build_recent_interviewer_scripts(
        items,
        job_description,
        resume_text,
        "Acme Systems",
        "Implementation Specialist",
    )
    assert_true(
        len(scripts) == len(items) and scripts[0][0] == items[0].prompt and scripts[1][0] == items[1].prompt,
        f"build_recent_interviewer_scripts() should pair each item's prompt with a resolved script; got {scripts}",
    )
    assert_true(
        "factual and brief" in scripts[0][1]
        and "My specialization has been" in scripts[1][1]
        and all("Fill in" not in script for _prompt, script in scripts),
        "build_recent_interviewer_scripts() should render complete evidence-safe scripts without placeholders; "
        f"got {scripts!r}",
    )


def test_startup_interview_false_positive_guard(build_interview_cheat_sheet: object) -> None:
    false_positive = build_interview_cheat_sheet.startup_interview_lines(
        "Fortune 500 company seeking someone comfortable in a fast-paced environment.",
        "",
    )
    assert_true(
        false_positive == [],
        "startup_interview_lines() should not trigger on a single startup-adjacent phrase at a Fortune 500 company",
    )
    operator_lines = build_interview_cheat_sheet.startup_interview_lines(
        "Role requires ownership, ambiguity, and autonomy across process lanes.",
        "",
    )
    assert_true(
        operator_lines and operator_lines[0].startswith("Operator fit:"),
        "startup_interview_lines() should use operator framing when broad-operator signals are present without true early-stage signals",
    )
    startup_lines = build_interview_cheat_sheet.startup_interview_lines(
        "Early-stage Series B startup needs a builder to build the function from scratch.",
        "",
    )
    assert_true(
        startup_lines and startup_lines[0].startswith("Startup fit:"),
        "startup_interview_lines() should keep startup framing for true early-stage contexts",
    )


def test_lead_burial_and_response_calibration(build_cover_letter: object, build_interview_cheat_sheet: object) -> None:
    calibration_lines = build_interview_cheat_sheet.response_calibration_lines()
    calibration_text = "\n".join(calibration_lines)
    assert_true(
        "Structural audit:" in calibration_text,
        "response_calibration_lines() should include a structural audit reminder",
    )
    assert_true(
        "Tonal audit:" in calibration_text,
        "response_calibration_lines() should include a tonal audit reminder",
    )

    buried_letter = "\n".join(
        [
            "Christian Estrada | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Hiring Manager,",
            "Acme Systems needs an Implementation Project Manager who can keep ERP adoption, workflow clarity, and customer confidence aligned through go-live.",
            "I work best in roles where implementation, migration, integration, and go-live workflows have to become usable under pressure without losing customer trust across enterprise systems.",
            "I have seen that kind of environment repeatedly, and I bring a practical style that keeps users, clients, and stakeholders aligned while the work moves.",
            "At East West Manufacturing, I supported 150+ users across five sites and reduced discrepancies by 22% while protecting migration readiness.",
            "I would welcome a conversation about how that pattern could help Acme Systems.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    lead_problem = build_cover_letter.lead_burial_check(buried_letter)
    assert_true(
        lead_problem is not None and "structural lead burial" in lead_problem.lower(),
        "lead_burial_check() should flag letters that hide concrete proof until after the second body paragraph",
    )


def test_application_checklist_debrief_lookup(build_application_checklist: object) -> None:
    with TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "debrief_history.txt"
        path.write_text(
            "\n\n".join(
                (
                    """POST-INTERVIEW DEBRIEF CAPTURED 2026-01-02 09:00:00

Interview date: 2026-01-02
Company name: Acme Systems
Round number: 1
Outcome: advance

Stories that generated follow-up questions:
Inventory workflow

Unexpected questions:
How did adoption hold?

Specific interviewer language about the role:
They need measurable implementation discipline.""",
                    """POST-INTERVIEW DEBRIEF CAPTURED 2026-01-03 09:00:00

Interview date: 2026-01-03
Company name: Other Company
Round number: 1
Outcome: pending

Stories that generated follow-up questions:
Other story""",
                )
            ),
            encoding="utf-8",
        )
        entries = build_application_checklist.find_debrief_entries_for_company("Acme", path)
    assert_true(len(entries) == 1, "find_debrief_entries_for_company() should filter entries by company")
    assert_true(
        "Inventory workflow" in build_application_checklist.section_value(entries[0], "Stories that generated follow-up questions"),
        "section_value() should extract debrief section content",
    )


def test_application_checklist_analysis_resume_prefers_tailored_output(build_application_checklist: object) -> None:
    with TemporaryDirectory(prefix="checklist_analysis_resume_") as temp_name:
        temp_root = Path(temp_name)
        original_output_dir = build_application_checklist.OUTPUT_DIR
        try:
            build_application_checklist.OUTPUT_DIR = temp_root
            selected_resume = build_minimal_docx(temp_root / "source_resume.docx", ["Source resume"])
            tailored_resume = build_minimal_docx(
                temp_root / "Christian Estrada - Smoke Test Systems - Implementation Consultant Resume.docx",
                ["Tailored resume"],
            )
            resolved = build_application_checklist.analysis_resume_path(DUMMY_JOB_DESCRIPTION, selected_resume)
            assert_true(
                resolved == tailored_resume,
                "analysis_resume_path() should prefer the newest tailored resume output when one exists",
            )
            assert_true(
                "Tailored resume" in build_application_checklist.analysis_basis_label(resolved, selected_resume),
                "analysis_basis_label() should describe when checklist analysis is based on a tailored output",
            )
        finally:
            build_application_checklist.OUTPUT_DIR = original_output_dir


def test_application_checklist_source_fallback_uses_alignment_grade(build_application_checklist: object) -> None:
    selected_resume = Path("C:/temp/source_resume.docx")
    original_fit_classification = build_application_checklist.fit_classification
    original_alignment_score_report = build_application_checklist.build_resume.alignment_score_report
    try:
        build_application_checklist.fit_classification = lambda *_args, **_kwargs: "FAIL (should not be used)"
        build_application_checklist.build_resume.alignment_score_report = lambda *_args, **_kwargs: {"grade": "Adjacent Fit"}
        result = build_application_checklist.fit_snapshot_status(
            DUMMY_JOB_DESCRIPTION,
            selected_resume,
            selected_resume,
            "Source resume text",
        )
    finally:
        build_application_checklist.fit_classification = original_fit_classification
        build_application_checklist.build_resume.alignment_score_report = original_alignment_score_report

    assert_true(
        result == "Adjacent Fit",
        "fit_snapshot_status() should use alignment grade instead of final-output audit rules when only the source resume is available",
    )


def test_general_advice_active_job_helpers(build_general_advice: object) -> None:
    presales_lines = build_general_advice.selected_coaching_lines(
        "presales_solution",
        LANE_JOB_DESCRIPTIONS["presales_solution"],
    )
    assert_true(len(presales_lines) <= 3, "selected_coaching_lines() should return at most three coaching items")
    assert_true(
        any("top-down" in line.lower() or "because" in line.lower() or "interest" in line.lower() for line in presales_lines),
        "selected_coaching_lines() should choose pre-sales-specific coaching when available",
    )


def test_general_advice_shared_sections(build_general_advice: object, job_search_guidance: object) -> None:
    sections = build_general_advice.general_advice_sections()
    assert_true("Search Strategy" in sections, "general_advice_sections() should include Search Strategy")
    assert_true("Job Search Diagnosis" in sections, "general_advice_sections() should include Job Search Diagnosis")
    assert_true("JD Wish List Reality" in sections, "general_advice_sections() should include JD Wish List Reality")
    assert_true("Recruiter Screen" in sections, "general_advice_sections() should include Recruiter Screen")
    assert_true(
        "Recruiter Call Research Calibration" in sections,
        "general_advice_sections() should include Recruiter Call Research Calibration",
    )
    assert_true("Follow-Up Discipline" in sections, "general_advice_sections() should include Follow-Up Discipline")
    search_text = " ".join(sections["Search Strategy"]).lower()
    diagnosis_text = " ".join(sections["Job Search Diagnosis"]).lower()
    wish_list_text = " ".join(sections["JD Wish List Reality"]).lower()
    resume_text = " ".join(sections["Resume Strategy"]).lower()
    followup_text = " ".join(sections["Follow-Up Discipline"]).lower()
    research_text = " ".join(sections["Recruiter Call Research Calibration"]).lower()
    communication_text = " ".join(sections["Communication Principles"]).lower()
    all_text = " ".join(" ".join(lines) for lines in sections.values()).lower()
    assert_true(
        "one-page career vision" in search_text and "job shopper" in search_text,
        "Search Strategy should include one-page vision and job-shopper guidance",
    )
    assert_true(
        "slow search" in diagnosis_text and "pipeline" in diagnosis_text,
        "Job Search Diagnosis should stay qualitative and pattern-based",
    )
    assert_true(
        "wish lists" in wish_list_text and "knockout" in wish_list_text,
        "JD Wish List Reality should explain broad-fit applications and real knockout terms",
    )
    assert_true(
        "ats reality" in resume_text and "autobiography" in resume_text,
        "Resume Strategy should include ATS myth-busting and autobiography guidance",
    )
    assert_true(
        "follow-up" in followup_text or "follow up" in followup_text,
        "Follow-Up Discipline should include timing guidance",
    )
    assert_true(
        "first recruiter call" in research_text and "later-round" in research_text,
        "Recruiter Call Research Calibration should distinguish early research from later-round research",
    )
    assert_true(
        "draft first, edit second" in communication_text,
        "Communication Principles should include the draft-first-edit-second rule",
    )
    recruiter_lines = job_search_guidance.recruiter_screen_lines(
        "Hybrid role with 30% travel. Must be authorized to work in the United States."
    )
    recruiter_text = " ".join(recruiter_lines).lower()
    assert_true(
        "compensation" in recruiter_text and "travel" in recruiter_text and "authorization" in recruiter_text,
        "recruiter_screen_lines() should cover compensation, travel, and work authorization",
    )
    banned_phrases = (
        "first 20 seconds determine 80%",
        "six months is the normal search",
        "half of hiring managers keep roles open",
        "apply online only as a last resort",
        "ghost jobs",
        "transparent notes",
        "confidence signals competence",
    )
    assert_true(
        not any(phrase in all_text for phrase in banned_phrases),
        "Shared job-search guidance should avoid hard stats, ghost-job framing, and sneaky interview tactics",
    )


def test_question_intent_framework(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    story = build_interview_cheat_sheet.StoryCard(
        title="Five-site ERP stabilization",
        story_types=("challenge", "implementation"),
        hook="Inherited a five-site ERP environment with inconsistent reporting and cutover risk.",
        takeaways=(
            "Mapped the real operating risk before acting.",
            "Aligned stakeholders around the next decision.",
            "Kept adoption visible during the change.",
        ),
        evidence="Led five-site ERP stabilization work with cross-functional owners, SQL validation, and user adoption follow-through.",
        level3_trait="Judgment under ambiguity",
        result="Stabilized reporting and cutover readiness across the active sites.",
        outcome="Protected adoption and decision quality during a complex transition.",
        evidence_terms=("ERP", "SQL", "adoption"),
        signals=("implementation", "stakeholder alignment"),
    )
    cards = build_interview_cheat_sheet.question_intent_framework(
        profile,
        "Smoke Test Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        [story],
    )
    prompts = {card.prompt for card in cards}
    assert_true(len(cards) == 14, f"question_intent_framework() should return 14 cards; got {len(cards)}")
    assert_true(
        "Why this company?" in prompts and "What would you do in the first 90 days?" in prompts,
        "question_intent_framework() should include company-fit and first-90-days prompts",
    )
    assert_true(
        all(card.hidden_assessment and card.bad_answer_trap and card.story_angle and card.clarify_when and card.ideal_length for card in cards),
        "question_intent_framework() should populate every card field",
    )


def test_likely_question_story_prefers_direct_example(build_interview_cheat_sheet: object) -> None:
    aptean_story = build_interview_cheat_sheet.StoryCard(
        title="Aptean lifecycle delivery",
        story_types=("implementation", "go-live"),
        hook="Supported full lifecycle delivery from discovery through go-live.",
        takeaways=("Scoped the work", "Aligned stakeholders", "Protected rollout quality"),
        evidence="Worked from discovery through go-live with business and technical teams.",
        level3_trait="Lifecycle judgment",
        result="Delivered a stable implementation path with stronger readiness.",
        outcome="Proved full implementation ownership and rollout discipline.",
        evidence_terms=("Aptean", "go-live"),
        signals=("implementation", "go-live", "Aptean"),
    )
    east_west_story = build_interview_cheat_sheet.StoryCard(
        title="East West ERP ownership",
        story_types=("enhancement", "upgrade"),
        hook="Managed changes in an already-live ERP environment.",
        takeaways=("Clarified the need", "Managed downstream risk", "Validated before release"),
        evidence="Owned customization and upgrade work inside an existing ERP environment.",
        level3_trait="Change-risk discipline",
        result="Kept changes controlled without creating avoidable disruption.",
        outcome="Proved enhancement and upgrade judgment in a live environment.",
        evidence_terms=("East West", "upgrade"),
        signals=("east west", "upgrade", "customization", "ERP"),
    )
    lifecycle_choice = build_interview_cheat_sheet.likely_question_story(
        build_interview_cheat_sheet.InterviewQuestion(
            "Walk me through an implementation you owned.",
            "Focus on lifecycle, configuration, and go-live readiness.",
        ),
        [east_west_story, aptean_story],
    )
    upgrade_choice = build_interview_cheat_sheet.likely_question_story(
        build_interview_cheat_sheet.InterviewQuestion(
            "Tell me about a customization or upgrade you managed.",
            "Focus on an already-live environment and downstream risk.",
        ),
        [aptean_story, east_west_story],
    )
    assert_true(
        lifecycle_choice.title == "Aptean lifecycle delivery",
        "likely_question_story() should prefer the direct lifecycle example for implementation questions",
    )
    assert_true(
        upgrade_choice.title == "East West ERP ownership",
        "likely_question_story() should prefer the live-environment example for upgrade questions",
    )


def test_likely_question_story_avoids_reuse_when_alternative_exists(build_interview_cheat_sheet: object) -> None:
    strong_match_story = build_interview_cheat_sheet.StoryCard(
        title="Aptean lifecycle delivery",
        story_types=("implementation", "go-live"),
        hook="Supported full lifecycle delivery from discovery through go-live.",
        takeaways=("Scoped the work", "Aligned stakeholders", "Protected rollout quality"),
        evidence="Worked from discovery through go-live with business and technical teams.",
        level3_trait="Lifecycle judgment",
        result="Delivered a stable implementation path with stronger readiness.",
        outcome="Proved full implementation ownership and rollout discipline.",
        evidence_terms=("Aptean", "go-live"),
        signals=("implementation", "go-live", "Aptean"),
    )
    weak_match_story = build_interview_cheat_sheet.StoryCard(
        title="200+ dashboards and decision visibility",
        story_types=("Analysis and Decision",),
        hook="Built reporting that made decisions visible to the business.",
        takeaways=("Found the missing signal", "Built the dashboard", "Made the decision visible"),
        evidence="Built 200+ dashboards used for day-to-day decisions.",
        level3_trait="Analytical judgment",
        result="Gave the business a faster, clearer decision signal.",
        outcome="Proved data-driven decision support.",
        evidence_terms=("200", "dashboards"),
        signals=("dashboard", "reporting", "analysis"),
    )
    question = build_interview_cheat_sheet.InterviewQuestion(
        "Walk me through an implementation you owned.",
        "Focus on lifecycle, configuration, and go-live readiness.",
    )
    fresh_choice = build_interview_cheat_sheet.likely_question_story(
        question, [strong_match_story, weak_match_story], used_titles=set()
    )
    assert_true(
        fresh_choice.title == "Aptean lifecycle delivery",
        "likely_question_story() should still prefer the strongest keyword match when nothing has been used yet",
    )
    repeat_choice = build_interview_cheat_sheet.likely_question_story(
        question, [strong_match_story, weak_match_story], used_titles={"Aptean lifecycle delivery"}
    )
    assert_true(
        repeat_choice.title == "200+ dashboards and decision visibility",
        "likely_question_story() should switch to an unused story rather than repeat one verbatim within "
        f"the same question list, even if the unused story is a weaker keyword match; got {repeat_choice.title!r}",
    )
    last_resort_choice = build_interview_cheat_sheet.likely_question_story(
        question,
        [strong_match_story, weak_match_story],
        used_titles={"Aptean lifecycle delivery", "200+ dashboards and decision visibility"},
    )
    assert_true(
        last_resort_choice.title in {"Aptean lifecycle delivery", "200+ dashboards and decision visibility"},
        "likely_question_story() should still return a story (best of the already-used ones) when every "
        f"story has been used; got {last_resort_choice.title!r}",
    )


def test_interview_addition_helpers(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    stories = [
        build_interview_cheat_sheet.StoryCard(
            title="Five-site ERP stabilization",
            story_types=("challenge", "implementation"),
            hook="Inherited a messy multi-site ERP environment.",
            takeaways=("Mapped the risk", "Aligned the owners", "Protected adoption"),
            evidence="Led validation, issue ownership, and cutover coordination across five sites and 150+ users.",
            level3_trait="Judgment under ambiguity",
            result="Stabilized cutover readiness and reporting clarity across the active sites.",
            outcome="Protected adoption and decision quality during the migration window.",
            evidence_terms=("ERP", "SQL", "adoption"),
            signals=("implementation", "risk", "adoption", "stakeholder alignment"),
        ),
        build_interview_cheat_sheet.StoryCard(
            title="At-risk account recovery",
            story_types=("persuasion", "customer"),
            hook="Inherited a frustrated high-risk customer relationship.",
            takeaways=("Created one recovery path", "Clarified ownership", "Restored trust"),
            evidence="Reframed the issue path, coordinated owners, and kept updates useful to the customer.",
            level3_trait="Calm ownership",
            result="Stabilized more than $1M in at-risk revenue and improved trust with the account.",
            outcome="Protected renewal confidence and customer credibility.",
            evidence_terms=("customer", "recovery", "renewal"),
            signals=("customer", "recovery", "trust", "account"),
        ),
        build_interview_cheat_sheet.StoryCard(
            title="Dashboard decision-support buildout",
            story_types=("analysis", "reporting"),
            hook="Leaders lacked a clean view of workflow issues and outcomes.",
            takeaways=("Found the metric gap", "Built the reporting path", "Made the output usable"),
            evidence="Built dashboards and reporting views that turned raw data into decision-ready visibility.",
            level3_trait="Decision clarity",
            result="Improved workflow visibility and made issue trends easier to act on.",
            outcome="Helped leaders make faster decisions with clearer reporting.",
            evidence_terms=("dashboard", "reporting", "workflow"),
            signals=("dashboard", "reporting", "decision", "visibility"),
        ),
    ]
    playbook = build_interview_cheat_sheet.phone_screen_first_round_playbook(profile, "Acme Health", "Implementation Consultant")
    translation = build_interview_cheat_sheet.recruiter_translation_lines(profile, "Acme Health", "Implementation Consultant")
    honesty = build_interview_cheat_sheet.interview_honesty_and_privacy_lines()
    closing = build_interview_cheat_sheet.closing_interview_drill_lines(
        "Acme Health",
        "Implementation Consultant",
        profile,
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
    )
    proof_themes = build_interview_cheat_sheet.three_supported_proof_theme_lines(profile, stories)
    combined_text = " ".join(playbook + translation + honesty + closing + proof_themes).lower()
    assert_true(
        any("30-minute pacing" in line.lower() for line in playbook) and any("compensation early" in line.lower() for line in playbook),
        "Phone screen playbook should cover pacing and salary deflection",
    )
    assert_true(
        any("translators and screeners" in line.lower() for line in translation),
        "Recruiter translation guidance should frame recruiters as translators and screeners",
    )
    assert_true(
        any("family plans" in line.lower() for line in honesty) and any("titles" in line.lower() for line in honesty),
        "Interview honesty and privacy guidance should distinguish honest scope from private topics",
    )
    assert_true(
        any("explicit interest" in line.lower() for line in closing) and any("next step" in line.lower() for line in closing),
        "Closing drill should cover explicit interest and a clear next-step signal",
    )
    assert_true(
        any("competitors" in line.lower() or "pain point" in line.lower() for line in playbook),
        "Phone screen playbook should push research beyond the company website",
    )
    assert_true(
        any("biggest gap between the experience you are seeing" in line.lower() for line in closing),
        "Closing drill should include the direct experience-gap question",
    )
    assert_true(
        proof_themes[0].lower().startswith("if an interviewer asks for three words")
        and any(" - " in line for line in proof_themes[1:]),
        "Three-word guidance should become supported proof themes with examples",
    )
    assert_true(
        "transparent notes" not in combined_text and "confidence signals competence" not in combined_text,
        "Interview helpers should avoid sneaky tactics and unsupported confidence claims",
    )


def test_search_progress_question_is_conditional(build_resume: object, build_interview_cheat_sheet: object) -> None:
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    regular_questions = build_interview_cheat_sheet.questions_to_ask("Acme Health", profile, LANE_JOB_DESCRIPTIONS["implementation_delivery"])
    late_stage_questions = build_interview_cheat_sheet.questions_to_ask(
        "Acme Health",
        profile,
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
        "Panel interview scheduled after the final round and references requested.",
    )
    assert_true(
        not any("other candidates" in question.lower() for question in regular_questions),
        "Search-progress question should stay out of default early-stage question sets",
    )
    assert_true(
        any(question.lower().startswith("what is the biggest gap between the experience you are seeing") for question in regular_questions),
        "Questions to ask should include the direct experience-gap diagnostic by default",
    )
    assert_true(
        any(question.lower().startswith("situational, later-stage only") for question in late_stage_questions),
        "Search-progress question should appear only when later-stage or long-open context exists",
    )


def test_salary_guide_helpers(build_salary_guide: object, job_search_guidance: object) -> None:
    walkaway = build_salary_guide.walkaway_criteria_lines("Remote hybrid role with 30% travel and bonus potential.")
    counteroffer = job_search_guidance.counteroffer_evaluation_lines()
    salary_text = " ".join(job_search_guidance.salary_research_lines("Remote implementation role"))
    assert_true(
        any("travel" in line.lower() for line in walkaway) and any("work-location policy" in line.lower() for line in walkaway),
        "Walkaway criteria should account for travel and work-location policy when relevant",
    )
    assert_true(
        any("original reason" in line.lower() or "began looking" in line.lower() for line in counteroffer)
        and any("full package" in line.lower() for line in counteroffer),
        "Counteroffer evaluation should stay neutral and checklist-based",
    )
    assert_true(
        "approved range" in salary_text and "walkaway point" in salary_text,
        "Shared salary guidance should emphasize range-first framing and walkaway criteria",
    )


def test_proof_text_rewrites_dense_cover_sentence() -> None:
    import proof_text

    cleaned = proof_text.sanitize_proof_sentence(
        "Owned a mission-critical enterprise system across five sites and 150+ users while supporting data migration, validation, training, and adoption.",
        max_words=24,
        artifact="cover_letter_proof",
    )
    assert_true(
        "five sites" in cleaned.lower() and "150+ users" in cleaned.lower(),
        f"sanitize_proof_sentence() should preserve the East West proof markers; got {cleaned!r}",
    )
    assert_true(
        "while supporting data migration, validation, training, and adoption" not in cleaned.lower(),
        f"sanitize_proof_sentence() should rewrite the dense Stord-style proof cluster; got {cleaned!r}",
    )
    assert_true(
        proof_text.word_count(cleaned) <= 24,
        f"sanitize_proof_sentence() should keep the rewritten proof inside the word budget; got {cleaned!r}",
    )


def test_followup_news_and_variants(
    build_resume: object,
    build_followup_email: object,
    interview_context: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    company_notes = (
        "BioTouch launched a Spectra Path workflow expansion for pathology teams.\n"
        "OtherCo raised a funding round for a different market."
    )
    news_line = interview_context.recent_company_news_line(company_notes, "BioTouch")
    assert_true(
        "Spectra Path workflow expansion" in news_line,
        f"recent_company_news_line() should keep the relevant company line; got {news_line!r}",
    )
    variants = build_followup_email.followup_email_variants(
        "BioTouch",
        "Implementation Consultant",
        profile,
        news_line,
    )
    labels = [label for label, _ in variants]
    variant_map = dict(variants)
    assert_true(
        labels == ["Brief Follow-Up", "Value-Forward Update", "News-Aware Follow-Up"],
        f"followup_email_variants() should append the news-aware variant when news exists; got {labels}",
    )
    assert_true(
        "I also noticed" in variant_map["News-Aware Follow-Up"],
        "News-aware follow-up should mention the verified company update",
    )


def test_followup_email_variants_pass_prose_quality(
    build_resume: object,
    build_followup_email: object,
    utils: object,
) -> None:
    profile = build_resume.job_problem_profile(DUMMY_JOB_DESCRIPTION)
    variants = build_followup_email.followup_email_variants(
        "BioTouch",
        "Implementation Consultant",
        profile,
        "BioTouch expanded its Spectra Path workflow support.",
    )
    for label, body in variants:
        with contextlib.redirect_stdout(io.StringIO()):
            report = utils.enforce_prose_quality(body, "followup_email_body", label=label)
        assert_true(report["passed"], f"{label} follow-up body should pass shared prose-quality checks; got {report}")


def test_resume_readiness_flags_finance_blockers(build_resume: object) -> None:
    job_description = """
Company: AuditCo
Role: Program Manager
This role owns budget management, EAC, ETC, and financial forecasting across a complex delivery program.
"""
    resume_text = """
Professional Summary
Implementation consultant with experience leading scope, reporting, stakeholder alignment, and launch readiness across enterprise systems.
"""
    readiness = build_resume.resume_readiness_report(
        job_description,
        resume_text,
        source_resume_text=resume_text,
    )
    assert_true(
        any(
            gap.blocker and gap.support_level == "unsupported-do-not-insert"
            and any(term in gap.label.lower() for term in ("budget", "eac", "etc"))
            for gap in readiness.hard_blockers
        ),
        f"resume_readiness_report() should expose finance ownership blockers without inserting them; got {readiness.hard_blockers}",
    )


def test_cover_close_uses_direct_ask(build_cover_letter: object) -> None:
    signals = build_cover_letter.CoverLetterSignals(
        company_mission="",
        role_core_function="implementation pace and customer confidence",
        top_accomplishment="Built 200+ dashboards that improved decision visibility.",
        fit_bridge="",
        jd_skill_terms=("implementation", "stakeholder coordination"),
        ambiguity_process="turn half-formed questions into structured plans",
        jd_test_environments=("UAT",),
        communication_metric="Facilitated 60+ executive workshops and QBRs.",
        partner_functions=("operations", "support"),
        jd_pain_area="implementation pace",
    )
    close, _terms = build_cover_letter.soft_close_paragraph("Guidehouse", signals)
    assert_true(
        close.startswith("I would welcome the chance to discuss"),
        f"soft_close_paragraph() should use a direct ask; got {close!r}",
    )
    assert_true(
        "would be valuable" not in close.lower(),
        f"soft_close_paragraph() should remove weak hypothetical closes; got {close!r}",
    )


def test_linkedin_guidance_helpers(build_resume: object, build_linkedin_update: object) -> None:
    job_description = LANE_JOB_DESCRIPTIONS["analytics_operations"]
    profile = build_resume.job_problem_profile(job_description)
    keywords = build_linkedin_update.recruiter_keywords(profile, job_description)
    themes = build_linkedin_update.thought_leadership_themes(profile, job_description)
    comments = build_linkedin_update.comment_strategy(profile, job_description)
    assert_true(
        any("analytics" in keyword.lower() or "report" in keyword.lower() for keyword in keywords),
        f"LinkedIn recruiter keywords should include lane-relevant search terms; got {keywords}",
    )
    assert_true(
        any("dashboard" in theme.lower() or "decision" in theme.lower() for theme in themes),
        f"LinkedIn themes should stay proof-based and lane-specific; got {themes}",
    )
    assert_true(
        any("observation" in line.lower() or "question" in line.lower() for line in comments),
        f"LinkedIn comment strategy should provide a practical comment structure; got {comments}",
    )


def test_cover_letter_signals_ollie_analytics(build_cover_letter: object) -> None:
    signals = build_cover_letter.build_cover_letter_signals(
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
    )
    assert_true(
        signals.company_mission.startswith("Ollie makes human-grade pet food"),
        f"build_cover_letter_signals() should extract a mission/product promise sentence; got {signals.company_mission!r}",
    )
    assert_true(
        "subscription retention and member loyalty" in signals.role_core_function,
        f"build_cover_letter_signals() should extract the role core function from the JD; got {signals.role_core_function!r}",
    )
    assert_true(
        all(term in signals.jd_skill_terms for term in ("retention", "lifecycle", "cohort analysis")),
        f"build_cover_letter_signals() should prefer exact analytics terms from the JD; got {signals.jd_skill_terms}",
    )
    assert_true(
        all(term in signals.jd_test_environments for term in ("lifecycle comms", "mobile apps", "member services")),
        f"build_cover_letter_signals() should extract test environments from the JD; got {signals.jd_test_environments}",
    )
    assert_true(
        all(term in signals.partner_functions for term in ("product", "marketing", "cx", "research")),
        f"build_cover_letter_signals() should extract cross-functional partners from the JD; got {signals.partner_functions}",
    )
    assert_true(
        "retention" in signals.jd_pain_area and "lifecycle" in signals.jd_pain_area,
        f"build_cover_letter_signals() should extract a concrete pain area for the close; got {signals.jd_pain_area!r}",
    )


def test_cover_lane_keeps_procore_in_implementation(build_resume: object, build_cover_letter: object) -> None:
    profile = build_resume.job_problem_profile(PROCORE_JOB_DESCRIPTION)
    lane_key = build_cover_letter.effective_lane_key(
        "Senior Solutions Architect",
        PROCORE_JOB_DESCRIPTION,
        profile,
    )
    assert_true(
        lane_key == "implementation_delivery",
        f"effective_lane_key() should keep implementation-heavy solution architect roles in implementation_delivery when Customer Success is only a partner team; got {lane_key!r}",
    )
    assert_true(
        not build_cover_letter.support_escalation_role("Senior Solutions Architect", PROCORE_JOB_DESCRIPTION),
        "support_escalation_role() should stay false for Procore-style solution architect postings that mention Customer Success only as a partner function",
    )


def test_cover_lane_keeps_explicit_implementation_roles_out_of_support(build_resume: object, build_cover_letter: object) -> None:
    job_description = """
Company: Infor
Job Title: CPQ Implementation Consultant, Senior

As a Functional Consultant supporting Infor's CPQ solutions, you will guide customers through end-to-end implementations of Configure, Price, and Quote capabilities.
You will collaborate cross-functionally with delivery, product, and support teams to drive successful implementations.
A Typical Day in the Life Includes:
Participating in full-cycle implementations, including Kickoff, Requirements Gathering, Functional Design and Documentation, System Configuration, Testing, and Training.
Partnering with Product Management, Product Development, and Customer Support to communicate customer needs and design effective solutions.
Basic Qualifications:
Experience with configure, price, quote software.
"""
    profile = build_resume.job_problem_profile(job_description)
    lane_key = build_cover_letter.effective_lane_key(
        "CPQ Implementation Consultant, Senior",
        job_description,
        profile,
    )
    assert_true(
        lane_key == "implementation_delivery",
        f"effective_lane_key() should keep explicit implementation roles out of the customer-support lane when support is only a partner function; got {lane_key!r}",
    )
    signals = build_cover_letter.build_cover_letter_signals(
        "Infor",
        "CPQ Implementation Consultant, Senior",
        job_description,
        PROCORE_RESUME_TEXT,
    )
    assert_true(
        "member issues" not in signals.role_core_function.lower()
        and ("implementation" in signals.role_core_function.lower() or "delivery" in signals.role_core_function.lower()),
        f"build_cover_letter_signals() should keep implementation roles in implementation-style cover language; got {signals.role_core_function!r}",
    )


def test_cover_lane_prefers_strategy_consulting_titles(build_resume: object, build_cover_letter: object) -> None:
    job_description = """
Company: Guidehouse
Job Title: Strategy & Transformation Senior Consultant

Guidehouse is a consulting firm serving public-sector and commercial clients.
The role helps clients solve complex business challenges from strategy through execution.
Responsibilities include requirements gathering and strategy sessions with senior-level clients,
developing findings and recommendation reports, root cause and gap analysis, process flows,
reporting dashboards, and stakeholder communication across project teams.
"""
    profile = build_resume.job_problem_profile(job_description)
    lane_key = build_cover_letter.effective_lane_key(
        "Strategy & Transformation Senior Consultant",
        job_description,
        profile,
    )
    assert_true(
        lane_key == "corporate_strategy",
        f"effective_lane_key() should prefer corporate_strategy for strategy-consulting titles even when the JD also mentions implementation-style responsibilities; got {lane_key!r}",
    )

    signals = build_cover_letter.build_cover_letter_signals(
        "Guidehouse",
        "Strategy & Transformation Senior Consultant",
        job_description,
        OLLIE_RESUME_TEXT,
    )
    assert_true(
        "technical scoping" not in signals.role_core_function.lower(),
        f"build_cover_letter_signals() should keep strategy-consulting roles out of technical-scoping cover language; got {signals.role_core_function!r}",
    )


def test_procore_bridge_hard_standard_outputs(
    build_resume: object,
    build_cover_letter: object,
    build_interview_cheat_sheet: object,
) -> None:
    summary = build_resume.build_problem_first_summary(PROCORE_JOB_DESCRIPTION, PROCORE_RESUME_TEXT)
    lowered_summary = summary.lower()
    assert_true(
        "technical scoping" in lowered_summary
        and "technical integration" in lowered_summary
        and "delivery leadership" in lowered_summary,
        f"Procore-style implementation summaries should make leadership, scoping, and integration explicit; got {summary}",
    )

    draft = build_cover_letter.compose_cover_letter_draft(
        "Procore",
        "Senior Solutions Architect",
        PROCORE_JOB_DESCRIPTION,
        PROCORE_RESUME_TEXT,
    )
    body = " ".join(draft.body_paragraphs)
    lowered_body = body.lower()
    assert_true(
        "use your technical consulting" not in lowered_body,
        f"standard Procore cover letter should avoid imperative JD copy as the mission sentence; got {body}",
    )
    assert_true(
        ("technical scoping" in lowered_body or "statements of work" in lowered_body)
        and "integration" in lowered_body
        and "delivery" in lowered_body,
        f"standard Procore cover letter should foreground scoping, integration, and delivery proof; got {body}",
    )

    profile = build_resume.job_problem_profile(PROCORE_JOB_DESCRIPTION, PROCORE_RESUME_TEXT)
    story = build_interview_cheat_sheet.StoryCard(
        title="Implementation Recovery",
        story_types=("Individual Achievement", "Rapid Learning", "Teamwork", "Persuasion", "Analysis and Decision"),
        hook="A complex implementation needed scope clarity and steadier execution.",
        takeaways=("Clarified the workflow", "Aligned owners", "Protected delivery"),
        evidence="Led scoping, validation, and stakeholder alignment across migration and testing work.",
        level3_trait="Separated the visible symptom from the delivery risk behind it.",
        result="Protected migration stability and kept implementation work moving through testing and launch readiness.",
        outcome="The team had a clearer path to delivery and fewer downstream surprises.",
        evidence_terms=("implementation", "testing", "migration"),
        signals=("implementation", "integration", "testing", "delivery"),
    )
    common_answers = {
        item.prompt: item.answer
        for item in build_interview_cheat_sheet.common_interview_answers(
            profile,
            "Procore",
            "Senior Solutions Architect",
            PROCORE_JOB_DESCRIPTION,
            [story],
            resume_text=PROCORE_RESUME_TEXT,
        )
    }
    gap_answer = common_answers["What is a gap we should know about?"]
    assert_true(
        "adjacent rather than deep" in gap_answer.lower()
        and "my strongest depth" in gap_answer.lower()
        and story.title not in gap_answer,
        f"a direct gap question should get an honest boundary, supported strength, and natural proof; got {gap_answer}",
    )

    keyword_answers = {
        item.prompt: item.answer
        for item in build_interview_cheat_sheet.keyword_ready_answers(
            profile,
            PROCORE_JOB_DESCRIPTION,
            [story],
        )
    }
    integration_prompt = next((prompt for prompt in keyword_answers if prompt.startswith("Integrations:")), "")
    integration_answer = keyword_answers.get(integration_prompt, "")
    assert_true(
        integration_prompt != ""
        and "technical integration work" in integration_answer.lower()
        and "i would be honest" not in integration_answer.lower(),
        f"integration answer should bridge hard through supported implementation proof; got {integration_answer}",
    )


def test_standard_cover_mode(build_cover_letter: object) -> None:
    application_responses = (
        build_cover_letter.question_prep.QualificationsResponse(
            "Why are you interested in joining this company?",
            "I am interested in joining Ollie because the Data Analyst, Retention opportunity sits at the intersection of retention analytics and customer-facing execution, which is where my background is strongest.",
        ),
        build_cover_letter.question_prep.QualificationsResponse(
            "Briefly describe your relevant experience.",
            "I bring approximately 10+ years of related experience. At East West Manufacturing, I owned a mission-critical enterprise platform across five sites and 150+ users. At Aptean, I managed 80+ international client engagements and led discovery, data migration, testing, training, and post-go-live support.",
        ),
        build_cover_letter.question_prep.QualificationsResponse(
            "Please describe your communication, public speaking, or training experience.",
            "My demonstrated level in this area is advanced and directly supported by experience facilitating 60+ executive workshops and QBRs, translating technical work into business language, and designing role-based training and adoption guides.",
        ),
        build_cover_letter.question_prep.QualificationsResponse(
            "What else may uniquely qualify you for this role?",
            "What may uniquely qualify me is the combination of solution consulting, operational analytics, and execution discipline. That includes 80+ international client engagements, five-site systems ownership for 150+ users, and 200+ reporting tools.",
        ),
    )
    draft = build_cover_letter.compose_cover_letter_draft(
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
        mode="standard",
        application_responses=application_responses,
    )
    min_words, max_words = build_cover_letter.cover_letter_word_range("standard")
    letter_text = "\n".join(
        [
            draft.salutation,
            *draft.body_paragraphs,
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    assert_true(
        draft.salutation == "Dear Ollie Team,",
        f"standard cover draft should default to company-team salutation; got {draft.salutation!r}",
    )
    assert_true(
        draft.paragraph_shape in {2, 3} and 2 <= len(draft.body_paragraphs) <= 3,
        f"standard cover draft should stay in the new 2-3 paragraph shape; got shape={draft.paragraph_shape}, paragraphs={len(draft.body_paragraphs)}",
    )
    assert_true(
        min_words <= word_count(letter_text) <= max_words,
        f"standard cover draft should stay within the hard concise range; got {word_count(letter_text)} words",
    )
    assert_true(
        4 <= sum(len(build_cover_letter.sentence_list(paragraph)) for paragraph in draft.body_paragraphs) <= 6,
        f"standard cover draft should stay within the 4-6 sentence contract; got {draft.body_paragraphs!r}",
    )
    assert_true(
        "Ollie" in draft.body_paragraphs[0]
        and any(title in draft.body_paragraphs[0] for title in ("Data Analyst, Retention", "Data Analyst Retention")),
        f"standard cover draft should open with direct company-and-role context; got {draft.body_paragraphs[0]!r}",
    )
    assert_true(
        bool(re.search(r"200\+|60\+|80\+|150\+ users", letter_text)),
        "standard cover draft should include at least one quantified credential",
    )
    assert_true(
        "executive workshops" in letter_text.lower() or "80+ international client engagements" in letter_text,
        f"standard cover draft should materially reflect the application-question response bank; got {letter_text}",
    )
    assert_true(
        "I would welcome the chance to discuss" in letter_text
        and any(term in letter_text.lower() for term in ("retention needs", "retention and lifecycle data")),
        f"standard cover draft should close with a direct ask tied to the job's pain area; got {letter_text}",
    )
    assert_true(
        not re.search(r"--|\b(?:I'm|I'd|I've|you're|we're|can't|don't|won't|it's|that's)\b", letter_text),
        f"standard cover draft should avoid contractions and double-dashes; got {letter_text}",
    )
    assert_true(
        "I bring i bring" not in letter_text and not re.search(r"(^|[.!?]\s+)That\b", letter_text),
        f"standard cover draft should not duplicate first-person lead-ins or leave a sentence starting with 'That'; got {letter_text}",
    )
    assert_true(
        all(
            signal not in letter_text
            for signal in ("The role moves", "The work needs", "Together, that background shows", "Where the role reaches beyond")
        ),
        f"standard cover draft should not reuse the old abstract concise-cover phrases; got {letter_text}",
    )


def test_force_bridge_standard_cover_stays_natural(build_cover_letter: object) -> None:
    application_responses = (
        build_cover_letter.question_prep.QualificationsResponse(
            "Why are you interested in joining this company?",
            "I am interested in joining AlphaSense because the Implementation Consultant, Large Deployments role sits close to onboarding quality, deployment scale, and customer-facing execution.",
        ),
        build_cover_letter.question_prep.QualificationsResponse(
            "Briefly describe your relevant experience.",
            "I bring approximately 10+ years of related experience across enterprise systems, analytics, and customer-facing delivery.",
        ),
        build_cover_letter.question_prep.QualificationsResponse(
            "What else may uniquely qualify you for this role?",
            "What may uniquely qualify me is the combination of implementation delivery, analytics, and execution discipline. That includes 80+ international client engagements, five-site systems ownership for 150+ users, and 200+ reporting tools.",
        ),
    )
    plan = build_cover_letter.build_cover_letter_plan(
        "AlphaSense",
        "Implementation Consultant, Large Deployments",
        DUMMY_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
        mode="standard",
        application_responses=application_responses,
        force_bridge=True,
    )
    letter_text = "\n".join(
        [
            f"Christian Estrada | christianj1914@gmail.com | {build_cover_letter.build_resume.LINKEDIN_URL}",
            plan.salutation,
            *plan.body_paragraphs,
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    build_cover_letter.validate_cover_letter_shape(letter_text, mode="standard")
    specificity_warnings, cover_warnings = build_cover_letter.validate_cover_letter_text(
        letter_text,
        DUMMY_JOB_DESCRIPTION,
        "AlphaSense",
        mode="standard",
    )
    assert_true(
        len(plan.body_paragraphs) == 3,
        f"force_bridge standard covers should keep the three-paragraph shape; got {plan.body_paragraphs}",
    )
    assert_true(
        "I bring i bring" not in letter_text and not re.search(r"(^|[.!?]\s+)That\b", letter_text),
        f"force_bridge standard covers should not duplicate first-person openers or leak raw 'That' starters; got {letter_text}",
    )
    assert_true(
        not any("abrupt first-person switch" in warning for warning in cover_warnings),
        f"force_bridge standard covers should stay natural under prose checks; got {cover_warnings}",
    )
    assert_true(
        isinstance(specificity_warnings, list),
        "force_bridge standard covers should still return normal specificity warnings instead of failing shape or prose checks",
    )


def test_cover_letter_validator_blocks_generic_experience_summary(build_cover_letter: object) -> None:
    bad_text = "\n".join(
        [
            f"Christian Estrada | Atlanta, GA | christianj1914@gmail.com | {build_cover_letter.build_resume.LINKEDIN_URL}",
            "Dear Hiring Manager,",
            "I am interested in the Implementation Consultant role at Smoke Test Systems because the team needs steady implementation delivery and stakeholder alignment.",
            "I bring approximately 10+ years of related experience across enterprise systems, analytics, and customer-facing delivery.",
            "I would welcome the chance to discuss how I could support Smoke Test Systems and its implementation priorities.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    try:
        build_cover_letter.validate_cover_letter_text(
            bad_text,
            DUMMY_JOB_DESCRIPTION,
            "Smoke Test Systems",
            mode=build_cover_letter.STANDARD_COVER_MODE,
        )
    except SystemExit as error:
        assert_true(
            "I bring approximately" in str(error),
            f"validate_cover_letter_text() should explain the banned generic experience summary; got {error}",
        )
        return
    raise SmokeFailure("validate_cover_letter_text() should fail on the banned generic experience-summary opener")


def test_cover_letter_qc_rejects_lowercase_proof_paragraph(build_cover_letter: object) -> None:
    text = "\n".join(
        [
            "Christian Estrada | Atlanta, GA | christianj1914@gmail.com | linkedin.com/in/cjne",
            "Dear Smoke Test Systems Team,",
            "I am interested in the Implementation Consultant role at Smoke Test Systems because the team needs steady implementation delivery, stakeholder alignment, and go-live readiness.",
            "led 80+ client engagements, coordinated data migration and testing, and supported five-site operations for 150+ users.",
            "I would welcome the chance to discuss how I could support Smoke Test Systems and its implementation priorities.",
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    try:
        build_cover_letter.assert_cover_letter_qc(
            text,
            "Smoke Test Systems",
            "Implementation Consultant",
            DUMMY_JOB_DESCRIPTION,
            mode=build_cover_letter.STANDARD_COVER_MODE,
        )
    except SystemExit as error:
        assert_true(
            "malformed lowercase fragment" in str(error),
            f"assert_cover_letter_qc() should explain lowercase proof failures; got {error}",
        )
        return
    raise SmokeFailure("assert_cover_letter_qc() should fail when the proof paragraph starts with a lowercase fragment")


def test_cover_opening_names_company_specific_role_context(build_cover_letter: object) -> None:
    biotouch_job_description = """
Company: BioTouch
Job Title: Solutions Implementation Consultant

Summary:
The Solutions Implementation Consultant serves as a strategic partner to clients, guiding them through the implementation and adoption of BioTouch's proprietary cloud-based platform.
"""
    braven_job_description = """
Company: Braven
Job Title: Logistics Innovation Project Specialist

About Braven
Together, our ambition is to help rebuild the middle class and revitalize the American Dream.

About the Role
The Logistics Innovation Project Specialist role helps manage grading operations, stakeholder communication, and program-quality data.
This individual is responsible for managing a pool of part-time Graders and supporting volunteer communication.
"""
    biotouch_signals = build_cover_letter.build_cover_letter_signals(
        "BioTouch",
        "Solutions Implementation Consultant",
        biotouch_job_description,
        OLLIE_RESUME_TEXT,
    )
    braven_signals = build_cover_letter.build_cover_letter_signals(
        "Braven",
        "Logistics Innovation Project Specialist",
        braven_job_description,
        OLLIE_RESUME_TEXT,
    )
    biotouch_opening = build_cover_letter.standard_mission_paragraph(
        "BioTouch",
        "Solutions Implementation Consultant",
        biotouch_signals,
    )
    braven_opening = build_cover_letter.standard_mission_paragraph(
        "Braven",
        "Logistics Innovation Project Specialist",
        braven_signals,
    )
    assert_true(
        "BioTouch" in biotouch_opening and "Solutions Implementation Consultant" in biotouch_opening,
        f"Standard mission paragraph should name company and role for BioTouch-style postings; got {biotouch_opening!r}",
    )
    assert_true(
        "Summary." not in biotouch_opening,
        f"Standard mission paragraph should not leak section labels into the opening; got {biotouch_opening!r}",
    )
    assert_true(
        "Braven" in braven_opening and "Logistics Innovation Project Specialist" in braven_opening,
        f"Standard mission paragraph should name company and role for Braven-style postings; got {braven_opening!r}",
    )
    assert_true(
        "Together, its ambition" not in braven_opening,
        f"Standard mission paragraph should normalize possessive mission language; got {braven_opening!r}",
    )
    assert_true(
        any(term in braven_opening.lower() for term in ("grader", "program", "student", "volunteer")),
        f"Standard mission paragraph should carry an education/program context term for Braven-style postings; got {braven_opening!r}",
    )


def test_cover_opening_without_mission_still_names_company(build_cover_letter: object) -> None:
    standard_draft = build_cover_letter.compose_cover_letter_draft(
        "Department of Veterans Affairs",
        "IT Program Manager",
        FEDERAL_DUMMY_JOB_DESCRIPTION,
        PROCORE_RESUME_TEXT,
        mode="standard",
    )
    long_draft = build_cover_letter.compose_cover_letter_draft(
        "Department of Veterans Affairs",
        "IT Program Manager",
        FEDERAL_DUMMY_JOB_DESCRIPTION,
        PROCORE_RESUME_TEXT,
        mode="long",
    )
    standard_opening = standard_draft.body_paragraphs[0]
    long_opening = long_draft.body_paragraphs[0]
    assert_true(
        "Department of Veterans Affairs" in standard_opening and "IT Program Manager" in standard_opening,
        f"Mission-light openings should still name the federal agency and role in standard mode; got {standard_opening!r}",
    )
    assert_true(
        standard_opening.lower().startswith("i am interested in the it program manager role at department of veterans affairs"),
        f"Mission-light standard openings should lead with the company or agency context; got {standard_opening!r}",
    )
    assert_true(
        "Department of Veterans Affairs" in long_opening and "IT Program Manager" in long_opening,
        f"Mission-light openings should still name the federal agency and role in long mode; got {long_opening!r}",
    )
    assert_true(
        not long_opening.startswith("Agency."),
        f"Long-mode openings should not collapse to generic placeholder mission text; got {long_opening!r}",
    )


def test_education_assessment_cover_draft(build_cover_letter: object) -> None:
    signals = build_cover_letter.build_cover_letter_signals(
        "Pearson",
        "Lead Specialist, Measurement",
        PEARSON_MEASUREMENT_JOB_DESCRIPTION,
        PROCORE_RESUME_TEXT,
    )
    opening = build_cover_letter.standard_mission_paragraph(
        "Pearson",
        "Lead Specialist, Measurement",
        signals,
    )
    draft = build_cover_letter.compose_cover_letter_draft(
        "Pearson",
        "Lead Specialist, Measurement",
        PEARSON_MEASUREMENT_JOB_DESCRIPTION,
        PROCORE_RESUME_TEXT,
    )
    letter_text = "\n".join(["Dear Hiring Manager,", *draft.body_paragraphs, "Thank you,", "Christian Estrada"])
    assert_true(
        "Pearson" in opening and "Lead Specialist, Measurement" in opening,
        f"education-assessment openings should name the company and role; got {opening!r}",
    )
    assert_true(
        any(term in opening.lower() for term in ("assessment", "learning", "validation", "quality")),
        f"education-assessment openings should surface role context instead of JD fragments; got {opening!r}",
    )
    assert_true(
        signals.jd_pain_area == "AI-assisted workflow validation, quality monitoring, and continuous improvement",
        f"education-assessment closes should use a real pain area, not team names; got {signals.jd_pain_area!r}",
    )
    assert_true(
        build_cover_letter.lead_burial_check(letter_text) is None,
        f"education-assessment cover drafts should pass lead-burial checks; got {letter_text!r}",
    )


def test_guidehouse_cover_opening_uses_consulting_context(build_cover_letter: object) -> None:
    signals = build_cover_letter.build_cover_letter_signals(
        "Guidehouse",
        "ERP Assessment Leader",
        GUIDEHOUSE_ERP_ASSESSMENT_JOB_DESCRIPTION,
        PROCORE_RESUME_TEXT,
    )
    opening = build_cover_letter.standard_mission_paragraph(
        "Guidehouse",
        "ERP Assessment Leader",
        signals,
    )
    assert_true(
        "Guidehouse" in opening and "ERP Assessment Leader" in opening,
        f"Guidehouse openings should still name the company and role; got {opening!r}",
    )
    assert_true(
        "schools, educators, and learners" not in opening.lower(),
        f"Guidehouse openings should not leak education-assessment company context; got {opening!r}",
    )
    assert_true(
        any(term in opening.lower() for term in ("financial modernization", "roadmap", "stakeholder", "delivery", "commercial and government sectors")),
        f"Guidehouse openings should surface consulting and modernization context from the posting; got {opening!r}",
    )


def test_education_assessment_accomplishment_prefers_measurable_bullet(build_cover_letter: object) -> None:
    resume_text = """
Professional Summary
Enterprise systems and analytics consultant with 10+ years turning AI-assisted workflow questions and data-trust issues into clearer decisions, measurable process improvement, and usable reporting.

Professional Experience
Built 200+ dashboards and KPI reporting tools that gave leaders clearer visibility into operational performance and decision-making.
Protected migration stability by leading implementation readiness, scope alignment, sandbox testing, UAT validation, and targeted training across concurrent program tracks.

Education
Sample University
"""
    accomplishment = build_cover_letter.extract_top_accomplishment(
        PEARSON_MEASUREMENT_JOB_DESCRIPTION,
        resume_text,
    )
    assert_true(
        accomplishment.startswith("Built 200+ dashboards and KPI reporting tools"),
        "education-assessment proof selection should prefer the measurable analytics bullet over the resume summary",
    )


def test_long_cover_mode(build_cover_letter: object) -> None:
    standard_draft = build_cover_letter.compose_cover_letter_draft(
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
        mode="standard",
    )
    draft = build_cover_letter.compose_cover_letter_draft(
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
        mode="long",
    )
    min_words, max_words = build_cover_letter.cover_letter_word_range("long")
    preferred_min, preferred_max = build_cover_letter.preferred_cover_word_range("long")
    letter_text = "\n".join(
        [
            draft.salutation,
            *draft.body_paragraphs,
            "Thank you for your time and consideration,",
            "Christian Estrada",
        ]
    )
    letter_words = build_cover_letter.word_count(letter_text)
    assert_true(
        len(draft.body_paragraphs) in {4, 5} and draft.paragraph_shape == len(draft.body_paragraphs),
        f"long cover draft should stay on the shared proof-first path with a stable paragraph count; got shape={draft.paragraph_shape}, paragraphs={len(draft.body_paragraphs)}",
    )
    assert_true(
        min_words <= letter_words <= max_words,
        f"long cover draft should stay within the configured long range; got {letter_words} words",
    )
    assert_true(
        preferred_min <= letter_words <= preferred_max,
        f"long cover draft should stay within the preferred long band; got {letter_words} words",
    )
    assert_true(
        "I am excited" not in draft.body_paragraphs[0] and "exciting opportunity" not in draft.body_paragraphs[0].lower(),
        f"long cover draft should avoid filler openers; got {draft.body_paragraphs[0]!r}",
    )
    assert_true(
        "I would welcome the chance to discuss" in letter_text and "My closest match is" in letter_text,
        f"long cover draft should keep the honest bridge and direct close on the shared path; got {letter_text!r}",
    )
    assert_true(
        letter_words > build_cover_letter.word_count("\n".join(standard_draft.body_paragraphs)),
        "long cover draft should stay fuller than the default concise version",
    )


def test_cover_letter_inherits_fail_resume_name(build_cover_letter: object) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        temp_root = Path(temp_name)
        job_path = temp_root / "job_description.txt"
        job_path.write_text(DUMMY_JOB_DESCRIPTION.strip(), encoding="utf-8")
        output_dir = temp_root / "output"
        output_dir.mkdir()
        fail_resume = temp_root / "Christian Estrada - Smoke Test Systems - Implementation Consultant FAIL Resume.docx"
        fail_resume.write_text("resume placeholder", encoding="utf-8")

        original_job = build_cover_letter.JOB_DESCRIPTION
        original_output = build_cover_letter.OUTPUT_DIR
        original_find_resume = build_cover_letter.find_resume_output
        original_extract_role = build_cover_letter.extract_role_title
        original_helper = build_cover_letter.build_cover_letter_for_inputs
        captured: dict[str, object] = {}

        try:
            build_cover_letter.JOB_DESCRIPTION = job_path
            build_cover_letter.OUTPUT_DIR = output_dir
            build_cover_letter.find_resume_output = lambda _job_description: fail_resume
            build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_cover_letter.build_cover_letter_for_inputs = lambda **kwargs: captured.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
                bullets_used=3,
                audit_status="FAIL",
                specificity_warnings=[],
                mode=kwargs["mode"],
            )

            build_cover_letter.build_cover_letter()
        finally:
            build_cover_letter.JOB_DESCRIPTION = original_job
            build_cover_letter.OUTPUT_DIR = original_output
            build_cover_letter.find_resume_output = original_find_resume
            build_cover_letter.extract_role_title = original_extract_role
            build_cover_letter.build_cover_letter_for_inputs = original_helper

        assert_true(
            " FAIL Cover Letter.docx" in captured["output_docx"].name,
            f"build_cover_letter() should propagate FAIL naming from the matched resume; got {captured['output_docx']}",
        )


def test_cover_letter_inherits_bridge_resume_name(build_cover_letter: object) -> None:
    with TemporaryDirectory(prefix="resume_smoke_") as temp_name:
        temp_root = Path(temp_name)
        job_path = temp_root / "job_description.txt"
        job_path.write_text(DUMMY_JOB_DESCRIPTION.strip(), encoding="utf-8")
        output_dir = temp_root / "output"
        output_dir.mkdir()
        bridge_resume = temp_root / "Christian Estrada - Smoke Test Systems - Implementation Consultant BRIDGE Resume.docx"
        bridge_resume.write_text("resume placeholder", encoding="utf-8")

        original_job = build_cover_letter.JOB_DESCRIPTION
        original_output = build_cover_letter.OUTPUT_DIR
        original_find_resume = build_cover_letter.find_resume_output
        original_extract_role = build_cover_letter.extract_role_title
        original_helper = build_cover_letter.build_cover_letter_for_inputs
        captured: dict[str, object] = {}

        try:
            build_cover_letter.JOB_DESCRIPTION = job_path
            build_cover_letter.OUTPUT_DIR = output_dir
            build_cover_letter.find_resume_output = lambda _job_description: bridge_resume
            build_cover_letter.extract_role_title = lambda _job_description: "Implementation Consultant"
            build_cover_letter.build_cover_letter_for_inputs = lambda **kwargs: captured.update(kwargs) or SimpleNamespace(
                company_name=kwargs["company_name"],
                role_title=kwargs["role_title"],
                resume_docx=kwargs["resume_docx"],
                output_docx=kwargs["output_docx"],
                bullets_used=3,
                audit_status="BRIDGE",
                specificity_warnings=[],
                mode=kwargs["mode"],
            )

            build_cover_letter.build_cover_letter()
        finally:
            build_cover_letter.JOB_DESCRIPTION = original_job
            build_cover_letter.OUTPUT_DIR = original_output
            build_cover_letter.find_resume_output = original_find_resume
            build_cover_letter.extract_role_title = original_extract_role
            build_cover_letter.build_cover_letter_for_inputs = original_helper

        assert_true(
            " BRIDGE Cover Letter.docx" in captured["output_docx"].name,
            f"build_cover_letter() should propagate BRIDGE naming from the matched resume; got {captured['output_docx']}",
        )


def test_ollie_cover_acceptance(build_cover_letter: object) -> None:
    standard_draft = build_cover_letter.compose_cover_letter_draft(
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
        mode="standard",
    )
    long_draft = build_cover_letter.compose_cover_letter_draft(
        "Ollie",
        "Data Analyst, Retention",
        OLLIE_ANALYTICS_JOB_DESCRIPTION,
        OLLIE_RESUME_TEXT,
        mode="long",
    )
    standard_text = "\n".join(standard_draft.body_paragraphs)
    long_text = "\n".join(long_draft.body_paragraphs)
    assert_true(
        "The role moves" not in standard_text and "The work needs" not in standard_text,
        "Ollie standard cover acceptance should reject the older abstract concise-cover phrasing",
    )
    assert_true(
        "Ollie" in standard_text
        and any(title in standard_text for title in ("Data Analyst, Retention", "Data Analyst Retention"))
        and "retention and lifecycle data" in standard_text,
        "Ollie standard cover acceptance should keep direct company-role context and a pain-specific close",
    )
    assert_true(
        "My closest match is" in long_text and "retention and lifecycle data" in long_text,
        "Ollie long cover acceptance should keep the fuller honest-bridge path without the removed legacy mission builder",
    )
    assert_true(
        word_count(long_text) > word_count(standard_text),
        "Ollie long cover acceptance should stay fuller than the default concise version",
    )


def test_cover_letter_validator_blocks_contractions_and_double_dashes(build_cover_letter: object) -> None:
    for bad_text, label in (
        (
            "\n".join(
                [
                    "Christian Estrada | Atlanta, GA | christianj1914@gmail.com | linkedin.com/in/cjne",
                    "Dear Hiring Manager,",
                    "I am interested in the Implementation Consultant role at Acme because it sits close to implementation delivery.",
                    "I have managed 80+ client engagements and I would welcome the chance to discuss the role because it's a clear fit.",
                    "Thank you for your time and consideration,",
                    "Christian Estrada",
                ]
            ),
            "contraction",
        ),
        (
            "\n".join(
                [
                    "Christian Estrada | Atlanta, GA | christianj1914@gmail.com | linkedin.com/in/cjne",
                    "Dear Hiring Manager,",
                    "I am interested in the Implementation Consultant role at Acme because it sits close to implementation delivery -- and the work aligns with my background.",
                    "I have managed 80+ client engagements and would welcome the chance to discuss the role.",
                    "Thank you for your time and consideration,",
                    "Christian Estrada",
                ]
            ),
            "double-dash",
        ),
    ):
        try:
            build_cover_letter.validate_cover_letter_text(
                bad_text,
                DUMMY_JOB_DESCRIPTION,
                "Smoke Test Systems",
                mode=build_cover_letter.STANDARD_COVER_MODE,
            )
        except SystemExit:
            continue
        raise SmokeFailure(f"validate_cover_letter_text() should fail on a {label} in standard cover mode")


def test_clean_but_bad_cover_regressions_fail_writing_eval(writing_eval: object) -> None:
    samples = (
        (
            Path("evals/examples/bad/Christian Estrada - Acme Software - Solutions Implementation Consultant Cover Letter.docx"),
            "cover_letter_opening",
            "role_summary_opening",
        ),
        (
            Path("evals/examples/bad/Christian Estrada - AlphaSense - Implementation Consultant, Large Deployments BRIDGE Cover Letter.docx"),
            "cover_letter_full",
            "abstract_role_moves",
        ),
        (
            Path("evals/examples/bad/Christian Estrada - BlackLine - Principal Presales Consultant BRIDGE Cover Letter.docx"),
            "cover_letter_full",
            "where_role_reaches_beyond",
        ),
    )
    for path, section, expected_code in samples:
        text = writing_eval.extract_docx_text(path, section)
        result = writing_eval.evaluate_text(section, text, sample_id=path.name)
        assert_true(not result.passed, f"{path.name} should fail writing_eval after the concise-cover reset")
        assert_true(
            any(issue.code == expected_code for issue in result.issues),
            f"{path.name} should flag {expected_code}; got {[issue.code for issue in result.issues]}",
        )


def test_registered_firm_profile_requires_real_bain_name(build_cover_letter: object) -> None:
    toast_match = build_cover_letter._registered_firm_profile(
        "Toast",
        "Senior Manager, Care Experience Management",
        (
            "Lead experience management and cross-functional delivery for care workflows. "
            "Partner with consulting operations, training, and support teams."
        ),
    )
    bain_match = build_cover_letter._registered_firm_profile(
        "Bain & Company",
        "Consultant",
        "Management consulting role with client delivery and executive alignment.",
    )
    assert_true(
        toast_match is None,
        f"_registered_firm_profile() should not treat generic consulting signals inside a non-consulting JD as Bain; got {toast_match}",
    )
    assert_true(
        bool(bain_match and bain_match[0] == "bain"),
        f"_registered_firm_profile() should still match actual Bain company names; got {bain_match}",
    )


def test_bigfour_cover_opening_avoids_aspiration_phrase(build_cover_letter: object) -> None:
    opening = build_cover_letter.bain_opening("Consultant", "Bain & Company")
    assert_true(
        "I want to do more of" not in opening,
        f"bigfour consulting opening should avoid the banned aspiration phrase; got {opening}",
    )


def test_reset_jobs_helpers(reset_jobs: object) -> None:
    with TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "debrief_history.txt"
        path.write_text(
            "POST-INTERVIEW DEBRIEF CAPTURED 2026-01-01 09:00:00\n\n"
            "POST-INTERVIEW DEBRIEF CAPTURED 2026-01-02 09:00:00\n",
            encoding="utf-8",
        )
        count = reset_jobs.debrief_entry_count(path)
    assert_true(count == 2, "reset_jobs.debrief_entry_count() should count captured debrief entries")


def test_job_context_archive_active_snapshot_includes_questions() -> None:
    import job_context_archive

    original_job_description = job_context_archive.JOB_DESCRIPTION
    original_application_questions = job_context_archive.APPLICATION_QUESTIONS
    original_jobs_dir = job_context_archive.JOBS_DIR
    original_archive_root = job_context_archive.SCRATCH_JD_LIBRARY
    original_index_path = job_context_archive.INDEX_PATH
    original_sync_complete = job_context_archive._SYNC_COMPLETE
    try:
        with TemporaryDirectory(prefix="jd_archive_") as temp_name:
            temp_root = Path(temp_name)
            jobs_dir = temp_root / "jobs"
            jobs_dir.mkdir()
            archive_root = temp_root / "scratch" / "jd_library"
            job_path = jobs_dir / "job_description.txt"
            questions_path = jobs_dir / "application_questions.txt"
            job_path.write_text(DUMMY_JOB_DESCRIPTION.strip(), encoding="utf-8")
            questions_path.write_text(
                "Application Questions\n\nBriefly describe your relevant experience.\n\nWhat else qualifies you?\n",
                encoding="utf-8",
            )

            job_context_archive.JOBS_DIR = jobs_dir
            job_context_archive.JOB_DESCRIPTION = job_path
            job_context_archive.APPLICATION_QUESTIONS = questions_path
            job_context_archive.SCRATCH_JD_LIBRARY = archive_root
            job_context_archive.INDEX_PATH = archive_root / "index.csv"
            job_context_archive._SYNC_COMPLETE = True

            snapshot = job_context_archive.archive_active_context(
                workflow_type="commercial",
                source_command="smoke-test",
                archive_reason="unit_test",
            )

            snapshot_exists = snapshot.path.exists()
            job_copy_exists = (snapshot.path / "job_description.txt").exists()
            questions_copy_exists = (snapshot.path / "application_questions.txt").exists()
            metadata = job_context_archive.metadata_for_snapshot(snapshot.snapshot_id)
            index_rows = job_context_archive.read_index()
    finally:
        job_context_archive.JOB_DESCRIPTION = original_job_description
        job_context_archive.APPLICATION_QUESTIONS = original_application_questions
        job_context_archive.JOBS_DIR = original_jobs_dir
        job_context_archive.SCRATCH_JD_LIBRARY = original_archive_root
        job_context_archive.INDEX_PATH = original_index_path
        job_context_archive._SYNC_COMPLETE = original_sync_complete

    assert_true(snapshot_exists, "archive_active_context() should create the snapshot folder")
    assert_true(
        job_copy_exists,
        "archive_active_context() should store job_description.txt in the snapshot folder",
    )
    assert_true(
        questions_copy_exists,
        "archive_active_context() should store application_questions.txt when active questions are present",
    )
    assert_true(
        metadata.get("questions_present") is True and int(metadata.get("question_count", 0)) == 2,
        f"archive metadata should record active application questions; got {metadata}",
    )
    assert_true(
        any(row.get("snapshot_id") == snapshot.snapshot_id for row in index_rows),
        "archive_active_context() should register the snapshot in index.csv",
    )


def test_cover_letter_trace_records_snapshot_and_selection_debug(build_cover_letter: object) -> None:
    application_responses = (
        build_cover_letter.question_prep.QualificationsResponse(
            "Why are you interested in joining this company?",
            "I am interested in joining micro1 because the Professional Services Consultant role sits close to structured client guidance, analytical problem solving, and stakeholder communication.",
        ),
        build_cover_letter.question_prep.QualificationsResponse(
            "Briefly describe your relevant experience.",
            "At Aptean, I managed 80+ international client engagements and led discovery, data migration, testing, training, and post-go-live support. At East West Manufacturing, I owned a five-site platform serving 150+ users and built reporting used in daily decision making.",
        ),
    )
    original_trace_dir = build_cover_letter.COVER_TRACE_DIR
    original_snapshot_id = os.environ.get(build_cover_letter.job_context_archive.SNAPSHOT_ID_ENV)
    original_metadata_reader = build_cover_letter.job_context_archive.current_snapshot_metadata
    try:
        with TemporaryDirectory(prefix="cover_trace_") as temp_name:
            temp_root = Path(temp_name)
            build_cover_letter.COVER_TRACE_DIR = temp_root
            os.environ[build_cover_letter.job_context_archive.SNAPSHOT_ID_ENV] = "smoke_snapshot"
            build_cover_letter.job_context_archive.current_snapshot_metadata = lambda: {
                "snapshot_id": "smoke_snapshot",
                "company": "micro1",
                "question_count": 2,
            }
            plan = build_cover_letter.build_cover_letter_plan(
                "micro1",
                "Professional Services Consultant",
                MICRO1_SMOKE_JOB_DESCRIPTION,
                OLLIE_RESUME_TEXT,
                mode=build_cover_letter.STANDARD_COVER_MODE,
                application_responses=application_responses,
            )
            draft = build_cover_letter.compose_cover_letter_from_plan(plan)
            trace_path = build_cover_letter.write_cover_letter_trace(
                plan,
                draft,
                temp_root / "Christian Estrada - micro1 - Professional Services Consultant Cover Letter.docx",
                body_text="\n".join([draft.salutation, *draft.body_paragraphs]),
            )
            payload = json.loads(trace_path.read_text(encoding="utf-8"))
    finally:
        build_cover_letter.COVER_TRACE_DIR = original_trace_dir
        build_cover_letter.job_context_archive.current_snapshot_metadata = original_metadata_reader
        if original_snapshot_id is None:
            os.environ.pop(build_cover_letter.job_context_archive.SNAPSHOT_ID_ENV, None)
        else:
            os.environ[build_cover_letter.job_context_archive.SNAPSHOT_ID_ENV] = original_snapshot_id

    assert_true(
        payload["job_context_snapshot"]["snapshot_id"] == "smoke_snapshot",
        f"write_cover_letter_trace() should record the active snapshot ID; got {payload['job_context_snapshot']}",
    )
    assert_true(
        payload["trace_summary"]["snapshot_id"] == "smoke_snapshot",
        f"write_cover_letter_trace() should repeat the snapshot ID in the trace summary; got {payload['trace_summary']}",
    )
    assert_true(
        payload["plan"]["selection_debug"]["paragraph_purposes"][1]["purpose"] == "proof",
        f"write_cover_letter_trace() should preserve paragraph-purpose metadata; got {payload['plan']['selection_debug']}",
    )
    assert_true(
        bool(payload["plan"]["selection_debug"]["proof_candidates"]),
        "write_cover_letter_trace() should preserve the ranked proof candidates for debugging",
    )


def test_run_level_erp_scrub_preserves_formatting(build_resume: object) -> None:
    from docx import Document
    from docx.shared import RGBColor

    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("ERP implementation")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    cleaned = build_resume.scrub_erp_language_for_non_erp_text(
        run.text,
        "Customer success role focused on adoption, reporting, renewal risk, and workflow improvement.",
    )
    if cleaned != run.text:
        run.text = cleaned

    assert_true(run.text == "software implementation", f"ERP scrub did not update run text as expected: {run.text}")
    assert_true(run.bold is True, "run-level ERP scrub should preserve bold formatting")
    assert_true(run.font.color.rgb == RGBColor(0x1F, 0x4E, 0x79), "run-level ERP scrub should preserve font color")


def test_supported_rewrite_scrubs_erp_language(build_resume: object) -> None:
    text = "Owned ERP migration from Aptean Intuitive to Epicor Kinetic for customer onboarding readiness."
    cleaned = build_resume.rewrite_supported_text(
        text,
        "Healthcare implementation role focused on onboarding, workflow automation, and stakeholder adoption.",
    )
    assert_true("ERP" not in cleaned, f"rewrite_supported_text() should scrub ERP for non-ERP roles: {cleaned}")
    assert_true(
        "Aptean Intuitive" in cleaned and "Epicor Kinetic" in cleaned,
        f"rewrite_supported_text() should preserve specific platform names for supported proof: {cleaned}",
    )


def test_named_erp_platform_scrub_clears_summary_audit(build_resume: object) -> None:
    samples = (
        "Led a five-site Aptean Intuitive rollout that reduced manual inventory work by 78%.",
        "Drove adoption of the Aptean Encompix platform across five manufacturing sites.",
        "Managed an Oracle ERP migration and Microsoft Dynamics 365 rollout for global teams.",
        "Implemented SAP across multiple business units while preserving SAP Crystal Reports dashboards.",
        "Owned the Epicor Kinetic rollout end to end.",
    )
    for sample in samples:
        cleaned = build_resume.scrub_named_erp_platforms_for_summary(sample)
        assert_true(
            build_resume.NON_ERP_VISIBLE_PLATFORM_PATTERN.search(cleaned) is None,
            f"scrub_named_erp_platforms_for_summary() left named ERP-platform language in: {cleaned!r}",
        )
    sap_crystal = build_resume.scrub_named_erp_platforms_for_summary(
        "Built dashboards using SAP Crystal Reports for executive reporting."
    )
    assert_true(
        "SAP Crystal Reports" in sap_crystal,
        f"scrub_named_erp_platforms_for_summary() should not strip SAP Crystal Reports: {sap_crystal!r}",
    )


def test_supported_text_keeps_named_platforms_outside_summary(build_resume: object) -> None:
    text = "Owned ERP migration from Aptean Intuitive to Epicor Kinetic for customer onboarding readiness."
    cleaned = build_resume.rewrite_supported_text(
        text,
        "Healthcare implementation role focused on onboarding, workflow automation, and stakeholder adoption.",
    )
    assert_true(
        "Aptean Intuitive" in cleaned and "Epicor Kinetic" in cleaned,
        f"rewrite_supported_text() should still preserve platform names for bullet-level supported proof: {cleaned}",
    )


def test_non_erp_audit_allows_sap_crystal_reports(build_resume: object) -> None:
    try:
        build_resume.assert_no_erp_language_for_non_erp_role(
            "Systems, Data, and Analytics: SQL | SAP Crystal Reports | Power BI",
            "Consulting role focused on analysis, stakeholder alignment, and executive communication.",
            label="resume",
        )
    except Exception as error:  # noqa: BLE001
        raise SmokeFailure(f"SAP Crystal Reports should not trigger the non-ERP blocker: {error}") from error


def test_company_profile_stub(build_resume: object) -> None:
    assert_true(
        build_resume.detect_company_profile("Salesforce", "") is None,
        "detect_company_profile() should remain a no-op stub until implemented",
    )


def test_hiring_manager_skim_lane_terms(build_resume: object) -> None:
    profiles = (
        build_resume.JobProblemProfile(
            primary_lane="process_improvement",
            lane_label="Process Improvement",
            core_problem="process improvement, root cause, workflow, and operational metrics",
            audience="operations leaders",
            outcomes=("efficiency",),
            direct_matches=(),
            adjacent_matches=(),
            unsupported_requirements=(),
            safe_terms=(),
        ),
        build_resume.JobProblemProfile(
            primary_lane="corporate_strategy",
            lane_label="Corporate Strategy",
            core_problem="strategy, recommendations, business case, and decision quality",
            audience="business leaders",
            outcomes=("stakeholder alignment",),
            direct_matches=(),
            adjacent_matches=(),
            unsupported_requirements=(),
            safe_terms=(),
        ),
    )
    summaries = (
        "Process improvement consultant with 10+ years improving workflow, root cause analysis, operational metrics, and efficiency across five sites.",
        "Corporate strategy analyst with 10+ years turning recommendations, business cases, and stakeholder alignment into better decisions across five sites.",
    )
    for profile, summary in zip(profiles, summaries):
        issues = build_resume.hiring_manager_skim_issues(summary, ["Delivered 80+ client implementations and 200+ dashboards."], profile)
        assert_true(isinstance(issues, list), f"hiring_manager_skim_issues() did not return a list for {profile.primary_lane}")


def test_direct_report_reporting_line_does_not_trigger_people_leadership(build_resume: object) -> None:
    reporting_line_jd = (
        "Customer Experience Solutions Lead\n"
        "Direct report \u200b Enterprise Solutions Manager\n"
        "Own Salesforce backlog, UAT, and reporting for NAM business users."
    )
    reporting_line_profile = build_resume.job_problem_profile(reporting_line_jd)
    assert_true(
        "Direct People Leadership" not in reporting_line_profile.unsupported_requirements,
        "job_problem_profile() should not treat a reporting-line reference as people leadership ownership",
    )

    management_jd = (
        "Customer Experience Solutions Lead\n"
        "Lead a team of five analysts and manage direct reports across support and reporting."
    )
    management_profile = build_resume.job_problem_profile(management_jd)
    assert_true(
        "Direct People Leadership" in management_profile.unsupported_requirements,
        "job_problem_profile() should still flag explicit people-management requirements",
    )


def test_east_west_solution_architecture_summary_preserves_aptean_ownership(build_resume: object) -> None:
    current_summary = (
        "Stabilized mission-critical enterprise systems during company reorganization and platform migration, "
        "keeping manufacturing, supply chain, finance, and engineering teams aligned across five sites and 150+ "
        "users while owning Aptean Intuitive and supporting final cutover to Epicor Kinetic. Position impacted by "
        "company reorganization."
    )
    job_description = (
        "Customer Experience Solutions Lead role focused on enterprise solutions, solution architecture, "
        "deployment strategies, stakeholder alignment, and digital operations."
    )
    summary = build_resume.optimized_role_summary("East West Manufacturing", current_summary, job_description)
    assert_true(
        "Aptean Intuitive" in summary,
        f"optimized_role_summary() should keep Aptean Intuitive explicit for East West source truth; got {summary}",
    )
    assert_true(
        "light administration" not in summary.lower(),
        f"optimized_role_summary() should not downplay East West responsibility with 'light administration'; got {summary}",
    )
    assert_true(
        "Epicor Kinetic" in summary,
        f"optimized_role_summary() should still mention Epicor Kinetic migration support; got {summary}",
    )
    assert_true(
        any(term in summary.lower() for term in ("transition planning", "transition readiness", "launch readiness")),
        f"optimized_role_summary() should keep the Epicor transition work explicit without softening it; got {summary}",
    )


def test_aptean_customer_success_role_summary_passes_prose_check(build_resume: object) -> None:
    import prose_engine
    import writing_eval as _writing_eval

    current_summary = (
        "Delivered 12 full-lifecycle ERP implementations for 80+ manufacturing clients across the Americas, "
        "Europe, and Asia, managing scope, testing, training, and post-go-live support. "
        "Protected $6M+ client book through QBRs, adoption planning, and escalation recovery."
    )
    job_description = (
        "Digital Customer Success Manager role focused on customer adoption, retention, "
        "at-risk account recovery, and expansion conversations."
    )
    summary = build_resume.optimized_role_summary("Aptean", current_summary, job_description)
    repair = prose_engine.repair_text(summary, "summary")
    assert_true(
        repair.converged,
        f"Aptean customer_success role summary must pass prose repair without PROSE_NESTED_LIST; "
        f"got rule_ids={[f.rule_id for f in repair.findings if f.severity == 'fail']!r}, text={summary!r}",
    )
    for sent in summary.split(". "):
        sent = sent.strip().rstrip(".")
        if not sent:
            continue
        density_issue = _writing_eval.list_density_issue(sent)
        assert_true(
            density_issue is None,
            f"Aptean CS summary sentence must also pass cover-letter list_density_issue "
            f"(fires at comma_count >= 3); got issue={density_issue!r} for {sent!r}",
        )


def test_proof_first_close_uses_discuss_and_has_no_first_person_switch(build_cover_letter: object) -> None:
    import question_prep

    brief = question_prep.PositioningBrief(
        company_name="Bitwarden",
        role_title="Customer Success Manager",
        primary_lane="customer_success",
        employer_type="b2b_saas",
        mission_or_context="",
        role_core_problem="protecting customer outcomes at scale",
        role_problem_phrase="protecting customer outcomes at scale",
        personal_reason_to_care="",
        personal_reason_source="",
        strongest_direct_proofs=["customer success and account health management"],
        strongest_bridge_theme="",
        top_proof_anchors=[],
        company_specific_fact="",
        gap_honesty_boundary="",
        selected_proof_sentences=[],
    )
    close = build_cover_letter.proof_first_close_paragraph(brief, "Bitwarden", "Customer Success Manager")
    assert_true("discuss" in close.lower(), f"close must use 'discuss' for first-person-switch exemption; got {close!r}")
    assert_true("talk through" not in close.lower(), f"old 'talk through' phrasing must be gone; got {close!r}")
    bitwarden_jd = (
        "Company: Bitwarden\n\nJob Title: Customer Success Manager\n\n"
        "Bitwarden is looking for a Customer Success Manager to drive product adoption, "
        "renewal, and expansion for enterprise accounts. You will conduct QBRs, "
        "manage at-risk account health, and partner with stakeholders to ensure retention."
    )
    linkedin_url = build_cover_letter.build_resume.LINKEDIN_URL
    opening = (
        "Bitwarden is growing its enterprise customer base. The Customer Success Manager role centers on "
        "protecting customer outcomes at scale."
    )
    proof = (
        "Managed customer success for 80+ manufacturing clients globally, protecting a $6M+ book of business "
        "through QBRs and adoption planning."
    )
    letter_text = f"{linkedin_url}\n\nDear Bitwarden Team,\n\n{opening}\n\n{proof}\n\n{close}"
    _, cover_warnings = build_cover_letter.validate_cover_letter_text(
        letter_text, bitwarden_jd, "Bitwarden"
    )
    switch_warnings = [w for w in cover_warnings if "abrupt first-person switch" in w]
    assert_true(not switch_warnings, f"proof-first close must not trigger first-person-switch warning; got {switch_warnings}")


def test_cs_enterprise_jd_extracts_cs_terms_not_analytics_fallback(build_cover_letter: object) -> None:
    bitwarden_jd = (
        "Company: Bitwarden\n\nJob Title: Customer Success Manager\n\n"
        "We are looking for a Customer Success Manager to drive product adoption, "
        "renewal, and expansion for our enterprise accounts. You will conduct QBRs, "
        "manage at-risk accounts, and partner with stakeholders across sales and product."
    )
    terms = build_cover_letter.extract_cover_letter_terms(bitwarden_jd, "customer_success", {})
    assert_true(
        "adoption" in terms,
        f"enterprise CS JD must extract 'adoption' from CS enterprise patterns; got {terms}",
    )
    assert_true(
        "data analysis" not in terms,
        f"enterprise CS JD must not fall back to generic 'data analysis' label; got {terms}",
    )


def test_role_bullet_budgets_meet_minimums(build_resume: object) -> None:
    job_description = (
        "Implementation project manager role focused on configuration, data migration, "
        "go-live readiness, customer onboarding, analytics, reporting, and adoption."
    )
    for company, minimum in build_resume.MIN_FINAL_BULLETS_BY_COMPANY.items():
        budget = build_resume.role_bullet_budget(company, "Representative Role", job_description)
        assert_true(
            budget >= minimum,
            f"role_bullet_budget() returned {budget} for {company}, below minimum {minimum}",
        )


def test_contact_constants(build_resume: object) -> None:
    assert_true(
        isinstance(build_resume.CONTACT_EMAIL, str) and "@" in build_resume.CONTACT_EMAIL and build_resume.CONTACT_EMAIL.strip(),
        "CONTACT_EMAIL must be a non-empty email string",
    )
    assert_true(
        isinstance(build_resume.LINKEDIN_URL, str) and "." in build_resume.LINKEDIN_URL and build_resume.LINKEDIN_URL.strip(),
        "LINKEDIN_URL must be a non-empty URL string",
    )


def test_job_title_label_stripping(build_resume: object) -> None:
    assert_true(
        build_resume.extract_job_title("Company: BioTouch\nJob Title: Solutions Implementation Consultant\nSummary: x")
        == "Solutions Implementation Consultant",
        "extract_job_title() should strip a Job Title: label",
    )
    assert_true(
        build_resume.extract_job_title("Company: BioTouch\nJob Title Solutions Implementation Consultant\nSummary: x")
        == "Solutions Implementation Consultant",
        "extract_job_title() should strip a Job Title label without a colon",
    )
    assert_true(
        build_resume.extract_job_title("Company: BioTouch\nRole: Solutions Implementation Consultant\nSummary: x")
        == "Solutions Implementation Consultant",
        "extract_job_title() should strip a Role: label",
    )


def test_contains_search_term_handles_simple_plural_forms(build_resume: object) -> None:
    assert_true(
        build_resume.contains_search_term("Leaders and stakeholders used dashboards to support recommendations.", "stakeholder"),
        "contains_search_term() should treat simple plural stakeholder forms as hits",
    )
    assert_true(
        build_resume.contains_search_term("Leaders and stakeholders used dashboards to support recommendations.", "dashboard"),
        "contains_search_term() should treat simple plural dashboard forms as hits",
    )
    assert_true(
        build_resume.contains_search_term("Leaders and stakeholders used dashboards to support recommendations.", "recommendation"),
        "contains_search_term() should treat simple plural recommendation forms as hits",
    )
    assert_true(
        build_resume.contains_search_term("Implementation processes improved after launch.", "process"),
        "contains_search_term() should treat -es plural process forms as hits",
    )
    assert_true(
        build_resume.contains_search_term("Customer businesses expanded after go-live.", "business"),
        "contains_search_term() should treat -es plural business forms as hits",
    )


def test_business_context_module(business_context: object) -> None:
    text = (
        "Company: BioTouch\n"
        "Job Title: Solutions Implementation Consultant\n"
        "This cloud-based platform supports B2B healthcare customers with 5-10 implementation projects, "
        "20-30 active customer accounts, APIs, automated file transfers, Jira, Azure, and HIPAA/PHI-sensitive workflows."
    )
    context = business_context.extract_business_context(text)
    assert_true(context.business_model == "cloud software / SaaS", "business context should detect cloud software / SaaS")
    assert_true("B2B" in context.customer_type, "business context should detect B2B customer type")
    assert_true("Jira" in context.technical_stack and "Azure" in context.technical_stack, "business context should detect technical stack")
    questions = business_context.business_interview_questions(text)
    assert_true(any("implementations usually lose momentum" in item.question for item in questions), "business interview questions should include implementation-risk questions")
    warnings = business_context.business_context_audit("Strategic results-driven professional.", text, "smoke")
    assert_true(any("too little objective business context" in warning for warning in warnings), "business context audit should flag generic low-context output")
    weak_questions = business_context.question_quality_warnings(["What is the culture like?"])
    assert_true(weak_questions, "question quality audit should flag weak culture-only questions")
    state_farm_text = (
        "Company: State Farm\n"
        "Job Title: Digital Experience Analytics Analyst\n"
        "Strong business acumen in insurance/financial services, including understanding of digital self-service capabilities.\n"
        "Proven collaboration skills to build strong working relationships with business partners and State Farm associates."
    )
    state_farm_context = business_context.extract_business_context(state_farm_text)
    assert_true(
        state_farm_context.industry == "financial services",
        f"business context should classify insurance postings as financial services; got {state_farm_context.industry!r}",
    )
    pearson_context = business_context.extract_business_context(PEARSON_MEASUREMENT_JOB_DESCRIPTION)
    assert_true(
        pearson_context.industry == "education / assessment",
        f"business context should classify Pearson-style postings as education / assessment; got {pearson_context.industry!r}",
    )
    assert_true(
        "educators" in pearson_context.customer_type and "learners" in pearson_context.customer_type,
        f"business context should surface education audience context; got {pearson_context.customer_type!r}",
    )
    pearson_questions = business_context.business_interview_questions(PEARSON_MEASUREMENT_JOB_DESCRIPTION)
    assert_true(
        any(item.signal == "education quality" for item in pearson_questions),
        "business interview questions should add an education-quality question for Pearson-style postings",
    )
    guidehouse_context = business_context.extract_business_context(GUIDEHOUSE_ERP_ASSESSMENT_JOB_DESCRIPTION)
    assert_true(
        guidehouse_context.industry != "education / assessment",
        f"business context should not treat ERP assessment roles as education/assessment just because the title contains 'assessment'; got {guidehouse_context.industry!r}",
    )
    assert_true(
        guidehouse_context.product_or_service != "assessment and learning systems",
        f"business context should not leak education-product context into Guidehouse ERP roles; got {guidehouse_context.product_or_service!r}",
    )
    m_tech_like_text = (
        "Company: M-Tech Systems\n"
        "Job Title: Implementation Specialist (Project Consultant)\n"
        "We are seeking an energetic, driven professional who thrives in dynamic, client-facing environments. "
        "The ideal candidate is a fast learner who can engage effectively with stakeholders ranging from executives to frontline users.\n"
        "This role contributes to SaaS implementation, user adoption, training, and go-live support."
    )
    m_tech_like_context = business_context.extract_business_context(m_tech_like_text)
    assert_true(
        m_tech_like_context.industry != "education / assessment",
        f"business context should not classify generic 'fast learner' implementation roles as education / assessment; got {m_tech_like_context.industry!r}",
    )
    clorox_like_text = (
        "Company: The Clorox Company\n"
        "Job Title: Senior Analyst, Global External Manufacturing\n"
        "This role supports external manufacturing operations with SAP expertise, process guidance, training, reporting, "
        "and continuous improvement across the network.\n"
        "We are proud to be in every corner of homes, schools, and offices through our beloved brands."
    )
    clorox_like_context = business_context.extract_business_context(clorox_like_text)
    assert_true(
        clorox_like_context.industry != "education / assessment",
        f"business context should not classify consumer/manufacturing postings as education / assessment just because the brand statement mentions schools; got {clorox_like_context.industry!r}",
    )
    assert_true(
        clorox_like_context.customer_type != "schools, educators, and learners",
        f"business context should not classify consumer/manufacturing postings as an education audience just because the brand statement mentions schools; got {clorox_like_context.customer_type!r}",
    )


def test_audit_keywords_filter_low_signal_quality_phrases(build_resume: object) -> None:
    keywords = build_resume.audit_keywords(PEARSON_MEASUREMENT_JOB_DESCRIPTION)
    assert_true(
        "compromising quality" not in keywords and "assessment quality" not in keywords and "validation quality" not in keywords,
        f"audit_keywords() should filter low-signal quality phrases instead of treating them as ATS terms; got {sorted(keywords)!r}",
    )
    assert_true(
        {"ai-assisted", "measurement", "validation"} <= keywords,
        f"audit_keywords() should keep real Pearson-style role signals; got {sorted(keywords)!r}",
    )
    assert_true(
        len(keywords) < 40,
        f"audit_keywords() should avoid keyword explosions on long education-assessment JDs; got {len(keywords)} keywords",
    )


def test_audit_keywords_filter_noisy_bigrams(build_resume: object) -> None:
    m_tech_jd = """
Company: M-Tech Systems
Job Title: Implementation Specialist (Project Consultant)
Drive assigned workstreams to completion, ensuring deliverables are completed on time, within scope, and at a high standard.
Participate in internal knowledge sharing and training to enhance team capabilities.
Participate in go-live activities, issue resolution, and user adoption post-implementation.
"""
    keywords = build_resume.audit_keywords(m_tech_jd)
    assert_true(
        "sharing training" not in keywords and "within scope" not in keywords and "contribute delivery" not in keywords,
        f"audit_keywords() should filter noisy JD bigrams before they reach placement audits; got {sorted(keywords)!r}",
    )
    assert_true(
        "user adoption" in keywords,
        f"audit_keywords() should keep supported implementation signals such as 'user adoption'; got {sorted(keywords)!r}",
    )


def test_guidehouse_audit_keywords_filter_bridge_noise(build_resume: object) -> None:
    keywords = build_resume.audit_keywords(GUIDEHOUSE_ERP_ASSESSMENT_JOB_DESCRIPTION)
    assert_true(
        "ability" not in keywords,
        f"audit_keywords() should filter low-value single-word fillers like 'ability'; got {sorted(keywords)!r}",
    )
    assert_true(
        "statu" not in keywords and "status" in keywords,
        f"audit_keywords() should preserve 'status' without broken singularization; got {sorted(keywords)!r}",
    )


def test_thank_you_contact_line_filter(build_thank_you: object) -> None:
    assert_true(
        build_thank_you.is_resume_contact_line("Atlanta, GA | 770-710-4216 | christianj1914@gmail.com | linkedin.com/in/cjne"),
        "thank-you proof extraction should identify resume contact lines",
    )
    assert_true(
        not build_thank_you.is_resume_contact_line("Built 200+ reporting tools that improved decision visibility for leaders."),
        "thank-you proof extraction should not reject normal proof points",
    )


def test_thank_you_proof_points_skip_summary(build_thank_you: object) -> None:
    from docx import Document

    with TemporaryDirectory(prefix="thank_you_resume_") as temp_name:
        resume_path = Path(temp_name) / "resume.docx"
        doc = Document()
        doc.add_paragraph("Christian Estrada")
        doc.add_paragraph("Atlanta, GA | 770-710-4216 | christianj1914@gmail.com | linkedin.com/in/cjne")
        doc.add_paragraph("PROFESSIONAL SUMMARY")
        doc.add_paragraph("Enterprise systems consultant with 10+ years turning implementation ambiguity into workflow clarity.")
        doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        doc.add_paragraph("Built 200+ dashboards that improved decision visibility for leaders.")
        doc.add_paragraph("Led 60+ executive workshops and QBRs across the Americas, Europe, and Asia.")
        doc.save(str(resume_path))

        proof_points = build_thank_you.proof_points_from_resume(resume_path)

    assert_true(
        proof_points and proof_points[0].startswith("Built 200+ dashboards"),
        f"proof_points_from_resume() should skip summary prose and pull concrete experience proof first; got {proof_points!r}",
    )


def test_company_constants(build_resume: object) -> None:
    for name in (
        "COMPANY_EAST_WEST",
        "COMPANY_APTEAN",
        "COMPANY_HOME_DEPOT",
        "COMPANY_ADERANT",
    ):
        value = getattr(build_resume, name)
        assert_true(isinstance(value, str) and value.strip(), f"{name} must be a non-empty string")
    assert_true(
        build_resume.normalize_compare(build_resume.COMPANY_EAST_WEST) == build_resume.normalize_compare("East West Manufacturing"),
        "COMPANY_EAST_WEST normalization should match optimized_role_summary comparisons",
    )


def test_role_heading_detection(build_resume: object) -> None:
    assert_true(
        build_resume.is_role_heading("Enterprise Systems Manager    March 2023 - Present"),
        "is_role_heading() should detect a role title followed by a month-year date",
    )
    assert_true(
        not build_resume.is_role_heading("Helped launch the Home Depot SMS pilot in March 2017 with cross-functional partners."),
        "is_role_heading() should not treat a dated bullet as a role heading",
    )
    assert_true(
        not build_resume.is_role_heading("March 2023 launch support across teams"),
        "is_role_heading() should not treat a short month-year phrase as a role heading",
    )


def test_collapse_redundant_role_blanks(build_resume: object) -> None:
    import resume_format
    from xml.etree import ElementTree as ET

    W = resume_format.W
    body = ET.Element(f"{W}body")
    section = ET.SubElement(body, f"{W}p")
    section_run = ET.SubElement(section, f"{W}r")
    section_text = ET.SubElement(section_run, f"{W}t")
    section_text.text = "Professional Experience"
    for _ in range(3):
        body.append(resume_format.make_separator_paragraph())
    role = ET.SubElement(body, f"{W}p")
    run = ET.SubElement(role, f"{W}r")
    text_node = ET.SubElement(run, f"{W}t")
    text_node.text = "Support Operations Analyst    November 2015 - November 2019"
    removed = resume_format.collapse_redundant_blank_paragraphs_before_role_headings(body)
    blank_count = sum(1 for child in body.findall(f"{W}p") if resume_format.is_blank_paragraph(child))
    assert_true(removed == 2, f"expected 2 redundant blanks removed, removed {removed}")
    assert_true(blank_count == 1, f"expected 1 blank before role heading, found {blank_count}")


def test_resume_experience_alignment(build_resume: object) -> None:
    w_ns = build_resume.W.strip("{}")
    document_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="{w_ns}">
  <w:body>
    <w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
    <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:t>Summary text.</w:t></w:r></w:p>
    <w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
    <w:p><w:pPr><w:jc w:val="both"/></w:pPr><w:r><w:t>Enterprise Systems Manager    March 2023 - Present</w:t></w:r></w:p>
    <w:p><w:pPr><w:jc w:val="both"/></w:pPr><w:r><w:t>Known Company | Knoxville, TN</w:t></w:r></w:p>
    <w:p><w:pPr><w:jc w:val="both"/></w:pPr></w:p>
    <w:p><w:pPr><w:jc w:val="both"/></w:pPr><w:r><w:t>Delivered a known role bullet with consistent alignment.</w:t></w:r></w:p>
    <w:p><w:r><w:t>Education</w:t></w:r></w:p>
    <w:p><w:pPr><w:jc w:val="both"/></w:pPr></w:p>
  </w:body>
</w:document>
"""
    with TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "document.xml"
        path.write_text(document_xml, encoding="utf-8")
        build_resume.apply_resume_alignment(path)
        root = ET.parse(path).getroot()
        paragraphs = root.findall(f".//{build_resume.W}p")

    def alignment(paragraph: object) -> str | None:
        jc = paragraph.find(f"{build_resume.W}pPr/{build_resume.W}jc")
        return jc.get(build_resume.w_attr("val")) if jc is not None else None

    role_heading, company_line, blank_separator, bullet, education_blank = paragraphs[3], paragraphs[4], paragraphs[5], paragraphs[6], paragraphs[8]
    assert_true(alignment(role_heading) == "left", "experience role heading should be explicitly left-aligned")
    assert_true(alignment(company_line) == "left", "experience company line should be explicitly left-aligned")
    assert_true(alignment(blank_separator) == "left", "blank separator inside Professional Experience should be explicitly left-aligned")
    assert_true(alignment(bullet) == "left", "experience bullet paragraph should be explicitly left-aligned")
    assert_true(alignment(education_blank) == "both", "blank separator outside Professional Experience should keep its existing alignment")


def test_writing_eval_flags_system_narration(writing_eval: object) -> None:
    text = (
        "Implementation consultant improving workflow clarity for enterprise software teams. "
        "Work includes automation reducing manual effort by 78% and reporting support across multiple sites. "
        "That experience also includes stakeholder alignment and data checks. "
        "This background fits teams that need cleaner decisions and stronger cross-functional execution."
    )
    result = writing_eval.evaluate_text("resume_summary", text, sample_id="templated_summary")
    codes = {issue.code for issue in result.issues}
    assert_true(not result.passed, "writing_eval should fail a summary that uses known system-narration patterns")
    assert_true(
        {"proof_opener_includes", "system_fit_closer", "sentence_starts_with_that", "summary_sentence_count"} <= codes,
        f"writing_eval should catch the known bad-summary patterns; got {sorted(codes)}",
    )


def test_writing_eval_passes_clean_summary(writing_eval: object) -> None:
    text = (
        "Implementation consultant improving data visibility, workflow clarity, and go-live readiness for enterprise software teams operating across complex customer and site environments. "
        "Delivered onboarding, reporting, and process improvements across 80+ clients, 150+ users, five sites, and 200+ KPI tools while reducing manual inventory work by 78% and discrepancies by 22% in a mission-critical ERP environment. "
        "Brings implementation, reporting, stakeholder alignment, and adoption experience across ERP, migration, dashboards, and process-improvement work for cleaner decisions, steadier launches, and durable process change."
    )
    result = writing_eval.evaluate_text("resume_summary", text, sample_id="clean_summary")
    assert_true(result.passed, f"writing_eval should pass a clean summary; got {[issue.code for issue in result.issues]}")
    assert_true(not result.issues, f"writing_eval should keep the clean sample issue-free; got {result.issues}")


def test_writing_eval_flags_weak_close_and_list_density(writing_eval: object) -> None:
    text = (
        "Hello. This background speaks directly to the role and covers implementation, reporting, stakeholder coordination, "
        "customer communication, testing, and training across complex programs. "
        "A conversation about where Acme needs faster early wins would be valuable."
    )
    result = writing_eval.evaluate_text("followup_email_body", text, sample_id="weak_followup")
    codes = {issue.code for issue in result.issues}
    assert_true(not result.passed, "writing_eval should fail weak bridge and close phrasing in outbound email copy")
    assert_true(
        {"direct_alignment_bridge", "weak_hypothetical_close", "list_density_overload"} <= codes,
        f"writing_eval should flag weak bridge, weak close, and list density issues; got {sorted(codes)}",
    )


def test_enforce_prose_quality_warn_mode_allows_prep_text(utils: object) -> None:
    sample = (
        "This background speaks directly to the role and covers implementation, reporting, stakeholder coordination, "
        "customer communication, testing, and training across complex programs."
    )
    stdout = io.StringIO()
    with contextlib.redirect_stdout(stdout):
        report = utils.enforce_prose_quality(
            sample,
            "interview_pitch",
            label="Interview pitch",
            mode="warn",
            check_template_leakage=False,
        )
    assert_true(not report["passed"], "warn-only prep prose should still report issues")
    assert_true(
        "PROSE WARNING" in stdout.getvalue(),
        f"warn-only prep prose should print explicit warnings; got {stdout.getvalue()!r}",
    )


def test_writing_eval_loads_file_backed_dataset(writing_eval: object) -> None:
    with TemporaryDirectory(prefix="writing_eval_") as temp_name:
        root = Path(temp_name)
        sample_file = root / "good_summary.txt"
        sample_file.write_text(
            "Implementation consultant improving data visibility, workflow clarity, and go-live readiness for enterprise software teams operating across complex customer and site environments. "
            "Delivered onboarding, reporting, and process improvements across 80+ clients, 150+ users, five sites, and 200+ KPI tools while reducing manual inventory work by 78% and discrepancies by 22% in a mission-critical ERP environment. "
            "Brings implementation, reporting, stakeholder alignment, and adoption experience across ERP, migration, dashboards, and process-improvement work for cleaner decisions, steadier launches, and durable process change.",
            encoding="utf-8",
        )
        dataset = root / "gold.jsonl"
        dataset.write_text(
            '{"id":"file_backed_summary","artifact":"resume_summary","file":"good_summary.txt","expected_outcome":"pass"}\n',
            encoding="utf-8",
        )
        samples = writing_eval.load_dataset(dataset)
        assert_true(len(samples) == 1, f"writing_eval.load_dataset() should return one file-backed sample; got {len(samples)}")
        assert_true(samples[0].source_path is not None, "writing_eval.load_dataset() should keep the resolved source path for file-backed samples")
        result = writing_eval.evaluate_text(samples[0].artifact, samples[0].text, samples[0].sample_id)
        assert_true(result.passed, f"file-backed good sample should pass writing_eval; got {[issue.code for issue in result.issues]}")


def build_minimal_docx(path: Path, paragraphs: list[str]) -> Path:
    body = "".join(f"<w:p><w:r><w:t>{escape(paragraph)}</w:t></w:r></w:p>" for paragraph in paragraphs)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}</w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("word/document.xml", document_xml)
    return path


def test_writing_eval_extracts_docx_sections(writing_eval: object) -> None:
    with TemporaryDirectory(prefix="writing_eval_docx_") as temp_name:
        temp_root = Path(temp_name)
        resume_docx = build_minimal_docx(
            temp_root / "resume.docx",
            [
                "Christian Estrada",
                "PROFESSIONAL SUMMARY",
                "Implementation consultant improving data visibility, workflow clarity, and go-live readiness for enterprise software teams operating across complex customer and site environments. Delivered onboarding, reporting, and process improvements across 80+ clients, 150+ users, five sites, and 200+ KPI tools while reducing manual inventory work by 78% and discrepancies by 22% in a mission-critical ERP environment. Brings implementation, reporting, stakeholder alignment, and adoption experience across ERP, migration, dashboards, and process-improvement work for cleaner decisions, steadier launches, and durable process change.",
                "PROFESSIONAL EXPERIENCE",
            ],
        )
        cover_docx = build_minimal_docx(
            temp_root / "cover.docx",
            [
                "Christian Estrada",
                "Dear Hiring Manager,",
                "Opening paragraph for the role.",
                "Proof paragraph with concrete scope and outcomes.",
                "Bridge paragraph for the role.",
                "Close paragraph asking for a conversation.",
                "Thank you for your time and consideration,",
                "Christian Estrada",
            ],
        )
        team_cover_docx = build_minimal_docx(
            temp_root / "team_cover.docx",
            [
                "Christian Estrada",
                "Dear Smoke Test Systems Team,",
                "Opening paragraph with company-specific context.",
                "Proof paragraph with concrete scope and outcomes.",
                "Close paragraph asking for a conversation.",
                "Thank you for your time and consideration,",
                "Christian Estrada",
            ],
        )
        assert_true(
            writing_eval.extract_docx_text(resume_docx, "resume_summary").startswith("Implementation consultant"),
            "writing_eval.extract_docx_text() should pull the resume summary from a DOCX file",
        )
        assert_true(
            writing_eval.extract_docx_text(cover_docx, "cover_letter_opening") == "Opening paragraph for the role.",
            "writing_eval.extract_docx_text() should pull the opening paragraph from a cover-letter DOCX file",
        )
        assert_true(
            writing_eval.extract_docx_text(cover_docx, "cover_letter_close") == "Close paragraph asking for a conversation.",
            "writing_eval.extract_docx_text() should pull the closing paragraph from a cover-letter DOCX file",
        )
        assert_true(
            writing_eval.extract_docx_text(team_cover_docx, "cover_letter_opening") == "Opening paragraph with company-specific context.",
            "writing_eval.extract_docx_text() should accept company-specific cover-letter salutations, not only 'Dear Hiring Manager'",
        )


def test_detailed_guide_notes_context_strips_leading_bullets(build_detailed_interview_guide: object) -> None:
    lines = build_detailed_interview_guide.notes_context(
        "POST-INTERVIEW NOTE [2026-06-12]: Smoke Test Systems is balancing modernization speed with compliance.",
        "Key signals:\n- Hiring manager cares about stakeholder alignment.\n- Team wants concise executive updates.\n",
    )
    assert_true(
        not any(line.startswith("-") for line in lines),
        f"notes_context() should strip leading bullet markers before the guide re-bullets them; got {lines}",
    )


def test_extract_writing_examples_writes_snippets(extract_writing_examples: object) -> None:
    with TemporaryDirectory(prefix="writing_extract_") as temp_name:
        temp_root = Path(temp_name)
        source_dir = temp_root / "good"
        source_dir.mkdir()
        docx_path = build_minimal_docx(
            source_dir / "Sample Resume.docx",
            [
                "Christian Estrada",
                "PROFESSIONAL SUMMARY",
                "Implementation consultant improving data visibility, workflow clarity, and go-live readiness for enterprise software teams operating across complex customer and site environments. Delivered onboarding, reporting, and process improvements across 80+ clients, 150+ users, five sites, and 200+ KPI tools while reducing manual inventory work by 78% and discrepancies by 22% in a mission-critical ERP environment. Brings implementation, reporting, stakeholder alignment, and adoption experience across ERP, migration, dashboards, and process-improvement work for cleaner decisions, steadier launches, and durable process change.",
                "PROFESSIONAL EXPERIENCE",
            ],
        )
        out_dir = temp_root / "out"
        written = extract_writing_examples.write_sections(docx_path, ("resume_summary",), out_dir)
        assert_true(len(written) == 1, f"extract_writing_examples.write_sections() should write one snippet; got {written}")
        content = written[0].read_text(encoding="utf-8").strip()
        assert_true(
            content.startswith("Implementation consultant"),
            f"extract_writing_examples.write_sections() should write the extracted summary text; got {content!r}",
        )
        assert_true(
            "good" in str(written[0].parent).lower(),
            f"extract_writing_examples.write_sections() should mirror good/bad labels in output folders; got {written[0]}",
        )


def test_track_row_for_active_job_prefers_tailored_resume_for_fit() -> None:
    import track_applications

    source_resume = Path("C:/temp/source_resume.docx")
    tailored_resume = Path("C:/temp/tailored_resume.docx")
    captured: dict[str, object] = {}

    original_read_job_description = track_applications.read_job_description
    original_extract_output_name = track_applications.resume_analysis.extract_output_name
    original_extract_job_title = track_applications.resume_analysis.extract_job_title
    original_choose_resume = track_applications.resume_analysis.choose_resume
    original_job_problem_profile = track_applications.resume_analysis.job_problem_profile
    original_latest_resume = track_applications.latest_resume
    original_active_job_fit_status = track_applications.active_job_fit_status
    original_build_resume = sys.modules.get("build_resume")
    try:
        track_applications.read_job_description = lambda: DUMMY_JOB_DESCRIPTION
        track_applications.resume_analysis.extract_output_name = lambda _job: "Smoke Test Systems"
        track_applications.resume_analysis.extract_job_title = lambda _job: "Implementation Consultant"
        track_applications.resume_analysis.choose_resume = lambda _job: source_resume
        track_applications.resume_analysis.job_problem_profile = (
            lambda _job, resume_text="": captured.setdefault("profile_resume_text", resume_text) and SimpleNamespace(lane_label="Implementation Delivery")
        )
        track_applications.latest_resume = lambda _job: tailored_resume
        track_applications.active_job_fit_status = lambda _job, resume_path: captured.setdefault("resume_path", resume_path) and "Adjacent Fit"
        sys.modules["build_resume"] = SimpleNamespace(docx_visible_text_from_path=lambda path: f"visible text from {path.name}")

        row = track_applications.row_for_active_job()
    finally:
        track_applications.read_job_description = original_read_job_description
        track_applications.resume_analysis.extract_output_name = original_extract_output_name
        track_applications.resume_analysis.extract_job_title = original_extract_job_title
        track_applications.resume_analysis.choose_resume = original_choose_resume
        track_applications.resume_analysis.job_problem_profile = original_job_problem_profile
        track_applications.latest_resume = original_latest_resume
        track_applications.active_job_fit_status = original_active_job_fit_status
        if original_build_resume is None:
            del sys.modules["build_resume"]
        else:
            sys.modules["build_resume"] = original_build_resume

    assert_true(
        captured.get("resume_path") == tailored_resume,
        "row_for_active_job() should score fit against the tailored output when one exists",
    )
    assert_true(
        captured.get("profile_resume_text") == "visible text from tailored_resume.docx",
        "row_for_active_job() should pass the selected resume text into job_problem_profile() for lane detection",
    )
    assert_true(
        row.get("output_file") == str(tailored_resume),
        "row_for_active_job() should keep storing the tailored output path in the tracker row",
    )


def test_track_add_row_refreshes_existing_metadata() -> None:
    import track_applications

    base_row = {
        "date_added": "2026-06-10",
        "company": "Smoke Test Systems",
        "role_title": "Implementation Consultant",
        "lane_label": "Implementation and Delivery",
        "fit_status": "Adjacent Fit",
        "audit_flag": "FAIL",
        "source_resume": "Estrada_Resume_Implementation.docx",
        "output_file": "C:/temp/resume.docx",
        "snapshot_id": "smoke_snapshot",
        "applied_date": "",
        "current_status": "draft",
        "last_round": "",
        "outcome": "",
        "notes": "fresh note",
    }
    captured: dict[str, object] = {}

    original_row_for_active_job = track_applications.row_for_active_job
    original_upsert_application = track_applications.upsert_application
    try:
        track_applications.row_for_active_job = lambda status, notes: dict(base_row, current_status=status, notes=notes)
        track_applications.upsert_application = (
            lambda company, role_title, updates, base_row=None: captured.update(
                {
                    "company": company,
                    "role_title": role_title,
                    "updates": updates,
                    "base_row": base_row,
                }
            ) or "Updated"
        )
        result = track_applications.add_row(SimpleNamespace(status="draft", notes="fresh note", applied_date=""))
    finally:
        track_applications.row_for_active_job = original_row_for_active_job
        track_applications.upsert_application = original_upsert_application

    updates = captured.get("updates", {})
    assert_true(result == 0, "add_row() should return success when the active job row is available")
    assert_true(isinstance(updates, dict), "add_row() should pass a metadata update dictionary into upsert_application()")
    assert_true(
        updates.get("lane_label") == "Implementation and Delivery",
        "add_row() should refresh lane metadata when updating an existing application row",
    )
    assert_true(
        updates.get("fit_status") == "Adjacent Fit",
        "add_row() should refresh fit metadata when updating an existing application row",
    )
    assert_true(
        updates.get("audit_flag") == "FAIL",
        "add_row() should refresh audit metadata when updating an existing application row",
    )
    assert_true(
        updates.get("source_resume") == "Estrada_Resume_Implementation.docx",
        "add_row() should refresh source resume metadata when updating an existing application row",
    )
    assert_true(
        updates.get("snapshot_id") == "smoke_snapshot",
        "add_row() should refresh the canonical snapshot ID when updating an existing application row",
    )


def test_track_refresh_uses_active_job_fit_status() -> None:
    import track_applications

    with TemporaryDirectory(prefix="tracker_fit_refresh_") as temp_name:
        output_path = Path(temp_name) / "resume.docx"
        output_path.write_text("placeholder", encoding="utf-8")
        row = {
            "company": "Smoke Test Systems",
            "role_title": "Implementation Consultant",
            "lane_label": "",
            "fit_status": "",
            "output_file": str(output_path),
        }
        original_row_job_description_text = track_applications.row_job_description_text
        original_active_job_fit_status = track_applications.active_job_fit_status
        original_job_problem_profile = track_applications.resume_analysis.job_problem_profile
        original_build_resume = sys.modules.get("build_resume")
        try:
            track_applications.row_job_description_text = lambda _row: DUMMY_JOB_DESCRIPTION
            track_applications.active_job_fit_status = lambda _jd, _path: "Strong Fit"
            track_applications.resume_analysis.job_problem_profile = lambda _jd, _text: SimpleNamespace(lane_label="Implementation Delivery")
            sys.modules["build_resume"] = SimpleNamespace(docx_visible_text_from_path=lambda _path: "resume text")
            refreshed, changed = track_applications.refresh_row_metadata(row, force=True)
        finally:
            track_applications.row_job_description_text = original_row_job_description_text
            track_applications.active_job_fit_status = original_active_job_fit_status
            track_applications.resume_analysis.job_problem_profile = original_job_problem_profile
            if original_build_resume is None:
                del sys.modules["build_resume"]
            else:
                sys.modules["build_resume"] = original_build_resume

    assert_true(changed, "refresh_row_metadata() should mark the row changed when it fills missing fit and lane metadata")
    assert_true(
        refreshed.get("fit_status") == "Strong Fit",
        f"refresh_row_metadata() should route fit refresh through active_job_fit_status(); got {refreshed.get('fit_status')!r}",
    )
    assert_true(
        refreshed.get("lane_label") == "Implementation Delivery",
        f"refresh_row_metadata() should still refresh the lane label alongside fit; got {refreshed.get('lane_label')!r}",
    )


def test_tracker_status_precedence_blocks_regressions() -> None:
    import track_applications

    rows = [
        {
            "date_added": "2026-06-10",
            "company": "Smoke Test Systems",
            "role_title": "Implementation Consultant",
            "lane_label": "Implementation Delivery",
            "fit_status": "Strong Fit",
            "audit_flag": "",
            "source_resume": "Estrada_Resume_Implementation.docx",
            "output_file": "C:/temp/resume.docx",
            "applied_date": "",
            "current_status": "final_round",
            "last_round": "",
            "outcome": "",
            "notes": "",
        }
    ]
    captured: dict[str, object] = {}
    original_read_rows = track_applications.read_rows
    original_write_rows = track_applications.write_rows
    try:
        track_applications.read_rows = lambda: [dict(row) for row in rows]
        track_applications.write_rows = lambda updated_rows: captured.setdefault("rows", updated_rows)
        track_applications.upsert_application(
            "Smoke Test Systems",
            "Implementation Consultant",
            {"current_status": "draft"},
        )
    finally:
        track_applications.read_rows = original_read_rows
        track_applications.write_rows = original_write_rows

    saved_rows = captured.get("rows", [])
    assert_true(
        bool(saved_rows),
        "upsert_application() should still write the row even when it preserves the existing later-stage status",
    )
    assert_true(
        saved_rows[0].get("current_status") == "final_round",
        f"upsert_application() should preserve a later-stage status instead of regressing it to draft; got {saved_rows[0].get('current_status')!r}",
    )


def test_interview_negotiation_section_avoids_bracket_placeholders(build_interview_cheat_sheet: object) -> None:
    sections = build_interview_cheat_sheet.negotiation_preparation_section(
        "Acme Health",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
        "implementation_delivery",
    )
    combined = " ".join(line for _title, lines in sections for line in lines)
    assert_true(
        "[" not in combined and "]" not in combined,
        f"Negotiation preparation guidance should not leak bracket placeholders into generated text; got {combined!r}",
    )


def test_tracker_audit_flag_derivation() -> None:
    import track_applications

    bridge_row = {"output_file": "C:/temp/Christian Estrada - Example BRIDGE Resume.docx"}
    fail_row = {"output_file": "C:/temp/Christian Estrada - Example FAIL Resume.docx"}
    poor_row = {"output_file": "C:/temp/Christian Estrada - Example POOR Resume.docx"}
    clean_row = {"output_file": "C:/temp/Christian Estrada - Example Resume.docx"}
    assert_true(
        track_applications.tracker_audit_flag(bridge_row) == "BRIDGE",
        "tracker_audit_flag() should derive BRIDGE from the output filename",
    )
    assert_true(
        track_applications.tracker_audit_flag(fail_row) == "FAIL",
        "tracker_audit_flag() should derive FAIL from the output filename",
    )
    assert_true(
        track_applications.tracker_audit_flag(poor_row) == "POOR",
        "tracker_audit_flag() should derive POOR from the output filename",
    )
    assert_true(
        track_applications.tracker_audit_flag(clean_row) == "",
        "tracker_audit_flag() should stay empty for clean resume filenames",
    )


def test_tracker_refresh_warning_without_matching_jd() -> None:
    import track_applications

    with TemporaryDirectory(prefix="tracker_refresh_") as temp_name:
        output_path = Path(temp_name) / "placeholder.docx"
        output_path.write_text("placeholder", encoding="utf-8")
        row = {
            "company": "No Match Company",
            "role_title": "No Match Role",
            "lane_label": "",
            "fit_status": "",
            "output_file": str(output_path),
        }
        stdout_buffer = io.StringIO()
        with contextlib.redirect_stdout(stdout_buffer):
            refreshed, changed = track_applications.refresh_row_metadata(row, force=False)
        output = stdout_buffer.getvalue()
    assert_true(
        not changed,
        "refresh_row_metadata() should leave unmatched rows unchanged when no current or archived JD is found",
    )
    assert_true(
        "SKIPPED refresh:" in output and "jd-archive" in output,
        f"refresh_row_metadata() should print a skip warning when no JD matches; got {output!r}",
    )
    assert_true(
        refreshed.get("company") == "No Match Company",
        "refresh_row_metadata() should return the original row data when no JD match exists",
    )


def test_tracker_row_job_description_prefers_snapshot_id() -> None:
    import track_applications

    original_read_job_description = track_applications.read_job_description
    original_snapshot_job_description_text = track_applications.job_context_archive.snapshot_job_description_text
    original_read_jd_library_rows = track_applications.read_jd_library_rows
    try:
        track_applications.read_job_description = lambda: ""
        track_applications.job_context_archive.snapshot_job_description_text = lambda snapshot_id: DUMMY_JOB_DESCRIPTION if snapshot_id == "smoke_snapshot" else ""
        track_applications.read_jd_library_rows = lambda: []
        text = track_applications.row_job_description_text(
            {
                "company": "Smoke Test Systems",
                "role_title": "Implementation Consultant",
                "snapshot_id": "smoke_snapshot",
            }
        )
    finally:
        track_applications.read_job_description = original_read_job_description
        track_applications.job_context_archive.snapshot_job_description_text = original_snapshot_job_description_text
        track_applications.read_jd_library_rows = original_read_jd_library_rows

    assert_true(
        "Smoke Test Systems" in text and "Implementation Consultant" in text,
        f"row_job_description_text() should prefer the stored snapshot ID when present; got {text!r}",
    )


def test_tracker_refresh_warning_when_output_missing() -> None:
    import track_applications

    row = {
        "company": "Missing Output Company",
        "role_title": "Missing Output Role",
        "lane_label": "",
        "fit_status": "",
        "output_file": "C:/temp/does_not_exist.docx",
    }
    stdout_buffer = io.StringIO()
    with contextlib.redirect_stdout(stdout_buffer):
        refreshed, changed = track_applications.refresh_row_metadata(row, force=True)
    output = stdout_buffer.getvalue()
    assert_true(not changed, "refresh_row_metadata() should leave the row unchanged when the output file is missing")
    assert_true(
        "SKIPPED refresh: output file not found for Missing Output Company - Missing Output Role." in output,
        f"refresh_row_metadata() should print a skip warning when the output file is missing; got {output!r}",
    )
    assert_true(
        refreshed.get("company") == "Missing Output Company",
        "refresh_row_metadata() should return the original row when the output file is missing",
    )


def test_debrief_outcome_mapping_and_notes() -> None:
    import post_interview_debrief
    import track_applications

    captured_updates: list[dict[str, str]] = []
    original_upsert_application = track_applications.upsert_application
    try:
        track_applications.upsert_application = lambda company, role_title, updates, base_row=None: captured_updates.append(dict(updates)) or "Updated"
        base_data = {
            "company_name": "Unmatched Smoke Test Company",
            "role_title": "Implementation Consultant",
            "interview_date": "2026-06-10",
            "round_number": "2",
            "feedback_received": "Need stronger workflow examples.",
            "role_language": "They emphasized launch readiness.",
            "company_intelligence": "Implementation team is growing.",
        }
        post_interview_debrief.update_tracker_from_debrief(dict(base_data, outcome="advance"))
        post_interview_debrief.update_tracker_from_debrief(dict(base_data, outcome="reject"))
        post_interview_debrief.update_tracker_from_debrief(dict(base_data, outcome="pending"))
        stdout_buffer = io.StringIO()
        with contextlib.redirect_stdout(stdout_buffer):
            post_interview_debrief.update_tracker_from_debrief(dict(base_data, outcome="withdrew"))
        warning_output = stdout_buffer.getvalue()
    finally:
        track_applications.upsert_application = original_upsert_application

    assert_true(
        [update.get("current_status") for update in captured_updates[:3]] == ["interview", "rejected", None],
        f"update_tracker_from_debrief() should map advance/reject and leave pending status untouched; got {captured_updates}",
    )
    assert_true(
        captured_updates[2].get("outcome") == "",
        "Pending debrief outcomes should keep the tracker outcome field blank",
    )
    assert_true(
        "current_status" not in captured_updates[2],
        "Pending debrief outcomes should preserve the existing tracker status instead of forcing a new one",
    )
    assert_true(
        "Feedback:" in captured_updates[0].get("notes", "")
        and "Role language:" in captured_updates[0].get("notes", "")
        and "Intel:" in captured_updates[0].get("notes", ""),
        f"Debrief tracker notes should keep labeled note sources; got {captured_updates[0].get('notes', '')!r}",
    )
    assert_true(
        len(captured_updates) == 3,
        "Unknown debrief outcomes should not write a tracker update",
    )
    assert_true(
        "WARNING: unrecognized debrief outcome 'withdrew'" in warning_output,
        f"Unknown debrief outcomes should print a warning; got {warning_output!r}",
    )


def test_workflow_tracks_after_required_steps_before_optional_failure() -> None:
    import run_resume_workflow

    order: list[str] = []
    original_parse_args = run_resume_workflow.parse_args
    original_workspace_health = run_resume_workflow.workspace_health.ensure_workspace_health_or_exit
    original_validate_job_description = run_resume_workflow.validate_job_description
    original_run_with_recovery = run_resume_workflow.run_with_recovery
    original_run_tracker_auto_add = run_resume_workflow.run_tracker_auto_add
    original_explain_unresolved = run_resume_workflow.explain_unresolved
    original_pairing_issue = run_resume_workflow.job_context_archive.application_question_pairing_issue
    try:
        run_resume_workflow.parse_args = lambda: SimpleNamespace(
            resume_only=False,
            include_cheat_sheet=True,
            include_detailed_guide=False,
            dry_run=False,
        )
        run_resume_workflow.workspace_health.ensure_workspace_health_or_exit = lambda _workflow_name: {}
        run_resume_workflow.validate_job_description = lambda _path: None
        run_resume_workflow.job_context_archive.application_question_pairing_issue = lambda *_args: ""

        def fake_run_with_recovery(step_name: str, script_name: str, *, can_rebuild_resume: bool = False):
            order.append(step_name)
            return run_resume_workflow.StepResult(
                name=step_name,
                returncode=1 if step_name == "Building interview cheat sheet" else 0,
                stdout="",
                stderr="",
                log_path=Path("dummy.log"),
                specificity_warnings=[],
                cover_warnings=[],
                preflight_warnings=[],
            )

        run_resume_workflow.run_with_recovery = fake_run_with_recovery
        run_resume_workflow.run_tracker_auto_add = lambda: order.append("tracker")
        run_resume_workflow.explain_unresolved = lambda result: order.append(f"explained:{result.name}")
        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            try:
                run_resume_workflow.main()
                raise SmokeFailure("run_resume_workflow.main() should exit when an optional step fails")
            except SystemExit as error:
                assert_true(
                    error.code == 1,
                    f"run_resume_workflow.main() should exit with the failing optional step code; got {error.code}",
                )
    finally:
        run_resume_workflow.parse_args = original_parse_args
        run_resume_workflow.workspace_health.ensure_workspace_health_or_exit = original_workspace_health
        run_resume_workflow.validate_job_description = original_validate_job_description
        run_resume_workflow.run_with_recovery = original_run_with_recovery
        run_resume_workflow.run_tracker_auto_add = original_run_tracker_auto_add
        run_resume_workflow.explain_unresolved = original_explain_unresolved
        run_resume_workflow.job_context_archive.application_question_pairing_issue = original_pairing_issue

    assert_true(
        order[:5] == [
            "Building resume",
            "Building cover letter",
            "Building qualifications statement",
            "tracker",
            "Building interview cheat sheet",
        ],
        f"Workflow should track immediately after required steps and before optional failures; got {order}",
    )


def test_workflow_parses_cover_warning_channels() -> None:
    import run_resume_workflow

    specificity_warnings, cover_warnings, preflight_warnings = run_resume_workflow.repair_generated_docx_outputs(
        "\n".join(
            [
                "SPECIFICITY WARNING: company appears only once",
                "COVER LETTER WARNING: abrupt first-person switch detected",
                "COVER LETTER PREFLIGHT: Cover letter has fewer than 4 job-description keyword hits.",
            ]
        )
    )
    assert_true(
        specificity_warnings == ["company appears only once"],
        f"Workflow warning parsing should preserve specificity warnings; got {specificity_warnings}",
    )
    assert_true(
        cover_warnings == ["abrupt first-person switch detected"],
        f"Workflow warning parsing should preserve cover warnings; got {cover_warnings}",
    )
    assert_true(
        preflight_warnings == ["Cover letter has fewer than 4 job-description keyword hits."],
        f"Workflow warning parsing should preserve preflight warnings; got {preflight_warnings}",
    )


def test_tasks_auto_archive_environment_for_commercial_command() -> None:
    import tasks

    original_should_auto_archive_command = tasks.should_auto_archive_command
    original_module = sys.modules.get("job_context_archive")

    class Snapshot:
        snapshot_id = "smoke_snapshot"
        path = Path("C:/temp/smoke_snapshot")
        questions_present = True
        question_count = 2

    try:
        tasks.should_auto_archive_command = lambda command_name: command_name == "resume"
        sys.modules["job_context_archive"] = SimpleNamespace(
            SNAPSHOT_ID_ENV="JOB_CONTEXT_SNAPSHOT_ID",
            SNAPSHOT_PATH_ENV="JOB_CONTEXT_SNAPSHOT_PATH",
            SOURCE_COMMAND_ENV="JOB_CONTEXT_SOURCE_COMMAND",
            WORKFLOW_TYPE_ENV="JOB_CONTEXT_WORKFLOW_TYPE",
            archive_active_context=lambda **_kwargs: Snapshot(),
        )
        stdout_buffer = io.StringIO()
        with contextlib.redirect_stdout(stdout_buffer):
            env = tasks.archive_environment_for_command("resume")
        output = stdout_buffer.getvalue()
    finally:
        tasks.should_auto_archive_command = original_should_auto_archive_command
        if original_module is None:
            sys.modules.pop("job_context_archive", None)
        else:
            sys.modules["job_context_archive"] = original_module

    assert_true(
        env.get("JOB_CONTEXT_SNAPSHOT_ID") == "smoke_snapshot",
        f"archive_environment_for_command() should set the snapshot ID for commercial output commands; got {env}",
    )
    assert_true(
        env.get("JOB_CONTEXT_SOURCE_COMMAND") == "resume" and env.get("JOB_CONTEXT_WORKFLOW_TYPE") == "commercial",
        f"archive_environment_for_command() should carry command metadata in the environment; got {env}",
    )
    assert_true(
        "Job context snapshot: smoke_snapshot" in output,
        f"archive_environment_for_command() should print the created snapshot ID; got {output!r}",
    )


def test_workflow_hard_fails_docx_validation_issues() -> None:
    import run_resume_workflow

    original_repair_docx_open_issues = run_resume_workflow.repair_docx_open_issues
    try:
        with TemporaryDirectory(prefix="workflow_docx_") as temp_name:
            docx_path = build_minimal_docx(
                Path(temp_name) / "Christian Estrada - Smoke Test Systems Resume.docx",
                ["Smoke output"],
            )
            run_resume_workflow.repair_docx_open_issues = lambda _path: (_ for _ in ()).throw(RuntimeError("broken package"))
            try:
                run_resume_workflow.repair_generated_docx_outputs(f"Output DOCX: {docx_path}")
                raise SmokeFailure("repair_generated_docx_outputs() should raise when a generated DOCX fails validation")
            except RuntimeError as error:
                assert_true(
                    "broken package" in str(error),
                    f"repair_generated_docx_outputs() should surface the DOCX validation failure; got {error}",
                )
    finally:
        run_resume_workflow.repair_docx_open_issues = original_repair_docx_open_issues


def test_run_resume_workflow_dry_run_skips_upfront_job_validation() -> None:
    import run_resume_workflow

    original_parse_args = run_resume_workflow.parse_args
    original_workspace_health = run_resume_workflow.workspace_health.ensure_workspace_health_or_exit
    original_validate_job_description = run_resume_workflow.validate_job_description
    original_run_dry_run = run_resume_workflow.run_dry_run
    try:
        run_resume_workflow.parse_args = lambda: SimpleNamespace(
            resume_only=False,
            include_cheat_sheet=False,
            include_detailed_guide=False,
            dry_run=True,
        )
        run_resume_workflow.workspace_health.ensure_workspace_health_or_exit = lambda _workflow_name: {}

        def should_not_validate(_path):
            raise SmokeFailure("run_resume_workflow.main() should not validate the live job file before --dry-run")

        run_resume_workflow.validate_job_description = should_not_validate
        run_resume_workflow.run_dry_run = lambda: 0
        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            try:
                run_resume_workflow.main()
                raise SmokeFailure("run_resume_workflow.main() should exit through the dry-run path")
            except SystemExit as error:
                assert_true(
                    error.code == 0,
                    f"run_resume_workflow.main() should exit cleanly for --dry-run; got {error.code}",
                )
    finally:
        run_resume_workflow.parse_args = original_parse_args
        run_resume_workflow.workspace_health.ensure_workspace_health_or_exit = original_workspace_health
        run_resume_workflow.validate_job_description = original_validate_job_description
        run_resume_workflow.run_dry_run = original_run_dry_run


def test_federal_workflow_supporting_flag_resolution() -> None:
    import run_federal_resume_workflow

    none_requested = run_federal_resume_workflow.requested_supporting_steps(
        SimpleNamespace(
            with_cover=False,
            with_interview=False,
            with_guide=False,
            with_supporting_docs=False,
        )
    )
    explicit_requested = run_federal_resume_workflow.requested_supporting_steps(
        SimpleNamespace(
            with_cover=True,
            with_interview=False,
            with_guide=True,
            with_supporting_docs=False,
        )
    )
    alias_requested = run_federal_resume_workflow.requested_supporting_steps(
        SimpleNamespace(
            with_cover=False,
            with_interview=False,
            with_guide=False,
            with_supporting_docs=True,
        )
    )
    assert_true(none_requested == (), f"No federal supporting-doc flags should yield no optional steps; got {none_requested}")
    assert_true(
        explicit_requested == (
            ("Building federal cover letter", "build_federal_cover_letter.py"),
            ("Building federal detailed interview guide", "build_federal_detailed_interview_guide.py"),
        ),
        f"Explicit federal supporting-doc flags should preserve the requested order; got {explicit_requested}",
    )
    assert_true(
        alias_requested == (
            ("Building federal cover letter", "build_federal_cover_letter.py"),
            ("Building federal interview cheat sheet", "build_federal_interview_cheat_sheet.py"),
            ("Building federal detailed interview guide", "build_federal_detailed_interview_guide.py"),
        ),
        f"--with-supporting-docs should expand to all federal supporting-doc steps; got {alias_requested}",
    )


def test_federal_workflow_runs_supporting_steps_after_resume() -> None:
    import run_federal_resume_workflow

    original_parse_args = run_federal_resume_workflow.parse_args
    original_workspace_health = run_federal_resume_workflow.workspace_health.ensure_workspace_health_or_exit
    original_validate_exists = run_federal_resume_workflow.validate_federal_job_description_exists
    original_run_step = run_federal_resume_workflow.run_step
    order: list[str] = []
    try:
        run_federal_resume_workflow.parse_args = lambda: SimpleNamespace(
            dry_run=False,
            with_cover=True,
            with_interview=False,
            with_guide=True,
            with_supporting_docs=False,
        )
        run_federal_resume_workflow.workspace_health.ensure_workspace_health_or_exit = lambda _workflow_name: {}
        run_federal_resume_workflow.validate_federal_job_description_exists = lambda: order.append("validate")
        run_federal_resume_workflow.run_step = lambda step_name, script_name: order.append(script_name) or SimpleNamespace(
            name=step_name,
            returncode=0,
            stdout="",
            stderr="",
            log_path=Path("smoke.log"),
            ok=True,
            output="",
        )
        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            run_federal_resume_workflow.main()
    finally:
        run_federal_resume_workflow.parse_args = original_parse_args
        run_federal_resume_workflow.workspace_health.ensure_workspace_health_or_exit = original_workspace_health
        run_federal_resume_workflow.validate_federal_job_description_exists = original_validate_exists
        run_federal_resume_workflow.run_step = original_run_step

    assert_true(
        order == [
            "validate",
            "build_federal_resume.py",
            "build_federal_cover_letter.py",
            "build_federal_detailed_interview_guide.py",
        ],
        f"Federal workflow should build the resume first and then requested supporting docs; got {order}",
    )


def test_run_federal_workflow_dry_run_skips_upfront_job_validation() -> None:
    import run_federal_resume_workflow

    original_parse_args = run_federal_resume_workflow.parse_args
    original_workspace_health = run_federal_resume_workflow.workspace_health.ensure_workspace_health_or_exit
    original_validate_exists = run_federal_resume_workflow.validate_federal_job_description_exists
    original_run_dry_run = run_federal_resume_workflow.run_dry_run
    try:
        run_federal_resume_workflow.parse_args = lambda: SimpleNamespace(dry_run=True)
        run_federal_resume_workflow.workspace_health.ensure_workspace_health_or_exit = lambda _workflow_name: {}

        def should_not_validate():
            raise SmokeFailure(
                "run_federal_resume_workflow.main() should not validate the live federal job file before --dry-run"
            )

        run_federal_resume_workflow.validate_federal_job_description_exists = should_not_validate
        run_federal_resume_workflow.run_dry_run = lambda: 0
        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            try:
                run_federal_resume_workflow.main()
                raise SmokeFailure("run_federal_resume_workflow.main() should exit through the dry-run path")
            except SystemExit as error:
                assert_true(
                    error.code == 0,
                    f"run_federal_resume_workflow.main() should exit cleanly for --dry-run; got {error.code}",
                )
    finally:
        run_federal_resume_workflow.parse_args = original_parse_args
        run_federal_resume_workflow.workspace_health.ensure_workspace_health_or_exit = original_workspace_health
        run_federal_resume_workflow.validate_federal_job_description_exists = original_validate_exists
        run_federal_resume_workflow.run_dry_run = original_run_dry_run


def test_run_federal_workflow_dry_run_labels_unverified_page_counts() -> None:
    import run_federal_resume_workflow

    original_job_path = run_federal_resume_workflow.FEDERAL_JOB_DESCRIPTION
    original_resume_source = run_federal_resume_workflow.FEDERAL_RESUME_SOURCE
    original_essay_source = run_federal_resume_workflow.FEDERAL_ESSAY_SOURCE
    original_matching_outputs = run_federal_resume_workflow.matching_federal_outputs
    original_matching_qualifications = run_federal_resume_workflow.matching_federal_qualifications_outputs
    original_build_federal_resume = sys.modules.get("build_federal_resume")
    try:
        with TemporaryDirectory(prefix="federal_dry_run_") as temp_name:
            temp_root = Path(temp_name)
            job_path = temp_root / "federal_job_description.txt"
            resume_source = temp_root / "Christian_Estrada_Federal_Source.json"
            essay_source = temp_root / "Christian_Estrada_Federal_Standard_Essays.json"
            job_path.write_text(FEDERAL_DUMMY_JOB_DESCRIPTION.strip(), encoding="utf-8")
            resume_source.write_text("{}", encoding="utf-8")
            essay_source.write_text("{}", encoding="utf-8")

            sys.modules["build_federal_resume"] = SimpleNamespace(
                validate_inputs=lambda text: text,
                extract_federal_agency_name=lambda _text: "Department of Veterans Affairs",
                extract_federal_role_title=lambda _text: "IT Program Manager",
                extract_federal_output_name=lambda _text: "Department of Veterans Affairs - IT Program Manager",
                load_federal_source=lambda: SimpleNamespace(),
                source_visible_text=lambda _source: "federal source text",
                job_problem_profile=lambda *_args, **_kwargs: SimpleNamespace(
                    primary_lane="implementation_delivery",
                    lane_label="Implementation Delivery",
                    core_problem="delivery governance and implementation risk",
                    direct_matches=("Implementation Delivery",),
                    adjacent_matches=("Change Enablement",),
                    unsupported_requirements=(),
                ),
                federal_requirement_audit=lambda *_args, **_kwargs: SimpleNamespace(keyword_targets=("implementation governance", "data migration")),
                audit_keywords=lambda _job_description: ("implementation governance", "data migration"),
                keyword_hits=lambda _job_description, keywords: 2 if "implementation governance" in keywords else 1,
                audit_keyword_sort_key=lambda _job_description, keyword: (1 if "implementation governance" in keyword else 0, len(keyword)),
                resume_plan=lambda *_args, **_kwargs: SimpleNamespace(
                    audit=SimpleNamespace(
                        warnings=(
                            "Federal page-count warning: selected the resume candidate with an unverified page count because the renderer was unavailable.",
                        )
                    ),
                    bullet_groups=(),
                    resume_layout=SimpleNamespace(name="federal_standard", font_size=10.0),
                    resume_page_count=None,
                    qualifications_layout=SimpleNamespace(name="qualifications_standard", font_size=10.0),
                    qualifications_page_count=None,
                ),
                requirement_report_lines=lambda _audit: ("Selective Factor [Direct] -> Example requirement",),
                selected_bullet_reference_lines=lambda _source, _groups: ("Department of Veterans Affairs: none",),
                federal_page_count_label=lambda page_count: "unverified (renderer unavailable)" if page_count is None else str(page_count),
            )

            run_federal_resume_workflow.FEDERAL_JOB_DESCRIPTION = job_path
            run_federal_resume_workflow.FEDERAL_RESUME_SOURCE = resume_source
            run_federal_resume_workflow.FEDERAL_ESSAY_SOURCE = essay_source
            run_federal_resume_workflow.matching_federal_outputs = lambda _company_name: []
            run_federal_resume_workflow.matching_federal_qualifications_outputs = lambda _company_name: []

            output_buffer = io.StringIO()
            with contextlib.redirect_stdout(output_buffer), contextlib.redirect_stderr(io.StringIO()):
                result = run_federal_resume_workflow.run_dry_run()
    finally:
        if original_build_federal_resume is None:
            sys.modules.pop("build_federal_resume", None)
        else:
            sys.modules["build_federal_resume"] = original_build_federal_resume
        run_federal_resume_workflow.FEDERAL_JOB_DESCRIPTION = original_job_path
        run_federal_resume_workflow.FEDERAL_RESUME_SOURCE = original_resume_source
        run_federal_resume_workflow.FEDERAL_ESSAY_SOURCE = original_essay_source
        run_federal_resume_workflow.matching_federal_outputs = original_matching_outputs
        run_federal_resume_workflow.matching_federal_qualifications_outputs = original_matching_qualifications

    dry_run_output = output_buffer.getvalue()
    assert_true(result == 0, f"run_dry_run() should succeed with the unverified-page-count stub; got {result}")
    assert_true(
        "Planned pages: unverified (renderer unavailable)" in dry_run_output,
        f"run_dry_run() should label unknown page counts as unverified; got {dry_run_output!r}",
    )
    assert_true(
        "Planned pages: None" not in dry_run_output,
        f"run_dry_run() should not print raw None for federal page counts; got {dry_run_output!r}",
    )


def test_tasks_register_federal_supporting_doc_commands() -> None:
    import tasks

    federal_resume = tasks.TASKS["federal-resume"]
    assert_true(
        "--with-cover" in federal_resume.description and "--with-supporting-docs" in federal_resume.description,
        "tasks.py should advertise the new optional federal supporting-doc flags on federal-resume",
    )
    for command_name, script_name in (
        ("federal-cover", "scripts/build_federal_cover_letter.py"),
        ("federal-interview", "scripts/build_federal_interview_cheat_sheet.py"),
        ("federal-guide", "scripts/build_federal_detailed_interview_guide.py"),
    ):
        task = tasks.TASKS[command_name]
        assert_true(task.args == (script_name,), f"{command_name} should point to {script_name}; got {task.args}")
        assert_true(
            not task.needs_job_description,
            f"{command_name} should bypass commercial job-description validation and use the federal file instead",
        )
        assert_true(task.maturity == "Experimental", f"{command_name} should be marked Experimental; got {task.maturity}")
        assert_true(not task.production_safe, f"{command_name} should start as review-heavy until it has more runtime coverage")
    assert_true(
        tasks.TASKS["cover"].args == ("scripts/build_cover_letter.py",),
        f"Commercial cover command registration should stay unchanged; got {tasks.TASKS['cover'].args}",
    )
    assert_true(
        tasks.TASKS["qualifications"].args == ("scripts/build_standard_qualifications_statement.py",),
        f"Commercial qualifications command should point to the standard qualifications builder; got {tasks.TASKS['qualifications'].args}",
    )


def test_tasks_register_interview_review_command() -> None:
    import tasks

    task = tasks.TASKS["interview-review"]
    assert_true(
        task.args == ("scripts/build_interview_review.py",),
        f"interview-review should point to the dedicated review builder; got {task.args}",
    )
    assert_true(not task.needs_job_description, "interview-review should be able to use structured debrief history even without an active JD")
    assert_true(task.maturity == "Experimental", f"interview-review should start as Experimental; got {task.maturity}")
    assert_true(not task.production_safe, "interview-review should remain review-heavy until it has more runtime coverage")
    assert_true(
        tasks.TASKS["interview"].args == ("scripts/build_interview_cheat_sheet.py",),
        f"Commercial interview command registration should stay unchanged; got {tasks.TASKS['interview'].args}",
    )
    assert_true(
        tasks.TASKS["guide"].args == ("scripts/build_detailed_interview_guide.py",),
        f"Commercial guide command registration should stay unchanged; got {tasks.TASKS['guide'].args}",
    )


def test_onedrive_run_guard_absent() -> None:
    import tasks

    with TemporaryDirectory(prefix="onedrive_guard_absent_") as temp_name:
        assert_true(
            tasks._onedrive_run_guard(Path(temp_name)) is None,
            "_onedrive_run_guard() should allow normal runs when the retirement sentinel is absent.",
        )


def test_onedrive_run_guard_present() -> None:
    import tasks

    with TemporaryDirectory(prefix="onedrive_guard_present_") as temp_name:
        root = Path(temp_name)
        (root / "DO_NOT_RUN_FROM_ONEDRIVE.txt").write_text("retired", encoding="utf-8")
        assert_true(
            tasks._onedrive_run_guard(root) == 2,
            "_onedrive_run_guard() should refuse runs when the retirement sentinel is present.",
        )


def test_build_interview_review_sections(build_interview_review: object, interview_context: object) -> None:
    record = interview_context.normalize_round_record(
        {
            "company_name": "Acme Systems",
            "role_title": "Implementation Consultant",
            "interview_date": "06/19/2026",
            "round_number": "1",
            "outcome": "reject",
            "role_language": ["Client-facing ownership matters."],
            "company_intelligence": ["A stronger consultant-style presence would help."],
            "imported_review_text": "\n".join(
                [
                    "Overall: relevant, but not consultative enough.",
                    '1) "Why should we move you to the next round?"',
                    "Questions at the end were tactical, not strategic.",
                ]
            ),
        }
    )
    sections = build_interview_review.interview_review_sections(record)
    flattened = "\n".join(line for _title, lines in sections for line in lines)
    titles = {title for title, _lines in sections}
    assert_true("Decision Signal" in titles, f"Interview review should include a Decision Signal section; got {titles}")
    assert_true("Answer Strategy" in titles, f"Interview review should include an Answer Strategy section; got {titles}")
    assert_true(
        "Why should we move you to the next round?" in flattened,
        f"Interview review should carry parsed interviewer questions into the answer-assets section; got {flattened!r}",
    )


def test_claude_packet_modes(build_claude_review_packet: object) -> None:
    broad_packet = build_claude_review_packet.build_packet("broad", skip_checks=True)
    tracker_packet = build_claude_review_packet.build_packet("tracker", skip_checks=True)
    federal_packet = build_claude_review_packet.build_packet("federal", skip_checks=True)
    claude_review_packet = build_claude_review_packet.build_packet("claude-review", skip_checks=True)
    assert_true(
        "# Broad Claude Review Packet" in broad_packet and "Packet mode: `broad`" in broad_packet,
        "build_packet() should label the broad packet mode clearly",
    )
    assert_true(
        "Questions For Claude" in broad_packet and "Relevant Code" in broad_packet,
        "build_packet() should include the required Claude packet headings",
    )
    assert_true(
        "# Tracker Claude Review Packet" in tracker_packet and "Packet mode: `tracker`" in tracker_packet,
        "build_packet() should label the tracker packet mode clearly",
    )
    assert_true(
        "track_applications.py" in tracker_packet and "update_tracker_from_debrief" in tracker_packet,
        "Tracker packets should include tracker and debrief code excerpts",
    )
    assert_true(
        "Packet mode: `federal`" in federal_packet and "Current System Contract" in federal_packet,
        "Federal packets should include the federal mode label and the current system contract block",
    )
    assert_true(
        "Packet mode: `claude-review`" in claude_review_packet
        and "proposal" in claude_review_packet.lower()
        and "Interview Feature Track" in claude_review_packet
        and "Commit Train Track" in claude_review_packet,
        "Claude-review packets should include the dedicated mode label, proposal-versus-live guidance, and dual-track progress-check contract",
    )


def test_claude_packet_run_command_keeps_stdout_and_stderr(build_claude_review_packet: object) -> None:
    original_run = build_claude_review_packet.subprocess.run
    try:
        build_claude_review_packet.subprocess.run = lambda *args, **kwargs: SimpleNamespace(
            stdout="validation summary\n",
            stderr="warning details\n",
            returncode=1,
        )
        output = build_claude_review_packet.run_command("validate")
    finally:
        build_claude_review_packet.subprocess.run = original_run

    assert_true(
        "[exit 1]" in output and "validation summary" in output and "warning details" in output,
        f"run_command() should preserve both stdout and stderr for failing commands; got {output!r}",
    )


def test_claude_packet_self_audit_checks_task_registration(build_claude_review_packet: object) -> None:
    original_tasks_path = build_claude_review_packet.TASKS_PATH
    try:
        with TemporaryDirectory(prefix="claude_packet_tasks_") as temp_name:
            tasks_path = Path(temp_name) / "tasks.py"
            tasks_path.write_text('TASKS = {"claude-packet": (), "validate": ()}\n', encoding="utf-8")
            build_claude_review_packet.TASKS_PATH = tasks_path
            warnings: list[str] = []
            build_claude_review_packet.packet_self_audit(
                build_claude_review_packet.PACKET_MODES["broad"],
                warnings,
            )
    finally:
        build_claude_review_packet.TASKS_PATH = original_tasks_path

    assert_true(
        any("claude-refresh" in warning for warning in warnings),
        f"packet_self_audit() should warn when expected packet task commands are missing; got {warnings}",
    )


def test_claude_packet_self_audit_accepts_unquoted_task_mentions(build_claude_review_packet: object) -> None:
    original_tasks_path = build_claude_review_packet.TASKS_PATH
    try:
        with TemporaryDirectory(prefix="claude_packet_tasks_") as temp_name:
            tasks_path = Path(temp_name) / "tasks.py"
            tasks_path.write_text(
                "# claude-packet claude-prompt claude-refresh validate integration-test track-report\nTASKS = {}\n",
                encoding="utf-8",
            )
            build_claude_review_packet.TASKS_PATH = tasks_path
            warnings: list[str] = []
            build_claude_review_packet.packet_self_audit(
                build_claude_review_packet.PACKET_MODES["broad"],
                warnings,
            )
    finally:
        build_claude_review_packet.TASKS_PATH = original_tasks_path

    assert_true(
        not any("does not visibly register" in warning for warning in warnings),
        f"packet_self_audit() should accept bare command mentions in tasks.py source; got {warnings}",
    )


def test_claude_packet_command_cache(build_claude_review_packet: object) -> None:
    original_run = build_claude_review_packet.subprocess.run
    call_counter = {"count": 0}

    def fake_run(*args, **kwargs):
        call_counter["count"] += 1
        return SimpleNamespace(stdout="cached output\n", stderr="", returncode=0)

    try:
        build_claude_review_packet.clear_command_output_cache()
        build_claude_review_packet.subprocess.run = fake_run
        first = build_claude_review_packet.run_command("validate")
        second = build_claude_review_packet.run_command("validate")
    finally:
        build_claude_review_packet.subprocess.run = original_run
        build_claude_review_packet.clear_command_output_cache()

    assert_true(first == second == "cached output", f"run_command() should reuse the cached command output; got {first!r} and {second!r}")
    assert_true(call_counter["count"] == 1, f"run_command() should only invoke subprocess once per cached command; got {call_counter['count']}")


def test_construction_engineering_context_and_bridge_coverage(build_resume: object) -> None:
    job_description = """
Company: Revive ERP
Role: Senior Implementation Consultant
Revive ERP supports construction and engineering firms with project-based ERP, project accounting, finance workflows,
close readiness, and client-side implementation delivery.
Lead implementation scope, client-side ERP ownership, data migration, training, go-live readiness, accounting
partnership, and stakeholder communication for construction and engineering customers.
"""
    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    profile = build_resume.job_problem_profile(job_description, resume_text)
    employer_context = build_resume.primary_employer_context(job_description)
    story_lens = build_resume.primary_story_lens(job_description)
    assert_true(
        employer_context is not None and employer_context.get("key") == "construction_engineering",
        f"primary_employer_context() should route project-based construction/engineering roles to the new employer context; got {employer_context}",
    )
    assert_true(
        story_lens is not None and story_lens.get("key") == "construction_engineering",
        f"primary_story_lens() should route project-based construction/engineering roles to the new story lens; got {story_lens}",
    )
    assert_true(
        len(profile.direct_matches) >= 2,
        f"construction/engineering bridge coverage should produce at least two direct matches on Revive-style input; got {profile.direct_matches}",
    )


def test_alignment_score_ignores_unsupported_named_platforms(build_resume: object) -> None:
    base_job_description = """
Role: Implementation Consultant
Lead requirements gathering, configuration, data migration, training, go-live readiness, stakeholder communication,
and post-launch adoption for enterprise software customers.
"""
    noisy_job_description = base_job_description + "\nExperience with Acumatica, NetSuite, Workday, PrismHR, and Smartsheet is preferred."
    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    base_report = build_resume.alignment_score_report(base_job_description, resume_text)
    noisy_report = build_resume.alignment_score_report(noisy_job_description, resume_text)
    assert_true(
        noisy_report["total_score"] == base_report["total_score"] and noisy_report["grade"] == base_report["grade"],
        f"Unsupported named platforms should not depress alignment on their own; base={base_report}, noisy={noisy_report}",
    )


def test_final_fit_audit_safe_none_alignment_grade(build_resume: object) -> None:
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p><w:r><w:t>Professional Summary</w:t></w:r></w:p>
<w:p><w:r><w:t>Implementation consultant leading implementation, data migration, go-live readiness, stakeholder communication, and reporting across enterprise software workflows. Reduced manual work by 78% across five sites for 150+ users while improving adoption and visibility. Brings practical execution, training, and measurable follow-through to customer-facing delivery work.</w:t></w:r></w:p>
<w:p><w:r><w:t>Professional Experience</w:t></w:r></w:p>
<w:p><w:r><w:t>Enterprise Systems Manager    March 2023 - Present</w:t></w:r></w:p>
<w:p><w:r><w:t>Known Company | Knoxville, TN</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Led implementation scope, data migration validation, training, and go-live readiness across five sites for 150+ users.</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>Built KPI reporting and stakeholder visibility that improved adoption, decision-making, and post-go-live follow-through.</w:t></w:r></w:p>
<w:p><w:r><w:t>Core Competencies</w:t></w:r></w:p>
<w:p><w:r><w:t>Implementation | Data Migration | Go-Live | Reporting | Training | Stakeholder Communication | KPI Reporting | Adoption | UAT | Change Enablement | Workflow Analysis | Process Improvement | Analytics | SQL | Power BI</w:t></w:r></w:p>
<w:p><w:r><w:t>Education</w:t></w:r></w:p>
</w:body></w:document>"""
    profile = SimpleNamespace(
        lane_label="Implementation Delivery",
        direct_matches=("Implementation Delivery",),
        unsupported_requirements=(),
        specialty_gaps=(),
        primary_lane="implementation_delivery",
    )
    original_job_problem_profile = build_resume.job_problem_profile
    original_alignment_score_report = build_resume.alignment_score_report
    original_poor_fit_requirements = build_resume.poor_fit_requirements
    original_keyword_placement_audit = build_resume.keyword_placement_audit
    original_hiring_manager_skim_issues = build_resume.hiring_manager_skim_issues
    original_resume_readiness_report = build_resume.resume_readiness_report
    original_extract_competency_items = build_resume.extract_competency_items
    try:
        build_resume.job_problem_profile = lambda *_args, **_kwargs: profile
        build_resume.alignment_score_report = lambda *_args, **_kwargs: {"total_score": 90}
        build_resume.poor_fit_requirements = lambda *_args, **_kwargs: ()
        build_resume.keyword_placement_audit = lambda *_args, **_kwargs: {"gaps": []}
        build_resume.hiring_manager_skim_issues = lambda *_args, **_kwargs: []
        build_resume.resume_readiness_report = lambda *_args, **_kwargs: SimpleNamespace(hard_blockers=())
        build_resume.extract_competency_items = lambda *_args, **_kwargs: tuple(f"Skill {index}" for index in range(1, 16))
        with TemporaryDirectory(prefix="final_fit_none_grade_") as temp_name:
            document_xml = Path(temp_name) / "document.xml"
            document_xml.write_text(xml, encoding="utf-8")
            status, notes = build_resume.final_fit_audit(document_xml, DUMMY_JOB_DESCRIPTION, alignment_grade=None)
    finally:
        build_resume.job_problem_profile = original_job_problem_profile
        build_resume.alignment_score_report = original_alignment_score_report
        build_resume.poor_fit_requirements = original_poor_fit_requirements
        build_resume.keyword_placement_audit = original_keyword_placement_audit
        build_resume.hiring_manager_skim_issues = original_hiring_manager_skim_issues
        build_resume.resume_readiness_report = original_resume_readiness_report
        build_resume.extract_competency_items = original_extract_competency_items

    assert_true(
        status == "PASS" and not any("Targeting bridge is weak" in note for note in notes),
        f"final_fit_audit() should not hard-fail the direct-match gate when alignment_grade is omitted; got status={status!r}, notes={notes}",
    )


def test_resume_builder_source_passes_alignment_grade_to_final_fit_audit(build_resume: object) -> None:
    source = Path(build_resume.__file__).read_text(encoding="utf-8")
    assert_true(
        'alignment_grade=str(alignment_report.get("grade", "")).strip() or None' in source,
        "build_resume.py should pass the computed alignment grade into final_fit_audit() on the main resume-build path",
    )


def test_application_checklist_fit_classification_passes_alignment_grade(build_application_checklist: object) -> None:
    original_final_fit_audit = build_application_checklist.build_resume.final_fit_audit
    original_alignment_score_report = build_application_checklist.build_resume.alignment_score_report
    captured: dict[str, object] = {}
    try:
        build_application_checklist.build_resume.alignment_score_report = lambda *_args, **_kwargs: {"grade": "Strong Fit"}

        def fake_final_fit_audit(document_xml, job_description, **kwargs):
            captured["alignment_grade"] = kwargs.get("alignment_grade")
            return ("PASS", [])

        build_application_checklist.build_resume.final_fit_audit = fake_final_fit_audit
        with TemporaryDirectory(prefix="checklist_grade_") as temp_name:
            resume_path = build_minimal_docx(
                Path(temp_name) / "resume.docx",
                [
                    "Professional Summary",
                    "Implementation consultant with measurable delivery outcomes.",
                    "Professional Experience",
                    "Enterprise Systems Manager    March 2023 - Present",
                    "Known Company | Knoxville, TN",
                    "Built rollout reporting and training support.",
                    "Education",
                ],
            )
            status = build_application_checklist.fit_classification(resume_path, DUMMY_JOB_DESCRIPTION)
    finally:
        build_application_checklist.build_resume.final_fit_audit = original_final_fit_audit
        build_application_checklist.build_resume.alignment_score_report = original_alignment_score_report

    assert_true(status == "PASS", f"fit_classification() should preserve the fake PASS return; got {status!r}")
    assert_true(
        captured.get("alignment_grade") == "Strong Fit",
        f"fit_classification() should pass a non-empty alignment grade into final_fit_audit(); got {captured}",
    )


def test_summary_anchor_retention_for_change_adoption(build_resume: object) -> None:
    import prose_engine

    job_description = """
Role: Change Enablement Manager
Lead change adoption, training, manager readiness, role-based enablement, stakeholder alignment,
communications, adoption measurement, and measurable follow-through across system changes.
"""
    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    summary = build_resume.build_problem_first_summary(job_description, resume_text, variant_index=0)
    sentences = build_resume.summary_sentences(summary)
    repair = prose_engine.repair_text(summary, "summary")
    assert_true(
        70 <= word_count(summary) <= 110,
        f"Change-adoption summary should stay within the 70-110 word contract after anchor retention; got {word_count(summary)} words in {summary!r}",
    )
    assert_true(
        len(sentences) == 3 and bool(re.search(r"\d|\$", sentences[1])),
        f"Quantified anchor should stay inside the existing proof sentence; got sentences={sentences}",
    )
    assert_true(
        repair.converged,
        f"Summary anchor guard should still converge through prose_engine.repair_text(); got {repair.findings}",
    )


def test_summary_quantified_anchor_detects_multi_digit_proof() -> None:
    import resume_content

    assert_true(
        resume_content.summary_has_quantified_anchor("Built 200+ dashboards and KPI tools for executive reporting."),
        "summary_has_quantified_anchor() should recognize multi-digit plus-count proof.",
    )
    assert_true(
        resume_content.summary_has_quantified_anchor("Cut manual inventory work by 78% through automated adjustments."),
        "summary_has_quantified_anchor() should recognize percentage proof.",
    )


def test_summary_repair_preserves_three_sentences_for_high_signal_variants(build_resume: object) -> None:
    import prose_engine
    import resume_content

    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    cases = (
        (
            "decision",
            """
Company: Dematic
Role: Solution Consultant
This role turns operations and engineering requirements into scoped recommendations, stakeholder-ready implementation plans,
workflow clarity, training support, and adoption follow-through across complex enterprise rollouts.
""",
            "QBRs",
        ),
        (
            "ai",
            """
Company: Dematic
Role: Solution Consultant
This role blends AI-assisted workflow automation, documentation, reporting, SQL troubleshooting, stakeholder communication,
training, and implementation clarity across enterprise software delivery.
""",
            "Claude",
        ),
        (
            "launch",
            """
Company: Dematic
Role: Solution Consultant
This role supports warehouse launch readiness, Amazon Robotics coordination, migration planning, testing, go-live execution,
and cross-functional implementation follow-through for complex customer rollouts.
""",
            "Amazon Robotics",
        ),
    )
    for proof_anchor, job_description, expected_term in cases:
        profile = build_resume.job_problem_profile(job_description, resume_text)
        emphasis = SimpleNamespace(proof_anchor=proof_anchor)
        positioning = resume_content.summary_positioning_sentence(profile, job_description, emphasis)
        proof = resume_content.summary_proof_sentence(profile, job_description, emphasis)
        close = resume_content.summary_fit_close_sentence(profile, job_description, emphasis)
        summary = resume_content.ensure_summary_minimum_words(positioning, proof, close, profile, job_description)
        summary = resume_content.neutralize_conflicting_region_lists(summary)
        findings = prose_engine.validate_text(summary, "summary")
        hard_failures = [finding.rule_id for finding in findings if finding.severity == "fail"]
        repair = prose_engine.repair_text(summary, "summary")
        sentences = build_resume.summary_sentences(repair.text)

        assert_true(
            not hard_failures,
            f"{proof_anchor} summary should avoid prose hard failures before repair; got {hard_failures} in {summary}",
        )
        assert_true(
            repair.converged,
            f"{proof_anchor} summary should still converge through prose repair; got {repair.findings}",
        )
        assert_true(
            len(sentences) == 3,
            f"{proof_anchor} summary should stay at three recruiter-friendly sentences after repair; got {sentences}",
        )
        assert_true(
            expected_term in repair.text,
            f"{proof_anchor} summary should preserve its core proof term after repair; got {repair.text}",
        )


def test_substitution_safety_respects_paragraph_boundaries() -> None:
    import text_safety

    safe_text = (
        "Christian Estrada Solution Consultant | Enterprise Systems\n"
        "Enterprise systems and change adoption consultant with 10+ years improving rollout clarity."
    )
    unsafe_text = (
        "Owned the enterprise platform while documenting the enterprise platform migration for users."
    )
    assert_true(
        "SUBSTITUTION_DUPLICATE_GENERIC_TERM" not in text_safety.substitution_safety_issues(safe_text),
        "substitution_safety_issues() should not merge header lines and the next paragraph into one duplicate-term sentence.",
    )
    assert_true(
        "SUBSTITUTION_DUPLICATE_GENERIC_TERM" in text_safety.substitution_safety_issues(unsafe_text),
        "substitution_safety_issues() should still flag duplicate generic platform terms inside one sentence.",
    )


def test_cover_selection_prefers_lane_direct_and_keeps_erp_for_erp_jd(build_cover_letter: object, build_resume: object) -> None:
    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    selected = build_cover_letter.selected_evidence_items(PROCORE_JOB_DESCRIPTION, resume_text)
    lane_terms = {"implementation", "go-live", "requirements", "configuration", "migration"}
    assert_true(
        bool(set(signal.lower() for signal in selected[0].signals) & lane_terms),
        f"selected_evidence_items() should keep the first proof lane-direct when possible; got first item={selected[0]}",
    )
    assert_true(
        any("erp" in item.text.lower() for item in selected),
        f"ERP-direct job descriptions should not strip ERP proof out of cover evidence selection; got {selected}",
    )


def test_cover_opening_quality_flags_article_and_circularity(build_cover_letter: object) -> None:
    article_problem = build_cover_letter.opening_quality_problem(
        "I am interested in a Implementation Consultant role because the work is concrete and measurable.",
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    circular_problem = build_cover_letter.opening_quality_problem(
        "The Implementation Consultant role fits my Implementation Consultant background because the Implementation Consultant work mirrors the Implementation Consultant work I already know.",
        "Acme Systems",
        "Implementation Consultant",
        DUMMY_JOB_DESCRIPTION,
    )
    assert_true(
        article_problem is not None and "article" in article_problem.lower(),
        f"opening_quality_problem() should reject article-agreement issues; got {article_problem!r}",
    )
    assert_true(
        circular_problem is not None and any(term in circular_problem.lower() for term in ("circular", "low-content", "generic", "direct, concrete")),
        f"opening_quality_problem() should reject circular or low-content openings; got {circular_problem!r}",
    )


def test_canonical_catalog_preserves_provenance_identity(commercial_resume_model: object) -> None:
    import evidence_engine

    catalog = evidence_engine.commercial_canonical_evidence()
    grouped: dict[tuple[str, str], list[object]] = {}
    for index, item in enumerate(catalog):
        grouped.setdefault((item.role, item.employer), []).append(
            commercial_resume_model.ProvenancedText(
                text=item.claim,
                provenance=(
                    commercial_resume_model.SourceRef(
                        source_path="catalog",
                        role_title=item.role,
                        employer=item.employer,
                        paragraph_kind="bullet",
                        paragraph_index=index,
                        source_text=item.source_text,
                        transformation="canonical-evidence",
                    ),
                ),
            )
        )
    summary_ref = catalog[0]
    model = commercial_resume_model.CommercialResumeModel(
        source_path="catalog",
        summary=(
            commercial_resume_model.ProvenancedText(
                text="Supported summary proof.",
                provenance=(
                    commercial_resume_model.SourceRef(
                        source_path="catalog",
                        role_title=summary_ref.role,
                        employer=summary_ref.employer,
                        paragraph_kind="professional-summary",
                        paragraph_index=0,
                        source_text=summary_ref.source_text,
                        transformation="canonical-evidence",
                    ),
                ),
            ),
        ),
        roles=tuple(
            commercial_resume_model.CommercialRoleModel(
                title=role_title,
                employer=employer,
                company_context="",
                summaries=(),
                bullets=tuple(items),
            )
            for (role_title, employer), items in grouped.items()
        ),
        content_hash="canonical-catalog",
    )
    commercial_resume_model.validate_content_model(model)
    assert_true(
        all(
            ref.employer == role.employer
            for role in model.roles
            for item in role.bullets
            for ref in item.provenance
        ),
        "Canonical evidence provenance should preserve same-employer mapping for commercial_resume_model validation.",
    )


def test_interview_filters_filler_and_claim_first_answers(build_resume: object, build_interview_cheat_sheet: object, interview_context: object) -> None:
    import question_prep

    questions = interview_context.extract_interviewer_questions(
        "\n".join(
            [
                '1) "Does that make sense?"',
                '2) "How do you prioritize competing work?"',
                '3) "Am I following that correctly?"',
            ]
        )
    )
    assert_true(
        questions == ("How do you prioritize competing work?",),
        f"extract_interviewer_questions() should filter filler questions while keeping substantive ones; got {questions}",
    )

    snapshot = question_prep.parse_resume(build_resume.IMPLEMENTATION_RESUME)
    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    answer = question_prep.answer_prompt(
        "How do you prioritize competing work?",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
        snapshot,
        resume_text,
    ).answer
    build_interview_cheat_sheet.assert_claim_then_prove_answer("prioritization", answer)

    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"], resume_text)
    desire_focus = build_interview_cheat_sheet.closing_desire_focus(
        profile,
        "Role: Implementation Consultant\nPriority: implementation scope, UAT, and go-live readiness.",
    )
    assert_true(
        "emphasis on" not in desire_focus.lower(),
        f"closing_desire_focus() should no longer emit the stale 'emphasis on {{role}}' artifact; got {desire_focus!r}",
    )


def test_candidate_facing_outputs_avoid_raw_core_problem(build_resume: object, build_cover_letter: object, build_interview_cheat_sheet: object) -> None:
    import question_prep

    resume_text = build_resume.docx_visible_text_from_path(build_resume.IMPLEMENTATION_RESUME)
    profile = build_resume.job_problem_profile(LANE_JOB_DESCRIPTIONS["implementation_delivery"], resume_text)
    brief = question_prep.build_positioning_brief(
        "Acme Systems",
        "Implementation Consultant",
        LANE_JOB_DESCRIPTIONS["implementation_delivery"],
        resume_text,
    )
    candidate_text = "\n".join(
        [
            build_cover_letter.proof_first_support_paragraph(brief),
            *build_interview_cheat_sheet.closing_fit_summary(profile),
            *build_interview_cheat_sheet.phone_screen_first_round_playbook(
                profile,
                "Acme Systems",
                "Implementation Consultant",
                LANE_JOB_DESCRIPTIONS["implementation_delivery"],
                resume_text,
            ),
        ]
    )
    assert_true(
        profile.core_problem not in candidate_text,
        f"Candidate-facing prose should use natural problem phrasing instead of the raw lane core_problem string; got {candidate_text!r}",
    )


def test_packet_excerpt_resolution_and_contract(build_claude_review_packet: object) -> None:
    warnings = build_claude_review_packet.packet_excerpt_resolution_warnings(
        build_claude_review_packet.PACKET_MODES["broad"]
    )
    contract = build_claude_review_packet.current_system_contract()
    assert_true(
        not warnings,
        f"Packet excerpt resolution guard should find live packet functions and files; got {warnings}",
    )
    assert_true(
        "DRAFT" in contract and "pages at" in contract and "10pt" in contract,
        f"current_system_contract() should document DRAFT plus the federal two-page/10pt minimum contract; got {contract!r}",
    )


def test_bridge_evidence_sweep_covers_multiple_lanes() -> None:
    from config import job_profiles

    labels = {item["label"] for item in job_profiles.BRIDGE_EVIDENCE_AREAS}
    assert_true(
        "Client-Side ERP Ownership and Finance Partnership" in labels
        and "Project-Based ERP Delivery" in labels
        and "Metrics and Decision Support" in labels
        and "Customer and Revenue Outcomes" in labels,
        f"Bridge evidence coverage should include the new client-side and project-based additions without collapsing to a single lane; got {sorted(labels)}",
    )


def test_cleanup_render_checks_limits_nested_cleanup_to_timestamped_folders() -> None:
    import cleanup_render_checks
    import os

    original_project_root = cleanup_render_checks.PROJECT_ROOT
    original_render_root = cleanup_render_checks.RENDER_ROOT
    try:
        with TemporaryDirectory(prefix="render_cleanup_") as temp_name:
            root = Path(temp_name)
            render_root = root / "render_check"
            render_root.mkdir()

            timestamped = render_root / "Example_Render_20260610_120000"
            timestamped.mkdir()
            (timestamped / "page-1.png").write_text("png", encoding="utf-8")

            manual_folder = render_root / "notes_for_manual_review"
            manual_folder.mkdir()
            (manual_folder / "keep.txt").write_text("keep", encoding="utf-8")

            old_timestamp = (datetime.now() - timedelta(hours=2)).timestamp()
            os.utime(timestamped, (old_timestamp, old_timestamp))
            os.utime(manual_folder, (old_timestamp, old_timestamp))

            cleanup_render_checks.PROJECT_ROOT = root
            cleanup_render_checks.RENDER_ROOT = render_root
            matches = cleanup_render_checks.old_render_folders(1)
    finally:
        cleanup_render_checks.PROJECT_ROOT = original_project_root
        cleanup_render_checks.RENDER_ROOT = original_render_root

    assert_true(
        timestamped in matches and manual_folder not in matches,
        f"cleanup_render_checks.old_render_folders() should only target timestamped nested render folders; got {matches}",
    )


def test_bootstrap_copy_filter_preserves_render_checks_module() -> None:
    import bootstrap_canonical_repo

    assert_true(
        not bootstrap_canonical_repo.should_exclude(Path("scripts/render_checks.py"), is_dir=False),
        "bootstrap copy filter should keep scripts/render_checks.py in the canonical repo",
    )
    assert_true(
        bootstrap_canonical_repo.should_exclude(Path("render_check_20260711"), is_dir=True),
        "bootstrap copy filter should still exclude render_check* artifact directories",
    )
    assert_true(
        bootstrap_canonical_repo.should_exclude(Path("review_updated_sources_20260711"), is_dir=True),
        "bootstrap copy filter should still exclude review_updated_sources* artifact directories",
    )


def test_bootstrap_gitignore_ignores_generated_bundle_manifest() -> None:
    import bootstrap_canonical_repo

    assert_true(
        "Claude Review/BUNDLE_MANIFEST.json" in bootstrap_canonical_repo.BOOTSTRAP_GITIGNORE,
        "bootstrap gitignore should exclude the generated Claude Review bundle manifest",
    )


def test_bootstrap_copy_filter_excludes_retired_onedrive_markers() -> None:
    import bootstrap_canonical_repo

    assert_true(
        bootstrap_canonical_repo.should_exclude(Path("DO_NOT_RUN_FROM_ONEDRIVE.txt"), is_dir=False),
        "bootstrap copy filter should exclude the retired OneDrive sentinel from the canonical repo copy",
    )
    assert_true(
        bootstrap_canonical_repo.should_exclude(Path("run_resume.bat"), is_dir=False),
        "bootstrap copy filter should exclude retired root run_*.bat launchers so canonical launchers can be regenerated",
    )


def test_bootstrap_writes_live_canonical_launchers() -> None:
    import bootstrap_canonical_repo

    with TemporaryDirectory(prefix="bootstrap_launchers_") as temp_name:
        root = Path(temp_name)
        bootstrap_canonical_repo.write_canonical_launchers(root)

        resume_text = (root / "run_resume.bat").read_text(encoding="utf-8")
        federal_text = (root / "run_federal_resume.bat").read_text(encoding="utf-8")
        refresh_text = (root / "run_claude_refresh.bat").read_text(encoding="utf-8")
        guide_text = (root / "run_detailed_interview_guide.bat").read_text(encoding="utf-8")
        debrief_text = (root / "run_post_interview_debrief.bat").read_text(encoding="utf-8")
        bootstrap_text = (root / "run_canonical_bootstrap.bat").read_text(encoding="utf-8")

        assert_true(
            "workspace_health.py" in resume_text
            and 'choice /c RFD' in resume_text
            and 'call :run_task resume --resume-only' in resume_text
            and 'call :run_task resume' in resume_text
            and 'call :run_task dry-run' in resume_text
            and 'Resume created, but the cover letter was saved as DRAFT.' in resume_text
            and 'if not exist ".\\output" mkdir ".\\output"' in resume_text,
            "run_resume.bat should offer resume-only, full, and dry-run modes while treating a DRAFT cover letter as a handled outcome",
        )
        assert_true(
            "call :run_task federal-resume" in federal_text
            and "call :run_task federal-dry-run" in federal_text
            and 'if not exist ".\\output" mkdir ".\\output"' in federal_text,
            "run_federal_resume.bat should offer both federal resume and federal dry-run commands and pre-create output",
        )
        assert_true(
            "call :run_task claude-refresh --skip-checks" in refresh_text and "call :run_task claude-refresh" in refresh_text,
            "run_claude_refresh.bat should offer both checked and skip-checks refresh commands",
        )
        assert_true(
            "call :run_task interview" in guide_text
            and "call :run_task guide --stage %GUIDE_STAGE%" in guide_text
            and "interview cheat sheet" in guide_text
            and "Selection [1-7]" in guide_text
            and "GUIDE_STAGE=hr_screen" in guide_text
            and "build_interview_companions.py" in guide_text
            and "Recruiter Screen Prep companion" in guide_text
            and "90 Day One-Pager companion" in guide_text
            and 'if not exist ".\\output" mkdir ".\\output"' in guide_text,
            "run_detailed_interview_guide.bat should build the cheat sheet, prompt for seven guide stages, and optionally offer the shared companion outputs after pre-creating output",
        )
        assert_true(
            debrief_text.count("choice /c YN") >= 2 and "prepare-company-notes" in debrief_text and "call :run_task debrief" in debrief_text,
            "run_post_interview_debrief.bat should preserve separate prompts for company notes and debrief capture",
        )
        assert_true(
            "already the canonical repo" in bootstrap_text,
            "run_canonical_bootstrap.bat should explain that bootstrap requires an explicit external source workspace",
        )


def test_run_resume_workflow_parse_args_accepts_resume_only() -> None:
    import run_resume_workflow

    parsed = run_resume_workflow.parse_args(["--resume-only"])

    assert_true(
        parsed.resume_only and not parsed.include_cheat_sheet and not parsed.include_detailed_guide and not parsed.dry_run,
        f"run_resume_workflow.parse_args() should accept --resume-only as a distinct mode; got {parsed}",
    )


def test_bootstrap_retire_source_launchers_skips_nested_canonical_root() -> None:
    import bootstrap_canonical_repo

    with TemporaryDirectory(prefix="bootstrap_retire_") as temp_name:
        source_root = Path(temp_name)
        canonical_root = source_root / ".tmp" / "canonical_repo"
        source_launcher = source_root / "run_resume.bat"
        canonical_launcher = canonical_root / "run_resume.bat"
        source_launcher.write_text("source launcher\n", encoding="utf-8")
        canonical_launcher.parent.mkdir(parents=True, exist_ok=True)
        canonical_launcher.write_text("canonical launcher\n", encoding="utf-8")

        bootstrap_canonical_repo.retire_source_launchers(source_root, canonical_root)

        assert_true(
            "retired and archival only" in source_launcher.read_text(encoding="utf-8"),
            "retire_source_launchers() should still stub launchers in the retired source tree",
        )
        assert_true(
            canonical_launcher.read_text(encoding="utf-8") == "canonical launcher\n",
            "retire_source_launchers() should not stub launchers inside the nested canonical destination",
        )


def test_bootstrap_reset_destination_removes_nested_tree() -> None:
    import bootstrap_canonical_repo

    original_safe_reset_root = bootstrap_canonical_repo.SAFE_RESET_ROOT
    try:
        with TemporaryDirectory(prefix="bootstrap_reset_") as temp_name:
            safe_root = Path(temp_name)
            destination = safe_root / "canonical_repo"
            nested_file = destination / "scratch" / "run_logs" / "latest.log"
            nested_file.parent.mkdir(parents=True)
            nested_file.write_text("log", encoding="utf-8")

            bootstrap_canonical_repo.SAFE_RESET_ROOT = safe_root
            bootstrap_canonical_repo.prepare_destination(destination, reset_destination=True)

            assert_true(
                destination.exists() and not any(destination.iterdir()),
                f"prepare_destination(..., reset_destination=True) should recreate an empty destination; got {list(destination.iterdir()) if destination.exists() else 'missing'}",
            )
    finally:
        bootstrap_canonical_repo.SAFE_RESET_ROOT = original_safe_reset_root


def test_cleanup_output_finders_and_selective_flag() -> None:
    import cleanup_output
    import os

    original_project_root = cleanup_output.PROJECT_ROOT
    original_output_dir = cleanup_output.OUTPUT_DIR
    original_tracker_csv = cleanup_output.TRACKER_CSV
    try:
        with TemporaryDirectory(prefix="cleanup_output_") as temp_name:
            root = Path(temp_name)
            output_dir = root / "output"
            tracker_csv = root / "scratch" / "applications.csv"
            render_root = root / "render_check"
            output_dir.mkdir()
            tracker_csv.parent.mkdir(parents=True)
            render_root.mkdir()

            tracker_csv.write_text(
                "\n".join(
                    [
                        "date_added,company,role_title,lane_label,fit_status,audit_flag,source_resume,output_file,applied_date,current_status,last_round,outcome,notes",
                        "2026-06-10,Interview Co,Implementation Consultant,Implementation Delivery,Strong Fit,,resume.docx,,2026-06-10,interview,,,",
                        "2026-06-10,Rejected Co,Implementation Consultant,Implementation Delivery,Strong Fit,,resume.docx,,2026-06-10,rejected,,,",
                    ]
                ),
                encoding="utf-8",
            )

            protected_file = output_dir / "Christian Estrada - Interview Co - Resume.docx"
            rejected_file = output_dir / "Christian Estrada - Rejected Co - Resume.docx"
            stale_file = output_dir / "Christian Estrada - Old Co - Resume.docx"
            fresh_file = output_dir / "Christian Estrada - Fresh Co - Resume.docx"
            for path in (protected_file, rejected_file, stale_file, fresh_file):
                path.write_text("docx", encoding="utf-8")

            stale_timestamp = (datetime.now() - timedelta(days=cleanup_output.OUTPUT_MAX_DAYS + 5)).timestamp()
            fresh_timestamp = (datetime.now() - timedelta(days=5)).timestamp()
            os.utime(protected_file, (stale_timestamp, stale_timestamp))
            os.utime(rejected_file, (stale_timestamp, stale_timestamp))
            os.utime(stale_file, (stale_timestamp, stale_timestamp))
            os.utime(fresh_file, (fresh_timestamp, fresh_timestamp))

            old_root_render = root / "render_check_20260601"
            old_root_render.mkdir()
            nested_render = render_root / "Example_Render_20260601_120000"
            nested_render.mkdir()
            unrelated_folder = root / "render_archive"
            unrelated_folder.mkdir()
            render_timestamp = (datetime.now() - timedelta(days=cleanup_output.RENDER_CHECK_MAX_DAYS + 2)).timestamp()
            os.utime(old_root_render, (render_timestamp, render_timestamp))
            os.utime(nested_render, (render_timestamp, render_timestamp))
            os.utime(unrelated_folder, (render_timestamp, render_timestamp))

            cleanup_output.PROJECT_ROOT = root
            cleanup_output.OUTPUT_DIR = output_dir
            cleanup_output.TRACKER_CSV = tracker_csv

            stale_outputs = cleanup_output.find_stale_output_files()
            stale_renders = cleanup_output.find_stale_render_folders()
            parsed_args = cleanup_output.parse_args(["--selective"])
    finally:
        cleanup_output.PROJECT_ROOT = original_project_root
        cleanup_output.OUTPUT_DIR = original_output_dir
        cleanup_output.TRACKER_CSV = original_tracker_csv

    assert_true(
        stale_file in stale_outputs and protected_file not in stale_outputs and rejected_file not in stale_outputs and fresh_file not in stale_outputs,
        f"cleanup_output.find_stale_output_files() should keep protected and recent files while flagging only stale unprotected files; got {stale_outputs}",
    )
    assert_true(
        old_root_render in stale_renders and nested_render in stale_renders and unrelated_folder not in stale_renders,
        f"cleanup_output.find_stale_render_folders() should reuse the safe render-folder policy; got {stale_renders}",
    )
    assert_true(
        parsed_args.selective,
        "cleanup_output.parse_args() should accept the --selective flag",
    )


def test_tasks_register_cleanup_command() -> None:
    import tasks

    cleanup_task = tasks.TASKS["cleanup"]
    assert_true(
        cleanup_task.args == ("scripts/cleanup_output.py",),
        f"tasks.py should register cleanup against scripts/cleanup_output.py; got {cleanup_task.args}",
    )
    assert_true(
        not cleanup_task.needs_job_description,
        "cleanup should not require an active job description",
    )
    output_buffer = io.StringIO()
    with contextlib.redirect_stdout(output_buffer):
        tasks.print_command_inventory()
    inventory_output = output_buffer.getvalue()
    assert_true(
        "cleanup:" in inventory_output and "scripts/cleanup_output.py" in inventory_output,
        f"Command inventory should include the cleanup task; got {inventory_output!r}",
    )


def test_tasks_check_prefers_latest_generated_resume() -> None:
    import tasks

    original_validate_job_description = tasks.validate_job_description
    original_load_build_resume = tasks.load_build_resume
    original_matching_resume_outputs = tasks.matching_resume_outputs
    original_docx_visible_text = tasks.docx_visible_text
    original_job_description = tasks.JOB_DESCRIPTION
    source_resume = Path("source_resume.docx")
    output_resume = Path("output_resume.docx")
    captured_profile_texts: list[str] = []
    visible_text_paths: list[Path] = []

    with TemporaryDirectory(prefix="tasks_check_smoke_") as temp_name:
        temp_job = Path(temp_name) / "job_description.txt"
        temp_job.write_text("Company: Acme\nJob Title: Implementation Consultant\nImplementation delivery and adoption.", encoding="utf-8")

        build_resume_stub = SimpleNamespace(
            PRESALES_SIGNALS=(),
            PRESALES_CSM_RESUME=Path("presales_resume.docx"),
            validate_inputs=lambda text: text,
            choose_resume=lambda _job_description: source_resume,
            job_problem_profile=lambda _job_description, resume_text: captured_profile_texts.append(resume_text) or SimpleNamespace(
                primary_lane="implementation_delivery",
                lane_label="Implementation and Delivery",
                core_problem="implementation scope, technical complexity, timeline risk, and adoption readiness",
                audience="clients, project teams, executives, operations, and technical stakeholders",
                outcomes=("go-live readiness", "delivery quality"),
                direct_matches=("Implementation Delivery",),
                adjacent_matches=(),
                unsupported_requirements=(),
            ),
            audit_keywords=lambda _job_description: set(),
            keyword_placement_audit=lambda _job_description, _resume_text: {"gaps": []},
            poor_fit_requirements=lambda _job_description, _resume_text: (),
            primary_story_lens=lambda _job_description: {},
        )

        try:
            tasks.validate_job_description = lambda: True
            tasks.load_build_resume = lambda: build_resume_stub
            tasks.matching_resume_outputs = lambda _build_resume, _job_description: [output_resume]
            tasks.docx_visible_text = lambda path: visible_text_paths.append(path) or ("generated resume text" if path == output_resume else "source resume text")
            tasks.JOB_DESCRIPTION = temp_job
            output_buffer = io.StringIO()
            with contextlib.redirect_stdout(output_buffer), contextlib.redirect_stderr(io.StringIO()):
                result = tasks.run_check()
        finally:
            tasks.validate_job_description = original_validate_job_description
            tasks.load_build_resume = original_load_build_resume
            tasks.matching_resume_outputs = original_matching_resume_outputs
            tasks.docx_visible_text = original_docx_visible_text
            tasks.JOB_DESCRIPTION = original_job_description

    assert_true(result == 0, f"run_check() should return success with a generated resume artifact; got {result}")
    assert_true(
        visible_text_paths and visible_text_paths[0] == output_resume,
        f"run_check() should audit the latest generated resume first when one exists; got {visible_text_paths}",
    )
    assert_true(
        captured_profile_texts and captured_profile_texts[0] == "generated resume text",
        f"run_check() should profile the generated resume text rather than the source resume text; got {captured_profile_texts}",
    )
    assert_true(
        "Audit artifact: output_resume.docx (latest generated output)" in output_buffer.getvalue(),
        f"run_check() should explain when it is auditing a generated resume; got {output_buffer.getvalue()!r}",
    )


def test_claude_prompt_templates(build_claude_prompt: object) -> None:
    review_prompt = build_claude_prompt.build_prompt("review", "tracker", "", "TEMP_FOR_REVIEW.md")
    plan_prompt = build_claude_prompt.build_prompt("plan", "tracker", "debrief and workflow sync", "TEMP_FOR_REVIEW.md")
    claude_review_prompt = build_claude_prompt.build_prompt("review", "claude-review", "", "TEMP_FOR_REVIEW_CLAUDE_REVIEW.md")
    assert_true(
        "1. Findings" in review_prompt and "Single Most Useful Next File To Inspect" in review_prompt,
        "Review prompt template should require the strict findings output structure",
    )
    assert_true(
        "Packet mode: `tracker`" in review_prompt,
        "Review prompt should embed the selected packet mode",
    )
    assert_true(
        "1. Recommended Fix Order" in plan_prompt and "4. Optional Follow-Up Improvements" in plan_prompt,
        "Plan prompt template should require the strict implementation-plan output structure",
    )
    assert_true(
        "Focus area: debrief and workflow sync" in plan_prompt,
        "Plan prompt should preserve explicit focus overrides",
    )
    assert_true(
        "1. Overall Status" in claude_review_prompt
        and "7. Next Exact Action" in claude_review_prompt
        and "Interview Feature Track" in claude_review_prompt
        and "Commit Train Track" in claude_review_prompt,
        "Claude-review prompts should use the dedicated dual-track progress-check output contract",
    )
    assert_true(
        "ON TRACK" in claude_review_prompt and "NEEDS ATTENTION" in claude_review_prompt and "proposal" in claude_review_prompt.lower(),
        "Claude-review prompts should preserve the top-level health language and live-versus-proposal guidance",
    )


def test_claude_prompt_rejects_residual_placeholders(build_claude_prompt: object) -> None:
    original_read_template = build_claude_prompt.read_prompt_template
    try:
        build_claude_prompt.read_prompt_template = lambda _path: "Packet {{PACKET_MODE}} {{UNKNOWN_TOKEN}}"
        try:
            build_claude_prompt.build_prompt("review", "tracker", "", "TEMP_FOR_REVIEW.md")
            raise SmokeFailure("build_prompt() should reject unreplaced placeholder tokens")
        except ValueError as error:
            assert_true(
                "{{UNKNOWN_TOKEN}}" in str(error),
                f"build_prompt() should report the leftover placeholder token; got {error}",
            )
    finally:
        build_claude_prompt.read_prompt_template = original_read_template


def test_claude_prompt_requires_current_default_packet_manifest(build_claude_prompt: object) -> None:
    original_dir = build_claude_prompt.CLAUDE_REVIEW_DIR
    try:
        with TemporaryDirectory(prefix="claude_prompt_manifest_") as temp_name:
            temp_dir = Path(temp_name)
            build_claude_prompt.CLAUDE_REVIEW_DIR = temp_dir
            packet_path = temp_dir / "TEMP_FOR_REVIEW_TRACKER.md"
            packet_path.write_text("packet\n", encoding="utf-8")
            try:
                build_claude_prompt.packet_manifest_status(packet_path)
            except Exception as error:  # noqa: BLE001
                raise SmokeFailure(f"packet_manifest_status() should not crash on a missing manifest: {error}") from error
            current, detail = build_claude_prompt.packet_manifest_status(packet_path)
            assert_true(
                not current and "Missing packet manifest" in detail,
                f"packet_manifest_status() should name the missing manifest; got current={current}, detail={detail!r}",
            )
    finally:
        build_claude_prompt.CLAUDE_REVIEW_DIR = original_dir


def test_claude_review_bundle_builders_include_expected_phrases(claude_review_bundle: object) -> None:
    claude_md = claude_review_bundle.build_review_claude_md()
    rules = claude_review_bundle.build_review_rules_for_claude()

    assert_true(
        "Commercial filename audit states live today are PASS, BRIDGE, FAIL, and POOR." in claude_md,
        "build_review_claude_md() should include the current live commercial audit-state guidance",
    )
    assert_true(
        "Federal runs now produce two Word documents by default" in claude_md,
        "build_review_claude_md() should include the federal two-document contract note",
    )
    assert_true(
        "Federal resume output stays at exactly two pages" in rules,
        "build_review_rules_for_claude() should include the federal two-page formatting rule",
    )
    assert_true(
        "If you recommend a new status, command, enum, packet mode, or public contract" in rules,
        "build_review_rules_for_claude() should include the live-vs-proposal guardrail",
    )


def test_claude_review_bundle_refresh(claude_review_bundle: object) -> None:
    original_dir = claude_review_bundle.CLAUDE_REVIEW_DIR
    try:
        with TemporaryDirectory(prefix="claude_review_bundle_") as temp_name:
            temp_dir = Path(temp_name)
            claude_review_bundle.CLAUDE_REVIEW_DIR = temp_dir
            written = claude_review_bundle.refresh_support_files()

            assert_true(
                {path.name for path in written} >= {
                    "CLAUDE.md",
                    "RESUME_SYSTEM_BRIEF.md",
                    "RULES_FOR_CLAUDE.md",
                    "COMMON_CHANGE_AREAS.md",
                    "CLAUDE_PROGRESS_CHECK_TEMPLATE.md",
                    "UPLOAD_GUIDE.md",
                    "BUNDLE_MANIFEST.json",
                },
                "refresh_support_files() should refresh the curated Claude Review support set",
            )

            claude_text = (temp_dir / "CLAUDE.md").read_text(encoding="utf-8")
            brief_text = (temp_dir / "RESUME_SYSTEM_BRIEF.md").read_text(encoding="utf-8")
            rules_text = (temp_dir / "RULES_FOR_CLAUDE.md").read_text(encoding="utf-8")
            common_text = (temp_dir / "COMMON_CHANGE_AREAS.md").read_text(encoding="utf-8")
            guide_text = (temp_dir / "UPLOAD_GUIDE.md").read_text(encoding="utf-8")
            manifest = json.loads((temp_dir / "BUNDLE_MANIFEST.json").read_text(encoding="utf-8"))
    finally:
        claude_review_bundle.CLAUDE_REVIEW_DIR = original_dir

    assert_true(
        "Keep these current workflow assumptions in mind during review:" in claude_text,
        "Claude Review CLAUDE.md should include the upload-specific review assumptions block",
    )
    assert_true(
        "Federal Qualifications Statement" in brief_text,
        "Claude Review RESUME_SYSTEM_BRIEF.md should include the federal qualifications statement output",
    )
    assert_true(
        "Federal resume output stays at exactly two pages" in rules_text,
        "Claude Review RULES_FOR_CLAUDE.md should include the federal output formatting guardrail",
    )
    assert_true(
        "## Federal Resume and Qualifications Statement" in common_text,
        "Claude Review COMMON_CHANGE_AREAS.md should include the federal inspection block",
    )
    assert_true(
        "python tasks.py claude-refresh" in guide_text,
        "Claude Review upload guide should advertise the one-step refresh command",
    )
    assert_true(
        "TEMP_FOR_REVIEW_FEDERAL.md" in guide_text and "TEMP_FOR_REVIEW_CLAUDE_REVIEW.md" in guide_text,
        "Claude Review upload guide should include the federal and claude-review packet variants",
    )
    assert_true(
        manifest.get("workspace_health", {}).get("project_root"),
        "Claude Review bundle manifest should record the local project root and workspace snapshot",
    )


def test_refresh_claude_review_bundle(refresh_claude_review_bundle: object, claude_review_bundle: object) -> None:
    import build_claude_prompt
    import build_claude_review_packet

    original_bundle_dir = claude_review_bundle.CLAUDE_REVIEW_DIR
    original_packet_dir = build_claude_review_packet.CLAUDE_REVIEW_DIR
    original_prompt_dir = build_claude_prompt.CLAUDE_REVIEW_DIR
    claude_review_prompt = ""
    try:
        with TemporaryDirectory(prefix="claude_refresh_") as temp_name:
            temp_dir = Path(temp_name)
            claude_review_bundle.CLAUDE_REVIEW_DIR = temp_dir
            build_claude_review_packet.CLAUDE_REVIEW_DIR = temp_dir
            build_claude_prompt.CLAUDE_REVIEW_DIR = temp_dir
            written = refresh_claude_review_bundle.refresh_bundle(("broad", "interview", "federal", "claude-review"), skip_checks=True)
            claude_review_prompt = next(path for path in written if path.name == "TEMP_CLAUDE_REVIEW_PROMPT_CLAUDE_REVIEW.txt").read_text(encoding="utf-8")
    finally:
        claude_review_bundle.CLAUDE_REVIEW_DIR = original_bundle_dir
        build_claude_review_packet.CLAUDE_REVIEW_DIR = original_packet_dir
        build_claude_prompt.CLAUDE_REVIEW_DIR = original_prompt_dir

    written_names = {path.name for path in written}
    assert_true(
        {
            "TEMP_FOR_REVIEW.md",
            "TEMP_FOR_REVIEW.manifest.json",
            "TEMP_FOR_REVIEW_INTERVIEW.md",
            "TEMP_FOR_REVIEW_INTERVIEW.manifest.json",
            "TEMP_FOR_REVIEW_FEDERAL.md",
            "TEMP_FOR_REVIEW_FEDERAL.manifest.json",
            "TEMP_FOR_REVIEW_CLAUDE_REVIEW.md",
            "TEMP_FOR_REVIEW_CLAUDE_REVIEW.manifest.json",
            "TEMP_CLAUDE_REVIEW_PROMPT_BROAD.txt",
            "TEMP_CLAUDE_PLAN_PROMPT_INTERVIEW.txt",
            "TEMP_CLAUDE_REVIEW_PROMPT_CLAUDE_REVIEW.txt",
        }.issubset(written_names),
        "refresh_bundle() should rebuild the common packet, manifest, and prompt files for the broad, interview, federal, and claude-review modes",
    )
    assert_true(
        "Interview Feature Track" in claude_review_prompt and "Commit Train Track" in claude_review_prompt,
        "refresh_bundle() should generate the dedicated dual-track Claude progress-check prompt for the claude-review mode",
    )


def main() -> None:
    checks = (
        ("import config files", import_config_files),
        ("import major scripts", lambda: import_major_scripts()),
        ("AGENTS word budget", None),
        ("validate dummy inputs", None),
        ("choose resume", None),
        ("lane profiles and summaries", None),
        ("first person detector ignores role level i", None),
        ("summary three sentence structure", None),
        ("boomi presales summary", None),
        ("customer success summary capitalization", None),
        ("customer success portfolio scope", None),
        ("general consulting lane isolation", None),
        ("strategy transformation consultant lane override", None),
        ("consulting audit keyword cleanup", None),
        ("audit keywords filter company affiliates and weak fragments", None),
        ("audit keywords filter boilerplate adjectives and normalize es plurals", None),
        ("keyword placement prefers role phrases", None),
        ("AI bridge evidence config", None),
        ("construction engineering context and bridge coverage", None),
        ("alignment score ignores unsupported named platforms", None),
        ("bridge evidence sweep covers multiple lanes", None),
        ("summary condense guard", None),
        ("ERP summary rebalance", None),
        ("summary anchor retention for change adoption", None),
        ("alignment score distinguishes lane and domain fit", None),
        ("cover letter colon smoothing", None),
        ("cover sentence score prioritizes signal density over length", None),
        ("mission or context sentence survives job label header", None),
        ("word budget trims opening filler before dense proof", None),
        ("cover selection prefers lane direct and keeps ERP for ERP JD", None),
        ("cover opening quality flags article and circularity", None),
        ("consulting story summary avoids list density overload", None),
        ("cover letter compaction", None),
        ("cover letter blocks JD artifacts and warns on switch", None),
        ("cover letter allows requirements in normal prose", None),
        ("cover letter sections recognize extended headers", None),
        ("cover letter plan normalizes slash titles", None),
        ("change consulting cover letter stays in change lane", None),
        ("communication metric fragment rejection", None),
        ("cover letter synthetic JD cleanup", None),
        ("role title dash not treated as bullet artifact", None),
        ("cover prompt leak pattern catches maps directly to", None),
        ("naturalness score and adverb cleanup", None),
        ("ownership language rewrites", None),
        ("build interview review sections", None),
        ("positive question framing", None),
        ("alignment score report", None),
        ("alignment gate decision", None),
        ("dynamic header title line", None),
        ("header dedupe avoids near-duplicate consultant titles", None),
        ("scope marker injection", None),
        ("competency relevance and page guards", None),
        ("final fit audit accepts presales top role heading", None),
        ("competency cap after additions", None),
        ("XML page estimate word guard", None),
        ("resume non-ERP audit ignores company context", None),
        ("keyword placement audit", None),
        ("build resume uses selected resume text for profile and alignment", None),
        ("obvious choice positioning", None),
        ("positioning statement output", None),
        ("future bridge summary and bullet clause", None),
        ("startup operator summary structure", None),
        ("natural top bullet meta penalty", None),
        ("offer blocker logic", None),
        ("ats plain text validation", None),
        ("moment in time context", None),
        ("customer success opening concrete context", None),
        ("customer success support paragraph separation", None),
        ("cover letter uses company specific context", None),
        ("gap address paragraph", None),
        ("first 90 days cover sentence", None),
        ("diagnose before selling framework", None),
        ("bold diagnostic questions", None),
        ("preloaded questions", None),
        ("slot based summary and interview answers", None),
        ("value compression opening", None),
        ("pitch variants reuse cover logic", None),
        ("pitch variants add 15-second claim", None),
        ("claim-then-prove validator catches delayed openings", None),
        ("natural voice validation and answer budgets", None),
        ("natural voice question coverage", None),
        ("keyword reference uses natural story intro", None),
        ("industry depth and company-scoped logistics", None),
        ("human motivation lane coverage", None),
        ("adjusted profile preserves non-lane fields", None),
        ("story answer parts preserve spoken alias", None),
        ("story answer constructors use full keyword", None),
        ("extended TMAY sections", None),
        ("extended TMAY uses module-level supersets", None),
        ("behavioral answer scripts empty story guard", None),
        ("story sample answer separates coaching note", None),
        ("story sample answer reuses claim sentence in full", None),
        ("delivery validator only scans scripted strings", None),
        ("scripted answer validator rejects unsupported metrics", None),
        ("AI customer work answer uses confirmed qualitative story", None),
        ("story sample answer does not call dead STAR selection", None),
        ("behavioral answer scripts use spoken answers", None),
        ("interview stage resolution and context parsing", None),
        ("stage filename suffix composes with detailed guide names", None),
        ("federal detailed guide wrapper keeps stage params optional", None),
        ("story adaptation and pre-interview routine helpers", None),
        ("spoken sentence split preserves leading characters", None),
        ("nonconverged spoken repairs are collected", None),
        ("interview filters filler and claim first answers", None),
        ("candidate facing outputs avoid raw core problem", None),
        ("first 90 day plan reuses shared stage source", None),
        ("application answers use written confirmation and automation boundary", None),
        ("detailed guide stage patterns and debrief overlay", None),
        ("interview companion documents reuse shared stage sources", None),
        ("federal version control scope gate is unconditional", None),
        ("Foundant summary uses human close", None),
        ("commercial resume model provenance and render", None),
        ("canonical catalog preserves provenance identity", None),
        ("application question pairing detects JD swap", None),
        ("repo guidance prefers rehearsed foundation", None),
        ("pitch helpers handle missing cover letter pitch parts", None),
        ("read the room opening", None),
        ("companies refer to same", None),
        ("debrief active job matching", None),
        ("compact context budget and diagnosis", None),
        ("great eight shared utility", None),
        ("interview cheat sheet static content guards", None),
        ("late-stage context detection", None),
        ("interview bullet marker is clean", None),
        ("why-company and negotiation context", None),
        ("communication audit context", None),
        ("story quality audit by type", None),
        ("lead burial and response calibration", None),
        ("four trust questions audit", None),
        ("executive presence signals", None),
        ("lane specific recording focus", None),
        ("debrief pattern analysis", None),
            ("interview followup body", None),
            ("post-round followup is lane aware", None),
            ("build skills database", None),
            ("default question when application questions file is empty", None),
            ("dry-run reports default question usage", None),
            ("default questions skip stale pairing check", None),
            ("explicit stale questions still flagged", None),
            ("qualifications builder uses question prep response engine", None),
            ("qualifications builder removes local shadow answer helpers", None),
            ("standard qualifications answers known questions", None),
            ("startup interview false-positive guard", None),
            ("application checklist debrief lookup", None),
            ("general advice active job helpers", None),
            ("general advice shared sections", None),
        ("question intent framework", None),
        ("interview addition helpers", None),
        ("search progress question is conditional", None),
        ("salary guide helpers", None),
        ("proof text rewrites dense cover sentence", None),
        ("follow-up news and variants", None),
        ("LinkedIn guidance helpers", None),
        ("cover letter signals ollie analytics", None),
        ("cover opening without mission still names company", None),
        ("cover lane prefers strategy consulting titles", None),
        ("standard cover mode", None),
        ("force-bridge standard cover stays natural", None),
        ("cover letter validator blocks generic experience summary", None),
        ("cover letter QC rejects lowercase proof paragraph", None),
        ("long cover mode", None),
        ("cover letter inherits fail resume name", None),
        ("ollie cover acceptance", None),
        ("cover letter validator blocks contractions and double-dashes", None),
        ("clean but bad cover regressions fail writing eval", None),
        ("registered firm profile requires real bain name", None),
        ("bigfour cover opening avoids aspiration phrase", None),
        ("reset jobs helpers", None),
        ("job context archive active snapshot includes questions", None),
        ("cover letter trace records snapshot and selection debug", None),
        ("run-level ERP scrub formatting", None),
        ("supported rewrite ERP scrub", None),
        ("non-ERP audit allows SAP Crystal Reports", None),
        ("company profile stub", None),
        ("hiring manager skim lane terms", None),
        ("Aptean customer success role summary passes prose check", None),
        ("proof first close uses discuss and has no first person switch", None),
        ("enterprise CS JD extracts CS terms not analytics fallback", None),
        ("role bullet budgets meet minimums", None),
        ("contact constants", None),
        ("business context module", None),
        ("audit keywords filter noisy bigrams", None),
        ("job title label stripping", None),
        ("thank-you contact line filter", None),
        ("thank-you proof points skip summary", None),
        ("company constants", None),
        ("role heading detection", None),
        ("writing eval flags system narration", None),
        ("writing eval passes clean summary", None),
        ("writing eval loads file dataset", None),
        ("writing eval extracts docx sections", None),
        ("summary detectors avoid member and operator false positives", None),
        ("implementation summary surfaces quality and process language", None),
        ("detailed guide notes context strips leading bullets", None),
        ("extract writing examples writes snippets", None),
        ("tracker refresh uses active-job fit status", None),
        ("tracker status precedence blocks regressions", None),
        ("tracker row job description prefers snapshot id", None),
        ("tracker refresh warning when output missing", None),
        ("collapse redundant role blanks", None),
        ("resume experience alignment", None),
        ("LinkedIn boilerplate cleanup", None),
        ("role requirement text resumes after about-us sections", None),
        ("planned competency trim integrity", None),
        ("federal agency extraction", None),
        ("federal source load", None),
        ("federal standard essay responses", None),
        ("federal summary structure", None),
        ("federal requirement audit and keywords", None),
        ("federal layouts stay at ten point", None),
        ("federal visibility report tracks selected requirements", None),
        ("federal standalone agency output name", None),
        ("federal supporting doc resolution", None),
        ("federal ATS plain-text warnings split hours and salary", None),
        ("federal resume plan warns on unverified page count", None),
        ("commercial builder entrypoints delegate to helpers", None),
        ("federal workflow supporting flags", None),
        ("federal workflow supporting step order", None),
        ("resume workflow dry-run skips upfront validation", None),
        ("federal workflow dry-run skips upfront validation", None),
        ("federal workflow dry-run labels unverified page counts", None),
        ("tasks register federal supporting doc commands", None),
        ("tasks register interview review command", None),
        ("onedrive run guard absent", None),
        ("onedrive run guard present", None),
        ("resume builder source passes alignment grade to final fit audit", None),
        ("application checklist fit classification passes alignment grade", None),
        ("final fit audit safe none alignment grade", None),
        ("claude packet run command keeps stdout and stderr", None),
        ("claude packet command cache", None),
        ("claude packet self-audit checks task registration", None),
        ("claude packet self-audit accepts unquoted task mentions", None),
        ("packet excerpt resolution and contract", None),
        ("tasks auto-archive environment for commercial command", None),
        ("workflow parses cover warning channels", None),
        ("workflow hard-fails docx validation issues", None),
        ("cleanup render checks only targets timestamped nested folders", None),
        ("bootstrap copy filter preserves render checks module", None),
        ("bootstrap gitignore ignores generated bundle manifest", None),
        ("bootstrap copy filter excludes retired onedrive markers", None),
        ("bootstrap writes live canonical launchers", None),
        ("bootstrap retire source launchers skips nested canonical root", None),
        ("bootstrap reset destination removes nested tree", None),
        ("cleanup output finders and selective flag", None),
        ("tasks check prefers latest generated resume", None),
        ("tasks register cleanup command", None),
        ("claude prompt rejects residual placeholders", None),
        ("claude prompt requires current default packet manifest", None),
        ("claude review bundle builders include expected phrases", None),
    )
    passed = 0
    executed_checks = 0
    failures: list[str] = []

    try:
        import_config_files()
        passed += 1
    except BaseException as error:  # noqa: BLE001
        failures.append(f"import config files: {error}")

    modules = {}
    if not failures:
        try:
            modules = import_major_scripts()
            passed += 1
        except BaseException as error:  # noqa: BLE001
            failures.append(f"import major scripts: {error}")

    if modules:
        build_resume = modules["build_resume"]
        build_claude_prompt = modules["build_claude_prompt"]
        build_claude_review_packet = modules["build_claude_review_packet"]
        claude_review_bundle = modules["claude_review_bundle"]
        refresh_claude_review_bundle = modules["refresh_claude_review_bundle"]
        build_federal_resume = modules["build_federal_resume"]
        build_federal_cover_letter = modules["build_federal_cover_letter"]
        build_federal_interview_cheat_sheet = modules["build_federal_interview_cheat_sheet"]
        build_federal_detailed_interview_guide = modules["build_federal_detailed_interview_guide"]
        build_followup_email = modules["build_followup_email"]
        build_cover_letter = modules["build_cover_letter"]
        commercial_resume_model = modules["commercial_resume_model"]
        build_standard_qualifications_statement = modules["build_standard_qualifications_statement"]
        build_interview_cheat_sheet = modules["build_interview_cheat_sheet"]
        build_detailed_interview_guide = modules["build_detailed_interview_guide"]
        build_interview_review = modules["build_interview_review"]
        build_debrief_analysis = modules["build_debrief_analysis"]
        build_interview_followup = modules["build_interview_followup"]
        build_post_round = modules["build_post_round"]
        build_linkedin_update = modules["build_linkedin_update"]
        build_salary_guide = modules["build_salary_guide"]
        build_skills_database = modules["build_skills_database"]
        build_application_checklist = modules["build_application_checklist"]
        build_thank_you = modules["build_thank_you"]
        business_context = modules["business_context"]
        build_general_advice = modules["build_general_advice"]
        federal_supporting_docs = modules["federal_supporting_docs"]
        interview_stage = modules["interview_stage"]
        interview_context = modules["interview_context"]
        job_search_guidance = modules["job_search_guidance"]
        reset_jobs = modules["reset_jobs"]
        extract_writing_examples = modules["extract_writing_examples"]
        utils = modules["utils"]
        writing_eval = modules["writing_eval"]
        question_prep = build_detailed_interview_guide.question_prep
        for label, check in (
            ("AGENTS word budget", test_agents_word_budget),
            ("validate dummy inputs", lambda: test_validate_inputs(build_resume)),
            ("LinkedIn boilerplate cleanup", lambda: test_validate_inputs_strips_linkedin_boilerplate(build_resume)),
            ("role requirement text resumes after about-us sections", lambda: test_role_requirement_text_resumes_after_about_us_sections(build_resume)),
            ("choose resume", lambda: test_choose_resume(build_resume)),
            ("lane profiles and summaries", lambda: test_lane_profiles_and_summaries(build_resume)),
            ("first person detector ignores role level i", lambda: test_first_person_detector_ignores_role_level_i(build_resume)),
            ("federal agency extraction", lambda: test_federal_agency_extraction(build_resume)),
            ("federal source load", lambda: test_federal_source_load(build_federal_resume)),
            ("federal standard essay responses", lambda: test_federal_standard_essay_responses(build_federal_resume)),
            ("federal qualifications append shared question prep", lambda: test_federal_qualifications_append_additional_questions_and_recent_interview_prep(build_federal_resume)),
            ("federal summary structure", lambda: test_federal_summary_structure(build_federal_resume)),
            ("federal ai summary and selection", lambda: test_federal_ai_summary_and_selection(build_federal_resume)),
            ("federal requirement audit and keywords", lambda: test_federal_requirement_audit_and_keywords(build_federal_resume)),
            ("federal layouts stay at ten point", lambda: test_federal_layouts_stay_at_ten_point(build_federal_resume)),
            ("federal visibility report tracks selected requirements", lambda: test_federal_visibility_report_tracks_selected_requirements(build_federal_resume)),
            ("federal standalone agency output name", lambda: test_federal_output_name_with_standalone_agency(build_federal_resume)),
            ("federal supporting doc resolution", lambda: test_federal_supporting_doc_resolution(build_federal_resume, federal_supporting_docs)),
            ("federal ATS plain-text warnings split hours and salary", lambda: test_federal_plain_text_validation_splits_hours_and_salary_warnings(build_federal_resume)),
            ("federal resume plan warns on unverified page count", lambda: test_federal_resume_plan_warns_when_page_count_is_unverified(build_federal_resume)),
            (
                "commercial builder entrypoints delegate to helpers",
                lambda: test_commercial_builder_entrypoints_delegate_to_input_helpers(
                    build_cover_letter,
                    build_interview_cheat_sheet,
                    build_detailed_interview_guide,
                ),
            ),
            (
                "interview outputs inherit BRIDGE resume name",
                lambda: test_interview_outputs_inherit_bridge_resume_name(
                    build_interview_cheat_sheet,
                    build_detailed_interview_guide,
                ),
            ),
            ("federal workflow supporting flags", test_federal_workflow_supporting_flag_resolution),
            ("federal workflow supporting step order", test_federal_workflow_runs_supporting_steps_after_resume),
            ("resume workflow dry-run skips upfront validation", test_run_resume_workflow_dry_run_skips_upfront_job_validation),
            ("federal workflow dry-run skips upfront validation", test_run_federal_workflow_dry_run_skips_upfront_job_validation),
            ("federal workflow dry-run labels unverified page counts", test_run_federal_workflow_dry_run_labels_unverified_page_counts),
            ("tasks register federal supporting doc commands", test_tasks_register_federal_supporting_doc_commands),
            ("tasks register interview review command", test_tasks_register_interview_review_command),
            ("onedrive run guard absent", test_onedrive_run_guard_absent),
            ("onedrive run guard present", test_onedrive_run_guard_present),
            ("resume builder source passes alignment grade to final fit audit", lambda: test_resume_builder_source_passes_alignment_grade_to_final_fit_audit(build_resume)),
            ("application checklist fit classification passes alignment grade", lambda: test_application_checklist_fit_classification_passes_alignment_grade(build_application_checklist)),
            ("final fit audit safe none alignment grade", lambda: test_final_fit_audit_safe_none_alignment_grade(build_resume)),
            ("claude packet run command keeps stdout and stderr", lambda: test_claude_packet_run_command_keeps_stdout_and_stderr(build_claude_review_packet)),
            ("claude packet command cache", lambda: test_claude_packet_command_cache(build_claude_review_packet)),
            ("packet excerpt resolution and contract", lambda: test_packet_excerpt_resolution_and_contract(build_claude_review_packet)),
            ("cleanup render checks only targets timestamped nested folders", test_cleanup_render_checks_limits_nested_cleanup_to_timestamped_folders),
            ("bootstrap copy filter preserves render checks module", test_bootstrap_copy_filter_preserves_render_checks_module),
            ("bootstrap gitignore ignores generated bundle manifest", test_bootstrap_gitignore_ignores_generated_bundle_manifest),
            ("bootstrap copy filter excludes retired onedrive markers", test_bootstrap_copy_filter_excludes_retired_onedrive_markers),
            ("bootstrap writes live canonical launchers", test_bootstrap_writes_live_canonical_launchers),
            ("bootstrap retire source launchers skips nested canonical root", test_bootstrap_retire_source_launchers_skips_nested_canonical_root),
            ("bootstrap reset destination removes nested tree", test_bootstrap_reset_destination_removes_nested_tree),
            ("summary three sentence structure", lambda: test_summary_three_sentence_structure(build_resume)),
            ("startup operator summary structure", lambda: test_startup_operator_summary_structure(build_resume)),
            ("summary detectors avoid member and operator false positives", lambda: test_summary_detectors_avoid_member_and_operator_false_positives(build_resume)),
            ("implementation summary surfaces quality and process language", lambda: test_implementation_summary_surfaces_quality_and_process_language(build_resume)),
            ("natural top bullet meta penalty", lambda: test_natural_top_bullet_meta_penalty(build_resume)),
            ("boomi presales summary", lambda: test_boomi_presales_summary(build_resume)),
            ("customer success summary capitalization", lambda: test_customer_success_summary_capitalization(build_resume)),
            ("customer success portfolio scope", lambda: test_customer_success_summary_clarifies_portfolio_scope(build_resume)),
            ("general consulting lane isolation", lambda: test_general_consulting_lane_isolated_from_other_consultant_roles(build_resume)),
            ("strategy transformation consultant lane override", lambda: test_strategy_transformation_consultant_prefers_corporate_strategy(build_resume)),
            ("consulting audit keyword cleanup", lambda: test_consulting_audit_keywords_filter_recruiting_fluff(build_resume)),
            ("audit keywords filter company affiliates and weak fragments", lambda: test_audit_keywords_filter_company_affiliates_and_weak_fragments(build_resume)),
            ("audit keywords filter boilerplate adjectives and normalize es plurals", lambda: test_audit_keywords_filter_boilerplate_adjectives_and_normalize_es_plurals(build_resume)),
            ("keyword placement prefers role phrases", lambda: test_keyword_placement_prefers_role_phrases(build_resume)),
            ("AI bridge evidence config", lambda: test_ai_evidence_area_config(build_resume)),
            ("construction engineering context and bridge coverage", lambda: test_construction_engineering_context_and_bridge_coverage(build_resume)),
            ("alignment score ignores unsupported named platforms", lambda: test_alignment_score_ignores_unsupported_named_platforms(build_resume)),
            ("bridge evidence sweep covers multiple lanes", test_bridge_evidence_sweep_covers_multiple_lanes),
            ("summary condense guard", lambda: test_summary_condense_guard(build_resume)),
            ("ERP summary rebalance", lambda: test_erp_summary_rebalance(build_resume)),
            ("summary anchor retention for change adoption", lambda: test_summary_anchor_retention_for_change_adoption(build_resume)),
            ("summary quantified anchor detects multi-digit proof", test_summary_quantified_anchor_detects_multi_digit_proof),
            ("summary repair preserves three sentences for high-signal variants", lambda: test_summary_repair_preserves_three_sentences_for_high_signal_variants(build_resume)),
            ("substitution safety respects paragraph boundaries", test_substitution_safety_respects_paragraph_boundaries),
            ("Aptean customer success role summary passes prose check", lambda: test_aptean_customer_success_role_summary_passes_prose_check(build_resume)),
            ("proof first close uses discuss and has no first person switch", lambda: test_proof_first_close_uses_discuss_and_has_no_first_person_switch(build_cover_letter)),
            ("enterprise CS JD extracts CS terms not analytics fallback", lambda: test_cs_enterprise_jd_extracts_cs_terms_not_analytics_fallback(build_cover_letter)),
            ("alignment score distinguishes lane and domain fit", lambda: test_alignment_score_report_distinguishes_lane_and_domain_fit(build_resume)),
            ("cover letter colon smoothing", lambda: test_cover_letter_colon_smoothing(build_cover_letter)),
            ("cover sentence score prioritizes signal density over length", lambda: test_cover_sentence_score_prioritizes_signal_density_over_length(build_cover_letter)),
            ("mission or context sentence survives job label header", lambda: test_mission_or_context_sentence_survives_job_label_header(build_cover_letter)),
            ("word budget trims opening filler before dense proof", lambda: test_word_budget_trims_opening_filler_before_dense_proof(build_cover_letter)),
            ("cover selection prefers lane direct and keeps ERP for ERP JD", lambda: test_cover_selection_prefers_lane_direct_and_keeps_erp_for_erp_jd(build_cover_letter, build_resume)),
            ("cover opening quality flags article and circularity", lambda: test_cover_opening_quality_flags_article_and_circularity(build_cover_letter)),
            ("consulting story summary avoids list density overload", lambda: test_consulting_story_summary_avoids_list_density_overload(writing_eval, build_resume)),
            ("cover letter compaction", lambda: test_cover_letter_compaction(build_cover_letter)),
            ("cover letter blocks JD artifacts and warns on switch", lambda: test_cover_letter_validator_blocks_jd_artifacts_and_warns_on_switch(build_cover_letter)),
            ("cover letter allows requirements in normal prose", lambda: test_cover_letter_validator_allows_requirements_in_normal_prose(build_cover_letter)),
            ("cover letter sections recognize extended headers", lambda: test_cover_letter_sections_recognize_extended_headers(build_cover_letter)),
            ("cover letter sections split typical day headers", lambda: test_cover_letter_sections_split_typical_day_headers(build_cover_letter)),
            ("cover letter plan normalizes slash titles", lambda: test_cover_letter_plan_normalizes_slash_role_titles(build_cover_letter)),
            ("cover letter blocks unsupported target context", lambda: test_cover_letter_blocks_unsupported_target_context(build_cover_letter)),
            ("extract company mission prefers clean intro line", lambda: test_extract_company_mission_prefers_clean_intro_line(build_cover_letter)),
            ("extract company mission prefers company descriptor over role intro", lambda: test_extract_company_mission_prefers_company_descriptor_over_role_intro(build_cover_letter)),
            ("extract company mission filters recruiting slogans", lambda: test_extract_company_mission_filters_recruiting_slogans(build_cover_letter)),
            ("extract company mission prefers business descriptor over culture line", lambda: test_extract_company_mission_prefers_business_descriptor_over_culture_line(build_cover_letter)),
            ("extract company mission ignores federal headers", lambda: test_extract_company_mission_ignores_federal_headers(build_cover_letter)),
            ("federal cover fallback does not invent education mission", lambda: test_federal_cover_fallback_does_not_invent_education_mission(build_cover_letter)),
            ("cover letter smoothing preserves job acronyms", lambda: test_cover_letter_smoothing_preserves_job_acronyms(build_cover_letter)),
            ("cover letter smoothing normalizes double dashes", lambda: test_cover_letter_smoothing_normalizes_double_dashes(build_cover_letter)),
            ("expand short cover opening reaches QC minimum", lambda: test_expand_short_cover_opening_reaches_qc_minimum(build_cover_letter)),
            ("cover communication metric stays lane relevant", lambda: test_cover_communication_metric_stays_lane_relevant(build_cover_letter)),
            ("short proof paragraph pulls supporting metric", lambda: test_short_proof_paragraph_pulls_supporting_metric(build_cover_letter)),
            ("cover compaction avoids incomplete clause fragments", lambda: test_compact_cover_sentence_avoids_incomplete_clause_fragments(build_cover_letter)),
            ("change consulting cover letter stays in change lane", lambda: test_change_consulting_cover_letter_stays_in_change_lane(build_cover_letter)),
            ("communication metric fragment rejection", lambda: test_finalize_communication_metric_candidate_rejects_fragment(build_cover_letter)),
            ("cover letter synthetic JD cleanup", lambda: test_cover_letter_no_jd_artifacts(build_resume, build_cover_letter)),
            ("role title dash not treated as bullet artifact", lambda: test_has_bullet_artifact_ignores_role_title_dash(build_cover_letter)),
            ("cover prompt leak pattern catches maps directly to", test_cover_prompt_leak_patterns_catch_maps_directly_to),
            ("cover letter unicode company rewrite", lambda: test_cover_letter_unicode_company_pronoun_rewrite(build_cover_letter)),
            ("naturalness score and adverb cleanup", lambda: test_naturalness_score_and_adverb_cleanup(build_resume)),
            ("ownership language rewrites", lambda: test_ownership_language_rewrites(build_resume, build_detailed_interview_guide, build_interview_cheat_sheet, build_cover_letter)),
            ("build interview review sections", lambda: test_build_interview_review_sections(build_interview_review, interview_context)),
            ("positive question framing", lambda: test_positive_question_framing(build_interview_cheat_sheet)),
            ("alignment score report", lambda: test_alignment_score_report(build_resume)),
            ("alignment gate decision", lambda: test_alignment_gate_decision(build_resume)),
            ("dynamic header title line", lambda: test_dynamic_header_title_line(build_resume)),
            ("header dedupe avoids near-duplicate consultant titles", lambda: test_header_dedupe_avoids_near_duplicate_consultant_titles(build_resume)),
            ("scope marker injection", lambda: test_scope_marker_injection(build_resume)),
            ("competency relevance and page guards", lambda: test_competency_relevance_and_page_guards(build_resume)),
            ("final fit audit accepts presales top role heading", lambda: test_final_fit_audit_accepts_presales_style_top_role_heading(build_resume)),
            ("Guidehouse final fit promotes BRIDGE", lambda: test_final_fit_audit_promotes_bridge_for_guidehouse_fixture(build_resume)),
            ("competency cap after additions", lambda: test_add_simple_core_competencies_respects_cap(build_resume)),
            ("XML page estimate word guard", lambda: test_xml_page_estimate_uses_word_guard(build_resume)),
            ("resume non-ERP audit ignores company context", lambda: test_resume_non_erp_audit_ignores_company_context(build_resume)),
            ("planned competency trim integrity", lambda: test_resume_integrity_allows_planned_competency_trim(build_resume)),
            ("keyword placement audit", lambda: test_keyword_placement_audit(build_resume)),
            ("build resume uses selected resume text for profile and alignment", lambda: test_build_resume_uses_selected_resume_text_for_profile_and_alignment_report(build_resume)),
            ("obvious choice positioning", lambda: test_obvious_choice_positioning(build_resume, build_cover_letter, build_interview_cheat_sheet)),
            ("positioning statement output", lambda: test_positioning_statement_output(build_resume, build_interview_cheat_sheet)),
            ("future bridge summary and bullet clause", lambda: test_future_bridge_summary_and_bullet_clause(build_resume)),
            ("offer blocker logic", lambda: test_offer_blocker_logic(build_resume, build_cover_letter, build_detailed_interview_guide)),
            ("ats plain text validation", lambda: test_ats_plain_text_validation(build_resume)),
            ("moment in time context", lambda: test_moment_in_time_context(build_resume, build_cover_letter)),
            ("customer success opening concrete context", lambda: test_customer_success_opening_surfaces_concrete_context(build_resume, build_cover_letter)),
            ("customer success support paragraph separation", lambda: test_customer_success_support_paragraph_stays_separate(build_cover_letter)),
            ("cover letter uses company specific context", lambda: test_cover_letter_uses_company_specific_context(build_cover_letter)),
            ("gap address paragraph", lambda: test_gap_address_paragraph(build_cover_letter)),
            ("first 90 days cover sentence", lambda: test_first_90_days_cover_sentence(build_cover_letter)),
            ("diagnose before selling framework", lambda: test_diagnose_before_selling_framework(build_resume, build_interview_cheat_sheet)),
            ("bold diagnostic questions", lambda: test_bold_diagnostic_questions(build_resume, build_interview_cheat_sheet)),
            ("preloaded questions", lambda: test_preloaded_questions(build_resume, build_interview_cheat_sheet)),
            ("slot based summary and interview answers", lambda: test_slot_based_summary_and_interview_answers(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("value compression opening", lambda: test_value_compression_opening(build_resume, build_interview_cheat_sheet)),
            ("pitch variants reuse cover logic", lambda: test_pitch_variants_reuse_cover_letter_logic(build_resume, build_interview_cheat_sheet)),
            ("pitch variants add 15-second claim", lambda: test_pitch_variants_add_fifteen_second_claim_and_validate_opening(build_resume, build_interview_cheat_sheet)),
            ("claim-then-prove validator catches delayed openings", lambda: test_claim_then_prove_validator_catches_delayed_openings(build_interview_cheat_sheet)),
            ("natural voice validation and answer budgets", lambda: test_natural_voice_validation_and_answer_budgets(build_interview_cheat_sheet)),
            ("natural voice question coverage", lambda: test_natural_voice_question_coverage(build_resume, build_interview_cheat_sheet)),
            ("keyword reference uses natural story intro", lambda: test_keyword_reference_uses_natural_story_intro(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("industry depth and company-scoped logistics", lambda: test_industry_depth_and_company_scoped_logistics(build_resume, build_interview_cheat_sheet)),
            ("human motivation lane coverage", lambda: test_human_motivation_sentence_has_lane_coverage(build_resume, build_interview_cheat_sheet)),
            ("adjusted profile preserves non-lane fields", lambda: test_adjusted_profile_for_role_preserves_non_lane_fields(build_resume, build_interview_cheat_sheet)),
            ("story answer parts preserve spoken alias", lambda: test_story_answer_parts_preserve_spoken_alias(build_detailed_interview_guide)),
            ("story answer constructors use full keyword", test_story_answer_constructors_use_full_keyword),
            ("extended TMAY sections", lambda: test_extended_tmay_sections_build_time_ladder(build_resume, build_detailed_interview_guide)),
            ("extended TMAY uses module-level supersets", lambda: test_extended_tmay_sections_use_module_level_superset(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("behavioral answer scripts empty story guard", lambda: test_behavioral_answer_scripts_empty_story_guard(build_resume, build_interview_cheat_sheet)),
            ("story sample answer separates coaching note", lambda: test_story_sample_answer_separates_coaching_note(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("story sample answer reuses claim sentence in full", lambda: test_story_sample_answer_reuses_claim_sentence_in_full(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("delivery validator only scans scripted strings", lambda: test_delivery_validator_only_scans_scripted_strings(build_detailed_interview_guide)),
            ("scripted answer validator rejects unsupported metrics", lambda: test_scripted_answer_validator_rejects_unsupported_metrics(build_detailed_interview_guide)),
            ("AI customer work answer uses confirmed qualitative story", lambda: test_ai_customer_work_answer_uses_confirmed_qualitative_story(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("story sample answer does not call dead STAR selection", lambda: test_story_sample_answer_does_not_call_dead_star_selection(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("behavioral answer scripts use spoken answers", lambda: test_behavioral_answer_scripts_use_spoken_answers_for_core_prompts(build_resume, build_interview_cheat_sheet)),
            ("interview stage resolution and context parsing", lambda: test_interview_stage_resolution_and_context_parsing(interview_stage)),
            ("stage filename suffix composes with detailed guide names", lambda: test_stage_filename_suffix_composes_with_existing_detailed_guide_names(interview_stage, build_detailed_interview_guide, question_prep)),
            ("federal detailed guide wrapper keeps stage params optional", lambda: test_federal_detailed_guide_wrapper_keeps_stage_params_optional(build_federal_detailed_interview_guide, build_detailed_interview_guide, federal_supporting_docs)),
            ("story adaptation and pre-interview routine helpers", lambda: test_story_adaptation_and_pre_interview_routine_helpers(build_resume, build_interview_cheat_sheet)),
            ("spoken sentence split preserves leading characters", test_spoken_sentence_split_preserves_leading_word_characters),
            ("nonconverged spoken repairs are collected", test_nonconverged_spoken_repairs_are_collected),
            ("interview filters filler and claim first answers", lambda: test_interview_filters_filler_and_claim_first_answers(build_resume, build_interview_cheat_sheet, interview_context)),
            ("candidate facing outputs avoid raw core problem", lambda: test_candidate_facing_outputs_avoid_raw_core_problem(build_resume, build_cover_letter, build_interview_cheat_sheet)),
            ("first 90 day plan reuses shared stage source", lambda: test_first_90_day_plan_reuses_shared_stage_source(build_resume, build_interview_cheat_sheet)),
            ("application answers use written confirmation and automation boundary", lambda: test_application_answers_use_written_confirmation_and_automation_boundary(question_prep)),
            ("detailed guide stage patterns and debrief overlay", lambda: test_detailed_guide_stage_patterns_and_debrief_overlay(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("interview companion documents reuse shared stage sources", lambda: test_interview_companion_documents_reuse_shared_stage_sources(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("federal version control scope gate is unconditional", lambda: test_federal_version_control_scope_gate_is_unconditional(build_federal_resume)),
            ("Foundant summary uses human close", lambda: test_foundant_summary_uses_human_close(build_resume)),
            ("commercial resume model provenance and render", lambda: test_commercial_resume_model_provenance_and_render(commercial_resume_model, build_resume)),
            ("canonical catalog preserves provenance identity", lambda: test_canonical_catalog_preserves_provenance_identity(commercial_resume_model)),
            ("application question pairing detects JD swap", test_application_question_pairing_detects_jd_swap),
            ("repo guidance prefers rehearsed foundation", test_repo_guidance_prefers_rehearsed_foundation),
            ("pitch helpers handle missing cover letter pitch parts", lambda: test_pitch_helpers_handle_missing_cover_letter_pitch_parts(build_resume, build_interview_cheat_sheet)),
            ("read the room opening", lambda: test_read_the_room_opening(build_resume, build_interview_cheat_sheet)),
            ("interview question prep sections render", lambda: test_interview_question_prep_sections_render(build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("companies refer to same", test_companies_refer_to_same),
            ("debrief active job matching", test_debrief_active_job_matching),
            ("compact context budget and diagnosis", lambda: test_interview_context_compact_context_budget_and_diagnosis(interview_context)),
            ("great eight shared utility", test_great_eight_utils),
            ("interview cheat sheet static content guards", lambda: test_interview_cheat_sheet_static_content_guards(build_interview_cheat_sheet)),
            ("late-stage context detection", lambda: test_late_stage_context_detection(build_interview_cheat_sheet)),
            ("interview bullet marker is clean", lambda: test_interview_bullet_marker_is_clean(build_interview_cheat_sheet)),
            ("why-company and negotiation context", lambda: test_why_company_and_negotiation_context(build_resume, build_interview_cheat_sheet)),
            ("communication audit context", lambda: test_communication_audit_context(build_interview_cheat_sheet)),
            ("story quality audit by type", lambda: test_story_quality_audit_by_type(build_resume, build_interview_cheat_sheet, build_detailed_interview_guide)),
            ("lead burial and response calibration", lambda: test_lead_burial_and_response_calibration(build_cover_letter, build_interview_cheat_sheet)),
            ("four trust questions audit", lambda: test_four_trust_questions_audit(build_resume, build_detailed_interview_guide)),
            ("executive presence signals", lambda: test_executive_presence_signals(build_resume, build_detailed_interview_guide, build_cover_letter)),
            ("lane specific recording focus", lambda: test_lane_specific_recording_focus(build_resume, build_detailed_interview_guide)),
            ("debrief pattern analysis", lambda: test_debrief_pattern_analysis(build_interview_cheat_sheet, build_debrief_analysis)),
            ("interview context review analysis", lambda: test_interview_context_review_analysis(interview_context)),
            ("interview context extracts exact interviewer questions", lambda: test_interview_context_extracts_exact_interviewer_questions(interview_context)),
            ("interview context structured round storage", lambda: test_interview_context_structured_round_storage(interview_context)),
            ("interview context review appendix storage", lambda: test_interview_context_review_appendix_storage(interview_context)),
            ("interview context global coaching fallback is lazy", lambda: test_interview_context_global_coaching_fallback_is_lazy(interview_context)),
            ("interview context recent question lookup", lambda: test_interview_context_recent_question_lookup_prefers_same_company(interview_context)),
            ("interview context legacy repair", lambda: test_interview_context_legacy_repair(interview_context)),
            ("interview context same-round dedupe", lambda: test_interview_context_dedupes_same_round_records(interview_context)),
            ("interview context scoped company filtering", lambda: test_interview_context_scopes_company_context_and_role_variants(interview_context)),
            ("discussion topic helpers", lambda: test_discussion_topic_helpers(utils)),
            ("template leakage validator", lambda: test_template_leakage_validator(utils)),
            ("thank-you body normalizes discussion topics", lambda: test_thank_you_body_normalizes_discussion_topics(build_thank_you, utils)),
            ("interview followup body", lambda: test_interview_followup_body(build_interview_followup)),
            ("post-round followup normalizes discussion topics", lambda: test_post_round_followup_email_normalizes_discussion_topics(build_post_round, utils)),
            ("post-round followup is lane aware", lambda: test_post_round_followup_email_is_lane_aware(build_post_round)),
            ("build skills database", lambda: test_build_skills_database(build_skills_database)),
            ("default question when application questions file is empty", lambda: test_standard_qualifications_default_question_when_file_empty(build_standard_qualifications_statement, build_standard_qualifications_statement.question_prep)),
            ("dry-run reports default question usage", lambda: test_dry_run_reports_default_question_usage(build_standard_qualifications_statement.question_prep)),
            ("default questions skip stale pairing check", lambda: test_default_questions_skip_stale_pairing_check(build_standard_qualifications_statement.question_prep)),
            ("explicit stale questions still flagged", lambda: test_explicit_stale_questions_still_flagged(build_standard_qualifications_statement.question_prep)),
            ("qualifications builder uses question prep response engine", test_qualifications_builder_uses_question_prep_response_engine),
            ("qualifications builder removes local shadow answer helpers", test_qualifications_builder_removes_local_shadow_answer_helpers),
            ("standard qualifications answers known questions", lambda: test_standard_qualifications_answers_known_questions(build_standard_qualifications_statement)),
            ("standard qualifications answers added company and implementation questions", lambda: test_standard_qualifications_answers_added_company_and_implementation_questions(build_standard_qualifications_statement)),
            ("standard application question parser", lambda: test_standard_application_question_parser_dedupes_blocks(build_standard_qualifications_statement)),
            ("standard qualifications render recent interview prep", lambda: test_standard_qualifications_document_renders_recent_interview_questions(build_standard_qualifications_statement)),
            ("standard qualifications recent interviewer scripts resolve factual script", lambda: test_standard_qualifications_recent_interviewer_scripts_resolve_factual_script(build_standard_qualifications_statement)),
            ("startup interview false-positive guard", lambda: test_startup_interview_false_positive_guard(build_interview_cheat_sheet)),
            ("application checklist debrief lookup", lambda: test_application_checklist_debrief_lookup(build_application_checklist)),
            ("application checklist prefers tailored analysis resume", lambda: test_application_checklist_analysis_resume_prefers_tailored_output(build_application_checklist)),
            ("application checklist source fallback uses alignment grade", lambda: test_application_checklist_source_fallback_uses_alignment_grade(build_application_checklist)),
            ("general advice active job helpers", lambda: test_general_advice_active_job_helpers(build_general_advice)),
            ("general advice shared sections", lambda: test_general_advice_shared_sections(build_general_advice, job_search_guidance)),
            ("question intent framework", lambda: test_question_intent_framework(build_resume, build_interview_cheat_sheet)),
            ("likely question story direct example preference", lambda: test_likely_question_story_prefers_direct_example(build_interview_cheat_sheet)),
            ("likely question story avoids reuse when alternative exists", lambda: test_likely_question_story_avoids_reuse_when_alternative_exists(build_interview_cheat_sheet)),
            ("interview addition helpers", lambda: test_interview_addition_helpers(build_resume, build_interview_cheat_sheet)),
            ("search progress question is conditional", lambda: test_search_progress_question_is_conditional(build_resume, build_interview_cheat_sheet)),
            ("salary guide helpers", lambda: test_salary_guide_helpers(build_salary_guide, job_search_guidance)),
            ("proof text rewrites dense cover sentence", test_proof_text_rewrites_dense_cover_sentence),
            ("follow-up news and variants", lambda: test_followup_news_and_variants(build_resume, build_followup_email, interview_context)),
            ("follow-up variants pass prose quality", lambda: test_followup_email_variants_pass_prose_quality(build_resume, build_followup_email, utils)),
            ("resume readiness flags finance blockers", lambda: test_resume_readiness_flags_finance_blockers(build_resume)),
            ("cover close uses direct ask", lambda: test_cover_close_uses_direct_ask(build_cover_letter)),
            ("LinkedIn guidance helpers", lambda: test_linkedin_guidance_helpers(build_resume, build_linkedin_update)),
            ("cover letter signals ollie analytics", lambda: test_cover_letter_signals_ollie_analytics(build_cover_letter)),
            ("cover lane keeps procore in implementation", lambda: test_cover_lane_keeps_procore_in_implementation(build_resume, build_cover_letter)),
            ("cover lane keeps explicit implementation roles out of support", lambda: test_cover_lane_keeps_explicit_implementation_roles_out_of_support(build_resume, build_cover_letter)),
            ("cover lane prefers strategy consulting titles", lambda: test_cover_lane_prefers_strategy_consulting_titles(build_resume, build_cover_letter)),
            ("Procore bridge-hard standard outputs", lambda: test_procore_bridge_hard_standard_outputs(build_resume, build_cover_letter, build_interview_cheat_sheet)),
            ("standard cover mode", lambda: test_standard_cover_mode(build_cover_letter)),
            ("force-bridge standard cover stays natural", lambda: test_force_bridge_standard_cover_stays_natural(build_cover_letter)),
            ("cover letter validator blocks generic experience summary", lambda: test_cover_letter_validator_blocks_generic_experience_summary(build_cover_letter)),
            ("cover letter QC rejects lowercase proof paragraph", lambda: test_cover_letter_qc_rejects_lowercase_proof_paragraph(build_cover_letter)),
            ("cover opening names company specific role context", lambda: test_cover_opening_names_company_specific_role_context(build_cover_letter)),
            ("cover opening without mission still names company", lambda: test_cover_opening_without_mission_still_names_company(build_cover_letter)),
            ("education assessment cover draft", lambda: test_education_assessment_cover_draft(build_cover_letter)),
            ("Guidehouse cover opening uses consulting context", lambda: test_guidehouse_cover_opening_uses_consulting_context(build_cover_letter)),
            ("education assessment measurable accomplishment", lambda: test_education_assessment_accomplishment_prefers_measurable_bullet(build_cover_letter)),
            ("long cover mode", lambda: test_long_cover_mode(build_cover_letter)),
            ("cover letter inherits fail resume name", lambda: test_cover_letter_inherits_fail_resume_name(build_cover_letter)),
            ("cover letter inherits BRIDGE resume name", lambda: test_cover_letter_inherits_bridge_resume_name(build_cover_letter)),
            ("ollie cover acceptance", lambda: test_ollie_cover_acceptance(build_cover_letter)),
            ("cover letter validator blocks contractions and double-dashes", lambda: test_cover_letter_validator_blocks_contractions_and_double_dashes(build_cover_letter)),
            ("clean but bad cover regressions fail writing eval", lambda: test_clean_but_bad_cover_regressions_fail_writing_eval(writing_eval)),
            ("registered firm profile requires real bain name", lambda: test_registered_firm_profile_requires_real_bain_name(build_cover_letter)),
            ("bigfour cover opening avoids aspiration phrase", lambda: test_bigfour_cover_opening_avoids_aspiration_phrase(build_cover_letter)),
            ("reset jobs helpers", lambda: test_reset_jobs_helpers(reset_jobs)),
            ("job context archive active snapshot includes questions", test_job_context_archive_active_snapshot_includes_questions),
            ("cover letter trace records snapshot and selection debug", lambda: test_cover_letter_trace_records_snapshot_and_selection_debug(build_cover_letter)),
            ("run-level ERP scrub formatting", lambda: test_run_level_erp_scrub_preserves_formatting(build_resume)),
            ("supported rewrite ERP scrub", lambda: test_supported_rewrite_scrubs_erp_language(build_resume)),
            ("named ERP platform scrub clears summary audit", lambda: test_named_erp_platform_scrub_clears_summary_audit(build_resume)),
            ("supported text keeps named platforms outside summary", lambda: test_supported_text_keeps_named_platforms_outside_summary(build_resume)),
            ("non-ERP audit allows SAP Crystal Reports", lambda: test_non_erp_audit_allows_sap_crystal_reports(build_resume)),
            ("company profile stub", lambda: test_company_profile_stub(build_resume)),
            ("hiring manager skim lane terms", lambda: test_hiring_manager_skim_lane_terms(build_resume)),
            ("direct report line avoids people leadership false positive", lambda: test_direct_report_reporting_line_does_not_trigger_people_leadership(build_resume)),
            ("east west summary preserves Aptean ownership", lambda: test_east_west_solution_architecture_summary_preserves_aptean_ownership(build_resume)),
            ("role bullet budgets meet minimums", lambda: test_role_bullet_budgets_meet_minimums(build_resume)),
            ("contact constants", lambda: test_contact_constants(build_resume)),
            ("business context module", lambda: test_business_context_module(business_context)),
            ("audit keywords filter low signal quality phrases", lambda: test_audit_keywords_filter_low_signal_quality_phrases(build_resume)),
            ("audit keywords filter noisy bigrams", lambda: test_audit_keywords_filter_noisy_bigrams(build_resume)),
            ("Guidehouse audit keywords filter bridge noise", lambda: test_guidehouse_audit_keywords_filter_bridge_noise(build_resume)),
            ("job title label stripping", lambda: test_job_title_label_stripping(build_resume)),
            ("thank-you contact line filter", lambda: test_thank_you_contact_line_filter(build_thank_you)),
            ("thank-you proof points skip summary", lambda: test_thank_you_proof_points_skip_summary(build_thank_you)),
            ("company constants", lambda: test_company_constants(build_resume)),
            ("role heading detection", lambda: test_role_heading_detection(build_resume)),
            ("writing eval flags system narration", lambda: test_writing_eval_flags_system_narration(writing_eval)),
            ("writing eval passes clean summary", lambda: test_writing_eval_passes_clean_summary(writing_eval)),
            ("writing eval flags weak closes and list density", lambda: test_writing_eval_flags_weak_close_and_list_density(writing_eval)),
            ("warn-only prose enforcement allows prep text", lambda: test_enforce_prose_quality_warn_mode_allows_prep_text(utils)),
            ("writing eval loads file dataset", lambda: test_writing_eval_loads_file_backed_dataset(writing_eval)),
            ("writing eval extracts docx sections", lambda: test_writing_eval_extracts_docx_sections(writing_eval)),
            ("detailed guide notes context strips leading bullets", lambda: test_detailed_guide_notes_context_strips_leading_bullets(build_detailed_interview_guide)),
            ("extract writing examples writes snippets", lambda: test_extract_writing_examples_writes_snippets(extract_writing_examples)),
            ("tracker row prefers tailored resume for fit", test_track_row_for_active_job_prefers_tailored_resume_for_fit),
            ("tracker add row refreshes existing metadata", test_track_add_row_refreshes_existing_metadata),
            ("tracker refresh uses active-job fit status", test_track_refresh_uses_active_job_fit_status),
            ("tracker status precedence blocks regressions", test_tracker_status_precedence_blocks_regressions),
            ("tracker audit flag derivation", test_tracker_audit_flag_derivation),
            ("tracker refresh warning without matching jd", test_tracker_refresh_warning_without_matching_jd),
            ("tracker row job description prefers snapshot id", test_tracker_row_job_description_prefers_snapshot_id),
            ("tracker refresh warning when output missing", test_tracker_refresh_warning_when_output_missing),
            ("debrief outcome mapping and notes", test_debrief_outcome_mapping_and_notes),
            ("workflow tracks before optional failure", test_workflow_tracks_after_required_steps_before_optional_failure),
            ("workflow parses cover warning channels", test_workflow_parses_cover_warning_channels),
            ("tasks auto-archive environment for commercial command", test_tasks_auto_archive_environment_for_commercial_command),
            ("workflow hard-fails docx validation issues", test_workflow_hard_fails_docx_validation_issues),
            ("claude packet modes", lambda: test_claude_packet_modes(build_claude_review_packet)),
            ("claude packet self-audit checks task registration", lambda: test_claude_packet_self_audit_checks_task_registration(build_claude_review_packet)),
            ("claude packet self-audit accepts unquoted task mentions", lambda: test_claude_packet_self_audit_accepts_unquoted_task_mentions(build_claude_review_packet)),
            ("claude prompt templates", lambda: test_claude_prompt_templates(build_claude_prompt)),
            ("claude prompt rejects residual placeholders", lambda: test_claude_prompt_rejects_residual_placeholders(build_claude_prompt)),
            ("claude prompt requires current default packet manifest", lambda: test_claude_prompt_requires_current_default_packet_manifest(build_claude_prompt)),
            ("claude review bundle builders include expected phrases", lambda: test_claude_review_bundle_builders_include_expected_phrases(claude_review_bundle)),
            ("claude review bundle refresh", lambda: test_claude_review_bundle_refresh(claude_review_bundle)),
            ("claude refresh bundle", lambda: test_refresh_claude_review_bundle(refresh_claude_review_bundle, claude_review_bundle)),
            ("cleanup output finders and selective flag", test_cleanup_output_finders_and_selective_flag),
            ("tasks check prefers latest generated resume", test_tasks_check_prefers_latest_generated_resume),
            ("tasks register cleanup command", test_tasks_register_cleanup_command),
            ("interview negotiation section avoids bracket placeholders", lambda: test_interview_negotiation_section_avoids_bracket_placeholders(build_interview_cheat_sheet)),
            ("collapse redundant role blanks", lambda: test_collapse_redundant_role_blanks(build_resume)),
            ("resume experience alignment", lambda: test_resume_experience_alignment(build_resume)),
        ):
            stdout_buffer = io.StringIO()
            stderr_buffer = io.StringIO()
            try:
                with contextlib.redirect_stdout(stdout_buffer), contextlib.redirect_stderr(stderr_buffer):
                    check()
                executed_checks += 1
                passed += 1
            except BaseException as error:  # noqa: BLE001
                executed_checks += 1
                details = stderr_buffer.getvalue().strip() or stdout_buffer.getvalue().strip()
                if details:
                    failures.append(f"{label}: {error}\n  Captured output: {details}")
                else:
                    failures.append(f"{label}: {error}")

    total = executed_checks + 2
    if failures:
        print(f"Smoke test FAILED: {passed}/{total} checks passed.")
        for failure in failures:
            print(f"- {failure}")
        raise SystemExit(1)

    print(f"Smoke test PASSED: {passed}/{total} checks passed.")


if __name__ == "__main__":
    main()
