"""Apply small, source-bound interview updates to maintained Study guide documents."""

from __future__ import annotations

from docx import Document

from config.paths import INTERVIEW_STORY_CARD, PERSONAL_OPERATING_WORKBOOK


def replace_paragraph(document: Document, old: str, new: str) -> None:
    paragraph = next((item for item in document.paragraphs if item.text == old), None)
    if paragraph is None:
        raise ValueError(f"Study guide anchor not found: {old[:70]}")
    if paragraph.text != new:
        paragraph.text = new


def update_story_card() -> None:
    document = Document(INTERVIEW_STORY_CARD)
    replace_paragraph(
        document,
        "3 Inventory automation:  cut manual work 78 percent and discrepancies 22 percent with an audit trail.",
        "3 Inventory automation:  cut manual work 78 percent and discrepancies 22 percent with an audit trail. Builder: repeatable auditable process.",
    )
    replace_paragraph(
        document,
        "5 SMS channel:  stood up a zero-to-one support channel at Home Depot and made it repeatable.",
        "5 SMS channel:  stood up a documented zero-to-one support channel at Home Depot. Builder: repeatable channel design.",
    )
    replace_paragraph(
        document,
        "20 Request to release:  turned email-thread requests into backlog-ready work with UAT and validation.",
        "20 Request to release:  turned email-thread requests into backlog-ready work with UAT and validation. Builder: reusable release mechanism.",
    )
    replace_paragraph(
        document,
        "Cross-lane rules:  Two full-mode stories per interview is the ceiling; everything else goes short or PREP. Do not reuse a story across rounds in the same loop, panels compare notes. Full bank with all four modes per story: interview_prep / Project Delivery Interview Stories.md",
        "Cross-lane rules: Two full-mode stories per interview is the ceiling; everything else goes short or PREP. Builder cue: 3 process, 5 channel, 20 request-to-release. Full bank: interview_prep / Project Delivery Interview Stories.md",
    )
    document.save(INTERVIEW_STORY_CARD)


def update_personal_workbook() -> None:
    document = Document(PERSONAL_OPERATING_WORKBOOK)
    replace_paragraph(
        document,
        "Interview them back: ask about management style, autonomy, and how decisions and changes are explained.",
        "Interview them back: ask about management style, autonomy, and how decisions and changes are explained. Builder cue for repeatability signals: Story 3 (auditable process), 5 (documented channel), or 20 (request-to-release).",
    )
    replace_paragraph(
        document,
        "Morning of: run the 10-Minute Pre-Interview Checklist, rehearse only the 30-second anchor, breathe, read your Evidence Log, and expect the pre-performance spike (it means it is working).",
        "Morning of: run the 10-Minute Pre-Interview Checklist, rehearse only the 30-second anchor, breathe, read your Evidence Log, and expect the pre-performance spike. CS boundary: no quota, NRR/GRR attainment, or closed expansion dollars; use at-risk recovery, QBRs/executive reviews, and expansion discovery.",
    )
    document.save(PERSONAL_OPERATING_WORKBOOK)


def build() -> None:
    update_story_card()
    update_personal_workbook()


if __name__ == "__main__":
    build()
